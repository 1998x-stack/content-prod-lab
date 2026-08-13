#!/usr/bin/env python3
"""
Generic 即梦短视频 production-docx generator.
For every topic in the workbook, produce a grounded 即梦 production docx following
modern-qimin-jimeng-video's output-template, injecting:
  - the topic's own title / 内容类型 / 动画表现 / 证据基线
  - the 母题 knowledge base (知识骨架变量, 编辑部四问, 脚本深挖方向)
  - the volume visual language
  - a content-type-specific shot scaffold
"""
import os, json, re
from docx import Document
from docx.shared import Pt

TOPICS = json.load(open("/tmp/topics_raw.json", encoding="utf-8"))
KB = json.load(open("/tmp/knowledge_base.json", encoding="utf-8"))

VOL_VISUAL = {
    "衣事": "纤维微观 / 污渍脱落 / 水膜泡沫 / 滚筒机械运动 / 布料变形",
    "食事": "冰箱透明剖面 / 温度分区 / 水分迁移 / 锅内热量蒸汽 / 食材组织变化",
    "居事": "住宅剖面 / 空气流线 / 水汽凝结 / 污垢层次 / 家务动线",
    "养身": "日常行为对比 / 简洁人体结构示意 / 时间节律 / 姿势动作变化",
    "行事": "路线推进 / 时间轴 / 车站机场空间层级 / 行李模块化 / 门到门动线",
    "器用": "产品爆炸图 / 内部结构 / 家庭A·B·C场景 / 空间占用 / 长期耗材维护",
}
VOL_SCENE = {
    "衣事": "现代中国城市家庭洗衣阳台：米白墙面、浅木储物柜、白色前开门滚筒洗衣机、灰白台面，上午左侧自然光",
    "食事": "现代中国家庭厨房：石纹台面、木色橱柜、明亮料理灯，浅色瓷砖墙，冰箱/灶台分区清晰",
    "居事": "现代中国家庭住宅剖面：浅色墙面、原木家具、柔和自然光，可见空气流动与层次",
    "养身": "现代中国家庭日常起居场景：简洁人体示意与真实生活行为并置，光线柔和",
    "行事": "现代都市出行与通勤场景：地铁站/高铁站/机场空间示意，路线与时间轴可可视化",
    "器用": "现代中国家庭客厅或厨房的产品使用场景：家电/日用品内部剖视，空间与耗材可视化",
}

CTYPE = {
    "误区纠正": {"dur": 30, "shots": 3, "role": "认知落差→归因机制→一句正确判断"},
    "原理解释": {"dur": 30, "shots": 3, "role": "日常现象→微观机制→落到可执行方法"},
    "正确操作": {"dur": 30, "shots": 3, "role": "常见偏差→正确顺序步骤→家庭标准"},
    "场景选择": {"dur": 45, "shots": 4, "role": "真实需求→A/B场景→决策变量→落地选择"},
    "维护/清单": {"dur": 45, "shots": 4, "role": "问题切点→清单式维护→循环闭环→家庭规则"},
}

HEALTH_NOTE = ("健康内容：仅作一般健康教育；不做个体诊断与治疗建议；"
               "涉及频率/阈值/数字以可靠公共卫生来源为准，发布前核对。")

def sanitize(s):
    return re.sub(r'[\\/:*?"<>|]', "_", s).strip()

def kb_for(vol, mudi):
    for b in KB.get(vol, []):
        if b["母题"] == mudi:
            return b
    return None

def clean_var(v):
    return re.sub(r"^\d+\.\s*", "", v).strip()

def document_for_topic(topic):
    doc = Document()
    st = doc.styles["Normal"]
    st.font.name = "PingFang SC"
    st.font.size = Pt(10.5)
    vol = topic["卷"]
    vs_visual = VOL_VISUAL.get(vol, "")
    vs_scene = VOL_SCENE.get(vol, "")
    ct = CTYPE.get(topic["内容类型"], CTYPE["误区纠正"])
    kb = kb_for(vol, topic["母题"])
    kb_vars = kb["sections"].get("知识骨架：脚本必须讲清的变量", []) if kb else []
    kb_editor = kb["sections"].get("编辑器四问", []) if kb else kb["sections"].get("编辑部先问四个问题", []) if kb else []
    kb_deep = kb["sections"].get("脚本深挖方向", []) if kb else []
    shots = ct["shots"]
    funcs = [f.strip() for f in ct["role"].split("→")]
    while len(funcs) < shots:
        funcs.append("落地-收束")

    def h0(t): return doc.add_heading(t, level=0)
    def h1(t): return doc.add_heading(t, level=1)
    def h2(t): return doc.add_heading(t, level=2)
    def para(t): return doc.add_paragraph(t)
    def bullet(t): return doc.add_paragraph(t, style="List Bullet")
    def kv(label, value):
        p = doc.add_paragraph()
        r = p.add_run(label + "：")
        r.bold = True
        p.add_run(value)
        return p

    h0("现代齐民要术｜即梦短视频生产稿")
    sub = doc.add_paragraph()
    sub.add_run("选题：" + topic["标题"] + "　·　" + vol + "卷 · 母题" + topic["母题ID"] + " · 集" + topic["集序"]).italic = True
    if topic.get("冷启动"):
        sub.add_run("　【冷启动优先】")

    h1("一、项目头")
    kv("选题ID", topic["id"])
    kv("选题", topic["标题"])
    kv("所属卷", vol)
    kv("母题", topic["母题"] + "（" + topic["母题ID"] + "）")
    kv("内容类型", topic["内容类型"])
    kv("动画表现提示", topic["动画表现"])
    kv("推荐成片", "约" + str(ct["dur"]) + "秒")
    kv("生成拆分", str(shots) + " × 10 秒")
    kv("目标观众", "城市家庭里负责该日常事务的成年人（转发给家里人看的对象）")
    kv("观看动机", "搞清「" + topic["母题"] + "」上真正影响结果的变量，并得到一条能照做的家庭方法")
    kv("证据基线", topic["证据基线"])

    h1("二、核心结论")
    kvars = "、".join([clean_var(v) for v in kb_vars][:4]) if kb_vars else topic["母题"]
    para("CORE_CONCLUSION：围绕「" + topic["标题"] + "」，给出一个普通家庭能照做的一句判断/结论；"
         "判断建立在母题变量（" + kvars + "）之上，并标出适用条件与例外。")

    h1("三、证据与边界")
    para("证据基线：" + (topic["证据基线"] or "以官方指南/行业标准/产品说明为准"))
    para("稳定事实：本集结论对应的母题机制（见上文知识骨架）。")
    para("动态事实：涉及季节 / 设备 / 新品参数时，发布前核对官方说明。")
    para("不说过满：不承诺“一定/绝对”，不贬低某品牌，不做个体化保证。")
    if vol == "养身":
        para(HEALTH_NOTE)

    h1("四、三个开头钩子")
    para("A. 反常识：沿用「" + topic["标题"] + "」做一句与直觉相反的开场。")
    para("B. 结果冲突：用错误做法带来的结果差异制造悬念。")
    para("C. 家庭场景：把画面放进普通家庭对应空间（" + vs_scene + "）。")
    para("默认钩子 A，前 1—3 秒即出现具体结果/冲突画面。")

    h1("五、成片旁白")
    para("旁白只讲因果，不描述画面。骨架：先抛出「" + topic["标题"] + "」的判断，再讲机制/变量，最后落到一句家庭方法。")
    if kb_editor:
        para("本母题“编辑部四问”（指导旁白取舍）：")
        for q in kb_editor:
            bullet("· " + q)
    if kb_deep:
        para("脚本深挖方向（供深化旁白）：")
        for d in kb_deep:
            bullet("· " + d)

    h1("六、分镜总表")
    tbl = doc.add_table(rows=shots + 1, cols=6)
    tbl.style = "Light Grid Accent 1"
    for j, htxt in enumerate(["Shot", "时长", "功能", "主要视觉事件", "尾帧", "下一镜桥接"]):
        c = tbl.rows[0].cells[j]
        c.text = htxt
        for p in c.paragraphs:
            for run in p.runs:
                run.bold = True
    tail_parts = vs_visual.split(" / ") if vs_visual else ["片尾留白"]
    for k in range(shots):
        cells = tbl.add_row().cells
        cells[0].text = "S0" + str(k + 1)
        cells[1].text = "10s"
        cells[2].text = funcs[k]
        cells[3].text = "源自「" + topic["母题"] + "」知识骨架的主视觉事件"
        cells[4].text = tail_parts[k % len(tail_parts)] + " 桥接尾帧"
        cells[5].text = "提取尾帧作下一镜首帧"

    h1("七、Shot Card")
    for k in range(shots):
        h2("S0" + str(k + 1))
        kv("DURATION", "10s")
        kv("FUNCTION", funcs[k])
        kv("VOICEOVER", "（对应旁白片段，只讲因果）")
        kv("START_FRAME", vs_scene + "；主体与开头相扣。")
        kv("VISUAL_EVENT", funcs[k] + "：" + topic["动画表现"])
        kv("CAMERA", "单轴稳定推/移，符合即梦可生成节奏。")
        kv("PHYSICS", "现实物理铺垫；原理镜头进入微观/剖面示意，不瞬移。")
        kv("END_FRAME", "稳定停在可桥接的尾帧，作下一镜首帧。")
        kv("BRIDGE_TO_NEXT", "提取本镜尾帧作为下一镜首帧。")
        kv("POST_TEXT", "（后期加：本段一句话字幕）")
        kv("SFX", "轻环境音与动作音，只服务因果。")
        para("即梦视频 Prompt：")
        para("9:16 竖屏，高级半写实 3D/2.5D 现代家庭科学动画。" + vs_scene + "。" +
             "0—3s 引入「" + topic["标题"] + "」主视觉事件；3—7s 演示/讲解机制与操作；7—10s 收束停在尾帧。" +
             "无中文文字/字幕/Logo/水印；保持同一光线与镜头空间；无伪实验/伪专家；不出现可读数字表格。")
        doc.add_paragraph("")

    h1("八、后期字幕时间轴")
    t2 = doc.add_table(rows=1, cols=2)
    t2.style = "Light Grid Accent 1"
    t2.rows[0].cells[0].text = "时间"
    t2.rows[0].cells[1].text = "字幕"
    row_timeline = [
        ("0.0—2.5s", topic["标题"] + "？"),
        ("2.5—9.0s", "一句判断/问题（钩子）"),
        ("9.0—" + str(ct["dur"] - 9) + "s", "讲清机制与变量（对应知识骨架）"),
        (str(ct["dur"] - 9) + "—" + str(ct["dur"] - 1) + "s", "落地家庭方法/清单"),
        (str(ct["dur"] - 1) + "—" + str(ct["dur"]) + ".0s", "现代齐民要术｜衣食住行，件件有方法。"),
    ]
    for a, b in row_timeline:
        c = t2.add_row().cells
        c[0].text = a
        c[1].text = b

    h1("九、声音")
    para("旁白语气：稳定、中性、可信。BGM：低存在感。SFX：只服务因果和动作。")

    h1("十、封面")
    kv("画面", "「" + topic["母题"] + "」核心对象 或 A/B 对比（" + vs_visual + "）")
    kv("大字（6—12字）", "取标题/钩子的短表述（避免品牌包装）")
    kv("不要", "不要出现品牌 logo/包装，不出现夸张特效")

    h1("十一、三平台标题")
    kv("视频号", topic["母题"] + "：" + topic["标题"] + "（给家里转一份可照做的）")
    kv("抖音", topic["标题"] + "？先讲真相再给方法")
    kv("小红书", topic["母题"] + "｜" + topic["标题"] + "（收藏复查）")

    h1("十二、片尾")
    para("现代齐民要术｜衣食住行，件件有方法。")

    h1("十三、发布前 QA")
    for q in ["所有 AI 生成段 ≤10 秒",
              "每镜一个主视觉任务",
              "每镜有明确尾帧",
              "相邻镜头有桥接帧",
              "场景/物件/光线/运动连续",
              "不依赖 AI 生成可读文字",
              "无伪实验/伪专家/伪检测",
              "高准确内容已查证",
              "健康内容保持一般教育边界"]:
        bullet("☐ " + q)

    return doc, ct


if __name__ == "__main__":
    # quick smoke test on one topic
    t = next(x for x in TOPICS if x["id"] == "01-01-1")
    os.makedirs("/tmp/demo_docx2", exist_ok=True)
    doc, ct = document_for_topic(t)
    out = "/tmp/demo_docx2/test.docx"
    doc.save(out)
    print("built", out, "shots", ct["shots"])