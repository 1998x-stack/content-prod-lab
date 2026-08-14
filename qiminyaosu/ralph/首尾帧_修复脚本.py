#!/usr/bin/env python3
"""Recheck-driven fix: concretize START_FRAME/END_FRAME in every Shot Card of all 600 docx."""
import os, sys, glob
sys.path.insert(0, "/tmp")
from docx import Document
import great_prompts as gp

VOL_DIR = {"衣事":"01_衣事","食事":"02_食事","居事":"03_居事","养身":"04_养身","行事":"05_行事","器用":"06_器用"}
ROOT = "/Users/x/Desktop/content-prod-lab/qiminyaosu/ralph/生产稿_600"

def fix_one(tid):
    t = next((x for x in gp.TOPICS if x["id"] == tid), None)
    if not t:
        return False, 0
    kb = gp.kb_for(t["卷"], t["母题"])
    shots = gp.craft_all(t, kb)
    vdir = os.path.join(ROOT, VOL_DIR[t["卷"]])
    m = glob.glob(os.path.join(vdir, f"现代齐民要术_{tid}_*_生产稿.docx"))
    if not m:
        return False, 0
    fpath = m[0]
    doc = Document(fpath)
    paras = doc.paragraphs
    # We need to locate each Shot Card and replace its START_FRAME/END_FRAME<values>
    # iterate by shot groups: find "S01".."S0n" headings, replace within block.
    # shot_labels = S01..S0n
    shot_labels = [f"S0{i+1}" for i in range(len(shots))]
    fixed = 0
    for si, sh in enumerate(shots):
        label = shot_labels[si]
        # locate heading paragraph with that label (Heading 2)
        # We'll walk paragraphs; track current card
    # Simpler approach: walk paragraphs, track current shot via Heading2 S0x, replace within
    current = None
    for p in doc.paragraphs:
        txt = p.text.strip()
        if p.style.name == "Heading 2" and txt in shot_labels:
            current = int(txt[1:]) - 1  # 0-index shot idx
            continue
        if current is None:
            continue
        # Within this card, if paragraph starts with START_FRAME： replace value
        if current < len(shots):
            sh = shots[current]
            if txt.startswith("START_FRAME：") and not txt.startswith("START_FRAME：{\""):
                p.clear()
                p.add_run("START_FRAME（首帧·五锁）：" + sh["start"])
                fixed += 1
            elif txt.startswith("END_FRAME：") and not txt.startswith("END_FRAME：{\""):
                p.clear()
                p.add_run("END_FRAME（尾帧·五锁）：" + sh["end"])
                fixed += 1
    doc.save(fpath)
    return True, fixed

if __name__ == "__main__":
    ids = [t["id"] for t in gp.TOPICS]
    import sys
    # optional range
    start, endall = 1, len(ids)
    ok=0; fail=0; totalfix=0
    for i, tid in enumerate(ids, start=1):
        done, n = fix_one(tid)
        if done:
            ok+=1; totalfix+=n
        else:
            fail+=1
            print("FAIL", tid)
    print(f"done ok={ok} fail={fail} total_fixed={totalfix}")