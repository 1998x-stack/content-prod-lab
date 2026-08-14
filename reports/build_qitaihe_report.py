# -*- coding: utf-8 -*-
"""Build 济宁市2025年政府工作报告 深度研究 DOCX, 参照地级市系列版式。"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DARK = RGBColor(0x1F, 0x2A, 0x44)
RED = RGBColor(0xB0, 0x1E, 0x2E)
GRAY = RGBColor(0x59, 0x60, 0x69)
BLUE = RGBColor(0x1F, 0x4E, 0x79)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
HEADER_FILL = "1F2A44"
K = "微软雅黑"


def set_run(run, size=11, bold=False, color=DARK, font=K):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    run.font.name = font
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), font)
    rFonts.set(qn("w:hAnsi"), font)
    rFonts.set(qn("w:eastAsia"), font)


def para(doc, text, size=11, bold=False, color=DARK, align=None,
         space_after=6, space_before=0, font=K):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        set_run(r, size=size, bold=bold or (i % 2 == 1), color=color, font=font)
    return p


def heading1(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(16)
    pf.space_after = Pt(8)
    r = p.add_run(text)
    set_run(r, size=16, bold=True, color=RED)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "B01E2E")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def heading2(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(10)
    pf.space_after = Pt(4)
    r = p.add_run(text)
    set_run(r, size=13, bold=True, color=BLUE)
    return p


def table(doc, headers, rows, widths=None, font_size=9.5):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_run(r, size=font_size, bold=True, color=WHITE)
        tcPr = hdr[i]._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), HEADER_FILL)
        tcPr.append(shd)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(val))
            set_run(r, size=font_size, color=DARK)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    return t


def quote_box(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(8)
    pf.left_indent = Pt(12)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), "B01E2E")
    pBdr.append(left)
    pPr.append(pBdr)
    r = p.add_run(text)
    set_run(r, size=11, color=DARK, font="楷体")
    return p


def bullet(doc, text, size=10.5):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.7)
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        set_run(r, size=size, bold=(i % 2 == 1), color=DARK)
    return p


# ================================================================= 文档主体
doc = Document()
sec = doc.sections[0]
sec.page_width = Cm(21)
sec.page_height = Cm(29.7)
sec.top_margin = Cm(2.4)
sec.bottom_margin = Cm(2.2)
sec.left_margin = Cm(2.5)
sec.right_margin = Cm(2.5)

st = doc.styles["Normal"]
st.font.name = K
st.font.size = Pt(11)
st.element.rPr.rFonts.set(qn("w:eastAsia"), K)


# ---- 封面 ----
para(doc, "七台河市2025年政府工作报告", size=24, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=60, space_after=4)
para(doc, "深度研究与\u201c容易被忽视的细节\u201d分析", size=16, bold=True, color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
para(doc, "从\u201c煤城转型、世界冠军之城（短道速滑）、煤化工、生物医药、现代农业\u201d重新理解七台河", size=12, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)
para(doc, "\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501", color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
para(doc, "研究对象：2025年七台河市政府工作报告及同期财政、国民经济和社会发展统计资料、2026年官方复盘", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
para(doc, "标定日期：2026-08-14", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
doc.add_page_break()

# ---- 目录 ----
catalog = [
    "执行摘要",
    "一、研究口径与阅读方法",
    "二、先看七台河的特别底盘：煤城、短道速滑冠军、农业、生物医药、转型",
    "三、最关键的宏观错位：GDP+5.5%增速全省第1、回升，但煤焦承压、体量小、财政弱",
    "四、容易被忽视、但信号很重的细节（15条）",
    "五、2025年GDP目标 vs 实际：对照表",
    "六、2025年增长是谁撑起来的？（结构归因）",
    "七、预算与财政的\u201c含金量\u201d",
    "八、民生底账：人口、收入与城乡",
    "九、城镇与农村：格局与均衡",
    "十、人口流入与流出",
    "十一、物价与货币环境",
    "十二、区域一体化：七台河在哈长城市群、黑龙江东部、煤城转型\u201c三圈\u201d里的位置",
    "十三、未来5—10年最值得观察的五条主线",
    "十四、最终结论：七台河在\u201c煤转绿+农业+冰雪\u201d里的增长逻辑",
    "附录A：主要资料来源与核验路径",
    "附录B：建议建立的年度跟踪仪表盘",
]
para(doc, "目　录", size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=10, space_after=10)
for item in catalog:
    para(doc, item, size=12, space_after=4)
doc.add_page_break()

# ---- 执行摘要 ----
heading1(doc, "执行摘要")
para(doc, "**核心判断**　2025年七台河最显著的是\u201cGDP 249.41亿/+5.5%（实际增速全省第1）、从2024年-7.2%大幅回升\u201d、\u201c但煤焦量价承压、体量小（黑龙江省最小地级市）、财政/固投/消费基数低\u201d。这说明七台河在\u201c煤城转型\u201d中，**GDP增速反超回升、但经济总量小、重工业依赖煤炭/焦炭**。")
para(doc, "把2025年目标（GDP+5.5%、规上+13%、社零+5%、固投+8%、进出口+4%、财收+10%）、2025年实际（GDP+5.5%达成全省第1、规上大幅回升、财收承压、进出口+17.6%）趋势看，七台河是\u201c煤焦+农业+冰雪\u201d路径：**煤炭洗选、焦化、煤化工、生物医药、现代农业、短道速滑冰雪（冠军之城）\u201d是支柱。")
para(doc, "最容易记住的一句话：**七台河是\u201c中国短道速滑之乡、煤城、全国产粮地\u201d，靠\u201c煤转绿+冰雪体育+现代农业\u201d驱动。**观察七台河，与其只看\u201cGDP 249亿\u201d，不如看\u201c增速全省第1（+5.5%）、短道速滑冠军、煤焦转型、大豆玉米农业\u201d。")
# ---- 一、研究口径与阅读方法 ----
heading1(doc, "一、研究口径与阅读方法")
para(doc, "本报告以\u201c七台河市政府工作报告（2025年1月，张涛作）\u201d为起点，把\u201c2025年GDP目标（5.5%）\u201d与\u201c官方2025年（249.41亿/+5.5%）\u201d并置对照，用\u201c七台河市统计公报\u201d和\u201c2026年政府工作报告\u201d横向核验。")
para(doc, "口径提示：\u201cGDP、规上工业增速\u201d按可比价（实际）；\u201c投资、消费、进出口、财政\u201d按现价（名义）。\u201c常住人口\u201d用统计公报（约72.6万户籍）。")
para(doc, "指标体系（与研究口径一致）：**总量与增速、产业动能（煤/农业/冰雪）、外贸、财政、民生与人口**。")
para(doc, "特别提示（不吃老本）：七台河2024年GDP 236.6亿/-7.2%（煤焦量价历史性冲击）、2025年249.41亿/+5.5%大幅回升；它不是\u201c只有煤\u201d——**短道速滑冠军之城、生物医药、现代农业、冰雪经济、煤焦转型\u201d才是真正底色。")
# ---- 二、先看七台河的特别底盘 ----
heading1(doc, "二、先看七台河的特别底盘：煤城、短道速滑冠军、农业、生物医药、转型")
para(doc, "七台河地处黑龙江省东部、东部湿润区，是**中国短道速滑冠军之城（世界冠军之乡，杨扬/王濛）、煤城（煤炭洗选焦化）、龙江东部粮仓/生猪基地**；是黑龙江省面积最小、人口最少的地级市。2025年GDP 249.41亿、户籍约72.6万、城镇化率约61%。")
para(doc, "五个底盘名词，先立框架：")
bullet(doc, "**煤炭/焦化（传统）**　煤炭开采洗选、焦炭、煤化工（精甲醇）——\u201c煤城\u201d。")
bullet(doc, "**短道速滑/冰雪（冠军之城）**　短道速滑国家队摇篮（王濛/鲍雪等）、冰雪装备——\u201c冠军之城\u201d。")
bullet(doc, "**现代农业/粮食**　粮食21.08亿斤（玉米/大豆/水稻）、粮食安全省基地——\u201c龙江粮仓\u201d。")
bullet(doc, "**生物医药/新能源**　医药制造、新能源（风/光）试点——\u201c产业转型\u201d。")
bullet(doc, "**文旅/冰雪/康养**　转型、冰雪/避暑、生态——\u201c绿色转型\u201d。")
para(doc, "这五根（煤+滑冰+农业+医药+转型）构成七台河独特底盘：**左手煤焦（量价压），右手冰雪冠军+农业**。理解七台河，先理解\u201c煤城转型、冠军之城、黑龙江人口最少的市\u201d。")
# ---- 三、最关键的宏观错位 ----
heading1(doc, "三、最关键的宏观错位：GDP+5.5%增速全省第1、回升，但煤焦承压、体量小、财政弱")
para(doc, "2025年七台河最需要辨析的一组\u201c错位\u201d：**GDP+5.5%（全省实际增速第1、较2024年-7.2%大幅回升）、但规模以上工业仍受煤焦拖累、第三产业占57.2%、经济总量仅249亿（黑龙江最小）\u201d、一般公共预算收入承压、固投/消费/进出口基数低**。")
para(doc, "为什么\u201cGDP增速全省第1\u201d但经济仍弱？三解释：")
para(doc, "**其一，低基数回升、增速最高**　2024年-7.2%（煤焦量价齐跌）、2025年+5.5%（同比回升、超哈/大庆）——\u201c触底反弹、低基数\u201d。")
para(doc, "**其二，产业/煤焦承压、体量小**　规上工业受煤炭洗选/焦化（量价弱）、医药+73.7%等补位；城市体量小（GDP 249亿）——\u201c结构单一、总量小\u201d。")
para(doc, "**其三，农业/冰雪/三产稳**　农业+（粮食21亿斤）、短道冰雪/三产占57.2%——\u201c农业+文旅\u201d。")
para(doc, "小结：七台河2025年是\u201c**低基数回升、转型中、体量小**\u201d：增速全省第1但总量小、煤焦弱、冰雪/农业/医药亮点。")

# ---- 四、容易被忽视、但信号很重的细节（15条） ----
heading1(doc, "四、容易被忽视、但信号很重的细节（15条）")
bullet(doc, "**1.GDP+5.5%（实际增速全省地级市第1）**\u201c触底回升、低基数。\u201d")
bullet(doc, "**2.从2024年-7.2%（煤焦量价）→2025年+5.5%**\u201c深度反弹。\u201d")
bullet(doc, "**3.第三产业占57.2%**\u201c服务业主导（冰雪/文旅/医疗）。\u201d")
bullet(doc, "**4.短道速滑冠军之城（杨扬/王濛）**\u201c冰雪/冠军IP、国家训练。\u201d")
bullet(doc, "**5.粮食21.08亿斤（玉米/大豆/水稻）**\u201c龙江粮仓、粮食安全。\u201d")
bullet(doc, "**6.医药制造业（+73.7%）、金属制品（+23.7%）**\u201c产业转型亮点。\u201d")
bullet(doc, "**7.煤焦（洗精煤281万t/焦炭291万t，2024降）**\u201c传统重在哪压。\u201d")
bullet(doc, "**8.进出口4.7亿/+17.6%（2024，出口+3.2%）**\u201c外贸小但稳。\u201d")
bullet(doc, "**9.一般公共预算收入29亿/-15.6%（2024）**\u201c财政弱、煤价拖累。\u201d")
bullet(doc, "**10.固定资产投资（工业/三产调整）**\u201c投资承压、中央投资增。\u201d")
bullet(doc, "**11.社零75.9亿/+2.2%、商品房销售-7.2%（2024）**\u201c消费/地产稳低。\u201d")
bullet(doc, "**12.常驻市域人口约72.6万、城镇化61%**\u201c人口少、县城化。\u201d")
bullet(doc, "**13.居民收入（城镇31595/农村20119，2024）**\u201c城镇+1.8%/农村+4.5%。\u201d")
bullet(doc, "**14.冰雪/避暑康养、转型（黑龙江东部科创）**\u201c绿色转型。\u201d")
bullet(doc, "**15.煤炭/能源安全、新能源试点**\u201c能源+转型。\u201d")

# ---- 五、2025年GDP目标 vs 实际：对照表 ----
heading1(doc, "五、2025年GDP目标 vs 实际：对照表")
table(doc,
    ["指标", "2025年目标", "2025年实际", "达成评价"],
    [
        ["地区生产总值(GDP)", "增长5.5%左右", "249.41亿/5.5%", "达成(全省第1)"],
        ["规模以上工业", "增长13%", "大幅回升", "煤焦承压"],
        ["社会消费品零售总额", "增长5%", "（2024: 75.9亿/+2.2%）", "基数低"],
        ["固定资产投资", "增长8%", "承压", "投资弱"],
        ["进出口总额", "增长4%", "（2024: +17.6%）", "稳"],
        ["一般公共预算收入", "可比口径增长10%", "承压", "财政弱"],
        ["城乡居民收入", "与经济增长同步", "（城镇+1.8%/农村+4.5%）", "农村快"],
    ],
)
para(doc, "注：GDP按可比价。**GDP（+5.5%）达成、增速全省第1（触底反弹）**；**规上（煤焦企稳）、财政、固投**仍承压（2024煤焦量价冲击后）。")
para(doc, "拆读：**低基数回升、第三（冰雪/文旅）、医药、农业是亮点**；**煤焦（量价）、财政（-15.6%）、总量小（249亿）、投资强**是短板——\u201c转型回升、体量小\u201d，是\u201c煤城转绿\u201d样本。")

# ---- 六、2025年增长是谁撑起来的？（结构归因） ----
heading1(doc, "六、2025年增长是谁撑起来的？（结构归因）")
para(doc, "把七台河GDP的+5.5%拆开（2024年：一产38.7亿/+2.4%、二产62.5亿/-29.7%、三产135.4亿/+3.4%，结构16.4：26.4：57.2）：**第三产业（+2.4%、占57.2%）是主要支撑、第一产业（农业）+稳、第二产业（煤焦）2024年-29.7%后回暖**。")
para(doc, "2026年七台河强调\u201c转型、冰雪经济、共建龙江东部科创中心、农业强市\u201d，聚焦**煤焦转型（煤化工/绿色）、短道冰雪（冠军之城）、生物医药、新能源、现代农业**——核心是\u201c减煤增收、冰雪农业\u201d。")
para(doc, "**第二产业（工业/煤焦）**：规上工业2024-45.9%（煤炭-55.1%）、2025回暖，医药+73.7%、金属+23.7%——\u201c煤焦压、新制造补\u201d。")
para(doc, "**第三产业（服务业/冰雪）**：占57.2%（冰雪体育、医疗、文旅）——\u201c冰雪/服务业\u201d。")
para(doc, "**第一产业（农业）**：+2.4%（粮食21亿斤、玉米/大豆/生猪）——\u201c龙江农业稳\u201d。")
para(doc, "一句话归因：**2025年七台河增长\u201c靠第三产业（冰雪/文旅）+农业+工业回暖\u201d**，低基数反弹；**煤焦转型、冰雪/冠军\u201d是独特底色。")

# ---- 七、预算与财政的\u201c含金量\u201d ----
heading1(doc, "七、预算与财政的\u201c含金量\u201d")
para(doc, "2024年七台河**一般公共预算收入29.0亿元（-15.6%）**；税收11.8亿（-21.2%）；支出112.7亿（+8.8%、社保/教育）。")
bullet(doc, "财收-15.6%（煤炭/焦炭税、量价）——\u201c财政承压、含金量低\u201d。")
bullet(doc, "民生支出：社会保障+11.8%、教育+5.4%、节能+74%（转型）。")
bullet(doc, "金融：存款878.9亿（+5.8%）、贷款444.2亿（+2.3%）——宽信用稳。")
para(doc, "**财政含金量小结**：财收承压（煤价）、民生维系、转型节能投入；财政对\u201c煤转绿、冰雪、农业、民生\u201d投入加大。")
# ---- 八、民生底账：人口、收入与城乡 ----
heading1(doc, "八、民生底账：人口、收入与城乡")
para(doc, "2024年七台河**城镇/农村居民人均可支配收入分别31595元（+1.8%）/20119元（+4.5%）**；城镇非私营单位就业人员平均工资80426元。")
para(doc, "人口画像：**户籍人口约72.6万（黑龙江最小）、城镇化率约61%**；老龄化（60岁+占28.2%）较重、人口外流。")
para(doc, "民生投入：社保/教育/医疗、节能环保+74%、冰雪体育设施——民生扎实、冠军之城培养。")

# ---- 九、城镇与农村：格局与均衡 ----
heading1(doc, "九、城镇与农村：格局与均衡")
para(doc, "七台河城镇化率约61%；农村收入增速（+4.5%）>城镇（+1.8%），**城乡差距缩小**；县域/粮食（勃利等）。")
para(doc, "农业底盘：**粮食21.08亿斤（玉米16.78/大豆1.85/水稻2.43）、生猪、禽蛋+11.6%**——\u201c龙江东部粮仓\u201d。")
para(doc, "一句话：\u201c七台河高农业比重、城镇化初期、城乡均衡推进\u201d。")

# ---- 十、人口流入与流出 ----
heading1(doc, "十、人口流入与流出")
para(doc, "七台河户籍约72.6万（黑龙江最小）、城镇化61%；\u201c煤城产业调整、人口外流（至哈市/大庆/省外）\u201d、冰雪体育/学校吸引。")
para(doc, "结构观察：**60岁+占28.2%（深度老龄化）、人口净流出**；冰雪后备人才。")
para(doc, "2026年目标：稳就业、冰雪产业育才——七台河靠\u201c冰雪+农业+转型\u201d聚人。")

# ---- 十一、物价与货币环境 ----
heading1(doc, "十一、物价与货币环境")
para(doc, "2024年七台河**CPI+0.6%（衣着+3.2%、居住+1.4%；食品烟酒-1.5%）**，PPI-3.4%（煤炭-1.9%）——\u201c低通胀、煤价弱\u201d。")
bullet(doc, "信贷：存款+5.8%、贷款+2.3%（转型/保障）——宽信用稳。")
bullet(doc, "消费：社零+2.2%（低）、以旧换新——消费弱。")
para(doc, "货币环境判断：**宽信用、低通胀**；七台河靠\u201c冰雪+农业+转型\u201d稳需求。")

# ---- 十二、区域一体化：七台河的位置 ----
heading1(doc, "十二、区域一体化：七台河在哈长城市群、黑龙江东部、煤城转型\u201c三圈\u201d里的位置")
para(doc, "七台河是**哈长城市群（黑龙江东部节点）、龙江东部产业科创中心共建、煤炭资源城市转型示范区**。")
bullet(doc, "**哈长城市群**　哈长城市群、对接哈尔滨（省会）、黑龙江东部。")
bullet(doc, "**龙江东部**　与佳木斯/双鸭山/鸡西/鹤岗（煤城）联动、东部产业科创中心。")
bullet(doc, "**煤炭转型**　国家煤城转型、东北振兴、冰雪/生态替代。")
para(doc, "一句话：**七台河在\u201c哈长+龙江东部+煤城转型\u201d里，最核心是\u201c冰雪冠军+煤炭转型\u201d**；短道速滑、区位、转型政策是优势。")

# ---- 十三、未来5—10年最值得观察的五条主线 ----
heading1(doc, "十三、未来5—10年最值得观察的五条主线")
bullet(doc, "**主线一：煤焦转型/绿色化**\u201c煤化工、煤焦升级、生态修复\u201d。")
bullet(doc, "**主线二：短道冰雪（冠军之城）**\u201c冰雪经济、训练基地、装备\u201d。")
bullet(doc, "**主线三：生物医药/新能源**\u201c医药+73.7%、风光、新材料\u201d转型。")
bullet(doc, "**主线四：现代农业/粮食**\u201c大豆玉米、生猪、食品加工\u201d。")
bullet(doc, "**主线五：人口/财政/体量**\u201c老龄化外流、财政增、总量跃升\u201d。")

# ---- 十四、最终结论 ----
heading1(doc, "十四、最终结论：七台河在\u201c煤转绿+冰雪冠军+农业\u201d里的增长逻辑")
para(doc, "把2025年闭环打穿：**七台河是\u201c短道冠军之城、煤城、龙江粮仓\u201d**：GDP 249.41亿/+5.5%（增速全省第1）、第三产业占57.2%、冰雪/医药/农业。")
para(doc, "七台河不是\u201c只有煤\u201d——它是**煤焦转型+短道冰雪+生物医药+现代农业**的复合，靠\u201c冰雪+农业+转型\u201d驱动；但体量小、煤焦/财政弱、老龄化。")
para(doc, "一句话结论：**七台河是\u201c冠军之城、煤转绿、东部粮仓\u201d；观察它先看\u201c增速回升、冰雪、医药、粮食\u201d，再看\u201c煤焦、财政、人口\u201d。**它是\u201c低基数回升、转型中\u201d的龙江煤城样本。")

# ---- 附录A ----
heading1(doc, "附录A：主要资料来源与核验路径")
bullet(doc, "《2025年七台河市政府工作报告》（2025年1月，张涛作，2025年目标、2024年回顾）")
bullet(doc, "《2024年七台河市国民经济和社会发展统计公报》（七台河市统计局，2025-07-26，参照基期）")
bullet(doc, "《2026年七台河市政府工作报告》（2026年1月，复盘+2026年目标）")
bullet(doc, "七台河市人民政府/统计局（qth.gov.cn）、黑龙江省人民政府")

# ---- 附录B：建议建立的年度跟踪仪表盘 ----
heading1(doc, "附录B：建议建立的年度跟踪仪表盘")
bullet(doc, "GDP总量/增速 vs 全年目标（全省排名）。")
bullet(doc, "规上工业/煤炭/焦炭（量价）。")
bullet(doc, "煤化工/医药/金属制造。")
bullet(doc, "冰雪产业/冠军/IP。")
bullet(doc, "粮食/玉米/大豆。")
bullet(doc, "固定资产/工业/地产投资。")
bullet(doc, "社零/进出口/以旧换新。")
bullet(doc, "财收/税收/民生。")
bullet(doc, "人口/老龄化/就业。")
bullet(doc, "CPI/PPI/存贷款。")

# ---------------- 保存 ----------------
out = "/Users/x/Desktop/content-prod-lab/reports/七台河市_2025年政府工作报告_深度研究_2026-08-14.docx"
doc.save(out)
print("SAVED OK: 七台河市", out)
