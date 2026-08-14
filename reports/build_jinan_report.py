# -*- coding: utf-8 -*-
"""Build 济南市2025年政府工作报告 深度研究 DOCX, 参照省会系列版式。"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DARK = RGBColor(0x1F, 0x2A, 0x44)
RED = RGBColor(0xB0, 0x1E, 0x2E)
GRAY = RGBColor(0x59, 0x60, 0x69)
BLUE = RGBColor(0x1F, 0x4E, 0x79)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
HEADER_FILL = "1F2A44"
K = "微软雅黑"


def set_run(run, size=11, bold=False, color=DARK, font=K):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    run.font.name = font
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), font)
    rFonts.set(qn("w:hAnsi"), font)
    rFonts.set(qn("w:eastAsia"), font)


def para(doc, text, size=11, bold=False, color=DARK, align=None,
         space_after=6, space_before=0, font=K):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        set_run(r, size=size, bold=bold or (i % 2 == 1), color=color, font=font)
    return p


def heading1(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(16)
    pf.space_after = Pt(8)
    r = p.add_run(text)
    set_run(r, size=16, bold=True, color=RED)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "B01E2E")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def heading2(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(10)
    pf.space_after = Pt(4)
    r = p.add_run(text)
    set_run(r, size=13, bold=True, color=BLUE)
    return p


def table(doc, headers, rows, widths=None, font_size=9.5):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_run(r, size=font_size, bold=True, color=WHITE)
        tcPr = hdr[i]._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), HEADER_FILL)
        tcPr.append(shd)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(val))
            set_run(r, size=font_size, color=DARK)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    return t


def quote_box(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(8)
    pf.left_indent = Pt(12)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), "B01E2E")
    pBdr.append(left)
    pPr.append(pBdr)
    r = p.add_run(text)
    set_run(r, size=11, color=DARK, font="楷体")
    return p


def bullet(doc, text, size=10.5):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.7)
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        set_run(r, size=size, bold=(i % 2 == 1), color=DARK)
    return p


# ================================================================= 文档主体
doc = Document()
sec = doc.sections[0]
sec.page_width = Cm(21)
sec.page_height = Cm(29.7)
sec.top_margin = Cm(2.4)
sec.bottom_margin = Cm(2.2)
sec.left_margin = Cm(2.5)
sec.right_margin = Cm(2.5)

st = doc.styles["Normal"]
st.font.name = K
st.font.size = Pt(11)
st.element.rPr.rFonts.set(qn("w:eastAsia"), K)


# ---- 封面 ----
para(doc, "济南市2025年政府工作报告", size=24, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=60, space_after=4)
para(doc, "深度研究与\u201c容易被忽视的细节\u201d分析", size=16, bold=True, color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
para(doc, "从\u201c新旧动能转换、经济强省省会、装备制造、电子信息与济南都市圈\u201d重新理解济南", size=12, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)
para(doc, "\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501", color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
para(doc, "研究对象：2025年济南市政府工作报告及同期财政、国民经济和社会发展统计资料、2026年官方复盘", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
para(doc, "标定日期：2026-08-13", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
doc.add_page_break()

# ---- 目录 ----
catalog = [
    "执行摘要",
    "一、研究口径与阅读方法",
    "二、先看济南的特殊底盘：新旧动能转换起步区、装备制造、电子信息与经济强省会",
    "三、最关键的宏观错位：GDP破1.4万亿但略低于目标，装备/出口强，投资转负",
    "四、容易被忽视、但信号很重的细节（15条）",
    "五、2025年GDP目标 vs 实际：对照表",
    "六、2025年增长是谁撑起来的？（结构归因）",
    "七、预算与财政的\u201c含金量\u201d",
    "八、民生底账：人口、收入与城乡",
    "九、城镇与农村：格局与均衡",
    "十、人口流入与流出",
    "十一、物价与货币环境",
    "十二、区域一体化：济南在\u201c山东省会+济南都市圈+黄河流域\u201d里的位置",
    "十三、未来5—10年最值得观察的五条主线",
    "十四、最终结论：济南在\u201c新旧动能+高新技术+装备制造\u201d里的增长逻辑",
    "附录A：主要资料来源与核验路径",
    "附录B：建议建立的年度跟踪仪表盘",
]
para(doc, "目　录", size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=10, space_after=10)
for item in catalog:
    para(doc, item, size=12, space_after=4)
doc.add_page_break()

# ---- 执行摘要 ----
heading1(doc, "执行摘要")
para(doc, "**核心判断**　2025年济南最显著的是\u201cGDP 14210亿元、增长5.4%（略低于5.5%目标）、三产占64.5%\u201d、\u201c规上工业+6.9%（装备+16.8%、电子+37.8%）、出口+24.2%\u201d、\u201c进出口2898亿/+24.7%\u201d、\u201c常住人口961.6万/城镇化77.3%\u201d。这说明济南在\u201c新旧动能转换+先进制造+开放\u201d下稳中提质，但**投资-12.9%、财政+1.0%偏低**是短板。")
para(doc, "把2025年目标（GDP+5.5%/规上+6.5%/财政+3%/社零+3%）、2025年统计、2026年前瞻一起看，济南是\u201c经济强省省会+新旧动能转换\u201d路径：**装备制造、电子信息、比亚迪整车、软件/数字经济、出口**是引擎。总量14210亿居山东第2、全国前20强。")
para(doc, "最容易记住的一句话：**济南是\u201c新旧动能转换+装备制造+数字经济\u201d的山东省会，靠\u201c先进制造（电子/汽车/装备）+软件/数字经济+外贸进出口+高端文旅\u201d增长。**观察济南，与其只看\u201cGDP 1.42万亿\u201d，不如看\u201c电子信息+37.8%、装备制造+16.8%、比亚迪整车36万辆、进出口+24.7%、软件业务6300亿\u201d。")# ---- 一、研究口径与阅读方法 ----
heading1(doc, "一、研究口径与阅读方法")
para(doc, "本报告以\u201c济南市2025年政府工作报告（2025年1月，于海田作）\u201d为起点，把\u201c2025年GDP目标（5.5%左右）\u201d与\u201c2025年国民经济和社会发展统计公报实际值（14210亿元/+5.4%）\u201d并置对照，再用2026年政府工作报告的复盘作事后验证。")
para(doc, "重点阅读三类证据：**一是指标目标是否达标**；**二是结构（谁贡献增长）是否匹配目标叙事**；**三是官方复盘如何自评、承认问题**。统计口径一般公共预算收入为地方级。")
para(doc, "统一采用\u201c同比增长%\u201d（GDP为不变价），财政与居民收入多为名义值，文中按原口径转录、不逐项换算。")

# ---- 二、特殊底盘 ----
heading1(doc, "二、先看济南的特殊底盘：新旧动能转换起步区、装备制造、电子信息与经济强省会")
para(doc, "**区位与身份**：济南是山东省会、副省级城市，\u201c经济强省省会\u201d、综合性国家科学中心承载地，济南新旧动能转换起步区、济南都市圈（获批）、黄河流域中心城市。")
para(doc, "**产业底盘**：一是装备制造（+16.8%）；二是电子信息/计算机（+37.8%）；三是汽车/比亚迪整车（新能源汽车36万辆、汽车制造+19.8%）；四是软件/数字经济（软件业务收入6300亿、算力4500-6451P）；五是医药健康（现代医药）；六是钢铁/先进材料。")
para(doc, "**人口底盘**：2025年末常住人口961.6万/+1.1%（+10.1万）、城镇化率77.3%（+1.1pct）；人才资源总量310万。")
para(doc, "**市场与出口**：2025年社零5547.9亿/+4.4%；进出口2898亿/+24.7%、出口+24.2%、进口+25.6%。开放快速增长是2025最大亮点。")

# ---- 三、核心宏观错位 ----
heading1(doc, "三、最关键的宏观错位：GDP破1.4万亿但略低于目标，装备/出口强，投资转负")
para(doc, "**第一组错位**：2025年GDP目标5.5%左右，实际14210亿/+5.4%**略低于目标0.1pct**，但\u201c十四五\u201d连跨4个千亿台阶、稳居全国前20强，人均GDP破2万美元。")
para(doc, "**第二组错位**：规上工业+6.9%、装备+16.8%、出口+24.2%、进出口+24.7%强劲；但**固投-12.9%（房地产-10.9%）**，\u201c工业、出口强、投资弱\u201d。")
para(doc, "**第三组错位**：人口净流入10.1万/城镇化77.3%但**自然增长率-0.2\u2030**；财政+1.0%偏低、低于5.4%GDP名义。")
para(doc, "一句话：**济南是\u201c先进制造+出口+数字强、投资与财政弱\u201d的山东省会**——靠新旧动能转换+外向经济支撑。")

# ---- 四、15条细节 ----
heading1(doc, "四、容易被忽视、但信号很重的细节（15条）")
bullet(doc, "1. **GDP破1.4万亿**：14210亿/+5.4%、全国前20强、人均破2万美元。")
bullet(doc, "2. **电子信息+37.8%**：占规上比重大、是新旧动能转换最强信号。")
bullet(doc, "3. **装备制造+16.8%**：快于规上9.9pct，先进制造高景气。")
bullet(doc, "4. **汽车制造+19.8%、比亚迪36万辆**：新能源整车放量。")
bullet(doc, "5. **出口+24.2%、进出口+24.7%**：外贸翻番后继续高增、外向度升至20.4%。")
bullet(doc, "6. **软件业务6300亿、算力4500P+**：数字济南、人工智能强省。")
bullet(doc, "7. **固投-12.9%**：房地产-10.9%、三产-14.4%，投资塌陷。")
bullet(doc, "8. **高新技术/规上工业利润542.1亿/+23.9%**：工业盈利强。")
bullet(doc, "9. **中欧班列1252列/+21.7%**：陆上通道。")
bullet(doc, "10. **常住961.6万/城镇化77.3%**：人口净流入10.1万、但自然-0.2‰。")
bullet(doc, "11. **收入：城镇68128元/+4.2%、农村28653元/+5.4%**，城乡2.38。")
bullet(doc, "12. **CPI+0.2%**：温和、微通缩边缘。")
bullet(doc, "13. **财政1093.4亿/+1.0%（税收+3.7%）**：税收强、总盘低。")
bullet(doc, "14. **人才310万、万人发明专利87.2件**：科创人才集聚。")
bullet(doc, "15. **中欧班列/国际友城93个**：开放枢纽。")# ---- 五、目标 vs 实际对照表 ----
heading1(doc, "五、2025年GDP目标 vs 实际：对照表")
table(doc,
    ["指标", "2025年目标", "2025年实际", "是否达标"],
    [
        ["地区生产总值增速", "5.5%左右", "14210亿元/+5.4%", "略低于目标"],
        ["规上工业增加值", "+6.5%左右", "+6.9%", "超目标"],
        ["固定资产投资", "合理增长", "-12.9%（转负）", "未达标"],
        ["社会消费品零售总额", "+3%左右", "5547.9亿元/+4.4%", "超目标"],
        ["一般公共预算收入", "+3%左右", "1093.4亿元/+1.0%", "未达标"],
        ["城镇/农村人均可支配收入", "+5.5%/+6%左右", "+4.2%/+5.4%", "均低于目标"],
        ["进出口", "高于全省", "2898亿/+24.7%", "高增超目标"],
    ],
    widths=[3.4, 2.5, 5.0, 3.1])
para(doc, "**简评**：处于\u201c工业、消费、出口达标、投资/财政/居民收入未达目标\u201d格局。最亮是**进出口+24.7%、装备+16.8%**；最弱是**投资-12.9%、财政+1.0%**。")

# ---- 六、结构归因 ----
heading1(doc, "六、2025年增长是谁撑起来的？（结构归因）")
para(doc, "**三大产业**：一产442.2亿/+3.6%、二产4605.2亿/+4.4%、三产9162.6亿/+6.0%；三产占GDP 64.5%。增长以三产为主、二产（工业）次之。")
para(doc, "**工业**：规上工业+6.9%；**计算机/电子+37.8%、汽车+19.8%、装备制造+16.8%、通用设备+7.3%**；高技术制造+14.7%。规上工业营收10786.5亿/+5.4%、利润542.1亿/+23.9%。")
para(doc, "**服务业**：金融、软件/数字经济、文旅拉动三产+6.0%；软件业务收入6300亿。")
para(doc, "**增长归因**：济南增长主要靠**三产（软件/数字/金融/文旅）+工业（电子/汽车/装备）+出口**；投资（-12.9%）、地产拖累。")

# ---- 七、财政 ----
heading1(doc, "七、预算与财政的\u201c含金量\u201d")
para(doc, "2025年一般公共预算收入1093.4亿元/+1.0%（税收812.8亿/+3.7%、占比74.3%）；一般公共预算支出1407.5亿/+0.7%，民生支出1127.6亿占80.1%。")
para(doc, "**结构性**：收入\u201c低增长\u201d（总盘+1.0%）但税收强（+3.7%）、质量高（税收占比74.3%全国前列）；支出+0.7%保民生。")
para(doc, "**含金量**：济南财政\u201c税收占比高、质量好\u201d但总量增速低，且投资靠政策资金（650亿）与市场，自主财政扩张有限。")

# ---------------- 八、民生 ----------------
heading1(doc, "八、民生底账：人口、收入与城乡")
para(doc, "常住人口961.6万/+1.1%（+10.1万）、城镇化率77.3%（+1.1pct），自然人小幅负增长-0.2\u2030。收入：**城镇68128元/+4.2%、农村28653元/+5.4%**，城乡比2.38。")
para(doc, "就业：城镇新增就业约21万、民生支出占80.1%。")
para(doc, "**民生结论**：收入农村快于城镇、差距略扩（2.38）但可控；人口净流入10万、城镇化提升；就业、社保、医保覆盖广。")

# ---------------- 九 ----------------
heading1(doc, "九、城镇与农村：格局与均衡")
para(doc, "济南城镇化率77.3%（全国前列），主城区（历下/市中/高新/起步区）承载数字/金融/装备；县域（章丘/莱芜/平阴/商河）发展双产/粮食/文旅。")
para(doc, "城乡收入比2.38（略高于全国平均），农村收入增速+5.4%快于城镇+4.2%，差距边际收敛。")

# =========== 十 ============
heading1(doc, "十、人口流入与流出")
para(doc, "济南常住961.6万、净流入10.1万，但自然-0.2\u2030，增长靠机械流入（高校、产业、人才）。人才资源310万、万人发明专利87件。")
para(doc, "**流入**：大学生/高端制造/数字经济人才；**流出**：部分中低技能劳动力。总体\u201c高素质净流入\u201d、人口质量优于数量。")

# =========== 十一 ============
heading1(doc, "十一、物价与货币环境")
para(doc, "2025年济南CPI**上涨0.2%**（温和），\u201c五涨三落\u201d（交通通信-2.6%、居住-0.8%、食品-1.4%）。")
para(doc, "金融：存款31958亿/+6.4%、贷款34513.2亿/+9.2%。\u201c宽货币、温和需求\u201d。")

# =========== 十二 ============
heading1(doc, "十二、区域一体化：济南在\u201c山东省会+济南都市圈+黄河流域\u201d里的位置")
para(doc, "济南是山东省会、济南都市圈核心（规划获批）、山东省新旧动能转换核心区、黄河流域国家中心城市。作为\u201c经济强省\u201d省会，承接齐鲁大区的集聚与辐射。")
para(doc, "开放：中欧班列、国际友城93个、外贸+24.7%、新建起步区/未来产业城。济南在\u201c强省会+黄河流域+都市圈+开放\u201d中扮演枢纽。")

# ============ 十三 =============
heading1(doc, "十三、未来5—10年最值得观察的五条主线")
bullet(doc, "1. **新旧动能转换起步区**：比亚迪/宾理高端新能源、中新未来产业城，\u201c起步区\u201d能否成第二增长极。")
bullet(doc, "2. **电子信息+集成电路**：电子+37.8%、集成电路目标300亿，制造强省主升。")
bullet(doc, "3. **软件/数字经济+AI**：软件6300亿、算力向5000P爬坡，“数字济南”。")
bullet(doc, "4. **开放+外贸**：进出口+24.7%、中欧班列，外向度20.4%提升。")
bullet(doc, "5. **投资修复+人口**：固投-12.9%后，靠都市圈、项目赋能修复；人口净流入与人才。")

# ============ 十四 =============
heading1(doc, "十四、最终结论：济南在\u201c新旧动能+高新技术+装备制造\u201d里的增长逻辑")
para(doc, "**结论**：济南2025年的\u201c真相\u201d是——**GDP+5.4%（略低于目标）、规上+6.9%、装备+16.8%、进出口+24.7%、财政+1.0%**。它是\u201c先进制造（电子/汽车/装备）+数字经济+开放\u201d驱动但**投资弱、财政缓**的山东省会。")
para(doc, "**对趋势判断**：装备制造/电子信息/软件/出口代表济南的\u201c新动能\u201d，投资/财政代表\u201c约束\u201d。**新旧动能转换+先进制造+外向经济**决定中后期潜力，**投资修复+人口/消费**决定韧性。")
para(doc, "**若只看一个指标**：看**电子信息制造业增速（+37.8%）+固定资产投资增速（-12.9%）**——济南是\u201c先进制造强、但投资缺口大\u201d的新旧动能转换枢纽，能否把产业亮点转化为投资与人口是关键。")

# ------------- 附录A -------------
heading1(doc, "附录A：主要资料来源与核验路径")
bullet(doc, "济南市人民政府《2025年济南市政府工作报告》（2025年1月，于海田）。")
bullet(doc, "济南市统计局《2025年济南市国民经济和社会发展统计公报》（2026年3月）。")
bullet(doc, "《2026年济南市政府工作报告》（2026年1月）及极简版/一图读懂。")
bullet(doc, "山东省2025年统计公报、济南市经济运行数据交叉核验。")

# ------------- 附录B -------------
heading1(doc, "附录B：建议建立的年度跟踪仪表盘")
bullet(doc, "GDP总量/增速 vs 全年目标（+/-个百分点）。")
bullet(doc, "规上工业增加值及分行业（电子/装备/汽车/医药/钢铁）增速。")
bullet(doc, "固定资产投资（总量/工业/民间/房地产）增速。")
bullet(doc, "社会消费品零售总额、新能源汽车、通讯器材。")
bullet(doc, "货物进出口（人民币）、出口、中欧班列。")
bullet(doc, "一般公共预算收入/税收/非税、财政支出、民生占比、税收占比。")
bullet(doc, "常住人口、自然增长率、城镇化率、人才总量。")
bullet(doc, "CPI/核心CPI、规上工业利润与营收。")
bullet(doc, "软件业务收入、算力（P）、AI大模型。")
bullet(doc, "比亚迪整车/新能源汽车产量、起步区投资。")

# ------------- 保存 -------------
out = "/Users/x/Desktop/content-prod-lab/reports/济南市_2025年政府工作报告_深度研究_2026-08-13.docx"
doc.save(out)
print("SAVED OK 济南市_2025年政府工作报告_深度研究_2026-08-13.docx")