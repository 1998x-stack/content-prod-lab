# -*- coding: utf-8 -*-
"""Build 济宁市2025年政府工作报告 深度研究 DOCX, 参照地级市系列版式。"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DARK = RGBColor(0x1F, 0x2A, 0x44)
RED = RGBColor(0xB0, 0x1E, 0x2E)
GRAY = RGBColor(0x59, 0x60, 0x69)
BLUE = RGBColor(0x1F, 0x4E, 0x79)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
HEADER_FILL = "1F2A44"
K = "微软雅黑"


def set_run(run, size=11, bold=False, color=DARK, font=K):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    run.font.name = font
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), font)
    rFonts.set(qn("w:hAnsi"), font)
    rFonts.set(qn("w:eastAsia"), font)


def para(doc, text, size=11, bold=False, color=DARK, align=None,
         space_after=6, space_before=0, font=K):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        set_run(r, size=size, bold=bold or (i % 2 == 1), color=color, font=font)
    return p


def heading1(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(16)
    pf.space_after = Pt(8)
    r = p.add_run(text)
    set_run(r, size=16, bold=True, color=RED)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "B01E2E")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def heading2(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(10)
    pf.space_after = Pt(4)
    r = p.add_run(text)
    set_run(r, size=13, bold=True, color=BLUE)
    return p


def table(doc, headers, rows, widths=None, font_size=9.5):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_run(r, size=font_size, bold=True, color=WHITE)
        tcPr = hdr[i]._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), HEADER_FILL)
        tcPr.append(shd)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(val))
            set_run(r, size=font_size, color=DARK)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    return t


def quote_box(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(8)
    pf.left_indent = Pt(12)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), "B01E2E")
    pBdr.append(left)
    pPr.append(pBdr)
    r = p.add_run(text)
    set_run(r, size=11, color=DARK, font="楷体")
    return p


def bullet(doc, text, size=10.5):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.7)
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        set_run(r, size=size, bold=(i % 2 == 1), color=DARK)
    return p


# ================================================================= 文档主体
doc = Document()
sec = doc.sections[0]
sec.page_width = Cm(21)
sec.page_height = Cm(29.7)
sec.top_margin = Cm(2.4)
sec.bottom_margin = Cm(2.2)
sec.left_margin = Cm(2.5)
sec.right_margin = Cm(2.5)

st = doc.styles["Normal"]
st.font.name = K
st.font.size = Pt(11)
st.element.rPr.rFonts.set(qn("w:eastAsia"), K)


# ---- 封面 ----
para(doc, "鄂尔多斯市2025年政府工作报告", size=24, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=60, space_after=4)
para(doc, "深度研究与\u201c容易被忽视的细节\u201d分析", size=16, bold=True, color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
para(doc, "从\u201c中国煤都、新能源之都、羊绒之都、现代煤化工、草原旅游\u201d重新理解鄂尔多斯", size=12, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)
para(doc, "\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501", color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
para(doc, "研究对象：2025年鄂尔多斯市政府工作报告及同期财政、国民经济和社会发展统计资料、2026年官方复盘", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
para(doc, "标定日期：2026-08-14", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
doc.add_page_break()

# ---- 目录 ----
catalog = [
    "执行摘要",
    "一、研究口径与阅读方法",
    "二、先看鄂尔多斯的特点底盘：煤都、新能源、绒都、现代煤化工、草原",
    "三、最关键的宏观错位：工业出口强（煤炭/新能源），但固投降速、地产销售-25.7%、三产弱",
    "四、容易被忽视、但信号很重的细节（15条）",
    "五、2025年GDP目标 vs 实际：对照表",
    "六、2025年增长是谁撑起来的？（结构归因）",
    "七、预算与财政的\u201c含金量\u201d",
    "八、民生底账：人口、收入与城乡",
    "九、城镇与农村：格局与均衡",
    "十、人口流入与流出",
    "十一、物价与货币环境",
    "十二、区域一体化：鄂尔多斯在呼包鄂榆、黄河流域、蒙晋陕\u201c三极\u201d里的位置",
    "十三、未来5—10年最值得观察的五条主线",
    "十四、最终结论：鄂尔多斯在\u201c煤炭+新能源+绒\u201d里的增长逻辑",
    "附录A：主要资料来源与核验路径",
    "附录B：建议建立的年度跟踪仪表盘",
]
para(doc, "目　录", size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=10, space_after=10)
for item in catalog:
    para(doc, item, size=12, space_after=4)
doc.add_page_break()

# ---- 执行摘要 ----
heading1(doc, "执行摘要")
para(doc, "**核心判断**　2025年鄂尔多斯最显著的是\u201cGDP 6122.2亿/+5.1%（煤炭8.9亿吨、全国1/5保供）、规上工业+7%、高技术制造+29.7%、新能源装机2779万千瓦（全区第1）、现代煤化工入选国家级集群\u201d、\u201c但一般公共预算收入-12.4%、商品房销售-25.7%、三产+2.1%、固投+1.7%\u201d。这说明鄂尔多斯在\u201c煤+新能源\u201d双引擎中，**工业制造强、财政/地产/三产弱、能源价格承压**。")
para(doc, "把2025年目标（GDP+6%、规上+6.5%、固投+10%、社零+5%、财收+2%）、2025年实际（GDP+5.1%、规上+7%、固投+1.7%、社零+4.4%、财收-12.4%）趋势看，鄂尔多斯是\u201c煤+风光氢储车\u201d路径：**煤炭（保供）、现代煤化工、新能源（风光氢储车）、硅/光伏、智能制造、羊绒**是支柱。")
para(doc, "最容易记住的一句话：**鄂尔多斯是\u201c中国煤都、能源大市、羊绒之乡\u201d，靠\u201c煤保供+新能源+煤化工\u201d驱动。**观察鄂尔多斯，与其只看\u201cGDP 6122亿\u201d，不如看\u201c煤炭8.9亿吨（全国1/5）、新能源装机2779万千瓦（全区第1）、高技术制造+29.7%、现代煤化工集群、人均GDP 27.3万\u201d。")
# ---- 一、研究口径与阅读方法 ----
heading1(doc, "一、研究口径与阅读方法")
para(doc, "本报告以\u201c鄂尔多斯市政府工作报告（2025年1月）\u201d为起点，把\u201c2025年GDP目标（6%）\u201d与\u201c官方2025年（6122.2亿/+5.1%）\u201d并置对照，用\u201c2025年鄂尔多斯市统计公报\u201d和\u201c2026年政府工作报告\u201d横向核验。")
para(doc, "口径提示：\u201cGDP、规上工业增速\u201d按可比价（实际）；\u201c投资、消费、进出口、财政\u201d按现价（名义）。\u201c常住人口\u201d用统计公报（224.5万）、城镇化率81.10%。")
para(doc, "指标体系（与研究口径一致）：**总量与增速、产业动能（煤/新能源/化工）、外贸、财政、民生与人口**。")
para(doc, "特别提示（不吃老本）：鄂尔多斯2024年GDP 6363亿/+6.4%、2025年6122.2亿/+5.1%（含五经普修订基）；它不是\u201c只有煤\u201d——**现代煤化工、新能源（风光氢储车）、光伏/硅、智能制造、羊绒、旅游\u201d才是真正底色；人均GDP 27.3万（全国前列）。")
# ---- 二、先看鄂尔多斯的特别底盘 ----
heading1(doc, "二、先看鄂尔多斯的特别底盘：煤、新能源、绒都、现代煤化工、草原旅游")
para(doc, "鄂尔多斯地处内蒙古中部、鄂尔多斯高原、黄河南岸，是**中国煤都（全国1/5保供）、国家级能源化工基地（现代煤化工入选国家级集群）、新能源之都（风光氢储车）、中国绒都（鄂尔多斯羊绒）、康巴什新城**。2025年GDP 6122.2亿（内蒙古第1、占全区1/3）、常住224.5万、城镇化率81.10%、人均GDP 27.3万（全国前列）。")
para(doc, "五个底盘名词，先立框架：")
bullet(doc, "**煤炭（保供）**　煤炭8.9亿吨/年（全国1/5保供）、煤价承压但量稳——\u201c中国煤都\u201d。")
bullet(doc, "**现代煤化工**　煤制甲醇、烯烃、油、化（国家先进制造业集群）——\u201c煤化工\u201d。")
bullet(doc, "**新能源（风光氢储车）**　新能源装机2779万（全区第1）、光伏/风电/氢能/储能/新能源车、光伏治沙——\u201c新能源之都\u201d。")
bullet(doc, "**羊绒/智能制造**　鄂尔多斯（Erdos）羊绒（中国绒都）、智能制造（专精特新）、硅/电池——\u201c绒+制造\u201d。")
bullet(doc, "**草原/沙漠旅游**　库布其沙漠/草原、康巴什、响沙湾（5A）——\u201c草原旅游\u201d。")
para(doc, "这五根（煤+煤化工+新能源+绒+旅游）构成鄂尔多斯独特底盘：**左手煤（保供），右手新能源+绒**。理解鄂尔多斯，先理解\u201c煤、新能源、绒\u201d。")
# ---- 三、最关键的宏观错位 ----
heading1(doc, "三、最关键的宏观错位：工业强（煤+新能源）、高技术+29.7%，但固投降速、地产销售-25.7%、三产/财政弱")
para(doc, "2025年鄂尔多斯最需要辨析的一组\u201c错位\u201d：**规上工业+7%、高技术制造+29.7%、新能源装备+24.2%、煤炭保供（8.9亿吨）、现代煤化工强，但固定资产投资+1.7%（较23.3%大幅降速）、商品房销售-25.7%、第三产业+2.1%、一般公共预算收入-12.4%**。")
para(doc, "为什么\u201c二产/工业\u201d强，投资/三产/财政却弱？三解释：")
para(doc, "**其一，二产/能源/制造强、点大**　二产+6.7%（占68.1%、贡献82.4%）、规上+7%（制造+23.9%）、高技术+29.7%、煤化工/单晶硅——\u201c工业强\u201d。")
para(doc, "**其二，投资/地产/三产/财政弱**　固投+1.7%（较2024的23.3%降速）、地产销售-25.7%、三产+2.1%、财收-12.4%（煤炭价格/增值税）——\u201c需求财政弱\u201d。")
para(doc, "**其三，消费/物价温、旅游强**　社零+4.4%、旅游总收入614亿+10.7%、CPI-0.1%——\u201c消费稳、物价低\u201d。")
para(doc, "小结：鄂尔多斯2025年是\u201c**工业制造强、投资三产财政弱**\u201d：煤/新能源/化工/绒强，地产、财政、三产慢。")

# ---- 四、容易被忽视、但信号很重的细节（15条） ----
heading1(doc, "四、容易被忽视、但信号很重的细节（15条）")
bullet(doc, "**1.煤炭8.9亿吨（全国1/5保供）**\u201c能源压舱石。\u201d")
bullet(doc, "**2.新能源装机2779万千瓦（全区第1）+655万/新能源发电257亿千瓦时+74%**\u201c新能源大爆发。\u201d")
bullet(doc, "**3.规上工业+7%、制造+23.9%（石油化工+15.6%）**\u201c工业制造强。\u201d")
bullet(doc, "**4.高技术制造+29.7%、单晶硅+147%、聚乙烯+125.6%、锂电+47.6%**\u201c新质制造。\u201d")
bullet(doc, "**5.现代煤化工入选国家级制造业集群、新增产能300万吨**\u201c鄂尔多斯煤炭新。\u201d")
bullet(doc, "**6.羊绒（中国绒都）、智能制造/专精特新**\u201c绒+智造。\u201d")
bullet(doc, "**7.一般公共预算收入803亿/-12.4%（煤炭价减）**\u201c能源依赖、财政承压。\u201d")
bullet(doc, "**8.固定资产投资+1.7%（亿元以上项目435个）**\u201c项目/投资稳中降。\u201d")
bullet(doc, "**9.商品房销售面积-22.7%**\u201c地产深度调整。\u201d")
bullet(doc, "**10.第三产业+2.1%（交通+5.7%、地产-1.8%）**\u201c三产慢。\u201d")
bullet(doc, "**11.社零698亿/+4.4%（新能源汽车+21.5%）**\u201c消费/新能源车旺。\u201d")
bullet(doc, "**12.旅游4326万人次/614亿（+10.1%/+10.7%）**\u201c草原旅游强。\u201d")
bullet(doc, "**13.居民收入55914元/+4.5%、城乡比1.79（缩）**\u201c农牧民+5.8%>城镇+3.8%。\u201d")
bullet(doc, "**14.CPI-0.1%、存贷款+2.7%/+2.4%**\u201c低通胀、信贷稳。\u201d")
bullet(doc, "**15.常住224.5万/+0.45万、城镇化81.10%**\u201c人口稳增、城镇化高。\u201d")
# ---- 五、2025年GDP目标 vs 实际：对照表 ----
heading1(doc, "五、2025年GDP目标 vs 实际：对照表")
table(doc,
    ["指标", "2025年目标", "2025年实际", "达成评价"],
    [
        ["地区生产总值(GDP)", "增长6%左右", "6122.2亿/5.1%", "差0.9pct"],
        ["规模以上工业", "增长6.5%以上", "+7%", "超额"],
        ["固定资产投资", "增长10%以上", "+1.7%", "大幅不及"],
        ["社会消费品零售总额", "增长5%以上", "698.0亿/+4.4%", "略低"],
        ["一般公共预算收入", "同口径增长2%左右", "803.0亿/-12.4%", "大降"],
        ["居民收入", "与增长同步", "55914元/+4.5%", "略低"],
    ],
)
para(doc, "注：GDP、规上工业按可比价。**规上（+7%）超额**；**GDP（+5.1%）、固投（+1.7%）、社零（+4.4%）略低/不及**；**财收（-12.4%）降幅大（煤炭价格）**。")
para(doc, "拆读：**煤（保供）、新能源、高技术制造、化工是亮色**；**财收（-12.4%）、地产（-25.7%）、固投（+1.7%）、三产（+2.1%）**是短板——\u201c工业强、能源价压、三产财政弱\u201d，是\u201c煤+新能源\u201d样本。")

# ---- 六、2025年增长是谁撑起来的？（结构归因） ----
heading1(doc, "六、2025年增长是谁撑起来的？（结构归因）")
para(doc, "把鄂尔多斯GDP的5.1%拆开：三次产业分别提5.3%、6.7%、2.1%（结构3.7：68.1：28.2；贡献一3.8%、二82.4%、三13.8%）。**第二产业（工业）+6.7%、贡献82.4%是绝对引擎**，第一产业稳、第三产业（+2.1%）慢。")
para(doc, "2026年鄂尔多斯强调\u201c四个世界级产业（煤化工/新能源/绒/智能制造）、能源保供与转化\u201d，聚焦**煤炭保供、现代煤化工、风光氢储车、绒纺、氢能（国家示范）、智能制造**——核心是\u201c煤稳+新能源\u201d。")
para(doc, "**第二产业（工业/能源）**：规上+7%（煤开采+5.7%、化工+39.7%、油气加工+15.6%、制造+23.9%）、煤炭8.9亿吨——\u201c煤+化工+新能源\u201d强。")
para(doc, "**第三产业（服务业）**：+2.1%（交通物流+5.7%、旅游、金融+0.6%；地产-1.8%）——\u201c服务/旅游温\u201d。")
para(doc, "**第一产业（农牧业）**：+5.3%（农林牧渔377亿、粮食232万吨、肉/绒）——\u201c农牧稳\u201d。")
para(doc, "一句话归因：**2025年鄂尔多斯增长\u201c靠第二产业（煤/化工/新能源/制造）\u201d**，三产/财政/地产弱；\u201c能源+新质\u201d是核心。")

# ---- 七、预算与财政的\u201c含金量\u201d ----
heading1(doc, "七、预算与财政的\u201c含金量\u201d")
para(doc, "2025年鄂尔多斯**一般公共预算收入803.0亿元（-12.4%）**；税收575.9亿；支出1001.7亿（-12.3%）、民生占71.7%（+4.6pct）。")
bullet(doc, "财收-12.4%（煤炭价格/增值税、资源税）、税收占比约72%——\u201c能源依赖、含金量靠煤\u201d。")
bullet(doc, "民生支出占71.7%（+4.6pct，教育/社保/医保）。")
bullet(doc, "金融：存款+2.7%、贷款+2.4%——宽信用稳、对煤化工/新能源支持。")
para(doc, "**财政含金量小结**：财收-12.4%（能源价拖累）、煤炭税为主、民生71.7%；财政对\u201c新能源、煤化工、绒纺、民生\u201d投入加大。")

# ---- 八、民生底账：人口、收入与城乡 ----
heading1(doc, "八、民生底账：人口、收入与城乡")
para(doc, "2025年鄂尔多斯**居民人均可支配收入55914元（+4.5%）**，其中城镇63946元（+3.8%）、农牧民30478元（+5.8%），城乡比2.1→1.79（缩小）。就业：城镇新增就业2.7万人。")
para(doc, "人口画像：**常住224.5万/+0.45万、城镇化率81.10%**；内蒙古人口净增/煤新能源吸引、民族地区。")
para(doc, "民生投入：财收民生71.7%、养老/医保、低保提高（城镇12240/农村9960）——民生保障扎实、城市化城区康巴什。")

# ---- 九、城镇与农村：格局与均衡 ----
heading1(doc, "九、城镇与农村：格局与均衡")
para(doc, "鄂尔多斯城镇化率81.10%（全区最高）；市区（东胜/康巴什/伊金霍洛）+旗县（达拉特羊绒/乌审天然气）；农牧民收入增速（+5.8%）>城镇（+3.8%），**城乡比缩小至1.79**。")
para(doc, "农牧业底盘：**粮食232万吨、羊肉、绒毛（阿尔巴斯白羊绒）、牧区**——\u201c绒/草原牧业\u201d。")
para(doc, "一句话：\u201c鄂尔多斯高城镇化、城区康巴+牧区绒、城乡比收窄\u201d。")

# ---- 十、人口流入与流出 ----
heading1(doc, "十、人口流入与流出")
para(doc, "鄂尔多斯常住224.5万（+0.45万）、城镇化81.1%；\u201c煤/新能源/制造业\u201d吸引（能源城康巴什/高新区）、十四五新增常住8.5万。")
para(doc, "结构观察：**人口净流入（能源经济/大基地）、自然增长-1.25‰（老龄化）**；能源/绒纺吸附就业。")
para(doc, "2026年目标：新增就业2万+、人才——鄂尔多斯靠\u201c能源+新能源+绒\u201d聚人。")

# ---- 十一、物价与货币环境 ----
heading1(doc, "十一、物价与货币环境")
para(doc, "2025年鄂尔多斯**CPI-0.1%**（食品烟酒-0.7%、交通通信-3.3%、居住-0.1%；衣着+0.8%、其他用品+11.8%）——\u201c低通胀、需求温\u201d。")
bullet(doc, "信贷扩张：存贷款+2.7%/+2.4%（制造业/新能源信贷）——宽信用稳。")
bullet(doc, "消费：社零+4.4%、新能源车+21.5%、通讯+46.5%——消费/新贵。")
para(doc, "货币环境判断：**宽信用、CPI-0.1%**；鄂尔多斯靠\u201c煤+新能源+旅游\u201d稳需求（2026 CPI平稳）。")

# ---- 十二、区域一体化：鄂尔多斯的位置 ----
heading1(doc, "十二、区域一体化：鄂尔多斯在呼包鄂榆、黄河流域、蒙晋陕\u201c三极\u201d里的位置")
para(doc, "鄂尔多斯是**呼包鄂榆城市群（呼和浩特-包头-鄂尔多斯-榆林）核心、黄河流域生态保护高质量发展带、蒙晋陕交界能源金三角**。")
bullet(doc, "**呼包鄂榆**　内蒙古呼包鄂城市群核心、能源经济带、鄂尔多斯（内蒙古第1）。")
bullet(doc, "**黄河流域**　黄河几字弯、能源保供、生态治理（光伏治沙/毛乌素）。")
bullet(doc, "**蒙晋陕能源金三角**　与榆林/忻州大同共筑国家能源基地、新通道。")
para(doc, "一句话：**鄂尔多斯在\u201c呼包鄂榆+黄河流域+能源金三角\u201d里，最核心是\u201c国家能源+绿色转型\u201d**；煤炭、新能源、区位是最大优势。")

# ---- 十三、未来5—10年最值得观察的五条主线 ----
heading1(doc, "十三、未来5—10年最值得观察的五条主线")
bullet(doc, "**主线一：现代煤化工（国家集群）**\u201c煤制油/烯烃/氢、化工\u201d能级。")
bullet(doc, "**主线二：新能源（风光氢储车）**\u201c装机2779万→800万+、氢能示范、储能\u201d。")
bullet(doc, "**主线三：光伏治沙/大基地/外送通道**\u201c库布其至上海/安徽/江苏3条外送通道、新能源并网\u201d。")
bullet(doc, "**主线四：绒纺/智能制造**\u201c鄂尔多斯绒、专精特新、硅\u201d。")
bullet(doc, "**主线五：财政/三产/人口韧性**\u201c能源价、地产、三产、聚人\u201d如何转型。")

# ---- 十四、最终结论 ----
heading1(doc, "十四、最终结论：鄂尔多斯在\u201c煤+新能源+绒\u201d里的增长逻辑")
para(doc, "把2025年闭环打穿：**鄂尔多斯是\u201c中国煤都、新能源之都、绒都\u201d**：GDP 6122.2亿/+5.1%、煤炭8.9亿吨（全国1/5）、新能源装机2779万（全区第1）、规上+7%、高技术+29.7%。")
para(doc, "鄂尔多斯不是\u201c只有煤\u201d——它是**现代煤化工+风光氢储新能源+绒纺+智能制造+生态旅游**的复合，靠\u201c二次产业\u201d驱动；但财收/三产/地产/能耗弱。")
para(doc, "一句话结论：**鄂尔多斯是\u201c煤都、新能源之城、绒之都\u201d；观察它先看\u201c煤炭保供、新能源装机、煤化工、高技术制造、绒\u201d，再看\u201c财政、三产、地产\u201d。**它是\u201c能源强、绿色转型中\u201d的内蒙样本。")

# ---- 附录A ----
heading1(doc, "附录A：主要资料来源与核验路径")
bullet(doc, "《2025年鄂尔多斯市政府工作报告》（2025年1月，2025年目标、2024年回顾+6.4%）")
bullet(doc, "《2025年鄂尔多斯市国民经济和社会发展统计公报》（鄂尔多斯市统计局，2026-04-17，2025年实际）")
bullet(doc, "《2026年鄂尔多斯市政府工作报告》（2026年1月，复盘+2026年目标）")
bullet(doc, "鄂尔多斯市人民政府/统计局（ordos.gov.cn）")

# ---- 附录B：建议建立的年度跟踪仪表盘 ----
heading1(doc, "附录B：建议建立的年度跟踪仪表盘")
bullet(doc, "GDP总量/增速 vs 全年目标。")
bullet(doc, "煤炭产量/价格、货物保供。")
bullet(doc, "新能源装机/发电/氢储车。")
bullet(doc, "现代煤化工/新产能。")
bullet(doc, "高技术/智能/绒纺。")
bullet(doc, "固定资产/工业/房地产投资。")
bullet(doc, "社零/旅游/CPI。")
bullet(doc, "财收/税收/民生。")
bullet(doc, "常住/城镇化/就业。")
bullet(doc, "能耗/固废/光伏治沙。")

# ---------------- 保存 ----------------
out = "/Users/x/Desktop/content-prod-lab/reports/鄂尔多斯市_2025年政府工作报告_深度研究_2026-08-14.docx"
doc.save(out)
print("SAVED OK: 鄂尔多斯市", out)
