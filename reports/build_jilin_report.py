# -*- coding: utf-8 -*-
"""Build 吉林省2025年政府工作报告 深度研究 DOCX, 参照省市系列版式。"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DARK = RGBColor(0x1F, 0x2A, 0x44)
RED = RGBColor(0xB0, 0x1E, 0x2E)
GRAY = RGBColor(0x59, 0x60, 0x69)
BLUE = RGBColor(0x1F, 0x4E, 0x79)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
HEADER_FILL = "1F2A44"
K = "微软雅黑"


def set_run(run, size=11, bold=False, color=DARK, font=K):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    run.font.name = font
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), font)
    rFonts.set(qn("w:hAnsi"), font)
    rFonts.set(qn("w:eastAsia"), font)


def para(doc, text, size=11, bold=False, color=DARK, align=None,
         space_after=6, space_before=0, font=K):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        set_run(r, size=size, bold=bold or (i % 2 == 1), color=color, font=font)
    return p


def heading1(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(16)
    pf.space_after = Pt(8)
    r = p.add_run(text)
    set_run(r, size=16, bold=True, color=RED)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "B01E2E")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def heading2(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(10)
    pf.space_after = Pt(4)
    r = p.add_run(text)
    set_run(r, size=13, bold=True, color=BLUE)
    return p


def table(doc, headers, rows, widths=None, font_size=9.5):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_run(r, size=font_size, bold=True, color=WHITE)
        tcPr = hdr[i]._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), HEADER_FILL)
        tcPr.append(shd)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(val))
            set_run(r, size=font_size, color=DARK)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    return t


def quote_box(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(8)
    pf.left_indent = Pt(12)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), "B01E2E")
    pBdr.append(left)
    pPr.append(pBdr)
    r = p.add_run(text)
    set_run(r, size=11, color=DARK, font="楷体")
    return p


def bullet(doc, text, size=10.5):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.7)
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        set_run(r, size=size, bold=(i % 2 == 1), color=DARK)
    return p


# ================================================================= 文档主体
doc = Document()
sec = doc.sections[0]
sec.page_width = Cm(21)
sec.page_height = Cm(29.7)
sec.top_margin = Cm(2.4)
sec.bottom_margin = Cm(2.2)
sec.left_margin = Cm(2.5)
sec.right_margin = Cm(2.5)

st = doc.styles["Normal"]
st.font.name = K
st.font.size = Pt(11)
st.element.rPr.rFonts.set(qn("w:eastAsia"), K)



# ---- 封面 ----
para(doc, "吉林省2025年政府工作报告", size=24, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=60, space_after=4)
para(doc, "深度研究与“容易被忽视的细节”分析", size=16, bold=True, color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
para(doc, "从“汽车、冰雪、人参医药、对俄开放与装备制造”重新理解吉林", size=12, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)
para(doc, "━━━━━━━━━━━━━━━━", color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
para(doc, "研究对象：2025年吉林省政府工作报告及同期财政、国民经济和社会发展统计资料、2026年官方复盘", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
para(doc, "标定日期：2026-08-13", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
doc.add_page_break()

# ---- 目录 ----
catalog = [
    "执行摘要",
    "一、研究口径与阅读方法",
    "二、先看吉林的特殊底盘：汽车、冰雪、人参医药、对俄开放与装备制造",
    "三、最关键的宏观错位：GDP近1.5万亿、制造/工业强，但投资/地产/人口弱",
    "四、容易被忽视、但信号很重的细节（15条）",
    "五、2025年GDP目标 vs 实际：对照表",
    "六、2025年增长是谁撑起来的？（结构归因）",
    "七、预算与财政的“含金量”",
    "八、民生底账：人口、收入与城乡",
    "九、城镇与农村：格局与均衡",
    "十、人口流入与流出",
    "十一、物价与货币环境",
    "十二、区域一体化：吉林在“对俄朝合作+中部城市群+东北振兴”里的位置",
    "十三、未来5—10年最值得观察的五条主线",
    "十四、最终结论：吉林在“汽车+冰雪+人参医药+对俄开放”里的增长逻辑",
    "附录A：主要资料来源与核验路径",
    "附录B：建议建立的年度跟踪仪表盘",
]
para(doc, "目　录", size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=10, space_after=10)
for item in catalog:
    para(doc, item, size=12, space_after=4)
doc.add_page_break()

# ---- 执行摘要 ----
heading1(doc, "执行摘要")
para(doc, "**核心判断**　如果只看新闻标题，2025年吉林最显眼的是“GDP近1.5万亿、增长5.0%（十四五收官）”、“规上工业+7.8%”、“新能源汽车产量+22.5%”、“游客4.68亿人次（+20.8%）”和“一般公共预算收入+13.3%”。但这份研究真正值得深读的，是这座“汽车/冰雪/人参医药/对俄开放”的东北省份，如何在固定资产投资（-13.1%）、房地产开发（-25.6%）与人口自然增长率（-6.37‰）的背景下，靠“汽车/装备/高技术制造+冰雪/文旅+财政收入”稳住增长。")
para(doc, "把2025年初设定的目标（GDP增长5.5%左右）、2025年《国民经济和社会发展统计公报》、以及2026年报告对2025年的复盘放在一起看，吉林呈现清晰暗线：**从“汽车/重工业”的旧底盘，向“汽车新能源化+装备/冰雪/人参医药+对俄开放”转型**。旧引擎（地产/基建/传统汽车）在调整；新引擎（新能源汽车、装备制造、冰雪文旅、人参医药、对俄贸易）被要求更快补位。")
para(doc, "因此，本报告不按政府工作报告逐段复述，而采用“显性表述—同期数据—制度含义—长期影响”的方式，专门提取容易被忽视、但对判断吉林未来5—10年发展模式更有价值的信号。")
para(doc, "最容易记住的一句话：**吉林是“汽车工业基地+冰雪旅游+人参医药+对俄开放”的东北样本，靠“汽车/装备+冰雪/旅游+粮食/人参”撑起增长。**观察吉林，与其看“GDP 1.5万亿”，不如看“汽车产量、新能源汽车、冰雪万亿产业、人参800亿、对俄/一带一路”这几张名片。")
heading2(doc, "一页速览：2025年吉林经济的“表与里”")
table(doc,
    ["维度", "表面现象", "更深层信号"],
    [
        ["增长", "GDP 14973.88亿、+5.0%", "一产11.2%、二产31.5%、三产57.3%"],
        ["产业", "规上工业+7.8%", "汽车/装备领先、新能源车+22.5%"],
        ["外贸", "进出口1638.63亿、-7.1%", "出口+11.8%、一带一路占62.1%"],
        ["投资", "固定资产投资-13.1%", "地产-25.6%、民间-7.2%"],
        ["财政", "一般公共预算收入1349.97亿、+13.3%", "税收+4.3%"],
        ["消费", "社零4396.67亿、+3.9%", "城镇+4.8%、餐饮-0.8%"],
        ["人口", "常住2296.69万、城镇化66.83%", "自然增-6.37‰、老龄化29.69%"],
        ["开放", "对俄对朝/一带一路", "冰雪/旅游/装备"],
    ],
    widths=[2.2, 5.4, 8.4])
para(doc, "", size=4, space_after=2)

# ---- 一、研究口径与阅读方法 ----
heading1(doc, "一、研究口径与阅读方法")
heading2(doc, "1.1 本报告的三份核心底稿")
bullet(doc, "**2025年《政府工作报告》**：2025年1月在省十四届人大三次会议上作，给出全年预期目标（GDP增长5.5%左右、固投5%左右、社零6%左右、一般公共预算收入增长3%左右）。")
bullet(doc, "**2025年《国民经济和社会发展统计公报》**：吉林省统计局2026年4月27日发布，给出全年实际执行数与增速，是“事后验证”的锚点。")
bullet(doc, "**2026年吉林省政府工作报告及官方复盘**：对2025执行结果的追认，交叉验证“目标—执行—再评价”三段式。")
heading2(doc, "1.2 阅读方法：显性—数据—制度—长期")
para(doc, "每一章按“**显性表述 → 同期数据 → 制度含义 → 长期影响**”四层展开。其中“制度含义”回答“政府为什么这么排优先序”，“长期影响”回答“这对未来5—10年意味着什么”。")
para(doc, "关键判别：**当报告使用的“增长词”和统计公报里的实际数不一致时，数据优先。**例如2025年GDP目标5.5%左右、实际5.0%略低于目标；社零目标6%左右、实际3.9%；固投目标5%左右、实际-13.1%。差异反映：吉林“工业/装备/冰雪/财政强，投资/地产/人口弱”。")
heading2(doc, "1.3 指标取舍与单位说明")
para(doc, "本报告统一以“2025年统计公报+2026年报告复盘”为口径，增速均为可比价，绝对数为人民币。GDP 14973.88亿元为全省初步核算数。")

# ---- 二、底盘 ----
heading1(doc, "二、先看吉林的特殊底盘：汽车、冰雪、人参医药、对俄开放与装备制造")
para(doc, "吉林的地盘，取决于它作为“**汽车工业基地+冰雪经济+人参医药+对俄/对朝开放**”的特殊定位。它是东北老工业基地，向“新质生产力”转型。")
bullet(doc, "**汽车**：2025年汽车产量146.13万辆（-3.5%）、新能源汽车18.36万辆（+22.5%）；红旗、奥迪PPE、比亚迪弗迪等项目支撑；汽车产业“电动化/智能化”转型。")
bullet(doc, "**冰雪/文旅**：2025年接待游客4.68亿人次（+20.8%）、入境游客+33%；长白山世界地质公园、查干湖/嫩江湾5A级；打造旅游万亿级产业。")
bullet(doc, "**人参医药**：人参产业综合产值2024年已破800亿元；医药/高技术制造增加值+13.0%；长白山人参资源独特。")
bullet(doc, "**对俄/对朝开放**：进出口对一带一路占62.1%、出口+11.8%；人参医药/汽车随之“走向东北亚”。")
para(doc, "**制度含义**：吉林不追求总量冒进，而是把“汽车、冰雪、人参医药、粮食、对俄开放”当核心资产，并在东北振兴中发掘“新质生产力”。")

# ---- 三、宏观错位 ----
heading1(doc, "三、最关键的宏观错位：GDP近1.5万亿、制造/工业强，但投资/地产/人口弱")
para(doc, "2025年吉林GDP 14973.88亿元、+5.0%（一产+4.6%、二产+5.5%、三产+4.8%）。表面看“稳中向好”，但拆开看是“**工业/装备/冰雪强、投资/地产/人口弱**”的错位：")
para(doc, "**强的部分**：规上工业+7.8%（制造业+8.8%、装备制造+13.3%、高技术制造+13.0%）；新能源汽车产量+22.5%、动车组+22.8%、城市轨道车辆+42.8%；游客4.68亿人次（+20.8%）；一般公共预算收入+13.3%。")
para(doc, "**弱的部分**：固定资产投资-13.1%（地产-25.6%、三产-18.4%）；房地产开发-25.6%；进出口1638.63亿、-7.1%（进口-18.9%）；人口自然增长率-6.37‰、老龄化29.69%；CPI-0.1%、PPI-3.5%。")
para(doc, "**核心错位一句话**：吉林“制造/装备/冰雪/财政强（汽车、装备+13%、冰雪+20.8%），但投资/地产/人口弱”。2026年“稳投资/地产、扩内需、续汽车/冰雪/边贸”是重点。")
table(doc,
    ["强引擎", "数据", "弱项", "数据"],
    [
        ["装备制造业", "+13.3%", "固定资产投资", "-13.1%"],
        ["高技术制造业", "+13.0%", "房地产开发投资", "-25.6%"],
        ["新能源汽车", "+22.5%", "进出口总额", "-7.1%"],
        ["接待游客", "+20.8%", "人口自然增长率", "-6.37‰"],
        ["一般公共预算收入", "+13.3%", "CPI", "-0.1%"],
    ],
    widths=[3.6, 3.0, 3.6, 3.0])
para(doc, "**错位结论**　吉林的增长“强工业/冰雪、弱投资/地产/人口”。2026年“稳工业/装备、补投资/地产、扩冰雪/对俄开放、稳人口”是重点。")

# ---- 四、被忽视细节 ----
heading1(doc, "四、容易被忽视、但信号很重的细节（15条）")
details = [
    ("1", "GDP 14973.88亿、+5.0%", "总量/稳增。"),
    ("2", "规上工业+7.8%、装备制造+13.3%", "工业/装备领跑。"),
    ("3", "高技术制造+13.0%", "新质生产力。"),
    ("4", "新能源汽车产量18.36万辆、+22.5%", "新能源车向。"),
    ("5", "汽车产量146.13万辆、-3.5%", "传统汽车调整。"),
    ("6", "动车组+22.8%、城市轨道车辆+42.8%", "轨道装备。"),
    ("7", "进出口1638.63亿、-7.1%、出口+11.8%", "出口强、进口弱。"),
    ("8", "对一带一路占62.1%、民企出口+24.7%", "外贸结构/民企。"),
    ("9", "游客4.68亿人次、+20.8%、入境+33.0%", "冰雪/旅游强。"),
    ("10", "旅游万亿产业/长白山世界地质公园", "文旅/冰雪。"),
    ("11", "人参产业破800亿、医药+13.0%", "人参医药。"),
    ("12", "常住2296.69万、城镇化率66.83%（+1.07pct）", "人口/城镇化。"),
    ("13", "自然增-6.37‰、60岁+29.69%", "老龄化加深。"),
    ("14", "居民收入城镇40817元/+4.2%、农村+5.8%", "农村快于城镇。"),
    ("15", "一般公共预算收入+13.3%（税收+4.3%）", "财政大幅增长。"),
]
for d in details:
    para(doc, "**%s**" % d[0], size=11)
    para(doc, d[1], size=10.5, color=GRAY)

    para(doc, "**备注**　这条“错位”在吉林尤其鲜明：增长靠制造/装备/冰雪/财政，但投资/地产/人口弱。2026年若投资/地产修复、对俄开放发力，增长可能从“工业单极”走向“工业+冰雪/内需”多极。这条细节，正是吉林2026年最难也最值得盯的“变量”。", size=10.5)

# ---- 五、2025年目标 vs 实际 对照 ----
heading1(doc, "五、2025年GDP目标 vs 实际：对照表")
t5 = [
    ["指标", "2025目标", "2025实际", "达标判定"],
    ["GDP增速", "5.5%左右", "+5.0%", "略低"],
    ["固定资产投资增速", "5%左右", "-13.1%", "大幅未达标"],
    ["社会消费品零售增速", "6%左右", "+3.9%", "未达标"],
    ["一般公共预算收入增速", "3%左右", "+13.3%", "大幅超预期"],
    ["粮食产量", "880亿斤以上", "853.2亿斤左右", "接近/微低"],
    ["城镇调查失业率", "控制在5.5%左右", "约5.5%", "达标"],
    ["居民消费价格(CPI)", "涨幅2%左右", "-0.1%", "远低于"],
]
table(doc, t5[0], t5[1:], widths=[3.4, 3.0, 3.0, 3.6])
para(doc, "**对照结论**　目标“偏进取”，实际“分化”：财政收入+13.3%大幅超预期（汽车/冰雪/装备带动），但固投-13.1%、社零3.9%未达（6%目标）。制造与财政接住，投资/地产/人口是短板。")

# ---- 六、增速分项支撑 ----
heading1(doc, "六、2025年增长是谁撑起来的？（结构归因）")
para(doc, "GDP+5.0%背后，是“**制造/装备/冰雪/财政强、投资/地产/人口弱**”的结构。把增长贡献拆开看：")
g6 = [
    ["引擎", "贡献/状态", "说明"],
    ["规上工业", "+7.8%", "装备+13.3%、高技术+13.0%"],
    ["汽车/新能源车", "-3.5%/+22.5%", "汽车转型、新能源向"],
    ["装备制造", "+13.3%", "动车组/城轨/装备"],
    ["冰雪旅游", "游客+20.8%", "4.68亿人次、万亿产业"],
    ["人参医药", "破800亿/医药+13.0%", "人参医药"],
    ["一般公共预算收入", "+13.3%", "财政大幅增长"],
    ["房地产开发", "-25.6%", "地产调整"],
    ["固定资产投资", "-13.1%", "地产/三产拖累"],
    ["进出口", "-7.1%", "出口+11.8%、进口-18.9%"],
    ["人口自然增长率", "-6.37‰", "人口/老龄化约束"],
]
table(doc, g6[0], g6[1:], widths=[3.6, 3.6, 6.0])
para(doc, "**一句话**　增长靠“工业（装备/高技术）+冰雪/文旅+财政”，但投资/地产/人口是最大拖累。2026年考验吉林“能不能让投资/地产/人口也跟着稳”。")

# ---- 七、预算与财政 ----
heading1(doc, "七、预算与财政的“含金量”")
para(doc, "2025年吉林一般公共预算收入**1349.97亿元、+13.3%**，其中税收收入**+4.3%**。财政收入大幅增长（汽车/冰雪/财政改革带动）。")
bullet(doc, "财政收入+13.3%、税收+4.3%，大幅改善。")
bullet(doc, "财政“稳收+民生+装备/冰雪/边贸支出优先”。")
bullet(doc, "产业投资占比超50%，支撑转型。")

# ---- 八、民生底账 ----
heading1(doc, "八、民生底账：人口、收入与城乡")
para(doc, "2025年吉林常住人口**2296.69万人**，城镇化率**66.83%、+1.07pct**。人口自然增长率**-6.37‰**，60岁及以上占29.69%。")
para(doc, "居民收入城镇**40817元、+4.2%**、农村**21911元、+5.8%**。")
g8 = [
    ["指标", "数值", "信号"],
    ["常住人口", "2296.69万", "人口负增压力"],
    ["人口自然增长率", "-6.37‰", "自然负增加深"],
    ["城镇化率", "66.83%/+1.07pct", "稳步城镇化"],
    ["农村居民收入", "21911元/+5.8%", "农村快于城镇"],
]
table(doc, g8[0], g8[1:], widths=[4.2, 4.0, 4.8])
para(doc, "**民生观察**：农村收入快于城镇、城镇化推进较快，但人口自然负增、老龄化（29.69%）是全国最深之一。")

# ---- 九、城镇与农村 ----
heading1(doc, "九、城镇与农村：格局与均衡")
para(doc, "吉林城镇化率66.83%、+1.07pct，城乡收入比缩小。")
bullet(doc, "农村居民收入+5.8%、快于城镇+4.2%。")
bullet(doc, "社零城镇+4.8%、乡村-3.4%，城镇消费引领。")
bullet(doc, "人口向长春都市圈/城市群集中。")

# ---- 十、人口流入与流出 ----
heading1(doc, "十、人口流入与流出")
para(doc, "吉林2025年常住2296.69万、自然增-6.37‰，人口负增长持续（出生率3.08‰、死亡率9.45‰）。")
para(doc, "未来看点：东北振兴+汽车/冰雪/人参能否留住人口；若“冰雪经济+汽车+对俄”成势，吉林有望减缓人口流出。")

# ---- 十一、物价与货币 ----
heading1(doc, "十一、物价与货币环境")
para(doc, "2025年吉林CPI**-0.1%**、PPI**-3.5%**，物价承压、工业品价格走弱。")
para(doc, "物价偏弱反映“工业/供给强、需求/消费弱”、农产品价格-6.0%，与全国低通胀一致。2026年“扩内需、稳物价”是主线。")

# ---- 十二、区域一体化 ----
heading1(doc, "十二、区域一体化：吉林在“对俄朝合作+中部城市群+东北振兴”里的位置")
para(doc, "吉林的核心战略坐标是“**对俄/对朝开放+中部城市群（长春都市圈）+东北振兴**”，是面向东北亚开放的前沿。")
bullet(doc, "对俄/对朝：一带一路占62.1%，延边/长吉图开发开放。")
bullet(doc, "长春都市圈：一汽集团总部、汽车产业集聚。")
bullet(doc, "东北振兴：新质生产力、装备制造、冰雪经济。")
bullet(doc, "生态：白山黑水、长白山、查干湖生态旅游。")
para(doc, "若“汽车/装备+冰雪/文旅+对俄开放”成势，吉林将作为东北振兴/东北亚开放的样本。")

# ---- 十三、五条主线 ----
heading1(doc, "十三、未来5—10年最值得观察的五条主线")
lines5 = [
    ("① 汽车新能源化/电动化", "新能源汽车能否成为新增长极。"),
    ("② 冰雪经济/文旅", "冰雪万亿产业、旅游能否持续升级。"),
    ("③ 人参医药/粮食", "人参医药、粮食(单产全国前列)能否壮大。"),
    ("④ 对俄朝开放/边贸", "一带一路、对俄/朝贸易能否放大。"),
    ("⑤ 工业/装备制造", "装备/高技术制造能否成为新动能。"),
]
for l in lines5:
    para(doc, "**%s**　%s" % l, space_after=6)

# ---- 十四、结论 ----
heading1(doc, "十四、最终结论：吉林在“汽车+冰雪+人参医药+对俄开放”里的增长逻辑")
para(doc, "吉林的2025年，本质上是“**汽车/装备/冰雪/人参医药/财政为核心，而投资/地产/人口弱**”的答卷：GDP14973.88亿、+5.0%，规上工业+7.8%、游客4.68亿人次（+20.8%）、财政+13.3%，但固投-13.1%、地产-25.6%、自然增-6.37‰、进出口-7.1%。")
para(doc, "只要汽车新能源化、冰雪经济、人参医药、对俄开放持续，吉林就站在“东北振兴+东北亚开放”的位；如果投资/地产/人口持续偏弱，吉林需承受“强工业/冰雪、弱投资/人口”的挑战。")
para(doc, "最稳观察信号：**一盯汽车/新能源（工业）、二盯冰雪/旅游（特色）、三盯人参医药/粮食（农业/医药）、四盯对俄朝/一带一路（开放）、五盯投资/人口（约束）。**吉林，是“汽车+冰雪+对俄开放”的新样本。")

# ---- 附录A ----
heading1(doc, "附录A：主要资料来源与核验路径")
bullet(doc, "吉林省2025年《政府工作报告》——目标来源。")
bullet(doc, "《2025年吉林省国民经济和社会发展统计公报》（省统计局，2026-04-27）——GDP、工业、外贸、人口实值。")
bullet(doc, "2026年吉林省政府工作报告——2025执行复盘+冰雪/人参/对俄/装备。")
bullet(doc, "长春海关、省财政厅——外贸/财政。")
heading2(doc, "核验说明")
para(doc, "本报告所有增速、总量、占比均以统计公报/官方口径为基准，涉“汽车/冰雪/人参医药/对俄朝”等以官方口径为准。")

# ---- 附录B ----
heading1(doc, "附录B：建议建立的年度跟踪仪表盘")
para(doc, "建议把下面10个“测脉搏”指标做成年度跟踪表：")
dash = [
    ["#", "指标", "2025基值", "为什么重要"],
    ["1", "GDP增速", "+5.0%", "总量与方向"],
    ["2", "规上工业增速", "+7.8%", "制造底盘"],
    ["3", "新能源汽车产量", "+22.5%", "工业转型"],
    ["4", "接待游客/增速", "4.68亿/+20.8%", "冰雪/旅游"],
    ["5", "固定资产投资/地产", "-13.1%/-25.6%", "投资结构"],
    ["6", "社零增速", "+3.9%", "内需消费"],
    ["7", "常住人口/自然增长率", "2296.69万/-6.37‰", "人口与城市"],
    ["8", "一般公共预算收入增速", "+13.3%", "财政质量"],
    ["9", "对一带一路占比", "62.1%", "东北亚开放"],
    ["10", "CPI/PPI", "-0.1%/-3.5%", "物价与需求"],
]
table(doc, dash[0], dash[1:], widths=[0.8, 4.6, 3.4, 4.4])
para(doc, "把这10个指标连起来看，任何一个汽车/新能源（3）、冰雪/旅游（4）、人参/粮食（农业）向上、投资/人口（5/7）修复，都说明吉林在真正换挡。")

# ===================================== 保存
out = "/Users/x/Desktop/content-prod-lab/reports/吉林省_2025年政府工作报告_深度研究_2026-08-13.docx"
doc.save(out)
print("SAVED:", out)
print("PARAGRAPHS:", len(doc.paragraphs))
print("TABLES:", len(doc.tables))
