# -*- coding: utf-8 -*-
"""Build 济宁市2025年政府工作报告 深度研究 DOCX, 参照地级市系列版式。"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DARK = RGBColor(0x1F, 0x2A, 0x44)
RED = RGBColor(0xB0, 0x1E, 0x2E)
GRAY = RGBColor(0x59, 0x60, 0x69)
BLUE = RGBColor(0x1F, 0x4E, 0x79)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
HEADER_FILL = "1F2A44"
K = "微软雅黑"


def set_run(run, size=11, bold=False, color=DARK, font=K):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    run.font.name = font
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), font)
    rFonts.set(qn("w:hAnsi"), font)
    rFonts.set(qn("w:eastAsia"), font)


def para(doc, text, size=11, bold=False, color=DARK, align=None,
         space_after=6, space_before=0, font=K):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        set_run(r, size=size, bold=bold or (i % 2 == 1), color=color, font=font)
    return p


def heading1(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(16)
    pf.space_after = Pt(8)
    r = p.add_run(text)
    set_run(r, size=16, bold=True, color=RED)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "B01E2E")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def heading2(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(10)
    pf.space_after = Pt(4)
    r = p.add_run(text)
    set_run(r, size=13, bold=True, color=BLUE)
    return p


def table(doc, headers, rows, widths=None, font_size=9.5):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_run(r, size=font_size, bold=True, color=WHITE)
        tcPr = hdr[i]._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), HEADER_FILL)
        tcPr.append(shd)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(val))
            set_run(r, size=font_size, color=DARK)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    return t


def quote_box(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(8)
    pf.left_indent = Pt(12)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), "B01E2E")
    pBdr.append(left)
    pPr.append(pBdr)
    r = p.add_run(text)
    set_run(r, size=11, color=DARK, font="楷体")
    return p


def bullet(doc, text, size=10.5):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.7)
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        set_run(r, size=size, bold=(i % 2 == 1), color=DARK)
    return p


# ================================================================= 文档主体
doc = Document()
sec = doc.sections[0]
sec.page_width = Cm(21)
sec.page_height = Cm(29.7)
sec.top_margin = Cm(2.4)
sec.bottom_margin = Cm(2.2)
sec.left_margin = Cm(2.5)
sec.right_margin = Cm(2.5)

st = doc.styles["Normal"]
st.font.name = K
st.font.size = Pt(11)
st.element.rPr.rFonts.set(qn("w:eastAsia"), K)


# ---- 封面 ----
para(doc, "三明市2025年政府工作报告", size=24, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=60, space_after=4)
para(doc, "深度研究与\u201c容易被忽视的细节\u201d分析", size=16, bold=True, color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
para(doc, "从\u201c三钢钢铁、客家祖地、沙县小吃、医改/林改国家样板、红色苏区\u201d重新理解三明", size=12, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)
para(doc, "\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501", color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
para(doc, "研究对象：2025年三明市政府工作报告及同期财政、国民经济和社会发展统计资料、2026年官方复盘", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
para(doc, "标定日期：2026-08-14", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
doc.add_page_break()

# ---- 目录 ----
catalog = [
    "执行摘要",
    "一、研究口径与阅读方法",
    "二、先看三明的特别底盘：三钢、客家、沙县小吃、医改林改、红色苏区",
    "三、最关键的宏观错位：GDP-4.3%（工业-15.7%拖累），但新材料/社零/医改林改稳",
    "四、容易被忽视、但信号很重的细节（15条）",
    "五、2025年GDP目标 vs 实际：对照表",
    "六、2025年增长是谁撑起来的？（结构归因）",
    "七、预算与财政的\u201c含金量\u201d",
    "八、民生底账：人口、收入与城乡",
    "九、城镇与农村：格局与均衡",
    "十、人口流入与流出",
    "十一、物价与货币环境",
    "十二、区域一体化：三明在海西经济区、闽西南协同区、客家文化圈\u201c三圈\u201d里的位置",
    "十三、未来5—10年最值得观察的五条主线",
    "十四、最终结论：三明在\u201c钢铁转型+小吃/文旅+林改医改\u201d里的增长逻辑",
    "附录A：主要资料来源与核验路径",
    "附录B：建议建立的年度跟踪仪表盘",
]
para(doc, "目　录", size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=10, space_after=10)
for item in catalog:
    para(doc, item, size=12, space_after=4)
doc.add_page_break()

# ---- 执行摘要 ----
heading1(doc, "执行摘要")
para(doc, "**核心判断**　2025年三明最显著的是\u201cGDP 2511.98亿/-4.3%（罕见负增长、第二产业-15.7%、规模以上工业-20.7%）、固定资产投资-17.9%、房地产开发-34.7%、进出口-28.8%\u201d、\u201c但新材料产业+13.4%、社零+2.9%、一般公共预算收入+2.6%、医改/林改国家样板\u201d。这说明三明在\u201c三钢钢铁老基地+客家/小吃\u201d中，**传统重工业深度调整是最大拖累，结构转型（医改/林改/文旅）稳**。")
para(doc, "把2025年目标（GDP+5%、财收+3%、固投+5%、社零+5%、出口+3%）、2025年实际（GDP-4.3%、财收+2.6%、固投-17.9%、社零+2.9%、出口-33.1%）趋势看，三明是\u201c钢铁+转型\u201d路径：**三钢（钢铁）、石墨烯/氟化工（新材料）、沙县小吃、纺织、林产、绿色食品\u201d是支柱；11条特色产业链6000亿目标。")
para(doc, "最容易记住的一句话：**三明是\u201c三钢钢铁城、客家祖地、沙县小吃之乡、中国绿都\u201d，靠\u201c钢铁+特色小吃+医改林改\u201d驱动。**观察三明，与其只看\u201cGDP 2512亿（-4.3%）\u201d，不如看\u201c规上工业-20.7%（钢铁深度调整）、新材料+13.4%、沙县小吃/扁食、医改林改、旅游6900万人次\u201d。")
# ---- 一、研究口径与阅读方法 ----
heading1(doc, "一、研究口径与阅读方法")
para(doc, "本报告以\u201c三明市政府工作报告（2025年2月）\u201d为起点，把\u201c2025年GDP目标（5%）\u201d与\u201c官方2025年（2511.98亿/-4.3%）\u201d并置对照，用\u201c2025年三明市统计公报\u201d和\u201c2026年计划执行报告\u201d横向核验。")
para(doc, "口径提示：\u201cGDP、规上工业增速\u201d按可比价（实际）；\u201c投资、消费、进出口、财政\u201d按现价（名义）。\u201c常住人口\u201d用统计公报（240.1万）、城镇化率68.26%。")
para(doc, "指标体系（与研究口径一致）：**总量与增速、产业动能（钢铁/新材料/小吃）、外贸、财政、民生与人口**。")
para(doc, "特别提示（不吃老本）：三明2024年GDP（+5.2%）、2025年2511.98亿/-4.3%（罕见负增长）；它不是\u201c只有小吃\u201d——**三钢（钢铁）、石墨烯/氟化工（新材料）、沙县小吃、林改医改、客家祖地、红色苏区\u201d才是真正底色。")
# ---- 二、先看三明的特别底盘 ----
heading1(doc, "二、先看三明的特别底盘：三钢、客家、沙县小吃、林改医改、红色苏区")
para(doc, "三明地处福建省中西部、闽江上游/武夷山南麓，是**中国钢铁基地（三钢集团）、客家祖地（宁化石壁）、沙县小吃之乡、中国绿都（国家林改试验区）、三明医改发源地**。2025年GDP 2511.98亿、常住240.1万、城镇化率68.26%、福建中部（福州/厦门/泉州/漳州之后）。")
para(doc, "五个底盘名词，先立框架：")
bullet(doc, "**三钢/钢铁（老基地）**　三钢（闽光钢铁）、钢铁冶炼、2025深度调整——\u201c钢铁城\u201d。")
bullet(doc, "**新材料/绿色工业**　石墨烯、氟化工（氟新材料）、生物医药、稀土——\u201c新材料+13.4%\u201d。")
bullet(doc, "**沙县小吃/美食经济**　沙县小吃（扁食/沙县）、沙县小吃产业园、全国网点——\u201c小吃经济\u201d。")
bullet(doc, "**林改/医改（国家样板）**　林改（林票/林权）、三明医改（全国推广）、集体林权——\u201c改革样板\u201d。")
bullet(doc, "**红色苏区/旅游业态**　建宁/宁化（中央苏区）、客家祖地、泰宁（丹霞）、绿都——\u201c红色+绿都+客家\u201d。")
para(doc, "这五根（钢铁+新材料+小吃+林改医改+红色）构成三明独特底盘：**左手三钢（重工业），右手小吃/林改医改（样板）**。理解三明，先理解\u201c三钢、沙县小吃、医改林改\u201d。")
# ---- 三、最关键的宏观错位 ----
heading1(doc, "三、最关键的宏观错位：GDP-4.3%（工业-15.7%拖累），但新材料/社零/医改林改稳")
para(doc, "2025年三明最需要辨析的一组\u201c错位\u201d：**GDP-4.3%（罕见负增长）、规上工业-20.7%（钢铁/制造深度调整）、固定资产投资-17.9%、房地产开发-34.7%、进出口-28.8%，但新材料+13.4%、社零+2.9%、地方财收+2.6%、医改林改/文旅（游客6900万/600亿）**。")
para(doc, "为什么\u201c钢铁/投资/外贸\u201d深度下滑，新材料/消费/改革却稳？三解释：")
para(doc, "**其一，钢铁/重周期深度调整、拖累总量**　规上-20.7%（钢铁、纺织、建材）、地产-34.7%、投资-17.9%——\u201c工业重周期\u201d。")
para(doc, "**其二，新材料/绿色转型稳**　石墨烯/氟化工等新材料+13.4%（占比11.6%、提高）、高技术产业投资+36.1%——\u201c新质亮点\u201d。")
para(doc, "**其三，消费/医改林改/文旅稳**　社零+2.9%、小吃/文旅6900万人次、医改林改（国家推广）、森林绿都——\u201c服务+改革\u201d。")
para(doc, "小结：三明2025年是\u201c**钢铁深度调整、新质消费改革稳**\u201d：新材料/小吃/文旅/医改林改稳，钢铁/投资/外贸弱。")

# ---- 四、容易被忽视、但信号很重的细节（15条） ----
heading1(doc, "四、容易被忽视、但信号很重的细节（15条）")
bullet(doc, "**1.GDP-4.3%（罕见负增长）**\u201c工业/钢铁深度调整。\u201d")
bullet(doc, "**2.规上工业-20.7%（轻-23.5%/重-19.4%）**\u201c三钢/重工拖累。\u201d")
bullet(doc, "**3.新材料产业+13.4%（石墨烯/氟化工/占比至11.6%）**\u201c新质亮点。\u201d")
bullet(doc, "**4.固定资产投资-17.9%但高技术产业投资+36.1%**\u201c转型投资。\u201d")
bullet(doc, "**5.地产-34.7%、商品房网签转正（居全省前列）**\u201c地产调整、销售回暖。\u201d")
bullet(doc, "**6.社零878亿/+2.9%（通讯+109.7%/手机+125%）**\u201c消费/智能。\u201d")
bullet(doc, "**7.进出口103.3亿/-28.8%（出口-33.1%/进口+19.8%）**\u201c外贸出口收缩。\u201d")
bullet(doc, "**8.三明医改（医疗收入占比>50%历史性突破、国家推广）**\u201c医改样板。\u201d")
bullet(doc, "**9.林改（林票2.0、森林四库、全国首个林业碳票）**\u201c林改样板、绿都。\u201d")
bullet(doc, "**10.沙县小吃（小吃经济/产业园）**\u201c美食IP、全国网点。\u201d")
bullet(doc, "**11.旅游6900万人次/600亿+、客家祖地/泰宁丹霞**\u201c文旅绿都。\u201d")
bullet(doc, "**12.一般公共预算收入120.12亿/+2.6%（总收167.11/+1.5%）**\u201c财政稳。\u201d")
bullet(doc, "**13.居民收入40752元/+4.8%、城乡比1.80（缩）**\u201c农村+5.1%>城镇+4.0%。\u201d")
bullet(doc, "**14.常住240.1万/-1.8万、城镇化68.26%、自然增长-3.86‰**\u201c人口净流出/老龄。\u201d")
bullet(doc, "**15.CPI-0.4%、中国绿都连续4年榜首**\u201c低通胀、生态。\u201d")
# ---- 五、2025年GDP目标 vs 实际：对照表 ----
heading1(doc, "五、2025年GDP目标 vs 实际：对照表")
table(doc,
    ["指标", "2025年目标", "2025年实际", "达成评价"],
    [
        ["地区生产总值(GDP)", "增长5%", "2511.98亿/-4.3%", "大幅不及"],
        ["规模以上工业", "——", "-20.7%", "深度下滑"],
        ["固定资产投资", "增长5%", "-17.9%", "大幅不及"],
        ["社会消费品零售总额", "增长5%", "878.08亿/+2.9%", "略低"],
        ["外贸出口", "增长3%", "89.3亿/-33.1%", "大幅负增"],
        ["地方一般公共预算收入", "增长3%", "120.12亿/+2.6%", "略低"],
        ["城镇/农村居民收入", "4%/5%", "+4.0%/+5.1%", "基本达成"],
    ],
)
para(doc, "注：GDP按不变价。**GDP（-4.3%）大幅不及5%、规上（-20.7%）、固投（-17.9%）、出口（-33.1%）深度下滑**；**新材料（+13.4%）、社零（+2.9%）、财收（+2.6%）、收入（+4.8%）**稳。")
para(doc, "拆读：**新材料、社零、医改林改、文旅是亮色**；**钢铁/规上（-20.7%）、固投（-17.9%）、地产（-34.7%）、出口（-33.1%）**是短板——\u201c钢铁深度调整、转型中\u201d，是\u201c老工业基地转绿\u201d样本。")

# ---- 六、2025年增长是谁撑起来的？（结构归因） ----
heading1(doc, "六、2025年增长是谁撑起来的？（结构归因）")
para(doc, "把三明GDP的-4.3%拆开：三次产业分别增3.6%、-15.7%、+3.3%（结构15.0：33.1：51.9）。**第二产业（钢铁/工业）-15.7%是最大拖累，第三产业（+3.3%）占51.9%是支撑**，第一产业（农业）+3.6%。")
para(doc, "2026年三明强调\u201c4.5%目标、钢铁智改数转、新材料、小吃文旅、医改林改、客属\u201d，聚焦**三钢智能化改造、石墨烯/氟化工、沙县小吃、医改林改（扩面）、红色绿色文旅**——核心是\u201c减钢铁、增新质\u201d。")
para(doc, "**第二产业（工业/钢铁）**：规上-20.7%（钢铁、纺织、建工），新材料+13.4%（占比11.6%）——\u201c钢铁压制、新质补位\u201d。")
para(doc, "**第三产业（服务业）**：+3.3%（社零、旅游6900万人次、医疗、现代服务）——\u201c服务+文旅+改革\u201d。")
para(doc, "**第一产业（农业）**：+3.6%（粮食、蔬菜222万吨、生猪、林产、茶叶）——\u201c农业+林业稳\u201d。")
para(doc, "一句话归因：**2025年三明增长\u201c靠第三产业（文旅/医改/消费）+农业林改\u201d**，钢铁/工业/投资/外贸深度调整；\u201c钢铁转型+生态绿都\u201d是主线。")

# ---- 七、预算与财政的\u201c含金量\u201d ----
heading1(doc, "七、预算与财政的\u201c含金量\u201d")
para(doc, "2025年三明**地方一般公共预算收入120.12亿元（+2.6%）**；财政总收入167.11亿（+1.5%）；支出381.44亿（-4.2%）。")
bullet(doc, "地方财收+2.6%（稳、逆GDP）——\u201c# ⚠ 财政含税稳\u201d。")
bullet(doc, "民生支出近八成（教育/医疗/社保）、医改。")
bullet(doc, "金融：绿色贷款+（林改/碳票）、政策性资金112.9亿（253个）、民营贷款+7.5%——宽信用+绿色。")
para(doc, "**财政含金量小结**：财收+2.6%（逆势）、民生八成、绿色金融强；财政对\u201c医改林改、新材料、民生\u201d投入加大。")
# ---- 八、民生底账：人口、收入与城乡 ----
heading1(doc, "八、民生底账：人口、收入与城乡")
para(doc, "2025年三明**居民人均可支配收入40752元（+4.8%）**，其中城镇50529元（+4.0%）、农村28001元（+5.1%），城乡比1.80（缩小）。消费：人均消费支出28672元（+4.9%）。就业：城镇新增就业0.92万人。")
para(doc, "人口画像：**常住240.1万/-1.8万、城镇化率68.26%、自然增长-3.86‰**；人口净流出（至沿海/省会）、老龄化。")
para(doc, "民生投入：民生支出近八成、三明医改（服务占比超50%）、学位/医疗/社保——民生扎实、改革惠民。")

# ---- 九、城镇与农村：格局与均衡 ----
heading1(doc, "九、城镇与农村：格局与均衡")
para(doc, "三明城镇化率68.26%；县域（沙县小吃/泰宁丹霞/建宁莲子/宁化客家）；农村收入增速（+5.1%）>城镇（+4.0%），**城乡比1.80缩小**；林改/竹林/乡村。")
para(doc, "农业底盘：**粮食/蔬菜222万吨/生猪/烟叶/茶叶/莲子、林产（林票/碳票）**——\u201c绿都+特色农业\u201d。")
para(doc, "一句话：\u201c三明是山区农业+林改、客家/小吃县域、城乡均衡\u201d。")

# ---- 十、人口流入与流出 ----
heading1(doc, "十、人口流入与流出")
para(doc, "三明常住240.1万（-1.8万）、城镇化68.26%；\u201c山区/钢铁调整、青年外流福州/厦门/珠三角\u201d、客家/小吃人口。")
para(doc, "结构观察：**自然增长-3.86‰（严重负）、深度老龄化**；小吃从业（遍布全国）外出。")
para(doc, "2026年目标：稳就业（0.75万）、健康城市——三明靠\u201c绿都+改革+文旅\u201d聚人。")

# ---- 十一、物价与货币环境 ----
heading1(doc, "十一、物价与货币环境")
para(doc, "2025年三明**市辖区CPI-0.4%**（消费品-0.8%、服务+0.4%）——\u201c低通胀、需求弱\u201d。")
bullet(doc, "信贷：绿色贷款+（林改碳票）、民营贷款+7.5%、政策性资金——宽信用+绿色。")
bullet(doc, "消费：社零+2.9%、通讯/手机高增、旅游——消费稳。")
para(doc, "货币环境判断：**宽信用/绿色、CPI-0.4%**；三明靠\u201c新材料+文旅+改革\u201d稳需求（2026 CPI 3%）。")

# ---- 十二、区域一体化：三明的位置 ----
heading1(doc, "十二、区域一体化：三明在海西经济区、闽西南协同区、客家文化圈\u201c三圈\u201d里的位置")
para(doc, "三明是**海西经济区（福建中西部）、闽西南协同发展区（互联福厦泉）、客家文化圈（宁化石壁祖地）、中央苏区（红色）**。")
bullet(doc, "**闽西南协同区**　对接厦门/福州/泉州、山区协作、绿色与产业链。")
bullet(doc, "**海西经济区**　福建中发展、福厦泉辐射、向海。")
bullet(doc, "**客家/红色/绿都**　客家祖地（宁化/客家公祭）、红色建宁、绿都林改。")
para(doc, "一句话：**三明在\u201c闽西南+海西+客家绿都\u201d里，最核心是\u201c钢铁转型+绿色+改革样板\u201d**；生态、小吃、医改林改、客家是优势。")

# ---- 十三、未来5—10年最值得观察的五条主线 ----
heading1(doc, "十三、未来5—10年最值得观察的五条主线")
bullet(doc, "**主线一：钢铁/三钢智改数转**\u201c三钢绿色/智能、去产能\u201d。")
bullet(doc, "**主线二：新材料（石墨烯/氟化工）**\u201c+13.4%、占比11.6%\u201d成新支柱。")
bullet(doc, "**主线三：沙县小吃/美食经济**\u201c标准化、产业园、IP出海\u201d。")
bullet(doc, "**主线四：医改/林改（国家样板）**\u201c林票/碳票、三明医改推广\u201d改革红利。")
bullet(doc, "**主线五：绿都/文旅/人口**\u201c泰宁/客家/红色、老龄化\u201d如何\u201c转型、聚人\u201d。")

# ---- 十四、最终结论 ----
heading1(doc, "十四、最终结论：三明在\u201c钢铁转型+小吃文旅+林改医改\u201d里的增长逻辑")
para(doc, "把2025年闭环打穿：**三明是\u201c三钢钢铁城、客家祖地、沙县小吃之乡、中国绿都\u201d**：GDP 2511.98亿/-4.3%、规上-20.7%、新材料+13.4%、医改林改、文旅600亿。")
para(doc, "三明不是\u201c只有小吃\u201d——它是**三钢+新材料+沙县小吃+林改医改+红色绿都**的复合，靠\u201c改革+新质+文旅\u201d驱动；但钢铁/投资/外贸深度调整、人口老龄化。")
para(doc, "一句话结论：**三明是\u201c钢铁转型、绿色改革、小吃之都\u201d；观察它先看\u201c新材料、医改林改、小吃、文旅、绿都\u201d，再看\u201c钢铁、投资、出口\u201d。**它是\u201c老工业基地深度调整、改革绿都支撑\u201d的福建样本。")

# ---- 附录A ----
heading1(doc, "附录A：主要资料来源与核验路径")
bullet(doc, "《2025年三明市政府工作报告》（2025年2月，2025年目标、2024年回顾+5.2%）")
bullet(doc, "《2025年三明市国民经济和社会发展统计公报》（三明市统计局，2026-06-15，2025年实际）")
bullet(doc, "《三明市2026年国民经济和社会发展计划草案报告》（2026年1月，复盘+2026目标）")
bullet(doc, "三明市人民政府/统计局（sm.gov.cn）")

# ---- 附录B：建议建立的年度跟踪仪表盘 ----
heading1(doc, "附录B：建议建立的年度跟踪仪表盘")
bullet(doc, "GDP总量/增速 vs 全年目标。")
bullet(doc, "规上工业（三钢/新材料）增速。")
bullet(doc, "沙县小吃/美食经济。")
bullet(doc, "医改/林改/碳票。")
bullet(doc, "文旅/红色/绿都（游客）。")
bullet(doc, "固定资产/工业/地产投资。")
bullet(doc, "进出口/外资/绿色贷款。")
bullet(doc, "财收/民生/改革。")
bullet(doc, "常住/老龄化/就业。")
bullet(doc, "CPI/存贷款/森林生态。")

# ---------------- 保存 ----------------
out = "/Users/x/Desktop/content-prod-lab/reports/三明市_2025年政府工作报告_深度研究_2026-08-14.docx"
doc.save(out)
print("SAVED OK: 三明市", out)
