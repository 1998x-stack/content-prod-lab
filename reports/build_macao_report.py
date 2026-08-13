# -*- coding: utf-8 -*-
"""Build 中国澳门特别行政区2025年深度研究 DOCX, 参照省系列版式。"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DARK = RGBColor(0x1F, 0x2A, 0x44)
RED = RGBColor(0xB0, 0x1E, 0x2E)
GRAY = RGBColor(0x59, 0x60, 0x69)
BLUE = RGBColor(0x1F, 0x4E, 0x79)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
HEADER_FILL = "1F2A44"
K = "微软雅黑"


def set_run(run, size=11, bold=False, color=DARK, font=K):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    run.font.name = font
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), font)
    rFonts.set(qn("w:hAnsi"), font)
    rFonts.set(qn("w:eastAsia"), font)


def para(doc, text, size=11, bold=False, color=DARK, align=None,
         space_after=6, space_before=0, font=K):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        set_run(r, size=size, bold=bold or (i % 2 == 1), color=color, font=font)
    return p


def heading1(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(16)
    pf.space_after = Pt(8)
    r = p.add_run(text)
    set_run(r, size=16, bold=True, color=RED)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "B01E2E")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def heading2(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(10)
    pf.space_after = Pt(4)
    r = p.add_run(text)
    set_run(r, size=13, bold=True, color=BLUE)
    return p


def table(doc, headers, rows, widths=None, font_size=9.5):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_run(r, size=font_size, bold=True, color=WHITE)
        tcPr = hdr[i]._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), HEADER_FILL)
        tcPr.append(shd)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(val))
            set_run(r, size=font_size, color=DARK)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    return t


def quote_box(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(8)
    pf.left_indent = Pt(12)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), "B01E2E")
    pBdr.append(left)
    pPr.append(pBdr)
    r = p.add_run(text)
    set_run(r, size=11, color=DARK, font="楷体")
    return p


def bullet(doc, text, size=10.5):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.7)
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        set_run(r, size=size, bold=(i % 2 == 1), color=DARK)
    return p


# ================================================================= 文档主体
doc = Document()
sec = doc.sections[0]
sec.page_width = Cm(21)
sec.page_height = Cm(29.7)
sec.top_margin = Cm(2.4)
sec.bottom_margin = Cm(2.2)
sec.left_margin = Cm(2.5)
sec.right_margin = Cm(2.5)

st = doc.styles["Normal"]
st.font.name = K
st.font.size = Pt(11)
st.element.rPr.rFonts.set(qn("w:eastAsia"), K)



# ---- 封面 ----
para(doc, "中国澳门特别行政区2025年深度研究", size=24, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=60, space_after=4)
para(doc, "经济与\u201c容易被忽视的细节\u201d分析", size=16, bold=True, color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
para(doc, "从\u201c世界旅游休闲中心、经济适度多元、横琴合作区\u201d重新理解澳门", size=12, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)
para(doc, "\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501", color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
para(doc, "研究对象：2025年中国澳门经济数据、施政报告/财政预算及经济适度多元、横琴粤澳深合区进展", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
para(doc, "标定日期：2026-08-13", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
doc.add_page_break()

# ---- 目录 ----
catalog = [
    "执行摘要",
    "一、研究口径与阅读方法",
    "二、先看澳门的特殊底盘：世界旅游休闲中心+经济适度多元+横琴",
    "三、最关键的宏观结构：GDP破4100亿澳门元、服务/博彩强，多元转型待加速",
    "四、容易被忽视、但信号很重的细节（15条）",
    "五、2025年GDP目标 vs 实际：对照表",
    "六、2025年增长是谁撑起来的？（结构归因）",
    "七、财政与博彩的\u201c含金量\u201d",
    "八、民生与人口",
    "九、博彩与文旅（收入结构）",
    "十、人口流动与横琴",
    "十一、物价与货币环境",
    "十二、区域一体化：横琴粤澳深合作区与大湾区",
    "十三、未来5—10年最值得观察的五条主线",
    "十四、最终结论",
    "附录A：主要资料来源与核验路径",
    "附录B：建议建立的年度跟踪仪表盘",
]
para(doc, "目　录", size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=10, space_after=10)
for item in catalog:
    para(doc, item, size=12, space_after=4)
doc.add_page_break()

# ---- 执行摘要 ----
heading1(doc, "执行摘要")
para(doc, "**核心判断**　2025年中国澳门最显著的是\u201c本地生产总值约4180亿澳门元、实质增长4.7%（恢复至2019年89.6%）\u201d、\u201c博彩服务出口+18%（Q4）\u201d、\u201c人均GDP 60.7万澳门元\u201d、\u201c访澳旅客第四季+15.4%\u201d。这说明澳门经济靠\u201c博彩+旅游服务\u201d持续复苏，但仍处\u201c依赖博彩、多元待加速\u201d阶段。")
para(doc, "把2025全年GDP（统计暨普查局）、2025施政报告（经济适度多元/横琴）放在一起看，澳门呈现清晰主线：**巩固世界旅游休闲中心（博彩+文旅）+推进经济适度多元（1+4：中医药大健康/现代金融/高新/会展文旅）+横琴一体化**。经济总量恢复至2019年约9成，多元发展是必答题。")
para(doc, "最容易记住的一句话：**澳门是\u201c世界旅游休闲中心+经济适度多元（1+4）+横琴\u201d，靠\u201c博彩/旅游/服务\u201d复苏，靠\u201c多元产业+横琴+大湾区\u201d升级转型。**观察澳门，不只是看\u201cGDP 4180亿澳门元\u201d，更应看\u201c博彩服务出口、旅客、多元产业（中医药/金融/高新）\u201d。")

# ---- 一、研究口径与阅读方法 ----
heading1(doc, "一、研究口径与阅读方法")
heading2(doc, "1.1 本报告的核心底稿")
bullet(doc, "**2025全年及第4季本地生产总值（统计暨普查局）**——+4.7%、人均60.7万澳门元。")
bullet(doc, "**2025财政年度施政报告/预算**——经济适度多元、横琴、就业。")
bullet(doc, "以官方统计/施政口径为准，博彩与旅游数据来自旅游/财政局。")
heading2(doc, "1.2 阅读方法：显性—数据—长期")
para(doc, "**关键判别**：数据优先。澳门2025年GDP+4.7%（恢复2019 89.6%）、博彩服务出口强、私人消费+1.3%、投资-7.8%。观察澳门穿透GDP看\u201c博彩/旅游/多元/横琴\u201d。")

# ---- 二、底盘 ----
heading1(doc, "二、先看澳门的特殊底盘：世界旅游休闲中心+经济适度多元+横琴")
para(doc, "澳门作为\u201c**世界旅游休闲中心+博彩娱乐+中葡平台+横琴深合区**\u201d，自由港、连接横琴与粤港澳大湾区。")
bullet(doc, "**旅游/博彩**：博彩+旅游服务，人均GDP高、博彩占财政主源。")
bullet(doc, "**经济适度多元（1+4）**：中医药大健康/现代金融/高新技术/会展及文旅。")
bullet(doc, "**横琴深合区**：横琴粤澳深度合作区，澳琴一体化。")
bullet(doc, "**自由港/中葡平台**：自由港、人民币/中葡语国家平台。")
para(doc, "这一底板决定2025年：**博彩/旅游/服务驱动总量、多元/横琴是转型抓手**；经济总量小、高度依赖博彩与旅客。")

# ---- 三、宏观结构 ----
heading1(doc, "三、最关键的宏观结构：GDP破4100亿澳门元、服务/博彩强，多元转型待加速")
para(doc, "澳门2025年最值得咀嚼的结构，是\u201c**博彩/旅游/服务强、投资-7.8%、多元待加速**\u201d。GDP约4180.4亿澳门元、+4.7%（恢复2019年89.6%）。")
bullet(doc, "**GDP**：4180.4亿澳门元、+4.7%；经济总量恢复至2019年约89.6%；人均60.7万澳门元。")
bullet(doc, "**服务出口/博彩**：服务出口+5.0%（第4季+9.8%）、博彩服务出口第4季+18.0%。")
bullet(doc, "**私人消费**：+1.3%（第4+0.7%）；政府消费+3.0%。")
bullet(doc, "**投资**：固定资本形成-7.8%（第4+1.3%）——投资疲软。")
bullet(doc, "**旅客/文旅**：访澳旅客第4季+15.4%；旅游/会展/娱乐。")
para(doc, "**为什么读这条**：澳门经济\u201c博彩/服务驱动总量、投资疲软、多元待加快\u201d，转型压力vs博彩依赖并存。")

# ---- 四、15条细节 ----
heading1(doc, "四、容易被忽视、但信号很重的细节（15条）")
para(doc, "以下15条在2025年官方数据/施政，但被\u201cGDP+4.7%\u201d等掩盖。它们是判断澳门2025之后5—10年的关键小信号。")
bullet(doc, "**1. 经济恢复2019年89.6%**：总量较疫情前仍有差距。")
bullet(doc, "**2. 人均GDP 60.7万澳门元（全球前列）**：高人均、博彩/金融。")
bullet(doc, "**3. 博彩服务出口第4季+18.0%**：博彩仍主导。")
bullet(doc, "**4. 投资-7.8%**：多元/基建投资仍弱。")
bullet(doc, "**5. 横琴深合区（澳琴）**：一体化租务、人才、产业。")
bullet(doc, "**6. 经济适度多元（1+4）**：中医药/金融/高新/会展。")
bullet(doc, "**7. 访澳旅客第4+15.4%**：旅游复苏。")
bullet(doc, "**8. 中葡平台/自由港**：中葡语国家桥头堡。")
bullet(doc, "**9. 金融/现代金融**：债券市场/人民币结算/数字澳门元。")
bullet(doc, "**10. 会展/演艺（新业态）**：世界旅游休闲中心+会展。")
bullet(doc, "**11. 失业率较低**：就业稳（施政：维持较低失业率）。")
bullet(doc, "**12. 横琴/琴澳融合**：推进一体。")
bullet(doc, "**13. 中医药大健康**：国家实验室/以医带药拓展葡语市场。")
bullet(doc, "**14. 高新/科技企业认证**：吸引科创。")
bullet(doc, "**15. 财政/博彩修复**：博彩税回补财政收入。")
para(doc, "", size=10, space_after=2)

# ---- 五、对照表 ----
heading1(doc, "五、2025年GDP目标 vs 实际：对照表")
para(doc, "澳门施政预期目标为\u201cGDP保持正增长、维持较低失业率\u201d，不作硬性GDP数字。")
tb = [
    ["指标", "2025年施政目标", "2025年实际（统计）", "达成"],
    ["GDP", "保持正增长", "+4.7%（4180.4亿澳门元）", "达成"],
    ["失业率", "维持较低", "约2%+低位", "达成"],
    ["服务出口", "复苏", "+5.0%（博彩Q4+18%）", "复苏"],
    ["私人消费", "改善", "+1.3%", "增进"],
    ["投资", "多元", "-7.8%", "偏弱"],
]
table(doc, tb[0], tb[1:], widths=[3.4, 3.4, 4.4, 3.6])
para(doc, "", size=10, space_after=2)
para(doc, "**要点**：GDP/就业/旅游达标，投资偏弱——澳门\u201c服务/消费强、投资/多元弱\u201d。")

# ---- 六、结构归因 ----
heading1(doc, "六、2025年增长是谁撑起来的？（结构归因）")
heading2(doc, "6.1 服务出口/博彩")
para(doc, "**博彩+旅游服务出口**（服务出口+5.0%、博彩Q4+18%）是绝对主动力。")
heading2(doc, "6.2 私人消费/文旅")
para(doc, "私人消费+1.3%、访澳旅客+15.4%，旅游/娱乐拉动。")
heading2(doc, "6.3 投资/多元")
para(doc, "投资-7.8%（多元基建/非博彩投入仍不足）。")
heading2(doc, "6.4 金融/资产")
para(doc, "现代金融/债券/数字元推进，中葡平台。")
para(doc, "**一句话归因**：澳门2025年靠\u201c**博彩/旅游/服务出口**\u201d实现+4.7%，总量依赖博彩旅客，多元/投资待加力。")

# ---- 七、财政与博彩 ----
heading1(doc, "七、财政与博彩的\u201c含金量\u201d")
para(doc, "澳门财政高度依赖\u201c**博彩毛收入税**\u201d，2025年博彩旅游修复下财政收入恢复，维持结构性盈余弹性。")
bullet(doc, "**博彩主导财政**：博彩税是主要税源，随旅客/博彩收入修复。")
bullet(doc, "**财政稳健**：施政以\u201c财政审慎、推动多元、支援民生\u201d，无高额赤字压力。")
bullet(doc, "**多元投入**：非博彩投入、政府产业基金、科创/中医药/金融。")
para(doc, "**财政含义**：澳门\u201c博彩税+旅游\u201d财政收入，随博彩修复而改善，需推动多元/非博彩可持续发展税源。")

# ---- 八、民生与人口 ----
heading1(doc, "八、民生与人口")
para(doc, "澳门人口约68.89万（统计暨普查局）；高人均收入、民生福利完善。")
bullet(doc, "**人口**：约68.9万人。")
bullet(doc, "**就业**：较低失业率（施政维持较低失业率）、就业稳。")
bullet(doc, "**民生**：惠民/福利、居住（经屋/横琴）、医疗。")
bullet(doc, "**人才**：三期人才引进、横琴人才联动。")
para(doc, "**民生含义**：澳门\u201c人口少、就业稳、福利高\u201d，靠人才/多元化拓宽空间。")

# ---- 九、博彩与文旅 ----
heading1(doc, "九、博彩与文旅（收入结构）")
bullet(doc, "**博彩**：博彩服务出口Q4+18%，博彩仍是经济/财政主轴。")
bullet(doc, "**旅客**：访澳旅客第4季+15.4%，全年持续回暖。")
bullet(doc, "**文旅/会展**：会展演艺、一程多站、国际客源。")
para(doc, "**博彩/文旅含义**：澳门经济高度依赖\u201c博彩+旅客\u201d，博彩修复是增长主动力，但也面临多元转型长期命题。")

# ---- 十、人口流动与横琴 ----
heading1(doc, "十、人口流动与横琴")
para(doc, "澳门依靠\u201c横琴+人才计划\u201d拓展生活与就业空间。")
bullet(doc, "**横琴**：横琴粤澳深度合作区，澳琴一体化、跨境/就业。")
bullet(doc, "**人才**：人才引进/高教留澳。")
bullet(doc, "**双城**：居民横琴置业/生活、通关便利。")
para(doc, "人口与横琴：澳门靠\u201c横琴+人才\u201d拓宽土地/劳动力，缓解本地空间局限。")

# ---- 十一、物价与货币 ----
heading1(doc, "十一、物价与货币环境")
para(doc, "澳门属\u201c自由港+美元/澳门元联系\u201d环境，2025整体通胀温和。")
bullet(doc, "**通胀**：温和、低；稳定民生供应。")
bullet(doc, "**货币**：澳门元与港元/美元联系，外储充足。")
para(doc, "**物价含义**：澳门\u201c通胀温和、联系汇率稳\u201d，助力旅游/金融环境。")

# ---- 十二、区域一体化 ----
heading1(doc, "十二、区域一体化：横琴粤澳深合作区与大湾区")
para(doc, "澳门与\u201c**横琴+粤港澳大湾区**\u201d深度融合，既是旅游休闲中心也推进多元。")
bullet(doc, "**横琴深合区**：粤澳共商共建共管共享、产业/人才/生活一体化。")
bullet(doc, "**大湾区**：互联互通（口岸/轨道/跨境金融）。")
bullet(doc, "**中葡平台**：葡语国家合作、一带一路。")
para(doc, "**区域含义**：澳门靠\u201c横琴+大湾区\u201d多元拓展，巩固旅游/金融/中葡门户。")

# ---- 十三、五主线 ----
heading1(doc, "十三、未来5—10年最值得观察的五条主线")
bullet(doc, "**主线1｜经济适度多元（1+4）**：中医药/金融/高新/会展文旅能否做大。")
bullet(doc, "**主线2｜横琴深合区**：澳琴一体能否成新极。")
bullet(doc, "**主线3｜博彩/旅游转型**：博彩服务出口+旅客。能否从依赖博彩转向综合文旅。")
bullet(doc, "**主线4｜现代金融**：债券/数字澳门元/人民币结算。能否拓金融。")
bullet(doc, "**主线5｜人才/人口/民生**：人才+横琴、民生福利。能否稳健。")
para(doc, "这五条，是澳门从\u201c博彩+旅游城市\u201d走向\u201c世界旅游休闲中心+适度多元+横琴\u201d的\u201c主赛道\u201d。")

# ---- 十四、结论 ----
heading1(doc, "十四、最终结论")
para(doc, "澳门2025年，本质上是\u201c**博彩/旅游/服务出口驱动、投资和多元偏弱**\u201d的答卷：GDP约4180.4亿澳门元、+4.7%（恢复2019 89.6%）、人均60.7万澳门元、博彩服务出口Q4+18%、私人消费+1.3%、投资-7.8%。")
para(doc, "只要博彩/旅游/服务持续，澳门经济稳固回复；但多元化/投资待加速是所有结构调整的核心。")
para(doc, "最稳观察信号：**一盯博彩/旅客（引擎）、二盯经济适度多元（1+4）、三盯横琴深合区（一体化）、四盯现代金融（新增长）、五盯投资/民生（质量）。**澳门，是\u201c世界旅游休闲中心+横琴\u201d独特样本。")

# ---- 附录A ----
heading1(doc, "附录A：主要资料来源与核验路径")
bullet(doc, "统计暨普查局《2025年全年及第四季本地生产总值》。")
bullet(doc, "2025财政年度施政报告/预算（经济适度多元/横琴）。")
bullet(doc, "澳门旅游/财政局、横琴深合区管委会口径。")
heading2(doc, "核验说明")
para(doc, "本报告以官方统计/施政口径为准；涉\u201c博彩/多元/横琴\u201d等以官方为准。")

# ---- 附录B ----
heading1(doc, "附录B：建议建立的年度跟踪仪表盘")
para(doc, "建议把下面10个\u201c测脉搏\u201d指标做成年度跟踪表：")
dash = [
    ["#", "指标", "2025基值", "为什么重要"],
    ["1", "GDP实质增速", "+4.7%", "总量与方向"],
    ["2", "人均GDP", "60.7万澳门元", "规模/结构"],
    ["3", "博彩服务出口", "+18%（Q4）", "博彩主导"],
    ["4", "服务出口", "+5.0%", "服务枢纽"],
    ["5", "私人消费", "+1.3%", "内需"],
    ["6", "本地投资", "-7.8%", "多元投资"],
    ["7", "访澳旅客", "第4季+15.4%", "旅游"],
    ["8", "常住人口/失业率", "68.9万/低位", "人口/就业"],
    ["9", "横琴深合区", "一体化中", "转型引擎"],
    ["10", "通胀/联系汇率", "温和/稳定", "货币"],
]
table(doc, dash[0], dash[1:], widths=[0.8, 4.6, 4.4, 4.0])
para(doc, "把这10个连起来看，博彩/旅客/服务（2/3/7）、多元/横琴（6/9），都说明澳门在重新扩张与转型。")

# ============ 保存 ============
out = "/Users/x/Desktop/content-prod-lab/reports/中国澳门_2025年深度研究_2026-08-13.docx"
doc.save(out)
print("SAVED:", out)
print("PARAGRAPHS:", len(doc.paragraphs))
print("TABLES:", len(doc.tables))
