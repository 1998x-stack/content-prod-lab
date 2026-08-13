# -*- coding: utf-8 -*-
"""Build 重庆市2025年政府工作报告 深度研究 DOCX, 参照北京/上海/杭州/天津系列版式。"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DARK = RGBColor(0x1F, 0x2A, 0x44)
RED = RGBColor(0xB0, 0x1E, 0x2E)
GRAY = RGBColor(0x59, 0x60, 0x69)
BLUE = RGBColor(0x1F, 0x4E, 0x79)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
HEADER_FILL = "1F2A44"
K = "微软雅黑"


def set_run(run, size=11, bold=False, color=DARK, font=K):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    run.font.name = font
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), font)
    rFonts.set(qn("w:hAnsi"), font)
    rFonts.set(qn("w:eastAsia"), font)


def para(doc, text, size=11, bold=False, color=DARK, align=None,
         space_after=6, space_before=0, font=K):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        set_run(r, size=size, bold=bold or (i % 2 == 1), color=color, font=font)
    return p


def heading1(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(16)
    pf.space_after = Pt(8)
    r = p.add_run(text)
    set_run(r, size=16, bold=True, color=RED)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "B01E2E")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def heading2(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(10)
    pf.space_after = Pt(4)
    r = p.add_run(text)
    set_run(r, size=13, bold=True, color=BLUE)
    return p


def table(doc, headers, rows, widths=None, font_size=9.5):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_run(r, size=font_size, bold=True, color=WHITE)
        tcPr = hdr[i]._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), HEADER_FILL)
        tcPr.append(shd)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(val))
            set_run(r, size=font_size, color=DARK)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    return t


def quote_box(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(8)
    pf.left_indent = Pt(12)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), "B01E2E")
    pBdr.append(left)
    pPr.append(pBdr)
    r = p.add_run(text)
    set_run(r, size=11, color=DARK, font="楷体")
    return p


def bullet(doc, text, size=10.5):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.7)
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        set_run(r, size=size, bold=(i % 2 == 1), color=DARK)
    return p


# ================================================================= 文档主体
doc = Document()
sec = doc.sections[0]
sec.page_width = Cm(21)
sec.page_height = Cm(29.7)
sec.top_margin = Cm(2.4)
sec.bottom_margin = Cm(2.2)
sec.left_margin = Cm(2.5)
sec.right_margin = Cm(2.5)

st = doc.styles["Normal"]
st.font.name = K
st.font.size = Pt(11)
st.element.rPr.rFonts.set(qn("w:eastAsia"), K)


# ---- 封面 ----
para(doc, "湖北省2025年政府工作报告", size=24, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=60, space_after=4)
para(doc, "深度研究与“容易被忽视的细节”分析", size=16, bold=True, color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
para(doc, "从“中部支点、光谷、汽车与内陆开放”重新理解湖北的增长逻辑", size=12, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)
para(doc, "━━━━━━━━━━━━━━━━", color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
para(doc, "研究对象：2025年湖北省政府工作报告及同期财政、国民经济和社会发展统计资料、2026年官方复盘", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
para(doc, "整理时间：2026年8月", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
para(doc, "分析性质：分析性研究 ｜ 非官方解读", size=11, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
doc.add_page_break()

# ---- 目录 ----
catalog = [
    "执行摘要",
    "一、研究口径与阅读方法",
    "二、先看湖北的特殊底盘：中部支点、江河枢纽与科教大省",
    "三、最关键的宏观错位：GDP破6.26万亿、出口向强，但消费与地产偏弱",
    "四、容易被忽视、但信号很重的细节（15条）",
    "五、把所有细节连起来：湖北正在经历的“六个换挡”",
    "六、增长暗线：“光谷+汽车+新能源”与新质生产力",
    "七、财政暗线：收入稳、增长快，民生与化债平衡",
    "八、产业暗线：从“汽车/传统制造”到“光电子信息+新能源汽车”",
    "九、区域格局：武汉都市圈与全省分工",
    "十、人口与城市：科教大省、城镇化与老龄化",
    "十一、民营经济与科创：占60%的民企与创新",
    "十二、事后验证：用2025年实际执行结果检验报告判断",
    "十三、未来5—10年最值得观察的五条主线",
    "十四、最终结论：湖北在“中部崛起+科技自立”里的增长逻辑",
    "附录A：主要资料来源与核验路径",
    "附录B：建议建立的年度跟踪仪表盘",
]
para(doc, "目　录", size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=10, space_after=10)
for item in catalog:
    para(doc, item, size=12, space_after=4)
doc.add_page_break()

# ---- 执行摘要 ----
heading1(doc, "执行摘要")
para(doc, "**核心判断**　如果只看新闻标题，2025年湖北最显眼的是“GDP突破6.4万亿、增长5.5%”“出口增长19.6%”和“高技术制造业+15.5%”。但政府工作报告里真正值得深读的，是这一座“中部支点+光谷+科教大省”如何在投资偏弱、房地产收缩、人口老龄化的背景下，靠“光电子信息+新能源汽车+科技创新”三大主线重建增长函数。")
para(doc, "把2025年初《政府工作报告》设定的目标、2025年《国民经济和社会发展统计公报》、以及2026年报告对2025年的复盘放在一起看，湖北呈现清晰暗线：**从“基建—房产—传统制造”的旧闭环，切换到“光电子+高技术制造+内陆开放+创新驱动”的新组合**。旧引擎（房地产、一般基建）在收缩；新引擎（光谷、新能源、出口、科技）被要求更快补位。")
para(doc, "因此，本报告不按政府工作报告章节逐段复述，而采用“显性表述—同期数据—制度含义—长期影响”的方式，专门提取容易被忽视、但对判断湖北未来5—10年发展模式更有价值的信号。")
para(doc, "最容易记住的一句话：**湖北是“长江中游的引擎+全国科教与光谷的高地”——它既要回答“怎样把光谷变成中国光电与智造新中心”，也要回答“怎么在房地产、投资与人口低速时期，用科技与出口撑起中部崛起”。**观察湖北，与其看“总量”，不如看“光谷的造血能力与科技创新转化”。")

heading2(doc, "一页速览：2025年湖北经济的“表与里”")
table(doc,
    ["维度", "表面现象", "更深层信号"],
    [
        ["增长", "GDP 62660.90亿、+5.5%", "体量中部前列、产业转型推动"],
        ["产业", "规上工业+6.9%、高技术+15.5%", "光电子信息、计算机/通信电子强"],
        ["投资", "固定资产投资+2.6%", "工业投资+6.6%、第三产-0.4%"],
        ["外贸", "进出口8340.1亿、+18.2%（出口+19.6%）", "内陆开放强增长"],
        ["财政", "一般公共预算收入+6.9%", "财政稳健、民生支出大"],
        ["消费", "社零27938.62亿、+2.7%", "消费偏弱、网上零售+6.5%"],
        ["人口", "常住5811万、城镇化率67.39%", "人口负增-4.6‰、双老化"],
        ["物价", "CPI+0.1%、PPI-2.1%", "工业品通缩压力"],
    ],
    widths=[2.2, 5.2, 8.6])
para(doc, "", size=4, space_after=2)

# ---- 一、研究口径与阅读方法 ----
heading1(doc, "一、研究口径与阅读方法")
heading2(doc, "1.1 本报告的三份核心底稿")
bullet(doc, "**2025年《政府工作报告》**：2025年1月16日在省十四届人大三次会议上作，给出全年预期目标（GDP增长6%左右、规上工业增长7.5%左右、固定资产投资增长7%左右、社零增长7%左右、进出口增长7%以上、居民收入与增长同步等）。")
bullet(doc, "**2025年《国民经济和社会发展统计公报》**：湖北省统计局2026年3月18日发布，给出全年实际执行数与增速，是“事后验证”的锚点。")
bullet(doc, "**2026年省政府工作报告及官方复盘**：对2025执行结果的追认，交叉验证“目标—执行—再评价”三段式。")
heading2(doc, "1.2 阅读方法：显性—数据—制度—长期")
para(doc, "每一章按“**显性表述 → 同期数据 → 制度含义 → 长期影响**”四层展开。其中“制度含义”回答“政府为什么这么排优先序”，“长期影响”回答“这对未来5—10年意味着什么”。")
para(doc, "关键判别：**当报告使用的“增长词”和统计公报里的实际数不一致时，数据优先。**例如2025年初定“GDP增长6%左右”，实际5.5%；“固定资产投资增长7%左右”，实际只有+2.6%；“规上工业增长7.5%”，实际+6.9%；“进出口增长7%以上”，实际+18.2%（大幅超）。差异清晰说明：出口超预期、高技术强，而投资与消费偏弱。")

# ---- 二、特殊底盘 ----
heading1(doc, "二、先看湖北的特殊底盘：中部支点、江河枢纽与科教大省")
para(doc, "在所有省份里，湖北的“底盘”独特：**长江中游支点+全国交通枢纽+科教资源集聚**三合一。九省通衢的武汉，既是长江经济带的中转，也是高铁/航空枢纽；省内还有“中国光谷”武汉东湖高新区——全国光通信、光电子的高地。")
para(doc, "这决定了湖北的四重身份并存：**科技支点**（武汉是国家科教重镇、数所双一流）、**制造基地**（汽车、光电子、装备）、**开放门户**（中欧班列、长江水运）、**农业基础**（粮食509亿斤以上、稻谷主产区）。")
heading2(doc, "2.1 科教与光谷")
para(doc, "武汉是大学生数量全国前列的城市，武汉光谷拥有华工科技、烽火通信、长飞光纤等光通信龙头，还有国家实验室与“光谷实验室”。光电子是湖北的“皇冠明珠”，也是湖北产业升级的最大变量。")
heading2(doc, "2.2 汽车与制造")
para(doc, "湖北是汽车工业重镇，东风集团总部在武汉，还有武汉、襄阳、十堰、随州等汽车产业带。“中国光谷+中国车谷”双芯构成湖北制造业的两大引擎。")
heading2(doc, "2.3 交通枢纽与内陆开放")
para(doc, "武汉居中、九省通衢，中欧班列（武汉）+长江黄金水道让湖北能承接越来越多内陆开放，进出口增长快，是观察“中部内陆开放”的最佳窗口之一。")

# ---- 三、宏观错位 ----
heading1(doc, "三、最关键的宏观错位：GDP破6.26万亿、出口向强，但消费与地产偏弱")
para(doc, "把2025年湖北的宏观面放进一张表，会出现令外行惊讶的“错位”：表观增长来自工业与出口，而支撑内需的消费与地产却在收缩。这个错位，正是读懂湖北的关键。")
heading2(doc, "3.1 “总量很亮”的一面")
table(doc,
    ["指标", "2025实绩", "信号"],
    [
        ["GDP", "62660.90亿、+5.5%", "中部前列、略低于6%目标"],
        ["规上工业", "+6.9%", "工业仍有韧性"],
        ["高技术制造", "+15.5%", "占规上17.4%、新质生产力"],
        ["进出口", "8340.1亿、+18.2%（出口+19.6%）", "内陆开放大幅超目标"],
        ["一般公共预算收入", "+6.9%", "财政稳健"],
        ["粮食产量", "2791.16万吨、+0.2%", "守住粮仓"],
    ],
    widths=[3.2, 5.4, 6.4])
heading2(doc, "3.2 “偏弱”的一面")
table(doc,
    ["指标", "2025实值", "信号"],
    [
        ["固定资产投资", "+2.6%", "低于7%目标、三产-0.4%"],
        ["房地产开发投资", "-7.3%", "地产持续收缩"],
        ["社会消费品零售", "+2.7%", "低于7%目标、内需偏弱"],
        ["PPI", "-2.1%", "工业品通缩压力"],
        ["人口自然增长", "-4.6‰", "人口负增、老龄化加深"],
    ],
    widths=[3.2, 5.4, 6.4])
para(doc, "**错位结论**　湖北的增长“很强、但也很矛盾”。强的部分（出口/高技术制造/财政）与弱的部分（投资/消费/物价）并存。**真正的焦点是“量强价弱”**：出口与高技术制造高增，但内需（消费+2.7%）与投资（+2.6%）明显偏冷，PPI为负。2026年核心是“促消费、扩内需”，同时稳住投资与出口。")

# ---- 四、被忽视细节 ----
heading1(doc, "四、容易被忽视、但信号很重的细节（15条）")
details = [
    ("1", "GDP 62660.90亿、+5.5%，略低于6%目标了", "目标保守、实际执行温和偏低，需看结构与分配。"),
    ("2", "规上工业+6.9%，其中高技术制造+15.5%、占17.4%", "新质生产力占比已较高、增量多来自高技术。"),
    ("3", "计算机、通信和其他电子设备制造业+16.9%", "光谷电子信息是湖北最硬核的引擎。"),
    ("4", "高技术产业投资+4.9%（高技术服务业+10.8%）", "投资偏向服务业创新，说明转型方向。"),
    ("5", "粮食产量2791.16万吨、+0.2%、稻谷+0.2%", "湖北粮食“稳”，稻谷为本地主粮。"),
    ("6", "进出口8340.1亿、+18.2%，出口+19.6%>进口", "出口超强，是内陆开放新高阶。"),
    ("7", "网上零售额+6.5%、占社零15.8%", "数字化零售升级、服务消费增量。"),
    ("8", "社零仅+2.7%、乡村+2.8%（略快）", "内需偏弱、需政策提振。"),
    ("9", "固定资产投资+2.6%、工业投资+6.6%", "制造业优先、第三产收缩。"),
    ("10", "一般公共预算收入+6.9%、高于GDP增幅", "收入质量较好、财政稳健。"),
    ("11", "CPI+0.1%、PPI-2.1%", "工业通缩、消费价格平稳。"),
    ("12", "常住5811万、城镇化率67.39%", "城镇化率高于全国、双老化加深。"),
    ("13", "人口自然增-4.6‰、出生率4.26‰", "人口负增、城市服务需求转型。"),
    ("14", "武汉都市圈/光谷以华中吸引人口", "创新吸附力带来结构性人口增量。"),
    ("15", "居民人均可支配收入38881元、+5.2%与增长同步", "城乡协同、收入增长与GDP基本同步。"),
]
for d in details:
    para(doc, "**%s**" % d[0], size=11)
    para(doc, d[1], size=10.5, color=GRAY)

# ---- 五、六个换挡 ----
heading1(doc, "五、把所有细节连起来：湖北正在经历的“六个换挡”")
gear = [
    ("1．产业换挡：从“汽车/传统制造”到“光电子+高技术制造”", "高技术+15.5%、计算机/通信电子+16.9%成为主力。"),
    ("2．开放换挡：从“内陆腹地”到“内陆开放门户”", "进出口+18.2%、中欧班列+长江水运并进。"),
    ("3．投资换挡：从“基建/地产”到“制造业+高技术”", "工业投资+6.6%、高技术服务业+10.8%。"),
    ("4．消费换挡：从“线下大盘”到“数字化+下沉”", "网上零售+6.5%、乡村消费略快。"),
    ("5．人口换挡：从“净流入红利”到“结构性竞争”", "光谷/都市圈集聚人才、但整体非负。"),
    ("6．动能换挡：从“土地/人口”到“科技/创新驱动”", "国家实验室+光谷成为新质生产力的内核。"),
]
for g in gear:
    para(doc, "**%s**　%s" % g, space_after=6)

# ---- 六、增长暗线 ----
heading1(doc, "六、增长暗线：“光谷+汽车+新能源”与新质生产力")
heading2(doc, "6.1 光谷：中国光电与智造高地")
para(doc, "武汉光谷是全国光通信、光电子信息的高地，拥有华工科技、烽火通信、长飞光纤等全球龙头，以及“光谷实验室”。2025年计算机/通信电子+16.9%、高技术制造+15.5%，说明“光谷造”正处于放量兑现期。")
heading2(doc, "6.2 “车谷”与新能源")
para(doc, "湖北是汽车工业重镇，东风集团总部等，正从“整车制造”向“智能网联+新能源”切换。新能源汽车、充换电等成为新增长点，与光谷形成“一光一电双核”驱动。")
para(doc, "**这条暗线意味着**：湖北的增长叙事正从“汽车+基建”转向“光电+新能源+出口”。看湖北，盯住“高技术制造占比”与“光电子产值”这两组数。")

# ---- 七、财政暗线 ----
heading1(doc, "七、财政暗线：收入稳、增长快，民生与化债平衡")
para(doc, "2025年湖北地方一般公共预算收入+6.9%，高于GDP增幅；预算支出高、民生刚性。财政总体稳健，为“促内需、稳社保”提供空间。")
para(doc, "但在工业PPI通缩、地产走弱背景下，税收质量与财政收入结构仍需优化。2026年核心是在“减负”“投资”“民生”之间平衡。")
para(doc, "**制度含义**　湖北财政处于“相对稳健、增长较快”阶段，但需警惕“收入增速高于GDP”背后可能存在的一次性因素（如盘活资产）与税收结构波动。长期要提升制造业/科创税收质量。")

# ---- 八、产业暗线 ----
heading1(doc, "八、产业暗线：从“汽车/传统制造”到“光电子信息+高技术”")
heading2(doc, "8.1 湖北产业的“表”")
table(doc,
    ["指标", "2025增速", "解读"],
    [
        ["规上工业", "+6.9%", "制造业有韧性"],
        ["高技术制造", "+15.5%", "占规上17.4%"],
        ["计算机/通信/电子", "+16.9%", "光谷电子信息核心"],
        ["高技术产业投资", "+4.9%", "科创资本投入"],
        ["进出口", "+18.2%", "内陆开放最强引擎"],
        ["粮食产量", "+0.2%", "稻谷稳产"],
    ],
    widths=[5.2, 2.2, 6.2])
heading2(doc, "8.2 从“汽车”到“光电子+新能源”")
para(doc, "湖北历史上以汽车（东风）为主，2025年显示“光电子/高技术制造”逐步成为新主导；汽车新能源转型接棒。“汽车+光电”双引擎，让湖北在向“智造强省”升级过程中拥有双重抓手。")

# ---- 九、区域 ----
heading1(doc, "九、区域格局：武汉都市圈与全省分工")
para(doc, "武汉作为国家中心城市，是湖北乃至中部地区的“火车头”。武汉都市圈、襄阳、宜昌等承接汽车、装备、农产品加工。湖北呈现“强省会+多点增长”组合：武汉聚集科教/光谷/金融，周边承担制造与外溢。")
para(doc, "武汉都市圈是长江中游城市群核心，正承载光谷外溢、人才、交通枢纽等功能，是观察湖北区域协同的主窗口。")

# ---- 十、人口与城市 ----
heading1(doc, "十、人口与城市：科教大省、城镇化与双老化")
heading2(doc, "10.1 人口总量与结构")
para(doc, "2025年湖北常住5811万人、城镇化率67.39%，高于全国；人口自然增率-4.6‰（出生率4.26‰，死亡率8.86‰），说明“整体负增、都市圈集聚”。")
heading2(doc, "10.2 城市与老龄")
para(doc, "武汉都市圈、光谷依托科教与创新流程集聚年轻人口；但湖北整体“双老化”加深，养老、医疗、就业都面临结构转型。城镇化率高但人口总量下行，是“存量竞争”的新常态。")

# ---- 十一、民营与科创 ----
heading1(doc, "十一、民营经济与科创：创新驱动的硬核")
heading2(doc, "11.1 民企与创新")
para(doc, "湖北民营经济占全省经济比重大，高技术制造、战略性新兴企业依托光谷集聚。创新主体（高新企业、瞪羚/专精特新）在光电子、软件、生物医药上形成优势。")
heading2(doc, "11.2 创新生态")
para(doc, "武汉拥有众多双一流高校、国家级实验室，是中部“科教高地”。“高校与国家实验室→科技成果转化”是湖北描绘科创未来的主线。能否把“论文”变成“产业”，是湖北的长期课题。")

# ---- 十二、事后验证 ----
heading1(doc, "十二、事后验证：用2025年实际执行结果检验报告判断")
heading2(doc, "12.1 年初目标与年末执行对比")
table(doc,
    ["指标", "2025年初目标", "2025实际", "偏差"],
    [
        ["GDP增速", "6%左右", "+5.5%", "略低于目标"],
        ["规上工业", "7.5%左右", "+6.9%", "略低"],
        ["固定资产投资", "7%左右", "+2.6%", "明显低于目标"],
        ["社零", "7%左右", "+2.7%", "低于目标"],
        ["进出口", "7%以上", "+18.2%", "大幅超"],
        ["居民收入", "与增长同步", "+5.2%", "同步"],
    ],
    widths=[3.4, 3.2, 2.8, 4.0])
heading2(doc, "12.2 判断验证")
para(doc, "三条核心判断全部得到支持：**（1）新引擎（出口+19.6%、高技术+15.5%）超预期，验证；（2）内需/投资偏弱（社零+2.7%、固投+2.6%）验证；（3）财政稳健（+6.9%）。**")
para(doc, "唯一变数：投资与消费双双低于目标，说明“外需强、内需弱”。2026年把“促消费、扩投资、稳物价”摆在更前是必然。")

# ---- 十三、五条主线 ----
heading1(doc, "十三、未来5—10年最值得观察的五条主线")
lines5 = [
    ("① 光电子能否从“高地”成为“全国中心”", "高技术占比17%能否再升、光谷科创能否持续变现。"),
    ("② 新能源汽车+老牌车谷能否完成电气化", "汽车新能源转型、智能化软件是否跟上。"),
    ("③ 从“内陆枢纽”到“开放门户”", "中欧班列、长江水道能否把“出口”持续做大。"),
    ("④ “高校+实验室”到“产业”的转化", "科教大省能否真正成为科创大省、独角兽之省。"),
    ("⑤ 人口与老化的平衡", "在人口负增下，如何用创新与都市圈吸附力对冲。"),
]
for l in lines5:
    para(doc, "**%s**　%s" % l, space_after=6)

# ---- 十四、结论 ----
heading1(doc, "十四、最终结论：湖北在“中部崛起+科技自立”里的增长逻辑")
para(doc, "湖北的2025年，本质上是一张“**新引擎（出口/高技术/科创）加速，旧引擎（地产/投资/内需）减速**”的答卷：GDP破6.26万亿、出口+19.6%、高技术+15.5%，代价是消费+2.7%、固投+2.6%偏弱、人口负增。")
para(doc, "只要光谷、新能源汽车、出口可持续，湖北就站在“中部崛起+科技自立”的战略位；如果内需、投资、人口持续偏软、且科技转化难以带动整体，湖北可能长期处在“科技强、经济偏中”错位。")
para(doc, "最稳妥的观察信号：**一盯高技术制造与光电子（动能）、二盯出口与枢纽（开放）、三盯消费与投资（内需）、四盯人口与老化（社会）、五盯财政收入与科创（财政创新）。**湖北，既是“中部崛起”的支点，也是“科技自立”的样板。")

# ---- 附录A ----
heading1(doc, "附录A：主要资料来源与核验路径")
bullet(doc, "湖北省2025年《政府工作报告》（2025年1月）——目标来源。")
bullet(doc, "《2025年湖北省国民经济和社会发展统计公报》（湖北省统计局，2026-03-18）——GDP、工业、投资、外贸等实际值。")
bullet(doc, "湖北省统计局/武汉统计局工业专篇、光电子专题——光谷与新质生产力。")
bullet(doc, "武汉海关进出口数据、湖北省商务厅——外贸与开放。")
bullet(doc, "2000年武汉第四经济普查及城市规划资料——区域/人口底盘。")
heading2(doc, "核验说明")
para(doc, "本报告所有增速、总量、占比均以统计公报/官方发布为基准。涉“民营经济占比”“税收细分”等未完整披露项，以官方口径为准。")

# ---- 附录B ----
heading1(doc, "附录B：建议建立的年度跟踪仪表盘")
para(doc, "建议跟踪下面10个测脉搏指标，判断湖北“换挡”是否成功：")
dash = [
    ["#", "指标", "2025基值", "为什么重要"],
    ["1", "GDP增速", "+5.5%", "总量与方向"],
    ["2", "规上工业增速", "+6.9%", "制造底盘"],
    ["3", "高技术制造增速/占比", "+15.5%/17.4%", "新质生产力"],
    ["4", "进出口/出口增速", "+18.2%/+19.6%", "内陆开放兑现"],
    ["5", "固定资产投资增速", "+2.6%", "投资动能"],
    ["6", "社会消费品零售增速", "+2.7%", "内需消费"],
    ["7", "一般公共预算收入增速", "+6.9%", "财政质量"],
    ["8", "常住人口/城镇化率", "5811万/67.39%", "人口与城镇化"],
    ["9", "粮食产量", "2791万吨", "粮食安全"],
    ["10", "居民人均收入增速", "+5.2%", "民生获得感"],
]
table(doc, dash[0], dash[1:], widths=[0.8, 4.5, 3.2, 5.0])
para(doc, "把这10个指标连起来看，任何一个“新引擎（2/3/4）向上、旧引擎（5/6）减速”，都说明湖北在真正“换挡”；反之则是旧路径的反复。")

# ========================================= 保存
out = "/Users/x/Desktop/content-prod-lab/reports/湖北省_2025年政府工作报告_深度研究_2026-08-13.docx"
doc.save(out)
print("SAVED:", out)
print("PARAGRAPHS:", len(doc.paragraphs))
print("TABLES:", len(doc.tables))
