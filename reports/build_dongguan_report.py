# -*- coding: utf-8 -*-
"""Build 东莞市2025年政府工作报告 深度研究 DOCX, 参照地级市系列版式。"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DARK = RGBColor(0x1F, 0x2A, 0x44)
RED = RGBColor(0xB0, 0x1E, 0x2E)
GRAY = RGBColor(0x59, 0x60, 0x69)
BLUE = RGBColor(0x1F, 0x4E, 0x79)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
HEADER_FILL = "1F2A44"
K = "微软雅黑"


def set_run(run, size=11, bold=False, color=DARK, font=K):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    run.font.name = font
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), font)
    rFonts.set(qn("w:hAnsi"), font)
    rFonts.set(qn("w:eastAsia"), font)


def para(doc, text, size=11, bold=False, color=DARK, align=None,
         space_after=6, space_before=0, font=K):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        set_run(r, size=size, bold=bold or (i % 2 == 1), color=color, font=font)
    return p


def heading1(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(16)
    pf.space_after = Pt(8)
    r = p.add_run(text)
    set_run(r, size=16, bold=True, color=RED)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "B01E2E")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def heading2(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(10)
    pf.space_after = Pt(4)
    r = p.add_run(text)
    set_run(r, size=13, bold=True, color=BLUE)
    return p


def table(doc, headers, rows, widths=None, font_size=9.5):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_run(r, size=font_size, bold=True, color=WHITE)
        tcPr = hdr[i]._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), HEADER_FILL)
        tcPr.append(shd)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(val))
            set_run(r, size=font_size, color=DARK)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    return t


def quote_box(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(8)
    pf.left_indent = Pt(12)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), "B01E2E")
    pBdr.append(left)
    pPr.append(pBdr)
    r = p.add_run(text)
    set_run(r, size=11, color=DARK, font="楷体")
    return p


def bullet(doc, text, size=10.5):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.7)
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        set_run(r, size=size, bold=(i % 2 == 1), color=DARK)
    return p


# ================================================================= 文档主体
doc = Document()
sec = doc.sections[0]
sec.page_width = Cm(21)
sec.page_height = Cm(29.7)
sec.top_margin = Cm(2.4)
sec.bottom_margin = Cm(2.2)
sec.left_margin = Cm(2.5)
sec.right_margin = Cm(2.5)

st = doc.styles["Normal"]
st.font.name = K
st.font.size = Pt(11)
st.element.rPr.rFonts.set(qn("w:eastAsia"), K)


# ---- 封面 ----
para(doc, "东莞市2025年政府工作报告", size=24, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=60, space_after=4)
para(doc, "深度研究与\u201c容易被忽视的细节\u201d分析", size=16, bold=True, color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
para(doc, "从\u201c世界工厂、电子信息、智能手机、先进制造与制造业转型\u201d重新理解东莞", size=12, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)
para(doc, "\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501", color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
para(doc, "研究对象：2025年东莞市政府工作报告及同期财政、国民经济和社会发展统计资料、2026年官方复盘", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
para(doc, "标定日期：2026-08-13", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
doc.add_page_break()

# ---- 目录 ----
catalog = [
    "执行摘要",
    "一、研究口径与阅读方法",
    "二、先看东莞的特殊底盘：世界工厂、电子信息/手机、先进制造与外贸依存",
    "三、最关键的宏观错位：GDP破1.27万亿但低于目标，制造/出口强，消费/房地产弱",
    "四、容易被忽视、但信号很重的细节（15条）",
    "五、2025年GDP目标 vs 实际：对照表",
    "六、2025年增长是谁撑起来的？（结构归因）",
    "七、预算与财政的\u201c含金量\u201d",
    "八、民生底账：人口、收入与城乡",
    "九、城镇与农村：格局与均衡",
    "十、人口流入与流出",
    "十一、物价与货币环境",
    "十二、区域一体化：东莞在\u201c粤港澳大湾区+广深港走廊+制造业\u201d里的位置",
    "十三、未来5—10年最值得观察的五条主线",
    "十四、最终结论：东莞在\u201c电子信息+先进制造+外贸\u201d里的增长逻辑",
    "附录A：主要资料来源与核验路径",
    "附录B：建议建立的年度跟踪仪表盘",
]
para(doc, "目　录", size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=10, space_after=10)
for item in catalog:
    para(doc, item, size=12, space_after=4)
doc.add_page_break()

# ---- 执行摘要 ----
heading1(doc, "执行摘要")
para(doc, "**核心判断**　2025年东莞最显著的是\u201cGDP 12760.20亿元、增长4.0%（低于5%目标）、二产占56.1%\u201d、\u201c规上工业+4.0%（电子信息+6.6%、汽车+32.4%）\u201d、\u201c进出口15794.3亿/+13.8%\u201d、\u201c常住人口1080.04万/城镇化93.63%\u201d。这说明东莞在\u201c世界工厂转向先进制造+外贸\u201d中稳中提质，但**房地产-45.6%、消费+2.8%**偏弱。")
para(doc, "把2025年目标（GDP+5%/规上+7%/社零+5%/财政+3%）、2025年统计、2026年前瞻一起看，东莞是\u201c万亿制造大市+外贸依存高\u201d路径：**电子信息、手机（近5成份额）、外贸+13.8%、工业投资+34.7%**是亮点。")
para(doc, "最容易记住的一句话：**东莞是\u201c世界工厂转向先进制造之城\u201d，靠\u201c电子信息+手机+高端装备+外贸+民营制造\u201d增长。**观察东莞，与其只看\u201cGDP 1.27万亿\u201d，不如看\u201c电子信息+6.6%、汽车+32.4%、进出口+13.8%、工业投资+34.7%占64%、手机出货量近5成\u201d。")# ---- 一、研究口径与阅读方法 ----
heading1(doc, "一、研究口径与阅读方法")
para(doc, "本报告以\u201c东莞市政府工作报告（2025年）\u201d为起点，把\u201c2025年GDP目标（5%以上）\u201d与\u201c2025年国民经济和社会发展统计公报实际值（12760.20亿元/+4.0%）\u201d并置对照，再用2026年计划和政府工作报告的复盘作事后验证。")
para(doc, "重点阅读三类证据：**一是指标目标是否达标**；**二是结构（谁贡献增长）是否匹配目标叙事**；**三是官方复盘如何自评、承认问题**。统计口径一般公共预算收入为地方级。")
para(doc, "统一采用\u201c同比增长%\u201d（GDP为不变价），财政与居民收入多为名义值，文中按原口径转录、不逐项换算。")

# ---- 二、特殊底盘 ----
heading1(doc, "二、先看东莞的特殊底盘：世界工厂、电子信息/手机、先进制造与外贸依存")
para(doc, "**区位与身份**：东莞是广东省地级市、粤港澳大湾区核心节点、广深港科创制造走廊重镇，\u201c世界制造业名城\u201d，采用不设区的镇街制（镇街经济）。")
para(doc, "**产业底盘**：一是电子信息/手机（计算机通信电子设备是最大行业，\u201c三大手机\u201d占国内智能手机近5成）；二是先进制造/高端装备（汽车+32.4%、航空航天+38.7%、生物医药器械+19.8%）；三是民营制造（22万家工业企业、市场主体195万户）；四是高新技术（10200家高企）。")
para(doc, "**人口底盘**：2025年末常住人口1080.04万/+22.96万、城镇化率93.63%，是制造业人口大市。")
para(doc, "**市场与出口**：2025年社零4446.00亿/+2.8%；进出口15794.3亿/+13.8%（出口+9.0%、进口+22.1%）。外贸依存度高是最大特征。")

# ---- 三、核心宏观错位 ----
heading1(doc, "三、最关键的宏观错位：GDP破1.27万亿但低于目标，制造/出口强，消费/房地产弱")
para(doc, "**第一组错位**：2025年GDP目标5%以上，实际12760.20亿/+4.0%**低于目标1pct**；但工业总产值超2.6万亿、规模跃居全省第2，总量破1.27万亿。")
para(doc, "**第二组错位**：规上工业+4.0%（电子信息+6.6%、汽车+32.4%、高技术+6.8%）、进出口+13.8%、工业投资+34.7%；但**固投仅+1.3%、房地产-45.6%、社零+2.8%**。\u201c制造/出口强、消费/地产弱\u201d。")
para(doc, "**第三组错位**：常住1080万/城镇化93.6%、净流入22.96万（人口强）；但**CPI-0.9%**通缩、居民收入增速放缓至+3.7%。")
para(doc, "一句话：**东莞是\u201c制造业强、外贸强、但内需/地产弱、价格通缩\u201d的万亿制造之城**——靠电子信息+出口+工业投资支撑。")

# ---- 四、15条细节 ----
heading1(doc, "四、容易被忽视、但信号很重的细节（15条）")
bullet(doc, "1. **GDP 12760.20亿/+4.0%**：破1.27万亿，但低于5%目标。")
bullet(doc, "2. **电子信息制造+6.6%**：占全国智能手机份额近5成。")
bullet(doc, "3. **汽车制造+32.4%**：比亚迪等新能源车零部件放量。")
bullet(doc, "4. **高技术制造+6.8%、先进制造+6.2%**：制造升级。")
bullet(doc, "5. **进出口+13.8%、出口+9.0%、进口+22.1%**：外贸强、进口电子中间品旺。")
bullet(doc, "6. **民营企业进出口+20.3%占62.9%**：民营外贸主导。")
bullet(doc, "7. **工业投资+34.7%占固投64%**：先进制造/高技术投资爆发。")
bullet(doc, "8. **房地产投资-45.6%塌陷**：房产销售面积-27.4%。")
bullet(doc, "9. **社零+2.8%（通讯器材+105%、网上零售+28.1%）**：线上强、线下弱。")
bullet(doc, "10. **常住1080.04万/城镇化93.63%**：净流入22.96万、人口强。")
bullet(doc, "11. **收入：城镇72137元/+3.4%、农村52133元/+5.5%**，城乡1.38。")
bullet(doc, "12. **CPI-0.9%**：通缩。")
bullet(doc, "13. **财政891.82亿/+3.2%（税收+1.5%）**：财政达标但税收缓。")
bullet(doc, "14. **高企10200家、小巨人236家**：科创制造。")
bullet(doc, "15. **专利76749件/每万人口76.8件**：专利大市。")# ---- 五、目标 vs 实际对照表 ----
heading1(doc, "五、2025年GDP目标 vs 实际：对照表")
table(doc,
    ["指标", "2025年目标", "2025年实际", "是否达标"],
    [
        ["地区生产总值增速", "5%以上", "12760.20亿元/+4.0%", "未达标"],
        ["规上工业增加值", "+7%左右", "+4.0%", "未达标"],
        ["固定资产投资", "提质提效", "+1.3%（工业+34.7%）", "整体低、工业高"],
        ["社会消费品零售总额", "+5%", "4446.00亿元/+2.8%", "未达标"],
        ["一般公共预算收入", "+3%", "891.82亿元/+3.2%", "达标"],
        ["进出口", "稳量增效", "15794.3亿元/+13.8%", "大幅超预期"],
        ["CPI涨幅", "2%左右", "-0.9%", "远低于目标(通缩)"],
    ],
    widths=[3.4, 2.5, 5.0, 3.1])
para(doc, "**简评**：唯一\u201c超目标\u201d的是**进出口+13.8%**与财政；**GDP、规上、消费、地产低于目标、CPI通缩**。这是\u201c制造/外贸强、内需弱\u201d的一年。")

# ---- 六、结构归因 ----
heading1(doc, "六、2025年增长是谁撑起来的？（结构归因）")
para(doc, "**三大产业**：一产36.90亿/+4.5%、二产7165.44亿/+4.4%、三产5557.87亿/+3.5%；二产占56.1%。增长以二产（制造）为主、三产（商贸物流）次之。")
para(doc, "**工业**：规上+4.0%、全部工业+4.8%；**电子信息+6.6%、电气机械+5.7%、汽车+32.4%、高技术+6.8%、先进制造+6.2%**；规上工业利润1209.16亿/-13.4%；工业总产值超2.6万亿。")
para(doc, "**服务业**：金融+6.9%、商贸（线上+28%）、物流快递+16.7%；房地产-3.3%拖累。")
para(doc, "**增长归因**：东莞GDP主要靠**二产（电子信息/汽车/高端制造）+外贸（进出口+13.8%）+工业投资**，房地产、传统消费为负。")

# ---- 七、财政 ----
heading1(doc, "七、预算与财政的\u201c含金量\u201d")
para(doc, "2025年一般公共预算收入891.82亿元/+3.2%（税收649.68亿/+1.5%、非税+9%左右）；支出921.00亿/+1.0%，教育+5.8%、社保+7.1%、公共安全+14.3%。")
para(doc, "**结构性**：收入\u201c达标、但税收缓增、靠非税\u201d；支出+1.0%保民生教育。显示制造业利润-13.4%对税收的拖累。")
para(doc, "**含金量**：东莞财政\u201c税收质量尚可但税收增速低于GDP\u201d，反映工业企业利润承压；支出让位于民生与教育。")

# ---------------- 八、民生 ----------------
heading1(doc, "八、民生底账：人口、收入与城乡")
para(doc, "常住人口1080.04万/+22.96万、城镇化率93.63%；出生率10.32\u2030/自然+6.8\u2030。收入：**城镇72137元/+3.4%、农村52133元/+5.5%**，城乡比1.38。")
para(doc, "消费支出：全体42932元/+3.1%、恩格尔32.2%。城镇新增就业14.45万。")
para(doc, "**民生结论**：收入农村快于城镇、城乡差距极小（1.38）；人口强流入、自然正增长；就业稳、普惠覆盖好。")

# ---------------- 九 ----------------
heading1(doc, "九、城镇与农村：格局与均衡")
para(doc, "东莞城镇化率93.63%（接近直辖市水平），中心城区（南城/东城/莞城）承载总部与商贸；镇街经济（虎门/长安/松山湖）支撑电子信息/手机的\u201c世界工厂\u201d。")
para(doc, "城乡收入比1.38（全国最低一档）、农村收入+5.5%，均衡度高；镇街经济避免\u201c城乡二元\u201d。")

# ---------------- 十 ----------------
heading1(doc, "十、人口流入与流出")
para(doc, "东莞常住1080.04万、净流入22.96万、自然正增长，是广东制造业人口聚集地；户籍345.28万/常住1080万的倒挂体现产业人口无户籍的特征。")
para(doc, "**流入**：制造业/电子/新能源岗位（比亚迪等）、大专蓝领/工程师；**流出**：向粤东西北回流部分低技能。总体\u201c强流入、结构年轻\u201d。")

# ---------------- 十一 ----------------
heading1(doc, "十一、物价与货币环境")
para(doc, "2025年东莞CPI**下降0.9%**（通缩），交通通信-3.2%、食品-1.0%；医疗保健+1.9%。")
para(doc, "金融：存款29664.35亿/+3.9%（住户存款+11.1%）、贷款20036.62亿/+4.0%（制造业+10.7%）；不良率1.56%。\u201c宽中稳、制造业融资旺\u201d。")

# ---------------- 十二 ----------------
heading1(doc, "十二、区域一体化：东莞在\u201c粤港澳大湾区+广深港走廊+制造业\u201d里的位置")
para(doc, "东莞是粤港澳大湾区珠三角制造业核心、广深港科创走廊的中枢：东邻深圳（华为/周边电子），北接广州，南连港澳。跨江有松山湖科学城、滨海湾新区承载。")
para(doc, "制造业协同：华为生态、比亚迪、OPPO/vivo；外贸（一带一路+23.9%、高新技术产品出口+17.8%）。东莞是\u201c制造+开放\u201d枢纽。")

# ---------------- 十三 ----------------
heading1(doc, "十三、未来5—10年最值得观察的五条主线")
bullet(doc, "1. **电子信息+智能手机迭代**：手机近5成份额、6G/算力/AI终端转型。")
bullet(doc, "2. **新能源/汽车+高端装备**：汽车+32.4%、半导体/新能源集群突破。")
bullet(doc, "3. **先进制造+工业技改/数字化转型**：22万企业智改数转。")
bullet(doc, "4. **外贸高依存的多元化**：一带一路/跨境电商/出口结构，对冲摩擦。")
bullet(doc, "5. **房地产与内需修复**：地产-45.6%、CPI-0.9%后的消费/地产企稳。")

# ---------------- 十四 ----------------
heading1(doc, "十四、最终结论：东莞在\u201c电子信息+先进制造+外贸\u201d里的增长逻辑")
para(doc, "**结论**：东莞2025年的\u201c真相\u201d是——**GDP+4%（低于5%）、规上+4%（电子+6.6%/汽车+32.4%）、进出口+13.8%、工业投资+34.7%、房地产-45.6%、CPI-0.9%**。它是\u201c制造强、外贸强、但内需/地产弱\u201d的万亿工厂。")
para(doc, "**对趋势判断**：电子信息/汽车/先进制造/外贸代表东莞的\u201c硬实力\u201d，消费/地产/价格代表\u201c软约束\u201d。**制造升级+外贸韧性**决定中期潜力，**内需修复+产业利润回升**决定韧性。")
para(doc, "**若只看一个指标**：看**工业投资增速（+34.7%）与电子信息制造增速（+6.6%）**——东莞正从\u201c代工/低附加\u201d走向\u201c先进制造+品牌\u201d，智造投入能否兑现为利润与内需是其转型核心。")

# ---------------- 附录A ----------------
heading1(doc, "附录A：主要资料来源与核验路径")
bullet(doc, "东莞市人民政府《2025年东莞市政府工作报告》（2025年）。")
bullet(doc, "东莞市统计局《2025年东莞市国民经济和社会发展统计公报》（2026年5月）。")
bullet(doc, "《2026年东莞国民经济和社会发展计划报告》（2026年，叶惠明）。")
bullet(doc, "广东省2025年统计公报、东莞统计年鉴与镇街数据交叉核验。")

# ---------------- 附录B ----------------
heading1(doc, "附录B：建议建立的年度跟踪仪表盘")
bullet(doc, "GDP总量/增速 vs 全年目标（+/-百分点）。")
bullet(doc, "规上工业增加值及分行业（电子/汽车/电气/高技术）增速。")
bullet(doc, "固定资产投资（总额/工业/先进制造/高技术/房地产）增速。")
bullet(doc, "社会消费品零售总额、网上零售、通讯器材。")
bullet(doc, "进出口总额、出口、进口、一带一路、跨境电商。")
bullet(doc, "一般公共预算收入/税收/非税、财政支出、教育/社保。")
bullet(doc, "常住人口、城镇化率、户籍/常住倒挂、净流入。")
bullet(doc, "CPI/核心CPI、规上工业利润与营收。")
bullet(doc, "手机出货量/三大品牌、高新技术企业、小巨人。")
bullet(doc, "工业总产值、市场主体（企业/个体）增量。")

# ---------------- 保存 ----------------
out = "/Users/x/Desktop/content-prod-lab/reports/东莞市_2025年政府工作报告_深度研究_2026-08-13.docx"
doc.save(out)
print("SAVED OK 东莞市_2025年政府工作报告_深度研究_2026-08-13.docx")