# -*- coding: utf-8 -*-
"""Build 福州市2025年政府工作报告 深度研究 DOCX, 参照省市系列版式。"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DARK = RGBColor(0x1F, 0x2A, 0x44)
RED = RGBColor(0xB0, 0x1E, 0x2E)
GRAY = RGBColor(0x59, 0x60, 0x69)
BLUE = RGBColor(0x1F, 0x4E, 0x79)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
HEADER_FILL = "1F2A44"
K = "微软雅黑"


def set_run(run, size=11, bold=False, color=DARK, font=K):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    run.font.name = font
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), font)
    rFonts.set(qn("w:hAnsi"), font)
    rFonts.set(qn("w:eastAsia"), font)


def para(doc, text, size=11, bold=False, color=DARK, align=None,
         space_after=6, space_before=0, font=K):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        set_run(r, size=size, bold=bold or (i % 2 == 1), color=color, font=font)
    return p


def heading1(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(16)
    pf.space_after = Pt(8)
    r = p.add_run(text)
    set_run(r, size=16, bold=True, color=RED)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "B01E2E")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def heading2(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(10)
    pf.space_after = Pt(4)
    r = p.add_run(text)
    set_run(r, size=13, bold=True, color=BLUE)
    return p


def table(doc, headers, rows, widths=None, font_size=9.5):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_run(r, size=font_size, bold=True, color=WHITE)
        tcPr = hdr[i]._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), HEADER_FILL)
        tcPr.append(shd)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(val))
            set_run(r, size=font_size, color=DARK)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    return t


def quote_box(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(8)
    pf.left_indent = Pt(12)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), "B01E2E")
    pBdr.append(left)
    pPr.append(pBdr)
    r = p.add_run(text)
    set_run(r, size=11, color=DARK, font="楷体")
    return p


def bullet(doc, text, size=10.5):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.7)
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        set_run(r, size=size, bold=(i % 2 == 1), color=DARK)
    return p


# ================================================================= 文档主体
doc = Document()
sec = doc.sections[0]
sec.page_width = Cm(21)
sec.page_height = Cm(29.7)
sec.top_margin = Cm(2.4)
sec.bottom_margin = Cm(2.2)
sec.left_margin = Cm(2.5)
sec.right_margin = Cm(2.5)

st = doc.styles["Normal"]
st.font.name = K
st.font.size = Pt(11)
st.element.rPr.rFonts.set(qn("w:eastAsia"), K)



# ---- 封面 ----
para(doc, "福州市2025年政府工作报告", size=24, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=60, space_after=4)
para(doc, "深度研究与“容易被忽视的细节”分析", size=16, bold=True, color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
para(doc, "从“数字经济、海峡两岸、福州都市圈与民营经济”重新理解福州", size=12, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)
para(doc, "━━━━━━━━━━━━━━━━", color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
para(doc, "研究对象：2025年福州市政府工作报告及同期财政、国民经济和社会发展统计资料、2026年官方复盘", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
para(doc, "标定日期：2026-08-13", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
doc.add_page_break()

# ---- 目录 ----
catalog = [
    "执行摘要",
    "一、研究口径与阅读方法",
    "二、先看福州的特殊底盘：数字经济、海峡两岸、都市圈与海洋经济",
    "三、最关键的宏观错位：GDP破1.51万亿、工业/消费强，但外贸/地产弱",
    "四、容易被忽视、但信号很重的细节（15条）",
    "五、2025年GDP目标 vs 实际：对照表",
    "六、2025年增长是谁撑起来的？（结构归因）",
    "七、预算与财政的“含金量”",
    "八、民生底账：人口、收入与城乡",
    "九、城镇与农村：格局与均衡",
    "十、人口流入与流出",
    "十一、物价与货币环境",
    "十二、区域一体化：福州在“福州都市圈+对台合作+海上丝绸之路”里的位置",
    "十三、未来5—10年最值得观察的五条主线",
    "十四、最终结论：福州在“数字经济+对台/海洋+民营/制造”里的增长逻辑",
    "附录A：主要资料来源与核验路径",
    "附录B：建议建立的年度跟踪仪表盘",
]
para(doc, "目　录", size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=10, space_after=10)
for item in catalog:
    para(doc, item, size=12, space_after=4)
doc.add_page_break()

# ---- 执行摘要 ----
heading1(doc, "执行摘要")
para(doc, "**核心判断**　如果只看新闻标题，2025年福州最显眼的是“GDP破1.51万亿、增长5.6%”、“规上工业+7.2%”、“装备制造+21.6%”和“跨境游客+83.8%”。但这份研究真正值得深读的，是这座“数字经济+海峡两岸+都市圈+海洋”的东南省会，如何在出口大幅下滑（-32.0%）的背景下，靠“工业（化工/电子/汽车）+消费/文旅+数字经济”实现5.6%的增长。")
para(doc, "把2025年初设定的目标（GDP增长5.5%左右）、2025年《国民经济和社会发展统计公报》、2026年复盘放在一起看，福州呈现清晰暗线：**从“传统制造/房地产”的旧底盘，向“数字经济+先进制造+对台/海洋+民营/消费”转型**。旧引擎（外贸/出口、房地产）在调整；新引擎（化工/电子/汽车、数字经济、消费/文旅、对台/海洋）被要求更快补位。")
para(doc, "因此，本报告不按政府工作报告逐段复述，而采用“显性表述—同期数据—制度含义—长期影响”的方式，专门提取容易被忽视、但对判断福州未来5—10年发展模式更有价值的信号。")
para(doc, "最容易记住的一句话：**福州是“数字经济重镇+海峡两岸门户+海洋经济大省+民营经济”的样本，靠“数字经济+制造/工业+消费+对台”撑起增长。**观察福州，与其看“GDP 1.51万亿”，不如看“数字经济7700亿、装备制造、化工/电子、对台、都市圈、海洋经济”这几张名片。")
heading2(doc, "一页速览：2025年福州经济的“表与里”")
table(doc,
    ["维度", "表面现象", "更深层信号"],
    [
        ["增长", "GDP 15112.32亿、+5.6%", "一产5.0%、二产36.0%、三产59.0%"],
        ["产业", "规上工业+7.2%", "化工+30.7%、电子+24.9%、装备+21.6%"],
        ["外贸", "进出口2664.3亿、-24.0%", "出口-32.0%、港口/对台"],
        ["投资", "固定资产投资持平", "工业投资+17.2%、地产-12.7%"],
        ["消费", "社零5969.1亿、+6.0%", "消费品+8.6%、餐饮+6.9%"],
        ["人口", "常住852.1万、城镇化74.63%", "人口+2.0万"],
        ["文旅", "游客14135万（+10.5%）、境外+83.8%", "数字经济+对台海洋"],
    ],
    widths=[2.2, 5.6, 8.2])
para(doc, "", size=10, space_after=2)

# ---- 一、研究口径与阅读方法 ----
heading1(doc, "一、研究口径与阅读方法")
heading2(doc, "1.1 本报告的核心底稿")
bullet(doc, "**2025年《政府工作报告》**（2025年1月）——目标：GDP 5.5%左右、规上工业5.8%、固投5.5%、社零6%+。")
bullet(doc, "**《2025年福州市统计公报》**（市统计局2026-04）——GDP、工业、外贸、人口实数。")
bullet(doc, "**2026年福州政府工作报告/复盘**——对2025执行追认。")
heading2(doc, "1.2 阅读方法：显性—数据—制度—长期")
para(doc, "每章按“**显性表述→同期数据→制度含义→长期影响**”四层展开。")
para(doc, "**关键判别**：数据优先。例2025年GDP目标5.5%、实际5.6%达标；工业目标5.8%、实际7.2%超预期。福州“工业/消费/数字经济强、出口偏弱”。")

# ---- 二、底盘 ----
heading1(doc, "二、先看福州的特殊底盘：数字经济、海峡两岸、都市圈与海洋经济")
para(doc, "福州的地盘，取决于它作为“**数字经济重镇+海峡两岸门户+海洋经济+民营经济大市**”的特殊定位。它是福建/东南省会、“闽都”。")
bullet(doc, "**数字经济**：数字经济规模超7700亿元（2024），数字人民币交易超3500亿；为国家数字经济创新发展试验区核心。")
bullet(doc, "**海峡两岸**：对台先行（新设台企465家）、平潭综合实验区对接口、马尾-马祖；榕台人才/产业/便利。")
bullet(doc, "**工业/制造**：规上工业+7.2%（化工+30.7%、电子+24.9%、汽车+20.1%、钢铁+18.8%、装备+21.6%、高技术+14.3%）；化工/电子“新质生产力”。")
bullet(doc, "**海洋经济/都市圈**：福州都市圈（福州-宁德-莆田-南平），海洋经济、渔业（海洋捕捞/养殖）；湾区（马尾/三江口）。")
para(doc, "**制度含义**：福州把“数字经济、民营经济、对台/海洋/都市圈”当核心资产，面向“数字经济/对台先行区+海上丝绸之路”。")

# ---- 三、宏观错位 ----
heading1(doc, "三、最关键的宏观错位：GDP破1.51万亿、工业/消费强，但外贸/地产弱")
para(doc, "2025年福州GDP 15112.32亿元、+5.6%（一产+4.7%、二产+5.6%、三产+5.6%）。表面“稳中向好”，但拆开看是“**工业/消费/数字经济强、出口/外贸/地产弱**”的错位：")
para(doc, "**强的部分**：规上工业+7.2%（化工+30.7%、电子+24.9%、汽车+20.1%、装备+21.6%）；社零+6.0%；消费/文旅强；数字经济7700亿。")
para(doc, "**弱的部分**：进出口2664.3亿、-24.0%（出口-32.0%）；房地产投资-12.7%；商品房销售下滑；外资-34.6%。")
para(doc, "**核心错位一句话**：福州“工业/消费/数字经济强（装备+21.6%），但出口/外贸/地产弱”。2026年“稳工业/数字经济、修复出口/内需、强对台/海洋”是重点。")
table(doc,
    ["强引擎", "数据", "弱项", "数据"],
    [
        ["规上工业", "+7.2%", "进出口总额", "-24.0%"],
        ["化工/电子制造", "+30.7%/+24.9%", "出口总额", "-32.0%"],
        ["装备制造", "+21.6%", "房地产投资", "-12.7%"],
        ["数字经济", "超7700亿", "实际利用外资", "-34.6%"],
        ["社零消费", "+6.0%", "港口货物吞吐量", "-2.7%"],
    ],
    widths=[3.6, 3.0, 3.6, 3.0])
para(doc, "**错位结论**　福州“强工业/数字经济/消费、弱出口/外贸/地产”。2026年“强数字经济/对台海洋、补出口/外贸”是重点。")

# ---- 四、被忽视细节 ----
heading1(doc, "四、容易被忽视、但信号很重的细节（15条）")
details = [
    ("1", "GDP 15112.32亿、+5.6%", "总量/稳增。"),
    ("2", "规上工业+7.2%（装备+21.6%）", "制造/装备强。"),
    ("3", "化工+30.7%、电子+24.9%、汽车+20.1%", "新质生产力。"),
    ("4", "高技术制造+14.3%、规上工业利润+10.2%", "高技术/效益。"),
    ("5", "数字经济超7700亿、数字人民币", "数字经济。"),
    ("6", "社零+6.0%、网络零售+16.8%", "内需/线上。"),
    ("7", "跨境游客+83.8%、旅游花费+19.4%", "文旅/入境。"),
    ("8", "进出口-24.0%、出口-32.0%", "外贸承压。"),
    ("9", "工业投资+17.2%、二产投资+17.1%", "工业投资。"),
    ("10", "新台资企业465家（对台）", "海峡两岸/对台。"),
    ("11", "海洋/渔业、远洋渔业+14.5%", "海洋经济。"),
    ("12", "常住852万、+2.0万、城镇化74.63%", "人口流入/城镇化。"),
    ("13", "收入54007元、+4.9%、农村+5.8%", "收入稳、农村快。"),
    ("14", "汽车保有/5G用户占比65.9%", "消费/基建。"),
    ("15", "CPI+0.2%（物价温和）", "物价。"),
]
for d in details:
    para(doc, "**%s**" % d[0], size=11)
    para(doc, d[1], size=10.5, color=GRAY)

    para(doc, "**备注**　这条“错位”在福州尤其鲜明：增长靠工业/数字经济/消费，但出口/外贸/地产弱。2026年若出口修复、对台/海洋放大，增长可能从“工业/数字经济”走向“工业+对外/海洋”多极。这条细节，正是福州2026年最难也最值得盯的“变量”。", size=10.5)

# ---- 五、2025年GDP目标 vs 实际 ----
heading1(doc, "五、2025年GDP目标 vs 实际：对照表")
t5 = [
    ["指标", "2025目标", "2025实际", "达标判定"],
    ["GDP增速", "5.5%左右", "+5.6%", "达标"],
    ["规模以上工业", "5.8%", "+7.2%", "超预期"],
    ["固定资产投资", "5.5%", "持平", "基本达标"],
    ["社会消费品零售", "6%以上", "+6.0%", "达标"],
    ["一般公共预算收入", "3%", "+0.01%", "未达标"],
    ["进出口", "正增长", "-24.0%", "大幅未达标"],
]
table(doc, t5[0], t5[1:], widths=[3.6, 3.0, 3.0, 3.6])
para(doc, "**对照结论**　目标“偏求稳”，实际“分化”：GDP 5.6%、工业7.2%、消费6.0%达标；但进出口-24.0%（未达正增长）、财政+0.01%（未达3%）。数字经济/工业/消费接住，外贸/财政是短板。")

# ---- 六、增速分项支撑 ----
heading1(doc, "六、2025年增长是谁撑起来的？（结构归因）")
para(doc, "GDP+5.6%背后，是“**工业/数字经济/消费强、出口/外贸弱**”的结构。把增长贡献拆开看：")
g6 = [
    ["引擎", "贡献/状态", "说明"],
    ["规上工业", "+7.2%", "化工+30.7%、电子+24.9%、装备+21.6%"],
    ["数字经济", "超7700亿", "数字人民币、数字产业"],
    ["消费/社零", "+6.0%", "汽车/文旅/线上"],
    ["文旅/入境", "+10.5%/+83.8%", "游客14135万"],
    ["工业投资", "+17.2%", "二产投资+17.1%"],
    ["出口", "-32.0%", "外贸承压"],
    ["房地产投资", "-12.7%", "地产调整"],
    ["港口吞吐量", "-2.7%", "港口弱"],
    ["人口自然", "常住+2.0万", "人口/总量稳"],
    ["CPI", "+0.2%", "物价温和"],
]
table(doc, g6[0], g6[1:], widths=[3.6, 3.6, 6.0])
para(doc, "**一句话**　增长靠“工业+数字经济+消费+文旅”，出口/外贸/地产是拖累。2026年考验福州“能否修复出口/外贸、放大数字经济/对台海洋”。")

# ---- 七、部门财政 ----
heading1(doc, "七、预算与财政的“含金量”")
para(doc, "2025年福州地方一般公共预算收入**750.55亿元、+0.01%**（支出+1.1%）。")
bullet(doc, "财政收入+0.01%（近零增长）、支出+1.1%。")
bullet(doc, "民生支出（教育/社保/卫生）占比高。")
bullet(doc, "财政“稳收+民生/数字经济/对台支持”。")

# ---- 八、民生与人口 ----
heading1(doc, "八、民生底账：人口、收入与城乡")
para(doc, "2025年福州常住人口**852.1万人、+2.0万**，城镇化率**74.63%**。")
para(doc, "居民人均可支配收入**54007元、+4.9%**，农村（+5.8%）快于城镇（+4.4%）。")

# ---- 九、城镇与农村 ----
heading1(doc, "九、城镇与农村：格局")
para(doc, "福州城镇化率74.63%、人口流入。")
bullet(doc, "农村居民收入+5.8%、快于城镇+4.4%。")
bullet(doc, "共同富裕/沿海平原/丘陵山区协同。")

# ---- 十、人口流入与流出 ----
heading1(doc, "十、人口流入与流出")
para(doc, "福州常住+2.0万（净流入），是东南人口吸引力城市（产业/都市圈/宜居）。")
para(doc, "未来看点：数字经济/对台/都市圈能否持续吸人。“七普”后持续净流入。")

# ---- 十一、物价与货币 ----
heading1(doc, "十一、物价与货币环境")
para(doc, "2025年福州CPI**+0.2%**，物价温和。5G用户占65.9%、宽带升级。")
para(doc, "物价温和、需求平稳，2026年“扩内需/促消费”是主线。")

# ---- 十二、区域一体化 ----
heading1(doc, "十二、区域一体化：福州在“福州都市圈+对台合作+海上丝绸之路”里的位置")
para(doc, "福州的核心战略坐标是“**福州都市圈+海峡两岸融合+海上丝绸之路（21世纪海丝中枢）**”。")
bullet(doc, "福州都市圈：福州-宁德-莆田-南平，引领闽东北。")
bullet(doc, "对台：平潭、马尾-马祖、对台先行区（台企465家）。")
bullet(doc, "海上丝绸之路：21世纪海丝中枢、海洋经济。")
bullet(doc, "数字经济：国家数字经济创新发展试验区。")
para(doc, "若“数字经济+对台/海洋+都市圈”成势，福州将是东南增长极与对台/海丝门户。")

# ===== 十三、五条主线 =====
heading1(doc, "十三、未来5—10年最值得观察的五条主线")
lines5 = [
    ("① 数字经济/数字人民币", "数字经济能否成为增长极。"),
    ("② 对台合作/海丝", "对台、平潭、海洋经济。"),
    ("③ 先进制造/化工/电子", "化工电子汽车装配能否壮大。"),
    ("④ 民营经济/都市圈", "民营、福州都市圈，人口协同。"),
    ("⑤ 出口/外贸/内需修复", "出口偏弱、外贸修复，扩内需。"),
]
for l in lines5:
    para(doc, "**%s**　%s" % l, space_after=6)

# ---- 十四、结论 ----
heading1(doc, "十四、最终结论：福州在“数字经济+对台/海洋+民营/制造”里的增长逻辑")
para(doc, "福州2025年，本质上是“**数字经济/工业/消费核心、出口/外贸/地产弱**”的答卷：GDP15112.32亿、+5.6%、工业+7.2%、数字经济超7700亿、社零+6.0%，但进出口-24.0%、出口-32.0%、地产-12.7%、外资-34.6%。")
para(doc, "只要数字经济、对台/海洋、先进制造持续，福州就站在“东南增长极+对台门户”的位；如果出口/地产偏弱，福州需承受“制造/消费强、外贸/楼市弱”的结构挑战。")
para(doc, "最稳观察信号：**一盯数字经济（引擎）、二盯对台/海丝（开放）、三盯制造/工业（底盘）、四盯出口/外贸（约束）、五盯都市圈人口（长期）。**福州，是“数字经济+对台门户”的新样本。")

# ---- 附录A ----
heading1(doc, "附录A：主要资料来源与核验路径")
bullet(doc, "福州市2025年政府工作报告（2025年1月）——目标来源。")
bullet(doc, "《2025年福州市国民经济和社会发展统计公报》（市统计局）——GDP、工业、外贸、人口实值。")
bullet(doc, "2026年福州市政府工作报告——2025复盘/数字经济/对台。")
bullet(doc, "福州海关（外贸）、市财政。")
heading2(doc, "核验说明")
para(doc, "本报告涉及以统计公报/官方口径为准。涉“数字经济/对台/海洋/都市圈”等以官方口径为准。")

# ---- 附录B ----
heading1(doc, "附录B：建议建立的年度跟踪仪表盘")
para(doc, "建议把下面10个‘测脉搏’指标做成年度跟踪表：")
dash = [
    ["#", "指标", "2025基值", "为什么重要"],
    ["1", "GDP增速", "+5.6%", "总量与方向"],
    ["2", "规上工业增速", "+7.2%", "制造底盘"],
    ["3", "数字经济规模", "超7700亿", "数字经济"],
    ["4", "进出口/出口增速", "-24.0%/-32.0%", "外贸韧性"],
    ["5", "固定资产投资/地产", "持平/-12.7%", "投资结构"],
    ["6", "社零增速", "+6.0%", "内需消费"],
    ["7", "常住人口/城镇化", "852万/74.63%", "人口与城市"],
    ["8", "地方财政收入", "+0.01%", "财政质量"],
    ["9", "跨境游客/对台", "+83.8%/台企465", "文旅/对台"],
    ["10", "CPI", "+0.2%", "物价"],
]
table(doc, dash[0], dash[1:], widths=[0.8, 4.6, 3.4, 4.4])
para(doc, "把这10个指标连起来看，数字经济/制造/对台（3/9）、对外修复（4）、人口（7），都说明福州在真正换挡。")

# ============ 保存 ============
out = "/Users/x/Desktop/content-prod-lab/reports/福州市_2025年政府工作报告_深度研究_2026-08-13.docx"
doc.save(out)
print("SAVED:", out)
print("PARAGRAPHS:", len(doc.paragraphs))
print("TABLES:", len(doc.tables))
