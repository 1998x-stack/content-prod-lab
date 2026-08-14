# -*- coding: utf-8 -*-
"""Build 曲靖市2025年政府工作报告 深度研究 DOCX, 参照地级市系列版式。"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DARK = RGBColor(0x1F, 0x2A, 0x44)
RED = RGBColor(0xB0, 0x1E, 0x2E)
GRAY = RGBColor(0x59, 0x60, 0x69)
BLUE = RGBColor(0x1F, 0x4E, 0x79)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
HEADER_FILL = "1F2A44"
K = "微软雅黑"


def set_run(run, size=11, bold=False, color=DARK, font=K):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    run.font.name = font
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), font)
    rFonts.set(qn("w:hAnsi"), font)
    rFonts.set(qn("w:eastAsia"), font)


def para(doc, text, size=11, bold=False, color=DARK, align=None,
         space_after=6, space_before=0, font=K):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        set_run(r, size=size, bold=bold or (i % 2 == 1), color=color, font=font)
    return p


def heading1(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(16)
    pf.space_after = Pt(8)
    r = p.add_run(text)
    set_run(r, size=16, bold=True, color=RED)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "B01E2E")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def heading2(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(10)
    pf.space_after = Pt(4)
    r = p.add_run(text)
    set_run(r, size=13, bold=True, color=BLUE)
    return p


def table(doc, headers, rows, widths=None, font_size=9.5):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_run(r, size=font_size, bold=True, color=WHITE)
        tcPr = hdr[i]._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), HEADER_FILL)
        tcPr.append(shd)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(val))
            set_run(r, size=font_size, color=DARK)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    return t


def quote_box(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(8)
    pf.left_indent = Pt(12)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), "B01E2E")
    pBdr.append(left)
    pPr.append(pBdr)
    r = p.add_run(text)
    set_run(r, size=11, color=DARK, font="楷体")
    return p


def bullet(doc, text, size=10.5):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.7)
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        set_run(r, size=size, bold=(i % 2 == 1), color=DARK)
    return p


# ================================================================= 文档主体
doc = Document()
sec = doc.sections[0]
sec.page_width = Cm(21)
sec.page_height = Cm(29.7)
sec.top_margin = Cm(2.4)
sec.bottom_margin = Cm(2.2)
sec.left_margin = Cm(2.5)
sec.right_margin = Cm(2.5)

st = doc.styles["Normal"]
st.font.name = K
st.font.size = Pt(11)
st.element.rPr.rFonts.set(qn("w:eastAsia"), K)


# ---- 封面 ----
para(doc, "曲靖市2025年政府工作报告", size=24, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=60, space_after=4)
para(doc, "深度研究与\u201c容易被忽视的细节\u201d分析", size=16, bold=True, color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
para(doc, "从\u201c滇东能源、烟草、煤化工、新能源、云南副中心\u201d重新理解曲靖", size=12, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)
para(doc, "\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501", color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
para(doc, "研究对象：2025年曲靖市政府工作报告及同期财政、国民经济和社会发展统计资料、2026年官方复盘", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
para(doc, "标定日期：2026-08-14", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
doc.add_page_break()

# ---- 目录 ----
catalog = [
    "执行摘要",
    "一、研究口径与阅读方法",
    "二、先看曲靖的特殊底盘：滇东能源、烟草、煤化工、新能源（光伏风电）与云南副中心",
    "三、最关键的宏观错位：GDP 3777.6亿/3.5%低于5%目标，新能源快但传统能源、烟草、地产弱",
    "四、容易被忽视、但信号很重的细节（15条）",
    "五、2025年GDP目标 vs 实际：对照表",
    "六、2025年增长是谁撑起来的？（结构归因）",
    "七、预算与财政的\u201c含金量\u201d",
    "八、民生底账：人口、收入与城乡",
    "九、城镇与农村：格局与均衡",
    "十、人口流入与流出",
    "十一、物价与货币环境",
    "十二、区域一体化：曲靖在\u201c滇东+云南副中心+面向南亚东南亚\u201d里的位置",
    "十三、未来5—10年最值得观察的五条主线",
    "十四、最终结论：曲靖在\u201c新能源+煤化工+烟草+磷\u201d里的增长逻辑",
    "附录A：主要资料来源与核验路径",
    "附录B：建议建立的年度跟踪仪表盘",
]
para(doc, "目　录", size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=10, space_after=10)
for item in catalog:
    para(doc, item, size=12, space_after=4)
doc.add_page_break()

# ---- 执行摘要 ----
heading1(doc, "执行摘要")
para(doc, "**核心判断**　2025年曲靖最显著的是\u201cGDP 3777.6亿元、增长3.5%（低于5%目标）、云南第2\u201d、\u201c规上工业+2.3%（新能源电池+45.9%、有色金属+9.4%）\u201d、\u201c进出口43.48亿/+50.2%\u201d、\u201c但固投-9.5%、社零+1.0%、财收-1.6%、CPI-0.1%\u201d、\u201c常住558.3万/城镇化53.77%\u201d。这说明曲靖在\u201c能源+烟草+新能源\u201d的转型中，**新能源、外贸亮但传统能源、地产、消费偏弱**。")
para(doc, "把2025年目标（GDP+5%左右/规上+6%以上/产业投资+7%/社零+6%以上/财收+2%左右）、2025年统计（GDP+3.5%大幅低于目标、规上+2.3%、固投-9.5%、社零+1.0%、财收-1.6%）、趋势一起看，曲靖是\u201c滇东能源+新能源\u201d路径：**能源（原煤3246万吨）、烟草、有色金属、煤化工、新能源电池、光伏风电**是支柱；2025年总量3777.6亿居云南第2（次于昆明）。")
para(doc, "最容易记住的一句话：**曲靖是\u201c云南副中心城市、滇东能源与烟草基地\u201d，靠\u201c煤+电+烟+新能源\u201d增长，正从\u201c煤电\u201d转向\u201c新能源+精细化工\u201d。**观察曲靖，与其只看\u201cGDP 3777亿\u201d，不如看\u201c锂电+418.6%、新能源电池+45.9%、进出口+50.2%、新能源装机895万千瓦\u201d。")

# ---- 一、研究口径与阅读方法 ----
heading1(doc, "一、研究口径与阅读方法")
para(doc, "本报告以\u201c曲靖市政府工作报告（2025年）\u201d为起点，把\u201c2025年GDP目标（5%左右）\u201d与\u201c官方2025年GDP（3777.6亿元/+3.5%）\u201d并置对照，并用\u201c2025年曲靖市统计公报\u201d和\u201c2026年政府工作报告复盘\u201d横向核验。")
para(doc, "口径提示：\u201cGDP、规上工业增加值增速\u201d按可比价（实际增速）；\u201c投资、消费、进出口、财政收入\u201d按现价（名义增速）。涉及\u201c常住人口\u201d用统计公报常住口径（558.3万），城镇化率用官方公布值（53.77%）。")
para(doc, "指标体系（与研究口径一致）：五大象限——**总量与增速（GDP）、产业动能（能源/烟草/煤化工/新能源）、固投、财政质量、民生与人口**。")
para(doc, "特别提示（不吃老本）：曲靖是**云南第2经济大市（副中心）、滇东能源基地（煤电+新能源）**，但2025年GDP增速明显放缓（3.5%）；它不只是\u201c煤+烟\u201d，正在\u201c新能源电池、硅铝精深、精细磷化工\u201d发力——真正要看的是\u201c传统能源承压、新能源接棒\u201d的转型阵痛。")
# ---- 二、先看曲靖的特殊底盘 ----
heading1(doc, "二、先看曲靖的特殊底盘：滇东能源、烟草、煤化工、新能源（光伏风电）与云南副中心")
para(doc, "曲靖地处滇东高原、云南东北，是**云南副中心城市、滇东能源基地、烟草基地**。2025年GDP 3777.6亿元、常住558.3万，云南第2，人均6.74万元。")
para(doc, "五个底盘名词，先立框架：")
bullet(doc, "**滇东能源基地**　2025年原煤3246.19万吨（云南主要产煤区）、发电量395.45亿千瓦时（火电196亿/水电86.6亿/风电82亿/光电30.8亿）——\u201c滇东煤电\u201d。")
bullet(doc, "**烟草**　曲靖卷烟厂（全国行业卓越级智能工厂），烟草是利税与财政大户。")
bullet(doc, "**有色金属/磷煤**　十种有色金属131.48万吨（+15.4%）、电解铝54.76万吨、锌64.36万吨；磷/煤化工。")
bullet(doc, "**新能源/新能源电池**　新能源电池+45.9%、锂离子电池530.69万千瓦时（+418.6%）；新能源装机895万千瓦（2025新增134.6万）。")
bullet(doc, "**云南副中心**　\u201c3815\u201d战略、\u201c一中心一样板两区\u201d、面向南亚东南亚辐射节点。")
para(doc, "这五根（能源+烟草+有色化工+新能源+副中心）构成曲靖独特底盘：**左手煤电/烟草（传统基盘），右手新能源+精细化工（接棒）**。理解曲靖，先理解\u201c滇东能源大市\u201d的禀赋与转型压力。")

# ---- 三、最关键的宏观错位 ----
heading1(doc, "三、最关键的宏观错位：GDP 3777.6亿/3.5%低于5%目标，工业、固投、财收、消费全面偏弱")
para(doc, "2025年曲靖最需要辨析的一组\u201c错位\u201d：**GDP 3.5%大幅低于5%目标、规上工业+2.3%、固投-9.5%、社零+1.0%、财收-1.6%、CPI-0.1%**。")
para(doc, "为什么\u201c新能源、外贸爆发\u201d，经济却在明显放缓？三个解释：")
para(doc, "**其一，工业整体平淡**　规上工业+2.3%：烟草+1.2%（个位）、能源+3.3%、非烟非能+2.0%；虽有色+9.4%、电子+10.0%、新能源+45.9%，但体量不足、煤电/火电-17.7%拖累。")
para(doc, "**其二，固投、地产、财政\u201c三弱\u201d**　固投-9.5%、地产投资-13.3%、财收-1.6%、财收总收-3.1%——投资与财政承压。")
para(doc, "**其三，消费物价弱**　社零+1.0%（低于6%目标）、CPI-0.1%——内需不足、物价走弱。")
para(doc, "小结：曲靖2025年是\u201c**传统能源承压、弱投资财政消费、新能源亮**\u201d的一年：新能源、外贸是亮点，但整体增长、投资、财政、消费全面低于目标/转负。")

# ---- 四、容易被忽视、但信号很重的细节（15条） ----
heading1(doc, "四、容易被忽视、但信号很重的细节（15条）")
bullet(doc, "**1.锂离子电池+418.6%、新能源电池+45.9%**　新能源电池（储能/动力）是2025年最大增长（能源/锂电集群）。")
bullet(doc, "**2.电解铝+16.8%、十种有色金属+15.4%**　有色/铝精深加工是第二增长极。")
bullet(doc, "**3.进出口43.48亿/+50.2%**　对东盟/湄公河、\u201cRCEP+80%\u201d，内陆开放点。")
bullet(doc, "**4.新能源装机895万千瓦（+145.6万/年）**　风电/太阳能+装机大增，绿色能源转型。")
bullet(doc, "**5.发电量395.45亿千瓦时（水电+61.5%、光电+21.6%，火电-17.7%）**　\u201c水绿增、火电减\u201d。")
bullet(doc, "**6.原煤3246万吨，能源保供稳**　滇东\u201c煤电\u201d基本盘。")
bullet(doc, "**7.烟草基本盘（规上+1.2%）**　烟草税利支撑财政/就业（全国行业级智能工厂）。")
bullet(doc, "**8.烟叶/食粮：生猪出栏1008.8万头、粮338.8万吨**　\u201c粮食+生猪\u201d农业底。")
bullet(doc, "**9.民营经济增加值占57.8%**\u201c民间投资渠道、产业投资占比54%。\u201d")
bullet(doc, "**10.社零+1.0%（1357.32亿）、乡村+1.3%**\u201c线上+75.5%亮，线下低迷。\u201d")
bullet(doc, "**11.财收164.21亿/-1.6%，税收+1.6%**\u201c财政收遇阻，但税收结构转正。\u201d")
bullet(doc, "**12.居民收入33224元/+3.9%、城乡比2.26**\u201c农村+6.0%快于城镇+2.8%\u201d城乡差大（2.26）。")
bullet(doc, "**13.常住558.3万/城镇化53.77%**\u201c云南人口大市，但城镇化低于全国（低于全国与全省）。\u201d")
bullet(doc, "**14.CPI-0.1%**\u201c物价弱、通缩隐忧。\u201d")
bullet(doc, "**15.生态：森林覆盖率47.1%、PM2.5 20μg、空气质量99.5%**\u201c绿电+生态\u201d底色。\u201d")
# ---- 五、2025年GDP目标 vs 实际：对照表 ----
heading1(doc, "五、2025年GDP目标 vs 实际：对照表")
table(doc,
    ["指标", "2025年目标", "2025年实际", "达成评价"],
    [
        ["地区生产总值(GDP)", "增长5%左右", "3777.6亿/3.5%", "未达成，差1.5pct"],
        ["规上工业增加值", "增长6%以上", "+2.3%", "未达成，差3.7pct"],
        ["固定资产投资", "——", "-9.5%", "大幅下行"],
        ["社会消费品零售总额", "增长6%以上", "1357.32亿/+1.0%", "未达成，差5pct"],
        ["地方一般公共预算收入", "增长2%左右", "164.21亿/-1.6%", "未达成，转负"],
        ["进出口总额", "——", "43.48亿/+50.2%", "大幅超额"],
        ["居民收入", "快于经济增长", "33224元/+3.9%", "快于GDP"],
    ],
)
para(doc, "注：GDP、规上工业按可比价；投资、消费、财收按现价。**进出口（+50.2%）大幅超额、居民收入（+3.9%）快于GDP**；**GDP、规上、社零、财收均低于目标**（固投-9.5%、财收-1.6%）。")
para(doc, "拆读：**新能源电池（+45.9%）、有色（+15.4%）、进出口（+50.2%）是亮色**；**传统能源（火电-17.7%）、固投（-9.5%）、社零（+1.0%）、财收（-1.6%）偏弱**——\u201c转型起步、整体承压\u201d，是滇东能源城市样本。")

# ---- 六、2025年增长是谁撑起来的？（结构归因） ----
heading1(doc, "六、2025年增长是谁撑起来的？（结构归因）")
para(doc, "把曲靖GDP的3.5%拆开：三次产业分别增1.0%、1.5%、5.7%（结构14.3：34.0：51.7）。**第三产业（服务业）是唯一较强项（+5.7%），第二产业（工业）弱，第一产业（农业）稳**。")
para(doc, "2026年曲靖强调\u201c十五五\u201d、云南副中心、新能源，聚焦**新能源电池、硅铝精深、精细磷化工、先进制造、文旅**——核心是\u201c稳能源基本盘、做好新能源接棒\u201d。")
para(doc, "**第二产业（工业）**：规上+2.3%、新能源电池+45.9%、有色/铝+15.4%、电子+10.0%；但烟草+1.2%、煤电/火电-17.7%、钢铁-8.3%——\u201c新能快、传统弱\u201d。")
para(doc, "**第三产业（服务业）**：+5.7%（交通/物流/商贸），外贸（进出口+50.2%）、文旅（旅游收入+3.8%）。")
para(doc, "**第一产业（农业）**：粮食338.8万吨、生猪+1008.8万头——农业稳。")
para(doc, "一句话归因：**2025年曲靖增长\u201c靠服务业+新能源电池+有色\u201d**，但传统能源、固投、财政、消费承压；\u201c转型阵痛期\u201d。")

# ---- 七、预算与财政的\u201c含金量\u201d ----
heading1(doc, "七、预算与财政的\u201c含金量\u201d")
para(doc, "2025年曲靖**一般公共预算收入164.21亿元（-1.6%）**，其中税收117.44亿元（+1.6%）、税收占比约71.5%；辖区财政总收入455.26亿元（-3.1%）；支出526.45亿元（-3.6%）。")
bullet(doc, "税收结构：税收+1.6%转正、烟草税利支撑；但非税/土地承压、财收整体-1.6%。")
bullet(doc, "民生支出：社保就业+12.1%、卫生健康+13.8%、节能环保+38.7%、住房保障+5.0%——民生投入加大。")
bullet(doc, "金融支撑：存款4075.21亿（+4.8%）、贷款2955.83亿（+9.4%，中长期+18.4%）——信贷宽对制造/基建。")
para(doc, "**财政含金量小结**：财收-1.6%转负（税收+1.6%），\u201c税收尚可、总量承压\u201d；民生支出增长快；财政对\u201c新能源、硅铝、有色\u201d投入加大。")

# ---- 八、民生底账：人口、收入与城乡 ----
heading1(doc, "八、民生底账：人口、收入与城乡")
para(doc, "2025年曲靖**居民人均可支配收入33224元（+3.9%）**，其中城镇48072元（+2.8%）、农村21281元（+6.0%），城比2.26。就业：城镇新增就业4.55万人。")
para(doc, "人口画像：**常住558.3万、城镇化53.77%（+0.82pct）**，云南人口第2大市；户籍598.24万（净流出），城镇化低于全国/全省。")
para(doc, "民生投入：医疗救助、养老、兜底；粮食338.8万吨、生猪1008.8万头保障\u201c菜篮子\u201d。")

# ---- 九、城镇与农村：格局与均衡 ----
heading1(doc, "九、城镇与农村：格局与均衡")
para(doc, "曲靖常住城镇化率53.77%（云南地州偏高但低于全国），农村体量大；农村收入增速（+6.0%）远高于城镇（+2.8%），但**城乡绝对差距大（约2.26倍）**。")
para(doc, "农业底盘：**粮食338.8万吨（云南粮仓）**、蔬菜、生猪出栏1008.8万头、烤烟——\u201c滇东大农业\u201d。")
para(doc, "一句话：\u201c曲靖农业粮猪烟强、农村收入快（+6%）\u201d，但\u201c城镇化低、城乡差大\u201d，需\u201c以城带乡、乡村振兴\u201d。")

# ---- 十、人口流入与流出 ----
heading1(doc, "十、人口流入与流出")
para(doc, "曲靖常住558.3万（户籍598.24万）、城镇化53.77%，是\u201c户籍>常住\u201d的劳务输出型人口大市；滇东劳动年龄人口外流至昆明、珠三角等。")
para(doc, "结构观察：**城镇化率低于全国/全省**，人口总体净流出、自然增长率-1.32‰（出生率8.04‰<死亡率9.36‰）。")
para(doc, "2026年目标：以\u201c副中心+新能源产业\u201d留人、聚人（2026新增就业4.2万）——\u201c人从县乡到城、从外流到回流\u201d是曲靖的人口命题。")

# ---- 十一、物价与货币环境 ----
heading1(doc, "十一、物价与货币环境")
para(doc, "2025年曲靖**CPI同比-0.1%**（交通通信-3.2%、居住-0.1%、食品烟酒-0.6%；衣着+1.3%、教育+1.8%）——**低通胀、需求偏弱**。")
bullet(doc, "信贷扩张：存款4075.21亿（+4.8%）、贷款2955.83亿（+9.4%、中长期+18.4%）——中长期贷款宽、支撑项目。")
bullet(doc, "新能源车保有6.87万辆（+50.8%）、燃油相关下行——\u201c绿电消费\u201d。")
para(doc, "货币环境判断：**宽信用、CPI负（-0.1%）**；\u201c中长期信贷宽、物价偏弱\u201d，曲靖需\u201c扩内需、稳物价\u201d。")

# ---- 十二、区域一体化：曲靖的位置 ----
heading1(doc, "十二、区域一体化：曲靖在\u201c滇东+云南副中心+面向南亚东南亚\u201d里的位置")
para(doc, "曲靖是**云南副中心城市、滇东城市群、面向南亚东南亚辐射\u201c滇东门户\u201d**，毗邻贵州、广西。")
bullet(doc, "**云南副中心**　\u201c3815\u201d战略、\u201c一中心一样板两区\u201d（制造+生态+示范），云南第2经济。")
bullet(doc, "**滇东门户**　濒临贵州、对接珠三角，高铁（沪昆/广昆）、可向西江/北部湾出海。")
bullet(doc, "**面向南亚东南亚**　进出口+50.2%、对东盟/湄公河流域、RCEP——\u201c沿边开放\u201d节点。")
para(doc, "一句话：**曲靖在\u201c滇东+云南副中心\u201d里，最核心的定位是\u201c云南第二大经济中心、滇东能源与制造枢纽\u201d**——副中心、能源、区位是最大优势。")

# ---- 十三、未来5—10年最值得观察的五条主线 ----
heading1(doc, "十三、未来5—10年最值得观察的五条主线")
bullet(doc, "**主线一：新能源接棒（硅锂/储能）**\u201c锂电+418%、新能源装机895万\u201d能否撑成\u201c绿色能源强市\u201d。")
bullet(doc, "**主线二：硅铝精深/高端制造**\u201c电解铝+16.8%、有色+15.4%\u201d能否走向\u201c新型材+装备\u201d。")
bullet(doc, "**主线三：精细磷/煤化工**\u201c磷、煤化工集群\u201d能否\u201c降碳增值\u201d。")
bullet(doc, "**主线四：能源保供与碳减排**\u201c火电-17.7%\u201d怎么在\u201c保供、转型\u201d间平衡。")
bullet(doc, "**主线五：人口与城镇化**\u201c558万人口、城镇化54%\u201d能否靠\u201c副中心+产业\u201d回流聚人。")

# ---- 十四、最终结论 ----
heading1(doc, "十四、最终结论：曲靖在\u201c新能源+煤化工+烟草+磷\u201d里的增长逻辑")
para(doc, "把2025年闭环打穿：**曲靖是\u201c云南副中心、滇东能源\u201d**：GDP 3777.6亿/+3.5%、规上+2.3%、新能源电池+45.9%、进出口+50.2%、财收-1.6%。")
para(doc, "曲靖不是\u201c只有煤和烟\u201d——它在\u201c新能源电池、硅铝、精细磷化工、面向东盟\u201d发力；但传统能源、投资、财政、消费承压，\u201c转型阵痛\u201d。")
para(doc, "一句话结论：**曲靖是\u201c滇东能源与烟草基地、云南副中心\u201d；观察它先看\u201c新能源电池、硅铝、光伏风电\u201d，再看\u201c煤电、固投、财收、消费\u201d。**它是\u201c新能源亮、传统弱、转型阵痛\u201d的云南样本。")

# ---- 附录A ----
heading1(doc, "附录A：主要资料来源与核验路径")
bullet(doc, "《2025年曲靖市政府工作报告》（2025年2月，2025年目标、2024年回顾）")
bullet(doc, "《2025年曲靖市国民经济和社会发展统计公报》（曲靖市统计局，2026-05-15，2025年实际数据）")
bullet(doc, "《2026年曲靖市政府工作报告》（2026年2月，2025年复盘+2026年目标）")
bullet(doc, "曲靖市政府官网、曲靖市统计局（qj.gov.cn）")

# ---- 附录B ----
heading1(doc, "附录B：建议建立的年度跟踪仪表盘")
bullet(doc, "GDP总量/增速 vs 全年目标（可比价）。")
bullet(doc, "规上工业增加值及分行业（烟草/能源/有色/新能源/电子）增速。")
bullet(doc, "新能源装机容量、新能源电池产能、锂电产量。")
bullet(doc, "原煤产量、发电量/结构（火电/水电/风光）。")
bullet(doc, "固定资产投资/产业/基建/房地产投资增速。")
bullet(doc, "社会消费品零售总额、线上、新能源车。")
bullet(doc, "进出口、对东盟/RCEP。")
bullet(doc, "一般公共预算收入、税收/非税、支出结构。")
bullet(doc, "常住人口、城镇化率、城镇新增就业。")
bullet(doc, "CPI、金融存贷款、中长期贷款。")

# ---------------- 保存 ----------------
out = "/Users/x/Desktop/content-prod-lab/reports/曲靖市_2025年政府工作报告_深度研究_2026-08-14.docx"
doc.save(out)
print("SAVED OK: 曲靖市", out)
