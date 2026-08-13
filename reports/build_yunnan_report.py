# -*- coding: utf-8 -*-
"""Build 云南省2025年政府工作报告 深度研究 DOCX, 参照省市系列版式。"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DARK = RGBColor(0x1F, 0x2A, 0x44)
RED = RGBColor(0xB0, 0x1E, 0x2E)
GRAY = RGBColor(0x59, 0x60, 0x69)
BLUE = RGBColor(0x1F, 0x4E, 0x79)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
HEADER_FILL = "1F2A44"
K = "微软雅黑"


def set_run(run, size=11, bold=False, color=DARK, font=K):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    run.font.name = font
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), font)
    rFonts.set(qn("w:hAnsi"), font)
    rFonts.set(qn("w:eastAsia"), font)


def para(doc, text, size=11, bold=False, color=DARK, align=None,
         space_after=6, space_before=0, font=K):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        set_run(r, size=size, bold=bold or (i % 2 == 1), color=color, font=font)
    return p


def heading1(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(16)
    pf.space_after = Pt(8)
    r = p.add_run(text)
    set_run(r, size=16, bold=True, color=RED)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "B01E2E")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def heading2(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(10)
    pf.space_after = Pt(4)
    r = p.add_run(text)
    set_run(r, size=13, bold=True, color=BLUE)
    return p


def table(doc, headers, rows, widths=None, font_size=9.5):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_run(r, size=font_size, bold=True, color=WHITE)
        tcPr = hdr[i]._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), HEADER_FILL)
        tcPr.append(shd)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(val))
            set_run(r, size=font_size, color=DARK)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    return t


def quote_box(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(8)
    pf.left_indent = Pt(12)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), "B01E2E")
    pBdr.append(left)
    pPr.append(pBdr)
    r = p.add_run(text)
    set_run(r, size=11, color=DARK, font="楷体")
    return p


def bullet(doc, text, size=10.5):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.7)
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        set_run(r, size=size, bold=(i % 2 == 1), color=DARK)
    return p


# ================================================================= 文档主体
doc = Document()
sec = doc.sections[0]
sec.page_width = Cm(21)
sec.page_height = Cm(29.7)
sec.top_margin = Cm(2.4)
sec.bottom_margin = Cm(2.2)
sec.left_margin = Cm(2.5)
sec.right_margin = Cm(2.5)

st = doc.styles["Normal"]
st.font.name = K
st.font.size = Pt(11)
st.element.rPr.rFonts.set(qn("w:eastAsia"), K)



# ---- 封面 ----
para(doc, "云南省2025年政府工作报告", size=24, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=60, space_after=4)
para(doc, "深度研究与“容易被忽视的细节”分析", size=16, bold=True, color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
para(doc, "从“绿色能源、生物多样性、旅游大省、面向南亚东南亚门户与乡村振兴”重新理解云南", size=12, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)
para(doc, "━━━━━━━━━━━━━━━━", color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
para(doc, "研究对象：2025年云南省政府工作报告及同期财政、国民经济和社会发展统计资料、2026年官方复盘", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
para(doc, "标定日期：2026-08-13", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
doc.add_page_break()

# ---- 目录 ----
catalog = [
    "执行摘要",
    "一、研究口径与阅读方法",
    "二、先看云南的特殊底盘：绿色能源、生物多样性、旅游/乡村振兴、面向南亚东南亚",
    "三、最关键的宏观错位：GDP破3.28万亿、绿能/电子强，但投资/地产/商贸偏弱",
    "四、容易被忽视、但信号很重的细节（15条）",
    "五、2025年GDP目标 vs 实际：对照表",
    "六、2025年增长是谁撑起来的？（结构归因）",
    "七、预算与财政的“含金量”",
    "八、民生底账：人口、收入与城乡",
    "九、城镇与农村：格局与均衡",
    "十、人口流入与流出",
    "十一、物价与货币环境",
    "十二、区域一体化：云南在“面向南亚东南亚+中老铁路+西部大开发”里的位置",
    "十三、未来5—10年最值得观察的五条主线",
    "十四、最终结论：云南在“绿色能源+旅游+面向南亚东南亚门户”里的增长逻辑",
    "附录A：主要资料来源与核验路径",
    "附录B：建议建立的年度跟踪仪表盘",
]
para(doc, "目　录", size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=10, space_after=10)
for item in catalog:
    para(doc, item, size=12, space_after=4)
doc.add_page_break()

# ---- 执行摘要 ----
heading1(doc, "执行摘要")
para(doc, "**核心判断**　如果只看新闻标题，2025年云南最显眼的是“GDP破3.28万亿、增长4.1%”、“进出口+10.2%（进口+15.1%）”、“清洁能源发电占87.6%”和“电子行业+22.7%”。但这份研究真正值得深读的，是这座“绿色能源+旅游+面向南亚东南亚+生物多样性”的边疆大省，如何在固定资产投资（-7.0%）、房地产开发（-2.9%）与人口自然负增（-1.25‰）的背景下，靠“绿电/铝硅/电子+旅游+进出口”实现4.1%的增长。")
para(doc, "把2025年初设定的目标（GDP增长5%左右）、2025年《国民经济和社会发展统计公报》、以及2026年报告对2025年的复盘放在一起看，云南呈现清晰暗线：**从“烟草/传统农业”的旧底盘，向“绿色能源+绿色铝/硅+电子/新材料+旅游+面向南亚东南亚”转型**。旧引擎（烟草、传统农业、地产）在调整；新引擎（绿电、铝/硅光伏、电子信息、旅游、RCEP/中老铁路）被要求更快补位。")
para(doc, "因此，本报告不按政府工作报告逐段复述，而采用“显性表述—同期数据—制度含义—长期影响”的方式，专门提取容易被忽视、但对判断云南未来5—10年发展模式更有价值的信号。")
para(doc, "最容易记住的一句话：**云南是“绿色能源+高原特色农业+旅游+面向南亚东南亚/一带一路”的样本，靠“绿电+制造业+旅游+开放”撑起增长。**观察云南，与其看“GDP 3.28万亿”，不如看“清洁能源、绿色铝、电子、旅游、中老铁路、面向南亚东南亚”这几张名片。")
heading2(doc, "一页速览：2025年云南经济的“表与里”")
table(doc,
    ["维度", "表面现象", "更深层信号"],
    [
        ["增长", "GDP 32765.78亿、+4.1%", "一产13.2%、二产32.0%、三产54.8%"],
        ["产业", "规上工业+4.5%", "电子+22.7%、绿铝破2000亿"],
        ["外贸", "进出口2737.36亿、+10.2%", "进口+15.1%、中老铁路+42.9%"],
        ["投资", "固定资产投资-7.0%", "产业占52.1%、民间-4.0%"],
        ["财政", "地方一般公共预算收入2243.3亿、+2.3%", "税收+6.6%"],
        ["消费", "社零12786.21亿、+2.4%", "通讯+46.9%、餐饮+4.2%"],
        ["人口", "常住4644万、城镇化55.02%", "自然增-1.25‰、脱贫县+6.7%"],
        ["开放", "面向南亚东南亚门户/中老铁路/黄金大通道", "进出口+10.2%"],
    ],
    widths=[2.2, 5.4, 8.4])
para(doc, "", size=10, space_after=2)

# ---- 一、研究口径与阅读方法 ----
heading1(doc, "一、研究口径与阅读方法")
heading2(doc, "1.1 本报告的三份核心底稿")
bullet(doc, "**2025年《政府工作报告》**：2025年1月在省十四届人大三次会议上作，给出全年预期目标（GDP增长5%左右、产业投资增长7%左右、地方一般公共预算收入增长2%）。")
bullet(doc, "**2025年《国民经济和社会发展统计公报》**：云南省统计局2026年3月31日发布，给出全年实际执行数与增速，是“事后验证”的锚点。")
bullet(doc, "**2026年云南省政府工作报告及官方复盘**：对2025执行结果的追认，交叉验证“目标—执行—再评价”三段式。")
heading2(doc, "1.2 阅读方法：显性—数据—制度—长期")
para(doc, "每一章按“**显性表述 → 同期数据 → 制度含义 → 长期影响**”四层展开。其中“制度含义”回答“政府为什么这么排优先序”，“长期影响”回答“这对未来5—10年意味着什么”。")
para(doc, "关键判别：**当报告使用的“增长词”和统计公报里的实际数不一致时，数据优先。**例如2025年GDP目标5%左右、实际4.1%略低于目标；产业投资目标7%左右、实际-5.1%、未达；地方财政收入目标2%、实际+2.3%达标。差异反映：云南“绿电/电子/旅游/边贸强，投资/地产/人口偏弱”。")
heading2(doc, "1.3 指标取舍与单位说明")
para(doc, "本报告统一以“2025年统计公报+2026年报告复盘”为口径，增速均为可比价，绝对数为人民币。GDP以初步核算为准。")

# ---- 二、底盘 ----
heading1(doc, "二、先看云南的特殊底盘：绿色能源、生物多样性、旅游/乡村振兴、面向南亚东南亚")
para(doc, "云南的地盘，取决于它作为“**绿色能源/清洁能源大省+高原特色农业+旅游/生物多样性+面向南亚东南亚门户**”的特殊定位。它是中国“绿色能源+旅游”的边疆样本。")
bullet(doc, "**绿色能源**：清洁能源发电占87.6%（水电71.7%、太阳能+33.7%）；绿色铝工业总产值2003亿元首破2000亿；硅光伏、绿铝、新能源电池产业合计+16.9%。")
bullet(doc, "**生物多样性/高原农业**：咖啡17.07万吨（+13.6%）、茶叶61.15万吨（+4.3%）、中药材+14.0%；高原特色农业、生物多样性（西双版纳等）。")
bullet(doc, "**旅游大省**：2024年旅游总花费1.14万亿元、接待入境过夜游客+168%；旅居389.75万人（+20.7%）（公报未单列2025人次但旅游业投资+5.0%）。")
bullet(doc, "**面向南亚东南亚**：中老铁路货物544.1万吨（+13.9%）、货值266.8亿（+42.9%）“黄金大通道”；RCEP/边民互市。")
para(doc, "**制度含义**：云南把“绿色能源、高原农业、旅游/生态、面向南亚东南亚开放”当核心资产，深入推进“绿色经济强省+面向南亚东南亚辐射中心”。民营经济占GDP 53.5%。")

# ---- 三、宏观错位 ----
heading1(doc, "三、最关键的宏观错位：GDP破3.28万亿、绿能/电子强，但投资/地产/商贸偏弱")
para(doc, "2025年云南GDP 32765.78亿元、+4.1%（一产+3.1%、二产+2.9%、三产+5.1%）。表面看“稳中有进”，但拆开看是“**绿电/电子/边贸强、投资/地产/商贸弱**”的错位：")
para(doc, "**强的部分**：规上工业+4.5%（电子+22.7%、有色冶炼+14.2%、绿铝+16.9%）；进出口+10.2%（进口+15.1%、中老铁路+42.9%）；清洁能源发电占87.6%；规上工业利润+5.0%。")
para(doc, "**弱的部分**：固定资产投资-7.0%（地产-2.9%、二产-4.9%、三产-9.0%）；房地产开发-2.9%；社零+2.4%（相对低）；人口自然增-1.25‰。")
para(doc, "**核心错位一句话**：云南“绿电/电子/铝硅/边贸强（电子+22.7%、中老铁路+42.9%），但投资/地产/商贸弱”。2026年“稳投资/地产、扩内需/旅游、续绿能/电子/面向南亚东南亚”是重点。")
table(doc,
    ["强引擎", "数据", "弱项", "数据"],
    [
        ["电子行业", "+22.7%", "固定资产投资", "-7.0%"],
        ["进出口总额", "+10.2%", "房地产开发投资", "-2.9%"],
        ["清洁能源发电占比", "87.6%", "社会消费品零售", "+2.4%"],
        ["绿色铝工业总产值", "超2000亿", "人口自然增长率", "-1.25‰"],
        ["规上工业利润", "+5.0%", "PPI", "-0.7%"],
    ],
    widths=[3.8, 3.0, 3.6, 3.0])
para(doc, "**错位结论**　云南的增长“强绿能/电子/边贸、弱投资/地产”。2026年“稳投资/扩内需、强能源/电子、拓南亚东南亚”是重点。")

# ---- 四、被忽视细节 ----
heading1(doc, "四、容易被忽视、但信号很重的细节（15条）")
details = [
    ("1", "GDP 32765.78亿、+4.1%", "总量/稳增。"),
    ("2", "规上工业+4.5%、电子+22.7%", "制造/电子强。"),
    ("3", "清洁能源发电占87.6%（太阳能+33.7%）", "绿色能源/绿电。"),
    ("4", "绿色铝工业总产值2003亿、首破2000亿", "绿铝/硅光伏。"),
    ("5", "硅光伏+绿铝+新能源电池材料+16.9%", "新材料集群。"),
    ("6", "规上工业利润+5.0%、营收利润率6.05%", "工业盈利好。"),
    ("7", "进出口2737.36亿、+10.2%、进口+15.1%", "外贸强、边贸。"),
    ("8", "中老铁路货物+13.9%、货值+42.9%", "黄金大通道/走向东南亚。"),
    ("9", "绿色铝/能源工业投资+9.4%", "能源投资。"),
    ("10", "粮食2001.89万吨、首破2000万吨", "粮食/农业强。"),
    ("11", "咖啡+13.6%、中药材+14.0%", "高原特色农业。"),
    ("12", "常住4644万、城镇化率55.02%（+0.91pct）", "人口/城镇化。"),
    ("13", "居民收入31311元、+4.6%、农村+6.1%", "收入稳、脱贫县+6.7%。"),
    ("14", "脱贫县农村居民收入+6.7%、快于全省", "乡村振兴。"),
    ("15", "民营经济占GDP 53.5%、旅游花费万亿", "民营/旅游/生态。"),
]
for d in details:
    para(doc, "**%s**" % d[0], size=11)
    para(doc, d[1], size=10.5, color=GRAY)

    para(doc, "**备注**　这条“错位”在云南尤其鲜明：增长靠绿能/电子/旅游/边贸，但投资/地产/商贸弱。2026年若投资修复、面向南亚东南亚/中老铁路深化，增长可能从“制造/边贸单极”走向“制造+内需/投资”多极。这条细节，正是云南2026年最难也最值得盯的“变量”。", size=10.5)

# ---- 五、2025年目标 vs 实际 对照 ----
heading1(doc, "五、2025年GDP目标 vs 实际：对照表")
t5 = [
    ["指标", "2025目标", "2025实际", "达标判定"],
    ["GDP增速", "5%左右", "+4.1%", "略低于目标"],
    ["产业投资增速", "7%左右", "-5.1%", "大幅未达标"],
    ["地方一般公共预算收入增速", "2%", "+2.3%", "达标"],
    ["城镇调查失业率", "5.5%以内", "5.2%", "达标"],
    ["城镇新增就业", "50万人以上", "51.1万人", "达标"],
    ["粮食产量", "1955万吨以上", "2001.89万吨", "大幅超额"],
]
table(doc, t5[0], t5[1:], widths=[3.4, 3.0, 3.0, 3.6])
para(doc, "**对照结论**　目标“求稳”，实际“有分化”：GDP 4.1%（略低于5%目标）、产业投资-5.1%（未达7%）；但财政收入+2.3%、失业率5.2%、就业51.1万、粮食2001.89万吨（超1955万吨）均达标。绿电/电子/边贸接住，投资是短板。")

# ---- 六、增速分项支撑 ----
heading1(doc, "六、2025年增长是谁撑起来的？（结构归因）")
para(doc, "GDP+4.1%背后，是“**绿电/电子/边贸强、投资/地产弱**”的结构。把增长贡献拆开看：")
g6 = [
    ["引擎", "贡献/状态", "说明"],
    ["规上工业", "+4.5%", "电子+22.7%、有色冶炼+14.2%"],
    ["绿色铝/硅光伏", "+16.9%", "绿铝破2000亿"],
    ["进出口/边贸", "+10.2%", "中老铁路货值+42.9%"],
    ["旅游/生态", "旅游花费万亿", "旅居389.75万、乡村振兴"],
    ["工业利润", "+5.0%", "盈利改善"],
    ["房地产开发", "-2.9%", "地产调整"],
    ["固定资产投资", "-7.0%", "三产/地产拖累"],
    ["社会消费品零售", "+2.4%", "商贸偏弱"],
    ["人口自然增长率", "-1.25‰", "人口约束"],
    ["PPI", "-0.7%", "物价偏弱"],
]
table(doc, g6[0], g6[1:], widths=[3.6, 3.6, 6.0])
para(doc, "**一句话**　增长靠“绿色铝/硅+电子+进出口+旅游”，投资/地产是最大拖累。2026年考验云南“能不能让投资稳定、扩大内需/旅游”。")

# ---- 七、预算与财政 ----
heading1(doc, "七、预算与财政的“含金量”")
para(doc, "2025年云南地方一般公共预算收入**2243.3亿元、+2.3%**，其中税收**+6.6%**。")
bullet(doc, "财政收入+2.3%、税收+6.6%，税收快于总量（质量改善）。")
bullet(doc, "财政“稳收+民生（节能环保+21.0%、社保+8.8%、卫生+7.2%）”。")
bullet(doc, "支持绿能、旅游、面向南亚东南亚开放。")

# ---- 八、民生底账 ----
heading1(doc, "八、民生底账：人口、收入与城乡")
para(doc, "2025年云南常住人口**4644万人**，城镇化率**55.02%、+0.91pct**。自然增长率**-1.25‰**。")
para(doc, "居民人均可支配收入**31311元、+4.6%**，农村（+6.1%）快于城镇（+3.1%），脱贫县**+6.7%**。")
g8 = [
    ["指标", "数值", "信号"],
    ["常住人口", "4644万", "人口规模大"],
    ["自然增长率", "-1.25‰", "自然负增"],
    ["城镇化率", "55.02%/+0.91pct", "稳步城镇化"],
    ["居民人均可支配收入", "31311元/+4.6%", "收入稳、脱贫县快"],
]
table(doc, g8[0], g8[1:], widths=[4.4, 4.0, 4.6])
para(doc, "**民生观察**：脱贫县农村收入增速（+6.7%）领跑，城乡差距缩小；但人口自然负增、城镇化率较低（55%）是短板。")

# ---- 九、城镇与农村 ----
heading1(doc, "九、城镇与农村：格局与均衡")
para(doc, "云南城镇化率55.02%、+0.91pct，城乡收入比缩小、农村快于城镇。")
bullet(doc, "农村居民收入+6.1%、快于城镇+3.1%。")
bullet(doc, "脱贫县农村居民收入+6.7%、快于全省平均。")
bullet(doc, "人口向昆明都市圈/沿边开放区集中。")

# ---- 十、人口流入与流出 ----
heading1(doc, "十、人口流入与流出")
para(doc, "云南2025年常住4644万、自然增-1.25‰但总量大（净流入相对少）。昆明、滇中都市圈、沿边口岸吸引人口。")
para(doc, "未来看点：绿色能源/高原农业/旅游/面向南亚东南亚能否留住人口。")

# ---- 十一、物价与货币 ----
heading1(doc, "十一、物价与货币环境")
para(doc, "2025年云南CPI**持平（0.0%）**、PPI**-0.7%**，物价温和、工业品价格偏弱。")
para(doc, "物价平稳反映“供需基本平衡”，与全国低通胀一致。2026年“扩内需、稳价格”是主线。")

# ---- 十二、区域一体化 ----
heading1(doc, "十二、区域一体化：云南在“面向南亚东南亚+中老铁路+西部大开发”里的位置")
para(doc, "云南的核心战略坐标是“**面向南亚东南亚辐射中心+中老铁路黄金大通道+西部大开发**”，是中国与东南亚开放的门户。")
bullet(doc, "面向南亚东南亚：RCEP、边民互市、口岸经济。")
bullet(doc, "中老铁路：货物544.1万吨、货值266.8亿（+42.9%）。")
bullet(doc, "绿色能源：向外输绿电、清洁能源。")
bullet(doc, "生态/生物多样性：高原生态屏障、旅游大省。")
para(doc, "若“绿能+面向南亚东南亚+旅游”闭环跑通，云南将作为中国西南开放与绿色经济强省。")

# ---- 十三、五条主线 ----
heading1(doc, "十三、未来5—10年最值得观察的五条主线")
lines5 = [
    ("① 绿色能源/绿铝/硅", "绿电、绿铝、硅光伏，能否持续壮大。"),
    ("② 面向南亚东南亚/中老铁路", "中老铁路、面向东盟能否放大开放。"),
    ("③ 旅游/生物多样性/养老旅居", "旅游万亿、旅居、生物多样性。"),
    ("④ 乡村振兴/高原农业", "脱贫县、咖啡/茶叶/中药材。"),
    ("⑤ 投资/内需/人口", "补投资/扩内需、稳人口。"),
]
for l in lines5:
    para(doc, "**%s**　%s" % l, space_after=6)

# ---- 十四、最终结论 ----
heading1(doc, "十四、最终结论：云南在“绿色能源+旅游+面向南亚东南亚”里的增长逻辑")
para(doc, "云南的2025年，本质上是“**绿色能源/绿铝/电子/旅游核心，投资/地产/商贸弱**”的答卷：GDP32765.78亿、+4.1%、规上工业+4.5%、电子+22.7%、进出口+10.2%、绿色能源占87.6%，但固投-7.0%、地产-2.9%、社零+2.4%。")
para(doc, "只要绿色能源、绿铝/硅、面向南亚东南亚、旅游、乡村振兴持续，云南就站在“西南绿色开放”的位；如果投资/地产/人口偏弱，云南需承受“产业/开放强、需求偏弱”的挑战。")
para(doc, "最稳观察信号：**一盯绿色能源/绿铝硅（绿能）、二盯中老铁路/南亚东南亚（开放）、三盯旅游/生物多样性（消费）、四盯投资/内需（约束）、五盯人口/乡村振兴（长期）。**云南，是“绿电+开放+旅游”的新样本。")

# ---- 附录A ----
heading1(doc, "附录A：主要资料来源与核验路径")
bullet(doc, "云南省2025年《政府工作报告》——目标来源。")
bullet(doc, "《2025年云南省国民经济和社会发展统计公报》（云南统计局，2026-03-31）——GDP、工业、外贸、人口实值。")
bullet(doc, "2026年云南省政府工作报告——2025执行复盘+绿能/旅游/边贸。")
bullet(doc, "昆明海关、省财政厅——外贸/财政。")
heading2(doc, "核验说明")
para(doc, "本报告所有增速、总量、占比均以统计公报/官方口径为基准，涉及“绿能/绿铝/硅/旅游/面向南亚东南亚”以官方为口径。")
para(doc, "（注）旅游总人次与旅游总花费、对南亚东南亚进出口、生物多样性专项，公报未单列，以文旅厅/海关/林草局官方为准。")

# ---- 附录B ----
heading1(doc, "附录B：建议建立的年度跟踪仪表盘")
para(doc, "建议把下面10个“测脉搏”指标做成年度跟踪表：")
dash = [
    ["#", "指标", "2025基值", "为什么重要"],
    ["1", "GDP增速", "+4.1%", "总量与方向"],
    ["2", "规上工业增速", "+4.5%", "制造底盘"],
    ["3", "电子/绿铝增速", "+22.7%/+16.9%", "新动能"],
    ["4", "进出口/中老铁路增速", "+10.2%/+42.9%", "外贸/开放"],
    ["5", "固定资产投资/地产", "-7.0%/-2.9%", "投资结构"],
    ["6", "社零增速", "+2.4%", "内需消费"],
    ["7", "常住人口/城镇化率", "4644万/55.02%", "人口与城市"],
    ["8", "地方财政收入增速", "+2.3%", "财政质量"],
    ["9", "清洁能源发电占比", "87.6%", "绿色能源"],
    ["10", "脱贫县农村收入增速", "+6.7%", "乡村振兴/民生"],
]
table(doc, dash[0], dash[1:], widths=[0.8, 4.6, 3.4, 4.4])
para(doc, "把这10个指标连起来看，任何一个“绿能/电子/农业（3/9）、开放（4）、旅游、投资（5）修复”，都说明云南在真正换挡。")

# ============ 保存 ============
out = "/Users/x/Desktop/content-prod-lab/reports/云南省_2025年政府工作报告_深度研究_2026-08-13.docx"
doc.save(out)
print("SAVED:", out)
print("PARAGRAPHS:", len(doc.paragraphs))
print("TABLES:", len(doc.tables))
