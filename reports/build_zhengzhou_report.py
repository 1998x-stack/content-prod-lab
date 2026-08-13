# -*- coding: utf-8 -*-
"""Build 郑州市2025年政府工作报告 深度研究 DOCX, 参照省市系列版式。"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DARK = RGBColor(0x1F, 0x2A, 0x44)
RED = RGBColor(0xB0, 0x1E, 0x2E)
GRAY = RGBColor(0x59, 0x60, 0x69)
BLUE = RGBColor(0x1F, 0x4E, 0x79)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
HEADER_FILL = "1F2A44"
K = "微软雅黑"


def set_run(run, size=11, bold=False, color=DARK, font=K):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    run.font.name = font
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), font)
    rFonts.set(qn("w:hAnsi"), font)
    rFonts.set(qn("w:eastAsia"), font)


def para(doc, text, size=11, bold=False, color=DARK, align=None,
         space_after=6, space_before=0, font=K):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        set_run(r, size=size, bold=bold or (i % 2 == 1), color=color, font=font)
    return p


def heading1(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(16)
    pf.space_after = Pt(8)
    r = p.add_run(text)
    set_run(r, size=16, bold=True, color=RED)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "B01E2E")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def heading2(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(10)
    pf.space_after = Pt(4)
    r = p.add_run(text)
    set_run(r, size=13, bold=True, color=BLUE)
    return p


def table(doc, headers, rows, widths=None, font_size=9.5):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_run(r, size=font_size, bold=True, color=WHITE)
        tcPr = hdr[i]._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), HEADER_FILL)
        tcPr.append(shd)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(val))
            set_run(r, size=font_size, color=DARK)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    return t


def quote_box(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(8)
    pf.left_indent = Pt(12)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), "B01E2E")
    pBdr.append(left)
    pPr.append(pBdr)
    r = p.add_run(text)
    set_run(r, size=11, color=DARK, font="楷体")
    return p


def bullet(doc, text, size=10.5):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.7)
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        set_run(r, size=size, bold=(i % 2 == 1), color=DARK)
    return p


# ================================================================= 文档主体
doc = Document()
sec = doc.sections[0]
sec.page_width = Cm(21)
sec.page_height = Cm(29.7)
sec.top_margin = Cm(2.4)
sec.bottom_margin = Cm(2.2)
sec.left_margin = Cm(2.5)
sec.right_margin = Cm(2.5)

st = doc.styles["Normal"]
st.font.name = K
st.font.size = Pt(11)
st.element.rPr.rFonts.set(qn("w:eastAsia"), K)



# ---- 封面 ----
para(doc, "郑州市2025年政府工作报告", size=24, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=60, space_after=4)
para(doc, "深度研究与\u201c容易被忽视的细节\u201d分析", size=16, bold=True, color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
para(doc, "从\u201c国家中心城市、交通枢纽、先进制造与中原龙头\u201d重新理解郑州", size=12, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)
para(doc, "\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501", color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
para(doc, "研究对象：2025年郑州市政府工作报告及同期财政、国民经济和社会发展统计资料、2026年官方复盘", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
para(doc, "标定日期：2026-08-13", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
doc.add_page_break()

# ---- 目录 ----
catalog = [
    "执行摘要",
    "一、研究口径与阅读方法",
    "二、先看郑州的特殊底盘：国家中心城市、交通枢纽、先进制造与中原龙头",
    "三、最关键的宏观错位：GDP破1.52万亿、工业/外贸强，但固投/地产/财政偏弱",
    "四、容易被忽视、但信号很重的细节（15条）",
    "五、2025年GDP目标 vs 实际：对照表",
    "六、2025年增长是谁撑起来的？（结构归因）",
    "七、预算与财政的\u201c含金量\u201d",
    "八、民生底账：人口、收入与城乡",
    "九、城镇与农村：格局与均衡",
    "十、人口流入与流出",
    "十一、物价与货币环境",
    "十二、区域一体化：郑州在\u201c郑州都市圈+中部崛起+空中/陆上丝路\u201d里的位置",
    "十三、未来5—10年最值得观察的五条主线",
    "十四、最终结论：郑州在\u201c制造+枢纽+国家中心城市\u201d里的增长逻辑",
    "附录A：主要资料来源与核验路径",
    "附录B：建议建立的年度跟踪仪表盘",
]
para(doc, "目　录", size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=10, space_after=10)
for item in catalog:
    para(doc, item, size=12, space_after=4)
doc.add_page_break()

# ---- 执行摘要 ----
heading1(doc, "执行摘要")
para(doc, "**核心判断**　如果只看新闻标题，2025年郑州最显著的是\u201cGDP破1.52万亿、增长5.4%\u201d、\u201c规上工业+9.0%（9个国家中心城市首位）\u201d、\u201c进出口6501.8亿、+16.8%（省会第4）\u201d、\u201c数字经济破8000亿\u201d。但这份研究真正值得深读的，是这座\u201c国家中心城市+交通枢纽+先进制造\u201d的中原龙头，如何在房地产投资-8.2%、外资-56.9%背景下，靠\u201c工业（电子/汽车/战新）+外贸+交通枢纽\u201d实现5.4%的增长。")
para(doc, "把2025年初设定的目标（GDP增长6%左右）、2025年《国民经济和社会发展统计公报》、2026年复盘放在一起看，郑州呈现清晰暗线：**从\u201c地产/传统制造\u201d的旧底盘，向\u201c电子信息/汽车/战新+交通枢纽+未来产业\u201d升级**。工业与外贸是亮点，地产/财政/总量（5.4%<6%）偏弱。")
para(doc, "因此，本报告不按政府工作报告逐段复述，而采用\u201c显性表述—同期数据—制度含义—长期影响\u201d的方式，专门提取容易被忽略、但对判断郑州未来5—10年发展模式更有价值的信号。")
para(doc, "最容易记住的一句话：**郑州是\u201c国家中心城市+交通枢纽+先进制造\u201d的中原龙头，靠\u201c工业（电子/汽车/战新）+外贸/枢纽+数字经济\u201d撑起增长。**观察郑州，与其看\u201cGDP 1.52万亿\u201d，不如看\u201c规上工业9%、进出口6500亿、数字经济8000亿、机场货邮100万吨、中欧班列\u201d这几张名片。")
heading2(doc, "一页速览：2025年郑州经济的\u201c表与里\u201d")
table(doc,
    ["维度", "表面现象", "更深层信号"],
    [
        ["增长", "GDP 15244.6亿、+5.4%", "一产1.2%、二产36.6%、三产62.2%"],
        ["产业", "规上工业+9.0%", "电子+16.2%、汽车+11.9%、战新+11.8%"],
        ["外贸", "进出口6501.8亿、+16.8%", "出口+20.3%、跨境电商1690亿"],
        ["投资", "固投/工业+17.9%、第三产-1.8%", "房投-8.2%、地产调整"],
        ["消费", "社零6629.4亿、+5.0%", "可穿戴+117%、智能手机+101.6%"],
        ["人口", "常住1313.8万、城镇化81.52%", "人口净增、自然增2.19‰"],
        ["枢纽", "机场货邮103.3万吨、中欧班列", "空中丝路+陆上丝路+国际港"],
    ],
    widths=[2.2, 5.6, 7.4])
para(doc, "", size=10, space_after=2)

# ---- 一、研究口径与阅读方法 ----
heading1(doc, "一、研究口径与阅读方法")
heading2(doc, "1.1 本报告的核心底稿")
bullet(doc, "**2025年《政府工作报告》**（2025年2月）——目标：GDP 6%左右、规上工业10%以上、固投8%、社零6%+、财政2%。")
bullet(doc, "**《2025年郑州市统计公报》**（市统计局2026-04）——GDP、工业、外贸、人口实数。")
bullet(doc, "**2026年郑州市政府工作报告/复盘**（2026年2月）——2025追认与电子信息/交通枢纽展望。")
heading2(doc, "1.2 阅读方法：显性—数据—制度—长期")
para(doc, "**关键判别**：数据优先。例2025年GDP目标6%、实际5.4%略欠；规上工业目标10%、实际9.0%略欠；但进出口+16.8%超预期。郑州\u201c工业/外贸强、地产/财政偏弱\u201d，穿透总量看制造与枢纽。")

# ---- 二、底盘 ----
heading1(doc, "二、先看郑州的特殊底盘：国家中心城市、交通枢纽、先进制造与中原龙头")
para(doc, "郑州的地盘取决于它作为\u201c**国家中心城市+综合交通枢纽+先进制造基地+中原龙头**\u201d的特殊定位。它是河南省会、人口大市（常住1313.8万）、中部崛起核心。")
bullet(doc, "**国家中心城市**：9大中心城市之一，中原城市群龙头，省域虹吸强。")
bullet(doc, "**交通枢纽**：米字形高铁、郑州机场货邮103.3万吨（冷链枢纽）、中欧班列3417班；航空+陆港双枢纽。")
bullet(doc, "**先进制造**：电子信息（富士康/超聚变）破9000亿、汽车（比亚迪/宇通/上汽）120万辆、超硬材料国家级集群。")
bullet(doc, "**数字经济**：数字经济规模突破8000亿、算力7.8万P、国家超算互联网核心节点。")
para(doc, "这一底板决定郑州2025成绩单\u201c底色\u201d：**只要电子/汽车/枢纽/数字经济持续，郑州就站在\u201c制造+枢纽\u201d增长极；若地产/财政承压，需承受\u201c工业外贸强、地产/收入弱\u201d的结构挑战。**")

# ---- 三、宏观错位 ----
heading1(doc, "三、最关键的宏观错位：GDP破1.52万亿、工业/外贸强，但固投/地产/财政偏弱")
para(doc, "郑州2025年最值得咀嚼的错位，是\u201c**工业/外贸/枢纽强、固投/地产/财政偏弱**\u201d。这种错位决定了对这座中部国家中心城市的观察不能只看GDP增速。")
bullet(doc, "**GDP**：15244.6亿、+5.4%（目标6%略欠）。一产189.9亿（+3.7%，占比1.2%）、二产5576.8亿（+5.4%，占比36.6%）、三产9477.8亿（+5.4%，占比62.2%）。")
bullet(doc, "**工业**：规上工业+9.0%（国家中心城市首位）；电子+16.2%、汽车+11.9%、战新+11.8%、高技术+14.9%。新能源汽车+16.9%。")
bullet(doc, "**外贸**：进出口6501.8亿、+16.8%（出口4275.5亿、+20.3%；进口+10.5%）；跨境电商1690亿+10.9%。")
bullet(doc, "**消费**：社零6629.4亿、+5.0%；可穿戴+117%、智能手机+101.6%、体育娱乐+86.1%。")
bullet(doc, "**固投/财政**：固投工业+17.9%、第三产-1.8%；房投-8.2%、地产调整；财政+2.3%。")
para(doc, "**为什么读这条**：郑州作为\u201c国家中心+枢纽\u201d，结构性矛盾是\u201c外贸/制造/枢纽强、地产/财政/总量偏弱\u201d。总量承压但新质制造弹性足。")

# ---- 四、15条细节 ----
heading1(doc, "四、容易被忽视、但信号很重的细节（15条）")
para(doc, "以下15条，都在2025年报告/公报，但被\u201cGDP 1.52万亿\u201d等掩盖。它们是判断郑州2025之后5—10年的关键小信号。")
bullet(doc, "**1. 规上工业+9.0%（国家中心城市居首）**：工业制造为主引擎。")
bullet(doc, "**2. 电子信息破9000亿**：富士康/超聚变/海康，向万亿迈进。")
bullet(doc, "**3. 汽车产量120.5万辆（新能源汽车高增）**：新能源车之城（比亚迪/宇通/上汽）。")
bullet(doc, "**4. 进出口6501.8亿、+16.8%（省会第4）**：中部开放高地、富士康出口链。")
bullet(doc, "**5. 跨境电商1690亿**：郑州跨境电商综试区全国10强。")
bullet(doc, "**6. 中欧班列（郑州）开行3417班、机场货邮103.3万吨**：空中+陆上丝路双枢纽。")
bullet(doc, "**7. 数字经济破8000亿、算力7.8万P**：超算互联网核心节点。")
bullet(doc, "**8. 常住1313.8万、自然增2.19‰**：人口大市+正增长。")
bullet(doc, "**9. 未来产业超2000亿**：量子/氢能/人形机器人/低空/脑机接口。")
bullet(doc, "**10. 高新技术产业+10.9%、战新+11.8%**：硬科技落地。")
bullet(doc, "**11. 集成电路+71%、传感器+54.2%、锂电+42.9%、服务器+19.2%**：新产业产品放量。")
bullet(doc, "**12. 社会研发强度2.8%+**：中部科创。")
bullet(doc, "**13. 空中/路丝**：中欧班列、跨境电商、双港枢纽（航空+陆港）。")
bullet(doc, "**14. 房投-8.2%、外资-56.9%**：地产/外资调整。")
bullet(doc, "**15. 财政+2.3%、税收占67.3%**：财政稳但质量待强。")
para(doc, "", size=10, space_after=2)

# ---- 五、对照表 ----
heading1(doc, "五、2025年GDP目标 vs 实际：对照表")
para(doc, "2025年报告中列出的主要预期目标，与2025年实际完成：")
tb = [
    ["指标", "2025年目标", "2025年实际（公报/报告）", "达标判定"],
    ["GDP增速", "6%左右", "+5.4%（15244.6亿）", "略欠"],
    ["规上工业增加值", "10%以上", "+9.0%", "略欠"],
    ["固定资产投资", "8%", "工业+17.9%、总量放缓", "分化"],
    ["社会消费品零售总额", "6%以上", "+5.0%（6629.4亿）", "略欠"],
    ["地方一般公共预算收入", "2%左右", "+2.3%", "达标"],
    ["进出口", "稳中提质", "+16.8%（6501.8亿）", "超预期"],
    ["城镇新增就业", "12.8万人", "14.17万（公报）", "达标"],
]
table(doc, tb[0], tb[1:], widths=[3.2, 3.4, 4.4, 4.0])
para(doc, "", size=10, space_after=2)
para(doc, "**要点**：进出口/就业/财政达标，GDP/工业/社零略欠——郑州\u201c外贸/民生强、总量/地产弱\u201d。")

# ---- 六、结构归因 ----
heading1(doc, "六、2025年增长是谁撑起来的？（结构归因）")
heading2(doc, "6.1 三次产业：二三产双稳")
para(doc, "**二产（+5.4%、占比36.6%）与三产（+5.4%、占比62.2%）双稳**：二产靠电子/汽车制造，三产靠金融/交通/服务。一产+3.7%（1.2%）。")
heading2(doc, "6.2 工业：电子/汽车/战新驱动")
para(doc, "规上工业+9.0%；电子+16.2%、汽车+11.9%、电气机械+17.9%、煤炭+18.6%；高技+14.9%、战新+11.8%、新一代信息+17.8%、新能源车+16.9%。")
heading2(doc, "6.3 外贸强")
para(doc, "进出口+16.8%（出口+20.3%），中部第一；富士康/跨境电商/中欧班列。")
heading2(doc, "6.4 消费企稳")
para(doc, "社零+5.0%；可穿戴/智能手机/体育超100%——绿色升级消费旺。")
heading2(doc, "6.5 投资分化")
para(doc, "工业投资+17.9%（制造高增）、但三产-1.8%、地产调整——制造业投资强、地产弱。")
para(doc, "**一句话归因**：郑州2025年\u201c**工业（电子/汽车/战新）+外贸/枢纽+数字经济**\u201d为主引擎，\u201c地产/总量\u201d偏弱——制造+枢纽驱动的中部中心。")

# ---- 七、财政 ----
heading1(doc, "七、预算与财政的\u201c含金量\u201d")
para(doc, "**地方财政一般公共预算收入1181.3亿元、+2.3%**；税收795.9亿、+2.6%，税占比67.4%；市本级184.5亿、-5.3%；一般公共预算支出1517.1亿、-0.6%。")
bullet(doc, "**收入稳、税占比中**：+2.3%、税收占67.4%；教育/社保/卫生健康支出增、科学支出-17.8%。")
bullet(doc, "**债务控制**：财政支出-0.6%，紧平衡；土地/契税走弱。")
bullet(doc, "**民生硬度**：社会保障+13.6%、卫生健康+8.0%；就业/教育稳。")
para(doc, "**财政含义**：郑州\u201c收入稳但税收中、支出缩、地产/契税弱\u201d，政策空间有限，需靠制造/枢纽创造税源。")

# ---- 八、民生 ----
heading1(doc, "八、民生底账：人口、收入与城乡")
para(doc, "郑州\u201c人口大市+正增长\u201d：**常住人口1313.8万人、自然增长率2.19‰（出生7.28‰、死亡5.09‰）；城镇化率81.52%（+0.52pct）**。居民人均可支配收入48021元、+4.4%。")
bullet(doc, "**收入**：人均可支配收入48021元、+4.4%；城镇52413元、+3.8%，农村33625元、+5.3%（城乡比缩小）。")
bullet(doc, "**就业**：城镇新增就业14.17万人；农村劳动力转移3.31万人。")
bullet(doc, "**社保**：养老参保550.44万、低保城市0.9万/农村3.27万；医保健全。")
bullet(doc, "**人口**：常住+稳定、自然增2.19‰；高教学生多（在校本专科144.89万）。")
para(doc, "**民生含义**：郑州\u201c人口大、收入稳、就业强\u201d，中部强人口吸附力。")

# ---- 九、城乡 ----
heading1(doc, "九、城镇与农村：格局与均衡")
para(doc, "**城镇化率81.52%（+0.52pct）**，郑州高度城镇化；县域（中牟/荥阳等）依托都市圈。")
bullet(doc, "**城市**：国家中心城市/省际综合枢纽（米字高铁、机场、中欧班列）。")
bullet(doc, "**农村**：粮食151.11万吨；农村收入+5.3%、城镇化转移就业3.31万。")
para(doc, "**城乡均衡**：郑州\u201c都市圈强、县域追赶\u201d，城郊融合+产业转移。")

# ---- 十、人口流 ----
heading1(doc, "十、人口流入与流出")
para(doc, "**郑州净流入**：常住1313.8万、自然增2.19‰。在全国人口总体收缩背景下，河南省内人口大市，靠省域虹吸+产业吸引。")
bullet(doc, "**流入**：国家中心城市/制造业/枢纽岗位；高校+劳动力。")
bullet(doc, "**竞争**：与其他省会（武汉/西安/长沙）争夺；郑州靠枢纽+制造+人口红利。")
para(doc, "人口方向决定中长期需求与增长；郑州的\u201c人口大市+枢纽+制造\u201d是其长逻辑。")

# ---- 十一、物价 ----
heading1(doc, "十一、物价与货币环境")
para(doc, "**2025年郑州CPI与上年持平**：食品-1.4%、交通通信-4.2%、居住+0.3%；衣着+2.4%、医疗+2.5%、其他+11.2%。")
bullet(doc, "**物价**：持平、温和，食品/交通负、服务略升。")
bullet(doc, "**货币/流动性**：本外币存款35505.7亿+7.2%、贷款+3.9%；金融稳健。")
para(doc, "**物价含义**：郑州通胀温和，内需压力可控，消费修复中。")

# ---- 十二、区域一体化 ----
heading1(doc, "十二、区域一体化：郑州在\u201c郑州都市圈+中部崛起+空中/陆上丝路\u201d里的位置")
para(doc, "郑州处于**中部崛起+郑州都市圈+国家中心城市**核心：既是河南省会，也是\u201c一带一路\u201d枢纽（中欧班列集结中心/航空港）。")
bullet(doc, "**郑州都市圈**：与开封/许昌/新乡协同，市域铁路/基础设施；中原城市群龙头。")
bullet(doc, "**中部崛起**：作为中部重要中心城市，承接产业梯度转移。")
bullet(doc, "**空中/陆上丝路**：中欧班列3417班、郑州-卢森堡空中丝路、国际航空货运枢纽（货邮100万吨）。")
bullet(doc, "**黄河战略**：黄河流域生态保护，京广/陇海廊道十字。")
para(doc, "**区域含义**：郑州作为\u201c中原枢纽+国家中心\u201d，靠\u201c枢纽+制造+都市圈\u201d带动河南与中部。")

# ---- 十三、五条主线 ----
heading1(doc, "十三、未来5—10年最值得观察的五条主线")
bullet(doc, "**主线1｜电子信息/半导体**：电子信息9000亿、富士康/超聚变/半导体。能否破万亿/卡位。")
bullet(doc, "**主线2｜新能源汽车/汽车**：汽车120万辆、比亚迪/宇通；能否建成\u201c新能源汽车之城\u201d。")
bullet(doc, "**主线3｜交通枢纽/开放**：中欧班列/机场货邮/跨境电商；能否持续扮演丝路枢纽。")
bullet(doc, "**主线4｜未来产业/数字经济**：未来产业2000亿、量子/氢能/机器人/7000+算力。能否孵化新增长极。")
bullet(doc, "**主线5｜都市圈/人口**：常住1313.8万、人口虹吸。能否把人口红利变成长效内需引擎。")
para(doc, "这五条，是郑州从\u201c制造+枢纽城市\u201d走向\u201c国家中心+新质制造+双枢纽\u201d的\u201c主赛道\u201d。")

# ---- 十四、结论 ----
heading1(doc, "十四、最终结论：郑州在\u201c制造+枢纽+国家中心城市\u201d里的增长逻辑")
para(doc, "郑州2025年，本质上是\u201c**工业/外贸/枢纽驱动、地产/财政/总量偏弱**\u201d的答卷：GDP15244.6亿、+5.4%、规上工业+9.0%（国家中心城市首位）、进出口6501.8亿+16.8%、数字经济8000亿、社零+5.0%、财政+2.3%。")
para(doc, "只要电子信息/汽车/枢纽持续，郑州就站在\u201c制造+枢纽\u201d增长极；若地产/总量承压，郑州需承受\u201c工业外贸强、地产/收入弱\u201d的结构挑战。")
para(doc, "最稳观察信号：**一盯电子/半导体（引擎）、二盯新能源汽车（制造）、三盯枢纽/外贸（开放）、四盯未来产业（新质）、五盯人口/都市圈（长期）。**郑州，是\u201c国家中心城市+制造枢纽\u201d的中原样本。")

# ---- 附录A ----
heading1(doc, "附录A：主要资料来源与核验路径")
bullet(doc, "郑州市2025年政府工作报告（2025年2月）——目标来源。")
bullet(doc, "《2025年郑州市国民经济和社会发展统计公报》（市统计局）——GDP、工业、外贸、人口实值。")
bullet(doc, "2026年郑州市政府工作报告（2026年2月）——2025追认/电子/汽车/枢纽/未来产业。")
bullet(doc, "郑州海关、市财政局、郑州航空港（外贸/贸易）。")
heading2(doc, "核验说明")
para(doc, "本报告涉及数据以统计公报/官方口径为准；\u201c国家中心/枢纽/制造/未来产业\u201d等以官方口径为准。")

# ---- 附录B ----
heading1(doc, "附录B：建议建立的年度跟踪仪表盘")
para(doc, "建议把下面10个\u201c测脉搏\u201d指标做成年度跟踪表：")
dash = [
    ["#", "指标", "2025基值", "为什么重要"],
    ["1", "GDP增速", "+5.4%（15244.6亿）", "总量与方向"],
    ["2", "规模以上工业增速", "+9.0%", "制造底盘"],
    ["3", "电子信息产业规模", "9000亿", "产业龙头"],
    ["4", "汽车/新能源车产量", "120.5万辆", "汽车制造"],
    ["5", "进出口/出口增速", "+16.8% / +20.3%", "开放/外贸"],
    ["6", "机场货邮/中欧班列", "103.3万吨/3417班", "枢纽"],
    ["7", "社零增速", "+5.0%（6629.4亿）", "内需消费"],
    ["8", "常住人口/城镇化", "1313.8万 / 81.5%", "人口与城市"],
    ["9", "地方财政收入/税占", "+2.3%(1181.3亿) / 67.4%", "财政质量"],
    ["10", "GDP增速 vs 目标", "5.4% vs 6%", "达标性"],
]
table(doc, dash[0], dash[1:], widths=[0.8, 4.8, 4.2, 3.6])
para(doc, "把这10个指标连起来看，电子/汽车/外贸/枢纽（3/4/5）、人口/都市圈（8），都说明郑州在真正换挡。")

# ============ 保存 ============
out = "/Users/x/Desktop/content-prod-lab/reports/郑州市_2025年政府工作报告_深度研究_2026-08-13.docx"
doc.save(out)
print("SAVED:", out)
print("PARAGRAPHS:", len(doc.paragraphs))
print("TABLES:", len(doc.tables))
