# -*- coding: utf-8 -*-
"""Build 洛阳市2025年政府工作报告 深度研究 DOCX, 参照地级市系列版式。"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DARK = RGBColor(0x1F, 0x2A, 0x44)
RED = RGBColor(0xB0, 0x1E, 0x2E)
GRAY = RGBColor(0x59, 0x60, 0x69)
BLUE = RGBColor(0x1F, 0x4E, 0x79)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
HEADER_FILL = "1F2A44"
K = "微软雅黑"


def set_run(run, size=11, bold=False, color=DARK, font=K):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    run.font.name = font
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), font)
    rFonts.set(qn("w:hAnsi"), font)
    rFonts.set(qn("w:eastAsia"), font)


def para(doc, text, size=11, bold=False, color=DARK, align=None,
         space_after=6, space_before=0, font=K):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        set_run(r, size=size, bold=bold or (i % 2 == 1), color=color, font=font)
    return p


def heading1(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(16)
    pf.space_after = Pt(8)
    r = p.add_run(text)
    set_run(r, size=16, bold=True, color=RED)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "B01E2E")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def heading2(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(10)
    pf.space_after = Pt(4)
    r = p.add_run(text)
    set_run(r, size=13, bold=True, color=BLUE)
    return p


def table(doc, headers, rows, widths=None, font_size=9.5):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_run(r, size=font_size, bold=True, color=WHITE)
        tcPr = hdr[i]._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), HEADER_FILL)
        tcPr.append(shd)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(val))
            set_run(r, size=font_size, color=DARK)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    return t


def quote_box(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(8)
    pf.left_indent = Pt(12)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), "B01E2E")
    pBdr.append(left)
    pPr.append(pBdr)
    r = p.add_run(text)
    set_run(r, size=11, color=DARK, font="楷体")
    return p


def bullet(doc, text, size=10.5):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.7)
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        set_run(r, size=size, bold=(i % 2 == 1), color=DARK)
    return p


# ================================================================= 文档主体
doc = Document()
sec = doc.sections[0]
sec.page_width = Cm(21)
sec.page_height = Cm(29.7)
sec.top_margin = Cm(2.4)
sec.bottom_margin = Cm(2.2)
sec.left_margin = Cm(2.5)
sec.right_margin = Cm(2.5)

st = doc.styles["Normal"]
st.font.name = K
st.font.size = Pt(11)
st.element.rPr.rFonts.set(qn("w:eastAsia"), K)


# ---- 封面 ----
para(doc, "洛阳市2025年政府工作报告", size=24, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=60, space_after=4)
para(doc, "深度研究与\u201c容易被忽视的细节\u201d分析", size=16, bold=True, color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
para(doc, "从\u201c十三朝古都、老工业重镇、装备制造、铝/钛钨、唐三彩与文旅\u201d重新理解洛阳", size=12, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)
para(doc, "\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501", color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
para(doc, "研究对象：2025年洛阳市政府工作报告及同期财政、国民经济和社会发展统计资料、2026年官方复盘", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
para(doc, "标定日期：2026-08-14      覆盖：2025年预期目标→2025年实际→2026目标", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
doc.add_page_break()

# ---- 目录 ----
catalog = [
    "执行摘要",
    "一、研究口径与阅读方法",
    "二、先看洛阳的特殊底盘：十三朝古都、老工业重工、装备制造、钨钼钛铝、文旅",
    "三、最关键的宏观错位：GDP 6164.52亿/6.0%增长、工业强（装备/高技术）但地产负、房地产-13.6%",
    "四、容易被忽视、但信号很重的细节（15条）",
    "五、2025年GDP目标 vs 实际：对照表",
    "六、2025年增长是谁撑起来的？（结构归因）",
    "七、预算与财政的\u201c含金量\u201d",
    "八、民生底账：人口、收入与城乡",
    "九、城镇与农村：格局与均衡",
    "十、人口流入与流出",
    "十一、物价与货币环境",
    "十二、区域一体化：洛阳在\u201c中原城市群+郑洛西+黄河经济带\u201d里的位置",
    "十三、未来5—10年最值得观察的五条主线",
    "十四、最终结论：洛阳在\u201c老工业转型+装备制造+文旅\u201d里的增长逻辑",
    "附录A：主要资料来源与核验路径",
    "附录B：建议建立的年度跟踪仪表盘",
]
para(doc, "目　录", size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=10, space_after=10)
for item in catalog:
    para(doc, item, size=12, space_after=4)
doc.add_page_break()

# ---- 执行摘要 ----
heading1(doc, "执行摘要")
para(doc, "**核心判断**　2025年洛阳最显著的是\u201cGDP 6164.52亿元、增长6.0%（高于全国全省）、河南第2\u201d、\u201c规上工业+8.5%（装备+20.7%、高技术+61.9%、战略新兴+32.8%）\u201d、\u201c进出口304.3亿/+20.9%、旅游1.6亿人次/1285.4亿\u201d、\u201c但房地产-13.6%、商品房销售-3.9%、人口自然增-1.55‰\u201d、\u201c财收421.8亿/+3.6%、常住708.3万\u201d。这说明洛阳在\u201c老工业+装备+文旅\u201d中，**工业科技、文旅、外贸强但地产、人口弱**。")
para(doc, "把2025年目标（GDP+5.5%左右）、2025年统计（GDP+5.3%达成超1pct）、趋势一起看，洛阳是\u201c老工业转型+装备制造+唐三彩文旅\u201d路径：**装备制造、石油化工（乙烯）、铝钨钼钛新材料、文旅（龙门/老君山/唐三彩）、新能源（中州时代动力电池）**是支柱；2025年总量6164.5亿居河南第2（次于郑州）。")
para(doc, "最容易记住的一句话：**洛阳是\u201c十三朝古都、老工业重镇、装备制造之都、唐三彩之都\u201d，靠\u201c装备+新材料+文旅\u201d增长。**观察洛阳，与其只看\u201cGDP 6164亿\u201d，不如看\u201c装备+20.7%、高技术制造+61.9%、进出口+20.9%、旅游1285.4亿、动力电池\u201d。")

# ---- 一、研究口径与阅读方法 ----
heading1(doc, "一、研究口径与阅读方法")
para(doc, "本报告以\u201c洛阳市政府工作报告（2025年，张玉杰作）\u201d为起点，把\u201c2025年GDP目标（5.5%左右）\u201d与\u201c官方2025年GDP（6164.52亿元/+6.0%）\u201d并置对照，并用\u201c2025年洛阳市统计公报\u201d和\u201c2026年政府工作报告\u201d横向核验。")
para(doc, "口径提示：\u201cGDP、规上工业增加值增速\u201d按可比价（实际增速）；\u201c投资、消费、进出口、财政收入\u201d按现价（名义增速）。涉及\u201c常住人口\u201d用统计公报常住口径（708.3万），城镇化率用公布值（69.08%）。")
para(doc, "指标体系（与研究口径一致）：五大象限——**总量与增速（GDP）、产业动能（装备/铝钛/石化/文旅）、投资、财政、民生与人口**。")
para(doc, "特别提示（不吃老本）：洛阳2024年GDP 5825.86亿/+5.3%，2025年+6.0%强回升；它不只是\u201c老工业\u201d——世界遗产**龙门石窟、牡丹、唐三彩（洛阳三彩）**，以及**装备制造、钨钼钛铝、新能源（中州时代动力电池/百万吨乙烯）**才是真正的增长逻辑。")
# ---- 二、先看洛阳的特殊底盘 ----
heading1(doc, "二、先看洛阳的特殊底盘：十三朝古都、老工业重工、装备制造、钨钼钛铝、文旅")
para(doc, "洛阳地处河南西部、黄河中下游，是**十三朝古都、共和国老工业基地、装备制造重镇、世界遗产（龙门石窟）文旅名城**。2025年GDP 6164.52亿元、常住708.3万，河南第2，人均约8.7万元。")
para(doc, "五个底盘名词，先立框架：")
bullet(doc, "**装备制造**　洛阳重工业（一拖、中信重工、洛轴），装备制造增加值+20.7%（占规上29.4%）——\u201c共和国长子\u201d。")
bullet(doc, "**钨钼钛铝/新材料**　洛阳钼业（世界级）、钨/钼/钛/铝金属深加工，新材料/有色。")
bullet(doc, "**石油化工（乙烯）**　百万吨乙烯项目，石化/化工板块。")
bullet(doc, "**新能源（动力电池）**　中州时代动力电池一、二期产值超200亿、新能源汽车/储能。")
bullet(doc, "**文旅（唐三彩/龙门/牡丹）**　龙门石窟、老君山、唐三彩（洛阳三彩）、牡丹，旅游1.6亿人次/1285.4亿。")
para(doc, "这五根（装备+钨钼钛+石化+新能源+文旅）构成洛阳独特底盘：**左手装备制造（工业脊梁），右手文旅（三彩/龙门），底子是唐宋文明**。理解洛阳，先理解\u201c老工业+古都\u201d。")

# ---- 三、最关键的宏观错位 ----
heading1(doc, "三、最关键的宏观错位：GDP 6164.52亿/6.0%增长、工业强（装备/高技术）但地产负、外贸、人口")
para(doc, "2025年洛阳最需要辨析的一组\u201c错位\u201d：**GDP 6.0%（强于目标）、规上工业+8.5%（装备+20.7%、高技术+61.9%强）但房地产-13.6%、固投+5.0%、人口自然增-1.55‰**。")
para(doc, "为什么\u201c工业这么强、经济却仍只是6%\u201d？三个解释：")
para(doc, "**其一，工业强、装备/高科技旺**　装备+20.7%（占规上29.4%）、高技术+61.9%、战略新兴+32.8%、新能源（中州时代）——\u201c制造业升级是总引擎\u201d。")
para(doc, "**其二，地产、传统行业仍拖累**　房地产-13.6%、商品房销售-13.9%；高耗能（占42.8%）仅+4.3%——\u201c地产/传统\u201d低增。")
para(doc, "**其三，人口负增、内需刚需**　自然增长-1.55‰、社零+5.7%（民生消费尚可）——\u201c人口少子化、内需温和\u201d。")
para(doc, "小结：洛阳2025年是\u201c**强工业科技文旅、弱地产/传统、势仍向上**\u201d的一年：装备、新材料、动力电池、文旅撑增速，地产弱。")

# ---- 四、容易被忽视、但信号很重的细节（15条） ----
heading1(doc, "四、容易被忽视、但信号很重的细节（15条）")
bullet(doc, "**1.装备制造+20.7%（占规上29.4%）**\u201c洛阳重工（一拖/洛轴/中信重工）是装备龙头。\u201d")
bullet(doc, "**2.高技术制造业+61.9%**\u201c中州时代动力电池、功率半导体、氢能——\u2018新质生产力\u2019。\u201d")
bullet(doc, "**3.工业战略性新兴产业+32.8%**\u201c新能源/新材料（钼钨钛）\u201d接棒。")
bullet(doc, "**4.进出口304.3亿/+20.9%、出口+19.8%**\u201c洛阳造（装备/机械）出口强（一带一路）。\u201d")
bullet(doc, "**5.中州时代动力电池一期二期投产、产值超200亿**\u201c新能源电池（动力/储能）基地。\u201d")
bullet(doc, "**6.文旅：旅游1.6亿人次/1285.4亿/+6.4%、入境恢复**\u201c龙门/老君山/唐三彩（三彩）。\u201d")
bullet(doc, "**7.财政421.8亿/+3.6%、税收+5.5%（占67.2%）**\u201c税收强、民生占74.7%。\u201d")
bullet(doc, "**8.社零2475.1亿/+5.7%**\u201c以旧换新+160亿、新能车/家电高。\u201d")
bullet(doc, "**9.固投+5.0%、工业投资+18.6%、约0.5万亿项目**\u201c设备/产业投资旺。\u201d")
bullet(doc, "**10.居民收入37240元/+4.7%、城乡比2.23**\u201c农村+6.1%快于城镇+4.1%。\u201d")
bullet(doc, "**11.常住708.3万/城镇化69.08%（+0.69pct）**\u201c河南人口大市、自然增长-1.55‰。\u201d")
bullet(doc, "**12.CPI持平（12月+1.6%、消费品+2.3%）**\u201c物价温和、年底走高。\u201d")
bullet(doc, "**13.粮食247.3万吨、高标准农田22万亩**\u201c黄河文化粮仓。\u201d")
bullet(doc, "**14.高新技术产业+11.7%占规上59.3%**\u201c高技术产业底座厚。\u201d")
bullet(doc, "**15.唐三彩/世界遗产（龙门石窟51处国保）**\u201c文旅、古都文化\u201d名片。")
# ---- 五、2025年GDP目标 vs 实际：对照表 ----
heading1(doc, "五、2025年GDP目标 vs 实际：对照表")
table(doc,
    ["指标", "2025年目标", "2025年实际", "达成评价"],
    [
        ["地区生产总值(GDP)", "增长5.5%左右", "6164.52亿/6.0%", "超额达成"],
        ["规上工业增加值", "——", "+8.5%", "工业强"],
        ["固定资产投资", "——", "+5.0%", "稳健"],
        ["社会消费品零售总额", "——", "2475.1亿/+5.7%", "稳健"],
        ["进出口总额", "——", "304.3亿/+20.9%", "大幅超额"],
        ["一般公共预算收入", "——", "421.8亿/+3.6%", "稳增"],
        ["居民收入", "与经济增长同步", "37240元/+4.7%", "略低GDP"],
    ],
)
para(doc, "注：GDP、规上工业按可比价。**GDP（+6.0%）、进出口（+20.9%）、规上（+8.5%）超额**；**社零、财收、固投稳增**；房地产（-13.6%）是短板。")
para(doc, "拆读：**装备（+20.7%）、高技术（+61.9%）、动力电池、文旅、进出口（+20.9%）是亮色**；**地产（-13.6%）、人口（-1.55‰）偏弱**——\u201c制造强、文旅强、地产人口弱\u201d，是文旅工业之城样本。")

# ---- 六、2025年增长是谁撑起来的？（结构归因） ----
heading1(doc, "六、2025年增长是谁撑起来的？（结构归因）")
para(doc, "把洛阳GDP的6.0%拆开：三次产业分别增3.9%、6.2%、6.0%（结构3.7：40.1：56.2）。**第二产业（工业）与第三产业（服务业+文旅）双撑，第一产业（农业）稳但体量小**。")
para(doc, "2026年洛阳强调\u201c1+2+4+N\u201d、装备制造、新能源、文旅，聚焦**装备、百万吨乙烯、动力电池、航天/氢能、龙门文旅**——核心是\u201c老工业转型+新质生产力\u201d。")
para(doc, "**第二产业（工业）**：规上+8.5%（装备+20.7%、高技术+61.9%、战新+32.8%）、动力电池/新材料——\u201c制造强\u201d。")
para(doc, "**第三产业（服务业）**：批发零售+6.3%、住宿餐饮+5.9%、文旅（旅游+6.4%）——\u201c服务业+文旅稳\u201d。")
para(doc, "**外贸（开放）**：进出口+20.9%（出口+19.8%）——\u201c洛阳造出海（装备/锂）\u201d。")
para(doc, "一句话归因：**2025年洛阳增长\u201c靠工业（装备/新材料/新能源）+文旅+出口\u201d**，地产、传统行业弱；\u201c新质+文旅\u201d是核心。")

# ---- 七、预算与财政的\u201c含金量\u201d ----
heading1(doc, "七、预算与财政的\u201c含金量\u201d")
para(doc, "2025年洛阳**一般公共预算收入421.8亿元（+3.6%）**，其中税收283.3亿元（+5.5%）、税收占比67.2%；一般公共预算支出725.2亿元（+3.3%）、民生支出占74.7%。")
bullet(doc, "税收结构：税收+5.5%（装备/科技税）强、占67.2%——\u201c税收质量高\u201d。")
bullet(doc, "民生支出：占74.7%、社保就业/教育/医疗优先。")
bullet(doc, "金融支撑：存款9128.02亿（+8.1%）、贷款7230.08亿（+1.1%，偏弱）、上市公司15家。")
para(doc, "**财政含金量小结**：财收+3.6%、税收+5.5%，\u201c增收靠税收（好）\u201d；民生占74.7%；财政对\u201c装备、新材料、动力电池、文旅\u201d投入加大。")

# ---- 八、民生底账：人口、收入与城乡 ----
heading1(doc, "八、民生底账：人口、收入与城乡")
para(doc, "2025年洛阳**居民人均可支配收入37240元（+4.7%）**，其中城镇49887元（+4.1%）、农村22340元（+6.1%），城乡比约2.23。就业：城镇新增就业10.6万人。")
para(doc, "人口画像：**常住708.3万、城镇化率69.08%（+0.69pct）**，河南第2大城；但自然增长-1.55‰（出生6.52‰<死亡8.08‰）、老龄化（少子化）。")
para(doc, "民生投入：公办园/义务校新建、养老、医疗床位——民生优先。")

# ---- 九、城镇与农村：格局与均衡 ----
heading1(doc, "九、城镇与农村：格局与均衡")
para(doc, "洛阳常住城镇化率69.08%（河南靠前），城乡较均衡；农村收入增速（+6.1%）高于城镇（+4.1%），**城乡比收敛（约2.23）**。")
para(doc, "农业底盘：**粮食247.3万吨（河南粮仓）**、蔬菜285.6万吨、水果129.3万吨、中药材14.7万吨——\u201c河洛农产\u201d。")
para(doc, "一句话：\u201c洛阳城镇工业强、农村农业稳、城乡比尚可\u201d，洛阳从\u201c老工业\u201d到\u201c装备+文旅小镇\u201d带县/乡。")

# ---- 十、人口流入与流出 ----
heading1(doc, "十、人口流入与流出")
para(doc, "洛阳常住708.3万、城镇化69.08%，是河南人口第2大城；但\u201c自然增长-1.55‰\u201d、流苏州/郑州打工，仅中心城区（涧西/洛龙）集聚。")
para(doc, "结构观察：**城镇化率仍低于全国（约69%）**、出生率低、老龄化；\u201c高校（在校23.1万）\u201d为留人多一点。")
para(doc, "2026年目标：留人/引才（装备/新材料/文旅）、人才南下——洛阳把\u201c古都+产业\u201d聚人。")

# ---- 十一、物价与货币环境 ----
heading1(doc, "十一、物价与货币环境")
para(doc, "2025年洛阳**CPI与上年持平（0%）**（12月+1.6%、消费品+2.3%、服务+0.5%）——\u201c物价低稳、年底走高\u201d。")
bullet(doc, "信贷：存款+8.1%、贷款+1.1%（投放放缓）；社零+5.7%。")
bullet(doc, "消费：以旧换新+160亿、汽车/家电等升级——\u201c扩内需\u201d。")
para(doc, "货币环境判断：**CPI平稳、宽存、贷款低增**；洛阳靠\u201c以旧换新、文旅、装备出海\u201d稳需求（2026 CPI目标2%）。")

# ---- 十二、区域一体化：洛阳的位置 ----
heading1(doc, "十二、区域一体化：洛阳在\u201c中原城市群+郑洛西+黄河经济带\u201d里的位置")
para(doc, "洛阳地处\u201c中原城市群副中心、郑洛西高质量发展带、黄河流域生态保护\u201d交汇，河南第2城。")
bullet(doc, "**中原城市群副中心**　郑洛一体（郑州-洛阳），装备/制造与郑州汽车金融互补。")
bullet(doc, "**黄河/文旅带**　黄河文化、洛阳古都、全域旅游（洛阳在黄河流域文旅核心）。")
bullet(doc, "**开放**　中欧班列、洛阳自贸/保税、面向中亚。")
para(doc, "一句话：**洛阳在\u201c中原城市群+黄河经济带\u201d中，最核心的定位是\u201c河南副中心+装备文旅\u201d**；区位、装备、文旅是大优势。")

# ---- 十三、未来5—10年最值得观察的五条主线 ----
heading1(doc, "十三、未来5—10年最值得观察的五条主线")
bullet(doc, "**主线一：装备制造高端化**\u201c装备+20.7%、洛轴/中信重工\u201d能否由\u201c重基础\u201d到\u201c智能装备\u201d。")
bullet(doc, "**主线二：新材料（钨钼钛）/新能源电池**\u201c中州时代、钼钨钛\u201d能否成\u201c新质增长点\u201d。")
bullet(doc, "**主线三：油气/乙烯化工**\u201c百万吨乙烯\u201d能否从\u201c炼油\u201d到\u201c化工新城\u201d。")
bullet(doc, "**主线四：文旅融合（唐三彩/龙门）**\u201c旅游1.6亿人次\u201d能否由\u201c走量\u201d到\u201c高客单、入境\u201d。")
bullet(doc, "**主线五：人口与旧城**\u201c708万、-1.55‰\u201d在\u201c强省会\u201d下能否靠\u201c装备+文旅\u201d留人。")

# ---- 十四、最终结论 ----
heading1(doc, "十四、最终结论：洛阳在\u201c老工业转型+装备制造+文旅\u201d里的增长逻辑")
para(doc, "把2025年闭环打穿：**洛阳是\u201c十三朝古都、老工业、文旅大市\u201d**：GDP 6164.52亿/+6.0%、规上+8.5%（装备+20.7%、高技术+61.9%）、文旅1.6亿人次/1285.4亿、进出口+20.9%。")
para(doc, "洛阳不是\u201c只有拖拉机和唐三彩\u201d——它是**装备+新材料+石化+新能源（中州时代）+文旅**的复合体，靠\u201c工业科技+古都文旅\u201d双轮；但地产、人口偏弱。")
para(doc, "一句话结论：**洛阳是\u201c十三朝古都、装备制造名城、世界遗产\u201d；观察它先看\u201c装备、中州时代电池、钨钼铝、龙门/唐三彩文旅、出口\u201d，再看\u201c地产、人口、固投\u201d。**它是\u201c工业强、文旅强、地产弱\u201d的中原转型样本。")

# ---- 附录A ----
heading1(doc, "附录A：主要资料来源与核验路径")
bullet(doc, "《2025年洛阳市政府工作报告》（2025年2月，张玉杰作，2025年目标、2024年回顾（2024年GDP 5825.86亿、增速6.5%））")
bullet(doc, "《2025年洛阳市国民经济和社会发展统计公报》（洛阳市统计局，2026-04-30，2025年实际数据）")
bullet(doc, "《2026年洛阳市政府工作报告》（2026年2月，2025年复盘+2026年目标）")
bullet(doc, "洛阳市人民政府/统计局（ly.gov.cn/lytjj）")

# ---- 附录B ----
heading1(doc, "附录B：建议建立的年度跟踪仪表盘")
bullet(doc, "GDP总量/增速 vs 全年目标（可比价）。")
bullet(doc, "规上工业增加值及分行业（装备/有色/化工/电力/制药）增速。")
bullet(doc, "钨钼铝产/价、乙烯/新材料。")
bullet(doc, "固定资产投资/工业/房地产投资增速。")
bullet(doc, "社零、以旧换新、旅游人数/收入/入境。")
bullet(doc, "进出口、出口、中欧班列。")
bullet(doc, "一般公共预算收入、税收/非税、民生%。")
bullet(doc, "常住人口、自然增长、城镇化率、高校。")
bullet(doc, "CPI、金融存贷款（贷款投放）。")
bullet(doc, "装备订单、动力电池、唐三彩文旅IP。")

# ---------------- 保存 ----------------
out = "/Users/x/Desktop/content-prod-lab/reports/洛阳市_2025年政府工作报告_深度研究_2026-08-14.docx"
doc.save(out)
print("SAVED OK: 洛阳市", out)
