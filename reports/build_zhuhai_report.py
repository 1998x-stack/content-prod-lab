# -*- coding: utf-8 -*-
"""Build 济宁市2025年政府工作报告 深度研究 DOCX, 参照地级市系列版式。"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DARK = RGBColor(0x1F, 0x2A, 0x44)
RED = RGBColor(0xB0, 0x1E, 0x2E)
GRAY = RGBColor(0x59, 0x60, 0x69)
BLUE = RGBColor(0x1F, 0x4E, 0x79)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
HEADER_FILL = "1F2A44"
K = "微软雅黑"


def set_run(run, size=11, bold=False, color=DARK, font=K):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    run.font.name = font
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), font)
    rFonts.set(qn("w:hAnsi"), font)
    rFonts.set(qn("w:eastAsia"), font)


def para(doc, text, size=11, bold=False, color=DARK, align=None,
         space_after=6, space_before=0, font=K):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        set_run(r, size=size, bold=bold or (i % 2 == 1), color=color, font=font)
    return p


def heading1(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(16)
    pf.space_after = Pt(8)
    r = p.add_run(text)
    set_run(r, size=16, bold=True, color=RED)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "B01E2E")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def heading2(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(10)
    pf.space_after = Pt(4)
    r = p.add_run(text)
    set_run(r, size=13, bold=True, color=BLUE)
    return p


def table(doc, headers, rows, widths=None, font_size=9.5):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_run(r, size=font_size, bold=True, color=WHITE)
        tcPr = hdr[i]._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), HEADER_FILL)
        tcPr.append(shd)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(val))
            set_run(r, size=font_size, color=DARK)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    return t


def quote_box(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(8)
    pf.left_indent = Pt(12)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), "B01E2E")
    pBdr.append(left)
    pPr.append(pBdr)
    r = p.add_run(text)
    set_run(r, size=11, color=DARK, font="楷体")
    return p


def bullet(doc, text, size=10.5):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.7)
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        set_run(r, size=size, bold=(i % 2 == 1), color=DARK)
    return p


# ================================================================= 文档主体
doc = Document()
sec = doc.sections[0]
sec.page_width = Cm(21)
sec.page_height = Cm(29.7)
sec.top_margin = Cm(2.4)
sec.bottom_margin = Cm(2.2)
sec.left_margin = Cm(2.5)
sec.right_margin = Cm(2.5)

st = doc.styles["Normal"]
st.font.name = K
st.font.size = Pt(11)
st.element.rPr.rFonts.set(qn("w:eastAsia"), K)


# ---- 封面 ----
para(doc, "珠海市2025年政府工作报告", size=24, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=60, space_after=4)
para(doc, "深度研究与\u201c容易被忽视的细节\u201d分析", size=16, bold=True, color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
para(doc, "从\u201c经济特区、横琴粤澳合作区、4+3产业、港珠澳大桥、中国航展\u201d重新理解珠海", size=12, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)
para(doc, "\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501", color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
para(doc, "研究对象：2025年珠海市政府工作报告及同期财政、国民经济和社会发展统计资料、2026年官方复盘", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
para(doc, "标定日期：2026-08-14", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
doc.add_page_break()

# ---- 目录 ----
catalog = [
    "执行摘要",
    "一、研究口径与阅读方法",
    "二、先看珠海的特殊底盘：特区、横琴粤澳合作区、4+3产业、港珠澳大桥",
    "三、最关键的宏观错位：GDP+2.7%、二产-2.9%、固投-31.6%大幅不及目标，三产、财收却强",
    "四、容易被忽视、但信号很重的细节（15条）",
    "五、2025年GDP目标 vs 实际：对照表",
    "六、2025年增长是谁撑起来的？（结构归因）",
    "七、预算与财政的\u201c含金量\u201d",
    "八、民生底账：人口、收入与城乡",
    "九、城镇与农村：格局与均衡",
    "十、人口流入与流出",
    "十一、物价与货币环境",
    "十二、区域一体化：珠海在粤港澳大湾区、横琴+澳门、珠江口西岸\u201c三极\u201d里的位置",
    "十三、未来5—10年最值得观察的五条主线",
    "十四、最终结论：珠海在\u201c横琴+4+3产业+大湾区枢纽\u201d里的增长逻辑",
    "附录A：主要资料来源与核验路径",
    "附录B：建议建立的年度跟踪仪表盘",
]
para(doc, "目　录", size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=10, space_after=10)
for item in catalog:
    para(doc, item, size=12, space_after=4)
doc.add_page_break()

# ---- 执行摘要 ----
heading1(doc, "执行摘要")
para(doc, "**核心判断**　2025年珠海最显著的是\u201cGDP 4573.10亿/+2.7%（大幅远低于6%目标）、第二产业-2.9%、固定资产投资-31.6%、房地产开发-39.5%\u201d、\u201c但第三产业+6.9%（占57.8%）、规上工业+4.1%、外贸3354亿/+3.4%、一般公共预算收入+4.0%（超额1pct）、实际使用外资+6.4%\u201d、\u201cCPI-0.7%\u201d。这说明珠海在\u201c经济特区+横琴\u201d中，**服务业强但制造业/投资/地产承压**。")
para(doc, "把2025年目标（GDP+6%、规上+8%、固投+3%、社零+5%、进出口+5%、财收+3%）、2025年实际（GDP+2.7%、规上+4.1%、固投-31.6%、社零+1.7%、进出口+3.4%、财收+4.0%）趋势看，珠海是\u201c4+3产业+横琴+开放\u201d路径：**集成电路（+26.5%）、高端装备（+21.6%）、新能源（+12.0%）、生物医药（-5.9%）**是支柱。")
para(doc, "最容易记住的一句话：**珠海是\u201c经济特区、横琴粤澳合作区、珠江口西岸核心城市\u201d，靠\u201c4+3新质产业+横琴+开放（港珠澳大桥）\u201d驱动。**观察珠海，与其只看\u201cGDP 4573亿\u201d，不如看\u201c集成电路+26.5%、高端装备+21.6%、横琴牌、口岸出入境2亿+、航展（近20年最热闹）\u201d。")
# ---- 一、研究口径与阅读方法 ----
heading1(doc, "一、研究口径与阅读方法")
para(doc, "本报告以\u201c珠海市政府工作报告（2025年2月，市十四届人大五次会议）\u201d为起点，把\u201c2025年GDP目标（6%）\u201d与\u201c官方2025年（4573.10亿/+2.7%）\u201d并置对照，用\u201c2025年珠海市统计公报\u201d和\u201c2026年政府工作报告\u201d横向核验。")
para(doc, "口径提示：\u201cGDP、规上工业增速\u201d按可比价（实际）；\u201c投资、消费、进出口、财政\u201d按现价（名义）。\u201c常住人口\u201d用统计公报（254.31万）、城镇化率90.96%。")
para(doc, "指标体系（与研究口径一致）：**总量与增速、产业动能（4+3产业）、外贸、财政、民生与人口**。")
para(doc, "特别提示（不吃老本）：珠海2024年GDP~4479.1亿（+3.5%）、2025年4573.10亿/+2.7%（远低于6%目标）；它不是\u201c只有海岛/航展\u201d——**集成电路、高端装备、新能源、生物医药、横琴粤澳合作区、港珠澳大桥\u201d才是真正底色；2025年二产下滑（-2.9%）是最大短板。")
# ---- 二、先看珠海的特殊底盘 ----
heading1(doc, "二、先看珠海的特殊底盘：特区、横琴粤澳合作区、4+3产业、港珠澳大桥")
para(doc, "珠海地处珠江口西岸、毗邻澳门，是**中国最早经济特区之一、横琴粤澳深度合作区所在地、珠江口西岸核心城市**；是以\u201c4+3产业\u201d、高校、宜居而著称。2025年GDP 4573.10亿、常住254.31万、城镇化率90.96%、广东第7（深广佛莞惠中山之后）。")
para(doc, "五个底盘名词，先立框架：")
bullet(doc, "**横琴粤澳深度合作区**　服务澳门经济适度多元化、跨境创新（集成电路/中医药/文旅会展）——\u201c横琴牌\u201d。")
bullet(doc, "**港珠澳大桥/开放枢纽**　港珠澳大桥经贸新通道（进出口2329亿+11.3%）、口岸出入境2亿+、拱北1.1亿——\u201c黄金内湾\u201d。")
bullet(doc, "**4+3产业**　集成电路、生物医药、新能源、新材料（四大主导）+高端装备、智能家电、精细化工（三优势）——\u201c新质生产\u201d。")
bullet(doc, "**中国航展**　珠海航展（第十五届1022家企业、近20年最热闹）——\u201c航空航天\u201d。")
bullet(doc, "**高校/创新**　研发投入占比4.06%（全省第2）、每万人口发明专利222.92件（全省第2）、高企865家——\u201c创新强市\u201d。")
para(doc, "这五根（横琴+大桥+4+3产业+航展+高校）构成珠海独特底盘：**左手横琴（服务澳门），右手4+4新质产业+开放**。理解珠海，先理解\u201c特区、横琴、珠江口西岸\u201d。")
# ---- 三、最关键的宏观错位 ----
heading1(doc, "三、最关键的宏观错位：GDP+2.7%、二产-2.9%、固投-31.6%大幅不及目标，三产、财收却强")
para(doc, "2025年珠海最需要辨析的一组\u201c错位\u201d：**GDP+2.7%（远低于6%）、第二产业-2.9%、固定资产投资-31.6%、房地产开发-39.5%、规上工业+4.1%（低8%目标），但第三产业+6.9%（占57.8%）、一般公共预算收入+4.0%（超额）、实际使用外资+6.4%**。")
para(doc, "为什么\u201cGDP/工业/投资\u201d大幅不及目标，服务业与财政却不错？三解释：")
para(doc, "**其一，制造业/建筑业/地产承压**　规上工业+4.1%（低目标8%）、建筑业-42.6%（总产值-46.9%）、地产-39.5%、高技术制造占比重但利润-17.8%——\u201c工业制造/建筑/地产深度调整\u201d。")
para(doc, "**其二，三产、横琴、服务强**　第三产业+6.9%（占57.8%）、高技术服务业+9.7%、租赁商务+17.3%、软件信息+11.9%、横琴粤澳合作——\u201c服务业+横琴强\u201d。")
para(doc, "**其三，投资/地产/外贸弱但财政外资好**　固投-31.6%、外贸+3.4%、CPI-0.7%；财收+4.0%（税收+15.1%）、外资+6.4%——\u201c财政外资稳、需求弱\u201d。")
para(doc, "小结：珠海2025年是\u201c**产业调整、投资地产大幅下行，但服务业/横琴/财政强**\u201d：集成电路/高端装备/新质强，制造建筑地产、固投CPI弱。")

# ---- 四、容易被忽视、但信号很重的细节（15条） ----
heading1(doc, "四、容易被忽视、但信号很重的细节（15条）")
bullet(doc, "**1.集成电路+26.5%、高端装备+21.6%、新能源+12.0%**\u201c4+3产业强（新质）。\u201d")
bullet(doc, "**2.生物医药-5.9%、利润总额-17.8%（盈-11.4%/亏损面31.4%）**\u201c医药、工业盈利承压。\u201d")
bullet(doc, "**3.固投-31.6%（工业-14.6%/基建-35.1%）**\u201c基建/扩产收缩。\u201d")
bullet(doc, "**4.房地产开发-39.2%、新开工-35.0%、竣工-73.8%**\u201c地产深度调整。\u201d")
bullet(doc, "**5.建筑业-42.6%、总产值-46.9%**\u201c建筑业/地产链条大缩。\u201d")
bullet(doc, "**6.第三产业+6.9%（高技术服务+9.7%）**\u201c服务/横琴强。\u201d")
bullet(doc, "**7.进出口3354.33亿/+3.4%、出口-0.2%**\u201c外贸稳、出口微降。\u201d")
bullet(doc, "**8.一般公共预算收入494.14亿/+4.0%（税收+15.1%）**\u201c财收超额（税收+15.1%）。\u201d")
bullet(doc, "**9.实际使用外资9.95亿美元/+6.4%（新设企业+21.7%）**\u201c外资稳、香港66.9%。\u201d")
bullet(doc, "**10.高新技术企业（通过认定865家）、研发投入占4.6%**\u201c创新强市（全省第2）。\u201d")
bullet(doc, "**11.横琴粤澳、港珠澳大桥（2329亿+11.3%）、口岸2亿**\u201c横琴-大桥。\u201d")
bullet(doc, "**12.中国航展（第十五届1022家企业、近20年最热闹）**\u201c航空航天+展会。\u201d")
bullet(doc, "**13.居民收入69532元/+3.3%、城乡比1.71（收窄）**\u201c农村+5.5%>城镇+3.2%。\u201d")
bullet(doc, "**14.CPI-0.7%（食品烟酒-0.6%、交通通信-3.2%）**\u201c低通胀、需求弱。\u201d")
bullet(doc, "**15.常住254.31万/+2.46万、城镇化90.96%**\u201c人口净增、城镇化极高。\u201d")
# ---- 五、2025年GDP目标 vs 实际：对照表 ----
heading1(doc, "五、2025年GDP目标 vs 实际：对照表")
table(doc,
    ["指标", "2025年目标", "2025年实际", "达成评价"],
    [
        ["地区生产总值(GDP)", "增长6%左右", "4573.10亿/2.7%", "大幅不及"],
        ["规模以上工业", "增长8%", "+4.1%", "差3.9pct"],
        ["固定资产投资", "增长9%", "-31.6%", "大幅下行"],
        ["社会消费品零售总额", "增长5%", "918.65亿/+1.7%", "差3.3pct"],
        ["进出口总额", "增长5%", "3354.33亿/+3.4%", "差1.6pct"],
        ["一般公共预算收入", "增长3%", "494.14亿/+4.0%", "超额"],
        ["居民收入", "与经济增长同步", "69532元/+3.3%", "略高GDP"],
    ],
)
para(doc, "注：GDP、规上工业按可比价。**财收（+4.0%）、外资（+6.4%）超额**；**GDP（+2.7%）、固投（-31.6%）、工业（+4.1%）、社零（+1.7%）、进出口（+3.4%）不及目标**。")
para(doc, "拆读：**4+3新质产业（集成电路/高端装备）、第三产业、横琴、财政、外资是亮色**；**二产（-2.9%）、固投（-31.6%）、地产（-39.5%）、建筑业（-42.6%）、CPI（-0.7%）**是短板——\u201c产业结构调整、投资地产深度下行\u201d，是\u201c特区转型\u201d样本。")

# ---- 六、2025年增长是谁撑起来的？（结构归因） ----
heading1(doc, "六、2025年增长是谁撑起来的？（结构归因）")
para(doc, "把珠海GDP的2.7%拆开：第一产业+6.2%（贡献3.3%）、第二产业-2.9%（贡献-46.7%）、第三产业+6.9%（贡献143.4%）（结构1.6：40.6：57.8）。**第三产业是唯一正拉动、贡献143.4%，第二产业（工业/建筑/地产）拖累-46.7%**。")
para(doc, "2026年珠海强调\u201c横琴+新质生产力+大湾区西岸\u201d，聚焦**集成电路/生物医药/新能源、横琴粤澳、高端装备、港珠澳大桥经贸、低空经济**——核心是\u201c新质+横琴+转型\u201d。")
para(doc, "**第二产业（工业/建筑）**：规上+4.1%（集成电路+26.5%、高端装备+21.6%、新能源+12.0%；生物医药-5.9%）、建筑业-42.6%——\u201c4+3强但建筑地产拖累\u201d。")
para(doc, "**第三产业（服务业）**：+6.9%（高技术服务+9.7%、租赁商务+17.3%、软件信息+11.9%、文化旅游/横琴）——\u201c服务业+横琴\u201d强。")
para(doc, "**第一产业（农业/海洋）**：+6.2%（水产品42.76万吨+9.2%、渔业/海洋牧场）——\u201c农业海洋稳\u201d。")
para(doc, "一句话归因：**2025年珠海增长\u201c靠第三产业（横琴+服务）+新质工业\u201d，二产/建筑/地产拖累**；是\u201c结构调整、增速换挡\u201d之年。")

# ---- 七、预算与财政的\u201c含金量\u201d ----
heading1(doc, "七、预算与财政的\u201c含金量\u201d")
para(doc, "2025年珠海**一般公共预算收入494.14亿元（+4.0%）**；税收收入336.96亿（+15.1%）；一般公共预算支出668.44亿（+2.8%）。")
bullet(doc, "税收+15.1%（增值税+29.6%、个税+40.0%、企业所得税+22.1%）——\u201c税收强、含金量高\u201d。")
bullet(doc, "财收超额（3%目标）——\u201c财政充裕\u201d。")
bullet(doc, "金融支撑：存款13230亿（+3.7%）、贷款12076亿（+4.1%）、境外存款/贷款高增长——跨境金融+横琴。")
para(doc, "**财政含金量小结**：财收+4.0%（高GDP增速）、税收+15.1%（含金量足）、支持横琴/新质/民生；财政是2025年珠海亮点。")

# ---- 八、民生底账：人口、收入与城乡 ----
heading1(doc, "八、民生底账：人口、收入与城乡")
para(doc, "2025年珠海**居民人均可支配收入69532元（+3.3%，实际+4.0%）**，其中城镇72259元（+3.2%）、农村42270元（+5.5%），城乡比1.71（收窄）。消费：人均消费支出43053元（+1.0%）。就业：城镇新增就业5.11万人。")
para(doc, "人口画像：**常住254.31万/+2.46万、城镇化率90.96%**（全国最高之一）、跨境/人才净流入。")
para(doc, "民生投入：12年免费义务教育、高校12.32万在校生、医疗、保障房12443套——民生投入扎实但收入增速偏低。")

# ---- 九、城镇与农村：格局与均衡 ----
heading1(doc, "九、城镇与农村：格局与均衡")
para(doc, "珠海城镇化率90.96%（全国地级市最高）；城乡协调发展：农村收入增速（+5.5%）>城镇（+3.2%），**城乡比1.71收窄**；斗门/海洋渔业、横琴。")
para(doc, "农业底盘：**水产品42.76万吨（+9.2%）、海洋捕捞/养殖、粮食稳**——\u201c向海要粮+海洋牧场\u201d。")
para(doc, "一句话：\u201c珠海高度城市化、农村总量小但增势好、海洋渔业\u201d。")

# ---- 十、人口流入与流出 ----
heading1(doc, "十、人口流入与流出")
para(doc, "珠海常住254.31万（净增+2.46万）、城镇化90.96%；\u201c横琴/产业/宜居\u201d吸引（粤港澳人口、港澳居民北移、人才），\u201c港珠澳大桥+横琴\u201d带来跨境人口。")
para(doc, "结构观察：**出生率7.86‰、自然增长+3.79‰、人口结构好**；老龄化低于全国平均。")
para(doc, "2026年目标：引才（研发/高企）、人才总量超90万——珠海靠\u201c横琴+产业+宜居\u201d聚人。")

# ---- 十一、物价与货币环境 ----
heading1(doc, "十一、物价与货币环境")
para(doc, "2025年珠海**CPI-0.7%**（食品烟酒-0.6%、交通通信-3.2%、医疗保健-3.1%；其他用品+13.1%、居住0.5%）——\u201c低通胀、需求弱\u201d。")
bullet(doc, "信贷扩张：贷款+4.1%、境外贷款+21%（跨境金融）、存款+3.7%。")
bullet(doc, "消费：社零+1.7%、家电+17.8%、通讯器材+80%（以旧换新）——结构消费。")
para(doc, "货币环境判断：**宽信用、CPI-0.7%**；珠海靠\u201c横琴+服务+新质\u201d稳需求（2026 CPI回升）。")

# ---- 十二、区域一体化：珠海的位置 ----
heading1(doc, "十二、区域一体化：珠海在大湾区、横琴-港澳、珠江口西岸\u201c三极\u201d里的位置")
para(doc, "珠海是**粤港澳大湾区重要极点、唯一与澳门陆地相连的城市、横琴粤澳深度合作区（服务澳门）、珠江口西岸核心城市（+中山共建）**。")
bullet(doc, "**横琴粤澳**　服务澳门经济多元、跨境创新（集成电路/中医药/文旅）——\u201c横琴配澳门\u201d。")
bullet(doc, "**港珠澳大桥/开放**　大桥经贸新通道、口岸出入境2亿+、跨境金融——内陆与港澳新枢纽。")
bullet(doc, "**西岸核心**　联动中山/江门、大湾区（广州/深圳东莞）、澳珠极点。")
para(doc, "一句话：**珠海在\u201c横琴+大桥+珠江口西岸\u201d里，最核心是\u201c服务澳门（横琴）+开放枢纽\u201d**；区位、特区、连接港澳是最大优势。")

# ---- 十三、未来5—10年最值得观察的五条主线 ----
heading1(doc, "十三、未来5—10年最值得观察的五条主线")
bullet(doc, "**主线一：横琴粤澳合作区**\u201c澳门多元、集成电路/中医药\u201d能否成国家级创新平台。")
bullet(doc, "**主线二：4+3产业/新质**\u201c集成电路/生物医药/新能源\u201d能否做大万亿产业。")
bullet(doc, "**主线三：港珠澳大桥/低空经济**\u201c大桥经贸、跨境/低空\u201d大湾区西岸枢纽。")
bullet(doc, "**主线四：海洋经济/绿色**\u201c海洋牧场、清洁能源、零碳\u201d。")
bullet(doc, "**主线五：人口/产业转型加减**\u201c制造业投资回升、地产止跌、聚才\u201d如何\u201c补短板、稳大盘\u201d。")

# ---- 十四、最终结论 ----
heading1(doc, "十四、最终结论：珠海在\u201c横琴+4+3产业+大湾区枢纽\u201d里的增长逻辑")
para(doc, "把2025年闭环打穿：**珠海是\u201c特区、横琴、大湾区西岸核心\u201d**：GDP 4573.10亿/+2.7%（转型）、集成电路+26.5%、第三产业+6.9%、财收+4.0%、外资+6.4%。")
para(doc, "珠海不是\u201c只有海岛/航展\u201d——它是**横琴（服务澳门）+4+3新质+港珠澳大桥（开放）**特区型城市，靠\u201c服务+新质\u201d驱动；但二产、投资、地产大幅承压。")
para(doc, "一句话结论：**珠海是\u201c特区、横琴粤澳、珠江口西岸\u201d；观察它先看\u201c集成电路、高端装备、横琴、港珠澳大桥、财政\u201d，再看\u201c二产、地产、固定投资\u201d。**它是\u201c结构调整、转型徘徊、服务+横琴强\u201d的特区样本。")

# ---- 附录A ----
heading1(doc, "附录A：主要资料来源与核验路径")
bullet(doc, "《2025年珠海市政府工作报告》（2025年2月，市十四届人大五次会议，2025年目标、2024年回顾4479.1亿/+3.5%）")
bullet(doc, "《2025年珠海市国民经济和社会发展统计公报》（珠海市统计局/国家统计局珠海调查队，2026-06-30，2025年实际）")
bullet(doc, "《2026年珠海市政府工作报告》（2026年2月，复盘+2026年目标）")
bullet(doc, "珠海市人民政府/统计局（zhuhai.gov.cn/tjj.zhuhai.gov.cn）")

# ---- 附录B ----
heading1(doc, "附录B：建议建立的年度跟踪仪表盘")
bullet(doc, "GDP总量/增速 vs 全年目标。")
bullet(doc, "规模以上工业（4+3产业）增速。")
bullet(doc, "集成电路/高端装备/生物医药。")
bullet(doc, "横琴粤澳、跨境金融、港澳。")
bullet(doc, "港珠澳大桥、口岸出入境。")
bullet(doc, "固定资产/工业/房地产投资。")
bullet(doc, "社零、以旧换新、文旅。")
bullet(doc, "进出口、外资、横琴贸易。")
bullet(doc, "常住/城镇化、青年。")
bullet(doc, "CPI、存贷款、财政/税收。")

# ---------------- 保存 ----------------
out = "/Users/x/Desktop/content-prod-lab/reports/珠海市_2025年政府工作报告_深度研究_2026-08-14.docx"
doc.save(out)
print("SAVED OK: 珠海市", out)
