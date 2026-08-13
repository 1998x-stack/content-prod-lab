# -*- coding: utf-8 -*-
"""Build 重庆市2025年政府工作报告 深度研究 DOCX, 参照北京/上海/杭州/天津系列版式。"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DARK = RGBColor(0x1F, 0x2A, 0x44)
RED = RGBColor(0xB0, 0x1E, 0x2E)
GRAY = RGBColor(0x59, 0x60, 0x69)
BLUE = RGBColor(0x1F, 0x4E, 0x79)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
HEADER_FILL = "1F2A44"
K = "微软雅黑"


def set_run(run, size=11, bold=False, color=DARK, font=K):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    run.font.name = font
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), font)
    rFonts.set(qn("w:hAnsi"), font)
    rFonts.set(qn("w:eastAsia"), font)


def para(doc, text, size=11, bold=False, color=DARK, align=None,
         space_after=6, space_before=0, font=K):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        set_run(r, size=size, bold=bold or (i % 2 == 1), color=color, font=font)
    return p


def heading1(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(16)
    pf.space_after = Pt(8)
    r = p.add_run(text)
    set_run(r, size=16, bold=True, color=RED)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "B01E2E")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def heading2(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(10)
    pf.space_after = Pt(4)
    r = p.add_run(text)
    set_run(r, size=13, bold=True, color=BLUE)
    return p


def table(doc, headers, rows, widths=None, font_size=9.5):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_run(r, size=font_size, bold=True, color=WHITE)
        tcPr = hdr[i]._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), HEADER_FILL)
        tcPr.append(shd)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(val))
            set_run(r, size=font_size, color=DARK)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    return t


def quote_box(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(8)
    pf.left_indent = Pt(12)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), "B01E2E")
    pBdr.append(left)
    pPr.append(pBdr)
    r = p.add_run(text)
    set_run(r, size=11, color=DARK, font="楷体")
    return p


def bullet(doc, text, size=10.5):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.7)
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        set_run(r, size=size, bold=(i % 2 == 1), color=DARK)
    return p


# ================================================================= 文档主体
doc = Document()
sec = doc.sections[0]
sec.page_width = Cm(21)
sec.page_height = Cm(29.7)
sec.top_margin = Cm(2.4)
sec.bottom_margin = Cm(2.2)
sec.left_margin = Cm(2.5)
sec.right_margin = Cm(2.5)

st = doc.styles["Normal"]
st.font.name = K
st.font.size = Pt(11)
st.element.rPr.rFonts.set(qn("w:eastAsia"), K)


# ---- 封面 ----
para(doc, "广州市2025年政府工作报告", size=24, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=60, space_after=4)
para(doc, "深度研究与“容易被忽视的细节”分析", size=16, bold=True, color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
para(doc, "从“南沙自贸区、广深联动、制造业立市与民营经济”重新理解广州", size=12, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)
para(doc, "━━━━━━━━━━━━━━━━", color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
para(doc, "研究对象：2025年广州市政府工作报告及同期财政、国民经济和社会发展统计资料、2026年官方复盘", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
para(doc, "整理时间：2026年8月", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
para(doc, "分析性质：分析性研究 ｜ 非官方解读", size=11, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
doc.add_page_break()

# ---- 目录 ----
catalog = [
    "执行摘要",
    "一、研究口径与阅读方法",
    "二、先看广州的特殊底盘：千年商都、湾区核心与省会门户",
    "三、最关键的宏观错位：GDP破3.2万亿、出口强，但工业与房产偏弱",
    "四、容易被忽视、但信号很重的细节（15条）",
    "五、把所有细节连起来：广州正在经历的“六个换挡”",
    "六、增长暗线：制造业立市、数字经济与南沙/中新知识城",
    "七、财政暗线：收入稳、税收占比高、民生/城中村改造",
    "八、产业暗线：从“商贸/汽车”到“先进制造+新质生产力”",
    "九、区域格局：广深联动、南沙与都市圈",
    "十、人口与城市：超1900万人、公共服务与城中村改造",
    "十一、民营经济：占43.1%、消费大市",
    "十二、事后验证：用2025年实际执行结果检验报告判断",
    "十三、未来5—10年最值得观察的五条主线",
    "十四、最终结论：广州在“粤港澳大湾区+国家中心城市”里的增长逻辑",
    "附录A：主要资料来源与核验路径",
    "附录B：建议建立的年度跟踪仪表盘",
]
para(doc, "目　录", size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=10, space_after=10)
for item in catalog:
    para(doc, item, size=12, space_after=4)
doc.add_page_break()

# ---- 执行摘要 ----
heading1(doc, "执行摘要")
para(doc, "**核心判断**　如果只看新闻标题，2025年广州最显眼的是“GDP突破3.2万亿、增长4.0%”、“出口增长17.8%”、“数字经济核心产业+6.0%”和“社零破1.1万亿”。但这份研究真正值得深读的，是这一座“千年商都+制造业立市”的城市如何在第二产业偏弱（规上工业+1.2%）、房地产投资（-2.8%）与办公楼/商业地产大幅收缩的背景下，靠“出口、消费、数字经济、战略性新兴产业与民营”守住基本盘。")
para(doc, "把2025年初设定的目标、2025年《国民经济和社会发展统计公报》、以及2026年报告对2025年的复盘放在一起看，广州呈现清晰暗线：**从“商贸+房地产+传统制造”的旧依赖，向“制造业立市+数字经济+新质生产力+南沙开放”转型**。旧引擎（房地产、传统商贸地产、一般制造）在调整；新引擎（出口、战略性新兴、数字经济、新能源汽车、集成电路）被要求更快补位。")
para(doc, "因此，本报告不按政府工作报告逐段复述，而采用“显性表述—同期数据—制度含义—长期影响”的方式，专门提取容易被忽视、但对判断广州未来5—10年发展模式更有价值的信号。")
para(doc, "最容易记住的一句话：**广州是全国性综合门户、商贸中心与国家中心城市，也是“粤港澳大湾区”的四大核心市之一——它以“万亿商业+开放+数字经济”为底色，既要做“制造业立市”，也要在房地产、工业结构调整中，靠出口、新质生产力与城市更新守住增长。**观察广州，与其看“3万亿 GDP”，不如看“制造业/数字经济、出口、战略性新兴产业、城中村改造与民营活力”这几张名片。")

heading2(doc, "一页速览：2025年广州经济的“表与里”")
table(doc,
    ["维度", "表面现象", "更深层信号"],
    [
        ["增长", "GDP 32039.46亿、+4.0%", "总量破3.2万亿、三产占比75%"],
        ["产业", "规上工业+1.2%、数字经济+6.0%", "战略新兴占GDP32.4%、先进制造占58.2%"],
        ["外贸", "出口+17.8%、进出口+10.4%", "用出口支撑增长"],
        ["投资", "固定资产投资-6.7%", "地产-2.8%、办公楼-37.2%"],
        ["财政", "一般公共预算收入2184.82亿、+3.1%", "税收占71.4%"],
        ["消费", "社零11032.38亿、+5.5%", "网络零售+13.1%"],
        ["人口", "常住1910.10万、城镇化率87.56%", "人口+12.3万、户籍增长"],
        ["民营", "增加值1.38万亿、占43.1%、+5.2%", "民营底盘厚"],
    ],
    widths=[2.2, 5.2, 8.6])
para(doc, "", size=4, space_after=2)

# ---- 一、研究口径与阅读方法 ----
heading1(doc, "一、研究口径与阅读方法")
heading2(doc, "1.1 本报告的三份核心底稿")
bullet(doc, "**2025年《政府工作报告》**：2025年2月在市十六届人大五次会议上作，给出全年预期目标（GDP增长5%左右、实际工作中力争更好结果等）。固定资产投资/社零/进出口等未设全部具体速率。")
bullet(doc, "**2025年《国民经济和社会发展统计公报》**：广州市统计局2026年5月10日发布，给出全年实际执行数与增速，是“事后验证”的锚点。")
bullet(doc, "**2026年广州市政府工作报告及官方复盘**：对2025执行结果的追认，交叉验证“目标—执行—再评价”三段式。")
heading2(doc, "1.2 阅读方法：显性—数据—制度—长期")
para(doc, "每一章按“**显性表述 → 同期数据 → 制度含义 → 长期影响**”四层展开。其中“制度含义”回答“政府为什么这么排优先序”，“长期影响”回答“这对未来5—10年意味着什么”。")
para(doc, "关键判别：**当报告使用的“增长词”和统计公报里的实际数不一致时，数据优先。**例如2025年GDP目标5%左右、实际4.0%（低于目标）；出口+17.8%超强、但规上工业+1.2%偏弱。差异清晰反映：广州“出口/消费/数字经济强、工业/地产偏弱”。")

# ---- 二、特殊底盘 ----
heading1(doc, "二、先看广州的特殊底盘：千年商都、湾区核心与省会门户")
para(doc, "在所有一线城市里，广州的“底盘”独特：**全国性商贸中心+千年商都+粤港澳大湾区核心市+省会门户**四合一。常住1910万、城镇化率87.56%，商贸、交通、医疗、科教枢纽地位突出。")
para(doc, "这决定广州的多重身份并存：**商贸中心**（社零1.1万亿）、**先进制造/汽车**（广州汽车）、**数字经济**（数字经济核心产业+6.0%）、**开放门户**（南沙自贸区、白云机场、广交会）、**省会服务**（全省金融/医疗/文化/枢纽）。")
heading2(doc, "2.1 制造业立市与汽车")
para(doc, "广州坚持“制造业立市”，先进制造占规上58.2%、装备占50.5%。新能源汽车产量+21.6%、集成电路+4.8%（模拟芯片+19.1%）、锂电+68.8%、充电桩+47.3%。广州制造在“新能源+高端制造”上升级。")
heading2(doc, "2.2 数字经济与战略新兴产业")
para(doc, "数字经济核心产业增加值增长6.0%（数字要素驱动+10.3%）；“3+5”战略性新兴产业增加值1.04万亿、占GDP 32.4%（+4.2%）。数字经济+战新是广州新质生产力的两大抓手。")
heading2(doc, "2.3 开放门户与南沙")
para(doc, "广州是“一带一路”与粤港澳大湾区门户，出口+17.8%、进出口+10.4%、实际使用外资+9.1%。南沙自贸区、中新知识城是制度型开放与科创平台。")

# ---- 三、宏观错位 ----
heading1(doc, "三、最关键的宏观错位：GDP破3.2万亿、出口强，但工业与房产偏弱")
para(doc, "把2025年广州的宏观面放进一张表，会出现令人惊讶的“错位”：表观增长来自出口/消费/数字经济，而工业与商业地产却在收缩。这个错位，正是读懂广州的关键。")
heading2(doc, "3.1 “强的一面”")
table(doc,
    ["指标", "2025实绩", "信号"],
    [
        ["GDP", "32039.46亿、+4.0%", "广州破3.2万亿"],
        ["出口", "+17.8%", "出口超强、稳外贸"],
        ["数字经济核心产业", "+6.0%", "数字新动能"],
        ["社零", "+5.5%", "消费大市"],
        ["战略性新兴", "占GDP 32.4%", "新质生产力"],
        ["民营经济", "占43.1%、+5.2%", "民营底盘厚"],
    ],
    widths=[3.2, 5.2, 6.0])
heading2(doc, "3.2 “弱的一面”")
table(doc,
    ["指标", "2025实值", "信号"],
    [
        ["规上工业", "+1.2%", "工业偏弱"],
        ["固定资产投资", "-6.7%", "投资收缩"],
        ["办公楼投资", "-37.2%", "商业地产大幅收缩"],
        ["房地产开发", "-2.8%", "地产调整"],
        ["GDP增速", "4.0%（低于5%目标）", "低于目标"],
    ],
    widths=[3.2, 5.2, 6.0])
para(doc, "**错位结论**　广州的增长“很强、但也很矛盾”。强的部分（出口/消费/数字经济/战新/民营）与弱的部分（工业/固定资产/商业地产）并存。**真正的焦点是“工业与地产承压，靠出口/消费/数字经济补位”**。2026年广州“制造业立市+稳住工业+城市更新/城中村改造+扩大内需”是重点。")

# ---- 四、被忽视细节 ----
heading1(doc, "四、容易被忽视、但信号很重的细节（15条）")
details = [
    ("1", "规上工业+1.2%、但先进制造占58.2%", "工业总量弱、但结构高端。"),
    ("2", "新能源汽车产量+21.6%、锂电+68.8%、充电桩+47.3%", "新能源车市加快。"),
    ("3", "集成电路产量+4.8%、模拟芯片+19.1%", "芯片补链。"),
    ("4", "数字经济核心产业+6.0%、数字产品服务+15.0%", "数字经济新动能。"),
    ("5", "“3+5”战略性新兴产业占GDP 32.4%", "新质生产力兑现。"),
    ("6", "出口+17.8%、进出口+10.4%", "出口支撑增长。"),
    ("7", "社会消费品零售11032.38亿、+5.5%", "消费大市。"),
    ("8", "网络零售+13.1%", "数字化消费升级。"),
    ("9", "固定资产投资-6.7%、房地产-2.8%", "投资/地产偏弱。"),
    ("10", "民营经济增加值1.38万亿、占43.1%", "民营底盘厚、+5.2%。"),
    ("11", "常住1910.10万、+12.3万、城镇化率87.56%", "人口净流入。"),
    ("12", "一般公共预算收入2184.82亿、+3.1%、税收占71.4%", "财政稳、税收高。"),
    ("13", "实际使用外资+9.1%", "外资逆势流入。"),
    ("14", "居民人均可支配收入80591元、+3.6%", "收入与增长低预期。"),
    ("15", "CPI+0.2%、PPI-2.2%", "物价稳、通缩压力。"),
]
for d in details:
    para(doc, "**%s**" % d[0], size=11)
    para(doc, d[1], size=10.5, color=GRAY)

# ---- 五、六个换挡 ----
heading1(doc, "五、把所有细节连起来：广州正在经历的“六个换挡”")
gear = [
    ("1．动能换挡：从“地产/商贸驱动”到“数字经济+出口+新质生产力”", "数字经济+6.0%、战新占32.4%、出口+17.8%。"),
    ("2．产业换挡：从“传统制造/商贸”到“制造业立市+新能源/高端制造”", "新能源车+21.6%、先进制造占58.2%。"),
    ("3．投资换挡：从“基建/办公楼地产”到“工业/城市更新”", "办公楼-37.2%、工业技改+3.0%。"),
    ("4．开放换挡：从“传统商贸”到“陆海/南沙开放+高新技术贸易”", "出口+17.8%、实际外资+9.1%。"),
    ("5．人口换挡：从“大规模流入”到“高质量集聚+存量”", "常住+12.3万、城镇化87.56%。"),
    ("6．动能换挡：从“总量扩张”到“高质量发展+新质生产力”", "战新占GDP32.4%、数字经济加速。"),
]
for g in gear:
    para(doc, "**%s**　%s" % g, space_after=6)

# ---- 六、增长暗线 ----
heading1(doc, "六、增长暗线：制造业立市、数字经济与南沙/中新知识城")
heading2(doc, "6.1 制造业立市+新质生产力")
para(doc, "广州坚持“制造业立市”，先进制造占58.2%、装备占50.5%、新能源汽车+21.6%、锂电+68.8%。数字经济和“3+5”战略性新兴产业（占GDP32.4%）是广州未来动能的引擎。")
heading2(doc, "6.2 南沙与中新知识城")
para(doc, "南沙自贸区、中新知识城等平台承载制度型开放与科创。广州以开放门户（出口+17.8%）+数字/科创双轮，构建“龙头城市”发展模式。")
para(doc, "**这条暗线意味着**：广州的增长已从“商贸+地产”转向“数字经济+战新+开放+制造业立市”。看广州，盯住“战新占GDP比重、数字经济、出口、城中村改造”这几组数。")

# ---- 七、财政暗线 ----
heading1(doc, "七、财政暗线：收入稳、税收占比高、民生/城中村改造")
para(doc, "2025年广州一般公共预算收入2184.82亿、+3.1%，税收占71.4%（+1.5%），收入质量高、主要来自商贸/制造/服务税源。支出向民生、城市更新（城中村改造）、科创倾斜。")
para(doc, "**制度含义**　广州财政“税收占比高、收入稳健”，关键是摆脱对地产依赖，靠“数字经济+商贸+制造+城市更新”产生稳健税源。")

# ---- 八、产业暗线 ----
heading1(doc, "八、产业暗线：从“商贸/汽车”到“先进制造+新质生产力”")
heading2(doc, "8.1 广州产业的“表”")
table(doc,
    ["指标", "2025增速/信号", "解读"],
    [
        ["规上工业", "+1.2%", "工业偏弱"],
        ["先进制造", "占58.2%", "制造业立市"],
        ["装备制造", "占50.5%", "高端装备"],
        ["新能源汽车", "+21.6%", "新能源车"],
        ["数字经济核心产业", "+6.0%", "数字动能"],
        ["战略性新兴", "占GDP32.4%", "新质生产力"],
        ["出口", "+17.8%", "外贸强"],
    ],
    widths=[4.6, 3.4, 5.2])
heading2(doc, "8.2 从“商贸/汽车/地产”到“制造+数字+战新”")
para(doc, "广州历史上以商贸、汽车、地产见长，2025年显示“数字经济+战新+新能源+先进制造”正成为新主导。广交会、商贸撑起基本盘，而新质生产力是未来弹性。")

# ---- 九、区域 ----
heading1(doc, "九、区域格局：广深联动、南沙与都市圈")
para(doc, "广州是粤港澳大湾区四大核心市（广深港澳）之一，与深圳“双城”联动，共同承接制造、科创、开放。广州都市圈（广佛等）带动广东西北/粤北。")
para(doc, "南沙自贸区、中新知识城、白云机场、广交会，构成广州制度型开放与枢纽。广深联动+湾区一体化，是广州重要增长点。")

# ---- 十、人口与城市 ----
heading1(doc, "十、人口与城市：超1900万人、公共服务与城中村改造")
heading2(doc, "10.1 人口总量")
para(doc, "2025年末广州常住1910.10万、城镇化率87.56%、净增12.3万；户籍人口1091.14万、户籍自然增长率4.83‰。广州仍是全国人口/人才流入高地。")
heading2(doc, "10.2 城市更新与城中村改造")
para(doc, "广州推进城中村改造与城市更新，盘活存量土地与商业地产。城市更新既是民生工程（老旧小区/城中村），也是广州“去地产化”后新的投资与空间来源。")

# ---- 十一、民营经济 ----
heading1(doc, "十一、民营经济：占43.1%、消费大市")
heading2(doc, "11.1 民企地位")
para(doc, "2025年广州民营经济增加值13798.05亿、占GDP 43.1%、+5.2%。民营企业是广州消费、商贸、制造与创新的主力。")
heading2(doc, "11.2 政策与服务")
para(doc, "广州持续优化营商环境、稳外资稳外贸、支持民营。民营+专精特新是广州未来10年的底盘。")

# ---- 十二、事后验证 ----
heading1(doc, "十二、事后验证：用2025年实际执行结果检验报告判断")
heading2(doc, "12.1 年初目标与年末执行对比")
table(doc,
    ["指标", "2025年初目标", "2025实际", "偏差"],
    [
        ["GDP增速", "5%左右", "+4.0%", "低于目标"],
        ["规上工业", "力争", "+1.2%", "偏弱"],
        ["社零", "力争5%+", "+5.5%", "达标"],
        ["进出口", "稳量提质", "+10.4%", "超预期"],
        ["固定资产投资", "稳", "-6.7%", "偏弱"],
        ["居民收入", "与增长同步", "+3.6%", "偏低"],
    ],
    widths=[3.0, 3.2, 3.0, 4.2])
heading2(doc, "12.2 判断验证")
para(doc, "三条核心判断得到支持：**（1）出口/消费/数字经济强（+17.8%/+5.5%/+6.0%），验证；（2）工业/地产偏弱（规上+1.2%/办公楼-37.2%），验证；（3）民营/战新/南沙强，验证。**")
para(doc, "核心观察：广州在“工业与地产偏弱、GDP低于5%目标”下，靠“出口+消费+数字经济+战新+民营”守住3.2万亿。2026年“制造业立市+稳住工业/地产+城市更新/城中村改造+扩内需”是重点。")

# ---- 十三、五条主线 ----
heading1(doc, "十三、未来5—10年最值得观察的五条主线")
lines5 = [
    ("① 制造业立市与先进制造/新质生产力", "先进占58%、战新占32%，能否把工业做大、稳住。"),
    ("② 数字经济与科创", "数字经济+6%，能否持续放量与产业化。"),
    ("③ 出口/开放/南沙/广深联动", "出口+17.8%能否持续、南沙/中新知识城可否兑现。"),
    ("④ 房地产/城市更新/城中村改造", "地产偏弱，能否用城市更新盘活空间与投资。"),
    ("⑤ 民营与消费", "民营43%、社零1.1万亿，能否在消费升级中扩大。"),
]
for l in lines5:
    para(doc, "**%s**　%s" % l, space_after=6)

# ---- 十四、结论 ----
heading1(doc, "十四、最终结论：广州在“粤港澳大湾区+国家中心城市”里的增长逻辑")
para(doc, "广州的2025年，本质上是“**出口+消费+数字经济+战新为核心，而对地产/工业依赖下降**”的答卷：GDP 3.2万亿、出口+17.8%、数字经济+6.0%，代价是规上工业+1.2%、地产/投资收缩、GDP低于目标。")
para(doc, "只要“制造业立市、数字经济、出口、民营、城市更新”能接住，广州就仍是粤港澳国家中心；如果工业、地产持续偏弱，广州需承受“转型阵痛”。")
para(doc, "最稳妥的观察信号：**一盯先进制造/战新/数字经济（动能）、二盯出口/开放（外贸）、三盯工业/地产/城市更新（结构）、四盯民营/消费（底座）、五盯财政/民生（财政民生）。**广州，是粤港澳大湾区的“南大门”与商都样本。")

# ---- 附录A ----
heading1(doc, "附录A：主要资料来源与核验路径")
bullet(doc, "广州市2025年《政府工作报告》（2025年2月）——目标来源。")
bullet(doc, "《2025年广州市国民经济和社会发展统计公报》（广州市统计局，2026-05-10）——GDP、工业、外贸、人口实值。")
bullet(doc, "广州市统计/运行分析、大湾区专报、数字经济专篇——区域与数字。")
bullet(doc, "2026年广州市政府工作报告——2025执行复盘。")
bullet(doc, "广州海关、市商务部（南沙），市财政局——外贸/财政。")
heading2(doc, "核验说明")
para(doc, "本报告所有增速、总量、占比均以统计公报/官方发布为基准。涉“南沙自贸区专项”“中新知识城”等以官方口径为准。")

# ---- 附录B ----
heading1(doc, "附录B：建议建立的年度跟踪仪表盘")
para(doc, "建议把下面10个“测脉搏”指标做成年度跟踪表：")
dash = [
    ["#", "指标", "2025基值", "为什么重要"],
    ["1", "GDP增速", "+4.0%", "总量与方向"],
    ["2", "规上工业增速", "+1.2%", "制造业立市"],
    ["3", "数字经济核心产业", "+6.0%", "数字动能"],
    ["4", "战略性新兴占GDP", "32.4%", "新质生产力"],
    ["5", "出口增速", "+17.8%", "外贸韧性"],
    ["6", "固定资产投资增速", "-6.7%", "投资结构"],
    ["7", "社零增速", "+5.5%", "消费大市"],
    ["8", "常住人口/城镇化率", "1910.10万/87.56%", "人口与城市"],
    ["9", "民营经济占比", "43.1%", "民营活力"],
    ["10", "居民人均可支配收入增速", "+3.6%", "民生获得感"],
]
table(doc, dash[0], dash[1:], widths=[0.8, 4.5, 3.2, 5.0])
para(doc, "把这10个指标连起来看，任何一个“新引擎（3/4/5）向上、旧引擎（6）修复”，都说明广州在真正换挡；反之则是旧路径的反复。")

# ========================================= 保存
out = "/Users/x/Desktop/content-prod-lab/reports/广州市_2025年政府工作报告_深度研究_2026-08-13.docx"
doc.save(out)
print("SAVED:", out)
print("PARAGRAPHS:", len(doc.paragraphs))
print("TABLES:", len(doc.tables))
