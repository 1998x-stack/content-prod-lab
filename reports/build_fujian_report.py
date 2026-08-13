# -*- coding: utf-8 -*-
"""Build 重庆市2025年政府工作报告 深度研究 DOCX, 参照北京/上海/杭州/天津系列版式。"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DARK = RGBColor(0x1F, 0x2A, 0x44)
RED = RGBColor(0xB0, 0x1E, 0x2E)
GRAY = RGBColor(0x59, 0x60, 0x69)
BLUE = RGBColor(0x1F, 0x4E, 0x79)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
HEADER_FILL = "1F2A44"
K = "微软雅黑"


def set_run(run, size=11, bold=False, color=DARK, font=K):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    run.font.name = font
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), font)
    rFonts.set(qn("w:hAnsi"), font)
    rFonts.set(qn("w:eastAsia"), font)


def para(doc, text, size=11, bold=False, color=DARK, align=None,
         space_after=6, space_before=0, font=K):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        set_run(r, size=size, bold=bold or (i % 2 == 1), color=color, font=font)
    return p


def heading1(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(16)
    pf.space_after = Pt(8)
    r = p.add_run(text)
    set_run(r, size=16, bold=True, color=RED)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "B01E2E")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def heading2(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(10)
    pf.space_after = Pt(4)
    r = p.add_run(text)
    set_run(r, size=13, bold=True, color=BLUE)
    return p


def table(doc, headers, rows, widths=None, font_size=9.5):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_run(r, size=font_size, bold=True, color=WHITE)
        tcPr = hdr[i]._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), HEADER_FILL)
        tcPr.append(shd)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(val))
            set_run(r, size=font_size, color=DARK)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    return t


def quote_box(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(8)
    pf.left_indent = Pt(12)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), "B01E2E")
    pBdr.append(left)
    pPr.append(pBdr)
    r = p.add_run(text)
    set_run(r, size=11, color=DARK, font="楷体")
    return p


def bullet(doc, text, size=10.5):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.7)
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        set_run(r, size=size, bold=(i % 2 == 1), color=DARK)
    return p


# ================================================================= 文档主体
doc = Document()
sec = doc.sections[0]
sec.page_width = Cm(21)
sec.page_height = Cm(29.7)
sec.top_margin = Cm(2.4)
sec.bottom_margin = Cm(2.2)
sec.left_margin = Cm(2.5)
sec.right_margin = Cm(2.5)

st = doc.styles["Normal"]
st.font.name = K
st.font.size = Pt(11)
st.element.rPr.rFonts.set(qn("w:eastAsia"), K)


# ---- 封面 ----
para(doc, "福建省2025年政府工作报告", size=24, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=60, space_after=4)
para(doc, "深度研究与“容易被忽视的细节”分析", size=16, bold=True, color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
para(doc, "从“海峡西岸、民营经济发达、海洋强省与福州厦门双核”重新理解福建", size=12, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)
para(doc, "━━━━━━━━━━━━━━━━", color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
para(doc, "研究对象：2025年福建省政府工作报告及同期财政、国民经济和社会发展统计资料、2026年官方复盘", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
para(doc, "整理时间：2026年8月", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
para(doc, "分析性质：分析性研究 ｜ 非官方解读", size=11, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
doc.add_page_break()

# ---- 目录 ----
catalog = [
    "执行摘要",
    "一、研究口径与阅读方法",
    "二、先看福建的特殊底盘：海峡西岸、民营强省与海洋大省",
    "三、最关键的宏观错位：GDP破6万亿、工业/出口/民营有，但外贸与地产偏弱",
    "四、容易被忽视、但信号很重的细节（15条）",
    "五、把所有细节连起来：福建正在经历的“六个换挡”",
    "六、增长暗线：民营经济、海洋经济与制造业升级（新质生产力）",
    "七、财政暗线：收入稳、税收尚可，民生/基建投入大",
    "八、产业暗线：从“轻工/食品”到“高技术+装备+海洋”",
    "九、区域格局：福州都市圈、厦漳泉与海峡西岸",
    "十、人口与城市：沿海人口、城镇化与代际",
    "十一、民营经济：占八成经营主体、活跃开放",
    "十二、事后验证：用2025年实际执行结果检验报告判断",
    "十三、未来5—10年最值得观察的五条主线",
    "十四、最终结论：福建在“海洋强省+民营+海峡西岸”里的增长逻辑",
    "附录A：主要资料来源与核验路径",
    "附录B：建议建立的年度跟踪仪表盘",
]
para(doc, "目　录", size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=10, space_after=10)
for item in catalog:
    para(doc, item, size=12, space_after=4)
doc.add_page_break()

# ---- 执行摘要 ----
heading1(doc, "执行摘要")
para(doc, "**核心判断**　如果只看新闻标题，2025年福建最显眼的是“GDP突破6万亿、增长5.0%”、“高技术制造+14.6%、装备制造+13.9%”和“社零破2.5万亿”。但这份研究真正值得深读的，是这座“海峡西岸+民营强省+海洋大省”的省份，如何在进出口（-5.4%）、房地产开发投资（-30.0%）偏弱的背景下，靠“高技术制造+装备+新能源车+民营+消费”守住增长。")
para(doc, "把2025年初设定的目标（GDP增长5%）、2025年《国民经济和社会发展统计公报》、以及2026年报告对2025年的复盘放在一起看，福建呈现清晰暗线：**从“轻工/食品+传统外贸”的旧底盘，向“高技术制造+装备+海洋经济+民营+新能源车”转型**。旧引擎（房地产、传统轻工外贸）在调整；新引擎（高技术、装备、新能源汽车、民营、海洋）被要求更快补位。")
para(doc, "因此，本报告不按政府工作报告逐段复述，而采用“显性表述—同期数据—制度含义—长期影响”的方式，专门提取容易被忽视、但对判断福建未来5—10年发展模式更有价值的信号。")
para(doc, "最容易记住的一句话：**福建是“民营经济（占比高）+海洋大省+海峡西岸”的样本，用民企制造、高技术、海洋与开放的高效，在“外向与内需、东部与沿海”间守住增长。**观察福建，与其看“GDP 6万亿”，不如看“民营占比、高技术/装备制造、海洋经济、进出口与居民收入”这几张名片。")
heading2(doc, "一页速览：2025年福建经济的“表与里”")
table(doc,
    ["维度", "表面现象", "更深层信号"],
    [
        ["增长", "GDP 60199.45亿、+5.0%", "二产/三产各+5.1%"],
        ["产业", "规上工业+6.5%、高技术+14.6%", "装备+13.9%、电子+16.9%"],
        ["外贸", "进出口18826.99亿、-5.4%", "机电/高新技术出口微增"],
        ["投资", "固定资产投资-3.3%", "地产-30.0%、工业+6.1%"],
        ["财政", "一般公共预算收入3723.35亿、+3.0%", "财政稳"],
        ["消费", "社零25433.59亿、+4.4%", "线上+14.2%、新能源车+27%"],
        ["人口", "常住4190万、城镇化率72.58%", "人口-1.5‰、城镇化"],
        ["海洋", "水产品965万吨、港口7.56亿吨", "海洋强省"],
    ],
    widths=[2.2, 5.2, 8.6])
para(doc, "", size=4, space_after=2)

# ---- 一、研究口径与阅读方法 ----
heading1(doc, "一、研究口径与阅读方法")
heading2(doc, "1.1 本报告的三份核心底稿")
bullet(doc, "**2025年《政府工作报告》**：2025年1月在省十四届人大三次会议上作，给出全年预期目标（GDP增长5%左右、一般公共预算收入增长等）。固定资产投资/进出口等未全部设具体速率。")
bullet(doc, "**2025年《国民经济和社会发展统计公报》**：福建省统计局2026年3月19日发布，给出全年实际执行数与增速，是“事后验证”的锚点。")
bullet(doc, "**2026年福建省政府工作报告及官方复盘**：对2025执行结果的追认，交叉验证“目标—执行—再评价”三段式。")
heading2(doc, "1.2 阅读方法：显性—数据—制度—长期")
para(doc, "每一章按“**显性表述 → 同期数据 → 制度含义 → 长期影响**”四层展开。其中“制度含义”回答“政府为什么这么排优先序”，“长期影响”回答“这对未来5—10年意味着什么”。")
para(doc, "关键判别：**当报告使用的“增长词”和统计公报里的实际数不一致时，数据优先。**例如2025年GDP目标5%左右、实际5.0%达标；但进出口-5.4%、地产投资-30%。差异反映：福建“工业/消费/民营强，外贸/地产偏弱”。")

# ---- 二、特殊底盘 ----
heading1(doc, "二、先看福建的特殊底盘：海峡西岸、民营强省与海洋大省")
para(doc, "在所有省份里，福建的“底盘”独特：**海峡西岸+全国民营经济高地+海洋大省**三合一。常住4190万、城镇化率72.58%，电子、鞋服、运动装备、新能源等制造业与海洋、对台经贸交织。")
para(doc, "这决定福建的多重身份并存：**民营强省**（经营主体770万）、**海洋大省**（水产品965万吨、港口7.56亿吨）、**制造/电子**（高技术+14.6%）、**海峡西岸开放**（对台、海上丝绸之路）、**侨乡**（海外闽商）。")
heading2(doc, "2.1 民营与制造")
para(doc, "福建以民营为“底色”，高技术制造+14.6%、装备+13.9%、电子+16.9%。新能源汽车4.41万辆、+78.1%。民营+制造是福建最活跃的引擎。")
heading2(doc, "2.2 海洋强省")
para(doc, "水产品总产量965万吨（全国前列）、沿海港口吞吐量7.56亿吨、集装箱1843万标箱。海洋渔业/海工是福建特色与“蓝色经济”增量。")
heading2(doc, "2.3 海峡西岸与侨乡")
para(doc, "福建地处海峡西岸，厦门、福州等对台开放前沿；海外闽商/侨汇密集。对台合作与“海丝”是福建开放双面向。")

# ---- 三、宏观错位 ----
heading1(doc, "三、最关键的宏观错位：GDP破6万亿、工业/出口有，但外贸与地产偏弱")
para(doc, "把2025年福建的宏观面放进一张表，会出现令人意外的“错位”：表观增长来自工业/消费/民营，而外贸与地产偏弱。这个错位，正是读懂福建的关键。")
heading2(doc, "3.1 “强的一面”")
table(doc,
    ["指标", "2025实绩", "信号"],
    [
        ["GDP", "60199.45亿、+5.0%", "破6万亿"],
        ["规上工业", "+6.5%", "制造业+6.9%"],
        ["高技术制造", "+14.6%", "新质生产力"],
        ["装备制造", "+13.9%", "高端装备"],
        ["新能源汽车零售/产量", "+27%/+78.1%", "新能源车"],
        ["社零", "+4.4%", "消费稳"],
    ],
    widths=[3.2, 5.4, 6.0])
heading2(doc, "3.2 “弱/调整的一面”")
table(doc,
    ["指标", "2025实值", "信号"],
    [
        ["进出口", "-5.4%（出口-6.7%）", "外贸转弱"],
        ["房地产开发投资", "-30.0%", "地产深度调整"],
        ["固定资产投资", "-3.3%", "投资偏弱"],
        ["第三产业投资", "-9.9%", "三产投资降"],
        ["CPI", "0.0%", "物价平"],
    ],
    widths=[3.2, 5.4, 6.0])
para(doc, "**错位结论**　福建的增长“很强、但也很矛盾”。强的部分（工业/高技术/消费/民营）与弱的部分（进出口/地产/投资）并存。**真正的焦点是“制造/内需强，外贸/地产弱”**：高技术+14.6%拉动增长，但进出口-5.4%、地产-30%。2026年福建“稳民营/制造+扩内需+修外贸/地产”是重点。")

# ---- 四、被忽视细节 ----
heading1(doc, "四、容易被忽视、但信号很重的细节（15条）")
details = [
    ("1", "规上工业+6.5%、制造业+6.9%", "工业/民生良。"),
    ("2", "高技术制造+14.6%、装备+13.9%", "高技术/装备领跑。"),
    ("3", "电子/医疗+16.9%/+14.2%", "电子信息/生物医药。"),
    ("4", "新能源汽车零售+27%、产量+78.1%", "新能源车爆发。"),
    ("5", "进出口18826.99亿、-5.4%（出口-6.7%）", "外贸转弱、需稳。"),
    ("6", "机电/高新技术产品出口微增", "出口结构向高技术。"),
    ("7", "固定资产投资-3.3%、工业+6.1%", "工业投资强。"),
    ("8", "房地产开发投资-30.0%", "地产深度调整。"),
    ("9", "社零25433.59亿、+4.4%、线上+14.2%", "内需/线上强。"),
    ("10", "水产品产量965.16万吨、+4.4%", "海洋渔业强。"),
    ("11", "沿海港口吞吐量7.56亿吨、集装箱1843万箱", "海洋枢纽。"),
    ("12", "常住4190万、城镇化率72.58%、自然增-1.52‰", "沿海人口/城镇化。"),
    ("13", "居民人均可支配收入50302元、+5.1%", "收入稳、城乡比2.04。"),
    ("14", "新登记主体105.61万家、民营经济活跃", "民营强。"),
    ("15", "CPI 0.0%、PPI-1.2%", "物价平、通缩缓。"),
]
for d in details:
    para(doc, "**%s**" % d[0], size=11)
    para(doc, d[1], size=10.5, color=GRAY)

# ---- 五、六个换挡 ----
heading1(doc, "五、把所有细节连起来：福建正在经历的“六个换挡”")
gear = [
    ("1．动能换挡：从“轻工/食品/传统外贸”到“高技术+装备+海洋”", "高技术+14.6%、装备+13.9%。"),
    ("2．产业换挡：从“鞋服/食品”到“电子信息+新能源车+海洋”", "电子+16.9%、新能源车+78.1%。"),
    ("3．投资换挡：从“基建/地产”到“工业/高技术/基建”", "工业+6.1%、地产-30%、基建+5.6%。"),
    ("4．开放换挡：从“传统外贸”到“对台/海丝+高技术出口”", "机电/高新微增、民营/海丝。"),
    ("5．人口换挡：从“劳务/侨乡”到“城镇化/新城集聚”", "常住4190万、城镇化72.58%。"),
    ("6．动能换挡：从“传统增长”到“民营+海洋+新质”", "民营主体770万、海洋强省。"),
]
for g in gear:
    para(doc, "**%s**　%s" % g, space_after=6)

# ---- 六、增长暗线 ----
heading1(doc, "六、增长暗线：民营经济、海洋经济与制造业升级（新质生产力）")
heading2(doc, "6.1 民营+高技术制造")
para(doc, "福建民营经济活跃，高技术制造+14.6%、装备+13.9%、电子+16.9%、新能源车+78.1%。民企制造+高技术，是福建“新质生产力”的内生引擎。")
heading2(doc, "6.2 海洋经济")
para(doc, "福建是海洋渔业/港口强省，水产品965万吨、港口7.56亿吨、集装箱1843万箱。海洋经济（渔业、海工、港口、装备）是福建“蓝色经济”增量。")
para(doc, "**这条暗线意味着**：福建的“民营+海洋+高技术制造”正接棒“轻工/传统外贸”。看福建，盯住“民企占比、高技术/装备、海洋经济、进出口与新质”这几组数。")

# ---- 七、财政暗线 ----
heading1(doc, "七、财政暗线：收入稳、税收尚可，民生/基建投入大")
para(doc, "2025年福建地方一般公共预算收入3723.35亿、+3.0%。收入稳，来自民营/制造/消费税源；支出向民生、基建、海洋倾斜。")
para(doc, "**制度含义**　福建财政“收入稳、结构民营为主”，关键是保持“制造业+海洋+民营”税源，平衡外贸波动对收入的传导。")

# ---- 八、产业暗线 ----
heading1(doc, "八、产业暗线：从“轻工/食品”到“高技术+装备+海洋”")
heading2(doc, "8.1 福建产业的“表”")
table(doc,
    ["指标", "2025增速/信号", "解读"],
    [
        ["规上工业", "+6.5%", "工业稳"],
        ["高技术制造", "+14.6%", "新质生产力"],
        ["装备制造", "+13.9%", "高端装备"],
        ["电子/医疗", "+16.9%", "电子信息/医药"],
        ["新能源车产量/零售", "+78.1%/+27%", "新能源车"],
        ["海洋/水产品", "965万吨/港口7.56亿吨", "海洋强省"],
        ["出口/机电", "机电+2.2%", "出口结构"],
    ],
    widths=[4.6, 3.8, 5.0])
heading2(doc, "8.2 从“轻工/食品”到“高技术/海洋”")
para(doc, "福建过去以鞋服、食品、轻工出口见长，2025年“高技术+装备+新能源车+海洋”成为新增长。民营+海洋+新质，是福建“蓝色/绿色”升级。")

# ---- 九、区域 ----
heading1(doc, "九、区域格局：福州都市圈、厦漳泉与海峡西岸")
para(doc, "福建“福州都市圈（强省会）+厦漳泉（闽南/制造）+海峡西岸（对台）”。福州（省会/产业）、厦门（特区/港口/金融）、泉州（制造/民企）。")
para(doc, "在“海峡西岸经济区”与“海丝/自贸”下，福建向台开放、连接海外闽商。双核+沿海开放，是福建增长布局。")

# ---- 十、人口与城市 ----
heading1(doc, "十、人口与城市：沿海人口、城镇化与代际")
heading2(doc, "10.1 人口总量")
para(doc, "2025年末福建常住4190万、城镇化率72.58%、自然增约-1.52‰。沿海人口集中、城镇化较高。")
heading2(doc, "10.2 城市/代际")
para(doc, "福州、厦门、泉州为核心城市。人口代际/侨乡、海外闽商，构成福建开放与消费的基础。")

# ---- 十一、民营经济 ----
heading1(doc, "十一、民营经济：占八成经营主体、活跃开放")
heading2(doc, "11.1 民企地位")
para(doc, "福建经营主体770.71万户、其中民营经济高度发达，主导鞋服、运动装备、新能源、海洋等。民营是福建最强底盘。")
heading2(doc, "11.2 政策/侨商")
para(doc, "福建支持民营、民营企业和侨商开放。民企、闽商、海丝，是福建未来10年底座。")

# ---- 十二、事后验证 ----
heading1(doc, "十二、事后验证：用2025年实际执行结果检验报告判断")
heading2(doc, "12.1 年初目标与年末执行对比")
table(doc,
    ["指标", "2025年初目标", "2025实际", "偏差"],
    [
        ["GDP增速", "5%左右", "+5.0%", "达标"],
        ["规上工业", "力争", "+6.5%", "超预期"],
        ["进出口", "稳", "-5.4%", "偏弱"],
        ["固定资产投资", "未设", "-3.3%", "偏弱"],
        ["一般公共预算收入", "未设", "+3.0%", "稳"],
        ["居民收入", "与增长同步", "+5.1%", "同步"],
    ],
    widths=[3.0, 3.0, 3.0, 4.4])
heading2(doc, "12.2 判断验证")
para(doc, "三条核心判断得到支持：**（1）工业/高技术/民营强（+6.5%/+14.6%），验证；（2）消费/内需强（社零+4.4%），验证；（3）外贸/地产弱（-5.4%/-30%），验证。**")
para(doc, "核心观察：福建靠“高技术+装备+民营+海洋+消费”守住5.0%增长，但外贸、地产偏弱。2026年“稳民营/制造+扩内需+修外贸/地产”是重点。")

# ---- 十三、五条主线 ----
heading1(doc, "十三、未来5—10年最值得观察的五条主线")
lines5 = [
    ("① 高技术/装备/新能源升级", "高技术、装备、新能源车能否持续壮大。"),
    ("② 海洋强国/海洋经济", "海洋渔业、港口、海工能否成为蓝色增长极。"),
    ("③ 民营经济与海丝/对台", "民营、对台、海丝开放能否持续。"),
    ("④ 外贸/地产再平衡", "-5.4%/-30%，能否用海洋/内需补。"),
    ("⑤ 人口/代际/城乡", "沿海集聚、人口、侨乡、城乡协调。"),
]
for l in lines5:
    para(doc, "**%s**　%s" % l, space_after=6)

# ---- 十四、结论 ----
heading1(doc, "十四、最终结论：福建在“海洋强省+民营+海峡西岸”里的增长逻辑")
para(doc, "福建的2025年，本质上是“**制造/高技术/民营/海洋为核心，而外贸/地产偏弱**”的答卷：GDP6万亿、高技术+14.6%、海洋强，但进出口-5.4%、地产-30%。")
para(doc, "只要民营、海洋、高技术制造、开放能接住，福建就站在“海峡西岸+海洋强省”的增长位；如果外贸/地产持续偏弱、内需不足，福建需承受“外向与内需”的结构挑战。")
para(doc, "最稳妥的观察信号：**一盯高技术/装备/新能源（动能）、二盯海洋/民营（蓝色/底座）、三盯外贸/对台开放（开放）、四盯地产/投资（约束）、五盯消费/居民（内需）。**福建，是“海洋+民营+海峡”的新样本。")

# ---- 附录A ----
heading1(doc, "附录A：主要资料来源与核验路径")
bullet(doc, "福建省2025年《政府工作报告》（2025年1月）——目标来源。")
bullet(doc, "《2025年福建省国民经济和社会发展统计公报》（福建省统计局，2026-03-19）——GDP、工业、外贸、人口实值。")
bullet(doc, "福建省统计/海洋/民营专题、自贸港——海洋与民营。")
bullet(doc, "2026年福建省政府工作报告——2025执行复盘。")
bullet(doc, "福州/厦门海关、省财政厅——外贸/财政。")
heading2(doc, "核验说明")
para(doc, "本报告所有增速、总量、占比均以统计公报/官方发布为基准。涉“海洋产业增加值”“民营经济占比”等以官方口径为准。")

# ---- 附录B ----
heading1(doc, "附录B：建议建立的年度跟踪仪表盘")
para(doc, "建议把下面10个“测脉搏”指标做成年度跟踪表：")
dash = [
    ["#", "指标", "2025基值", "为什么重要"],
    ["1", "GDP增速", "+5.0%", "总量与方向"],
    ["2", "规上工业增速", "+6.5%", "制造底盘"],
    ["3", "高技术制造增速", "+14.6%", "新质生产力"],
    ["4", "水产品产量/港口", "965万吨/7.56亿吨", "海洋强省"],
    ["5", "进出口/出口增速", "-5.4%/-6.7%", "外贸韧性"],
    ["6", "固定资产投资/地产", "-3.3%/-30.0%", "投资结构"],
    ["7", "社零增速", "+4.4%", "内需消费"],
    ["8", "常住人口/城镇化率", "4190万/72.58%", "人口与城市"],
    ["9", "一般公共预算收入增速", "+3.0%", "财政质量"],
    ["10", "居民人均可支配收入增速", "+5.1%", "民生获得感"],
]
table(doc, dash[0], dash[1:], widths=[0.8, 4.5, 3.6, 4.4])
para(doc, "把这10个指标连起来看，任何一个“新引擎（2/3/4）向上、旧引擎（5/6）修复”，都说明福建在真正换挡。")

# ========================================= 保存
out = "/Users/x/Desktop/content-prod-lab/reports/福建省_2025年政府工作报告_深度研究_2026-08-13.docx"
doc.save(out)
print("SAVED:", out)
print("PARAGRAPHS:", len(doc.paragraphs))
print("TABLES:", len(doc.tables))
