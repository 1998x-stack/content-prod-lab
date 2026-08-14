# -*- coding: utf-8 -*-
"""Build 包头市2025年政府工作报告 深度研究 DOCX, 参照地级市系列版式。"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DARK = RGBColor(0x1F, 0x2A, 0x44)
RED = RGBColor(0xB0, 0x1E, 0x2E)
GRAY = RGBColor(0x59, 0x60, 0x69)
BLUE = RGBColor(0x1F, 0x4E, 0x79)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
HEADER_FILL = "1F2A44"
K = "微软雅黑"


def set_run(run, size=11, bold=False, color=DARK, font=K):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    run.font.name = font
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), font)
    rFonts.set(qn("w:hAnsi"), font)
    rFonts.set(qn("w:eastAsia"), font)


def para(doc, text, size=11, bold=False, color=DARK, align=None,
         space_after=6, space_before=0, font=K):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        set_run(r, size=size, bold=bold or (i % 2 == 1), color=color, font=font)
    return p


def heading1(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(16)
    pf.space_after = Pt(8)
    r = p.add_run(text)
    set_run(r, size=16, bold=True, color=RED)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "B01E2E")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def heading2(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(10)
    pf.space_after = Pt(4)
    r = p.add_run(text)
    set_run(r, size=13, bold=True, color=BLUE)
    return p


def table(doc, headers, rows, widths=None, font_size=9.5):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_run(r, size=font_size, bold=True, color=WHITE)
        tcPr = hdr[i]._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), HEADER_FILL)
        tcPr.append(shd)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(val))
            set_run(r, size=font_size, color=DARK)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    return t


def quote_box(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(8)
    pf.left_indent = Pt(12)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), "B01E2E")
    pBdr.append(left)
    pPr.append(pBdr)
    r = p.add_run(text)
    set_run(r, size=11, color=DARK, font="楷体")
    return p


def bullet(doc, text, size=10.5):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.7)
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        set_run(r, size=size, bold=(i % 2 == 1), color=DARK)
    return p


# ================================================================= 文档主体
doc = Document()
sec = doc.sections[0]
sec.page_width = Cm(21)
sec.page_height = Cm(29.7)
sec.top_margin = Cm(2.4)
sec.bottom_margin = Cm(2.2)
sec.left_margin = Cm(2.5)
sec.right_margin = Cm(2.5)

st = doc.styles["Normal"]
st.font.name = K
st.font.size = Pt(11)
st.element.rPr.rFonts.set(qn("w:eastAsia"), K)


# ---- 封面 ----
para(doc, "包头市2025年政府工作报告", size=24, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=60, space_after=4)
para(doc, "深度研究与\u201c容易被忽视的细节\u201d分析", size=16, bold=True, color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
para(doc, "从\u201c稀土之都、钢铁、晶硅光伏、铝业与新能源装备\u201d重新理解包头", size=12, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)
para(doc, "\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501", color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
para(doc, "研究对象：2025年包头市政府工作报告及同期财政、国民经济和社会发展统计资料、2026年官方复盘", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
para(doc, "标定日期：2026-08-13", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
doc.add_page_break()

# ---- 目录 ----
catalog = [
    "执行摘要",
    "一、研究口径与阅读方法",
    "二、先看包头的特殊底盘：稀土之都、钢铝、晶硅光伏、风电装备与工业转型",
    "三、最关键的宏观错位：GDP低于8%目标、工业强，但外贸/财政/地产大幅下行",
    "四、容易被忽视、但信号很重的细节（15条）",
    "五、2025年GDP目标 vs 实际：对照表",
    "六、2025年增长是谁撑起来的？（结构归因）",
    "七、预算与财政的\u201c含金量\u201d",
    "八、民生底账：人口、收入与城乡",
    "九、城镇与农村：格局与均衡",
    "十、人口流入与流出",
    "十一、物价与货币环境",
    "十二、区域一体化：包头在\u201c呼包鄂榆城市群+京津冀+口岸\u201d里的位置",
    "十三、未来5—10年最值得观察的五条主线",
    "十四、最终结论：包头在\u201c稀土+硅+风电装备+工业升级\u201d里的增长逻辑",
    "附录A：主要资料来源与核验路径",
    "附录B：建议建立的年度跟踪仪表盘",
]
para(doc, "目　录", size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=10, space_after=10)
for item in catalog:
    para(doc, item, size=12, space_after=4)
doc.add_page_break()

# ---- 执行摘要 ----
heading1(doc, "执行摘要")
para(doc, "**核心判断**　2025年包头最显著的是\u201cGDP 4711.9亿元、增长5.4%（低于8%目标）、三产占50.8%\u201d、\u201c规上工业+11.1%（稀土+27.4%、装备+20.5%）\u201d、\u201c进出口250.5亿/-22.8%、财政-11.9%\u201d、\u201c常住人口277.40万/城镇化88.82%\u201d。这说明包头在\u201c稀土+硅+风电\u201d的工业升级中增长，但**出口、财政、地产大幅下行**是短板。")
para(doc, "把2025年目标（GDP+8%/规上+14%/固投+18%/财政+5%）、2025年统计（GDP+5.4%远低于目标、规上+11.1%）、趋势一起看，包头是\u201c草原工业城市、稀土+硅业\u201d路径：**稀土产业、晶硅光伏、风电装备、铝、钢**是支柱；但2025年增速未达雄心目标。总量4711.9亿居内蒙第2。")
para(doc, "最容易记住的一句话：**包头是\u201c稀土之都+晶硅光伏+钢铁铝\u201d的内蒙古工业重镇，靠\u201c稀土、硅料（多晶硅）、风电、铝、钢铁\u201d增长。**观察包头，与其只看\u201cGDP 4711亿\u201d，不如看\u201c规上工业+11.1%、稀土+27.4%、风电装机+43.3%、晶硅101.8万吨+35%、固投+6.2%\u201d。")# ---- 一、研究口径与阅读方法 ----
heading1(doc, "一、研究口径与阅读方法")
para(doc, "本报告以\u201c包头市政府工作报告（2025年，孟庆维作）\u201d为起点，把\u201c2025年GDP目标（8%左右）\u201d与\u201c官方2025年GDP（4711.9亿元/+5.4%）\u201d并置对照，并用经济运行分析复盘。")
para(doc, "注：本报告以\u201c全市\u201d口径为主；公报/运行口径（GDP与统计局核算）稍异，均以包头市统计局为准。")
para(doc, "统一采用\u201c同比增长%\u201d（GDP按不变价），一般公共预算收入为地方级。")

# ---- 二、研究背景 ----
heading1(doc, "二、先看包头的特殊底盘：稀土之都、钢铝、晶硅光伏、风电装备与工业转型")
para(doc, "**区位与身份**：包头是内蒙古最大工业城市、呼包鄂榆经济带核心、\u201c草原钢城\u201d\u201c稀土之都\u201d，面向中蒙/口岸；全国百强城市。")
para(doc, "**产业底盘**：一是稀土（全球最重要稀土基地，稀土产业产值+27.4%、北方稀土）；二是钢铝（钢铁/铝业产值千亿）；三是晶硅光伏（多晶硅101.8万吨、单晶59.3万吨、光伏+20.3%）；四是风电装备（风电装机+43.3%、轴承/电机）。")
para(doc, "**人口底盘**：2025年末常住277.40万/城镇化88.82%，内蒙人口枢纽之一。")
para(doc, "**市场与贸易**：2025年社零939.5/+5.4%；进出口250.5亿/-22.8%（出口-12.7%、进口-32%）。")

# ---- 三、核心宏观错位 ----
heading1(doc, "三、最关键的宏观错位：GDP低于8%目标、工业强，但外贸/财政/地产大幅下行")
para(doc, "**第一组错位**：2025年GDP目标8%左右/5000亿，实际4711.9亿/+5.4%**远低于目标**（较2024+7.9%回落）；工业+11.1但总量缺口大。")
para(doc, "**第二组错位**：规上工业+11.1%、稀土+27.4%、装备+20.5%，工业强；但**进出口-22.8%、财政-11.9%、地产-54.1%**大幅下行。\u201c工业增、贸易/财政落\u201d。")
para(doc, "**第三组错位**：常住277.4万/城镇化88.8%、人口微调（出生少）；但收入+3.6%、CPI-0.1%。")
para(doc, "一句话：**包头是\u201c稀土/硅/风电装备强、但出口/财政/地产弱\u201d的工业城市**——2025年工业高增但未达GDP雄心目标。")

# ---- 四、15条细节 ----
heading1(doc, "四、容易被忽视、但信号很重的细节（15条）")
bullet(doc, "1. **GDP 4711.9亿/+5.4%**：低于8%目标（2024+7.9%），内蒙第2。")
bullet(doc, "2. **规上工业+11.1%**：高于全国5.2pct、高于全区4.4pct。")
bullet(doc, "3. **稀土产业+27.4%**：全球稀土基地、永磁电机/氧化钇。")
bullet(doc, "4. **装备制造业+20.5%**：新能源/风电装备。")
bullet(doc, "5. **晶硅光伏+7.7%、多晶硅+35%**：硅产业链。")
bullet(doc, "6. **风电装机+43.3%、储能+4倍**：绿色能源。")
bullet(doc, "7. **钢铁-2.4%**：唯一负增的五大产业。")
bullet(doc, "8. **固投+6.2%（工业+12.2%占79%）**：工业投强。")
bullet(doc, "9. **房地产-54.1%（销售-33.6%）**：地产腰斩。")
bullet(doc, "10. **社零+5.4%（文化办公+3.9倍）**：消费支出回稳。")
bullet(doc, "11. **进出口-22.8%、出口-12.7%**：外贸承压。")
bullet(doc, "12. **财政183.7亿/-11.9%、税收-8.8%**：财政下行。")
bullet(doc, "13. **常住277.40万/城镇化88.82%**：人口稳。")
bullet(doc, "14. **收入城镇63083元/+3.2%、农村28609元/+5.1%**：城乡约2.2。")
bullet(doc, "15. **风电装机+43.3%、电网储能+4倍**：绿色能源。")
# ---- 五、目标 vs 实际对照表 ----
heading1(doc, "五、2025年GDP目标 vs 实际：对照表")
table(doc,
    ["指标", "2025年目标", "2025年实际", "是否达标"],
    [
        ["地区生产总值增速", "8%左右/5000亿", "4711.9亿元/+5.4%", "远低于目标"],
        ["规上工业增加值", "+14%以上", "+11.1%", "未达标(接近)"],
        ["固定资产投资", "+18%以上", "+6.2%", "未达标"],
        ["社会消费品零售总额", "+5%以上", "939.5亿元/+5.4%", "达标"],
        ["一般公共预算收入", "+5%左右", "183.7亿元/-11.9%", "大幅未达标"],
        ["进出口", "（稳）", "250.5亿元/-22.8%", "大幅下行"],
        ["居民收入", "与GDP同步", "全体+3.6%", "低于GDP"],
    ],
    widths=[3.4, 2.5, 5.0, 3.1])
para(doc, "**简评**：包头2025年**仅社零达标**；GDP、规上、投资、财政、进出口均低于目标。**稀土/风电装备/工业高增**是亮点，**出口、财政、贸易**大幅下行。")

# ---- 六、结构归因 ----
heading1(doc, "六、2025年增长是谁撑起来的？（结构归因）")
para(doc, "**三大产业**：一产150.1亿/+5.0%、二产2166.1亿/+6.5%、三产2395.7亿/+4.5%；三产占50.8%。增长以二产（工业，拉动2.8pct）+三产（营利性服务）支撑。")
para(doc, "**工业**：规上工业+11.1%；**稀土+27.4%、装备+20.5%、化工+28.2%、煤炭+47.6%、电子+17%、有色+7.7%**；钢铁-2.4%、晶硅+7.7%。多晶硅101.8万吨+35%。规上利润207.3亿。")
para(doc, "**服务业**：批发零售/物流/营利性服务较快；金融-11.9%、房地产-4.1%拖累。")
para(doc, "**增长归因**：包头GDP主要靠**稀土+硅业+风电装备+有色+工业**；固定资产投资、出口、财政为负，金融/地产下拖。")

# ---- 七、财政 ----
heading1(doc, "七、预算与财政的\u201c含金量\u201d")
para(doc, "2025年一般公共预算收入183.7亿元/-11.9%（税收133.2亿/-8.8%、非税-19.3%）；支出431.8亿/-7.2%。")
para(doc, "**结构性**：收入\u201c大幅下行（-11.9%）\u201d，税收、非税双降（工业利润/出口走弱）；支出-7.2%收缩化债。")
para(doc, "**含金量**：包头财政\u201c税基受工业/贸易/地产冲击\u201d，2025年是财政压力年；依赖稀土/硅业波动与转移支付。")

# ---------------- 八、民生 ----------------
heading1(doc, "八、民生底账：人口、收入与城乡")
para(doc, "常住人口277.40万/城镇化88.82%（+0.3pct）；出生4.29\u2030/死8.29\u2030（自然负）。收入：**全体59067元/+3.6%、城镇63083元/+3.2%、农村28609元/+5.1%**，城乡比约2.2。")
para(doc, "就业：城镇新增40770人。")
para(doc, "**民生结论**：收入农村快于城镇、城乡差距约2.2；人口稳、城镇化高；就业、收入温和、普惠扎实。")

# ---------------- 九 ----------------
heading1(doc, "九、城镇与农村：格局与均衡")
para(doc, "包头城镇化率88.82%（很高），中心城区（昆都仑/青山/稀土高新）承载钢铝/稀土/硅；旗县承接工业与草原牧业。")
para(doc, "城乡收入比约2.2（内蒙古中等），农村增速快、在收敛；稀土配套与县域协同。")

# ---------------- 十 ----------------
heading1(doc, "十、人口流入与流出")
para(doc, "包头常住277.4万、自然负（出生少），城镇化高；工业就业吸引部分机械流入。")
para(doc, "**流入**：稀土/硅/新能源岗位；**流出**：青壮年赴呼包/一线。总体\u201c微净、城镇化饱和度高\u201d。")

# ---------------- 十一 ----------------
heading1(doc, "十一、物价与货币环境")
para(doc, "2025年包头CPI**下降0.1%**（通缩），食品-1.1%、交通-2.9%；其他用品+10.5%。")
para(doc, "金融：存款4916.6亿/+1.0%、贷款3404.3亿/+0.2%。\u201c宽中低增\u201d。")

# ---------------- 十二 ----------------
heading1(doc, "十二、区域一体化：包头在\u201c呼包鄂榆城市群+京津冀+中蒙口岸\u201d里的位置")
para(doc, "包头是呼包鄂榆城市群核心（与呼和浩特/鄂尔多斯/榆林）、内蒙古西部工业枢纽，面向中蒙/欧亚（满都拉口岸过货1000万吨）。")
para(doc, "产业协同：稀土（全球）、钢铝、硅-风电，京津冀产业转移承接；呼和浩特都市圈。包头是\u201c草原工业+开放口岸\u201d枢纽。")

# ---------------- 十三 ----------------
heading1(doc, "十三、未来5—10年最值得观察的五条主线")
bullet(doc, "1. **稀土产业链**：北方稀土、磁材、永磁电机，全球稀土话语权。")
bullet(doc, "2. **晶硅光伏+风电装备**：硅（多晶硅/单晶）+风电，绿色能源装备。")
bullet(doc, "3. **钢铁/铝业转型**：钢铁-2.4%后能不能向特钢/铝轻量化升级。")
bullet(doc, "4. **固投/房地产开发修复**：地产-54%、固投+6%后的新项目。")
bullet(doc, "5. **外贸/财政修复**：进出口-22.8%、财政-11.9%后的口岸/税收。")

# ---------------- 十四 ----------------
heading1(doc, "十四、最终结论：包头在\u201c稀土+硅+风电装备+工业升级\u201d里的增长逻辑")
para(doc, "**结论**：包头2025年的\u201c真相\u201d是——**GDP+5.4%（远低8%目标）、规上+11.1%、稀土+27.4%、固投+6.2%、进出口-22.8%、财政-11.9%**。它是\u201c稀土+硅+风电\u201d强、但**贸易/财政/投资/地产弱**的草原工业城市。")
para(doc, "**对趋势判断**：稀土/硅/风电装备代表包头的\u201c核心竞争力\u201d，出口/财政/投资代表\u201c近期压力\u201d。**工业+绿色装备**决定中期潜力，**外贸/财政修复+地产企稳**决定近期韧性。")
para(doc, "**若只看一个指标**：看**稀土产业增速（+27.4%）与出口增速（-12.7%）**——包头靠\u201c稀土+晶硅+风电\u201d的制造业升级很强，但金属/硅料出口与财政的波动，决定其增长的成色与对外部价格（稀土/硅价）的高度敏感。")

# ---------------- 附录A ----------------
heading1(doc, "附录A：主要资料来源与核验路径")
bullet(doc, "包头市人民政府《2025年包头市政府工作报告》（2025年，孟庆维）。")
bullet(doc, "包头市统计局《2025年包头市国民经济和社会发展统计公报》(2026年)。")
bullet(doc, "《2025年包头经济运行情况分析》（包头统计局）。")
bullet(doc, "内蒙古2025年统计公报、满都拉口岸/稀土指数交叉核验。")

# ---------------- 附录B ----------------
heading1(doc, "附录B：建议建立的年度跟踪仪表盘")
bullet(doc, "GDP总量/增速 vs 全年目标。")
bullet(doc, "规上工业增加值及分行业（稀土/晶硅/钢/铝/风电/装备）增速。")
bullet(doc, "固定资产投资（总量/工业/技改/房地产）增速。")
bullet(doc, "社会消费品零售总额、限上/网络。")
bullet(doc, "进出口（人民币）、出口、稀土/硅出口价。")
bullet(doc, "一般公共预算收入/税收/非税、财政支出。")
bullet(doc, "常住人口、城镇化率、自然增长率。")
bullet(doc, "CPI/核心CPI、规上工业利润。")
bullet(doc, "稀土产量/产值/价格、多晶硅/单晶硅。")
bullet(doc, "风电装机、储能、满都拉口岸过货。")

# ---------------- 保存 ----------------
out = "/Users/x/Desktop/content-prod-lab/reports/包头市_2025年政府工作报告_深度研究_2026-08-13.docx"
doc.save(out)
print("SAVED OK 包头市_2025年政府工作报告_深度研究_2026-08-13.docx")