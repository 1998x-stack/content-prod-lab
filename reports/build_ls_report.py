# -*- coding: utf-8 -*-
"""Build 济宁市2025年政府工作报告 深度研究 DOCX, 参照地级市系列版式。"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DARK = RGBColor(0x1F, 0x2A, 0x44)
RED = RGBColor(0xB0, 0x1E, 0x2E)
GRAY = RGBColor(0x59, 0x60, 0x69)
BLUE = RGBColor(0x1F, 0x4E, 0x79)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
HEADER_FILL = "1F2A44"
K = "微软雅黑"


def set_run(run, size=11, bold=False, color=DARK, font=K):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    run.font.name = font
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), font)
    rFonts.set(qn("w:hAnsi"), font)
    rFonts.set(qn("w:eastAsia"), font)


def para(doc, text, size=11, bold=False, color=DARK, align=None,
         space_after=6, space_before=0, font=K):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        set_run(r, size=size, bold=bold or (i % 2 == 1), color=color, font=font)
    return p


def heading1(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(16)
    pf.space_after = Pt(8)
    r = p.add_run(text)
    set_run(r, size=16, bold=True, color=RED)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "B01E2E")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def heading2(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(10)
    pf.space_after = Pt(4)
    r = p.add_run(text)
    set_run(r, size=13, bold=True, color=BLUE)
    return p


def table(doc, headers, rows, widths=None, font_size=9.5):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_run(r, size=font_size, bold=True, color=WHITE)
        tcPr = hdr[i]._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), HEADER_FILL)
        tcPr.append(shd)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(val))
            set_run(r, size=font_size, color=DARK)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    return t


def quote_box(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(8)
    pf.left_indent = Pt(12)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), "B01E2E")
    pBdr.append(left)
    pPr.append(pBdr)
    r = p.add_run(text)
    set_run(r, size=11, color=DARK, font="楷体")
    return p


def bullet(doc, text, size=10.5):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.7)
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        set_run(r, size=size, bold=(i % 2 == 1), color=DARK)
    return p


# ================================================================= 文档主体
doc = Document()
sec = doc.sections[0]
sec.page_width = Cm(21)
sec.page_height = Cm(29.7)
sec.top_margin = Cm(2.4)
sec.bottom_margin = Cm(2.2)
sec.left_margin = Cm(2.5)
sec.right_margin = Cm(2.5)

st = doc.styles["Normal"]
st.font.name = K
st.font.size = Pt(11)
st.element.rPr.rFonts.set(qn("w:eastAsia"), K)


# ---- 封面 ----
para(doc, "丽水市2025年政府工作报告", size=24, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=60, space_after=4)
para(doc, "深度研究与\u201c容易被忽视的细节\u201d分析", size=16, bold=True, color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
para(doc, "从\u201c中国生态第一市、青瓷宝剑、制药、五金汽配、健康医药、黄帝缙云\u201d重新理解丽水", size=12, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)
para(doc, "\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501", color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
para(doc, "研究对象：2025年丽水市政府工作报告及同期财政、国民经济和社会发展统计资料、2026年官方复盘", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
para(doc, "标定日期：2026-08-14", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
doc.add_page_break()

# ---- 目录 ----
catalog = [
    "执行摘要",
    "一、研究口径与阅读方法",
    "二、先看丽水的特别底盘：生态、青瓷宝剑、五金、健康、茶、山城",
    "三、最关键的宏观错位：规上+9.1%、GDP+6.4%、但固投-1.8%、社零+4.4%、对美出口-23.1%",
    "四、容易被忽视、但信号很重的细节（15条）",
    "五、2025年GDP目标 vs 实际：对照表",
    "六、2025年增长是谁撑起来的？（结构归因）",
    "七、预算与财政的\u201c含金量\u201d",
    "八、民生底账：人口、收入与城乡",
    "九、城镇与农村：格局与均衡",
    "十、人口流入与流出",
    "十一、物价与货币环境",
    "十二、区域一体化：丽水在浙江西南、金衢、杭州都市圈、生态屏障里的位置",
    "十三、未来5\u201310年最值得观察的五条主线",
    "十四、最终结论：丽水在\u201c生态+制造+侨乡\u201d里的增长逻辑",
    "附录A：主要资料来源与核验路径",
    "附录B：建议建立的年度跟踪仪表盘",
]
para(doc, "目　录", size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=10, space_after=10)
for item in catalog:
    para(doc, item, size=12, space_after=4)
doc.add_page_break()

# ---- 执行摘要 ----
heading1(doc, "执行摘要")
para(doc, "**核心判断**　2025年丽水最显著的是\u201cGDP 2301.4亿/+6.4%（达成6.5%目标附近）、规上工业+9.1%（达成9%目标）、三产+7.0%、高技术制造+23.6%、数字经济核心制造+21.2%\u201d、\u201c但固投-1.8%、房地产-21.8%、社零+4.4%（差6.5%目标）、对美出口-23.1%、财政+2.5%\u201d。这说明丽水在\u201c浙江生态绿谷+制造业（青瓷宝剑五金）+健康医药\u201d中，**工业制造与三产强、投资与内需/外贸（对美）偏弱**。")
para(doc, "把丽水2025年目标（GDP+6.5%、规上+9%、固投+7%以上、社零+6.5%、财收+3.8%、出口+6.5%）与实际（GDP+6.4%、规上+9.1%、固投-1.8%、社零+4.4%、财收+2.5%、出口+7.2%）对照，丽水是\u201c生态+制造\u201d路径：**青瓷宝剑、五金汽配、合成革、健康医药、绿色能源、生态旅游**是支柱。")
para(doc, "最容易记住的一句话：**丽水是\u201c中国生态第一市、青瓷之都、宝剑之乡、黄帝缙云\u201d，靠\u201c生态+制造+健康\u201d驱动。**观察丽水，与其只看\u201cGDP 2301亿/+6.4%\u201d，不如看\u201c规上工业+9.1%、高技术制造+23.6%、青瓷宝剑、生态屏障、侨乡消费\u201d。")

# ---- 一、研究口径与阅读方法 ----
heading1(doc, "一、研究口径与阅读方法")
para(doc, "**本文的最终目标**，不是复述丽水市2025年政府工作报告的原文，而是把它放到\u201c目标→实际\u201d和\u201c长期底盘→年度对撞\u201d的坐标里重新读一遍，找出丽水2025年到底靠什么增长、哪里在换挡、哪里是暗线。")
para(doc, "**三组资料互证**：（1）《2025年丽水市国民经济和社会发展统计公报》（丽水市统计局，2026-05发布）提供全部2025年**实际完成**数据；（2）《丽水市2026年政府工作报告》（2026-01市人大常委会）既复盘2025年、又给出2025-2026主线；（3）《2025年丽水市国民经济和社会发展计划》提供2025年**目标设定**（GDP+6.5%、规上+9%、固投+7%、社零+6.5%、财收+3.8%）。三份材料交叉核验，避免单源失真。")
para(doc, "**核心框架**：先交代丽水的\u201c底盘\u201d（生态+青瓷宝剑+五金+健康+茶的山城禀赋），再用\u201c目标vs实际\u201d对照表定位最关键的\u201c宏观错位\u201d，接着用15条细节捕捉暗线，最后从财政、民生、城乡、人口、物价五个侧面把错位坐实，落到2026年及\u201c十五五\u201d主线与结论。")
para(doc, "**口径说明**：GDP 2301.4亿+6.4%、人均GDP 90769元（+6.1%）。三次产业结构5.6:35.4:59.0。常住人口253.9万人、城镇化率67.0%（+1.0pct）。固定资产投资-1.8%（扣房地产+6.4%）、房地产投资260.5亿-21.8%。一般公共预算收入198.0亿+2.5%（税收148.3亿+3.6%、占比74.9%）。进出口449.6亿+6.1%（出口+7.2%）。")
para(doc, "**为什么值得单独研究丽水**：GDP+6.4%、规上+9.1%，是浙江山区/生态市里增速靠前；但固投仅-1.8%、社零+4.4%、对美出口-23.1%。\u201c工业制造强、内需投资外贸分化\u201d的组合，是观察\u201c生态屏障城市在高质量发展与制造业升级\u201d的上佳切片。")
# ---- 二、先看丽水的特别底盘 ----
heading1(doc, "二、先看丽水的特别底盘：生态、青瓷宝剑、五金、健康、茶、山城")
para(doc, "丽水的成长逻辑，不能只看GDP数字，而要先把\u201c底盘\u201d摊开。底盘决定了几十年、上百年的禀赋，也决定了2025年增长来源和约束。")
bullet(doc, "**中国生态第一市·生态屏障**：丽水是浙江西南山区、全国生态文明示范，被誉为\u201c中国生态第一市\u201d，GEP/生态屏障是国家生态资产。生态+文旅=城市名片。")
bullet(doc, "**青瓷之都·宝剑之乡**：龙泉青瓷（世界非遗）、龙泉宝剑、庆元香菇、云和木制玩具，是丽水的文化+轻工业特色。")
bullet(doc, "**五金汽配/合成革制造**：缙云/仙都一带的机械五金、汽车零部件、合成革等制造业，是丽水工业与出口的主体。")
bullet(doc, "**健康医药/绿色能源**：畲族医药、健康生物医药、绿色水电/风电，是丽水产业转型的新方向。")
bullet(doc, "**黄帝缙云·生态旅游**：缙云仙都（黄帝祭典）、古堰画乡、云和梯田，丽水是浙江重要的生态文旅目的地。")
heading2(doc, "底盘的产业钟摆")
para(doc, "把底盘归纳成一句话：**丽水是\u201c生态+制造\u201d的山区山海城市**。青瓷宝剑五金定制造、生态绿谷供永续、健康医药添新、黄帝缙云供旅游。城市资产在\u201c生态文明+特色制造\u201d链条上。工业强（规上+9.1%、高技术+23.6%）是2025亮点，但投资与内需（固投-1.8%、社零+4.4%）偏弱。")

# ---- 三、最关键的宏观错位 ----
heading1(doc, "三、最关键的宏观错位：规上+9.1%、GDP+6.4%达成、但固投-1.8%、社零+4.4%、对美出口-23.1%")
para(doc, "把2025年数据摊开，丽水最醒目的不是\u201c+6.4%\u201d，而是**工业强、投资与内需外贸偏弱**的三重错位。")
quote_box(doc, "宏观错位一号：**工业制造强 vs 投资偏弱**。规上工业+9.1%、高技术制造+23.6%、数字经济核心+21.2%，制造在旺；但固定资产投资-1.8%、房地产投资-21.8%（扣地产才+6.4%）。工业增长来自技术/设备，但整体投资与地产拖累。")
quote_box(doc, "宏观错位二号：**资本消费强 vs 内需外贸分化**。社零1035.8亿+4.4%（差6.5%目标）、限上+6.9%、网络零售+33.6%；但对美出口-23.1%、机电/高科出口走弱、内需新消费（家电以旧换新）强。")
quote_box(doc, "宏观错位三号：**总量高增 vs 财政低增、人口低速**。GDP+6.4%、规上+9.1%，但财收+2.5%、税收+3.6%；人口+0.7万（自然-3.12‰）。增长好、财税与人口改善有限。")
para(doc, "三条错位加在一起，指向同一个结论：丽水2025年是\u201c以规上制造+高技术+三产为核心、投资/地产/外贸/内需偏弱\u201d的生态制造型增长。工业技术创新驱动是亮点，但固投-1.8%、社零+4.4%、对美-23.1%是短板。丽水要把制造优势转化为财政、人口与内需增量。")

# ---- 四、容易被忽视、但信号很重的细节（15条） ----
heading1(doc, "四、容易被忽视、但信号很重的细节（15条）")
para(doc, "下面这15条，散落在统计公报的角落里，单独看无关紧要，合在一起却拼出丽水2025年真实的结构肌理。")
bullet(doc, "**1. 规上工业+9.1%（汽车制造+35.3%、计算机通信+13.7%、通用设备+11.3%）**：出口型制造（汽车/通信/通用装备）是工业引擎。")
bullet(doc, "**2. 高技术制造+23.6%、数字经济核心制造+21.2%**：新质制造快于整体、是丽水升级的确定性。")
bullet(doc, "**3. 规上工业研发费用86.9亿+9.4%、约为营收3.4%**：民企研发强、创新驱动好。")
bullet(doc, "**4. 工业增加值658.5亿+7.0%、建筑业158.2亿持平**：工业是第二产业主力、建筑缓。")
bullet(doc, "**5. 货币/存款6229.0亿+6.1%、贷款6074.3亿+10.7%**：存贷两旺、贷款更猛、金融支持足。")
bullet(doc, "**6. 家具+8.4%、家用电器+23.9%**：以旧换新政策直接拉动家装家电消费。")
bullet(doc, "**7. 出口393.9亿+7.2%、一带一路+13.3%、但对美-23.1%**：市场多元化对冲对美下滑。")
bullet(doc, "**8. 机电产品占出口51.2%**：机电设备是出口绝对主力。")
bullet(doc, "**9. 一般公共预算收入198.0亿+2.5%（税收+3.6%、占74.9%）**：财税慢增但税收质量中上。")
bullet(doc, "**10. 民生支出占比78.5%、住房保障+17.9%**：支出向民生、住房倾斜。")
bullet(doc, "**11. 常住253.9万、+0.7万**：人口在浙江山区中仍正增长（靠县域+吸引）。")
bullet(doc, "**12. 城镇化率67.0%（+1.0pct）**：城镇化率处于浙江中上、持续提升。")
bullet(doc, "**13. 城镇低收入农户收入24006元+11.9%**：低收入农户补短板、共富推进。")
bullet(doc, "**14. 生态/文化：GEP 5000亿（目标）、生态屏障+黄帝缙云**：生态资产与文旅结合。")
bullet(doc, "**15. 城乡收入比1.85（缩小0.02）**：城乡差距收窄、共富。")
# ---- 五、2025年GDP目标 vs 实际：对照表 ----
heading1(doc, "五、2025年GDP目标 vs 实际：对照表")
para(doc, "把丽水市2025年政府工作报告设定的目标，与统计公报披露的实际值逐项对照：")
table(doc,
    ["指标", "2025年目标", "2025年实际", "判定"],
    [
        ["地区生产总值增速", "+6.5%左右", "+6.4%", "基本达成"],
        ["规上工业增加值增速", "+9%", "+9.1%", "达成，超额0.1pct"],
        ["固定资产投资增速", "+7%以上", "-1.8%", "未达标"],
        ["社会消费品零售总额", "+6.5%以上", "+4.4%", "未达标"],
        ["一般公共预算收入", "+3.8%（自然）", "+2.5%", "略低"],
        ["出口总额增速", "+6.5%", "+7.2%", "达成"],
    ],
    widths=[4.6, 3.2, 4.4, 3.8],
)
para(doc, "这张表透露丽水2025年的\u201c成色\u201d：**规上、出口达成，GDP基本达成6.5%**；但**固投-1.8%（差7%目标）、社零+4.4%（差6.5%）、财收+2.5%是主要短板**。\u201c制造强、投资内需财政偏弱\u201d是对2025年最简练的总结。")

# ---- 六、2025年增长是谁撑起来的？（结构归因） ----
heading1(doc, "六、2025年增长是谁撑起来的？（结构归因）")
para(doc, "丽水2025年GDP+6.4%的功劳簿，大致可以拆成四块：")
bullet(doc, "**第一功：规上制造+高技术（工业引擎）**。规上+9.1%、高技术制造+23.6%、数字经济核心+21.2%；汽车制造+35.3%、通信电子+13.7%、通用设备+11.3%。制造是绝对主力。")
bullet(doc, "**第二功：三产+7.0%（消费与前服务）**。第三产业1356.9亿、占比59%，生态旅游、健康、批发零售、金融支撑。")
bullet(doc, "**第三功：第二产业+5.6%、一产+4.6%**。农林、制造双稳。")
bullet(doc, "**拖累：投资/地产**。固投-1.8、房地产-21.8%拖累固定资产投资，但扣地产仍+6.4%（工业/基建投资在补）。")
para(doc, "**结构归因结论**：2025年的丽水增长是\u201c规上制造+高技术+三产\u201d双轮驱动，工业最亮（+9.1%）；投资因地产-21.8%被拖、社零+4.4%偏弱。丽水的增长靠制造创新，但投资转化、内需与财政余额待加强。")

# ---- 七、预算与财政的\u201c含金量\u201d ----
heading1(doc, "七、预算与财政的\u201c含金量\u201d")
para(doc, "看丽水财政，收入低增、但税收质量尚可。")
table(doc, ["财政指标", "2025年数值", "同比", "点评"],
    [
        ["一般公共预算收入", "198.0亿元", "+2.5%", "低增"],
        ["＃税收收入", "148.3亿元", "+3.6%", "高于收入增速"],
        ["＃税收占比", "74.9%", "—", "质量中上"],
        ["一般公共预算支出", "617.7亿元", "-5.0%", "支出收缩"],
        ["＃民生支出占比", "78.5%", "住房+17.9%", "民生优先"],
    ],
    widths=[4.6, 3.2, 3.2, 4.6],
)
para(doc, "财政核心判断：**收入低增（+2.5%）但税收占比74.9%（质量较高）**。税收+3.6%、由制造/服务税基带动；支出-5.0%收缩、民生占比78.5%优先（住房+17.9%、社保+6.7%）。财政在生态/转移省份里造血较弱但含金量不错。")
quote_box(doc, "财政与宏观的勾连：**GDP高、财税低**。GDP+6.4%但财收+2.5%，源于丽水作为生态屏障省对GDP转移制且山区税基企业规模小；虽税收质量尚可，但财政自给有限，需要中央/省转移与生态补偿支撑。")

# ---- 八、民生底账 ----
heading1(doc, "八、民生底账：人口、收入与城乡")
bullet(doc, "**居民收入**：居民人均可支配收入52979元+5.9%；城镇64895元+5.3%、农村35119元+6.4%；城乡比1.85（缩小0.02）。")
bullet(doc, "**就业**：城镇新增就业2.2万人。")
bullet(doc, "**低收入人口**：低收入农户人均可支配收入24006元+11.9%（补短板、共富明显）。")
bullet(doc, "**社保/民生**：民生支出484.8亿占78.5%、住房+17.9%、社保+6.7%。")
para(doc, "民生综合评价：**收入中高速、农村快于城镇、低收入农户高增、共富推进**；人口+0.7万在浙江山区仍正。生活水平在生态市里较均衡。")

# ================= 九、城镇与农村 =================
heading1(doc, "九、城镇与农村：格局与均衡")
bullet(doc, "**城镇化率67.0%（+1.0pct）**：城镇常住约170万、农村约84万，中等水平。")
bullet(doc, "**收入城乡**：城镇64895/农村35119、城乡比1.85（缩小0.02），相对均衡。")
bullet(doc, "**消费**：社零1035.8亿+4.4%，家电/家具以旧换新拉动强。")
bullet(doc, "**生态/文旅**：生态旅游、黄帝缙云、侨乡返馈，支撑农村旅游与消费。")
para(doc, "丽水城乡相对均衡（城乡比1.85低、低收入农户高增、共富成效），是浙江山区城乡协调样板。")

# ============ 十、人口流入与流出 ============
heading1(doc, "十、人口流入与流出")
bullet(doc, "**常住+0.7万**：253.9万人（户籍268.5万），在浙江山区罕见正增长。")
bullet(doc, "**自然增长-3.12‰**：出生1.15万（4.52‰）、死亡1.94万（7.64‰），自然负增、靠回流/流入补。")
bullet(doc, "**侨乡人口**：丽水华侨（青田侨乡）众多，回流人口与经济带来增量。")
bullet(doc, "**对经济含义**：253.9万人口微增、侨乡回流，是浙西南中人口较稳的市。")
quote_box(doc, "人口：常住+0.7万、侨乡回流，是山区市少有的人口正增长样本；若靠生态宜居+产业与侨乡经济留人，人口可稳住。")

# ============ 十一、物价与货币 ============
heading1(doc, "十一、物价与货币环境")
bullet(doc, "**CPI-0.1%**：消费品-1.1%、服务+1.2%，通缩轻微、消费偏冷。")
bullet(doc, "**金融**：存款6229.0亿+6.1%、贷款6074.3亿+10.7%——贷款增速远超存款、金融支持实体/基建好。")
para(doc, "通缩+信贷高增：资金便宜、金融活跃，对工业制造与基建有利。")

# ========== 十二、区域一体化 ==========
heading1(doc, "十二、区域一体化：丽水在浙江西南、金衢、杭州都市圈、生态屏障里的位置")
bullet(doc, "**浙江生态屏障·山海协作**：丽水是浙江的生态屏障（水/空气/GEP），承担\u201c两山\u201d理念转化样板。")
bullet(doc, "**金衢城市群辐射**：丽水位于金衢丽、温州/福建交界，承接东部制造业梯度转移。")
bullet(doc, "**侨乡/山海协作**：青田华侨经济、山海协作（丽水—杭州/温州），回流产业与资本。")
bullet(doc, "**县域**：龙泉（青瓷/宝剑）、缙云（制造/黄帝）、青田（侨乡/石雕）、云和（木玩/梯田）、庆元/景宁（畲乡），构成生态+特色制造骨架。")
para(doc, "丽水的区域坐标：**在\u201c浙江生态屏障+金衢丽\u201d双轮下，丽水是长三角南翼的生态绿谷与特色制造高地**。若生态补偿、山海协作、侨乡资本做实，丽水的\u201c生态+制造\u201d弹性更大。")

# ============ 十三、五大主线 ============
heading1(doc, "十三、未来5\u201310年最值得观察的五条主线")
bullet(doc, "**主线一：生态价值转化（GEP→GDP）**。生态产品价值实现、碳汇/生态补偿、\u201c两山\u201d转化，能否把生态变为财政与收入，是核心命题。")
bullet(doc, "**主线二：高技术制造与数字经济**。汽车/通信/通用设备+高技术（+23.6%）、数字经济核心（+21.2%），若制造继续升级，工业再上台阶。")
bullet(doc, "**主线三：五金/青瓷/宝剑特色产业的品牌化**。把龙泉青瓷/宝剑、特色轻工做成品牌+工业旅游，增利增收。")
bullet(doc, "**主线四：投资地产的修复**。固投-1.8%、地产-21.8%，需基建、制造业、保障房投资补。")
bullet(doc, "**主线五：侨乡回流与人口**。常住+0.7万、侨乡回流，若把这些转化为消费与产业，人口红利可期。")
para(doc, "五条主线里，**最值得盯的是主线一（生态价值化）与主线二（高技术制造）**——丽水要崛起，关键是\u201c生态\u201d能不能转成\u201c产业+财源\u201d、\u201c制造\u201d能不能持续升维。")


# ============ 十四、最终结论 ============
heading1(doc, "十四、最终结论：丽水在\u201c生态+制造+侨乡\u201d里的增长逻辑")
para(doc, "综合2025年的开局与结构，给丽水一个平衡的结论：")
para(doc, "**一、增长不错、制造最亮。**GDP+6.4%、规上+9.1%、高技术制造+23.6%、三产+7.0%，制造业与高技术驱动是最大亮点；生态之城走出了一条\u201c制造+数字\u201d的增长。")
para(doc, "**二、但投资/内需/财政偏弱。**固投仅有-1.8%、社零+4.4%、财收+2.5%，房地产-21.8%拖累；汽车/机电出口强但内需消费和投资转化不足。")
para(doc, "**三、生态与侨乡是变量。**GEP生态价值、侨乡回流（常住+0.7万），有机会转为产业与消费。")
para(doc, "**四、2026年前瞻。**在\u201c稳制造、提内需、修投资、强财政\u201d约束下，丽水2026年边际变量是**高技术制造、生态价值化（GEP）、侨乡资本、固投地产修复、数字经济**。若用制造+生态+侨乡三张牌拉动投资与财政，丽水增速可望更均衡。")
para(doc, "**一句话总结：丽水2025年交出了\u201c+6.4%、规上+9.1%、高技术+23.6%\u201d的亮眼答卷，靠制造与数字经济把山河生态市盘活了；但更该问的是——在生态屏障与山区禀赋下，这座城市能否把\u201c生态+青瓷宝剑+五金\u201d的牌炼成\u201c制造+生态价值+侨乡经济\u201d的均衡竞争力，让增长不只靠工业、也落到投资、消费与人口。**当下看，工业强、生态好，但投资内需财政的短板提醒：补齐这三块，丽水的\u201c生态制造\u201d之路才完整。")

# ============ 附录A ============

heading1(doc, "附录A：主要资料来源与核验路径")
bullet(doc, "《2025年丽水市国民经济和社会发展统计公报》（丽水市统计局，2026-05）：GDP、三次产业、工业、投资、外贸、财政、收入、人口等全部数据。")
bullet(doc, "《丽水市2026年政府工作报告》（2026-01市人大）：2025年回顾、2026年目标主线。")
bullet(doc, "《2025年丽水市国民经济和社会发展计划》（丽政函〔2025〕13号）：2025年目标设定（GDP+6.5%、规上+9%、固投+7%、社零+6.5%）。")
bullet(doc, "核验方式：以统计公报为准、以政府工作报告/计划为框架，交叉比对目标-实际，关注可比口径。")

# ============ 附录B ============
heading1(doc, "附录B：建议建立年度跟踪仪表盘")
table(doc, ["维度", "跟踪指标", "用途/预警"],
    [
        ["总量", "GDP、人均GDP", "守住+6.5%"],
        ["工业", "规上/高技术/数字经济、研发", "升级进度"],
        ["生态", "GEP、碳汇、生态补偿", "生态价值化"],
        ["投资", "固投、房地产、制造业投资", "投资修复"],
        ["外贸", "进出口、出口、对美", "外贸多元"],
        ["财政", "收入、税收占比、转移支付", "财政质量"],
        ["人口", "常住、侨乡回流", "人口动向"],
    ],
    widths=[2.6, 6.2, 6.0],
)

# ========== 保存
doc.save("/Users/x/Desktop/content-prod-lab/reports/丽水市_2025年政府工作报告_深度研究_2026-08-14.docx")
print("SAVED OK: 丽水市 /Users/x/Desktop/content-prod-lab/reports/丽水市_2025年政府工作报告_深度研究_2026-08-14.docx")
