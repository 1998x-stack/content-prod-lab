# -*- coding: utf-8 -*-
"""Build 济宁市2025年政府工作报告 深度研究 DOCX, 参照地级市系列版式。"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DARK = RGBColor(0x1F, 0x2A, 0x44)
RED = RGBColor(0xB0, 0x1E, 0x2E)
GRAY = RGBColor(0x59, 0x60, 0x69)
BLUE = RGBColor(0x1F, 0x4E, 0x79)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
HEADER_FILL = "1F2A44"
K = "微软雅黑"


def set_run(run, size=11, bold=False, color=DARK, font=K):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    run.font.name = font
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), font)
    rFonts.set(qn("w:hAnsi"), font)
    rFonts.set(qn("w:eastAsia"), font)


def para(doc, text, size=11, bold=False, color=DARK, align=None,
         space_after=6, space_before=0, font=K):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        set_run(r, size=size, bold=bold or (i % 2 == 1), color=color, font=font)
    return p


def heading1(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(16)
    pf.space_after = Pt(8)
    r = p.add_run(text)
    set_run(r, size=16, bold=True, color=RED)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "B01E2E")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def heading2(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(10)
    pf.space_after = Pt(4)
    r = p.add_run(text)
    set_run(r, size=13, bold=True, color=BLUE)
    return p


def table(doc, headers, rows, widths=None, font_size=9.5):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_run(r, size=font_size, bold=True, color=WHITE)
        tcPr = hdr[i]._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), HEADER_FILL)
        tcPr.append(shd)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(val))
            set_run(r, size=font_size, color=DARK)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    return t


def quote_box(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(8)
    pf.left_indent = Pt(12)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), "B01E2E")
    pBdr.append(left)
    pPr.append(pBdr)
    r = p.add_run(text)
    set_run(r, size=11, color=DARK, font="楷体")
    return p


def bullet(doc, text, size=10.5):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.7)
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        set_run(r, size=size, bold=(i % 2 == 1), color=DARK)
    return p


# ================================================================= 文档主体
doc = Document()
sec = doc.sections[0]
sec.page_width = Cm(21)
sec.page_height = Cm(29.7)
sec.top_margin = Cm(2.4)
sec.bottom_margin = Cm(2.2)
sec.left_margin = Cm(2.5)
sec.right_margin = Cm(2.5)

st = doc.styles["Normal"]
st.font.name = K
st.font.size = Pt(11)
st.element.rPr.rFonts.set(qn("w:eastAsia"), K)


# ---- 封面 ----
para(doc, "赣州市2025年政府工作报告", size=24, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=60, space_after=4)
para(doc, "深度研究与\u201c容易被忽视的细节\u201d分析", size=16, bold=True, color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
para(doc, "从\u201c省域副中心、稀土、赣南脐橙、深赣对口合作、苏区振兴、客家摇篮\u201d重新理解赣州", size=12, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)
para(doc, "\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501", color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
para(doc, "研究对象：2025年赣州市政府工作报告及同期财政、国民经济和社会发展统计资料、2026年官方复盘", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
para(doc, "标定日期：2026-08-14", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
doc.add_page_break()

# ---- 目录 ----
catalog = [
    "执行摘要",
    "一、研究口径与阅读方法",
    "二、先看赣州的特别底盘：省域副中心、稀土、脐橙、深赣合作、客家摇篮",
    "三、最关键的宏观错位：破5000亿、工业外贸强，但财收+1.9%、人口-0.44万、消费一般",
    "四、容易被忽视、但信号很重的细节（15条）",
    "五、2025年GDP目标 vs 实际：对照表",
    "六、2025年增长是谁撑起来的？（结构归因）",
    "七、预算与财政的\u201c含金量\u201d",
    "八、民生底账：人口、收入与城乡",
    "九、城镇与农村：格局与均衡",
    "十、人口流入与流出",
    "十一、物价与货币环境",
    "十二、区域一体化：赣州在深圳对口合作、粤港澳大湾区、中部\u201c三圈\u201d里的位置",
    "十三、未来5—10年最值得观察的五条主线",
    "十四、最终结论：赣州在\u201c稀土+家具+脐橙+接珠融湾\u201d里的增长逻辑",
    "附录A：主要资料来源与核验路径",
    "附录B：建议建立的年度跟踪仪表盘",
]
para(doc, "目　录", size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=10, space_after=10)
for item in catalog:
    para(doc, item, size=12, space_after=4)
doc.add_page_break()

# ---- 执行摘要 ----
heading1(doc, "执行摘要")
para(doc, "**核心判断**　2025年赣州最显著的是\u201cGDP 5221.29亿/+5.5%（破5000亿、达成目标、江西第2）、规上工业+9.2%（超额8.5%）、高技术制造+14.7%、进出口+14.4%破1200亿（出口+10%）、装备制造+15%\u201d、\u201c但一般公共预算收入+1.9%、人口-0.44万、CPI+0.2%\u201d。这说明赣州在\u201c省域副中心+对接大湾区\u201d中，**工业外贸强、财政人口消费偏弱**。")
para(doc, "把2025年目标（GDP+5.5%、规上+8.5%、固投+6%、社零+6%、财收+2%、城市收入+5%/农村+7.5%）、2025年实际（GDP+5.5%达成、规上+9.2%、固投+4.6%、社零+5.6%、财收+1.9%）趋势看，赣州在\u201c工业+稀土+脐橙\u201d路径：**稀土/钨、电子信息/家电、汽车零部件、脐橙、新材料**是支柱；非公经济+7.1%。")
para(doc, "最容易记住的一句话：**赣州是\u201c红色故都、省域副中心、稀土之都、脐橙之乡\u201d，靠\u201c对接大湾区（深赣合作）+稀土产业+红色文旅\u201d驱动。**观察赣州，与其只看\u201cGDP 5221亿\u201d，不如看\u201c规上+9.2%、进出口破1270亿、稀土/机电出口、深赣对口合作、脐橙品牌全国第1\u201d。")
# ---- 一、研究口径与阅读方法 ----
heading1(doc, "一、研究口径与阅读方法")
para(doc, "本报告以\u201c赣州市政府工作报告（2025年2月）\u201d为起点，把\u201c2025年GDP目标（5.5%）\u201d与\u201c官方2025年（5221.29亿/+5.5%）\u201d并置对照，用\u201c2025年赣州市统计公报\u201d和\u201c2026年政府工作报告\u201d横向核验。")
para(doc, "口径提示：\u201cGDP、规上工业增速\u201d按可比价（实际）；\u201c投资、消费、进出口、财政\u201d按现价（名义）。\u201c常住人口\u201d用统计公报（895.62万）、城镇化率60.27%。")
para(doc, "指标体系（与研究口径一致）：**总量与增速、产业动能（稀土/家电/脐橙）、外贸、财政、民生与人口**。")
para(doc, "特别提示（不吃老本）：赣州2024年GDP（+5.4%）、2025年5221.29亿/+5.5%（破5000亿）；它不是\u201c只有脐橙\u201d——**稀土（中国稀土集团）、电子信息（家电/整机）、汽车零部件、家具（南康）、红色苏区\u201d才是真正底色；深赣对口合作。")
# ---- 二、先看赣州的特别底盘 ----
heading1(doc, "二、先看赣州的特别底盘：省域副中心、稀土、脐橙、深赣合作、客家摇篮")
para(doc, "赣州地处江西南部、赣江源头、粤闽赣四省交界，是**江西省域副中心城市（江西第二城）、全国唯一的省域副中心红色故都（瑞金/苏区）、稀土之都（中国稀土集团总部）、脐橙之乡、深赣对口合作地**。2025年GDP 5221.29亿（破5000亿）、常住895.62万（江西人口第2）、城镇化率60.27%。")
para(doc, "五个底盘名词，先立框架：")
bullet(doc, "**稀土/钨（战略资源）**　中国稀土集团（赣州总部）、稀土（世界最大离子型稀土）、钨、新材料——\u201c稀土之都\u201d。")
bullet(doc, "**电子信息/家电制造**　家电（格力/美的）、整机、电子元器件、立讯——\u201c赣州制造\u201d。")
bullet(doc, "**汽车零部件/装备**　新能源车零部件、机电产品（出口600亿）、装备制造+15%——\u201c装备制造\u201d。")
bullet(doc, "**家具/脐橙（南康/赣南）**　南康家具（千亿）、赣南脐橙（品牌全国第1）、富硒农业——\u201c家具+脐橙\u201d。")
bullet(doc, "**红色文旅/客家文化/枢纽**　瑞金/长征红色旅游、客家摇篮、对接大湾区（深圳）——\u201c红色+融湾\u201d。")
para(doc, "这五根（稀土+家电+装备+家具脐橙+红色）构成赣州独特底盘：**左手战略资源（稀土）+右手接珠（深圳）**。理解赣州，先理解\u201c省域副中心、稀土、脐橙、深赣合作\u201d。")
# ---- 三、最关键的宏观错位 ----
heading1(doc, "三、最关键的宏观错位：破5000亿、工业外贸强，但财收+1.9%、人口-0.44万、消费一般")
para(doc, "2025年赣州最需要辨析的一组\u201c错位\u201d：**GDP+5.5%（破5000亿）、规上+9.2%、高技术+14.7%、进出口+14.4%、装备+15%强，但一般公共预算收入+1.9%、人口-0.44万、固投+4.2%、社零+5.6%、CPI+0.2%**。")
para(doc, "为什么\u201c工业/出口/装备\u201d强，财政/人口/投资却不温？三解释：")
para(doc, "**其一，工业/出口/新质强、量大**　规上+9.2%（家电/电子/稀土）、进出口+14.4%、非公+7.1%、装备+15%——\u201c制造出口强\u201d。")
para(doc, "**其二，财政/人口/固投偏弱**　财收+1.9%（低2%目标）、人口-0.44万（净流出）、固投+2.2%/房地产-7.5%——\u201c财政/人口地产弱\u201d。")
para(doc, "**其三，消费/物价温**　社零+5.6%、CPI+0.2%、消费弱修复——\u201c需求修复、量价\u201d。")
para(doc, "小结：赣州2025年是\u201c**工业出口强、财政人口地产弱**\u201d：稀土/家电/装备、外贸强，财收、人口、地产弱。")

# ---- 四、容易被忽视、但信号很重的细节（15条） ----
heading1(doc, "四、容易被忽视、但信号很重的细节（15条）")
bullet(doc, "**1.规上工业+9.2%（超额8.5%）**\u201c赣州制造强。\u201d")
bullet(doc, "**2.高技术制造+14.7%、装备制造+15%**\u201c新质制造。\u201d")
bullet(doc, "**3.进出口1272.05亿/+14.4%（出口+10%/进口+22.5%）**\u201c外贸强、机电出口600亿、汽车/锂电池+7.9/2.8倍。\u201d")
bullet(doc, "**4.稀土（中国稀土集团）、钨、新材料**\u201c稀土战略资源。\u201d")
bullet(doc, "**5.南康家具（千亿）、赣南脐橙（品牌第1）**\u201c家具+脐橙。\u201d")
bullet(doc, "**6.深赣对口合作（全国最优等次、对标深圳规制）**\u201c接珠融湾。\u201d")
bullet(doc, "**7.非公有+7.1%、民营企业营收+10.7%**\u201c民营活跃。\u201d")
bullet(doc, "**8.固定资产投资+4.2%（基础设施+14%）**\u201c基建强、地产弱。\u201d")
bullet(doc, "**9.房地产开发投资-7.5%、商品房销售面积-1.0%**\u201c地产调整。\u201d")
bullet(doc, "**10.一般公共预算收入333.89亿/+1.9%（税收-1.1%）**\u201c财收稳、税收待升。\u201d")
bullet(doc, "**11.社零2352.42亿/+5.6%、以旧换新破100亿**\u201c消费/以旧换新。\u201d")
bullet(doc, "**12.规模以上服务业+8.5%、旅游+9.6%人次**\u201c服务/红色旅游。\u201d")
bullet(doc, "**13.居民收入33813元/+5.2%、城乡比2.43（缩）**\u201c农村+6.3%>城镇+4.4%。\u201d")
bullet(doc, "**14.CPI+0.2%、常住895.62万(净流出-0.44万)**\u201c低通胀、人口微流。\u201d")
bullet(doc, "**15.营商环境全省第一、规上企业3117家（全省第1）、高企1148家**\u201c产业生态。\u201d")
# ---- 五、2025年GDP目标 vs 实际：对照表 ----
heading1(doc, "五、2025年GDP目标 vs 实际：对照表")
table(doc,
    ["指标", "2025年目标", "2025年实际", "达成评价"],
    [
        ["地区生产总值(GDP)", "增长5.5%左右", "5221.29亿/5.5%", "达成(破5000)"],
        ["规模以上工业", "增长8.5%左右", "+9.2%", "超额"],
        ["固定资产投资", "增长6%左右", "+4.6%", "略低"],
        ["社会消费品零售总额", "增长6%左右", "2352.42亿/+5.6%", "略低"],
        ["进出口总额", "促稳提质", "1272.05亿/+14.4%", "大幅超额"],
        ["一般公共预算收入", "增长2%左右", "333.89亿/+1.9%", "基本达成"],
        ["城镇居民人均收入", "增长5%左右", "+4.4%", "略低"],
        ["农村居民人均收入", "增长7.5%左右", "+6.3%", "略低"],
    ],
)
para(doc, "注：GDP、规上工业按可比价。**GDP（+5.5%）、规上（+9.2%）、进出口（+14.4%）达成/超额**；**固投（+4.6%）、社零（+5.6%）、财收（+1.9%）**略低目标。")
para(doc, "拆读：**工业（稀土/家电），出口/机电、非公是亮色**；**财收（+1.9%）、人口（-0.44万）、地产（-7.5%）、CPI（+0.2%）**是短板——\u201c制造外贸强、财政人口弱\u201d，是\u201c对接大湾区+红色苏区\u201d样本。")

# ---- 六、2025年增长是谁撑起来的？（结构归因） ----
heading1(doc, "六、2025年增长是谁撑起来的？（结构归因）")
para(doc, "把赣州GDP的5.5%拆开：三次产业分别+3.9%、+6.9%、+4.8%（贡献：一产7.3%、二产47.9%、三产44.8%）（结构9.6：36.4：54.0）。**第二产业（工业）+6.9%、贡献47.9%最强，第三产业（服务）贡献44.8%**。")
para(doc, "2026年赣州强调\u201c省域副中心、7510计划、接珠\u201d，聚焦**稀土、电子信息（家电）、汽车、家具、脐橙、红色文旅、深赣对口**——核心是\u201c工业+融湾\u201d。")
para(doc, "**第二产业（工业）**：规上+9.2%（家电/电子/稀土/汽车零部件）、高技术+14.7%、装备+15%——\u201c制造强\u201d。")
para(doc, "**第三产业（服务业）**：+4.8%（商贸、物流、旅游+9.6%、金融）、规上服务业+8.5%——\u201c服务+红色文旅\u201d。")
para(doc, "**第一产业（农业）**：+3.9%（脐橙/富硒、粮食、生猪）——\u201c农业稳\u201d。")
para(doc, "一句话归因：**2025年赣州增长\u201c靠工业（稀土/家电/装备）+服务业（红色旅游）+进出口\u201d**，财政人口弱；\u201c制造+新质\u201d是核心。")

# ---- 七、预算与财政的\u201c含金量\u201d ----
heading1(doc, "七、预算与财政的\u201c含金量\u201d")
para(doc, "2025年赣州**一般公共预算收入333.89亿元（+1.9%）**；税收191.97亿（-1.1%、占比57.5%）；支出1126.44亿（-3.6%）、民生占82.9%。")
bullet(doc, "税收-1.1%、占比57.5%——\u201c含金量待升（依赖非税）\u201d。")
bullet(doc, "民生支出占82.9%（教育/社保/医疗高投入）。")
bullet(doc, "金融支撑：存贷款近2万亿、国企资产破1.25万亿、信贷支持稀土/制造/民营——宽信用。")
para(doc, "**财政含金量小结**：财收+1.9%（低GDP、税收-1.1%）、民生82.9%；财政对\u201c稀土产业、红色文旅、乡村振兴\u201d投入加大。")
# ---- 八、民生底账：人口、收入与城乡 ----
heading1(doc, "八、民生底账：人口、收入与城乡")
para(doc, "2025年赣州**居民人均可支配收入33813元（+5.2%）**，其中城镇48188元（+4.4%）、农村19839元（+6.3%），城乡比2.43（缩小0.04）。就业：城镇新增就业（完成112.5%）。")
para(doc, "人口画像：**常住895.62万/-0.44万、城镇化率60.27%（+1.39pct）**；江西人口第2、净流出（赴大湾区/珠三角）。")
para(doc, "民生投入：民生支出82.9%、10件民生实事、医保/教育、托位3.1万——民生扎实、红色保障。")

# ---- 九、城镇与农村：格局与均衡 ----
heading1(doc, "九、城镇与农村：格局与均衡")
para(doc, "赣州城镇化率60.27%（+1.39pct）；县域经济（南康家具/于都纺织/龙南电子信息/赣县稀土）；农村收入增速（+6.3%）>城镇（+4.4%），**城乡比2.43缩小**；乡村振兴、富硒消费。")
para(doc, "农业底盘：**脐橙（品牌全国第1）、粮食、富硒、蔬菜**——\u201c赣南脐橙+富硒农业\u201d。")
para(doc, "一句话：\u201c赣州是山地农业大市、县域产业强、乡村振兴加速、城乡比大但收窄\u201d。")

# ---- 十、人口流入与流出 ----
heading1(doc, "十、人口流入与流出")
para(doc, "赣州常住895.62万（-0.44万净流出）、城镇化60.27%；\u201c劳动力外流珠三角/大湾区、红色苏区\u201d；中心城区+经开区＋稀土产业吸引回流。")
para(doc, "结构观察：**外出务工多（厂深圳/东莞）、返乡创业**；深赣合作引才。")
para(doc, "2026年目标：城镇新增就业6万、对接深圳（同事同标）——赣州靠\u201c产业+赣港\u201d聚人。")

# ---- 十一、物价与货币环境 ----
heading1(doc, "十一、物价与货币环境")
para(doc, "2025年赣州**CPI+0.2%**（其他用品+14.2%、衣着+1.4%；交通通信-2.9%、食品烟酒-0.7%）——\u201c低通胀、需求温\u201d。")
bullet(doc, "信贷扩张：存贷款近2万亿、贷款支持产业——宽信用。")
bullet(doc, "消费：社零+5.6%、以旧换新破100亿、红色旅游——消费稳。")
para(doc, "货币环境判断：**宽信用、CPI+0.2%**；赣州靠\u201c制造+红色+农产品\u201d稳需求（2026 CPI 2%）。")

# ---- 十二、区域一体化：赣州的位置 ----
heading1(doc, "十二、区域一体化：赣州在深圳对口合作、粤港澳大湾区、中部\u201c三圈\u201d里的位置")
para(doc, "赣州是**江西省域副中心城市、对接粤港澳大湾区（接珠融湾）、深赣对口合作（全国最优等次）、中部地区重要城市**。")
bullet(doc, "**深赣对口合作**　对标深圳规制（619项同事同标）、深赣高铁）、电子信息/稀土/汽车产业承接——\u201c接珠融入\u201d。")
bullet(doc, "**粤港澳大湾区**　赣深高铁（2小时）、大湾区后花园、返乡兴业、铁海联运。")
bullet(doc, "**中部崛起**　江西中南、大余/信丰、苏区振兴（对口支援至2030）、红色苏区。")
para(doc, "一句话：**赣州在\u201c深圳+大湾区+中部\u201d里，最核心是\u201c深度对接大湾区（稀土/家电/承接）\u201d**；区位、深赣对口、红色资源是最大优势。")

# ---- 十三、未来5—10年最值得观察的五条主线 ----
heading1(doc, "十三、未来5—10年最值得观察的五大主线")
bullet(doc, "**主线一：稀土/钨战略产业（中国稀土集团）**\u201c稀土/新材料、战略资源\u201d能否成世界级。")
bullet(doc, "**主线二：电子信息/家电（格力美的立讯）**\u201c赣州制造、承接大湾区\u201d产业转移。")
bullet(doc, "**主线三：汽车零部件/机电出口**\u201c出口600亿、新能源零部件\u201d。")
bullet(doc, "**主线四：家具/脐橙/富硒农业**\u201c南康家具、脐橙品牌、红色\u201d。")
bullet(doc, "**主线五：深赣/融湾/人口**\u201c对标深圳、人口回、江西副中心\u201d。")

# ---- 十四、最终结论 ----
heading1(doc, "十四、最终结论：赣州在\u201c稀土+家电制造+脐橙/红色\u201d里的增长逻辑")
para(doc, "把2025年闭环打穿：**赣州是\u201c省域副中心、稀土之都、脐橙之乡\u201d**：GDP 5221.29亿/+5.5%（破5000）、规上+9.2%、进出口+14.4%、非公+7.1%。")
para(doc, "赣州不是\u201c只有脐橙\u201d——它是**稀土+家电+汽车+南康家具+红色文旅**的复合，靠\u201c制造+融湾\u201d驱动；但财政/人口/地产/消费弱。")
para(doc, "一句话结论：**赣州是\u201c红色故都、稀土之都、省域副中心\u201d；观察它先看\u201c稀土、家电、进出口、深赣、脐橙\u201d，再看\u201c财收、人口、地产\u201d。**它是\u201c制造融湾、财政人口待强\u201d的江西样本。")

# ---- 附录A ----
heading1(doc, "附录A：主要资料来源与核验路径")
bullet(doc, "《2025年赣州市政府工作报告》（2025年1月，2025年目标、2024年回顾+5.4%）")
bullet(doc, "《2025年赣州市国民经济和社会发展统计公报》（赣州市统计局，2026-04-30，2025年实际）")
bullet(doc, "《2026年赣州市政府工作报告》（2026年1月，复盘+2026年目标，由3703→5221）")
bullet(doc, "赣州市人民政府/统计局（ganzhou.gov.cn）")

# ---- 附录B：建议建立的年度跟踪仪表盘 ----
heading1(doc, "附录B：建议建立的年度跟踪仪表盘")
bullet(doc, "GDP总量/增速 vs 全年目标。")
bullet(doc, "规模以上工业（稀土/家电）增速。")
bullet(doc, "稀土产业/中国稀土集团。")
bullet(doc, "固定资产/工业/房地产投资。")
bullet(doc, "家电/汽车/家具/电子信息。")
bullet(doc, "进出口/机电/新能源汽车。")
bullet(doc, "社零/以旧换新/红色旅游。")
bullet(doc, "财收/税收/民生%。")
bullet(doc, "常住/城镇化/回流。")
bullet(doc, "CPI/存贷款/工业用电。")

# ---------------- 保存 ----------------
out = "/Users/x/Desktop/content-prod-lab/reports/赣州市_2025年政府工作报告_深度研究_2026-08-14.docx"
doc.save(out)
print("SAVED OK: 赣州市", out)
