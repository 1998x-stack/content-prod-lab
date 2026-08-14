# -*- coding: utf-8 -*-
"""Build 济宁市2025年政府工作报告 深度研究 DOCX, 参照地级市系列版式。"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DARK = RGBColor(0x1F, 0x2A, 0x44)
RED = RGBColor(0xB0, 0x1E, 0x2E)
GRAY = RGBColor(0x59, 0x60, 0x69)
BLUE = RGBColor(0x1F, 0x4E, 0x79)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
HEADER_FILL = "1F2A44"
K = "微软雅黑"


def set_run(run, size=11, bold=False, color=DARK, font=K):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    run.font.name = font
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), font)
    rFonts.set(qn("w:hAnsi"), font)
    rFonts.set(qn("w:eastAsia"), font)


def para(doc, text, size=11, bold=False, color=DARK, align=None,
         space_after=6, space_before=0, font=K):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        set_run(r, size=size, bold=bold or (i % 2 == 1), color=color, font=font)
    return p


def heading1(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(16)
    pf.space_after = Pt(8)
    r = p.add_run(text)
    set_run(r, size=16, bold=True, color=RED)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "B01E2E")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def heading2(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(10)
    pf.space_after = Pt(4)
    r = p.add_run(text)
    set_run(r, size=13, bold=True, color=BLUE)
    return p


def table(doc, headers, rows, widths=None, font_size=9.5):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_run(r, size=font_size, bold=True, color=WHITE)
        tcPr = hdr[i]._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), HEADER_FILL)
        tcPr.append(shd)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(val))
            set_run(r, size=font_size, color=DARK)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    return t


def quote_box(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(8)
    pf.left_indent = Pt(12)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), "B01E2E")
    pBdr.append(left)
    pPr.append(pBdr)
    r = p.add_run(text)
    set_run(r, size=11, color=DARK, font="楷体")
    return p


def bullet(doc, text, size=10.5):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.7)
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        set_run(r, size=size, bold=(i % 2 == 1), color=DARK)
    return p


# ================================================================= 文档主体
doc = Document()
sec = doc.sections[0]
sec.page_width = Cm(21)
sec.page_height = Cm(29.7)
sec.top_margin = Cm(2.4)
sec.bottom_margin = Cm(2.2)
sec.left_margin = Cm(2.5)
sec.right_margin = Cm(2.5)

st = doc.styles["Normal"]
st.font.name = K
st.font.size = Pt(11)
st.element.rPr.rFonts.set(qn("w:eastAsia"), K)


# ---- 封面 ----
para(doc, "九江市2025年政府工作报告", size=24, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=60, space_after=4)
para(doc, "深度研究与\u201c容易被忽视的细节\u201d分析", size=16, bold=True, color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
para(doc, "从\u201c长江经济带节点、庐山、鄱阳湖、石化钢铁、九江港\u201d重新理解九江", size=12, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)
para(doc, "\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501", color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
para(doc, "研究对象：2025年九江市政府工作报告及同期财政、国民经济和社会发展统计资料、2026年官方复盘", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
para(doc, "标定日期：2026-08-14", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
doc.add_page_break()

# ---- 目录 ----
catalog = [
    "执行摘要",
    "一、研究口径与阅读方法",
    "二、先看九江的特别底盘：庐山、鄱阳湖、石化钢铁、九江港、长江",
    "三、最关键的宏观错位：规上+7.2%、固投+4.3%、旅游9090万人次、但财收-1.2%、石化炼量-17.2%",
    "四、容易被忽视、但信号很重的细节（15条）",
    "五、2025年GDP目标 vs 实际：对照表",
    "六、2025年增长是谁撑起来的？（结构归因）",
    "七、预算与财政的\u201c含金量\u201d",
    "八、民生底账：人口、收入与城乡",
    "九、城镇与农村：格局与均衡",
    "十、人口流入与流出",
    "十一、物价与货币环境",
    "十二、区域一体化：九江在长江经济带、赣江、九江都市圈/武汉长沙南昌里的位置",
    "十三、未来5\u201310年最值得观察的五条主线",
    "十四、最终结论：九江在\u201c长江+庐山+石化\u201d里的增长逻辑",
    "附录A：主要资料来源与核验路径",
    "附录B：建议建立的年度跟踪仪表盘",
]
para(doc, "目　录", size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=10, space_after=10)
for item in catalog:
    para(doc, item, size=12, space_after=4)
doc.add_page_break()

# ---- 执行摘要 ----
heading1(doc, "执行摘要")
para(doc, "**核心判断**　2025年九江最显著的是\u201cGDP 4246.51亿/+5.2%（长江节点城市）、规上工业+7.2%（电子信息+12.9%、新能源+16.2%、钢铁+9.1%）、旅游9090.20万人次/817.26亿（+13%/+33.9%）、常住448.42万\u201d、\u201c但财收-1.2%、原油加工量-17.2%、进出口-1.0%\u201d。这说明九江在\u201c长江经济带+庐山+石化\u201d中，**工业制造/旅游强、港口开放、但财政与石化炼量承压**。")
para(doc, "把九江2025年目标与实际对照，九江是\u201c港口+旅游+工业\u201d路径：**石油化工、电子信息、新能源、装备制造、纺织服装、庐山旅游、九江港**是支柱。")
para(doc, "最容易记住的一句话：**九江是\u201c长江经济带节点城市、庐山、鄱阳湖、石化之城、九江港\u201d，靠\u201c工业+文旅+港口\u201d驱动。**观察九江，与其只看\u201cGDP 4247亿/+5.2%\u201d，不如看\u201c规上工业+7.2%、旅游+33.9%、港口吞吐、枢纽港口\u201d。")

# ---- 一、研究口径与阅读方法 ----
heading1(doc, "一、研究口径与阅读方法")
para(doc, "**本文的最终目标**，不是复述九江市2025年政府工作报告的原文，而是把它放到\u201c目标→实际\u201d和\u201c长期底盘→年度对撞\u201d的坐标里重新读一遍，找出九江2025年到底靠什么增长、哪里在换挡、哪里是暗线。")
para(doc, "**三组资料互证**：（1）《九江市2025年国民经济和社会发展统计公报》（市统计局，2026-05-29发布）提供全部2025年**实际完成**数据；（2）《九江市2026年政府工作报告》（2026-02市十六届人大七次会议）复盘2025年；（3）《九江市2025年政府工作报告》提供**目标设定**。三份材料交叉核验，避免单源失真。")
para(doc, "**核心框架**：先交代九江的\u201c底盘\u201d（长江+庐山+鄱阳湖+石化港口的禀赋），再用\u201c目标vs实际\u201d对照表定位宏观错位，接着用15条细节捕捉暗线，最后从财政、民生、城乡、人口、物价坐实，落到2026/十五五主线与结论。")
para(doc, "**口径说明**：GDP 4246.51亿元+5.2%、人均94544元+5.8%。三次产业结构6.7:38.3:55.0。常住448.42万人、城镇化率65.06%、自然增长-2.45‰。固投+4.3%、房地产-21.2%。社零1788.14亿+4.6%。一般公共预算收入322.02亿-1.2%（地方税收194.67+3.5%、占60.5%）。进出口398.7亿-1.0%。")
para(doc, "**为什么值得单独研究九江**：GDP+5.2%、规上+7.2%、旅游+33.9%，是长江经济带节点城市；但财收-1.2%、原油加工-17.2%、人口-1.48万。\u201c工业旅游强、财政石化炼量弱\u201d，是观察\u201c长江港口城市在石化调整+文旅升级\u201d的上佳切片。")
# ---- 二、先看九江的特别底盘 ----
heading1(doc, "二、先看九江的特别底盘：庐山、鄱阳湖、石化钢铁、九江港、长江")
para(doc, "九江的成长逻辑，不能只看GDP数字，而要先把\u201c底盘\u201d摊开。底盘决定了几十年、上百年的禀赋，也决定了2025年增长的来源和约束。")
bullet(doc, "**庐山·世界文化景观**：庐山（世界文化景观遗产）是世界级名山，鄱阳湖（中国最大淡水湖）环绕，九江是江西的旅游/生态名片（旅游9090万人次+13%、总收入+33.9%）。")
bullet(doc, "**长江经济带节点·九江港**：九江是长江经济带重要的港口节点城市（长江中游），九江港是长江中游枢纽，铁公水多式联运。")
bullet(doc, "**石化基地**：九江是全国重要的石化基地（九江石化/央企），2025年原油加工量641.09万吨（-17.2%），石化化工是工业龙头（增加值+6.5%）。")
bullet(doc, "**制造业集群（电子信息/新能源/装备）**：电子信息（+12.9%）、新能源（+16.2%）、钢铁（+9.1%）、有色（+4.5%）、装备（+5.0%）等九大产业链，是江西工业重镇。")
bullet(doc, "**赣北农业（粮食/油料/水产品）**：粮食150.73万吨+1.8%、油料+2.5%、水产品+2.8%，赣北粮仓。")
heading2(doc, "底盘的产业钟摆")
para(doc, "把底盘归纳成一句话：**九江是\u201c长江港口+石化+旅游\u201d的临江重工业/文旅城市**。长江+九江港定枢纽、石化定工业、庐山鄱阳湖供文旅。城市资产在\u201c港口+工业+文旅\u201d链条上。2025年规上+7.2%、旅游+33.9%亮眼，但财收-1.2%、炼量-17.2%承压。")

# ---- 三、最关键的宏观错位 ----
heading1(doc, "三、最关键的宏观错位：规上+7.2%、旅游9090万人次、但财收-1.2%、石化炼量-17.2%、人口-1.48万")
para(doc, "把2025年数据摊开，九江最醒目的不是\u201c+5.2%\u201d，而是**工业旅游强、财政与石化炼量弱**的三重错位。")
quote_box(doc, "宏观错位一号：**工业旅游强 vs 财政弱**。规上+7.2%、旅游+33.9%、社零+4.6%，经济在增；但一般公共预算收入-1.2%（税收+3.5%勉强）、支出+4.5%——财政收入负增长，与经济脱节。")
quote_box(doc, "宏观错位二号：**石化制造高 vs 炼量低**。原油加工量641.09万吨-17.2%（减炼量/油品调整），但石化化工增加值+6.5%？——石化向化工深加工转、炼油在减。")
quote_box(doc, "宏观错位三号：**港口/文旅强 vs 房地产/人口弱**。旅游+33.9%、港口货运增长；但房地产投资-21.2%、常住人口-1.48万（自然-2.45‰）。")
para(doc, "三条错位加在一起，指向同一个结论：九江2025年是\u201c港口+制造+文旅驱动、财政与地产人口偏弱\u201d的长江经济带节点城市。增长靠工业制造与旅游，但财政负增（石化利润/税基）、地产下滑与人口自然减是短板。")

# ---- 四、容易被忽视、但信号很重的细节（15条） ----
heading1(doc, "四、容易被忽视、但信号很重的细节（15条）")
para(doc, "下面这15条，散落在统计公报的角落里，单独看无关紧要，合在一起却拼出九江2025年真实的结构肌理。")
bullet(doc, "**1. 电子信息+12.9%、新能源+16.2%**：新兴产业在规上工业中领跑。")
bullet(doc, "**2. 钢铁+9.1%、不完全金属、装备+5.0%**：重工业强、装备稳。")
bullet(doc, "**3. 规上工业利润255.57亿+5.5%**：工业创利、效益稳。")
bullet(doc, "**4. 旅游9090.20万人次+13%、总收入817.26亿+33.9%**：庐山鄱阳湖带旺文旅、人均花费提升。")
bullet(doc, "**5. 原油加工量641.09万吨-17.2%**：炼量下行、向深加工转。")
bullet(doc, "**6. 发电量307.56亿千瓦时+11.3%**：电力/能源稳增。")
bullet(doc, "**7. 食品+13.2%、轻工业+6.8%**：食品饮料、纺织（纱/服装）分化。")
bullet(doc, "**8. 固投+4.3%（二产+13.2%/非国有+5.9%/一产+39.1%）**：工业与农业投资好。")
bullet(doc, "**9. 制造业投资+18.0%、有色冶炼投资+64.4%**：制造业扩张。")
bullet(doc, "**10. 存款6197.33亿/贷款5957.94亿（平稳）**：金融稳。")
bullet(doc, "**11. 实际使用外资9545万美元+39.9%**：外资高增、开放提升。")
bullet(doc, "**12. 旅游总人数9090万人次、人均花费约899元**：文旅量价齐升。")
bullet(doc, "**13. 通用机场/航空旅客+24.7%**：交通/开放增强。")
bullet(doc, "**14. 常住448.42万-1.48万、城镇化65.06%（+0.58）**：人口城镇化提升但总量略减。")
bullet(doc, "**15. 城镇居民51298元+4.3%/农村24171元+5.8%**：农村快于城镇。")
# ---- 五、2025年GDP目标 vs 实际：对照表 ----
heading1(doc, "五、2025年GDP目标 vs 实际：对照表")
para(doc, "把九江市2025年预期目标与实际值逐项对照：")
table(doc,
    ["指标", "2025年目标", "2025年实际", "判定"],
    [
        ["地区生产总值增速", "+5.5%左右", "4246.51亿，+5.2%", "基本达成"],
        ["规上工业增加值", "+8%左右", "+7.2%", "略低"],
        ["固定资产投资", "+5%以上", "+4.3%", "略低"],
        ["社会消费品零售总额", "+6%左右", "1788.14亿，+4.6%", "略低"],
        ["进出口总额", "（稳外贸）", "398.7亿，-1.0%", "略降"],
        ["一般公共预算收入", "（+增）", "322.02亿，-1.2%", "未达、负增"],
    ],
    widths=[4.6, 3.2, 4.4, 3.8],
)
para(doc, "这张表透露九江2025年的\u201c成色\u201d：**GDP、规上、社零等均略低目标、财收-1.2%负增**；旅游、制造业等结构亮眼。\u201c总量稳、财政石化弱\u201d是对2025年最简练的总结。")

# ---- 六、2025年增长是谁撑起来的？ ----
heading1(doc, "六、2025年增长是谁撑起来的？（结构归因）")
para(doc, "九江2025年GDP+5.2%的功劳簿，大致可以拆成几块：")
bullet(doc, "**第一功：工业/规上+7.2%**。电子信息+12.9%、新能源+16.2%、钢铁+9.1%、有色+4.5%、装备+5.0%，制造业是主力。")
bullet(doc, "**第二功：三产+4.8%（旅游引擎）**。旅游+33.9%（817亿）、住宿餐饮/批发零售、庐山鄱阳湖文旅。")
bullet(doc, "**第三功：消费/港口**。社零1788亿+4.6%、港口与开放（外资+39.9%）。")
bullet(doc, "**拖累：石化炼量、财政**。原油加工-17.2%、财收-1.2%拖GDP增速。")
para(doc, "**结构归因结论**：2025年的九江增长是\u201c工业制造+旅游\u201d双轮，制造业（电子信息/新能源）与文旅是主引擎；但石化炼量、财政与人口是拖累。九江下一程在\u201c炼化深加工（精细化工）与文旅消费转化\u201d。")


# ---- 七、预算与财政的\u201c含金量\u201d ----
heading1(doc, "七、预算与财政的\u201c含金量\u201d")
table(doc, ["财政指标", "2025年数值", "同比", "点评"],
    [
        ["一般公共预算收入", "322.02亿元", "-1.2%", "负增"],
        ["＃地方税收", "194.67亿元", "+3.5%", "税收增、占60.5%"],
        ["＃非税", "约127亿", "下降", "非税拖累"],
        ["一般公共预算支出", "737.54亿元", "+4.5%", "支出扩张"],
        ["存/贷款", "6197.33/5957.94亿", "平稳", "金融稳"],
    ],
    widths=[4.6, 3.4, 3.0, 4.6],
)
para(doc, "财政核心判断：**总收入-1.2%、税收+3.5%（占60.5%）、非税下滑**。财政负增主因非税与石化（原油加工-17.2%）税基弱；支出+4.5%扩张。财收与GDP+5.2%脱节，反映石化/地产税基收缩。")
quote_box(doc, "财政与宏观的勾连：**经济增、财政负**。GDP+5.2%但财收-1.2%，石化（九江石化炼量-17.2%）与地产税收拖累。九江要靠制造业/旅游/新税基补财政。")

# ================= 八、民生底账 =================
heading1(doc, "八、民生底账：人口、收入与城乡")
bullet(doc, "**居民收入**：城镇51298元+4.3%、农村24171元+5.8%；农村快于城镇。")
bullet(doc, "**就业**：新增城镇就业5.34万、下岗再就业2.49万、困难群体1.1万。")
bullet(doc, "**收入与消费**：居民消费（城镇31477元/农村21215元）。")
bullet(doc, "**社保**：城/乡养老86.86+192.44万、医保458.47万。")
para(doc, "民生综合评价：**收入中速、农村快于城镇、就业稳**；人口自然-2.45‰、老龄化（60+占22.15%）是隐忧。")

# ================= 九、城镇与农村 =================
heading1(doc, "九、城镇与农村：格局与均衡")
bullet(doc, "**城镇化率65.06%**：常住448.42万、城镇291.74万，江西较高。")
bullet(doc, "**收入城乡**：城51298/农24171、比2.12，农村+5.8%快于城。")
bullet(doc, "**农业**：粮食150.73万吨、水产品55.2万吨+2.8%、茶叶/油料。")
para(doc, "九江城乡相对均衡（2.12）、赣北粮仓/水乡农业，城镇化中高。")

# ============ 十、人口流入与流出 ============
heading1(doc, "十、人口流入与流出")
bullet(doc, "**总量-1.48万**：常住448.42万、出生率5.16‰、死亡率7.61‰、自然-2.45‰。")
bullet(doc, "**结构老化**：60+占22.15%、65+15.38%，老龄加速。")
bullet(doc, "**城镇化提升**：+0.58pct、人口向城镇集中。")
quote_box(doc, "人口：常住448.42万、自然负增、老龄化，九江要留人靠制造/旅游/港口就业。")

# ============ 十一、物价与货币 ============
heading1(doc, "十一、物价与货币环境")
bullet(doc, "**CPI+0.3%**：六升二降（其他用品+12.7%/食品-0.7%/交通-3.1%），温和、无明显通缩。")
bullet(doc, "**金融**：存款6197.33亿/贷款5957.94亿（平稳），金融稳健。")
para(doc, "CPI+0.3%温和+存款贷款平衡：九江价格与金融环境稳定。")

# ========== 十二、区域一体化 ==========
heading1(doc, "十二、区域一体化：九江在长江经济带、赣江、九江都市圈/武汉长沙南昌里的位置")
bullet(doc, "**长江经济带节点**：九江是长江中游港口节点、长江经济带重要节点城市（三个区域中心）。")
bullet(doc, "**武汉/长沙/南昌三角**：九江在武汉-长沙-南昌、鄱阳湖生态经济区，铁公水联运。")
bullet(doc, "**九江港/江海联运**：九江港长江中游枢纽、通江达海。")
bullet(doc, "**县域**：浔阳/濂溪（城区/石化）、庐山（文旅）、湖口/彭泽（沿江/冶金）、瑞昌（工业），构成\u201c石化+港口+文旅\u201d。")
para(doc, "九江坐标：**在\u201c长江经济带+鄱阳湖+赣鄱\u201d下，九江是江西门户与长江节点**。港口、制造、文旅联动定能级。")

# ============ 十三、未来5-10年主线 ============
heading1(doc, "十三、未来5\u201310年最值得观察的五条主线")
bullet(doc, "**主线一：石化深加工（炼化一体）**。原油加工-17.2%，若向精细化工/新材料（PX/芳烃）转，工业增值。")
bullet(doc, "**主线二：电子信息/新能源制造**。电子信息+12.9%、新能源+16.2%，若放大为第二支柱。")
bullet(doc, "**主线三：庐山鄱阳湖文旅升级**。旅游+33.9%、量价齐升，酒店/度假/会展。")
bullet(doc, "**主线四：九江港/江海联运开放**。港口+外资+39.9%、开放能级。")
bullet(doc, "**主线五：人口与财政修复**。财政-1.2%、人口-1.48万，需新税基与留人。")
para(doc, "五条主线里，**最值得盯的是主线一（炼化深加工）与主线二（电子信息/新能源）**——九江从\u201c石化炼油\u201d到\u201c化工新材料+文旅\u201d，是长坡与增量。")

# ============ 附录A ============
heading1(doc, "附录A：主要资料来源与核验路径")
bullet(doc, "《九江市2025年国民经济和社会发展统计公报》（市统计局，2026-05）：GDP、三次产业、工业、投资、消费、财政、人口等。")
bullet(doc, "《九江市2026年政府工作报告》（2026-02）：2025年回顾、2026年目标。")
bullet(doc, "《九江市2025年政府工作报告》：目标设定。")
bullet(doc, "核验方式：以统计公报为准、以政府工作报告为框架，交叉比对目标-实际。")

# ============ 附录B ============
heading1(doc, "附录B：建议建立年度跟踪仪表盘")
table(doc, ["维度", "跟踪指标", "用途/预警"],
    [
        ["总量", "GDP、人均、三产占比", "稳总量"],
        ["工业", "规上、石化/电子信息/新能源", "制造升级"],
        ["旅游", "旅游人次/收入、庐山", "文旅变现"],
        ["投资", "固投、房地产、制造业", "投资质量"],
        ["港口", "九江港吞吐、外资", "开放能级"],
        ["财政", "收入、税收", "财税修复"],
        ["人口", "常住/自然增长率", "人口"],
    ],
    widths=[2.6, 6.2, 6.0],
)

# ========== 保存
doc.save("/Users/x/Desktop/content-prod-lab/reports/九江市_2025年政府工作报告_深度研究_2026-08-14.docx")
print("SAVED OK: 九江市 /Users/x/Desktop/content-prod-lab/reports/九江市_2025年政府工作报告_深度研究_2026-08-14.docx")
