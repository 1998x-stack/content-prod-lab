# -*- coding: utf-8 -*-
"""Build 辽宁省2025年政府工作报告 深度研究 DOCX, 参照省市系列版式。"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DARK = RGBColor(0x1F, 0x2A, 0x44)
RED = RGBColor(0xB0, 0x1E, 0x2E)
GRAY = RGBColor(0x59, 0x60, 0x69)
BLUE = RGBColor(0x1F, 0x4E, 0x79)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
HEADER_FILL = "1F2A44"
K = "微软雅黑"


def set_run(run, size=11, bold=False, color=DARK, font=K):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    run.font.name = font
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), font)
    rFonts.set(qn("w:hAnsi"), font)
    rFonts.set(qn("w:eastAsia"), font)


def para(doc, text, size=11, bold=False, color=DARK, align=None,
         space_after=6, space_before=0, font=K):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        set_run(r, size=size, bold=bold or (i % 2 == 1), color=color, font=font)
    return p


def heading1(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(16)
    pf.space_after = Pt(8)
    r = p.add_run(text)
    set_run(r, size=16, bold=True, color=RED)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "B01E2E")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def heading2(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(10)
    pf.space_after = Pt(4)
    r = p.add_run(text)
    set_run(r, size=13, bold=True, color=BLUE)
    return p


def table(doc, headers, rows, widths=None, font_size=9.5):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_run(r, size=font_size, bold=True, color=WHITE)
        tcPr = hdr[i]._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), HEADER_FILL)
        tcPr.append(shd)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(val))
            set_run(r, size=font_size, color=DARK)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    return t


def quote_box(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(8)
    pf.left_indent = Pt(12)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), "B01E2E")
    pBdr.append(left)
    pPr.append(pBdr)
    r = p.add_run(text)
    set_run(r, size=11, color=DARK, font="楷体")
    return p


def bullet(doc, text, size=10.5):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.7)
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        set_run(r, size=size, bold=(i % 2 == 1), color=DARK)
    return p


# ================================================================= 文档主体
doc = Document()
sec = doc.sections[0]
sec.page_width = Cm(21)
sec.page_height = Cm(29.7)
sec.top_margin = Cm(2.4)
sec.bottom_margin = Cm(2.2)
sec.left_margin = Cm(2.5)
sec.right_margin = Cm(2.5)

st = doc.styles["Normal"]
st.font.name = K
st.font.size = Pt(11)
st.element.rPr.rFonts.set(qn("w:eastAsia"), K)



# ---- 封面 ----
para(doc, "辽宁省2025年政府工作报告", size=24, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=60, space_after=4)
para(doc, "深度研究与“容易被忽视的细节”分析", size=16, bold=True, color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
para(doc, "从“老工业基地、沿海经济带、冰雪、装备制造与人口/财政”重新理解辽宁", size=12, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)
para(doc, "━━━━━━━━━━━━━━━━", color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
para(doc, "研究对象：2025年辽宁省政府工作报告及同期财政、国民经济和社会发展统计资料、2026年官方复盘", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
para(doc, "标定日期：2026-08-13", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
doc.add_page_break()

# ---- 目录 ----
catalog = [
    "执行摘要",
    "一、研究口径与阅读方法",
    "二、先看辽宁的特殊底盘：老工业基地、沿海经济带、装备制造与冰雪",
    "三、最关键的宏观错位：GDP破3.32万亿、利润/出口强，但投资/地产/人口弱",
    "四、容易被忽视、但信号很重的细节（15条）",
    "五、2025年GDP目标 vs 实际：对照表",
    "六、2025年增长是谁撑起来的？（结构归因）",
    "七、预算与财政的“含金量”",
    "八、民生底账：人口、收入与城乡",
    "九、城镇与农村：格局与均衡",
    "十、人口流入与流出",
    "十一、物价与货币环境",
    "十二、区域一体化：辽宁在“沿海经济带+东北振兴+沈阳都市圈”里的位置",
    "十三、未来5—10年最值得观察的五条主线",
    "十四、最终结论：辽宁在“装备制造+沿海+工业利润+冰雪”里的增长逻辑",
    "附录A：主要资料来源与核验路径",
    "附录B：建议建立的年度跟踪仪表盘",
]
para(doc, "目　录", size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=10, space_after=10)
for item in catalog:
    para(doc, item, size=12, space_after=4)
doc.add_page_break()

# ---- 执行摘要 ----
heading1(doc, "执行摘要")
para(doc, "**核心判断**　如果只看新闻标题，2025年辽宁最显眼的是“GDP破3.32万亿、增长3.7%”、“规上工业利润+45.3%”、“铁路/船舶/航空航天装备+41.8%”和“出口+8.0%”。但这份研究真正值得深读的，是这座“老工业基地+沿海经济带+装备制造”的东北强省，如何在固定资产投资（-19.0%）、房地产开发（-32.3%）、人口自然增（-6.88‰）的背景下，靠“工业利润修复+装备/冶金+出口/沿海”实现3.7%的增长。")
para(doc, "把2025年初设定的目标（GDP增长5%以上）、2025年《国民经济和社会发展统计公报》、以及2026年报告对2025年的复盘放在一起看，辽宁呈现清晰暗线：**从“钢铁/石化/传统重工”的旧底盘，向“装备制造（航天/船舶）+战略性新兴产业+出口/沿海”转型**。旧引擎（汽车、专用设备、地产）在调整；新引擎（航天/船舶装备、冶金、战略新兴产业、出口）被要求更快补位。")
para(doc, "因此，本报告不按政府工作报告逐段复述，而采用“显性表述—同期数据—制度含义—长期影响”的方式，专门提取容易被忽视、但对判断辽宁未来5—10年发展模式更有价值的信号。")
para(doc, "最容易记住的一句话：**辽宁是“老工业基地+沿海经济带+装备制造”的东北样本，靠“工业利润+装备+出口”撑起3.7%增长。**观察辽宁，与其看“GDP 3.32万亿”，不如看“歼-35A、大连造船、冶金/钢铁、造船/航天装备、冰雪6金”这几张名片。")
heading2(doc, "一页速览：2025年辽宁经济的“表与里”")
table(doc,
    ["维度", "表面现象", "更深层信号"],
    [
        ["增长", "GDP 33182.9亿、+3.7%", "一产8.3%、二产33.2%、三产58.4%"],
        ["产业", "规上工业+0.6%", "船舶/航天+41.8%、冶金+8.3%"],
        ["外贸", "进出口7485.6亿、-2.0%", "出口+8.0%、东盟+23.9%"],
        ["投资", "固定资产投资-19.0%", "地产-32.3%、民间-23.4%"],
        ["财政", "地方一般公共预算收入2918.2亿、+0.4%", "税收-0.6%"],
        ["消费", "社零10371.3亿、+1.4%", "电子出版物+26.1%、通讯+23.1%"],
        ["人口", "常住4131万、城镇化74.63%", "自然增-6.88‰、净流入4.5万"],
        ["开放", "沿海港口7.3亿吨/集装箱1368万标箱", "出口/装备/沿海"],
    ],
    widths=[2.2, 5.4, 8.4])
para(doc, "", size=10, space_after=2)

# ---- 一、研究口径与阅读方法 ----
heading1(doc, "一、研究口径与阅读方法")
heading2(doc, "1.1 本报告的三份核心底稿")
bullet(doc, "**2025年《政府工作报告》**：2025年1月在省十四届人大三次会议上作，给出预期目标(GDP 5%以上、规上工业4.5%以上、固投8%左右、社零5%以上、进出口5%左右、财政4%)。")
bullet(doc, "**2025年《国民经济和社会发展统计公报》**：辽宁省统计局2026年3月31日发布，给出全年实际执行数，是“事后验证”锚。")
bullet(doc, "**2026年政府工作报告/复盘**：对2025执行的追认。")
heading2(doc, "1.2 阅读方法：显性—数据—制度—长期")
para(doc, "每章按“**显性表述→同期数据→制度含义→长期影响**”四层展开。")
para(doc, "**关键判别**：数据优先。例如2025年GDP目标5%以上、实际3.7%未达；固投目标8%左右、实际-19.0%大幅未达；进出口目标5%左右、实际-2.0%。差异反映：辽宁“利润/装备/出口强，投资/地产/工业增速弱”。")

# ---- 二、底盘 ----
heading1(doc, "二、先看辽宁的特殊底盘：老工业基地、沿海经济带、装备制造与冰雪")
para(doc, "辽宁的地盘，取决于它作为“**老工业基地+沿海经济带+装备制造强省+冰雪经济**”的特殊定位。它是中国重工业/装备制造的核心。")
bullet(doc, "**装备制造**：铁路/船舶/航空航天+41.8%为最大亮点；歼-35A空军亮相、大连造船、沈大工业母机、全球最大甲醇双燃料船用曲轴；3个国家级先进制造业集群。")
bullet(doc, "**沿海经济带**：港口吞吐量7.3亿吨、集装箱1368.2万标箱（+1.5%）；大连、营口、锦州等沿海港口；水产品540万吨（+4.4%）。")
bullet(doc, "**工业利润修复**：规上工业利润1174.7亿、+45.3%；冶金工业+8.3%（黑色金属矿采选+14.1%、有色冶炼+8.1%）；营业收入利润率改善。")
bullet(doc, "**冰雪/科创**：亚冬会辽宁冰雪健儿6金2银2铜（全国第3）；高铁2387公里（+173公里）；研发投入增速超全国、技术合同+17.8%。")
para(doc, "**制度含义**：辽宁把“装备制造、工业利润、沿海开放、科技创新”当核心资产，深入推进“东北振兴/辽宁全面振兴”。3个国家级先进制造业集群，向“新质生产力”转型。")

# ---- 三、宏观错位 ----
heading1(doc, "三、最关键的宏观错位：GDP破3.32万亿、利润/出口强，但投资/地产/人口弱")
para(doc, "2025年辽宁GDP 33182.9亿元、+3.7%（一产+3.4%、二产+0.7%、三产+5.2%）。表面看“稳中带忧”，但拆开看是“**利润/装备/出口强、工业/投资/地产/人口弱**”的错位：")
para(doc, "**强的部分**：规上工业利润+45.3%、出口+8.0%（东盟+23.9%、民营企业+6.7%）；装备制造（航天/船舶）+41.8%；冶金工业+8.3%；高技术投资+1.6%。")
para(doc, "**弱的部分**：固定资产投资-19.0%（地产-32.3%、民间-23.4%、三产-24.9%）；房地产开发-32.3%；工业增加值+0.6%；人口自然增-6.88‰（常住4131万、60岁+占32.58%）。")
para(doc, "**核心错位一句话**：辽宁“装备/利润/出口强（利润+45.3%、航天+41.8%、出口+8.0%），但投资/地产/人口弱”。2026年“稳投资/地产、强装备/出海、补人口/内需”是重点。")
table(doc,
    ["强引擎", "数据", "弱项", "数据"],
    [
        ["规上工业利润", "+45.3%", "固定资产投资", "-19.0%"],
        ["航天/船舶装备", "+41.8%", "房地产开发投资", "-32.3%"],
        ["冶金工业", "+8.3%", "民间投资", "-23.4%"],
        ["出口", "+8.0%", "人口自然增长率", "-6.88‰"],
        ["高技术制造投资", "+1.6%", "规上工业增速", "+0.6%"],
    ],
    widths=[3.8, 3.0, 3.6, 3.0])
para(doc, "**错位结论**　辽宁的增长“强利润/装备/出口、弱投资/地产/人口”。2026年“稳装备/利润、修复投资、补人口/内需”是重点。")

# ---- 四、被忽视细节 ----
heading1(doc, "四、容易被忽视、但信号很重的细节（15条）")
details = [
    ("1", "GDP 33182.9亿、+3.7%", "总量/增速偏低。"),
    ("2", "规上工业利润+45.3%", "效益大幅改善。"),
    ("3", "航天/船舶/铁路装备+41.8%", "装备制造强。"),
    ("4", "冶金工业+8.3%（铁矿+14.1%）", "钢铁/冶金。"),
    ("5", "高技术制造投资+1.6%", "高技术投资。"),
    ("6", "出口+8.0%、东盟+23.9%", "外贸出口强。"),
    ("7", "民营企业进出口+6.7%、占比52.1%", "民企/外贸。"),
    ("8", "进出口7485.6亿、-2.0%", "外贸-进口弱。"),
    ("9", "歼-35A、大连造船、沈大机床", "国之重器/装备。"),
    ("10", "港口吞吐量7.3亿吨、集装箱+1.5%", "沿海经济带。"),
    ("11", "冰雪6金2银2铜、全国第3", "冰雪/体育。"),
    ("12", "常住4131万、城镇化74.63%、自然增-6.88‰", "人口/老龄化32.58%。"),
    ("13", "居民收入41703元、+4.7%、农村+5.3%", "收入稳、农村快。"),
    ("14", "城镇新增就业48.0万、就业稳", "就业。"),
    ("15", "CPI-0.1%、PPI-4.1%", "通缩/工业品弱。"),
]
for d in details:
    para(doc, "**%s**" % d[0], size=11)
    para(doc, d[1], size=10.5, color=GRAY)

    para(doc, "**备注**　这条“错位”在辽宁尤其鲜明：增长靠利润/装备/出口，但投资/地产/人口弱。2026年若投资/地产修复、装备出海深化，增长可能从“利润/出口单极”走向“利润/装备+投资/内需”多极。这条细节，正是辽宁2026年最难也最值得盯的“变量”。", size=10.5)

# ---- 五、2025年目标 vs 实际 对照 ----
heading1(doc, "五、2025年GDP目标 vs 实际：对照表")
t5 = [
    ["指标", "2025目标", "2025实际", "达标判定"],
    ["GDP增速", "5%以上", "+3.7%", "未达标"],
    ["规模以上工业增速", "4.5%以上", "+0.6%", "未达标"],
    ["固定资产投资增速", "8%左右", "-19.0%", "大幅未达标"],
    ["社会消费品零售增速", "5%以上", "+1.4%", "未达标"],
    ["进出口增速", "5%左右", "-2.0%", "未达标"],
    ["一般公共预算收入增速", "4%", "+0.4%", "未达标"],
    ["居民消费价格(CPI)", "涨幅2%左右", "-0.1%", "远低于"],
]
table(doc, t5[0], t5[1:], widths=[3.4, 3.0, 3.0, 3.6])
para(doc, "**对照结论**　目标“进取”，实际“大缺口”：GDP 3.7%、规上工业0.6%、固投-19.0%、社零1.4%均未达标、财政0.4%（低于4%）。唯“工业利润+45.3%”亮眼。辽宁面临“利润改善但总量/投资/需求偏弱”的转型阵痛。")

# ========= 六、增速分项支撑 =========
heading1(doc, "六、2025年增长是谁撑起来的？（结构归因）")
para(doc, "GDP+3.7%背后，是“**装备/利润/出口强、投资/地产/人口弱**”的结构。把增长贡献拆开看：")
g6 = [
    ["引擎", "贡献/状态", "说明"],
    ["装备制造(航天/船舶)", "+41.8%", "歼-35A、大连造船、沈大机床"],
    ["规上工业利润", "+45.3%", "效益大幅改善"],
    ["冶金工业", "+8.3%", "铁矿+14.1%、有色+8.1%"],
    ["出口", "+8.0%", "东盟+23.9%、民企+6.7%"],
    ["三产/服务业", "+5.2%", "拉动主力"],
    ["高技术制造投资", "+1.6%", "新动能投资"],
    ["规上工业增加值", "+0.6%", "工业总量弱"],
    ["固定资产投资", "-19.0%", "投资大幅拖累"],
    ["房地产开发", "-32.3%", "地产深度调整"],
    ["人口自然增长率", "-6.88‰", "人口/老龄化约束"],
]
table(doc, g6[0], g6[1:], widths=[3.6, 3.6, 6.0])
para(doc, "**一句话**　增长靠“装备/利润/出口+三产”，投资/地产/工业总量是拖累。2026年考验辽宁“装备利润能否转化为投资与总量”。")

# ---- 七、预算与财政 ----
heading1(doc, "七、预算与财政的“含金量”")
para(doc, "2025年辽宁地方一般公共预算收入**2918.2亿元、+0.4%**，税收-0.6%。")
bullet(doc, "财政收入+0.4%（2024年5.5%），增速回落。")
bullet(doc, "政府减税退税降费670亿（2024），支持企业。")
bullet(doc, "支出-3.6%，过紧日子。")

# ---- 八、民生底账 ----
heading1(doc, "八、民生底账：人口、收入与城乡")
para(doc, "2025年辽宁常住人口**4131万人**，城镇化率**74.63%**。自然增长率**-6.88‰**，60岁+占32.58%。")
para(doc, "居民人均可支配收入**41703元、+4.7%**，农村（+5.3%）快于城镇（+4.3%）。")
g8 = [
    ["指标", "数值", "信号"],
    ["常住人口", "4131万", "人口总量大"],
    ["人口自然增长率", "-6.88‰", "自然负增"],
    ["城镇化率", "74.63%", "高城镇化"],
    ["居民人均可支配收入", "41703元/+4.7%", "收入稳、农村快"],
]
table(doc, g8[0], g8[1:], widths=[4.4, 4.0, 4.6])
para(doc, "**民生观察**：收入增速+4.7%不错，但人口自然负增、老龄化（32.58%）是全国最深之一。")

# ---- 九、城镇与农村 ----
heading1(doc, "九、城镇与农村：格局与均衡")
para(doc, "辽宁城镇化率74.63%（较高），农村居民收入快于城镇。")
bullet(doc, "农村居民收入+5.3%、快于城镇+4.3%。")
bullet(doc, "社零城镇+1.5%、乡村+1.0%。")
bullet(doc, "人口向沈阳都市圈/大连都市圈集中。")

# ---- 十、人口流入与流出 ----
heading1(doc, "十、人口流入与流出")
para(doc, "辽宁常住4131万、自然增-6.88‰，但**实现省际净流入+4.5万人**（东北少数人口回流省份）。")
para(doc, "未来看点：装备制造/沿海经济/科创能否持续吸引人口回流。")

# ---- 十一、物价与货币环境 ----
heading1(doc, "十一、物价与货币环境")
para(doc, "2025年辽宁CPI**-0.1%**、PPI**-4.1%**，物价微降、工业品价格承压。")
para(doc, "物价偏弱反映“工业价格下行”，与全国低通胀一致。2026年“扩内需、稳价格”是主线。")

# ---- 十二、区域一体化 ----
heading1(doc, "十二、区域一体化：辽宁在“沿海经济带+东北振兴+沈阳都市圈”里的位置")
para(doc, "辽宁的核心战略坐标是“**沿海经济带+东北振兴+沈阳/大连都市圈**”，是东北振兴与沿海开放的重镇。")
bullet(doc, "沿海经济带：大连、营口、葫芦岛、锦州等沿海港口群。")
bullet(doc, "东北振兴：装备制造、老工业基地转型。")
bullet(doc, "沈阳都市圈：沈大工业母机、科技创新。")
bullet(doc, "冰雪体育：亚冬会6金、冰雪经济。")
para(doc, "若“装备+沿海+科创”成势，辽宁将作为东北振兴与沿海开放的强省。")

# ---------- 十三、五条主线 ----------
heading1(doc, "十三、未来5—10年最值得观察的五条主线")
lines5 = [
    ("① 装备制造/国之重器", "歼-35A、大连造船、沈大机床能否持续。"),
    ("② 工业利润/结构", "利润能否转化为工业总量、投资。"),
    ("③ 沿海经济带/港口", "港口、外贸、海洋经济能否壮大。"),
    ("④ 冰雪/文体旅", "冰雪、文旅、赛后经济。"),
    ("⑤ 人口/内需/投资", "人口回流、内需、地产修复。"),
]
for l in lines5:
    para(doc, "**%s**　%s" % l, space_after=6)

# ---- 十四、结论 ----
heading1(doc, "十四、最终结论：辽宁在“装备制造+沿海+工业利润+冰雪”里的增长逻辑")
para(doc, "辽宁的2025年，本质上是“**装备/利润/出口为核心，而投资/地产/人口弱**”的答卷：GDP33182.9亿、+3.7%，规上工业利润+45.3%、航天/船舶+41.8%、出口+8.0%，但固投-19.0%、地产-32.3%、工业+0.6%、自然增-6.88‰。")
para(doc, "只要装备制造、工业利润、沿海开放、冰雪文旅持续，辽宁就站在“东北振兴重镇”的位置；如果投资/地产/人口持续偏弱，辽宁需承受“利润好、总量/需求弱”的结构挑战。")
para(doc, "最稳观察信号：**一盯装备/装备利润（制造）、二盯沿海/港口（开放）、三盯出口/民企（外贸）、四盯投资/地产（约束）、五盯人口/老龄化（长期）。**辽宁，是“装备+沿海+振兴”的新样本。")

# ---- 附录A ----
heading1(doc, "附录A：主要资料来源与核验路径")
bullet(doc, "辽宁省2025年《政府工作报告》——目标来源。")
bullet(doc, "《辽宁省2025年国民经济和社会发展统计公报》（省统计局，2026-03-31）——GDP、工业、外贸、人口实值。")
bullet(doc, "2026年辽宁省政府工作报告——2025复盘+装备/沿边/冰雪。")
bullet(doc, "大连海关、省财政厅——外贸/财政。")
heading2(doc, "核验说明")
para(doc, "本报告所有增速、总量、占比均以统计公报/官方口径为基准，涉“装备/沿海/冰雪”等以官方为口径。")

# ---- 附录B ----
heading1(doc, "附录B：建议建立的年度跟踪仪表盘")
para(doc, "建议把下面10个“测脉搏”指标做成年度跟踪表：")
dash = [
    ["#", "指标", "2025基值", "为什么重要"],
    ["1", "GDP增速", "+3.7%", "总量与方向"],
    ["2", "规上工业利润增速", "+45.3%", "工业效益"],
    ["3", "装备制造(航天/船舶)", "+41.8%", "装备制造"],
    ["4", "出口/东盟增速", "+8.0%/+23.9%", "外贸"],
    ["5", "固定资产投资/地产", "-19.0%/-32.3%", "投资结构"],
    ["6", "社零增速", "+1.4%", "内需消费"],
    ["7", "常住人口/自然增长率", "4131万/-6.88‰", "人口与城市"],
    ["8", "地方财政收入增速", "+0.4%", "财政质量"],
    ["9", "港口集装箱吞吐量", "1368万标箱", "沿海经济带"],
    ["10", "CPI/PPI", "-0.1%/-4.1%", "物价与需求"],
]
table(doc, dash[0], dash[1:], widths=[0.8, 4.6, 3.4, 4.4])
para(doc, "把这10个指标连起来看，装备/利润/出口（2/3/4）、港口（9）向上、投资/人口（5/7）修复，都说明辽宁在真正换挡。")

# ============ 保存 ============
out = "/Users/x/Desktop/content-prod-lab/reports/辽宁省_2025年政府工作报告_深度研究_2026-08-13.docx"
doc.save(out)
print("SAVED:", out)
print("PARAGRAPHS:", len(doc.paragraphs))
print("TABLES:", len(doc.tables))
