# -*- coding: utf-8 -*-
"""Build 贵阳市2025年政府工作报告 深度研究 DOCX, 参照省会系列版式。"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DARK = RGBColor(0x1F, 0x2A, 0x44)
RED = RGBColor(0xB0, 0x1E, 0x2E)
GRAY = RGBColor(0x59, 0x60, 0x69)
BLUE = RGBColor(0x1F, 0x4E, 0x79)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
HEADER_FILL = "1F2A44"
K = "微软雅黑"


def set_run(run, size=11, bold=False, color=DARK, font=K):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    run.font.name = font
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), font)
    rFonts.set(qn("w:hAnsi"), font)
    rFonts.set(qn("w:eastAsia"), font)


def para(doc, text, size=11, bold=False, color=DARK, align=None,
         space_after=6, space_before=0, font=K):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        set_run(r, size=size, bold=bold or (i % 2 == 1), color=color, font=font)
    return p


def heading1(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(16)
    pf.space_after = Pt(8)
    r = p.add_run(text)
    set_run(r, size=16, bold=True, color=RED)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "B01E2E")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def heading2(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(10)
    pf.space_after = Pt(4)
    r = p.add_run(text)
    set_run(r, size=13, bold=True, color=BLUE)
    return p


def table(doc, headers, rows, widths=None, font_size=9.5):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_run(r, size=font_size, bold=True, color=WHITE)
        tcPr = hdr[i]._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), HEADER_FILL)
        tcPr.append(shd)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(val))
            set_run(r, size=font_size, color=DARK)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    return t


def quote_box(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(8)
    pf.left_indent = Pt(12)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), "B01E2E")
    pBdr.append(left)
    pPr.append(pBdr)
    r = p.add_run(text)
    set_run(r, size=11, color=DARK, font="楷体")
    return p


def bullet(doc, text, size=10.5):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.7)
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        set_run(r, size=size, bold=(i % 2 == 1), color=DARK)
    return p


# ================================================================= 文档主体
doc = Document()
sec = doc.sections[0]
sec.page_width = Cm(21)
sec.page_height = Cm(29.7)
sec.top_margin = Cm(2.4)
sec.bottom_margin = Cm(2.2)
sec.left_margin = Cm(2.5)
sec.right_margin = Cm(2.5)

st = doc.styles["Normal"]
st.font.name = K
st.font.size = Pt(11)
st.element.rPr.rFonts.set(qn("w:eastAsia"), K)


# ---- 封面 ----
para(doc, "贵阳市2025年政府工作报告", size=24, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=60, space_after=4)
para(doc, "深度研究与\u201c容易被忽视的细节\u201d分析", size=16, bold=True, color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
para(doc, "从\u201c大数据之都、磷化工、电子信息制造与科教人才\u201d重新理解贵阳", size=12, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)
para(doc, "\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501", color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
para(doc, "研究对象：2025年贵阳市政府工作报告及同期财政、国民经济和社会发展统计资料、2026年官方复盘", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
para(doc, "标定日期：2026-08-13", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
doc.add_page_break()

# ---- 目录 ----
catalog = [
    "执行摘要",
    "一、研究口径与阅读方法",
    "二、先看贵阳的特殊底盘：大数据之都、磷化工/铝加工、电子信息与人口净流入",
    "三、最关键的宏观错位：GDP破6000亿但增速回落，工业/数据中心强，投资却转负",
    "四、容易被忽视、但信号很重的细节（15条）",
    "五、2025年GDP目标 vs 实际：对照表",
    "六、2025年增长是谁撑起来的？（结构归因）",
    "七、预算与财政的\u201c含金量\u201d",
    "八、民生底账：人口、收入与城乡",
    "九、城镇与农村：格局与均衡",
    "十、人口流入与流出",
    "十一、物价与货币环境",
    "十二、区域一体化：贵阳在\u201c强省会+黔中城市群+\u2018一带一路\u2019大数据走廊\u201d里的位置",
    "十三、未来5—10年最值得观察的五条主线",
    "十四、最终结论：贵阳在\u201c大数据+产业+科教人才\u201d里的增长逻辑",
    "附录A：主要资料来源与核验路径",
    "附录B：建议建立的年度跟踪仪表盘",
]
para(doc, "目　录", size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=10, space_after=10)
for item in catalog:
    para(doc, item, size=12, space_after=4)
doc.add_page_break()

# ---- 执行摘要 ----
heading1(doc, "执行摘要")
para(doc, "**核心判断**　2025年贵阳最显著的是\u201cGDP首破6038.15亿元、增长5.1%、第三产占66.8%\u201d、\u201c规上工业+8.8%但固定资产投资-1.9%\u201d、\u201c外贸进出口92.64亿美元、-2.1%\u201d、\u201c常住人口666.89万、自然增长率-0.12\u2030\u201d。这说明贵阳经济在\u201c产业（电子信息/化学制造/汽车）+数字经济（算力）+文旅\u201d驱动下稳中有进，但投资转负、外贸下滑、人口自然负增长是三大弯道。")
para(doc, "把2025年初设定的目标（GDP 6%左右/规上工业10%左右/固投5%左右/社零5.5%/财政+3%）、2025年统计、2026年前瞻一起看，贵阳呈现\u201c省会型城市\u201d的典型路径：**产业+数字经济+科教人才是引擎，大数据产业位居全国第一梯队**。总量6038亿为全省第2，增速5.1%。")
para(doc, "最容易记住的一句话：**贵阳是\u201c大数据之都+强省会\u201d的贵州省会，靠\u201c产业（电子信息/化工/铝）+数字经济和算力+文旅+科教人才\u201d实现稳增长。**观察贵阳，与其只看\u201cGDP 6038亿\u201d，不如看\u201c规上工业+8.8%、电子信息+48.8%、民间投资+6.6%、数字经济+算力53EFLOPS\u201d。")# ---- 一、研究口径与阅读方法 ----
heading1(doc, "一、研究口径与阅读方法")
para(doc, "本报告以\u201c贵阳市2025年政府工作报告（2025年1月，王宏作）\u201d为起点，把\u201c2025年GDP目标（6%左右/总量破6000亿）\u201d与\u201c2025年国民经济和社会发展统计公报实际值（6038.15亿元/+5.1%）\u201d并置对照，再用2026年政府工作报告的复盘作事后验证。")
para(doc, "重点阅读三类证据：**一是指标目标（计划）是否达标**；**二是结构（谁贡献增长）是否匹配目标叙事**；**三是官方复盘（2026年报告）如何自评、承认问题**。统计口径一般公共预算收入为地方级。")
para(doc, "为便于跨年比较，统一采用\u201c同比增长%\u201d（多数为可比价/不变价），财政与居民收入多为名义值，文中已按原始口径转录、不再逐项换算。")

# ---- 二、先看贵阳的特殊底盘 ----
heading1(doc, "二、先看贵阳的特殊底盘：大数据之都、磷化工/铝加工、电子信息与人口净流入")
para(doc, "**区位与身份**：贵阳是贵州省省会、黔中城市群核心，全国首个国家级大数据综合试验区核心区，被称\u201c中国数谷\u201d\u201c大数据之都\u201d。2025年贵阳贵安一体化推进，贵安新区为全国第八个国家级新区。")
para(doc, "**产业底盘**：一是大数据/数字经济（算力53.12EFLOPS、智算占95.6%、软件信息服务营收破980亿、数据交易额破50亿）；二是磷及磷化工、铝及铝加工（占规上工业比重大，磷化行业+27.5%）；三是电子信息制造业（计算机/通信/电子设备制造业占规上工业7.4%但增速+48.8%）；四是汽车（新能源汽车2025年产量破10万辆、对全省汽车制造业贡献超92%）。")
para(doc, "**人口底盘**：2025年末常住人口666.89万、+1.0%（净增6.64万），城镇化率81.4%；全省第1人口大市，人口规模与\u201c985/211\u201d高校群（贵大、贵州医科大等）构成人才基础。")
para(doc, "**市场与数据**：2025年社会消费品零售总额2576.34亿元/+3.4%；外贸进出口92.64亿美元/-2.1%。内需缓、外需弱是贵阳的结构性短板。")

# ---- 三、最关键的宏观错位 ----
heading1(doc, "三、最关键的宏观错位：GDP破6000亿但增速回落，工业/数据中心强，投资却转负")
para(doc, "**最值得关注的一组错位**：2024年GDP增速6%，2025年目标设为6%左右；2025年实际GDP 6038.15亿元/+5.1%，增速回落约1个百分点且**低于目标**，但**总量从5310亿跳升到6038亿、一举突破6000亿大关**，是\u201c达标值里含金量最高的部分\u201d。")
para(doc, "**第二组错位**：规上工业+8.8%（快于GDP），但固定资产投资**-1.9%（转负）**；社会消费品零售总额+3.4%；外贸进出口**-2.1%**。即在\u201c工业强、投资弱、消费平、外贸降\u201d的格局里，贵阳的增长更依赖存量产业与第三产业。")
para(doc, "**第三组错位**：数字经济/算力高速扩张（智算95.6%）、电子信息+48.8%很亮眼，但这些新兴产业占整个经济盘子的比重仍相对有限，传统行业（烟草27.8%、电力9.7%、非金属矿9.0%）仍是税收与就业基本盘。")
para(doc, "一句话：**贵阳是\u201c产业升级正在进行、但总量仍靠传统+第三产\u201d的省会**——规模越壮大，内部的产业新旧动能切换越需要时间。")

# ---- 四、15条细节 ----
heading1(doc, "四、容易被忽视、但信号很重的细节（15条）")
bullet(doc, "1. **总量破6000亿**：GDP6038.14亿首破6000亿，贵州全省第2，仅次于遵义（但作为省会人均更高）。")
bullet(doc, "2. **增速目标未达标**：2025年GDP目标6%左右，实际+5.1%，低于目标约0.9个百分点，为回落年。")
bullet(doc, "3. **规上工业跑赢GDP**：+8.8%，明显高于GDP的5.1%，第二产仍是动能支撑。")
bullet(doc, "4. **电子信息制造业增速爆炸**：占规上工业7.4%但增加值+48.8%，是贵阳产业升级的最强信号。")
bullet(doc, "5. **化学原料与化学制品+25.4%**：磷化工（及新能源材料）高景气，支撑\u201c富矿精开\u201d。")
bullet(doc, "6. **数字经济\u201c算力-数据-应用\u201d全链**：算力53.12EFLOPS、智算占95.6%、数据交易额破50亿。")
bullet(doc, "7. **民间投资+6.6%**：在固投-1.4%背景下逆势正增长，说明民营经济活跃度尚可。")
bullet(doc, "8. **房地产开发投资+8.7%**：2025年贵阳地产投资正增长，与全国多数城市相悖。")
bullet(doc, "9. **网络零售+56.4%**：限上网络商品零售额同比+56.4%，线上消费爆发。")
bullet(doc, "10. **外贸分化**：出口-12.8%、进口+21.0%，净出口走弱、进口高增。")
bullet(doc, "11. **人口自然负增长**：自然增长率-0.12\u2030，出生率6.85\u2030、死亡率6.96\u2030，人口主要靠机械流入。")
bullet(doc, "12. **收入增速放缓**：城镇/农村人均可支配收入名义增速仅+4.4%/+5.6%，低于2023—2024。")
bullet(doc, "13. **城乡收入比降至1.98**：连续两年缩小，均衡性改善。")
bullet(doc, "14. **旅游人次/花费上行**：旅游总人数+9.7%、总收入+9.2%，避暑/大数据双名片。")
bullet(doc, "15. **财政增收但支出低速**：一般公共预算收入+4.0%、支出+1.1%，“三保”与债务化解压力仍在。")# ---- 五、目标 vs 实际对照表 ----
heading1(doc, "五、2025年GDP目标 vs 实际：对照表")
table(doc,
    ["指标", "2025年目标", "2025年实际", "是否达标"],
    [
        ["地区生产总值增速", "6%左右/破6000亿", "6038.15亿元/+5.1%", "总量达标、增速略低于目标"],
        ["规上工业增加值", "+10%左右", "+8.8%", "接近但略低于目标"],
        ["固定资产投资", "+5%左右", "-1.9%（转负）", "未达标"],
        ["社会消费品零售总额", "+5.5%左右", "2576.34亿元/+3.4%", "低于目标"],
        ["一般公共预算收入", "+3%左右", "490.76亿元/+4.0%", "略超目标"],
        ["城镇/农村人均可支配收入", "+5%/+7%左右", "+4.4%/+5.6%", "均低于目标"],
        ["城镇新增就业", "15万人", "（2025全年约15.9万，同比口径）", "约达标"],
    ],
    widths=[3.6, 2.6, 4.6, 3.2])
para(doc, "**简评**：\u201c总量破6000亿\u201d兑现是最大亮点，但**投资转负、消费/外贸偏弱**，多重指标的增速落在目标值下方，是\u201c量与速\u201d的典型错位年。")

# ---- 六、结构归因 ----
heading1(doc, "六、2025年增长是谁撑起来的？（结构归因）")
para(doc, "**三大产业**：一产219.28亿/+4.7%，二产1783.14亿/+4.9%，三产4035.72亿/+5.2%；三产占GDP 66.8%。增长主要由**三产+二产共同贡献**，其中三产是总量主体。")
para(doc, "**工业内部**：规上工业+8.8%，其中制造业+10.0%、采矿业+6.4%；六大重点产业+15.3%占规上工业55.5%；**电子信息（计算机/通信/电子设备）+48.8%、化学原料及化学制品+25.4%** 是升得快的新动能；烟草（27.8%占规上工业）、电力（9.7%）构成稳定基本盘。")
para(doc, "**投资驱动**：固投-1.9%，但**民间投资+6.6%、高技术产业投资+4.9%、房地产开发投资+8.7%**；基础设施+6.7%是主要正向支撑。")
para(doc, "**总量归因结论**：贵阳的增长来自\u201c**三产（含数字经济+文旅金融）+工业（新能源汽车/电子信息/化工）+民间投资**\u201d，主体是第三产业与制造业转型升级，投资贡献边际走弱。")

# ---- 七、财政 ----
heading1(doc, "七、预算与财政的\u201c含金量\u201d")
para(doc, "2025年财政总收入1027.69亿元/+4.0%，**一般公共预算收入490.76亿元/+4.0%**，一般公共预算支出814.63亿元/+1.1%。收入增速高于GDP名义增速，财税质量尚可。")
para(doc, "**结构性观察**：支出低速增长（+1.1%），反映\u201c三保\u201d与政府主动克制、化债背景下支出让位于化解风险；收入端靠烟草等税源稳定与数字经济、新产业税收增量支撑。")
para(doc, "**含金量**：收入与支出增速差（收+4.0% vs 支+1.1%）意味着结余/平衡空间收窄，财政主要矛盾是\u201c稳收入+控制支出节奏\u201d。")

# ---- 八、民生底账 ----
heading1(doc, "八、民生底账：人口、收入与城乡")
para(doc, "2025年末常住人口**666.89万**（+1.0%、净增6.64万），城镇化率81.4%；出生率6.85\u2030、死亡率6.96\u2030，**自然增长率-0.12\u2030**，人口属\u201c低出生+净流入\u201d结构。")
para(doc, "收入：**城镇52778元/+4.4%、农村26671元/+5.6%**，城乡收入比**1.98**（连续缩小）。城镇新增就业、城镇/农村人均消费支出（城镇38942元、农村21362元）均温和上行。")
para(doc, "**民生结论**：收入增速放缓但农村快于城镇、城乡差距改善；人口净流入依托\u201c大数据产业+大学城+避暑旅居\u201d，但自然负增长是长期慢变量。")

# ---- 九、城乡 ----
heading1(doc, "九、城镇与农村：格局与均衡")
para(doc, "贵阳城镇化率81.4%，为全国省会前列，城区/主城区承载主要服务业与数字经济；农村/县域的\u201c磷化+铝加工+生态农业\u201d提供就业。")
para(doc, "县域经济（开阳磷化工、清镇铝加工等）与主城\u201c大数据+文旅\u201d形成梯度：**总量看城、增长看新动能**。城乡收入比1.98连续缩小，是均衡改善的信号。")

# =========== 十、人口流入流出（并入贵阳线索） =============
heading1(doc, "十、人口流入与流出")
para(doc, "贵阳人口总量上升、自然负增长，增长来自净流入。结合\u201c新增常住人口约13万\u201d（2025年政府工作报告口径）、\u201c高校毕业生留筑超10万\u201d\u201c新增产业人才10万+\u201d，贵阳是贵州及周边省份人口净流入+高校集聚地。")
para(doc, "**流入**：云锦人才（大数据/电子信息）、大学生留筑、避暑旅居人口；**流出**：部分传统制造业/低技能劳动力和公务员编制吸引力相对有限。净流入为正，但强度低于成都、长沙等强省会。")

# =========== 十一、物价与货币 =============
heading1(doc, "十一、物价与货币环境")
para(doc, "2025年贵阳市CPI**上涨0.5%**（核心CPI+1.0%），物价总体温和、处于低位，食品+0.1%、非食品+0.5%，服务+0.9%。环境表现为通胀温和、价格整体稳定。")
para(doc, "信贷环境：全市人民币存款余额约1.6万亿、贷款余额约2.39万亿（银行口径），广义流动性充裕，但固定资产投资的转负更多源于项目与产业预期而非货币短缺。")

# =========== 十二、区域一体化 =============
heading1(doc, "十二、区域一体化：贵阳在\u201c强省会+黔中城市群+\u2018一带一路\u2019大数据走廊\u201d里的位置")
para(doc, "贵阳作为**贵州省会+黔中城市群核心**，是贵州省\u201c强省会\u201d战略、贵阳贵安一体化（贵安新区全国第八个国家级新区）的载体。大数据产业（茅台、华为云、车路协同）使贵阳成为\u201c东数西算\u201d国家战略枢纽节点（算力53EFLOPS、智算全国前列）。")
para(doc, "对外开放上，贵阳地处西南，依托贵南高铁、中欧/加密国际班列与大湄公河次区域，正推进\u201c一带一路\u201d铁路、跨境数据与跨境电商（外贸进出口92.6亿美元）。区域位置是\u201c出海偏弱、但大数据+算力全国链接强\u201d的特殊形态。")

# =========== 十三、五条主线 =============
heading1(doc, "十三、未来5—10年最值得观察的五条主线")
bullet(doc, "1. **数字经济与算力**：算力53EFLOPS→190EFLOPS（2026规划），能否从\u201c算力规模\u201d变成\u201c应用+数据交易+软件营收\u201d的产业收益，决定贵阳能否跳出\u201c看数据不看产业\u201d。")
bullet(doc, "2. **新能源汽车+电子信息制造**：新能源汽车产量10万→25万辆、电子信息+48.8%，是贵阳工业换挡的主升浪。")
bullet(doc, "3. **磷化/铝加工“富矿精开”与新能源材料**：磷化、铝都、新能源电池材料价值链上移。")
bullet(doc, "4. **强省会+贵安一体化+黔中城市群**：地铁/高铁/科创城、避暑旅居与人口集聚，决定省会吸附力。")
bullet(doc, "5. **人口与少子化**：自然增长率-0.12\u2030下，能否维持净流入、撑起消费与地产，是长期变量。")

# =========== 十四、最终结论 =============
heading1(doc, "十四、最终结论：贵阳在\u201c大数据+产业+科教人才\u201d里的增长逻辑")
para(doc, "**结论**：贵阳2025年的\u201c真相\u201d是——**总量破6000亿、增速5.1%（低于目标）、工业/电子信息/数字经济亮眼，但投资转负、外贸-2.1%、人口自然负增长**。它正从\u201c投资+基建+烟草\u201d驱动，切换到\u201c产业（电子信息/新能源+化工/汽车）+数字经济+第三产业（文旅/新兴服务）+科教人才\u201d驱动的换挡期。")
para(doc, "**对趋势判断的关键**：经济的代表量是\u201c若干新动能（电子信息、民间投资）\u201d而非传统量（固投整体）。投资转负是与全国一致的主动去杠杆+房地产调整，不能据此看空贵阳；反而**民间投资+6.6%、电子信息+48.8%、数字经济算力、大学城+人才**是中期的结构性亮点。")
para(doc, "**若只看一个指标**：看**规上工业中的电子信息+48.8%这一增速**——它比GDP更能说明贵阳的新动能真实成色。贵阳不再是\u201c只靠大数据概念\u201d，而是把算力、数字产业、电子信息、磷化/铝、新能源材料真正落到工业增加值与居民收入上。")

# ---------------- 附录A ----------------
heading1(doc, "附录A：主要资料来源与核验路径")
bullet(doc, "贵阳市人民政府《2025年贵阳市政府工作报告》（2025年1月）。")
bullet(doc, "贵阳市统计局《2025年贵阳市国民经济和社会发展统计公报》（2026年5月）。")
bullet(doc, "《2026年贵阳市人民政府工作报告》（2026年2月）及\u201c数读\u201d解读。")
bullet(doc, "中国国民经济核算与贵州统计年鉴、贵阳市统计公报交叉核验。")

# ---------------- 附录B ----------------
heading1(doc, "附录B：建议建立的年度跟踪仪表盘")
bullet(doc, "GDP总量/增速 vs 全年目标、财报对比（+/-个百分点）。")
bullet(doc, "规上工业增加值，及其分行业（电子信息/化学制造/汽车/烟草/电力）增速。")
bullet(doc, "固定资产投资（总量/民间/基建/工业/房地产）增速。")
bullet(doc, "社会消费品零售总额与2026年目标（城市限额以上）。")
bullet(doc, "贸易进出口总额（美元口径）与2026年目标。")
bullet(doc, "一般公共预算收入/支出增速、财政自给率。")
bullet(doc, "常住人口增量、自然增长率、城镇化率、高校留筑人数。")
bullet(doc, "CPI、规上工业企业利润总额与营收。")
bullet(doc, "旅游人数/旅游总收入、避暑旅居人数。")
bullet(doc, "数据中心算力（EFLOPS）、数据交易额、数字经济核心产业占GDP比重。")

# ---------------- 保存 ----------------
out = "/Users/x/Desktop/content-prod-lab/reports/贵阳市_2025年政府工作报告_深度研究_2026-08-13.docx"
doc.save(out)
print("SAVED OK 贵阳市_2025年政府工作报告_深度研究_2026-08-13.docx")