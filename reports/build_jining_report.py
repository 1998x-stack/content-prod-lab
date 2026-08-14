# -*- coding: utf-8 -*-
"""Build 济宁市2025年政府工作报告 深度研究 DOCX, 参照地级市系列版式。"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DARK = RGBColor(0x1F, 0x2A, 0x44)
RED = RGBColor(0xB0, 0x1E, 0x2E)
GRAY = RGBColor(0x59, 0x60, 0x69)
BLUE = RGBColor(0x1F, 0x4E, 0x79)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
HEADER_FILL = "1F2A44"
K = "微软雅黑"


def set_run(run, size=11, bold=False, color=DARK, font=K):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    run.font.name = font
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), font)
    rFonts.set(qn("w:hAnsi"), font)
    rFonts.set(qn("w:eastAsia"), font)


def para(doc, text, size=11, bold=False, color=DARK, align=None,
         space_after=6, space_before=0, font=K):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        set_run(r, size=size, bold=bold or (i % 2 == 1), color=color, font=font)
    return p


def heading1(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(16)
    pf.space_after = Pt(8)
    r = p.add_run(text)
    set_run(r, size=16, bold=True, color=RED)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "B01E2E")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def heading2(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(10)
    pf.space_after = Pt(4)
    r = p.add_run(text)
    set_run(r, size=13, bold=True, color=BLUE)
    return p


def table(doc, headers, rows, widths=None, font_size=9.5):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_run(r, size=font_size, bold=True, color=WHITE)
        tcPr = hdr[i]._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), HEADER_FILL)
        tcPr.append(shd)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(val))
            set_run(r, size=font_size, color=DARK)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    return t


def quote_box(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(8)
    pf.left_indent = Pt(12)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), "B01E2E")
    pBdr.append(left)
    pPr.append(pBdr)
    r = p.add_run(text)
    set_run(r, size=11, color=DARK, font="楷体")
    return p


def bullet(doc, text, size=10.5):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.7)
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        set_run(r, size=size, bold=(i % 2 == 1), color=DARK)
    return p


# ================================================================= 文档主体
doc = Document()
sec = doc.sections[0]
sec.page_width = Cm(21)
sec.page_height = Cm(29.7)
sec.top_margin = Cm(2.4)
sec.bottom_margin = Cm(2.2)
sec.left_margin = Cm(2.5)
sec.right_margin = Cm(2.5)

st = doc.styles["Normal"]
st.font.name = K
st.font.size = Pt(11)
st.element.rPr.rFonts.set(qn("w:eastAsia"), K)


# ---- 封面 ----
para(doc, "济宁市2025年政府工作报告", size=24, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=60, space_after=4)
para(doc, "深度研究与\u201c容易被忽视的细节\u201d分析", size=16, bold=True, color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
para(doc, "从\u201c孔孟之乡（曲阜）、京杭运河港、煤炭、装备制造、文旅\u201d重新理解济宁", size=12, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)
para(doc, "\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501", color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
para(doc, "研究对象：2025年济宁市政府工作报告及同期财政、国民经济和社会发展统计资料、2026年官方复盘", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
para(doc, "标定日期：2026-08-14", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
doc.add_page_break()

# ---- 目录 ----
catalog = [
    "执行摘要",
    "一、研究口径与阅读方法",
    "二、先看济宁的特殊底盘：孔孟之乡、水运煤、装备制造、港航与文旅",
    "三、最关键的宏观错位：GDP 6128亿/5.8%达成，工业外贸强但固投弱、地产、煤炭调",
    "四、容易被忽视、但信号很重的细节（15条）",
    "五、2025年GDP目标 vs 实际：对照表",
    "六、2025年增长是谁撑起来的？（结构归因）",
    "七、预算与财政的\u201c含金量\u201d",
    "八、民生底账：人口、收入与城乡",
    "九、城镇与农村：格局与均衡",
    "十、人口流入与流出",
    "十一、物价与货币环境",
    "十二、区域一体化：济宁在\u201c淮海经济区+山东都市圈+运河经济带\u201d里的位置",
    "十三、未来5—10年最值得观察的五条主线",
    "十四、最终结论：济宁在\u201c煤炭转型+装备制造+孔孟文旅\u201d里的增长逻辑",
    "附录A：主要资料来源与核验路径",
    "附录B：建议建立的年度跟踪仪表盘",
]
para(doc, "目　录", size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=10, space_after=10)
for item in catalog:
    para(doc, item, size=12, space_after=4)
doc.add_page_break()

# ---- 执行摘要 ----
heading1(doc, "执行摘要")
para(doc, "**核心判断**　2025年济宁最显著的是\u201cGDP 6128亿元、增长5.8%（达5.5%目标、山东第4\u201d、\u201c规上工业+7.9%（制造业+10.4%、新能源链+59.7%）\u201d、\u201c文化（孔孟·曲阜）、港航（运河亿吨大港）、煤炭\u201d、\u201c但固投+0.3%弱、CPI-0.2%、财收+1.3%、进出口+7.4%\u201d、\u201c常住812.93万/城镇化65.37%\u201d。这说明济宁在\u201c煤炭+制造+港航+文旅\u201d中，**工业外贸强、固投弱、地产煤炭调**。")
para(doc, "把2025年目标（GDP+5.5%以上/财收+3.5%左右）、2025年统计（GDP+5.8%达成、规上+7.9%、固投+0.3%、财收+1.3%、社零+6.2%超额）、趋势一起看，济宁是\u201c孔孟之乡（文化）+港航+制造\u201d路径：**煤炭（水运）、装备制造、新能源（电池）、港航物流（1.16亿吨）、文旅（孔孟/微山湖）**是支柱；2025年总量6128亿居山东第4（次于青岛/济南/烟台）。")
para(doc, "最容易记住的一句话：**济宁是\u201c孔孟之乡（曲阜三孔）、北方内河亿吨港、煤炭转型制造\u201d，靠\u201c港航+制造+文旅\u201d增长。**观察济宁，与其只看\u201cGDP 6128亿\u201d，不如看\u201c规上+7.9%、煤港吞吐1.16亿吨（亿吨大港）、外贸1274亿/+8.5%、新能源+59.7%、孔孟文旅950亿\u201d。")

# ---- 一、研究口径与阅读方法 ----
heading1(doc, "一、研究口径与阅读方法")
para(doc, "本报告以\u201c济宁市政府工作报告（2025年，张海波作）\u201d为起点，把\u201c2025年GDP目标（5.5%以上）\u201d与\u201c官方2025年GDP（6128亿元/+5.8%）\u201d并置，并用\u201c2025年济宁市统计公报\u201d和\u201c2026年政府工作报告\u201d横向核验。")
para(doc, "口径提示：\u201cGDP、规上工业增速\u201d按可比价（实际）；\u201c投资、消费、进出口、财政\u201d按现价（名义）。\u201c常住人口\u201d用统计公报（812.93万）、城镇化率65.37%。")
para(doc, "指标体系（与研究口径一致）：五大象限——**总量与增速（GDP）、产业动能（煤炭/制造/港航/文旅）、投资、财政、民生与人口**。")
para(doc, "特别提示（不吃老本）：济宁是\u201c孔孟之乡（曲阜三孔）、运河、煤炭\u201d之地，2024年GDP约5800亿/+6.2%；2025年+5.8%略放缓到6%内，但**港口亿吨、新能源链+59.7%、装备制造**是新一轮增长——真正要看\u201c煤炭转型+亿吨大港+文化品牌\u201d。")
# ---- 二、先看济宁的特殊底盘 ----
heading1(doc, "二、先看济宁的特殊底盘：孔孟之乡、水运煤、装备制造、亿吨大港与文旅")
para(doc, "济宁地处山东西南部、孔孟故里，是**孔孟之乡（曲阜三孔）、京杭运河重镇、煤炭/装备制造城市**。2025年GDP 6128亿元、常住812.93万，山东第4，人均约7.5万元。")
para(doc, "五个底盘名词，先立框架：")
bullet(doc, "**孔孟之乡/文旅**　曲阜三孔（孔子故里）、微山湖、水浒，文旅收入950亿、2025游客破1亿。")
bullet(doc, "**煤炭（+水运）**　济宁是山东煤炭重镇，原煤外运、坑口/煤化工。")
bullet(doc, "**装备制造**　山推（推土机全球）、鲁抗、机械装备，装备新兴产业强。")
bullet(doc, "**京杭运河亿吨大港**　港航营收破1000亿、货物吞吐1.16亿吨（北方内河首个亿吨大港）、集装箱52万标箱。")
bullet(doc, "**新能源（电池/氢能）**　新能源产业链+59.7%、电池/储能、新一代信息技术+23%。")
para(doc, "这五根（文化+煤+装备+港航+新能源）构成济宁独特底盘：**左右孔孟（黄河文化）、右手运河亿吨港（物流/煤炭）**。理解济宁，先理解\u201c孔孟文化+煤运\u201d。")

# ---- 三、最关键的宏观错位 ----
heading1(doc, "三、最关键的宏观错位：GDP 6128亿/5.8%达成，工业外贸强但固投弱、地产、煤炭调")
para(doc, "2025年济宁最需要辨析的一组\u201c错位\u201d：**GDP 5.8%（达5.5%目标）、规上+7.9%（制造+10.4%）强但固投+0.3%、CPI-0.2%、地产/煤运调、财收+1.3%**。")
para(doc, "为什么\u201c工业这么强、外贸/消费不错\u201d，投资与财政却平淡？三个解释：")
para(doc, "**其一，工业强但在地产/煤炭调整下**　制造业+10.4%（三产快），但煤炭、地产投资弱；\u201c制造强、煤炭转型\u201d。")
para(doc, "**其二，内需韧性但物价负**　社零+6.2%（好）、以旧换新+居，但CPI-0.2%（食品/交通降）、\u201c供稳需弱\u201d。")
para(doc, "**其三，财政（+1.3%）低于目标（3.5%）**　税收-1.7%、占60.9%；地产/煤炭税调、靠非税/国企。")
para(doc, "小结：济宁2025年是\u201c**强工业港航文旅、弱固投/煤/财政**\u201d的一年：制造、港、文化撑增长，煤炭、地产、投资弱。")

# ---- 四、容易被忽视、但信号很重的细节（15条） ----
heading1(doc, "四、容易被忽视、但信号很重的细节（15条）")
bullet(doc, "**1.规上工业+7.9%、制造业+10.4%（近三年新高）**\u201c装备、新能源链、食品\u201d强。\u201d")
bullet(doc, "**2.新能源产业链+59.7、新一代信息技术+23%**\u201c电池/储能/氢能\u201d新动力。\u201d")
bullet(doc, "**3.港航业营收破1000亿（亿吨大港）**\u201c京杭运河北方内河第一港、吞吐1.16亿吨。\u201d")
bullet(doc, "**4.煤炭/运维\u201c双吨\u201d**（原煤、煤化工）基本盘。")
bullet(doc, "**5.进出口1274.4亿/+7.4%、出口+7%**\u201c机电/装备/（碳中和/机械）\u201d；对东盟/巴西高增。\u201d")
bullet(doc, "**6.社零3233.6亿/+6.2%/以旧换新+（居民消费强）**")
bullet(doc, "**7.\u201c232\u201d集群突破4000亿、规上工业527家**\u201c产业/企业梯队扩容。\u201d")
bullet(doc, "**8.文旅游客破1亿/入境翻番、收入950亿**\u201c孔孟·微山湖文旅爆红。\u201d")
bullet(doc, "**9.财收502.7亿/+1.3%、税收-1.7%（占60.9%）**\u201c税收承压（煤炭/地产）、民生80.8%。\u201d")
bullet(doc, "**10.居民收入39778元/+4.9%、城乡比2.23**\u201c农村+5.8%>城镇+4.2%。\u201d")
bullet(doc, "**11.常住812.93万/城镇化65.37%（+1.24pct）**\u201c人口大市（山东第3）、自然增负。\u201d")
bullet(doc, "**12.CPI-0.2%（食品-0.9%/交通-）**\u201c低通胀、需求偏弱。\u201d")
bullet(doc, "**13.粮食超100亿斤（山东粮仓）**、高标准农田/良种繁育")
bullet(doc, "**14.制造业单项冠军/专精特新1378家**\u201c制造底座厚。\u201d")
bullet(doc, "**15.雄商高铁通车、7条高速在建**\u201c枢纽——交通/物流带。\u201d")
# ---- 五、2025年GDP目标 vs 实际：对照表 ----
heading1(doc, "五、2025年GDP目标 vs 实际：对照表")
table(doc,
    ["指标", "2025年目标", "2025年实际", "达成评价"],
    [
        ["地区生产总值(GDP)", "增长5.5%以上", "6128亿/5.8%", "达成"],
        ["规上工业增加值", "增长7.5%左右", "+7.9%", "超额"],
        ["固定资产投资", "——", "+0.3%", "偏低"],
        ["社会消费品零售总额", "——", "3233.6亿/+6.2%", "稳健"],
        ["进出口总额", "——", "1274.4亿/+7.4%", "稳健"],
        ["一般公共预算收入", "增长3.5%左右", "502.7亿/+1.3%", "未达"],
        ["居民收入", "与经济增长同步", "39778元/+4.9%", "略低GDP"],
    ],
)
para(doc, "注：GDP、规上工业按可比价。**GDP（+5.8%）、规上（+7.9%）、社零（+6.2%）达成/超额**；**财收（+1.3%）未达3.5%目标（税收-1.7%）、固投（+0.3%）弱**。")
para(doc, "拆读：**规上（+7.9%/制造+10.4%）、港航（亿吨大港）、文旅（1亿人次/950亿）、外贸（+7.4%）是亮色**；**固投（+0.3%）、CPI（-0.2%）、财收（+1.3%）偏弱**——\u201c制造港航强、投资财政弱\u201d，是孔孟煤炭转型样本。")

# ---- 六、2025年增长是谁撑起来的？（结构归因） ----
heading1(doc, "六、2025年增长是谁撑起来的？（结构归因）")
para(doc, "把济宁GDP的5.8%拆开：第一产业、第二产业、第三产业，其中制造业+10.4%（工业升级）、港航文旅（三产）。**第二产业（制造）与第三产业（港航/文旅/服务）双撑，第一产业（农业）稳**。")
para(doc, "2026年济宁强调\u201c工业经济头号工程\u201d、新能源、港航、文旅，聚焦**装备、新能源电池、煤炭转型、亿吨港、孔孟文旅**——核心是\u201c制造强市+运河经济带\u201d。")
para(doc, "**第二产业（工业）**：规上+7.9%（制造业+10.4%）、新能源链+59.7%、新一代信息+23%、232集群4000亿——\u201c制造强\u201d。")
para(doc, "**第三产业（服务业）**：港航（营收破千亿）、物流（亿吨港）、文旅（1亿人次/950亿）、数字经济——\u201c港航+文旅强\u201d。")
para(doc, "**外贸（开放）**：进出口1274.4亿/+7.4%（对东盟/巴西高增）、中欧班列——\u201c外向稳\u201d。")
para(doc, "一句话归因：**2025年济宁增长\u201c靠工业（装备/新能源）+港航+文旅\u201d**，煤炭、地产、投资弱；\u201c港航文旅\u201d是新引擎。")

# ---- 七、预算与财政的\u201c含金量\u201d ----
heading1(doc, "七、预算与财政的\u201c含金量\u201d")
para(doc, "2025年济宁**一般公共预算收入502.7亿元（+1.3%）**，其中税收305.9亿元（-1.7%）、税收占比60.9%；支出843.4亿元、民生占80.8%。")
bullet(doc, "税收结构：税收-1.7%（煤炭/地产/石化税）、占60.9%；财收靠非税/国企上缴——\u201c税收承压、含金量需看\u201d。")
bullet(doc, "民生支出占80.8%，社保/教育/医疗倾斜。")
bullet(doc, "金融支撑：存款破万亿、贷款+10.6%、不良率0.86%（全省低）——信贷宽、资产质量好。")
para(doc, "**财政含金量小结**：财收+1.3%（低目标）、税收-1.7%；\u201c税收弱、民生高\u201d；财政对\u201c制造、港航、文旅\u201d投入加大。")

# ---- 八、民生底账：人口、收入与城乡 ----
heading1(doc, "八、民生底账：人口、收入与城乡")
para(doc, "2025年济宁**居民人均可支配收入39778元（+4.9%）**，其中城镇49825元（+4.2%）、农村30320元左右（+5.2%），城乡比约2.1、民生投入80.8%。就业：城镇新增就业（十四五33.8万）。")
para(doc, "人口画像：**常住812.93万、城镇化65.37%（+1.24pct）**，山东人口第3大市；人口自然增长小幅负、老龄化。")
para(doc, "民生投入：城乡低保提、集采降价、养老床位6.8万张、教卫新改扩建293所——民生扎实。")

# ---- 九、城镇与农村：格局与均衡 ----
heading1(doc, "九、城镇与农村：格局与均衡")
para(doc, "济宁常住城镇化率65.37%（山东中上），城乡较均衡；农村收入增速（+5%+）高于城镇（+4.2%）——**城乡差收敛**。")
para(doc, "农业底盘：**粮食超100亿斤（山东粮仓）**、高标准农田、良种繁育70万亩——\u201c济宁粮仓+县域（金乡大蒜/微山湖渔）\u201d。")
para(doc, "一句话：\u201c济宁农业粮仓、农村收入快、县域特色（大蒜/水乡）\u201d，\u201c以港带城、以城带乡\u201d。")

# ---- 十、人口流入与流出 ----
heading1(doc, "十、人口流入与流出")
para(doc, "济宁常住812.93万、城镇化65.37%，是山东人口第3（（山东人口第二大市之一），因孔孟/济宁都市圈、制造/煤运/文旅留人；但自然增长负（出生<死亡）、少子化老龄化。")
para(doc, "结构观察：**高校（济宁医学院/学院等）、制造/文旅岗位**聚人，城镇化率还低于全省平均（约75%）。")
para(doc, "2026年目标：青年人才5万、留济引才——济宁靠\u201c港航+装备+文旅\u201d聚人。")

# ---- 十一、物价与货币环境 ----
heading1(doc, "十一、物价与货币环境")
para(doc, "2025年济宁**CPI-0.2%**（食品烟酒-0.9%、交通通信-衣着+1%）——**低通胀、需求偏弱**。")
bullet(doc, "信贷：存款破万亿、贷款+10.6%（制造业/基建）、不良贷款低（0.86%）——宽信用、质量好。")
bullet(doc, "消费：社零+6.2%、以旧换新政策券——\u201c促内需\u201d。")
para(doc, "货币环境判断：**宽信用、CPI-0.2%低通胀**；济宁靠\u201c以旧换新、港航、文旅\u201d稳需求（2026 CPI低于1）。")

# ---- 十二、区域一体化：济宁的位置 ----
heading1(doc, "十二、区域一体化：济宁在\u201c淮海经济区+山东都市圈+运河经济带\u201d里的位置")
para(doc, "济宁地处\u201c淮海经济区核心、山东南半、西朗都市圈、京杭运河北部枢纽\u201d，是\u201c运河经济带+孔孟文化带\u201d交汇。")
bullet(doc, "**淮海协同**　济宁-临沂-徐州（淮海）、徐州都市圈，装备/物流协作。")
bullet(doc, "**运河枢纽/亿吨大港**　京杭运河（济宁-微山湖），吞吐1.16亿吨、面向华东/长江。")
bullet(doc, "**文化带**　孔孟（曲阜）、水浒、微山湖——黄河南北文化走廊。")
para(doc, "一句话：**济宁在\u201c淮海+运河\u201d中，最核心是\u201c亿吨内河港+孔孟文化\u201d**；区位（运河）、港口是大优势。")

# ---- 十三、未来5—10年最值得观察的五条主线 ----
heading1(doc, "十三、未来5—10年最值得观察的五条主线")
bullet(doc, "**主线一：亿吨大港/运河经济带**\u201c吞吐1.16亿吨、集装箱52万\u201d能否成\u201c北方内河枢纽\u201d。")
bullet(doc, "**主线二：制造（装备/新能源链+59.7%）**\u201c山推、电池/氢能\u201d能否成\u201c新质\u201d。")
bullet(doc, "**主线三：煤炭转型/能源**\u201c煤矿\u201d向\u201c光伏、新能源\u201d转（装机800万目标）。")
bullet(doc, "**主线四：孔孟文旅（1亿人次）**\u201c三孔、微山湖\u201d能否高端/入境。")
bullet(doc, "**主线五：人口与县域**“813万人口、运河新城”“金乡/兖州县域”“港产城融合”。")

# ---- 十四、最终结论 ----
heading1(doc, "十四、最终结论：济宁在\u201c煤炭转型+装备制造+孔孟文旅\u201d里的增长逻辑")
para(doc, "把2025年闭环打穿：**济宁是\u201c孔孟之乡、亿吨大港、制造强\u201d**：GDP 6128亿/+5.8%、规上+7.9%（制造+10.4%）、港航破千亿、文旅950亿、外贸+7.4%。")
para(doc, "济宁不是\u201c只有煤（孔孟）\u201d——它是**煤炭+装备+港航+新能源+文旅**的复合，靠\u201c制造、港、文化\u201d驱动；但煤炭、地产、固投、财政弱。")
para(doc, "一句话结论：**济宁是\u201c孔孟之乡、碧水运河（吨港）、装备文旅\u201d 观察它先看\u201c亿吨大港、装备/新能源、孔孟文旅、外贸\u201d，再看\u201c煤炭、固投、财政、人口\u201d。**它是\u201c港航文旅强、煤炭转型\u201d的山东样本。")

# ---- 附录A ----
heading1(doc, "附录A：主要资料来源与核验路径")
bullet(doc, "《2025年济宁市政府工作报告》（2025年2月，2025年目标、2024年回顾+6.2%）")
bullet(doc, "《2025年济宁市国民经济和社会发展统计公报》（济宁市统计局，2026-04-01 PDF，2025年实际）")
bullet(doc, "《2026年济宁市政府工作报告》（2026年2月，2025年复盘+2026年目标）")
bullet(doc, "济宁市/统计局（jining.gov.cn/tjj）")

# ---- 附录B ----
heading1(doc, "附录B：建议建立的年度跟踪仪表盘")
bullet(doc, "GDP总量/增速 vs 全年目标。")
bullet(doc, "规上工业（装备/新能源链）增速。")
bullet(doc, "港航吞吐/集装箱/营收。")
bullet(doc, "煤炭产量/煤化工。")
bullet(doc, "固定资产/工业/房地产投资。")
bullet(doc, "社零、以旧换新。")
bullet(doc, "进出口、中欧班列。")
bullet(doc, "财收/税收、民生%。")
bullet(doc, "常住/城镇化、高校。")
bullet(doc, "CPI、存款/贷款。")

# ---------------- 保存 ----------------
out = "/Users/x/Desktop/content-prod-lab/reports/济宁市_2025年政府工作报告_深度研究_2026-08-14.docx"
doc.save(out)
print("SAVED OK: 济宁市", out)
