# -*- coding: utf-8 -*-
"""Build 重庆市2025年政府工作报告 深度研究 DOCX, 参照北京/上海/杭州/天津系列版式。"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DARK = RGBColor(0x1F, 0x2A, 0x44)
RED = RGBColor(0xB0, 0x1E, 0x2E)
GRAY = RGBColor(0x59, 0x60, 0x69)
BLUE = RGBColor(0x1F, 0x4E, 0x79)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
HEADER_FILL = "1F2A44"
K = "微软雅黑"


def set_run(run, size=11, bold=False, color=DARK, font=K):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    run.font.name = font
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), font)
    rFonts.set(qn("w:hAnsi"), font)
    rFonts.set(qn("w:eastAsia"), font)


def para(doc, text, size=11, bold=False, color=DARK, align=None,
         space_after=6, space_before=0, font=K):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        set_run(r, size=size, bold=bold or (i % 2 == 1), color=color, font=font)
    return p


def heading1(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(16)
    pf.space_after = Pt(8)
    r = p.add_run(text)
    set_run(r, size=16, bold=True, color=RED)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "B01E2E")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def heading2(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(10)
    pf.space_after = Pt(4)
    r = p.add_run(text)
    set_run(r, size=13, bold=True, color=BLUE)
    return p


def table(doc, headers, rows, widths=None, font_size=9.5):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_run(r, size=font_size, bold=True, color=WHITE)
        tcPr = hdr[i]._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), HEADER_FILL)
        tcPr.append(shd)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(val))
            set_run(r, size=font_size, color=DARK)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    return t


def quote_box(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(8)
    pf.left_indent = Pt(12)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), "B01E2E")
    pBdr.append(left)
    pPr.append(pBdr)
    r = p.add_run(text)
    set_run(r, size=11, color=DARK, font="楷体")
    return p


def bullet(doc, text, size=10.5):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.7)
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        set_run(r, size=size, bold=(i % 2 == 1), color=DARK)
    return p


# ================================================================= 文档主体
doc = Document()
sec = doc.sections[0]
sec.page_width = Cm(21)
sec.page_height = Cm(29.7)
sec.top_margin = Cm(2.4)
sec.bottom_margin = Cm(2.2)
sec.left_margin = Cm(2.5)
sec.right_margin = Cm(2.5)

st = doc.styles["Normal"]
st.font.name = K
st.font.size = Pt(11)
st.element.rPr.rFonts.set(qn("w:eastAsia"), K)


# ---- 封面 ----
para(doc, "山东省2025年政府工作报告", size=24, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=60, space_after=4)
para(doc, "深度研究与“容易被忽视的细节”分析", size=16, bold=True, color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
para(doc, "从“北方经济大省、海洋强省、新旧动能转换与乡村振兴”重新理解山东", size=12, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)
para(doc, "━━━━━━━━━━━━━━━━", color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
para(doc, "研究对象：2025年山东省政府工作报告及同期财政、国民经济和社会发展统计资料、2026年官方复盘", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
para(doc, "整理时间：2026年8月", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
para(doc, "分析性质：分析性研究 ｜ 非官方解读", size=11, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
doc.add_page_break()

# ---- 目录 ----
catalog = [
    "执行摘要",
    "一、研究口径与阅读方法",
    "二、先看山东的特殊底盘：北方经济第一大省、海洋强省与乡村振兴大省",
    "三、最关键的宏观错位：GDP破10万亿、工业强，但投资与进口偏弱",
    "四、容易被忽视、但信号很重的细节（15条）",
    "五、把所有细节连起来：山东正在经历的“六个换挡”",
    "六、增长暗线：新旧动能转换、绿色低碳与新质生产力",
    "七、财政暗线：收入稳、税收平，民生与转型平衡",
    "八、产业暗线：从“重化工业”到“十强产业+高端制造”",
    "九、区域格局：山东半岛城市群与海洋强省",
    "十、人口与城市：过亿人口大省、城镇化与老龄化",
    "十一、民营经济：占七成多进出，市场主体活跃",
    "十二、事后验证：用2025年实际执行结果检验报告判断",
    "十三、未来5—10年最值得观察的五条主线",
    "十四、最终结论：山东在“新旧动能转换+绿色低碳”里的增长逻辑",
    "附录A：主要资料来源与核验路径",
    "附录B：建议建立的年度跟踪仪表盘",
]
para(doc, "目　录", size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=10, space_after=10)
for item in catalog:
    para(doc, item, size=12, space_after=4)
doc.add_page_break()

# ---- 执行摘要 ----
heading1(doc, "执行摘要")
para(doc, "**核心判断**　如果只看新闻标题，2025年山东最显眼的是“GDP突破10.32万亿、成为全国第三个/北方第一个过10万亿的省份”、“规上工业增长7.6%”和“新能源汽车产量110万辆、+50.3%”。但政府工作报告里真正值得深读的，是这一座“北方经济第一大省+海洋强省+农业大省”如何在固定资产投资（-8.6%）与制造业投资偏弱的背景下，靠“新旧动能转换、绿色低碳与新质生产力”稳住增长。")
para(doc, "把2025年初《政府工作报告》设定的目标、2025年《国民经济和社会发展统计公报》、以及2026年报告对2025年的复盘放在一起看，山东呈现清晰暗线：**从“重化工业+基建”的旧底盘，向“新旧动能转换+绿色低碳+高端制造（新质生产力）”转型**。旧引擎（钢铁、化工、房地产、传统基建）在调整；新引擎（装备制造、新能源、锂电、高端制造、海洋经济、民营）被要求更快补位。")
para(doc, "因此，本报告不按政府工作报告逐段复述，而采用“显性表述—同期数据—制度含义—长期影响”的方式，专门提取容易被忽视、但对判断山东未来5—10年发展模式更有价值的信号。")
para(doc, "最容易记住的一句话：**山东是全国“新旧动能转换+绿色低碳”最重要的试验田之一，也是北方经济的“压舱石”——它既要用“10万亿GDP+海洋+大农业”挑好“挑起北方经济大梁”的担子，也要在投资、房地产与高碳工业调整中，用“转型+绿色+新动能”稳住基本盘。**观察山东，与其看“总量”，不如看“新旧动能替代、绿色低碳转型、海洋经济、民营活力与城乡共富”这几张名片。")

heading2(doc, "一页速览：2025年山东经济的“表与里”")
table(doc,
    ["维度", "表面现象", "更深层信号"],
    [
        ["增长", "GDP 103197.5亿、+5.5%", "北方第一个破10万亿省、工业强"],
        ["产业", "规上工业+7.6%、装备+11.4%", "新能源车+50.3%、动能转换"],
        ["外贸", "进出口3.53万亿（出口2.16万亿）", "民企占76.3%"],
        ["投资", "固定资产投资-8.6%", "制造业投资-0.5%、基建-7.7%"],
        ["财政", "一般公共预算收入+2.0%（税收+1.4%）", "收入低增、税收偏弱"],
        ["消费", "社零42082.9亿、+5.1%", "网上零售+10.1%"],
        ["人口", "常住10043万、城镇化率67.54%", "过亿人口大省、自然增-3.05‰"],
        ["海洋", "海水产品产量+3.4%、港口吞吐量+5.7%", "海洋强省、港口枢纽"],
    ],
    widths=[2.2, 5.2, 8.6])
para(doc, "", size=4, space_after=2)

# ---- 一、研究口径与阅读方法 ----
heading1(doc, "一、研究口径与阅读方法")
heading2(doc, "1.1 本报告的三份核心底稿")
bullet(doc, "**2025年《政府工作报告》**：2025年1月20日在省十四届人大三次会议上作，给出全年预期目标（GDP增长5%以上、一般公共预算收入增长3%、城镇调查失业率5.5%左右、城镇新增就业110万人以上、居民收入与增长同步、CPI保持合理水平、节能减排降碳完成约束性指标等）。固定资产投资/社零/进出口未设具体速率目标。")
bullet(doc, "**2025年《国民经济和社会发展统计公报》**：山东省统计局2026年3月3日发布，给出全年实际执行数与增速，是“事后验证”的锚点。")
bullet(doc, "**2026年山东省政府工作报告及官方复盘**：对2025执行结果的追认，交叉验证“目标—执行—再评价”三段式。")
heading2(doc, "1.2 阅读方法：显性—数据—制度—长期")
para(doc, "每一章按“**显性表述 → 同期数据 → 制度含义 → 长期影响**”四层展开。其中“制度含义”回答“政府为什么这么排优先序”，“长期影响”回答“这对未来5—10年意味着什么”。")
para(doc, "关键判别：**当报告使用的“增长词”和统计公报里的实际数不一致时，数据优先。**例如2025年初定“GDP增长5%以上”、实际+5.5%达标；“一般公共预算收入3%”，实际+2.0%；固定资产投资则实际约-8.6%。差异反映：山东在投资/房地产调整下，仍靠工业与动能转换守住增长，但财政收入增速低于预期。")

# ---- 二、特殊底盘 ----
heading1(doc, "二、先看山东的特殊底盘：北方经济第一大省、海洋强省与乡村振兴大省")
para(doc, "在所有省份里，山东的“底盘”独特：**北方经济第一大省+全国过亿人口大省+海洋强省+农业大省（粮食/农产品出口）**四合一。面积、人口、产业都居全国前列，是“北方经济的压舱石”。")
para(doc, "这决定山东的多重身份并存：**制造重镇**（装备、化工、汽车、新能源）、**海洋强省**（沿海港口21.9亿吨吞吐量）、**农业大省**（粮食产量全国前列、现代农业示范）、**绿色低碳转型前沿**（新旧动能转换综合试验区）。")
heading2(doc, "2.1 工业与动能转换")
para(doc, "山东是我国工业门类最齐全的省份之一，2025年规上工业+7.6%、装备制造成长。新旧动能转换从“重化+基建”转向“十强产业（新一代信息、高端装备、新能源、海洋等）”。")
heading2(doc, "2.2 海洋与开放")
para(doc, "山东半岛有青烟威等沿海城市，港口、海洋经济、外海资源丰富。进出口3.53万亿，实际使用外资投向高技术产业较多。海洋经济是山东“蓝海”优势。")
heading2(doc, "2.3 农业与乡村振兴")
para(doc, "山东是粮食/蔬菜/农产品大国，正推进“打造乡村振兴齐鲁样板”。在“粮食安全+现代农业”下，山东既大且重，也夹带着向“高附加值/品牌化”升级的任务。")

# ---- 三、宏观错位 ----
heading1(doc, "三、最关键的宏观错位：GDP破10万亿、工业强，但投资与进口偏弱")
para(doc, "把2025年山东的宏观面放进一张表，会出现令外行惊讶的“错位”：表观增长来自工业/消费/出口，而投资与进口/税收相对偏弱。这个错位，正是读懂山东的关键。")
heading2(doc, "3.1 “强的一面”")
table(doc,
    ["指标", "2025实绩", "信号"],
    [
        ["GDP", "103197.5亿、+5.5%", "北方第一个破10万亿省"],
        ["规上工业", "+7.6%", "制造业强、动能转换"],
        ["装备制造", "+11.4%", "高端制造加速"],
        ["新能源车", "110万辆、+50.3%", "新能源放量"],
        ["社零", "+5.1%", "内需消费稳"],
        ["民营进出口", "占76.3%", "民营外贸强"],
    ],
    widths=[3.2, 5.2, 6.0])
heading2(doc, "3.2 “弱的一面”")
table(doc,
    ["指标", "2025实值", "信号"],
    [
        ["固定资产投资", "-8.6%", "投资收缩"],
        ["制造业投资", "-0.5%", "制造投资平淡"],
        ["一般公共预算收入", "+2.0%（税收+1.4%）", "收入低增"],
        ["基建投资", "-7.7%", "基建放缓"],
        ["人口自然增率", "-3.05‰", "人口负增、老龄化"],
    ],
    widths=[3.2, 5.2, 6.0])
para(doc, "**错位结论**　山东的增长“很强、但也很矛盾”。强的部分（工业/装备/消费/新能源）与弱的部分（固投/基建/税收/人口）并存。**真正的焦点不是“有没有增长”，而是“量的持续性”**：工业与能耗转换撑起增长，但投资/基建偏弱、税收增速低。2026年“新旧动能+绿色低碳+海洋+新质生产力”继续发力，同时补投资与稳税收。")

# ---- 四、被忽视细节 ----
heading1(doc, "四、容易被忽视、但信号很重的细节（15条）")
details = [
    ("1", "装备制造业增加值+11.4%、占规上26.6%", "高端制造加速提升。"),
    ("2", "新能源车产量110.5万辆、+50.3%", "新能源车翻倍式放量。"),
    ("3", "锂离子电池产量+103.1%", "储能/电池景气。"),
    ("4", "工业机器人产量2.8万套、+26.6%", "智能制造新赛道。"),
    ("5", "高新技术企业3.5万家、科技型中小企业5.1万家", "创新主体壮大。"),
    ("6", "进出口3.53万亿、民营企业占76.3%", "民企外贸压舱。"),
    ("7", "海水产品产量+3.4%、港口吞吐量21.9亿吨+5.7%", "海洋与港口枢纽。"),
    ("8", "固定资产投资-8.6%、制造业投资-0.5%", "投资偏弱、需补动能。"),
    ("9", "一般公共预算收入+2.0%、税收+1.4%", "收入低增、税收弱。"),
    ("10", "常住10043万、城镇化率67.54%", "过亿人口大省。"),
    ("11", "人口自然增-3.05‰、出生率5.16‰", "自然负增、老龄化。"),
    ("12", "居民人均可支配收入44180元、+5.0%", "收入稳增、与增长同步。"),
    ("13", "网上零售额+10.1%", "数字化消费升级。"),
    ("14", "农村居民收入+5.4%（>城镇4.4%）", "乡村振兴/城乡差距缩小。"),
    ("15", "新登记市场主体150.5万户、总数1453.3万户", "市场主体、民营活力。"),
]
for d in details:
    para(doc, "**%s**" % d[0], size=11)
    para(doc, d[1], size=10.5, color=GRAY)

# ---- 五、六个换挡 ----
heading1(doc, "五、把所有细节连起来：山东正在经历的“六个换挡”")
gear = [
    ("1．动能换挡：从“重化/基建驱动”到“新旧动能转换+新质生产力”", "装备+11.4%、工业机器人+26.6%、新能源车+50.3%。"),
    ("2．产业换挡：从“钢铁化工铅”到“十强产业+高端制造+绿色”", "高端装备、新材料、新能源在放量。"),
    ("3．投资换挡：从“基建/地产”到“制造业+设备更新/绿色投资”", "制造投资偏弱、技改+5.3%，绿色投资较大。"),
    ("4．开放换挡：从“传统外贸”到“民营外贸+海洋+一带一路”", "民企占76.3%、港口吞吐量增、海洋/外海。"),
    ("5．人口换挡：从“净流入红利”到“过亿人口存量+老龄化”", "常住1.00亿、自然增-3.05‰。"),
    ("6．动能换挡：从“高碳增长”到“绿色低碳+海洋经济”", "新能源车/锂电/绿色能源/海洋。"),
]
for g in gear:
    para(doc, "**%s**　%s" % g, space_after=6)

# ---- 六、增长暗线 ----
heading1(doc, "六、增长暗线：新旧动能转换、绿色低碳与新质生产力")
heading2(doc, "6.1 新旧动能转换综合试验区")
para(doc, "山东是全国“新旧动能转换”先行区，从“重化+基建”转向“十强产业”（高端装备、新能源、新一代信息、海洋）。装备制造+11.4%、新能源车+50.3%、锂电+103.1%，说明“新动能”正在放量。")
heading2(doc, "6.2 绿色低碳与海洋")
para(doc, "山东推进“绿色低碳高质量发展先行区”，新能源与可再生能源装机新增较快；海洋经济（港口、海水产品、海洋高端装备）是山东“蓝海”优势。")
para(doc, "**这条暗线意味着**：山东的增长正从“高碳重化”转向“绿色低碳+新动能+海洋”。看山东，盯住“装备制造/新能源占比”与“绿色转型、海洋经济”这几组数。")

# ---- 七、财政暗线 ----
heading1(doc, "七、财政暗线：收入稳、税收平，民生与转型平衡")
para(doc, "2025年山东一般公共预算收入7864.3亿元、+2.0%，税收5109.7亿元、+1.4%。收入低增长，与“减税降费+地产弱+工业通缩”有关；收入结构偏税收弱化。")
para(doc, "支出主向民生、乡村振兴、绿色转型。山东在“稳增长—稳民生—促转型”间平衡，需提升制造业/新动能的税源质量。")
para(doc, "**制度含义**　山东财政“收入稳、税收平”，长期要靠“新旧动能转换+海洋经济+新质生产力”带来的真实税源，摆脱对地产与重化的依赖。")

# ---- 八、产业暗线 ----
heading1(doc, "八、产业暗线：从“重化工业”到“十强产业+高端制造”")
heading2(doc, "8.1 山东产业的“表”")
table(doc,
    ["指标", "2025增速/信号", "解读"],
    [
        ["规上工业", "+7.6%", "制造强省"],
        ["装备制造", "+11.4%", "高端制造加速"],
        ["新能源车", "110万辆、+50.3%", "新能源放量"],
        ["锂离子电池", "+103.1%", "储能/电池景气"],
        ["工业机器人", "+26.6%", "智能装备"],
        ["民营进出口", "占76.3%", "民营外贸强"],
        ["海洋", "港口21.9亿吨", "海洋强省"],
    ],
    widths=[5.2, 4.0, 5.0])
heading2(doc, "8.2 从“重化”到“十强/高端”")
para(doc, "山东过去以钢铁、化工、重化见长，2025年显示“装备+新能源+锂电+高端制造”正加速成为新增长。“新旧动能转换”从口号到数据（装备、机器人、锂电高增），是理解山东产业升级的主线索。")

# ---- 九、区域 ----
heading1(doc, "九、区域格局：山东半岛城市群与海洋强省")
para(doc, "山东半岛城市群是全国重要的沿海城市群，青岛、济南双核驱动，威海、烟台、潍坊等环渤海/黄海联动。以“海洋强省+半岛城市群”为轴，青岛港/烟台港等港口支撑开放。")
para(doc, "山东也是全国农业重镇，半岛/鲁中/鲁西等农业高产区与乡村振兴协同。基建设施+新材料+装备在各地承建。")

# ---- 十、人口与城市 ----
heading1(doc, "十、人口与城市：过亿人口大省、城镇化与老龄化")
heading2(doc, "10.1 人口总量")
para(doc, "2025年末山东常住10043万人、城镇化率67.54%，自然增长率-3.05‰、出生率5.16‰。山东已进入人口负增，但仍是过亿人口大省。")
heading2(doc, "10.2 城市与老龄")
para(doc, "青岛、济南等是半岛城市群核心，吸引人口；但全省老龄化加深（预计65岁以上约17%），养老、医疗、就业面临转型。山东要在“人口存量+产业升级”中寻找新均衡。")

# ---- 十一、民营经济 ----
heading1(doc, "十一、民营经济：占七成多进出，市场主体活跃")
heading2(doc, "11.1 民企地位")
para(doc, "山东民营企业进出口占76.3%，市场主体1453.3万户、新登记150.5万户。民营企业是山东外贸与经济的绝对主力。")
heading2(doc, "11.2 政策与服务")
para(doc, "山东持续优化营商环境、支持民营经济与专精特新。民营+个体+科技型中小企业，是山东高质量发展的底座。")

# ---- 十二、事后验证 ----
heading1(doc, "十二、事后验证：用2025年实际执行结果检验报告判断")
heading2(doc, "12.1 年初目标与年末执行对比")
table(doc,
    ["指标", "2025年初目标", "2025实际", "偏差"],
    [
        ["GDP增速", "5%以上", "+5.5%", "达标"],
        ["一般公共预算收入", "3%", "+2.0%", "低于目标"],
        ["城镇新增就业", "110万以上", "充裕", "达标"],
        ["居民收入", "与增长同步", "+5.0%", "同步"],
        ["CPI", "合理水平", "偏低", "待观察"],
        ["固定资产投资", "未设", "约-8.6%", "弱"],
    ],
    widths=[3.0, 3.4, 3.0, 4.0])
heading2(doc, "12.2 判断验证")
para(doc, "三条核心判断得到支持：**（1）工业/装备/新能源强（规上+7.6%、装备+11.4%），验证；（2）投资/基建偏弱、税收低增，验证；（3）海洋/民营强，验证。**")
para(doc, "核心观察：山东靠“新旧动能转换+工业+民营”守住5.5%增长、成为北方首个破10万亿，但投资/基建、税收偏弱。2026年看“绿色低碳+海洋+民营+新质生产力”能否续力。")

# ---- 十三、五条主线 ----
heading1(doc, "十三、未来5—10年最值得观察的五条主线")
lines5 = [
    ("① 新旧动能转换能否完成", "装备/新能源车/锂电，能否彻底替代重化/基建。"),
    ("② 绿色低碳与海洋经济", "绿色能源、海洋深蓝，能否成第二增长。"),
    ("③ 制造业/民营/出口", "民营占76%、装备出海，能否保持韧性。"),
    ("④ 投资与税收再平衡", "固投-8.6%，能否用产业投资找回。"),
    ("⑤ 人口/城乡/老龄化", "过亿人口、负增，如何用创新与乡村振兴稳住。"),
]
for l in lines5:
    para(doc, "**%s**　%s" % l, space_after=6)

# ---- 十四、结论 ----
heading1(doc, "十四、最终结论：山东在“新旧动能转换+绿色低碳”里的增长逻辑")
para(doc, "山东的2025年，本质上是“**新旧动能转换+工业+绿色低碳为核心，而对重化/基建依赖下降**”的答卷：GDP破10.32万亿、规上工业+7.6%、装备/新能源加速，代价是固投-8.6%、税收低增、人口负增。")
para(doc, "只要动能转换、绿色、海洋、民营能接住，山东就站在“北方经济强省”头排；如果投资、税收、人口持续偏弱，山东需承受“转型阵痛”。")
para(doc, "最稳妥的观察信号：**一盯装备/新能源动能（动能）、二盯绿色低碳/海洋（绿色）、三盯民营/出口（底座）、四盯财政税收（财政）、五盯投资/人口（约束）。**山东，是中国“新旧动能转换”的试验田。")

# ---- 附录A ----
heading1(doc, "附录A：主要资料来源与核验路径")
bullet(doc, "山东省2025年《政府工作报告》（2025年1月）——目标来源。")
bullet(doc, "《2025年山东省国民经济和社会发展统计公报》（山东省统计局，2026-03-03）——GDP、工业、外贸、人口等实值。")
bullet(doc, "山东省统计/运行分析、新旧动能转换专报——产业与动能。")
bullet(doc, "2026年山东省政府工作报告——2025执行复盘。")
bullet(doc, "山东海关、海洋渔业、财政厅——外贸/海洋/财政。")
heading2(doc, "核验说明")
para(doc, "本报告所有增速、总量、占比均以统计公报/官方发布为基准。涉“海洋经济总量”“新旧动能转换细分”等以官方口径为准。")

# ---- 附录B ----
heading1(doc, "附录B：建议建立的年度跟踪仪表盘")
para(doc, "建议把下面10个“测脉搏”指标做成年度跟踪表：")
dash = [
    ["#", "指标", "2025基值", "为什么重要"],
    ["1", "GDP增速", "+5.5%", "总量与方向"],
    ["2", "规上工业增速", "+7.6%", "制造底盘"],
    ["3", "装备制造业增速", "+11.4%", "高端制造"],
    ["4", "新能源汽车产量", "110.5万辆+50.3%", "新动能"],
    ["5", "民营进出口占比", "76.3%", "民营活力"],
    ["6", "固定资产投资增速", "-8.6%", "投资动能"],
    ["7", "一般公共预算收入/税收", "+2.0%/+1.4%", "财政质量"],
    ["8", "常住人口/城镇化率", "10043万/67.54%", "人口与城市"],
    ["9", "社零增速", "+5.1%", "内需消费"],
    ["10", "居民人均可支配收入增速", "+5.0%", "民生获得感"],
]
table(doc, dash[0], dash[1:], widths=[0.8, 4.5, 3.6, 4.6])
para(doc, "把这10个指标连起来看，任何一个“新引擎（2/3/4/5）向上、旧引擎（6）调整”，都说明山东在真正换挡；反之则是旧路径的反复。")

# ========================================= 保存
out = "/Users/x/Desktop/content-prod-lab/reports/山东省_2025年政府工作报告_深度研究_2026-08-13.docx"
doc.save(out)
print("SAVED:", out)
print("PARAGRAPHS:", len(doc.paragraphs))
print("TABLES:", len(doc.tables))
