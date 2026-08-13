# -*- coding: utf-8 -*-
"""Build 宁波市2025年政府工作报告 深度研究 DOCX, 参照省市系列版式。"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DARK = RGBColor(0x1F, 0x2A, 0x44)
RED = RGBColor(0xB0, 0x1E, 0x2E)
GRAY = RGBColor(0x59, 0x60, 0x69)
BLUE = RGBColor(0x1F, 0x4E, 0x79)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
HEADER_FILL = "1F2A44"
K = "微软雅黑"


def set_run(run, size=11, bold=False, color=DARK, font=K):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    run.font.name = font
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), font)
    rFonts.set(qn("w:hAnsi"), font)
    rFonts.set(qn("w:eastAsia"), font)


def para(doc, text, size=11, bold=False, color=DARK, align=None,
         space_after=6, space_before=0, font=K):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        set_run(r, size=size, bold=bold or (i % 2 == 1), color=color, font=font)
    return p


def heading1(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(16)
    pf.space_after = Pt(8)
    r = p.add_run(text)
    set_run(r, size=16, bold=True, color=RED)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "B01E2E")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def heading2(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(10)
    pf.space_after = Pt(4)
    r = p.add_run(text)
    set_run(r, size=13, bold=True, color=BLUE)
    return p


def table(doc, headers, rows, widths=None, font_size=9.5):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_run(r, size=font_size, bold=True, color=WHITE)
        tcPr = hdr[i]._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), HEADER_FILL)
        tcPr.append(shd)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(val))
            set_run(r, size=font_size, color=DARK)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    return t


def quote_box(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(8)
    pf.left_indent = Pt(12)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), "B01E2E")
    pBdr.append(left)
    pPr.append(pBdr)
    r = p.add_run(text)
    set_run(r, size=11, color=DARK, font="楷体")
    return p


def bullet(doc, text, size=10.5):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.7)
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        set_run(r, size=size, bold=(i % 2 == 1), color=DARK)
    return p


# ================================================================= 文档主体
doc = Document()
sec = doc.sections[0]
sec.page_width = Cm(21)
sec.page_height = Cm(29.7)
sec.top_margin = Cm(2.4)
sec.bottom_margin = Cm(2.2)
sec.left_margin = Cm(2.5)
sec.right_margin = Cm(2.5)

st = doc.styles["Normal"]
st.font.name = K
st.font.size = Pt(11)
st.element.rPr.rFonts.set(qn("w:eastAsia"), K)



# ---- 封面 ----
para(doc, "宁波市2025年政府工作报告", size=24, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=60, space_after=4)
para(doc, "深度研究与\u201c容易被忽视的细节\u201d分析", size=16, bold=True, color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
para(doc, "从\u201c港口经济、外贸大市、先进制造与计划单列\u201d重新理解宁波", size=12, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)
para(doc, "\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501", color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
para(doc, "研究对象：2025年宁波市政府工作报告及同期财政、国民经济和社会发展统计资料、2026年官方复盘", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
para(doc, "标定日期：2026-08-13", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
doc.add_page_break()

# ---- 目录 ----
catalog = [
    "执行摘要",
    "一、研究口径与阅读方法",
    "二、先看宁波的特殊底盘：港口经济、外贸大市、先进制造与计划单列",
    "三、最关键的宏观错位：GDP破1.87万亿、外贸/制造强，但固投/地产大幅下行",
    "四、容易被忽视、但信号很重的细节（15条）",
    "五、2025年GDP目标 vs 实际：对照表",
    "六、2025年增长是谁撑起来的？（结构归因）",
    "七、预算与财政的\u201c含金量\u201d",
    "八、民生底账：人口、收入与城乡",
    "九、城镇与农村：格局与均衡",
    "十、人口流入与流出",
    "十一、物价与货币环境",
    "十二、区域一体化：宁波在\u201c宁波舟山港+长三角港口群+浙江海洋经济\u201d里的位置",
    "十三、未来5—10年最值得观察的五条主线",
    "十四、最终结论：宁波在\u201c港口/外贸+先进制造+计划单列\u201d里的增长逻辑",
    "附录A：主要资料来源与核验路径",
    "附录B：建议建立的年度跟踪仪表盘",
]
para(doc, "目　录", size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=10, space_after=10)
for item in catalog:
    para(doc, item, size=12, space_after=4)
doc.add_page_break()

# ---- 执行摘要 ----
heading1(doc, "执行摘要")
para(doc, "**核心判断**　如果只看新闻标题，2025年宁波最显著的是\u201cGDP破1.87万亿、增长5%左右\u201d、\u201c进出口1.46万亿、出口居全国城市第四\u201d、\u201c宁波舟山港货物吞吐量14.3亿吨（全球第一）、集装箱4387万标箱（全球第三）\u201d、\u201c规上工业+5.3%\u201d。但这份研究真正值得深读的，是这座\u201c港口经济+外贸大市+先进制造\u201d的计划单列城市，如何在出口高位、房地产大降（开发投资-45.9%、固投-21.4%）的背景下，靠\u201c制造业+外贸+港口物流\u201d实现约4.9%的落地增长。")
para(doc, "把2025年初设定的目标（GDP增长5.5%以上）、2025年《国民经济和社会发展统计公报》、2026年复盘放在一起看，宁波呈现清晰暗线：**从\u201c外贸+重化\u201d的既有底盘，向\u201c先进制造+外贸新业态+港口枢纽+海洋经济\u201d的升级**。优先于总量，宁波的核心竞争力在\u201c港口+制造+民营外贸\u201d。")
para(doc, "因此，本报告不按政府工作报告逐段复述，而采用\u201c显性表述—同期数据—制度含义—长期影响\u201d的方式，专门提取容易被忽略、但对判断宁波未来5—10年发展模式更有价值的信号。")
para(doc, "最容易记住的一句话：**宁波是\u201c港口+外贸+先进制造\u201d的计划单列城市，靠\u201c宁波舟山港+外贸出口+高端制造\u201d撑起增长。**观察宁波，与其看\u201cGDP 1.87万亿\u201d，不如看\u201c港口14.3亿吨、外贸出口全国第四、制造业单项冠军104家、生产性服务业\u201d这几张名片。")
heading2(doc, "一页速览：2025年宁波经济的\u201c表与里\u201d")
table(doc,
    ["维度", "表面现象", "更深层信号"],
    [
        ["增长", "GDP 18716亿、+4.9%", "一产2.5%、二产42.0%、三产55.5%"],
        ["产业", "规上工业+5.3%", "石油+23.2%、计算机通信+13.7%、高技术+11.9%"],
        ["外贸", "进出口14562亿、+2.6%", "出口+3.7%(全国第四)、新三样+76.3%"],
        ["投资", "固投-21.4%、扣除地产-6.9%", "房投-45.9%、地产大幅拖累"],
        ["消费", "社零5703亿、+1.8%", "网络零售+8.8%、档期偏弱"],
        ["人口", "常住983.3万、城镇化81.5%", "人口+5.6万、自然增-1.45‰"],
        ["港口", "宁波舟山港14.3亿吨、4387万标箱", "连续17年货物全球第一/集装箱全球第三"],
    ],
    widths=[2.2, 5.6, 7.4])
para(doc, "", size=10, space_after=2)

# ---- 一、研究口径与阅读方法 ----
heading1(doc, "一、研究口径与阅读方法")
heading2(doc, "1.1 本报告的核心底稿")
bullet(doc, "**2025年《政府工作报告》**（2025年1月）——目标：GDP 5.5%以上、一般公共预算收入2%、城镇就业20万、社零6%。")
bullet(doc, "**《2025年宁波市统计公报》**（市统计局2026-04）——GDP、工业、外贸、港口、人口实数。")
bullet(doc, "**2026年宁波市政府工作报告/复盘**——对2025执行追认与港口/先进制造展望。")
heading2(doc, "1.2 阅读方法：显性—数据—制度—长期")
para(doc, "**关键判别**：数据优先。例2025年GDP目标5.5%以上、实际4.9%（公报）/报告称\u201c5%左右\u201d；社零目标6%、实际1.8%明显偏弱。宁波\u201c外贸制造/港口强、消费/固投/地产弱\u201d，需穿透GDP总量看港口与制造。")

# ---- 二、底盘 ----
heading1(doc, "二、先看宁波的特殊底盘：港口经济、外贸大市、先进制造与计划单列")
para(doc, "宁波的地盘取决于它作为\u201c**宁波舟山港枢纽+外贸大市+先进制造基地+计划单列市**\u201d的特殊定位。它是进口外贸与海洋经济重镇，浙东工业与港口中心。")
bullet(doc, "**港口经济**：宁波舟山港货物吞吐量14.3亿吨（连续17年全球第一）、集装箱4387.2万标箱（连续8年全球第三）、集装箱航线309条；国际航运中心指数全球第七。")
bullet(doc, "**外贸大市**：进出口14562亿、出口居全国城市第四、进口第六；有进出口实绩企业3.15万家、贸易伙伴234国；\u201c新三样\u201d出口+76.3%。")
bullet(doc, "**先进制造**：规上工业+5.3%，高技术制造+11.9%、战略性新兴产业+4.9%；国家级制造业单项冠军累计全国第一（首破百）。")
bullet(doc, "**计划单列**：经济总量全国城市第十一；财政、金融、外贸权限独立，海洋经济与开放门户地位突出。")
para(doc, "这一底板决定了宁波2025成绩单的\u201c底色\u201d：**只要外贸/港口/先进制造持续，宁波就站在\u201c开放枢纽+制造强市\u201d增长极；若地产/传统“重化”调整，宁波需承受\u201c外贸制造强、消费投资弱\u201d的结构挑战。**")

# ---- 三、宏观错位 ----
heading1(doc, "三、最关键的宏观错位：GDP破1.87万亿、外贸/制造强，但固投/地产大幅下行")
para(doc, "宁波2025年最值得咀嚼的错位，是\u201c**出口/制造/港口强，消费/固投/地产却大幅走弱**\u201d。这种错位决定了对这座\u201c开放型计划单列\u201d城市的观察不能只看GDP总量。")
bullet(doc, "**GDP**：18716亿、+4.9%。一产472亿（+3.9%，占比2.5%）、二产7866亿（+4.0%，占比42.0%）、三产10378亿（+5.6%，占比55.5%）。人均190879元。")
bullet(doc, "**工业**：规上工业+5.3%；石油+23.2%、计算机通信+13.7%、电子元件+75.0%、机器人+46.4%、集成电路+10.6%；高技术制造+11.9%。")
bullet(doc, "**消费**：社零5703亿、+1.8%；网络零售2341亿、+8.8%（不足目标6%）。")
bullet(doc, "**外贸**：进出口14562亿、+2.6%（出口+3.7%、居全国第四；进口+0.2%）；\u201c新三样\u201d出口+76.3%。")
bullet(doc, "**固投/地产**：固投-21.4%、扣除地产后-6.9%；房地产投资-45.9%、商品房销售下行。")
para(doc, "**为什么读这条**：宁波作为\u201c计划单列+外贸/制造+港口\u201d城市，结构性矛盾是\u201c外部开放与制造强、内部投资/消费/地产弱\u201d。增长靠外贸/制造/三产，但固投和地产大幅收缩。")

# ---- 四、15条细节 ----
heading1(doc, "四、容易被忽视、但信号很重的细节（15条）")
para(doc, "以下15条，都在2025年报告/公报里，但常被\u201cGDP 1.87万亿\u201d、\u201c外贸1.46万亿\u201d等总量掩盖。它们是判断宁波2025之后5—10年的关键小信号。")
bullet(doc, "**1. 港口全球第一（宁波舟山港）**：货物吞吐量14.3亿吨（连续17年全球第一）、集装箱4387.2万标箱（全球第三）；国际航运中心指数全球第七。")
bullet(doc, "**2. 出口全国第四**：进出口14562亿、出口居全国城市第四；\u201c新三样\u201d出口+76.3%（电动车、锂电池、光伏）。")
bullet(doc, "**3. 石油加工+23.2%**：石化（镇海炼化）是宁波制造第一引擎，与化工/油气链强相关。")
bullet(doc, "**4. 计算机通信电子+13.7%、电子元件+75%**：电子信息/集成电路是先进制造的第二引擎。")
bullet(doc, "**5. 数字经济核心+12.1%、占GDP 8.5%**：数字经济制造业增加值1599亿、增长12.1%。")
bullet(doc, "**6. 制造业单项冠军全国第一**：累计104家（全国首个破百），专精特新小巨人居前。")
bullet(doc, "**7. 工业营收27266亿、利润+5.0%、利润率5.52%**：利润强于营收（+1.6%）体现结构升级。")
bullet(doc, "**8. 城乡收入倍差1.60**：缩小0.02，长三角高质量均衡。")
bullet(doc, "**9. 民营外贸**：3.15万家有进出口实绩企业、贸易伙伴234个——民营外贸活跃。")
bullet(doc, "**10. 规划/新质**：海铁联运200万标箱（2026目标215万）、梅山港万吨级泊位2个、航运集聚区营收2000亿。")
bullet(doc, "**11. 未来产业**：低空经济、人形机器人争创未来产业先导区；人工智能核心产业营收目标900亿。")
bullet(doc, "**12. 全球贸易伙伴/东盟**：对东盟、拉美、非洲出口两位数增长，新兴市场多元化。")
bullet(doc, "**13. 财政质量**：一般公共预算收入1795亿、+0.3%，税收1476亿、+2.9%，税占比82.2%——质量好。")
bullet(doc, "**14. 绿色/疫源**：全社会用电/减排；大气/水环境持续改善（优良率保持高位）。")
bullet(doc, "**15. 计划单列/自主权**：全国城市经济第十一，财政/金融/外资/外贸权限独立，海洋经济门户。")
para(doc, "", size=10, space_after=2)

# ---- 五、对照表 ----
heading1(doc, "五、2025年GDP目标 vs 实际：对照表")
para(doc, "2025年报告中列出的主要预期目标，与2025年《统计公报》实际完成情况：")
tb = [
    ["指标", "2025年目标", "2025年实际（公报/报告）", "达标判定"],
    ["GDP增速", "5.5%以上", "+4.9%（公报）/约5%左右（报告）", "略低"],
    ["一般公共预算收入", "+2%", "+0.3%（1795亿）", "不及预期"],
    ["社会消费品零售总额", "+6%", "+1.8%（5703亿）", "明显偏低"],
    ["进出口", "保持全国占比", "+2.6%（14562亿）", "平稳"],
    ["城镇新增就业", "20万以上", "26.4万人（公报）", "达标"],
    ["CPI", "2%左右", "+0.4%", "温和"],
]
table(doc, tb[0], tb[1:], widths=[3.2, 3.2, 4.6, 4.0])
para(doc, "", size=10, space_after=2)
para(doc, "**要点**：就业/外贸达标，GDP/收入/社零偏弱——宁波\u201c总量放缓、地产/固投拖累、外贸底盘稳\u201d。")

# ---- 六、结构归因 ----
heading1(doc, "六、2025年增长是谁撑起来的？（结构归因）")
heading2(doc, "6.1 三次产业：三产带动、二产稳住")
para(doc, "**三产（+5.6%、占比55.5%）支撑**：港口物流/外贸服务/金融/软件现代服务业强。二产+4.0%（占比42%）稳住盘子；一产+3.9%（2.5%）。")
heading2(doc, "6.2 工业：石油+汽车高端驱动")
para(doc, "规上工业+5.3%；石油+23.2%、计算机通信+13.7%、高技术制造+11.9%、电子元件+75.0%、机器人+46.4%。**石化（镇海炼化）+电子信息（集成电路/机器人）**是双引擎。")
heading2(doc, "6.3 外贸开放：第一引擎")
para(doc, "进出口14562亿、+2.6%（出口3.7%/进口0.2%）；新三样+76.3%。宁波是典型\u201c出口导向+港口牵引\u201d模式。")
heading2(doc, "6.4 消费偏弱")
para(doc, "社零+1.8%、明显低于6%目标；网络零售+8.8%是亮点，但整体需求与地产/收入偏弱。")
heading2(doc, "6.5 投资：地产深度调整")
para(doc, "固投-21.4%、房投-45.9%；扣除地产后-6.9%；工业/制造业投资-12%/-15%——投资端地产与制造同步承压。")
para(doc, "**一句话归因**：宁波2025年\u201c**外贸+港口+高端制造/三产**\u201d为主引擎，\u201c固投/地产/消费\u201d大幅走弱——典型的\u201c外向型、港口+制造强、内部投资/地产调整\u201d计划单列。")

# ---- 七、财政 ----
heading1(doc, "七、预算与财政的\u201c含金量\u201d")
para(doc, "**地方一般公共预算收入1795亿元、+0.3%**；税收1476亿、+2.9%，税占比82.2%；一般公共预算支出2104亿元、-5.6%。")
bullet(doc, "**收入质量高**：税占比82.2%，财政依赖外贸/制造/现代服务高质量税基。")
bullet(doc, "**支出收缩**：支出-5.6%，反映紧平衡/积极财政空间有限，民生与债务统筹。")
bullet(doc, "**民生硬度**：居民人均可支配收入77779元、+4.0%；城乡收入倍差1.60；低保、就业26.4万。")
para(doc, "**财政含义**：宁波\u201c收入稳、质量高，但增长放缓、支出收缩\u201d——政策空间有限，需靠外贸/制造/港口与债务防风险平衡。")

# ---- 八、民生 ----
heading1(doc, "八、民生底账：人口、收入与城乡")
para(doc, "宁波\u201c十四五\u201d人口净增：**常住人口983.3万人、比上年增加5.6万；城镇化率81.5%（+0.6pct）**。自然增长率-1.45‰（出生4.91‰、死亡6.36‰，人口结构老化），依赖净流入。")
bullet(doc, "**收入**：居民人均可支配收入77779元、+4.0%；城镇86088元、+3.6%，农村56016元、+4.8%（城乡倍差1.60、缩小）。")
bullet(doc, "**就业**：城镇新增就业26.4万人、登记失业率2.14%；提供岗位充足。")
bullet(doc, "**社保**：低保/养老/医保体系完善，低保标准统筹。")
bullet(doc, "**民生**：万人有效发明专利77.3件（每万人），创新与民生并重。")
para(doc, "**民生含义**：宁波\u201c收入/就业\u201d稳、人口净流入（-1.45‰自然增长靠机械增长补），高城镇化+高收入是浙江/宁波特色。")

# ---- 九、城乡 ----
heading1(doc, "九、城镇与农村：格局与均衡")
para(doc, "**城镇化率81.5%（+0.6pct）**，宁波高度城镇化/都市化；城乡收入倍差缩小至1.60、居全国领先。")
bullet(doc, "**城市**：轨道交通、智慧城市、海洋经济门户；海铁联运、港口物流驱动城市化。")
bullet(doc, "**农村**：县域强县（慈溪、余姚、宁海等）制造/外贸结合；农村收入+4.8%、低收入农户+10.6%。")
para(doc, "**城乡均衡**：宁波\u201c城市强、县域强\u201d，县域制造/外贸贡献显著，城乡收入均衡居全国前列。")

# ---- 十、人口流 ----
heading1(doc, "十、人口流入与流出")
para(doc, "**宁波呈净流入**：2025年常住+5.6万、自然增长-1.45‰（出生4.91‰、死亡6.36‰）。在人口总体收缩背景下，依赖净流入（外来务工/外贸/制造）。")
bullet(doc, "**流入**：制造业岗位+外贸/码头/新产业，吸引省内外劳动力；计划单列对人才有吸引力。")
bullet(doc, "**竞争**：与杭州/苏州/上海争夺；需靠制造岗位、港口产业、收入留住人口。")
para(doc, "人口方向决定中长期需求与增长；宁波的\u201c高质量发展+强县\u201d是其长逻辑之一，但需警惕老龄化/自然负增长。")

# ---- 十一、物价 ----
heading1(doc, "十一、物价与货币环境")
para(doc, "**2025年宁波市区CPI上涨0.4%**（略高于全国/浙江的0）：食品烟酒+0.1%、衣着+4.3%、居住+0.2%、交通通信-2.5%、其他+12.3%。")
bullet(doc, "**物价**：CPI+0.4%、温和，主要受衣着/其他（服务）拉动；交通通信负（原油/出行）拖累。")
bullet(doc, "**货币/流动性**：本外币存款余额增长约10%、贷款+8.5%；金融活跃（新增上市公司）。")
para(doc, "**物价含义**：宁波通胀温和可控，居民购买力稳定；关注服务价格与交通能源走势。")

# ---- 十二、区域一体化 ----
heading1(doc, "十二、区域一体化：宁波在\u201c宁波舟山港+长三角港口群+浙江海洋经济\u201d里的位置")
para(doc, "宁波处于**长三角世界级港口群+浙江海洋经济+宁波舟山港**核心：既是计划单列市，也是长江经济带龙头和长三角一体化的重要门户。")
bullet(doc, "**宁波舟山港**：货物吞吐量14.3亿吨（连续17年全球第一）、集装箱4387.2万标箱（全球第三）；国际航运中心指数全球第七；309条航线。")
bullet(doc, "**长三角一体化**：共建大飞机集群、港口群，宁镇扬/杭甬都市圈互动；外贸/港口协同长三角。")
bullet(doc, "**海洋经济/浙江**：作为计划单列、海洋门户，支撑浙江/长江带对外开放。")
bullet(doc, "**海铁联运**：200万标箱（2026目标215万），提升内陆腹地联动。")
para(doc, "**区域含义**：宁波靠\u201c宁波舟山港+外贸开放+制造\u201d承东启西、服务长三角与长江经济带，是浙江挑大梁的开放引擎。")

# ---- 十三、五条主线 ----
heading1(doc, "十三、未来5—10年最值得观察的五条主线")
bullet(doc, "**主线1｜港口与航运枢纽**：宁波舟山港全球第一/全球第三；能否持续守住全球集装箱前三、国际航运中心进位。")
bullet(doc, "**主线2｜外贸新业态**：新三样+76.3%、跨境电商、海外仓、新兴市场（东盟/拉美/非洲）。能否把出口全国第四变成更高附加值。")
bullet(doc, "**主线3｜先进制造/单项冠军**：制造业单项冠军104家、石化/电子信息/新能源车。能否在高端制造/海洋装备卡位。")
bullet(doc, "**主线4｜数字经济与未来产业**：数字经济核心+12.1%、低空/具身智能/人工智能900亿。能否孵化新增长极。")
bullet(doc, "**主线5｜民营经济与新质**：民营外贸/制造强、单项冠军+专精小巨人。能否持续激活民营资本与创新。")
para(doc, "这五条，是宁波从\u201c外贸+制造+港口城市\u201d走向\u201c高端制造+开放枢纽+海洋经济强市\u201d的\u201c主赛道\u201d。")

# ---- 十四、结论 ----
heading1(doc, "十四、最终结论：宁波在\u201c港口/外贸+先进制造+计划单列\u201d里的增长逻辑")
para(doc, "宁波2025年，本质上是\u201c**港口/外贸+高端制造/三产驱动，固投/地产/消费大幅走弱**\u201d的答卷：GDP18716亿、+4.9%、进出口14562亿（出口全国第四、新三样+76.3%）、宁波舟山港全球第一、规上工业+5.3%（高技术+11.9%）、社零+1.8%、财政+0.3%（税占82.2%）。")
para(doc, "只要港口/外贸/高端制造持续，宁波就站在\u201c开放枢纽+制造强市\u201d增长极；若地产/固投深度调整拖累，宁波需承受\u201c外向强、内需/投资弱\u201d的结构挑战。")
para(doc, "最稳观察信号：**一盯港口吞吐量（枢纽）、二盯出口/新三样（开放）、三盯单项冠军/高端制造（制造）、四盯社零/消费（内需）、五盯人口/都市圈（长期）。**宁波，是\u201c港口+外贸+先进制造\u201d的新样本。")

# ---- 附录A ----
heading1(doc, "附录A：主要资料来源与核验路径")
bullet(doc, "宁波市2025年政府工作报告（2025年1月）——目标来源。")
bullet(doc, "《2025年宁波市国民经济和社会发展统计公报》（市统计局）——GDP、工业、外贸、港口、人口实值。")
bullet(doc, "2026年宁波市政府工作报告（2026年1月）——2025复盘/港口/先进制造/对外。")
bullet(doc, "宁波舟山港、市财政/商务等。")
heading2(doc, "核验说明")
para(doc, "本报告涉及数据以统计公报/官方口径为准；\u201c港口/外贸/先进制造/计划单列\u201d等以官方口径为准。")

# ---- 附录B ----
heading1(doc, "附录B：建议建立的年度跟踪仪表盘")
para(doc, "建议把下面10个\u201c测脉搏\u201d指标做成年度跟踪表：")
dash = [
    ["#", "指标", "2025基值", "为什么重要"],
    ["1", "GDP增速", "+4.9%（18716亿）", "总量与方向"],
    ["2", "规模以上工业增速", "+5.3%", "制造底盘"],
    ["3", "进出口/出口增速", "+2.6%/+3.7%（全国第四）", "开放韧性"],
    ["4", "港口吞吐量/集装箱", "14.3亿吨/4387万标箱", "港口枢纽"],
    ["5", "货物贸易伙伴/新三样", "234国 / +76.3%", "外贸结构"],
    ["6", "固定资产投资/工业", "-21.4% / -12.4%", "投资结构"],
    ["7", "社零增速", "+1.8%（5703亿）", "内需消费"],
    ["8", "常住人口/城镇化", "983.3万 / 81.5%", "人口与城市"],
    ["9", "地方财政收入/税占", "+0.3%(1795亿) / 82.2%", "财政质量"],
    ["10", "CPI", "+0.4%", "物价"],
]
table(doc, dash[0], dash[1:], widths=[0.8, 4.6, 4.2, 3.6])
para(doc, "把这10个指标连起来看，港口/出口/新三样（3/4/5）、制造/单项冠军（2）、人口（8），都说明宁波在真正换挡。")

# ============ 保存 ============
out = "/Users/x/Desktop/content-prod-lab/reports/宁波市_2025年政府工作报告_深度研究_2026-08-13.docx"
doc.save(out)
print("SAVED:", out)
print("PARAGRAPHS:", len(doc.paragraphs))
print("TABLES:", len(doc.tables))
