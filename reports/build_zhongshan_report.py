# -*- coding: utf-8 -*-
"""Build 中山市2025年政府工作报告 深度研究 DOCX, 参照地级市系列版式。"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DARK = RGBColor(0x1F, 0x2A, 0x44)
RED = RGBColor(0xB0, 0x1E, 0x2E)
GRAY = RGBColor(0x59, 0x60, 0x69)
BLUE = RGBColor(0x1F, 0x4E, 0x79)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
HEADER_FILL = "1F2A44"
K = "微软雅黑"


def set_run(run, size=11, bold=False, color=DARK, font=K):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    run.font.name = font
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), font)
    rFonts.set(qn("w:hAnsi"), font)
    rFonts.set(qn("w:eastAsia"), font)


def para(doc, text, size=11, bold=False, color=DARK, align=None,
         space_after=6, space_before=0, font=K):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        set_run(r, size=size, bold=bold or (i % 2 == 1), color=color, font=font)
    return p


def heading1(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(16)
    pf.space_after = Pt(8)
    r = p.add_run(text)
    set_run(r, size=16, bold=True, color=RED)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "B01E2E")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def heading2(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(10)
    pf.space_after = Pt(4)
    r = p.add_run(text)
    set_run(r, size=13, bold=True, color=BLUE)
    return p


def table(doc, headers, rows, widths=None, font_size=9.5):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_run(r, size=font_size, bold=True, color=WHITE)
        tcPr = hdr[i]._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), HEADER_FILL)
        tcPr.append(shd)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(val))
            set_run(r, size=font_size, color=DARK)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    return t


def quote_box(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(8)
    pf.left_indent = Pt(12)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), "B01E2E")
    pBdr.append(left)
    pPr.append(pBdr)
    r = p.add_run(text)
    set_run(r, size=11, color=DARK, font="楷体")
    return p


def bullet(doc, text, size=10.5):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.7)
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        set_run(r, size=size, bold=(i % 2 == 1), color=DARK)
    return p


# ================================================================= 文档主体
doc = Document()
sec = doc.sections[0]
sec.page_width = Cm(21)
sec.page_height = Cm(29.7)
sec.top_margin = Cm(2.4)
sec.bottom_margin = Cm(2.2)
sec.left_margin = Cm(2.5)
sec.right_margin = Cm(2.5)

st = doc.styles["Normal"]
st.font.name = K
st.font.size = Pt(11)
st.element.rPr.rFonts.set(qn("w:eastAsia"), K)


# ---- 封面 ----
para(doc, "中山市2025年政府工作报告", size=24, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=60, space_after=4)
para(doc, "深度研究与\u201c容易被忽视的细节\u201d分析", size=16, bold=True, color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
para(doc, "从\u201c粤港澳大湾区、家电制造、灯具古镇、装备制造与深中协同\u201d重新理解中山", size=12, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)
para(doc, "\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501", color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
para(doc, "研究对象：2025年中山市政府工作报告及同期财政、国民经济和社会发展统计资料、2026年官方复盘", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
para(doc, "标定日期：2026-08-13", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
doc.add_page_break()

# ---- 目录 ----
catalog = [
    "执行摘要",
    "一、研究口径与阅读方法",
    "二、先看中山的特殊底盘：深中通道、家电/灯饰制造、装备制造与大湾区",
    "三、最关键的宏观错位：GDP破4200亿但低于目标，装备制造强，固投/地产转负",
    "四、容易被忽视、但信号很重的细节（15条）",
    "五、2025年GDP目标 vs 实际：对照表",
    "六、2025年增长是谁撑起来的？（结构归因）",
    "七、预算与财政的\u201c含金量\u201d",
    "八、民生底账：人口、收入与城乡",
    "九、城镇与农村：格局与均衡",
    "十、人口流入与流出",
    "十一、物价与货币环境",
    "十二、区域一体化：中山在\u201c粤港澳大湾区+深中同城+黄金内湾\u201d里的位置",
    "十三、未来5—10年最值得观察的五条主线",
    "十四、最终结论：中山在\u201c家电+装备+深中协同\u201d里的增长逻辑",
    "附录A：主要资料来源与核验路径",
    "附录B：建议建立的年度跟踪仪表盘",
]
para(doc, "目　录", size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=10, space_after=10)
for item in catalog:
    para(doc, item, size=12, space_after=4)
doc.add_page_break()

# ---- 执行摘要 ----
heading1(doc, "执行摘要")
para(doc, "**核心判断**　2025年中山最显著的是\u201cGDP 4260.56亿元、增长3.8%（低于5.5%目标）、二产占51.4%\u201d、\u201c规上工业+5.4%（装备+11.1%、高技术+18.6%）\u201d、\u201c进出口2872.1亿/+1.1%\u201d、\u201c常住人口457.23万/城镇化87.95%\u201d。这说明中山在\u201c深中同城+产业升级\u201d下稳中向好，但**投资-22.6%、地产-30.5%、消费+3.4%**偏弱。")
para(doc, "把2025年目标（GDP+5.5%/规上+6%/财政+3%）、2025年统计、2026年前瞻一起看，中山是\u201c制造大市+深中通道效应\u201d路径：**家电（空调/冰箱）+灯具+装备+先进制造+深中协同**是引擎。GDP连跨3500/4000亿、2025站上4200亿。")
para(doc, "最容易记住的一句话：**中山是\u201c家电制造+深中通道+大湾区\u201d的广东城市，靠\u201c家电/灯具、装备制造、先进制造、深中同城协同\u201d增长。**观察中山，与其只看\u201cGDP 4260亿\u201d，不如看\u201c装备制造+11.1%、高技术+18.6%、洗碗机/冰箱产量、进出口2872亿、深中通道\u201d。")# ---- 一、研究口径与阅读方法 ----
heading1(doc, "一、研究口径与阅读方法")
para(doc, "本报告以\u201c中山市政府工作报告（2025年）\u201d为起点，把\u201c2025年GDP目标（5.5%）\u201d与\u201c2025年国民经济和社会发展统计公报实际值（4260.56亿元/+3.8%）\u201d并置对照，再用2026年政府工作报告的复盘作事后验证。")
para(doc, "重点阅读三类证据：**一是指标目标是否达标**；**二是结构（谁贡献增长）是否匹配目标叙事**；**三是官方复盘如何自评、承认问题**。统计口径一般公共预算收入为地方级。")
para(doc, "统一采用\u201c同比增长%\u201d（GDP为不变价），财政与居民收入多为名义值，文中按原口径转录、不逐项换算。")

# ---- 二、特殊底盘 ----
heading1(doc, "二、先看中山的特殊底盘：深中通道、家电/灯饰制造、装备制造与大湾区")
para(doc, "**区位与身份**：中山是广东省地级市、粤港澳大湾区中心节点，2024年深中通道建成通车（中山尽享深中同城红利）、黄金内湾核心、\u201c孙中山故里\u201d。")
para(doc, "**产业底盘**：一是家电（空调2247万台、冰箱1511万台、微波炉4170万台，冰洗空/餐厨）；二是灯饰/灯具（古镇灯饰之都，灯具及照明装置10.98亿套）；三是装备制造（+11.1%、先进制造+8.9%、高技术+18.6%）；四是工业母机（+17.9%）。")
para(doc, "**人口底盘**：2025年末常住人口457.23万/+7.77万、城镇化率87.95%（提高0.17pct）；广东人口/制造大市。")
para(doc, "**市场与出口**：2025年社零1604.29亿/+3.4%；进出口2872.1亿/+1.1%。")

# ---- 三、核心宏观错位 ----
heading1(doc, "三、最关键的宏观错位：GDP破4200亿但低于目标，装备制造强，固投/地产转负")
para(doc, "**第一组错位**：2025年GDP目标5.5%以上，实际4260.56亿/+3.8%**低于目标1.7pct**；但GDP破4200亿、\u201c十四五\u201d连跨3500/4000亿，年均+4.5%。")
para(doc, "**第二组错位**：规上工业+5.4%、装备+11.1%、高技术+18.6%增长好；但**固投-22.6%、房地产-30.5%**塌陷，\u201c产业增、投资落\u201d。")
para(doc, "**第三组错位**：常住+7.77万、自然+3.42\u2030（人口正流入）；但**CPI-0.4%**、收入实际+2.7%放缓。")
para(doc, "一句话：**中山是\u201c制造（家电/装备/灯饰）强、投资/地产弱、深中红利于此\u201d的城市**——靠产业+深中同城+出口，投资收缩是最大短板。")

# ---- 四、15条细节 ----
heading1(doc, "四、容易被忽视、但信号很重的细节（15条）")
bullet(doc, "1. **GDP 4260.56亿/+3.8%**：破4200亿、\u201c十四五\u201d连跨3500/4000亿。")
bullet(doc, "2. **装备制造+11.1%**：先进制造占规上52.2%。")
bullet(doc, "3. **高技术制造+18.6%、生物医药/器械+40.8%**：新质生产力。")
bullet(doc, "4. **工业母机+17.9%**：中山特色装备母机。")
bullet(doc, "5. **家电/冰洗/空调产量大**：空调2247万台、冰箱1511万台、洗碗机全球重要基地。")
bullet(doc, "6. **进出口2872.1亿/+1.1%（东盟+21%）**：外贸韧性、顺差2159亿。")
bullet(doc, "7. **固投-22.6%**：工业投资-14.6%、基础设施-26%。")
bullet(doc, "8. **房地产-30.5%（销售-19%）**：地产深度调整。")
bullet(doc, "9. **社零+3.4%（家电+78%、网络+56.8%）**：线上、家电强。")
bullet(doc, "10. **常住457.23万/城镇化87.95%**：人口净流入、自然正增长。")
bullet(doc, "11. **收入：城镇69138元/+2.5%、农村50326元/+4.6%**，城乡1.37。")
bullet(doc, "12. **CPI-0.4%**：通缩。")
bullet(doc, "13. **财政380.01亿/+3.5%（税收+3.0%）**：达标。")
bullet(doc, "14. **R&D强度3.00%、高企2917家**：科创制造。")
bullet(doc, "15. **规上工业利润347.41亿/+14.4%**：盈利改善。")# ---- 五、目标 vs 实际对照表 ----
heading1(doc, "五、2025年GDP目标 vs 实际：对照表")
table(doc,
    ["指标", "2025年目标", "2025年实际", "是否达标"],
    [
        ["地区生产总值增速", "5.5%以上", "4260.56亿元/+3.8%", "未达标"],
        ["规上工业增加值", "+6%以上", "+5.4%", "略低于目标"],
        ["固定资产投资", "提质提效", "-22.6%（转负）", "未达标"],
        ["社会消费品零售总额", "+5%", "1604.29亿元/+3.4%", "未达标"],
        ["一般公共预算收入", "+3%", "380.01亿元/+3.5%", "达标"],
        ["进出口", "稳量提质", "2872.1亿元/+1.1%", "稳中偏低"],
        ["CPI涨幅", "2%左右", "-0.4%", "低于目标(通缩)"],
    ],
    widths=[3.4, 2.5, 5.0, 3.1])
para(doc, "**简评**：仅**财政达标**；GDP、规上、消费、投资均低于目标。**制造升级（装备+11%、高技术+18.6%）与进出口顺差**是亮点，**投资/地产**是短板。")

# ---- 六、结构归因 ----
heading1(doc, "六、2025年增长是谁撑起来的？（结构归因）")
para(doc, "**三大产业**：一产98.18亿/+6.3%、二产2189.06亿/+3.9%、三产1973.32亿/+3.6%；二产占51.4%。增长以二产（工业）+三产（商贸物流）双支撑。")
para(doc, "**工业**：规上+5.4%；**先进制造、装备+11.1%、高技术+18.6%、工业母机+17.9%**；家电（空调/冰箱/洗碗机）+灯饰+装备是主体。规上工业利润347.41亿/+14.4%。")
para(doc, "**服务业**：批发零售（家具/家电）/物流；网络零售+56.8%。")
para(doc, "**增长归因**：中山GDP主要靠**二产（家电+装备+高技术制造）+工业升级+进出口顺差**；投资、地产为负贡献。")

# ---- 七、财政 ----
heading1(doc, "七、预算与财政的\u201c含金量\u201d")
para(doc, "2025年一般公共预算收入380.01亿元/+3.5%（税收229.18亿/+3.0%）；支出441.89亿/+4.6%，民生支出279.49亿占63.2%。")
para(doc, "**结构性**：收入\u201c达标、税收+3%同步\u201d；支出保民生。规上工业利润+14.4%支撑税基。")
para(doc, "**含金量**：中山财政\u201c税收质量尚可、民生占比高\u201d；制造业盈利回暖（+14.4%）是税收韧性来源。")

# ---------------- 八、民生 ----------------
heading1(doc, "八、民生底账：人口、收入与城乡")
para(doc, "常住人口457.23万/+7.77万、城镇化率87.95%；出生7.35\u2030/死亡3.93\u2030，自然+3.42\u2030（正自然增长）。收入：**城镇69138元/+2.5%（实际）、农村50326元/+4.6%（实际）**，城乡比1.37。")
para(doc, "消费：全体39922元/+0.5（实际）、城镇40912。城镇新增就业66130人/+20.3%。")
para(doc, "**民生结论**：收入农村快于城镇、城乡差距极小（1.37）；人口正自然增长+净流入；就业、收入温和、普惠好。")

# ---------------- 九、 ----------------
heading1(doc, "九、城镇与农村：格局与均衡")
para(doc, "中山城镇化率87.95%（地处珠三角、很高），中心（石岐/坦洲）与镇街（古镇/小榄/横栏/南头）承载制造；家电、灯饰、装备集群分布。")
para(doc, "城乡收入比1.37（全国最低一档），农村收入快、极大均衡；镇街经济避免二元。")

# ============ 十、人口流入与流出 ============
heading1(doc, "十、人口流入与流出")
para(doc, "中山常住457万、净流入+自然正增长（+7.77万），是珠三角制造业人口聚集地。")
para(doc, "**流入**：制造业/家电/装备岗位、深中同城带来深圳产业/人口外溢；**流出**：部分低技能回流。总体\u201c净流入、质量提升\u201d。")

# ============ 十一 ============
heading1(doc, "十一、物价与货币环境")
para(doc, "2025年中山CPI**下降0.4%**（通缩），交通通信-4.0%、居住-0.5%。")
para(doc, "金融：存款9740.10亿/+3.5%、贷款8177.45亿/+3.7%。\u201c宽中稳、实体融资\u201d。")

# ============ 十二 ============
heading1(doc, "十二、区域一体化：中山在\u201c粤港澳大湾区+深中同城+黄金内湾\u201d里的位置")
para(doc, "中山是粤港澳大湾区中心节点、深中通道（2024年通车）直接受益的\u201c桥头堡\u201d，深中同城红利巨大（承接深圳产业/创新外溢）；作为南中城际、广中珠澳高铁、黄金内湾核心。")
para(doc, "制造业协同：家电/家电制造对接深圳创新链；中开高速/深中通道支撑“黄金内湾”物流带。中山是\u201c制造+协同\u201d的枢纽城市。")

# ============ 十三 ============
heading1(doc, "十三、未来5—10年最值得观察的五条主线")
bullet(doc, "1. **深中同城红利兑现**：深中通道后深圳产值/产业/资金外溢能否持续放大。")
bullet(doc, "2. **家电制造升级**：智能家电/数字经济，培育链主升级。")
bullet(doc, "3. **装备制造+先进制造**：装备+11.1%、先进制造52%，打造制造强市。")
bullet(doc, "4. **工改+土地整备**：\u201c工改\u201d拆除4.6万亩，释放产业空间。")
bullet(doc, "5. **投资与地产修复**：固投-22.6%、地产-30.5%后，靠深中/基建/工业投资修复。")

# ============ 十四 ============
heading1(doc, "十四、最终结论：中山在\u201c家电+装备+深中协同\u201d里的增长逻辑")
para(doc, "**结论**：中山2025年的\u201c真相\u201d是——**GDP+3.8%（低于目标）、规上+5.4%（装备/高技术强）、进出口+1.1%、固投-22.6%、财政+3.5%**。它是\u201c制造（家电+装备）+深中红利\u201d驱动但**投资/地产弱、体量7000亿集群**的珠三角城市。")
para(doc, "**对趋势判断**：家电/装备/先进制造/深中同城代表中山\u201c动能\u201d，投资/地产代表\u201c约束\u201d。**深中红利+制造升级**决定中期潜力，**投资修复+产业利润**决定韧性。")
para(doc, "**若只看一个指标**：看**装备/高技术制造增速（+11.1%/+13.9%）与固定资产投资增速（-22.6%）**——中山正借助深中通道与\u201c工改\u201d释放的产业空间换挡，制造升级能否把深中红利转化为投资与增量是关键。")

# ---------------- 附录A ----------------
heading1(doc, "附录A：主要资料来源与核验路径")
bullet(doc, "中山市人民政府《2025年中山市政府工作报告》（2025年）。")
bullet(doc, "中山市统计局《2025年中山市国民经济和社会发展统计公报》（2026年5月）。")
bullet(doc, "《2026年中山市政府工作报告》（2026年2月，尹念红）及极简版。")
bullet(doc, "广东省2025年统计公报与中山统计年鉴交叉核验。")

# ---------------- 附录B ----------------
heading1(doc, "附录B：建议建立的年度跟踪仪表盘")
bullet(doc, "GDP总量/增速 vs 全年目标（+/-百分点）。")
bullet(doc, "规上工业增加值及分行业（装备/高技术/工业母机/家电）增速。")
bullet(doc, "固定资产投资（总量/工业/高技术/房地产）增速。")
bullet(doc, "社会消费品零售总额、网络零售、家电类。")
bullet(doc, "进出口（人民币）、出口、东盟/香港。")
bullet(doc, "一般公共预算收入/税收、财政支出、民生占比。")
bullet(doc, "常住人口、城镇化率、净流入。")
bullet(doc, "CPI/核心CPI、规上工业利润与营收。")
bullet(doc, "工业母机、空调/冰箱/洗碗机/灯饰产量。")
bullet(doc, "深中通道车流量、深企入驻、\u201c工改\u201d面积。")

# ---------------- 保存 ----------------
out = "/Users/x/Desktop/content-prod-lab/reports/中山市_2025年政府工作报告_深度研究_2026-08-13.docx"
doc.save(out)
print("SAVED OK 中山市_2025年政府工作报告_深度研究_2026-08-13.docx")