# -*- coding: utf-8 -*-
"""Build 德州市2025年政府工作报告 深度研究 DOCX, 参照地级市系列版式。"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DARK = RGBColor(0x1F, 0x2A, 0x44)
RED = RGBColor(0xB0, 0x1E, 0x2E)
GRAY = RGBColor(0x59, 0x60, 0x69)
BLUE = RGBColor(0x1F, 0x4E, 0x79)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
HEADER_FILL = "1F2A44"
K = "微软雅黑"


def set_run(run, size=11, bold=False, color=DARK, font=K):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    run.font.name = font
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), font)
    rFonts.set(qn("w:hAnsi"), font)
    rFonts.set(qn("w:eastAsia"), font)


def para(doc, text, size=11, bold=False, color=DARK, align=None,
         space_after=6, space_before=0, font=K):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        set_run(r, size=size, bold=bold or (i % 2 == 1), color=color, font=font)
    return p


def heading1(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(16)
    pf.space_after = Pt(8)
    r = p.add_run(text)
    set_run(r, size=16, bold=True, color=RED)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "B01E2E")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def heading2(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(10)
    pf.space_after = Pt(4)
    r = p.add_run(text)
    set_run(r, size=13, bold=True, color=BLUE)
    return p


def table(doc, headers, rows, widths=None, font_size=9.5):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_run(r, size=font_size, bold=True, color=WHITE)
        tcPr = hdr[i]._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), HEADER_FILL)
        tcPr.append(shd)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(val))
            set_run(r, size=font_size, color=DARK)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    return t


def quote_box(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(8)
    pf.left_indent = Pt(12)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), "B01E2E")
    pBdr.append(left)
    pPr.append(pBdr)
    r = p.add_run(text)
    set_run(r, size=11, color=DARK, font="楷体")
    return p


def bullet(doc, text, size=10.5):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.7)
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        set_run(r, size=size, bold=(i % 2 == 1), color=DARK)
    return p


# ================================================================= 文档主体
doc = Document()
sec = doc.sections[0]
sec.page_width = Cm(21)
sec.page_height = Cm(29.7)
sec.top_margin = Cm(2.4)
sec.bottom_margin = Cm(2.2)
sec.left_margin = Cm(2.5)
sec.right_margin = Cm(2.5)

st = doc.styles["Normal"]
st.font.name = K
st.font.size = Pt(11)
st.element.rPr.rFonts.set(qn("w:eastAsia"), K)


# ---- 封面 ----
para(doc, "德州市2025年政府工作报告", size=24, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=60, space_after=4)
para(doc, "深度研究与\u201c容易被忽视的细节\u201d分析", size=16, bold=True, color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
para(doc, "从\u201c德州扒鸡（食品）、装备制造、太阳能（中国太阳城）、京津冀协同、绿色低碳\u201d重新理解德州", size=12, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)
para(doc, "\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501", color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
para(doc, "研究对象：2025年德州市政府工作报告及同期财政、国民经济和社会发展统计资料、2026年官方复盘", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
para(doc, "标定日期：2026-08-14", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
doc.add_page_break()

# ---- 目录 ----
catalog = [
    "执行摘要",
    "一、研究口径与阅读方法",
    "二、先看德州的特殊底盘：中国太阳城（太阳能）、装备制造、食品（扒鸡）、京津冀协同节点",
    "三、最关键的宏观错位：GDP 4214.6亿/5.3%达标，工业新能源强但固投低、CPI负、需求不足",
    "四、容易被忽视、但信号很重的细节（15条）",
    "五、2025年GDP目标 vs 实际：对照表",
    "六、2025年增长是谁撑起来的？（结构归因）",
    "七、预算与财政的\u201c含金量\u201d",
    "八、民生底账：人口、收入与城乡",
    "九、城镇与农村：格局与均衡",
    "十、人口流入与流出",
    "十一、物价与货币环境",
    "十二、区域一体化：德州在\u201c京津冀协同+黄河流域+山东半岛\u201d里的位置",
    "十三、未来5—10年最值得观察的五条主线",
    "十四、最终结论：德州在\u201c太阳能+装备制造+京津冀协同\u201d里的增长逻辑",
    "附录A：主要资料来源与核验路径",
    "附录B：建议建立的年度跟踪仪表盘",
]
para(doc, "目　录", size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=10, space_after=10)
for item in catalog:
    para(doc, item, size=12, space_after=4)
doc.add_page_break()

# ---- 执行摘要 ----
heading1(doc, "执行摘要")
para(doc, "**核心判断**　2025年德州最显著的是\u201cGDP 4214.6亿元、增长5.3%（达到5.5%以上目标）、人均约7.8万\u201d、\u201c规上工业+7.5%（装备制造、太阳能光伏强）\u201d、\u201c进出口767.9亿/+4.6%（出口+7.1%）\u201d、\u201c但固投+0.2%低位、CPI-0.4%、外贸总量偏小\u201d、\u201c财收271.24亿/+3%、常住542.66万\u201d。这说明德州在\u201c太阳能+装备+食品\u201d的绿色制造升级中达标稳增，但**投资偏弱、内需不足**是短板。")
para(doc, "把2025年目标（GDP+5.5%以上/财收+3%左右）、2025年统计（GDP+5.3%略低于目标、规上+7.5%、固投+0.2%、社零+4.5%、财收+3.0%达成、进出口+4.6%）、趋势一起看，德州是\u201c太阳城+制造+京津冀\u201d路径：**太阳能（中国太阳城、光伏）、装备制造、绿色食品、化工新材料、京津冀承接**是支柱；2025年总量4214亿居山东第9。")
para(doc, "最容易记住的一句话：**德州是\u201c中国太阳城（太阳能/光伏）、京津冀协同带、食品加工与扒鸡之乡\u201d，靠\u201c太阳能+装备制造+京津冀承接\u201d增长。**观察德州，与其只看\u201cGDP 4214亿\u201d，不如看\u201c新能源装机1063.79万千瓦占76.6%、装备制造、规上+7.5%、京津冀项目\u201d。")

# ---- 一、研究口径与阅读方法 ----
heading1(doc, "一、研究口径与阅读方法")
para(doc, "本报告以\u201c德州市政府工作报告（2025年，朱开国作）\u201d为起点，把\u201c2025年GDP目标（5.5%以上）\u201d与\u201c官方2025年GDP（4214.6亿元/+5.3%）\u201d并置对照，并用\u201c2025年德州市统计公报\u201d和\u201c2026年政府工作报告复盘\u201d横向核验。")
para(doc, "口径提示：\u201cGDP、规上工业增加值增速\u201d按可比价（实际增速）；\u201c投资、消费、进出口、财政收入\u201d按现价（名义增速）。涉及\u201c常住人口\u201d用统计公报常住口径（542.66万），城镇化率用官方公布值（58.4%）。")
para(doc, "指标体系（与研究口径一致）：五大象限——**总量与增速（GDP）、产业动能（太阳能/装备/食品/京津冀）、外贸、财政质量、民生与人口**。")
para(doc, "特别提示（不吃老本）：德州2024年GDP 4047.7亿/+5.6%、规上+8.4%，2025年放缓到+5.3%、规上+7.1%；德州最有辨识度的是**中国太阳城（太阳能光伏装机占全省前列）+京津冀协同承接（北汽、央企）+食品加工（扒鸡/粮食）**——真正要看的是\u201c太阳能制造+京津产业承接\u201d的转型成色。")
# ---- 二、先看德州的特殊底盘 ----
heading1(doc, "二、先看德州的特殊底盘：中国太阳城（太阳能）、装备制造、食品（扒鸡）、京津冀协同节点")
para(doc, "德州位于山东黄河下游、鲁西北，是\u201c中国太阳城\u201d、京津冀协同发展承接区、山东省\u201c黄河流域生态保护和高质量发展\u201d节点。2025年GDP 4214.6亿、常住人口542.66万，人均约7.8万元。")
para(doc, "四个底盘名词，先立框架：")
bullet(doc, "**中国太阳城（太阳能/光伏）**　德州是全球\u201c太阳能之城\u201d之一（皇明/景津等），2025年新能源和可再生能源装机1063.79万千瓦（+11.86%）占全部装机76.6%，清洁能源发电量181.39亿千瓦时（+17.81%）居全省前列。")
bullet(doc, "**装备制造**　景津过滤成套装备、电梯设备制造（入选省级产业集群）、机床、机械等，装备制造是工业主力。")
bullet(doc, "**食品加工（扒鸡/粮食）**　德州扒鸡、金锣等食品加工业；粮食总产量787.85万吨（连续多年高位）、蔬菜731万吨——\u201c扒鸡+粮仓\u201d。")
bullet(doc, "**京津冀协同承接**　德州对接京津冀（北汽/央企产业转移、京津项目落地100个以上），是鲁北\u201c京津冀辐射节点\u201d。")
para(doc, "这四根支柱（太阳能+装备+食品+京津冀）构成德州独特底盘：**左手太阳城（绿色能源），右手装备制造+食品，腹地是京津冀协同带**。理解德州，先理解它\u201c兼有山东与京津冀\u201d的双重区位。")

# ---- 三、最关键的宏观错位 ----
heading1(doc, "三、最关键的宏观错位：GDP 4214.6亿/5.3%达标，工业新能源强但固投低、CPI负、需求不足")
para(doc, "2025年德州最需要辨析的一组\u201c错位\u201d：**GDP 5.3%（略低5.5%目标）、规上工业+7.5%、新能源装机占76.6%强，但固投+0.2%近零、CPI-0.4%、进出口767.9亿**。")
para(doc, "为什么\u201c工业强、太阳能大红\u201d，投资却平淡？三个解释：")
para(doc, "**其一，工业快但投资弱**　规上工业+7.5%，装备/新能源（光伏）制造强；但固投仅+0.2%，说明**企业主体投资动能不足、制造业扩张在放缓**。")
para(doc, "**其二，内需偏弱、物价负**　社零+4.5%但CPI-0.4%（食品-3.4%）、粮食降价；有效需求不足是报告自承挑战。")
para(doc, "**其三，外贸总量小**　进出口767.9亿/+4.6%，出口+7.1%尚可，但总量在山东偏小、进口-9.0%。")
para(doc, "小结：德州2025年是\u201c**稳总量、强工业&新能源、弱投资&需求**\u201d的一年：太阳能+装备+京津冀撑增量，但**投资、物价、内需**偏软。")

# ---- 四、容易被忽视、但信号很重的细节（15条） ----
heading1(doc, "四、容易被忽视、但信号很重的细节（15条）")
bullet(doc, "**1.太阳城/新能源装机1063.79万千瓦占比76.6%**　风电/太阳能发电181.39亿千瓦时/+17.8%，清洁能源制造与发电并进。")
bullet(doc, "**2.规上工业+7.5%**\u201c装备制造、电梯、太阳能/光伏、化工\u201d拉动（制造业强）。")
bullet(doc, "**3.装备制造（景津过滤、电梯集群）**\u201c电梯设备入选省级产业集群\u201d，装备是工业主力。")
bullet(doc, "**4.社会责任\u201c绿色制造\u201d**\u201c新增省级绿色工业园3个、绿色工厂17家\u201d，绿色低碳转型。")
bullet(doc, "**5.进出口767.9亿/+4.6%（出口+7.1%）**\u201c出口好于进口\u201d，跨境电商企业新增130家。")
bullet(doc, "**6.京津冀协同承接**\u201c京津冀落地项目100个以上（目标）、北汽/央企带\u201d。")
bullet(doc, "**7.财收271.24亿/+3.0%、民生支出占82.7%**\u201c财政稳、民生占比高\u201d。")
bullet(doc, "**8.社零+4.5%（1617.61亿）**\u201c以旧换新拉动59.5亿、智能家电+34.7\u201d新消费。")
bullet(doc, "**9.居民收入31792元/+4.5%**\u201c城镇38237/+4.0%、农村24276/+4.8%\u201d城乡差缩小。")
bullet(doc, "**10.常住542.66万/城镇化58.4%（+1.16pct）**\u201c山东人口大市、城镇化加速\u201d。")
bullet(doc, "**11.CPI-0.4%**\u201c物价负、粮食/食品降价\u201d，流通良好。")
bullet(doc, "**12.太阳能发电/光伏制造**\u201c1.7GW/光伏组件\u201d、太阳城产能、光伏玻璃。")
bullet(doc, "**13.新增规上工业156家/专精特新183家**\u201c企业梯队扩容\u201d。")
bullet(doc, "**14.\u201c氢能试点\u201d+地热供暖800万㎡**\u201c绿色能源（氢/地热）试点\u201d。")
bullet(doc, "**15.城镇新增就业5.4万人（完成108.5%）**\u201c就业稳定\u201d。")
# ---- 五、2025年GDP目标 vs 实际：对照表 ----
heading1(doc, "五、2025年GDP目标 vs 实际：对照表")
table(doc,
    ["指标", "2025年目标", "2025年实际", "达成评价"],
    [
        ["地区生产总值(GDP)", "增长5.5%以上", "4214.6亿/5.3%", "略低于目标"],
        ["一般公共预算收入", "增长3%左右", "271.24亿/+3.0%", "达成"],
        ["规上工业增加值", "——", "+7.5%", "工业强"],
        ["固定资产投资", "——", "+0.2%", "低位"],
        ["社会消费品零售总额", "——", "1617.61亿/+4.5%", "稳健"],
        ["进出口总额", "——", "767.9亿/+4.6%", "增长"],
        ["居民收入", "与经济增长基本同步", "31792元/+4.5%", "基本同步"],
    ],
)
para(doc, "注：GDP、规上工业按可比价；投资、消费、进出口、财收按现价。**GDP（+5.3%）略低于\u201c5.5%以上\u201d、财收（+3.0%）达成**；**规上工业（+7.5%）、社零（+4.5%）稳健**，固投仅+0.2%低位。")
para(doc, "拆读：**规上工业（+7.5%）、新能源装机（76.6%）、进出口（+4.6%）是亮色**，**固投（+0.2%）、CPI（-0.4%）偏弱**；\u201cGDP目标5.5%以上\u201d实际5.3%——\u201c大致达标、需求待启\u201d，是\u201c太阳城+京津冀\u201d的普通样本。")

# ---- 六、2025年增长是谁撑起来的？（结构归因） ----
heading1(doc, "六、2025年增长是谁撑起来的？（结构归因）")
para(doc, "把德州GDP的5.3%拆开：三次产业结构9.4：42.5：48.1，第二产业（工业）与第三产业（服务业）双撑。**第二产业（装备/太阳能制造）是工业引擎，第三产业（物流/商贸/服务）是增量，第一产业（粮仓）稳**。")
para(doc, "2026年德州强调\u201c十五五\u201d开局、京津冀协同、绿色低碳，聚焦**太阳能/光伏、装备制造、绿色食品、新能源（氢能/地热）**——核心是\u201c太阳城+京津产业承接\u201d。")
para(doc, "**第二产业（工业制造）**：规上工业+7.5%、装备制造（景津/电梯）、太阳能光伏、化工新材料——\u201c制造业强\u201d。")
para(doc, "**第三产业（服务业）**：现代物流、商贸（京津冀枢纽）、电商（网络零售210亿）、旅游——\u201c服务业稳\u201d。")
para(doc, "**外贸（开放型）**：进出口+4.6%（出口+7.1%）、跨境电商——\u201c外向量小但稳\u201d。")
para(doc, "一句话归因：**2025年德州增长\u201c靠工业（装备+太阳能）+京津冀承接\u201d**；投资偏弱、内需不足是短板。")

# ---- 七、预算与财政的\u201c含金量\u201d ----
heading1(doc, "七、预算与财政的\u201c含金量\u201d")
para(doc, "2025年德州**一般公共预算收入271.24亿元（+3.0%）**，民生支出占比82.7%；财政稳、民生优先。")
bullet(doc, "收入结构：GDP与财收增速匹配（财收+3%、GDP+5.3%），\u201c增收靠工业税（装备/新能源）\u201d。")
bullet(doc, "民生支出占82.7%，投向教育、社保、医疗、养老；城镇新增就业5.4万人（完成108.5%）。")
bullet(doc, "金融支撑：新增贷款406亿、政策性资金262亿、政府性融资担保24.6亿——金融支持制造/小微。")
para(doc, "**财政含金量小结**：财收+3%稳增、民生占比82.7%高，\u201c保民生、保产业\u201d；财政对太阳能、装备、京津冀项目投入加大。")

# ---- 八、民生底账：人口、收入与城乡 ----
heading1(doc, "八、民生底账：人口、收入与城乡")
para(doc, "2025年德州**居民人均可支配收入31792元（+4.5%）**，其中城镇38237元（+4.0%）、农村24276元（+4.8%），城乡差1.58（较上年缩小）。就业：城镇新增就业5.4万人。")
para(doc, "人口画像：**常住人口542.66万、城镇化率58.4%（+1.16pct）**，是山东人口大市；城镇化率低于山东平均（约65%）但提升快，中心城区集聚。")
para(doc, "民生投入：普惠托位5487个、养老/医养、居家养老服务——民生保障扎实。")

# ---- 九、城镇与农村：格局与均衡 ----
heading1(doc, "九、城镇与农村：格局与均衡")
para(doc, "德州常住城镇化率58.4%（仍低于山东平均），农村体量大；农村收入增速（+4.8%）高于城镇（+4.0%），**城乡差缩小（1.58）**。")
para(doc, "农业底盘：**粮食总产787.85万吨（连续多年高位、山东粮仓）**、蔬菜731万吨、生猪出栏——\u201c德州是全国重要的商品粮/蔬菜基地\u201d。")
para(doc, "一句话：\u201c德州农业粮仓大、农村收入快、城乡差收敛\u201d，但\u201c城镇化率仍低、农业人口多\u201d，需\u201c以产业带乡、以城带乡\u201d。")

# ---- 十、人口流入与流出 ----
heading1(doc, "十、人口流入与流出")
para(doc, "德州常住542.66万、城镇化率58.4%（山东人口大市），是\u201c户籍多于常住\u201d的劳务输出型城市，但也承接\u201c京津冀\u201d产业转移（京津冀项目100+）反吸人口。")
para(doc, "结构观察：**城镇化率高于陕甘、仍低于山东平均**，中心城区、京津冀产业带形成集聚；就地城镇化是主要路径。")
para(doc, "2026年目标：城镇新增就业5万以上、集聚青年大学生3.5万——德州以\u201c太阳城+京津产业\u201d聚人。")

# ---- 十一、物价与货币环境 ----
heading1(doc, "十一、物价与货币环境")
para(doc, "2025年德州**CPI同比-0.4%**（食品-3.4%、消费品-0.7%；扣除食品能源+0.5%）——**低通胀、需求不足**，与\u201c有效需求不足\u201d的报告挑战一致。")
bullet(doc, "信贷：新增贷款406亿、金融支撑制造/项目；政策性资金262亿。")
bullet(doc, "以旧换新/促消费：以旧换新拉动59.5亿、智能家电+34.7%、网络零售210亿——政策稳消费。")
para(doc, "货币环境判断：**宽信用、CPI负（-0.4%）**；\u201c资金充裕、物价偏低\u201d，德州通过\u201c以旧换新、扩内需\u201d稳消费（2026年扩内需）。")

# ---- 十二、区域一体化：德州的位置 ----
heading1(doc, "十二、区域一体化：德州在\u201c京津冀协同+黄河流域+山东半岛\u201d里的位置")
para(doc, "德州地济南与北京之间，是\u201c京津冀协同发展\u201d承接区、山东\u201c黄河流域生态保护和高质量发展\u201d、山东半岛城市群节点，\u201c1小时到京津\u201d。")
bullet(doc, "**京津冀协同**　德州承接北京/天津产业转移（北汽、央企、\u201c京津冀落地项目100+\u201d），是\u201c鲁北融京\u201d门户。")
bullet(doc, "**黄河流域**　山东黄河生态保护/高质量发展机会，发展\u201c绿色低碳、太阳城\u201d。")
bullet(doc, "**区域商务**　高铁（京沪/石济/京九）、太阳城、京津冀商贸辐射——区域物流/消费节点。")
para(doc, "一句话：**德州在\u201c京津冀协同+黄河流域\u201d里，最核心的定位是\u201c京津冀辐射山东的承接节点+中国太阳城\u201d**——近京、交通、绿色是最大优势。")

# ---- 十三、未来5—10年最值得观察的五条主线 ----
heading1(doc, "十三、未来5—10年最值得观察的五条主线")
bullet(doc, "**主线一：太阳城产业升级**\u201c新能源装机76.6%、光伏制造\u201d能否走向\u201c光伏+储能+氢能\u201d高端。")
bullet(doc, "**主线二：装备制造高端化**\u201c景津/电梯集群\u201d能否升级\u201c智能装备、锂电/储能装备\u201d。")
bullet(doc, "**主线三：京津冀承接成色**\u201c京津项目100+\u201d能否把\u201c承接\u201d变成\u201c高端制造+人才\u201d。")
bullet(doc, "**主线四：食品加工升级**\u201c扒鸡+粮食\u201d能否从\u201c原料\u201d走向\u201c品牌/预制菜（百亿）\u201d。")
bullet(doc, "**主线五：投资与需求**\u201c固投+0.2%&CPI-0.4%\u201d能否靠\u201c设备更新、扩内需\u201d转正。")

# ---- 十四、最终结论 ----
heading1(doc, "十四、最终结论：德州在\u201c太阳能+装备制造+京津冀协同\u201d里的增长逻辑")
para(doc, "把2025年闭环打穿：**德州是\u201c中国太阳城、京津冀协同承接区\u201d**：GDP 4214.6亿/+5.3%、规上工业+7.5%、新能源装机1063.79万千瓦占76.6%、进出口+4.6%、财收+3%。")
para(doc, "德州不是\u201c只有扒鸡\u201d——它是**太阳能（全球太阳城）+装备制造+京津冀协同+食品加工**的复合经济，靠\u201c绿色能源+制造+区位\u201d驱动；但投资、消费、物价偏弱，需求待启。")
para(doc, "一句话结论：**德州是\u201c中国太阳城、京津冀协同带\u201d；观察它先看\u201c太阳能/光伏、装备制造、京津冀承接\u201d，再看\u201c固投、物价、需求\u201d。**它是\u201c制造优、能源绿、需求弱\u201d的山东\u201c绿色制造\u201d样本。")

# ---- 附录A ----
heading1(doc, "附录A：主要资料来源与核验路径")
bullet(doc, "《2025年德州市政府工作报告》（2025年2月，朱开国作，2025年目标、2024年回顾）")
bullet(doc, "《2025年德州市国民经济和社会发展统计公报》（德州市统计局，2026-04-03，2025年实际数据）")
bullet(doc, "《2026年德州市政府工作报告》（德州市政府，2026年1月，2025年复盘+2026年目标）")
bullet(doc, "德州市政府官网、德州市统计局（dezhou.gov.cn/dztj）")
bullet(doc, "《2025年德州市经济运行情况》（德州市统计局，2026-02-12）")

# ---- 附录B ----
heading1(doc, "附录B：建议建立的年度跟踪仪表盘")
bullet(doc, "GDP总量/增速 vs 全年目标（可比价）。")
bullet(doc, "规上工业增加值及分行业（装备/太阳能/食品/化工）增速。")
bullet(doc, "新能源装机容量/占比、太阳能/光伏发电量。")
bullet(doc, "固定资产投资/工业/基建投资增速。")
bullet(doc, "社会消费品零售总额、以旧换新、网络零售。")
bullet(doc, "进出口、出口/进口、跨境电商。")
bullet(doc, "一般公共预算收入、税收/非税、民生支出占比。")
bullet(doc, "常住人口、城镇化率、城镇新增就业。")
bullet(doc, "CPI、金融存贷款、设备更新贷款。")
bullet(doc, "京津冀落地项目、装备制造订单、太阳城产能。")

# ---------------- 保存 ----------------
out = "/Users/x/Desktop/content-prod-lab/reports/德州市_2025年政府工作报告_深度研究_2026-08-14.docx"
doc.save(out)
print("SAVED OK: 德州市", out)
