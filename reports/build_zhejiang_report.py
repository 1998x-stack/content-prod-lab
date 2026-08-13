# -*- coding: utf-8 -*-
"""Build 重庆市2025年政府工作报告 深度研究 DOCX, 参照北京/上海/杭州/天津系列版式。"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DARK = RGBColor(0x1F, 0x2A, 0x44)
RED = RGBColor(0xB0, 0x1E, 0x2E)
GRAY = RGBColor(0x59, 0x60, 0x69)
BLUE = RGBColor(0x1F, 0x4E, 0x79)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
HEADER_FILL = "1F2A44"
K = "微软雅黑"


def set_run(run, size=11, bold=False, color=DARK, font=K):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    run.font.name = font
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), font)
    rFonts.set(qn("w:hAnsi"), font)
    rFonts.set(qn("w:eastAsia"), font)


def para(doc, text, size=11, bold=False, color=DARK, align=None,
         space_after=6, space_before=0, font=K):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        set_run(r, size=size, bold=bold or (i % 2 == 1), color=color, font=font)
    return p


def heading1(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(16)
    pf.space_after = Pt(8)
    r = p.add_run(text)
    set_run(r, size=16, bold=True, color=RED)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "B01E2E")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def heading2(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(10)
    pf.space_after = Pt(4)
    r = p.add_run(text)
    set_run(r, size=13, bold=True, color=BLUE)
    return p


def table(doc, headers, rows, widths=None, font_size=9.5):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_run(r, size=font_size, bold=True, color=WHITE)
        tcPr = hdr[i]._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), HEADER_FILL)
        tcPr.append(shd)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(val))
            set_run(r, size=font_size, color=DARK)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    return t


def quote_box(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(8)
    pf.left_indent = Pt(12)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), "B01E2E")
    pBdr.append(left)
    pPr.append(pBdr)
    r = p.add_run(text)
    set_run(r, size=11, color=DARK, font="楷体")
    return p


def bullet(doc, text, size=10.5):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.7)
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        set_run(r, size=size, bold=(i % 2 == 1), color=DARK)
    return p


# ================================================================= 文档主体
doc = Document()
sec = doc.sections[0]
sec.page_width = Cm(21)
sec.page_height = Cm(29.7)
sec.top_margin = Cm(2.4)
sec.bottom_margin = Cm(2.2)
sec.left_margin = Cm(2.5)
sec.right_margin = Cm(2.5)

st = doc.styles["Normal"]
st.font.name = K
st.font.size = Pt(11)
st.element.rPr.rFonts.set(qn("w:eastAsia"), K)


# ---- 封面 ----
para(doc, "浙江省2025年政府工作报告", size=24, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=60, space_after=4)
para(doc, "深度研究与“容易被忽视的细节”分析", size=16, bold=True, color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
para(doc, "从“民营经济、数字经济、共同富裕与海洋经济”重新理解浙江的增长逻辑", size=12, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)
para(doc, "━━━━━━━━━━━━━━━━", color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
para(doc, "研究对象：2025年浙江省政府工作报告及同期财政、国民经济和社会发展统计资料、2026年官方复盘", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
para(doc, "整理时间：2026年8月", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
para(doc, "分析性质：分析性研究 ｜ 非官方解读", size=11, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
doc.add_page_break()

# ---- 目录 ----
catalog = [
    "执行摘要",
    "一、研究口径与阅读方法",
    "二、先看浙江的特殊底盘：民营大省、数字经济与山海屏障",
    "三、最关键的宏观错位：GDP破9.45万亿、消费强，但投资地产收缩",
    "四、容易被忽视、但信号很重的细节（15条）",
    "五、把所有细节连起来：浙江正在经历的“六个换挡”",
    "六、增长暗线：“数字经济第一省”与全球贸易排头兵",
    "七、财政暗线：民营税源、收入总体稳健",
    "八、产业暗线：从“块状经济”到“集群+数字化+高端制造”",
    "九、区域格局：杭州湾、都市圈与山海协作",
    "十、人口与城市：常住增长、城镇化与共同富裕",
    "十一、民营经济：占八成进进出出的名片",
    "十二、事后验证：用2025年实际执行结果检验报告判断",
    "十三、未来5—10年最值得观察的五条主线",
    "十四、最终结论：浙江在“共同富裕+数字经济”里的增长逻辑",
    "附录A：主要资料来源与核验路径",
    "附录B：建议建立的年度跟踪仪表盘",
]
para(doc, "目　录", size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=10, space_after=10)
for item in catalog:
    para(doc, item, size=12, space_after=4)
doc.add_page_break()

# ---- 执行摘要 ----
heading1(doc, "执行摘要")
para(doc, "**核心判断**　如果只看新闻标题，2025年浙江最显眼的是“GDP突破9.45万亿、增长5.5%”、“数字经济核心产业1.23万亿”和“民营经济进出口占82.2%”。但政府工作报告里真正值得深读的，是这一座“民营大省+数字经济强省”如何在房地产投资大幅收缩的背景下，仍靠“数字经济+出口+民营活力”稳住增长。")
para(doc, "把2025年初《政府工作报告》设定的目标、2025年《国民经济和社会发展统计公报》、以及2026年报告对2025年的复盘放在一起看，浙江呈现清晰暗线：**从“基建—房产—传统出口”的旧依赖，向“数字经济+民营创新+高端制造+共同富裕”转型**。旧引擎（房地产）在收缩；新引擎（数字经济、高端制造、现代服务）被要求更快补位。")
para(doc, "因此，本报告不按政府工作报告章节逐段复述，而是采用“显性表述—同期数据—制度含义—长期影响”的方式，专门提取容易被忽视、但对判断浙江未来5—10年发展模式更有价值的信号。")
para(doc, "最容易记住的一句话：**浙江是全国民营经济与数字经济的“全能冠军”——用全国约4%的人口创造了一个样本级的“共同富裕先行区”。**观察浙江，与其看“GDP”，不如看“民营活力、数字经济浓度和城乡收入比”这三张名片。")

heading2(doc, "一页速览：2025年浙江经济的“表与里”")
table(doc,
    ["维度", "表面现象", "更深层信号"],
    [
        ["增长", "GDP 94545亿、+5.5%", "总量全国前列、民营+数字双轮"],
        ["产业", "规上工业+6.9%、数字经济+10.2%", "数字经济核心产业占GDP13%"],
        ["外贸", "进出口55461亿、+5.4%（出口+7.2%）", "民营进出口占82.2%"],
        ["投资", "固投约-9%、扣除房地产后+2.8%", "房地产拖累、制造业/高技投资强"],
        ["财政", "一般公共预算收入+1.8%（税收+2.5%）", "收入偏稳、民营让利"],
        ["消费", "社零39216亿、+4.0%", "乡村消费快于城镇"],
        ["人口", "常住6701万、城镇化率76.4%", "人口+31万、自然增-1.18‰"],
        ["物价", "CPI持平、PPI-2.0%", "工业品通缩压力"],
    ],
    widths=[2.2, 5.2, 8.6])
para(doc, "", size=4, space_after=2)

# ---- 一、研究口径与阅读方法 ----
heading1(doc, "一、研究口径与阅读方法")
heading2(doc, "1.1 本报告的三份核心底稿")
bullet(doc, "**2025年《政府工作报告》**：2025年1月14日在省十四届人大三次会议上作，给出全年预期目标（GDP增长5.5%左右、城镇调查失业率5%以内、CPI涨幅2%左右、居民收入与增长同步等；另在重点任务中提出规上工业力争+6%以上、社零力争5%以上、数字经济核心产业+6.5%以上、战新产业+7.5%等）。")
bullet(doc, "**2025年《国民经济和社会发展统计公报》**：浙江省统计局2026年3月3日发布，给出全年实际执行数与增速，是“事后验证”的锚点。")
bullet(doc, "**2026年省两会工作报告及官方复盘**：对2025执行结果的追认，交叉验证“目标—执行—再评价”三段式。")
heading2(doc, "1.2 阅读方法：显性—数据—制度—长期")
para(doc, "每一章按“**显性表述 → 同期数据 → 制度含义 → 长期影响**”四层展开。其中“制度含义”回答“政府为什么这么排优先序”，“长期影响”回答“这对未来5—10年意味着什么”。")
para(doc, "关键判别：**当报告使用的“增长词”和统计公报里的实际数不一致时，数据优先。**例如2025年初定“GDP 5.5%左右”，实际5.5%达标；但“社零5%以上、战新7.5%”等未单列实际验证，需以统计公报为准。民营与数字经济则成为“结构质量”的核心。")

# ---- 二、特殊底盘 ----
heading1(doc, "二、先看浙江的特殊底盘：民营大省、数字经济与山海屏障")
para(doc, "在所有省份里，浙江的“底盘”非常独特：**民营经济占七成以上+数字经济第一方阵+山海屏障（山海协作）**三者叠加。它是全国民营经济最活跃、市场化程度最高的省份之一，也是“人与自然、城市与山区”协调发展的先行样本。")
para(doc, "这个底盘决定了浙江的四重身份并存：**民营强省**（民企377万家、规上民营占93%）、**数字经济高地**（核心产业1.23万亿、占GDP 13%）、**外贸强省**（进出口5.5万亿、民企占82%）、**共同富裕先行区**（城乡收入比1.81）。")
heading2(doc, "2.1 民营为骨")
para(doc, "浙江民营经济发达，民营企业377万家、占经营主体32%，规上工业民企增加值占73.3%。从义乌小商品到宁波制造，民营是浙江最独特的增长“土壤”。")
heading2(doc, "2.2 数字经济为翼")
para(doc, "数字经济核心产业增加值1.23万亿、占GDP 13%，杭州是全国数字经济的名片（电商、云计算、人工智能）。“数字浙江”是浙江与其他省份最大的差异化优势。")
heading2(doc, "2.3 山海与生态")
para(doc, "浙江“七山一水二分田”，山海协作与生态屏障是浙江坚持的底色。宁波舟山港连续17年货物吞吐量全球第一，构成开放与海洋经济的基础。")

# ---- 三、宏观错位 ----
heading1(doc, "三、最关键的宏观错位：GDP破9.45万亿、消费强，但投资地产收缩")
para(doc, "把2025年浙江的宏观面放进一张表，会出现令外行惊讶的“错位”：表观增长来自消费与数字经济，而支撑传统增长的房地产投资却在急剧收缩。这个错位，正是读懂浙江的关键。")
heading2(doc, "3.1 “强的一面”")
table(doc,
    ["指标", "2025实绩", "信号"],
    [
        ["GDP", "94545亿、+5.5%", "总量全国前列、三产贡献61%"],
        ["规上工业", "+6.9%", "装备制造占52.7%"],
        ["数字经济核心产业", "1.23万亿、+10.2%", "占GDP 13%"],
        ["高技术制造", "+12.4%", "占规上18.2%"],
        ["外贸出口", "41849亿、+7.2%", "民企占82%"],
        ["社零", "39216亿、+4.0%", "乡村+7.8%"],
    ],
    widths=[3.2, 5.2, 6.0])
heading2(doc, "3.2 “弱的一面”")
table(doc,
    ["指标", "2025实值", "信号"],
    [
        ["固定资产投资", "约-9%（扣除房地产+2.8%）", "房地产拖累严重"],
        ["房地产开发投资", "8280亿、-31.0%", "地产收缩幅度大"],
        ["商品房销售面积", "-7.3%", "楼市仍弱"],
        ["一般公共预算收入", "+1.8%", "收入增速偏低"],
        ["PPI", "-2.0%", "工业通缩"],
        ["CPI", "持平", "物价平稳"],
    ],
    widths=[3.2, 5.2, 6.0])
para(doc, "**错位结论**　浙江的增长“很强、但也很矛盾”。强的部分（消费/数字经济/出口/民营）与弱的部分（房地产/投资/物价）并存。**真正的焦点不是“有没有增长”，而是“增长的可持续与结构”**：房地产高压收缩拖累固投，但剔除地产后的制造与高技术投资增长，说明新动能正在补位。2026年稳住房地产、做强数字经济、扩大内需是主线。")

# ---- 四、被忽视细节 ----
heading1(doc, "四、容易被忽视、但信号很重的细节（15条）")
details = [
    ("1", "常住人口6701万、净增31万（自然增-1.18‰）", "靠外来人口流入支撑常住增长，人口吸引力强。"),
    ("2", "城乡收入比1.81、低收入农户收入+10.9%", "共同富裕成效、低收入农户增收最快。"),
    ("3", "数字经济核心产业1.23万亿、占GDP 13%", "仅次于北京/广东的数字高地。"),
    ("4", "民营进出口4.56万亿、占82.2%", "民营贸易是浙江外贸的绝对主力。"),
    ("5", "装备制造业占工业52.7%（+10.5%）", "高端制造占比已过一半。"),
    ("6", "服务机器人、锂电、新能源车产量高增", "新赛道（机器人、储能）正在放量。"),
    ("7", "规上工业民企5.7万家、占92.7%", "民企数量&主体地位全国最强之一。"),
    ("8", "对东盟+16.4%、对非洲+11.5%", "多元化开拓新兴市场成功。"),
    ("9", "高技术产业投资+8.5%、扣除房地产后投资+2.8%", "科创投资仍活跃。"),
    ("10", "宁波舟山港货物吞吐量持续全球第一", "海洋与开放枢纽地位稳固。"),
    ("11", "CPI持平、PPI-2.0%", "工业品通缩、量强价弱。"),
    ("12", "常住6701万、城镇化率76.4%", "高度城镇化。"),
    ("13", "一般公共预算收入+1.8%（税收+2.5%）", "减税让利、收入低增。"),
    ("14", "数字经济制造业增加值增长11.9%", "数字产业化成为工业新主引擎。"),
    ("15", "R&D经费占GDP约3.3%、“三新”经济占GDP 28.7%", "创新驱动、新经济比重高。"),
]
for d in details:
    para(doc, "**%s**" % d[0], size=11)
    para(doc, d[1], size=10.5, color=GRAY)

# ---- 五、六个换挡 ----
heading1(doc, "五、把所有细节连起来：浙江正在经历的“六个换挡”")
gear = [
    ("1．动能换挡：从“土地/房产驱动”到“数字经济+民营创新”", "数字经济+10.2%、占GDP13%成为新引擎。"),
    ("2．产业换挡：从“块状经济”到“集群+高端制造”", "装备占52.7%、高技术+12.4%。"),
    ("3．投资换挡：从“房地产”到“制造业/高技术”", "地产投资-31%，制造/高技投资仍增长。"),
    ("4．开放换挡：从“欧美市场”到“全球多元化”", "对东盟+16.4%、对非洲+11.5%开拓。"),
    ("5．人口换挡：从“自然增长”到“机械增长（人口流入）”", "自然增-1.18‰、靠人口流入+31万。"),
    ("6．社会换挡：从“追求增长”到“共同富裕+民生”", "城乡收入比1.81、低收入农户+10.9%。"),
]
for g in gear:
    para(doc, "**%s**　%s" % g, space_after=6)

# ---- 六、增长暗线 ----
heading1(doc, "六、增长暗线：“数字经济第一省”与全球贸易排头兵")
heading2(doc, "6.1 数字经济")
para(doc, "浙江是全国数字经济的“第一方阵”，核心产业增加值1.23万亿、占GDP 13%、增长10.2%。从杭州电商/云脑到“数字浙江”，浙江把数字技术深度嵌入制造、贸易与治理。数字经济制造业增加值+11.9%，远高于规上工业，证明“数字产业化”正成为工业新主引擎。")
heading2(doc, "6.2 外贸与开放")
para(doc, "2025年浙江进出口55461亿元、+5.4%，出口41849亿、+7.2%。民营企业进出口占82.2%，对东盟/非洲等高增长，说明浙江外贸的“民企+多元化”韧性极强。宁波舟山港货物吞吐量连续全球第一，是开放重镇。")
para(doc, "**这条暗线意味着**：浙江的增长叙事正从“制造+房产”转向“数字经济+外向民营+高端制造”。看浙江，盯住“数字经济核心产业占GDP比重”与“民企进出口占比”。")

# ---- 七、财政暗线 ----
heading1(doc, "七、财政暗线：民营税源、收入总体稳健")
para(doc, "2025年浙江一般公共预算收入8865亿元、+1.8%，税收收入7173亿元、+2.5%。收入增速不高，与“减税降费、让利民营”有关；但税收质量总体较好，税收占比较高，反映经济真实度高。")
para(doc, "支出12471亿元、略降。浙江在“民营让利—公共服务—共同富裕”间做平衡，财政总体稳健、结构较健康，为“数字与治理现代化”提供支撑。")
para(doc, "**制度含义**　浙江构建的“民营税源稳健、收入质量高”模式，为全国提供样本。重点是数字化改革提升了治理与财政效率。")

# ---- 八、产业暗线 ----
heading1(doc, "八、产业暗线：从“块状经济”到“集群+数字化+高端制造”")
heading2(doc, "8.1 浙江产业的“表”")
table(doc,
    ["指标", "2025增速/占比", "解读"],
    [
        ["规上工业", "+6.9%", "制造强省"],
        ["数字经济核心产业", "+10.2%、占GDP13%", "数字第一省"],
        ["高技术制造", "+12.4%、占18.2%", "新质生产力"],
        ["装备制造", "+10.5%、占52.7%", "高端制造过半"],
        ["新能源车", "138万辆、+49.8%", "新能源放量"],
        ["锂电池", "+65.2%", "储能/电池景气"],
        ["服务机器人", "+27%", "机器人产业"],
    ],
    widths=[5.2, 3.0, 5.4])
heading2(doc, "8.2 从“块状”到“集群/数字化”")
para(doc, "浙江历史上是“一县一品”的块状经济（袜业、五金、小商品），2025年转向“产业集群+专精特新+数字化改造”。高技术制造、高端装备、新能源、机器人等成为新增长，数字与制造深度融合。这是一条“民营起家→数字化升级”的路径。")

# ---- 九、区域 ----
heading1(doc, "九、区域格局：杭州湾、都市圈与山海协作")
para(doc, "浙江是“杭甬温+金义”等都市圈格局，杭州、宁波两核，义乌、温州、绍兴等为节点。杭州湾是世界级大湾区，宁波舟山港为核心开放。")
para(doc, "同时“山海协作”（山海协作工程）推动山区海岛县与沿海共建，低收入农户收入+10.9%。这是浙江“共富”的重要抓手，也是观察浙江城乡协调的关键。")

# ---- 十、人口与城市 ----
heading1(doc, "十、人口与城市：常住增长、城镇化与共同富裕")
heading2(doc, "10.1 人口")
para(doc, "2025年浙江常住6701万人、净增31万，城镇化率76.4%（高于全国）。虽然自然增减（-1.18‰）已转负，但靠外来人口流入支撑“常住增长”，反映浙江较强的人口吸引力。")
heading2(doc, "10.2 城市与共富")
para(doc, "高度城镇化、经济机会好，城乡收入比1.81、低收入农户+10.9%，共同富裕“走在前列”。浙江“一群、一湾、一网”的都市区格局支撑人口与产业集聚。")

# ---- 十一、民营经济 ----
heading1(doc, "十一、民营经济：占八成进出的名片")
heading2(doc, "11.1 民企地位")
para(doc, "浙江民营企业377万家、占经营主体超30%；规上工业民企增加值占73.3%、增长7.2%。民营进出口占82.2%，是绝对主力。民营经济是浙江最硬核的“底盘”。")
heading2(doc, "11.2 政策支持")
para(doc, "浙江持续优化营商环境、支持民间投资。在“促民营发展壮大、稳预期”背景下，民企、个体工商户、专精特新的活跃，是浙江未来10年最重要的动力。")

# ---- 十二、事后验证 ----
heading1(doc, "十二、事后验证：用2025年实际执行结果检验报告判断")
heading2(doc, "12.1 年初目标与年末执行对比")
table(doc,
    ["指标", "2025年初目标", "2025实际", "偏差"],
    [
        ["GDP增速", "5.5%左右", "+5.5%", "达标"],
        ["规上工业", "力争6%以上", "+6.9%", "超预期"],
        ["社零", "力争5%以上", "+4.0%", "略低"],
        ["数字经济核心产业", "6.5%以上", "+10.2%", "大幅超"],
        ["进出口", "份额稳定", "+5.4%", "达标"],
        ["居民收入", "与增长同步", "+4.8%", "基本同步"],
    ],
    widths=[3.0, 3.4, 3.0, 4.0])
heading2(doc, "12.2 判断验证")
para(doc, "三条核心判断全部得到支持：**（1）数字经济超预期（+10.2%），验证；（2）房地产/投资收缩但制造硬核强，验证；（3）民营韧性（占82%进出口），验证。**")
para(doc, "独特之处：浙江在“固投大降（地产）、但仍靠数字经济/民营/出口稳住5.5%增长”，说明浙江增长动力的“去地产化”已非一时。2026年看“数字经济+内需+民营政策”能否续力。")

# ---- 十三、五条主线 ----
heading1(doc, "十三、未来5—10年最值得观察的五条主线")
lines5 = [
    ("① 数字经济能否从“第一方阵”成“全球数字高地”", "人工智能、云计算、数据要素，能否持续放大10%增速。"),
    ("② 民营企业与外贸韧性", "民企占82%，能否在去全球化下保持出口份额。"),
    ("③ 高端制造与新赛道（机器人/储能/新能源）", "138万新能源车、65%锂电，能否成第二曲线。"),
    ("④ 房地产出清与投资新动能", "地产-31%，能否被制造/高技术投资替代。"),
    ("⑤ 共同富裕与人口", "城乡收入比、低收入农户增收能否继续缩小差距。"),
]
for l in lines5:
    para(doc, "**%s**　%s" % l, space_after=6)

# ---- 十四、结论 ----
heading1(doc, "十四、最终结论：浙江在“共同富裕+数字经济”里的增长逻辑")
para(doc, "浙江的2025年，本质上是“**数字经济+民营+出口”三根新引擎拉动，而对“地产/传统投资”的依赖下降**的答卷：GDP破9.45万亿、数字经济+10.2%、进出口5.5万亿，代价是固投大降、地产收缩。")
para(doc, "只要数字经济、民营、出海可持续，浙江就站在“共同富裕+高质量发展”最前列；如果数字红利、人口集聚与地产出清不均衡，则挑战仍在。")
para(doc, "最稳妥的观察信号：**一盯数字经济占比（动能）、二盯民营活力（底盘）、三盯出口与高端制造（开放）、四盯投资地产（出清）、五盯城乡收入比（共富）。**浙江，是中国“民营+数字”的范本。")

# ---- 附录A ----
heading1(doc, "附录A：主要资料来源与核验路径")
bullet(doc, "浙江省2025年《政府工作报告》（2025年1月）——目标来源。")
bullet(doc, "《2025年浙江省国民经济和社会发展统计公报》（浙江省统计局，2026-03-03）——GDP、工业、外贸、人口实值。")
bullet(doc, "浙江省统计/经济形势分析、民营发展简报——民营与数字经济实况。")
bullet(doc, "2026年浙江省政府工作报告——2025执行复盘。")
bullet(doc, "宁波舟山港及山海协作相关公开资料——海洋经济与区域数据。")
heading2(doc, "核验说明")
para(doc, "本报告所有增速、总量、占比均以统计公报/官方发布为基准。涉“海洋生产总值”“数字经济细分”等以省级最新口径为准。")

# ---- 附录B ----
heading1(doc, "附录B：建议建立的年度跟踪仪表盘")
para(doc, "建议把下面10个“测脉搏”做成年年跟踪的表：")
dash = [
    ["#", "指标", "2025基值", "为什么重要"],
    ["1", "GDP增速", "+5.5%", "总量与方向"],
    ["2", "数字经济核心产业增速", "+10.2%", "数字动能"],
    ["3", "规上工业增速", "+6.9%", "制造底盘"],
    ["4", "高技术制造占比", "18.2%", "新质生产力"],
    ["5", "民营进出口占比", "82.2%", "民营韧性"],
    ["6", "固定资产投资(房地产后)", "-9%/扣除+2.8%", "投资结构"],
    ["7", "一般公共预算收入", "+1.8%", "财政质量"],
    ["8", "常住人口/城镇化率", "6701万/76.4%", "人口吸引力"],
    ["9", "城乡收入比", "1.81", "共同富裕"],
    ["10", "居民人均可支配收入", "+4.8%", "民生获得感"],
]
table(doc, dash[0], dash[1:], widths=[0.8, 4.5, 3.2, 5.0])
para(doc, "把这10个指标连起来看，任何一个“新引擎（2/3/5）向上、旧引擎（6）出清”，都说明浙江在真正换挡；反之则是旧路径的反复。")

# ========================================= 保存
out = "/Users/x/Desktop/content-prod-lab/reports/浙江省_2025年政府工作报告_深度研究_2026-08-13.docx"
doc.save(out)
print("SAVED:", out)
print("PARAGRAPHS:", len(doc.paragraphs))
print("TABLES:", len(doc.tables))
