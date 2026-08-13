# -*- coding: utf-8 -*-
"""Build 青海省2025年政府工作报告 深度研究 DOCX, 参照省系列版式。"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DARK = RGBColor(0x1F, 0x2A, 0x44)
RED = RGBColor(0xB0, 0x1E, 0x2E)
GRAY = RGBColor(0x59, 0x60, 0x69)
BLUE = RGBColor(0x1F, 0x4E, 0x79)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
HEADER_FILL = "1F2A44"
K = "微软雅黑"


def set_run(run, size=11, bold=False, color=DARK, font=K):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    run.font.name = font
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), font)
    rFonts.set(qn("w:hAnsi"), font)
    rFonts.set(qn("w:eastAsia"), font)


def para(doc, text, size=11, bold=False, color=DARK, align=None,
         space_after=6, space_before=0, font=K):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        set_run(r, size=size, bold=bold or (i % 2 == 1), color=color, font=font)
    return p


def heading1(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(16)
    pf.space_after = Pt(8)
    r = p.add_run(text)
    set_run(r, size=16, bold=True, color=RED)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "B01E2E")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def heading2(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(10)
    pf.space_after = Pt(4)
    r = p.add_run(text)
    set_run(r, size=13, bold=True, color=BLUE)
    return p


def table(doc, headers, rows, widths=None, font_size=9.5):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_run(r, size=font_size, bold=True, color=WHITE)
        tcPr = hdr[i]._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), HEADER_FILL)
        tcPr.append(shd)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(val))
            set_run(r, size=font_size, color=DARK)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    return t


def quote_box(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(8)
    pf.left_indent = Pt(12)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), "B01E2E")
    pBdr.append(left)
    pPr.append(pBdr)
    r = p.add_run(text)
    set_run(r, size=11, color=DARK, font="楷体")
    return p


def bullet(doc, text, size=10.5):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.7)
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        set_run(r, size=size, bold=(i % 2 == 1), color=DARK)
    return p


# ================================================================= 文档主体
doc = Document()
sec = doc.sections[0]
sec.page_width = Cm(21)
sec.page_height = Cm(29.7)
sec.top_margin = Cm(2.4)
sec.bottom_margin = Cm(2.2)
sec.left_margin = Cm(2.5)
sec.right_margin = Cm(2.5)

st = doc.styles["Normal"]
st.font.name = K
st.font.size = Pt(11)
st.element.rPr.rFonts.set(qn("w:eastAsia"), K)



# ---- 封面 ----
para(doc, "青海省2025年政府工作报告", size=24, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=60, space_after=4)
para(doc, "深度研究与\u201c容易被忽视的细节\u201d分析", size=16, bold=True, color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
para(doc, "从\u201c三江源生态、盐湖化工、清洁能源与国家生态屏障\u201d重新理解青海", size=12, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)
para(doc, "\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501", color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
para(doc, "研究对象：2025年青海省政府工作报告及同期财政、国民经济和社会发展统计资料、2026年官方复盘", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
para(doc, "标定日期：2026-08-13", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
doc.add_page_break()

# ---- 目录 ----
catalog = [
    "执行摘要",
    "一、研究口径与阅读方法",
    "二、先看青海的特殊底盘：三江源生态、盐湖化工、清洁能源与西部生态屏障",
    "三、最关键的宏观错位：GDP破4100亿、工业强，但总量小/人口/固投/财政偏弱",
    "四、容易被忽视、但信号很重的细节（15条）",
    "五、2025年GDP目标 vs 实际：对照表",
    "六、2025年增长是谁撑起来的？（结构归因）",
    "七、预算与财政的\u201c含金量\u201d",
    "八、民生底账：人口、收入与城乡",
    "九、城镇与农村：格局与均衡",
    "十、人口流入与流出",
    "十一、物价与货币环境",
    "十二、区域一体化：青海在\u201c黄河流域生态保护+兰西城市群+西部大开发\u201d里的位置",
    "十三、未来5—10年最值得观察的五条主线",
    "十四、最终结论：青海在\u201c生态+盐湖+清洁能源\u201d里的增长逻辑",
    "附录A：主要资料来源与核验路径",
    "附录B：建议建立的年度跟踪仪表盘",
]
para(doc, "目　录", size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=10, space_after=10)
for item in catalog:
    para(doc, item, size=12, space_after=4)
doc.add_page_break()

# ---- 执行摘要 ----
heading1(doc, "执行摘要")
para(doc, "**核心判断**　如果只看新闻标题，2025年青海最显著的是\u201cGDP破4121亿、增长4.1%\u201d、\u201c规上工业+7.6%、新材料+39.3%\u201d、\u201c盐湖化工+15.7%、清洁能源发电占比86.6%\u201d、\u201c三江源生态屏障\u201d。但这份研究真正值得深读的，是这座\u201c生态屏障+盐湖+清洁能源\u201d省份，如何在总量小、人口592万、财政下滑背景下，靠\u201c工业（盐湖/新材料/有色）+清洁能源+数字经济\u201d拉动增长。")
para(doc, "把2025年初设定的目标、2025年《国民经济和社会发展统计公报》、2026年复盘放在一起看，青海呈现清晰暗线：**从\u201c盐湖/水电\u201d旧底盘，向\u201c清洁能源+新材料+数字经济生态\u201d升级**。生态与能源是青海最硬的名片。")
para(doc, "因此，本报告不按政府工作报告逐段复述，而采用\u201c显性表述—同期数据—制度含义—长期影响\u201d的方式，专门提取容易被忽略、但对判断青海未来5—10年发展模式更有价值的信号。")
para(doc, "最容易记住的一句话：**青海是\u201c中华水塔+盐湖+清洁能源\u201d的西部省份，靠\u201c盐湖化工/新材料+清洁能源+生态旅游\u201d撑起。**观察青海，与其只看\u201cGDP 4121亿\u201d，不如看\u201c清洁能源86.6%、盐湖、三江源、清洁能源\u201d这几张名片。")
heading2(doc, "一页速览：2025年青海经济的\u201c表与里\u201d")
table(doc,
    ["维度", "表面现象", "更深层信号"],
    [
        ["增长", "GDP 4121.84亿、+4.1%", "一产9.7%、二产41.7%、三产48.6%"],
        ["产业", "规上工业+7.6%", "新材料+39.3%、盐湖+15.7%、有色+21.1%"],
        ["外贸", "进出口70.7亿、+17.6%", "出口+41.9%、一带一路+33.7%"],
        ["投资", "固投-7.9%", "工业-16.4%、房投-20.1%"],
        ["消费", "社零1048.99亿、+2.0%", "体量小、旅游+20.1%"],
        ["人口", "常住592万、城镇化64.7%", "人口-1万、自然增0.68‰"],
        ["能源", "清洁能源发电占86.6%", "生态屏障+沙/水风电"],
    ],
    widths=[2.2, 5.6, 7.4])
para(doc, "", size=10, space_after=2)

# ---- 一、研究口径与阅读方法 ----
heading1(doc, "一、研究口径与阅读方法")
heading2(doc, "1.1 本报告的核心底稿")
bullet(doc, "**2025年《政府工作报告》**（2025年1月）——目标：GDP 4.5%左右、规上工业、固投等。")
bullet(doc, "**《2025年青海省统计公报》**（省统计局2026-02）——GDP、工业、能源、人口实数。")
bullet(doc, "**2026年青海省政府工作报告**（2026年1月）——2025追认/2026目标GDP 4.5%。")
heading2(doc, "1.2 阅读方法：显性—数据—制度—长期")
para(doc, "**关键判别**：数据优先。青海2025年GDP增长4.1%（2026目标4.5%），工业+7.6%亮眼。青海\u201c工业/清洁能源/盐湖强、总量/人口/固投/财政弱\u201d，穿透总量看生态能源牌。")

# ---- 二、底盘 ----
heading1(doc, "二、先看青海的特殊底盘：生态屏障、盐湖化工、清洁能源、生态旅游")
para(doc, "青海的地盘取决于它作为\u201c**中华水塔/三江源+盐湖化工+清洁能源+西部生态**\u201d的特殊定位。它是黄河/长江/澜沧江发源地、国家生态屏障。")
bullet(doc, "**三江源/生态**：三江源国家公园、长江黄河澜沧江源头，生态保护\u201c国之大者\u201d。")
bullet(doc, "**盐湖化工**：察尔汗/柴达木盐湖（锂钾），盐湖化工+15.7%；全国最大的盐湖基地。")
bullet(doc, "**清洁能源**：水/风/光电；清洁能源发电占比86.6%；沙戈荒大基地、大数据中心（绿电算力）。")
bullet(doc, "**西部生态区**：黄河几字弯、祁连山生态、草地牧业。")
para(doc, "这一底板决定青海2025成绩单：**生态保护是根本前提，盐湖+清洁能源+旅游是增长源**；人口/经济总量小是短板。")

# ---- 三、宏观错位 ----
heading1(doc, "三、最关键的宏观错位：GDP破4100亿、工业强，但总量小/人口/固投/财政偏弱")
para(doc, "青海2025年最值得咀嚼的错位，是\u201c**工业/清洁能源/盐湖强、总量小/人口/固投/财政偏弱**\u201d。这种错位决定了对青海经济的观察不能只看GDP增速。")
bullet(doc, "**GDP**：4121.84亿、+4.1%。一产398.51亿（+5.5%、占比9.7%）、二产1719.86亿（+4.7%、占比41.7%）、三产2003.48亿（+3.3%、占比48.6%）。")
bullet(doc, "**工业**：规上工业+7.6%；新材料+39.3%、盐湖化工+15.7%、有色+21.1%、装备+10.7%、高技术+9.5%。")
bullet(doc, "**投资**：固投-7.9%；工业-16.4%、房投-20.1%（投资回落）。")
bullet(doc, "**消费**：社零1048.99亿、+2.0%（体量小）；旅游游客+20.1%。")
bullet(doc, "**外贸**：进出口70.7亿、+17.6%（出口+41.9%）；一带一路+33.7%。")
bullet(doc, "**财政**：总收入575.27亿、-3.8%；地方354.95亿、-4.2%（财政下滑）。")
para(doc, "**为什么读这条**：青海作为\u201c人口小省+生态大区\u201d，结构性矛盾是\u201c工业/能源/旅游强、总量/人口/固投/财政弱\u201d。稳经济靠工业与生态转型。")

# ---- 四、15条细节 ----
heading1(doc, "四、容易被忽视、但信号很重的细节（15条）")
para(doc, "以下15条，都在2025年报告/公报，但被\u201cGDP 4100亿\u201d等掩盖。它们是判断青海2025之后5—10年的关键小信号。")
bullet(doc, "**1. 清洁能源发电占比86.6%**：水风光电主力，全国能源绿色名片。")
bullet(doc, "**2. 盐湖化工+15.7%、有色+21.1%**：盐湖锂钾/有色，资源高景气。")
bullet(doc, "**3. 新材料+39.3%**：新材料产业爆发的第一增长极。")
bullet(doc, "**4. 三江源国家公园**：生态屏障、\u201c中华水塔\u201d。")
bullet(doc, "**5. 旅游游客+20.1%、入境+209.1%**：生态/文旅爆发。")
bullet(doc, "**6. 绿电算力/大数据中心**：清洁能源+数据中心（东数西算/青海）。")
bullet(doc, "**7. 出口+41.9%、一带一路+33.7%**：外向开放。")
bullet(doc, "**8. 规上工业利润337.08亿（制造+56.1%）**：盐湖/新材料盈利强。")
bullet(doc, "**9. 常住592万、城镇化64.7%（+0.8pct）**：人口总量小、城市微增。")
bullet(doc, "**10. 空气优良率97.9%、地表水100%**：生态质量全国领先。")
bullet(doc, "**11. 兰西城市群**：西宁-兰州协同、西部城市群节点。")
bullet(doc, "**12. 农牧/草牧**：粮食/牧业稳定性。")
bullet(doc, "**13. 数字经济/清洁能源基地**：绿电+算力试点。")
bullet(doc, "**14. 中央财政/转移支付**：青海财政转移依赖大。")
bullet(doc, "**15. 科技创新/盐湖研究院**：盐湖锂循环。")
para(doc, "", size=10, space_after=2)

# ---- 五、对照表 ----
heading1(doc, "五、2025年GDP目标 vs 实际：对照表")
para(doc, "2025年报告中列出的主要预期目标，与2025年实际完成（2026目标为据）：")
tb = [
    ["指标", "2026年目标", "2025年实际（公报）", "达标/方向"],
    ["GDP增速", "4.5%左右", "+4.1%（4121.84亿）", "偏低"],
    ["规上工业增加值", "—", "+7.6%", "亮眼"],
    ["社会消费品零售总额", "—", "+2.0%（1048.99亿）", "偏弱"],
    ["进出口", "—", "+17.6%（70.7亿）", "超预期"],
    ["地方一般公共预算收入", "—", "-4.2%（354.95亿）", "下滑"],
    ["城镇新增就业", "6万人以上", "—", "—"],
]
table(doc, tb[0], tb[1:], widths=[3.2, 3.2, 4.4, 4.4])
para(doc, "", size=10, space_after=2)
para(doc, "**要点**：青海GDP+4.1%偏低、工业/旅游/进出口好、财政下滑——总量动能待修复但结构亮。")

# ---- 六、结构归因 ----
heading1(doc, "六、2025年增长是谁撑起来的？（结构归因）")
heading2(doc, "6.1 三次产业：工业强、三产稳")
para(doc, "**二产（+4.7%）**驱动、三产+3.3%、一产+5.5%。规上工业+7.6%。")
heading2(doc, "6.2 工业：盐湖/新材料/有色")
para(doc, "规上工业+7.6%；新材料+39.3%、有色+21.1%、盐湖+15.7%；利润337亿+。")
heading2(doc, "6.3 消费小、旅游旺")
para(doc, "社零+2.0%体量小；旅游+20%、入境爆发。")
heading2(doc, "6.4 投资/财政弱")
para(doc, "固投-7.9%、财政-3.8%——投资与财政承压。")
heading2(doc, "6.5 外贸体量小")
para(doc, "进出口70.7亿、+17.6%（出口+41.9%），弹性大而规模小。")
para(doc, "**一句话归因**：青海2025年\u201c**盐湖/新材料/清洁能源工业+旅游**\u201d拉动，\u201c固投/财政\u201d压力大——生态+资源型的西部增长模式。")

# ---- 七、财政 ----
heading1(doc, "七、预算与财政的\u201c含金量\u201d")
para(doc, "**一般公共预算总收入575.27亿元、-3.8%**；地方354.95亿元、-4.2%（税254.54亿、-3.4%）；一般公共预算支出2164.19亿元、与上年持平。")
bullet(doc, "**收入下滑**：-3.8%/-4.2%，主要受大宗商品/资源税收/企业利润拖累；高度依赖中央转移支付。")
bullet(doc, "**支出高企**：支出2164亿、约为地方收入的6倍，财政高度依赖中央。")
bullet(doc, "**民生/生态**：三江源/生态补偿、民生支出刚性。")
para(doc, "**财政含义**：青海\u201c收入下滑、转移依赖大\u201d，靠中央支持+资源税收支撑，需盐湖/能源/旅游创造税源。")

# ---- 八、民生 ----
heading1(doc, "八、民生底账：人口、收入与城乡")
para(doc, "**常住人口592万人（-1万）、城镇化率64.7%（+0.8pct）**；居民人均可支配收入31661元、+5.1%（农村+6.0%）。")
bullet(doc, "**收入**：居民31661元、+5.1%；城镇43879元、+4.0%，农村17721元、+6.0%（城乡比2.48、缩小）。")
bullet(doc, "**就业**：城镇新增就业（2026目标6万人以上）；失业率5.5%左右。")
bullet(doc, "**社保**：养老参保458.24万、医保547.97万；低保健全。")
bullet(doc, "**人口**：常住592万、城镇化微升；自然增0.68‰。")
para(doc, "**民生含义**：青海\u201c收入增速高于全国、城乡缩小\u201d，但总量小、人口少，靠转移支付保民生。")

# ---- 九、城乡 ----
heading1(doc, "九、城镇与农村：格局与均衡")
para(doc, "**城镇化率64.7%（+0.8pct）**，青海城镇化低于全国但提升；城乡收入比2.48、缩小。")
bullet(doc, "**城市**：西宁（兰西城市群）、格尔木（盐湖/能源）。")
bullet(doc, "**农村**：农牧业（粮食110万吨、牧区）；生态补偿/乡村振兴。")
para(doc, "**城乡均衡**：青海\u201c城市/县域差异大\u201d，生态补偿+特色农业带动乡村。")

# ---- 十、人口流 ----
heading1(doc, "十、人口流入与流出")
para(doc, "**青海常住592万、-1万**：人口总体平稳略降、生态区人口少，依赖西宁城市群。")
bullet(doc, "**结构**：牧区向城市集中、生态移民，人口总量小。")
bullet(doc, "**竞争**：劳动力外流倾向（相对沿海），靠盐湖/能源/旅游留人。")
para(doc, "人口方向决定中长期需求；青海的人口红利有限，靠绿色产业/算力创造岗位。")

# ---- 十一、物价 ----
heading1(doc, "十一、物价与货币环境")
para(doc, "**2025年青海CPI下降0.1%**（城市-0.1、农村-0.2）；PPI-1.6%、购进-2.7%。")
bullet(doc, "**物价**：CPI低位/微降；PPI负（资源/工业品价格）。")
bullet(doc, "**货币/流动性**：本外币存款8643.3亿+4.4%、贷款+3.7%。")
para(doc, "**物价含义**：青海\u201c温和/通缩压力\u201d，内需偏弱，关注资源价格。")

# ---- 十二、区域一体化 ----
heading1(doc, "十二、区域一体化：青海在\u201c黄河流域生态保护+兰西城市群+西部大开发\u201d里的位置")
para(doc, "青海处于**黄河流域生态保护+兰西城市群+西部大开发**交汇：既是\u201c中华水塔\u201d/三江源生态屏障，也是向西开放/能源要地。")
bullet(doc, "**黄河流域生态**：黄河源头保护、水土保持。")
bullet(doc, "**兰西城市群**：西宁-兰州协同、西部唯一城市群节点。")
bullet(doc, "**西部大开发/一带一路**：清洁能源外送、绿电算力、格库铁路等。")
bullet(doc, "**生态补偿/转移支付**：国家财政与碳汇支持。")
para(doc, "**区域含义**：青海作为\u201c生态屏障+能源要地\u201d，靠中央支持+清洁能源+兰西城市群融入西部。")

# ---- 十三、五条主线 ----
heading1(doc, "十三、未来5—10年最值得观察的五条主线")
bullet(doc, "**主线1｜清洁能源**：清洁能源发电86.6%、沙戈荒大基地。能否建成全国绿电/算力高地。")
bullet(doc, "**主线2｜盐湖化工/新材料**：盐湖锂钾、新材料+39.3%。能否做全球锂/钾盐基地。")
bullet(doc, "**主线3｜生态旅游/低碳**：三江源、入境+209%。能否把生态变旅游钱。")
bullet(doc, "**主线4｜绿电算力/数字经济**：东数西算、数据中心。能否把绿电变算力增长极。")
bullet(doc, "**主线5｜民生/人口**：转移支付、乡村振兴。能否靠民生稳人口。")
para(doc, "这五条，是青海从\u201c生态屏障+资源省份\u201d走向\u201c清洁能源+盐湖+生态旅游强省\u201d的\u201c主赛道\u201d。")

# ---- 十四、结论 ----
heading1(doc, "十四、最终结论：青海在\u201c生态+盐湖+清洁能源\u201d里的增长逻辑")
para(doc, "青海2025年，本质上是\u201c**盐湖/新材料/清洁能源工业+旅游驱动、总量小/固投/财政偏弱**\u201d的答卷：GDP4121.84亿、+4.1%、规上工业+7.6%、清洁能源86.6%、盐湖+15.7%、新材料+39.3%、进出口+17.6%、常住592万。")
para(doc, "只要清洁能源/盐湖/生态旅游持续，青海就站在\u201c绿电+盐湖+生态\u201d增长极；若人口/总量/财政收入受限，需承受\u201c规模小、投资/财政弱\u201d的结构挑战。")
para(doc, "最稳观察信号：**一盯清洁能源发电/装机（引擎）、二盯盐湖/新材料（资源）、三盯生态旅游（文旅）、四盯绿电算力（产业）、五盯中央财政/民生（支撑）。**青海，是\u201c生态+能源+盐湖\u201d西部新样本。")

# ---- 附录A ----
heading1(doc, "附录A：主要资料来源与核验路径")
bullet(doc, "青海省2025年政府工作报告（2025年1月）——目标来源。")
bullet(doc, "《2025年青海省国民经济和社会发展统计公报》——GDP、工业、能源、人口实数。")
bullet(doc, "2026年青海省政府工作报告（2026年1月）——2025追认/2026目标/生态/盐湖/清洁能源。")
bullet(doc, "省财政厅、省统计局、三江源国家公园管理局。")
heading2(doc, "核验说明")
para(doc, "本报告涉及数据以统计公报/官方口径为准；\u201c生态/盐湖/清洁能源\u201d等以官方口径为准。")

# ---- 附录B ----
heading1(doc, "附录B：建议建立的年度跟踪仪表盘")
para(doc, "建议把下面10个\u201c测脉搏\u201d指标做成年度跟踪表：")
dash = [
    ["#", "指标", "2025基值", "为什么重要"],
    ["1", "GDP增速", "+4.1%（4121.84亿）", "总量与方向"],
    ["2", "规模以上工业增速", "+7.6%", "工业底盘"],
    ["3", "清洁能源发电占比", "86.6%", "绿色能源"],
    ["4", "盐湖化工/新材料", "+15.7% / +39.3%", "资源产业"],
    ["5", "进出口/出口", "+17.6% / +41.9%", "开放/外贸"],
    ["6", "固定资产投资", "-7.9%", "投资结构"],
    ["7", "社零增速", "+2.0%（1048.99亿）", "内需消费"],
    ["8", "常住人口/城镇化", "592万 / 64.7%", "人口与城市"],
    ["9", "地方财政收入", "-4.2%（354.95亿）", "财政质量"],
    ["10", "CPI", "-0.1%", "物价"],
]
table(doc, dash[0], dash[1:], widths=[0.8, 4.6, 4.4, 4.0])
para(doc, "把这10个指标连起来看，清洁能源/盐湖/新材料（2/3/4）、固投/财政（6/9），都说明青海在真正换挡。")

# ============ 保存 ============
out = "/Users/x/Desktop/content-prod-lab/reports/青海省_2025年政府工作报告_深度研究_2026-08-13.docx"
doc.save(out)
print("SAVED:", out)
print("PARAGRAPHS:", len(doc.paragraphs))
print("TABLES:", len(doc.tables))
