# -*- coding: utf-8 -*-
"""Build 重庆市2025年政府工作报告 深度研究 DOCX, 参照北京/上海/杭州/天津系列版式。"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DARK = RGBColor(0x1F, 0x2A, 0x44)
RED = RGBColor(0xB0, 0x1E, 0x2E)
GRAY = RGBColor(0x59, 0x60, 0x69)
BLUE = RGBColor(0x1F, 0x4E, 0x79)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
HEADER_FILL = "1F2A44"
K = "微软雅黑"


def set_run(run, size=11, bold=False, color=DARK, font=K):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    run.font.name = font
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), font)
    rFonts.set(qn("w:hAnsi"), font)
    rFonts.set(qn("w:eastAsia"), font)


def para(doc, text, size=11, bold=False, color=DARK, align=None,
         space_after=6, space_before=0, font=K):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        set_run(r, size=size, bold=bold or (i % 2 == 1), color=color, font=font)
    return p


def heading1(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(16)
    pf.space_after = Pt(8)
    r = p.add_run(text)
    set_run(r, size=16, bold=True, color=RED)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "B01E2E")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def heading2(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(10)
    pf.space_after = Pt(4)
    r = p.add_run(text)
    set_run(r, size=13, bold=True, color=BLUE)
    return p


def table(doc, headers, rows, widths=None, font_size=9.5):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_run(r, size=font_size, bold=True, color=WHITE)
        tcPr = hdr[i]._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), HEADER_FILL)
        tcPr.append(shd)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(val))
            set_run(r, size=font_size, color=DARK)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    return t


def quote_box(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(8)
    pf.left_indent = Pt(12)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), "B01E2E")
    pBdr.append(left)
    pPr.append(pBdr)
    r = p.add_run(text)
    set_run(r, size=11, color=DARK, font="楷体")
    return p


def bullet(doc, text, size=10.5):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.7)
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        set_run(r, size=size, bold=(i % 2 == 1), color=DARK)
    return p


# ================================================================= 文档主体
doc = Document()
sec = doc.sections[0]
sec.page_width = Cm(21)
sec.page_height = Cm(29.7)
sec.top_margin = Cm(2.4)
sec.bottom_margin = Cm(2.2)
sec.left_margin = Cm(2.5)
sec.right_margin = Cm(2.5)

st = doc.styles["Normal"]
st.font.name = K
st.font.size = Pt(11)
st.element.rPr.rFonts.set(qn("w:eastAsia"), K)


# ---- 封面 ----
para(doc, "内蒙古自治区2025年政府工作报告", size=24, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=60, space_after=4)
para(doc, "深度研究与“容易被忽视的细节”分析", size=16, bold=True, color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
para(doc, "从“能源重镇、绿色能源、稀土、乳业与草原生态”重新理解内蒙古", size=12, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)
para(doc, "━━━━━━━━━━━━━━━━", color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
para(doc, "研究对象：2025年内蒙古自治区政府工作报告及同期财政、国民经济和社会发展统计资料、2026年官方复盘", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
para(doc, "整理时间：2026年8月", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
para(doc, "分析性质：分析性研究 ｜ 非官方解读", size=11, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
doc.add_page_break()

# ---- 目录 ----
catalog = [
    "执行摘要",
    "一、研究口径与阅读方法",
    "二、先看内蒙古的特殊底盘：能源大区、奶业与稀土、草原屏障",
    "三、最关键的宏观错位：GDP破2.67万亿、制造业强，但煤价/PPI与财政偏弱",
    "四、容易被忽视、但信号很重的细节（15条）",
    "五、把所有细节连起来：内蒙古正在经历的“六个换挡”",
    "六、增长暗线：绿色能源（风电/光伏）、稀土/新材料与新质生产力",
    "七、财政暗线：收入下滑、税收偏弱、能源价格拖累",
    "八、产业暗线：从“产煤大区”到“能源+稀土+高端制造”",
    "九、区域格局：重点城市群与“兴边/向北开放”",
    "十、人口与生态：人口负增、草原生态屏障",
    "十一、民营经济：民间投资与市场主体",
    "十二、事后验证：用2025年实际执行结果检验报告判断",
    "十三、未来5—10年最值得观察的五条主线",
    "十四、最终结论：内蒙古在“能源转型+富民兴边”里的增长逻辑",
    "附录A：主要资料来源与核验路径",
    "附录B：建议建立的年度跟踪仪表盘",
]
para(doc, "目　录", size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=10, space_after=10)
for item in catalog:
    para(doc, item, size=12, space_after=4)
doc.add_page_break()

# ---- 执行摘要 ----
heading1(doc, "执行摘要")
para(doc, "**核心判断**　如果只看新闻标题，2025年内蒙古最显眼的是“GDP突破2.67万亿、增长4.7%”、“规上工业+6.7%”、“新能源装机（风电破1亿千瓦）”。但这份研究真正值得深读的，是这块“煤炭、风电、光伏、稀土、乳业”的能源与资源重镇如何在煤价下行、PPI大幅下滑（-7.0%）与一般公共预算收入负增长（-4.6%）的背景下，向“新能源+新材料+高端制造”转型。")
para(doc, "把2025年初设定的目标（GDP增长6%左右）、2025年《国民经济和社会发展统计公报》、以及2026年报告对2025年的复盘放在一起看，内蒙古呈现清晰暗线：**从“产煤、改能源+卖能源”的旧依赖，向“绿色能源+稀土/新材料+高端制造+富民”转型**。旧引擎（煤、电力、价格）在波动；新引擎（风电光伏、新能源车相关、稀土/新材料、装备制造、消费）被要求更快补位。这也是“一产（奶/农）+二产（能源制造）+三产（生态/旅游）”协调、处在“兴边+生态屏障”双重定位的省份。")
para(doc, "因此，本报告不按政府工作报告逐段复述，而采用“显性表述—同期数据—制度含义—长期影响”的方式，专门提取容易被忽视、但对判断内蒙古未来5—10年发展模式更有价值的信号。")
para(doc, "最容易记住的一句话：**内蒙古是全国最典型的“能源+资源+草原”省份，一半靠煤电油、一半要“绿能+稀土+生态”，在“抓大（能源）与求巧（转型）”“经济增长与生态保护”间找平衡。**观察内蒙古，与其看“总量GDP 2.67万亿”，不如看“新能源装机、稀土/新材料、煤价对财政的传导、乳业与草原生态”这几张名片。")
heading2(doc, "一页速览：2025年内蒙古经济的“表与里”")
table(doc,
    ["维度", "表面现象", "更深层信号"],
    [
        ["增长", "GDP 26710.3亿、+4.7%", "第二产业贡献46%、能源盘"],
        ["产业", "规上工业+6.7%、制造业+9.1%", "高端制造/新能源提速"],
        ["外贸", "进出口2209.6亿、+6.5%（出口+5.4%）", "向北开放、一带一路+3.1%"],
        ["投资", "固定资产投资+4.0%、工业+7.3%", "基建投资+8.9%、民间+8.8%"],
        ["财政", "一般公共预算收入-4.6%（税收-4.4%）", "煤价/PPI拖累财政"],
        ["消费", "社零5375.5亿、+4.3%", "乡村/线上增速快"],
        ["人口", "常住2374万、城镇化率71.48%", "自然增-4.58‰、人口下降"],
        ["物价", "CPI-0.1%、PPI-7.0%", "工业品通缩严重"],
    ],
    widths=[2.2, 5.2, 8.6])
para(doc, "", size=4, space_after=2)

# ---- 一、研究口径与阅读方法 ----
heading1(doc, "一、研究口径与阅读方法")
heading2(doc, "1.1 本报告的三份核心底稿")
bullet(doc, "**2025年《政府工作报告》**：2025年1月在自治区十四届人大三次会议上作，给出全年预期目标（GDP增长6%左右、规上工业增长7%左右等）。固定资产投资/进出口等未全部设具体速率。")
bullet(doc, "**2025年《国民经济和社会发展统计公报》**：内蒙古自治区统计局2026年4月9日发布，给出全年实际执行数与增速，是“事后验证”的锚点。")
bullet(doc, "**2026年内蒙古自治区政府工作报告及官方复盘**：对2025执行结果的追认，交叉验证“目标—执行—再评价”三段式。")
heading2(doc, "1.2 阅读方法：显性—数据—制度—长期")
para(doc, "每一章按“**显性表述 → 同期数据 → 制度含义 → 长期影响**”四层展开。其中“制度含义”回答“政府为什么这么排优先序”，“长期影响”回答“这对未来5—10年意味着什么”。")
para(doc, "关键判别：**当报告使用的“增长词”和统计公报里的实际数不一致时，数据优先。**例如2025年GDP目标6%左右、实际4.7%；规上工业目标7%、实际6.7%。而一般公共预算收入-4.6%（能源价格拖累）。差异反映：内蒙古“能源/制造业强，煤价/PPI/财政弱”。")

# ---- 二、特殊底盘 ----
heading1(doc, "二、先看内蒙古的特殊底盘：能源大区、奶业与稀土、草原屏障")
para(doc, "在所有省份里，内蒙古的“底盘”独特：**全国能源大区（煤炭/风电/光伏）+奶业与稀土+草原生态屏障**三合一。面积全国第三、常住2374万，是国家重要的能源保障、粮/奶基地与生态屏障。")
para(doc, "这决定内蒙古的多重身份并存：**能源重镇**（煤/电/风电/光伏）、**资源大区**（稀土、煤、奶）、**奶业与农畜**（奶业全国前列）、**开放沿边**（向北开放、中蒙边境、一带一路）、**生态屏障**（草原/沙地/森林）。")
heading2(doc, "2.1 能源与“绿能”")
para(doc, "内蒙古是煤炭第一、风电/光伏重地，风电装机破1亿千瓦、占全部超52%。新能源及相关产业+18%，既保障能源安全，也走“绿色能源”转型。")
heading2(doc, "2.2 稀土/新材料")
para(doc, "内蒙古是全球稀土主产区（包头），稀土产业增加值+32.0%。稀土/新材料是国家战略资源与内蒙高端制造的增量。")
heading2(doc, "2.3 奶业与生态")
para(doc, "牛奶产量758.5万吨、乳制品473.6万吨。同时作为生态屏障，草原、环保、碳汇是内蒙重要定位。")

# ---- 三、宏观错位 ----
heading1(doc, "三、最关键的宏观错位：GDP破2.67万亿、制造业强，但煤价/PPI与财政偏弱")
para(doc, "把2025年内蒙古的宏观面放进一张表，会出现令人惊讶的“错位”：表观增长来自二产（能源制造），而PPI/财政/价格大幅走弱。这个错位，正是读懂内蒙的关键。")
heading2(doc, "3.1 “强的一面”")
table(doc,
    ["指标", "2025实绩", "信号"],
    [
        ["GDP", "26710.3亿、+4.7%", "二产贡献46%"],
        ["规上工业", "+6.7%", "制造业+9.1%"],
        ["装备制造", "+20.6%", "高端制造"],
        ["高技术制造", "+17.5%", "新质生产力"],
        ["新能源及产业", "+18.0%", "绿色能源"],
        ["进出口", "+6.5%", "向北开放"],
    ],
    widths=[3.2, 5.4, 6.0])
heading2(doc, "3.2 “偏弱的一面”")
table(doc,
    ["指标", "2025实值", "信号"],
    [
        ["PPI", "-7.0%", "工业品通缩严重"],
        ["一般公共预算收入", "-4.6%", "财政收入下滑"],
        ["税收收入", "-4.4%", "税收下降"],
        ["房地产开发投资", "-7.3%", "地产调整"],
        ["CPI", "-0.1%", "需求偏弱"],
    ],
    widths=[3.2, 5.4, 6.0])
para(doc, "**错位结论**　内蒙古的增长“很强、但也很矛盾”。强的部分（制造业/新能源/高技术/稀土）与弱的部分（PPI/财政/房地产）并存。**真正的焦点是“量增价跌、财政受煤价拖累”**：制造业强但PPI-7%，是一般工业品通缩对能源/资源大省的典型冲击。2026年“绿色能源+高端制造+稳价格/财政收入”是重点。")

# ---- 四、被忽视细节 ----
heading1(doc, "四、容易被忽视、但信号很重的细节（15条）")
details = [
    ("1", "规上工业+6.7%、制造业+9.1%", "制造升级。"),
    ("2", "装备制造+20.6%、高技术制造+17.5%", "高端/高技术高增。"),
    ("3", "新能源及产业+18.0%、风电装机破1亿千瓦", "绿色能源爆发。"),
    ("4", "太阳能/风电压正、装机占比超52%", "新能源主导装机。"),
    ("5", "稀土产业增加值+32.0%", "稀土战略/高端。"),
    ("6", "原煤产量12.87亿吨、-1.0%", "煤量仍稳、价格跌。"),
    ("7", "多晶硅+34.1%、单晶硅+9.8%", "硅料/新能源上游放量。"),
    ("8", "进出口2209.6亿、+6.5%、出口+5.4%", "向北开放。"),
    ("9", "固定资产投资+4.0%、民间+8.8%、基建+8.9%", "民间/基建投资强。"),
    ("10", "制造业投资-7.3%、地产投资-7.3%", "制造业/地产投资回落。"),
    ("11", "社零5375.5亿、+4.3%、乡村/线上快", "内需/消费升级。"),
    ("12", "牛奶758.5万吨、乳制品473.6万吨", "奶业/奶源地。"),
    ("13", "常住2374万、城镇化率71.48%、自然增-4.58‰", "人口负增。"),
    ("14", "居民人均可支配收入41921元、+4.6%", "收入稳、城镇/农村协调。"),
    ("15", "一般公共预算收入-4.6%、税收-4.4%", "能源价格拖累财政。"),
]
for d in details:
    para(doc, "**%s**" % d[0], size=11)
    para(doc, d[1], size=10.5, color=GRAY)

# ---- 五、六个换挡 ----
heading1(doc, "五、把所有细节连起来：内蒙古正在经历的“六个换挡”")
gear = [
    ("1．动能换挡：从“卖煤/卖电”到“绿色能源+高端制造”", "新能源+18%、装备+20.6%、稀土+32%。"),
    ("2．产业换挡：从“产煤大区”到“新能源+稀土/新材料+高端制造”", "风电/光伏放量、多晶硅/稀土。"),
    ("3．投资换挡：从“地产/传统基建”到“新能源/制造业/基建”", "基建+8.9%、民间+8.8%、地产-7.3%。"),
    ("4．开放换挡：从“内贸”到“向北开放/一带一路”", "进出口+6.5%、一带一路+3.1%。"),
    ("5．生态换挡：从“开发”到“草原保护/碳汇”", "生态屏障、草原修复、绿色。"),
    ("6．动能换挡：从“煤价驱动”到“能源+生态+富民”", "价格波动、但产业多元。"),
]
for g in gear:
    para(doc, "**%s**　%s" % g, space_after=6)

# ---- 六、增长暗线 ----
heading1(doc, "六、增长暗线：绿色能源（风电/光伏）、稀土/新材料与新质生产力")
heading2(doc, "6.1 绿色能源")
para(doc, "内蒙古风电装机破1亿千瓦、新能源占比超52%，新能源及产业+18%。在“双碳+绿能”下，内蒙从“煤电”转向“风/光/储能+绿电外送”。")
heading2(doc, "6.2 稀土/新材料+高端制造")
para(doc, "包头稀土、多晶硅/单晶硅、新能源车相关等，稀土产业+32%、装备+20.6%、高技术+17.5%。稀土/新材料+高端制造，是内蒙从“卖资源”到“造材料”的增量。")
para(doc, "**这条暗线意味着**：内蒙古的增长叙事正从“产煤/卖电”迈向“绿色能源+稀土/新材料+高端制造”。看内蒙，盯住“新能源装机、稀土/新材料、装备/高技术增速、煤价对财政传导”这几组数。")

# ---- 七、财政暗线 ----
heading1(doc, "七、财政暗线：收入下滑、税收偏弱、能源价格拖累")
para(doc, "2025年内蒙古一般公共预算收入3005.2亿、-4.6%，税收2115.6亿、-4.4%。煤价/PPI高、价格下行直接拖累能源税基、财政收入。支出向基建、民生、生态倾斜。")
para(doc, "**制度含义**　内蒙财政高度依赖能源（煤电税源），能源价格波动直接冲击收入。长期要靠“新能源/稀土/高端制造”带来的新税源，降低对单一煤价的依赖。")

# ---- 八、产业暗线 ----
heading1(doc, "八、产业暗线：从“产煤大区”到“能源+稀土+高端制造”")
heading2(doc, "8.1 内蒙古的“表”")
table(doc,
    ["指标", "2025增速/信号", "解读"],
    [
        ["规上工业", "+6.7%", "制造升级"],
        ["装备制造", "+20.6%", "高端制造"],
        ["高技术制造", "+17.5%", "新质生产力"],
        ["新能源及产业", "+18.0%", "绿色能源"],
        ["稀土产业", "+32.0%", "稀土战略"],
        ["多晶硅/单晶硅", "+34.1%/+9.8%", "新能源上游"],
        ["原煤产量", "12.87亿吨、-1.0%", "煤量稳、价跌"],
    ],
    widths=[4.6, 3.6, 5.0])
heading2(doc, "8.2 从“煤”到“能/科/高端”")
para(doc, "内蒙古过去是“煤/电大户”，2025年“风电光伏+稀土/新材料+装备制造+高技术”正在崛起。新能源+高端制造，是内蒙“从资源到产业”的未来。")

# ---- 九、区域 ----
heading1(doc, "九、区域格局：重点城市群与“兴边/向北开放”")
para(doc, "呼和浩特、包头、鄂尔多斯构成“呼包鄂”城市群：呼和浩特（乳业/大数据）、包头（稀土/制造）、鄂尔多斯（能源/煤化工），是内蒙经济核心。")
para(doc, "作为“向北开放”前沿，内蒙古对接蒙古、俄罗斯、中蒙经济走廊、一带一路，是沿边开放+边疆发展的节点。")

# ---- 十、人口与生态 ----
heading1(doc, "十、人口与生态：人口负增、草原生态屏障")
heading2(doc, "10.1 人口")
para(doc, "2025年末内蒙古常住2374万、城镇化率71.48%、自然增-4.58‰。人口负增，与全国人口变化一致。")
heading2(doc, "10.2 生态屏障")
para(doc, "内蒙古是北方生态屏障，肩负草原、沙地、森林保护与碳汇。生态保护与能源开发要平衡，是“兴边+绿色”关键。")

# ---- 十一、民营经济 ----
heading1(doc, "十一、民营经济：民间投资与市场主体")
heading2(doc, "11.1 民间投资")
para(doc, "2025年内蒙古民间投资+8.8%（扣除房地产后更高），民间主体在能源、新能源、消费、贸易上活跃。")
heading2(doc, "11.2 政策/生态")
para(doc, "内蒙古优化营商环境、支持民营/进入新能源、推动奶业/农畜。民营是能源系+非煤产业的新活力。")

# ---- 十二、事后验证 ----
heading1(doc, "十二、事后验证：用2025年实际执行结果检验报告判断")
heading2(doc, "12.1 年初目标与年末执行对比")
table(doc,
    ["指标", "2025年初目标", "2025实际", "偏差"],
    [
        ["GDP增速", "6%左右", "+4.7%", "低于目标"],
        ["规上工业", "7%左右", "+6.7%", "略低"],
        ["固定资产投资", "未设", "+4.0%", "略低"],
        ["进出口", "未设", "+6.5%", "达标"],
        ["一般公共预算收入", "未设", "-4.6%", "下滑"],
        ["居民收入", "与增长同步", "+4.6%", "同步"],
    ],
    widths=[3.0, 3.0, 3.0, 4.4])
heading2(doc, "12.2 判断验证")
para(doc, "三条核心判断得到支持：**（1）制造业/新能源/稀土强（+9.1%/+18%/+32%），验证；（2）煤价/PPI/财政弱（PPI-7%、财政-4.6%），验证；（3）向北开放/生态/民间强，验证。**")
para(doc, "核心观察：内蒙古靠“能源+绿色+制造”守住4.7%增长，但煤价/PPI通缩直接冲击财政与收入。2026年“绿色能源+高端制造+稳价格/财政”是重点。")

# ---- 十三、五条主线 ----
heading1(doc, "十三、未来5—10年最值得观察的五条主线")
lines5 = [
    ("① 绿色能源（绿电外送/储能）", "风电光伏能否更高效、储能/绿电消纳。"),
    ("② 稀土/新材料/高端制造", "能否把稀土、硅料做成全球新材料中心。"),
    ("③ 煤价/能源价格与财政", "能否减少单一煤价冲击、财政转型。"),
    ("④ 奶业/农畜/富民", "奶业、农畜、兴边富民。"),
    ("⑤ 生态屏障与碳汇", "草原保护、碳汇、生态价值变现。"),
]
for l in lines5:
    para(doc, "**%s**　%s" % l, space_after=6)

# ---- 十四、结论 ----
heading1(doc, "十四、最终结论：内蒙古在“能源转型+富民兴边”里的增长逻辑")
para(doc, "内蒙古的2025年，本质上是“**能源（煤+绿）+高端制造（稀土/新材料）+生态**”的答卷：GDP2.67万亿、规上+6.7%、新能源装机破亿，代价是PPI-7%、财政-4.6%。")
para(doc, "只要绿色能源、稀土/新材料、高端制造、向北开放能接住，内蒙就既保能源安全又走兴边富民；如果煤价长期低位、财政退坡，内蒙需承受“转型+外部价格”压力。")
para(doc, "最稳妥的观察信号：**一盯新能源/稀土/高端制造（动能）、二盯煤价/财政（约束）、三盯奶业/农畜（富民）、四盯生态/开放（定位）、五盯基建/民间（投资）。**内蒙古，是“能源强区+绿色转型”的样本。")

# ---- 附录A ----
heading1(doc, "附录A：主要资料来源与核验路径")
bullet(doc, "内蒙古自治区2025年《政府工作报告》（2025年1月）——目标来源。")
bullet(doc, "《2025年内蒙古自治区国民经济和社会发展统计公报》（内蒙古统计局，2026-04-09）——GDP、工业、外贸、人口实值。")
bullet(doc, "自治区统计局/能源、稀土、奶业专题——能源/资源产业。")
bullet(doc, "2026年内蒙古自治区政府工作报告——2025执行复盘。")
bullet(doc, "呼和浩特海关、财政厅、发改委——外贸/财政/能源。")
heading2(doc, "核验说明")
para(doc, "本报告所有增速、总量、占比均以统计公报/官方发布为基准。涉“稀土产量”“新能源细分”等以官方口径为准。")

# ---- 附录B ----
heading1(doc, "附录B：建议建立的年度跟踪仪表盘")
para(doc, "建议把下面10个“测脉搏”指标做成年度跟踪表：")
dash = [
    ["#", "指标", "2025基值", "为什么重要"],
    ["1", "GDP增速", "+4.7%", "总量与方向"],
    ["2", "规上工业增速", "+6.7%", "制造底盘"],
    ["3", "新能源及产业增速", "+18.0%", "绿色能源"],
    ["4", "稀土产业增速", "+32.0%", "战略资源"],
    ["5", "装备/高技术制造", "+20.6%/+17.5%", "新质生产力"],
    ["6", "PPI/一般公共预算收入", "-7.0%/-4.6%", "价格/财政"],
    ["7", "社零增速", "+4.3%", "内需消费"],
    ["8", "常住人口/城镇化率", "2374万/71.48%", "人口与城市"],
    ["9", "民间投资", "+8.8%", "民营活力"],
    ["10", "居民人均可支配收入增速", "+4.6%", "民生获得感"],
]
table(doc, dash[0], dash[1:], widths=[0.8, 4.5, 3.2, 4.6])
para(doc, "把这10个指标连起来看，任何一个“新引擎（3/4/5）向上、旧引擎（6）修复”，都说明内蒙古在真正换挡。")

# ========================================= 保存
out = "/Users/x/Desktop/content-prod-lab/reports/内蒙古自治区_2025年政府工作报告_深度研究_2026-08-13.docx"
doc.save(out)
print("SAVED:", out)
print("PARAGRAPHS:", len(doc.paragraphs))
print("TABLES:", len(doc.tables))
