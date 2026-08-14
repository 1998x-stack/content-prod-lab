# -*- coding: utf-8 -*-
"""Build 济宁市2025年政府工作报告 深度研究 DOCX, 参照地级市系列版式。"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DARK = RGBColor(0x1F, 0x2A, 0x44)
RED = RGBColor(0xB0, 0x1E, 0x2E)
GRAY = RGBColor(0x59, 0x60, 0x69)
BLUE = RGBColor(0x1F, 0x4E, 0x79)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
HEADER_FILL = "1F2A44"
K = "微软雅黑"


def set_run(run, size=11, bold=False, color=DARK, font=K):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    run.font.name = font
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), font)
    rFonts.set(qn("w:hAnsi"), font)
    rFonts.set(qn("w:eastAsia"), font)


def para(doc, text, size=11, bold=False, color=DARK, align=None,
         space_after=6, space_before=0, font=K):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        set_run(r, size=size, bold=bold or (i % 2 == 1), color=color, font=font)
    return p


def heading1(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(16)
    pf.space_after = Pt(8)
    r = p.add_run(text)
    set_run(r, size=16, bold=True, color=RED)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "B01E2E")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def heading2(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(10)
    pf.space_after = Pt(4)
    r = p.add_run(text)
    set_run(r, size=13, bold=True, color=BLUE)
    return p


def table(doc, headers, rows, widths=None, font_size=9.5):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_run(r, size=font_size, bold=True, color=WHITE)
        tcPr = hdr[i]._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), HEADER_FILL)
        tcPr.append(shd)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(val))
            set_run(r, size=font_size, color=DARK)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    return t


def quote_box(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(8)
    pf.left_indent = Pt(12)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), "B01E2E")
    pBdr.append(left)
    pPr.append(pBdr)
    r = p.add_run(text)
    set_run(r, size=11, color=DARK, font="楷体")
    return p


def bullet(doc, text, size=10.5):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.7)
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        set_run(r, size=size, bold=(i % 2 == 1), color=DARK)
    return p


# ================================================================= 文档主体
doc = Document()
sec = doc.sections[0]
sec.page_width = Cm(21)
sec.page_height = Cm(29.7)
sec.top_margin = Cm(2.4)
sec.bottom_margin = Cm(2.2)
sec.left_margin = Cm(2.5)
sec.right_margin = Cm(2.5)

st = doc.styles["Normal"]
st.font.name = K
st.font.size = Pt(11)
st.element.rPr.rFonts.set(qn("w:eastAsia"), K)


# ---- 封面 ----
para(doc, "乌海市2025年政府工作报告", size=24, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=60, space_after=4)
para(doc, "深度研究与\u201c容易被忽视的细节\u201d分析", size=16, bold=True, color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
para(doc, "从\u201c黄河明珠、乌海湖、葡萄之乡、煤焦精细化工、沙漠葡萄酒\u201d重新理解乌海", size=12, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)
para(doc, "\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501", color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
para(doc, "研究对象：2025年乌海市政府工作报告及同期财政、国民经济和社会发展统计资料、2026年官方复盘", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
para(doc, "标定日期：2026-08-14", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
doc.add_page_break()
catalog = [
    "执行摘要",
    "一、研究口径与阅读方法",
    "二、先看乌海的特别底盘：焦化、PVC精细化工、葡萄、乌海湖、沙漠",
    "三、最关键的宏观错位：GDP-4.3%、规上-10.4%、但高技术制造+4.5%、BDO+7%、固投+0.6%",
    "四、容易被忽视、但信号很重的细节（15条）",
    "五、2025年GDP目标 vs 实际：对照表",
    "六、2025年增长是谁撑起来的？（结构归因）",
    "七、预算与财政的\u201c含金量\u201d",
    "八、民生底账：人口、收入与城乡",
    "九、城镇与农村：格局与均衡",
    "十、人口流入与流出",
    "十一、物价与货币环境",
    "十二、区域一体化：乌海在呼包鄂、黄河几字湾、蒙宁、乌金之地里的位置",
    "十三、未来5\u201310年最值得观察的五条主线",
    "十四、最终结论：乌海在\u201c煤化工+新产业+葡萄文旅\u201d里的转型逻辑",
    "附录A：主要资料来源与核验路径",
    "附录B：建议建立的年度跟踪仪表盘",
]
para(doc, "目　录", size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=10, space_after=10)
for item in catalog:
    para(doc, item, size=12, space_after=4)
doc.add_page_break()

# ---- 执行摘要 ----
heading1(doc, "执行摘要")
para(doc, "**核心判断**　2025年乌海最显著的是\u201cGDP 540.75亿/-4.3%（承压下行、逐季收窄）、二产-7.3%、规上工业-10.4%、社零-1.4%\u201d、\u201c但高技术制造+4.5%、BDO一体+7.0%、固态电池+151.7%、硅基+13.8%、外贸+60.1%、固投+0.6%、技改+173.7%\u201d。这说明乌海在\u201c煤焦精细化工之城\u201d中，**旧煤焦产业下、新质/高技术/外贸在上，转型阵痛期但结构在优化**。")
para(doc, "把乌海2025年目标与实际对照，乌海是\u201c煤化工+新兴产业链\u201d路径：**煤焦化、PVC/氯碱、BDO一体化、精细化工、新材料（硅基/碳纤维）、氢能**是支柱。经济总量下滑，但技改、高技、外向在回升。")
para(doc, "最容易记住的一句话：**乌海是\u201c黄河明珠、乌金之都、葡萄之乡、精细化工之城\u201d，靠\u201c煤化工+新材料+葡萄文旅\u201d转型。**观察乌海，与其只看\u201cGDP 541亿/-4.3%\u201d，不如看\u201cBDO固态电池一体化+151.7%、技改+173.7%、硅基碳纤维、乌海葡萄\u201d。")

# ---- 一、研究口径与阅读方法 ----
heading1(doc, "一、研究口径与阅读方法")
para(doc, "**本文的最终目标**，不是复述乌海市2025年政府工作报告的原文，而是把它放到\u201c目标→实际\u201d和\u201c长期底盘→年度对撞\u201d的坐标里重新读一遍，找出乌海2025年到底靠什么增长、哪里在换挡、哪里是暗线。")
para(doc, "**三组资料互证**：（1）《乌海市2025年国民经济和社会发展统计公报》与《2025年乌海经济承压前行总体平稳》（市统计局，2026-01发布）提供2025年**实际完成**数据；（2）《乌海市2026年政府工作报告》（2026-01-28）复盘2025年；（3）《乌海市2025年政府工作报告》提供**目标设定**。三份材料交叉核验，避免单源失真。")
para(doc, "**核心框架**：先交代乌海的\u201c底盘\u201d（煤焦精细化工+新材料+葡萄文旅+乌海湖的产业与地理禀赋），再用\u201c目标vs实际\u201d对照表定位宏观错位，接着用15条细节捕捉暗线，最后从财政、民生、城乡、人口、物价坐实，落到2026/十五五主线与结论。")
para(doc, "**口径说明**：GDP 540.75亿元-4.3%（2025年是收缩年）、三次产业增加值8.64/247.31/284.80（占比1.6:45.7:52.7）。500万元以上固投+0.6%、房地产-18.7%。社零-1.4%。一般公共预算收入+0.2%。规上工业-10.4%。居民收入城镇+3.0%/农村+4.8%。CPI-0.3%左右。")
para(doc, "**为什么值得单独研究乌海**：GDP-4.3%、规上-10.4%，是2025年典型的\u201c收缩转型\u201d城市；但高技术制造+4.5%、BDO+7%、固态电池+151.7%、技改+173.7%、外贸+60.1%。\u201c旧产业阵痛、新产业起跑\u201d的组合，是观察\u201c煤焦化工城市绿色低碳转型\u201d的上佳切片。")
# ---- 二、先看乌海的特别底盘 ----
heading1(doc, "二、先看乌海的特别底盘：煤焦化工、PVC、新材、葡萄、乌海湖")
para(doc, "乌海的成长逻辑，不能只看GDP数字，而要先把\u201c底盘\u201d摊开。底盘决定了几十年、上百年的禀赋，也决定了2025年增长的来源和约束。")
bullet(doc, "**煤焦精细化工之城**：乌海是内蒙古煤焦化、氯碱/PVC、精细化工的重要基地，煤焦化（焦炭1660.7万吨、原煤4156.3万吨）、PVC（71.4万吨+3.9%）、BDO一体化、精细化工集群是工业主干。")
bullet(doc, "**新兴产业链（固态电池/BDO/硅基）**：BDO/可降解塑料、固态电池材料、硅基新材料、碳纤维是新材料新赛道；BDO一体化+7%、固态电池+151.7%、硅基+13.8%。")
bullet(doc, "**黄河明珠·乌海湖**：黄河穿城、乌海湖（2号治理+拦河大坝形成），\u201c沙漠看海\u201d、葡萄印象，黄河生态与文旅名片。")
bullet(doc, "**葡萄之乡·沙漠葡萄酒**：乌海葡萄（1.14万吨）、\u201c乌海葡萄\u201d国家地理标志，沙漠葡萄酒业与旅游（沙漠葡萄酒文化旅游节）。")
bullet(doc, "**乌金/绿电转型**：原煤、焦化基地在绿色化（技改+173.7%、新能源发电15.1亿千瓦时+1.2%、氢能），向低碳精细化工转型。")
heading2(doc, "底盘的产业钟摆")
para(doc, "把底盘归纳成一句话：**乌海是\u201c煤焦精细化工+新材料\u201d的工业城市，叠加黄河乌海湖+葡萄**。煤焦定旧基、新材PVC/BDO撑第二曲线、乌海湖沙漠葡萄酒供文旅。城市资产在\u201c煤化工→新材料\u201d链条上，2025年正处换挡阵痛：旧煤焦（原煤-17.1%、焦炭-0.2%、采矿业-14.6%）下、新质（BDO/硅基/固态电池）上。")

# ---- 三、最关键的宏观错位 ----
heading1(doc, "三、最关键的宏观错位：GDP-4.3%、规上-10.4%、但高技术+4.5%、BDO+7%、固投+0.6%")
para(doc, "把2025年数据摊开，乌海最醒目的不是\u201c-4.3%\u201d，而是**旧动能下、新动能上**的结构错位。")
quote_box(doc, "宏观错位一号：**总量的负 vs 结构的分化**。GDP-4.3%、二产-7.3%、规上-10.4%（采掘-14.6%、制造-5.5%）；但高技术制造+4.5%、BDO+7%、固态电池+151.7%、硅基+13.8%——旧煤焦在跌、新产业链在涨，转型阵痛。")
quote_box(doc, "宏观错位二号：**煤焦收缩 vs 投资/技改正**。原煤-17.1%、焦炭-0.2%、发电-5.7%下行；但固投+0.6%、基础设施+75.1%、技改+173.7%（电力燃气水+57.6%）——投资在向风电/绿电/技改与新型基础设施转。")
quote_box(doc, "宏观错位三号：**内需弱 vs 外贸强**。社零-1.4%（消费弱）；但外贸+60.1%（顺差/出口回升）——外需在回暖、内需在收缩。")
para(doc, "三条错位加在一起，指向同一个结论：乌海2025年是\u201c煤焦化工收缩、新型材料与技改投资上、外向回暖\u201d的换挡之年。总量下滑是旧产业去产能的代价，但高技术/BDO/固态电池/技改/基建在积蓄下一程的动能。")

# ---- 四、容易被忽视、但信号很重的细节（15条） ----
heading1(doc, "四、容易被忽视、但信号很重的细节（15条）")
para(doc, "下面这15条，散落在统计公报的角落里，单独看无关紧要，合在一起却拼出乌海2025年真实的结构肌理。")
bullet(doc, "**1. 高技术制造+4.5%**：在规上-10.4%大背景下，高技术制造翻正，是新质的最亮点。")
bullet(doc, "**2. BDO一体化+7.0%、固态电池链路+151.7%、硅基新材料+13.8%**：卡三大新材料链在跃升、工业第二曲线成型。")
bullet(doc, "**3. 精甲醇+18.5%、BDO+5.9%、PTMEG+19.7%**：化工品向BDO/PTMEG高端延链增效。")
bullet(doc, "**4. PVC产量71.4万吨+3.9%**：氯碱/PVC在旧煤基中仍正增长。")
bullet(doc, "**5. 技改投资+173.7%、基建+75.1%、电力燃气水+57.6%**：投资向技改/绿电/基建转、转型投入加大。")
bullet(doc, "**6. 能源：新能源发电15.1亿千瓦时、固投转风电/光伏/氢**：绿电初成。")
bullet(doc, "**7. 粮食1.1亿斤+12.9%、玉米亩均600公斤全区第一**：农业高产、粮食六连增。")
bullet(doc, "**8. 葡萄1.14万吨、蔬菜5.74万吨**：葡萄/蔬菜市郊农业。")
bullet(doc, "**9. 外贸+60.1%（进口或出口回升）**：外向商贸大幅回暖。")
bullet(doc, "**10. 规上工业营收1102.9亿、利润25.0亿**：营收/利润在煤价下行中仍维持。")
bullet(doc, "**11. 服务业-1.4%（信息软件+7.2%、文体娱+14.8%）**：服务业承压但信息/文体在增。")
bullet(doc, "**12. 一产+5.7%**：农业高增、是少数正增长的产业。")
bullet(doc, "**13. 焦化超低排放/环保改造完成**：煤化工绿色低碳改造、去污减排。")
bullet(doc, "**14. 工业产品从100余种扩至300余种（十四五）**：精细化工产业链延伸。")
bullet(doc, "**15. 财收增速+0.2%（微增）**：经济收缩下财政仍稳住。")
# ---- 五、2025年GDP目标 vs 实际：对照表 ----
heading1(doc, "五、2025年GDP目标 vs 实际：对照表")
para(doc, "把乌海市2025年预期目标与实际值逐项对照：")
table(doc,
    ["指标", "2025年目标", "2025年实际", "判定"],
    [
        ["地区生产总值增速", "（稳增长）", "540.75亿，-4.3%", "未达、负增长"],
        ["规上工业增加值增速", "（+稳）", "-10.4%", "收缩"],
        ["固定资产投资增速", "+8%以上", "+0.6%", "远低目标"],
        ["社会消费品零售总额", "（稳）", "-1.4%", "下"],
        ["一般公共预算收入", "（+稳）", "+0.2%", "微增"],
        ["外贸总额", "（稳外贸）", "+60.1%", "大幅超额"],
    ],
    widths=[4.6, 3.2, 4.4, 3.8],
)
para(doc, "这张表透露乌海2025年的\u201c成色\u201d：**除外贸+60.1%大幅超额外，GDP、规上、固投、社零全部负增长/远低于目标**。这是乌海自煤焦下行以来最困难的一年——\u201c总量收缩、外贸独撑、新产业在孵化\u201d。")

# ---- 六、2025年增长是谁撑起来的？（结构归因） ----
heading1(doc, "六、2025年增长是谁撑起来的？（结构归因）")
para(doc, "乌海2025年GDP-4.3%的\u201c负增长功劳簿\u201d（收缩下的结构性亮点），大致拆成几块：")
bullet(doc, "**拖累：二产-7.3%（煤焦/采掘/制造）**。规上工业-10.4%、采矿业-14.6%、制造业-5.5%、原煤-17.1%，煤焦去产能与价格下行拖垮工业。")
bullet(doc, "**三产-1.4%（微降，但信息软件+7.2%）**：服务业承压，但大数据/软件信息在增。")
bullet(doc, "**正增：一产+5.7%（农业）**：粮食/葡萄/蔬菜高增，是GDP少数正产业。")
bullet(doc, "**引擎：高技术/外贸/技改投资**。高技术制造+4.5%、BDO+7%、固态电池+151.7%、硅基+13.8%、外贸+60.1%、技改+173.7%——新质与外向在积蓄。")
para(doc, "**结构归因结论**：2025年的乌海是\u201c旧煤焦塌方、新质/外贸/技改起跑\u201d。总量下降由煤焦与采矿业拖累，但高技术、BDO、固态电池、技改投资与外贸证明\u201c换挡方向对\u201d——曙光在技术上、代价在存量上。")

# ---- 七、预算与财政的\u201c含金量\u201d ----
heading1(doc, "七、预算与财政的\u201c含金量\u201d")
table(doc, ["财政指标", "2025年情况", "点评"],
    [
        ["一般公共预算收入", "（微增+0.2%）", "经济收缩下稳"],
        ["税收结构", "工业/煤焦化工/基建税基", "承压"],
        ["财政支出", "技改/绿电/民生", "转型投入"],
        ["转型投入", "技改+173.7%、电力+57.6%", "绿色转型"],
    ],
    widths=[4.6, 5.6, 4.8],
)
para(doc, "财政核心判断：**经济收缩、财政微增（+0.2%）**。乌海一般公共预算收入在GDP-4.3%下仍+0.2%，靠非税/基建/技改与存量税基，是承压中含金量尚可。支出侧重于技改（+173.7%）、绿电、环保与民生，体现\u201c以技改带转型\u201d。")
quote_box(doc, "财政与宏观的勾连：**经济负、财政稳**。GDP-4.3%但财收+0.2%，说明乌海财政韧性较好（税收基础厚、非税补）也靠中央/自治区转移支付支撑；技改/绿电的投入是换挡的关键。")

# ================= 八、民生底账 =================
heading1(doc, "八、民生底账：人口、收入与城乡")
bullet(doc, "**居民收入**：城镇+3.0%、农村+4.8%（全体居民约5.6万、十四五末）。城乡比相对均衡。")
bullet(doc, "**就业**：规上工业收缩、部分就业承压；服务业/农业吸就业。")
bullet(doc, "**民生**：财政民生支出、社保医保覆盖（中西部工业城、保障较全）。")
para(doc, "民生综合评价：**收入中低速（城镇+3%/农村+4.8%）**、就业受工业收缩影响；靠农业/服务/社保托底，民生在收缩中保持稳定。")

# ================= 九、城镇与农村 =================
heading1(doc, "九、城镇与农村：格局与均衡")
bullet(doc, "**城镇化高**：乌海是河谷小城、城镇化率高（约90%）、农村人口少。")
bullet(doc, "**城乡收入**：农村+4.8%快于城镇+3.0%，城乡差距小。")
bullet(doc, "**农业**：粮食1.1亿斤+12.9%、葡萄/蔬菜，都市农业。")
para(doc, "乌海城镇化极高、城乡差异小、农业为城市菜篮子，是\u201c小而均衡\u201d的河谷城市。")

# ============ 十、人口流入与流出 ============
heading1(doc, "十、人口流入与流出")
bullet(doc, "**常住约56万（2017后）**：乌海人口规模小、相对稳定（工业城）。")
bullet(doc, "**人口结构**：工业人口、矿区人口，转型期年轻人外流/GDP收缩。")
quote_box(doc, "人口：常住约56万小城、依赖工业；转型期就业/人口压力缓显，需靠新产业留人。")

# ============ 十一、物价与货币 ============
heading1(doc, "十一、物价与货币环境")
bullet(doc, "**CPI约-0.3%**：工业收缩下通缩、消费价格偏弱。")
bullet(doc, "**金融**：存款相对充裕、贷款配合技改绿电投资。")
para(doc, "微通缩+信贷对转型投资：乌海资金环境平、支持技改绿电。")

# ========== 十二、区域一体化 ==========
heading1(doc, "十二、区域一体化：乌海在呼包鄂、黄河几字湾、蒙宁、乌金之地里的位置")
bullet(doc, "**呼包鄂榆城市群**：乌海是呼包鄂大都市圈西端节点，与呼和浩特、包头、鄂尔多斯联动（煤化/电力）。")
bullet(doc, "**黄河几字湾/蒙宁交界**：位居黄河几字湾与蒙宁交界，黄河生态与区域协作。")
bullet(doc, "**乌金（煤）之都向绿电转型**：依托原煤焦化+新能源，是内蒙古西部工业/电力枢纽。")
bullet(doc, "**县域**：海勃湾（主城/精细化工）、海南（焦化/新材料）、乌达（园区），构成\u201c煤化+新材\u201d骨架。")
para(doc, "乌海坐标：**在\u201c呼包鄂城市群+黄河/蒙西\u201d下，乌海是内蒙古西部的煤化工与新材基地**。向绿氢/精细高附加值转型是其能级关键。")

# ============ 十三、未来5-10年主线 ============
heading1(doc, "十三、未来5\u201310年最值得观察的五条主线")
bullet(doc, "**主线一：BDO/可降解塑料与固态电池链**。BDO+7%、固态电池+151.7%，若这些新材料链放量，工业第二曲线成。")
bullet(doc, "**主线二：精细化工/煤化工延链（300余种产品）**。从\u201c煤焦化\u201d到精细化工/新材料附加值。")
bullet(doc, "**主线三：绿电/氢能转型**。技改+173.7%、新能源+光伏风电氢，向绿色低碳化工转。")
bullet(doc, "**主线四：外贸持续高增**。+60.1%外贸，若拓展下游链出口，外向增量。")
bullet(doc, "**主线五：葡萄/乌海湖文旅**。沙漠葡萄酒、乌海湖国家旅游度假区，三产与留人。")
para(doc, "五条主线里，**最值得盯的是主线一（新材料/BDO）与主线三（绿电解氢）**——乌海能否走出-4.3%，卡在把BDO/固态电池链和绿氢做起来、对冲煤焦收缩。")

# ============ 十四、最终结论 ============
heading1(doc, "十四、最终结论：乌海在\u201c煤化工→新材料\u201d里的转型逻辑")
para(doc, "综合2025年的开局与结构，给乌海一个平衡的结论：")
para(doc, "**一、总量收缩、但转型在积蓄。**GDP-4.3%、规上-10.4%，是煤焦化去产能与价格下行的阵痛；但高技术+4.5%、BDO+7%、固态电池+151.7%、技改+173.7%、外贸+60.1，证明\u201c换挡\u201d已在路上。")
para(doc, "**二、煤焦与新材料是两条腿。**旧煤焦（原煤-17.1%、采掘-14.6%）拖总量，新材（BDO/硅基/固态电池）、技改投资、外贸在补。乌海要\u201c破旧立新\u201d，靠新产业链放大。")
para(doc, "**三、财政稳、但内需弱。**财收+0.2%稳住，但社零-1.4%、CPI微通缩，消费与房地产待修复。")
para(doc, "**四、2026年前瞻。**在\u201c稳总量、放大新材料、促绿氢、稳财政、提消费\u201d约束下，乌海2026年边际变量是**BDO/固态电池放量、新能源/绿氢、精化工延长链、外贸、技改投资**。若能靠新材料与绿电对冲煤焦，乌海有望\u201c触底回升\u201d。")
para(doc, "**一句话总结：乌海2025年交出了\u201c-4.3%、规上-10.4%\u201d的收缩答卷，煤焦化工的阵痛吞噬了总量；但更该问的是——在这个乌金之地的城市，这座城市能否把\u201cBDO、固态电池、硅基、绿氢\u201d的新材料与绿电牌打出来，炼成\u201c新材料+绿色化工+外贸\u201d的新曲线，让乌海从\u201c靠煤吃煤\u201d走向\u201c换轨起飞\u201d。**当下看，阵痛是真、转型也是真，坚持新材料与绿电，乌海终将破旧立新。")

# ============ 附录A ============
heading1(doc, "附录A：主要资料来源与核验路径")
bullet(doc, "《乌海市2025年国民经济和社会发展统计公报》/《2025年乌海经济承压前行总体平稳》（市统计局，2026-01）：GDP、三次产业、工业、投资、收入等。")
bullet(doc, "《乌海市2026年政府工作报告》（2026-01-28齐海斌）：2025年回顾、2026年目标。")
bullet(doc, "《乌海市2025年政府工作报告》：2025年预期目标。")
bullet(doc, "核验方式：以统计公报/经济运行情况为准、以政府工作报告为框架，交叉比对目标-实际。")

# ============ 附录B ============
heading1(doc, "附录B：建议建立年度跟踪仪表盘")
table(doc, ["维度", "跟踪指标", "用途/预警"],
    [
        ["总量", "GDP、人均GDP", "触底回升"],
        ["工业", "规上工业、采矿、高技术制造", "新旧动能"],
        ["新材料", "BDO/固态电池/硅基产业链", "第二曲线"],
        ["投资", "固投、技改、绿电/基建", "转型投入"],
        ["外贸", "进出口增速", "外向回暖"],
        ["财政", "收入、税收、技改支出", "财政质量"],
        ["消费", "社零、CPI", "内需修复"],
    ],
    widths=[2.6, 6.2, 6.0],
)

# ========== 保存
doc.save("/Users/x/Desktop/content-prod-lab/reports/乌海市_2025年政府工作报告_深度研究_2026-08-14.docx")
print("SAVED OK: 乌海市 /Users/x/Desktop/content-prod-lab/reports/乌海市_2025年政府工作报告_深度研究_2026-08-14.docx")
