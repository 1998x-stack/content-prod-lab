# -*- coding: utf-8 -*-
"""Build 南宁市2025年政府工作报告 深度研究 DOCX, 参照省会系列版式。"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DARK = RGBColor(0x1F, 0x2A, 0x44)
RED = RGBColor(0xB0, 0x1E, 0x2E)
GRAY = RGBColor(0x59, 0x60, 0x69)
BLUE = RGBColor(0x1F, 0x4E, 0x79)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
HEADER_FILL = "1F2A44"
K = "微软雅黑"


def set_run(run, size=11, bold=False, color=DARK, font=K):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    run.font.name = font
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), font)
    rFonts.set(qn("w:hAnsi"), font)
    rFonts.set(qn("w:eastAsia"), font)


def para(doc, text, size=11, bold=False, color=DARK, align=None,
         space_after=6, space_before=0, font=K):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        set_run(r, size=size, bold=bold or (i % 2 == 1), color=color, font=font)
    return p


def heading1(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(16)
    pf.space_after = Pt(8)
    r = p.add_run(text)
    set_run(r, size=16, bold=True, color=RED)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "B01E2E")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def heading2(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(10)
    pf.space_after = Pt(4)
    r = p.add_run(text)
    set_run(r, size=13, bold=True, color=BLUE)
    return p


def table(doc, headers, rows, widths=None, font_size=9.5):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_run(r, size=font_size, bold=True, color=WHITE)
        tcPr = hdr[i]._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), HEADER_FILL)
        tcPr.append(shd)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(val))
            set_run(r, size=font_size, color=DARK)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    return t


def quote_box(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(8)
    pf.left_indent = Pt(12)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), "B01E2E")
    pBdr.append(left)
    pPr.append(pBdr)
    r = p.add_run(text)
    set_run(r, size=11, color=DARK, font="楷体")
    return p


def bullet(doc, text, size=10.5):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.7)
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        set_run(r, size=size, bold=(i % 2 == 1), color=DARK)
    return p


# ================================================================= 文档主体
doc = Document()
sec = doc.sections[0]
sec.page_width = Cm(21)
sec.page_height = Cm(29.7)
sec.top_margin = Cm(2.4)
sec.bottom_margin = Cm(2.2)
sec.left_margin = Cm(2.5)
sec.right_margin = Cm(2.5)

st = doc.styles["Normal"]
st.font.name = K
st.font.size = Pt(11)
st.element.rPr.rFonts.set(qn("w:eastAsia"), K)


# ---- 封面 ----
para(doc, "南宁市2025年政府工作报告", size=24, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=60, space_after=4)
para(doc, "深度研究与\u201c容易被忽视的细节\u201d分析", size=16, bold=True, color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
para(doc, "从\u201c中国—东盟门户、平陆运河、跨境合作、新能源汽车与电子信息\u201d重新理解南宁", size=12, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)
para(doc, "\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501", color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
para(doc, "研究对象：2025年南宁市政府工作报告及同期财政、国民经济和社会发展统计资料、2026年官方复盘", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
para(doc, "标定日期：2026-08-13", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
doc.add_page_break()

# ---- 目录 ----
catalog = [
    "执行摘要",
    "一、研究口径与阅读方法",
    "二、先看南宁的特殊底盘：东盟门户、平陆运河、新能源电池与面向东盟合作",
    "三、最关键的宏观错位：GDP破6200亿但偏低，工业/高技术/出口强，投资转负",
    "四、容易被忽视、但信号很重的细节（15条）",
    "五、2025年GDP目标 vs 实际：对照表",
    "六、2025年增长是谁撑起来的？（结构归因）",
    "七、预算与财政的\u201c含金量\u201d",
    "八、民生底账：人口、收入与城乡",
    "九、城镇与农村：格局与均衡",
    "十、人口流入与流出",
    "十一、物价与货币环境",
    "十二、区域一体化：南宁在\u201c珠江—西江经济带+中国东盟自贸区3.0+平陆运河\u201d里的位置",
    "十三、未来5—10年最值得观察的五条主线",
    "十四、最终结论：南宁在\u201c东盟门户+新能源电池+跨境电商\u201d里的增长逻辑",
    "附录A：主要资料来源与核验路径",
    "附录B：建议建立的年度跟踪仪表盘",
]
para(doc, "目　录", size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=10, space_after=10)
for item in catalog:
    para(doc, item, size=12, space_after=4)
doc.add_page_break()

# ---- 执行摘要 ----
heading1(doc, "执行摘要")
para(doc, "**核心判断**　2025年南宁最显著的是\u201cGDP 6212.46亿元、增长4.7%（低于5%目标）、三产占70.1%\u201d、\u201c规上工业+9.1%但高技术制造+66%\u201d、\u201c货物进出口1044.37亿元/+12.4%、出口+19.3%\u201d、\u201c常住人口901.66万、自然增长率-0.49\u2030\u201d。这说明南宁经济在\u201c东盟门户+平陆运河+新能源电池\u201d驱动下稳健增长，但**投资-15.2%大幅转负、人口自然负增长**是短板。")
para(doc, "把2025年目标（GDP +5%/规上工业+8%/固投+1%/财政+3%）、2025年统计、2026年前瞻一起看，南宁是\u201c面向东盟的开放门户+制造强市\u201d路径：**工业（新能源电池/电子）+出口+面向东盟口岸**是引擎，但内需偏弱、投资塌陷。总量6212亿居广西首位、占全区约1/4。")
para(doc, "最容易记住的一句话：**南宁是\u201c中国—东盟门户+平陆运河出海口\u201d的广西省会，靠\u201c新能源电池+电子信息+贸易+面向东盟+跨境电商\u201d增长。**观察南宁，与其只看\u201cGDP 6212亿\u201d，不如看\u201c规上工业+9.1%、高技术制造+66%、出口+19.3%、平陆运河投资完成90%、跨境人民币结算6966.95亿\u201d。")# ---- 一、研究口径与阅读方法 ----
heading1(doc, "一、研究口径与阅读方法")
para(doc, "本报告以\u201c南宁市2025年政府工作报告（2025年，侯刚作）\u201d为起点，把\u201c2025年GDP目标（5%左右）\u201d与\u201c2025年国民经济和社会发展统计公报实际值（6212.46亿元/+4.7%）\u201d并置对照，再用2026年政府工作报告的复盘作事后验证。")
para(doc, "重点阅读三类证据：**一是指标目标是否达标**；**二是结构（谁贡献增长）是否匹配目标叙事**；**三是官方复盘如何自评、承认问题**。统计口径一般公共预算收入为地方级。")
para(doc, "统一采用\u201c同比增长%\u201d（GDP为不变价），财政与居民收入多为名义值，文中按原口径转录、不逐项换算。")

# ---- 二、特殊底盘 ----
heading1(doc, "二、先看南宁的特殊底盘：东盟门户、平陆运河、新能源电池与面向东盟合作")
para(doc, "**区位与身份**：南宁是广西壮族自治区省会，中国—东盟博览会（东博会）永久举办地，面向东盟\u201c桥头堡\u201d，中国西南出海口门户。国家重大工程\u201c平陆运河\u201d从南宁起（南宁段投资完成90%），将让南宁直连北部湾港东盟。定位\u201c3枢纽+1基地\u201d综合物流枢纽城市。")
para(doc, "**产业底盘**：一是新能源/电池（投产产能超100GWh、电池产值累计破1300亿）；二是电子信息/计算机（高技术制造+66%）；三是造纸、铝加工（动力电池铝箔全国前三）；四是跨境电商与面向东盟贸易（对东盟出口+19.2%）；五是数字经济（人工智能/低空经济/五象新区）。")
para(doc, "**人口底盘**：2025年末常住人口901.66万/+0.5%（+4.47万），城镇化率72.4%（提高0.52个百分点）；是广西人口大市、广西最大城市。")
para(doc, "**市场与出口**：2025年社零2186.7亿元/+4.1%；货物进出口1044.37亿元/+12.4%、出口+19.3%；跨境人民币结算6966.95亿元/+46.6%。对外开放是南宁最强特色。")

# ---- 三、核心宏观错位 ----
heading1(doc, "三、最关键的宏观错位：GDP破6200亿但偏低，工业/高技术/出口强，投资转负")
para(doc, "**第一组错位**：2025年GDP目标5%左右，实际6212.46亿/+4.7%**略低于目标**，但\u201c十四五\u201d连续突破5000/6000亿，达6212亿、广西首位。与之对应，投资-15.2%（目标+1%**严重未达标**）。")
para(doc, "**第二组错位**：规上工业+9.1%（约达目标+8%）、**高技术制造+66%、装备制造+20.3%**爆发；但固投（不含农户）-15.2%、房地产-30.7%。\u201c工业新增、投资收缩\u201d明显。")
para(doc, "**第三组错位**：城镇化72.4%但**自然增长率-0.49\u2030（出生6.59\u2030<死亡7.08\u2030）**人口自然微弱负增长，靠机械流入维持总量。")
para(doc, "一句话：**南宁是\u201c对外开放+制造业转型强、投资塌陷+内需偏弱\u201d的广西省会**——增长靠外贸、新能源电池和面向东盟口岸。")

# ---- 四、15条细节 ----
heading1(doc, "四、容易被忽视、但信号很重的细节（15条）")
bullet(doc, "1. **GDP破6200亿**：6212.46亿/+4.7%，\u201c十四五\u201d连破5000/6000亿。")
bullet(doc, "2. **高技术制造业+66%**：占规上工业比重提升至25.7%（+8.5pct）。")
bullet(doc, "3. **新能电池/铝（电池产能100GWh）**：新能源电池产业产能破千亿级、铝加工110万吨。")
bullet(doc, "4. **出口+19.3%、对东盟出口+19.2%**：面向东盟贸易强。")
bullet(doc, "5. **平陆运河南宁段投资完成90%**：南宁到北部湾港、出海口正打通。")
bullet(doc, "6. **规上工业+9.1%、利润143.35亿/+34.2%**：工业盈利强但结构分化。")
bullet(doc, "7. **固投-15.2%**：房地产-30.7%、建安、基建也走弱，是最大隐患。")
bullet(doc, "8. **限上网络零售+39.5%、可穿戴智能+539%**：线上消费爆发。")
bullet(doc, "9. **汽车锂离子动力电池产量+72.4%**：新能源电池主线。")
bullet(doc, "10. **电子元件+17.4%、智能手机+63.7%**：电子信息制造。")
bullet(doc, "11. **常住人口901.66万/城镇化72.4%**：但自然增长率-0.49‰。")
bullet(doc, "12. **收入：城镇48375元/+4.3%、农村23146元/+6.1%**：城乡速度快。")
bullet(doc, "13. **CPI -0.4%**：低通胀/微通缩。")
bullet(doc, "14. **一般公共预算收入381.69亿/+0.1%**：基本零增长、税收+2.0%。")
bullet(doc, "15. **旅游1.97亿人次/+12%、收入2215.94亿/+12.5%**：文旅热，兼作东盟会展。")# ---- 五、目标 vs 实际对照表 ----
heading1(doc, "五、2025年GDP目标 vs 实际：对照表")
table(doc,
    ["指标", "2025年目标", "2025年实际", "是否达标"],
    [
        ["地区生产总值增速", "5%左右", "6212.46亿元/+4.7%", "略低于目标"],
        ["规上工业增加值", "+8%", "+9.1%", "超目标"],
        ["固定资产投资", "+1%", "-15.2%（转负）", "严重未达标"],
        ["社会消费品零售总额", "（未列明）", "2186.7亿元/+4.1%", "—"],
        ["一般公共预算收入", "+3%", "381.69亿元/+0.1%", "未达标(近零增长)"],
        ["居民人均可支配收入", "与GDP同步", "全体+5.1%", "略高于GDP"],
        ["CPI涨幅", "2%左右", "-0.4%", "远低于目标"],
    ],
    widths=[3.4, 2.5, 5.0, 3.1])
para(doc, "**简评**：唯一达标且超的是**规上工业**；**GDP、投资、财政明显低于目标**、消费与收入温和。这是南宁\u201c工业与出口好、投资塌陷、财政零增长\u201d的一年。")

# ---- 六、结构归因 ----
heading1(doc, "六、2025年增长是谁撑起来的？（结构归因）")
para(doc, "**三大产业**：一产655.84亿/+4.2%、二产1200.63亿/+5.2%、三产4355.99亿/+4.6%；三产占GDP 70.1%（贡献率66%）。增长以\u201c三产+二产（工业）\u201d双支撑。")
para(doc, "**工业**：全部工业增加值922.72亿/+8.0%、规上+9.1%；**高技术制造+66%、装备制造+20.3%**（占规上25.7%）；电气机械+25.4%、电池制造+26.6%、烟草+4%。规上工业利润143.35亿/+34.2%。")
para(doc, "**服务业**：信息传输软件+12.1%、金融+1.6%、批零+3.4%；电子商务/数字经济扩张。")
para(doc, "**开放**：货物进出口1044.37亿/+12.4%（出口+19.3%、进口+5.7%），对东盟+15.4%、RCEP+9.5%；跨境人民币结算6966.95亿/+46.6%。")
para(doc, "**总量归因**：南宁增长主要靠**三产（贸易/金融/文旅/数字经济）+工业（新能源电池/电子信息）+出口**；投资、房地产为负贡献。")

# ---- 七、财政 ----
heading1(doc, "七、预算与财政的\u201c含金量\u201d")
para(doc, "2025年一般公共预算收入381.69亿元/+0.1%（税收248.46亿/+2.0%）；一般公共预算支出822.34亿元/+3.2%，民生支出664.83亿/+4.5%占80.8%。")
para(doc, "**结构性**：收入\u201c近零增长\u201d（税收微升+2%、非税承压），支出稳增（+3.2%）支撑民生。\u201c减收保民生、基建强依赖转移支付\u201d。")
para(doc, "**含金量**：财政\u201c量微增、民生占比高\u201d，但工业/出口强而财政弱，反映南宁自主财力有限、更依赖中央转移与贸易。")

# ---------------- 八、民生 ----------------
heading1(doc, "八、民生底账：人口、收入与城乡")
para(doc, "常住人口901.66万/+0.5%（+4.47万）、城镇化率72.4%；出生率6.59\u2030/死亡率7.08\u2030，**自然增长率-0.49\u2030**。收入：**城镇48375元/+4.3%、农村23146元/+6.1%**，城乡比2.09（缩小0.04）。")
para(doc, "消费支出全体23858元/+8.4%（城镇27342/+9.8%）。城镇新增就业7.92万人。")
para(doc, "**民生结论**：收入农村快于城镇、城乡差距收敛；人口总量靠机械流入维持、自然轻微负增长；就业、消费、医保保障稳定。")

# ---------------- 九、城乡 ----------------
heading1(doc, "九、城镇与农村：格局与均衡")
para(doc, "南宁城镇化率72.4%（广西前列），主城区承载贸易/金融/数字经济；县域（横州茉莉花、武鸣农业、宾阳）发展特色农业与食品加工。")
para(doc, "城乡收入比2.09（缩小），农村收入增速快、差距在边际收敛；县域特色产业（茉莉花茶品牌226.69亿、武鸣国家产业园）支撑乡村振兴。")

# =========== 十、人口流入流出 =============
heading1(doc, "十、人口流入与流出")
para(doc, "南宁常住901.66万、微增4.47万，但自然负增长（出生6.59\u2030<死亡7.08\u2030，-0.49\u2030），增长靠机械流入（高校+东盟人才+产业劳动力+新市民）。")
para(doc, "**流入**：东博会/东盟经贸人才、制造业（新能源电池）、大学生留邕；**流出**：部分低收入/基层劳动力外迁。净流入有限、人口趋稳，是长期主线。")

# =========== 十一、物价 =============
heading1(doc, "十一、物价与货币环境")
para(doc, "2025年南宁CPI**下降0.4%**（通缩），八大类\u201c三升四降一平\u201d。为低通胀/微通缩。")
para(doc, "金融：存款余额15766.69亿、贷款24740.83亿（省会贷存比高）；跨境人民币结算6966.95亿/+46.6%是\u201c货币\u201d亮点。\u201c宽货币、需求弱\u201d。")

# =========== 十二、区域一体化 ============
heading1(doc, "十二、区域一体化：南宁在\u201c珠江—西江经济带+东盟自贸区3.0+平陆运河\u201d里的位置")
para(doc, "南宁是珠江—西江经济带核心、中国东盟自贸区3.0版\u201c桥头堡\u201d、面向东盟的大通道。**平陆运河（南宁段投资完成90%）**将让南宁经北部湾港直达东盟，是\u201c出海口+临港\u201d的关键工程。")
para(doc, "对外开放：东博会永久举办地、中越班列加密（+86%）、跨境电商、跨境人民币结算6966.95亿、对东盟贸易强。南宁是\u201c门户型\u201d核心：面向东盟+西部陆海新通道+平陆运河。")

# ============ 十三、五条主线 =============
heading1(doc, "十三、未来5—10年最值得观察的五条主线")
bullet(doc, "1. **平陆运河与出海港口**：2026年建平陆运河首发港，南宁\u201c内陆变出海口\u201d能否兑现。")
bullet(doc, "2. **新能源电池+电子信息**：投产产能超100GWh、锂电/铝/电子，制造强市主升浪。")
bullet(doc, "3. **面向东盟贸易+跨境电商**：东博会、跨境人民币、自由贸易（跨境）电商，出口第二增长极。")
bullet(doc, "4. **投资修复**：固投-15.2%后能否靠\u201c平陆+东部新城+五象\u201d重大项目修复（2026目标700亿）。")
bullet(doc, "5. **人口与AI/低空**：901万人口/自然负增长，AI+低空经济+五象新区催生新产业与就业。")

# ============ 十四、最终结论 =============
heading1(doc, "十四、最终结论：南宁在\u201c东盟门户+新能源电池+跨境电商\u201d里的增长逻辑")
para(doc, "**结论**：南宁2025年的\u201c真相\u201d是——**GDP+4.7%（略低于5%）、规上工业+9.1%、高技术制造+66%、出口+19.3%，但固投-15.2%、财政+0.1%、人口自然-0.49\u2030**。它是\u201c开放（东盟）+制造（新能源/电子）强、投资/财政/人口弱\u201d的广西省会。")
para(doc, "**对趋势判断**：工业与开放代表南宁的\u201c新动能\u201d，投资/财政代表\u201c约束\u201d。**新兴电池/电子/东盟贸易/平陆运河**决定未来潜力；**投资盘活+人口/消费**决定长期韧性。")
para(doc, "**若只看一个指标**：看**固投增速（-15.2%）+高技术制造增速（+66%）**——南宁是\u201c投资与产业严重错配、靠对外开放+制造升级\u201d的省会，平陆运河与投资修复是南宁能否摆脱\u201c低效投资\u201d的关键。")

# ---------------- 附录A ----------------
heading1(doc, "附录A：主要资料来源与核验路径")
bullet(doc, "南宁市人民政府《2025年南宁市政府工作报告及2024年回顾》（2025年）。")
bullet(doc, "南宁市统计局《2025年南宁市国民经济和社会发展统计公报》（2026年5月）。")
bullet(doc, "《2026年南宁市政府工作报告》（2026年2月）及\u201c政府工作报告速览\u201d。")
bullet(doc, "南宁市发展和改革委员会《2026年国民经济和社会发展计划报告》（2026年3月）。")

# ---------------- 附录B ----------------
heading1(doc, "附录B：建议建立的年度跟踪仪表盘")
bullet(doc, "GDP总量/增速 vs 全年目标（+/-个百分点）。")
bullet(doc, "规上工业增加值及分行业（高技术/装备/电池/电子信息）增速。")
bullet(doc, "固定资产投资（总量/工业/基建/民间/房地产）增速。")
bullet(doc, "社会消费品零售总额、限上网络零售。")
bullet(doc, "外贸进出口（人民币）、出口、对东盟出口、跨境人民币结算。")
bullet(doc, "一般公共预算收入/税收/非税、财政支出、民生占比。")
bullet(doc, "常住人口、自然增长率、城镇化率、高校留邕。")
bullet(doc, "CPI/核心CPI、规上工业企业利润与营收。")
bullet(doc, "旅游人数/旅游总收入、东博会/东盟合作项目。")
bullet(doc, "平陆运河南宁段投资/进展、新能源电池产能/GWh。")

# ---------------- 保存 ----------------
out = "/Users/x/Desktop/content-prod-lab/reports/南宁市_2025年政府工作报告_深度研究_2026-08-13.docx"
doc.save(out)
print("SAVED OK 南宁市_2025年政府工作报告_深度研究_2026-08-13.docx")