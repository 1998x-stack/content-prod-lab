# -*- coding: utf-8 -*-
"""Build 宜宾市2025年政府工作报告 深度研究 DOCX, 参照地级市系列版式。"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DARK = RGBColor(0x1F, 0x2A, 0x44)
RED = RGBColor(0xB0, 0x1E, 0x2E)
GRAY = RGBColor(0x59, 0x60, 0x69)
BLUE = RGBColor(0x1F, 0x4E, 0x79)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
HEADER_FILL = "1F2A44"
K = "微软雅黑"


def set_run(run, size=11, bold=False, color=DARK, font=K):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    run.font.name = font
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), font)
    rFonts.set(qn("w:hAnsi"), font)
    rFonts.set(qn("w:eastAsia"), font)


def para(doc, text, size=11, bold=False, color=DARK, align=None,
         space_after=6, space_before=0, font=K):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        set_run(r, size=size, bold=bold or (i % 2 == 1), color=color, font=font)
    return p


def heading1(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(16)
    pf.space_after = Pt(8)
    r = p.add_run(text)
    set_run(r, size=16, bold=True, color=RED)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "B01E2E")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def heading2(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(10)
    pf.space_after = Pt(4)
    r = p.add_run(text)
    set_run(r, size=13, bold=True, color=BLUE)
    return p


def table(doc, headers, rows, widths=None, font_size=9.5):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_run(r, size=font_size, bold=True, color=WHITE)
        tcPr = hdr[i]._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), HEADER_FILL)
        tcPr.append(shd)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(val))
            set_run(r, size=font_size, color=DARK)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    return t


def quote_box(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(8)
    pf.left_indent = Pt(12)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), "B01E2E")
    pBdr.append(left)
    pPr.append(pBdr)
    r = p.add_run(text)
    set_run(r, size=11, color=DARK, font="楷体")
    return p


def bullet(doc, text, size=10.5):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.7)
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        set_run(r, size=size, bold=(i % 2 == 1), color=DARK)
    return p


# ================================================================= 文档主体
doc = Document()
sec = doc.sections[0]
sec.page_width = Cm(21)
sec.page_height = Cm(29.7)
sec.top_margin = Cm(2.4)
sec.bottom_margin = Cm(2.2)
sec.left_margin = Cm(2.5)
sec.right_margin = Cm(2.5)

st = doc.styles["Normal"]
st.font.name = K
st.font.size = Pt(11)
st.element.rPr.rFonts.set(qn("w:eastAsia"), K)


# ---- 封面 ----
para(doc, "宜宾市2025年政府工作报告", size=24, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=60, space_after=4)
para(doc, "深度研究与\u201c容易被忽视的细节\u201d分析", size=16, bold=True, color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
para(doc, "从\u201c白酒之都（五粮液）、动力电池之都、晶硅光伏、智能终端与新能源\u201d重新理解宜宾", size=12, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)
para(doc, "\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501", color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
para(doc, "研究对象：2025年宜宾市政府工作报告及同期财政、国民经济和社会发展统计资料、2026年官方复盘", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
para(doc, "标定日期：2026-08-14", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
doc.add_page_break()

# ---- 目录 ----
catalog = [
    "执行摘要",
    "一、研究口径与阅读方法",
    "二、先看宜宾的特殊底盘：白酒五粮液、动力电池、晶硅光伏、智能终端与交通物流",
    "三、最关键的宏观错位：GDP 4000亿级/5.5%低于6.5%目标，工业强但固投消费双降、物价走弱",
    "四、容易被忽视、但信号很重的细节（15条）",
    "五、2025年GDP目标 vs 实际：对照表",
    "六、2025年增长是谁撑起来的？（结构归因）",
    "七、预算与财政的\u201c含金量\u201d",
    "八、民生底账：人口、收入与城乡",
    "九、城镇与农村：格局与均衡",
    "十、人口流入与流出",
    "十一、物价与货币环境",
    "十二、区域一体化：宜宾在\u201c成渝双城经济圈+长江经济带上游\u201d里的位置",
    "十三、未来5—10年最值得观察的五条主线",
    "十四、最终结论：宜宾在\u201c白酒+动力电池+新能源\u201d里的增长逻辑",
    "附录A：主要资料来源与核验路径",
    "附录B：建议建立的年度跟踪仪表盘",
]
para(doc, "目　录", size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=10, space_after=10)
for item in catalog:
    para(doc, item, size=12, space_after=4)
doc.add_page_break()

# ---- 执行摘要 ----
heading1(doc, "执行摘要")
para(doc, "**核心判断**　2025年宜宾最显著的是\u201cGDP 4000亿级、增长5.5%（低于6.5%目标）、四川第3位\u201d、\u201c规上工业+7.7%（动力电池176GWh/+37.5%、白酒产业营业收入2040亿/+1.5%、光伏+24.9%）\u201d、\u201c但社零1187.63亿/-1.5%、固投-9.3%、CPI-0.6%\u201d、\u201c一般公共预算收入333.9亿/+3.0%\u201d。这说明宜宾在\u201c白酒+动力电池+新能源\u201d的产业升级中增长，但**消费、投资、物价下行**是短板。")
para(doc, "把2025年目标（GDP+6.5%左右/固投+6%/社零+6%/财收+3%）、2025年统计（GDP+5.5%低于目标、规上+7.7%、固投-9.3%、社零-1.5%、财收+3.0%）、趋势一起看，宜宾是\u201c白酒+制造业\u201d路径：**五粮液白酒、动力电池（全国15%/全球11%）、晶硅光伏、智能网联汽车**是支柱；2025年工业强、财政稳，但内需不足。总量居四川第3（次于成都、绵阳）。")
para(doc, "最容易记住的一句话：**宜宾是\u201c中国白酒之都（五粮液）+动力电池之都\u201d，靠\u201c白酒+动力电池+新能源\u201d增长。**观察宜宾，与其只看\u201cGDP 4000亿\u201d，不如看\u201c动力电池176GWh/+37.5%、白酒2040亿、光伏+24.9%、锂电+49.2%\u201d。")

# ---- 一、研究口径与阅读方法 ----
heading1(doc, "一、研究口径与阅读方法")
para(doc, "本报告以\u201c宜宾市政府工作报告（2025年，廖文彬作）\u201d为起点，把\u201c2025年GDP目标（6.5%左右）\u201d与\u201c官方2025年GDP增速（5.5%，2024年现价总量4008.01亿最终核实）\u201d并置对照，并用\u201c2025年宜宾经济运行情况\u201d和\u201c2026年政府工作报告复盘\u201d作为横向核验。")
para(doc, "口径提示：\u201cGDP、规上工业增加值增速\u201d按可比价（实际增速）；\u201c投资、消费、财政收入\u201d按现价（名义增速）。本报告GDP总量以\u201c2024年最终核实4008.01亿+2025年5.5%\u201d推算约4230亿元量级。涉及\u201c常住人口\u201d用常住口径约461万人，城镇化率约55%。")
para(doc, "指标体系（与研究口径一致）：核心看五个象限——**总量与增速（GDP）、产业动能（白酒/电池/光伏/汽车）、财政质量、民生底账、区域一体化（成渝+长江经济带）**。")
para(doc, "特别提示（不吃老本）：宜宾近年靠\u201c动力电池超级工厂（宁德时代四川基地）\u201d逆转经济能级，是\u201c工业强市\u201d的成功样板；但2025年GDP增速放缓到5.5%、消费投资走弱，**要看\u201c制造业强、内生需求弱\u201d的矛盾**。")
# ---- 二、先看宜宾的特殊底盘 ----
heading1(doc, "二、先看宜宾的特殊底盘：白酒五粮液、动力电池、晶硅光伏、智能终端与交通物流")
para(doc, "宜宾地处四川南部、金沙江与岷江汇合的长江零公里处，是**中国白酒之都、动力电池之都**，也是成渝双城经济圈的南翼支点。2025年GDP 4000亿级、常住人口约461万，居四川第3。")
para(doc, "四个底盘名词，先立框架：")
bullet(doc, "**白酒之都（五粮液）**　2025年优质白酒产业营业收入2040亿元（+1.5%），以五粮液为龙头的产业集群，占全国比重提高6.1pct。")
bullet(doc, "**动力电池之都**　2025年动力电池产量176GWh（+37.5%）、产销量约占全国15%、全球11%，全产业链规模1250亿（+18.1%）；宁德时代四川基地是核心引擎。")
bullet(doc, "**晶硅光伏**　2025年光伏产业规模370亿（+28.5%），单晶硅+41.5%、晶硅光伏产业+24.9%。")
bullet(doc, "**交通物流/智能终端**　长江黄金水道+宜宾港（吞吐量瞄准1100万吨）；智能终端、智能网联新能源汽车（凯翼汽车破10万辆）、数字经济630亿。")
para(doc, "这三根支柱（白酒+电池+光伏）叠加\u201c长江经济带区位\u201d，构成宜宾独特的底盘：**左手五粮液（稳现金），右手宁德时代（高成长）**。理解宜宾，先理解它从\u201c酒城\u201d到\u201c动力电池之城\u201d的产业跃迁。")

# ---- 三、最关键的宏观错位 ----
heading1(doc, "三、最关键的宏观错位：GDP 4000亿级/5.5%低于6.5%目标，工业强但固投消费双降、物价走弱")
para(doc, "2025年宜宾最需要辨析的一组\u201c错位\u201d：**GDP 5.5%未达6.5%目标（差1个百分点）、工业强（规上+7.7%）但固投-9.3%、社零-1.5%、CPI-0.6%、房地产-7.6%**。")
para(doc, "为什么\u201c工业这么强、经济却在放缓\u201d？三个解释：")
para(doc, "**其一，工业强但当量见顶**　规上工业+7.7%，动力电池+24.1%、光伏+24.9%、汽车+100.9%增长亮眼；但**工业投资-17.9%、固投-9.3%**说明产能扩张在放缓，\u201c电池热\u201d向\u201c消化存量\u201d切换。")
para(doc, "**其二，消费物价双弱**　社零1187.63亿/-1.5%（餐饮-6.3%），CPI全年-0.6%（食品烟酒-2.6%、交通通信-2.2%）——\u201c供强需弱、通缩隐忧\u201d，是报告自承的挑战。")
para(doc, "**其三，地产下行、投资转负**　固投-9.3%（工业-17.9%、三产-5.3%）、房地产投资-7.6%、商品房销售面积-17.1%——内需与地产拖累投资。")
para(doc, "小结：宜宾2025年是\u201c**强产业、弱需求**\u201d的一年：制造（电池/光伏/车）强、财政稳，但**投资、消费、地产、物价**全面走弱，需\u201c扩内需+稳投资\u201d。")

# ---- 四、容易被忽视、但信号很重的细节（15条） ----
heading1(doc, "四、容易被忽视、但信号很重的细节（15条）")
bullet(doc, "**1.动力电池产量176GWh/+37.5%，产销量占全国15%**　 宁德时代四川基地是全球电源电池重镇，全产业链规模1250亿（+18.1%）。")
bullet(doc, "**2.白酒产业营业收入2040亿（+1.5%）**　受白酒行业调整影响仅微增，但占全国比重提升6.1pct。")
bullet(doc, "**3.晶硅光伏+24.9%、单晶硅产量+41.5%**　光伏/硅料是第二增长极。")
bullet(doc, "**4.锂离子电池产量+49.2%、汽车产量+100.9%**　 高端锂电、智能网联汽车（凯翼超10万辆）爆发。")
bullet(doc, "**5.高技术制造业+14.5%**　 比规上工业平均高6.8pct——制造业升级质量高。")
bullet(doc, "**6.规上工业营收4795.85亿/+7.8%、产品销售率97.2%**　工业\u201c产销两旺\u201d。")
bullet(doc, "**7.软件信息服务+16.3%、租赁商务+9.7%**　现代服务业快于传统（住宿餐饮+4.4%）。")
bullet(doc, "**8.贷款余额6229.5亿/+17.2%**　金融有力支撑制造业；存款6367.35亿/+9.6%。")
bullet(doc, "**9.财政333.9亿/+3.0%，税收净增1.6%**　财收稳，但税收增长慢于财收（非税依赖略升）。")
bullet(doc, "**10.社零-1.5%、餐饮-6.3%**　居民消费弱；线上商品零售+99.4%（新业态爆发）。")
bullet(doc, "**11.汽车/智能终端**　凯翼汽车产销量首破10万辆，新能源整车起步。")
bullet(doc, "**12.常住人口约461万/城镇化率约55%**　四川第8人口大市；农村收入增速高于城镇。")
bullet(doc, "**13.高新技术企业**　高企及科技型中小企业目标2300家、全社会研发70亿（目标）。")
bullet(doc, "**14.环境卫生**　PM2.5浓度34.2微克/立米（首次达标二级），生态屏障改善。")
bullet(doc, "**15.\u201c新三样\u201d出口**　新能源车、锂电池、光伏\u201c新三样\u201d出口额占四川\u201c超30%\u201d——绿色智造出口强。")
# ---- 五、2025年GDP目标 vs 实际：对照表 ----
heading1(doc, "五、2025年GDP目标 vs 实际：对照表")
table(doc,
    ["指标", "2025年目标", "2025年实际", "达成评价"],
    [
        ["地区生产总值(GDP)", "增长6.5%左右", "约4230亿/5.5%", "未达成，差1pct"],
        ["固定资产投资", "力争增长6%", "-9.3%", "大幅未达成"],
        ["社会消费品零售总额", "——", "1187.63亿/-1.5%", "负增长"],
        ["一般公共预算收入", "增长3%左右", "333.9亿/+3.0%", "达成"],
        ["规上工业增加值", "——", "+7.7%", "工业强项"],
        ["进出口", "正增长(2026)", "345亿级", "结构改善"],
        ["居民收入", "与经济增长同步", "城镇+4.2%/农村+5.2%", "总体同步"],
    ],
)
para(doc, "注：GDP按可比价，5.5%未达6.5%目标；投资、消费按现价。\u201cGDP总量以2024年最终核实4008.01亿为基础推算约4230亿\u201d\u201c规上工业+7.7%高于目标预期\u201d\u201c工业强、需求弱\u201d是对照表最突出特征。")
para(doc, "拆读：**规上工业（7.7%）与财政（3.0%）达成**，**固投（-9.3%）、社零（-1.5%）大幅落后**；\u201cGDP目标6.5%\u201d实际5.5%——\u201c产业升级成、内生需求弱\u201d，是四川\u201c制造强市\u201d的典型样本。")

# ---- 六、2025年增长是谁撑起来的？（结构归因） ----
heading1(doc, "六、2025年增长是谁撑起来的？（结构归因）")
para(doc, "把宜宾GDP的5.5%拆开：三次产业增加值分别增长3.6%、5.9%、5.4%，2024年三次产业结构比约10：50：40。**第二产业（工业）是主引擎，服务业（数字经济/物流）次之，第一产业（农渔）稳定。**")
para(doc, "2026年宜宾强调\u201c生态优先绿色低碳先行区\u201d，八大产业\u201c4+4+4\u201d布局（白酒、电池、光伏、数字经济等），核心是\u201c以绿为底、以智赋能\u201d。")
para(doc, "**第二产业（工业+制造业）**：规上工业+7.7%，高技术制造+14.5%；动力电池（+24.1%）把供应链（锂电+49.2%）一起带起；光伏+24.9%、酒+1.5%——工业升级质量亮眼。")
para(doc, "**第三产业（服务业）**：软件信息+16.3%、租赁商务+9.7%、金融+6.9%、住宿餐饮+4.4%；电商线上零售+99.4%（新业态）。")
para(doc, "**外贸（开放型经济）**：\u201c新三样\u201d（新能车/锂电/光伏）出口占四川30%+；进出口虽受全球贸易扰动，但绿色出口结构优。")
para(doc, "一句话归因：**2025年宜宾增长\u201c靠工业制造（电池/光伏/车）+新质生产力\u201d**，消费/投资/地产在调整；\u201c制造强、内需弱\u201d是宜宾的核心结构性矛盾。")

# ---- 七、预算与财政的\u201c含金量\u201d ----
heading1(doc, "七、预算与财政的\u201c含金量\u201d")
para(doc, "2025年宜宾**地方一般公共预算收入333.9亿元（+3.0%）**，其中税收性收入193.92亿元（+1.6%）；一般公共预算支出702.01亿元（+4.6%）。财收与GDP、工业企业利润匹配，税收占比约58%。")
bullet(doc, "税收结构：主体税种中增值税/所得税走稳，税收+1.6%慢于财收——非税依赖略有上升，财政含金量还需观察。")
bullet(doc, "民生支出：支出602.01亿、+4.6%，投向教育、社保、医疗、农业等；育儿补贴3.4亿、困难群体补助4270万覆盖20.2万。")
bullet(doc, "金融支撑：贷款余额6229.5亿（+17.2%），存款6367.35亿（+9.6%），金融对制造业/工业设备更新支持力度大。")
para(doc, "**财政含金量小结**：财收3.0%与GDP匹配、支出稳增，**\u201c税收稳、金融宽\u201d**；财政对\u201c4+4+4\u201d产业、民生（育儿/养老）、环保的投入在加大，是\u201c制造业强市\u201d财源结构从\u201c酒税\u201d向\u201c工业税\u201d切换的体现。")

# ---- 八、民生底账：人口、收入与城乡 ----
heading1(doc, "八、民生底账：人口、收入与城乡")
para(doc, "2025年宜宾**城乡居民人均可支配收入**城镇+4.2%、农村+5.2%（与经济增长同步），农村快于城镇；常住人口约461万、城镇化率约55%。就业：城镇新增就业超目标（目标5.8万）、留宜工作毕业生增长。")
para(doc, "人口画像：宜宾是四川人口大市（第8）且**人口仍处净流入**（城镇化率提升），是四川\u201c劳动力+消费\u201d重要腹地；**白酒+电池产业创造了大量高收入岗位**，收纳农村转移劳力。")
para(doc, "民生投入：幼儿补贴3.4亿、困难群众一次性补贴4270万覆盖20.2万人、新增基本养老参保10.7万人——民生保底扎实。")

# ---- 九、城镇与农村：格局与均衡 ----
heading1(doc, "九、城镇与农村：格局与均衡")
para(doc, "宜宾常住城镇化率约55%（提升中），城乡格局均衡；农村收入增速（+5.2%）高于城镇（+4.2%），**城乡差距小幅收窄**。")
para(doc, "农业底盘：粮食总产量264.67万吨（+0.9%）、稳定在52.9亿斤以上；蔬菜364.62万吨、水果98.3万吨、茶叶11.94万吨、生猪出栏491.98万头——四川粮油猪茶农林净调出大市。")
para(doc, "一句话：\u201c宜宾农业稳、农村收入快、城镇化推进\u201d，但\u201c农业立市、制造兴市\u201d并行，城乡在产业（白酒/电池/农业）与收入（+5.2%/+4.2%）上双线并进。")

# ---- 十、人口流入与流出 ----
heading1(doc, "十、人口流入与流出")
para(doc, "宜宾常住人口约461万、**城镇化率约55%**，是四川盆地\u201c人口腹地\u201d；第三/第四产业的电池、酒制造提供了大量岗位，**人口总体平稳、劳动力净流入**（在四川地级市中较突出）。")
para(doc, "结构与挑战：**高密度、低老龄化且仍有劳动力富余**；但\u201c常住人口城镇化率约55%\u201d仍低于四川或成渝平均水平，**城镇化还有提升空间**，也是\u201c扩内需\u201d的潜力所在。")
para(doc, "2026年目标：城镇新增就业5.8万以上、加快人才落地（本科以上不少于2.5万）——宜宾把\u201c人口+就业\u201d作为城市量能的关键。")

# ---- 十一、物价与货币环境 ----
heading1(doc, "十一、物价与货币环境")
para(doc, "2025年宜宾CPI**全年下降0.6%**，其中食品烟酒-2.6%、交通通信-2.2%、居住-1.3%（支撑：银行/网购降价、产能过剩）；\u201c本地低通胀\u201d与\u201c国内食品弱势\u201d一致。")
bullet(doc, "信贷扩张：贷款余额6229.5亿（+17.2%）、存款6367.35亿（+9.6%）——融资旺盛、企业信贷充足。")
bullet(doc, "政策刺激：退出\u201c以旧换新\u201d、设备更新资金14.7亿、消费券1.9亿拉动消费超24亿。")
para(doc, "货币环境判断：**流通性充足、物价走弱**（CPI为负）、\u201c宽信用、低物价\u201d并存；2026年目标CPI+2%左右，政策收紧\u201c稳物价\u201d、扩内需。")

# ---- 十二、区域一体化：宜宾的位置 ----
heading1(doc, "十二、区域一体化：宜宾在\u201c成渝双城经济圈+长江经济带上游\u201d里的位置")
para(doc, "宜宾地处长江零公里、成渝南翼，是**成渝地区双城经济圈南部重要支点、长江经济带绿色示范，西部陆海新通道节点**。")
bullet(doc, "**交通枢纽**　内昆铁路、成贵高铁、宜宾港（长江黄金水道上游枢纽吞吐量瞄准1100万吨），陆海联动。")
bullet(doc, "**成渝分工**　宜宾承接成渝产业配套（宁德动力电池、五粮液白酒、光伏），打造\u201c川南经济区\u201d主角。")
bullet(doc, "**开放带**　对东盟、中亚、\u201c一带一路\u201d（\u201c新三样\u201d出口），港航+铁路\u201c出海\u201d。")
para(doc, "一句话：**宜宾在\u201c成渝双圈+长江经济带上游\u201d里，最核心的定位是\u201c川南制造业中心+长江港口物流枢纽\u201d**——区位与\u201c白酒+电池\u201d禀赋是它的最大护城河。")

# ---- 十三、未来5—10年最值得观察的五条主线 ----
heading1(doc, "十三、未来5—10年最值得观察的五条主线")
bullet(doc, "**主线一：动力电池产能与价格周期**\u201c240GWh产能\u201d能否消化为营收/利润（+18.1%），守住\u201c电池之都\u201d。")
bullet(doc, "**主线二：白酒穿越周期**\u201c五粮液+2040亿\u201d在需求收缩中能否稳住高端、开拓新商务/自贸。")
bullet(doc, "**主线三：光伏与新能源车（第二曲线）**\u201c光伏+24.9%、车+100.9%\u201d能否在产业内卷中向前。")
bullet(doc, "**主线四：内需与消费复苏**\u201c社零-1.5%\u201d能否在\u201c以旧换新+电商+人口导入\u201d下企稳。")
bullet(doc, "**主线五：人口与城镇化**\u201c常住461万、城镇化55%\u201d能否继续\u201c引才回流、强城\u201d。")

# ---- 十四、最终结论 ----
heading1(doc, "十四、最终结论：宜宾在\u201c白酒+动力电池+新能源\u201d里的增长逻辑")
para(doc, "把2025年闭环打穿：**宜宾是\u201c白酒之都+动力电池之都\u201d的四川制造强市**：GDP 4000亿级/+5.5%、规上工业+7.7%、财收333.9亿/+3%；酒、电池、光伏、新能源车是\u201c四柱\u201d，但消费、投资、地产、物价偏弱。")
para(doc, "宜宾不是\u201c只靠五粮液\u201d——它从\u201c酒城\u201d跃迁为\u201c电池之城\u201d，再到\u201c绿色智造\u201d（光伏/储能/数字经济），但**内生需求与投资动能承压**。")
para(doc, "一句话结论：**宜宾是\u201c长江零公里的白酒+动力电池之城\u201d；观察它先看\u201c电池、白酒、光伏\u201d，再看\u201c社零、固投、物价、人口\u201d。**它是\u201c制造强、财政稳、需求弱\u201d的四川\u201c产业强市\u201d样本。")

# ---- 附录A ----
heading1(doc, "附录A：主要资料来源与核验路径")
bullet(doc, "《2025年宜宾市政府工作报告》（2025年1月7日，廖文彬作，2025年目标、2024年回顾）")
bullet(doc, "《2025年宜宾经济运行情况》（宜宾市统计局，2026-01-24，2025年实际数据）")
bullet(doc, "《2026年宜宾市政府工作报告》（宜宾市政府，2026年1月，2025年复盘+2026年目标）")
bullet(doc, "宜宾市人民政府官网、宜宾市统计局/统计公报官方URL（yibin.gov.cn/tjj）")
bullet(doc, "聚汇数据·宜宾人口（常住人口约461万、城镇化率约55%）")

# ---- 附录B ----
heading1(doc, "附录B：建议建立的年度跟踪仪表盘")
bullet(doc, "GDP总量/增速 vs 全年目标（可比价）。")
bullet(doc, "规上工业增加值及分产业（白酒/电池/光伏/车/高技术）增速。")
bullet(doc, "动力电池产量/产能（GWh）、碳酸锂价、产能利用率。")
bullet(doc, "固定资产投资/工业投资/房地产投资增速。")
bullet(doc, "社会消费品零售总额、餐饮+线上零售。")
bullet(doc, "外贸、\u201c新三样\u201d（车/锂电/光伏）出口、进出口总额。")
bullet(doc, "一般公共预算收入、税收/非税、支出。")
bullet(doc, "常住人口、城镇化率、城镇新增就业、留宜毕业生。")
bullet(doc, "CPI、核心CPI、金融存贷款。")
bullet(doc, "五粮液营收/利润、宁德时代宜宾基地、新增专精特新、能耗。")

# ---------------- 保存 ----------------
out = "/Users/x/Desktop/content-prod-lab/reports/宜宾市_2025年政府工作报告_深度研究_2026-08-14.docx"
doc.save(out)
print("SAVED OK: 宜宾市", out)
