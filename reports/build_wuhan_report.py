# -*- coding: utf-8 -*-
"""Build 重庆市2025年政府工作报告 深度研究 DOCX, 参照北京/上海/杭州/天津系列版式。"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DARK = RGBColor(0x1F, 0x2A, 0x44)
RED = RGBColor(0xB0, 0x1E, 0x2E)
GRAY = RGBColor(0x59, 0x60, 0x69)
BLUE = RGBColor(0x1F, 0x4E, 0x79)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
HEADER_FILL = "1F2A44"
K = "微软雅黑"


def set_run(run, size=11, bold=False, color=DARK, font=K):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    run.font.name = font
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), font)
    rFonts.set(qn("w:hAnsi"), font)
    rFonts.set(qn("w:eastAsia"), font)


def para(doc, text, size=11, bold=False, color=DARK, align=None,
         space_after=6, space_before=0, font=K):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        set_run(r, size=size, bold=bold or (i % 2 == 1), color=color, font=font)
    return p


def heading1(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(16)
    pf.space_after = Pt(8)
    r = p.add_run(text)
    set_run(r, size=16, bold=True, color=RED)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "B01E2E")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def heading2(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(10)
    pf.space_after = Pt(4)
    r = p.add_run(text)
    set_run(r, size=13, bold=True, color=BLUE)
    return p


def table(doc, headers, rows, widths=None, font_size=9.5):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_run(r, size=font_size, bold=True, color=WHITE)
        tcPr = hdr[i]._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), HEADER_FILL)
        tcPr.append(shd)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(val))
            set_run(r, size=font_size, color=DARK)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    return t


def quote_box(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(8)
    pf.left_indent = Pt(12)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), "B01E2E")
    pBdr.append(left)
    pPr.append(pBdr)
    r = p.add_run(text)
    set_run(r, size=11, color=DARK, font="楷体")
    return p


def bullet(doc, text, size=10.5):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.7)
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        set_run(r, size=size, bold=(i % 2 == 1), color=DARK)
    return p


# ================================================================= 文档主体
doc = Document()
sec = doc.sections[0]
sec.page_width = Cm(21)
sec.page_height = Cm(29.7)
sec.top_margin = Cm(2.4)
sec.bottom_margin = Cm(2.2)
sec.left_margin = Cm(2.5)
sec.right_margin = Cm(2.5)

st = doc.styles["Normal"]
st.font.name = K
st.font.size = Pt(11)
st.element.rPr.rFonts.set(qn("w:eastAsia"), K)


# ---- 封面 ----
para(doc, "武汉市2025年政府工作报告", size=24, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=60, space_after=4)
para(doc, "深度研究与“容易被忽视的细节”分析", size=16, bold=True, color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
para(doc, "从“光谷、九省通衢、中部崛起与产业投资”重新理解武汉", size=12, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)
para(doc, "━━━━━━━━━━━━━━━━", color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
para(doc, "研究对象：2025年武汉市政府工作报告及同期财政、国民经济和社会发展统计资料、2026年官方复盘", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
para(doc, "整理时间：2026年8月", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
para(doc, "分析性质：分析性研究 ｜ 非官方解读", size=11, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
doc.add_page_break()

# ---- 目录 ----
catalog = [
    "执行摘要",
    "一、研究口径与阅读方法",
    "二、先看武汉的特殊底盘：九省通衢、光谷与科教大市",
    "三、最关键的宏观错位：GDP破2.21万亿、工业与投资稳，但地产偏弱",
    "四、容易被忽视、但信号很重的细节（15条）",
    "五、把所有细节连起来：武汉正在经历的“六个换挡”",
    "六、增长暗线：光谷/光电子信息/科创与新质生产力",
    "七、财政暗线：收入稳、税收平，科教/民生投入大",
    "八、产业暗线：从“汽车/钢铁”到“光谷+新兴产业+新能源”",
    "九、区域格局：中部崛起、武汉都市圈与枢纽",
    "十、人口与城市：超1370万、科教大城与都市更新",
    "十一、民营经济：市场主体活跃、科创企业多",
    "十二、事后验证：用2025年实际执行结果检验报告判断",
    "十三、未来5—10年最值得观察的五条主线",
    "十四、最终结论：武汉在“中部崛起+科技创新”里的增长逻辑",
    "附录A：主要资料来源与核验路径",
    "附录B：建议建立的年度跟踪仪表盘",
]
para(doc, "目　录", size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=10, space_after=10)
for item in catalog:
    para(doc, item, size=12, space_after=4)
doc.add_page_break()

# ---- 执行摘要 ----
heading1(doc, "执行摘要")
para(doc, "**核心判断**　如果只看新闻标题，2025年武汉最显眼的是“GDP突破2.21万亿、增长5.6%、稳居中部第一”、“进出口+12.7%”和“高技术制造业+16.6%”。但这份研究真正值得深读的，是这一座“九省通衢+光谷+科教大市”的城市如何在房地产调整（地产投资-4.1%）与一般公共预算收入低增（+4.5%、税收+0.1%）背景下，靠“光电子信息+高技术制造+出口+连锁投资+科创”撑起中部增长。")
para(doc, "把2025年初设定的目标（GDP增长6%左右）、2025年《国民经济和社会发展统计公报》、以及2026年报告对2025年的复盘放在一起看，武汉呈现清晰暗线：**从“基建/房产/传统制造”的旧依赖，向“光谷/高技术制造+科创+出口+新质生产力”转型**。旧引擎（房地产、一般基建、传统钢铁）在调整；新引擎（光电子信息、高技术制造、芯/屏/车、科创、出口）被要求更快补位。")
para(doc, "因此，本报告不按政府工作报告逐段复述，而采用“显性表述—同期数据—制度含义—长期影响”的方式，专门提取容易被忽视、但对判断武汉未来5—10年发展模式更有价值的信号。")
para(doc, "最容易记住的一句话：**武汉是“光谷+科教大城+枢纽”的中部极核城市，靠“光电子信息+科创+出口”做增量——要用“大区大学大城”的科教势能，把“武汉光谷”做成中国光谷。**观察武汉，与其看“GDP 2.2万亿”，不如看“光谷/高技术制造、科创转化、出口、产业投资与人口净流入”这几张名片。")

heading2(doc, "一页速览：2025年武汉经济的“表与里”")
table(doc,
    ["维度", "表面现象", "更深层信号"],
    [
        ["增长", "GDP 22147.35亿、+5.6%", "中部第一、服务业主导"],
        ["产业", "规上工业+6.2%、高技术+16.6%", "电子+18.9%、汽车+11.7%"],
        ["外贸", "进出口4548.5亿、+12.7%（出口+13.2%）", "出口强、一带一路+10.4%"],
        ["投资", "固定资产投资+2.0%、工业+9.7%", "工业/制造业投资强"],
        ["财政", "一般公共预算收入1743.06亿、+4.5%", "税收+0.1%偏低"],
        ["消费", "社零9013.96亿、+4.5%", "线上+12.2%、金银珠宝强"],
        ["人口", "常住1386.19万、城镇化率85.22%", "科教大市/人才强"],
        ["科创", "高新企业超1.7万家、光谷/武汉", "光谷全国一流"],
    ],
    widths=[2.2, 5.2, 8.6])
para(doc, "", size=4, space_after=2)

# ---- 一、研究口径与阅读方法 ----
heading1(doc, "一、研究口径与阅读方法")
heading2(doc, "1.1 本报告的三份核心底稿")
bullet(doc, "**2025年《政府工作报告》**：2025年1月在市十五届人大五次会议上作，给出全年预期目标（GDP增长6%左右、固定资产投资增长5.3%左右、社零增长5%左右等）。进出口等其他未全部设具体速率。")
bullet(doc, "**2025年《国民经济和社会发展统计公报》**：武汉市统计局2026年4月9日发布，给出全年实际执行数与增速，是“事后验证”的锚点。")
bullet(doc, "**2026年武汉市政府工作报告及官方复盘**：对2025执行结果的追认，交叉验证“目标—执行—再评价”三段式。")
heading2(doc, "1.2 阅读方法：显性—数据—制度—长期")
para(doc, "每一章按“**显性表述 → 同期数据 → 制度含义 → 长期影响**”四层展开。其中“制度含义”回答“政府为什么这么排优先序”，“长期影响”回答“这对未来5—10年意味着什么”。")
para(doc, "关键判别：**当报告使用的“增长词”和统计公报里的实际数不一致时，数据优先。**例如2025年GDP目标6%左右、实际5.6%；固定资产投资目标5.3%、实际2.0%；税收仅+0.1%。差异反映：武汉“高技术/出口强、地产/税收偏弱”。")

# ---- 二、特殊底盘 ----
heading1(doc, "二、先看武汉的特殊底盘：九省通衢、光谷与科教大市")
para(doc, "在所有中部城市里，武汉的“底盘”独特：**九省通衢交通枢纽+中国光谷+科教大城（高校/院所密集）**三合一。常住1386万、城镇化率85%，是“大区大学大城”的中部极核。")
para(doc, "这决定武汉的多重身份并存：**光谷（光电子信息、集成电路）**、**科教与科创（高校、两院院士）**、**汽车与制造（东风、智能网联汽车）**、**枢纽（九省通衢、中欧班列枢纽）**、**中部金融/商贸中心**。")
heading2(doc, "2.1 光谷与新质")
para(doc, "武汉光谷是全国光电子/半导体高地，高技术制造+16.6%、电子+18.9%。芯/屏/光谷/车等是“新质生产力”主力（显示器、电子计算机整机、光纤、成套装备）。")
heading2(doc, "2.2 科教/科创")
para(doc, "武汉高校/科研院所密集（两院院士92人、重点实验室41个、高新技术企业超1.7万家）。技术合同成交2728亿，光谷/武汉科创转化是全国标杆。")
heading2(doc, "2.3 枢纽与开放")
para(doc, "武汉九省通衢，中欧班列（武汉）、多式联运连接“一带一路”；进出口4548.5亿、+12.7%、出口+13.2%，是中部内陆开放门户。")

# ---- 三、宏观错位 ----
heading1(doc, "三、最关键的宏观错位：GDP破2.21万亿、工业与投资稳，但地产偏弱")
para(doc, "把2025年武汉的宏观面放进一张表，会出现令人意外的“错位”：表观增长来自工业/出口/高技术，而地产与税收偏弱。这个错位，正是读懂武汉的关键。")
heading2(doc, "3.1 “强的一面”")
table(doc,
    ["指标", "2025实绩", "信号"],
    [
        ["GDP", "22147.35亿、+5.6%", "中部第一"],
        ["规上工业", "+6.2%", "电子+18.9%、汽车+11.7%"],
        ["高技术制造", "+16.6%", "占总26.2%"],
        ["进出口", "+12.7%（出口+13.2%）", "出口强"],
        ["固投", "+2.0%（工业+9.7%）", "工业投资强"],
        ["高新技术产业", "占GDP32.5%", "新质生产力"],
    ],
    widths=[3.2, 5.2, 6.0])
heading2(doc, "3.2 “弱的一面”")
table(doc,
    ["指标", "2025实值", "信号"],
    [
        ["房地产开发投资", "-4.1%", "地产调整"],
        ["税收收入", "+0.1%", "税收增长弱"],
        ["一般公共预算收入", "+4.5%", "收入低增"],
        ["三产/服务业", "+5.9%", "主引擎仍偏服务"],
        ["全社会用电", "+4.7%", "工业偏弱相关"],
    ],
    widths=[3.2, 5.2, 6.0])
para(doc, "**错位结论**　武汉的增长“很强、但也很矛盾”。强的部分（高技术/出口/固投/科创）与弱的部分（地产/税收）并存。**真正的焦点是“科创/制造强，但地产/税收承压”**：高技术+16.6%拉动增长，但地产-4.1%、税收+0.1%。2026年武汉“强化科教/科创+稳定地产/税收+扩内需”是重点。")

# ---- 四、被忽视细节 ----
heading1(doc, "四、容易被忽视、但信号很重的细节（15条）")
details = [
    ("1", "规上工业+6.2%、电子/通信+18.9%", "光电子信息全国领先。"),
    ("2", "高技术制造+16.6%、占规上26.2%", "新质生产力占比高。"),
    ("3", "汽车制造+11.7%（产量81万辆）", "汽车/新能源回暖。"),
    ("4", "化学原料/电气机械+10.1%/+11.5%", "材料/装备景气。"),
    ("5", "规模工业利润总额+46.8%", "利润大幅修复。"),
    ("6", "进出口4548.5亿、+12.7%、出口+13.2%", "出口强、稳外贸。"),
    ("7", "对一带一路+10.4%、对外工程+24.0%", "一带一路走深。"),
    ("8", "固定资产投资+2.0%、工业+9.7%", "工业投资强。"),
    ("9", "高新区/产业投资强（制造业+10.1%）", "先进制造投资。"),
    ("10", "社零9013.96亿、+4.5%、线上+12.2%", "消费/线上存力。"),
    ("11", "实际使用外资9.01亿美元、企业新批+5.2%", "外资稳。"),
    ("12", "常住1386.19万、城镇化率85.22%", "科教大市。"),
    ("13", "两院院士92人、高新企业超1.7万家", "科创能级强。"),
    ("14", "技术合同登记成交2728亿、+4.6%", "科创转化。"),
    ("15", "市场主体269.89万户、新登记61.48万（+10.8%）", "市场主体活跃、新增量多。"),
]
for d in details:
    para(doc, "**%s**" % d[0], size=11)
    para(doc, d[1], size=10.5, color=GRAY)

# ---- 五、六个换挡 ----
heading1(doc, "五、把所有细节连起来：武汉正在经历的“六个换挡”")
gear = [
    ("1．动能换挡：从“基建/房产驱动”到“光谷+高技术制造+科创”", "高技术+16.6%、光电子/芯屏车放量。"),
    ("2．产业换挡：从“汽车/钢铁”到“光电子/集成电路+新质生产力”", "电子+18.9%、显示器/计算机/光纤。"),
    ("3．投资换挡：从“基建/地产”到“工业/制造业/高技术”", "工业+9.7%、制造业+10.1%、地产-4.1%。"),
    ("4．开放换挡：从“内贸”到“一带一路/出口”", "出口+13.2%、一带一路+10.4%。"),
    ("5．人口换挡：从“零散流入”到“科教/人才集聚”", "常住1386万、科教大市吸附。"),
    ("6．动能换挡：从“传统增长”到“新质生产力+市场化”", "高新企业1.7万、技术合同2728亿。"),
]
for g in gear:
    para(doc, "**%s**　%s" % g, space_after=6)

# ---- 六、增长暗线 ----
heading1(doc, "六、增长暗线：光谷/光电子信息/科创与新质生产力")
heading2(doc, "6.1 光谷+高技术")
para(doc, "武汉光谷是全国光电子/半导体高地，高技术制造+16.6%、电子/通信+18.9%。芯（集成电路）、屏（显示器）、光谷、车（智能网联）构成“新质生产力”主线。")
heading2(doc, "6.2 科教与科创转合")
para(doc, "武汉两院院士92人、高校密集、国家实验室/重点实验室，高新技术企业超1.7万家、技术合同成交2728亿。科教势能→产业转化，是武汉从“制造”到“智造/科创”的增量。")
para(doc, "**这条暗线意味着**：武汉的增长叙事正从“基建/房产”转向“光谷+高技术制造+科创转化”。看武汉，盯住“高技术制造占比、光电/芯片、技术合同、出口、人口/人才”这几组数。")

# ---- 七、财政暗线 ----
heading1(doc, "七、财政暗线：收入稳、税收平，科教/民生投入大")
para(doc, "2025年武汉一般公共预算收入1743.06亿、+4.5%，税收1225.10亿、+0.1%。收入低增、税收近乎持平，反映地产弱/价格通缩对税源影响；支出向科教、民生、都市更新倾斜。")
para(doc, "**制度含义**　武汉财政“收入稳、税收平”，长期要靠“科创+新质生产力+出口”带来真实税源，摆脱对地产依赖。")

# ---- 八、产业暗线 ----
heading1(doc, "八、产业暗线：从“汽车/钢铁”到“光谷+新兴产业+新能源”")
heading2(doc, "8.1 武汉产业的“表”")
table(doc,
    ["指标", "2025增速/信号", "解读"],
    [
        ["规上工业", "+6.2%", "制造底盘"],
        ["计算机/通信电子", "+18.9%", "光电子/芯片"],
        ["高技术制造", "+16.6%、占26.2%", "新质生产力"],
        ["汽车制造", "+11.7%", "汽车/新能源"],
        ["规模工业利润", "+46.8%", "利润修复"],
        ["高新技术产业", "占GDP32.5%", "科创驱动"],
        ["进出口", "+12.7%", "出口强"],
    ],
    widths=[4.6, 3.6, 5.0])
heading2(doc, "8.2 从“汽车/钢铁”到“光电/科创”")
para(doc, "武汉过去以汽车（东风）、钢铁见长，2025年显示“光电子/集成电路+高技术制造+新质生产力”正成为新主导。光谷、芯/屏/车、科创，是武汉产业升级的新底盘。")

# ---- 九、区域 ----
heading1(doc, "九、区域格局：中部崛起、武汉都市圈与枢纽")
para(doc, "武汉是中部最大城市、国家中心城市，武汉都市圈带动鄂州、黄石、孝感等鄂东。九省通衢的交通枢纽（中欧班列、多式联运）是中部开放门户。")
para(doc, "在“中部崛起”战略下，武汉以“强省会+都市圈+枢纽”拉动湖北及中部腹地。光谷+武汉都市圈，是中部增长极。")

# ---- 十、人口与城市 ----
heading1(doc, "十、人口与城市：超1370万、科教大城与都市更新")
heading2(doc, "10.1 人口总量")
para(doc, "2025年末武汉常住1386.19万人、城镇化率85.22%，户籍961.34万。武汉依托科教/就业吸引中原人口、人才。")
heading2(doc, "10.2 城市更新/都市圈")
para(doc, "武汉推进城中村改造（3.22万户）、老旧小区改造（322个）、保障性住房，是中部都市更新/韧性城市先行。")

# ---- 十一、民营经济 ----
heading1(doc, "十一、民营经济：市场主体活跃、科创企业多")
heading2(doc, "11.1 市场主体")
para(doc, "武汉市场主体269.89万户、新登记61.48万（+10.8%），企业127.29万、增长20.7%。民营与科创活跃。")
heading2(doc, "11.2 政策/科创")
para(doc, "武汉优化营商环境、支持科技成果转化。光谷民营企业、专精特新、高新技术企业，是武汉未来10年底盘。")

# ---- 十二、事后验证 ----
heading1(doc, "十二、事后验证：用2025年实际执行结果检验报告判断")
heading2(doc, "12.1 年初目标与年末执行对比")
table(doc,
    ["指标", "2025年初目标", "2025实际", "偏差"],
    [
        ["GDP增速", "6%左右", "+5.6%", "略低"],
        ["固定资产投资", "5.3%左右", "+2.0%", "低于目标"],
        ["社零", "5%左右", "+4.5%", "略低"],
        ["进出口", "未设", "+12.7%", "超预期"],
        ["规上工业", "力争", "+6.2%", "达预期"],
        ["高新技术产业", "—", "占GDP32.5%", "强"],
    ],
    widths=[3.0, 3.2, 3.0, 4.2])
heading2(doc, "12.2 判断验证")
para(doc, "三条核心判断得到支持：**（1）高技术/出口/固投强（高技术+16.6%、出口+13.2%），验证；（2）地产/税收偏弱（地产-4.1%、税收+0.1%），验证；（3）科教/科创/新质强（高新企业1.7万、光谷），验证。**")
para(doc, "核心观察：武汉靠“光谷+高技术+出口+科创”守住中部第一（2.21万亿），但地产、税收偏弱。2026年“科教/科创+稳地产/税收+扩内需”是重点。")

# ---- 十三、五条主线 ----
heading1(doc, "十三、未来5—10年最值得观察的五条主线")
lines5 = [
    ("① 光谷/光电/集成电路升级", "高技术占26%、光电/芯片，能否向高水平自主。"),
    ("② 科教→科创转化", "两院院士、高校、光谷实验室，能否孵化更多独角兽。"),
    ("③ 出口/一带一路/内陆开放", "出口+13%、中欧班列，能否把“九省通衢”变“开放门户”。"),
    ("④ 汽车/新能源/新质", "汽车回暖，能否在新能源/智能化上做强。"),
    ("⑤ 地产/财政/人口平衡", "地产弱、税收平，能否在都市圈/科教中稳住。"),
]
for l in lines5:
    para(doc, "**%s**　%s" % l, space_after=6)

# ---- 十四、结论 ----
heading1(doc, "十四、最终结论：武汉在“中部崛起+科技创新”里的增长逻辑")
para(doc, "武汉的2025年，本质上是“**光谷+科教科创+高技术制造+出口为核心，而地产/税收偏弱**”的答卷：GDP中部第一、高技术+16.6%、出口+13.2%，代价是地产-4.1%、税收+0.1%。")
para(doc, "只要光电子、科教科创、出口、高技术制造能接住，武汉就站牢中部极核；如果地产/税收/内需持续偏弱，武汉需承受“科创强、地产弱”的结构缺口。")
para(doc, "最稳妥的观察信号：**一盯高技术/光电/芯屏（动能）、二盯科教/科创（长长期）、三盯出口/枢纽（开放）、四盯地产/税收（约束）、五盯人口/人才（底座）。**武汉，是“光谷+科教”的中部极核。")

# ---- 附录A ----
heading1(doc, "附录A：主要资料来源与核验路径")
bullet(doc, "武汉市2025年《政府工作报告》（2025年1月）——目标来源。")
bullet(doc, "《2025年武汉市国民经济和社会发展统计公报》（武汉市统计局，2026-04-09）——GDP、工业、外贸、人口实值。")
bullet(doc, "武汉市统计局/光谷专报、科创专题——光电子与科创。")
bullet(doc, "2026年武汉市政府工作报告——2025执行复盘。")
bullet(doc, "武汉海关、市商务局、财政局——外贸/财政。")
heading2(doc, "核验说明")
para(doc, "本报告所有增速、总量、占比均以统计公报/官方发布为基准。涉“光谷具体产值”“常住精确值”等以官方口径为准。")

# ---- 附录B ----
heading1(doc, "附录B：建议建立的年度跟踪仪表盘")
para(doc, "建议把下面10个“测脉搏”指标做成年度跟踪表：")
dash = [
    ["#", "指标", "2025基值", "为什么重要"],
    ["1", "GDP增速", "+5.6%", "总量与方向"],
    ["2", "规上工业增速", "+6.2%", "工业底盘"],
    ["3", "高技术制造增速/占比", "+16.6%/26.2%", "新质生产力"],
    ["4", "进出口/出口增速", "+12.7%/+13.2%", "外贸韧性"],
    ["5", "固定资产投资/工业投资", "+2.0%/+9.7%", "投资结构"],
    ["6", "社零增速", "+4.5%", "内需消费"],
    ["7", "一般公共预算收入/税收", "+4.5%/+0.1%", "财政质量"],
    ["8", "常住人口/城镇化率", "1386.19万/85.22%", "人口与城市"],
    ["9", "高新技术企业", "超1.7万家", "科创能级"],
    ["10", "居民人均可支配收入增速", "+4.7%", "民生获得感"],
]
table(doc, dash[0], dash[1:], widths=[0.8, 4.5, 3.2, 4.6])
para(doc, "把这10个指标连起来看，任何一个“新引擎（3/4）向上、旧引擎（7）修复”，都说明武汉在真正“换挡”。")

# ========================================= 保存
out = "/Users/x/Desktop/content-prod-lab/reports/武汉市_2025年政府工作报告_深度研究_2026-08-13.docx"
doc.save(out)
print("SAVED:", out)
print("PARAGRAPHS:", len(doc.paragraphs))
print("TABLES:", len(doc.tables))
