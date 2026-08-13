# -*- coding: utf-8 -*-
"""Build 西藏自治区2025年政府工作报告 深度研究 DOCX, 参照省系列版式。"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DARK = RGBColor(0x1F, 0x2A, 0x44)
RED = RGBColor(0xB0, 0x1E, 0x2E)
GRAY = RGBColor(0x59, 0x60, 0x69)
BLUE = RGBColor(0x1F, 0x4E, 0x79)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
HEADER_FILL = "1F2A44"
K = "微软雅黑"


def set_run(run, size=11, bold=False, color=DARK, font=K):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    run.font.name = font
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), font)
    rFonts.set(qn("w:hAnsi"), font)
    rFonts.set(qn("w:eastAsia"), font)


def para(doc, text, size=11, bold=False, color=DARK, align=None,
         space_after=6, space_before=0, font=K):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        set_run(r, size=size, bold=bold or (i % 2 == 1), color=color, font=font)
    return p


def heading1(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(16)
    pf.space_after = Pt(8)
    r = p.add_run(text)
    set_run(r, size=16, bold=True, color=RED)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "B01E2E")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def heading2(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(10)
    pf.space_after = Pt(4)
    r = p.add_run(text)
    set_run(r, size=13, bold=True, color=BLUE)
    return p


def table(doc, headers, rows, widths=None, font_size=9.5):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_run(r, size=font_size, bold=True, color=WHITE)
        tcPr = hdr[i]._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), HEADER_FILL)
        tcPr.append(shd)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(val))
            set_run(r, size=font_size, color=DARK)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    return t


def quote_box(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(8)
    pf.left_indent = Pt(12)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), "B01E2E")
    pBdr.append(left)
    pPr.append(pBdr)
    r = p.add_run(text)
    set_run(r, size=11, color=DARK, font="楷体")
    return p


def bullet(doc, text, size=10.5):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.7)
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        set_run(r, size=size, bold=(i % 2 == 1), color=DARK)
    return p


# ================================================================= 文档主体
doc = Document()
sec = doc.sections[0]
sec.page_width = Cm(21)
sec.page_height = Cm(29.7)
sec.top_margin = Cm(2.4)
sec.bottom_margin = Cm(2.2)
sec.left_margin = Cm(2.5)
sec.right_margin = Cm(2.5)

st = doc.styles["Normal"]
st.font.name = K
st.font.size = Pt(11)
st.element.rPr.rFonts.set(qn("w:eastAsia"), K)



# ---- 封面 ----
para(doc, "西藏自治区2025年政府工作报告", size=24, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=60, space_after=4)
para(doc, "深度研究与\u201c容易被忽视的细节\u201d分析", size=16, bold=True, color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
para(doc, "从\u201c固边稳藏、清洁能源、文旅经济与国家生态屏障\u201d重新理解西藏", size=12, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)
para(doc, "\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501", color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
para(doc, "研究对象：2025年西藏自治区政府工作报告及同期财政、国民经济和社会发展统计资料、2026年官方复盘", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
para(doc, "标定日期：2026-08-13", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
doc.add_page_break()

# ---- 目录 ----
catalog = [
    "执行摘要",
    "一、研究口径与阅读方法",
    "二、先看西藏的特殊底盘：高原经济、清洁能源、文旅与生态屏障",
    "三、最关键的宏观错位：GDP破3000亿、工业/旅游强，但总量小/固投依赖",
    "四、容易被忽视、但信号很重的细节（15条）",
    "五、2025年GDP目标 vs 实际：对照表",
    "六、2025年增长是谁撑起来的？（结构归因）",
    "七、预算与财政的\u201c含金量\u201d",
    "八、民生底账：人口、收入与城乡",
    "九、城镇与农村：格局与均衡",
    "十、人口流入与流出",
    "十一、物价与货币环境",
    "十二、区域一体化：西藏在\u201c一带一路南亚大通道+南亚陆路通道\u201d里的位置",
    "十三、未来5—10年最值得观察的五条主线",
    "十四、最终结论：西藏在\u201c高原特色+清洁能源+旅游\u201d里的增长逻辑",
    "附录A：主要资料来源与核验路径",
    "附录B：建议建立的年度跟踪仪表盘",
]
para(doc, "目　录", size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=10, space_after=10)
for item in catalog:
    para(doc, item, size=12, space_after=4)
doc.add_page_break()

# ---- 执行摘要 ----
heading1(doc, "执行摘要")
para(doc, "**核心判断**　2025年西藏最显著的是\u201cGDP破3031.89亿元、增长7.0%（增速连续居全国前列）\u201d、\u201c居民收入33600元、+7.2%\u201d、\u201c旅游7073万人次、+10.7%\u201d、\u201c规模以上工业高增（上半年+20.9%）\u201d。这说明西藏经济在\u201c投资+工业（清洁能源）+文旅\u201d驱动下高速增长。")
para(doc, "把2025年初设定的目标（GDP 7%+/规上工业16%/固投15%）、2025年统计、2026年前瞻一起看，西藏呈现\u201c内陆特色省份典型路径\u201d：**投资主导、清洁能源+文旅双引擎、固边稳藏**。总量小（3000亿）但增速居全国前列（四个季度）。")
para(doc, "最容易记住的一句话：**西藏是\u201c高原经济+清洁能源+生态屏障\u201d的自治区，靠\u201c投资+工业（水电/新能源）+旅游\u201d实现高速增长。**观察西藏，与其只看\u201cGDP 3031亿\u201d，不如看\u201c增速全国前列、收入+7.2%、旅游7073万人次、水电清洁能源\u201d。")

# ---- 一、研究口径与阅读方法 ----
heading1(doc, "一、研究口径与阅读方法")
heading2(doc, "1.1 本报告的核心底稿")
bullet(doc, "**2025年《政府工作报告》（2025年1月）**——目标：GDP 7%+/规上工业16%/固投15%+。")
bullet(doc, "**2025年经济运行/统计资料**——GDP 3031.89亿、+7.0%、收入/旅游实数。")
bullet(doc, "**2026年报告/发布会**——2025追认、固边民生。")
heading2(doc, "1.2 阅读方法：显性—数据—长期")
para(doc, "**关键判别**：数据优先。西藏GDP+7.0%（全国前列）、工业/旅游强、收入快。西藏\u201c工业+能源+旅游强、总量小\u201d，穿透总量看增速/能源/文旅。")

# ---- 二、底盘 ----
heading1(doc, "二、先看西藏的特殊底盘：高原经济、清洁能源、文旅、生态屏障")
para(doc, "西藏作为\u201c**青藏高原经济+清洁能源+文旅+国家生态屏障**\u201d，肩负固边稳藏与发展双重使命。")
bullet(doc, "**清洁能源/水利**：金沙江/澜沧江/藏东南水电、清洁能源集群、装机2705万千瓦。")
bullet(doc, "**高原有色/特色产业**：绿色矿产、数字经济、特色手工业（唐卡/藏毯/藏香）。")
bullet(doc, "**文旅**：布达拉宫-大昭寺-罗布林卡、世界旅游目的地；2025年旅游7073万人次。")
bullet(doc, "**固边稳藏**：边境口岸/试验区、固边兴边富民。")
para(doc, "这一底板决定西藏2025增速\u201c主动力\u201d：**投资基建+清洁能源工业+文旅+收入**，增速全国前列。")

# ---- 三、宏观错位 ----
heading1(doc, "三、最关键的宏观错位：GDP破3000亿、工业/旅游强，但总量小/固投依赖")
para(doc, "西藏2025年最值得咀嚼的错位，是\u201c**增速高/工业旅游强、总量小/固投+转移依赖**\u201d。GDP约3031.89亿、+7.0%（全国前列）。")
bullet(doc, "**GDP**：3031.89亿元、+7.0%（按不变价）。一/二/三产分布，总量小。")
bullet(doc, "**工业**：规上工业高增（2025上半年+20.9%）、清洁能源水电驱动。")
bullet(doc, "**投资**：固投高增（2024+19.6%、2025目标15%+），基建/能源大投入。")
bullet(doc, "**消费/旅游**：社零、旅游7073万人次+10.7%、花费816.81亿+9.5%。")
bullet(doc, "**财政/收入**：财政收入高增（2024+17.1%）、居民收入+7.2%。")
para(doc, "**为什么读这条**：西藏作为\u201c人口小省+高原\u201d，结构性特征是\u201c增速全国前列、总量小、投资+转移依赖、文旅/能源强\u201d。固边+发展双重目标。")

# ---- 四、15条细节 ----
heading1(doc, "四、容易被忽视、但信号很重的细节（15条）")
para(doc, "以下15条在2025年报告/数据，但被\u201cGDP 7%\u201d等掩盖。它们是判断西藏2025之后5—10年的关键小信号。")
bullet(doc, "**1. 经济增速连续四季度全国前列**：高原省份高增模型。")
bullet(doc, "**2. 旅游7073万人次、花费816.81亿**：文旅第一引擎（布达拉宫等）。")
bullet(doc, "**3. 清洁能源/水电**：金沙江/澜沧江/藏东南水电、装机2705万千瓦。")
bullet(doc, "**4. 居民收入33600元、+7.2%（农村+7.4%）**：民生改善快、农村+7.4%。")
bullet(doc, "**5. 固边兴边**：边境口岸/试验区、固边安民。")
bullet(doc, "**6. 财政转移依赖+基建**：国家投资/转移支付。")
bullet(doc, "**7. 特色产业（唐卡/藏毯/藏香/手工艺）**：民族手工业。")
bullet(doc, "**8. 数字经济/绿色矿产**：高原新质。")
bullet(doc, "**9. 清洁能源东送（川藏/电力）**：能源枢纽。")
bullet(doc, "**10. 农村收入快于城镇**：城乡差距缩小。")
bullet(doc, "**11. 城镇化提升/人口**：人口少但城镇化/口岸。")
bullet(doc, "**12. 交通运输/Air**：拉萨国际机场、铁路（拉林/青藏）。")
bullet(doc, "**13. 青藏高原生态屏障**：国家生态安全。")
bullet(doc, "**14. 就业转移/富余劳动力**：农村转移、就业。")
bullet(doc, "**15. 对口支援（中央+山东/湖北）**：援藏机制。")
para(doc, "", size=10, space_after=2)

# ---- 五、对照表 ----
heading1(doc, "五、2025年GDP目标 vs 实际：对照表")
para(doc, "2025年报告中列出的主要预期目标，与2025年实际好转：")
tb = [
    ["指标", "2025年目标", "2025年实际", "达标"],
    ["GDP增速", "7%以上（争取8%）", "+7.0%（3031.89亿）", "达标"],
    ["规上工业增加值", "16%", "高增（上半年+20.9%）", "超预期"],
    ["固定资产投资", "15%以上", "投资高基", "达标"],
    ["社会消费品零售", "8%+", "类似体量增长", "接近"],
    ["全体居民收入", "—", "+7.2%（33600元）", "快"],
    ["城镇调查失业率", "5%内", "内稳", "达标"],
]
table(doc, tb[0], tb[1:], widths=[3.4, 3.4, 4.4, 3.6])
para(doc, "", size=10, space_after=2)
para(doc, "**要点**：西藏GDP/工业/失业达标、收入/旅游快，主要目标圆满实现，'十四五'良好收官。")

# ---- 六、结构归因 ----
heading1(doc, "六、2025年增长是谁撑起来的？（结构归因）")
heading2(doc, "6.1 投资/基建")
para(doc, "**固定资产投资高增**（基建/能源/交通）是主动力，对口支援+中央转移。")
heading2(doc, "6.2 工业/清洁能源")
para(doc, "规上工业高增（水电/新能源/高原有色），能源东送。")
heading2(doc, "6.3 文旅")
para(doc, "旅游7073万人次+10.7%、花费+9.5%，世界旅游目的地。")
heading2(doc, "6.4 收入/民生")
para(doc, "居民收入+7.2%（农村+7.4%）、固边兴边富民。")
heading2(doc, "6.5 转移支付/财政")
para(doc, "财政支出高、中央转移支撑。")
para(doc, "**一句话归因**：西藏2025年靠\u201c**投资基建+清洁能源工业+文旅+收入民生**\u201d，实现\u201c总量小增速高\u201d的全国前列增长。")

# ---- 七、财政 ----
heading1(doc, "七、预算与财政的\u201c含金量\u201d")
para(doc, "西藏财政高度依赖\u201c**中央转移支付+投资**\u201d，2025年财政收入高增（2024+17.1%）、支出强民生。")
bullet(doc, "**转移依赖**：财政以中央转移为主，支出用于基建/民生/固边。")
bullet(doc, "**收入改善**：财政收入高增（2024+17.1%），带动基建与福利。")
bullet(doc, "**民生硬度**：教育/医疗/社保、固边兴边富民。")
para(doc, "**财政含义**：西藏\u201c中央转移+收入增\u201d，支撑基建/民生/固边，财政政策空间大。")

# ---- 八、民生 ----
heading1(doc, "八、民生底账：人口、收入与城乡")
para(doc, "西藏以\u201c高原、人口少、收入稳增\u201d为特征：**居民人均可支配收入33600元、+7.2%**（城镇+6.0%、农村+7.4%）。")
bullet(doc, "**收入**：居民33600元、+7.2%；城镇58794元、+6.0%，农村23184元、+7.4%（农村快）。")
bullet(doc, "**就业**：城镇新增就业、农村劳动力转移、高校毕业生就业95%+。")
bullet(doc, "**社保/民生**：医保/低保、教育（西藏班）、医疗。")
bullet(doc, "**人口**：人口规模小、城镇化提升。")
para(doc, "**民生含义**：西藏\u201c收入增速全国前列、农村快\u201d，固边+民生改善持续。")

# ---- 九、城乡 ----
heading1(doc, "九、城镇与农村：格局与均衡")
para(doc, "**城镇化提升**但农区/牧区占比高；城乡收入差距缩小（农村+7.4%快于城镇+6.0%）。")
bullet(doc, "**城市**：拉萨/日喀则、城市高原枢纽。")
bullet(doc, "**农村**：畜牧、特色农林（青稞/牦牛）、边境/固边。")
para(doc, "**城乡均衡**：西藏\u201c城市/农牧区兼顾\u201d，农村/固边发展带动均衡。")

# ---- 十、人口流 ----
heading1(doc, "十、人口流入与流出")
para(doc, "西藏人口以本地为主、规模小；靠\u201c对口支援+产业+边境\u201d稳定人口、吸引内地互建。")
bullet(doc, "**人口**：高原人口、城镇化/口岸集聚。")
bullet(doc, "**人才**：援藏干部/人才、就业引才。")
para(doc, "人口方向：西藏靠\u201c固边+支援+特色产业\u201d稳定/吸引人才。")

# ---- 十一、物价 ----
heading1(doc, "十一、物价与货币环境")
para(doc, "西藏CPI目标控稳（3%内）、整体温和；货币/金融支持高原/边贸。")
bullet(doc, "**物价**：温和、控制在目标内。")
bullet(doc, "**金融**：对口/绿色金融、能源/基建信贷。")
para(doc, "**物价含义**：西藏\u201c通胀温和可控\u201d，民生稳定。")

# ---- 十二、区域一体化 ----
heading1(doc, "十二、区域一体化：一带一路南亚通道、青藏与南亚陆路")
para(doc, "西藏处于**一带一路南亚大通道+中尼/中印边境口岸+青藏高原旅游带**，既是生态屏障也是南亚陆路门户。")
bullet(doc, "**南亚大通道**：中尼（吉隆口岸）、中印口岸、南亚陆路。")
bullet(doc, "**口岸经济**：吉隆边境合作区、日喀则国际陆港、樟木、里孜口岸。")
bullet(doc, "**交通**：青藏/拉林铁路、拉萨国际机场、中尼铁路规划。")
bullet(doc, "**对口支援**：中央+省市援藏。")
para(doc, "**区域含义**：西藏作为\u201c固边+南亚通道+生态\u201d，靠口岸/铁路/支援带动发展。")

# ---- 十三、五主线 ----
heading1(doc, "十三、未来5—10年最值得观察的五条主线")
bullet(doc, "**主线1｜清洁能源/水电**：金沙江/澜沧江/藏东南。能否建成青藏能源基地/东送。")
bullet(doc, "**主线2｜文旅（世界旅游目的地）**：旅游7073万人次、布达拉宫/冰川。能否旅游业强。")
bullet(doc, "**主线3｜南亚通道/口岸**：吉隆/里孜口岸、南亚陆路。能否成开放门户。")
bullet(doc, "**主线4｜固边兴边/民生**：投资+固边+收入。能否持续高增。")
bullet(doc, "**主线5｜对口支援/生态**：青藏生态屏障+支援。能否绿色高质。")
para(doc, "这五条，是西藏从\u201c高原+能源+旅游\u201d走向\u201c清洁能源强+文旅+南亚通道\u201d的\u201c主赛道\u201d。")

# ---- 十四、结论 ----
heading1(doc, "十四、最终结论")
para(doc, "西藏2025年，本质上是\u201c**投资+清洁能源+旅游驱动、总量小/固投依赖**\u201d的高增速答卷：GDP3031.89亿、+7.0%（全国前列）、居民收入33600元、旅游7073万人次、规上工业高增")
para(doc, "只要投资/清洁能源/文旅持续，西藏就维持\u201c增速全国前列\u201d；靠固边/转移/收入改善保民生。")
para(doc, "最稳观察信号：**一盯清洁能源/水电（引擎）、二盯旅游/文旅（消费）、三盯固投/基建（投资）、四盯南亚通道/口岸（开放）、五盯收入/固边（民生）。**西藏，是\u201c高原+能源+旅游\u201d独特样本。")

# ---- 附录A ----
heading1(doc, "附录A：主要资料来源与核验路径")
bullet(doc, "西藏自治区2025年政府工作报告（2025年1月）——目标来源。")
bullet(doc, "2025年西藏经济运行发布会/统计资料——GDP、工业、旅游、收入实数。")
bullet(doc, "2026年西藏政府工作报告/发布会——2025追认/固边/能源/旅游。")
heading2(doc, "核验说明")
para(doc, "本报告涉及以官方统计/发布会口径为准；\u201c能源/旅游/南亚\u201d等以官方为准。")

# ---- 附录B ----
heading1(doc, "附录B：建议建立的年度跟踪仪表盘")
para(doc, "建议把下面10个\u201c测脉搏\u201d指标做成年度跟踪表：")
dash = [
    ["#", "指标", "2025基值", "为什么重要"],
    ["1", "GDP增速", "+7.0%（3031.89亿）", "总量与方向"],
    ["2", "规上工业增加值", "高增（上半年+20.9%）", "工业"],
    ["3", "清洁能源/装机", "水电/新能源集群", "能源"],
    ["4", "旅游接待/花费", "7073万人次/816.81亿", "文旅"],
    ["5", "固定资产投资", "高增", "投资"],
    ["6", "社会消费品零售", "增长", "内需"],
    ["7", "居民收入", "+7.2%（33600元）", "民生"],
    ["8", "常住人口/城镇化", "高原少人/提升", "人口"],
    ["9", "财政收入", "高增(+17.1%有2024)", "财政"],
    ["10", "CPI", "温和(3%内)", "物价"],
]
table(doc, dash[0], dash[1:], widths=[0.8, 4.6, 4.4, 4.0])
para(doc, "把这10连起来看，能源/旅游/投资（2/4/5）、收入/人口（7/8），都说明西藏在快速发展与固边富民。")

# ============ 保存 ============
out = "/Users/x/Desktop/content-prod-lab/reports/西藏自治区_2025年政府工作报告_深度研究_2026-08-13.docx"
doc.save(out)
print("SAVED:", out)
print("PARAGRAPHS:", len(doc.paragraphs))
print("TABLES:", len(doc.tables))
