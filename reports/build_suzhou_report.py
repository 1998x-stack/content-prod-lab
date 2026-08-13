# -*- coding: utf-8 -*-
"""Build 重庆市2025年政府工作报告 深度研究 DOCX, 参照北京/上海/杭州/天津系列版式。"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DARK = RGBColor(0x1F, 0x2A, 0x44)
RED = RGBColor(0xB0, 0x1E, 0x2E)
GRAY = RGBColor(0x59, 0x60, 0x69)
BLUE = RGBColor(0x1F, 0x4E, 0x79)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
HEADER_FILL = "1F2A44"
K = "微软雅黑"


def set_run(run, size=11, bold=False, color=DARK, font=K):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    run.font.name = font
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), font)
    rFonts.set(qn("w:hAnsi"), font)
    rFonts.set(qn("w:eastAsia"), font)


def para(doc, text, size=11, bold=False, color=DARK, align=None,
         space_after=6, space_before=0, font=K):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        set_run(r, size=size, bold=bold or (i % 2 == 1), color=color, font=font)
    return p


def heading1(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(16)
    pf.space_after = Pt(8)
    r = p.add_run(text)
    set_run(r, size=16, bold=True, color=RED)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "B01E2E")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def heading2(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(10)
    pf.space_after = Pt(4)
    r = p.add_run(text)
    set_run(r, size=13, bold=True, color=BLUE)
    return p


def table(doc, headers, rows, widths=None, font_size=9.5):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_run(r, size=font_size, bold=True, color=WHITE)
        tcPr = hdr[i]._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), HEADER_FILL)
        tcPr.append(shd)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(val))
            set_run(r, size=font_size, color=DARK)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    return t


def quote_box(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(8)
    pf.left_indent = Pt(12)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), "B01E2E")
    pBdr.append(left)
    pPr.append(pBdr)
    r = p.add_run(text)
    set_run(r, size=11, color=DARK, font="楷体")
    return p


def bullet(doc, text, size=10.5):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.7)
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        set_run(r, size=size, bold=(i % 2 == 1), color=DARK)
    return p


# ================================================================= 文档主体
doc = Document()
sec = doc.sections[0]
sec.page_width = Cm(21)
sec.page_height = Cm(29.7)
sec.top_margin = Cm(2.4)
sec.bottom_margin = Cm(2.2)
sec.left_margin = Cm(2.5)
sec.right_margin = Cm(2.5)

st = doc.styles["Normal"]
st.font.name = K
st.font.size = Pt(11)
st.element.rPr.rFonts.set(qn("w:eastAsia"), K)


# ---- 封面 ----
para(doc, "苏州市2025年政府工作报告", size=24, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=60, space_after=4)
para(doc, "深度研究与“容易被忽视的细节”分析", size=16, bold=True, color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
para(doc, "从“工业立市、生物医药、纳米/智能制造与外企转型”重新理解苏州", size=12, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)
para(doc, "━━━━━━━━━━━━━━━━", color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
para(doc, "研究对象：2025年苏州市政府工作报告及同期财政、国民经济和社会发展统计资料、2026年官方复盘", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
para(doc, "整理时间：2026年8月", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
para(doc, "分析性质：分析性研究 ｜ 非官方解读", size=11, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
doc.add_page_break()

# ---- 目录 ----
catalog = [
    "执行摘要",
    "一、研究口径与阅读方法",
    "二、先看苏州的特殊底盘：万亿工业强市、园区经济与外向经济",
    "三、最关键的宏观错位：GDP破2.77万亿、工业强，但消费与收入偏弱",
    "四、容易被忽视、但信号很重的细节（15条）",
    "五、把所有细节连起来：苏州正在经历的“六个换挡”",
    "六、增长暗线：工业立市、生物医药/纳米/未来产业与新质生产力",
    "七、财政暗线：收入稳、税收质量高、高科技投入大",
    "八、产业暗线：从“代工制造”到“高端智造+生物医药+未来产业”",
    "九、区域格局：县域经济、园区与长三角红利",
    "十、人口与城市：1300万人口大市、城镇化与人才",
    "十一、民营经济：外资与民企双强的外向经济",
    "十二、事后验证：用2025年实际执行结果检验报告判断",
    "十三、未来5—10年最值得观察的五条主线",
    "十四、最终结论：苏州在“工业立市+外向+创新”里的增长逻辑",
    "附录A：主要资料来源与核验路径",
    "附录B：建议建立的年度跟踪仪表盘",
]
para(doc, "目　录", size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=10, space_after=10)
for item in catalog:
    para(doc, item, size=12, space_after=4)
doc.add_page_break()

# ---- 执行摘要 ----
heading1(doc, "执行摘要")
para(doc, "**核心判断**　如果只看新闻标题，2025年苏州最显眼的是“GDP突破2.77万亿、增长5.4%”、“进出口28119亿、再创新高、增长7.4%”和“高新技术产业占规上56.2%”。但这份研究真正值得深读的，是这一座“工业立市、万亿多区、外向型经济”的城市如何在社零（+2.0%）、CPI（-0.3%）与居民收入（+4.2%）偏弱的背景下，靠“工业+外贸+生物医药/纳米/智造+外资”守住增长与基本盘。")
para(doc, "把2025年初设定的目标（GDP增长6%左右）、2025年《国民经济和社会发展统计公报》、以及2026年报告对2025年的复盘放在一起看，苏州呈现清晰暗线：**从“外企代工+传统制造”的旧依赖，向“高端智造+生物医药/纳米/未来产业+内资与外资兼修”转型**。旧引擎（一般外贸代工、传统制造、地产）在调整；新引擎（智能/集成电路/生物医药/新能源/未来产业/科创）被要求更快补位。")
para(doc, "因此，本报告不按政府工作报告逐段复述，而采用“显性表述—同期数据—制度含义—长期影响”的方式，专门提取容易被忽视、但对判断苏州未来5—10年发展模式更有价值的信号。")
para(doc, "最容易记住的一句话：**苏州是全国“工业立市+外向型经济”的天花板——它以“万亿工业+万亿外贸/外资”为底盘，用“工业园区+生物医药+纳米/智造”做增量，是中国最强“工业+开放”的万亿地级市。**观察苏州，与其看“GDP 2.7万亿”，不如看“工业/外资在新质、进出口、生物医药与科技投入”这几张名片。")

heading2(doc, "一页速览：2025年苏州经济的“表与里”")
table(doc,
    ["维度", "表面现象", "更深层信号"],
    [
        ["增长", "GDP 27695.1亿、+5.4%", "第二产业+5.6%、工业立市"],
        ["产业", "规上工业+7.6%、高新技术占56.2%", "电子/医疗/航空航天高增"],
        ["外贸", "进出口28119.3亿、+7.4%", "出口+7.6%、一带一路+17.3%"],
        ["投资", "固定资产投资-6.5%、工业+6.1%", "工业投资支撑、设备更新+17%"],
        ["财政", "一般公共预算收入2490.2亿、+1.3%", "税收占83.4%、质量高"],
        ["消费", "社零9092.2亿、+2.0%", "线上+5.4%、以旧换新强"],
        ["人口", "常住1304.77万、城镇化率82.9%", "人口+0.5%、外资/人才强"],
        ["外资", "实际使用外资58.1亿美元、+12.3%", "外资逆势流入"],
    ],
    widths=[2.2, 5.2, 8.6])
para(doc, "", size=4, space_after=2)

# ---- 一、研究口径与阅读方法 ----
heading1(doc, "一、研究口径与阅读方法")
heading2(doc, "1.1 本报告的三份核心底稿")
bullet(doc, "**2025年《政府工作报告》**：2025年2月在市十七届人大五次会议上作，给出全年预期目标（GDP增长6%左右、一般公共预算收入增长5%左右等）。进出口/社零等未全部设具体速率。")
bullet(doc, "**2025年《国民经济和社会发展统计公报》**：苏州市统计局2026年4月30日发布，给出全年实际执行数与增速，是“事后验证”的锚点。")
bullet(doc, "**2026年苏州市政府工作报告及官方复盘**：对2025执行结果的追认，交叉验证“目标—执行—再评价”三段式。")
heading2(doc, "1.2 阅读方法：显性—数据—制度—长期")
para(doc, "每一章按“**显性表述 → 同期数据 → 制度含义 → 长期影响**”四层展开。其中“制度含义”回答“政府为什么这么排优先序”，“长期影响”回答“这对未来5—10年意味着什么”。")
para(doc, "关键判别：**当报告使用的“增长词”和统计公报里的实际数不一致时，数据优先。**例如2025年GDP目标6%左右，实际5.4%；一般公共预算收入目标5%，实际1.3%。差异反映：苏州工业/外贸强，但消费/收入/财政偏弱、低于目标。")

# ---- 二、特殊底盘 ----
heading1(doc, "二、先看苏州的特殊底盘：万亿工业强市、园区经济与外向经济")
para(doc, "在所有地级市里，苏州的“底盘”独特：**全国最强工业地级市+万亿外贸/外资+园区经济（苏州工业园区）**三合一。以全国约千分之一的土地创造巨大的工业与外贸体量。")
para(doc, "这决定苏州的多重身份并存：**工业立市**（规上工业4.9万亿）、**外向型**（进出口2.8万亿、外资58亿美元）、**园区经济**（苏州工业园区、高新区）、**县域经济**（昆山/张家港/常熟/太仓等强县）、**科技创新**（生物医药/纳米/半导体）。")
heading2(doc, "2.1 工业为本")
para(doc, "苏州工业总规模全国前列，高新技术产业占规上56.2%、制造业+7.6%。集成电路、医疗设备、航空航天、机器人、生物医药是新动能。")
heading2(doc, "2.2 开放与外资")
para(doc, "实际使用外资58.1亿美元、+12.3%，进出口2.8万亿、+7.4%，是“外资+外贸”双强。数字化人民币交易量占全国近三分之二，是开放的先锋。")
heading2(doc, "2.3 县域/园区与科创")
para(doc, "昆山、张家港等县域经济发达，苏州工业园区是全国开发区样板。生物医药、纳米、量子等未来产业是苏州科创增量。")

# ---- 三、宏观错位 ----
heading1(doc, "三、最关键的宏观错位：GDP破2.77万亿、工业强，但消费与收入偏弱")
para(doc, "把2025年苏州的宏观面放进一张表，会出现令人意外的“错位”：表观增长来自工业与外贸，而消费/收入/财政偏弱。这个错位，正是读懂苏州的关键。")
heading2(doc, "3.1 “强的一面”")
table(doc,
    ["指标", "2025实绩", "信号"],
    [
        ["GDP", "27695.1亿、+5.4%", "第二产业+5.6%、工业立市"],
        ["规上工业", "+7.6%", "工业总规模4.9万亿"],
        ["高新技术占比", "56.2%", "高新技术主导"],
        ["进出口", "+7.4%（出口+7.6%）", "外贸韧性"],
        ["实际使用外资", "+12.3%", "外资逆势流入"],
        ["高技术产业投资", "+4.3%", "科创投资"],
    ],
    widths=[3.2, 5.2, 6.0])
heading2(doc, "3.2 “弱的一面”")
table(doc,
    ["指标", "2025实值", "信号"],
    [
        ["社零", "+2.0%", "消费偏弱"],
        ["固定资产投资", "-6.5%", "投资回落"],
        ["一般公共预算收入", "+1.3%", "财政低增"],
        ["居民人均可支配收入", "+4.2%（低于GDP）", "收入偏慢"],
        ["CPI", "-0.3%", "通缩压力"],
    ],
    widths=[3.2, 5.2, 6.0])
para(doc, "**错位结论**　苏州的增长“很强、但也很矛盾”。强的部分（工业/外贸/外资）与弱的部分（消费/投资/收入/财政）并存。**真正的焦点是“生产旺、需求弱”**：工业与外贸强，但内需（社零+2.0%）与收入（+4.2%）偏弱。2026年苏州“稳住工业与外贸+扩大内需+提升收入/消费+继续科创”是重点。")

# ---- 四、被忽视细节 ----
heading1(doc, "四、容易被忽视、但信号很重的细节（15条）")
details = [
    ("1", "高新技术产业占规上56.2%、电子/医疗/航空航天高增", "新质生产力主导工业。"),
    ("2", "集成电路圆片+13.6%、工业自动调节仪表+33.0%", "半导体/智能制造。"),
    ("3", "工业机器人产量+20.6%", "智能制造/机器人。"),
    ("4", "国家级专精特新“小巨人”产值+8.9%、累计848家", "专精特新集群。"),
    ("5", "进出口28119.3亿、+7.4%、创历史新高", "外贸韧性、再创新高。"),
    ("6", "对一带一路+17.3%、对东盟+25.2%", "多元化开拓。"),
    ("7", "跨境电商进出口+92.1%", "新业态爆发。"),
    ("8", "实际使用外资58.1亿美元、+12.3%", "外资逆势流入、强。"),
    ("9", "民营外贸占比提升（民企进出口10924亿、+5.7%）", "民企与外资双强。"),
    ("10", "工业投资2250亿、+6.1%、占固投39.4%", "工业投资强。"),
    ("11", "固定资产投资-6.5%、但设备更新+17%", "投资结构升级。"),
    ("12", "常住1304.77万、城镇化率82.9%，人口+0.5%", "人口稳、人才强。"),
    ("13", "居民人均可支配收入80796元、+4.2%、城乡比清晰", "收入高、城乡差距小。"),
    ("14", "R&D投入强、上海-苏州创新集群全球第六", "科创能级强。"),
    ("15", "数字人民币交易量占全国近2/3", "数字金融创新先锋。"),
]
for d in details:
    para(doc, "**%s**" % d[0], size=11)
    para(doc, d[1], size=10.5, color=GRAY)

# ---- 五、六个换挡 ----
heading1(doc, "五、把所有细节连起来：苏州正在经历的“六个换挡”")
gear = [
    ("1．动能换挡：从“代工制造”到“高端智造+生物医药/纳米+未来产业”", "高新技术占56.2%、机器人/半导体/生物医药放量。"),
    ("2．开放换挡：从“传统代工外贸”到“一带一路/数字化+外资”", "一带一路+17.3%、跨境电商+92.1%、外资+12.3%。"),
    ("3．投资换挡：从“基建/地产”到“工业/设备更新/高技术”", "工业投资+6.1%、设备更新+17%、高技术投资+4.3%。"),
    ("4．结构换挡：从“外企主导”到“民企/内资与外资兼修”", "民企外贸占提升、外资研发中心增多。"),
    ("5．人口换挡：从“人力流入”到“人才/高学历集聚”", "人才425万、高层次47万、引才28.3万毕业生。"),
    ("6．社会换挡：从“追求规模”到“收入/消费/科创升级”", "消费偏弱、R&D、数字金融。"),
]
for g in gear:
    para(doc, "**%s**　%s" % g, space_after=6)

# ---- 六、增长暗线 ----
heading1(doc, "六、增长暗线：工业立市、生物医药/纳米/未来产业与新质生产力")
heading2(doc, "6.1 工业立市+新质")
para(doc, "苏州是工业强市，高新技术产业占规上56.2%；集成电路圆片、半导体、机器人、生物医药（医疗器械）等新动能放量。国家级专精特新“小巨人”848家。")
heading2(doc, "6.2 科创与未来产业")
para(doc, "上海—苏州创新集群全球第六、财政科技投入占10.8%、有效发明专利18.4万件。量子、生物医药、纳米等未来产业，是苏州从“制造”到“智造/科创”的增量。")
para(doc, "**这条暗线意味着**：苏州的增长叙事正从“代工外贸”转向“工业+科创+未来产业+外资与民企”。看苏州，盯住“高新技术占比、集成电路/生物制药、外贸/外资、科技投入”这几组数。")

# ---- 七、财政暗线 ----
heading1(doc, "七、财政暗线：收入稳、税收质量高、高科技投入大")
para(doc, "2025年苏州一般公共预算收入2490.2亿、+1.3%，税收2076亿、+3.6%，税收占83.4%——收入质量极高。财政科技投入274.5亿、占10.8%，支撑科创。")
para(doc, "**制度含义**　苏州财政“收入稳、税收质量高”，靠“工业+科创”真实税源。重点是支持科创转化、摆脱对地产依赖、平衡投资与民生。")

# ---- 八、产业暗线 ----
heading1(doc, "八、产业暗线：从“代工制造”到“高端智造+生物医药+未来产业”")
heading2(doc, "8.1 苏州产业的“表”")
table(doc,
    ["指标", "2025增速/信号", "解读"],
    [
        ["规上工业", "+7.6%", "工业强市"],
        ["高新技术产业", "占56.2%", "高新技术主导"],
        ["集成电路圆片", "+13.6%", "半导体/芯片"],
        ["生物医药/医疗器械", "医疗设备+9.6%", "生物医药赛道"],
        ["航空航天设备", "+14.0%", "航空航天"],
        ["工业机器人", "+20.6%", "智能制造"],
        ["进出口", "+7.4%", "外贸韧性"],
    ],
    widths=[4.6, 3.6, 5.0])
heading2(doc, "8.2 从“代工”到“高端智造/科创”")
para(doc, "苏州从外企代工/传统制造，升级为“高端智造+生物医药+纳米+未来产业”。集成电路、生物医药、机器人、航空航天放量，是苏州“工业立市+科创”的新底盘。")

# ---- 九、区域 ----
heading1(doc, "九、区域格局：县域经济、园区与长三角红利")
para(doc, "苏州“一核四园”+县域经济发达：苏州工业园区、高新区是开放样板，昆山、张家港、常熟、太仓是全国强县。苏州深度融入长三角一体化，借力上海。")
para(doc, "在长三角辐射下，苏州承接制造业、外资与科创外溢，同时以“工业+外贸+科创”支撑全省。县域+园区经济，是苏州独特优势。")

# ---- 十、人口与城市 ----
heading1(doc, "十、人口与城市：1300万人口大市、城镇化与人才")
heading2(doc, "10.1 人口总量")
para(doc, "2025年末苏州常住1304.77万人、城镇化率82.9%、人口+0.5%。苏州因产业吸引力，人口稳、人才强。")
heading2(doc, "10.2 人才与就业")
para(doc, "各类人才总量425万、高层次人才47万；引留高校毕业生28.3万、发布青年人才岗位31.8万。苏州在“就业+人才+收入”上保持吸引力。")

# ---- 十一、民营经济 ----
heading1(doc, "十一、民营经济：外资与民企双强的外向经济")
heading2(doc, "11.1 民企/外资并存")
para(doc, "苏州民营企业进出口10924亿、+5.7%，同时实际使用外资58亿美元、+12.3%。民企与外资“双强”，是苏州外向经济的底色。")
heading2(doc, "11.2 政策与服务")
para(doc, "苏州优化营商环境（“换位跑一次”、AI政策明白卡），连续获“最佳口碑城市”。民企、专精特新、科创企业，是苏州未来10年底盘。")

# ---- 十二、事后验证 ----
heading1(doc, "十二、事后验证：用2025年实际执行结果检验报告判断")
heading2(doc, "12.1 年初目标与年末执行对比")
table(doc,
    ["指标", "2025年初目标", "2025实际", "偏差"],
    [
        ["GDP增速", "6%左右", "+5.4%", "略低"],
        ["一般公共预算收入", "5%左右", "+1.3%", "明显低"],
        ["进出口", "稳量提质", "+7.4%", "达标"],
        ["规上工业", "力争", "+7.6%", "达预期"],
        ["社零", "未设", "+2.0%", "偏弱"],
        ["居民收入", "与增长同步", "+4.2%", "偏低"],
    ],
    widths=[3.0, 3.2, 3.0, 4.2])
heading2(doc, "12.2 判断验证")
para(doc, "三条核心判断得到支持：**（1）工业/外贸/科创强（规上+7.6%、进出口+7.4%、高新技术56%），验证；（2）消费/收入/财政偏弱（社零+2.0%、收入+4.2%、财政+1.3%），验证；（3）外资/民企/未来产业强，验证。**")
para(doc, "核心观察：苏州靠“工业+外贸+科创+外资”守住2.77万亿，但消费、收入、财政增速低于目标。2026年“稳工业/外贸+扩内需+提升收入/资产”是重点。")

# ---- 十三、五条主线 ----
heading1(doc, "十三、未来5—10年最值得观察的五条主线")
lines5 = [
    ("① 高端智造/集成电路/生物医药/未来产业", "能否从“制造”升级为“技术/智造先行区”。"),
    ("② 外贸/外资/一带一路", "进出口、一带一路、外资能否持续扩张。"),
    ("③ 科创与人才", "生物医药、纳米、量子、人才，能否孵化更多独角兽。"),
    ("④ 内需/收入/消费升级", "社零弱，能否在消费与收入提振中补齐。"),
    ("⑤ 民营/知识产权/绿色", "民营、专精特新、绿色低碳、数字金融。"),
]
for l in lines5:
    para(doc, "**%s**　%s" % l, space_after=6)

# ---- 十四、结论 ----
heading1(doc, "十四、最终结论：苏州在“工业立市+外向+创新”里的增长逻辑")
para(doc, "苏州的2025年，本质上是“**工业/外贸/科创为核心，而消费/收入/财政偏弱**”的答卷：GDP 2.77万亿、工业+7.6%、进出口新高，代价是内需与财政收入增速低于目标。")
para(doc, "只要工业、外贸、外资、科创能接住，苏州就仍是“万亿+工业+开放”的头号城市；如果内需、收入与税源持续偏弱，苏州需在“生产旺、需求弱”间平衡。")
para(doc, "最稳妥的观察信号：**一盯高新技术/集成电路/生物医药（动能）、二盯外贸/外资（外向）、三盯科创/人才（创新）、四盯消费/收入（内需）、五盯民营/数字（底座）。**苏州，是中国最强“工业+开放”的万亿地级市。")

# ---- 附录A ----
heading1(doc, "附录A：主要资料来源与核验路径")
bullet(doc, "苏州市2025年《政府工作报告》（2025年2月）——目标来源。")
bullet(doc, "《2025年苏州市国民经济和社会发展统计公报》（苏州市统计局，2026-04-30）——GDP、工业、外贸、人口实值。")
bullet(doc, "苏州市统计/运行分析、工业园区专报、生物医药专篇——区域与科创。")
bullet(doc, "2026年苏州市政府工作报告——2025执行复盘。")
bullet(doc, "苏州海关、市商务局、财政局——外贸/财政。")
heading2(doc, "核验说明")
para(doc, "本报告所有增速、总量、占比均以统计公报/官方发布为基准。涉“生物医药具体产值”“常住精确值”等以官方口径为准。")

# ---- 附录B ----
heading1(doc, "附录B：建议建立的年度跟踪仪表盘")
para(doc, "建议把下面10个“测脉搏”指标做成年度跟踪表：")
dash = [
    ["#", "指标", "2025基值", "为什么重要"],
    ["1", "GDP增速", "+5.4%", "总量与方向"],
    ["2", "规上工业增速", "+7.6%", "工业底盘"],
    ["3", "高新技术产业占比", "56.2%", "新质生产力"],
    ["4", "进出口/出口增速", "+7.4%/+7.6%", "外贸韧性"],
    ["5", "实际使用外资", "+12.3%", "外资信心"],
    ["6", "固定资产投资/工业投资", "-6.5%/+6.1%", "投资结构"],
    ["7", "社零增速", "+2.0%", "内需消费"],
    ["8", "常住人口/城镇化率", "1304.77万/82.9%", "人口与城市"],
    ["9", "一般公共预算收入/税收占比", "+1.3%/83.4%", "财政质量"],
    ["10", "居民人均可支配收入增速", "+4.2%", "民生获得感"],
]
table(doc, dash[0], dash[1:], widths=[0.8, 4.5, 3.2, 4.6])
para(doc, "把这10个指标连起来看，任何一个“新引擎（2/3/5）向上、旧引擎（7）修复”，都说明苏州在真正换挡；反之则是旧路径的反复。")

# ========================================= 保存
out = "/Users/x/Desktop/content-prod-lab/reports/苏州市_2025年政府工作报告_深度研究_2026-08-13.docx"
doc.save(out)
print("SAVED:", out)
print("PARAGRAPHS:", len(doc.paragraphs))
print("TABLES:", len(doc.tables))
