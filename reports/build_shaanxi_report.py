# -*- coding: utf-8 -*-
"""Build 陕西省2025年政府工作报告 深度研究 DOCX, 参照省市系列版式。"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DARK = RGBColor(0x1F, 0x2A, 0x44)
RED = RGBColor(0xB0, 0x1E, 0x2E)
GRAY = RGBColor(0x59, 0x60, 0x69)
BLUE = RGBColor(0x1F, 0x4E, 0x79)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
HEADER_FILL = "1F2A44"
K = "微软雅黑"


def set_run(run, size=11, bold=False, color=DARK, font=K):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    run.font.name = font
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), font)
    rFonts.set(qn("w:hAnsi"), font)
    rFonts.set(qn("w:eastAsia"), font)


def para(doc, text, size=11, bold=False, color=DARK, align=None,
         space_after=6, space_before=0, font=K):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        set_run(r, size=size, bold=bold or (i % 2 == 1), color=color, font=font)
    return p


def heading1(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(16)
    pf.space_after = Pt(8)
    r = p.add_run(text)
    set_run(r, size=16, bold=True, color=RED)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "B01E2E")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def heading2(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(10)
    pf.space_after = Pt(4)
    r = p.add_run(text)
    set_run(r, size=13, bold=True, color=BLUE)
    return p


def table(doc, headers, rows, widths=None, font_size=9.5):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_run(r, size=font_size, bold=True, color=WHITE)
        tcPr = hdr[i]._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), HEADER_FILL)
        tcPr.append(shd)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(val))
            set_run(r, size=font_size, color=DARK)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    return t


def quote_box(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(8)
    pf.left_indent = Pt(12)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), "B01E2E")
    pBdr.append(left)
    pPr.append(pBdr)
    r = p.add_run(text)
    set_run(r, size=11, color=DARK, font="楷体")
    return p


def bullet(doc, text, size=10.5):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.7)
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        set_run(r, size=size, bold=(i % 2 == 1), color=DARK)
    return p


# ================================================================= 文档主体
doc = Document()
sec = doc.sections[0]
sec.page_width = Cm(21)
sec.page_height = Cm(29.7)
sec.top_margin = Cm(2.4)
sec.bottom_margin = Cm(2.2)
sec.left_margin = Cm(2.5)
sec.right_margin = Cm(2.5)

st = doc.styles["Normal"]
st.font.name = K
st.font.size = Pt(11)
st.element.rPr.rFonts.set(qn("w:eastAsia"), K)



# ---- 封面 ----
para(doc, "陕西省2025年政府工作报告", size=24, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=60, space_after=4)
para(doc, "深度研究与“容易被忽视的细节”分析", size=16, bold=True, color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
para(doc, "从“硬科技、能源化工、陕北、关中平原与生态文明”重新理解陕西", size=12, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)
para(doc, "━━━━━━━━━━━━━━━━", color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
para(doc, "研究对象：2025年陕西省政府工作报告及同期财政、国民经济和社会发展统计资料、2026年官方复盘", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
para(doc, "标定日期：2026-08-13", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
doc.add_page_break()

# ---- 目录 ----
catalog = [
    "执行摘要",
    "一、研究口径与阅读方法",
    "二、先看陕西的特殊底盘：硬科技、能源化工、陕北/关中/陕南、生态文明",
    "三、最关键的宏观错位：GDP破3.65万亿、出口强，但财政/投资/工业利润弱",
    "四、容易被忽视、但信号很重的细节（15条）",
    "五、2025年GDP目标 vs 实际：对照表",
    "六、2025年增长是谁撑起来的？（结构归因）",
    "七、预算与财政的“含金量”",
    "八、民生底账：人口、收入与城乡",
    "九、城镇与农村：格局与均衡",
    "十、人口流入与流出",
    "十一、物价与货币环境",
    "十二、区域一体化：陕西在“关中平原城市群+一带一路+西部大开发”里的位置",
    "十三、未来5—10年最值得观察的五条主线",
    "十四、最终结论：陕西在“硬科技+能源化工+汽车/电子+生态文明”里的增长逻辑",
    "附录A：主要资料来源与核验路径",
    "附录B：建议建立的年度跟踪仪表盘",
]
para(doc, "目　录", size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=10, space_after=10)
for item in catalog:
    para(doc, item, size=12, space_after=4)
doc.add_page_break()

# ---- 执行摘要 ----
heading1(doc, "执行摘要")
para(doc, "**核心判断**　如果只看新闻标题，2025年陕西最显眼的是“GDP破3.65万亿、增长5.1%”、“规上工业+7.3%”、“进出口+18.5%（出口+22.0%）”和“新能源汽车产量109.47万辆”。但这份研究真正值得深读的，是这座“硬科技+能源化工+汽车+秦岭生态”的西部强省，如何在工业利润（-20.5%）、财政（-3.1%）与人口负增长（-2.54‰）的背景下，靠“出口+制造+能源”撑起5.1%的增长。")
para(doc, "把2025年初设定的目标（GDP增长5%左右）、2025年《国民经济和社会发展统计公报》、以及2026年报告对2025年的复盘放在一起看，陕西呈现清晰暗线：**从“能源/重工业”的旧底盘，向“硬科技/新能源+汽车+出口/一带一路”转型**。旧引擎（粗钢/水泥/传统能源）在调整；新引擎（汽车/电气机械/光伏/半导体/硬科技）被要求更快补位。")
para(doc, "因此，本报告不按政府工作报告逐段复述，而采用“显性表述—同期数据—制度含义—长期影响”的方式，专门提取容易被忽视、但对判断陕西未来5—10年发展模式更有价值的信号。")
para(doc, "最容易记住的一句话：**陕西是“硬科技+能源化工+汽车/装备+秦岭生态”的西部样本，靠“出口+能源+汽车/硬科技”撑起增长。**观察陕西，与其看“GDP 3.65万亿”，不如看“硬科技、新能源汽车、出口+22%、能源保供、关中城市群”这几张名片。")
heading2(doc, "一页速览：2025年陕西经济的“表与里”")
table(doc,
    ["维度", "表面现象", "更深层信号"],
    [
        ["增长", "GDP 36551.10亿、+5.1%", "一产7.5%、二产39.7%、三产52.8%"],
        ["产业", "规上工业+7.3%", "采矿业+9.5%、汽车+20.2%"],
        ["外贸", "进出口5379.75亿、+18.5%", "出口+22.0%、对中亚增"],
        ["投资", "固定资产投资-2.8%", "民间+4.6%、地产-7.9%"],
        ["财政", "一般公共预算收入3289.44亿、-3.1%", "税收-7.4%"],
        ["消费", "社零11587.99亿、+6.0%", "通讯+40.9%、家电+27.1%"],
        ["人口", "常住3936万、-17万、城镇化67.17%", "自然增-2.54‰、老龄化23.22%"],
        ["开放", "一带一路/中欧班列/西部陆海新通道", "硬科技/汽车/出口"],
    ],
    widths=[2.2, 5.4, 8.4])
para(doc, "", size=4, space_after=2)

# ---- 一、研究口径与阅读方法 ----
heading1(doc, "一、研究口径与阅读方法")
heading2(doc, "1.1 本报告的三份核心底稿")
bullet(doc, "**2025年《政府工作报告》**：2025年1月在省十四届人大三次会议上作，给出全年预期目标（GDP增长5%左右、进出口总值增长5%）。")
bullet(doc, "**2025年《国民经济和社会发展统计公报》**：陕西省统计局2026年4月15日发布，给出全年实际执行数与增速，是“事后验证”的锚点。")
bullet(doc, "**2026年陕西省政府工作报告及官方复盘**：对2025执行结果的追认，交叉验证“目标—执行—再评价”三段式。")
heading2(doc, "1.2 阅读方法：显性—数据—制度—长期")
para(doc, "每一章按“**显性表述 → 同期数据 → 制度含义 → 长期影响**”四层展开。其中“制度含义”回答“政府为什么这么排优先序”，“长期影响”回答“这对未来5—10年意味着什么”。")
para(doc, "关键判别：**当报告使用的“增长词”和统计公报里的实际数不一致时，数据优先。**例如2025年GDP目标5%左右、实际5.1%达标；进出口目标5%、实际+18.5%显著超预期。差异反映：陕西“出口/汽车/硬科技强，财政/工业利润/能源价格弱”。")
heading2(doc, "1.3 指标取舍与单位说明")
para(doc, "本报告统一以“2025年统计公报+2026年报告复盘”为口径，增速均为可比价，绝对数为人民币。GDP 36551.10亿元为全省初步核算数。")

# ---- 二、底盘 ----
heading1(doc, "二、先看陕西的特殊底盘：硬科技、能源化工、陕北/关中/陕南、生态文明")
para(doc, "陕西的地盘，取决于它作为“**硬科技+能源化工+汽车/装备+秦岭生态**”的特殊定位，是拥有西安硬科技之都（秦创原）与陕北能源化工基地的西部强省。")
bullet(doc, "**硬科技**：西安“硬科技之都”，秦创原产业创新聚集区21个、国家科学技术奖47项（全国第2）、技术合同成交额5260.36亿元；半导体/光子/航空航天/北斗等“卡脖子”突破。")
bullet(doc, "**能源化工**：原煤80461.72万吨（+2.9%）、天然原油2430.70万吨、天然气378.95亿m³（+5.2%）、原油加工1985.55万吨；陕北能源化工基地（煤/油/气/电）保供。")
bullet(doc, "**汽车/装备**：汽车产量172.50万辆（-1.6%）、新能源汽车109.47万辆（-8.6%）；电气机械+26.3%、汽车+20.2%、发动机+63.0%、金属切削机床+18.1%、变压器+17.1%。")
bullet(doc, "**陕北/关中/陕南+生态文明**：陕北能源、关中平原城市群（西安都市圈）、陕南绿色；秦岭生态屏障、黄河流域生态保护（生态文明）。")
para(doc, "**制度含义**：陕西把“硬科技（秦创原）、能源安全、汽车/装备、秦岭生态”当核心资产，面向“一带一路+西部大开发+关中平原城市群”。民营经济占GDP 50.3%、战略性新兴产业占11.4%。")

# ---- 三、宏观错位 ----
heading1(doc, "三、最关键的宏观错位：GDP破3.65万亿、出口强，但财政/投资/工业利润弱")
para(doc, "2025年陕西GDP 36551.10亿元、+5.1%（一产+3.8%、二产+5.6%、三产+4.9%）。表面看“稳中有进”，但拆开看是“**出口/制造强、财政/工业利润/投资弱**”的错位：")
para(doc, "**强的部分**：规上工业+7.3%（采矿业+9.5%、煤炭+10.8%、电气机械+26.3%、汽车+20.2%）；进出口+18.5%（出口+22.0%、外资企业+27.6%）；社零+6.0%（通讯+40.9%）；民间投资+4.6%。")
para(doc, "**弱的部分**：规上工业利润-20.5%、工业营收-4.2%；一般公共预算收入3289.44亿、-3.1%（税收-7.4%）；固定资产投资-2.8%（基建-19.0%、地产-7.9%）；人口自然增长率-2.54‰。")
para(doc, "**核心错位一句话**：陕西“出口/汽车/硬科技/能源开采强，但工业利润/财政/投资/人口弱”。2026年“稳出口/制造、补投资/财政、强能源化工+硬科技”是重点。")
table(doc,
    ["强引擎", "数据", "弱项", "数据"],
    [
        ["进出口总值", "+18.5%", "一般公共预算收入", "-3.1%"],
        ["规上工业", "+7.3%", "规上工业利润", "-20.5%"],
        ["社零总额", "+6.0%", "固定资产投资", "-2.8%"],
        ["汽车制造业", "+20.2%", "房地产开发投资", "-7.9%"],
        ["电气机械制造", "+26.3%", "人口自然增长率", "-2.54‰"],
    ],
    widths=[3.6, 3.0, 3.6, 3.0])
para(doc, "**错位结论**　陕西的增长“强出口/制造/能源，弱财政/工业利润/投资/人口”。2026年“稳出口/硬科技、补财政/投资、强能源安全”是重点。")

# ---- 四、被忽视细节 ----
heading1(doc, "四、容易被忽视、但信号很重的细节（15条）")
details = [
    ("1", "GDP 36551.10亿、+5.1%", "总量/稳增。"),
    ("2", "规上工业+7.3%、煤炭开采+10.8%", "能源支撑。"),
    ("3", "电气机械+26.3%、汽车+20.2%", "高端制造。"),
    ("4", "新能源汽车109.47万辆（-8.6%）", "新能源车回调。"),
    ("5", "发动机+63.0%、机床+18.1%、变压器+17.1%", "装备/器械强。"),
    ("6", "进出口5379.75亿、+18.5%、出口+22.0%", "外贸强、一带一路+10.5%。"),
    ("7", "硬科技/秦创原/技术合同5260亿", "硬科技。"),
    ("8", "民营经济占GDP 50.3%、战新占11.4%", "民营/战新。"),
    ("9", "社零+6.0%、通讯/家电高增", "消费升级。"),
    ("10", "工业项目华电等投资+9.4%", "工业投资。"),
    ("11", "秦岭生态/流域生态修复", "生态文明。"),
    ("12", "常住3936万、-17万、城镇化67.17%、自然增-2.54‰", "人口/城镇化。"),
    ("13", "居民收入35790元、+5.6%、农村+6.0%", "收入稳、农村快。"),
    ("14", "快递业务量+38.6%、航空旅客+4.4%", "物流/交通。"),
    ("15", "CPI+0.2%、PPI-5.8%", "物价/PPI大降。"),
]
for d in details:
    para(doc, "**%s**" % d[0], size=11)
    para(doc, d[1], size=10.5, color=GRAY)

    para(doc, "**备注**　这条“错位”在陕西尤其鲜明：增长靠出口/汽车/能源/硬科技，但财政/工业利润/投资弱。2026年若能源价格/工业利润修复、投资回暖，增长可能从“出口单极”走向“出口+内需/投资”多极。这条细节，正是陕西2026年最难也最值得盯的“变量”。", size=10.5)

# ---- 五、2025年目标 vs 实际 对照 ----
heading1(doc, "五、2025年GDP目标 vs 实际：对照表")
t5 = [
    ["指标", "2025目标", "2025实际", "达标判定"],
    ["GDP增速", "5%左右", "+5.1%", "达标"],
    ["进出口总值增速", "5%", "+18.5%", "大幅超预期"],
    ["一般公共预算收入增速", "3%", "-3.1%", "未达标"],
    ["城镇/农村居民收入增速", "5%左右/6%左右", "+4.8%/+6.0%", "城镇略低"],
    ["城镇调查失业率", "5.5%左右", "约5.3%", "达标"],
    ["粮食产量", "1300万吨以上", "1352万吨", "达标/超"],
    ["居民消费价格(CPI)", "涨幅2%左右", "+0.2%", "远低于"],
]
table(doc, t5[0], t5[1:], widths=[3.4, 3.0, 3.0, 3.6])
para(doc, "**对照结论**　目标“求稳”，实际“分化”：GDP 5.1%、进出口+18.5%（大幅超预期）、粮食超额；但财政收入-3.1%（未达3%目标）、CPI仅+0.2%。出口/制造/能源接住，财政/工业利润/投资是短板。")

# ---- 六、增速分项支撑 ----
heading1(doc, "六、2025年增长是谁撑起来的？（结构归因）")
para(doc, "GDP+5.1%背后，是“**出口/制造/能源强、财政/工业利润/投资弱**”的结构。把增长贡献拆开看：")
g6 = [
    ["引擎", "贡献/状态", "说明"],
    ["规上工业", "+7.3%", "煤炭+10.8%、电气机械+26.3%"],
    ["外贸出口", "+22.0%", "进出口+18.5%、外资企业+27.6%"],
    ["汽车/装备制造", "+20.2%/+26.3%", "新能源汽车、发动机+63%"],
    ["能源开采", "煤炭+2.9%/气+5.2%", "能源保供/化工"],
    ["硬科技/秦创原", "技术合同5260亿", "硬科技/半导体"],
    ["社零消费", "+6.0%", "通讯+40.9%、家电+27.1%"],
    ["规上工业利润", "-20.5%", "工业效益/价格弱"],
    ["一般公共预算收入", "-3.1%", "财政压力"],
    ["固定资产投资", "-2.8%", "基建/地产拖累"],
    ["人口自然增长率", "-2.54‰", "人口/老龄化约束"],
]
table(doc, g6[0], g6[1:], widths=[3.6, 3.6, 6.0])
para(doc, "**一句话**　增长靠“外贸+工业（汽车/电气/能源）+硬科技”，但工业利润/财政/投资是拖累。2026年考验陕西“能不能让财政/工业利润/投资跟上来”。")

# ---- 七、预算与财政 ----
heading1(doc, "七、预算与财政的“含金量”")
para(doc, "2025年陕西一般公共预算收入**3289.44亿元、-3.1%**，其中税收收入**-7.4%**。财政收入承压（能源价格/企业利润回落）。")
bullet(doc, "财政收入-3.1%、税收-7.4%，承压。")
bullet(doc, "财政“过紧日子+保民生（支出-2.8%）”，支持硬科技/能源。")
bullet(doc, "规上工业利润-20.5%、营收-4.2%，影响税收。")

# ---- 八、民生底账 ----
heading1(doc, "八、民生底账：人口、收入与城乡")
para(doc, "2025年陕西常住人口**3936万人、-17万**，城镇化率**67.17%**。人口自然增长率**-2.54‰**，60岁及以上占23.22%。")
para(doc, "居民人均可支配收入**35790元、+5.6%**，农村（+6.0%）快于城镇（+4.8%），城乡比2.54。")
g8 = [
    ["指标", "数值", "信号"],
    ["常住人口", "3936万/-17万", "人口净流出"],
    ["人口自然增长率", "-2.54‰", "自然负增"],
    ["城镇化率", "67.17%", "较高城镇化"],
    ["居民人均可支配收入", "35790元/+5.6%", "收入稳、农村快"],
]
table(doc, g8[0], g8[1:], widths=[4.2, 4.0, 4.8])
para(doc, "**民生观察**：收入增速+5.6%不错、农村快于城镇，但人口净流出、老龄化（23.22%）是重要约束。")

# ---- 九、城镇与农村 ----
heading1(doc, "九、城镇与农村：格局与均衡")
para(doc, "陕西城镇化率67.17%、城乡收入比2.54（缩小0.03），城乡相对协调。")
bullet(doc, "农村居民收入+6.0%、快于城镇+4.8%。")
bullet(doc, "社零城镇+6.0%、乡村+5.5%。")
bullet(doc, "人口向西安都市圈/关中城市群集中。")

# ---- 十、人口流入与流出 ----
heading1(doc, "十、人口流入与流出")
para(doc, "陕西2025年常住3936万人、-17万，人口净流出、自然负增（-2.54‰）。西安/关中吸引人口/人才。")
para(doc, "未来看点：硬科技/秦创原+汽车/能源能否留住人口/吸引回流；若“硬科技+教育科研高地”成势，陕西有望减缓人口流出。")

# ---- 十一、物价与货币 ----
heading1(doc, "十一、物价与货币环境")
para(doc, "2025年陕西CPI**+0.2%**、PPI**-5.8%**，物价温和、工业品出厂价深度走弱。")
para(doc, "物价偏弱反映“大宗/能源价格回落”、采矿业PPI-12.6%，与全国低通胀一致。2026年“稳价格、扩内需”是主线。")

# ---- 十二、区域一体化 ----
heading1(doc, "十二、区域一体化：陕西在“关中平原城市群+一带一路+西部大开发”里的位置")
para(doc, "陕西的核心战略坐标是“**关中平原城市群（西安都市圈）+一带一路枢纽+西部大开发**”，是西北开放与硬科技的高地。")
bullet(doc, "关中城市群：西安都市圈、硬科技/大学高地。")
bullet(doc, "一带一路：中欧班列长安号、西部陆海新通道。")
bullet(doc, "西部大开发：陕北能源、陕南绿色、关中制造。")
bullet(doc, "生态文明：秦岭生态屏障、黄河流域保护。")
para(doc, "若“硬科技+能源+一带一路+城市群”成势，陕西将作为西部开放与科技强省的样本。")

# ========= 十三、五条主线 =========
heading1(doc, "十三、未来5—10年最值得观察的五条主线")
lines5 = [
    ("① 硬科技/秦创原/半导体", "硬科技、半导体/光子/航天能否成极。"),
    ("② 汽车/新能源汽车/装备", "汽车、新能源、装备制造能否壮大。"),
    ("③ 能源化工/陕北", "能源保供、化工能否走向高端。"),
    ("④ 财政/工业利润/投资", "财政承压，能否用硬科技/出口补。"),
    ("⑤ 城市群/一带一路/生态文明", "关中城市群、对外开放、秦岭生态。"),
]
for l in lines5:
    para(doc, "**%s**　%s" % l, space_after=6)

# ---- 十四、结论 ----
heading1(doc, "十四、最终结论：陕西在“硬科技+能源化工+汽车/电子+生态文明”里的增长逻辑")
para(doc, "陕西的2025年，本质上是“**出口/汽车/能源/硬科技为核心，而财政/工业利润/投资偏弱**”的答卷：GDP36551.10亿、+5.1%，规上工业+7.3%、进出口+18.5%、汽车+20.2%、社零+6.0%，但财政-3.1%、工业利润-20.5%、固投-2.8%、自然增-2.54‰。")
para(doc, "只要硬科技、能源化工、汽车/装备、出口持续，陕西就站在“西部强省/硬科技之都”的位；如果工业利润/财政/人口持续偏弱，陕西需承受“强出口/科技、弱财政/人口”的挑战。")
para(doc, "最稳观察信号：**一盯硬科技/秦创原（科技）、二盯能源化工（底座）、三盯汽车/出口（动能）、四盯财政/投资（约束）、五盯人口/生态（长期）。**陕西，是“硬科技+能源+西部开放”的新样本。")

# ---- 附录A ----
heading1(doc, "附录A：主要资料来源与核验路径")
bullet(doc, "陕西省2025年《政府工作报告》——目标来源。")
bullet(doc, "《2025年陕西省国民经济和社会发展统计公报》（省统计局）——GDP、工业、外贸、人口实值。")
bullet(doc, "2026年陕西省政府工作报告——2025执行复盘+硬科技/能源/汽车。")
bullet(doc, "西安海关、省财政厅——外贸/财政。")
heading2(doc, "核验说明")
para(doc, "本报告所有增速、总量、占比均以统计公报/官方口径为基准，涉“硬科技/能源化工/关中/生态文明”等以官方口径为准。")

# ---- 附录B ----
heading1(doc, "附录B：建议建立的年度跟踪仪表盘")
para(doc, "建议把下面10个“测脉搏”指标做成年度跟踪表：")
dash = [
    ["#", "指标", "2025基值", "为什么重要"],
    ["1", "GDP增速", "+5.1%", "总量与方向"],
    ["2", "规上工业增速", "+7.3%", "制造底盘"],
    ["3", "新能源汽车产量", "109.47万辆", "工业转型"],
    ["4", "进出口/出口增速", "+18.5%/+22.0%", "外贸韧性"],
    ["5", "固定资产投资/地产", "-2.8%/-7.9%", "投资结构"],
    ["6", "社零增速", "+6.0%", "内需消费"],
    ["7", "常住人口/自然增长率", "3936万/-2.54‰", "人口与城市"],
    ["8", "一般公共预算收入增速", "-3.1%", "财政质量"],
    ["9", "硬科技/技术合同", "5260亿", "硬科技"],
    ["10", "CPI/PPI", "+0.2%/-5.8%", "物价与需求"],
]
table(doc, dash[0], dash[1:], widths=[0.8, 4.6, 3.4, 4.4])
para(doc, "把这10个指标连起来看，任何一个硬科技（9）、汽车/出口（3/4）、能源向上、财政/人口（8/7）修复，都说明陕西在真正换挡。")

# ===================================== 保存
out = "/Users/x/Desktop/content-prod-lab/reports/陕西省_2025年政府工作报告_深度研究_2026-08-13.docx"
doc.save(out)
print("SAVED:", out)
print("PARAGRAPHS:", len(doc.paragraphs))
print("TABLES:", len(doc.tables))
