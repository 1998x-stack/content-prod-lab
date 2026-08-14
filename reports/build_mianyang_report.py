# -*- coding: utf-8 -*-
"""Build 济宁市2025年政府工作报告 深度研究 DOCX, 参照地级市系列版式。"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DARK = RGBColor(0x1F, 0x2A, 0x44)
RED = RGBColor(0xB0, 0x1E, 0x2E)
GRAY = RGBColor(0x59, 0x60, 0x69)
BLUE = RGBColor(0x1F, 0x4E, 0x79)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
HEADER_FILL = "1F2A44"
K = "微软雅黑"


def set_run(run, size=11, bold=False, color=DARK, font=K):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    run.font.name = font
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), font)
    rFonts.set(qn("w:hAnsi"), font)
    rFonts.set(qn("w:eastAsia"), font)


def para(doc, text, size=11, bold=False, color=DARK, align=None,
         space_after=6, space_before=0, font=K):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        set_run(r, size=size, bold=bold or (i % 2 == 1), color=color, font=font)
    return p


def heading1(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(16)
    pf.space_after = Pt(8)
    r = p.add_run(text)
    set_run(r, size=16, bold=True, color=RED)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "B01E2E")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def heading2(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(10)
    pf.space_after = Pt(4)
    r = p.add_run(text)
    set_run(r, size=13, bold=True, color=BLUE)
    return p


def table(doc, headers, rows, widths=None, font_size=9.5):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_run(r, size=font_size, bold=True, color=WHITE)
        tcPr = hdr[i]._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), HEADER_FILL)
        tcPr.append(shd)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(val))
            set_run(r, size=font_size, color=DARK)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    return t


def quote_box(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(8)
    pf.left_indent = Pt(12)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), "B01E2E")
    pBdr.append(left)
    pPr.append(pBdr)
    r = p.add_run(text)
    set_run(r, size=11, color=DARK, font="楷体")
    return p


def bullet(doc, text, size=10.5):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.7)
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        set_run(r, size=size, bold=(i % 2 == 1), color=DARK)
    return p


# ================================================================= 文档主体
doc = Document()
sec = doc.sections[0]
sec.page_width = Cm(21)
sec.page_height = Cm(29.7)
sec.top_margin = Cm(2.4)
sec.bottom_margin = Cm(2.2)
sec.left_margin = Cm(2.5)
sec.right_margin = Cm(2.5)

st = doc.styles["Normal"]
st.font.name = K
st.font.size = Pt(11)
st.element.rPr.rFonts.set(qn("w:eastAsia"), K)


# ---- 封面 ----
para(doc, "绵阳市2025年政府工作报告", size=24, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=60, space_after=4)
para(doc, "深度研究与\u201c容易被忽视的细节\u201d分析", size=16, bold=True, color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
para(doc, "从\u201c中国科技城、电子信息/军工科研、五市战略、省域经济副中心\u201d重新理解绵阳", size=12, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)
para(doc, "\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501", color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
para(doc, "研究对象：2025年绵阳市政府工作报告及同期财政、国民经济和社会发展统计资料、2026年官方复盘", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
para(doc, "标定日期：2026-08-14", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
doc.add_page_break()

# ---- 目录 ----
catalog = [
    "执行摘要",
    "一、研究口径与阅读方法",
    "二、先看绵阳的特殊底盘：中国科技城、电子信息军工、省域经济副中心",
    "三、最关键的宏观错位：GDP+6.5%达标、社零破2000亿强，但规上+6.5%低8%、外贸+0.7%",
    "四、容易被忽视、但信号很重的细节（15条）",
    "五、2025年GDP目标 vs 实际：对照表",
    "六、2025年增长是谁撑起来的？（结构归因）",
    "七、预算与财政的\u201c含金量\u201d",
    "八、民生底账：人口、收入与城乡",
    "九、城镇与农村：格局与均衡",
    "十、人口流入与流出",
    "十一、物价与货币环境",
    "十二、区域一体化：绵阳在成都都市圈、成渝双城经济圈、中国科技城\u201c三圈\u201d里的位置",
    "十三、未来5—10年最值得观察的五条主线",
    "十四、最终结论：绵阳在\u201c科技城+电子信息+省域副中心\u201d里的增长逻辑",
    "附录A：主要资料来源与核验路径",
    "附录B：建议建立的年度跟踪仪表盘",
]
para(doc, "目　录", size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=10, space_after=10)
for item in catalog:
    para(doc, item, size=12, space_after=4)
doc.add_page_break()

# ---- 执行摘要 ----
heading1(doc, "执行摘要")
para(doc, "**核心判断**　2025年绵阳最显著的是\u201cGDP 4600.66亿/+6.5%（达成6.5%、增速全省领先）、社零破2000亿/+7.1%（西部非省会率先）、服务业+7.8%、数字经济核心+11%、高科技制造+10.4%\u201d、\u201c但规上工业+6.5%（低8%目标）、进出口312.5亿/+0.7%\u201d、\u201cCPI低位\u201d。这说明绵阳在\u201c中国科技城+五市战略\u201d中，**科技服务业/数字经济/消费强，但工业增速、外贸弱**。")
para(doc, "把2025年目标（GDP+6.5%、规上+8%、固投+4%、社零+5.5%、服务业+7.5%、CPI+2%）、2025年实际（GDP+6.5%达成、规上+6.5%、固投+2.6%、社零+7.1%、服务业+7.8%、财收+5.1%）趋势看，绵阳是\u201c科技城+军工+电子信息\u201d路径：**电子信息（1500亿）、五大特色产业、国防科研、科技服务（+26.4%）\u201d是支柱；GDP占全省比重+0.1pct。")
para(doc, "最容易记住的一句话：**绵阳是\u201c中国科技城、国防军工重镇、省域副中心\u201d，靠\u201c科技+军工+电子信息\u201d驱动。**观察绵阳，与其只看\u201cGDP 4601亿\u201d，不如看\u201c研发强度全国前列/全省第1、规上科技服务业+26.4%、高技术制造+10.4%、军工科研院所20家\u201d。")
# ---- 一、研究口径与阅读方法 ----
heading1(doc, "一、研究口径与阅读方法")
para(doc, "本报告以\u201c绵阳市政府工作报告（2025年1月，李云作）\u201d为起点，把\u201c2025年GDP目标（6.5%）\u201d与\u201c官方2025年（4600.66亿/+6.5%）\u201d并置对照，用\u201c2025年绵阳市统计公报\u201d和\u201c2026年政府工作报告\u201d横向核验。")
para(doc, "口径提示：\u201cGDP、规上工业增速\u201d按可比价（实际）；\u201c投资、消费、进出口、财政\u201d按现价（名义）。\u201c常住人口\u201d用统计公报（530万+）、城镇化率。")
para(doc, "指标体系（与研究口径一致）：**总量与增速、产业动能（科技/电子信息/军工）、外贸、财政、民生与人口**。")
para(doc, "特别提示（不吃老本）：绵阳2024年GDP 4344亿（+7%）、2025年4600.66亿/+6.5%；它不是\u201c只有长虹\u201d——**中国科技城（绵阳科技城）、国防科研院所（中物院/中国工程物理研究院）、电子信息（京东方）、军工转换\u201d才是真正底色；研发强度全省第1。")
# ---- 二、先看绵阳的特殊底盘 ----
heading1(doc, "二、先看绵阳的特殊底盘：中国科技城、电子信息军工、省域经济副中心")
para(doc, "绵阳地处四川盆地西北、涪江上中游，是**党中央、国务院批准的中国唯一科技城、国防军工重镇（中物院/中国工程物理研究院所在地）、四川省域经济副中心**；以\u201c长虹、京东方、九洲、军工科研院所\u201d著称。2025年GDP 4600.66亿（四川第2、成都之后）、常住530万+、全省领先。")
para(doc, "五个底盘名词，先立框架：")
bullet(doc, "**中国科技城**　全国唯一\u201c科技城\u201d、国家级科研院所20家、研发强度全省第1（全国前列）——\u201c科技立市\u201d。")
bullet(doc, "**电子信息（千亿）**　电子信息产值1500亿、京东方/长虹/九洲——\u201c成渝西电子信息\u201d。")
bullet(doc, "**国防军工/军民融合**　中物院/中国工程物理研究院、军转民、核技术——\u201c军工重镇\u201d。")
bullet(doc, "**先进制造/新材料**　高端装备、特种材料、锂电（绵阳SiC/电池）、医药——\u201c制造强\u201d。")
bullet(doc, "**五市战略/副中心**　产业强市/科技立市/人才兴市/开放活市/生态美市、省域经济副中心——\u201c战略定位\u201d。")
para(doc, "这五根（科技城+电子信息+军工+先进制造+副中心战略）构成绵阳独特底盘：**左手科技军工（科研院所），右手电子信息制造**。理解绵阳，先理解\u201c中国科技城、国防军工、省域副中心\u201d。")
# ---- 三、最关键的宏观错位 ----
heading1(doc, "三、最关键的宏观错位：GDP+6.5%达标、社零破2000亿强，但规上+6.5%低8%、外贸+0.7%")
para(doc, "2025年绵阳最需要辨析的一组\u201c错位\u201d：**GDP+6.5%（达标）、社零+7.1%（破2000亿、西部非省会率先）、服务业+7.8%、数字经济+11%、科技服务+26.4%强，但规上工业+6.5%（低8%目标）、进出口+0.7%、固投+2.6%（低4%）**。")
para(doc, "为什么\u201cGDP/消费/科技\u201d强，工业与外贸却不快？三解释：")
para(doc, "**其一，科技/消费/服务业强、体量大**　服务+7.8%、科技服务+26.4%、数字经济+11%、社零+7.5%、高技术制造+10.4%（计算机+68.2%）——\u201c科技+消费强\u201d。")
para(doc, "**其二，工业/制造业增速降档**　规上工业+6.5%（低8%目标）、电子信息/军工-民待转换——\u201c工业增速换挡\u201d。")
para(doc, "**其三，外贸/投资弱**　进出口+0.7%（312.5亿）、固投+2.6%、但省市项目完成率125.7%——\u201c外贸投资弱、项目好\u201d。")
para(doc, "小结：绵阳2025年是\u201c**科技服务/数字经济/消费强、工业外贸弱**\u201d：研发/科技/军工强，规上、外贸、CPI弱。")

# ---- 四、容易被忽视、但信号很重的细节（15条） ----
heading1(doc, "四、容易被忽视、但信号很重的细节（15条）")
bullet(doc, "**1.社零破2000亿/+7.1%（西部非省会率先）**\u201c消费旺、内需强。\u201d")
bullet(doc, "**2.规上科技服务业+26.4%、数字经济核心+11%**\u201c科技服务业/数字。\u201d")
bullet(doc, "**3.高技术制造+10.4%（计算+68.2%/医疗仪器+14%）**\u201c新质制造。\u201d")
bullet(doc, "**4.电子信息产值1500亿、五大特色产业+8.7%**\u201c电子/绵阳制造。\u201d")
bullet(doc, "**5.研发强度全国前列/全省第1、技术合同231亿/+15.5%**\u201c科技城/研发。\u201d")
bullet(doc, "**6.国防科研院所20家、中物院/中国工程物理研究院**\u201c军工/军民融合。\u201d")
bullet(doc, "**7.进出口312.5亿/+0.7%、跨境电商+16%**\u201c外贸弱但结构优。\u201d")
bullet(doc, "**8.规上工业+6.5%（低8%目标）**\u201c工业增速降档。\u201d")
bullet(doc, "**9.省市重点项目/1115.7亿（完成率125.7%）、新签约128个**\u201c项目/招商好。\u201d")
bullet(doc, "**10.新引进人才3万+、人才总量91.7万、科技城人才卡**\u201c人才强市。\u201d")
bullet(doc, "**11.粮食47.7亿斤、森林覆盖56.13%、生态绿色**\u201c农业+生态。\u201d")
bullet(doc, "**12.中心城区185.5万人/建成区200平方公里**\u201c副中心/城市扩张。\u201d")
bullet(doc, "**13.居民收入（城镇+4.6%/农村+5.6%）**\u201c增收、城乡比缩。\u201d")
bullet(doc, "**14.CPI低位/2%、低通胀**\u201c物价稳、需求温。\u201d")
bullet(doc, "**15.常住/城镇化率年均+1pct、省域副中心**\u201c人口/城市集聚。\u201d")
# ---- 五、2025年GDP目标 vs 实际：对照表 ----
heading1(doc, "五、2025年GDP目标 vs 实际：对照表")
table(doc,
    ["指标", "2025年目标", "2025年实际", "达成评价"],
    [
        ["地区生产总值(GDP)", "增长6.5%左右", "4600.66亿/6.5%", "达成"],
        ["规模以上工业", "增长8%", "+6.5%", "差1.5pct"],
        ["固定资产投资", "增长4%左右", "+2.6%(全省第2)", "略低"],
        ["社会消费品零售总额", "增长5.5%", "破2000亿/+7.1%", "超额"],
        ["进出口总额", "保持300亿+", "312.5亿/+0.7%", "达标"],
        ["地方一般公共预算收入", "与经济发展一致", "+5.1%", "超额"],
        ["居民收入", "与经济增长同步", "城镇+4.6%/农村+5.6%", "略低/同步"],
        ["居民消费价格", "2%左右", "低位", "低位"],
    ],
)
para(doc, "注：GDP、规上工业按可比价。**GDP（+6.5%）、社零（+7.1%）达成/超额**；**规上工业（+6.5%）、固投（+2.6%）略低目标**；服务业（+7.8%）、财收（+5.1%）优。")
para(doc, "拆读：**科技服务、数字经济、社零、财收是亮色**；**规上工业（+6.5%）、外贸（+0.7%）、固投（+2.6%）**是短板——\u201c科技城+军工\u201d是核心，\u201c科技服务强、工业外贸降档\u201d，是\u201c转型副中心\u201d样本。")

# ---- 六、2025年增长是谁撑起来的？（结构归因） ----
heading1(doc, "六、2025年增长是谁撑起来的？（结构归因）")
para(doc, "把绵阳GDP的6.5%拆开：三次产业分别增3.9%、5.6%、7.8%（结构7.8：40.0：52.2）。**第三产业（服务业）+7.8%最快、占比过半**，第二产业（制造）+5.6%、第一产业（农业）稳\u2014\u2014\u201c三产领跑\u201d。")
para(doc, "2026年绵阳强调\u201c省域副中心、中国科技城、五市战略\u201d，聚焦**电子信息、先进制造、核产业（军转民）、数字经济、低空/人工智能**——核心是\u201c科技+制造\u201d。")
para(doc, "**第二产业（工业/制造）**：规上+6.5%、电子信息1500亿、五大特色产业+8.7%、高技术制造+10.4%——\u201c信息/军工制造\u201d。")
para(doc, "**第三产业（服务业）**：+7.8%（科技服务+26.4%、数字经济+11%、金融/商贸）——\u201c科技服务业+消费\u201d强。")
para(doc, "**第一产业（农业）**：+3.9%（粮食47.7亿斤、生猪、特色农业）——\u201c农业稳\u201d。")
para(doc, "一句话归因：**2025年绵阳增长\u201c靠第三产业（科技服务+消费）+电子信息制造\u201d**，工业增速换挡、外贸弱；\u201c科技+消费\u201d是核心。")

# ---- 七、预算与财政的\u201c含金量\u201d ----
heading1(doc, "七、预算与财政的\u201c含金量\u201d")
para(doc, "2025年绵阳**地方一般公共预算收入+5.1%**；税占比稳定在50%以上；民生支出占比较稳定（65%以上）。")
bullet(doc, "财收+5.1%（高GDP增速）、税占比50%+——\u201c财政稳健\u201d。")
bullet(doc, "民生支出占65%+（教育/科技/医疗）。")
bullet(doc, "金融支撑：存款+9.7%、贷款+14.7%（科技贷款+12%）——信贷密集支持科技/制造业。")
para(doc, "**财政含金量小结**：财收+5.1%（超额）、税占比50%+、科技信贷强；财政对\u201c科技城、军工转换、数字经济、民生\u201d投入加大。")

# ---- 八、民生底账：人口、收入与城乡 ----
heading1(doc, "八、民生底账：人口、收入与城乡")
para(doc, "2025年绵阳**城镇/农村居民人均可支配收入分别+4.6%/+5.6%**；城乡居民收入比缩小；就业：城镇新增就业6.2万人（超4.5万目标）。")
para(doc, "人口画像：**常住约530万/560万、中心城区185.5万、城镇化率年均+1pct**；省域副中心、军工/科技吸纳人才、九个大县。")
para(doc, "民生投入：低保/特困、科技城人才、教育医疗、保障房、养老托育——民生支出占65%+、扎实。")

# ---- 九、城镇与农村：格局与均衡 ----
heading1(doc, "九、城镇与农村：格局与均衡")
para(doc, "绵阳城镇化率（中心城区185.5万/全市约55%）；县域经济（三台/江油/梓潼/盐亭农业）；农村收入增速（+5.6%）>城镇（+4.6%），**城乡比缩小**；五市战略、全域旅游。")
para(doc, "农业底盘：**粮食47.7亿斤（创历史新高）、生猪、油菜、林果/中药材**——\u201c天府粮仓+特色农业\u201d。")
para(doc, "一句话：\u201c绵阳是科技城+农业大市、城乡融合、县域经济强\u201d。")

# ---- 十、人口流入与流出 ----
heading1(doc, "十、人口流入与流出")
para(doc, "绵阳常住约530-560万、省域副中心、中心城区185.5万（较十三五+13.4万）；\u201c科技城/军工/高校\u201d吸纳人才（人才91.7万、新引进3万+）；科教文卫资源辐射川西北。")
para(doc, "结构观察：**科技/大学城（西南科大/绵师）聚精英**；县域人口向中心城区集中。")
para(doc, "2026年目标：新引人才3万+、科技城建设——绵阳靠\u201c科技城+军工+副中心\u201d聚人。")

# ---- 十一、物价与货币环境 ----
heading1(doc, "十一、物价与货币环境")
para(doc, "2025年绵阳**CPI低位（2%内）**、PPI/需求弱——\u201c低通胀、需求温\u201d。")
bullet(doc, "信贷扩张：存款+9.7%、贷款+14.7%（科技贷款+12%）——宽信用支持科技/制造。")
bullet(doc, "消费：社零破2000亿/+7.1%、以旧换新/新能源汽车——消费旺。")
para(doc, "货币环境判断：**宽信用、CPI低位**；绵阳靠\u201c科技+消费\u201d稳需求（2026 CPI 2%）。")

# ---- 十二、区域一体化：绵阳的位置 ----
heading1(doc, "十二、区域一体化：绵阳在成都都市圈、成渝双城经济圈、中国科技城\u201c三圈\u201d里的位置")
para(doc, "绵阳是**成都都市圈/川北重要城市、成渝地区双城经济圈重要节点、中国唯一科技城、省域经济副中心城市**。")
bullet(doc, "**成都都市圈**　毗邻成都、承接产业辐射（四川第2极）、成绵一体化。")
bullet(doc, "**成渝双城经济圈**　成渝中部崛起、科技城融入成渝科创走廊。")
bullet(doc, "**科技城/军工**　国家战略科技力量、国防军工、核技术、军民融合。")
para(doc, "一句话：**绵阳在\u201c成都+成渝+科技城\u201d里，最核心是\u201c科技军工+省域副中心\u201d**；科技、军工、区位是最大优势。")

# ---- 十三、未来5—10年最值得观察的五条主线 ----
heading1(doc, "十三、未来5—10年最值得观察的五条主线")
bullet(doc, "**主线一：中国科技城（研发/平台）**\u201c国家战略科技、绵阳科技城新区\u201d能否辐射西南创新。")
bullet(doc, "**主线二：军工/国防科研（中物院）**\u201c军转民、核技术/激光\u201d新赛道。")
bullet(doc, "**主线三：电子信息（京东方）**\u201c1500亿\u201d能否冲、新型显示/半导体。")
bullet(doc, "**主线四：先进制造/锂电/低空**\u201c先进制造、低空经济、固投\u201d强链。")
bullet(doc, "**主线五：副中心/人口/外贸**\u201c中心城区聚人、外贸破300亿\u201d如何\u201c提快工业、扩开放\u201d。")

# ---- 十四、最终结论 ----
heading1(doc, "十四、最终结论：绵阳在\u201c科技城+电子信息+省域副中心\u201d里的增长逻辑")
para(doc, "把2025年闭环打穿：**绵阳是\u201c中国科技城、国防军工重镇、省域副中心\u201d**：GDP 4600.66亿/+6.5%、社零破2000亿、研发强度全省第1、电子信息1500亿。")
para(doc, "绵阳不是\u201c只有长虹\u201d——它是**科技城（研发）+电子信息+军工+数字经济+先进制造**的复合，靠\u201c科技+军工\u201d驱动；但规上工业增速、外贸、固投弱。")
para(doc, "一句话结论：**绵阳是\u201c中国科技城、军工重镇、省域副中心\u201d；观察它先看\u201c研发强度、科技服务、电子信息、军工转换、社零\u201d，再看\u201c规上、外贸、固投\u201d。**它是\u201c科技强、工业外贸降档\u201d的西部转型样本。")

# ---- 附录A ----
heading1(doc, "附录A：主要资料来源与核验路径")
bullet(doc, "《2025年绵阳市人民政府工作报告》（2025年1月，李云作，2025年目标、2024年回顾+7%）")
bullet(doc, "《2025年绵阳市国民经济和社会发展统计公报》（绵阳市统计局，2026-05-31，2025年实际）")
bullet(doc, "《2026年绵阳市人民政府工作报告》（2026年3月，江彬，复盘+2026年目标）")
bullet(doc, "绵阳市人民政府/统计局（my.gov.cn）")

# ---- 附录B ----
heading1(doc, "附录B：建议建立的年度跟踪仪表盘")
bullet(doc, "GDP总量/增速 vs 全年目标。")
bullet(doc, "规模以上工业/高技术制造增速。")
bullet(doc, "科技城/研发/技术合同。")
bullet(doc, "电子信息/军工/新材料。")
bullet(doc, "社零/数字经济/科技服务。")
bullet(doc, "进出口/跨境电商。")
bullet(doc, "固定资产/重大项目/项目投资。")
bullet(doc, "财收/税收/民生%。")
bullet(doc, "常住/城镇化/人才。")
bullet(doc, "CPI/存贷款/工业用电。")

# ---------------- 保存 ----------------
out = "/Users/x/Desktop/content-prod-lab/reports/绵阳市_2025年政府工作报告_深度研究_2026-08-14.docx"
doc.save(out)
print("SAVED OK: 绵阳市", out)
