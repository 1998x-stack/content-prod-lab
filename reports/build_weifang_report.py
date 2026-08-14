# -*- coding: utf-8 -*-
"""Build 济宁市2025年政府工作报告 深度研究 DOCX, 参照地级市系列版式。"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DARK = RGBColor(0x1F, 0x2A, 0x44)
RED = RGBColor(0xB0, 0x1E, 0x2E)
GRAY = RGBColor(0x59, 0x60, 0x69)
BLUE = RGBColor(0x1F, 0x4E, 0x79)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
HEADER_FILL = "1F2A44"
K = "微软雅黑"


def set_run(run, size=11, bold=False, color=DARK, font=K):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    run.font.name = font
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), font)
    rFonts.set(qn("w:hAnsi"), font)
    rFonts.set(qn("w:eastAsia"), font)


def para(doc, text, size=11, bold=False, color=DARK, align=None,
         space_after=6, space_before=0, font=K):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        set_run(r, size=size, bold=bold or (i % 2 == 1), color=color, font=font)
    return p


def heading1(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(16)
    pf.space_after = Pt(8)
    r = p.add_run(text)
    set_run(r, size=16, bold=True, color=RED)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "B01E2E")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def heading2(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(10)
    pf.space_after = Pt(4)
    r = p.add_run(text)
    set_run(r, size=13, bold=True, color=BLUE)
    return p


def table(doc, headers, rows, widths=None, font_size=9.5):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_run(r, size=font_size, bold=True, color=WHITE)
        tcPr = hdr[i]._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), HEADER_FILL)
        tcPr.append(shd)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(val))
            set_run(r, size=font_size, color=DARK)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    return t


def quote_box(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(8)
    pf.left_indent = Pt(12)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), "B01E2E")
    pBdr.append(left)
    pPr.append(pBdr)
    r = p.add_run(text)
    set_run(r, size=11, color=DARK, font="楷体")
    return p


def bullet(doc, text, size=10.5):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.7)
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        set_run(r, size=size, bold=(i % 2 == 1), color=DARK)
    return p


# ================================================================= 文档主体
doc = Document()
sec = doc.sections[0]
sec.page_width = Cm(21)
sec.page_height = Cm(29.7)
sec.top_margin = Cm(2.4)
sec.bottom_margin = Cm(2.2)
sec.left_margin = Cm(2.5)
sec.right_margin = Cm(2.5)

st = doc.styles["Normal"]
st.font.name = K
st.font.size = Pt(11)
st.element.rPr.rFonts.set(qn("w:eastAsia"), K)


# ---- 封面 ----
para(doc, "潍坊市2025年政府工作报告", size=24, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=60, space_after=4)
para(doc, "深度研究与\u201c容易被忽视的细节\u201d分析", size=16, bold=True, color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
para(doc, "从\u201c制造业强市、现代农业样板、民营经济、潍柴歌尔、海洋经济\u201d重新理解潍坊", size=12, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)
para(doc, "\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501", color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
para(doc, "研究对象：2025年潍坊市政府工作报告及同期财政、国民经济和社会发展统计资料、2026年官方复盘", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
para(doc, "标定日期：2026-08-14", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
doc.add_page_break()

# ---- 目录 ----
catalog = [
    "执行摘要",
    "一、研究口径与阅读方法",
    "二、先看潍坊的特殊底盘：制造业强市、现代农业、民营经济、潍柴歌尔",
    "三、最关键的宏观错位：制造业+消费外贸强，但固投-4.5%、地产-18%、财收+1.8%",
    "四、容易被忽视、但信号很重的细节（15条）",
    "五、2025年GDP目标 vs 实际：对照表",
    "六、2025年增长是谁撑起来的？（结构归因）",
    "七、预算与财政的\u201c含金量\u201d",
    "八、民生底账：人口、收入与城乡",
    "九、城镇与农村：格局与均衡",
    "十、人口流入与流出",
    "十一、物价与货币环境",
    "十二、区域一体化：潍坊在济南都市圈、胶东经济圈、山东半岛\u201c三圈\u201d里的位置",
    "十三、未来5—10年最值得观察的五条主线",
    "十四、最终结论：潍坊在\u201c高端制造+现代农业+民营经济\u201d里的增长逻辑",
    "附录A：主要资料来源与核验路径",
    "附录B：建议建立的年度跟踪仪表盘",
]
para(doc, "目　录", size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=10, space_after=10)
for item in catalog:
    para(doc, item, size=12, space_after=4)
doc.add_page_break()

# ---- 执行摘要 ----
heading1(doc, "执行摘要")
para(doc, "**核心判断**　2025年潍坊最显著的是“GDP 8587.36亿/+5.5%（达成目标、冲万亿、山东前列）、规上工业+7.5%、社零+6.3%、进出口+6.8%、农林牧渔总产值1409.7亿（全省第1）”、“但固定资产投资-4.5%、房地产开发-18.1%、一般公共预算收入+1.8%（差3%目标）、CPI-0.4%”。这说明潍坊在“制造强市+农业样板”中，**工业出口农业强、投资地产财政弱**。")
para(doc, "把2025年目标（GDP+5.5%、规上+6.5%、财收+3%、服务业高于全省）、2025年实际（GDP+5.5%达成、规上+7.5%超额、社零+6.3%、进出口+6.8%、财收+1.8%）趋势看，潍坊是\u201c制造业+现代农业\u201d路径：**动力装备（潍柴）、电子信息（歌尔）、化工高端、农业产业化\u201d是支柱；2025年冲万亿。")
para(doc, "最容易记住的一句话：**潍坊是\u201c世界风筝都、制造业强市、现代农业样板\u201d，靠\u201c潍柴歌尔+动力/光电/化工/农业\u201d增长。**观察潍坊，与其只看\u201cGDP 8587亿\u201d，不如看\u201c民营经济山东百强第1、现代农业全省第1、进出口破3800亿、潍柴雷沃青啤竣工\u201d。")
# ---- 一、研究口径与阅读方法 ----
heading1(doc, "一、研究口径与阅读方法")
para(doc, "本报告以\u201c潍坊市政府工作报告（2025年1月，刘建军作）\u201d为起点，把\u201c2025年GDP目标（5.5%）\u201d与\u201c官方2025年（8587.36亿/+5.5%）\u201d并置对照，用\u201c2025年潍坊市统计公报\u201d和\u201c2026年政府工作报告\u201d横向核验。")
para(doc, "口径提示：\u201cGDP、规上工业增速\u201d按可比价（实际）；\u201c投资、消费、进出口、财政\u201d按现价（名义）。\u201c常住人口\u201d用统计公报（927.12万）、城镇化率68.09%。")
para(doc, "指标体系（与研究口径一致）：**总量与增速、产业动能（制造/动力/光电/化工/农业）、外贸、财政、民生与人口**。")
para(doc, "特别提示（不吃老本）：潍坊2024年GDP突破8100亿（+6%）、2025年8587.36亿/+5.5%（冲万亿）；它不是\u201c只有风筝\u201d——**潍柴（动力）、歌尔（声光电）、豪迈（精密）、寿光/诸城农业、民营经济\u201d才是真正底色。")
# ---- 二、先看潍坊的特殊底盘 ----
heading1(doc, "二、先看潍坊的特殊底盘：制造业强市、现代农业、民营经济、潍柴歌尔")
para(doc, "潍坊地处山东半岛中部，是**国家制造业强市、世界风筝都（国际风筝联合会）、中国农业产业化发源地**；是\u201c潍柴动力、歌尔声学、豪迈科技\u201d的代名词。2025年GDP 8587.36亿（冲万亿）、常住927.12万、城镇化率68.09%、山东第4（青岛/济南/烟台之后）。")
para(doc, "五个底盘名词，先立框架：")
bullet(doc, "**动力装备（潍柴）**　潍柴动力、雷沃农业装备——\u201c动力之都、发动机\u201d。")
bullet(doc, "**电子信息（歌尔）**　声学/光学/智能穿戴（VR/AR）、歌尔——\u201c声光电\u201d。")
bullet(doc, "**现代化工/精密**　高端化工（石化/新材料）、豪迈精密机械、纺织——\u201c制造强\u201d。")
bullet(doc, "**现代农业/农业产业化**　寿光蔬菜、诸城肉鸡、青州花卉、农业总产值全省第1（1409.7亿）——\u201c中国农业样板\u201d。")
bullet(doc, "**民营经济/海洋**　民营经济（潍坊民企百强山东第1）、海洋牧场/盐化工——\u201c民营+海洋\u201d。")
para(doc, "这五根（动力+光电+化工+农业+民营）构成潍坊独特底盘：**左手制造业（潍柴/歌尔），右手现代农业+民营**。理解潍坊，先理解\u201c风筝都、制造强市、农业样板\u201d。")
# ---- 三、最关键的宏观错位 ----
heading1(doc, "三、最关键的宏观错位：制造业+消费外贸强，但固投-4.5%、地产-18%、财收+1.8%")
para(doc, "2025年潍坊最需要辨析的一组\u201c错位\u201d：**规上工业+7.5%、社零+6.3%、进出口+6.8%、民营经济山东第1强，但固定资产投资-4.5%、房地产开发-18.1%、财收+1.8%（差3%目标）、CPI-0.4%**。")
para(doc, "为什么\u201c工业/消费/外贸这么强\u201d，投资与财政却不温不火？三解释：")
para(doc, "**其一，工业/民营旺、经济体量大**　规上+7.5%（化学制造+11.4%、通用设备+13.3%、食品+16%）、民营百强山东第1、农林牧渔总产值全省第1——\u201c制造+农业强\u201d。")
para(doc, "**其二，投资、地产、财政弱**　固投-4.5%（三产-14.2%）、地产-18.1%（商品房销售走弱）、财收+1.8%——\u201c投资地产财政不温\u201d。")
para(doc, "**其三，消费/外贸好、物价弱**　社零+6.3%（家电+25.1%）、进出口+6.8%（进口+12.1%）—但CPI-0.4%、PPI弱——\u201c量好价弱、需求修复\u201d。")
para(doc, "小结：潍坊2025年是\u201c**制造农业外贸强、投资地产财政弱**\u201d：潍柴/歌尔/农业、消费好，地产投资、财政、物价弱。")

# ---- 四、容易被忽视、但信号很重的细节（15条） ----
heading1(doc, "四、容易被忽视、但信号很重的细节（15条）")
bullet(doc, "**1.规上工业+7.5%、化学制造+11.4%/食品制造+16%**\u201c动力/化工/食品拉动。\u201d")
bullet(doc, "**2.民营经济山东百强第1（十强3家/百强17家/就业百强14家）**\u201c民营绝对主力。\u201d")
bullet(doc, "**3.农林牧渔总产值1409.7亿（全省第1/+4.6%）**\u201c现代农业样板（寿光蔬菜/诸城肉鸡）。\u201d")
bullet(doc, "**4.进出口3797.6亿/+6.8%（进口+12.6%）**\u201c外贸强、进口原料旺。\u201d")
bullet(doc, "**5.社零3584.9亿/+6.3%（家电+25.1%、餐饮-小幅）**\u201c以旧换新、消费旺。\u201d")
bullet(doc, "**6.固定资产投资-4.5%（三产-14.2%）**\u201c投资靠二产（+2.8%）。\u201d")
bullet(doc, "**7.房地产开发-355.6亿/-18.1%**\u201c地产深度调整、拖累。\u201d")
bullet(doc, "**8.一般公共预算收入630.5亿/+1.8%（税收占60.9%）**\u201c财收弱、税收占比待提高。\u201d")
bullet(doc, "**9.潍柴雷沃高端农业装备、青啤生产基地等291个重点项目竣工**\u201c工业大项目、增量445亿。\u201d")
bullet(doc, "**10.新增国家级制造业单项冠军、昌盛高新（豪迈）**\u201c\u2018小巨人\u2019、精密高端制造。\u201d")
bullet(doc, "**11.潍坊港全国首个\u2018零碳港口\u2019、海洋经济**\u201c零碳、向海图强。\u201d")
bullet(doc, "**12.中国人民银行存款+10.2%、贷款+7.7%**\u201c宽信用、信贷密集。\u201d")
bullet(doc, "**13.居民收入45794元/+4.6%、城乡比1.82（收窄）**\u201c农村+4.9%>城镇+4.0%。\u201d")
bullet(doc, "**14.CPI-0.4%（食品烟酒-1.5%）**\u201c低通胀、量价。\u201d")
bullet(doc, "**15.常住927.12万、城镇化68.09%**\u201c山东人口大市、城镇化稳升。\u201d")
# ---- 五、2025年GDP目标 vs 实际：对照表 ----
heading1(doc, "五、2025年GDP目标 vs 实际：对照表")
table(doc,
    ["指标", "2025年目标", "2025年实际", "达成评价"],
    [
        ["地区生产总值(GDP)", "增长5.5%左右", "8587.36亿/5.5%", "达成(冲万亿)"],
        ["规模以上工业", "增长6.5%左右", "+7.5%", "超额"],
        ["一般公共预算收入", "增长3%", "630.5亿/+1.8%", "差1.2pct"],
        ["固定资产投资", "高于全省平均", "-4.5%", "下行"],
        ["社会消费品零售总额", "高于全省平均", "3584.9亿/+6.3%", "超额"],
        ["进出口总额", "促稳提质", "3797.6亿/+6.8%", "超额"],
        ["居民收入", "与经济增长同步", "45794元/+4.6%", "略低GDP"],
    ],
)
para(doc, "注：GDP、规上工业按可比价。**GDP（+5.5%）、规上（+7.5%）、社零（+6.3%）、进出口（+6.8%）达成/超额**；**固投（-4.5%）、财收（+1.8%差3%）、房价**偏弱。")
para(doc, "拆读：**民营（山东第1）、农业（全省第1）、制造（+7.5%）、消费/外贸是亮色**；**固投（-4.5%）、地产（-18.1%）、财收（+1.8%）**是短板——\u201c制造农业强、投资财政弱\u201d，是\u201c产业强市\u201d样本。")

# ---- 六、2025年增长是谁撑起来的？（结构归因） ----
heading1(doc, "六、2025年增长是谁撑起来的？（结构归因）")
para(doc, "把潍坊GDP的5.5%拆开：三次产业分别增4.0%、4.6%、6.5%（结构8.5：41.2：50.3）。**第三产业（服务业）增速最快（+6.5%）、占比过半**，第二产业（制造业）总量最大，第一产业（农业）稳\u2014\u2014\u201c三产+二产\u201d双轮。")
para(doc, "2026年潍坊强调\u201c冲万亿、区域副中心、北方增长极\u201d，聚焦**动力装备升级、声光电（歌尔）、高端化工新材料、现代农业、海洋经济**——核心是\u201c制造强市+农业样板\u201d。")
para(doc, "**第二产业（工业）**：规上+7.5%（化学制造+11.4%、通用设备+13.3%、食品+16%、农副食品+12.1%）、潍柴/歌尔/豪迈——\u201c制造业+新质生产力\u201d强。")
para(doc, "**第三产业（服务业）**：+6.5%（商贸、物流、会展、文旅）、社零+6.3%——\u201c服务业+扩张需求\u201d。")
para(doc, "**第一产业（农业）**：+4.0%（农林牧渔总产值1409.7亿全省第1、粮食451.7万吨、蔬菜/肉鸡）——\u201c农业样板、稳\u201d。")
para(doc, "一句话归因：**2025年潍坊增长\u201c靠制造业（动力/光电/化工）+服务业（消费）+农业\u201d**，地产投资、财政弱；\u201c产业+民营\u201d是核心。")

# ---- 七、预算与财政的\u201c含金量\u201d ----
heading1(doc, "七、预算与财政的\u201c含金量\u201d")
para(doc, "2025年潍坊**一般公共预算收入630.5亿元（+1.8%）**；税收收入383.85亿（占比60.9%）；一般公共预算支出909亿（+0.8%）、民生占82.2%。")
bullet(doc, "税收占60.9%（较开发区/税基、制造税）——\u201c税收结构待优\u201d。")
bullet(doc, "民生支出占82.2%（连续10年超80%、教育/社保/医疗）。")
bullet(doc, "金融支撑：存款+10.2%、贷款+7.7%——信贷充裕、支持制造/农业/民营。")
para(doc, "**财政含金量小结**：财收+1.8%（低GDP增速）、靠制造/农业税，民生82.2%；财政对\u201c新质生产力、农业、民生\u201d投入加大。")
# ---- 八、民生底账：人口、收入与城乡 ----
heading1(doc, "八、民生底账：人口、收入与城乡")
para(doc, "2025年潍坊**居民人均可支配收入45794元（+4.6%）**，其中城镇55875元（+4.0%）、农村30643元（+4.9%），城乡比1.82（收窄）。消费：人均消费支出26601元（+3.0%）。就业：城镇新增就业10.61万人。")
para(doc, "人口画像：**常住927.12万、城镇化68.09%（山东人口大市）**；潍坊工业/农业/民营吸纳就业、出生率偏低。")
para(doc, "民生投入：退休职工养老金人均月增103元、居民养老金最低190元/月、医保住院——民生扎实（连续10年80%+）。")

# ---- 九、城镇与农村：格局与均衡 ----
heading1(doc, "九、城镇与农村：格局与均衡")
para(doc, "潍坊城镇化率68.09%，县域经济极强（寿光蔬菜/诸城肉鸡/青州花卉/高密装备/临朐铝材）；农村收入增速（+4.9%）>城镇（+4.0%），**城乡比1.82收窄**；乡村振兴、高标准农田30.4万亩。")
para(doc, "农业底盘：**粮食451.7万吨（连续2年450万吨+）、蔬菜、肉鸡、畜牧、水产；农林牧渔总产值1409.7亿全省第1**——\u201c中国农业产业化发源地/菜篮子（粤港澳大湾区4个基地）\u201d。")
para(doc, "一句话：\u201c潍坊是工业强市+农业样板、县域经济发达\u201d，城乡融合全国前列。")

# ---- 十、人口流入与流出 ----
heading1(doc, "十、人口流入与流出")
para(doc, "潍坊常住927.12万（山东第3人口大市）、城镇化68.09%；\u201c制造业/农业/民营\u201d吸纳就业，但部分青年外流至青岛/济南，主城区（奎文/高新/寒亭）+峡山吸引。")
para(doc, "结构观察：**出生约4.3万人（偏低）、农业人口与转移**；青壮年进产业、县城化明显。")
para(doc, "2026年目标：新引育人才、支持就业——潍坊靠\u201c产业+农业+海洋\u201d聚人。")

# ---- 十一、物价与货币环境 ----
heading1(doc, "十一、物价与货币环境")
para(doc, "2025年潍坊**CPI-0.4%**（食品烟酒-1.5%、交通通信-3.1%；衣着+1.5%、教育文娱+2.5%）——\u201c低通胀、需求修复\u201d。")
bullet(doc, "信贷扩张：存款+10.2%、贷款+7.7%——宽信用支持制造/农业/民营。")
bullet(doc, "消费：社零+6.3%（家电+25.1%、新消费）——扩张。")
para(doc, "货币环境判断：**宽信用、CPI-0.4%**；潍坊靠\u201c制造+消费+农业\u201d稳需求（2026 CPI温和回升）。")

# ---- 十二、区域一体化：潍坊的位置 ----
heading1(doc, "十二、区域一体化：潍坊在济南都市圈、胶东经济圈、山东半岛\u201c三圈\u201d里的位置")
para(doc, "潍坊是**济南都市圈、胶东经济圈交汇节点、山东半岛中部枢纽、环渤海经济带重要城市**。")
bullet(doc, "**胶东经济圈**　对接青岛（港口）/烟台（海洋）、胶东一体化、向海。")
bullet(doc, "**济南都市圈/省会**　西接济南、山东半岛中轴、胶济铁路。")
bullet(doc, "**开放一带一路**　潍坊港（零碳）、海陆空立体、RCEP——区域开放枢纽。")
para(doc, "一句话：**潍坊在\u201c胶东+济南双圈\u201d里，最核心是\u201c制造业+农业+临港\u201d**；产业基础、民营、农业是最大优势。")

# ---- 十三、未来5—10年最值得观察的五条主线 ----
heading1(doc, "十三、未来5—10年最值得观察的五条主线")
bullet(doc, "**主线一：动力装备/新能源（潍柴雷沃/氢能）**\u201c发动机、智能农机\u201d能否冲击万亿/千亿。")
bullet(doc, "**主线二：声光电/VR（歌尔）**\u201c歌尔歌、元宇宙\u201d智造。")
bullet(doc, "**主线三：高端化工/新材料/精密**\u201c豪迈、新和成\u201d新质。")
bullet(doc, "**主线四：现代农业/食品（寿光诸城）**\u201c农业产业化升级、菜篮子\u201d。")
bullet(doc, "**主线五：海洋经济/零碳+人口**\u201c潍坊港零碳、海洋牧场\u201d、如何\u201c聚人、冲万亿\u201d。")

# ---- 十四、最终结论 ----
heading1(doc, "十四、最终结论：潍坊在\u201c高端制造+现代农业+民营经济\u201d里的增长逻辑")
para(doc, "把2025年闭环打穿：**潍坊是\u201c制造强市、农业样板、民营大市\u201d**：GDP 8587.36亿/+5.5%（冲万亿）、规上+7.5%、农林牧渔全省第1、民营山东第1、进出口3797.6亿。")
para(doc, "潍坊不是\u201c只有风筝\u201d——它是**动力装备+声光电+高端化工+现代农业+民营**的复合，靠\u201c产业+农业\u201d驱动；但地产投资、财政、物价弱。")
para(doc, "一句话结论：**潍坊是\u201c世界风筝都、制造强市、农业样板\u201d；观察它先看\u201c民营、制造业（潍柴歌尔）、农业、进出口\u201d，再看\u201c地产、财政、物价\u201d。**它是\u201c产业强、投资弱、财收待优\u201d的山东样本。")

# ---- 附录A ----
heading1(doc, "附录A：主要资料来源与核验路径")
bullet(doc, "《2025年潍坊市政府工作报告》（2025年1月，刘建军作，2025年目标、2024年回顾+6%突破8100亿）")
bullet(doc, "《2025年潍坊市国民经济和社会发展统计公报》（潍坊市统计局，2026-03-19，2025年实际）")
bullet(doc, "《2026年潍坊市政府工作报告》（2026年1月，复盘+2026年目标，冲万亿）")
bullet(doc, "潍坊市人民政府/统计局、中新网山东、潍坊日报等")

# ---- 附录B ----
heading1(doc, "附录B：建议建立的年度跟踪仪表盘")
bullet(doc, "GDP总量/增速 vs 全年目标。")
bullet(doc, "规上工业（动力/光电/化工）增速。")
bullet(doc, "民营经济、百强/就业。")
bullet(doc, "固定资产/工业/房地产投资。")
bullet(doc, "农林牧渔/粮食/农业产业链。")
bullet(doc, "社零、以旧换新、家电。")
bullet(doc, "进出口、家电/电子、外资。")
bullet(doc, "一般公共预算/税收占比/民生%。")
bullet(doc, "常住/城镇化、青年人才。")
bullet(doc, "CPI、存贷款、海洋/零碳港口。")

# ---------------- 保存 ----------------
out = "/Users/x/Desktop/content-prod-lab/reports/潍坊市_2025年政府工作报告_深度研究_2026-08-14.docx"
doc.save(out)
print("SAVED OK: 潍坊市", out)
