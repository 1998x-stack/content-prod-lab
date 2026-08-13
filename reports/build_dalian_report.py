# -*- coding: utf-8 -*-
"""Build 大连市2025年政府工作报告 深度研究 DOCX, 参照省市系列版式。"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DARK = RGBColor(0x1F, 0x2A, 0x44)
RED = RGBColor(0xB0, 0x1E, 0x2E)
GRAY = RGBColor(0x59, 0x60, 0x69)
BLUE = RGBColor(0x1F, 0x4E, 0x79)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
HEADER_FILL = "1F2A44"
K = "微软雅黑"


def set_run(run, size=11, bold=False, color=DARK, font=K):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    run.font.name = font
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), font)
    rFonts.set(qn("w:hAnsi"), font)
    rFonts.set(qn("w:eastAsia"), font)


def para(doc, text, size=11, bold=False, color=DARK, align=None,
         space_after=6, space_before=0, font=K):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        set_run(r, size=size, bold=bold or (i % 2 == 1), color=color, font=font)
    return p


def heading1(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(16)
    pf.space_after = Pt(8)
    r = p.add_run(text)
    set_run(r, size=16, bold=True, color=RED)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "B01E2E")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def heading2(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(10)
    pf.space_after = Pt(4)
    r = p.add_run(text)
    set_run(r, size=13, bold=True, color=BLUE)
    return p


def table(doc, headers, rows, widths=None, font_size=9.5):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_run(r, size=font_size, bold=True, color=WHITE)
        tcPr = hdr[i]._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), HEADER_FILL)
        tcPr.append(shd)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(val))
            set_run(r, size=font_size, color=DARK)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    return t


def quote_box(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(8)
    pf.left_indent = Pt(12)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), "B01E2E")
    pBdr.append(left)
    pPr.append(pBdr)
    r = p.add_run(text)
    set_run(r, size=11, color=DARK, font="楷体")
    return p


def bullet(doc, text, size=10.5):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.7)
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        set_run(r, size=size, bold=(i % 2 == 1), color=DARK)
    return p


# ================================================================= 文档主体
doc = Document()
sec = doc.sections[0]
sec.page_width = Cm(21)
sec.page_height = Cm(29.7)
sec.top_margin = Cm(2.4)
sec.bottom_margin = Cm(2.2)
sec.left_margin = Cm(2.5)
sec.right_margin = Cm(2.5)

st = doc.styles["Normal"]
st.font.name = K
st.font.size = Pt(11)
st.element.rPr.rFonts.set(qn("w:eastAsia"), K)



# ---- 封面 ----
para(doc, "大连市2025年政府工作报告", size=24, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=60, space_after=4)
para(doc, "深度研究与“容易被忽视的细节”分析", size=16, bold=True, color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
para(doc, "从“东北开放门户、港口、石化、海洋与软件IT”重新理解大连", size=12, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)
para(doc, "━━━━━━━━━━━━━━━━", color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
para(doc, "研究对象：2025年大连市政府工作报告及同期财政、国民经济和社会发展统计资料、2026年官方复盘", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
para(doc, "标定日期：2026-08-13", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
doc.add_page_break()

# ---- 目录 ----
catalog = [
    "执行摘要",
    "一、研究口径与阅读方法",
    "二、先看大连的特殊底盘：东北开放门户、港口/石化、海洋经济与软件IT",
    "三、最关键的宏观错位：GDP破万亿、工业/出口强，但投资/地产/人口弱",
    "四、容易被忽视、但信号很重的细节（15条）",
    "五、2025年GDP目标 vs 实际：对照表",
    "六、2025年增长是谁撑起来的？（结构归因）",
    "七、预算与财政的“含金量”",
    "八、民生底账：人口、收入与城乡",
    "九、城镇与农村：格局与均衡",
    "十、人口流入与流出",
    "十一、物价与货币环境",
    "十二、区域一体化：大连在“东北开放门户+沿海/港口+大连都市圈”里的位置",
    "十三、未来5—10年最值得观察的五条主线",
    "十四、最终结论：大连在“万亿GDP+港口/石化+软件/海洋”里的增长逻辑",
    "附录A：主要资料来源与核验路径",
    "附录B：建议建立的年度跟踪仪表盘",
]
para(doc, "目　录", size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=10, space_after=10)
for item in catalog:
    para(doc, item, size=12, space_after=4)
doc.add_page_break()

# ---- 执行摘要 ----
heading1(doc, "执行摘要")
para(doc, "**核心判断**　如果只看新闻标题，2025年大连最显眼的是“GDP破万亿（10002.1亿）、增长5.7%”、“规上工业+11.7%”、“民用钢质船舶+74.6%”和“港口集装箱548万标箱”。但这份研究真正值得深读的，是这座“东北开放门户+港口+石化+软件IT”的计划单列市如何实现万亿突破，并在固定资产投资（-17.0%）、房地产（-30.1%）与人口负增背景下，靠“工业/出口/海洋/高端制造”实现彰显东北振兴的“大连升级”。")
para(doc, "把2025年初设定的目标（GDP 5.5%以上、规上工业7%、固投5%左右）、2025年统计公报、2026年复盘放在一起看，大连呈现清晰暗线：**从“石化/传统重工”的旧底盘，向“高技术制造/装备/船舶+软件IT+海洋经济+开放门户”转型**。旧引擎（地产/传统基建）在调整；新引擎（装备、船舶、化工新材料、软件、期货、海洋）被要求更快补位。")
para(doc, "因此，本报告不按政府工作报告逐段复述，而采用“显性表述—同期数据—制度含义—长期影响”的方式，专门提取容易被忽视、但对判断大连未来5—10年发展模式更有价值的信号。")
para(doc, "最容易记住的一句话：**大连是“东北开放门户+港口/石化+海洋经济+软件IT”的计划单列样本，靠“工业+出口+海洋+高端制造”撑起万亿GDP。**观察大连，与其看“GDP 1万亿”，不如看“港口、石化、船舶、软件IT、期货、海洋牧场、装备制造”这几张名片。")
heading2(doc, "一页速览：2025年大连经济的“表与里”")
table(doc,
    ["维度", "表面现象", "更深层信号"],
    [
        ["增长", "GDP 10002.1亿、+5.7%（首次破万亿）", "一产6.6%、二产35.3%、三产58.0%"],
        ["产业", "规上工业+11.7%", "高技术+13.9%、装备+15.4%"],
        ["外贸", "进出口4492.6亿、-0.3%（出口+10.6%）", "港口/船舶/石化出海"],
        ["投资", "固定资产投资-17.0%", "地产-30.6%、民间-14.1%"],
        ["财政", "地方一般公共预算收入749.5亿、-3.3%", "税收+3.4%"],
        ["消费", "社零2180.8亿、+2.1%", "家电+109.7%、通讯+80.5%"],
        ["人口", "常住607.8万、自然增-5.46‰", "户籍人口-1.4万"],
        ["开放", "东北开放门户/港口3亿吨/集装箱548万箱", "软件/期货/海洋"],
    ],
    widths=[2.2, 5.6, 8.2])
para(doc, "", size=10, space_after=2)

# ---- 一、研究口径与阅读方法 ----
heading1(doc, "一、研究口径与阅读方法")
heading2(doc, "1.1 本报告的核心底稿")
bullet(doc, "**2025年《政府工作报告》**：2025年1月在市十七届人大5次会议上作（目标：GDP 5.5%以上、财政3%、规上工业7%、固投5%左右）。")
bullet(doc, "**《2025年大连市国民经济和社会发展统计公报》**（市统计局2026-04）——GDP、工业、贸易、港口实数。")
bullet(doc, "**2026年大连市政府工作报告/复盘**——对2025执行追认。")
heading2(doc, "1.2 阅读方法：显性—数据—制度—长期")
para(doc, "每章按“**显性表述→同期数据→制度含义→长期影响**”四层展开。")
para(doc, "**关键判别**：数据优先。例2025年GDP目标5.5%、实际5.7%达标（破万亿）；规上工业目标7%、实际11.7%超预期。大连“工业/出口/海洋强，投资/地产弱”。")

# ---- 二、底盘 ----
heading1(doc, "二、先看大连的特殊底盘：东北开放门户、港口/石化、海洋经济与软件IT")
para(doc, "大连的地盘，取决于它作为“**东北开放门户+北方港口+石化基地+软件IT+海洋经济**”的特殊定位。它是东北唯一的副省级/计划单列市。")
bullet(doc, "**港口/开放门户**：大连港（港口货物吞吐量3.0亿吨、集装箱548.2万标箱+1.5%）；中国东北出海海口、中欧班列、自贸区/综合保税区；面向日本、韩国、远东/东北亚。")
bullet(doc, "**石化**：恒力石化、绿色石化集群（国家级）；PTA 1643.6万吨、聚酯+20.9%、乙烯146.6万吨；石化工业+8.9%。")
bullet(doc, "**装备/高端制造**：规上工业+11.7%（高技术+13.9%、装备+15.4%）；民用钢质船舶439.5万载重吨（+74.6%）、汽车30万辆（+29.5%）、铁路机车+20.1%。")
bullet(doc, "**软件IT/期货/海洋**：软件和服务外包（外包企业1420家、从业人员20.1万）；大连商品交易所期货占全国28.7%；海洋牧场32处、海洋经济总产值4500亿元（+5%）。")
para(doc, "**制度含义**：大连把“港口/开放、石化、装备/船舶、软件/期货、海洋”当核心资产，面向东北振兴与东北亚开放。高新技术制造+13.9%、海洋牧场居全国首位。")

# ---- 三、宏观错位 ----
heading1(doc, "三、最关键的宏观错位：GDP破万亿、工业/出口强，但投资/地产/人口弱")
para(doc, "2025年大连GDP 10002.1亿元、+5.7%（首次破万亿）（一产+3.6%、二产+7.7%、三产+4.8%）。表面看“向上突破”，但拆开看是“**工业/出口/海洋强、投资/地产/人口弱**”的错位：")
para(doc, "**强的部分**：规上工业+11.7%（高技术+13.9%、装备+15.4%、石化+8.9%）；利润560亿+21.3%；出口+10.6%；高新技术制造投资+114%。")
para(doc, "**弱的部分**：固定资产投资-17.0%（地产-30.1%、三产-22.8%）；商品房销售面积-12.4%；财政-3.3%；人口自然增-5.46‰（户籍-1.4万）。")
para(doc, "**核心错位一句话**：大连“工业/装备/船舶/出口/海洋强（规上工业+11.7%、船舶+74.6%），但投资/地产/人口弱”。2026年“稳投资/地产、强工业/港口/软件/海洋、补人口”是重点。")
table(doc,
    ["强引擎", "数据", "弱项", "数据"],
    [
        ["规上工业", "+11.7%", "固定资产投资", "-17.0%"],
        ["民用船舶/船舶", "+74.6%", "房地产开发投资", "-30.1%"],
        ["装备制造", "+15.4%", "商品房销售面积", "-12.4%"],
        ["出口", "+10.6%", "人口自然增长率", "-5.46‰"],
        ["高技术制造投资", "+114.0%", "地方财政收入", "-3.3%"],
    ],
    widths=[3.8, 3.0, 3.6, 3.0])
para(doc, "**错位结论**　大连的增长“强工业/港口/海洋、弱投资/地产/人口”。万亿之上，2026年“破万亿/强港口/软件/海洋、补投资/人口”是重点。")

# ---- 四、被忽视细节 ----
heading1(doc, "四、容易被忽视、但信号很重的细节（15条）")
details = [
    ("1", "GDP 10002.1亿、+5.7%（首次破万亿）", "万亿突破。"),
    ("2", "规上工业+11.7%（东北第一）", "工业强。"),
    ("3", "民用钢质船舶+74.6%、铁路机车+20.1%", "装备/船舶领跑。"),
    ("4", "高技术制造+13.9%、装备+15.4%", "新质生产力。"),
    ("5", "采矿业+72.7%、石化+8.9%", "石化/矿业。"),
    ("6", "利润560亿+21.3%、利税+21.2%", "工业盈利强。"),
    ("7", "出口+10.6%、高技术制造投资+114%", "出口/制造投资。"),
    ("8", "港口3亿吨、集装箱548万标箱+1.5%", "港口/开放门户。"),
    ("9", "软件外包企业1420家、从业20.1万", "软件IT/服务外包。"),
    ("10", "大商所期货占全国28.7%", "期货/金融。"),
    ("11", "海洋牧场32处、海洋经济破4500亿", "海洋经济。"),
    ("12", "常住607.8万、自然增-5.46‰、户籍-1.4万", "人口负增/流出。"),
    ("13", "收入52007元、+4.1%、农村+4.5%", "收入稳、农村快。"),
    ("14", "实际使用外资+43.8%（三产+58.9%）", "外资/开放。"),
    ("15", "CPI-0.1%、PPI承压", "物价/通缩。"),
]
for d in details:
    para(doc, "**%s**" % d[0], size=11)
    para(doc, d[1], size=10.5, color=GRAY)

    para(doc, "**备注**　这条“错位”在大连尤其鲜明：增长靠工业/港口/海洋/软件，但投资/地产/人口弱。2026年若地产/投资修复，增长可能从“工业/出口单极”走向“工业/海洋+投资”多极。万亿之上，这条细节正是大连最该盯的“变量”。", size=10.5)

# ---- 五、2025年GDP目标 vs 实际 对照 ----
heading1(doc, "五、2025年GDP目标 vs 实际：对照表")
t5 = [
    ["指标", "2025目标", "2025实际", "达标判定"],
    ["GDP增速", "5.5%以上", "+5.7%", "达标/破万亿"],
    ["规模以上工业", "7%", "+11.7%", "大幅超预期"],
    ["固定资产投资", "5%左右", "-17.0%", "大幅未达标"],
    ["一般公共预算收入", "3%", "-3.3%", "未达标"],
    ["社会消费品零售", "高于全国", "+2.1%", "基本达标"],
    ["粮食产量", "25亿斤以上", "136.3万吨(约27亿斤)", "达标"],
]
table(doc, t5[0], t5[1:], widths=[3.6, 3.0, 3.0, 3.6])
para(doc, "**对照结论**　目标“冲刺万亿”：GDP 5.7%达标（破万亿）、规上工业11.7%大幅超预期；但固投-17.0%（未达5%）、财政-3.3%（未达3%）。工业/出口/港口接住，投资/地产/财政是短板。破万亿是最大成果。")

# ---- 六、增速分项支撑 ----
heading1(doc, "六、2025年增长是谁撑起来的？（结构归因）")
para(doc, "GDP破万亿，是“**工业/装备/船舶/出口强、投资/地产弱**”的结构。把增长贡献拆开看：")
g6 = [
    ["引擎", "贡献/状态", "说明"],
    ["规上工业", "+11.7%", "高技术+13.9%、装备+15.4%"],
    ["船舶/铁路装备", "+74.6%/+20.1%", "民用船舶领跑"],
    ["石化", "+8.9%", "恒力/绿色石化集群"],
    ["出口/高技术投资", "+10.6%/+114%", "出口+高新制造"],
    ["港口/海洋", "3亿吨/海洋4500亿", "港口/海洋牧场32处"],
    ["消费/社零", "+2.1%", "家电+109.7%"],
    ["房地产开发", "-30.1%", "地产深度调整"],
    ["固定资产投资", "-17.0%", "投资拖累"],
    ["人口自然增长率", "-5.46‰", "人口约束"],
    ["财政", "-3.3%", "财税承压"],
]
table(doc, g6[0], g6[1:], widths=[3.6, 3.6, 6.0])
para(doc, "**一句话**　增长靠“工业+装备/船舶+出口+港口/海洋”，投资/地产是拖累。破万亿之后，2026年考验大连“能否让港口/软件/海洋进一步放大”。")

# ---- 七、预算与财政 ----
heading1(doc, "七、预算与财政的“含金量”")
para(doc, "2025年大连地方一般公共预算收入**749.5亿元、-3.3%**，其中税收**505.9亿元、+3.4%**。")
bullet(doc, "财政收入-3.3%、税收+3.4%，税收转正。")
bullet(doc, "民生支出969.1亿、+5.3%，占86.4%。")
bullet(doc, "财政“稳收+民生/科创/港口支出优先”。")

# ---- 八、民生底账 ----
heading1(doc, "八、民生底账：人口、收入与城乡")
para(doc, "2025年大连户籍人口**607.8万人、-1.4万**，自然率**-5.46‰**。居民可支配收入**52007元、+4.1%**，农村（+4.5%）快于城镇（+4.0%）。")
g8 = [
    ["指标", "数值", "信号"],
    ["户籍人口", "607.8万/-1.4万", "人口净流出"],
    ["自然增长率", "-5.46‰", "深度负增"],
    ["全体收入", "52007元/+4.1%", "收入稳"],
    ["农村收入", "29185元/+4.5%", "农村快于城镇"],
]
table(doc, g8[0], g8[1:], widths=[4.2, 4.0, 4.8])
para(doc, "**民生观察**：收入增速+4.1%不错、农村快于城镇；但人口自然负增、户籍流出（-1.4万）是约束。")

# ---- 九、城镇与农村 ----
heading1(doc, "九、城镇与农村：格局与均衡")
para(doc, "大连城镇化率较高（约80%），农村收入快于城镇。")
bullet(doc, "农村居民收入+4.5%、快于城镇+4.0%。")
bullet(doc, "金普新区/自贸、软件园聚集。")

# ---- 十、人口流入与流出 ----
heading1(doc, "十、人口流入与流出")
para(doc, "大连户籍607.8万、-1.4万，人口仍在净流出（东北人口外流共性）。")
para(doc, "未来看点：港口/软件/海洋/装备能否留人；若“万亿+东北亚开放”成势有望回流。")

# ---- 十一、物价与货币 ----
heading1(doc, "十一、物价与货币环境")
para(doc, "2025年大连CPI**-0.1%**、PPI承压，物价微降、工业品价格偏弱。")
para(doc, "物价偏弱与全国低通胀一致。2026年“扩内需、稳价格”是主线。")

# ---- 十二、区域一体化 ----
heading1(doc, "十二、区域一体化：大连在“东北开放门户+沿海/港口+大连都市圈”里的位置")
para(doc, "大连的核心战略坐标是“**东北开放门户+沿海/港口+大连都市圈**”。")
bullet(doc, "东北开放门户：面向日韩、远东、东北亚，大连自贸区。")
bullet(doc, "港口：大连港3亿吨、集装箱548万箱。")
bullet(doc, "海洋经济：海洋牧场32处、海洋4500亿。")
bullet(doc, "软件/期货：大连商品交易所、软件园。")
para(doc, "若“港口+石化+软件+海洋+开放”成势，大连是东北振兴与东北亚开放的“龙头”。")

# ---- 十三、五条主线 ----
heading1(doc, "十三、未来5—10年最值得观察的五条主线")
lines5 = [
    ("① 港口/开放门户", "港口、自贸区、东北亚物流能否壮大。"),
    ("② 石化/绿色石化", "恒力、绿色石化集群能否延伸。"),
    ("③ 装备/船舶/船舶装备", "船舶、装备、氢能能否成极。"),
    ("④ 软件/期货/海洋", "软件IT、期货、海洋经济。"),
    ("⑤ 人口/地产/内需", "人口回流、地产修复、扩内需。"),
]
for l in lines5:
    para(doc, "**%s**　%s" % l, space_after=6)

# ---- 十四、结论 ----
heading1(doc, "十四、最终结论：大连在“万亿GDP+港口/石化+软件/海洋”里的增长逻辑")
para(doc, "大连的2025年，本质上是“工业/港口/海洋/出口进阶突破万亿，投资/地产/人口弱”的答卷：GDP10002.1亿、+5.7%、规上工业+11.7%、高技术+13.9%、出口船舶+74.6%、港口3亿吨，但固投-17.0%、地产-30.1%、财政-3.3%、自然增-5.46‰。")
para(doc, "只要港口/石化、装备/船舶、软件/期货、海洋经济持续，大连就站在“东北开放门户”的位置；如果投资/地产/人口持续偏弱，大连需承受“工业强、内需/楼市/人口弱”的挑战。")
para(doc, "最稳观察信号：**一盯港口/石化（底座）、二盯装备/船舶（制造）、三盯软件/期货/海洋（新动能）、四盯投资/地产（约束）、五盯人口/民生（长期）。**大连，是“东北开放+万亿港口”的新样本。")

# ---- 附录A ----
heading1(doc, "附录A：主要资料来源与核验路径")
bullet(doc, "大连市2025年《政府工作报告》——目标来源。")
bullet(doc, "《大连市2025年国民经济和社会发展统计公报》（市统计局）——GDP、工业、外贸、港口人力。")
bullet(doc, "2026年大连市政府工作报告——2025复盘/港口/石化/软件。")
bullet(doc, "大连海关（外贸）、市财政（财政）。")
heading2(doc, "核验说明")
para(doc, "本报告所有增速、总量、占比均以统计公报/官方口径为基准，涉“港口/石化/海洋/软件”等以官方口径为准。")

# ---- 附录B ----
heading1(doc, "附录B：建议建立的年度跟踪仪表盘")
para(doc, "建议把下面10个‘测脉搏’指标做成年度跟踪表：")
dash = [
    ["#", "指标", "2025基值", "为什么重要"],
    ["1", "GDP增速/总量", "+5.7%/10002亿", "破万亿"],
    ["2", "规上工业增速", "+11.7%", "工业底盘"],
    ["3", "船舶/装备增速", "+74.6%/+15.4%", "高端制造"],
    ["4", "出口增速", "+10.6%", "出口"],
    ["5", "固定资产投资/地产", "-17.0%/-30.1%", "投资"],
    ["6", "港口集装箱", "548万标箱", "港口/开放"],
    ["7", "常住户籍/(自然增长)", "607.8万/-5.46‰", "人口与城市"],
    ["8", "地方财政收入", "-3.3%", "财政质量"],
    ["9", "海洋经济/软件外包", "4500亿/32亿美元", "海洋/软件"],
    ["10", "CPI", "-0.1%", "物价"],
]
table(doc, dash[0], dash[1:], widths=[0.8, 4.6, 3.4, 4.4])
para(doc, "把这10个指标连起来看，港口/石化、装备/海洋、软件（2/3/9）向上、人口/地产（5/7）修复，都说明大连在真正换挡。")

# ============ 保存 ============
out = "/Users/x/Desktop/content-prod-lab/reports/大连市_2025年政府工作报告_深度研究_2026-08-13.docx"
doc.save(out)
print("SAVED:", out)
print("PARAGRAPHS:", len(doc.paragraphs))
print("TABLES:", len(doc.tables))
