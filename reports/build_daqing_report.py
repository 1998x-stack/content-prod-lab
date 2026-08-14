# -*- coding: utf-8 -*-
"""Build 大庆市2025年政府工作报告 深度研究 DOCX, 参照地级市系列版式。"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DARK = RGBColor(0x1F, 0x2A, 0x44)
RED = RGBColor(0xB0, 0x1E, 0x2E)
GRAY = RGBColor(0x59, 0x60, 0x69)
BLUE = RGBColor(0x1F, 0x4E, 0x79)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
HEADER_FILL = "1F2A44"
K = "微软雅黑"


def set_run(run, size=11, bold=False, color=DARK, font=K):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    run.font.name = font
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), font)
    rFonts.set(qn("w:hAnsi"), font)
    rFonts.set(qn("w:eastAsia"), font)


def para(doc, text, size=11, bold=False, color=DARK, align=None,
         space_after=6, space_before=0, font=K):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        set_run(r, size=size, bold=bold or (i % 2 == 1), color=color, font=font)
    return p


def heading1(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(16)
    pf.space_after = Pt(8)
    r = p.add_run(text)
    set_run(r, size=16, bold=True, color=RED)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "B01E2E")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def heading2(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(10)
    pf.space_after = Pt(4)
    r = p.add_run(text)
    set_run(r, size=13, bold=True, color=BLUE)
    return p


def table(doc, headers, rows, widths=None, font_size=9.5):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_run(r, size=font_size, bold=True, color=WHITE)
        tcPr = hdr[i]._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), HEADER_FILL)
        tcPr.append(shd)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(val))
            set_run(r, size=font_size, color=DARK)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    return t


def quote_box(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(8)
    pf.left_indent = Pt(12)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), "B01E2E")
    pBdr.append(left)
    pPr.append(pBdr)
    r = p.add_run(text)
    set_run(r, size=11, color=DARK, font="楷体")
    return p


def bullet(doc, text, size=10.5):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.7)
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        set_run(r, size=size, bold=(i % 2 == 1), color=DARK)
    return p


# ================================================================= 文档主体
doc = Document()
sec = doc.sections[0]
sec.page_width = Cm(21)
sec.page_height = Cm(29.7)
sec.top_margin = Cm(2.4)
sec.bottom_margin = Cm(2.2)
sec.left_margin = Cm(2.5)
sec.right_margin = Cm(2.5)

st = doc.styles["Normal"]
st.font.name = K
st.font.size = Pt(11)
st.element.rPr.rFonts.set(qn("w:eastAsia"), K)


# ---- 封面 ----
para(doc, "大庆市2025年政府工作报告", size=24, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=60, space_after=4)
para(doc, "深度研究与\u201c容易被忽视的细节\u201d分析", size=16, bold=True, color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
para(doc, "从\u201c石油城、石化转型、能源装备、页岩油与新能源\u201d重新理解大庆", size=12, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)
para(doc, "\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501", color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
para(doc, "研究对象：2025年大庆市政府工作报告及同期财政、国民经济和社会发展统计资料、2026年官方复盘", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
para(doc, "标定日期：2026-08-14", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
doc.add_page_break()

# ---- 目录 ----
catalog = [
    "执行摘要",
    "一、研究口径与阅读方法",
    "二、先看大庆的特殊底盘：石油城、石化炼化、能源装备、页岩油与新能源",
    "三、最关键的宏观错位：GDP 2693.2亿/3.5%低于目标，财政暴增但工业弱、固投大跌、出口下行",
    "四、容易被忽视、但信号很重的细节（15条）",
    "五、2025年GDP目标 vs 实际：对照表",
    "六、2025年增长是谁撑起来的？（结构归因）",
    "七、预算与财政的\u201c含金量\u201d",
    "八、民生底账：人口、收入与城乡",
    "九、城镇与农村：格局与均衡",
    "十、人口流入与流出",
    "十一、物价与货币环境",
    "十二、区域一体化：大庆在\u201c哈大齐工业走廊+东北振兴+中俄合作\u201d里的位置",
    "十三、未来5—10年最值得观察的五条主线",
    "十四、最终结论：大庆在\u201c油城转型+石化炼化+新能源\u201d里的增长逻辑",
    "附录A：主要资料来源与核验路径",
    "附录B：建议建立的年度跟踪仪表盘",
]
para(doc, "目　录", size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=10, space_after=10)
for item in catalog:
    para(doc, item, size=12, space_after=4)
doc.add_page_break()

# ---- 执行摘要 ----
heading1(doc, "执行摘要")
para(doc, "**核心判断**　2025年大庆最显著的是\u201cGDP 2693.2亿元、增长3.5%（低于5%目标）、人均GDP 10.16万元\u201d、\u201c规上工业+1.5%（原油产量2911.9万吨、石化炼化1474万吨+1.1%）\u201d、\u201c地方财政总收入648.4亿/+16.0%、一般公共预算收入206.1亿/+22.8%\u201d、\u201c但固投-19.8%、出口-44.7%、常住人口降到263.7万\u201d。这说明大庆在\u201c石油石化+页岩油+新能源\u201d的转型中，**财政高增但工业、投资、外贸偏弱，人口持续流失**是突出矛盾。")
para(doc, "把2025年目标（GDP+5%左右/规上+5%/固投+6%/社零+5%/财收+5%）、2025年统计（GDP+3.5%远低于目标、规上+1.5%、固投-19.8%、财收+22.8%远超）、趋势一起看，大庆是\u201c油城转型\u201d路径：**石油（原油约2911.9万吨）、石化炼化（1474万吨）、能源装备、页岩油、新能源**是支柱；2025年财政因中石油集团税收上划等大增，但工业、投资、外贸走弱。总量2693.2亿居黑龙江第2。")
para(doc, "最容易记住的一句话：**大庆是\u201c中国石油长子、共和国明珠\u201d，靠\u201c原油+石化+新能源\u201d增长，正从\u201c油城\u201d转向\u201c综合性资源转型城市\u201d。**观察大庆，与其只看\u201cGDP 2693亿\u201d，不如看\u201c原油2911.9万吨、石化炼化1474万吨+1.1%、页岩油突破百万吨、财政+22.8%、城镇化率75.62%\u201d。")

# ---- 一、研究口径与阅读方法 ----
heading1(doc, "一、研究口径与阅读方法")
para(doc, "本报告以\u201c大庆市政府工作报告（2025年，李岩松作）\u201d为起点，把\u201c2025年GDP目标（5%左右）\u201d与\u201c官方2025年GDP（2693.2亿元/+3.5%）\u201d并置对照，并用\u201c2025年大庆市统计公报\u201d和\u201c2026年政府工作报告复盘\u201d作为横向核验。")
para(doc, "口径提示：\u201cGDP、规上工业增加值增速\u201d按可比价（实际增速）；\u201c投资、消费、进出口、财政收入\u201d按现价（名义增速）。涉及\u201c常住人口\u201d用统计公报常住口径（263.7万人），城镇化率用官方公布值（75.62%）。")
para(doc, "指标体系（与研究口径一致）：核心看五个象限——**总量与增速（GDP）、动能（工业+投资+消费+外贸）、财政质量（一般公共预算/税收）、民生底账（收入/就业/人口）、转型主线（油城→新能源/页岩油）**。")
para(doc, "特别提示（不吃老本）：大庆是\u201c因油而生\u201d的资源型城市，过去靠中石油带动；2025年GDP 2693亿但增速仅3.5%，**财政却大增22.8%**——这背后有财政口径/转移支付因素，本报告会特别辨析\u201c财政高增≠经济高增\u201d。真正要盯的是**产业转型、人口去留、非油经济占比**。")
# ---- 二、先看大庆的特殊底盘 ----
heading1(doc, "二、先看大庆的特殊底盘：石油城、石化炼化、能源装备、页岩油与新能源")
para(doc, "大庆是中国最大的油田城市、\u201c大庆精神/铁人精神\u201d发源地，位于黑龙江省西部、松嫩平原。2025年GDP 2693.2亿元、常住人口263.7万，是黑龙江第2大经济体，也是中石油体系内炼化与石化重镇。")
para(doc, "四个底盘名词，先立框架：")
bullet(doc, "**石油城/原油**　2025年原油产量2911.9万吨（连续多年保持高位但略降），大庆油田开创中国工业文明\u201c石油长子\u201d地位。")
bullet(doc, "**石化炼化**　2025年加工原油1474.1万吨（+1.1%），汽油373.4万吨、柴油387.3万吨；大庆石化、大庆炼化是东北炼化枢纽，乙烯产能\u201c十四五\u201d+48.3%。")
bullet(doc, "**能源装备/页岩油**　油田开采辅助业+7.6%、页岩油年产突破100万吨，成为接续开采新战场。")
bullet(doc, "**新能源**　新能源和可再生能源装机606万千瓦；数字经济规模1200亿；沃尔沃大庆工厂累计生产整车64万辆，\u201c油城转智造\u201d。")
para(doc, "大庆的底色是\u201c靠油吃饭、以石化立市\u201d：理解大庆，先理解它\u201c业因油建、城随油兴\u201d的资源禀赋与\u201c必须转型\u201d的现实压力。")

# ---- 三、最关键的宏观错位 ----
heading1(doc, "三、最关键的宏观错位：GDP 2693.2亿/3.5%低于目标，财政暴增但工业弱、固投大跌、出口下行")
para(doc, "2025年大庆最需要辨析的一组\u201c错位\u201d：**GDP 3.5%未达5%目标（差1.5个百分点）、规上工业仅+1.5%、固投-19.8%、出口-44.7%、进出口-8.8%**；但**一般公共预算收入+22.8%大幅超额**。")
para(doc, "为什么\u201c财政暴增、经济放缓、投资外贸大跌\u201d同时出现？三个解释：")
para(doc, "**其一，财政口径的\u201c虚高\u201d**　一般公共预算收入206.1亿/+22.8%、地方财政总收入648.4亿/+16.0%，但税收仅127.4亿/+3.1%——**增收主要来自非税与上划财政（中石油税收/国企利润上缴），不等于经济活力**。")
para(doc, "**其二，工业在低基数上微增**　规上工业+1.5%，其中采油辅助业+7.6%（石化-5.9%、装备-21.2%、电力热力-7.2%）；因油价、炼化景气与修井影响，工业整体只有1.5%的低个位数增长。")
para(doc, "**其三，投资出口双弱、人口流失**　固投-19.8%（制造业-29.5%、基础设施-30.3%、外资-48.7%）；进出口-8.8%、出口-44.7%；常住人口263.7万、比上年-2.6万——\u201c产业投资不足+外贸下行+人口流出\u201d三重承压。")
para(doc, "小结：大庆2025年是\u201c**稳财政、弱工业、弱投资&外贸**\u201d的一年：财政靠国企上划\u201c表面好看\u201d，但**工业、投资、出口、人口**是短板，转型压力大。")

# ---- 四、容易被忽视、但信号很重的细节（15条） ----
heading1(doc, "四、容易被忽视、但信号很重的细节（15条）")
bullet(doc, "**1.财政+22.8%但税收仅+3.1%**　增收靠非税/上划财政；税收占比61.8%——\u201c财政高增\u201d含金量需谨慎。")
bullet(doc, "**2.规上工业+1.5%，石化-5.9%、装备-21.2%**　工业仅微增，装备、石化、电力都在降。")
bullet(doc, "**3.原油2911.9万吨、石化炼化1474.1万吨**　油田稳产、炼化微升+1.1%是工业\u201c压舱石\u201d。")
bullet(doc, "**4.页岩油年产突破100万吨**　非常规油气开采（页岩油）是2025年最大新增长极。")
bullet(doc, "**5.新能源装机606万千瓦**　风光+新能源发电扩容，新能源产业加速。")
bullet(doc, "**6.沃尔沃累计64万辆、汽车产业**\u201c油城转智造\u201d的样板（2024年留庆就业毕业生+20%）。")
bullet(doc, "**7.民营经济活**\u201c市场主体累计33.1万户/+3.5%\u201d，虽然民营规模小于沿海，但活力在长。")
bullet(doc, "**8.社零+4.2%（640.2亿）**消费企稳、以旧换新+16.6%带动家电/汽车；但总量有限。")
bullet(doc, "**9.进出口-8.8%、出口-44.7%**\u201c进口多（1379亿）、出口少（37.8亿）\u201d，贸易逆差巨大，出口依赖进口中间品。")
bullet(doc, "**10.居民收入：城镇49430元/+4%  农村21888元/+6%，城乡比2.26**\u201c农村工资性收入快于城镇\u201d但基数低。")
bullet(doc, "**11.常住人口263.7万、-2.6万**\u201c人口净流出趋势未止\u201d，常住城镇化率75.62%（+0.66pct）。")
bullet(doc, "**12.CPI+0.1%**\u201c物价基本平稳，与全国低通胀一致；食品烟酒-0.7%。\u201d")
bullet(doc, "**13.农村居民收入+6%快于城镇+4%**\u201c脱贫人口年人均1.9万、城乡低保标准城镇800元/月、农村700元/月。\u201d")
bullet(doc, "**14.高等教育人口**\u201c驻庆高校应届毕业生留庆率18.6%（创历史新高）、就业城镇化。\u201d")
bullet(doc, "**15.“智改数转网联”、数字经济1200亿**“159户企业完成智改数转、数字经济“十四五”累计1200亿元——大庆在“油城求变”。”")
# ---- 五、2025年GDP目标 vs 实际：对照表 ----
heading1(doc, "五、2025年GDP目标 vs 实际：对照表")
table(doc,
    ["指标", "2025年目标", "2025年实际", "达成评价"],
    [
        ["地区生产总值(GDP)", "增长5%左右", "2693.2亿元/3.5%", "未达成，差1.5pct"],
        ["规上工业增加值", "增长5%左右", "+1.5%", "未达成，差3.5pct"],
        ["固定资产投资", "增长6%左右", "-19.8%", "大幅未达成"],
        ["社会消费品零售总额", "增长5%左右", "640.2亿/4.2%", "基本达到"],
        ["一般公共预算收入", "增长5%左右", "206.1亿/+22.8%", "大幅超额"],
        ["进出口", "——", "1417.7亿/-8.8%", "负增长"],
        ["城乡居民收入", "与经济增长同步", "城镇+4%/农村+6%", "总体同步"],
    ],
)
para(doc, "注：GDP、规上工业按可比价；投资、消费、财收按现价。GDP 3.5%、规上+1.5%均低于目标；**财政+22.8%靠非税/上划财政\u201c虚高\u201d**，工业、投资、外贸均未达目标。")
para(doc, "拆读：**财政（22.8%）是唯一大幅超额项**，**工业/投资/外贸全面走弱**；\u201cGDP目标5%\u201d实际3.5%、\u201c固投目标6%\u201d实际-19.8%——\u201c进取得分、现实失分\u201d，是资源型城市转型阵痛期样本。")

# ---- 六、2025年增长是谁撑起来的？（结构归因） ----
heading1(doc, "六、2025年增长是谁撑起来的？（结构归因）")
para(doc, "把大庆GDP的3.5%拆开：第一产业244.8亿（+3.1%）、第二产业1289.3亿（+1.0%）、第三产业1159.1亿（+6.1%），三次产业结构比9.1：47.9：43.0。**第三产业（服务业）是唯一强项，第二产业（工业+建筑）是拖累项**。")
para(doc, "2026年大庆强调\u201c守好石油基本盘、做大做强接续产业\u201d，围绕**油气稳产、页岩油、高端制造、新能源**部署，核心抓手是\u201c油经济+非油经济\u201d双轮。")
para(doc, "**第二产业（石油石化+制造）**：原油2911.9万吨、石化炼化1474.1万吨，工业增加值+1.5%低位；石化-5.9%、装备-21.2%、电力热力-7.2%，制造业投资-29.5%——工业转型阵痛明显。")
para(doc, "**第三产业（服务业）**：旅游+13.3%（游客花费+15%）、票房+51%、社零+4.2%；物流、零售、文旅是增长点。")
para(doc, "**财政与国企（上划）**：财政+22.8%、实际利用内资+39.1%，但\u201c国企利润上缴、中石油税收\u201d贡献大——\u201c财政强、实体经济弱\u201d的结构观察。")
para(doc, "一句话归因：**2025年大庆增长\u201c靠财政+服务业\u201d稳住基本盘，**工业、投资、外贸、人口\u201d四大短板拖累；转型在页岩油、新能源、数字经济中缓慢推进。**")

# ---- 七、预算与财政的\u201c含金量\u201d ----
heading1(doc, "七、预算与财政的\u201c含金量\u201d")
para(doc, "2025年大庆**一般公共预算收入206.1亿元（+22.8%）**，其中税收收入127.4亿元（+3.1%）；地方财政总收入648.4亿元（+16.0%）。**财政增速远超GDP，但税收仅+3.1%——含金量需打折**。")
bullet(doc, "增收结构：税收+3.1%、非税大幅高增（上划、国企利润），税收占一般公共预算收入比例约61.8%——非税依赖偏高。")
bullet(doc, "支出民生：社会保障、教育、卫生仍是主投向；城镇低保800元/月、农村低保600元/月，救助金3.3亿惠及6.16万人。")
bullet(doc, "金融支撑：全市存贷款规模较大，贷款稳定支持石化、装备、新能源等主导产业。")
para(doc, "**财政含金量小结**：财收22.8%高增需警惕\u201c虚高\u201d（非税/上划为主），真实税收仅+3.1%；**\u201c增收靠国企、实体税收弱\u201d**是大庆财政的底色；落到产业、民生、转型的\u201c实打实投入\u201d仍需观察。")

# ---- 八、民生底账：人口、收入与城乡 ----
heading1(doc, "八、民生底账：人口、收入与城乡")
para(doc, "2025年大庆**城镇居民人均可支配收入49430元（+4.0%）、农村21888元（+6.0%）**，城乡收入比约2.26。就业总体平稳：城镇新增就业3.3万人（完成省定任务123%），发放创业贷款3.1亿元。")
para(doc, "人口画像：**常住人口263.7万、比上年-2.6万**，**人口净流出趋势未止**；但城镇化率75.62%（+0.66pct）仍高，驻庆高校应届毕业生留庆率18.6%（创历史新高）显示青年人才留存改善。")
para(doc, "民生投入：城乡居民基本医保、养老保险参保率均超95%；发放育儿补贴2307万元；新增优质学位9692个；农村低保提标。")
para(doc, "人口战略判断：\u201c兼顾保民生、稳就业、引人才\u201d，但人口微降、老龄化与城镇化的双轨，仍是\u201c城市可持续\u201d的隐变量。")

# ---- 九、城镇与农村：格局与均衡 ----
heading1(doc, "九、城镇与农村：格局与均衡")
para(doc, "大庆常住城镇化率75.62%（居黑龙江前列），城乡格局较均衡；农村居民收入增速（+6%）快于城镇（+4%），但**城乡收入差距仍大（约2.26倍）**，共同富裕任重道远。")
para(doc, "农业底盘：**粮食产量超95亿斤（实现\u201c27连丰\u201d）**、高标准农田162万亩；大庆是黑龙江重要粮食/生猪生产地，\u201c农业稳、粮仓牢是底盘\u201d。")
para(doc, "一句话：\u201c大庆城镇化率高、农业产粮稳、城乡收入差收窄\u201d，但**农村基数低、人口外流**仍是制约农村发展的真问题。")

# ---- 十、人口流入与流出 ----
heading1(doc, "十、人口流入与流出")
para(doc, "大庆常住人口263.7万（-2.6万），**延续净流出趋势**，是东北资源型城市\u201c人口收缩\u201d的典型样本；但**城镇人口199.4万、常住城镇化率75.62%（+0.66pct）**显示人口往中心城区集中。")
para(doc, "结构观察：**老龄化+少子化**叠加\u201c留庆就业率\u201d的上升（驻庆毕业生留庆率18.6%创历史新高），说明有高等教育人才留在本地、但整体人口基数在降。")
para(doc, "2026年目标：稳住人口、导入青年产业人才、降生育成本（育儿补贴2307万）——大庆把\u201c人口\u201d作为转型关键变量。")

# ---- 十一、物价与货币环境 ----
heading1(doc, "十一、物价与货币环境")
para(doc, "2025年大庆**全年城市CPI同比+0.1%**（居民消费价格基本平稳），其中食品烟酒-0.7%、衣着+2.2%、交通通信-3%——**物价低位运行、需求偏弱**，与\u201c消费偏弱、人口流失\u201d相互印证。")
bullet(doc, "信贷：全市金融机构存贷款平稳，贷款更多投向原油/石化、装备、新能源等主导产业与基础设施。")
bullet(doc, "利率/流动：稳信用支持下，重大项目融资、设备更新（技改投资+30%）等有资金支撑。")
para(doc, "货币环境判断：**低通胀、宽信用**并存，\u201c资金供给尚可、物价走平\u201d；大庆需\u201c扩内需、稳投资、促消费\u201d，尤其把\u201c设备更新+技改+制造业投资\u201d做实。")

# ---- 十二、区域一体化：大庆的位置 ----
heading1(doc, "十二、区域一体化：大庆在\u201c哈大齐工业走廊+东北振兴+中俄合作\u201d里的位置")
para(doc, "大庆地处\u201c哈大齐（哈尔滨-大庆-齐齐哈尔）工业走廊\u201d，是黑龙江资源重镇，承担**东北振兴、能源保障、中俄能源合作**多重角色。")
bullet(doc, "**能源走廊**　大庆油田+大庆石化/炼化，支撑国家油气安全；页岩油、天然气（61亿m³）配合\u201c油化并举\u201d。")
bullet(doc, "**哈大齐协同**　哈大齐工业走廊串联装备/石化；大庆承接哈尔滨科技与齐齐哈尔装备的联动。")
bullet(doc, "**中俄合作**　黑龙江对俄经贸、口岸；大庆\u201c进口原油-炼化-出口成品\u201d与俄罗斯油气合作紧密。")
para(doc, "一句话：**大庆在\u201c哈大齐走廊+东北振兴+中俄能源\u201d里，最核心的定位是\u201c国家能源安全\u2019压舱石＋工业走廊节点\u201d**——区位/能源禀赋是大庆区别于普通地级市的最大底色。")

# ---- 十三、未来5—10年最值得观察的五条主线 ----
heading1(doc, "十三、未来5—10年最值得观察的五条主线")
bullet(doc, "**主线一：油城转型（页岩油/页岩气）**\u201c非常规油气（页岩油破100万吨）\u201d能否接续稳产，决定大庆\u201c石油长子\u201d的基本盘。")
bullet(doc, "**主线二：石化炼化升级**\u201c中石油炼化一体化、乙烯扩能（+48.3%）\u201d能否从\u201c炼油\u201d走向\u201c化工新材料\u201d。")
bullet(doc, "**主线三：新能源与绿色转型**\u201c新能源装机606万千瓦、风光\u201d能否撑起\u201c非油经济\u201d的第二曲线。")
bullet(doc, "**主线四：先进制造/智造（沃尔沃/装备）**\u201c油城转智造、高端装备、汽车（64万辆）\u201d对冲装备业的下降。")
bullet(doc, "**主线五：人口与城市可持续**\u201c人口净流出、老龄化\u201d能否被\u201c青年留庆、产业导入\u201d扭转——这是大庆长治的百年课题。")

# ---- 十四、最终结论 ----
heading1(doc, "十四、最终结论：大庆在\u201c油城转型+石化炼化+新能源\u201d里的增长逻辑")
para(doc, "把2025年闭环打穿：**大庆是\u201c石油+石化为基、油城转型进行中\u201d的资源型城市**：GDP 2693.2亿/+3.5%、规上工业+1.5%、财政+22.8%；服务业（旅游/人口）与财政是2025年的\u201c托底\u201d，但工业、投资、外贸、人口四大短板突出。")
para(doc, "大庆不是\u201c只有一座油田\u201d——它在炼化、页岩油、新能源、智造（沃尔沃）、数字经济逐步延伸；同时**投资、外贸、人口承压**、转型爬坡。")
para(doc, "一句话结论：**大庆是\u201c石油长子、东北石化重镇、转型攻坚\u201d的资源型城市；观察它先看\u201c原油炼化+页岩油+新能源\u201d，再看\u201c工业投资外贸人口四大短板\u201d。**它是\u201c财政强、服务业稳、工业弱、人口减\u201d的东北油城样本。")

# ---- 附录A ----
heading1(doc, "附录A：主要资料来源与核验路径")
bullet(doc, "《2025年大庆市政府工作报告》（2025年1月6日，李岩松作，2025年预期目标、2024年回顾）")
bullet(doc, "《2025年大庆市国民经济和社会发展统计公报》（大庆市统计局，2026年发布，2025年实际数据）")
bullet(doc, "《2026年大庆市政府工作报告》（大庆市政府，2026年1月，2025年复盘+2026年目标）")
bullet(doc, "大庆市人民政府官网、大庆市统计局官方URL（daqing.gov.cn）")

# ---- 附录B ----
heading1(doc, "附录B：建议建立的年度跟踪仪表盘")
bullet(doc, "GDP总量/增速 vs 全年目标（可比价口径）。")
bullet(doc, "规上工业增加值及分行业（原油/石化/装备/页岩油/新能源）增速。")
bullet(doc, "原油、天然气、页岩油产量；石化加工量；成品油/化工品价。")
bullet(doc, "固定资产投资/制造业/基础设施/房地产投资增速。")
bullet(doc, "社会消费品零售总额、旅游、票房、以旧换新补贴。")
bullet(doc, "外贸（人民币）、出口/进口、原油-成品油贸易。")
bullet(doc, "一般公共预算收入、税收/非税占比、税收增速。")
bullet(doc, "常住人口、城镇化率、自然增长率、留庆毕业生率。")
bullet(doc, "CPI、金融存贷款。")
bullet(doc, "新能源装机、页岩油/页岩气、数字经济规模、沃尔沃产量。")

# ---------------- 保存 ----------------
out = "/Users/x/Desktop/content-prod-lab/reports/大庆市_2025年政府工作报告_深度研究_2026-08-14.docx"
doc.save(out)
print("SAVED OK: 大庆市", out)
