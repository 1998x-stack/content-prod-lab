# -*- coding: utf-8 -*-
"""Build 济宁市2025年政府工作报告 深度研究 DOCX, 参照地级市系列版式。"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DARK = RGBColor(0x1F, 0x2A, 0x44)
RED = RGBColor(0xB0, 0x1E, 0x2E)
GRAY = RGBColor(0x59, 0x60, 0x69)
BLUE = RGBColor(0x1F, 0x4E, 0x79)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
HEADER_FILL = "1F2A44"
K = "微软雅黑"


def set_run(run, size=11, bold=False, color=DARK, font=K):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    run.font.name = font
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), font)
    rFonts.set(qn("w:hAnsi"), font)
    rFonts.set(qn("w:eastAsia"), font)


def para(doc, text, size=11, bold=False, color=DARK, align=None,
         space_after=6, space_before=0, font=K):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        set_run(r, size=size, bold=bold or (i % 2 == 1), color=color, font=font)
    return p


def heading1(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(16)
    pf.space_after = Pt(8)
    r = p.add_run(text)
    set_run(r, size=16, bold=True, color=RED)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "B01E2E")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def heading2(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(10)
    pf.space_after = Pt(4)
    r = p.add_run(text)
    set_run(r, size=13, bold=True, color=BLUE)
    return p


def table(doc, headers, rows, widths=None, font_size=9.5):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_run(r, size=font_size, bold=True, color=WHITE)
        tcPr = hdr[i]._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), HEADER_FILL)
        tcPr.append(shd)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(val))
            set_run(r, size=font_size, color=DARK)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    return t


def quote_box(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(8)
    pf.left_indent = Pt(12)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), "B01E2E")
    pBdr.append(left)
    pPr.append(pBdr)
    r = p.add_run(text)
    set_run(r, size=11, color=DARK, font="楷体")
    return p


def bullet(doc, text, size=10.5):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.7)
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        set_run(r, size=size, bold=(i % 2 == 1), color=DARK)
    return p


# ================================================================= 文档主体
doc = Document()
sec = doc.sections[0]
sec.page_width = Cm(21)
sec.page_height = Cm(29.7)
sec.top_margin = Cm(2.4)
sec.bottom_margin = Cm(2.2)
sec.left_margin = Cm(2.5)
sec.right_margin = Cm(2.5)

st = doc.styles["Normal"]
st.font.name = K
st.font.size = Pt(11)
st.element.rPr.rFonts.set(qn("w:eastAsia"), K)


# ---- 封面 ----
para(doc, "聊城市2025年政府工作报告", size=24, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=60, space_after=4)
para(doc, "深度研究与\u201c容易被忽视的细节\u201d分析", size=16, bold=True, color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
para(doc, "从\u201c江北水城、钢铁铜铝（信发/祥光）、新能源汽车、轴承/化工、农业\u201d重新理解聊城", size=12, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)
para(doc, "\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501", color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
para(doc, "研究对象：2025年聊城市政府工作报告及同期财政、国民经济和社会发展统计资料、2026年官方复盘", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
para(doc, "标定日期：2026-08-14", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
doc.add_page_break()

# ---- 目录 ----
catalog = [
    "执行摘要",
    "一、研究口径与阅读方法",
    "二、先看聊城的特别底盘：江北水城、钢铁铜铝、新能源汽车、农业强市",
    "三、最关键的宏观错位：工业外贸强（进出口+15.5%），但固投+0.3%、地产-16.6%偏弱",
    "四、容易被忽视、但信号很重的细节（15条）",
    "五、2025年GDP目标 vs 实际：对照表",
    "六、2025年增长是谁撑起来的？（结构归因）",
    "七、预算与财政的\u201c含金量\u201d",
    "八、民生底账：人口、收入与城乡",
    "九、城镇与农村：格局与均衡",
    "十、人口流入与流出",
    "十一、物价与货币环境",
    "十二、区域一体化：聊城在济南都市圈、中原经济区、黄河流域\u201c三圈\u201d里的位置",
    "十三、未来5—10年最值得观察的五条主线",
    "十四、最终结论：聊城在\u201c江北水城+制造贸易+农业强市\u201d里的增长逻辑",
    "附录A：主要资料来源与核验路径",
    "附录B：建议建立的年度跟踪仪表盘",
]
para(doc, "目　录", size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=10, space_after=10)
for item in catalog:
    para(doc, item, size=12, space_after=4)
doc.add_page_break()

# ---- 执行摘要 ----
heading1(doc, "执行摘要")
para(doc, "**核心判断**　2025年聊城最显著的是\u201cGDP 3304.15亿/+5.6%、规上工业+7.7%（金属/化工/汽车强）、进出口728.07亿/+15.5%（进口+33.1%）、农林牧渔总产值854.67亿\u201d、\u201c但固定资产投资+0.3%、房地产开发-16.6%、CPI-0.3%、社零+5.9%\u201d。这说明聊城在\u201c江北水城+工业制造+农业强市\u201d中，**工业外贸农业强、投资地产弱**。")
para(doc, "把2025年目标（GDP+5.5%、财收+3.5%、固投+6%）、2025年实际（GDP+5.6%、规上+7.7%、固投+0.3%、财收+3.0%）趋势看，聊城是\u201c制造+农业+贸易\u201d路径：**钢铁（钢材）+有色金属（铜铝）+新能源汽车/汽车制造、农机、轴承、农产品加工（食品）**是支柱；进出口破700亿。")
para(doc, "最容易记住的一句话：**聊城是\u201c江北水城、中国轴承之城、钢铁铜铝强市、农业大市\u201d，靠\u201c金属制造+外贸+食品农业\u201d增长。**观察聊城，与其只看\u201cGDP 3304亿\u201d，不如看\u201c规上+7.7%、进出口破700亿（+15.5%）、金属制品+43.1%、新能源汽车+75.9%、农业强市\u201d。")
# ---- 一、研究口径与阅读方法 ----
heading1(doc, "一、研究口径与阅读方法")
para(doc, "本报告以\u201c聊城市政府工作报告（2025年2月）\u201d为起点，把\u201c2025年GDP目标（5.5%）\u201d与\u201c官方2025年（3304.15亿/+5.6%）\u201d并置对照，用\u201c2025年聊城市统计公报\u201d和\u201c2026年政府工作报告\u201d横向核验。")
para(doc, "口径提示：\u201cGDP、规上工业增速\u201d按可比价（实际）；\u201c投资、消费、进出口、财政\u201d按现价（名义）。\u201c常住人口\u201d用统计公报（575.86万）、城镇化率58.34%。")
para(doc, "指标体系（与研究口径一致）：**总量与增速、产业动能（金属/农业/汽车）、外贸、财政、民生与人口**。")
para(doc, "特别提示（不吃老本）：聊城2024年GDP~3129亿（+5.5%）、2025年3304.15亿/+5.6%；它不是\u201c只有江北水城\u201d——**钢铁（信发）、有色金属（祥光/铜铝）、汽车（时风农机）、轴承（临清）、农产品加工、轴承\u201d才是真正底色；装备制造+22.5%。")
# ---- 二、先看聊城的特别底盘 ----
heading1(doc, "二、先看聊城的特别底盘：江北水城、钢铁铜铝、新能源汽车、农业强市")
para(doc, "聊城地处山东省西部、黄河北岸、冀鲁豫三省交界，是**江北水城（京杭大运河/东昌湖）、中国轴承之城、钢铁铜铝之都、农业大市**；以\u201c信发集团、祥光铜业、时风农机制造\u201d著称。2025年GDP 3304.15亿、常住575.86万、城镇化率58.34%、山东中西部。")
para(doc, "五个底盘名词，先立框架：")
bullet(doc, "**钢铁/有色金属（铜铝）**　信发（铝/电力/化工）、祥光铜（铜加工）、十种有色金属207.4万吨——\u201c金属制造\u201d。")
bullet(doc, "**汽车/农机制造**　时风（低速载货汽车）、新能源汽车（+75.9%）、轴承（临清轴承）、机械——\u201c装备制造+22.5%\u201d。")
bullet(doc, "**化工/建材**　化学原料及制品+12.9%、烧碱/尿素、建材——\u201c绿色化工\u201d。")
bullet(doc, "**农业强市**　粮食589万吨、农产品加工（食品）、农牧渔总产值854.67亿、蔬菜/肉蛋——**农业大市\u201d。")
bullet(doc, "**江北水城/现代服务业**　运河文化、文旅（东昌湖）、商贸物流、医养——\u201c服务业+旅游\u201d。")
para(doc, "这五根（金属+汽车农机+化工+农业+水城服务）构成聊城独特底盘：**左手重工业（钢铁铜铝），右手农业+水城文化**。理解聊城，先理解\u201c江北水城、轴承之城、农业强市\u201d。")
# ---- 三、最关键的宏观错位 ----
heading1(doc, "三、最关键的宏观错位：工业外贸强（进出口+15.5%），但固投+0.3%、地产-16.6%偏弱")
para(doc, "2025年聊城最需要辨析的一组\u201c错位\u201d：**规上工业+7.7%（金属/化工/汽车强）、进出口+15.5%（进口+33.1%）、装备制造+22.5%、新能源汽车+75.9%强，但固定资产投资+0.3%、房地产开发-16.6%、社零+5.9%、财收+3.0%（低3.5%）**。")
para(doc, "为什么\u201c工业/外贸这么强\u201d，投资与地产却偏弱？三解释：")
para(doc, "**其一，工业/外贸/农业旺、量大**　规上+7.7%（金属制品+43.1%、化原+12.9%、汽车+27.2%）、进出口+15.5%、农林牧渔+4.5%——\u201c制造外贸强\u201d。")
para(doc, "**其二，投资/地产偏弱**　固投+0.3%、房地产-16.6%/销售-7.2%、工业技改+18.6%但总量有限——\u201c地产拖累投资\u201d。")
para(doc, "**其三，消费/物价温**　社零+5.9%（金银珠宝+75%）、但CPI-0.3%、汽车/家电-——\u201c消费修复、物价低\u201d。")
para(doc, "小结：聊城2025年是\u201c**工业外贸农业强、投资地产弱**\u201d：金属/汽车/化工、出口强，地产、CPI弱。")

# ---- 四、容易被忽视、但信号很重的细节（15条） ----
heading1(doc, "四、容易被忽视、但信号很重的细节（15条）")
bullet(doc, "**1.规上工业+7.7%（金属制品+43.1%/汽车+27.2%）**\u201c聊城制造强。\u201d")
bullet(doc, "**2.进出口728.07亿/+15.5%（进口+33.1%/加工贸易+1.9倍）**\u201c外贸强、进口原料旺。\u201d")
bullet(doc, "**3.装备制造业+22.5%、高新技术占比58.4%**\u201c新质制造。\u201d")
bullet(doc, "**4.新能源汽车+75.9%（时风/农机制造）**\u201c新能源车、农机强。\u201d")
bullet(doc, "**5.钢铁/十种有色金属207.4万吨（信发/祥光）**\u201c金属制造。\u201d")
bullet(doc, "**6.农业：粮食589万吨、农林牧渔854.67亿、蔬菜/肉蛋**\u201c农业强市。\u201d")
bullet(doc, "**7.民营企业进出口370.6亿（占51%）**\u201c民营外贸主力。\u201d")
bullet(doc, "**8.固定资产投资+0.3%（制造业+21.8%/民间+15.3%）**\u201c项目/民间投资旺、地产弱。\u201d")
bullet(doc, "**9.房地产开发293.7亿/-16.6%、销售-7.2%**\u201c地产深度调整。\u201d")
bullet(doc, "**10.一般公共预算收入264.80亿/+3.0%（税收占59.7%）**\u201c财收稳增。\u201d")
bullet(doc, "**11.社零1388.01亿/+5.9%、金银珠宝+75%**\u201c消费旺、结构。\u201d")
bullet(doc, "**12.R&D经费97.97亿/占3.09%、高企超1000家**\u201c创新、新质。\u201d")
bullet(doc, "**13.可再生能源装机691.5万千瓦/+32.1%、发电72.24亿千瓦时**\u201c绿色低碳。\u201d")
bullet(doc, "**14.居民收入30349元/+5.1%、城乡比1.73（缩）**\u201c农村+5.5%>城镇+4.6%。\u201d")
bullet(doc, "**15.常住575.86万、城镇化率58.34%（+1.29pct）**\u201c人口稳、城镇化快升。\u201d")
# ---- 五、2025年GDP目标 vs 实际：对照表 ----
heading1(doc, "五、2025年GDP目标 vs 实际：对照表")
table(doc,
    ["指标", "2025年目标", "2025年实际", "达成评价"],
    [
        ["地区生产总值(GDP)", "增长5.5%以上", "3304.15亿/5.6%", "达成"],
        ["规模以上工业", "——", "+7.7%", "稳增"],
        ["固定资产投资", "增长6%", "+0.3%", "差5.7pct"],
        ["社会消费品零售总额", "——", "1388.01亿/+5.9%", "稳增"],
        ["进出口总额", "促稳提质", "728.07亿/+15.5%", "超额"],
        ["一般公共预算收入", "增长3.5%", "264.80亿/+3.0%", "差0.5pct"],
        ["居民收入", "与经济增长同步", "30349元/+5.1%", "略低GDP(名义)"],
    ],
)
para(doc, "注：GDP、规上工业按可比价。**GDP（+5.6%）、进出口（+15.5%）、规上（+7.7%）**较好；**固投（+0.3%）、财收（+3.0%）**略低目标。")
para(doc, "拆读：**金属/汽车/化工制造、外贸、农业是亮色**；**地产（-16.6%）、固投（+0.3%）、CPI（-0.3%）、财收（+3.0%低3.5%）**是短板——\u201c工业外贸农业强、投资地产弱\u201d，是\u201c江北水城制造贸易\u201d样本。")

# ---- 六、2025年增长是谁撑起来的？（结构归因） ----
heading1(doc, "六、2025年增长是谁撑起来的？（结构归因）")
para(doc, "把聊城GDP的5.6%拆开：三次产业分别增4.0%、5.2%、6.3%（结构13.3：40.6：46.1）。**第三产业（服务业）+6.3%最快、第二产业（制造）+5.2%、第一产业（农业/粮食）+4.0%**——\u201c三产+二产\u201d双轮。")
para(doc, "2026年聊城强调\u201c工业强市、江北水城、黄河战略\u201d，聚焦**钢铁/有色金属升级、新能源汽车、轴承（临清）、食品农业、化工新材料、跨境电商**——核心是\u201c制造+农业+水城\u201d。")
para(doc, "**第二产业（工业）**：规上+7.7%（金属制品+43.1%、汽车+27.2%、化原+12.9%）、钢铁铜铝、新能源车——\u201c金属/制造强\u201d。")
para(doc, "**第三产业（服务业）**：+6.3%（商贸、物流、文旅、医养、跨境电商）——\u201c服务+水城\u201d。")
para(doc, "**第一产业（农业）**：+4.0%（粮食589万吨、蔬菜/肉蛋、农产品加工）——\u201c农业强市、稳\u201d。")
para(doc, "一句话归因：**2025年聊城增长\u201c靠工业（金属/汽车）+服务业（商贸/文旅）+农业\u201d**，地产投资弱；\u201c制造+外贸+农业\u201d是核心。")

# ---- 七、预算与财政的\u201c含金量\u201d ----
heading1(doc, "七、预算与财政的\u201c含金量\u201d")
para(doc, "2025年聊城**一般公共预算收入264.80亿元（+3.0%）**；税收收入157.96亿（+3.8%、占比59.7%）；支出554.92亿（-3.7%）、民生占74.4%。")
bullet(doc, "税收+3.8%（占比59.7%）——\u201c税收稳、含金量待提高\u201d。")
bullet(doc, "民生支出占74.4%（就业社保/教育环保增）。")
bullet(doc, "金融支撑：R&D 97.97亿（占3.09%、+0.09pct）、贷款/存款——信贷支持制造/农业。")
para(doc, "**财政含金量小结**：财收+3.0%（与GDP匹配、略低目标）、税收占59.7%、民生74.4%；财政对\u201c新质、农业、民生\u201d投入加大。")

# ---- 八、民生底账：人口、收入与城乡 ----
heading1(doc, "八、民生底账：人口、收入与城乡")
para(doc, "2025年聊城**居民人均可支配收入30349元（+5.1%）**，其中城镇38864元（+4.6%）、农村22437元（+5.5%），城乡比1.73（缩小）。消费：人均消费支出18293元（+3.7%）。就业：城镇新增就业4.37万人（完成109.2%）。")
para(doc, "人口画像：**常住575.86万、城镇化率58.34%（+1.29pct）**；山东人口大市、城镇化快速提升、出生率偏低。")
para(doc, "民生投入：就业社保/教育/节能环保、医疗养老——民生支出74.4%、扎实。")

# ---- 九、城镇与农村：格局与均衡 ----
heading1(doc, "九、城镇与农村：格局与均衡")
para(doc, "聊城城镇化率58.34%；县域经济强（临清轴承/高唐农机/茌平铝材/莘县蔬菜）；农村收入增速（+5.5%）>城镇（+4.6%），**城乡比1.73缩小**；乡村振兴（一镇一策）、和美乡村128个。")
para(doc, "农业底盘：**粮食589万吨、蔬菜/肉蛋、农林牧渔854.67亿、产粮油果蔬**——\u201c山东农业大市、菜篮子\u201d。")
para(doc, "一句话：\u201c聊城是工业+农业大市、县域经济强、城乡融合推进\u201d。")

# ---- 十、人口流入与流出 ----
heading1(doc, "十、人口流入与流出")
para(doc, "聊城常住575.86万（山东人口大市）、城镇化58.34%；\u201c制造/农业/轴承\u201d吸纳就业，但青年外流济南/青岛，主城（东昌府/茌平）+县域农业。")
para(doc, "结构观察：**人口较稳（农业县、总量）、自然增长低**；城镇化快速提升（+1.29pct）。")
para(doc, "2026年目标：稳就业、育人才（高企1000+、科技型1490家）——聊城靠\u201c制造+农业\u201d聚人。")

# ---- 十一、物价与货币环境 ----
heading1(doc, "十一、物价与货币环境")
para(doc, "2025年聊城**CPI-0.3%**（食品烟酒-2.1%、交通通信-2.7%；教育文娱+2.1%、其他用品+12.4%）——\u201c低通胀、需求温\u201d。")
bullet(doc, "信贷扩张：存款/贷款（R&D/科技贷）——宽信用支持制造/农业。")
bullet(doc, "消费：社零+5.9%、金银珠宝+75%、以旧换新——消费旺。")
para(doc, "货币环境判断：**宽信用、CPI-0.3%**；聊城靠\u201c制造+消费+农业\u201d稳需求（2026 CPI回升）。")

# ---- 十二、区域一体化：聊城的位置 ----
heading1(doc, "十二、区域一体化：聊城在济南都市圈、中原经济区、黄河流域\u201c三圈\u201d里的位置")
para(doc, "聊城是**济南都市圈重要成员、中原经济区/鲁西门户、黄河流域生态保护与高质量发展带、冀鲁豫交界中心城市**。")
bullet(doc, "**济南都市圈**　接济（省会）、济南都市圈西翼、高铁（济郑高铁）。")
bullet(doc, "**中原经济区/黄河流域**　冀鲁豫交界、黄河战略、中原城市群节点。")
bullet(doc, "**开放/一带一路**　聊城口岸、单县/莘县农业出口、跨境电商、京杭大运河。")
para(doc, "一句话：**聊城在\u201c济南+黄河流域+中原\u201d里，最核心是\u201c制造+农业+水城门户\u201d**；区位（三省交界、运河）、农业是优势。")

# ---- 十三、未来5—10年最值得观察的五条主线 ----
heading1(doc, "十三、未来5—10年最值得观察的五条主线")
bullet(doc, "**主线一：钢铁/有色金属升级（信发/祥光）**\u201c铜铝/大宗商品、绿色转型\u201d攻势。")
bullet(doc, "**主线二：新能源汽车/农机（时风）**\u201c新能源车+75.9%、高端农机\u201d强链。")
bullet(doc, "**主线三：轴承/装备制造（临清）**\u201c轴承之城、装备+22.5%\u201d。")
bullet(doc, "**主线四：农业强市/食品加工**\u201c粮食589万吨、蔬菜/肉蛋、预制菜\u201d。")
bullet(doc, "**主线五：江北水城/人口/外贸**\u201c文旅运河、外贸破700亿、聚人\u201d。")

# ---- 十四、最终结论 ----
heading1(doc, "十四、最终结论：聊城在\u201c江北水城+制造贸易+农业强市\u201d里的增长逻辑")
para(doc, "把2025年闭环打穿：**聊城是\u201c江北水城、工业制造之都、农业大市\u201d**：GDP 3304.15亿/+5.6%、规上+7.7%、进出口破700亿（+15.5%）、农林牧渔854.67亿。")
para(doc, "聊城不是\u201c只有水城\u201d——它是**钢铁铜铝（信发/祥光）+汽车农机+轴承+化工+农业**的复合，靠\u201c工业+外贸+农业\u201d驱动；但地产、投资、CPI弱。")
para(doc, "一句话结论：**聊城是\u201c江北水城、中国轴承之乡、制造农业强市\u201d；观察它先看\u201c金属/汽车/制造、进出口、农业\u201d，再看\u201c地产、投资、财收\u201d。**它是\u201c工业外贸强、投资地产弱\u201d的山东样本。")

# ---- 附录A ----
heading1(doc, "附录A：主要资料来源与核验路径")
bullet(doc, "《2025年聊城市人民政府工作报告》（2025年2月，2025年目标、2024年回顾+5.5%）")
bullet(doc, "《2025年聊城市国民经济和社会发展统计公报》（聊城市统计局，2026-03-27，2025年实际）")
bullet(doc, "《2026年聊城市政府工作报告》（2026年，复盘+2026年目标）")
bullet(doc, "聊城市人民政府/统计局（liaocheng.gov.cn）")

# ---- 附录B ----
heading1(doc, "附录B：建议建立的年度跟踪仪表盘")
bullet(doc, "GDP总量/增速 vs 全年目标。")
bullet(doc, "规模以上工业（钢铁/有色/汽车/轴承）增速。")
bullet(doc, "进出口/进口原料/加工贸易。")
bullet(doc, "固定资产/工业/房地产投资。")
bullet(doc, "农林牧渔/粮食/食品农业。")
bullet(doc, "社零/金银珠宝/以旧换新。")
bullet(doc, "新兴制造/高端、R&D。")
bullet(doc, "财收/税收占比/民生%。")
bullet(doc, "常住/城镇化/热补。")
bullet(doc, "CPI/存贷款/可再生能源。")

# ---------------- 保存 ----------------
out = "/Users/x/Desktop/content-prod-lab/reports/聊城市_2025年政府工作报告_深度研究_2026-08-14.docx"
doc.save(out)
print("SAVED OK: 聊城市", out)
