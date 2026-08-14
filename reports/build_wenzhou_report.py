# -*- coding: utf-8 -*-
"""Build 温州市2025年政府工作报告 深度研究 DOCX, 参照地级市系列版式。"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DARK = RGBColor(0x1F, 0x2A, 0x44)
RED = RGBColor(0xB0, 0x1E, 0x2E)
GRAY = RGBColor(0x59, 0x60, 0x69)
BLUE = RGBColor(0x1F, 0x4E, 0x79)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
HEADER_FILL = "1F2A44"
K = "微软雅黑"


def set_run(run, size=11, bold=False, color=DARK, font=K):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    run.font.name = font
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), font)
    rFonts.set(qn("w:hAnsi"), font)
    rFonts.set(qn("w:eastAsia"), font)


def para(doc, text, size=11, bold=False, color=DARK, align=None,
         space_after=6, space_before=0, font=K):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        set_run(r, size=size, bold=bold or (i % 2 == 1), color=color, font=font)
    return p


def heading1(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(16)
    pf.space_after = Pt(8)
    r = p.add_run(text)
    set_run(r, size=16, bold=True, color=RED)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "B01E2E")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def heading2(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(10)
    pf.space_after = Pt(4)
    r = p.add_run(text)
    set_run(r, size=13, bold=True, color=BLUE)
    return p


def table(doc, headers, rows, widths=None, font_size=9.5):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_run(r, size=font_size, bold=True, color=WHITE)
        tcPr = hdr[i]._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), HEADER_FILL)
        tcPr.append(shd)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(val))
            set_run(r, size=font_size, color=DARK)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    return t


def quote_box(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(8)
    pf.left_indent = Pt(12)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), "B01E2E")
    pBdr.append(left)
    pPr.append(pBdr)
    r = p.add_run(text)
    set_run(r, size=11, color=DARK, font="楷体")
    return p


def bullet(doc, text, size=10.5):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.7)
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        set_run(r, size=size, bold=(i % 2 == 1), color=DARK)
    return p


# ================================================================= 文档主体
doc = Document()
sec = doc.sections[0]
sec.page_width = Cm(21)
sec.page_height = Cm(29.7)
sec.top_margin = Cm(2.4)
sec.bottom_margin = Cm(2.2)
sec.left_margin = Cm(2.5)
sec.right_margin = Cm(2.5)

st = doc.styles["Normal"]
st.font.name = K
st.font.size = Pt(11)
st.element.rPr.rFonts.set(qn("w:eastAsia"), K)


# ---- 封面 ----
para(doc, "温州市2025年政府工作报告", size=24, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=60, space_after=4)
para(doc, "深度研究与\u201c容易被忽视的细节\u201d分析", size=16, bold=True, color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
para(doc, "从\u201c温州模式、民营经济、鞋服电气、小商品、浙南第三极\u201d重新理解温州", size=12, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)
para(doc, "\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501", color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
para(doc, "研究对象：2025年温州市政府工作报告及同期财政、国民经济和社会发展统计资料、2026年官方复盘", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
para(doc, "标定日期：2026-08-14", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
doc.add_page_break()

# ---- 目录 ----
catalog = [
    "执行摘要",
    "一、研究口径与阅读方法",
    "二、先看温州的特殊底盘：民营经济、鞋服电气、小商品、港口与浙南第三极",
    "三、最关键的宏观错位：GDP 10213.9亿/6.1%破万亿，民营工业强但固投-14.7%、楼市城调、CPI-0.3%",
    "四、容易被忽视、但信号很重的细节（15条）",
    "五、2025年GDP目标 vs 实际：对照表",
    "六、2025年增长是谁撑起来的？（结构归因）",
    "七、预算与财政的\u201c含金量\u201d",
    "八、民生底账：人口、收入与城乡",
    "九、城镇与农村：格局与均衡",
    "十、人口流入与流出",
    "十一、物价与货币环境",
    "十二、区域一体化：温州在\u201c海西、浙南、长三角、闽北\u201d里的位置",
    "十三、未来5—10年最值得观察的五条主线",
    "十四、最终结论：温州在\u201c民营制造+鞋服电气+浙南枢纽\u201d里的增长逻辑",
    "附录A：主要资料来源与核验路径",
    "附录B：建议建立的年度跟踪仪表盘",
]
para(doc, "目　录", size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=10, space_after=10)
for item in catalog:
    para(doc, item, size=12, space_after=4)
doc.add_page_break()

# ---- 执行摘要 ----
heading1(doc, "执行摘要")
para(doc, "**核心判断**　2025年温州最显著的是\u201cGDP 10213.9亿元、增长6.1%（突破万亿、浙江第3、全国地级市10强）\u201d、\u201c规上工业+10.3%（全省第1）、民营占规上90.7%\u201d、\u201c进出口3116.5亿/+6.1%（首破3000亿）\u201d、\u201c但固投-14.7%、三产-23.3%、CPI-0.3%\u201d、\u201c财收647.0亿/+2.3%、常住990.2万\u201d。这说明温州在\u201c民营经济+温州模式\u201d中，**工业出口强、民营好但投资楼市偏弱**。")
para(doc, "把2025年目标（GDP+6%左右）、2025年统计（GDP+6.1%达成、规上+10.7%超额、固投-14.7%、社零+4.2%、财收+2.3%）趋势看，温州是\u201c民营制造+鞋服+电气+小商品\u201d路径：**鞋业、服装、电气（低压电器）、汽摩配、五金、数字经济**是支柱；2025年突破万亿，全国地级市第10强。")
para(doc, "最容易记住的一句话：**温州是\u201c温州模式（民营经济）、轻工商贸之都、浙南第三极\u201d，靠\u201c民营微企+鞋服电气出口\u201d增长。**观察温州，与其只看\u201cGDP 10214亿\u201d，不如看\u201c民营占规上90.7%、规上+10.7%、进出口破3000亿、出口东盟/中东+13%+\u201d。")

# ---- 一、研究口径与阅读方法 ----
heading1(doc, "一、研究口径与阅读方法")
para(doc, "本报告以\u201c温州市政府工作报告（2025年，张文杰作）\u201d为起点，把\u201c2025年GDP目标（6%左右）\u201d与\u201c官方2025年（10213.9亿元/+6.1%）\u201d并置对照，用\u201c2025年温州市统计公报\u201d和\u201c2026年政府工作报告\u201d横向核验。")
para(doc, "口径提示：\u201cGDP、规上工业增速\u201d按可比价（实际）；\u201c投资、消费、进出口、财政\u201d按现价（名义）。\u201c常住人口\u201d用统计公报（990.2万）、城镇化率76.7%。")
para(doc, "指标体系（与研究口径一致）：**总量与增速、产业动能（民营/鞋服/电气）、外贸、财政、民生与人口**。")
para(doc, "特别提示（不吃老本）：温州2024年GDP最终核实9671亿，2025年破万亿（+6.1%）；它不只是\u201c小商品\u201d——**民营经济（占规上90.7%）、低压电器（正泰/德力西）、鞋服、浙商资本\u201d才是真正底色。")
# ---- 二、先看温州的特殊底盘 ----
heading1(doc, "二、先看温州的特殊底盘：民营经济、鞋服电气、小商品、港口与浙南第三极")
para(doc, "温州地处浙江东南、瓯江口，是**民营经济发源地（温州模式）、中国鞋都/低压电器之都、浙南枢纽（全省第三极）**。2025年GDP 10213.9亿元（破万亿）、常住990.2万，浙江第3，人均10.34万。")
para(doc, "五个底盘名词，先立框架：")
bullet(doc, "**温州模式/民营经济**　在册市场经营主体161.1万户（民营+个体占97.6%），民营占规上90.7%\uff5c浙商·温商80万（民营500强14家）。")
bullet(doc, "**鞋服**　中国鞋都；温州鞋靴出口263.1亿、服装143.1亿，轻工/鞋都。")
bullet(doc, "**电气（低压电器）**　中国低压电器之都（正泰/德力西在乐清柳市），电气入选国家先进制造业集群。")
bullet(doc, "**小商品/五金/汽摩配**　温州五金、汽摩配、制笔/打火机/眼镜等\u201c块状经济\u201d小商品；温州医科大学/师范。")
bullet(doc, "**港口/浙南第三极**　温州港（集装箱目标300万TEU）、面向海西/闽北/长三角，\u201c浙南枢纽+侨商\u201d。")
para(doc, "这五根（民营+鞋服+电气+小商品+港口）构成温州独特底盘：**左手民营轻工（鞋服电气），右手浙商资本/港口**。理解温州，先理解\u201c温州模式、敢为天下先\u201d。")

# ---- 3、最关键的宏观错位 ----
heading1(doc, "三、最关键的宏观错位：GDP 10213.9亿/6.1%破万亿，民营工业强但固投-14.7%、楼市调整、CPI弱")
para(doc, "2025年温州最需要辨析的一组\u201c错位\u201d：**GDP 6.1%（破万亿、达6%目标）、规上+10.3%（民营90.7%）强，但固投-14.7%、三产投资-23.3%、CPI-0.3%**。")
para(doc, "为什么\u201c工业/出口这么强、民营这么好\u201d，投资与消费却偏弱？三解释：")
para(doc, "**其一，工业出口强、民营旺**　规上+10.3%（全省第1）、出口+9.6%、\u201c三样\u201d+60%、对东盟/中东两位数——\u201c温州造\u201d外销旺。")
para(doc, "**其二，投资、地产、三产基建弱**　固投-14.7%（房地产/三产基建-23.3%）；房屋销售641万㎡、楼市调。")
para(doc, "**其三，消费/物价弱**　社零+4.2%（温州消费可但低）、CPI-0.3%、PPI-0.5%——\u201c需求偏弱、内需-消费\u201d。")
para(doc, "小结：温州2025年是\u201c**强工业出口民营、弱投资楼市消费略弱**\u201d：鞋服电气、外销强，固投（房地产拖累）、CPI弱。")

# ---- 四、容易被忽视、但信号很重的细节（15条） ----
heading1(doc, "四、容易被忽视、但信号很重的细节（15条）")
bullet(doc, "**1.民营占规上90.7%、民营规上+11.1%**\u201c温州模式（浙商民营绝对主力）。\u201d")
bullet(doc, "**2.规上工业+10.3%全省第1**\u201c鞋服/低压电器/汽摩配\u201d制造旺。\u201d")
bullet(doc, "**3.进出口3116.5亿破3000亿/+6.1%，对东盟+16.8%、中东+18%**\u201c\u2019一带一路\u2019\u2019+13.4%。\u201d")
bullet(doc, "**4.鞋出口263.1亿/-7.3%、服装-5.5%**\u201c传统鞋服出口承压（受欧盟/需求）。\u201d")
bullet(doc, "**5.旅游、体育会展、首店104个/演出1.3万场**\u201c文旅、世泰大会\u201d文旅消费。")
bullet(doc, "**6.财收647.0亿/+2.3%（税收）**\u201c财政中规中矩、民生79.2%。\u201d")
bullet(doc, "**7.社零5143.9亿/+4.2%、线上+20.8%**\u201c电商/新消费（鞋服/家居高）。\u201d")
bullet(doc, "**8.居民收入73660元/+4.0%、城乡比1.82**\u201c农村+4.9%>城镇+3.5%。\u201d")
bullet(doc, "**9.常住990.2万/+5万/城镇化76.7%（+0.7pct）**\u201c人口为净流入大市（侨乡/外来）。\u201d")
bullet(doc, "**10.中国民营500强温企14家（10家总部）、浙商**\u201c商帮、财团（正泰/华峰）。\u201d")
bullet(doc, "**11.高新技术/专精特新**、\u201c新三样\u201d出口+60%\u201c创新制造。\u201d")
bullet(doc, "**12.数字经济核心产业（+8.5%）、AI（220亿目标）**\u201c新质。\u201d")
bullet(doc, "**13.CPI-0.3%、PPI-0.5%**\u201c物价低、工业品弱。\u201d")
bullet(doc, "**14.乐清柳市（低压电器）、瑞安（汽摩配）、平阳（印刷）**\u201c县域块状经济（\u201c温州智造）。\u201d")
bullet(doc, "**15.温州港集装箱\u201c温货温运\u201d30万标箱、对马尼拉/迪拜航线**\u201c临港/外贸航线。\u201d")
# ---- 五、2025年GDP目标 vs 实际：对照表 ----
heading1(doc, "五、2025年GDP目标 vs 实际：对照表")
table(doc,
    ["指标", "2025年目标", "2025年实际", "达成评价"],
    [
        ["地区生产总值(GDP)", "增长6%左右", "10213.9亿/6.1%", "达成(破万亿)"],
        ["规上工业增加值", "——", "+10.3%(全省第1)", "大幅超额"],
        ["固定资产投资", "——", "-14.7%", "大幅下行"],
        ["社会消费品零售总额", "——", "5143.9亿/+4.2%", "稳健偏低"],
        ["进出口总额", "——", "3116.5亿/+6.1%", "超额(破3000亿)"],
        ["一般公共预算收入", "——", "647.0亿/+2.3%", "稳增"],
        ["居民收入", "与经济增长同步", "73660元/+4.0%", "略低GDP"],
    ],
)
para(doc, "注：GDP、规上工业按可比价。**GDP（+6.1%破万亿）、规上（+10.3%）、进出口（破3000亿/+6.1%）超额**；**固投（-14.7%）、CPI（-0.3%）偏弱**。")
para(doc, "拆读：**民营（占90.7%/规上+10.3%全省第1）、出口（破3000亿/东盟+16.8%）、财收是亮色**；**固投（-14.7%）、CPI（-0.3%）**是短板——\u201c民营制造强、固投弱\u201d，是\u201c温州模式\u201d样本。")

# ---- 六、2025年增长是谁撑起来的？（结构归因） ----
heading1(doc, "六、2025年增长是谁撑起来的？（结构归因）")
para(doc, "把温州GDP的6.1%拆开：三次产业分别增3.6%、5.3%、6.6%（结构2.0：36.3：61.7）。**第三产业（服务业）是主引擎（+6.6%），第二产业（民营制造）强，第一产业（农业）稳**。")
para(doc, "2026年温州强调\u201c全省第三极、民营经济、港通\u201d，聚焦**鞋服电气升级、新能源（新三样）、数字经济/AI、临港、侨商回归**——核心是\u201c民营强市+浙南枢纽\u201d。")
para(doc, "**第二产业（工业）**：规上+10.3%（全省第1）、民营工业+11.1%、低压电器/鞋服/汽摩配/五金——\u201c民营制造强\u201d。")
para(doc, "**第三产业（服务业）**：+6.6%（商贸、物流、会展、旅游）、数字经济核心+8.5%——\u201c服务业+数字经济\u201d。")
para(doc, "**外贸（开放）**：进出口+6.1%破3000亿（出口+9.6%、对东盟/中东/一带一路+13%+）、\u201c新三样\u201d+60%、实际使用外资+15.1%——\u201c外向强\u201d。")
para(doc, "一句话归因：**2025年温州增长\u201c靠民营工业（鞋服电气）+服务业+出口\u201d**，固投（地产）、CPI弱；\u201c民营+外向\u201d是核心。")

# ---- 七、预算与财政的\u201c含金量\u201d ----
heading1(doc, "七、预算与财政的\u201c含金量\u201d")
para(doc, "2025年温州**一般公共预算收入647.0亿元（+2.3%）**；支出1206.1亿（+2.4%）、民生占79.2%。")
bullet(doc, "税收（大体）、其中民营/鞋服出口税、制造业税——\u201c税收靠民营经济\u201d。")
bullet(doc, "民生支出占79.2%（教育/社保/医疗）。")
bullet(doc, "金融支撑：存款24754亿（+9.0%）、贷款24170亿（+7.2%）——信贷充裕、支持民营小微/外贸。")
para(doc, "**财政含金量小结**：财收+2.3%（与GDP匹配偏低）、靠民营制造税，民生79.2%；财政对\u201c民营经济、鞋服电气升级、临港\u201d投入加大。")

# ---- 八、民生底账：人口、收入与城乡 ----
heading1(doc, "八、民生底账：人口、收入与城乡")
para(doc, "2025年温州**居民人均可支配收入73660元（+4.0%）**，其中城镇84347元（+3.5%）、农村46446元（+4.9%），城乡比1.82（收窄）。就业：城镇新增就业10.2万人。")
para(doc, "人口画像：**常住990.2万/+5万、城镇化76.7%（+0.7pct）**，浙江人口第3；温州为\u201c侨乡+外来打工\u201d净流入大市、出生率偏低。")
para(doc, "民生投入：公办普高学位、普惠园96.3%、医疗教育——民生较好但生活消费占比高。")

# ---- 九、城镇与农村：格局与均衡 ----
heading1(doc, "九、城镇与农村：格局与均衡")
para(doc, "温州常住城镇化率76.7%，乡镇/县域块状经济发达（乐清/瑞安/平阳/苍南）；农村收入增速（+4.9%）高于城镇（+3.5%），**城乡比缩至1.82**。")
para(doc, "农业底盘：**粮食约14亿斤、渔业（瓯江/渔获）、水果、杨梅/茶**——\u201c温台温菜、海产\u201d。")
para(doc, "一句话：\u201c温州城镇化高、县域民营（小商品厂）强、城乡比良好\u201d，但\u201c山区（文成/泰顺）振兴需造\u201d。")

# ---- 十、人口流入与流出 ----
heading1(doc, "十、人口流入与流出")
para(doc, "温州常住990.2万（净流入+5万）、城镇化76.7%，是浙江省人口第2大市；\u201c民营工厂/鞋服/侨乡\u201d吸引外来（江西/贵州/四川）劳动力、\u201c温商侨二代\u201d回流。")
para(doc, "结构观察：**侨乡（海外侨胞约140万）、外来务工**；14大人口、自然增长-1.4‰（低生育）。")
para(doc, "2026年目标：新引育人才20万、留温——温州靠\u201c民营+侨乡+临港\u201d聚人。")

# ---- 十一、物价与货币环境 ----
heading1(doc, "十一、物价与货币环境")
para(doc, "2025年温州**CPI-0.3%**（食品烟酒-0.4%、交通通信-2.8%、居住-0.6%；衣着+2.0%）——\u201c低通胀、需求弱\u201d。")
bullet(doc, "信贷扩张：贷款+7.2%、存款+9.0%，宽信用对民营/外贸。")
bullet(doc, "消费：线上+20.8%、鞋服/家居/智能手机高——新消费。")
para(doc, "货币环境判断：**宽信用、CPI-0.3%**；温州靠\u201c外贸+民营制造+数字经济\u201d稳需求（2026 CPI+2%）。")

# ---- 十二、区域一体化：温州的位置 ----
heading1(doc, "十二、区域一体化：温州在\u201c长江、浙南第三极、海西、闽北\u201d里的位置")
para(doc, "温州是**浙江省\u201c第三极\u201d（省委定位、杭州/宁波之后）、浙南闽北中心城市、海峡两岸门户枢纽**。")
bullet(doc, "**全省第三极**　温州在浙江产业（鞋服/电气）对杭甬的承接、浙南增长极。")
bullet(doc, "**海西枢纽**　闽北（宁德）、海坛海峡，面向台湾海峡、海上丝路。")
bullet(doc, "**港口/侨商**　温州港（集箱300万）、瓯江，侨商（温商）链接全球、一带一路。")
para(doc, "一句话：**温州在\u201c浙南+海西\u201d里，最核心是\u201c民营制造+浙南枢纽（全省第三极）\u201d**；区位、侨力、民营企业是大优势。")

# ---- 十三、未来5—10年最值得观察的五条主线 ----
heading1(doc, "十三、未来5—10年最值得观察的五条主线")
bullet(doc, "**主线一：民营经济/小微制造升级（温州模式）**\u201c鞋服/低压电器\u201d能否智造/品牌（鞋都、电器之都）。")
bullet(doc, "**主线二：新能源/新材料（碳电池）**\u201c新三样\u201d出口+60%、储能电池。")
bullet(doc, "**主线三：数字经济/AI**\u201c数字经济+8.5%、AI核心220亿\u201d新质。")
bullet(doc, "**主线四：浙南枢纽/临港（300万TEU）**\u201c温州港、侨商出海\u201d（对东盟/中东）。")
bullet(doc, "**主线五：人口与侨乡**\u201c990万、140万侨\u201d如何\u201c引才、回流、放手\u201d。")

# ---- 十四、最终结论 ----
heading1(doc, "十四、最终结论：温州在\u201c民营制造+鞋服电气+浙南枢纽\u201d里的增长逻辑")
para(doc, "把2025年闭环打穿：**温州是\u201c温州模式、民营经济、浙南第三极\u201d**：GDP 10213.9亿/+6.1%（破万亿）、规上+10.3%全省第1、进出口破3000亿、民营占90.7%。")
para(doc, "温州不是\u201c只有鞋\u201d——它是**民营轻工（鞋服/电气）+数字经济+侨商/港（浙南枢纽）**的民营强市，靠\u201c民营制造+外向\u201d驱动；固投（地产）、CPI、三产基建弱。")
para(doc, "一句话结论：**温州是\u201c民营经济之都、鞋服电气强、浙南第三极\u201d；观察它先看\u201c民营占比、鞋服/电气/汽摩配出口、数字经济、港口\u201d，再看\u201c固投、CPI、外资\u201d。**它是\u201c制造强、民营活、投资弱\u201d的浙南样本。")

# ---- 附录A ----
heading1(doc, "附录A：主要资料来源与核验路径")
bullet(doc, "《2025年温州市政府工作报告》（2025年1月，张文杰作，2025年目标、2024年回顾9671.0亿）")
bullet(doc, "《2025年温州市国民经济和社会发展统计公报》（温州市统计局，2026-04-24，2025年实际）")
bullet(doc, "《2026年温州市政府工作报告》（2026年1月，复盘+2026年目标）")
bullet(doc, "温州市人民政府/统计局（wenzhou.gov.cn/wztjj）")

# ---- 附录B ----
heading1(doc, "附录B：建议建立的年度跟踪仪表盘")
bullet(doc, "GDP总量/增速 vs 全年目标。")
bullet(doc, "规上工业（民营/鞋服/低压电器）增速。")
bullet(doc, "民营业主、500强、县域块状。")
bullet(doc, "固定资产/工业/房地产投资。")
bullet(doc, "社零、线上、以旧换新。")
bullet(doc, "进出口、\u201c新三样\u201d、东盟/中东。")
bullet(doc, "一般公共预算收入、税收、民生%。")
bullet(doc, "常住人口/城镇化、大学生人才。")
bullet(doc, "CPI、存贷款。")
bullet(doc, "港口箱量、温侨商回归/新城。")

# ---------------- 保存 ----------------
out = "/Users/x/Desktop/content-prod-lab/reports/温州市_2025年政府工作报告_深度研究_2026-08-14.docx"
doc.save(out)
print("SAVED OK: 温州市", out)
