# -*- coding: utf-8 -*-
"""Build 青岛市2025年政府工作报告 深度研究 DOCX, 参照省市系列版式。"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DARK = RGBColor(0x1F, 0x2A, 0x44)
RED = RGBColor(0xB0, 0x1E, 0x2E)
GRAY = RGBColor(0x59, 0x60, 0x69)
BLUE = RGBColor(0x1F, 0x4E, 0x79)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
HEADER_FILL = "1F2A44"
K = "微软雅黑"


def set_run(run, size=11, bold=False, color=DARK, font=K):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    run.font.name = font
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), font)
    rFonts.set(qn("w:hAnsi"), font)
    rFonts.set(qn("w:eastAsia"), font)


def para(doc, text, size=11, bold=False, color=DARK, align=None,
         space_after=6, space_before=0, font=K):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        set_run(r, size=size, bold=bold or (i % 2 == 1), color=color, font=font)
    return p


def heading1(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(16)
    pf.space_after = Pt(8)
    r = p.add_run(text)
    set_run(r, size=16, bold=True, color=RED)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "B01E2E")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def heading2(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(10)
    pf.space_after = Pt(4)
    r = p.add_run(text)
    set_run(r, size=13, bold=True, color=BLUE)
    return p


def table(doc, headers, rows, widths=None, font_size=9.5):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_run(r, size=font_size, bold=True, color=WHITE)
        tcPr = hdr[i]._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), HEADER_FILL)
        tcPr.append(shd)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(val))
            set_run(r, size=font_size, color=DARK)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    return t


def quote_box(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(8)
    pf.left_indent = Pt(12)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), "B01E2E")
    pBdr.append(left)
    pPr.append(pBdr)
    r = p.add_run(text)
    set_run(r, size=11, color=DARK, font="楷体")
    return p


def bullet(doc, text, size=10.5):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.7)
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        set_run(r, size=size, bold=(i % 2 == 1), color=DARK)
    return p


# ================================================================= 文档主体
doc = Document()
sec = doc.sections[0]
sec.page_width = Cm(21)
sec.page_height = Cm(29.7)
sec.top_margin = Cm(2.4)
sec.bottom_margin = Cm(2.2)
sec.left_margin = Cm(2.5)
sec.right_margin = Cm(2.5)

st = doc.styles["Normal"]
st.font.name = K
st.font.size = Pt(11)
st.element.rPr.rFonts.set(qn("w:eastAsia"), K)



# ---- 封面 ----
para(doc, "青岛市2025年政府工作报告", size=24, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=60, space_after=4)
para(doc, "深度研究与\u201c容易被忽视的细节\u201d分析", size=16, bold=True, color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
para(doc, "从\u201c港口经济、海洋强市、先进制造与计划单列\u201d重新理解青岛", size=12, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)
para(doc, "\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501", color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
para(doc, "研究对象：2025年青岛市政府工作报告及同期财政、国民经济和社会发展统计资料、2026年官方复盘", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
para(doc, "标定日期：2026-08-13", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
doc.add_page_break()

# ---- 目录 ----
catalog = [
    "执行摘要",
    "一、研究口径与阅读方法",
    "二、先看青岛的特殊底盘：港口经济、海洋强市、先进制造与计划单列",
    "三、最关键的宏观错位：GDP破1.75万亿、工业/消费/港口强，但出口/外资偏弱",
    "四、容易被忽视、但信号很重的细节（15条）",
    "五、2025年GDP目标 vs 实际：对照表",
    "六、2025年增长是谁撑起来的？（结构归因）",
    "七、预算与财政的\u201c含金量\u201d",
    "八、民生底账：人口、收入与城乡",
    "九、城镇与农村：格局与均衡",
    "十、人口流入与流出",
    "十一、物价与货币环境",
    "十二、区域一体化：青岛在\u201c青岛都市圈+胶东经济圈+港口开放\u201d里的位置",
    "十三、未来5—10年最值得观察的五条主线",
    "十四、最终结论：青岛在\u201c海洋/港口+先进制造+计划单列\u201d里的增长逻辑",
    "附录A：主要资料来源与核验路径",
    "附录B：建议建立的年度跟踪仪表盘",
]
para(doc, "目　录", size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=10, space_after=10)
for item in catalog:
    para(doc, item, size=12, space_after=4)
doc.add_page_break()

# ---- 执行摘要 ----
heading1(doc, "执行摘要")
para(doc, "**核心判断**　如果只看新闻标题，2025年青岛最显著的是\u201cGDP破1.76万亿、增长5.4%（高于全国0.4pct）\u201d、\u201c规上工业+7.9%\u201d、\u201c青岛港货物/集装箱全球第4/第5\u201d、\u201c海洋经济+7.5%\u201d。但这份研究真正值得深读的，是这座\u201c海洋强市+港口经济+先进制造\u201d的计划单列城市，如何在出口低增（进出口+0.6%）、外资下滑的背景下，靠\u201c工业（汽车+25.9%/化工+15.9%/船舶+12.8%）+海洋经济+消费/文旅\u201d实现5.4%的增长。")
para(doc, "把2025年初设定的目标（GDP增长5.5%左右）、2025年《国民经济和社会发展统计公报》、2026年复盘放在一起看，青岛呈现清晰暗线：**从\u201c传统制造+港口\u201d向\u201c海洋经济+先进制造(汽车/船舶/新一代信息技术)+品牌消费\u201d升级**。工业是超预期的亮点，外贸/外资则偏弱待修。")
para(doc, "因此，本报告不按政府工作报告逐段复述，而采用\u201c显性表述—同期数据—制度含义—长期影响\u201d的方式，专门提取容易被忽略、但对判断青岛未来5—10年发展模式更有价值的信号。")
para(doc, "最容易记住的一句话：**青岛是\u201c海洋经济+港口+先进制造\u201d的计划单列城市，靠\u201c工业（汽车/船舶/化工）+海洋经济+消费/文旅\u201d撑起增长。**观察青岛，与其看\u201cGDP 1.76万亿\u201d，不如看\u201c海洋经济、青岛港全球第4/第5、车辆/船舶/化工、品牌城市\u201d这几张名片。")
heading2(doc, "一页速览：2025年青岛经济的\u201c表与里\u201d")
table(doc,
    ["维度", "表面现象", "更深层信号"],
    [
        ["增长", "GDP 17560.67亿、+5.4%", "一产2.9%、二产33.5%、三产63.6%"],
        ["产业", "规上工业+7.9%", "汽车+25.9%、化工+15.9%、船舶+12.8%"],
        ["外贸", "进出口9128.9亿、+0.6%", "出口+2.7%、对一带一路+6.2%"],
        ["投资", "固投/港口投资稳健", "海洋项目、青岛港全球前列"],
        ["消费", "社零6960.7亿、+3.2%", "乡村+4.5%、可穿戴+130%"],
        ["人口", "常住约1050万、城镇化79%", "人口+5.9万、收入+4.7%"],
        ["海洋", "海洋生产总值+7.5%", "海洋中心城市、港口全球第4/第5"],
    ],
    widths=[2.2, 5.6, 7.4])
para(doc, "", size=10, space_after=2)

# ---- 一、研究口径与阅读方法 ----
heading1(doc, "一、研究口径与阅读方法")
heading2(doc, "1.1 本报告的核心底稿")
bullet(doc, "**2025年《政府工作报告》**（2025年2月）——目标：GDP 5.5%左右、一般公共预算收入3%、城镇就业35万。")
bullet(doc, "**《2025年青岛市统计公报》**（市统计局2026-04）——GDP、工业、外贸、海洋、人口实数。")
bullet(doc, "**2026年青岛市政府工作报告/复盘**（2026年1月）——对2025执行追认与海洋/先进制造展望。")
heading2(doc, "1.2 阅读方法：显性—数据—制度—长期")
para(doc, "**关键判别**：数据优先。例2025年GDP目标5.5%、实际5.4%基本达标；规上工业目标、实际7.9%超预期。青岛\u201c工业/港口落地强、外贸/外资偏弱\u201d，需穿透总量看海洋与制造。")

# ---- 二、底盘 ----
heading1(doc, "二、先看青岛的特殊底盘：海洋经济、港口经济、先进制造与计划单列")
para(doc, "青岛的地盘取决于它作为\u201c**海洋经济强市+青岛港枢纽+先进制造基地+计划单列市**\u201d的特殊定位。它是山东/胶东经济龙头，长江以北沿海开放门户。")
bullet(doc, "**海洋经济**：海洋生产总值+7.5%，\u201c4+4+2\u201d现代海洋产业体系；大型散货船/超大型矿砂船订单全球第一；海葵一号交付。")
bullet(doc, "**青岛港**：港口货物吞吐量7.1亿吨（全球第4）、集装箱3087万标箱（全球第5）；中欧班列开行超1000列；海铁联运连续10年全国第一。")
bullet(doc, "**先进制造**：规上工业+7.9%；汽车/化工/铁路船舶航空运输设备增加值高速增长；动车组产量约占全国50%；\u201c10+1\u201d创新型产业体系。")
bullet(doc, "**计划单列**：经济总量副省级城市前列（GDP 1.76万亿）；财政、外贸、海洋权限独立，制造业/品牌聚集。")
para(doc, "这一底板决定了青岛2025成绩单的\u201c底色\u201d：**只要先进制造/海洋/港口持续，青岛就站在\u201c海洋强+制造强+开放门户\u201d增长极；若外贸/外资/地产承压，青岛需承受\u201c工业强、外需与地产偏弱\u201d的结构挑战。**")

# ---- 三、宏观错位 ----
heading1(doc, "三、最关键的宏观错位：GDP破1.75万亿、工业/消费/港口强，但出口/外资偏弱")
para(doc, "青岛2025年最值得咀嚼的错位，是\u201c**工业+海洋+消费落地强、出口/外资/地产相对弱**\u201d。这种错位决定了对这座\u201c海洋计划单列\u201d城市的观察不能只看GDP总量。")
bullet(doc, "**GDP**：17560.67亿、+5.4%（高于全国0.4pct）。一产516.21亿（+4.0%，占比2.9%）、二产5873.83亿（+4.9%，占比33.5%）、三产11170.63亿（+5.6%，占比63.6%）。")
bullet(doc, "**工业**：规上工业+7.9%；汽车+25.9%、化工+15.9%、铁路船舶航空运输设备+12.8%、装备+11.8%、高技术+8.4%。")
bullet(doc, "**消费**：社零6960.7亿、+3.2%；乡村+4.5%、可穿戴智能设备+130%、公共网络零售+15.5%。")
bullet(doc, "**外贸**：进出口9128.9亿、+0.6%（出口5418.7亿、+2.7%）；对一带一路+6.2%、民营占72.4%。")
bullet(doc, "**海洋/港口**：海洋生产总值+7.5%；青岛港货物7.1亿吨（全球第4）、集装箱3087万标箱（全球第5）。")
para(doc, "**为什么读这条**：青岛作为\u201c计划单列+海洋强市\u201d，结构性矛盾是\u201c工业/海洋/港口强，出口/外资偏弱\u201d。经济靠工业+海洋+消费落地，外循环依赖与走强尚需观察。")

# ---- 四、15条细节 ----
heading1(doc, "四、容易被忽视、但信号很重的细节（15条）")
para(doc, "以下15条，都在2025年报告/公报里，但常被\u201cGDP 1.76万亿\u201d、\u201c规上工业+7.9%\u201d等总量掩盖。它们是判断青岛2025之后5—10年的关键小信号。")
bullet(doc, "**1. 海洋经济+7.5%**：海洋是青岛第一特色\uff0c\u201c国家海洋中心城市\u201d定位。")
bullet(doc, "**2. 青岛港全球第4/第5**：货物7.1亿吨（全球第4）、集装箱3087万标箱（第5）；海铁联运全国第一。")
bullet(doc, "**3. 动车组产量约占全国50%**：高端轨道交通制造（中车四方等）。")
bullet(doc, "**4. 大型散货船/超大型矿砂船订单全球第一**：造船/海工装备领先，海葵一号交付。")
bullet(doc, "**5. 集成电路圆片+51.9%、光电子器件+36.1%**：新一代信息技术/半导体增速亮眼。")
bullet(doc, "**6. 数字经济核心产业占GDP 12%以上**：\u201c10+1\u201d创新型产业、AI核心产业目标1000亿。")
bullet(doc, "**7. 城镇新增就业37.3万（民营占84.4%）**：民营经济吸纳就业主力。")
bullet(doc, "**8. 常住约1050万、城镇化79%**：\u201c十四五\u201d城镇化率+2.5pct。")
bullet(doc, "**9. 收入62738元、+4.7%、农村+5.5%**：收入增长快于多数副省级城市。")
bullet(doc, "**10. 品牌/消费城市**：\u201c乐购青岛\u201d1100场、品牌首店100家、国际消费中心城市。")
bullet(doc, "**11. 文旅**：游客/旅游收入均增长11%+，入境游+48.5%、海上游+52.4%。")
bullet(doc, "**12. 财政民生占比76.3%**：科学技术支出+43.7%，研发投入强度目标3.5%。")
bullet(doc, "**13. 港口开放/中欧班列**：开行超1000列、+17.8%；上合示范、自贸制度创新。")
bullet(doc, "**14. 品牌/制造**：国家级制造业单项冠军、专精特新居前；家电（海尔/海信）集群强。")
bullet(doc, "**15. 都市圈/济青一体化**：青岛都市圈50个重点项目、济青高铁公交化；胶东经济圈。")
para(doc, "", size=10, space_after=2)

# ---- 五、对照表 ----
heading1(doc, "五、2025年GDP目标 vs 实际：对照表")
para(doc, "2025年报告中列出的主要预期目标，与2025年实际完成情况：")
tb = [
    ["指标", "2025年目标", "2025年实际（公报/报告）", "达标判定"],
    ["GDP增速", "5.5%左右", "+5.4%（17560.67亿）", "基本达标"],
    ["一般公共预算收入", "3%", "+0.1%（1340.7亿）", "明显偏低"],
    ["城镇新增就业", "35万人以上", "37.3万人", "达标"],
    ["居民收入", "与经济增长同步", "+4.7%（62738元）", "达标"],
    ["进出口", "——", "9128.9亿、+0.6%", "平缓"],
]
table(doc, tb[0], tb[1:], widths=[3.2, 3.2, 4.6, 4.0])
para(doc, "", size=10, space_after=2)
para(doc, "**要点**：GDP/就业/收入达标，财政/外贸偏弱——青岛\u201c工业/民生强、财政/贸易待修\u201d。")

# ---- 六、结构归因 ----
heading1(doc, "六、2025年增长是谁撑起来的？（结构归因）")
heading2(doc, "6.1 三次产业：三产主导、二产稳")
para(doc, "**三产（+5.6%、占比63.6%，贡献率66.5%）主导**：批发零售+7.3%、营运互联网+44.3%、科技推广+19.2%。二产+4.9%（占比33.5%）、一产+4.0%。")
heading2(doc, "6.2 工业：汽车/化工/船舶驱动")
para(doc, "规上工业+7.9%；汽车+25.9%、化工+15.9%、铁路/船舶/航空航天+12.8%、装备+11.8%、高技术+8.4%。**汽车+海洋装备+新能源**是高景气引擎。")
heading2(doc, "6.3 消费企稳")
para(doc, "社零+3.2%；可穿戴+130%、公共网络零售+15.5%、乡村+4.5%——新型消费/下沉市场亮眼。")
heading2(doc, "6.4 外贸平缓")
para(doc, "进出口+0.6%（出口+2.7%）、对一带一路+6.2%、民营占72.4%——外贸韧性但增速低。")
heading2(doc, "6.5 海洋/港口强")
para(doc, "海洋生产总值+7.5%，港口全球前列——独有差异化增长极。")
para(doc, "**一句话归因**：青岛2025年\u201c**工业(汽车/船舶/化工)+海洋经济+消费/文旅**\u201d为主引擎，\u201c外贸/外资\u201d偏弱——靠内需制造+海洋驱动的计划单列。")

# ---- 七、财政 ----
heading1(doc, "七、预算与财政的\u201c含金量\u201d")
para(doc, "**地方一般公共预算收入1340.7亿元、+0.1%**；一般公共预算支出1718.5亿元；民生支出占76.3%（+0.5pct）。")
bullet(doc, "**收入放缓**：+0.1%，反映减税降费/土地/企业利润承压；但税基仍靠工业/海洋/品牌。")
bullet(doc, "**民生硬度高**：民生占比76.3%，科技支出+43.7%、社保就业+5.3%、农林水+7.9%。")
bullet(doc, "**债务/金融稳**：本外币存款3.09万亿、+7.8%，贷款3.4万亿、+6.6%；金融稳健。")
para(doc, "**财政含义**：青岛\u201c收入放缓但民生/科技投入加大\u201d，靠财政结构支撑创新与民生，政策空间有限、需海洋/制造创造税源。")

# ---- 八、民生 ----
heading1(doc, "八、民生底账：人口、收入与城乡")
para(doc, "青岛\u201c十四五\u201d人口稳增：**常住人口增至约1050万人、城镇化率79%（五年+2.5pct）**；居民人均可支配收入62738元、+4.7%（快于多数副省级城市），城乡收入倍差缩小。")
bullet(doc, "**收入**：居民人均可支配收入62738元、+4.7%；城镇71703元、+4.2%，农村33296元、+5.5%（农村增速快于城镇）。")
bullet(doc, "**就业**：城镇新增就业37.3万人（民营吸纳31.5万、占84.4%）。")
bullet(doc, "**社保**：人均预期寿命82.2岁；低保/医保完善。")
bullet(doc, "**人口**：自然增长平稳、依赖净流入集聚。")
para(doc, "**民生含义**：青岛\u201c就业（民营主导）+收入+康寿\u201d全面向好，区域中心强人口吸引力。")

# ---- 九、城乡 ----
heading1(doc, "九、城镇与农村：格局与均衡")
para(doc, "**城镇化率约79%（+2.5pct/五年）**，青岛高度城镇化；农村收入增速（+5.5%）快于城镇（+4.2%），城乡差距收敛。")
bullet(doc, "**城市**：副省级/都市圈核心，海洋城市、品牌消费、国际城市品质。")
bullet(doc, "**农村**：农业生产稳（粮食321.4万吨、水产品110.6万吨）；乡村振兴/县域（即墨、胶州等）制造强。")
para(doc, "**城乡均衡**：青岛\u201c城市强、县域+沿海\u201d，以都市圈+乡村振兴推动城乡融合。")

# ---- 十、人口流 ----
heading1(doc, "十、人口流入与流出")
para(doc, "**青岛呈净流入**：2025年常住约1050万人，在全国人口总体收缩背景下，靠机械增长（省内外迁入）保持增长。")
bullet(doc, "**流入**：制造/海洋/港口/文旅岗位+高品质城市，吸引人才；民营就业84.4%。")
bullet(doc, "**竞争**：与济南/胶东/长三角争夺；青岛靠海洋、收入、港口生态吸引。")
para(doc, "人口方向决定中长期需求与增长；青岛的\u201c高收入+海洋/制造\u201d是其长逻辑。")

# ---- 十一、物价 ----
heading1(doc, "十一、物价与货币环境")
para(doc, "**2025年青岛CPI上涨0.1%**：温和低位，与全国/山东吻合；食品、交通等价格收缩。")
bullet(doc, "**物价**：CPI+0.1%，整体通缩压力可控、需求端偏弱。")
bullet(doc, "**货币/流动性**：本外币存款+7.8%、贷款+6.6%；金融稳健（上市+储贷增长）。")
para(doc, "**物价含义**：青岛通胀温和，消费修复有限；关注服务价格与收入驱动的再通胀。")

# ---- 十二、区域一体化 ----
heading1(doc, "十二、区域一体化：青岛在\u201c青岛都市圈+胶东经济圈+港口开放\u201d里的位置")
para(doc, "青岛处于**山东省会/青岛都市圈+胶东经济圈+沿海开放**核心：既是计划单列/副省级，也是山东/胶东的经济龙头与开放门户。")
bullet(doc, "**青岛都市圈**：总投资3274亿元的52个重点项目；共建园区（济青双圈联动）；市域铁路/公路打通。")
bullet(doc, "**济青一体化**：济青高铁公交化；济济青双圈联合机制，形成省会+港口双子核。")
bullet(doc, "**港口开放**：青岛港全球前列、中欧班列超1000列、上合示范/自贸制度创新；海铁联运全国第一。")
bullet(doc, "**沿黄/一带一路**：融入黄河战略、提升沿黄物流走廊；面向一带一路（占进出口59.8%）。")
para(doc, "**区域含义**：青岛作为\u201c山东港城双核+海洋+开放门户\u201d，靠\u201c港口+海洋+都市圈\u201d带动山东，辐射沿黄与一带一路。")

# ---- 十三、五条主线 ----
heading1(doc, "十三、未来5—10年最值得观察的五条主线")
bullet(doc, "**主线1｜海洋经济/海洋城市**：海洋生产总值+7.5%、2100亿海洋项目；能否建成国家海洋经济发展示范区/海洋强市。")
bullet(doc, "**主线2｜港口/开放**：青岛港全球第4/第5、中欧班列、一带一路。能否持续提升枢纽与航运中心地位。")
bullet(doc, "**主线3｜先进制造/新能源汽车**：汽车+25.9%、船舶/化工；\u201c10+1\u201d产业、新一代信息技术2000亿。能否在高端制造/新能源卡位。")
bullet(doc, "**主线4｜AI+/数字经济**：数字经济核心占GDP 12%+、AI核心产业目标1000亿、1000个AI场景。能否孵化新增长极。")
bullet(doc, "**主线5｜品牌消费/都市圈**：国际消费中心城市、文旅11%+、青岛都市圈。能否把\u201c人口/消费\u201d变成长效引擎。")
para(doc, "这五条，是青岛从\u201c海洋+港口+制造城市\u201d走向\u201c海洋强市+先进制造+开放枢纽\u201d的\u201c主赛道\u201d。")

# ---- 十四、结论 ----
heading1(doc, "十四、最终结论：青岛在\u201c海洋/港口+先进制造+计划单列\u201d里的增长逻辑")
para(doc, "青岛2025年，本质上是\u201c**工业(汽车/船舶/化工)+海洋经济+消费落地、外贸/外资偏弱**\u201d的答卷：GDP17560.67亿、+5.4%、规上工业+7.9%、海洋+7.5%、港口全球第4/第5、社零+3.2%、进出口+0.6%、财政+0.1%。")
para(doc, "只要先进制造/海洋/港口持续，青岛就站在\u201c海洋+制造+开放门户\u201d增长极；若外贸/外资/地产承压，青岛需承受\u201c工业强、外需与地产偏弱\u201d的结构挑战。")
para(doc, "最稳观察信号：**一盯海洋经济/港口（引擎）、二盯汽车/船舶/高端装备（制造）、三盯外贸/外资（开放）、四盯消费/文旅（内需）、五盯人口/都市圈（长期）。**青岛，是\u201c海洋经济+先进制造\u201d的新样本。")

# ---- 附录A ----
heading1(doc, "附录A：主要资料来源与核验路径")
bullet(doc, "青岛市2025年政府工作报告（2025年2月）——目标来源。")
bullet(doc, "《2025年青岛市国民经济和社会发展统计公报》（市统计局）——GDP、工业、海洋、外贸、人口实值。")
bullet(doc, "2026年青岛市政府工作报告（2026年1月）——2025复盘/海洋/先进制造/都市圈。")
bullet(doc, "青岛港、市财政/海洋发展局等。")
heading2(doc, "核验说明")
para(doc, "本报告涉及数据以统计公报/官方口径为准；\u201c海洋/港口/先进制造/计划单列\u201d等以官方口径为准。")

# ---- 附录B ----
heading1(doc, "附录B：建议建立的年度跟踪仪表盘")
para(doc, "建议把下面10个\u201c测脉搏\u201d指标做成年度跟踪表：")
dash = [
    ["#", "指标", "2025基值", "为什么重要"],
    ["1", "GDP增速", "+5.4%（17560.67亿）", "总量与方向"],
    ["2", "规模以上工业增速", "+7.9%", "制造底盘"],
    ["3", "海洋经济增速", "+7.5%", "海洋强市"],
    ["4", "港口吞吐量/集装箱", "7.1亿吨/3087万标箱", "港口枢纽"],
    ["5", "进出口增速", "+0.6%（9128.9亿）", "外贸韧性"],
    ["6", "固定资产投资/工业", "稳健/船舶汽车高速", "投资结构"],
    ["7", "社零增速", "+3.2%（6960.7亿）", "内需消费"],
    ["8", "常住人口/城镇化", "约1050万 / 79%", "人口与城市"],
    ["9", "地方财政收入", "+0.1%（1340.7亿）", "财政质量"],
    ["10", "CPI", "+0.1%", "物价"],
]
table(doc, dash[0], dash[1:], widths=[0.8, 4.6, 4.2, 3.6])
para(doc, "把这10个指标连起来看，海洋/港口/制造（3/4/6）、消费/人口（7/8），都说明青岛在真正换挡。")

# ============ 保存 ============
out = "/Users/x/Desktop/content-prod-lab/reports/青岛市_2025年政府工作报告_深度研究_2026-08-13.docx"
doc.save(out)
print("SAVED:", out)
print("PARAGRAPHS:", len(doc.paragraphs))
print("TABLES:", len(doc.tables))
