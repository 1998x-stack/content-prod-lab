# -*- coding: utf-8 -*-
"""Build 西宁市2025年政府工作报告 深度研究 DOCX, 参照省会系列版式。"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DARK = RGBColor(0x1F, 0x2A, 0x44)
RED = RGBColor(0xB0, 0x1E, 0x2E)
GRAY = RGBColor(0x59, 0x60, 0x69)
BLUE = RGBColor(0x1F, 0x4E, 0x79)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
HEADER_FILL = "1F2A44"
K = "微软雅黑"


def set_run(run, size=11, bold=False, color=DARK, font=K):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    run.font.name = font
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), font)
    rFonts.set(qn("w:hAnsi"), font)
    rFonts.set(qn("w:eastAsia"), font)


def para(doc, text, size=11, bold=False, color=DARK, align=None,
         space_after=6, space_before=0, font=K):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        set_run(r, size=size, bold=bold or (i % 2 == 1), color=color, font=font)
    return p


def heading1(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(16)
    pf.space_after = Pt(8)
    r = p.add_run(text)
    set_run(r, size=16, bold=True, color=RED)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "B01E2E")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def heading2(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(10)
    pf.space_after = Pt(4)
    r = p.add_run(text)
    set_run(r, size=13, bold=True, color=BLUE)
    return p


def table(doc, headers, rows, widths=None, font_size=9.5):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_run(r, size=font_size, bold=True, color=WHITE)
        tcPr = hdr[i]._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), HEADER_FILL)
        tcPr.append(shd)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(val))
            set_run(r, size=font_size, color=DARK)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    return t


def quote_box(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(8)
    pf.left_indent = Pt(12)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), "B01E2E")
    pBdr.append(left)
    pPr.append(pBdr)
    r = p.add_run(text)
    set_run(r, size=11, color=DARK, font="楷体")
    return p


def bullet(doc, text, size=10.5):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.7)
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        set_run(r, size=size, bold=(i % 2 == 1), color=DARK)
    return p


# ================================================================= 文档主体
doc = Document()
sec = doc.sections[0]
sec.page_width = Cm(21)
sec.page_height = Cm(29.7)
sec.top_margin = Cm(2.4)
sec.bottom_margin = Cm(2.2)
sec.left_margin = Cm(2.5)
sec.right_margin = Cm(2.5)

st = doc.styles["Normal"]
st.font.name = K
st.font.size = Pt(11)
st.element.rPr.rFonts.set(qn("w:eastAsia"), K)


# ---- 封面 ----
para(doc, "西宁市2025年政府工作报告", size=24, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=60, space_after=4)
para(doc, "深度研究与\u201c容易被忽视的细节\u201d分析", size=16, bold=True, color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
para(doc, "从\u201c高原生态屏障、绿色算力、清洁能源、盐湖化工与民族文化旅游\u201d重新理解西宁", size=12, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)
para(doc, "\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501", color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
para(doc, "研究对象：2025年西宁市政府工作报告及同期财政、国民经济和社会发展统计资料、2026年官方复盘", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
para(doc, "标定日期：2026-08-13", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
doc.add_page_break()

# ---- 目录 ----
catalog = [
    "执行摘要",
    "一、研究口径与阅读方法",
    "二、先看西宁的特殊底盘：高原生态、绿色算力、清洁能源与民族文化旅游",
    "三、最关键的宏观错位：GDP增速回升但工业/电力温和，投资/财政/地方收入双降",
    "四、容易被忽视、但信号很重的细节（15条）",
    "五、2025年GDP目标 vs 实际：对照表",
    "六、2025年增长是谁撑起来的？（结构归因）",
    "七、预算与财政的\u201c含金量\u201d",
    "八、民生底账：人口、收入与城乡",
    "九、城镇与农村：格局与均衡",
    "十、人口流入与流出",
    "十一、物价与货币环境",
    "十二、区域一体化：西宁在\u201c兰西城市群+青藏高原门户+绿电算力\u201d里的位置",
    "十三、未来5—10年最值得观察的五条主线",
    "十四、最终结论：西宁在\u201c绿电算力+清洁能源+高原文旅\u201d里的增长逻辑",
    "附录A：主要资料来源与核验路径",
    "附录B：建议建立的年度跟踪仪表盘",
]
para(doc, "目　录", size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=10, space_after=10)
for item in catalog:
    para(doc, item, size=12, space_after=4)
doc.add_page_break()

# ---- 执行摘要 ----
heading1(doc, "执行摘要")
para(doc, "**核心判断**　2025年西宁最显著的是\u201cGDP 1914.8亿元、增长3.3%（低于5%目标）、三产占58.7%\u201d、\u201c规上工业+3.2%（制造业+3.9%）\u201d、\u201c进出口56.4亿元/+14.8%、出口+48.8%\u201d、\u201c常住人口247.56万、城镇化81.2%\u201d。这说明西宁在\u201c绿色算力+清洁能源+高原文旅\u201d上求转型，但**总量小、投资-19.1%、财政-22.9%地方**是硬约束。")
para(doc, "把2025年目标（GDP+5%/规上工业、城镇就业3.4万）、2025年统计、2026年前瞻一起看，西宁是\u201c青藏高原门户+绿电算力城市\u201d路径：**绿色算力（算力5300P）、生态、文旅**是特色，但工业与投资双弱。总量1914亿居青海第1。")
para(doc, "最容易记住的一句话：**西宁是\u201c高原生态+绿色算力+清洁能源+民族文化建设\u201d的青海省会，靠\u201c绿电、算力、清洁能源、文旅\u201d转型。**观察西宁，与其只看\u201cGDP 1914亿\u201d，不如看\u201c绿色算力5300P、出口+48.8%、清洁能源、民族文旅（旅客+21.2%）\u201d。")# ---- 一、研究口径与阅读方法 ----
heading1(doc, "一、研究口径与阅读方法")
para(doc, "本报告以\u201c西宁市2025年政府工作报告（2025年1月，石建平作）\u201d为起点，把\u201c2025年GDP目标（5%左右）\u201d与\u201c2025年国民经济和社会发展统计公报实际值（1914.8亿元/+3.3%）\u201d并置对照，再用2026年官方口径作事后验证。")
para(doc, "重点阅读三类证据：**一是指标目标是否达标**；**二是结构（谁贡献增长）是否匹配目标叙事**；**三是官方复盘如何自评、承认问题**。统计口径一般公共预算收入分\u201c全口径\u201d与\u201c地方\u201d。")
para(doc, "统一采用\u201c同比增长%\u201d（GDP为不变价、多数指标可比口径），财政与居民收入多为名义值，文中按原口径转录、不逐项换算。")

# ---- 二、特殊底盘 ----
heading1(doc, "二、先看西宁的特殊底盘：高原生态、绿色算力、清洁能源与民族文化旅游")
para(doc, "**区位与身份**：西宁是青海省省会、青藏高原最大城市、兰西城市群核心之一，是青藏铁路、进藏门户。定位\u201c高原生态屏障\u201d\u201c绿色算力\u201d\u201c清洁能源基地\u201d。")
para(doc, "**产业底盘**：一是绿色算力（2024年算力5300P、标准机架10778架），依托青海绿电和凉爽气候发展\u201c东数西算\u201d；二是清洁能源（光伏/风电、绿电消纳）；三是盐湖化工（盐湖知识产权运营中心、甘河中试基地）；四是民族文化旅游（塔尔寺/湟水中游、入境游客+56%）。")
para(doc, "**人口底盘**：2025年末常住人口247.56万（-0.13万）、城镇化率81.18%（提高0.39个百分点）；是青海第一人口城市。")
para(doc, "**市场与出口**：2025年社零663.4亿/+2.2%；进出口56.4亿元/+14.8%、出口+48.8%；进口-55.9%。")

# ---- 三、核心宏观错位 ----
heading1(doc, "三、最关键的宏观错位：GDP增速回升但工业/总量温和，投资、地方财政双降")
para(doc, "**第一组错位**：2025年GDP目标5%左右，实际1914.8亿/+3.3%（较上年3.1%回升0.2pct）但仍低于目标约1.7pct；总量1914.8亿、青海第一。总量小、增速低是\u201c高原小体量\u201d常态。")
para(doc, "**第二组错位**：规上工业+3.2%、第三产业+3.7%（快于工业）；但**固投-19.1%（工业投资-37.3%）**，投资是最大缺口。")
para(doc, "**第三组错位**：**地方一般公共预算收入103.7亿/-22.9%（骤降）**，财政收入断崖，而全口径213.4亿/-3.9%相对温和——地方财力严重承压。")
para(doc, "一句话：**西宁是\u201c生态+绿电算力、但投资弱、地方财政差\u201d的高原省会**——转型在、总量小、财政紧。")

# ---- 四、15条细节 ----
heading1(doc, "四、容易被忽视、但信号很重的细节（15条）")
bullet(doc, "1. **GDP 1914.8亿/+3.3%**：青海之首，但低于5%目标、增速温和。")
bullet(doc, "2. **规上工业+3.2%、制造业+3.9%**：重工业（有色冶炼+5%）但总盘小。")
bullet(doc, "3. **采矿业+75%**：非金属矿采选是关键增量。")
bullet(doc, "4. **电气机械+24.6%、锂电+56.9%**：新能源电池制造亮点。")
bullet(doc, "5. **固投-19.1%**：工业投资-37.3%、一产-12.3%，投资塌陷。")
bullet(doc, "6. **地方一般公共预算收入/地方103.7亿-22.9%**：财政骤降、税收增值税-52.4%。")
bullet(doc, "7. **出口+48.8%、进口-55.9%**：贸易净回顺、铜箔/光纤。")
bullet(doc, "8. **进出口56.4亿/+14.8%**：总量小但活。")
bullet(doc, "9. **绿色算力5300P**：\u201c东数西算\u201d、算力集群、低空地。")
bullet(doc, "10. **常住247.56万/城镇化81.18%**：城镇化高但总量小。")
bullet(doc, "11. **收入：城镇45605元/+4.1%、农村19369元/+6.0%**，城乡2.35。")
bullet(doc, "12. **CPI -0.2%**：低通胀。")
bullet(doc, "13. **旅游3859.5万人次/+21.2%、收入406.1亿/+20.1%**：入境+56%。")
bullet(doc, "14. **养老参保168.7万**：民生保障稳。")
bullet(doc, "15. **多晶硅-24.5%/单晶硅-15.6%**：硅料价格承压拖累电子材料工业。")# ---- 五、目标 vs 实际对照表 ----
heading1(doc, "五、2025年GDP目标 vs 实际：对照表")
table(doc,
    ["指标", "2025年目标", "2025年实际", "是否达标"],
    [
        ["地区生产总值增速", "5%左右", "1914.8亿元/+3.3%", "未达标"],
        ["规上工业增加值", "（目标未细化）", "+3.2%", "—"],
        ["固定资产投资", "（强化投资）", "-19.1%（转负）", "未达标"],
        ["社会消费品零售总额", "（促消费）", "663.4亿元/+2.2%", "温和"],
        ["一般公共预算收入(地方)", "（未设量化）", "103.7亿元/-22.9%", "大幅下滑"],
        ["城镇/农村人均可支配收入", "与GDP同步", "+4.1%/+6.0%", "农村快于GDP"],
        ["CPI涨幅", "省定目标", "-0.2%", "偏低"],
    ],
    widths=[3.4, 2.5, 5.0, 3.1])
para(doc, "**简评**：2025年西宁目标偏重在生态/民生/就业，硬性经济目标温和。**实际GDP 3.3%低于5%目标**、固投-19.1%、地方财政-22.9%骤降；唯有出口+48.8%、旅游+21%亮眼。")

# ---- 六、结构归因 ----
heading1(doc, "六、2025年增长是谁撑起来的？（结构归因）")
para(doc, "**三大产业**：一产63.7亿/+6.1%、二产728.2亿/+2.3%、三产1122.9亿/+3.7%；三产占GDP 58.7%。增长主要由\u201c三产（金融/物流/信息）+一产\u201d支撑，**二产（工业）只+2.3%**。")
para(doc, "**工业**：全部工业增加值595.3亿/+3.1%、规上+3.2%；**采矿业（非金属矿）+75%、电气机械+24.6%、金属制品修理+13.1%、锂电+56.9%**；硅料（单晶/多晶硅）承压。规上工业利润49.6亿。")
para(doc, "**服务业**：信息传输/软件+11.2%、租赁商务+15.6%、运输仓储邮政+8.8%、科研+3.2%；金融-0.1%、房地产+3.0%。")
para(doc, "**增长归因**：西宁GDP增长主要靠**三产（含绿色算力/信息/物流）+一产+出口、少量工业**；投资（-19.1%）与地产为负。")

# ---- 七、财政 ----
heading1(doc, "七、预算与财政的\u201c含金量\u201d")
para(doc, "2025年全口径一般公共预算收入213.4亿元/-3.9%；**地方一般公共预算收入103.7亿元/-22.9%（骤降）**，其中税收86.4亿/-25.8%、增值税-52.4%、企业所得税-17.7%；一般公共预算支出375.4亿/-6.4%（压缩支出）。")
para(doc, "**结构性**：地方财力\u201c断崖下滑\u201d（税收、增值税 双降），支出同步收缩，反映\u201c减收+化债+清税\u201d。上划中央收入87.1亿/-3.6%也降。")
para(doc, "**含金量**：西宁财政\u201c高度依赖转移支付、地方自给能力弱\u201d，2025年收入/支出双降是财政风险的集中体现。")

# ---------------- 八、民生 ----------------
heading1(doc, "八、民生底账：人口、收入与城乡")
para(doc, "常住人口247.56万/-0.05%（-0.13万）、城镇化率81.18%（提高0.39pct）；城镇新增就业3.7万、登记失业率0.9%。收入：**城镇45605元/+4.1%、农村19369元/+6.0%**，城乡比2.35（缩小0.05）。")
para(doc, "消费支出：全体25951元/+3.8%（城镇29418/+2.8%、农村16193/+7.3%）。")
para(doc, "**民生结论**：收入农村快于城镇、差距收敛；人口总量趋稳、城镇化高；就业稳、社保覆盖面高。")

# ---------------- 九 ----------------
heading1(doc, "九、城镇与农村：格局与均衡")
para(doc, "西宁城镇化率81.2%（青藏高原前列），主城区承载服务/绿电算力；县域（湟中、大通、湟源）发展农业与特色文旅。")
para(doc, "城乡收入比2.35（缩小），农村增速快、均衡边际改善；高原农业（青稞/油菜/牛羊肉）+文旅支撑乡村。")

# =========== 十、人口流入流出 =============
heading1(doc, "十、人口流入与流出")
para(doc, "西宁常住247.56万、微降0.13万，人口规模小、趋稳。自然与机械流动均不显著（青海吸附力有限），高校（青海大学）与绿电算力岗位提供部分吸引。")
para(doc, "**流入**：绿电算力/数字/文旅人才、大学生；**流出**：青壮年外迁至内地。整体\u201c低流动、稳定\u201d，人口红利有限。")

# =========== 十一 -------------=
heading1(doc, "十一、物价与货币环境")
para(doc, "2025年西宁CPI**下降0.2%**（通缩），食品烟酒-1.8%、交通通信-2.3%、衣着+0.8%。")
para(doc, "金融：存款5772.4亿/+4.7%（住户+8.5%）、贷款6172.8亿/+1.9%（住户+10.2%）。“宽中稳、需求弱”。")

# =========== 十二 ===========
heading1(doc, "十二、区域一体化：西宁在\u201c兰西城市群+青藏高原门户+绿电算力\u201d里的位置")
para(doc, "西宁是兰西城市群（兰州—西宁）核心城市、青藏高原门户（青藏铁路/公路进藏）。作为青海省会和最大城市，承接省会功能与高原文旅/物流。")
para(doc, "特色定位：**绿色算力（东数西算、算力5300P）+清洁能源（光伏/风电/绿电）+民族文化旅游**，是\u201c一带一路\u201d与\u201c高原门户\u201d的重要节点。")

# ============ 十三 =============
heading1(doc, "十三、未来5—10年最值得观察的五条主线")
bullet(doc, "1. **绿色算力（东数西算）**：算力5300P→29000P（2025规划）、十万卡集群，绿电×算力能否兑现。")
bullet(doc, "2. **清洁能源+储能**：光伏/风电/虚拟电厂、光储充检，绿电基地。")
bullet(doc, "3. **盐湖化工+生态工业**：盐湖中试、锂电/光纤、资源精深加工。")
bullet(doc, "4. **民族文化旅游**：湟水中游、塔尔寺、入境+56%，高原文旅。")
bullet(doc, "5. **投资与财政修复**：固投-19%后、地方财政-22.9%后，靠绿电算力/基建/项目能否修复。")

# ============ 十四 =============
heading1(doc, "十四、最终结论：西宁在\u201c绿电算力+清洁能源+高原文旅\u201d里的增长逻辑")
para(doc, "**结论**：西宁2025年的\u201c真相\u201d是——**GDP+3.3%（低于目标）、规上工业+3.2%、固投-19.1%、地方财政-22.9%、出口+48.8%**。它正从\u201c传统工业+生态\u201d向\u201c**绿电算力+清洁能源+高原文旅**\u201d转型，但总量小、投资弱、财政压力大。")
para(doc, "**对趋势判断**：绿色算力/清洁能源代表西宁的\u201c未来\u201d，内需与财政代表\u201c约束\u201d。**绿电算力、盐湖化工、文旅**决定潜力；**投资修复+财政/收入**决定韧性。")
para(doc, "**若只看一个指标**：看**地方一般公共预算收入增速（-22.9%）+绿色算力规模（5300P→）**——西宁是\u201c生态与算力强、财政与投资弱\u201d的高原省会，算力与项目的持续落地是扭转低增长的关键。")

# ------------- 附录A -------------
heading1(doc, "附录A：主要资料来源与核验路径")
bullet(doc, "西宁市人民政府《2025年西宁市政府工作报告》（2025年1月，政府工作报告）。")
bullet(doc, "西宁市统计局《2025年西宁市国民经济和社会发展统计公报》（2026年5月）。")
bullet(doc, "2026年西宁市政府工作报告（2026年2月）及绿色算力规划。")
bullet(doc, "青海省GDP与西宁市统计公报、兰西城市群规划交叉核验。")

# ------------- 附录B -------------
heading1(doc, "附录B：建议建立的年度跟踪仪表盘")
bullet(doc, "GDP总量/增速 vs 全年目标（+/-个百分点）。")
bullet(doc, "规上工业增加值及分行业（采掘/电气/锂电/硅料/有色）增速。")
bullet(doc, "固定资产投资（总量/工业/基建/民间/房地产）增速。")
bullet(doc, "社会消费品零售总额、网络零售。")
bullet(doc, "外贸进出口（人民币）、出口、财政对。")
bullet(doc, "一般公共预算收入（全口径/地方）/税收/增值税、财政支出。")
bullet(doc, "常住人口、城镇化率、劳动力流动。")
bullet(doc, "CPI/核心CPI、规上工业企业利润。")
bullet(doc, "旅游人数/旅游总花费、入境游客。")
bullet(doc, "绿色算力（P）、标准机架、储能、光伏装机。")

# ------------- 保存 -------------
out = "/Users/x/Desktop/content-prod-lab/reports/西宁市_2025年政府工作报告_深度研究_2026-08-13.docx"
doc.save(out)
print("SAVED OK 西宁市_2025年政府工作报告_深度研究_2026-08-13.docx")