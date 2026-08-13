# -*- coding: utf-8 -*-
"""Build 中国香港特别行政区2025年深度研究 DOCX, 参照省系列版式。"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DARK = RGBColor(0x1F, 0x2A, 0x44)
RED = RGBColor(0xB0, 0x1E, 0x2E)
GRAY = RGBColor(0x59, 0x60, 0x69)
BLUE = RGBColor(0x1F, 0x4E, 0x79)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
HEADER_FILL = "1F2A44"
K = "微软雅黑"


def set_run(run, size=11, bold=False, color=DARK, font=K):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    run.font.name = font
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), font)
    rFonts.set(qn("w:hAnsi"), font)
    rFonts.set(qn("w:eastAsia"), font)


def para(doc, text, size=11, bold=False, color=DARK, align=None,
         space_after=6, space_before=0, font=K):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        set_run(r, size=size, bold=bold or (i % 2 == 1), color=color, font=font)
    return p


def heading1(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(16)
    pf.space_after = Pt(8)
    r = p.add_run(text)
    set_run(r, size=16, bold=True, color=RED)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "B01E2E")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def heading2(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(10)
    pf.space_after = Pt(4)
    r = p.add_run(text)
    set_run(r, size=13, bold=True, color=BLUE)
    return p


def table(doc, headers, rows, widths=None, font_size=9.5):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_run(r, size=font_size, bold=True, color=WHITE)
        tcPr = hdr[i]._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), HEADER_FILL)
        tcPr.append(shd)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(val))
            set_run(r, size=font_size, color=DARK)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    return t


def quote_box(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(8)
    pf.left_indent = Pt(12)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), "B01E2E")
    pBdr.append(left)
    pPr.append(pBdr)
    r = p.add_run(text)
    set_run(r, size=11, color=DARK, font="楷体")
    return p


def bullet(doc, text, size=10.5):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.7)
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        set_run(r, size=size, bold=(i % 2 == 1), color=DARK)
    return p


# ================================================================= 文档主体
doc = Document()
sec = doc.sections[0]
sec.page_width = Cm(21)
sec.page_height = Cm(29.7)
sec.top_margin = Cm(2.4)
sec.bottom_margin = Cm(2.2)
sec.left_margin = Cm(2.5)
sec.right_margin = Cm(2.5)

st = doc.styles["Normal"]
st.font.name = K
st.font.size = Pt(11)
st.element.rPr.rFonts.set(qn("w:eastAsia"), K)



# ---- 封面 ----
para(doc, "中国香港特别行政区2025年深度研究", size=24, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=60, space_after=4)
para(doc, "经济与\u201c容易被忽视的细节\u201d分析", size=16, bold=True, color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
para(doc, "从\u201c国际金融中心、港深合作、北部都会区与新质经济\u201d重新理解香港", size=12, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)
para(doc, "\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501", color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
para(doc, "研究对象：2025年中国香港经济数据、施政报告/预算案及国际金融中心定位、北部都会区进展", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
para(doc, "标定日期：2026-08-13", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
doc.add_page_break()

# ---- 目录 ----
catalog = [
    "执行摘要",
    "一、研究口径与阅读方法",
    "二、先看香港的特殊底盘：国际金融/贸易/航运中心+港深联动",
    "三、最关键的宏观结构：GDP破3.3万亿港元、出口/投资强，消费/楼市趋稳",
    "四、容易被忽视、但信号很重的细节（15条）",
    "五、2025年GDP目标 vs 实际：对照表",
    "六、2025年增长是谁撑起来的？（结构归因）",
    "七、财政与预算的\u201c含金量\u201d",
    "八、民生与人口",
    "九、楼宇与资产（楼市/金融）",
    "十、人口流入与流出",
    "十一、物价与货币环境（联系汇率）",
    "十二、区域一体化：港深合作与北部都会区",
    "十三、未来5—10年最值得观察的五条主线",
    "十四、最终结论",
    "附录A：主要资料来源与核验路径",
    "附录B：建议建立的年度跟踪仪表盘",
]
para(doc, "目　录", size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=10, space_after=10)
for item in catalog:
    para(doc, item, size=12, space_after=4)
doc.add_page_break()

# ---- 执行摘要 ----
heading1(doc, "执行摘要")
para(doc, "**核心判断**　2025年中国香港最显著的是\u201c本地生产总值(GDP)约3.31万亿港元、实质增长3.5%（连续第3年扩张，快于2024年2.6%）\u201d、\u201c货品出口+12.0%\u201d、\u201c新股上市集资额全球第一、恒指+28%\u201d、\u201c2025年访港旅客约4500万人次\u201d。这说明香港经济在\u201c外部出口强劲+本地投资/消费转稳\u201d下持续复苏。")
para(doc, "把2025年首三季/全年官方统计（政府统计处）、财政预算、施政报告2025放在一起看，香港呈现出\u201c全球枢纽+中国内地融合（北部都会区/港深）\u201d双引擎。国际（货品出口+12%）、本地（投资+4.5%、私人消费+1.6%）同时改善，是连续第三年扩张。")
para(doc, "最容易记住的一句话：**香港是\u201c国际金融/贸易/航运中心+港深/北部都会\u201d，靠\u201c出口+服务+金融\u201d恢复，靠\u201c创科+北部都会区+中国内地融合\u201d升级。**观察香港，不只是看\u201cGDP 3.3万亿港元\u201d，更应看\u201c出口+12%、恒指+28%、新股集资全球第一、房价回稳\u201d。")

# ---- 一、研究口径与阅读方法 ----
heading1(doc, "一、研究口径与阅读方法")
heading2(doc, "1.1 本报告的核心底稿")
bullet(doc, "**2025全年GDP预先估计（政府统计处2026-01）**——3.5%、第4季3.8%。")
bullet(doc, "**2025年《施政报告》及2026-27财政预算案**——方向与财政。")
bullet(doc, "「2025年中后期经济展望」——简略，以官方口径为准。")
heading2(doc, "1.2 阅读方法：显性—数据—长期")
para(doc, "**关键判别**：数据优先。香港2025年GDP约3.5%（4季3.8%）连续扩张，出口/投资/服务强、私人消费温和（+1.6%）。观察香港穿透GDP看\u201c枢纽/出口/金融/港深\u201d。")

# ---- 二、底盘 ----
heading1(doc, "二、先看香港的特殊底盘：国际金融/贸易/航运中心+港深联动")
para(doc, "香港作为\u201c**国际金融中心+贸易+航运中心+离岸人民币枢纽**\u201d、全球少有的自由港，连接中国内地与全球。")
bullet(doc, "**国际金融中心**：全球前三金融中心，金融市场活跃（恒指+28%、新股集资全球第一）。")
bullet(doc, "**贸易/航运**：货品出口+12%、集装箱枢纽、航空/航运中心。")
bullet(doc, "**服务输出**：+6.3%，金融/旅游/专业服务。")
bullet(doc, "**港深融合**：北部都会区、河套深港创科、大湾区一体化。")
para(doc, "这一底板决定2025年增长\u201c主动力\u201d：**出口（电子/亚洲）+金融/服务+投资（+4.5%）**层层加力；长远靠\u201c创科+北部都会区+中国内地融合\u201d。")

# ---- 三、宏观结构 ----
heading1(doc, "三、最关键的宏观结构：GDP破3.3万亿港元、出口/投资强，消费/楼市趋稳")
para(doc, "香港2025年最值得咀嚼的结构，是\u201c**货品出口/服务+投资强、私人消费/楼市温和**\u201d。GDP约3.31万亿港元、实质+3.5%（连续第3年扩张）。")
bullet(doc, "**GDP**：3.31万亿港元（约3.05万亿人民币）、+3.5%（2024年2.6%）；第4季+3.8%。")
bullet(doc, "**货品出口**：+12.0%（电子/亚洲需求）；进口+12.6%。")
bullet(doc, "**服务输出**：+6.3%（访港旅游/金融）。")
bullet(doc, "**投资（固定资本形成）**：+4.5%（第4季+10.9%）；私人消费+1.6%。")
bullet(doc, "**资产/文旅**：恒指+28%、新股集资全球第一；房价+3.3%（结束跌势）、租金+4.3%。")
para(doc, "**为什么读这条**：香港经济\u201c外部出口+金融/投资强、本地消费/楼市趋稳\u201d，外部需求与金融是主动力。")

# ---- 四、15条细节 ----
heading1(doc, "四、容易被忽视、但信号很重的细节（15条）")
para(doc, "以下15条在2025年官方数据/报告里，但常被\u201cGDP 3.5%\u201d等掩盖。它们是判断香港2025之后5—10年的关键小信号。")
bullet(doc, "**1. 货品出口+12.0%、电子主导**：出口回稳是最大拉动力。")
bullet(doc, "**2. 服务输出+6.3%**：金融/旅游/专业服务，枢纽含金量。")
bullet(doc, "**3. 新股集资额全球第一（超2800亿港元）**：IPO回暖、国际资金回流。")
bullet(doc, "**4. 恒指全年+28%、日均成交近2500亿港元**：资本市场活跃。")
bullet(doc, "**5. 房价+3.3%（结束三年跌势）、租金+4.3%**：楼市企稳。")
bullet(doc, "**6. 访港旅客2025年约4500万人次**：旅游复苏。")
bullet(doc, "**7. 私人消费+1.6%（第2季起改善）**：内需温和修复。")
bullet(doc, "**8. 投资+4.5%（第4季+10.9%）**：资本开支回稳。")
bullet(doc, "**9. 北部都会区/港深创科（河套）**：未来增长极。")
bullet(doc, "**10. 离岸人民币枢纽/互联互通**：金融再深化。")
bullet(doc, "**11. 2026预测2.5%-3.5%**：增长中枢换档。")
bullet(doc, "**12. 中国内地/大湾区一体**：市场联动。")
bullet(doc, "**13. 创科（第三代半导体/AI）**：新质生产力。")
bullet(doc, "**14. 失业率约3%+低位**：就业稳定。")
bullet(doc, "**15. 国际航运/航空枢纽**：转口+航运服务。")
para(doc, "", size=10, space_after=2)

# ---- 五、对照表 ----
heading1(doc, "五、2025年GDP 目标 vs 实际：对照表")
para(doc, "香港框架与其他省不同：政府统计处按季度发布、不设硬性年度GDP目标（预算案给预测）。")
tb = [
    ["指标", "2024实际", "2025实际（预先估计）", "方向"],
    ["GDP实质增速", "+2.6%", "+3.5%（第4季+3.8%）", "加"],
    ["货品出口", "低增", "+12.0%", "大幅回升"],
    ["服务输出", "—", "+6.3%", "升"],
    ["私人消费", "—", "+1.6%", "升"],
    ["本地固定资本形成", "—", "+4.5%", "加快"],
]
table(doc, tb[0], tb[1:], widths=[3.4, 3.2, 4.6, 3.6])
para(doc, "", size=10, space_after=2)
para(doc, "**要点**：香港GDP连续第3年扩张、出口/投资/服务恢复，增长动能增强。")

# ---- 六、结构归因 ----
heading1(doc, "六、2025年增长是谁撑起来的？（结构归因）")
heading2(doc, "6.1 出口/外部")
para(doc, "**货品出口+12.0%**（电子/亚洲）是主动力；服务输出+6.3%（金融/旅游）。")
heading2(doc, "6.2 投资/资本")
para(doc, "固定资本形成+4.5%（第4季+10.9%），投资/科创/基建带动。")
heading2(doc, "6.3 本地消费")
para(doc, "私人消费+1.6%（消费市道自第2季改善），温和。")
heading2(doc, "6.4 资产/金融周期")
para(doc, "恒指+28%、新股全球第一、房价企稳——金融与资产周期回暖。")
para(doc, "**一句话归因**：香港2025年靠\u201c**出口+服务+投资（金融/资产）**\u201d实现连续第3年扩张，外部/金融双引擎。")

# ---- 七、财政 ----
heading1(doc, "七、财政与预算的\u201c含金量\u201d")
para(doc, "2025-26财政年度，特区政府以\u201c巩固复元\u201d为财政主线（库房在变动后略有赤字/存续弹性），财政政策靠\u201c地盘/派息/税基\u201d平衡。")
bullet(doc, "**财政框架**：以\u201c量入为出、巩固复元\u201d为原则，靠稳定收入（利得税/印花税/地价）+支出纪律。")
bullet(doc, "**资产周期**：恒指+28%、IPO全球第一、楼市企稳，印花税/地价收益改善。")
bullet(doc, "**财政储备**：充足、可持续，支持创科/民生/北部都会。")
para(doc, "**财政含义**：香港财政\u201c靠金融/资产周期回暖\u201d，在外部出口与资本市场回暖下改善，空间仍有余力。")

# ---- 八、民生与人口 ----
heading1(doc, "八、民生与人口")
para(doc, "香港人口约749万（统计处2025年中）、全球高收入城市；民生关注\u201c住房/社会治安/基层\u201d。")
bullet(doc, "**人口**：约750万，人口结构老龄化、依赖人才计划。")
bullet(doc, "**就业**：失业率低位（约3%+）、就业稳定。")
bullet(doc, "**住房**：楼市企稳+房价3.3%、完善置业/入住（简约公屋/青年宿舍）。")
bullet(doc, "**民生**：交通补助、医疗、北部都会/拓展生活空间。")
para(doc, "**民生含义**：香港\u201c就业稳、收入高、楼市企稳\u201d，民生以\u201c住/基层\u201d与资本周期联动。")

# ---- 九、楼宇与资产 ----
heading1(doc, "九、楼宇与资产（楼市/金融）")
bullet(doc, "**楼市**：住宅房价2025年升3.3%（结束前三年跌势）、租金+4.3%——资产回稳。")
bullet(doc, "**股市**：恒指全年+28%、日均成交近2500亿港元、新股集资全球第一（超2800亿港元）。")
bullet(doc, "**金融**：国际金融中心，离岸人民币/互联互通，美债/IPO等。")
para(doc, "**资产含义**：2025年香港\u201c股楼双升、IPO全球第一\u201d，资产与金融周期强劲改善。")

# ---- 十、人口流 ----
heading1(doc, "十、人口流入与流出")
para(doc, "香港人口结构稳定（约750万750万、高收入），通过\u201c人才计划+港深融合\u201d吸纳人才；本地老龄化。")
bullet(doc, "**人才**：高才/优才/人才计划、国际人才。")
bullet(doc, "**北上/大湾区**：居民北上消费/康居，但人才流入仍增。")
para(doc, "人口方向：香港靠人才计划与大湾区融合稳人口、扩大劳动力。")

# ---- 十一、物价与货币 ----
heading1(doc, "十一、物价与货币环境（联系汇率）")
para(doc, "**2025年香港整体通胀（综合CPI）温和**（全年约略正、受外围压抑/服务成本），联系汇率下随美元/环境。")
bullet(doc, "**通胀**：综合消费物价指数温和、约1%上下。")
bullet(doc, "**货币**：联系汇率制、金管局管理，外汇储备充足。")
para(doc, "**物价含义**：香港\u201c通胀轻丶联系汇率稳\u201d，助力资金与资产周期。")

# ---- 十二、区域一体化 ----
heading1(doc, "十二、区域一体化：港深合作与北部都会区")
para(doc, "香港与\u201c**深港/大湾区+北部都会区**\u201d深度融合，既是国际枢纽也强化内地市场联动。")
bullet(doc, "**港深/河套**：深港科技创新合作区（河套）、创科一体化。")
bullet(doc, "**北部都会区**：拓展土地+创科（北部之都），新增长极。")
bullet(doc, "**大湾区**：互联互通（口岸/高铁）、跨境金融/人才/基建一体化。")
bullet(doc, "**中国内地市场**：离岸人民币、互联互通、北向资金。")
para(doc, "**区域含义**：香港靠\u201c港深+北部都会区+大湾区\u201d双引擎，巩固国际与内地枢纽。")

# ---- 十三、五主线 ----
heading1(doc, "十三、未来5—10年最值得观察的五条主线")
bullet(doc, "**主线1｜国际金融/资本市场**：新股第一、恒指/IPO。能否重回全球IPO/金融地位。")
bullet(doc, "**主线2｜创科/第三代半导体**：新质生产力、第三代半导体/AI。能否孵新产业。")
bullet(doc, "**主线3｜北部都会区/港深**：科创土地+融合。能否成新增长极。")
bullet(doc, "**主线4｜出口/服务枢纽**：电子出口/金融服务。能否保枢纽。")
bullet(doc, "**主线5｜楼市/资产周期**：价格+租金回稳。能否稳财政/消费。")
para(doc, "这五条，是香港从\u201c金融/贸易/楼市周期\u201d走向\u201c创科+港深+多元\u201d的\u201c主赛道\u201d。")

# ---- 十四、结论 ----
heading1(doc, "十四、最终结论")
para(doc, "香港2025年，本质上是\u201c**出口/服务+金融/资产+投资回稳**\u201d的连续第3年扩张：GDP约3.31万亿港元、+3.5%、货品出口+12%、服务+6.3%、恒指+28%、新股全球第一、房价+3.3%。")
para(doc, "只要国际枢纽/金融/出口持续，香港就站在\u201c全球中心+大湾区\u201d增长极；若外部不确/流动性波动，需承受高开放度波动。")
para(doc, "最稳观察信号：**一盯出口/服务（外需）、二盯恒指/IPO（金融）、三盯楼市（资产）、四盯北部都会区/港深（新极）、五盯人才/人口（长期）。**香港，是\u201c国际枢纽+内地融合\u201d独特样本。")

# ---- 附录A ----
heading1(doc, "附录A：主要资料来源与核验路径")
bullet(doc, "政府统计处《2025年第四季及全年本地生产总值预先估计》（2026-01）。")
bullet(doc, "《2025年施政报告》、2026-27财政预算案。")
bullet(doc, "中国香港金融管理局/投资推广署口径。")
heading2(doc, "核验说明")
para(doc, "本报告以官方统计/政府口径为准；涉\u201c金融/北部都会/港深\u201d等以官方为准。")

# ---- 附录B ----
heading1(doc, "附录B：建议建立的年度跟踪仪表盘")
para(doc, "建议把下面10个\u201c测脉搏\u201d指标做成年度跟踪表：")
dash = [
    ["#", "指标", "2025基值", "为什么重要"],
    ["1", "GDP实质增速", "+3.5%", "总量与方向"],
    ["2", "货品出口", "+12.0%", "外需"],
    ["3", "服务输出", "+6.3%", "枢纽"],
    ["4", "恒指/IPO", "+28%/全球第一", "金融"],
    ["5", "住宅价格", "+3.3%", "资产周期"],
    ["6", "私人消费", "+1.6%", "内需"],
    ["7", "本地投资", "+4.5%", "资本开支"],
    ["8", "常住人口/失业率", "約750万/约3%", "人口/就业"],
    ["9", "综合CPI", "温和", "通胀"],
    ["10", "访港旅客", "约4500万人次", "旅游"],
]
table(doc, dash[0], dash[1:], widths=[0.8, 4.6, 4.4, 4.0])
para(doc, "把这10个连起来看，出口/金融/资产（2/4/5）、北部/港深（8），都说明香港在重新扩张。")

# ============ 保存 ============
out = "/Users/x/Desktop/content-prod-lab/reports/中国香港_2025年深度研究_2026-08-13.docx"
doc.save(out)
print("SAVED:", out)
print("PARAGRAPHS:", len(doc.paragraphs))
print("TABLES:", len(doc.tables))
