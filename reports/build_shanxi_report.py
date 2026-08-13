# -*- coding: utf-8 -*-
"""Build 山西省2025年政府工作报告 深度研究 DOCX, 参照省市系列版式。"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DARK = RGBColor(0x1F, 0x2A, 0x44)
RED = RGBColor(0xB0, 0x1E, 0x2E)
GRAY = RGBColor(0x59, 0x60, 0x69)
BLUE = RGBColor(0x1F, 0x4E, 0x79)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
HEADER_FILL = "1F2A44"
K = "微软雅黑"


def set_run(run, size=11, bold=False, color=DARK, font=K):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    run.font.name = font
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), font)
    rFonts.set(qn("w:hAnsi"), font)
    rFonts.set(qn("w:eastAsia"), font)


def para(doc, text, size=11, bold=False, color=DARK, align=None,
         space_after=6, space_before=0, font=K):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        set_run(r, size=size, bold=bold or (i % 2 == 1), color=color, font=font)
    return p


def heading1(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(16)
    pf.space_after = Pt(8)
    r = p.add_run(text)
    set_run(r, size=16, bold=True, color=RED)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "B01E2E")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def heading2(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(10)
    pf.space_after = Pt(4)
    r = p.add_run(text)
    set_run(r, size=13, bold=True, color=BLUE)
    return p


def table(doc, headers, rows, widths=None, font_size=9.5):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_run(r, size=font_size, bold=True, color=WHITE)
        tcPr = hdr[i]._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), HEADER_FILL)
        tcPr.append(shd)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(val))
            set_run(r, size=font_size, color=DARK)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    return t


def quote_box(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(8)
    pf.left_indent = Pt(12)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), "B01E2E")
    pBdr.append(left)
    pPr.append(pBdr)
    r = p.add_run(text)
    set_run(r, size=11, color=DARK, font="楷体")
    return p


def bullet(doc, text, size=10.5):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.7)
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        set_run(r, size=size, bold=(i % 2 == 1), color=DARK)
    return p


# ================================================================= 文档主体
doc = Document()
sec = doc.sections[0]
sec.page_width = Cm(21)
sec.page_height = Cm(29.7)
sec.top_margin = Cm(2.4)
sec.bottom_margin = Cm(2.2)
sec.left_margin = Cm(2.5)
sec.right_margin = Cm(2.5)

st = doc.styles["Normal"]
st.font.name = K
st.font.size = Pt(11)
st.element.rPr.rFonts.set(qn("w:eastAsia"), K)



# ---- 封面 ----
para(doc, "山西省2025年政府工作报告", size=24, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=60, space_after=4)
para(doc, "深度研究与“容易被忽视的细节”分析", size=16, bold=True, color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
para(doc, "从“煤炭与能源、新型综合能源基地、生态修复与产业转型”重新理解山西", size=12, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)
para(doc, "━━━━━━━━━━━━━━━━", color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
para(doc, "研究对象：2025年山西省政府工作报告及同期财政、国民经济和社会发展统计资料、2026年官方复盘", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
para(doc, "标定日期：2026-08-13", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
doc.add_page_break()

# ---- 目录 ----
catalog = [
    "执行摘要",
    "一、研究口径与阅读方法",
    "二、先看山西的特殊底盘：煤炭能源、新型综合能源、生态修复与产业转型",
    "三、最关键的宏观错位：GDP破2.55万亿、新能源/装备强，但工业利润/外贸/人口弱",
    "四、容易被忽视、但信号很重的细节（15条）",
    "五、2025年GDP目标 vs 实际：对照表",
    "六、2025年增长是谁撑起来的？（结构归因）",
    "七、预算与财政的“含金量”",
    "八、民生底账：人口、收入与城乡",
    "九、城镇与农村：格局与均衡",
    "十、人口流入与流出",
    "十一、物价与货币环境",
    "十二、区域一体化：山西在“国家能源基地+京津冀对接+中部崛起”里的位置",
    "十三、未来5—10年最值得观察的五条主线",
    "十四、最终结论：山西在“煤炭能源+新型综合能源+生态修复+产业转型”里的增长逻辑",
    "附录A：主要资料来源与核验路径",
    "附录B：建议建立的年度跟踪仪表盘",
]
para(doc, "目　录", size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=10, space_after=10)
for item in catalog:
    para(doc, item, size=12, space_after=4)
doc.add_page_break()

# ---- 执行摘要 ----
heading1(doc, "执行摘要")
para(doc, "**核心判断**　如果只看新闻标题，2025年山西最显眼的是“GDP破2.55万亿、增长4.0%（较上年加快1.8pct）”、“原煤产量13.09亿吨（+2.4%）”和“新能源/清洁能源装机占55.1%”。但这份研究真正值得深读的，是这座“煤炭能源+新型综合能源+生态修复”的能源大省，如何在煤价回落、工业利润（-36.3%）、财政（-9.1%）与人口自然负增（-3.32‰）的背景下，靠“能源保供+新质生产力+装备制造”实现转型回升。")
para(doc, "把2025年初设定的目标（GDP增长5%左右）、2025年《国民经济和社会发展统计公报》、以及2026年报告对2025年的复盘放在一起看，山西呈现清晰暗线：**从“煤炭/资源”的旧底盘，向“新型综合能源+装备/新材料+生态修复/绿色转型”转型**。旧引擎（煤价/传统能源）在调整；新引擎（新能源/机器人/光伏、装备制造、新材料、软件）被要求更快补位。")
para(doc, "因此，本报告不按政府工作报告逐段复述，而采用“显性表述—同期数据—制度含义—长期影响”的方式，专门提取容易被忽视、但对判断山西未来5—10年发展模式更有价值的信号。")
para(doc, "最容易记住的一句话：**山西是“煤炭能源保供+新型综合能源+产业转型”的能源样本，靠“能源安全+新质生产力+生态修复”撑起转型。**观察山西，与其看“GDP 2.55万亿”，不如看“原煤13亿、新能源装机55.1%、能源革命、装备/新材料、生态修复”这几张名片。")
heading2(doc, "一页速览：2025年山西经济的“表与里”")
table(doc,
    ["维度", "表面现象", "更深层信号"],
    [
        ["增长", "GDP 25495.68亿、+4.0%", "一产5.5%、二产40.4%、三产54.0%"],
        ["产业", "规上工业+4.9%", "采掘+6.3%、计算机+1.2倍"],
        ["能源", "原煤13.09亿吨+2.4%/焦炭+10.4%", "新能源装机占55.1%"],
        ["投资", "固定资产投资-0.4%", "民间+2.8%、地产-9.9%"],
        ["财政", "一般公共预算收入3218.6亿、-9.1%", "税收-6.9%"],
        ["消费", "社零8030.9亿、+4.7%", "通讯+40.6%、家电+19.7%"],
        ["人口", "常住3424.01万、-21.95万、城镇化67.35%", "自然增-3.32‰"],
        ["开放", "能源保供/京津冀协同/中部城市群", "新质生产力/装备"],
    ],
    widths=[2.2, 5.4, 8.4])
para(doc, "", size=10, space_after=2)

# ---- 一、研究口径与阅读方法 ----
heading1(doc, "一、研究口径与阅读方法")
heading2(doc, "1.1 本报告的三份核心底稿")
bullet(doc, "**2025年《政府工作报告》**：2025年1月在省十四届人大三次会议上作，给出全年预期目标（GDP增长5%左右、固投6%以上、社零5%以上、城镇新增就业45万人）。")
bullet(doc, "**2025年《国民经济和社会发展统计公报》**：山西省统计局2026年3月25日发布，给出全年实际执行数与增速，是“事后验证”的锚点。")
bullet(doc, "**2026年山西省政府工作报告及官方复盘**：对2025执行结果的追认，交叉验证“目标—执行—再评价”三段式。")
heading2(doc, "1.2 阅读方法：显性—数据—制度—长期")
para(doc, "每一章按“**显性表述 → 同期数据 → 制度含义 → 长期影响**”四层展开。其中“制度含义”回答“政府为什么这么排优先序”，“长期影响”回答“这对未来5—10年意味着什么”。")
para(doc, "关键判别：**当报告使用的“增长词”和统计公报里的实际数不一致时，数据优先。**例如2025年GDP目标5%左右、实际4.0%略低于目标；固投目标6%以上、实际-0.4%、未达标；社零目标5%以上、实际4.7%。差异反映：山西“能源/装备/新质生产力强，煤价/工业利润/外贸弱”。")
heading2(doc, "1.3 指标取舍与单位说明")
para(doc, "本报告统一以“2025年统计公报+2026年官方复盘”为口径，增速均为可比价，绝对数为人民币，能源/煤炭以万吨/亿m³计。GDP以初步核算为准。")

# ---- 二、底盘 ----
heading1(doc, "二、先看山西的特殊底盘：煤炭能源、新型综合能源、生态修复与产业转型")
para(doc, "山西的地盘，取决于它作为“**国家能源保供基地+新型综合能源+生态修复/产业转型**”的特殊定位。它是中国煤炭/能源的“压舱石”。")
bullet(doc, "**煤炭/能源**：原煤产量130912.1万吨（+2.4%）、焦炭+10.4%、非常规天然气182.3亿m³（+8.9%）、发电量4704.5亿千瓦时；煤炭先进产能占比84%。")
bullet(doc, "**新型综合能源**：发电装机16414.3万千瓦（+11.7%），新能源/清洁能源装机占55.1%（太阳能+42.4%、风电+8.3%）；外送绿电交易98.9亿千瓦时（+31.3%）、万元GDP能耗-2.0%。")
bullet(doc, "**产业转型/新质生产力**：计算机及办公设备制造+1.2倍、新能源产业+25.3%、新材料+6.2%；工业机器人+3.7倍、光伏电池+2.3倍、服务器+47.8%；装备制造利润+1.1倍。")
bullet(doc, "**生态修复/绿色**：PM2.5平均30μg/m³、优良天数81.4%；94个国考断面水质优良比例98.9%；造林227.3千公顷、森林覆盖率22.8%，推进生态修复。")
para(doc, "**制度含义**：山西把“能源安全（煤炭保供）、新型综合能源、装备/新材料、生态修复”当核心资产，并深入推进“能源革命/转型”。")

# ---- 三、宏观错位 ----
heading1(doc, "三、最关键的宏观错位：GDP破2.55万亿、新能源/装备强，但工业利润/外贸/人口弱")
para(doc, "2025年山西GDP 25495.68亿元、+4.0%（一产+4.5%、二产+3.1%、三产+4.5%），增速较上年加快1.8pct。表面看“转型回升”，但拆开看是“**新能源/装备/煤炭保供强、工业利润/外贸/人口弱**”的错位：")
para(doc, "**强的部分**：规上工业+4.9%（采掘+6.3%、计算机+1.2倍）、新能源发电+24.2%；重点景区游客+14.8%；民间投资+2.8%；亿元以上项目4121个。")
para(doc, "**弱的部分**：规上工业利润-36.3%、营收-9.7%；进出口1372.4亿、-21.2%（出口-34.3%）；财政收入3218.6亿、-9.1%；人口自然增-3.32‰；PPI-9.8%。")
para(doc, "**核心错位一句话**：山西“能源保供/新质生产力/装备强（原煤+2.4%、机器人+3.7倍），但煤价/工业利润/外贸/人口弱”。2026年“高端能源化工+装备/新材料+生态修复、稳就业/外贸”是重点。")
table(doc,
    ["强引擎", "数据", "弱项", "数据"],
    [
        ["计算机及办公设备", "+1.2倍", "规上工业利润", "-36.3%"],
        ["新能源产业", "+25.3%", "进出口总额", "-21.2%"],
        ["工业机器人产量", "+3.7倍", "一般公共预算收入", "-9.1%"],
        ["煤炭先进产能占比", "84%", "人口自然增长率", "-3.32‰"],
        ["民间投资", "+2.8%", "PPI", "-9.8%"],
    ],
    widths=[3.8, 3.0, 3.6, 3.0])
para(doc, "**错位结论**　山西的增长“强能源依赖、弱工业利润/外贸”。转型成效明显（装备/新材料/机器人/光伏），但煤价回落带来财政/利润压力。2026年“高端制造+能源转型+稳外贸/就业”是重点。")

# ---- 四、被忽视细节 ----
heading1(doc, "四、容易被忽视、但信号很重的细节（15条）")
details = [
    ("1", "GDP 25495.68亿、+4.0%（加快1.8pct）", "总量/增速回升。"),
    ("2", "原煤13.09亿吨+2.4%、焦炭+10.4%", "能源保供。"),
    ("3", "新能源/清洁能源装机占55.1%（太阳能+42.4%）", "能源转型。"),
    ("4", "新能源发电量+24.2%、外送绿电+31.3%", "绿色电力。"),
    ("5", "计算机及办公设备制造+1.2倍", "新质生产力。"),
    ("6", "工业机器人+3.7倍、光伏电池+2.3倍", "先进制造。"),
    ("7", "装备制造利润+1.1倍、新能源发电利润+1.2倍", "装备/新能源盈利。"),
    ("8", "亿元以上项目4121个（新兴）", "重大项目。"),
    ("9", "社零+4.7%、通讯+40.6%、家电+19.7%", "消费升级。"),
    ("10", "网上零售额+15.0%、快递+32.4%", "数字/物流。"),
    ("11", "PM2.5降至30μg/m³、优良天数81.4%", "生态改善。"),
    ("12", "常住3424.01万、-21.95万、城镇化67.35%", "人口/城镇化。"),
    ("13", "居民收入33923元、+4.6%、农村+5.5%", "收入稳、农村快。"),
    ("14", "旅游/重点景区游客+14.8%", "文旅/旅游。"),
    ("15", "外贸出口-34.3%、进口-0.7%", "外贸承压。"),
]
for d in details:
    para(doc, "**%s**" % d[0], size=11)
    para(doc, d[1], size=10.5, color=GRAY)

    para(doc, "**备注**　这条“错位”在山西尤其鲜明：增长靠能源保供/新质生产力/装备，但煤价/工业利润/外贸/人口弱。2026年若煤价企稳、外贸修复，增长可能从“能源单极”走向“能源+装备/新能源+内外需”多极。这条细节，正是山西2026年最难也最值得盯的“变量”。", size=10.5)

# ---- 五、2025年目标 vs 实际 对照 ----
heading1(doc, "五、2025年GDP目标 vs 实际：对照表")
t5 = [
    ["指标", "2025目标", "2025实际", "达标判定"],
    ["GDP增速", "5%左右", "+4.0%", "略低于目标"],
    ["固定资产投资增速", "6%以上", "-0.4%", "大幅未达标"],
    ["社会消费品零售增速", "5%以上", "+4.7%", "基本达标"],
    ["城镇新增就业", "45万人", "47.3万人", "超额"],
    ["居民消费价格(CPI)", "涨幅2%左右", "+0.1%", "远低于"],
    ["进出口", "力争正增长", "-21.2%", "未达标"],
]
table(doc, t5[0], t5[1:], widths=[3.4, 3.0, 3.0, 3.6])
para(doc, "**对照结论**　目标“偏进取”，实际“有分化”：GDP 4.0%（低于5%目标）、固投-0.4%（未达6%）、进出口-21.2%（未达）；就业47.3万超额、CPI仅0.1%。能源/装备/新质生产力接住，煤价/外贸/财政是短板。")

# ---- 六、增速分项支撑 ----
heading1(doc, "六、2025年增长是谁撑起来的？（结构归因）")
para(doc, "GDP+4.0%背后，是“**能源保供/新质生产力强、工业利润/外贸弱**”的结构。把增长贡献拆开看：")
g6 = [
    ["引擎", "贡献/状态", "说明"],
    ["规上工业", "+4.9%", "采掘+6.3%、计算机+1.2倍"],
    ["煤炭保供", "原煤+2.4%/焦炭+10.4%", "能源安全"],
    ["新能源", "装机占55.1%/发电+24.2%", "能源革命"],
    ["装备/新材料", "装备利润+1.1倍/新材料+6.2%", "产业转型"],
    ["消费/社零", "+4.7%", "通讯+40.6%、家电+19.7%"],
    ["旅游/文旅", "重点景区+14.8%", "文旅"],
    ["规上工业利润", "-36.3%", "煤价/盈利弱"],
    ["一般公共预算收入", "-9.1%", "财政压力"],
    ["进出口", "-21.2%", "外贸承压"],
    ["人口自然增长率", "-3.32‰", "人口约束"],
]
table(doc, g6[0], g6[1:], widths=[3.6, 3.6, 6.0])
para(doc, "**一句话**　增长靠“煤炭保供+新能源+装备+消费”，但工业利润/财政/外贸是拖累。2026年考验山西“转型动能能否对冲煤价周期”。")

# ---- 七、预算与财政 ----
heading1(doc, "七、预算与财政的“含金量”")
para(doc, "2025年山西一般公共预算收入**3218.6亿元、-9.1%**，其中税收**-6.9%**。财政收入回落（煤价/企业利润下行）。")
bullet(doc, "财政收入-9.1%、税收-6.9%，承压。")
bullet(doc, "财政“过紧日子+保民生（社保+2.4%等）+支持转型”。")
bullet(doc, "支出6058.9亿元、-4.0%，服务能源保供/转型。")

# ---- 八、民生底账 ----
heading1(doc, "八、民生底账：人口、收入与城乡")
para(doc, "2025年山西常住人口**3424.01万人、-21.95万**，城镇化率**67.35%、+1.03pct**。人口自然增长率**-3.32‰**。")
para(doc, "居民人均可支配收入**33923元、+4.6%**，农村（+5.5%）快于城镇（+3.7%），城乡比2.26。")
g8 = [
    ["指标", "数值", "信号"],
    ["常住人口", "3424.01万/-21.95万", "人口净流出"],
    ["人口自然增长率", "-3.32‰", "自然负增"],
    ["城镇化率", "67.35%/+1.03pct", "稳步城镇化"],
    ["居民人均可支配收入", "33923元/+4.6%", "收入稳、农村快"],
]
table(doc, g8[0], g8[1:], widths=[4.2, 4.0, 4.8])
para(doc, "**民生观察**：农村收入快于城镇、城镇化推进，但人口净流出、自然负增是约束。")

# ---- 九、城镇与农村 ----
heading1(doc, "九、城镇与农村：格局与均衡")
para(doc, "山西城镇化率67.35%、+1.03pct，城乡收入比2.26（缩小0.04），城乡相对协调。")
bullet(doc, "农村居民收入+5.5%、快于城镇+3.7%。")
bullet(doc, "社零城镇+4.9%、乡村+4.6%。")
bullet(doc, "人口向太原都市圈/城市群集中。")

# ---- 十、人口流入与流出 ----
heading1(doc, "十、人口流入与流出")
para(doc, "山西2025年常住3424.01万人、-21.95万，人口净流出、自然负增（-3.32‰）。")
para(doc, "未来看点：产业转型/新能源/装备能否留住人口；若“能源+制造+生态”成势，山西有望减缓人口流出。")

# ---- 十一、物价与货币 ----
heading1(doc, "十一、物价与货币环境")
para(doc, "2025年山西CPI**+0.1%**、PPI**-9.8%**（购进-6.9%），工业品出厂价深度走弱、煤价回落。")
para(doc, "物价偏弱反映“大宗/能源价格下行”，与全国低通胀一致。2026年“稳煤价/扩内需”是主线。")

# ---- 十二、区域一体化 ----
heading1(doc, "十二、区域一体化：山西在“国家能源基地+京津冀对接+中部崛起”里的位置")
para(doc, "山西的核心战略坐标是“**国家能源基地（煤炭保供）+京津冀协同对接+中部崛起**”，是能源安全的“压舱石”。")
bullet(doc, "能源保供：原煤13亿、输电力1503.6亿千瓦时。")
bullet(doc, "京津冀对接：向京津输送电力/清洁能源/冷链。")
bullet(doc, "中部城市群：太原都市圈、晋东冀西都市圈。")
bullet(doc, "生态屏障：京津冀生态支撑区、黄河流域生态修复。")
para(doc, "若“能源安全+产业转型+生态修复”成势，山西将作为全国能源绿色转型与中部崛起的样本。")

# ========= 十三、五条主线 =========
heading1(doc, "十三、未来5—10年最值得观察的五条主线")
lines5 = [
    ("① 煤炭/能源安全", "煤炭保供、能源安全能否持续稳固。"),
    ("② 新型综合能源/新能源", "新能源装机、绿电/氢能能否壮大。"),
    ("③ 产业转型/装备/新材料", "装备、新材料、机器人/光伏能否成极。"),
    ("④ 生态修复/绿色转型", "PM2.5、生态修复、碳达峰。"),
    ("⑤ 人口/就业/外贸", "人口净流出、外贸能否修复。"),
]
for l in lines5:
    para(doc, "**%s**　%s" % l, space_after=6)

# ---- 十四、结论 ----
heading1(doc, "十四、最终结论：山西在“煤炭能源+新型综合能源+生态修复+产业转型”里的增长逻辑")
para(doc, "山西的2025年，本质上是“**煤炭保供/新能源/装备/生态为核心，而工业利润/外贸/人口弱**”的答卷：GDP25495.68亿、+4.0%，原煤13.09亿吨、新能源装机占55.1%、GDP能耗-2.0%，但财政-9.1%、工业利润-36.3%、进出口-21.2%、自然增-3.32‰。")
para(doc, "只要能源安全保供、新能源/装备转型、生态修复持续，山西就站在“能源绿色转型”的位；如果煤价周期/工业利润/人口持续偏弱，山西需承受“能源依赖、弱经济/利润”的挑战。")
para(doc, "最稳观察信号：**一盯能源安全（煤炭保供）、二盯新能源/装备（转型）、三盯生态修复（绿色）、四盯工业利润/财政（约束）、五盯人口/外贸（长期）。**山西，是“能源转型+生态修复”的新样本。")

# ---- 附录A ----
heading1(doc, "附录A：主要资料来源与核验路径")
bullet(doc, "山西省2025年《政府工作报告》——目标来源。")
bullet(doc, "《2025年山西省国民经济和社会发展统计公报》（省统计局，2026-03-25）——GDP、工业、外贸、人口实值。")
bullet(doc, "2026年山西省政府工作报告——2025执行复盘+能源/生态/转型。")
bullet(doc, "太原海关、省财政厅——外贸/财政。")
heading2(doc, "核验说明")
para(doc, "本报告所有增速、总量、占比均以统计公报/官方口径为基准，涉“煤炭/能源/生态/转型”等以官方口径为准。")

# ---- 附录B ----
heading1(doc, "附录B：建议建立的年度跟踪仪表盘")
para(doc, "建议把下面10个“测脉搏”指标做成年度跟踪表：")
dash = [
    ["#", "指标", "2025基值", "为什么重要"],
    ["1", "GDP增速", "+4.0%", "总量与方向"],
    ["2", "规上工业增速", "+4.9%", "制造底盘"],
    ["3", "原煤产量", "13.09亿吨", "能源安全"],
    ["4", "新能源装机占比", "55.1%", "能源转型"],
    ["5", "固定资产投资/地产", "-0.4%/-9.9%", "投资结构"],
    ["6", "社零增速", "+4.7%", "内需消费"],
    ["7", "常住人口/自然增长率", "3424.01万/-3.32‰", "人口与城市"],
    ["8", "一般公共预算收入增速", "-9.1%", "财政质量"],
    ["9", "PM2.5浓度", "30μg/m³", "生态环保"],
    ["10", "CPI/PPI", "+0.1%/-9.8%", "物价与需求"],
]
table(doc, dash[0], dash[1:], widths=[0.8, 4.6, 3.4, 4.4])
para(doc, "把这10个指标连起来看，任何一个“新能源（4）、煤炭（3）、装备/机器人（生态环保）向上、财政/人口（8/7）修复”，都说明山西在真正换挡。")

# ============ 保存 ============
out = "/Users/x/Desktop/content-prod-lab/reports/山西省_2025年政府工作报告_深度研究_2026-08-13.docx"
doc.save(out)
print("SAVED:", out)
print("PARAGRAPHS:", len(doc.paragraphs))
print("TABLES:", len(doc.tables))
