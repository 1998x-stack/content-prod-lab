# -*- coding: utf-8 -*-
"""Build 海口市2025年政府工作报告 深度研究 DOCX, 参照省会系列版式。"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DARK = RGBColor(0x1F, 0x2A, 0x44)
RED = RGBColor(0xB0, 0x1E, 0x2E)
GRAY = RGBColor(0x59, 0x60, 0x69)
BLUE = RGBColor(0x1F, 0x4E, 0x79)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
HEADER_FILL = "1F2A44"
K = "微软雅黑"


def set_run(run, size=11, bold=False, color=DARK, font=K):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    run.font.name = font
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), font)
    rFonts.set(qn("w:hAnsi"), font)
    rFonts.set(qn("w:eastAsia"), font)


def para(doc, text, size=11, bold=False, color=DARK, align=None,
         space_after=6, space_before=0, font=K):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        set_run(r, size=size, bold=bold or (i % 2 == 1), color=color, font=font)
    return p


def heading1(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(16)
    pf.space_after = Pt(8)
    r = p.add_run(text)
    set_run(r, size=16, bold=True, color=RED)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "B01E2E")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def heading2(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(10)
    pf.space_after = Pt(4)
    r = p.add_run(text)
    set_run(r, size=13, bold=True, color=BLUE)
    return p


def table(doc, headers, rows, widths=None, font_size=9.5):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_run(r, size=font_size, bold=True, color=WHITE)
        tcPr = hdr[i]._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), HEADER_FILL)
        tcPr.append(shd)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(val))
            set_run(r, size=font_size, color=DARK)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    return t


def quote_box(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(8)
    pf.left_indent = Pt(12)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), "B01E2E")
    pBdr.append(left)
    pPr.append(pBdr)
    r = p.add_run(text)
    set_run(r, size=11, color=DARK, font="楷体")
    return p


def bullet(doc, text, size=10.5):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.7)
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        set_run(r, size=size, bold=(i % 2 == 1), color=DARK)
    return p


# ================================================================= 文档主体
doc = Document()
sec = doc.sections[0]
sec.page_width = Cm(21)
sec.page_height = Cm(29.7)
sec.top_margin = Cm(2.4)
sec.bottom_margin = Cm(2.2)
sec.left_margin = Cm(2.5)
sec.right_margin = Cm(2.5)

st = doc.styles["Normal"]
st.font.name = K
st.font.size = Pt(11)
st.element.rPr.rFonts.set(qn("w:eastAsia"), K)


# ---- 封面 ----
para(doc, "海口市2025年政府工作报告", size=24, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=60, space_after=4)
para(doc, "深度研究与\u201c容易被忽视的细节\u201d分析", size=16, bold=True, color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
para(doc, "从\u201c海南自贸港核心引领区、离岛免税、现代服务业与旅游贸易\u201d重新理解海口", size=12, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)
para(doc, "\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501", color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
para(doc, "研究对象：2025年海口市政府工作报告及同期财政、国民经济和社会发展统计资料、2026年官方复盘", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
para(doc, "标定日期：2026-08-13", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
doc.add_page_break()

# ---- 目录 ----
catalog = [
    "执行摘要",
    "一、研究口径与阅读方法",
    "二、先看海口的特殊底盘：自贸港核心引领区、离岛免税、现代服务业与海洋经济",
    "三、最关键的宏观错位：GDP增但低于目标，规上工业爆发，投资与外贸却双弱",
    "四、容易被忽视、但信号很重的细节（15条）",
    "五、2025年GDP目标 vs 实际：对照表",
    "六、2025年增长是谁撑起来的？（结构归因）",
    "七、预算与财政的\u201c含金量\u201d",
    "八、民生底账：人口、收入与城乡",
    "九、城镇与农村：格局与均衡",
    "十、人口流入与流出",
    "十一、物价与货币环境",
    "十二、区域一体化：海口在\u201c自贸港+琼州海峡+海洋经济\u201d里的位置",
    "十三、未来5—10年最值得观察的五条主线",
    "十四、最终结论：海口在\u201c免税消费+自贸港制度+现代服务业\u201d里的增长逻辑",
    "附录A：主要资料来源与核验路径",
    "附录B：建议建立的年度跟踪仪表盘",
]
para(doc, "目　录", size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=10, space_after=10)
for item in catalog:
    para(doc, item, size=12, space_after=4)
doc.add_page_break()

# ---- 执行摘要 ----
heading1(doc, "执行摘要")
para(doc, "**核心判断**　2025年海口最显著的是\u201cGDP 2562.85亿元、增长4.8%（低于6.5%目标）、三产占79.1%\u201d、\u201c规上工业增加值+18.7%但固投-28.5%\u201d、\u201c货物进出口932.92亿元/+8.5%\u201d、\u201c常住人口305.63万/城镇化83.5%\u201d。这说明海口经济在\u201c自贸港制度红利+离岛免税+旅游/海洋经济+现代服务业\u201d下维持增长，但**投资与外贸是最大短腿**。")
para(doc, "把2025年初设定的目标（GDP 6.5%以上/规上工业+7%/固投+8.5%/社零+10%/财政+3%）、2025年统计、2026年前瞻一起看，海口呈现\u201c自贸港核心城市\u201d的典型路径：**增速目标年年设定偏高、实际靠免税+旅游+海洋+园区支撑**。总量2563亿为海南全省最大、占比超三成。")
para(doc, "最容易记住的一句话：**海口是“海南自贸港核心引领区+离岛免税城市”，靠“免税消费+旅游+现代服务业+自贸港制度红利+海洋经济”增长。**观察海口，与其只看“GDP 2562亿”，不如看“规上工业+18.7%、离岛免税、旅游2853万人次、海洋经济、‘六之城’规划”。")# ---- 一、研究口径与阅读方法 ----
heading1(doc, "一、研究口径与阅读方法")
para(doc, "本报告以\u201c海口市2025年政府工作报告（2025年1月，丁晖作）\u201d为起点，把\u201c2025年GDP目标（6.5%以上）\u201d与\u201c2025年国民经济和社会发展统计公报实际值（2562.85亿元/+4.8%）\u201d并置对照，再用2026年政府工作报告的复盘作事后验证。")
para(doc, "重点阅读三类证据：**一是指标目标（计划）是否达标**；**二是结构（谁贡献增长）是否匹配目标叙事**；**三是官方复盘（2026年报告）如何自评、承认问题**。统计口径一般公共预算收入为地方级，进出口以海口海关公布为口径。")
para(doc, "为便于跨年比较，统一采用\u201c同比增长%\u201d（GDP为不变价，其余多为现价），财政与居民收入多为名义值，文中按原口径转录、不再逐项换算。")

# ---- 二、特殊底盘 ----
heading1(doc, "二、先看海口的特殊底盘：自贸港核心引领区、离岛免税、现代服务业与海洋经济")
para(doc, "**区位与身份**：海口是海南省会、全省政治经济文化中心，海南自贸港\u201c核心引领区\u201d，2025年底如期完成封关运作任务（二线口岸建成启用、三张清单落地）。海口定位为\u201c两区一地\u201d：自贸港集中展示区、国际化现代服务业集聚区、双基地重要承载地。")
para(doc, "**产业底盘**：一是离岛免税+旅游（游客2853.51万人次/+9.3%、总花费469.44亿元/+8.2%、离岛免税购物金额环比+39.6%）；二是现代服务业（三产占GDP 79.1%，金融、租赁、会展、数字经济）；三是海洋经济（2024年GDP占比44.6%、2025年海洋生产总值目标+10%）；四是园区经济（江东新区现代商贸服务产业群营收破万亿、综保区消费精品/加工产业、复兴城数字经济）。")
para(doc, "**人口底盘**：2025年末常住人口305.63万/+（城镇化率83.55%、出生率7.82\u2030、死亡率6.21\u2030，自然正增长）；高校（海南大学等）+自贸港政策吸引人才落户（\u201c十四五\u201d引进人才29万）。")
para(doc, "**市场与出口**：2025年社零1185.92亿元/+3.2%但增速远低于目标+10%；货物进出口932.92亿元/+8.5%、服务进出口296.27亿元/+18.9%。外贸正增长但以进口为主（逆差268.86亿元）。")

# ---- 三、核心宏观错位 ----
heading1(doc, "三、最关键的宏观错位：GDP增但低于目标，规上工业爆发，投资与外贸却双弱")
para(doc, "**第一组错位**：2025年GDP目标6.5%以上，实际2562.85亿元/+4.8%，**连续两年低于目标**（2024年实际3.9%），但规上工业增加值+18.7%却爆发式增长——工业快于GDP 4倍。**增长更依赖第二产加持的三产服务业+园区**。")
para(doc, "**第二组错位**：固定资产投资**-28.5%（断崖式转负）**，与2025年目标+8.5%相差35.9个百分点；2024年还是+3%左右，2025年大幅回落。投资是海口最惨淡的引擎。")
para(doc, "**第三组错位**：社零+3.2%（目标+10%大幅不及）、CPI持平；但**免税+旅游+会展+数字经济**这些\u201c更高进阶\u201d服务业仍扩张（会展效益+19.3%、复兴城数字经济破700亿、离岛免税环比回暖）。")
para(doc, "一句话：**海口是\u201c服务业+免税+制度红利\u201d强、而投资和传统贸易弱的省会**——增长韧性靠自贸港红利与第三产，但投资缺口极大。")

# ---- 四、15条细节 ----
heading1(doc, "四、容易被忽视、但信号很重的细节（15条）")
bullet(doc, "1. **GDP低于目标**：2562.85亿元/+4.8%<6.5%目标，连续第二年被目标吊打。")
bullet(doc, "2. **规上工业+18.7%**：远高于GDP，是海口少见的工业高光。")
bullet(doc, "3. **石油和天然气开采业+47.5%**：采矿业+55%、石油开采贡献最大，反映\u201c油气+海洋\u201d。")
bullet(doc, "4. **医药制造-13%/非金属矿-19.5%**：传统制造输血不足，制造业高端化待补。")
bullet(doc, "5. **固投-28.5%**：是2025年最大塌陷，房地产-31.5%、建安重挫。")
bullet(doc, "6. **房地产开发投资303.31亿/-31.5%，但销售面积+19.4%、销售额+1.2%**：竣工去库存、投资收缩。")
bullet(doc, "7. **货物进口+14.4%、出口-0.8%**：自贸港\u201c进口为主\u201d，贸易逆差268.86亿。")
bullet(doc, "8. **服务进出口+18.9%（顺差112.91亿）**：现代服务业出口顺差是亮点。")
bullet(doc, "9. **离岛免税购物金额+39.6%**：封关前消费政策利好集中释放。")
bullet(doc, "10. **旅游2853万人次/+9.3%、入境过夜+21.9%**：入境游恢复、国际航班/航线扩张。")
bullet(doc, "11. **常住人口305.63万/城镇化83.6%**：人口自然正增长、居海南首位。")
bullet(doc, "12. **收入：城镇49741元/+3.3%、农村24862元/+5.2%**：城乡差距较全国更大。")
bullet(doc, "13. **CPI持平&核心CPI（未披露核心值）**：物价持续低迷、通缩压力。")
bullet(doc, "14. **财政：地方一般公共预算收入253.80亿元/-1.7%（全国罕见负增长）**，税收+2.2%/非税-18.3%。")
bullet(doc, "15. **停车：全年新增市场主体194万户（\u201c十四五\u201d累计）**、会展综合收入+17.3%。")# ---- 五、目标 vs 实际对照表 ----
heading1(doc, "五、2025年GDP目标 vs 实际：对照表")
table(doc,
    ["指标", "2025年目标", "2025年实际", "是否达标"],
    [
        ["地区生产总值增速", "6.5%以上", "2562.85亿元/+4.8%", "未达标"],
        ["规上工业增加值", "+7%", "+18.7%", "大幅超目标"],
        ["固定资产投资", "+8.5%以上", "-28.5%（转负）", "严重未达标"],
        ["社会消费品零售总额", "+10%以上", "1185.92亿元/+3.2%", "未达标"],
        ["地方一般公共预算收入", "+3%左右", "253.80亿元/-1.7%", "未达标(负增长)"],
        ["城镇/农村人均可支配收入", "+6.5%/+约7%", "+3.3%/+5.2%", "均低于目标"],
        ["货物/服务进出口", "均+20%以上", "+8.5%/+18.9%", "服务接近、货物低于"],
    ],
    widths=[3.4, 2.5, 5.0, 3.1])
para(doc, "**简评**：唯一\u201c大幅超目标\u201d的是规上工业；**GDP、投资、消费、财政全部未达目标，外贸货物亦打折**。这是海口\u201c目标设高、实际收敛\u201d的年份，增长更多来自工业与服务业的内在修复而非目标工程。")

# ---- 六、结构归因 ----
heading1(doc, "六、2025年增长是谁撑起来的？（结构归因）")
para(doc, "**三大产业**：一产111.39亿/+3.1%、二产425.28亿/+4.9%、三产2026.18亿/+4.9%；**三产占GDP 79.1%**，是绝对主体。增长主要由\u201c三产（免税/旅游/金融/会展/数字经济）+二产（工业）\u201d共同贡献。")
para(doc, "**工业**：规上工业增加值+18.7%，其中石油和天然气开采业+47.5%、采矿业+55%、汽车制造+23.7%、农副食品+8.1%；医药-13%、非金属矿-19.6%、电力-8.2%拖累。规上工业利润总额+35.1%。")
para(doc, "**服务业/消费**：接待游客2853.51万人次+9.3%、总花费469.44亿元+8.2%、会展效益178.79亿元+19.3%、离岛免税环比+39.6%；跨境电商/数字经济（复兴城）扩张。")
para(doc, "**总量归因结论**：海口GDP增长主要靠**第三产业（旅游/免税/服务/园区）+工业高景气（油气/汽车）**支撑；投资、房地产、传统出口贡献为负。")

# ---- 七、财政 ----
heading1(doc, "七、预算与财政的\u201c含金量\u201d")
para(doc, "2025年全口径一般公共预算收入569.56亿元/-1.3%；**地方一般公共预算收入253.80亿元/-1.7%（负增长，全国省会少见的财税收缩）**，其中税收213.44亿元/+2.2%、非税40.36亿元/-18.3%。地方一般公共预算支出336.74亿元/-9.5%。")
para(doc, "**结构性反转**：税收正增长但非税大幅压降、支出大幅收缩（-9.5%），反映\u201c减收+收紧支出+盘活存量（含封关税制改革）\u201d。收入与支出双降、支出降幅更大，财政多地位于\u201c稳税基、压非税、化债\u201d。")
para(doc, "**含金量**：收入靠税收质量（+2.2%的税）而非非税；但支出负增长意味着基建与民生（部分）让渡给化债与收紧，需留意财政乘数回退。")

# ---------------- 八、民生 ----------------
heading1(doc, "八、民生底账：人口、收入与城乡")
para(doc, "常住人口305.63万（城镇化83.6%、出生率7.82\u2030/死亡率6.21\u2030，自然正增长）。收入：**城镇49741元/+3.3%、农村24862元/+5.2%**，城乡比值约2.0，高于全国平均、城乡差距偏大。居民人均消费支出32597元/+4.5%。")
para(doc, "**民生结论**：人口平稳净自然增长、公共服务扩张（18个街心公园、社区服务中心），但收入增速放缓（城镇仅+3.3%）、财政支出收缩对民生项目或有掣肘。")

# ---------------- 九、城乡 ----------------
heading1(doc, "九、城镇与农村：格局与均衡")
para(doc, "海口城镇化率约84%（常住中），主城区承载自贸港服务业、免税、金融、数字经济；农村县域发展热带农业（园林水果+8.1%、生猪/水产品正增）与和美乡村。")
para(doc, "**城乡收入比**约2.0，高于全国平均，均衡度偏弱；但农村收入增速（+5.2%）快于城镇（+3.3%），差距在边际收敛（趋势改善）。")

# ---------------- 十、人口流入流出 ----------------
heading1(doc, "十、人口流入与流出")
para(doc, "海口常住人口305.63万、自然正增长（出生7.7\u2030>死亡6.2\u2030），并叠加自贸港人才落户（\u201c十四五\u201d引进人才约29万）。多元人才+高校（海南大学）+免税/旅游/数字经济就业吸引\u201c产业移民\u201d。")
para(doc, "**流入**：自贸港高端服务业/免税/旅游/园区人才+大学生；**流出**：部分传统制造、低技能劳动力与房租压力外迁。净流入为正，是全国少有的\u201c人口+税-产业\u201d韧性型省会。")

# ---------------- 十一、物价 ----------------
heading1(doc, "十一、物价与货币环境")
para(doc, "2025年海口CPI**持平（0%）**，食品+0.1%、交通通信-2.0%；为低通胀/准通缩。广义流动性充足（银行资产总额13945.62亿元/+15%）。")
para(doc, "货币环境：存款贷款规模扩张而物价低迷，反映\u201c宽货币、弱需求\u201d；免税/会展等可选消费靠政策（补贴/免税）拉动而非民众收入。")

# ---------------- 十二、区域一体化 ----------------
heading1(doc, "十二、区域一体化：海口在\u201c自贸港+琼州海峡+海洋经济\u201d里的位置")
para(doc, "海口是海南自贸港\u201c核心引领区\u201d与全省人流物流资金流中枢：依托琼州海峡（与广东湛江、经济圈）、海澄文定（海口都市圈）一体化。\u201c海澄文定\u201d、海口经济圈推进8条城际干道、12个项目。")
para(doc, "海洋经济是特色：海洋生产总值2024年占GDP 44.6%、2025年目标+10%；美兰国际机场（国际航线47条）+港口集装箱区位\u201c自贸港海、空枢纽\u201d。一体化上一靠琼州海峡、二是南海、三是大湄公河次区域，海口是\u201c门户型\u201d核心。")

# ---------------- 十三、五条主线 ----------------
heading1(doc, "十三、未来5—10年最值得观察的五条主线")
bullet(doc, "1. **封关运作红利落地**：2026年全面封关，三张清单+零关税+加工增值能否把\u201c免税消费+免税货物+人才\u201d落到实处。")
bullet(doc, "2. **离岛免税与消费升级**：离岛免税店、首店/演艺/赛事，能否在收入放缓下维持高增长。")
bullet(doc, "3. **海洋经济+油气加工**：石油开采+47.5%、海洋生产总值+10%，特色的\u201c蓝色矿\u201d。")
bullet(doc, "4. **投资修复**：固投-28.5%后能否靠\u201c自贸港重大项目+新一轮招商\u201d修复（2026年目标+7%）。")
bullet(doc, "5. **人口与城市能级**：常住城镇化+高校科创、抢占\u201c自贸港人口红利\u201d与\u201c海洋/服务人才\u201d。")

# ---------------- 十四、最终结论 ----------------
heading1(doc, "十四、最终结论：海口在\u201c免税消费+自贸港制度+现代服务业\u201d里的增长逻辑")
para(doc, "**结论**：海口2025年的\u201c真相\u201d是——**GDP+4.8%低于目标、规上工业+18.7%、固投-28.5%塌陷、社零+3.2%、地方财政收入-1.7%负增长**。它是\u201c服务业+免税+自贸港制度红利\u201d驱动但**缺投资、弱第二产业**的省会。")
para(doc, "**最值得跟踪的不是GDP**，而是\u201c**规上工业（+18.7%）+离岛免税+服务出口顺差+海洋经济**\u201d——它们比总量更能刻画自贸港真实成色。**固投-28.5%与财政-1.7%是海口最大的两处暗礁**，决定未来空间。")
para(doc, "**若只看一个指标**：看**固定资产投资增速（-28%左右）+地方财政收入（-1.7%）**——海口是\u201c免税收入好、投资与财政弱\u201d的矛盾省会，投资修复是自贸港成色改善的关键变量。")

# ---------------- 附录A ----------------
heading1(doc, "附录A：主要资料来源与核验路径")
bullet(doc, "海口市人民政府《2025年海口市政府工作报告》（2025年1月，丁晖作）。")
bullet(doc, "海口市统计局《2025年海口市国民经济和社会发展统计公报》（2026年4月）。")
bullet(doc, "海口市统计局《2025年海口市全年经济运行情况》（2026年1月）。")
bullet(doc, "《2026年海口市政府工作报告》（2026年2月，张勇）及百度百科/官方解读。")

# ---------------- 附录B ----------------
heading1(doc, "附录B：建议建立的年度跟踪仪表盘")
bullet(doc, "GDP总量/增速 vs 全年目标（+0.5个百分点口径）。")
bullet(doc, "规上工业增加值及分行业（石油开采/医药/汽车/非金属矿）增速。")
bullet(doc, "固定资产投资（总量/房地产/基建/民间）增速。")
bullet(doc, "社会消费品零售总额、离岛免税店销售额。")
bullet(doc, "货物/服务进出口额（海澄）、贸易逆差/顺差。")
bullet(doc, "地方一般公共预算收入/税收/非税、财政支出。")
bullet(doc, "常住人口增量、城镇化率、高校/迁入人才。")
bullet(doc, "CPI/核心CPI、规上工业利润总额与营收。")
bullet(doc, "旅游人次/旅游总收入、会展综合收入、入境游客。")
bullet(doc, "海洋生产总值、江东/高新区园区、高校数。")

# ---------------- 保存 ----------------
out = "/Users/x/Desktop/content-prod-lab/reports/海口市_2025年政府工作报告_深度研究_2026-08-13.docx"
doc.save(out)
print("SAVED OK 海口市_2025年政府工作报告_深度研究_2026-08-13.docx")