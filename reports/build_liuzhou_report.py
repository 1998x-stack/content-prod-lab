# -*- coding: utf-8 -*-
"""Build 柳州市2025年政府工作报告 深度研究 DOCX, 参照地级市系列版式。"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DARK = RGBColor(0x1F, 0x2A, 0x44)
RED = RGBColor(0xB0, 0x1E, 0x2E)
GRAY = RGBColor(0x59, 0x60, 0x69)
BLUE = RGBColor(0x1F, 0x4E, 0x79)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
HEADER_FILL = "1F2A44"
K = "微软雅黑"


def set_run(run, size=11, bold=False, color=DARK, font=K):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    run.font.name = font
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), font)
    rFonts.set(qn("w:hAnsi"), font)
    rFonts.set(qn("w:eastAsia"), font)


def para(doc, text, size=11, bold=False, color=DARK, align=None,
         space_after=6, space_before=0, font=K):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        set_run(r, size=size, bold=bold or (i % 2 == 1), color=color, font=font)
    return p


def heading1(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(16)
    pf.space_after = Pt(8)
    r = p.add_run(text)
    set_run(r, size=16, bold=True, color=RED)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "B01E2E")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def heading2(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(10)
    pf.space_after = Pt(4)
    r = p.add_run(text)
    set_run(r, size=13, bold=True, color=BLUE)
    return p


def table(doc, headers, rows, widths=None, font_size=9.5):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_run(r, size=font_size, bold=True, color=WHITE)
        tcPr = hdr[i]._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), HEADER_FILL)
        tcPr.append(shd)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(val))
            set_run(r, size=font_size, color=DARK)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    return t


def quote_box(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(8)
    pf.left_indent = Pt(12)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), "B01E2E")
    pBdr.append(left)
    pPr.append(pBdr)
    r = p.add_run(text)
    set_run(r, size=11, color=DARK, font="楷体")
    return p


def bullet(doc, text, size=10.5):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.7)
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        set_run(r, size=size, bold=(i % 2 == 1), color=DARK)
    return p


# ================================================================= 文档主体
doc = Document()
sec = doc.sections[0]
sec.page_width = Cm(21)
sec.page_height = Cm(29.7)
sec.top_margin = Cm(2.4)
sec.bottom_margin = Cm(2.2)
sec.left_margin = Cm(2.5)
sec.right_margin = Cm(2.5)

st = doc.styles["Normal"]
st.font.name = K
st.font.size = Pt(11)
st.element.rPr.rFonts.set(qn("w:eastAsia"), K)


# ---- 封面 ----
para(doc, "柳州市2025年政府工作报告", size=24, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=60, space_after=4)
para(doc, "深度研究与\u201c容易被忽视的细节\u201d分析", size=16, bold=True, color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
para(doc, "从\u201c汽车之城（上汽通用五菱）、螺蛳粉、机械、新能源、西南智谷\u201d重新理解柳州", size=12, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)
para(doc, "\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501", color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
para(doc, "研究对象：2025年柳州市政府工作报告及同期财政、国民经济和社会发展统计资料、2026年官方复盘", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
para(doc, "标定日期：2026-08-14", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
doc.add_page_break()

# ---- 目录 ----
catalog = [
    "执行摘要",
    "一、研究口径与阅读方法",
    "二、先看柳州的特殊底盘：汽车之城、螺蛳粉、机械制造、新能源（三电）与西南智谷",
    "三、最关键的宏观错位：GDP 3096.91亿/6.2%低于8%目标，工业强但社零低、人口外流、CPI负",
    "四、容易被忽视、但信号很重的细节（15条）",
    "五、2025年GDP目标 vs 实际：对照表",
    "六、2025年增长是谁撑起来的？（结构归因）",
    "七、预算与财政的\u201c含金量\u201d",
    "八、民生底账：人口、收入与城乡",
    "九、城镇与农村：格局与均衡",
    "十、人口流入与流出",
    "十一、物价与货币环境",
    "十二、区域一体化：柳州在\u201c广西+西部陆海新通道+粤港澳\u201d里的位置",
    "十三、未来5—10年最值得观察的五条主线",
    "十四、最终结论：柳州在\u201c新能源汽车+螺蛳粉+机械\u201d里的增长逻辑",
    "附录A：主要资料来源与核验路径",
    "附录B：建议建立的年度跟踪仪表盘",
]
para(doc, "目　录", size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=10, space_after=10)
for item in catalog:
    para(doc, item, size=12, space_after=4)
doc.add_page_break()

# ---- 执行摘要 ----
heading1(doc, "执行摘要")
para(doc, "**核心判断**　2025年柳州最显著的是\u201cGDP 3096.91亿元、增长6.2%（低于8%目标）、广西第2\u201d、\u201c规上工业+10.3%（汽车制造+14.8%、装备+17.8%、高技术+31.1%）\u201d、\u201c进出口442.23亿/+17.9%（电动汽车出口+264.5%）\u201d、\u201c财收167亿/+12.0%\u201d、\u201c但社零+1.2%、常住412.58万/-2.02万、CPI-0.2%\u201d。这说明柳州在\u201c汽车+螺蛳粉+制造\u201d的转型中，**工业、外贸、财收亮但消费、人口、物价弱**。")
para(doc, "把2025年目标（GDP+8%以上、规上工业约9%）、2025年统计（GDP+6.2%未达、规上+10.4%超额、财收+12.0%超额、社零+1.2%）、趋势一起看，柳州是\u201c汽车主导+螺蛳粉\u201d路径：**上汽通用五菱（新能源车75.91万辆）、机械（柳工）、螺蛳粉（全产业链800亿+）、汽车电子**是支柱；2025年总量3096.9亿居广西第2（次于南宁）。")
para(doc, "最容易记住的一句话：**柳州是\u201c汽车之城（上汽通用五菱）、螺蛳粉之都、西南制造\u201d，靠\u201c新能源汽车+机械+螺蛳粉+出口\u201d增长。**观察柳州，与其只看\u201cGDP 3096亿\u201d，不如看\u201c新能源汽车75.91万辆、电动车出口+264.5%、螺蛳粉全产业链800亿+、规上+10.3%+\u201d。")

# ---- 一、研究口径与阅读方法 ----
heading1(doc, "一、研究口径与阅读方法")
para(doc, "本报告以\u201c柳州市政府工作报告（2025年）\u201d为起点，把\u201c2025年GDP目标（8%以上）\u201d与\u201c官方2025年GDP（3096.91亿元/+6.2%）\u201d并置对照，并用\u201c2025年柳州市统计公报\u201d和\u201c2026年政府工作报告复盘\u201d横向核验。")
para(doc, "口径提示：\u201cGDP、规上工业增加值增速\u201d按可比价（实际增速）；\u201c投资、消费、进出口、财政收入\u201d按现价（名义增速）。涉及\u201c常住人口\u201d用统计公报常住口径（412.58万），城镇化率用官方公布值（72.0%）。")
para(doc, "指标体系（与研究口径一致）：五大象限——**总量与增速（GDP）、产业动能（汽车/螺蛳粉/机械/高技术）、外贸、财政质量、民生与人口**。")
para(doc, "特别提示（不吃老本）：柳州2024年GDP+1.5%（规上+3.2%）偏低，2025年大幅回升到+6.2%、规上+10.4%——它是\u201c汽车（上汽通用五菱）重启+螺蛳粉爆发\u201d带动回升；真正要看的是\u201c汽车电动化、螺蛳粉产业化、人口外流\u201d的矛盾。")
# ---- 二、先看柳州的特殊底盘 ----
heading1(doc, "二、先看柳州的特殊底盘：汽车之城、螺蛳粉、机械制造、新能源（三电）与西南智谷")
para(doc, "柳州地处广西中北部、柳江之滨，是**汽车之城（上汽通用五菱、东风柳汽）、螺蛳粉之都、机械制造重镇（柳工）、西南智谷**。2025年GDP 3096.91亿元、常住412.58万，广西第2，人均7.49万元。")
para(doc, "五个底盘名词，先立框架：")
bullet(doc, "**汽车之城**　2025年汽车产量115.19万辆（+4.1%）、新能源汽车75.91万辆；上汽通用五菱（全球小型电动车龙头）、东风柳汽，\u201c三电\u201d核心布局。")
bullet(doc, "**螺蛳粉之都**　柳州螺蛳粉（全球网红食品），2023全产业链669.9亿、预计2025达889.9亿，带动原料/包装/电商/出海。")
bullet(doc, "**机械制造**　柳工（工程机械、装载机/挖掘机）全球，机械/装备制造（+17.8%）。")
bullet(doc, "**新能源（三电）**　锂离子电池+56.6%、新能源汽车出口+264.5%，\u201c三电\u201d（电池/电机/电控）布局。")
bullet(doc, "**西南智谷**　广西工业最大城市、\u201c机器之城\u201d（工业机器人、服务机器人+23%）、面向东盟。")
para(doc, "这五根（汽车+螺蛳粉+机械+新能源+智谷）构成柳州独特底盘：**左手汽车（上汽通用五菱），右手螺蛳粉，腹地是西南制造**。理解柳州，先理解\u201c一袋粉、一辆车\u201d的产业逻辑。")

# ---- 三、最关键的宏观错位 ----
heading1(doc, "三、最关键的宏观错位：GDP 3096.91亿/6.2%低于8%目标，工业强但社零低、人口外流、CPI负")
para(doc, "2025年柳州最需要辨析的一组\u201c错位\u201d：**GDP 6.2%未达8%目标、工业强（规上+10.3%、高技术+31.1%）但社零+1.2%、常住-2.02万、CPI-0.2%**。")
para(doc, "为什么\u201c工业这么强、经济却未达标、人口还流失\u201d？三个解释：")
para(doc, "**其一，工业高价但当量有限**　规上+10.3%（汽车+14.8%、装备+17.8%、高技术+31.1%），但汽车业对就业/内生带动边际下滑，工业增长未完全转化为\u201c数量和需求\u201d。")
para(doc, "**其二，消费、人口、物价\u201c三弱\u201d**　社零+1.2%（低于全区/目标）、常住-2.02万（广西三线/工业城市人口外流）、CPI-0.2%——\u201c需求不足+人口流失\u201d。")
para(doc, "**其三，投资/外资偏弱**　固投+4.6%偏低（但高技术投资+35.9%结构优）；实际利用外资（商务口径）4822万美元/+134%小。")
para(doc, "小结：柳州2025年是\u201c**强工业（汽车/螺蛳粉/机械）、弱消费人口、回升中**\u201d的一年：制造+出口+财收亮，但**社零、人口、物价**是短板。")

# ---- 四、容易被忽视、但信号很重的细节（15条） ----
heading1(doc, "四、容易被忽视、但信号很重的细节（15条）")
bullet(doc, "**1.新能源汽车75.91万辆、MPV+52.2%**　上汽通用五菱/新能源是柳州第一增长极。")
bullet(doc, "**2.电动车出口+264.5%、整车出口+27.4%**　汽车/新能源出海（面向东盟）爆发。")
bullet(doc, "**3.螺蛳粉全产业链（预计889.9亿、进出口/跨境带动）**\u201c小吃大产业\u201d，原料/工业/出口。")
bullet(doc, "**4.锂离子电池+56.6%、服务机器人+23.0%**\u201c新能源三电\u201d、智能机器人崛起。")
bullet(doc, "**5.高技术制造业+31.1%、装备+17.8%**\u201c新质生产力\u201d成色。")
bullet(doc, "**6.财收167亿/+12.0%、税收+5.1%**\u201c财政大幅增长（非税/国企为主）、税收稳。\u201d")
bullet(doc, "**7.互联网+101.4%、软件信息+49.4%**\u201c数字经济爆发\u201d。")
bullet(doc, "**8.对东盟出口+35.2%、RCEP**\u201c面向东南亚\u201d外向窗口。\u201d")
bullet(doc, "**9.社零+1.2%但线上+10.8%、家具+118.4%**\u201c新消费/在家宅\u201d分化。\u201d")
bullet(doc, "**10.居民收入39412元/+5.0%、城乡比2.13**\u201c农村+6.0%快于城镇+4.2%但差大。\u201d")
bullet(doc, "**11.常住412.58万/城镇化72.0%、自然增长-6.49‰**\u201c人口外流、深度老龄化少子化。\u201d")
bullet(doc, "**12.CPI-0.2%（食品烟酒-1.2%）**\u201c低通胀、需求不足。\u201d")
bullet(doc, "**13.社会消费品+1.2%、粮74.94万吨**\u201c农业稳（螺蛳粉原料——大米/螺）。\u201d")
bullet(doc, "**14.广西工业最大城市、工业产值占全区前列**“西南工业脊梁”。")
bullet(doc, "**15.一般公共预算支出479.39亿、民生占80.9%**\u201c财政对民生倾斜。\u201d")
# ---- 五、2025年GDP目标 vs 实际：对照表 ----
heading1(doc, "五、2025年GDP目标 vs 实际：对照表")
table(doc,
    ["指标", "2025年目标", "2025年实际", "达成评价"],
    [
        ["地区生产总值(GDP)", "增长8%以上", "3096.91亿/6.2%", "未达成，差1.8pct"],
        ["规上工业增加值", "增长9%左右", "+10.3%", "超额"],
        ["社会消费品零售总额", "——", "+1.2%", "低位"],
        ["一般公共预算收入", "——", "167.00亿/+12.0%", "大幅超额"],
        ["固定资产投资", "——", "+4.6%", "偏低"],
        ["进出口总额", "——", "442.23亿/+17.9%", "超额"],
        ["居民收入", "与经济增长同步", "39412元/+5.0%", "总体同步"],
    ],
)
para(doc, "注：GDP、规上工业按可比价。**规上工业（+10.3%）、进出口（+17.9%）、财收（+12.0%）超额**；**GDP（+6.2%）未达8%目标、社零（+1.2%）低位**。")
para(doc, "拆读：**工业（规上+10.3%/高技术+31.1%）、电动车出口（+264.5%）、螺蛳粉、财收（+12%）是亮色**；**社零（+1.2%）、人口（-2万）、CPI（-0.2%）偏弱**——\u201c制造强、需求弱\u201d，是工业城市样本。")

# ---- 六、2025年增长是谁撑起来的？（结构归因） ----
heading1(doc, "六、2025年增长是谁撑起来的？（结构归因）")
para(doc, "把柳州GDP的6.2%拆开：三次产业分别增4.2%、7.7%、5.0%（工业增加值+8.5%、结构约9.9：40.4：49.7）。**第二产业（工业）是主引擎（+8.6%），第三产业（服务业）稳，第一产业（农业）稳**。")
para(doc, "2026年柳州强调\u201c工业强市+面向东盟\u201d，聚焦**新能源汽车、螺蛳粉、机械、未来产业、智能机器人**——核心是\u201c先进制造+新质\u201d。")
para(doc, "**第二产业（工业）**：规上工业+10.3%、汽车制造+14.8%、装备+17.8%、高技术+31.1%、锂电+56.6%——\u201c汽车+机械+新能源\u201d支撑。")
para(doc, "**第三产业（服务业）**：交通运输+7.3%、批发零售+5.6%、住宿餐饮+5.1%；互联网+101.4%、软信+49.4%——\u201c数字经济爆发\u201d。")
para(doc, "**外贸（开放型）**：进出口+17.9%（出口+27.4%）、电动车出口+264.5%、对东盟+35.2%——\u201c面向东南亚\u201d外向强。")
para(doc, "一句话归因：**2025年柳州增长\u201c靠工业（汽车+螺蛳粉+机械）+\u2018外循环\u2019（出口）\u201d**，短板在\u201c消费、人口\u201d。")

# ---- 七、预算与财政的\u201c含金量\u201d ----
heading1(doc, "七、预算与财政的\u201c含金量\u201d")
para(doc, "2025年柳州**一般公共财政预算收入167.00亿元（+12.0%）**，其中税收103.91亿元（+5.1%）、税收占比约62%；一般公共预算支出479.39亿元（+3.8%）、民生支出占80.9%。")
bullet(doc, "税收结构：税收+5.1%，财收+12%更多靠非税（国企/收费）——\u201c增收含金量需看税收\u201d。")
bullet(doc, "民生支出占80.9%（社保就业+12.7%、教育+11.5%、卫生+9.8%）。")
bullet(doc, "金融支撑：存款5413.31亿（+4.7%）、贷款6187.00亿（+7.6%）——信贷宽对制造/设备。")
para(doc, "**财政含金量小结**：财收+12%与GDP/工业匹配、税收+5.1%；\u201c增收主、民生优\u201d；财政对\u201c新能源、螺蛳粉、智造\u201d投入加大。")

# ---- 八、民生底账：人口、收入与城乡 ----
heading1(doc, "八、民生底账：人口、收入与城乡")
para(doc, "2025年柳州**居民人均可支配收入39412元（+5.0%）**，其中城镇47689元（+4.2%）、农村22427元（+6.0%），城乡比约2.1。就业：城镇新增就业4.75万人。")
para(doc, "人口画像：**常住412.58万、城镇化率72.0%（+0.58pct）**，但自然增长-6.49‰、常住-2.02万——\u201c人口外流、深度老龄化少子化\u201d是柳州城市收缩的真问题。")
para(doc, "民生投入：支出占80.9%、社保/就业/教育优先——\u201c保民生\u201d。")

# ---- 九、城镇与农村：格局与均衡 ----
heading1(doc, "九、城镇与农村：格局与均衡")
para(doc, "柳州常住城镇化率72.0%（广西前列、素有\u201c工业城市\u201d），城乡较均衡；农村收入增速（+6.0%）高于城镇（+4.2%），但**城乡绝对差（约2.1倍）**仍大。")
para(doc, "农业底盘：**粮食74.94万吨（+0.01%）**、蔬菜/水果（+4.0%/+3.7%）、螺蛳粉原料（大米/田螺）——\u201c农业支撑螺蛳粉\u201d。")
para(doc, "一句话：\u201c柳州城镇化较高、螺蛳粉带动农村（原料/就业）\u201d；但\u201c人口外流、城乡收入差\u201d仍需治理。")

# ---- 十、人口流入与流出 ----
heading1(doc, "十、人口流入与流出")
para(doc, "柳州常住412.58万（比上-2.02万），是\u201c人口净流出\u201d的广西工业城市；受\u201c南宁聚集、沿海虹吸、三线制造业\u201d影响，青年/劳动力外流至珠三角、南宁等。")
para(doc, "结构观察：**自然增长-6.49‰、深度老龄化**，城镇化72%但人口总量收缩；\u201c柳州（人口）流失\u201d与\u201c广东（人口）流入\u201d形成反差。")
para(doc, "2026年目标：稳就业（新增4.7万+）、引才/流动人口服务——柳州需以\u201c新能源+螺蛳粉产业留人\u201d。")

# ---- 十一、物价与货币环境 ----
heading1(doc, "十一、物价与货币环境")
para(doc, "2025年柳州**CPI-0.2%**（食品烟酒-1.2%、交通通信-3.0%；教育文娱+2.6%、医疗+2.3%）——**低通胀、需求偏弱**。")
bullet(doc, "信贷扩张：贷款+7.6%、存款+4.7%，宽信用对制造。")
bullet(doc, "汽车/家电以旧换新（家电+）、线上+10.8%、家具+118.4%——\u201c新消费\u201d。")
para(doc, "货币环境判断：**宽信用、CPI负**（-0.2%）；柳州靠\u201c汽车以旧换新、螺蛳粉+文旅\u201d扩内需（2026稳物价）。")

# ---- 十二、区域一体化：柳州的位置 ----
heading1(doc, "十二、区域一体化：柳州在\u201c桂柳邕、西部陆海新通道、粤港澳\u201d里的位置")
para(doc, "柳州是**广西工业第一城（工业总值全区前列）、柳北湾/西江经济带、西南出海新通道节点**。")
bullet(doc, "**广西格局**　与南宁（首府）、桂林（旅游）构成\u201c柳邕桂\u201d三角，柳州=工业、面向东盟。")
bullet(doc, "**连接粤港澳/黔桂**　柳工/五菱辐射西南+东南亚，高铁（衡柳/柳广）、西江航运、拟通过珠江/西江出海。")
bullet(doc, "**面向东盟**　对东盟/越南出口+35.2%、RCEP——东盟窗口。")
para(doc, "一句话：**柳州在\u201c广西+面向东盟\u201d里，最核心的定位是\u201c广西工业引擎、汽车+螺蛳粉制造\u201d**——工业、区位是大优势。")

# ---- 十三、未来5—10年最值得观察的五条主线 ----
heading1(doc, "十三、未来5—10年最值得观察的五条主线")
bullet(doc, "**主线一：汽车电动化（上汽通用五菱）**\u201c新能源车75.91万辆、电动车出口+264.5%\u201d能否\u201c卖全球\u201d。")
bullet(doc, "**主线二：螺蛳粉产业化（千亿）**\u201c全产业链889.9亿\u201d能否\u201c一碗粉、一个千亿产业\u201d。")
bullet(doc, "**主线三：高性能机械/智造（柳工）**\u201c装备+17.8%、智能机器人\u201d能否撑\u201c新质\u201d。")
bullet(doc, "**主线四：新能源三电/电池**\u201c锂电+56.6%\u201d能否形成\u201c西南电池之都\u201d。")
bullet(doc, "**主线五：人口与城市收缩**\u201c412万、-2万\u201d能否扭转\u201c人口外流\u201d、把\u201c制造业+螺蛳粉\u201d变\u201c留人\u201d。")

# ---- 十四、最终结论 ----
heading1(doc, "十四、最终结论：柳州在\u201c新能源汽车+螺蛳粉+机械\u201d里的增长逻辑")
para(doc, "把2025年闭环打穿：**柳州是\u201c广西工业引擎、汽车之城\u201d**：GDP 3096.91亿/+6.2%、规上+10.3%（汽车+14.8%、高技术+31.1%）、进出口+17.9%、财收+12%、螺蛳粉/新能源。")
para(doc, "柳州不是\u201c只靠汽车\u201d——它是**汽车+螺蛳粉+机械+新能源+外向**的西南制造业复合体，靠\u201c工业+出口\u201d驱动；但**社零、人口、物价**偏弱，\u201c制造强、需求弱\u201d。")
para(doc, "一句话结论：**柳州是\u201c广西工业引擎、汽车之城\u201d；观察它先看\u201c新能源汽车、螺蛳粉、机械、东盟出口\u201d，再看\u201c社零、人口、物价、外资\u201d。**它是\u201c制造强国、外向强、需求弱\u201d的西南工业样本。")

# ---- 附录A ----
heading1(doc, "附录A：主要资料来源与核验路径")
bullet(doc, "《2025年柳州市政府工作报告》（2025年2月，2025年目标、2024年回顾）")
bullet(doc, "《2025年柳州市国民经济和社会发展统计公报》（柳州市统计局，2026-05-15，2025年实际数据）")
bullet(doc, "《2026年柳州市政府工作报告》（2026年3月，2025年复盘+2026年目标）")
bullet(doc, "柳州市政府/发改委官网（liuzhou.gov.cn/fgw）")

# ---- 附录B ----
heading1(doc, "附录B：建议建立的年度跟踪仪表盘")
bullet(doc, "GDP总量/增速 vs 全年目标（可比价）。")
bullet(doc, "规上工业增加值及分行业（汽车工业+装备+高技术）增速。")
bullet(doc, "汽车产量/新能源汽车出口、上汽五菱订单。")
bullet(doc, "螺蛳粉全产业链销售收入/出口。")
bullet(doc, "固定资产投资/产业/高技术投资增速。")
bullet(doc, "社会消费品零售总额、线上、以旧换新。")
bullet(doc, "进出口、对东盟/RCEP、电动车出口。")
bullet(doc, "一般公共预算收入、税收/非税、民生支出占比。")
bullet(doc, "常住人口/自然增长、城镇化率、城镇新增就业。")
bullet(doc, "CPI、金融存贷款。")

# ---------------- 保存 ----------------
out = "/Users/x/Desktop/content-prod-lab/reports/柳州市_2025年政府工作报告_深度研究_2026-08-14.docx"
doc.save(out)
print("SAVED OK: 柳州市", out)
