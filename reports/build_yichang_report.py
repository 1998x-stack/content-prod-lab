# -*- coding: utf-8 -*-
"""Build 宜昌市2025年政府工作报告 深度研究 DOCX, 参照地级市系列版式。"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DARK = RGBColor(0x1F, 0x2A, 0x44)
RED = RGBColor(0xB0, 0x1E, 0x2E)
GRAY = RGBColor(0x59, 0x60, 0x69)
BLUE = RGBColor(0x1F, 0x4E, 0x79)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
HEADER_FILL = "1F2A44"
K = "微软雅黑"


def set_run(run, size=11, bold=False, color=DARK, font=K):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    run.font.name = font
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), font)
    rFonts.set(qn("w:hAnsi"), font)
    rFonts.set(qn("w:eastAsia"), font)


def para(doc, text, size=11, bold=False, color=DARK, align=None,
         space_after=6, space_before=0, font=K):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        set_run(r, size=size, bold=bold or (i % 2 == 1), color=color, font=font)
    return p


def heading1(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(16)
    pf.space_after = Pt(8)
    r = p.add_run(text)
    set_run(r, size=16, bold=True, color=RED)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "B01E2E")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def heading2(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(10)
    pf.space_after = Pt(4)
    r = p.add_run(text)
    set_run(r, size=13, bold=True, color=BLUE)
    return p


def table(doc, headers, rows, widths=None, font_size=9.5):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_run(r, size=font_size, bold=True, color=WHITE)
        tcPr = hdr[i]._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), HEADER_FILL)
        tcPr.append(shd)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(val))
            set_run(r, size=font_size, color=DARK)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    return t


def quote_box(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(8)
    pf.left_indent = Pt(12)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), "B01E2E")
    pBdr.append(left)
    pPr.append(pBdr)
    r = p.add_run(text)
    set_run(r, size=11, color=DARK, font="楷体")
    return p


def bullet(doc, text, size=10.5):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.7)
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        set_run(r, size=size, bold=(i % 2 == 1), color=DARK)
    return p


# ================================================================= 文档主体
doc = Document()
sec = doc.sections[0]
sec.page_width = Cm(21)
sec.page_height = Cm(29.7)
sec.top_margin = Cm(2.4)
sec.bottom_margin = Cm(2.2)
sec.left_margin = Cm(2.5)
sec.right_margin = Cm(2.5)

st = doc.styles["Normal"]
st.font.name = K
st.font.size = Pt(11)
st.element.rPr.rFonts.set(qn("w:eastAsia"), K)


# ---- 封面 ----
para(doc, "宜昌市2025年政府工作报告", size=24, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=60, space_after=4)
para(doc, "深度研究与\u201c容易被忽视的细节\u201d分析", size=16, bold=True, color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
para(doc, "从\u201c三峡工程、水电清洁能源、磷化工、旅游之都、两坝一峡\u201d重新理解宜昌", size=12, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)
para(doc, "\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501", color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
para(doc, "研究对象：2025年宜昌市政府工作报告及同期财政、国民经济和社会发展统计资料、2026年官方复盘", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
para(doc, "标定日期：2026-08-14", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
doc.add_page_break()

# ---- 目录 ----
catalog = [
    "执行摘要",
    "一、研究口径与阅读方法",
    "二、先看宜昌的特殊底盘：三峡、水电清洁能源、磷化工、旅游之都、先进制造",
    "三、最关键的宏观错位：GDP 6464.42亿/6.1%低于7%目标，工业旅游强但地产负、消费偏弱",
    "四、容易被忽视、但信号很重的细节（15条）",
    "五、2025年GDP目标 vs 实际：对照表",
    "六、2025年增长是谁撑起来的？（结构归因）",
    "七、预算与财政的\u201c含金量\u201d",
    "八、民生底账：人口、收入与城乡",
    "九、城镇与农村：格局与均衡",
    "十、人口流入与流出",
    "十一、物价与货币环境",
    "十二、区域一体化：宜昌在\u201c湖北省域副中心+长江经济带\u201d里的位置",
    "十三、未来5—10年最值得观察的五条主线",
    "十四、最终结论：宜昌在\u201c水电+磷化工+旅游+先进制造\u201d里的增长逻辑",
    "附录A：主要资料来源与核验路径",
    "附录B：建议建立的年度跟踪仪表盘",
]
para(doc, "目　录", size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=10, space_after=10)
for item in catalog:
    para(doc, item, size=12, space_after=4)
doc.add_page_break()

# ---- 执行摘要 ----
heading1(doc, "执行摘要")
para(doc, "**核心判断**　2025年宜昌最显著的是\u201cGDP 6464.42亿元、增长6.1%不足7%目标、人均16.52万、湖北第2\u201d、\u201c规上工业+7.3%（高技术制造+15.9%、储能锂电+61.1%）\u201d、\u201c旅游1.41亿人次/1427亿/22.5%\u201d、\u201c财收327.05亿/+11.2%、进出口544.5亿/+9.3%\u201d、\u201c但固投+5.6%、社零+4.4%、CPI-0.1%\u201d。这说明宜昌在\u201c水电+磷化工+旅游+先进制造\u201d的升级中稳健增长，但**消费、地产偏弱**是短板。")
para(doc, "把2025年目标（GDP+7%/固投+9.5%/规上+9%/社零+7.5%/进出口+7.5%）、2025年统计（GDP+6.1%低于目标、规上+7.3%、固投+5.6%、社零+4.4%低于目标、财收+11.2%超额、进出口+9.3%超额）、趋势一起看，宜昌是\u201c清洁能源+磷化工+旅游\u201d路径：**水电（三峡/葛洲坝）、磷化工（精细化工47.8%）、旅游、先进制造（储能锂电）**是支柱；2025年总量6464亿居湖北第2（次于武汉）。")
para(doc, "最容易记住的一句话：**宜昌是\u201c三峡门户、世界水电之都、磷化工与旅游大市\u201d，靠\u201c水电（清洁能源）+磷化工+旅游+先进制造\u201d增长。**观察宜昌，与其只看\u201cGDP 6464亿\u201d，不如看\u201c旅游收入1427亿/+22.5%、储能锂电+61.1%、精细化工47.8%、财收+11.2%、发电量1293.98亿千瓦时\u201d。")

# ---- 一、研究口径与阅读方法 ----
heading1(doc, "一、研究口径与阅读方法")
para(doc, "本报告以\u201c宜昌市政府工作报告（2025年，陈红辉作）\u201d为起点，把\u201c2025年GDP目标（7%）\u201d与\u201c官方2025年GDP（6464.42亿元/+6.1%）\u201d并置对照，并用\u201c2025年宜昌市统计公报\u201d和\u201c2026年政府工作报告复盘\u201d作为横向核验。")
para(doc, "口径提示：\u201cGDP、规上工业增加值增速\u201d按可比价（实际增速）；\u201c投资、消费、进出口、财政收入\u201d按现价（名义增速）。涉及\u201c常住人口\u201d用统计公报常住口径（390.06万），城镇化率用官方公布值（67.31%）。")
para(doc, "指标体系（与研究口径一致）：核心看五个象限——**总量与增速（GDP）、动能（工业+投资+消费+进出口+旅游）、财政质量、民生底账、区域一体化**。")
para(doc, "特别提示（不吃老本）：宜昌是湖北\u201c省域副中心\u201d、\u201c水电之都\u201d，2024年GDP 6191亿、增速6.5%；2025年略放缓到6.1%，但**旅游、出海口、磷化工**升级仍是最大看点。真正要盯的是**生态环保约束下的产业转型、旅游质效、磷化工高端化**。")
# ---- 二、先看宜昌的特殊底盘 ----
heading1(doc, "二、先看宜昌的特殊底盘：三峡、水电清洁能源、磷化工、旅游之都、先进制造")
para(doc, "宜昌地处长江中上游接合部、三峡工程所在地，是**世界水电之都、湖北省域副中心**。2025年GDP 6464.42亿、常住人口390.06万，居湖北第2（次于武汉）。")
para(doc, "五个底盘名词，先立框架：")
bullet(doc, "**三峡工程/水电**　三峡大坝、葛洲坝电站在此，2025年发电量1293.98亿千瓦时（+13.3%），清洁能源枢纽。")
bullet(doc, "**磷化工**　全国磷矿/磷化工重镇，精细化工占比47.8%（134家沿江化工企业\u201c关改搬转\u201d全面完成）；磷酸铁/磷酸铁锂+新能源电池全产业链闭环。")
bullet(doc, "**旅游之都（两坝一峡）**　2025年接待游客1.41亿人次（+20.4%）、旅游综合收入1427亿（+22.5%）、入境游客40.4万人次（+29%）。")
bullet(doc, "**先进制造/新能源**　储能锂电+61.1%、汽车用动力锂电+85.9%、新能源船舶138艘、算力及大数据产值破千亿。")
para(doc, "这三根支柱（水电+磷化工+旅游）叠加\u201c长江经济带区位\u201d，构成宜昌独特底盘：**左手三峡水电（能源禀赋），右手磷化工（工业基础），腹地是旅游金山**。理解宜昌，先理解\u201c靠水（水电+水运）、靠磷、靠旅游\u201d。")

# ---- 三、最关键的宏观错位 ----
heading1(doc, "三、最关键的宏观错位：GDP 6464.42亿/6.1%低于7%目标，工业旅游强但地产负、消费偏弱")
para(doc, "2025年宜昌最需要辨析的一组\u201c错位\u201d：**GDP 6.1%未达7%目标（差0.9pct），规上工业+7.3%低于9%目标，固投+5.6%、社零+4.4%（低于7.5%）、CPI-0.1%、地产投资-16.4%**。")
para(doc, "为什么\u201c旅游爆发、工业不错\u201d，GDP却没达目标？三个解释：")
para(doc, "**其一，第二产业增长平淡**　工业+7.3%（制造业+8.6%），磷肥-3.2%、磷酸一铵-25.2%、水泥-12.5%等化工受产能/价格压制，\u201c总量大但结构分化\u201d。")
para(doc, "**其二，消费地产偏弱**　社零+4.4%（低于7.5%目标）、烟酒-17.1%、服装-5.4%；地产投资-16.4%、商品房销售-2.2%——**内需与地产仍在调整期**。")
para(doc, "**其三，投资结构性放缓**　固投+5.6%低于9.5%目标，其中第一产业投资-37.3%、民间投资仅+1.0%——私营投资与农业投资走弱。")
para(doc, "小结：宜昌2025年是\u201c**稳总量、强工业&旅游、弱地产&消费**\u201d的一年：水电、磷化工、旅游、先进制造撑底，但**投资节奏、消费内需、地产**偏弱，\u201c进取得分、现实失分\u201d。")

# ---- 四、容易被忽视、但信号很重的细节（15条） ----
heading1(doc, "四、容易被忽视、但信号很重的细节（15条）")
bullet(doc, "**1.储能锂电+61.1%、车用锂电+85.9%**　新能源电池（储能/动力）是除磷化工外最大新引擎，全产业链闭环。")
bullet(doc, "**2.高技术制造业+15.9%、占规上工业17.1%**　制造升级质量高（装备/电子/生物医药强）。")
bullet(doc, "**3.发电量1293.98亿千瓦时/+13.3%**　三峡、葛洲坝水电满发，清洁能源枢纽地位强化。")
bullet(doc, "**4.旅游收入1427亿/+22.5%、入境游客+29%**　\u201c两坝一峡\u201d文旅客流量质双升。")
bullet(doc, "**5.新能源船舶138艘****\u201c以内河绿色智能装备\u201d，锂电+船坞带动\u201c水路运输+制造\u201d。")
bullet(doc, "**6.精细化工占比47.8%**\u201c关改搬转\u201d后磷化工高端化（精细磷化工/电子化学品）。")
bullet(doc, "**7.进出口544.5亿/+9.3%、出口+10.8%**\u201c新能源/磷化/化工\u201d出口带动。")
bullet(doc, "**8.财收327.05亿/+11.2%、税收+6.9%**\u201c财收高增（非税+），一个安全的\u2018财政缓冲垫\u2019。\u201d")
bullet(doc, "**9.社零+4.4%、餐饮/烟酒-17.1%**\u201c消费偏弱；但线上商品零售+12.0%、化妆品/金银珠宝增速快。\u201d")
bullet(doc, "**10.居民收入：城镇50874元/+4.8%/农村27524元/+5.9%，城乡比1.85**\u201c农村快于城镇、城乡差缩小（约1.85）。\u201d")
bullet(doc, "**11.常住人口390.06万/+城镇化67.31%**\u201c湖北人口大市、城镇化率提升、常住人口与户籍相当。\u201d")
bullet(doc, "**12.港口吞吐量1.56亿吨、长江黄金水道**\u201c水运枢纽、港口吞吐量大、三峡双线扩能。\u201d")
bullet(doc, "**13.高新技术产业增加值1245.02亿/**\u201c高企突破2000家、研发投入加、“高端制造”成色足。\u201d")
bullet(doc, "**14.生磷高效利用、磷石膏消纳**\u201c磷酸铁锂正极、磷石膏建材消纳——资源循环利用领先。\u201d")
bullet(doc, "**15.\u201c十五五\u201d能耗/双碳**\u201c单位GDP能耗下降、\u201c三峡清洁能源\u201d与\u201c碳达峰\u201d协同。\u201d")
# ---- 五、2025年GDP目标 vs 实际：对照表 ----
heading1(doc, "五、2025年GDP目标 vs 实际：对照表")
table(doc,
    ["指标", "2025年目标", "2025年实际", "达成评价"],
    [
        ["地区生产总值(GDP)", "增长7%", "6464.42亿/6.1%", "未达成，差0.9pct"],
        ["规上工业增加值", "增长9%", "+7.3%", "未达成，差1.7pct"],
        ["固定资产投资", "增长9.5%", "+5.6%", "未达成，差3.9pct"],
        ["社会消费品零售总额", "增长7.5%", "2376.24亿/+4.4%", "未达成，差3.1pct"],
        ["进出口总额", "增长7.5%", "544.5亿/+9.3%", "超额"],
        ["一般公共预算收入", "——", "327.05亿/+11.2%", "财收强"],
        ["居民收入", "与经济增长同步", "城镇+4.8%/农村+5.9%", "总体同步"],
    ],
)
para(doc, "注：GDP、规上工业按可比价；投资、消费、进出口、财收按现价。**进出口（+9.3%）、财收（+11.2%）达成**，**GDP/工业/固投/社零均未达目标**；\u201c量稳质升但增速有回落\u201d是宜昌2025年的特征。")
para(doc, "拆读：**旅游（1427亿/+22.5%）、进出口（+9.3%）、财收（+11.2%）是亮色**，**GDP、工业、固投、社零、CPI偏弱**；\u201cGDP目标7%\u201d实际6.1%——\u201c进取目标、稳健收官\u201d，是湖北\u201c副中心\u201d的普通正面样本。")

# ---- 六、2025年增长是谁撑起来的？（结构归因） ----
heading1(doc, "六、2025年增长是谁撑起来的？（结构归因）")
para(doc, "把宜昌GDP的6.1%拆开：三次产业增加值分别为第一产业696.34亿（+3.8%）、第二产业2554亿（+5.8%）、第三产业3214.08亿（+6.8%），三次产业结构比10.8：39.5：49.7。**第三产业（服务业/旅游）是最大增量，第二产业（工业）是中流砥柱，第一产业（农渔）稳**。")
para(doc, "2026年宜昌强调\u201c三个重要\u201d新使命、\u201c六个奋勇争先\u201d，聚焦**磷化工高端化、新能源电池、算力/大数据、文旅、先进制造**——核心是\u201c绿色低碳转型+新质生产力\u201d。")
para(doc, "**第二产业（工业+制造业）**：规上工业+7.3%、高技术制造+15.9%；储能锂电+61.1%、发电量+13.3%、磷化工精细占比47.8%——工业升级、能源与制造并进。")
para(doc, "**第三产业（服务业）**：旅游收入+22.5%（1.41亿人次）、算力/大数据破千亿、物流水运（港口吞吐量1.56亿吨）——服务业/文旅是增长极。")
para(doc, "**外贸（开放型经济）**：进出口+9.3%（出口+10.8%）、实际使用外资+8.6%——全球化、新能源/化工出口带动。")
para(doc, "一句话归因：**2025年宜昌增长\u201c靠旅游（三产）+工业（水电/磷化工/锂电）+外贸\u201d三驾马车**，短板在\u201c投资节奏（固投+5.6%）与消费、地产\u201d。")

# ---- 七、预算与财政的\u201c含金量\u201d ----
heading1(doc, "七、预算与财政的\u201c含金量\u201d")
para(doc, "2025年宜昌**地方一般公共预算收入327.05亿元（+11.2%）**，其中税收收入236.14亿元（+6.9%）；一般公共预算支出694.98亿元（-3.6%）。财收增速远超GDP，税收占比约72%、质量较高。")
bullet(doc, "税收结构：增值税/所得税主体税种正增长，财收+11.2%一半来自非税（国企/土地/收费），财政\u201c量增\u201d但\u201c质\u201d需看税收+6.9%。")
bullet(doc, "民生与产业：争取上级资金661.07亿（+14.4%）、超长期特别国债61.4亿，专项投向新能源、磷化工、城市更新。")
bullet(doc, "金融支撑：存款7401.17亿（+9.4%）、贷款7319.14亿（+10.4%），金融充裕支撑制造/基建。")
para(doc, "**财政含金量小结**：财收+11.25%远超GDP增速、税收+6.9%，\u201c增收出量\u201d但需关注**非税依赖与财政支出下降（-3.6%）**；财政对新能源、磷化工、水利投入是大看点。")

# ---- 八、民生底账：人口、收入与城乡 ----
heading1(doc, "八、民生底账：人口、收入与城乡")
para(doc, "2025年宜昌**全体居民人均可支配收入42714元（+5.4%）**，其中城镇50874元（+4.8%）、农村27524元（+5.9%），城乡收入比约1.85（缩小）。就业：城镇新增就业9.97万人（+1.1%）、高校毕业生创业就业3.6万人。")
para(doc, "人口画像：**常住人口390.06万、城镇化率67.31%（+3.2pct）**，是湖北人口大市且城镇化率高；在校大学生突破10万，教育消费力强。")
para(doc, "民生投入：新增保障性租赁住房3710套、新改扩建学校32所/学位8900个、婴幼儿托位18789个、育儿补贴2.1亿——民生保障扎实。")

# ---- 九、城镇与农村：格局与均衡 ----
heading1(doc, "九、城镇与农村：格局与均衡")
para(doc, "宜昌常住城镇化率67.31%（湖北上游水平），城乡格局总体均衡；农村收入增速（+5.9%）高于城镇（+4.8%），**城乡收入比缩小（约1.85）**。")
para(doc, "农业底盘：粮食产量154.32万吨（+0.4%）、蔬菜586.15万吨、园林水果472.26万吨、茶叶12.98万吨、生猪出栏632.21万头——湖北粮/果/茶/猪农林大市。")
para(doc, "一句话：\u201c宜昌农业稳、农村收入快、城镇化推进\u201d，但\u201c农业基数小、城乡收入差（1.85）\u201d仍需持续收敛。")

# ---- 十、人口流入与流出 ----
heading1(doc, "十、人口流入与流出")
para(doc, "宜昌常住人口390.06万、**城镇化率67.31%**，是湖北\u201c人口大市+教育重镇\u201d；在湖北人口总体收缩背景下，**宜昌城镇化率3.2pct的\u201c十四五\u201d提升显示人口向主城集聚**。")
para(doc, "结构观察：**湖北人口向武汉集中，宜昌作为副中心承接周边地市人口产业**；高校在校10万人（三峡大学等）、岗位多于区县，城镇化与就业带。")
para(doc, "2026年目标：城镇新增就业9万以上、高校毕业生创业就业、返乡创业6000个——宜昌把\u201c人口+就业\u201d作为\u201c省域副中心\u201d量能的基础。")

# ---- 十一、物价与货币环境 ----
heading1(doc, "十一、物价与货币环境")
para(doc, "2025年宜昌CPI**全年下降0.1%**，其中交通通信-2.9%、食品烟酒-0.4%、居住-0.1%；衣着+3.7%、教育文娱+0.4%——**低通胀、需求温和**，与全国趋势一致。")
bullet(doc, "信贷：存款7401.41亿（+9.4%）、贷款7319.14亿（+10.4%），信贷充裕支持制造/基建。")
bullet(doc, "电价优势：水电绿电、全社会用电量292.69亿千瓦时（+8.2%）——能源/工业景气。")
para(doc, "货币环境判断：**低通胀、宽信用**并存，\u201c资金充裕、物价温和\u201d；宜昌依托水电清洁能源、工业景气，2026年CPI目标2%左右、扩内需。")

# ---- 十二、区域一体化：宜昌的位置 ----
heading1(doc, "十二、区域一体化：宜昌在\u201c湖北省域副中心+长江经济带\u201d里的位置")
para(doc, "宜昌是**湖北省域副中心城市、长江中上游区域性中心城市**，处于\u201c长江经济带+中部崛起+省域\u2018一主两副\u2019\u201d的枢纽位置。")
bullet(doc, "**长江黄金水道**　港口吞吐量1.56亿吨，三峡工程与双线水运新通道（可研获批、前期启动）放大\u201c水利枢纽+物流\u201d。")
bullet(doc, "**省域分工**　宜昌与襄阳并列湖北\u201c两大省域副中心\u201d，承载武汉产业疏解与中部制造业承接。")
bullet(doc, "**开放带**　高铁（沿江/郑万）、宜昌三峡国际机场（旅客破340万），对接成渝、中部与全国。")
para(doc, "一句话：**宜昌在\u201c湖北省域副中心+长江经济带\u201d里，最核心的定位是\u201c水电枢纽+磷化工基地+文旅标杆\u201d**——区位、水陆资源是宜昌的最大护城河。")

# ---- 十三、未来5—10年最值得观察的五条主线 ----
heading1(doc, "十三、未来5—10年最值得观察的五条主线")
bullet(doc, "**主线一：磷化工高端化**\u201c精细磷化工占比47.8%、磷酸铁锂\u201d能否把\u201c磷肥\u201d拉到\u201c新能源/电子化学品\u201d。")
bullet(doc, "**主线二：新能源电池与储能（锂电+）**\u201c储能锂电+61.1%、磷酸铁锂53万吨\u201d能否撑起\u201c动力/储能\u201d新曲线。")
bullet(doc, "**主线三：水电与绿电经济**\u201c三峡、葛洲坝满发+三峡双线扩能\u201d能否把\u201c绿电低碳\u201d变成产业与碳汇优势。")
bullet(doc, "**主线四：文旅质效**\u201c游客1.41亿/收入1427亿\u201d能否从\u201c走量\u201d走向\u201c人均消费、入境化\u201d。")
bullet(doc, "**主线五：生态与制造业平衡**\u201c关改搬转\u201d后如何\u201c磷石膏消纳、能耗双控、双碳\u201d，守住景区/长江生态。")

# ---- 十四、最终结论 ----
heading1(doc, "十四、最终结论：宜昌在\u201c水电+磷化工+旅游+先进制造\u201d里的增长逻辑")
para(doc, "把2025年闭环打穿：**宜昌是\u201c三峡水电之都、磷化工与旅游强市\u201d**：GDP 6464.42亿/+6.1%、规上工业+7.3%、旅游收入1427亿/+22.5%、财收+11.2%、进出口+9.3%。")
para(doc, "宜昌不是\u201c只有一个三峡\u201d——它是**水电+磷化工+旅游+先进制造（锂电）**的复合经济，靠\u201c绿色能源+高端化工+文旅\u201d驱动；同时投资、消费、地产偏弱，需\u201c扩内需+强投资\u201d。")
para(doc, "一句话结论：**宜昌是\u201c世界水电之都、磷化工与旅游大市\u201d；观察它先看\u201c水电绿电、磷化工、旅游、储能锂电\u201d，再看\u201c固投、社零、物价、人口\u201d。**它是\u201c绿电菜单、文旅爆发、需求偏弱\u201d的湖北\u201c省域副中心\u201d样本。")

# ---- 附录A ----
heading1(doc, "附录A：主要资料来源与核验路径")
bullet(doc, "《2025年宜昌市政府工作报告》（2025年1月5日，陈红辉作，2025年目标、2024年回顾\u201cGDP 6191、增速6.5%\u201d）")
bullet(doc, "《2025年宜昌市国民经济和社会发展统计公报》（宜昌市统计局，2026-04，2025年实际数据）")
bullet(doc, "《2026年宜昌市政府工作报告》（宜昌市政府，2026年1月，2025年复盘+2026年目标）")
bullet(doc, "宜昌市人民政府官网、宜昌市统计局官方URL（yichang.gov.cn）")
bullet(doc, "宜昌市国民经济和社会发展统计公报（统计公报库/PDF）")

# ---- 附录B ----
heading1(doc, "附录B：建议建立的年度跟踪仪表盘")
bullet(doc, "GDP总量/增速 vs 全年目标（可比价）。")
bullet(doc, "规上工业增加值及分行业（水电/磷化工/锂电/高技术）增速。")
bullet(doc, "发电量、三峡/葛洲坝电力、清洁能源占比。")
bullet(doc, "磷化工产量（磷肥/磷酸铁锂）、精细磷化工占比、磷石膏利用率。")
bullet(doc, "固定资产投资/工业/基建/房地产投资增速。")
bullet(doc, "旅游人数/收入/入境游客、交通吞吐量。")
bullet(doc, "进出口、出口、实际使用外资。")
bullet(doc, "一般公共预算收入、税收/非税、支出。")
bullet(doc, "常住人口、城镇化率、城镇新增就业。")
bullet(doc, "CPI、金融存贷款、电价/用电量。")

# ---------------- 保存 ----------------
out = "/Users/x/Desktop/content-prod-lab/reports/宜昌市_2025年政府工作报告_深度研究_2026-08-14.docx"
doc.save(out)
print("SAVED OK: 宜昌市", out)
