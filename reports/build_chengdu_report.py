# -*- coding: utf-8 -*-
"""Build 重庆市2025年政府工作报告 深度研究 DOCX, 参照北京/上海/杭州/天津系列版式。"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DARK = RGBColor(0x1F, 0x2A, 0x44)
RED = RGBColor(0xB0, 0x1E, 0x2E)
GRAY = RGBColor(0x59, 0x60, 0x69)
BLUE = RGBColor(0x1F, 0x4E, 0x79)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
HEADER_FILL = "1F2A44"
K = "微软雅黑"


def set_run(run, size=11, bold=False, color=DARK, font=K):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    run.font.name = font
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), font)
    rFonts.set(qn("w:hAnsi"), font)
    rFonts.set(qn("w:eastAsia"), font)


def para(doc, text, size=11, bold=False, color=DARK, align=None,
         space_after=6, space_before=0, font=K):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        set_run(r, size=size, bold=bold or (i % 2 == 1), color=color, font=font)
    return p


def heading1(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(16)
    pf.space_after = Pt(8)
    r = p.add_run(text)
    set_run(r, size=16, bold=True, color=RED)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "B01E2E")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def heading2(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(10)
    pf.space_after = Pt(4)
    r = p.add_run(text)
    set_run(r, size=13, bold=True, color=BLUE)
    return p


def table(doc, headers, rows, widths=None, font_size=9.5):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_run(r, size=font_size, bold=True, color=WHITE)
        tcPr = hdr[i]._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), HEADER_FILL)
        tcPr.append(shd)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(val))
            set_run(r, size=font_size, color=DARK)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    return t


def quote_box(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(8)
    pf.left_indent = Pt(12)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), "B01E2E")
    pBdr.append(left)
    pPr.append(pBdr)
    r = p.add_run(text)
    set_run(r, size=11, color=DARK, font="楷体")
    return p


def bullet(doc, text, size=10.5):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.7)
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        set_run(r, size=size, bold=(i % 2 == 1), color=DARK)
    return p


# ================================================================= 文档主体
doc = Document()
sec = doc.sections[0]
sec.page_width = Cm(21)
sec.page_height = Cm(29.7)
sec.top_margin = Cm(2.4)
sec.bottom_margin = Cm(2.2)
sec.left_margin = Cm(2.5)
sec.right_margin = Cm(2.5)

st = doc.styles["Normal"]
st.font.name = K
st.font.size = Pt(11)
st.element.rPr.rFonts.set(qn("w:eastAsia"), K)


# ---- 封面 ----
para(doc, "成都市2025年政府工作报告", size=24, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=60, space_after=4)
para(doc, "深度研究与“容易被忽视的细节”分析", size=16, bold=True, color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
para(doc, "从“雪山下的公园城市、成渝双城、电子信息与人口虹吸”重新理解成都", size=12, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)
para(doc, "━━━━━━━━━━━━━━━━", color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
para(doc, "研究对象：2025年成都市政府工作报告及同期财政、国民经济和社会发展统计资料、2026年官方复盘", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
para(doc, "整理时间：2026年8月", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
para(doc, "分析性质：分析性研究 ｜ 非官方解读", size=11, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
doc.add_page_break()

# ---- 目录 ----
catalog = [
    "执行摘要",
    "一、研究口径与阅读方法",
    "二、先看成都的特殊底盘：西部极核、公园城市与人口虹吸",
    "三、最关键的宏观错位：GDP破2.48万亿、工业稳，但地产/外贸与物价偏弱",
    "四、容易被忽视、但信号很重的细节（15条）",
    "五、把所有细节连起来：成都正在经历的“六个换挡”",
    "六、增长暗线：电子信息+新能源+民生的“新质生产力”",
    "七、财政暗线：收入稳、税收稳，民生支出大",
    "八、产业暗线：从“电子信息/白酒”到“先进制造+新质生产力”",
    "九、区域格局：成渝双城经济圈与成都都市圈",
    "十、人口与城市：人口虹吸、公园城市与老龄化",
    "十一、民营经济：占50.6%、消费强市",
    "十二、事后验证：用2025年实际执行结果检验报告判断",
    "十三、未来5—10年最值得观察的五条主线",
    "十四、最终结论：成都在“成渝双城+西部极核”里的增长逻辑",
    "附录A：主要资料来源与核验路径",
    "附录B：建议建立的年度跟踪仪表盘",
]
para(doc, "目　录", size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=10, space_after=10)
for item in catalog:
    para(doc, item, size=12, space_after=4)
doc.add_page_break()

# ---- 执行摘要 ----
heading1(doc, "执行摘要")
para(doc, "**核心判断**　如果只看新闻标题，2025年成都最显眼的是“GDP突破2.48万亿、增长5.8%”、“常住人口2153万、净增6.1万”和“五大先进制造业+8.1%”。但这份政府工作报告调研真正值得深读的，是这一座“雪山下的公园城市”如何既做“成渝双城经济圈的极核”，又靠“电子信息+新能源车+消费+人口虹吸”撑起西部第一大中心城市。")
para(doc, "把2025年初设定的目标、2025年《国民经济和社会发展统计公报》、以及2026年报告对2025年的复盘放在一起看，成都呈现清晰暗线：**从“人口+地产+基建”的旧依赖，向“新质生产力+消费+西南开放枢纽+人口虹吸”转型**。旧引擎（地产、一般基建）在调整；新引擎（电子信息、装备、新能源、软件服务、消费、科技创新）被要求更快补位。")
para(doc, "因此，本报告不按政府工作报告逐段复述，而采用“显性表述—同期数据—制度含义—长期影响”的方式，专门提取容易被忽视、但对判断成都未来5—10年发展模式更有价值的信号。")
para(doc, "最容易记住的一句话：**成都是“西部极核+公园城市”双标签的城市，持续吸引西部人口与人才，也在电子信息、新质生产力、消费与西南开放上做文章。**观察成都，与其看“GDP 2.48亿”，不如看“人口净流入、高技术制造、消费、装配与新质生产力兑现”。")
heading2(doc, "一页速览：2025年成都经济的“表与里”")
table(doc,
    ["维度", "表面现象", "更深层信号"],
    [
        ["增长", "GDP 24763.6亿、+5.8%", "西部中心城市、服务业引领"],
        ["产业", "规上工业+7.0%、高技术+8.9%", "装备+16.1%、电子信息+9.8%"],
        ["消费", "社零11434.1亿、+5.5%", "新能源汽车零售+45.8%、通讯+70.6%"],
        ["外贸", "进出口8502.3亿、+1.4%", "高新技术出口+11.0%"],
        ["投资", "固定资产投资+2.2%、工业+19.7%", "工业/高技术制造投资强"],
        ["财政", "一般公共预算收入2000.7亿、+2.6%", "财政稳健"],
        ["人口", "常住2153.5万、城镇化率81.46%", "人口+6.1万、西部虹吸"],
        ["物价", "CPI+0.1%", "物价平稳"],
    ],
    widths=[2.2, 5.2, 8.6])
para(doc, "", size=4, space_after=2)

# ---- 一、研究口径与阅读方法 ----
heading1(doc, "一、研究口径与阅读方法")
heading2(doc, "1.1 本报告的三份核心底稿")
bullet(doc, "**2025年《政府工作报告》**：2025年2月在市十八届人大三次会议上作，给出全年预期目标（GDP增长5.7%以上、一产3%以上等）。部分指标以方向性要求投入。")
bullet(doc, "**2025年《国民经济和社会发展统计公报》**：成都市统计局2026年4月15日发布，给出全年实际执行数与增速，是“事后验证”的锚点。")
bullet(doc, "**2026年成都市政府工作报告及官方复盘**：对2025执行结果的追认，交叉验证“目标—执行—再评价”三段式。")
heading2(doc, "1.2 阅读方法：显性—数据—制度—长期")
para(doc, "每一章按“**显性表述 → 同期数据 → 制度含义 → 长期影响**”四层展开。其中“制度含义”回答“政府为什么这么排优先序”，“长期影响”回答“这对未来5—10年意味着什么”。")
para(doc, "关键判别：**当报告使用的“增长词”和统计公报里的实际数不一致时，数据优先。**例如2025年GDP目标5.7%以上，实际+5.8%略超；固定资产投资、进出口是西部城市里相对弱的环节。")

# ---- 二、特殊底盘 ----
heading1(doc, "二、先看成都的特殊底盘：西部极核、公园城市与人口虹吸")
para(doc, "在所有全国主要城市里，成都的“底盘”独特：**西部极核+公园城市+人口虹吸高地**三合一。常住2153.5万、城镇化率81.46%，是全国人口最多的超大城市之一，也是“雪山下的公园城市”与“天府之国”的文旅与宜居名片。")
para(doc, "这决定成都的多重身份并存：**西部极核**（成渝双城经济圈、国家中心城市）、**制造与电子基地**（电子信息、装备、新能源）、**消费强市**（社零1.14万亿）、**科技文创**（高校66所、高新技术企业破1.5万家）、**西南门户**（双机场、中欧班列枢纽）。")
heading2(doc, "2.1 电子信息与制造业")
para(doc, "成都以电子信息（京东方、成都集成电路）见长，2025年电子信息产业+9.8%、装备+16.1%、汽车制造+17.8%、工业机器人、服务器等新赛道放量。制造业转型升级是成都的重心。")
heading2(doc, "2.2 人口虹吸与公园城市")
para(doc, "成都持续吸引西部人口/人才，常住人口净增（+6.1万），城镇化81.5%。同时推进“雪山下的公园城市”，生态/文旅/宜居成为聚人聚产的名片。")
heading2(doc, "2.3 成渝双城与门户")
para(doc, "成都领衔成渝地区双城经济圈，航空（天府/双流）、中欧班列（成都）连接“一带一路”。是西部大开发与西南开放的重要门户。")

# ---- 三、宏观错位 ----
heading1(doc, "三、最关键的宏观错位：GDP破2.48万亿、工业稳，但地产/外贸与物价偏弱")
para(doc, "把2025年成都的宏观面放进一张表，会出现令人惊讶的“错位”：表观增长来自工业/消费/服务，而外贸、地产与价格偏弱。这个错位，正是读懂成都的关键。")
heading2(doc, "3.1 “强的一面”")
table(doc,
    ["指标", "2025实绩", "信号"],
    [
        ["GDP", "24763.6亿、+5.8%", "西部中心城市、服务业+6.1%"],
        ["规上工业", "+7.0%", "汽车+17.8%、装备+16.1%"],
        ["高技术制造", "+8.9%", "新质生产力"],
        ["社零", "+5.5%", "消费强市、新能源车+45.8%"],
        ["工业投资", "+19.7%", "工业/制造投资强"],
        ["民营经济", "占50.6%、+6.2%", "民营底盘厚"],
    ],
    widths=[3.2, 5.2, 6.0])
heading2(doc, "3.2 “弱的一面”")
table(doc,
    ["指标", "2025实值", "信号"],
    [
        ["进出口", "+1.4%（增速低）", "外贸/电子外需波动"],
        ["房地产开发", "+0.9%（销售-4.4%）", "地产开工/销售偏弱"],
        ["房地产增加值", "-1.0%", "地产拖累"],
        ["一般公共预算收入", "+2.6%", "收入低增"],
        ["人口", "+6.1万（少于往年）", "人口吸纳放缓"],
    ],
    widths=[3.2, 5.2, 6.0])
para(doc, "**错位结论**　成都的增长“很强、但也很矛盾”。强的部分（工业/消费/服务业/民营）与弱的部分（地产/外贸/物价/人口增速放缓）并存。**真正的焦点是“经济增长靠消费/服务业，但外贸地产承压”**。2026年成都重点在“新质生产力+成渝协同+扩大内需+稳住外贸”。")

# ---- 四、被忽视细节 ----
heading1(doc, "四、容易被忽视、但信号很重的细节（15条）")
details = [
    ("1", "规上工业+7.0%，其中汽车制造+17.8%", "新能源汽车、汽车制造景气。"),
    ("2", "五大先进制造业+8.1%，装备制造+16.1%", "制造业转型升级主引擎。"),
    ("3", "电子信息产业+9.8%", "成都电子信息全国领先。"),
    ("4", "高技术制造+8.9%、高技术制造投资+23.4%", "新质生产力投资猛。"),
    ("5", "新能源汽车零售额+45.8%、通讯器材+70.6%", "消费升级、新能源车热销。"),
    ("6", "社会消费品零售11434.1亿、+5.5%", "西部消费强市。"),
    ("7", "进出口8502.3亿、+1.4%，高新技术出口+11.0%", "外贸稳增、高新出口强。"),
    ("8", "固定资产投资+2.2%、工业投资+19.7%", "工业投资强、地产偏弱。"),
    ("9", "常住2153.5万、城镇化率81.46%", "人口虹吸、西部城区极强。"),
    ("10", "民营经济占50.6%、+6.2%", "民营强、市场活跃。"),
    ("11", "在蓉高校66所、在校131.6万人", "人才培养/消费双引擎。"),
    ("12", "高新技术企业突破1.5万家", "科创主体壮大。"),
    ("13", "轨道交通运营里程739.5公里", "超大轨道交通网络。"),
    ("14", "一般公共预算收入+2.6%、税收+2.6%", "财政稳健。"),
    ("15", "CPI+0.1%", "物价平稳。"),
]
for d in details:
    para(doc, "**%s**" % d[0], size=11)
    para(doc, d[1], size=10.5, color=GRAY)

# ---- 五、六个换挡 ----
heading1(doc, "五、把所有细节连起来：成都正在经历的“六个换挡”")
gear = [
    ("1．动能换挡：从“人口/地产驱动”到“新质生产力+消费”", "高技术制造+8.9%、工业投资+19.7%、社零+5.5%。"),
    ("2．产业换挡：从“电子信息/白酒”到“先进制造+新质生产力”", "装备+16.1%、汽车+17.8%、电子信息+9.8%。"),
    ("3．投资换挡：从“基建/地产”到“工业/高技术制造投资”", "工业投资+19.7%、高技术制造投资+23.4%、地产偏弱。"),
    ("4．开放换挡：从“传统电子外贸”到“高新技术+西南门户”", "高新出口+11.0%、中欧/航空枢纽。"),
    ("5．人口换挡：从“大规模流入”到“高质量集聚+存量”", "常住+6.1万、城镇化81.5%。"),
    ("6．社会换挡：从“追求规模”到“公园城市+民生/消费”", "文旅、宜居、消费与生态协同。"),
]
for g in gear:
    para(doc, "**%s**　%s" % g, space_after=6)

# ---- 六、增长暗线 ----
heading1(doc, "六、增长暗线：电子信息+新能源+消费的“新质生产力”")
heading2(doc, "6.1 电子信息与新能源")
para(doc, "成都是西部电子信息/半导体高地和新能源车消费重镇。电子信息+9.8%、汽车+17.8%、新能源车零售+45.8%、工业机器人/服务器放量，说明“新质生产力”正从制造与消费两端兑现。")
heading2(doc, "6.2 消费强市")
para(doc, "社零1.14万亿、+5.5%，新能源汽车、通讯、化妆品、金银珠宝等升级类消费高增。成都是“公园城市+消费强市”双标签，人口与消费相互强化。")
para(doc, "**这条暗线意味着**：成都的增长叙事正从“人口+地产”转向“新质生产力+消费+西南门户”。看成都，盯住“高技术制造、新能源车、消费、人口净流入”这几组数。")

# ---- 七、财政暗线 ----
heading1(doc, "七、财政暗线：收入稳、税收稳，民生支出大")
para(doc, "2025年成都一般公共预算收入2000.7亿元、+2.6%，税收1408.9亿元、+2.6%。收入与税收同步稳增、质量尚可，来自制造/消费/服务税源。支出2680亿、民生投入大。")
para(doc, "**制度含义**　成都财政“收入稳、税收稳”，但要摆脱对地产/一次性收入依赖，靠“新质生产力+消费+城市价值”产生真实税源。")

# ---- 八、产业暗线 ----
heading1(doc, "八、产业暗线：从“电子信息/白酒”到“先进制造+新质生产力”")
heading2(doc, "8.1 成都产业的“表”")
table(doc,
    ["指标", "2025增速/信号", "解读"],
    [
        ["规上工业", "+7.0%", "制造强市"],
        ["汽车制造", "+17.8%", "汽车/新能源"],
        ["装备制造", "+16.1%", "高端制造"],
        ["电子信息", "+9.8%", "电子高地"],
        ["高技术制造", "+8.9%", "新质生产力"],
        ["新能源车零售", "+45.8%", "消费升级"],
        ["社会消费品零售", "+5.5%", "消费强市"],
    ],
    widths=[4.6, 3.4, 5.2])
heading2(doc, "8.2 从“电子/消费”到“制造+新质”")
para(doc, "成都过去以电子信息、白酒、文旅见长，2025年显示“先进制造（装备、汽车、半导体）+新质生产力”正加速成为新增长。科技创新、高企1.5万家、轨道交通是大盘。")

# ---- 九、区域 ----
heading1(doc, "九、区域格局：成渝双城经济圈与成都都市圈")
para(doc, "成都领衔“成渝地区双城经济圈”，与重庆联动承接西部内陆开放、先进制造、文旅、人才。成都都市圈（成德眉资）带动四川经济高质量发展。")
para(doc, "天府新区、成都高新区等创新平台，加上机场群/中欧班列，构建西部枢纽。成都—重庆“双核”并进，是西部开发与双循环的重要支点。")

# ---- 十、人口与城市 ----
heading1(doc, "十、人口与城市：人口虹吸、公园城市与潜在老龄化")
heading2(doc, "10.1 人口总量")
para(doc, "2025年末成都常住2153.5万、城镇化率81.46%，净增6.1万。成都持续吸引西部人口/人才，是西部最大超大城市；户籍人口1637.3万（+13.8万，回流）。")
heading2(doc, "10.2 城市与高质量")
para(doc, "成都推进“雪山下的公园城市”，生态、文旅、消费、地铁（739.5公里）等支撑城市的人居吸引力。高素质户籍+人才回流，是“高质量人口”的路子。")

# ---- 十一、民营经济 ----
heading1(doc, "十一、民营经济：占50.6%、市场活跃")
heading2(doc, "11.1 民企地位")
para(doc, "成都民营经济增加值12528.9亿元、占GDP 50.6%、+6.2%。民营是成都最活跃的“引擎”，尤其在消费、软件、创新上。")
heading2(doc, "11.2 政策与服务")
para(doc, "成都持续优化营商环境、支持创业创新。民营+专精特新，是成都未来10年高质量与消费的关键。")

# ---- 十二、事后验证 ----
heading1(doc, "十二、事后验证：用2025年实际执行结果检验报告判断")
heading2(doc, "12.1 年初目标与年末执行对比")
table(doc,
    ["指标", "2025年初目标", "2025实际", "偏差"],
    [
        ["GDP增速", "5.7%以上", "+5.8%", "略超"],
        ["一产", "3%以上", "+3.3%", "达标"],
        ["社零", "未设具体", "+5.5%", "达预期"],
        ["进出口", "未设", "+1.4%", "偏低"],
        ["固定资产投资", "未设", "+2.2%", "偏低"],
        ["一般公共预算收入", "未设", "+2.6%", "稳"],
    ],
    widths=[3.0, 3.2, 3.0, 4.2])
heading2(doc, "12.2 判断验证")
para(doc, "三条核心判断得到支持：**（1）工业/制造/高技术强（规上+7.0%、装备+16.1%），验证；（2）消费强市（社零+5.5%、新能源车+45.8%），验证；（3）地产/外贸/人口放缓偏弱，验证。**")
para(doc, "核心观察：成都靠“新质生产力+消费+人口+民营”守住5.8%增长、西部第一。地产、进出口、人口吸纳放缓是2026年的重点变量。")

# ---- 十三、五条主线 ----
heading1(doc, "十三、未来5—10年最值得观察的五条主线")
lines5 = [
    ("① 电子信息/装备制造（新质生产力）", "能否从制造/代工升级为“智造+科创”。"),
    ("② “成渝双城”协同与西部门户", "双极联动能否把成都/重庆做强、西南开放。"),
    ("③ 消费与公园城市", "社零、新能源车/消费升级能否持续，文旅生态能否变现。"),
    ("④ 人口虹吸与公共服务", "人口净流入、高质量人才、市场化能不能维持。"),
    ("⑤ 民营/科创新动能", "高企1.5万家、民营50.6%，能否孵化更多独角兽。"),
]
for l in lines5:
    para(doc, "**%s**　%s" % l, space_after=6)

# ---- 十四、结论 ----
heading1(doc, "十四、最终结论：成都在“成渝双城+西部极核”里的增长逻辑")
para(doc, "成都的2025年，是“**新质生产力+消费+西南门户”为核心，而对地产/外贸/人口依赖略降**的答卷：GDP 2.48万亿、西部第一、高技术/消费表现好，地产/进出口/人口增速放缓。")
para(doc, "只要电子信息、装备、新能源车、消费/民营/人口能接住，成都就仍是西部极核城市；如果地产、外贸、人口偏弱，成都同样面临结构调整。")
para(doc, "最稳妥的观察信号：**一盯高技术/装备（动能）、二盯消费/新能源（内需）、三盯人口/民营（底座）、四盯外贸/地产（约束）、五盯成渝协同（区域）。**成都，是西部“公园城市+极核增长”的新样版。")

# ---- 附录A ----
heading1(doc, "附录A：主要资料来源与核验路径")
bullet(doc, "成都市2025年《政府工作报告》（2025年2月）——目标来源。")
bullet(doc, "《2025年成都市国民经济和社会发展统计公报》（成都市统计局，2026-04-15）——GDP、工业、人口实值。")
bullet(doc, "成都市统计/运行分析、成渝专报——区域与消费。")
bullet(doc, "2026年成都市政府工作报告——2025执行复盘。")
bullet(doc, "成都海关、市商务部、市财政局——贸易与财政实况。")
heading2(doc, "核验说明")
para(doc, "本报告所有增速、总量、占比均以统计公报/官方发布为基准。涉“电子信息产业增加值”等以官方口径为准。")

# ---- 附录B ----
heading1(doc, "附录B：建议建立的年度跟踪仪表盘")
para(doc, "建议把下面10个“测脉搏”指标做成年度仪表盘：")
dash = [
    ["#", "指标", "2025基值", "为什么重要"],
    ["1", "GDP增速", "+5.8%", "总量与方向"],
    ["2", "规上工业增速", "+7.0%", "制造底盘"],
    ["3", "高技术制造增速", "+8.9%", "新质生产力"],
    ["4", "社零增速", "+5.5%", "消费强市"],
    ["5", "出口增速", "+1.4%", "外贸韧性"],
    ["6", "固定资产投资/工业投资", "+2.2%/+19.7%", "投资结构"],
    ["7", "常住人口/城镇化率", "2153.5万/81.46%", "人口虹吸"],
    ["8", "一般公共预算收入增速", "+2.6%", "财政质量"],
    ["9", "民营经济占比", "50.6%", "民营活力"],
    ["10", "居民人均可支配收入增速", "+4.5%", "民生获得感"],
]
table(doc, dash[0], dash[1:], widths=[0.8, 4.5, 3.2, 5.0])
para(doc, "把这10个指标连起来看，任何一个“新引擎（2/3/4）向上、旧引擎（5/6）修复”，都说明成都真正在换挡。")

# ========================================= 保存
out = "/Users/x/Desktop/content-prod-lab/reports/成都市_2025年政府工作报告_深度研究_2026-08-13.docx"
doc.save(out)
print("SAVED:", out)
print("PARAGRAPHS:", len(doc.paragraphs))
print("TABLES:", len(doc.tables))
