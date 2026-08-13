# -*- coding: utf-8 -*-
"""Build 甘肃省2025年政府工作报告 深度研究 DOCX, 参照省系列版式。"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DARK = RGBColor(0x1F, 0x2A, 0x44)
RED = RGBColor(0xB0, 0x1E, 0x2E)
GRAY = RGBColor(0x59, 0x60, 0x69)
BLUE = RGBColor(0x1F, 0x4E, 0x79)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
HEADER_FILL = "1F2A44"
K = "微软雅黑"


def set_run(run, size=11, bold=False, color=DARK, font=K):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    run.font.name = font
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), font)
    rFonts.set(qn("w:hAnsi"), font)
    rFonts.set(qn("w:eastAsia"), font)


def para(doc, text, size=11, bold=False, color=DARK, align=None,
         space_after=6, space_before=0, font=K):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        set_run(r, size=size, bold=bold or (i % 2 == 1), color=color, font=font)
    return p


def heading1(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(16)
    pf.space_after = Pt(8)
    r = p.add_run(text)
    set_run(r, size=16, bold=True, color=RED)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "B01E2E")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def heading2(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(10)
    pf.space_after = Pt(4)
    r = p.add_run(text)
    set_run(r, size=13, bold=True, color=BLUE)
    return p


def table(doc, headers, rows, widths=None, font_size=9.5):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_run(r, size=font_size, bold=True, color=WHITE)
        tcPr = hdr[i]._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), HEADER_FILL)
        tcPr.append(shd)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(val))
            set_run(r, size=font_size, color=DARK)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    return t


def quote_box(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(8)
    pf.left_indent = Pt(12)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), "B01E2E")
    pBdr.append(left)
    pPr.append(pBdr)
    r = p.add_run(text)
    set_run(r, size=11, color=DARK, font="楷体")
    return p


def bullet(doc, text, size=10.5):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.7)
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        set_run(r, size=size, bold=(i % 2 == 1), color=DARK)
    return p


# ================================================================= 文档主体
doc = Document()
sec = doc.sections[0]
sec.page_width = Cm(21)
sec.page_height = Cm(29.7)
sec.top_margin = Cm(2.4)
sec.bottom_margin = Cm(2.2)
sec.left_margin = Cm(2.5)
sec.right_margin = Cm(2.5)

st = doc.styles["Normal"]
st.font.name = K
st.font.size = Pt(11)
st.element.rPr.rFonts.set(qn("w:eastAsia"), K)



# ---- 封面 ----
para(doc, "甘肃省2025年政府工作报告", size=24, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=60, space_after=4)
para(doc, "深度研究与\u201c容易被忽视的细节\u201d分析", size=16, bold=True, color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
para(doc, "从\u201c河西走廊、新能源、中医药与兰白都市圈\u201d重新理解甘肃", size=12, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)
para(doc, "\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501", color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
para(doc, "研究对象：2025年甘肃省政府工作报告及同期财政、国民经济和社会发展统计资料、2026年官方复盘", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
para(doc, "标定日期：2026-08-13", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
doc.add_page_break()

# ---- 目录 ----
catalog = [
    "执行摘要",
    "一、研究口径与阅读方法",
    "二、先看甘肃的特殊底盘：河西走廊、新能源、中医药与兰白都市圈",
    "三、最关键的宏观错位：GDP破1.37万亿、工业/外贸强，但消费/人口/固投偏弱",
    "四、容易被忽视、但信号很重的细节（15条）",
    "五、2025年GDP目标 vs 实际：对照表",
    "六、2025年增长是谁撑起来的？（结构归因）",
    "七、预算与财政的\u201c含金量\u201d",
    "八、民生底账：人口、收入与城乡",
    "九、城镇与农村：格局与均衡",
    "十、人口流入与流出",
    "十一、物价与货币环境",
    "十二、区域一体化：甘肃在\u201c兰白都市圈+河西走廊+黄河几字弯\u201d里的位置",
    "十三、未来5—10年最值得观察的五条主线",
    "十四、最终结论：甘肃在\u201c新能源+工业+中医药\u201d里的增长逻辑",
    "附录A：主要资料来源与核验路径",
    "附录B：建议建立的年度跟踪仪表盘",
]
para(doc, "目　录", size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=10, space_after=10)
for item in catalog:
    para(doc, item, size=12, space_after=4)
doc.add_page_break()

# ---- 执行摘要 ----
heading1(doc, "执行摘要")
para(doc, "**核心判断**　如果只看新闻标题，2025年甘肃最显著的是\u201cGDP破1.37万亿、增长5.8%（全国第2）\u201d、\u201c规上工业+9.5%（全国第3）\u201d、\u201c新能源装机超1.2亿千瓦\u201d、\u201c进出口首破700亿、+16.2%\u201d。但这份研究真正值得深读的，是这座\u201c河西走廊+新能源+中医药\u201d的西部省份，如何在常住人口自然负增长（-2.29‰）、固投-1.3%背景下，靠\u201c工业（有色/电力/制造）+新能源+农业/中医药\u201d实现5.8%的增长。")
para(doc, "把2025年初设定的目标（GDP 5.5%左右）、2025年《国民经济和社会发展统计公报》、2026年复盘放在一起看，甘肃呈现清晰暗线：**从\u201c有色/能源\u201d既有底盘，向\u201c新能源+装备+中医药+算力\u201d升级**。GDP/工业超预期、新能源是最大亮点。")
para(doc, "因此，本报告不按政府工作报告逐段复述，而采用\u201c显性表述—同期数据—制度含义—长期影响\u201d的方式，专门提取容易被忽略、但对判断甘肃未来5—10年发展模式更有价值的信号。")
para(doc, "最容易记住的一句话：**甘肃是\u201c河西走廊+新能源+中医药+兰白\u201d的西部省份，靠\u201c工业+新能源+农业医药\u201d撑起增长。**观察甘肃，与其只看\u201cGDP 1.37万亿\u201d，不如看\u201c规上工业+9.5%、新能源装机、陇药、河西走廊\u201d这几张名片。")
heading2(doc, "一页速览：2025年甘肃经济的\u201c表与里\u201d")
table(doc,
    ["维度", "表面现象", "更深层信号"],
    [
        ["增长", "GDP 13697.5亿、+5.8%（全国第2）", "一产12.9%、二产33.3%、三产53.8%"],
        ["产业", "规上工业+9.5%", "有色+19.5%、电力+18.5%、制造+9.3%"],
        ["外贸", "进出口711.7亿、+16.2%", "出口+44.5%、一带一路占70.7%"],
        ["投资", "固投-1.3%、基建+14.5%", "工业+2.9%、民间-5.6%"],
        ["消费", "社零4237.6亿、+2.5%", "网上零售+36.0%"],
        ["人口", "常住2443万、城镇化58.08%", "自然增-2.29‰"],
        ["能源", "装机12442万千瓦、风/光", "新能源+中药+算力"],
    ],
    widths=[2.2, 5.6, 7.4])
para(doc, "", size=10, space_after=2)

# ---- 一、研究口径与阅读方法 ----
heading1(doc, "一、研究口径与阅读方法")
heading2(doc, "1.1 本报告的核心底稿")
bullet(doc, "**2025年《政府工作报告》**（2025年1月）——目标：GDP 5.5%左右。")
bullet(doc, "**《2025年甘肃省统计公报》**（省统计局2026-03）——GDP、工业、新能源、人口实数。")
bullet(doc, "**2026年甘肃省政府工作报告**（2026年2月）——2025追认/2026目标/河西/新能源。")
heading2(doc, "1.2 阅读方法：显性—数据—制度—长期")
para(doc, "**关键判别**：数据优先。甘肃2025年GDP+5.8%（目标5.5%超）、规上工业+9.5%。甘肃\u201c工业/新能源/中医药强、人口/固投/消费偏弱\u201d，穿透总量看工业与能源。")

# ---- 二、底盘 ----
heading1(doc, "二、先看甘肃的特殊底盘：河西走廊、新能源、中医药、兰白都市圈")
para(doc, "甘肃的地盘取决于它作为\u201c**河西走廊+新能源基地+中医药/特色农业+兰白都市圈**\u201d的特殊定位。它是向西开放门户、丝绸之路通道、生态屏障。")
bullet(doc, "**河西走廊**：酒泉/张掖/武威，风电光伏大基地、绿氢走廊；丝绸之路节点。")
bullet(doc, "**新能源**：新能源装机超1亿千瓦（风/光）、陇电入浙/入川、国家风光基地。")
bullet(doc, "**中医药/农业**：陇药/中药材全国第一，八大产业集群、高原夏菜。")
bullet(doc, "**兰白都市圈**：兰州-白银一体化、兰白定临城市群、兰白千万吨炼化。")
para(doc, "这一底板决定甘肃2025增速\u201c底色\u201d：**只要新能源/有色/制造持续，甘肃就站在\u201c工业+能源\u201d增长极；若人口/消费偏弱，需承受\u201c农业稳产业强、民生改善快\u201d的结构。**")

# ---- 三、宏观错位 ----
heading1(doc, "三、最关键的宏观错位：GDP破1.37万亿、工业/外贸强，但消费/人口/固投偏弱")
para(doc, "甘肃2025年最值得咀嚼的错位，是\u201c**工业/新能源/农业强、消费/人口/固投偏弱**\u201d。这种错位决定了对甘肃经济的观察不能只看总量增速。")
bullet(doc, "**GDP**：13697.5亿、+5.8%（全国第2）。一产1773.0亿（+5.5%、占12.9%）、二产4558.2亿（+6.7%、占33.3%）、三产7366.3亿（+5.3%、占53.8%）。")
bullet(doc, "**工业**：规上工业+9.5%（全国第3）；有色+19.5%、电力+18.5%、制造+9.3%；私营+13.4%。")
bullet(doc, "**投资**：固投-1.3%、第三产-5.1%、工业+2.9%、基建+14.5%（投资偏弱）。")
bullet(doc, "**消费**：社零4237.6亿、+2.5%；网上零售+36.0%（体量中等）。")
bullet(doc, "**外贸**：进出口711.7亿、+16.2%（出口+44.5%）、一带一路70.7%。")
bullet(doc, "**人口/财政**：收入+6.1%、财政+5.7%、人口自然增-2.29‰。")
para(doc, "**为什么读这条**：甘肃作为\u201c能源/农业大省+人口流出\u201d，结构性矛盾是\u201c产业/能源/外贸强、人口/消费/固投偏弱\u201d。稳增长靠工业/能源，民生改善靠收入。")

# ---- 四、15条细节 ----
heading1(doc, "四、容易被忽视、但信号很重的细节（15条）")
para(doc, "以下15条，都在2025年报告/公报，但被\u201cGDP 1.37万亿\u201d等掩盖。它们是判断甘肃2025之后5—10年的关键小信号。")
bullet(doc, "**1. 新能源装机12442万千瓦（风+光8000+）**：风电+27.3%、光伏+22.7%，全国新能源高地。")
bullet(doc, "**2. 规上工业+9.5%（全国第3）**：有色/电力/制造引擎。")
bullet(doc, "**3. 中药/陇药（中药材全国第一）**：中医药产业链1109亿。")
bullet(doc, "**4. 进出口+16.2%、出口+44.5%**：西部外贸高增。")
bullet(doc, "**5. 河西走廊**：丝绸之路、风光大基地、绿氢走廊。")
bullet(doc, "**6. 兰白一体**：兰白/定临城市群、兰白千万吨炼化、城市更新。")
bullet(doc, "**7. 庆阳算力/东数西算**：算力11.4万P、数据中心。")
bullet(doc, "**8. 规上工业利润+3.4%（520.7亿）**：有色/能源盈利。")
bullet(doc, "**9. 常住2443万、城镇化58.08%（+1.25pct）**：城镇化快速提升、劳动力外流。")
bullet(doc, "**10. 收入28224元、+6.1%（全国第3）**：民生改善快。")
bullet(doc, "**11. 省级化债化险/自然灾害承压。")
bullet(doc, "**12. 网上零售+36%**：数字经济渗透。")
bullet(doc, "**13. 旅游5.02亿人次、花费4036亿元**：文旅大省。")
bullet(doc, "**14. 12个中央转移大省**：财政转移依赖。")
bullet(doc, "**15. 兰白科创（黄河生态）**：都市圈+黄河治理。")
para(doc, "", size=10, space_after=2)

# ---- 五、对照表 ----
heading1(doc, "五、2025年GDP目标 vs 实际：对照表")
para(doc, "2025年报告中列出的主要预期目标，与2025年实际完成：")
tb = [
    ["指标", "2025年目标", "2025年实际（公报）", "达标判定"],
    ["GDP增速", "5.5%左右", "+5.8%（13697.5亿）", "超预期"],
    ["规上工业增加值", "—", "+9.5%", "亮眼"],
    ["社会消费品零售总额", "—", "+2.5%（4237.6亿）", "偏弱"],
    ["外贸进出口", "—", "+16.2%（711.7亿）", "超预期"],
    ["地方一般公共预算收入", "—", "+5.7%（1112.4亿）", "稳健"],
    ["居民人均可支配收入", "—", "+5.8%（28224元）", "快"],
]
table(doc, tb[0], tb[1:], widths=[3.2, 3.2, 4.4, 4.0])
para(doc, "", size=10, space_after=2)
para(doc, "**要点**：GDP/工业/出口/财政全面超，社零偏弱——甘肃\u201c生产强、消费弱\u201d。")

# ---- 六、结构归因 ----
heading1(doc, "六、2025年增长是谁撑起来的？（结构归因）")
heading2(doc, "6.1 三次产业：工业强（二产+6.7%）")
para(doc, "**二产（+6.7%）主导**、一产+5.5%、三产+5.3%。规上工业是最高景气。")
heading2(doc, "6.2 工业：有色/电力/制造")
para(doc, "规上工业+9.5%；有色+19.5%、电力+18.5%、制药+3.5%。新能源带动发电。")
heading2(doc, "6.3 消费/旅游")
para(doc, "社零+2.5%、旅游5.02亿人次（+11.3%）。")
heading2(doc, "6.4 投资/财政")
para(doc, "固投-1.3%，基建+14.5%；财政+5.7%。")
heading2(doc, "6.5 外贸/工业")
para(doc, "出口+44.5%（新能源/有色），一带一路70.7%。")
para(doc, "**一句话归因**：甘肃2025年靠\u201c**工业（有色/电力/制造）+新能源+农业/中医药**\u201d，支撑\u201cGDP结构改善、民生快\u201d。")

# ---- 七、财政 ----
heading1(doc, "七、预算与财政的\u201c含金量\u201d")
para(doc, "**一般公共预算收入1112.4亿元、+5.7%**；税收705.4亿、+3.3%，非税406.9亿、+10.3%。")
bullet(doc, "**收入改善**：+5.7%，主要靠工业/非税/资源税收；全国第5。")
bullet(doc, "**民生/化债**：化债化险，民生支出/基层三保优先。")
bullet(doc, "**转移依赖**：中央转移支付占比高。")
para(doc, "**财政含义**：甘肃\u201c收入+5.7%、财政稳\u201d，靠工业/税收支撑，民生与化债统筹。")

# ---- 八、民生 ----
heading1(doc, "八、民生底账：人口、收入与城乡")
para(doc, "**常住人口2443万人、城镇化率58.08%（+1.25pct）**；居民人均可支配收入28224元、+6.1%（全国第3，快于GDP）。")
bullet(doc, "**收入**：居民28224元、+6.1%；城镇43910元、+4.9%，农村15001元、+6.3%（城乡比2.93、缩小）。")
bullet(doc, "**就业**：城镇新增31.7万人；输转富余劳动力515.3万人（外出务工）。")
bullet(doc, "**社保**：低保/医保完善，困难家庭资助、白内障免费手术等民生实事。")
bullet(doc, "**人口**：常住2443万、城镇化+1.25pct、自然增-2.29‰（人口流出）。")
para(doc, "**民生含义**：甘肃\u201c收入增速全国前列、城乡缩小\u201d，但人口自然负增长、劳动力外流。")

# ---- 九、城乡 ----
heading1(doc, "九、城镇与农村：格局与均衡")
para(doc, "**城镇化率58.08%（+1.25pct）**，甘肃城镇化快速上升；城乡收入比2.93、逐年缩小。")
bullet(doc, "**城市**：兰州/兰白都市圈，河西走廊城市带（酒泉/张掖）、陇东。")
bullet(doc, "**农村**：农业（中药材/高原夏菜/玉米）、牧区、劳务输出。")
para(doc, "**城乡均衡**：甘肃\u201c城市强、县域弱\u201d，靠劳务/产业转移补城乡均衡。")

# ---- 十、人口流 ----
heading1(doc, "十、人口流入与流出")
para(doc, "**甘肃常住2443万、自然增-2.29‰**：人口总体流出（出生6.65‰、死亡8.94‰），劳动力输出515.3万人。")
bullet(doc, "**流出**：劳动力赴沿海/建筑/服务，依赖劳务转移。")
bullet(doc, "**流入**：兰州/兰白都市圈，产业/新能源留人。")
para(doc, "人口方向决定中长期需求；甘肃靠新能源/产业/就业减少流出。")

# ---- 十一、物价 ----
heading1(doc, "十一、物价与货币环境")
para(doc, "**2025年甘肃CPI下降0.1%**：食品/农产品负、工业PRI负（-3.2%）；需求温和。")
bullet(doc, "**物价**：CPI-0.1%低位、PPI-3.2%工业品降价。")
bullet(doc, "**货币/金融**：存贷款稳，A股/上市市值+27%。")
para(doc, "**物价含义**：甘肃\u201c通缩压力\u201d温和、农产品/能源价格弱。")

# ---- 十二、区域一体化 ----
heading1(doc, "十二、区域一体化：甘肃在\u201c兰白都市圈+河西走廊+黄河几字弯\u201d里的位置")
para(doc, "甘肃处于**兰白都市圈+河西走廊（丝绸之路）+黄河几字弯生态**交汇：既是向西开放门户、能源/农业大省，也是生态屏障。")
bullet(doc, "**兰白都市圈**：兰州-白银一体化、定临城市群、1000万吨炼化。")
bullet(doc, "**河西走廊**：兰新铁路、绿氢走廊、新能源基地。")
bullet(doc, "**黄河几字弯**：黄河上游生态、兰西城市群协同。")
bullet(doc, "**向西开放**：中欧/中亚通道、兰新铁路。")
para(doc, "**区域含义**：甘肃作为\u201c河西/兰白+向西开放门户\u201d，靠新能源/医药/走廊带辐射。")

# ---- 十三、五主线 ----
heading1(doc, "十三、未来5—10年最值得观察的五条主线")
bullet(doc, "**主线1｜新能源**：装机12443万千瓦、陇电入浙/入川。能否建成1.5亿千瓦国家基地。")
bullet(doc, "**主线2｜有色/先进制造**：有色+19.5%、装备/兰白。能否在有色化工/新能源装备卡位。")
bullet(doc, "**主线3｜中医药/特色农业**：1109亿中医药产业、中药材第一。能否做强陇药产业。")
bullet(doc, "**主线4｜算力/数据中心**：庆阳11.4万P。能否把\u201c绿电+兰白\u201d变算力。")
bullet(doc, "**主线5｜人口/民生/劳务**：2443万、务工515万。能否靠产业把人口留住。")
para(doc, "这五条，是甘肃从\u201c能源+农业大省\u201d走向\u201c新能源+中医药+算力强省\u201d的\u201c主赛道\u201d。")

# ---- 十四、结论 ----
heading1(doc, "十四、最终结论：甘肃在\u201c新能源+工业+中医药\u201d里的增长逻辑")
para(doc, "甘肃2025年，本质上是\u201c**工业（有色/电力/制造）+新能源+农业快速、固投/消费偏弱**\u201d的答卷：GDP13697.5亿、+5.8%（全国第2）、规上工业+9.5%、新能源装机12443万千瓦、进出口+16.2%、收入+6.1%、财政+5.7%。")
para(doc, "只要新能源/工业/中医药持续，甘肃就站在\u201c工业+新能源\u201d增长极；若人口/固投/消费受限，需承受\u201c生产强、需求弱\u201d的结构。")
para(doc, "最稳观察信号：**一盯新能源装机（引擎）、二盯有色/制造（工业）、三盯中医药/农业（特色）、四盯人口/民生（长期）、五盯固投/财政（质量）。**甘肃，是\u201c河西走廊+新能源\u201d新样本。")

# ---- 附录A ----
heading1(doc, "附录A：主要资料来源与核验路径")
bullet(doc, "甘肃省2025年政府工作报告（2025年1月）——目标来源。")
bullet(doc, "《2025年甘肃省国民经济和社会发展统计公报》——GDP、工业、新能源、人口实数。")
bullet(doc, "2026年甘肃省政府工作报告（2026年2月）——2025追认/河西/新能源/中医药。")
bullet(doc, "省财政厅、省统计局、兰州海关。")
heading2(doc, "核验说明")
para(doc, "本报告涉及数据以统计公报/官方口径为准；\u201c新能源/中医药/兰白\u201d等以官方口径为准。")

# ---- 附录B ----
heading1(doc, "附录B：建议建立的年度跟踪仪表盘")
para(doc, "建议把下面10个\u201c测脉搏\u201d指标做成年度跟踪表：")
dash = [
    ["#", "指标", "2025基值", "为什么重要"],
    ["1", "GDP增速", "+5.8%（13697.5亿）", "总量与方向"],
    ["2", "规模以上工业增速", "+9.5%", "工业底盘"],
    ["3", "新能源装机", "12443万千瓦(风/光)", "绿色能源"],
    ["4", "中医药/农业", "陇药/中药材第一", "特色产业"],
    ["5", "进出口/出口", "+16.2% / +44.5%", "开放/外贸"],
    ["6", "固定资产投资", "-1.3%", "投资结构"],
    ["7", "社零增速", "+2.5%（4237.6亿）", "内需消费"],
    ["8", "常住人口/城镇化", "2443万 / 58.08%", "人口与城市"],
    ["9", "地方财政收入", "+5.7%（1112.4亿）", "财政质量"],
    ["10", "居民收入", "+6.1%（28224元）", "民生"],
]
table(doc, dash[0], dash[1:], widths=[0.8, 4.6, 4.4, 4.0])
para(doc, "把这10个指标连起来看，新能源/工业/中医药（2/3/4）、人口/收入（8/10），都说明甘肃在真正换挡。")

# ============ 保存 ============
out = "/Users/x/Desktop/content-prod-lab/reports/甘肃省_2025年政府工作报告_深度研究_2026-08-13.docx"
doc.save(out)
print("SAVED:", out)
print("PARAGRAPHS:", len(doc.paragraphs))
print("TABLES:", len(doc.tables))
