# -*- coding: utf-8 -*-
"""Build 重庆市2025年政府工作报告 深度研究 DOCX, 参照北京/上海/杭州/天津系列版式。"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DARK = RGBColor(0x1F, 0x2A, 0x44)
RED = RGBColor(0xB0, 0x1E, 0x2E)
GRAY = RGBColor(0x59, 0x60, 0x69)
BLUE = RGBColor(0x1F, 0x4E, 0x79)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
HEADER_FILL = "1F2A44"
K = "微软雅黑"


def set_run(run, size=11, bold=False, color=DARK, font=K):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    run.font.name = font
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), font)
    rFonts.set(qn("w:hAnsi"), font)
    rFonts.set(qn("w:eastAsia"), font)


def para(doc, text, size=11, bold=False, color=DARK, align=None,
         space_after=6, space_before=0, font=K):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        set_run(r, size=size, bold=bold or (i % 2 == 1), color=color, font=font)
    return p


def heading1(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(16)
    pf.space_after = Pt(8)
    r = p.add_run(text)
    set_run(r, size=16, bold=True, color=RED)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "B01E2E")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def heading2(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(10)
    pf.space_after = Pt(4)
    r = p.add_run(text)
    set_run(r, size=13, bold=True, color=BLUE)
    return p


def table(doc, headers, rows, widths=None, font_size=9.5):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_run(r, size=font_size, bold=True, color=WHITE)
        tcPr = hdr[i]._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), HEADER_FILL)
        tcPr.append(shd)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(val))
            set_run(r, size=font_size, color=DARK)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    return t


def quote_box(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(8)
    pf.left_indent = Pt(12)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), "B01E2E")
    pBdr.append(left)
    pPr.append(pBdr)
    r = p.add_run(text)
    set_run(r, size=11, color=DARK, font="楷体")
    return p


def bullet(doc, text, size=10.5):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.7)
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        set_run(r, size=size, bold=(i % 2 == 1), color=DARK)
    return p


# ================================================================= 文档主体
doc = Document()
sec = doc.sections[0]
sec.page_width = Cm(21)
sec.page_height = Cm(29.7)
sec.top_margin = Cm(2.4)
sec.bottom_margin = Cm(2.2)
sec.left_margin = Cm(2.5)
sec.right_margin = Cm(2.5)

st = doc.styles["Normal"]
st.font.name = K
st.font.size = Pt(11)
st.element.rPr.rFonts.set(qn("w:eastAsia"), K)


# ---- 封面 ----
para(doc, "合肥市2025年政府工作报告", size=24, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=60, space_after=4)
para(doc, "深度研究与“容易被忽视的细节”分析", size=16, bold=True, color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
para(doc, "从“芯屏汽合、政府基金、量子与人口虹吸”重新理解合肥", size=12, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)
para(doc, "━━━━━━━━━━━━━━━━", color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
para(doc, "研究对象：2025年合肥市政府工作报告及同期财政、国民经济和社会发展统计资料、2026年官方复盘", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
para(doc, "整理时间：2026年8月", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
para(doc, "分析性质：分析性研究 ｜ 非官方解读", size=11, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
doc.add_page_break()

# ---- 目录 ----
catalog = [
    "执行摘要",
    "一、研究口径与阅读方法",
    "二、先看合肥的特殊底盘：科创之城、芯屏汽合与人口虹吸",
    "三、最关键的宏观错位：GDP破1.42万亿、工业强，但投资与消费偏弱",
    "四、容易被忽视、但信号很重的细节（15条）",
    "五、把所有细节连起来：合肥正在经历的“六个换挡”",
    "六、增长暗线：芯屏汽合、政府基金与量子/战新",
    "七、财政暗线：收入稳、税收质量高、民生支出大",
    "八、产业暗线：从“家电/白电”到“芯屏汽合+战新产业”",
    "九、区域格局：长三角一体化与合肥都市圈",
    "十、人口与人才：人口净流入、高学历集聚",
    "十一、民营经济与科创：基金招商与产业链",
    "十二、事后验证：用2025年实际执行结果检验报告判断",
    "十三、未来5—10年最值得观察的五条主线",
    "十四、最终结论：合肥在“科技+产业+合肥模式”里的增长逻辑",
    "附录A：主要资料来源与核验路径",
    "附录B：建议建立的年度跟踪仪表盘",
]
para(doc, "目　录", size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=10, space_after=10)
for item in catalog:
    para(doc, item, size=12, space_after=4)
doc.add_page_break()

# ---- 执行摘要 ----
heading1(doc, "执行摘要")
para(doc, "**核心判断**　如果只看新闻标题，2025年合肥最显眼的是“GDP突破1.42万亿、增长6.1%”、“规上工业+17.6%、创近4年新高”和“新能源汽车产量137万辆、占全国8.3%”。但这份研究真正值得深读的，是这座“科创之城、芯屏汽合”的城市如何在固定资产投资（-16.5%）与社零（+3.2%）偏弱的背景下，靠“战新产业（占规上60.4%）+新型显示/集成电路/新能源车+出口+政府基金”撑起高增长。")
para(doc, "把2025年初设定的目标、2025年《国民经济和社会发展统计公报》、以及2026年报告对2025年的复盘放在一起看，合肥呈现清晰暗线：**从“家电/传统制造”的旧底盘，向“芯屏汽合+战新产业+新质生产力+出口”转型**。旧引擎（房地产、一般基建、传统家电制造）在调整；新引擎（新型显示、集成电路、新能源汽车、量子、战新、出口）成为增长主线。")
para(doc, "因此，本报告不按政府工作报告逐段复述，而采用“显性表述—同期数据—制度含义—长期影响”的方式，专门提取容易被忽视、但对判断合肥未来5—10年发展模式更有价值的信号。")
para(doc, "最容易记住的一句话：**合肥是“科创之城+政府基金招商”的样本，靠“芯屏汽合”和战新产业实现制造强市的跃迁——它用“政府基金+科技转化”把一个中部省会做成全国产业新贵。**观察合肥，与其看“GDP 1.4万亿”，不如看“战新占规上比重、集成电路/显示/新能源车产量、出口增速与人口净流入”这几张名片。")
heading2(doc, "一页速览：2025年合肥经济的“表与里”")
table(doc,
    ["维度", "表面现象", "更深层信号"],
    [
        ["增长", "GDP 14210亿、+6.1%", "第二产业+8.7%拉动、科创强"],
        ["产业", "规上工业+17.6%、战新+16.6%", "战新占规上60.4%、电子+60.6%"],
        ["外贸", "进出口4551.8亿、+20.2%（出口+23.6%）", "新三样出口+34.8%"],
        ["投资", "固定资产投资-16.5%", "高端装备+17.4%、新能源投资+10.1%"],
        ["财政", "一般公共预算收入+2.3%（税收+3.4%）", "财政稳、税收质量高"],
        ["消费", "社零5320.9亿、+3.2%", "线上+16.5%、以旧换新强"],
        ["人口", "常住约1000万+、人口净流入", "人口虹吸、人才集聚"],
    ],
    widths=[2.2, 5.2, 8.6])
para(doc, "", size=4, space_after=2)

# ---- 一、研究口径与阅读方法 ----
heading1(doc, "一、研究口径与阅读方法")
heading2(doc, "1.1 本报告的三份核心底稿")
bullet(doc, "**2025年《政府工作报告》**：2025年1月在市十七届人大三次会议上作，给出全年预期目标（GDP增长5.5%以上、规上工业增长10%左右等）。进出口/社零等未全部设具体速率。")
bullet(doc, "**2025年《国民经济和社会发展统计公报》**：合肥市统计局发布，给出全年实际执行数与增速，是“事后验证”的锚点。")
bullet(doc, "**2026年合肥市政府工作报告及官方复盘**：对2025执行结果的追认，交叉验证“目标—执行—再评价”三段式。")
heading2(doc, "1.2 阅读方法：显性—数据—制度—长期")
para(doc, "每一章按“**显性表述 → 同期数据 → 制度含义 → 长期影响**”四层展开。其中“制度含义”回答“政府为什么这么排优先序”，“长期影响”回答“这对未来5—10年意味着什么”。")
para(doc, "关键判别：**当报告使用的“增长词”和统计公报里的实际数不一致时，数据优先。**例如2025年GDP目标5.5%以上、实际6.1%大幅超；规上工业目标10%、实际17.6%超预期；而固定资产投资-16.5%（回调）。差异反映：合肥“战新/出口强、投资/消费偏弱”。")

# ---- 二、特殊底盘 ----
heading1(doc, "二、先看合肥的特殊底盘：科创之城、芯屏汽合与人口虹吸")
para(doc, "在所有省会城市里，合肥的“底盘”独特：**科创之城+芯屏汽合+人口净流入**三合一。作为长三角副中心城市，合肥以“科大硅谷”、中国科大、量子信息等科创资源著称。")
para(doc, "这决定合肥的多重身份并存：**芯屏汽合（芯片、显示、汽车/BYD、综合）**、**科创与转化（量子、AI、政府基金）**、**战新产业高地（占规上60.4%）**、**长三角一体化**、**人口虹吸高地（产业+就业）**。")
heading2(doc, "2.1 芯屏汽合")
para(doc, "合肥以“芯屏汽合”著称（芯片、显示屏（京东方）、汽车（蔚来/江淮/比亚迪）、光伏储能）。2025年战新产业占规上60.4%、电子+60.6%、存储器等多重放量。")
heading2(doc, "2.2 政府基金与科创")
para(doc, "合肥以“政府引导基金+产业链招商”闻名，孵化出汽车、显示、半导体等产业集群。科创（中科大、量子、科学岛）→产业转化，是合肥独特的发展模式。")
heading2(doc, "2.3 人口虹吸")
para(doc, "合肥因产业与科创吸引力，人口净流入、高学历人才集聚，成为中部/长三角的人口与人才枢纽。")

# ---- 三、宏观错位 ----
heading1(doc, "三、最关键的宏观错位：GDP破1.42万亿、工业强，但投资与消费偏弱")
para(doc, "把2025年合肥的宏观面放进一张表，会出现令人意外的“错位”：表观增长来自工业/出口/战新，而投资与消费偏弱。这个错位，正是读懂合肥的关键。")
heading2(doc, "3.1 “强的一面”")
table(doc,
    ["指标", "2025实绩", "信号"],
    [
        ["GDP", "14210亿、+6.1%", "第二产业+8.7%、科创强"],
        ["规上工业", "+17.6%", "创近4年新高"],
        ["战新产业", "+16.6%、占规上60.4%", "战新主导"],
        ["电子设备制造", "+60.6%", "半导体/显示放量"],
        ["出口", "+23.6%", "新三样+34.8%"],
        ["新能源汽车", "137万辆", "占全国8.3%"],
    ],
    widths=[3.2, 5.2, 6.0])
heading2(doc, "3.2 “弱的一面”")
table(doc,
    ["指标", "2025实值", "信号"],
    [
        ["固定资产投资", "-16.5%", "投资大幅回落"],
        ["社零", "+3.2%", "消费偏弱"],
        ["一产", "占GDP低", "工业化城市"],
        ["第三产业", "+4.4%", "服务业偏慢"],
    ],
    widths=[3.2, 5.2, 6.0])
para(doc, "**错位结论**　合肥的增长“很强、但也很极端”。强的部分（工业/战新/出口）与弱的部分（投资/消费）并存。**真正的焦点是“制造强、内需弱”**：工业+17.6%、占规上60%战新，但固投-16.5%、社零仅+3.2%。2026年合肥在“工业强+扩内需+稳投资”间平衡。")

# ---- 四、被忽视细节 ----
heading1(doc, "四、容易被忽视、但信号很重的细节（15条）")
details = [
    ("1", "规上工业+17.6%、创近4年新高", "制造业空前景气。"),
    ("2", "战新产业+16.6%、占规上60.4%", "战新主导、量质齐升。"),
    ("3", "计算机/通信/电子+60.6%", "电子制造爆发。"),
    ("4", "新能源汽车产量137万辆、占全国8.3%", "汽车出口/内销强。"),
    ("5", "微型计算机2772万台、占全国8.4%", "显示/计算设备。"),
    ("6", "半导体分立器件+42.7%、工业机器人+1.1倍", "半导体/机器人。"),
    ("7", "集成电路设计营业收入+36.5%", "芯片补链、设计放量。"),
    ("8", "进出口4551.8亿、+20.2%、出口+23.6%", "出口超强、新三样+34.8%。"),
    ("9", "高端装备制造业投资+17.4%", "高端制造投资强。"),
    ("10", "新能源产业投资+10.1%", "绿色投资。"),
    ("11", "社零+3.2%、线上+16.5%", "内需弱、线上强。"),
    ("12", "以旧换新强、通讯/家电高增", "促消费政策见效。"),
    ("13", "常住人口约1000万级、人口净流入", "人口虹吸与人才集聚。"),
    ("14", "居民人均可支配收入58930元、+5.5%、城乡比1.86", "收入稳、城乡差距小。"),
    ("15", "一般公共预算收入+2.3%（税收+3.4%）", "财政稳、税收质量高。"),
]
for d in details:
    para(doc, "**%s**" % d[0], size=11)
    para(doc, d[1], size=10.5, color=GRAY)

# ---- 五、六个换挡 ----
heading1(doc, "五、把所有细节连起来：合肥正在经历的“六个换挡”")
gear = [
    ("1．动能换挡：从“家电/传统制造”到“战新+新质生产力”", "战新占60.4%、电子+60.6%成为新引擎。"),
    ("2．产业换挡：从“白电/服装”到“芯屏汽合+量子/半导体”", "新型显示、集成电路、新能源车放量。"),
    ("3．投资换挡：从“基建/地产”到“高端制造/新能源投资”", "高端装备+17.4%、新能源投资+10.1%、地产弱。"),
    ("4．开放换挡：从“内销”到“出口+新三样”", "出口+23.6%、新三样+34.8%。"),
    ("5．人口换挡：从“外流”到“净流入+人才集聚”", "人口约1000万、虹吸+科创。"),
    ("6．动能换挡：从“政府招商”到“政府基金+科技转化”", "科大硅谷、量子、基金招商形成生态。"),
]
for g in gear:
    para(doc, "**%s**　%s" % g, space_after=6)

# ---- 六、增长暗线 ----
heading1(doc, "六、增长暗线：芯屏汽合、政府基金与量子/战新")
heading2(doc, "6.1 芯屏汽合+战新")
para(doc, "合肥以“芯屏汽合”著称，战新产业占规上60.4%、电子+60.6%、集成电路设计营收+36.5%、新能源汽车137万辆（占全国8.3%）。战新+新质是合肥最硬的核心。")
heading2(doc, "6.2 政府基金+科创")
para(doc, "合肥“政府引导基金+产业链招商+科创转化”模式（科大硅谷、量子、科学岛）孵化出多个数千亿级产业集群。政府基金+科技，是合肥独特的“合肥模式”。")
para(doc, "**这条暗线意味着**：合肥的增长叙事已从“家电/传统制造”转向“芯屏汽合+战新+科创转化”。看合肥，盯住“战新占规上比重、电子/新能车产量、出口、政府基金/科创转化”。")

# ---- 七、财政暗线 ----
heading1(doc, "七、财政暗线：收入稳、税收质量高、民生支出大")
para(doc, "2025年合肥一般公共预算收入977.35亿、+2.3%，税收698.19亿、+3.4%。税收质量高，来自战新/制造/服务业税源。支出民生占比超八成、增长7.1%。")
para(doc, "**制度含义**　合肥财政“收入稳、税收质量高”，支撑科创转化与民生。关键是保持“产业+科创”真实税源，平衡投资与化债。")

# ---- 八、产业暗线 ----
heading1(doc, "八、产业暗线：从“家电/白电”到“芯屏汽合+战新产业”")
heading2(doc, "8.1 合肥产业的“表”")
table(doc,
    ["指标", "2025增速/信号", "解读"],
    [
        ["规上工业", "+17.6%", "制造业爆发"],
        ["战新产业", "占规上60.4%", "战新主导"],
        ["电子设备制造", "+60.6%", "电子爆发"],
        ["新能源汽车", "137万辆、占8.3%", "汽车/新能源"],
        ["集成电路设计", "+36.5%", "芯片补链"],
        ["高端装备投资", "+17.4%", "高端制造"],
        ["出口", "+23.6%", "新三样强"],
    ],
    widths=[4.6, 3.6, 5.0])
heading2(doc, "8.2 从“家电”到“芯屏汽合/战新”")
para(doc, "合肥过去以家电（白电）见长，2025年显示“芯片/显示/汽车/光伏/量子”正成为新引擎。战新产业占60%以上，是“政府基金+科技转化”造就的产业跃迁。")

# ---- 九、区域 ----
heading1(doc, "九、区域格局：长三角一体化与合肥都市圈")
para(doc, "合肥是长三角副中心（群）城市之一，通过“长三角一体化”承接产业、创新与人才外溢。合肥都市圈（合滁/合芜等）带动皖中/皖北。")
para(doc, "在长三角协同下，合肥借力上海/江苏外溢，同时向中西部辐射产业。合肥的“长三角承接+中部辐射”，是区域位势。")

# ---- 十、人口与人才 ----
heading1(doc, "十、人口与人才：人口净流入、高学历集聚")
heading2(doc, "10.1 人口总量")
para(doc, "合肥常住人口约1000万级（七普约937万、逐年增长），因产业与科创吸引力而净流入、高学历人才集聚。合肥是中部/长三角人口虹吸点。")
heading2(doc, "10.2 人才与就业")
para(doc, "中科大、合肥工业大学、科学岛等科教资源+产业集群，吸引与培养大量人才。城镇新增就业24万，人才集聚推动“科技+制造”。")

# ---- 十一、民营经济与科创 ----
heading1(doc, "十一、民营经济与科创：基金招商与产业链")
heading2(doc, "11.1 民企与基金")
para(doc, "合肥以“政府引导基金+头部企业（蔚来/比亚迪/京东方）”孵化产业链，民企与科技企业（专精特新）活跃。基金招商是合肥独特模式。")
heading2(doc, "11.2 科创生态")
para(doc, "“科大硅谷”“科学岛”“量子信息”等是合肥科创名片。政府基金+科技转化+产业落地，构成合肥发展闭环。")

# ---- 十二、事后验证 ----
heading1(doc, "十二、事后验证：用2025年实际执行结果检验报告判断")
heading2(doc, "12.1 年初目标与年末执行对比")
table(doc,
    ["指标", "2025年初目标", "2025实际", "偏差"],
    [
        ["GDP增速", "5.5%以上", "+6.1%", "超预期"],
        ["规上工业", "10%左右", "+17.6%", "大幅超"],
        ["固定资产投资", "未设", "-16.5%", "回落"],
        ["社零", "未设", "+3.2%", "偏弱"],
        ["进出口", "未设", "+20.2%", "超预期"],
        ["居民收入", "与增长同步", "+5.5%", "同步"],
    ],
    widths=[3.0, 3.2, 3.0, 4.2])
heading2(doc, "12.2 判断验证")
para(doc, "三条核心判断得到支持：**（1）工业/战新/出口强（规上+17.6%、战新占60.4%、出口+23.6%），验证；（2）投资/消费偏弱（固投-16.5%、社零+3.2%），验证；（3）财政稳/民生（收入+2.3%、税收+3.4%），验证。**")
para(doc, "核心观察：合肥靠“战新+出口+政府基金”实现高增长、制造强市，但投资与内需偏弱。2026年“稳投资+扩内需+继续战新/科创”是重点。")

# ---- 十三、五条主线 ----
heading1(doc, "十三、未来5—10年最值得观察的五条主线")
lines5 = [
    ("① 战新产业/芯屏汽合能否持续升级", "占规上60%、电子，能否向半导体高端/自主升级。"),
    ("② 政府基金与科创转化", "科大硅谷、量子、基金链，能否孵化更多独角兽/千亿龙头。"),
    ("③ 出口/新能源汽车出海", "新三样+34.8%、汽车占8.3%，能否全球做强。"),
    ("④ 投资与消费再平衡", "固投-16.5%、社零弱，能否用产业投资/内需找回。"),
    ("⑤ 人口/人才集聚", "净流入、高学历，能否在城市更新/公共服务留住。"),
]
for l in lines5:
    para(doc, "**%s**　%s" % l, space_after=6)

# ---- 十四、结论 ----
heading1(doc, "十四、最终结论：合肥在“科技+产业+合肥模式”里的增长逻辑")
para(doc, "合肥的2025年，本质上是“**战新/科创/出口为核心，而房地产/内需承压**”的答卷：GDP1.42万亿、规上工业+17.6%、战新占60.4%，实现制造强市跃迁，代价是固投-16.5%、社零偏弱。")
para(doc, "只要“芯屏汽合、政府基金、科创转化、出口”能接住，合肥就保持“科创+制造”新贵；如果投资、内需、战新利润率偏弱，合肥需承受“重制造轻内需”的挑战。")
para(doc, "最稳妥的观察信号：**一盯战新/新/出口（动能）、二盯半导体/显示/新能源（产业）、三盯政府基金/科创（模式）、四盯投资/内需（平衡）、五盯人口/人才（底座）。**合肥，是中国“政府+科创=制造强市”的典型样本。")

# ---- 附录A ----
heading1(doc, "附录A：主要资料来源与核验路径")
bullet(doc, "合肥市2025年《政府工作报告》（2025年1月）——目标来源。")
bullet(doc, "《2025年合肥市国民经济和社会发展统计公报》（合肥市统计局）——GDP、工业、外贸、人口实值。")
bullet(doc, "合肥市统计局/科创专题、战新企业专报——战新与科创。")
bullet(doc, "2026年合肥市政府工作报告——2025执行复盘。")
bullet(doc, "安徽/合肥海关、市发改委、财政局——外贸/财政。")
heading2(doc, "核验说明")
para(doc, "本报告所有增速、总量、占比均以统计公报/官方发布为基准。涉“常住人口精确值”“科创细分”等以官方口径为准。")

# ---- 附录B ----
heading1(doc, "附录B：建议建立的年度跟踪仪表盘")
para(doc, "建议把下面10个“测脉搏”指标做成年度跟踪表：")
dash = [
    ["#", "指标", "2025基值", "为什么重要"],
    ["1", "GDP增速", "+6.1%", "总量与方向"],
    ["2", "规上工业增速", "+17.6%", "制造第二引擎"],
    ["3", "战新产业占比", "60.4%", "新质生产力"],
    ["4", "集成电路设计营收", "+36.5%", "芯片产业"],
    ["5", "出口增速", "+23.6%", "外向/新三样"],
    ["6", "固定资产投资增速", "-16.5%", "投资动能"],
    ["7", "社零增速", "+3.2%", "内需消费"],
    ["8", "一般公共预算收入/税收", "+2.3%/+3.4%", "财政质量"],
    ["9", "常住人口/人才", "约1000万+", "人口虹吸"],
    ["10", "居民人均可支配收入增速", "+5.5%", "民生获得感"],
]
table(doc, dash[0], dash[1:], widths=[0.8, 4.5, 3.2, 5.0])
para(doc, "把这10个指标连起来看，任何一个“新引擎（2/3/5）向上、旧引擎（6/7）修复”，都说明合肥在真正换挡；反之则是旧路径（基建/内需强弱）的反复。")

# ========================================= 保存
out = "/Users/x/Desktop/content-prod-lab/reports/合肥市_2025年政府工作报告_深度研究_2026-08-13.docx"
doc.save(out)
print("SAVED:", out)
print("PARAGRAPHS:", len(doc.paragraphs))
print("TABLES:", len(doc.tables))
