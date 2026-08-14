# -*- coding: utf-8 -*-
"""Build 保定市2025年政府工作报告 深度研究 DOCX, 参照地级市系列版式。"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DARK = RGBColor(0x1F, 0x2A, 0x44)
RED = RGBColor(0xB0, 0x1E, 0x2E)
GRAY = RGBColor(0x59, 0x60, 0x69)
BLUE = RGBColor(0x1F, 0x4E, 0x79)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
HEADER_FILL = "1F2A44"
K = "微软雅黑"


def set_run(run, size=11, bold=False, color=DARK, font=K):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    run.font.name = font
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), font)
    rFonts.set(qn("w:hAnsi"), font)
    rFonts.set(qn("w:eastAsia"), font)


def para(doc, text, size=11, bold=False, color=DARK, align=None,
         space_after=6, space_before=0, font=K):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        set_run(r, size=size, bold=bold or (i % 2 == 1), color=color, font=font)
    return p


def heading1(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(16)
    pf.space_after = Pt(8)
    r = p.add_run(text)
    set_run(r, size=16, bold=True, color=RED)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "B01E2E")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def heading2(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(10)
    pf.space_after = Pt(4)
    r = p.add_run(text)
    set_run(r, size=13, bold=True, color=BLUE)
    return p


def table(doc, headers, rows, widths=None, font_size=9.5):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_run(r, size=font_size, bold=True, color=WHITE)
        tcPr = hdr[i]._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), HEADER_FILL)
        tcPr.append(shd)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(val))
            set_run(r, size=font_size, color=DARK)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    return t


def quote_box(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(8)
    pf.left_indent = Pt(12)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), "B01E2E")
    pBdr.append(left)
    pPr.append(pBdr)
    r = p.add_run(text)
    set_run(r, size=11, color=DARK, font="楷体")
    return p


def bullet(doc, text, size=10.5):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.7)
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        set_run(r, size=size, bold=(i % 2 == 1), color=DARK)
    return p


# ================================================================= 文档主体
doc = Document()
sec = doc.sections[0]
sec.page_width = Cm(21)
sec.page_height = Cm(29.7)
sec.top_margin = Cm(2.4)
sec.bottom_margin = Cm(2.2)
sec.left_margin = Cm(2.5)
sec.right_margin = Cm(2.5)

st = doc.styles["Normal"]
st.font.name = K
st.font.size = Pt(11)
st.element.rPr.rFonts.set(qn("w:eastAsia"), K)


# ---- 封面 ----
para(doc, "保定市2025年政府工作报告", size=24, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=60, space_after=4)
para(doc, "深度研究与\u201c容易被忽视的细节\u201d分析", size=16, bold=True, color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
para(doc, "从\u201c京津冀协同、装备制造、汽车零部件、白洋淀与雄安协同\u201d重新理解保定", size=12, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)
para(doc, "\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501", color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
para(doc, "研究对象：2025年保定市政府工作报告及同期财政、国民经济和社会发展统计资料、2026年官方复盘", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
para(doc, "标定日期：2026-08-13", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
doc.add_page_break()

# ---- 目录 ----
catalog = [
    "执行摘要",
    "一、研究口径与阅读方法",
    "二、先看保定的特殊底盘：京津冀协同、装备制造、汽车（长城）与白洋淀生态",
    "三、最关键的宏观错位：GDP破5000亿达标、制造强，但固投大降、人口自然负增",
    "四、容易被忽视、但信号很重的细节（15条）",
    "五、2025年GDP目标 vs 实际：对照表",
    "六、2025年增长是谁撑起来的？（结构归因）",
    "七、预算与财政的\u201c含金量\u201d",
    "八、民生底账：人口、收入与城乡",
    "九、城镇与农村：格局与均衡",
    "十、人口流入与流出",
    "十一、物价与货币环境",
    "十二、区域一体化：保定在\u201c京津冀协同+雄安新区+保定都市圈\u201d里的位置",
    "十三、未来5—10年最值得观察的五条主线",
    "十四、最终结论：保定在\u201c装备+汽车+京津冀协同\u201d里的增长逻辑",
    "附录A：主要资料来源与核验路径",
    "附录B：建议建立的年度跟踪仪表盘",
]
para(doc, "目　录", size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=10, space_after=10)
for item in catalog:
    para(doc, item, size=12, space_after=4)
doc.add_page_break()

# ---- 执行摘要 ----
heading1(doc, "执行摘要")
para(doc, "**核心判断**　2025年保定最显著的是\u201cGDP 5002.5亿元（含定州）、增长6.0%（达标6%目标）、三产占55.4%\u201d、\u201c规上工业+9.4%（装备+13.6%、汽车+13.8%）\u201d、\u201c进出口848.1亿/+6.8%\u201d、\u201c常住人口900.07万/城镇化62.27%\u201d。这说明保定在\u201c京津冀协同+装备制造\u201d下高速增长，但**固投-25.2%、人口自然-3.92\u2030**是短板。")
para(doc, "把2025年目标（GDP+6%、规上+7%、固投+7.5%、社零+6%）、2025年统计（GDP+6%达标、规上+9.4%超、固投-25.2%大幅未达）一起看，保定是\u201c京津冀+制造强市\u201d路径：**装备制造（占规上43%）、汽车（长城）、光伏、白洋淀+雄安协同**是引擎。GDP含定州破5000亿。")
para(doc, "最容易记住的一句话：**保定是\u201c京津冀协同+装备制造/汽车之城\u201d的河北地级大市，靠\u201c装备、汽车（长城）、光伏新能源、白洋淀生态、雄安协同\u201d增长。**观察保定，与其只看\u201cGDP 5002亿\u201d，不如看\u201c规上工业+9.4%、装备制造+13.6%占43%、汽车+13.8%、新能源汽车17万辆、出口+7.1%\u201d。")# ---- 一、研究口径与阅读方法 ----
heading1(doc, "一、研究口径与阅读方法")
para(doc, "本报告以\u201c保定市政府工作报告（2025年，闫继红作）\u201d为起点，把\u201c2025年GDP目标（6%左右）\u201d与\u201c2025年国民经济和社会发展统计公报实际值（含定州5002.5亿/+6.0%）\u201d并置对照，并用官方复盘作验证。")
para(doc, "注：保定统计含\u201c定州\u201d与\u201c不含定州\u201d两套口径，本报告主口径为含定州（与考核一致），文中标注。一般公共预算收入为地方级。")
para(doc, "统一采用\u201c同比增长%\u201d（GDP为不变价），居民收入按原口径转录。")

# ---- 二、特殊底盘 ----
heading1(doc, "二、先看保定的特殊底盘：京津冀协同、装备制造、汽车（长城）与白洋淀生态")
para(doc, "**区位与身份**：保定是河北省地级市、京津冀协同核心，毗邻雄安新区（2017年设），\u201c雄保协同\u201d、白洋淀所在；河北人口/制造大市。")
para(doc, "**产业底盘**：一是装备制造（占规上43%、电力/新能源高端装备）；二是汽车（长城汽车，2025年新能源汽车17万辆/海外销量45万辆）；三是光伏/新能源（太阳能电池+20.3%）；四是京津冀协同产业（疏解/协作）。")
para(doc, "**人口底盘**：2025年末常住900.07万（不含定州）/城镇化62.27%；河北人口大市。")
para(doc, "**市场与出口**：社零含定州1851.2亿/+4.7%；进出口848.1亿/+6.8%（出口807亿），民营占83.9%。")

# =========== 三、核心宏观错位 ============
heading1(doc, "三、最关键的宏观错位：GDP破5000亿、制造强，但固投大降、人口自然负增")
para(doc, "**第一组错位**：2025年GDP目标6%，实际含定州5002.5亿/+6.0%**精准达标**、破5000亿；但固投-25.2%（含定州口径，目标+7.5%大幅未达）。")
para(doc, "**第二组错位**：规上工业+9.4%（超7%目标）、装备+13.6%占43%、汽车（长城）强；但**固投-25.2%、三产投资-29.4%**。\u201c工业强、投资弱\u201d。")
para(doc, "**第三组错位**：常住900万/城镇化62.3%、但**自然-3.92\u2030**（出生4.55/死亡8.08万）人口负自然增长。")
para(doc, "一句话：**保定是\u201c装备/汽车制造强、固投弱、人口自然负增\u201d的京津冀协同大市**——靠制造+雄安协同+汽车支撑。")

# =========== 四、15条细节 ============
heading1(doc, "四、容易被忽视、但信号很重的细节（15条）")
bullet(doc, "1. **GDP破5000亿（含定州）、+6.0%**：全省第二、达标。")
bullet(doc, "2. **规上工业+9.4%**：制造业+10.9%。")
bullet(doc, "3. **装备制造+13.6%（占规上43%）**：电力新能源/高端装备主线。")
bullet(doc, "4. **汽车制造+13.8%、新能源汽车17万辆**：长城强（海外45万辆）。")
bullet(doc, "5. **光伏电池+20.3%、电子计算机+22.9%**：新能源/电子。")
bullet(doc, "6. **固投-25.2%**：但设备工器具+32.3%、高技术投资+9.7%。")
bullet(doc, "7. **房地产-14.4%、销售面积+0.7%**：地产调整收窄。")
bullet(doc, "8. **进出口+6.8%、出口+7.1%（机电占74.3%）**：出口强。")
bullet(doc, "9. **民营进出口；+82.2%占84%**：民营贸易主导。")
bullet(doc, "10. **常住900.07万/城镇化62.27%**：人口大市、自然-3.92‰。")
bullet(doc, "11. **收入：城镇44900元/+4.1%、农村24283元/+5.5%**，城乡1.85。")
bullet(doc, "12. **CPI+0.3%**：温和。")
bullet(doc, "13. **财政362.7亿/+1.8%（税收192.1/+3.2%）**：达标。")
bullet(doc, "14. **雄安协同+京津冀承接**：地缘。")
bullet(doc, "15. **规上利润204.3亿/+48.6%**：盈利大增。")# ---- 五、目标 vs 实际对照表 ----
heading1(doc, "五、2025年GDP目标 vs 实际：对照表")
table(doc,
    ["指标", "2025年目标", "2025年实际", "是否达标"],
    [
        ["地区生产总值增速", "6%左右", "含定州5002.5亿/+6.0%", "精准达标"],
        ["规上工业增加值", "+7%左右", "+9.4%", "超目标"],
        ["固定资产投资", "+7.5%左右", "-25.2%（转负）", "严重未达标"],
        ["社会消费品零售总额", "+6%左右", "含定州1851.2亿/+4.7%", "略低目标"],
        ["一般公共预算收入", "+3%左右", "含定州362.7亿/+1.8%", "略低目标"],
        ["居民人均可支配收入", "+6%左右", "全体+4.9%", "略低目标"],
        ["进出口", "（稳增）", "含定州848.1亿/+6.8%", "达标"],
    ],
    widths=[3.4, 2.5, 5.0, 3.1])
para(doc, "**简评**：**GDP、规上工业、进出口达标**；投资-25.2%大幅未达、社零/财政/收入略低。制造强是保定最大亮点，投资弱是最大约束。")

# ---- 六、结构归因 ----
heading1(doc, "六、2025年增长是谁撑起来的？（结构归因）")
para(doc, "**三大产业（含定州）**：一产584.5亿/+3.7%、二产1646.2亿/+5.6%、三产2771.8亿/+6.8%；三产占55.4%。增长以三产+二产双驱。")
para(doc, "**工业**：规上工业+9.4%、制造业+10.9%；**装备制造+13.6%（占43%）、汽车+13.9%、金属制品+32.6%、电气机械+11.9%、医药+11.4%、光伏+20.3%**；规上利润204.3亿/+48.6%。")
para(doc, "**服务业/商贸**：批发零售+物流、餐饮+37%（报复性）、网络零售+56.5%。")
para(doc, "**增长归因**：保定GDP主要靠**装备制造+汽车+出口、三产（批发/餐饮）**；固投、地产拖累。")

# ---- 七、财政 ----
heading1(doc, "七、预算与财政的\u201c含金量\u201d")
para(doc, "2025年一般公共预算收入（含定州）362.7亿/+1.8%（税收192.1亿/+3.2%）；支出（含定州）1086.5亿/-14.7%。")
para(doc, "**结构性**：收入\u201c达标、税收+3.2%\u201d；但支出-14.7%（降幅大），反映债务化解/缩支。")
para(doc, "**含金量**：保定财政\u201c税收稳增、支出收缩过大\u201d，依赖转移支付；工业盈利+48.6%缓解税基。")

# ---------------- 八、民生 ----------------
heading1(doc, "八、民生底账：人口、收入与城乡")
para(doc, "常住（不含定州）900.07万/城镇化62.27%（较上年+1.17pct）；自然-3.92\u2030。收入：**城镇44900元/+4.1%、农村24283元/+5.5%**，城乡1.85。")
para(doc, "就业：城镇新增9.3万；民营GDP占比67.2%。")
para(doc, "**民生结论**：收入农村快于城镇、城乡差距1.85中等偏小；人口自然负增、城镇化提升；就业稳、社保覆盖面广。")

# ---------------- 九 ----------------
heading1(doc, "九、城镇与农村：格局与均衡")
para(doc, "保定城镇化率约62%（河北中低水平），中心城区（竞秀/莲池）承接京津疏解；县域（雄安相关/白洋淀）与农业並重。")
para(doc, "城乡收入比1.85（中等），农村+5.5%快于城镇、差距在收敛；雄安协同+乡村振兴。")

# ---------------- 十 ----------------
heading1(doc, "十、人口流入与流出")
para(doc, "保定常住900万、自然-3.92\u2030（净死亡大于出生），靠机械流入（雄安/京津疏解带来的）+本地青壮年外出部分。")
para(doc, "**流入**：雄安协同承接、京津疏解产业人口、高校；**流出**：青壮年赴京津。总体\u201c净流入趋势、结构改善\u201d但自然负增是压力。")

# ---------------- 十一 ----------------
heading1(doc, "十一、物价与货币环境")
para(doc, "2025年保定CPI**上涨0.3%**（温和），食品-0.4%、交通-2.1%。")
para(doc, "金融：存款14681.3亿/+7.3%（住户+9.5%）、贷款9049.4亿/+8.1%；存贷比61.6%。\u201c宽中向上\u201d。")

# ---------------- 十二 ----------------
heading1(doc, "十二、区域一体化：保定在\u201c京津冀协同+雄安+保定都市圈\u201d里的位置")
para(doc, "保定是京津冀协同核心、白洋淀/雄安新区（2017年）临近，\u201c雄保一体化\u201d、保定都市圈获批；承接北京/天津功能疏解（京津冀协同）。")
para(doc, "制造业协同：长城汽车（海外45万辆）、装备制造对接京津/雄安，华北物流枢纽。保定是\u201c一保定雄安+京津冀协同\u201d枢纽市。")

# ---------------- 十三 ----------------
heading1(doc, "十三、未来5—10年最值得观察的五条主线")
bullet(doc, "1. **装备制造升级**：装备占43%、电力/高端装备，制造强市。")
bullet(doc, "2. **汽车（长城/新能源汽车）**：新能源17万辆、海外45万辆、汽车链条。")
bullet(doc, "3. **雄安协同/京津冀**：疏解承接、白洋淀生态。")
bullet(doc, "4. **固定资产投资修复**：固投-25.2%后，靠基建/产业/雄安协同修复。")
bullet(doc, "5. **人口/消费/城镇化**：自然负增下低城镇化率提升空间。")

# ---------------- 十四 ----------------
heading1(doc, "十四、最终结论：保定在\u201c装备+汽车+京津冀协同\u201d里的增长逻辑")
para(doc, "**结论**：保定2025年的\u201c真相\u201d是——**GDP+6%（达标）、规上+9.4%、装备/汽车强、进出口+6.8%、固投-25.2%、财政+1.8%**。它是\u201c制造+京津冀协同+汽车\u201d驱动，但投资弱、人口自然负增的城市。")
para(doc, "**对趋势判断**：装备制造+汽车+出口代表保定\u201c动态\u201d，投资/人口代表\u201c约束\u201d。**制造升级+雄安协同**决定中期潜力，**投资修复+京津冀红化**决定韧性。")
para(doc, "**若只看一个指标**：看**装备制造增速（+13.4%）与固投增速（-25.2%）**——保定靠\u201c装备+汽车\u201d强工业，但投资大幅收缩是隐忧；若固投能依托雄安/京津冀修复，保定可从中期规模增长。")

# ---------------- 附录A ----------------
heading1(doc, "附录A：主要资料来源与核验路径")
bullet(doc, "保定市人民政府《2025年保定市政府工作报告》（2025年，张宝亮）。")
bullet(doc, "保定市统计局《2025年保定市国民经济和社会发展统计公报》（2026年5月）。")
bullet(doc, "河北省2025年统计公报、返回京津冀协同规划交叉核验。")

# ---------------- 附录B ----------------
heading1(doc, "附录B：建议建立的年度跟踪仪表盘")
bullet(doc, "GDP（含/不含定州）增速 vs 目标。")
bullet(doc, "规上工业增加值及分行业（装备/汽车/光伏/医药）增速。")
bullet(doc, "固定资产投资（总量/工业/房地产/设备）增速。")
bullet(doc, "社会消费品零售总额、餐饮、网络零售。")
bullet(doc, "进出口、出口、民营企业占比。")
bullet(doc, "一般公共预算收入/税收、财政支出、民生占比。")
bullet(doc, "常住人口、城镇化率、自然增长率。")
bullet(doc, "CPI/核心CPI、规上工业利润。")
bullet(doc, "长城汽车销量/新能源汽车、光伏。")
bullet(doc, "雄安协同项目、京津冀疏解。")

# ---------------- 保存 ----------------
out = "/Users/x/Desktop/content-prod-lab/reports/保定市_2025年政府工作报告_深度研究_2026-08-13.docx"
doc.save(out)
print("SAVED OK 保定市_2025年政府工作报告_深度研究_2026-08-13.docx")