# -*- coding: utf-8 -*-
"""Build 昆明市2025年政府工作报告 深度研究 DOCX, 参照省会系列版式。"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DARK = RGBColor(0x1F, 0x2A, 0x44)
RED = RGBColor(0xB0, 0x1E, 0x2E)
GRAY = RGBColor(0x59, 0x60, 0x69)
BLUE = RGBColor(0x1F, 0x4E, 0x79)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
HEADER_FILL = "1F2A44"
K = "微软雅黑"


def set_run(run, size=11, bold=False, color=DARK, font=K):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    run.font.name = font
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), font)
    rFonts.set(qn("w:hAnsi"), font)
    rFonts.set(qn("w:eastAsia"), font)


def para(doc, text, size=11, bold=False, color=DARK, align=None,
         space_after=6, space_before=0, font=K):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        set_run(r, size=size, bold=bold or (i % 2 == 1), color=color, font=font)
    return p


def heading1(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(16)
    pf.space_after = Pt(8)
    r = p.add_run(text)
    set_run(r, size=16, bold=True, color=RED)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "B01E2E")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def heading2(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(10)
    pf.space_after = Pt(4)
    r = p.add_run(text)
    set_run(r, size=13, bold=True, color=BLUE)
    return p


def table(doc, headers, rows, widths=None, font_size=9.5):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_run(r, size=font_size, bold=True, color=WHITE)
        tcPr = hdr[i]._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), HEADER_FILL)
        tcPr.append(shd)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(val))
            set_run(r, size=font_size, color=DARK)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    return t


def quote_box(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(8)
    pf.left_indent = Pt(12)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), "B01E2E")
    pBdr.append(left)
    pPr.append(pBdr)
    r = p.add_run(text)
    set_run(r, size=11, color=DARK, font="楷体")
    return p


def bullet(doc, text, size=10.5):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.7)
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        set_run(r, size=size, bold=(i % 2 == 1), color=DARK)
    return p


# ================================================================= 文档主体
doc = Document()
sec = doc.sections[0]
sec.page_width = Cm(21)
sec.page_height = Cm(29.7)
sec.top_margin = Cm(2.4)
sec.bottom_margin = Cm(2.2)
sec.left_margin = Cm(2.5)
sec.right_margin = Cm(2.5)

st = doc.styles["Normal"]
st.font.name = K
st.font.size = Pt(11)
st.element.rPr.rFonts.set(qn("w:eastAsia"), K)


# ---- 封面 ----
para(doc, "昆明市2025年政府工作报告", size=24, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=60, space_after=4)
para(doc, "深度研究与\u201c容易被忽视的细节\u201d分析", size=16, bold=True, color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
para(doc, "从\u201c面向南亚东南亚辐射中心、旅游、生物医药与区域中心\u201d重新理解昆明", size=12, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)
para(doc, "\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501", color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
para(doc, "研究对象：2025年昆明市政府工作报告及同期财政、国民经济和社会发展统计资料、2026年官方复盘", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
para(doc, "标定日期：2026-08-13", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
doc.add_page_break()

# ---- 目录 ----
catalog = [
    "执行摘要",
    "一、研究口径与阅读方法",
    "二、先看昆明的特殊底盘：面向南亚东南亚辐射中心、旅游、生物医药与磨憨口岸",
    "三、最关键的宏观错位：GDP破8600亿但偏低，装备/电子强，投资/消费弱",
    "四、容易被忽视、但信号很重的细节（15条）",
    "五、2025年GDP目标 vs 实际：对照表",
    "六、2025年增长是谁撑起来的？（结构归因）",
    "七、预算与财政的\u201c含金量\u201d",
    "八、民生底账：人口、收入与城乡",
    "九、城镇与农村：格局与均衡",
    "十、人口流入与流出",
    "十一、物价与货币环境",
    "十二、区域一体化：昆明在\u201c面向南亚东南亚+中老铁路+滇中城市群\u201d里的位置",
    "十三、未来5—10年最值得观察的五条主线",
    "十四、最终结论：昆明在\u201c对外开放+制造转型+生物医药\u201d里的增长逻辑",
    "附录A：主要资料来源与核验路径",
    "附录B：建议建立的年度跟踪仪表盘",
]
para(doc, "目　录", size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=10, space_after=10)
for item in catalog:
    para(doc, item, size=12, space_after=4)
doc.add_page_break()

# ---- 执行摘要 ----
heading1(doc, "执行摘要")
para(doc, "**核心判断**　2025年昆明最显著的是\u201cGDP 8637.45亿元、增长4.2%（低于5%目标）、三产占67.2%\u201d、\u201c规上工业+5.1%但装备+22.3%、电子+39.4%、冶金+37.1%\u201d、\u201c进出口1647.81亿/+15.2%\u201d、\u201c常住人口874.4万/城镇化84.0%\u201d。这说明昆明在\u201c面向南亚东南亚+装备制造+磨憨口岸\u201d下稳中提质，但**投资-9.9%、消费+1.3%、财政+5.1%**分化。")
para(doc, "把2025年目标（GDP+5%/规上+6%/产业投资+7%/社零+5%）、2025年统计、2026年前瞻一起看，昆明是\u201c面向南亚东南亚辐射中心+制造强市+旅游\u201d路径：**装备/电子/冶金/生物医药+中老铁路+磨憨口岸**是引擎。总量8637亿居云南第一。")
para(doc, "最容易记住的一句话：**昆明是\u201c面向南亚东南亚辐射中心+旅游+生物医药\u201d的云南省会，靠\u201c装备制造、电子信息、冶金、旅游、中老铁路/磨憨口岸\u201d增长。**观察昆明，与其只看\u201cGDP 8637亿\u201d，不如看\u201c电子信息+39.4%、装备制造+22.3%、进出口+15.2%、磨憨/中老铁路、旅游（游客+10.6%）\u201d。")# ---- 一、研究口径与阅读方法 ----
heading1(doc, "一、研究口径与阅读方法")
para(doc, "本报告以\u201c昆明市2025年政府工作报告（2025年）\u201d为起点，把\u201c2025年GDP目标（5%左右）\u201d与\u201c2025年国民经济和社会发展统计公报实际值（8637.45亿元/+4.2%）\u201d并置对照，再用2026年计划和政府工作报告的复盘作事后验证。")
para(doc, "重点阅读三类证据：**一是指标目标是否达标**；**二是结构（谁贡献增长）是否匹配目标叙事**；**三是官方复盘如何自评、承认问题**。统计口径一般公共预算收入为地方级。")
para(doc, "统一采用\u201c同比增长%\u201d（GDP为不变价），财政与居民收入多为名义值，文中按原口径转录、不逐项换算。")

# ---- 二、特殊底盘 ----
heading1(doc, "二、先看昆明的特殊底盘：面向南亚东南亚辐射中心、旅游、生物医药与磨憨口岸")
para(doc, "**区位与身份**：昆明是云南省会、\u201c面向南亚东南亚辐射中心\u201d、面向东南亚的开放枢纽，中老铁路始发地，磨憨—磨丁经济合作区（昆明托管）、滇中城市群核心。")
para(doc, "**产业底盘**：一是装备制造（+22.3%、电子+39.4%、冶金+37.1%）；二是生物医药（\u201c5+1\u201d全产业链、康乐卫士HPV疫苗）；三是旅游（游客+10.6%、\u201c旅居昆明\u201d）；四是开放（磨憨、中老铁路、进出口+15.2%）；五是数字经济（算力4600P、数据口岸）。")
para(doc, "**人口底盘**：2025年末常住人口874.4万/+（城镇化84.02%、提高0.48pct）；是云南最大人口城市。")
para(doc, "**市场与出口**：2025年社零3817.69亿/+1.3%；进出口1647.81亿/+15.2%（进口+23.4%、出口+0.8%）。")

# ---- 三、核心宏观错位 ----
heading1(doc, "三、最关键的宏观错位：GDP破8600亿但偏低，装备/电子强，投资/消费弱")
para(doc, "**第一组错位**：2025年GDP目标5%左右，实际8637.45亿/+4.2%**低于目标0.8pct**，但为2020年以来最高增速、高于全省0.1pct；对全省贡献率27%。")
para(doc, "**第二组错位**：规上工业+5.1%（目标+6%略低），但**冶金+37.1%、装备+22.3%、电子+39.4%、高技术制造+21.6%**高增；**固投-9.9%、社零+1.3%**偏弱。\u201c工业转型强、内需弱\u201d。")
para(doc, "**第三组错位**：三产占67.2%（贡献率81.3%）、城镇化84%；人口自然-0.10\u2030，靠机械流入。")
para(doc, "一句话：**昆明是\u201c面向东盟+装备/电子/冶金强、投资与消费弱\u201d的云南省会**——靠制造与开放、依托旅游支撑。")

# ---- 四、15条细节 ----
heading1(doc, "四、容易被忽视、但信号很重的细节（15条）")
bullet(doc, "1. **GDP破8600亿**：8637.45亿/+4.2%、2020年以来最高增速。")
bullet(doc, "2. **装备制造+22.3%、电子信息+39.4%**：制造转型最强信号。")
bullet(doc, "3. **冶金工业+37.1%**：稀贵金属/有色（滇中稀贵金属）走强。")
bullet(doc, "4. **高技术制造+21.6%（占规上11.5%）**：新质生产力。")
bullet(doc, "5. **进出口+15.2%、进口+23.4%**：面向南亚东南亚开放强。")
bullet(doc, "6. **磨憨口岸：公路出入境破200万创高/榴莲集散中心**：中老通道。")
bullet(doc, "7. **固投-9.9%（产业投资-7.2%、民间-4.3%）**：投资承压。")
bullet(doc, "8. **房地产-1.0%但商品住宅+8.1%、销售面积+1.0%**：见底回稳。")
bullet(doc, "9. **旅游游客+10.6%、花费+12.1%、旅居150万**：旅游复苏。")
bullet(doc, "10. **常住874.4万/城镇化84.02%**：人口净流入但自然-0.10‰。")
bullet(doc, "11. **收入：城镇58938元/+2.6%、农村24935元/+5.8%**：城乡2.36。")
bullet(doc, "12. **CPI-0.1%**：低通胀。")
bullet(doc, "13. **财政575.00亿/+5.1%、税收+4.8%**：财政稳增、优于GDP。")
bullet(doc, "14. **民营占GDP41.7%**：民营经济韧性。")
bullet(doc, "15. **空气优良362天/99.2%、清洁能源78.7%**：绿色生态。")# ---- 五、目标 vs 实际对照表 ----
heading1(doc, "五、2025年GDP目标 vs 实际：对照表")
table(doc,
    ["指标", "2025年目标", "2025年实际", "是否达标"],
    [
        ["地区生产总值增速", "5%左右", "8637.45亿元/+4.2%", "未达标"],
        ["规上工业增加值", "+6%左右", "+5.1%", "略低于目标"],
        ["固定资产投资", "（产业投资+7%）", "-9.9%（转负）", "未达标"],
        ["社会消费品零售总额", "+5%左右", "3817.69亿元/+1.3%", "未达标"],
        ["一般公共预算收入", "+2.1%", "575.00亿元/+5.1%", "超目标"],
        ["城镇/农村人均可支配收入", "与GDP同步", "+2.6%/+5.8%", "农村高、城镇低"],
        ["进出口", "+5%左右", "1647.81亿/+15.2%", "大幅超目标"],
    ],
    widths=[3.4, 2.5, 5.0, 3.1])
para(doc, "**简评**：**进出口、财政超目标**；GDP、规上、消费、投资均低于目标。这是\u201c制造与开放强、内需与投资弱\u201d的一年。")

# ---- 六、结构归因 ----
heading1(doc, "六、2025年增长是谁撑起来的？（结构归因）")
para(doc, "**三大产业**：一产355.75亿/+2.9%、二产2478.21亿/+2.3%、三产5803.49亿/+5.0%；三产占67.2%（贡献率81.3%）。增长以三产为主。")
para(doc, "**工业**：规上+5.1%；**冶金+37.1%、装备+22.3%、电子+39.4%、高技术+21.6%**；烟草+1.5%、石化-7.8%、医药-3.4%。规上营收6798.02亿/+7.7%、利税-4.6%。")
para(doc, "**服务业**：互联网+27.9%、信息软件+9.6%、金融+2.5%；旅游/会展/数字经济拉三产。")
para(doc, "**开放**：进出口1647.81亿/+15.2%（进口+23.4%）、磨憨/中老铁路。")
para(doc, "**增长归因**：昆明主要靠**三产（旅游/金融/数字/物流）+工业（装备/电子/冶金）+开放**；消费、地产相对弱。")

# ---- 七、财政 ----
heading1(doc, "七、预算与财政的\u201c含金量\u201d")
para(doc, "2025年地方一般公共预算收入575.00亿元/+5.1%（税收475.66亿/+4.8%、占比82.9%）；支出843.47亿/+0.1%，民生支出626.23亿占74.2%。")
para(doc, "**结构性**：财政收入\u201c稳增、税收质量高\u201d（税收占比82.9%）优于GDP；支出受制于化债（+0.1%）保民生。")
para(doc, "**含金量**：昆明财政\u201c质量高、增速优于GDP\u201d（+5.1%名义），是2025亮点；但化债（存量债务压降）约束支出，自主扩支有限。")

# ---------------- 八、民生 ----------------
heading1(doc, "八、民生底账：人口、收入与城乡")
para(doc, "常住人口874.4万/+（城镇化84.02%）、自然-0.10\u2030。收入：**城镇58938元/+2.6%、农村24935元/+5.8%**，城乡比2.36。消费支出全体38296元/+2.6%。")
para(doc, "就业：城镇新增就业19.89万、失业再就业9.45万。")
para(doc, "**民生结论**：收入农村快于城镇、差距2.36偏大；人口净流入、城镇化84%；就业稳、医保/养老95%。")

# ---------------- 九 ----------------
heading1(doc, "九、城镇与农村：格局与均衡")
para(doc, "昆明城镇化率84%（西南前列），主城区（五华/盘龙/官渡/西山）承载金融/旅游/总部；县域（安宁/晋宁/嵩明）发展装备/冶金/农业。")
para(doc, "城乡收入比2.36（偏高），农村增速快、边际收敛；县域工业+乡村振兴支撑。")

# =========== 十 ============
heading1(doc, "十、人口流入与流出")
para(doc, "昆明常住874.4万、微增约6万（自然-0.10\u2030+机械流入），是云南人口/人才中心。高校（云大/昆工）与东盟/生物医药岗位提供吸引力。")
para(doc, "**流入**：大学生、面向东盟/磨憨人才、旅居人口；**流出**：部分中低技能劳动力。总体\u201c净流入、质量提升\u201d。")

# =========== 十一 ============
heading1(doc, "十一、物价与货币环境")
para(doc, "2025年昆明CPI**下降0.1%**（通缩），食品-0.8%、交通通信-2.3%。")
para(doc, "金融：存款20349.38亿/+5.5%、贷款27141.42亿/+3.4%；上市29家/市值7808.37亿。\u201c宽中稳\u201d。")

# =========== 十二 ============
heading1(doc, "十二、区域一体化：昆明在\u201c面向南亚东南亚+中老铁路+滇中城市群\u201d里的位置")
para(doc, "昆明是云南省会、\u201c面向南亚东南亚辐射中心\u201d、滇中城市群核心；中老铁路（昆明-万象，\u201c一地两检\u201d）、磨憨—磨丁合作区、全国唯一同时拥有边境口岸和数据口岸的省会。")
para(doc, "对外开放：进出口+15.2%、中欧（亚）班列集结中心、南亚东南亚大通道。昆明是\u201c门户+枢纽\u201d型核心城市。")

# ============ 十三 =============
heading1(doc, "十三、未来5—10年最值得观察的五条主线")
bullet(doc, "1. **面向南亚东南亚辐射中心**：中老铁路+磨憨口岸+中欧班列，人流物流在昆集聚。")
bullet(doc, "2. **装备/电子/冶金转型**：装备+22%、电子+39%、滇中稀贵金属集群。")
bullet(doc, "3. **生物医药+大健康**：\u201c5+1\u201d产业链、HPV疫苗、生命科学。")
bullet(doc, "4. **旅游+旅居**：游客+10.6%、避暑旅居150万、\u201c旅居昆明\u201d。")
bullet(doc, "5. **投资修复+财政**：固投-9.9%后、化债，靠制造/开放项目修复。")

# ============ 十四 =============
heading1(doc, "十四、最终结论：昆明在\u201c对外开放+制造转型+生物医药\u201d里的增长逻辑")
para(doc, "**结论**：昆明2025年的\u201c真相\u201d是——**GDP+4.2%（低于目标）、规上+5.1%、装备/电子高增、进出口+15.2%、财政+5.1%、固投-9.9%**。它是\u201c面向南亚东南亚+制造转型+旅游\u201d驱动，但内需/投资偏弱的云南省会。")
para(doc, "**对趋势判断**：装备/电子/冶金/生物医药/开放代表昆明\u201c动能\u201d，投资/消费代表\u201c约束\u201d。**制造+开放+旅游**决定中期潜力，**投资修复+收入改善**决定韧性。")
para(doc, "**若只看一个指标**：看**电子信息增速（+39.4%）+进出口增速（+15.2%）**——昆明是\u201c开放+制造强、但投资弱\u201d的辐射中心省会，能否把东盟通道与先进制造变成投资与消费是关键。")

# ------------- 附录A -------------
heading1(doc, "附录A：主要资料来源与核验路径")
bullet(doc, "昆明市人民政府《2025年昆明市政府工作报告》（2025年1月）。")
bullet(doc, "昆明市统计局《2025年昆明市国民经济和社会发展统计公报》（2026年5月）。")
bullet(doc, "昆明市发改委《2026年国民经济和社会发展计划报告》（2026年2月）。")
bullet(doc, "云南省2025年统计公报与昆明统计年鉴交叉核验。")

# ------------- 附录B -------------
heading1(doc, "附录B：建议建立的年度跟踪仪表盘")
bullet(doc, "GDP总量/增速 vs 全年目标（+/-百分点）。")
bullet(doc, "规上工业增加值及分行业（装备/电子/冶金/医药/烟草）增速。")
bullet(doc, "固定资产投资（总量/产业/工业/民间/房地产）增速。")
bullet(doc, "社会消费品零售总额、旅游/旅居人数。")
bullet(doc, "进出口总额（人民币）、出口、对南亚东南亚增速。")
bullet(doc, "一般公共预算收入/税收/非税、财政支出、民生占比。")
bullet(doc, "常住人口、自然增长率、城镇化率、磨憨口岸人流。")
bullet(doc, "CPI/核心CPI、规上工业利润与营收。")
bullet(doc, "中老铁路客货、跨境贸易、数据口岸。")
bullet(doc, "生物医药/大健康、算力（P）、绿色美谷。")

# ------------- 保存 -------------
out = "/Users/x/Desktop/content-prod-lab/reports/昆明市_2025年政府工作报告_深度研究_2026-08-13.docx"
doc.save(out)
print("SAVED OK 昆明市_2025年政府工作报告_深度研究_2026-08-13.docx")