# -*- coding: utf-8 -*-
"""Build 厦门市2025年政府工作报告 深度研究 DOCX, 参照省市系列版式。"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DARK = RGBColor(0x1F, 0x2A, 0x44)
RED = RGBColor(0xB0, 0x1E, 0x2E)
GRAY = RGBColor(0x59, 0x60, 0x69)
BLUE = RGBColor(0x1F, 0x4E, 0x79)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
HEADER_FILL = "1F2A44"
K = "微软雅黑"


def set_run(run, size=11, bold=False, color=DARK, font=K):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    run.font.name = font
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), font)
    rFonts.set(qn("w:hAnsi"), font)
    rFonts.set(qn("w:eastAsia"), font)


def para(doc, text, size=11, bold=False, color=DARK, align=None,
         space_after=6, space_before=0, font=K):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        set_run(r, size=size, bold=bold or (i % 2 == 1), color=color, font=font)
    return p


def heading1(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(16)
    pf.space_after = Pt(8)
    r = p.add_run(text)
    set_run(r, size=16, bold=True, color=RED)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "B01E2E")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def heading2(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(10)
    pf.space_after = Pt(4)
    r = p.add_run(text)
    set_run(r, size=13, bold=True, color=BLUE)
    return p


def table(doc, headers, rows, widths=None, font_size=9.5):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_run(r, size=font_size, bold=True, color=WHITE)
        tcPr = hdr[i]._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), HEADER_FILL)
        tcPr.append(shd)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(val))
            set_run(r, size=font_size, color=DARK)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    return t


def quote_box(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(8)
    pf.left_indent = Pt(12)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), "B01E2E")
    pBdr.append(left)
    pPr.append(pBdr)
    r = p.add_run(text)
    set_run(r, size=11, color=DARK, font="楷体")
    return p


def bullet(doc, text, size=10.5):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.7)
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        set_run(r, size=size, bold=(i % 2 == 1), color=DARK)
    return p


# ================================================================= 文档主体
doc = Document()
sec = doc.sections[0]
sec.page_width = Cm(21)
sec.page_height = Cm(29.7)
sec.top_margin = Cm(2.4)
sec.bottom_margin = Cm(2.2)
sec.left_margin = Cm(2.5)
sec.right_margin = Cm(2.5)

st = doc.styles["Normal"]
st.font.name = K
st.font.size = Pt(11)
st.element.rPr.rFonts.set(qn("w:eastAsia"), K)



# ---- 封面 ----
para(doc, "厦门市2025年政府工作报告", size=24, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=60, space_after=4)
para(doc, "深度研究与\u201c容易被忽视的细节\u201d分析", size=16, bold=True, color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
para(doc, "从\u201c经济特区、外向型经济、金砖基地与台海融合\u201d重新理解厦门", size=12, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)
para(doc, "\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501", color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
para(doc, "研究对象：2025年厦门市政府工作报告及同期财政、国民经济和社会发展统计资料、2026年官方复盘", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
para(doc, "标定日期：2026-08-13", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
doc.add_page_break()

# ---- 目录 ----
catalog = [
    "执行摘要",
    "一、研究口径与阅读方法",
    "二、先看厦门的底盘：经济特区、外向型、金砖基地与台海融合",
    "三、最关键的宏观错位：GDP破8900亿、工业/出口/文旅强，但地产/外资偏弱",
    "四、容易被忽视、但信号很重的细节（15条）",
    "五、2025年 GDP/目标 vs 实际：对照表",
    "六、2025年增长是谁撑起来的？（结构归因）",
    "七、预算与财政的\u201c含金量\u201d",
    "八、民生底账：人口、收入与城乡",
    "九、城镇与农村：格局与均衡",
    "十、人口流入与流出",
    "十一、物价与货币环境",
    "十二、区域一体化：厦门在\u201c闽西南协同+两岸融合+海上丝路·金砖\u201d里的位置",
    "十三、未来5—10年最值得观察的五条主线",
    "十四、最终结论：厦门在\u201c外向+制造+台海+文旅\u201d里的增长逻辑",
    "附录A：主要资料来源与核验路径",
    "附录B：建议建立的年度跟踪仪表盘",
]
para(doc, "目　录", size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=10, space_after=10)
for item in catalog:
    para(doc, item, size=12, space_after=4)
doc.add_page_break()

# ---- 执行摘要 ----
heading1(doc, "执行摘要")
para(doc, "**核心判断**　如果只看新闻标题，2025年厦门最显著的是\u201cGDP破8980亿、增长5.7%（高于全国）\u201d、\u201c规上工业+10.2%、高技术+18.1%\u201d、\u201c出口+8.6%、对金砖+16.5%\u201d、\u201c文旅接待游客1.44亿人次\u201d。但这份研究真正值得深读的，是这座\u201c经济特区+外向型+金砖/台海\u201d的计划单列城市，如何在房地产大降（-41.0%）、出口虽稳但进口-3.5%的背景下，靠\u201c工业（电子/电气/金属制品+18%+）+出口/文旅+先进制造\u201d实现5.7%的增长。")
para(doc, "把2025年初设定的目标（GDP增长5.0%—5.5%，按5.5%推进）、2025年《国民经济和社会发展统计公报》、2026年复盘放在一起看，厦门呈现清晰暗线：**从\u201c外向型+房地产\u201d的旧底盘，向\u201c高技术制造+数字经济+文创文旅+金砖/台海开放\u201d升级**。工业是高景气亮点，地产/外资则待修复。")
para(doc, "因此，本报告不按政府工作报告逐段复述，而采用\u201c显性表述—同期数据—制度含义—长期影响\u201d的方式，专门提取容易被忽略、但对判断厦门未来5—10年发展模式更有价值的信号。")
para(doc, "最容易记住的一句话：**厦门是\u201c经济特区+外向型+金砖/台海门户\u201d，靠\u201c高技术工业+出口+文旅+两岸\u201d撑起增长。**观察厦门，与其看\u201cGDP 8980亿\u201d，不如看\u201c规上工业+10.2%、出口+8.6%、旅游1.44亿人次、金砖+台海、两岸融合\u201d这几张名片。")
heading2(doc, "一页速览：2025年厦门经济的\u201c表与里\u201d")
table(doc,
    ["维度", "表面现象", "更深层信号"],
    [
        ["增长", "GDP 8980.37亿、+5.7%", "一产0.3%、二产37.8%、三产61.9%"],
        ["产业", "规上工业+10.2%", "高技术+18.1%、金属制品+54.7%"],
        ["外贸", "进出口9600.22亿、+3.0%", "出口+8.6%、金砖+16.5%"],
        ["投资", "固投-9.0%、房投-41.0%", "基建+11.2%、住宅大降"],
        ["消费", "社零3448.6亿、+3.6%", "餐饮+7.4%、网络+10.8%"],
        ["人口", "常住536.5万、城镇化91.31%", "人口自然增1.68‰"],
        ["文旅", "游客14447万人次（+12.7%）", "入境+54.5%、创汇+42.4%"],
    ],
    widths=[2.2, 5.6, 7.4])
para(doc, "", size=10, space_after=2)

# ---- 一、研究口径与阅读方法 ----
heading1(doc, "一、研究口径与阅读方法")
heading2(doc, "1.1 本报告的核心底稿")
bullet(doc, "**2025年《政府工作报告》**（2025年1月）——目标：GDP 5.0—5.5%、规上工业9%、固投7%、财政3%、就业13万以上。")
bullet(doc, "**《2025年厦门市统计公报》**（市统计局2026-04）——GDP、工业、外贸、文旅、人口实数。")
bullet(doc, "**2026年厦门市政府工作报告/复盘**（2026年1月）——对2025执行追认与金砖/台海/先进制造展望。")
heading2(doc, "1.2 阅读方法：显性—数据—制度—长期")
para(doc, "**关键判别**：数据优先。例2025年GDP目标5-5.5%按5.5%推进、实际5.7%达标；规上工业目标9%、实际10.2%超预期。厦门\u201c工业/出口/文旅强、固投/地产/外资偏弱\u201d，需穿透总量看高技术制造与两岸开放。")

# ---- 二、底盘 ----
heading1(doc, "二、先看厦门的底盘：经济特区、外向型、金砖基地与台海融合")
para(doc, "厦门的地盘取决于它作为\u201c**经济特区+外向型经济+金砖创新基地+台海融合门户**\u201d的特殊定位。它是全国少有的计划单列经济特区、面向金砖与台港澳的开放门户。")
bullet(doc, "**经济特区/开放**：工资、金融、贸易自由化度高；投洽会、自贸片区、两岸融合先行区。")
bullet(doc, "**外向型制造**：规上工业+10.2%，计算机通信/电气机械/金属制品占规上工业48%+；高教制造业占41.3%。")
bullet(doc, "**金砖基地**：金砖创新基地、江济金砖出海综合服务港；\u201c一带一路\u201d对沿线+6.8%。")
bullet(doc, "**台海融合**：厦金大桥、翔安国际机场共用金门、小三通；对台贸易523.21亿。")
para(doc, "这一底板决定了厦门2025成绩单的\u201c底色\u201d：**只要高技术制造/出口/金砖/两岸持续，厦门就站在\u201c开放型+制造+两岸\u201d增长极；若地产/外资/传统出口承压，厦门需承受\u201c制造文旅强、固投/外资弱\u201d的结构挑战。**")

# ---- 三、宏观错位 ----
heading1(doc, "三、最关键的宏观错位：GDP破8900亿、工业/出口/文旅强，但地产/外资偏弱")
para(doc, "厦门2025年最值得咀嚼的错位，是\u201c**工业/出口/文旅落地强、固投/地产/外资偏弱**\u201d。这种错位决定了对这座\u201c外向型经济特区\u201d城市的观察不能只看GDP总量。")
bullet(doc, "**GDP**：8980.37亿、+5.7%（高于全国/全省）。一产24.24亿（+2.5%，占比0.3%）、二产3394.74亿（+7.1%，占比37.8%）、三产5561.39亿（+4.8%，占比61.9%）。")
bullet(doc, "**工业**：规上工业+10.2%；计算机通信+15.3%、电气机械+19.0%、金属制品+54.7%、医药+17.6%、高技术制造+18.1%（占41.3%）。")
bullet(doc, "**外贸**：进出口9600.22亿、+3.0%（出口5407.63亿、+8.6%；进口4192.58亿、-3.5%）；对金砖+16.5%、一带一路+6.8%。")
bullet(doc, "**消费/文旅**：社零3448.6亿、+3.6%；游客14447万人次+12.7%、旅游总花费2135.77亿+11.8%。")
bullet(doc, "**固投/地产**：固投-9.0%、房地产开发投资-41.0%；基建+11.2%、工业投资+3.0%。")
para(doc, "**为什么读这条**：厦门作为\u201c人口小城+经济特区\u201d，结构性矛盾是\u201c工业/出口/文旅强、固投/地产/外资弱\u201d。增长靠高技术制造+出口+文旅，但地产与投资大幅收缩。")

# ---- 四、15条细节 ----
heading1(doc, "四、容易被忽视、但信号很重的细节（15条）")
para(doc, "以下15条，都在2025年报告/公报里，但常被\u201c规上工业+10.2%\u201d、\u201cGDP+5.7%\u201d等总量掩盖。它们是判断厦门2025之后5—10年的关键小信号。")
bullet(doc, "**1. 规上工业+10.2%**：工业是厦门的绝对引擎，增速远高于GDP。")
bullet(doc, "**2. 高技术制造+18.4%、占规上工业41.3%**：电子信息/半导体/高端制造是高阶名片。")
bullet(doc, "**3. 出口+8.6%**：出口好于进口（-3.5%），外向需求是支撑。")
bullet(doc, "**4. 对金砖+16.5%、一带一路+6.8%**：金砖基地/多元市场。")
bullet(doc, "**5. 民营进出口+17.4%、占54%**：民营外贸活跃。")
bullet(doc, "**6. 房地产投资-41.0%**：地产深度回落，固投-9%主因。")
bullet(doc, "**7. 集成电路产量-16.6%（但锂离子电池/印制电路板高增）**：半导体部分承压。")
bullet(doc, "**8. 常住536.5万、城镇化91.31%**：国内少有的超高城镇化、岛外人口占61.2%。")
bullet(doc, "**9. 收入77762元、+4.7%（农村+6.0%）**：居民收入居全国前列，农村增速高。")
bullet(doc, "**10. 文创/会展/赛事**：投洽会、金鸡百花、马拉松/钻石联赛，会展经济287.95亿。")
bullet(doc, "**11. 旅游入境+54.5%、创汇+42.4%**：入境游爆发、国际名片。")
bullet(doc, "**12. 金砖出海综服港/台海**：厦金大桥、小三通；两岸融合。")
bullet(doc, "**13. 财政：总收入1638.35亿/+3.4%、地方961.08亿/+3.0%**：财政稳健（好于多数城市）。")
bullet(doc, "**14. 环境PM2.5 17.4微克、优良率99.2%全国第2**：宜居生态。")
bullet(doc, "**15. 空气质量全国第8、世界一流旅游城市**：城市品牌。")
para(doc, "", size=10, space_after=2)

# ---- 五、对照表 ----
heading1(doc, "五、2025年 GDP/目标 vs 实际：对照表")
para(doc, "2025年报告中列出的主要预期目标，与2025年实际完成：")
tb = [
    ["指标", "2025年目标", "2025年实际（公报/报告）", "达标判定"],
    ["GDP增速", "5.0%—5.5%", "+5.7%（8980.37亿）", "超标"],
    ["规上工业增加值", "9%左右", "+10.2%", "超预期"],
    ["固定资产投资", "7%左右", "-9.0%", "明显负"],
    ["一般公共预算总收入", "3%左右", "+3.4%（1638.35亿）", "达标"],
    ["地方一般公共预算收入", "3%左右", "+3.0%（961.08亿）", "达标"],
    ["城镇新增就业", "13万人以上", "18.01万人", "达标"],
    ["居民人均可支配收入", "与经济增长同步", "+4.7%（77762元）", "达标"],
    ["CPI", "2%以内", "持平(0%)", "达标"],
]
table(doc, tb[0], tb[1:], widths=[3.2, 3.2, 4.8, 4.0])
para(doc, "", size=10, space_after=2)
para(doc, "**要点**：GDP/工业/财政/就业均达标或超预期，固投/地产负增长——厦门\u201c工业/民生强、投资/地产弱\u201d。")

# ---- 六、结构归因 ----
heading1(doc, "六、2025年增长是谁撑起来的？（结构归因）")
heading2(doc, "6.1 三次产业：二产最猛、三产稳")
para(doc, "**二产（+7.1%、占比37.8%）最强**：规上工业+10.2%、全部工业增加值+8.7%。三产+4.8%（占比61.9%）、一产+2.5%（占比0.3%）。工业是绝对引擎。")
heading2(doc, "6.2 工业：电子/电气/金属制品驱动")
para(doc, "规上工业+10.2%；电气机械+19.0%、计算机通信+15.3%、金属制品+54.7%、铁路船舶+17.5%、医药+17.6%、高技术制造+18.4%。三大行业合计占规上工业49.7%。")
heading2(doc, "6.3 出口/文旅强")
para(doc, "出口+8.6%、金砖+16.5%、跨境电商；游客1.44亿人次+12.7%、旅游花费+11.8%——外需+文旅双轮。")
heading2(doc, "6.4 消费企稳")
para(doc, "社零+3.6%；餐饮+7.4%、日用品+13.2%、化妆品+41.1%、网络零售+10.8%；汽车-6.3%。")
heading2(doc, "6.5 投资/地产走弱")
para(doc, "固投-9.0%、房地产-41.0%、住宅-40.5%；基建+11.2%、社会事业+24.2%。")
para(doc, "**一句话归因**：厦门2025年\u201c**高技术制造+出口/文旅+现代服务**\u201d为主引擎，\u201c固投/地产/外资\u201d走弱——靠工业+文旅的开放型经济特区。")

# ---- 七、财政 ----
heading1(doc, "七、预算与财政的\u201c含金量\u201d")
para(doc, "**一般公共预算总收入1638.35亿元、+3.4%**；地方级961.08亿元、+3.0%；一般公共预算支出1076.56亿元、+1.6%。")
bullet(doc, "**收入质地稳定**：地方收入+3.0%、税收入+3.3%（增值税+9.8%个税+14.0%）；财政质量好于多数同类城市。")
bullet(doc, "**民生硬度高**：教育支出223.44亿、科学技术79.08亿（+8.8%）；一级支出结构向民生/科创倾斜。")
bullet(doc, "**债务风险可控**：金融稳健，本外币存款1.95万亿+9.2%、贷款+2.8%；政府债风险低。")
para(doc, "**财政含义**：厦门\u201c收入稳、科创投入强、风险可控\u201d，政策空间好于地产依赖型城市，需靠高技术制造/金砖创造税源。")

# ---- 八、民生 ----
heading1(doc, "八、民生底账：人口、收入与城乡")
para(doc, "厦门\u201c高城镇化、人口结构年轻\u201d：**常住人口536.50万人、城镇化率91.31%、自然增长率1.68‰**（出生5.60‰、死亡3.92‰），超高城镇化+岛外占61.2%。")
bullet(doc, "**收入**：居民人均可支配收入77762元、+4.7%（全国前列）；城镇79719元、+4.7%，农村38512元、+6.0%。")
bullet(doc, "**就业**：城镇新增就业18.01万人；高校毕业生去向落实率95.9%。")
bullet(doc, "**社保**：基本养老参保622.11万（外来占395.67万、+11%）；低保/医保健全。")
bullet(doc, "**人口**：自然增1.68‰（正增长），人口年轻、城市吸引力强。")
para(doc, "**民生含义**：厦门\u201c高收入+高城镇化+人口正增长\u201d，是少数保持自然增长的高收入城市。")

# ---- 九、城乡 ----
heading1(doc, "九、城镇与农村：格局与均衡")
para(doc, "**城镇化率91.31%**（全国最高之一），厦门高度都市化；岛外人口占比61.2%（超岛内），新城/岛内外一体化。")
bullet(doc, "**城市**：岛内（思明/湖里）品质，岛外（海沧/集美/同安/翔安）产业新城；轨道交通/地铁4、6号线。")
bullet(doc, "**农村**：农业总产值53.13亿、样本小；农村收入+6.0%快于城镇。")
para(doc, "**城乡均衡**：厦门\u201c中心极强+岛外高发展\u201d，农村收入增速高，城乡融合程度全国领先。")

# ---- 十、人口流 ----
heading1(doc, "十、人口流入与流出")
para(doc, "**厦门净流入**：2025年自然增长+1.68‰、城镇化91.31%（岛外占61.2%）。在全国人口总体收缩背景下，厦门人口保持增长。")
bullet(doc, "**流入**：高校（厦门大学等）+产业/先进制造/金砖/文旅；外来从业人员约395.67万。")
bullet(doc, "**竞争**：与福州/泉州/珠三角争夺；靠高收入、宜居、两岸环境吸引。")
para(doc, "人口方向决定中长期需求与增长；厦门的\u201c高收入+宜居+两岸\u201d是其长逻辑。")

# ---- 十一、物价 ----
heading1(doc, "十一、物价与货币环境")
para(doc, "**2025年厦门CPI与上年持平**：食品-0.8%、非食品+0.2%；服务-0.2%、消费品+0.2%。")
bullet(doc, "**物价**：温和、低位，反映内需压力可控。")
bullet(doc, "**货币/流动性**：本外币存款1.95万亿+9.2%、贷款+2.8%；证券交易额15.05万亿+58.1%活跃。")
para(doc, "**物价含义**：厦门\u201c通缩压力可控\u201d，消费修复有限；关注收入与服务价格。")

# ---- 十二、区域一体化 ----
heading1(doc, "十二、区域一体化：厦门在\u201c闽西南协同+两岸融合+海上丝路·金砖\u201d里的位置")
para(doc, "厦门处于**闽西南协同发展区+两岸融合发展+海上丝绸之路·金砖创新基地**核心：既是经济特区/计划单列，也是面向台港澳与金砖的开放枢纽。")
bullet(doc, "**闽西南协同**：与漳州/泉州/龙岩/三明协同，港口/产业协作；都市圈一体化。")
bullet(doc, "**两岸融合**：厦金大桥（厦门段）、金门共用翔安国际机场、小三通；建设两岸（厦泉金）合作发展区；对台贸易523.21亿。")
bullet(doc, "**金砖/一带一路**：金砖创新基地、金砖出海综服港；对一带一路出口+6.8%。")
bullet(doc, "**海上丝路**：外贸9600亿（出口全国前列）、枢纽港+机场，国际开放门户。")
para(doc, "**区域含义**：厦门作为\u201c两岸门户+海上丝路/金砖\u201d枢纽，靠\u201c两岸+金砖+闽西南\u201d带动福建与丝路。")

# ---- 十三、五条主线 ----
heading1(doc, "十三、未来5—10年最值得观察的五条主线")
bullet(doc, "**主线1｜高技术制造/集成电路**：规上工业+10.2%、高技术+18.4%。能否在半导体/电子信息/高端装备站稳。")
bullet(doc, "**主线2｜外向型/金砖/台海**：出口+8.6%、金砖+16.5%、两岸融合。能否在外部市场多元化/两岸融合卡位。")
bullet(doc, "**主线3｜文创文旅/会展赛事**：游客+12.7%、投洽会/金鸡/马拉松。能否建成世界一流旅游城市。")
bullet(doc, "**主线4｜数字经济/新质**：数字经济核心6600亿目标、海洋生产总值2800亿。能否孵化新增长极。")
bullet(doc, "**主线5｜经济特区改革**：综合改革试点、自贸片区。能否持续扮演\u201c改革先行\u201d角色。")
para(doc, "这五条，是厦门从\u201c外向经济特区\u201d走向\u201c科创+文旅+两岸开放高地\u201d的\u201c主赛道\u201d。")

# ---- 十四、结论 ----
heading1(doc, "十四、最终结论：厦门在\u201c外向+制造+台海+文旅\u201d里的增长逻辑")
para(doc, "厦门2025年，本质上是\u201c**高技术制造+出口/文旅驱动、固投/地产/外资偏弱**\u201d的答卷：GDP8980.37亿、+5.7%、规上工业+10.2%、高技术+18.4%、出口+8.6%、文旅+12.7%、财政总收入+3.4%、常住城镇化91.31%。")
para(doc, "只要高技术制造/出口/金砖/两岸持续，厦门就站在\u201c开放型+制造+两岸\u201d增长极；若地产/固投深度调整，厦门需承受\u201c工业文旅强、投资/外资弱\u201d的结构挑战。")
para(doc, "最稳观察信号：**一盯规上工业/高技术（引擎）、二盯出口/金砖（开放）、三盯文旅/入境（消费）、四盯地产/固投（投资）、五盯两岸/改革（长期）。**厦门，是\u201c外向+制造+两岸\u201d的新样本。")

# ---- 附录A ----
heading1(doc, "附录A：主要资料来源与核验路径")
bullet(doc, "厦门市2025年政府工作报告（2025年1月）——目标来源。")
bullet(doc, "《2025年厦门市国民经济和社会发展统计公报》（市统计局）——GDP、工业、外贸、文旅、人口实值。")
bullet(doc, "2026年厦门市政府工作报告（2026年1月）——2025复盘/金砖/台海/先进制造。")
bullet(doc, "厦门海关、市财政局（外贸/财政）。")
heading2(doc, "核验说明")
para(doc, "本报告涉及数据以统计公报/官方口径为准；\u201c外向/金砖/台海/文旅/制造\u201d等以官方口径为准。")

# ---- 附录B ----
heading1(doc, "附录B：建议建立的年度跟踪仪表盘")
para(doc, "建议把下面10个\u201c测脉搏\u201d指标做成年度跟踪表：")
dash = [
    ["#", "指标", "2025基值", "为什么重要"],
    ["1", "GDP增速", "+5.7%（8980.37亿）", "总量与方向"],
    ["2", "规模以上工业增速", "+10.2%", "制造底盘"],
    ["3", "高技术制造占比", "+18.4% / 41.3%", "产业升级"],
    ["4", "出口/金砖", "+8.6% / +16.5%", "开放韧性"],
    ["5", "文旅/入境", "+12.7% / 入境+54.5%", "消费文旅"],
    ["6", "固定资产投资/地产", "-9.0% / -41.0%", "投资结构"],
    ["7", "社零增速", "+3.6%（3448.6亿）", "内需消费"],
    ["8", "常住人口/城镇化", "536.5万 / 91.31%", "人口与城市"],
    ["9", "地方财政收入", "+3.0%（961.08亿）", "财政质量"],
    ["10", "CPI", "持平(0%)", "物价"],
]
table(doc, dash[0], dash[1:], widths=[0.8, 4.6, 4.6, 3.6])
para(doc, "把这10个指标连起来看，高技术制造/出口/文旅（2/3/4）、人口/城镇化（8），都说明厦门在真正换挡。")

# ============ 保存 ============
out = "/Users/x/Desktop/content-prod-lab/reports/厦门市_2025年政府工作报告_深度研究_2026-08-13.docx"
doc.save(out)
print("SAVED:", out)
print("PARAGRAPHS:", len(doc.paragraphs))
print("TABLES:", len(doc.tables))
