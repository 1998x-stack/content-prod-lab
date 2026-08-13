# -*- coding: utf-8 -*-
"""Build 宁夏回族自治区2025年政府工作报告 深度研究 DOCX, 参照省系列版式。"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DARK = RGBColor(0x1F, 0x2A, 0x44)
RED = RGBColor(0xB0, 0x1E, 0x2E)
GRAY = RGBColor(0x59, 0x60, 0x69)
BLUE = RGBColor(0x1F, 0x4E, 0x79)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
HEADER_FILL = "1F2A44"
K = "微软雅黑"


def set_run(run, size=11, bold=False, color=DARK, font=K):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    run.font.name = font
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), font)
    rFonts.set(qn("w:hAnsi"), font)
    rFonts.set(qn("w:eastAsia"), font)


def para(doc, text, size=11, bold=False, color=DARK, align=None,
         space_after=6, space_before=0, font=K):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        set_run(r, size=size, bold=bold or (i % 2 == 1), color=color, font=font)
    return p


def heading1(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(16)
    pf.space_after = Pt(8)
    r = p.add_run(text)
    set_run(r, size=16, bold=True, color=RED)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "B01E2E")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def heading2(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(10)
    pf.space_after = Pt(4)
    r = p.add_run(text)
    set_run(r, size=13, bold=True, color=BLUE)
    return p


def table(doc, headers, rows, widths=None, font_size=9.5):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_run(r, size=font_size, bold=True, color=WHITE)
        tcPr = hdr[i]._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), HEADER_FILL)
        tcPr.append(shd)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(val))
            set_run(r, size=font_size, color=DARK)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    return t


def quote_box(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(8)
    pf.left_indent = Pt(12)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), "B01E2E")
    pBdr.append(left)
    pPr.append(pBdr)
    r = p.add_run(text)
    set_run(r, size=11, color=DARK, font="楷体")
    return p


def bullet(doc, text, size=10.5):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.7)
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        set_run(r, size=size, bold=(i % 2 == 1), color=DARK)
    return p


# ================================================================= 文档主体
doc = Document()
sec = doc.sections[0]
sec.page_width = Cm(21)
sec.page_height = Cm(29.7)
sec.top_margin = Cm(2.4)
sec.bottom_margin = Cm(2.2)
sec.left_margin = Cm(2.5)
sec.right_margin = Cm(2.5)

st = doc.styles["Normal"]
st.font.name = K
st.font.size = Pt(11)
st.element.rPr.rFonts.set(qn("w:eastAsia"), K)



# ---- 封面 ----
para(doc, "宁夏回族自治区2025年政府工作报告", size=24, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=60, space_after=4)
para(doc, "深度研究与\u201c容易被忽视的细节\u201d分析", size=16, bold=True, color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
para(doc, "从\u201c新能源、特色农业、数字经济与民族自治区\u201d重新理解宁夏", size=12, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)
para(doc, "\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501", color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
para(doc, "研究对象：2025年宁夏回族自治区政府工作报告及同期财政、国民经济和社会发展统计资料、2026年官方复盘", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
para(doc, "标定日期：2026-08-13", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
doc.add_page_break()

# ---- 目录 ----
catalog = [
    "执行摘要",
    "一、研究口径与阅读方法",
    "二、先看宁夏的特殊底盘：新能源、特色农业、绿电算力与民族自治区",
    "三、最关键的宏观错位：GDP破5600亿、工业/投资强，但消费/外贸体量小",
    "四、容易被忽视、但信号很重的细节（15条）",
    "五、2025年GDP目标 vs 实际：对照表",
    "六、2025年增长是谁撑起来的？（结构归因）",
    "七、预算与财政的\u201c含金量\u201d",
    "八、民生底账：人口、收入与城乡",
    "九、城镇与农村：格局与均衡",
    "十、人口流入与流出",
    "十一、物价与货币环境",
    "十二、区域一体化：宁夏在\u201c沿黄城市群+黄河几字弯+西部陆海新通道\u201d里的位置",
    "十三、未来5—10年最值得观察的五条主线",
    "十四、最终结论：宁夏在\u201c新能源+特色农业+算力\u201d里的增长逻辑",
    "附录A：主要资料来源与核验路径",
    "附录B：建议建立的年度跟踪仪表盘",
]
para(doc, "目　录", size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=10, space_after=10)
for item in catalog:
    para(doc, item, size=12, space_after=4)
doc.add_page_break()

# ---- 执行摘要 ----
heading1(doc, "执行摘要")
para(doc, "**核心判断**　如果只看新闻标题，2025年宁夏最显著的是\u201cGDP破5696亿、增长5.3%（连续居全国前六）\u201d、\u201c规上工业+6.9%、装备制造+18.1%\u201d、\u201c新能源装机5796万千瓦（占63.6%）\u201d、\u201c农业+6.0%（全国第二）\u201d。但这份研究真正值得深读的，是这座\u201c新能源+特色农业+绿电算力\u201d的民族自治区，如何在消费体量小（社零1449亿）、外贸体量小（209亿）背景下，靠\u201c能源/工业/投资+农业/乡村振兴+数字经济\u201d实现5.3%的增长。")
para(doc, "把2025年初设定的目标（GDP增长5.5%左右）、2025年《国民经济和社会发展统计公报》、2026年复盘放在一起看，宁夏呈现清晰暗线：**从\u201c煤/能源\u201d的既有底盘，向\u201c新能源+绿电算力+特色农业+现代煤化工\u201d升级**。能源转型是超预期亮点，二三产趋稳。")
para(doc, "因此，本报告不按政府工作报告逐段复述，而采用\u201c显性表述—同期数据—制度含义—长期影响\u201d的方式，专门提取容易被忽略、但对判断宁夏未来5—10年发展模式更有价值的信号。")
para(doc, "最容易记住的一句话：**宁夏是\u201c新能源+特色农业+绿电算力\u201d的自治区，靠\u201c能源/工业+农业+投资\u201d撑起增长。**观察宁夏，与其只看\u201cGDP 5696亿\u201d，不如看\u201c新能源63.6%、装备制造+18%、葡萄酒/枸杞、算力枢纽\u201d这几张名片。")
heading2(doc, "一页速览：2025年宁夏经济的\u201c表与里\u201d")
table(doc,
    ["维度", "表面现象", "更深层信号"],
    [
        ["增长", "GDP 5696.49亿、+5.3%", "一产8.3%、二产41.4%、三产50.3%"],
        ["产业", "规上工业+6.9%", "装备+18.1%、高技术+13.8%"],
        ["外贸", "进出口209.93亿、+3.4%", "一带一路+8.8%"],
        ["投资", "固投+6.0%、工业+18.5%", "新能源投资+70%、房投-26.9%"],
        ["消费", "社零1449.1亿、+2.1%", "体量小、餐饮+3.4%"],
        ["人口", "常住732万、城镇化69.13%", "人口+3万、自然增"],
        ["能源", "新能源装机5796万千瓦占63.6%", "首个绿电自足省区"],
    ],
    widths=[2.2, 5.6, 7.4])
para(doc, "", size=10, space_after=2)

# ---- 一、研究口径与阅读方法 ----
heading1(doc, "一、研究口径与阅读方法")
heading2(doc, "1.1 本报告的核心底稿")
bullet(doc, "**2025年《政府工作报告》**（2025年1月）——目标：GDP 5.5%左右、规上工业6.5%、固投6%、财政3%。")
bullet(doc, "**《2025年宁夏统计公报》**（统计局2026-04）——GDP、工业、新能源、人口实数。")
bullet(doc, "**2026年宁夏政府工作报告/复盘**（2026年2月）——2025追认与新能源/特色农业展望。")
heading2(doc, "1.2 阅读方法：显性—数据—制度—长期")
para(doc, "**关键判别**：数据优先。例2025年GDP目标5.5%、实际5.3%基本达标；工业目标6.5%、实际6.9%（超）。宁夏\u201c能源/农业/投资强、消费/外贸体量小\u201d，穿透总量看能源结算。")

# ---- 二、底盘 ----
heading1(doc, "二、先看宁夏的特殊底盘：新能源、特色农业、绿电算力与民族自治区")
para(doc, "宁夏的地盘取决于它作为\u201c**新能源+特色农业+绿电算力+民族自治区**\u201d的特殊定位。它是西部面积最小、人口偏少的自治区，但能源/农业特色突出。")
bullet(doc, "**新能源/绿电**：新能源装机5796万千瓦、占63.6%；\u201c宁电入湘\u201d外送；全国首个\u201c绿电自足\u201d省区。")
bullet(doc, "**特色农业**：葡萄酒/枸杞/牛奶/滩羊等\u201c六特\u201d国家产业集群；农产品加工转化率75.8%。")
bullet(doc, "**绿电算力**：\u201c中国算力之都\u201d、数据中心绿电自足/东数西算枢纽。")
bullet(doc, "**民族自治区**：民生倾斜、区域协作、经济总量小但增速稳。")
para(doc, "这一底板决定宁夏2025成绩单\u201c底色\u201d：**只要新能源/特色农业/算力持续，宁夏就站在\u201c绿电+农业+算力\u201d增长极；若内需/外贸体量不变，宁夏需靠政策与能源拉动。**")

# ---- 三、宏观错位 ----
heading1(doc, "三、最关键的宏观错位：GDP破5600亿、工业/投资强，但消费/外贸体量小")
para(doc, "宁夏2025年最值得咀嚼的错位，是\u201c**能源/工业/投资强、消费/外贸体量小**\u201d。这种错位决定了对这座民族自治区经济的观察不能只看GDP增速。")
bullet(doc, "**GDP**：5696.49亿、+5.3%。一产469.86亿（+6.0%，占比8.3%）、二产2360.38亿（+5.2%，占比41.4%）、三产2866.25亿（+5.2%，占比50.3%）。")
bullet(doc, "**工业**：规上工业+6.9%；制造业+11.2%、采矿业+2.5%；装备制造+18.1%、高技术+13.8%。")
bullet(doc, "**投资**：全社会固投+3.7%、不含农户+6.0%；工业投资+18.5%（占62.6%）、新能源投资+70%；房投-26.9%。")
bullet(doc, "**消费**：社零1449.1亿、+2.1%（体量小）；餐饮+3.4%。")
bullet(doc, "**外贸**：进出口209.93亿、+3.4%（出口+2.2%）；一带一路+8.8%。")
bullet(doc, "**财政**：总收入897.8亿/+3.4%、地方540.57亿/+4.6%（税收+7.9%、占71.7%）。")
para(doc, "**为什么读这条**：宁夏作为\u201c人口小省/能源农业大区\u201d，结构性矛盾是\u201c投资/能源/农业强、消费/外贸小\u201d。增速靠投资与能源，内需外需体量有限。")

# ---- 四、15条细节 ----
heading1(doc, "四、容易被忽视、但信号很重的细节（15条）")
para(doc, "以下15条，都在2025年报告/公报，但被\u201cGDP 5600亿\u201d等总量掩盖。它们是判断宁夏2025之后5—10年的关键小信号。")
bullet(doc, "**1. 新能源装机5796万千瓦、占63.6%、太阳能+59.3%**：能源转型第一名片。")
bullet(doc, "**2. 装备制造+18.1%、高技术+13.8%**：制造升级（煤化工/新材料）。")
bullet(doc, "**3. 规上工业利润290亿、+7.8%**：能源/化工利润率好。")
bullet(doc, "**4. 农业+6.0%（全国第二）**：特色农业强。")
bullet(doc, "**5. 葡萄酒/枸杞/牛奶国家级集群**：\u201c六特\u201d产业全部入选。")
bullet(doc, "**6. 宁夏算力/东数西算**：\u201c中国算力之都\u201d、绿电数据中心。")
bullet(doc, "**7. 煤化工超4000万吨、煤制油全国第一**：现代煤化工龙头。")
bullet(doc, "**8. 城镇/居民收入**：居民35184元+5.5%、农村+5.9%（缩小）。")
bullet(doc, "**9. 固投+6%（全国第四）、新能源投资+70%**：投资驱动+能源。")
bullet(doc, "**10. 旅游游客+13.6%、花费+11.8%、入境+25%**：西部文旅增长。")
bullet(doc, "**11. PM2s 25.9微克、优良生态**：环境。")
bullet(doc, "**12. 常住732万、+3万、城镇化69.13%**：人口净流入+城镇。")
bullet(doc, "**13. 财政税收+7.9%**：财政质量好。")
bullet(doc, "**14. 首次年供/外送电量双过1000亿度——西北电力枢纽。")
bullet(doc, "**15. 特色农业/农产品远销**：外销多国。")
para(doc, "", size=10, space_after=2)

# ---- 五、对照表 ----
heading1(doc, "五、2025年GDP目标 vs 实际：对照表")
para(doc, "2025年报告中列出的主要预期目标，与2025年实际完成：")
tb = [
    ["指标", "2025年目标", "2025年实际（公报）", "达标判定"],
    ["GDP增速", "5.5%左右", "+5.3%（5696.49亿）", "基本达标"],
    ["规上工业增加值", "6.5%左右", "+6.9%", "超预期"],
    ["固定资产投资", "6%以上", "+6.0%（不含农户）", "达标"],
    ["社会消费品零售总额", "5%左右", "+2.1%（1449.1亿）", "明显偏低"],
    ["地方一般公共预算收入", "3%左右", "+4.6%（540.57亿）", "超预期"],
    ["居民人均可支配收入", "6%左右", "+5.5%（35184元）", "略欠"],
    ["城镇调查失业率", "5.5%左右", "5.2%", "达标"],
]
table(doc, tb[0], tb[1:], widths=[3.2, 3.4, 4.4, 4.0])
para(doc, "", size=10, space_after=2)
para(doc, "**要点**：工业/固投/财政达标或超预期，社零/收入略欠——宁夏\u201c投资能源强、消费偏弱\u201d。")

# ---- 六、结构归因 ----
heading1(doc, "六、2025年增长是谁撑起来的？（结构归因）")
heading2(doc, "6.1 三次产业：农业/工业双稳")
para(doc, "**一产+6.0%、二产与三产+5.2%**：农业（特色）稳健、工业（能源/化工）支撑、服务业居半。")
heading2(doc, "6.2 工业：能源/装备驱动")
para(doc, "规上工业+6.9%；装备制造+18.1%、高技术+13.8%、制造业+11.2%。新能源/煤化工是高景气。")
heading2(doc, "6.3 投资强、消费弱")
para(doc, "固投+6%、工业+18.5%、新能源投资+70%；但社零仅+2.1%（内需体量小）。")
heading2(doc, "6.4 财政/收入改善")
para(doc, "地方财政+4.6%、税收+7.9%；居民收入+5.5%。")
heading2(doc, "6.5 外贸小")
para(doc, "进出口+3.4%、体量小，不是增长主引擎。")
para(doc, "**一句话归因**：宁夏2025年\u201c**能源/工业/投资（新能源/煤化工）+农业**\u201d为主引擎，\u201c消费/外贸体量小\u201d——投资能源驱动的西部增长模式。")

# ---- 七、财政 ----
heading1(doc, "七、预算与财政的\u201c含金量\u201d")
para(doc, "**一般公共预算总收入897.80亿元、+3.4%**；地方540.57亿元、+4.6%；税收387.40亿元、+7.9%，占地方71.7%。")
bullet(doc, "**收入质地好**：地方+4.6%、税收+7.9%，税占比71.7%——财政质量居西部前列。")
bullet(doc, "**民生/区域倾斜**：民生支出占比高，支持特色农业/算力/民生。")
bullet(doc, "**债务控制**：化解隐性债务、财政可持续。")
para(doc, "**财政含义**：宁夏\u201c收入稳+税质高\u201d，靠能源/特色农业/机电创造税源，财政支撑西部追赶。")

# ---- 八、民生 ----
heading1(doc, "八、民生底账：人口、收入与城乡")
para(doc, "**常住人口732万人（+3万）、城镇化率69.13%（+0.91pct）**；居民人均可支配收入35184元、+5.5%（农村+5.9%快于城镇+4.8%）。")
bullet(doc, "**收入**：居民35184元、+5.5%；城镇46593、+4.8%，农村20142、+5.9%（城乡比2.31、缩小）。")
bullet(doc, "**就业**：城镇新增8.55万、城镇调查失业率5.2%。")
bullet(doc, "**社保**：低保/医保完善，乡村振兴投入。")
bullet(doc, "**人口**：常住+3万、城镇化提升——转入人口/城镇化增量。")
para(doc, "**民生含义**：宁夏\u201c收入增速全国前列、城镇化提升\u201d，是西部民生追赶样本。")

# ---- 九、城乡 ----
heading1(doc, "九、城镇与农村：格局与均衡")
para(doc, "**城镇化率69.13%（+0.91pct）**，宁夏城镇化尚低于全国但提升快；城乡收入比2.31、逐年缩小。")
bullet(doc, "**城市**：银川\u201c强首府\u201d，石嘴山/吴忠/固原/中卫差异发展。")
bullet(doc, "**农村**：粮食381.44万吨、葡萄酒/枸杞/牛奶国家级集群、乡村振兴。")
para(doc, "**城乡均衡**：宁夏\u201c强首府+五市协同+县域特色\u201d，农文旅/特色产业带动乡村。")

# ---- 十、人口流 ----
heading1(doc, "十、人口流入与流出")
para(doc, "**宁夏常住732万（+3万）**：在全国人口总体收缩背景下，宁夏保持小幅净流入/净增长（城镇化提升）。")
bullet(doc, "**流入**：银川强首府/新能源/算力/特色农业吸引；民生好于周边。")
bullet(doc, "**竞争**：与周边（陕西/甘肃/内蒙古）争夺劳动力；宁夏靠能源/农业/政策。")
para(doc, "人口方向决定中长期需求；宁夏的\u201c特色农业/能源区位+\u201d是其长逻辑。")

# ---- 十一、物价 ----
heading1(doc, "十一、物价与货币环境")
para(doc, "**2025年宁夏CPI与上年持平（0%）、PPI出厂/购进均降5.7%**：终端需求偏、工业品价格下行。")
bullet(doc, "**物价**：CPI持平、低位；PPI负、工业品价格压力（能源/化工）。")
bullet(doc, "**货币/流动性**：金融为主、居民存贷稳定。")
para(doc, "**物价含义**：宁夏\u201c通缩压力可控\u201d，内需偏弱、工业品束线，关注消费与能源价格。")

# ---- 十二、区域一体化 ----
heading1(doc, "十二、区域一体化：宁夏在\u201c沿黄城市群+黄河几字弯+西部陆海新通道\u201d里的位置")
para(doc, "宁夏处于**沿黄城市群+黄河几字弯（黄河生态保护/高质量发展）+西部陆海新通道**交汇：既是西部能源/农业特色省区，也是沿黄经济带节点。")
bullet(doc, "**沿黄城市群**：\u201c1+1+9\u201d体系、银川\u201c强首府\u201d、宁东基地。")
bullet(doc, "**黄河几字弯**：黄河流域生态保护，能源/农业绿色发展。")
bullet(doc, "**西部陆海新通道**：宝中铁路、包兰铁路、银巴铁路、河东机场改造。")
bullet(doc, "**跨区域协作**：与沿海结对（宁闽协作）、飞地经济。")
para(doc, "**区域含义**：宁夏作为\u201c沿黄能源农业大区\u201d，靠\u201c能源+农业+黄河\uff0c协作\u201d融入西部新通道。")

# ---- 十三、五条主线 ----
heading1(doc, "十三、未来5—10年最值得观察的五条主线")
bullet(doc, "**主线1｜新能源/绿电**：新能源装机5796万千瓦、宁电外送2000万千瓦。能否建成全国绿电高地。")
bullet(doc, "**主线2｜特色农业（六特）**：葡萄酒/枸杞/牛奶国家级、\u201c六特\u201d。能否做强国家农业集群。")
bullet(doc, "**主线3｜绿电算力/东数西算**：数据中心100%绿电、算力之都。能否把\u201c电力\u201d变\u201c算力\u201d。")
bullet(doc, "**主线4｜现代煤化工**：煤化工4000万吨、煤制油第一。能否绿色升级。")
bullet(doc, "**主线5｜文旅/黄河**：各地游客+13.6%，黄河/西夏陵。能否创世界旅游目的地。")
para(doc, "这五条，是宁夏从\u201c能源+农业省区\u201d走向\u201c绿电+算力+特色农业强区\u201d的\u201c主赛道\u201d。")

# ---- 十四、结论 ----
heading1(doc, "十四、最终结论：宁夏在\u201c新能源+特色农业+绿电算力\u201d里的增长逻辑")
para(doc, "宁夏2025年，本质上是\u201c**能源/工业/投资（新能源+煤化工）+农业驱动、消费/外贸体量小**\u201d的答卷：GDP5696.49亿、+5.3%、规上工业+6.9%、能源新能源占63.6%、投资+6%、社零+2.1%、地方财政+4.6%。")
para(doc, "只要新能源/特色农业/算力持续，宁夏就站在\u201c绿电+农业+算力\u201d增长极；若内需/外需体量开发不足，需承受\u201c投资强、消费小\u201d的结构挑战。")
para(doc, "最稳观察信号：**一盯新能源装机/外送（引擎）、二盯葡萄枸杞/六特（农业）、三盯算力/数字经济（产业）、四盯固投/财政（质量）、五盯人口/城镇化（长期）。**宁夏，是\u201c绿电+特色农业\u201d西部新样本。")

# ---- 附录A ----
heading1(doc, "附录A：主要资料来源与核验路径")
bullet(doc, "宁夏回族自治区2025年政府工作报告（2025年1月）——目标来源。")
bullet(doc, "《2025年宁夏回族自治区国民经济和社会发展统计公报》——GDP、工业、新能源、人口实数。")
bullet(doc, "2026年宁夏政府工作报告（2026年2月）——2025追认/新能源/特色农业/算力。")
bullet(doc, "自治区财政厅、能源局。")
heading2(doc, "核验说明")
para(doc, "本报告涉及数据以统计公报/官方口径为准；\u201c新能源/特色农业/算力\u201d等以官方口径为准。")

# ---- 附录B ----
heading1(doc, "附录B：建议建立的年度跟踪仪表盘")
para(doc, "建议把下面10个\u201c测脉搏\u201d指标做成年度跟踪表：")
dash = [
    ["#", "指标", "2025基值", "为什么重要"],
    ["1", "GDP增速", "+5.3%（5696.49亿）", "总量与方向"],
    ["2", "规模以上工业增速", "+6.9%", "工业底盘"],
    ["3", "新能源装机/占比", "5796万千瓦 / 63.6%", "绿电/能源"],
    ["4", "固投/工业投资", "+6.0% / 工业+18.5%", "投资结构"],
    ["5", "进出口增速", "+3.4%（209.93亿）", "外贸体量"],
    ["6", "社零增速", "+2.1%（1449.1亿）", "内需消费"],
    ["7", "地方财政收入/税占", "+4.6% / 71.7%", "财政质量"],
    ["8", "常住人口/城镇化", "732万 / 69.13%", "人口与城市"],
    ["9", "居民收入", "+5.5%（35184元）", "民生"],
    ["10", "CPI", "持平(0%)", "物价"],
]
table(doc, dash[0], dash[1:], widths=[0.8, 4.8, 4.0, 4.0])
para(doc, "把这10个指标连起来看，新能源/工业/投资（2/3/4）、农业/民生（7/9），都说明宁夏在真正换挡。")

# ============ 保存 ============
out = "/Users/x/Desktop/content-prod-lab/reports/宁夏回族自治区_2025年政府工作报告_深度研究_2026-08-13.docx"
doc.save(out)
print("SAVED:", out)
print("PARAGRAPHS:", len(doc.paragraphs))
print("TABLES:", len(doc.tables))
