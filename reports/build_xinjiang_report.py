# -*- coding: utf-8 -*-
"""Build 新疆维吾尔自治区2025年政府工作报告 深度研究 DOCX, 参照省系列版式。"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DARK = RGBColor(0x1F, 0x2A, 0x44)
RED = RGBColor(0xB0, 0x1E, 0x2E)
GRAY = RGBColor(0x59, 0x60, 0x69)
BLUE = RGBColor(0x1F, 0x4E, 0x79)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
HEADER_FILL = "1F2A44"
K = "微软雅黑"


def set_run(run, size=11, bold=False, color=DARK, font=K):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    run.font.name = font
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), font)
    rFonts.set(qn("w:hAnsi"), font)
    rFonts.set(qn("w:eastAsia"), font)


def para(doc, text, size=11, bold=False, color=DARK, align=None,
         space_after=6, space_before=0, font=K):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        set_run(r, size=size, bold=bold or (i % 2 == 1), color=color, font=font)
    return p


def heading1(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(16)
    pf.space_after = Pt(8)
    r = p.add_run(text)
    set_run(r, size=16, bold=True, color=RED)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "B01E2E")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def heading2(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(10)
    pf.space_after = Pt(4)
    r = p.add_run(text)
    set_run(r, size=13, bold=True, color=BLUE)
    return p


def table(doc, headers, rows, widths=None, font_size=9.5):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_run(r, size=font_size, bold=True, color=WHITE)
        tcPr = hdr[i]._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), HEADER_FILL)
        tcPr.append(shd)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(val))
            set_run(r, size=font_size, color=DARK)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    return t


def quote_box(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(8)
    pf.left_indent = Pt(12)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), "B01E2E")
    pBdr.append(left)
    pPr.append(pBdr)
    r = p.add_run(text)
    set_run(r, size=11, color=DARK, font="楷体")
    return p


def bullet(doc, text, size=10.5):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.7)
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        set_run(r, size=size, bold=(i % 2 == 1), color=DARK)
    return p


# ================================================================= 文档主体
doc = Document()
sec = doc.sections[0]
sec.page_width = Cm(21)
sec.page_height = Cm(29.7)
sec.top_margin = Cm(2.4)
sec.bottom_margin = Cm(2.2)
sec.left_margin = Cm(2.5)
sec.right_margin = Cm(2.5)

st = doc.styles["Normal"]
st.font.name = K
st.font.size = Pt(11)
st.element.rPr.rFonts.set(qn("w:eastAsia"), K)



# ---- 封面 ----
para(doc, "新疆维吾尔自治区2025年政府工作报告", size=24, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=60, space_after=4)
para(doc, "深度研究与\u201c容易被忽视的细节\u201d分析", size=16, bold=True, color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
para(doc, "从\u201c丝绸之路经济带、能源产业、棉花农业与旅游兴疆\u201d重新理解新疆", size=12, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)
para(doc, "\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501", color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
para(doc, "研究对象：2025年新疆维吾尔自治区政府工作报告及同期财政、国民经济和社会发展统计资料、2026年官方复盘", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
para(doc, "标定日期：2026-08-13", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
doc.add_page_break()

# ---- 目录 ----
catalog = [
    "执行摘要",
    "一、研究口径与阅读方法",
    "二、先看新疆的特殊底盘：丝绸之路经济带核心区、能源、棉花与旅游",
    "三、最关键的宏观错位：GDP破2.14万亿、工业/进出口强，但消费/PPI偏弱",
    "四、容易被忽视、但信号很重的细节（15条）",
    "五、2025年GDP目标 vs 实际：对照表",
    "六、2025年增长是谁撑起来的？（结构归因）",
    "七、预算与财政的\u201c含金量\u201d",
    "八、民生底账：人口、收入与城乡",
    "九、城镇与农村：格局与均衡",
    "十、人口流入与流出",
    "十一、物价与货币环境",
    "十二、区域一体化：新疆在\u201c丝绸之路经济带+中欧/中巴走廊+口岸开放\u201d里的位置",
    "十三、未来5—10年最值得观察的五条主线",
    "十四、最终结论：新疆在\u201c能源+棉花+丝绸之路\u201d里的增长逻辑",
    "附录A：主要资料来源与核验路径",
    "附录B：建议建立的年度跟踪仪表盘",
]
para(doc, "目　录", size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=10, space_after=10)
for item in catalog:
    para(doc, item, size=12, space_after=4)
doc.add_page_break()

# ---- 执行摘要 ----
heading1(doc, "执行摘要")
para(doc, "**核心判断**　如果只看新闻标题，2025年新疆最显著的是\u201cGDP破2.15万亿、增长5.5%\u201d、\u201c规上工业+7.7%、有色金属+18.5%\u201d、\u201c进出口727.22亿美元、+19.2%\u201d、\u201c棉花616.5万吨（占全国92.8%）\u201d。但这份研究真正值得深读的，是这座\u201c丝绸之路经济带+能源农业\u201d的自治区，如何在社零+3.6%、PPI-5.1%背景下，靠\u201c工业（能源/纺织/有色）+外贸（中亚/一带一路）+棉花/旅游\u201d实现5.5%的增长。")
para(doc, "把2025年初设定的目标、2025年《国民经济和社会发展统计公报》、2026年复盘放在一起看，新疆呈现清晰暗线：**从\u201c能源/棉花\u201d既有底盘，向\u201c丝绸之路枢纽+能源化工+口岸开放+旅游\u201d升级**。新增新能源6482万千瓦、外贸破5000亿（人民币）是最大亮点。")
para(doc, "因此，本报告不按政府工作报告逐段复述，而采用\u201c显性表述—同期数据—制度含义—长期影响\u201d的方式，专门提取容易被忽略、但对判断新疆未来5—10年发展模式更有价值的信号。")
para(doc, "最容易记住的一句话：**新疆是\u201c丝绸之路经济带核心区+能源+棉花\u201d，靠\u201c能源/工业+进出口+农业\u201d撑起增长，靠\u201c枢纽开放+旅游+新能源\u201d升级。**观察新疆，与其只看\u201cGDP 2.15万亿\u201d，不如看\u201c外贸+19%、棉花92.8%、新能源、驰名口岸\u201d这几张名片。")

# ---- 一、研究口径与阅读方法 ----
heading1(doc, "一、研究口径与阅读方法")
heading2(doc, "1.1 本报告的核心底稿")
bullet(doc, "**2025年《政府工作报告》**（2025年1月）与2026年报告——目标与复盘。")
bullet(doc, "**《2025年新疆维吾尔自治区统计公报》**（统计局2026-03）——GDP、工业、棉花、外贸实数。")
bullet(doc, "涉进出口以美元为主（727.22亿美元）。")
heading2(doc, "1.2 阅读方法：显性—数据—长期")
para(doc, "**关键判别**：数据优先。新疆2025年GDP+5.5%、工业+7.7%、外贸+19.2%亮眼，但消费偏弱。市场\u201c工业外贸强、内需弱\u201d，穿透总量看能源/棉花/枢纽。")

# ---- 二、底盘 ----
heading1(doc, "二、先看新疆的特殊底盘：丝绸之路经济带核心区、能源、棉花、旅游")
para(doc, "新疆作为\u201c**丝绸之路经济带核心区+国家能源基地+棉花大省+旅游兴疆**\u201d，向西开放门户。")
bullet(doc, "**丝绸之路核心区**：乌鲁木齐国际陆港、喀什/霍尔果斯经济开发区、中欧班列。")
bullet(doc, "**能源**：原煤5.53亿吨、油气、新增新能源6482万千瓦、沙戈荒基地。")
bullet(doc, "**棉花/农业**：棉花616.5万吨、粮食、特色林果。")
bullet(doc, "**旅游**：2025年接待3.23亿人次、花费3700亿元（+8%）。")
para(doc, "这一底板决定新疆2025增速：**外贸/能源/棉花/旅游多轮驱动**，向西开放是其最突出的战略牌。")

# ---- 三、宏观错位 ----
heading1(doc, "三、最关键的宏观错位：GDP破2.14万亿、工业/进出口强，但消费/PPI偏弱")
para(doc, "新疆2025年最值得咀嚼的错位，是\u201c**工业/外贸/农业强、消费/PPI偏弱**\u201d。这种错位决定了对这座丝绸之路枢纽的观察不能只看GDP增速。")
bullet(doc, "**GDP**：21462.14亿、+5.5%。一产2816.27亿（+5.2%、占13.1%）、二产8189.91亿（+5.5%、占38.2%）、三产10455.95亿（+5.6%、占48.7%）。")
bullet(doc, "**工业**：规上工业+7.7%；有色冶炼+18.5%、纺织+26.1%、化学+11.2%。原煤5.53亿吨、发电5714亿千瓦时。")
bullet(doc, "**投资**：固投+7.2%；基建+15.6%、民间+12.6%。新增新能源装机6482万千瓦。")
bullet(doc, "**消费**：社零4107.82亿、+3.6%；家电+50.4%、通讯+78.4%。")
bullet(doc, "**外贸**：进出口727.22亿美元、+19.2%（出口+24.2%）、中亚/一带一路。")
bullet(doc, "**财政/物价**：财政收入+10.5%；但PPI-5.1%。")
para(doc, "**为什么读这条**：新疆作为\u201c能源+向西开放\u201d枢纽，结构性矛盾是\u201c工业/外贸/农业强、消费/价格（能源）弱\u201d。稳增长靠工业+外贸+能源。")

# ---- 四、15条细节 ----
heading1(doc, "四、容易被忽视、但信号很重的细节（15条）")
para(doc, "以下15条，都在2025年报告/公报，但被\u201cGDP 2.15万亿\u201d等掩盖。它们是判断新疆2025之后5—10年的关键小信号。")
bullet(doc, "**1. 外贸破5000亿元（人民币）、出口+24.2%**：向西开放核心成果。")
bullet(doc, "**2. 中亚五国占全国36.4%**：对中亚贸易门户。")
bullet(doc, "**3. 棉花616.5万吨、占全国92.8%**：棉花绝对主产区。")
bullet(doc, "**4. 新能源新增6482万千瓦**：沙戈荒基地、电力绿色。")
bullet(doc, "**5. 丝绸之路（中欧/中巴走廊、中亚班列）**：陆权大通道。")
bullet(doc, "**6. 纺织业+26.1%**：棉花深加工/纺织转移，潜力大。")
bullet(doc, "**7. 粮食486.5亿斤、单产全国第一**：粮食安全。")
bullet(doc, "**8. 旅游3.23亿人次、花费3700亿**：旅游兴疆成效。")
bullet(doc, "**9. 口岸开放（乌鲁木齐/霍尔果斯/喀什）**：枢纽经济。")
bullet(doc, "**10. 常住2639万（+16.2万）、城镇化61.61%**：人口/城镇化提升。")
bullet(doc, "**11. 农村收入20793元（+7%，首破2万）**：民生改善。")
bullet(doc, "**12. 对一带一路+14.7%**：多元开放。")
bullet(doc, "**13. 有色金属+18.5%、有色矿选+34.2%**：资源禀赋。")
bullet(doc, "**14. 财政+10.5%**：能源/工业税源。")
bullet(doc, "**15. 油气增储/煤化工**：未来能源深加工。")
para(doc, "", size=10, space_after=2)

# ---- 五、对照表 ----
heading1(doc, "五、2025年GDP目标 vs 实际：对照表")
para(doc, "2025年报告中列出的主要预期目标，与2025年实际（2026目标为据）：")
tb = [
    ["指标", "2026年目标", "2025年实际（公报）", "达标"],
    ["GDP增速", "5.5%-6%", "+5.5%（21462.14亿）", "达标"],
    ["规上工业增加值", "7.5%左右", "+7.7%", "达标"],
    ["固定资产投资", "8%左右", "+7.2%", "偏近"],
    ["社会消费品零售总额", "6%左右", "+3.6%", "偏低"],
    ["外贸进出口", "10%左右", "+19.4%（727.22亿美元）", "超预期"],
    ["地方财政", "10%左右", "+10.5%", "达标"],
    ["城镇新增就业", "47万人以上", "48.81万人", "达标"],
]
table(doc, tb[0], tb[1:], widths=[3.2, 3.4, 4.4, 3.6])
para(doc, "", size=10, space_after=2)
para(doc, "**要点**：GDP/工业/外贸/财政/就业达标，社零偏低——新疆\u201c生产/外贸强、内需偏弱\u201d。")

# ---- 六、结构归因 ----
heading1(doc, "六、2025年增长是谁撑起来的？（结构归因）")
heading2(doc, "6.1 工业/能源")
para(doc, "规上工业+7.7%；有色/纺织/化工/能源驱动。")
heading2(doc, "6.2 外贸/开放")
para(doc, "进出口+19.2%（出口+24.2%）、对中亚/一带一路贸易促，向西开放主动力。")
heading2(doc, "6.3 棉花/农业")
para(doc, "棉花616.5万吨、粮食单产第一；农业稳。")
heading2(doc, "6.4 能源/新能源")
para(doc, "煤炭/油气+新能源，能源产业支撑。")
heading2(doc, "6.5 旅游/消费")
para(doc, "旅游3.23亿人次、社零+3.6%（体量大、增速温和）。")
para(doc, "**一句话归因**：新疆2025年靠\u201c**工业（能源/纺织/有色）+外贸（中亚）+棉花/能源**\u201d，支撑\u201c枢纽+农业+文旅\u201d结构改善、总量+5.5%。")

# ---- 七、财政 ----
heading1(doc, "七、预算与财政的\u201c含金量\u201d")
para(doc, "**一般公共预算收入2662.8亿元、+10.5%**；税收1513.5亿、+7.3%，非税1149.3亿、+15.0%；一般公共预算支出6617.9亿元、+6.1%。")
bullet(doc, "**收入高增**：+10.5%（全国前列），主要靠能源/工业/资源税收。")
bullet(doc, "**民生/转移**：支出6617.9亿、民生优先，中央转移支持。")
bullet(doc, "**财政结构**：税收+非税双增，能源/外贸税源强。")
para(doc, "**财政含义**：新疆\u201c收入+10.5%强、财政稳健\u201d，靠能源/工业支撑，民生与基建投入大。")

# ---- 八、民生 ----
heading1(doc, "八、民生底账：人口、收入与城乡")
para(doc, "**常住人口2639.0万人（+16.2万）、城镇化率61.61%（+1.25pct）**；居民人均可支配收入32881元、+6.4%（农村+7.0%，首破2万）。")
bullet(doc, "**收入**：居民32881元、+6.4%；城镇45106元、+5.3%，农村20793元、+7.0%（城乡比缩小）。")
bullet(doc, "**就业**：城镇新增就业48.81万人；农村劳动力外务。")
bullet(doc, "**社保/民生**：低保/医保、乡村振兴、教育医疗。")
bullet(doc, "**人口**：常住+16.2万、城镇化61.61%、自然增2.70‰。")
para(doc, "**民生含义**：新疆\u201c人口/收入/城镇化提升\u201d，民生改善快、十四五良好收官。")

# ---- 九、城乡 ----
heading1(doc, "九、城镇与农村：格局与均衡")
para(doc, "**城镇化率61.61%（+1.25pct）**，新疆城镇化快速提升；城乡收入比缩小（农村快于城镇）。")
bullet(doc, "**城市**：乌鲁木齐/克拉玛依/喀什、城市群/陆港。")
bullet(doc, "**农村**：棉花/粮食/林果/畜牧，乡村振兴。")
para(doc, "**城乡均衡**：新疆\u201c城市枢纽+县域农业\u201d，靠棉花/旅游带动乡村。")

# ---- 十、人口流 ----
heading1(doc, "十、人口流入与流出")
para(doc, "**新疆常住2639万、+16.2万、自然增2.70‰**：人口正增长+净流入（较多于西部）。")
bullet(doc, "**流入**：能源/边境贸易/旅游，吸引省内外。")
bullet(doc, "**竞争**：生活/收入提升，人口向城镇/口岸集中。")
para(doc, "人口方向：新疆靠\u201c产业+枢纽+旅游\u201d留住/吸引人口，城镇化提升。")

heading1(doc, "十一、物价与货币环境")
para(doc, "**2025年新疆CPI与上年持平（0%）**、PPI-5.1%（能源/工业品价降）。")
bullet(doc, "**物价**：CPI持平、PPI-5.1%——能源价格弱、通胀温和。")
bullet(doc, "**货币/金融**：金融支持枢纽/农业。")
para(doc, "**物价含义**：新疆\u201cPPI降、CPI平\u201d，能源价格走低、内需弱，关注价格修复。")

heading1(doc, "十二、区域一体化：丝绸之路经济带核心区、中欧/中巴走廊与口岸开放")
para(doc, "新疆处于**丝绸之路经济带核心区+中欧/中巴走廊+向西开放门户**：乌鲁木齐/霍尔果斯/喀什等口岸枢纽。")
bullet(doc, "**丝绸之路**：中欧（亚）班列、乌鲁木齐国际陆港。")
bullet(doc, "**口岸/边境**：霍尔果斯/喀什/阿拉山口/塔城重点开发开放试验区。")
bullet(doc, "**基建**：新铁路/公路/机场，跨境互联互通。")
bullet(doc, "**亚欧博览会**：中国-亚欧博览会，中亚/欧合作。")
para(doc, "**区域含义**：新疆作为\u201c核心区+口岸\u201d，靠丝路/边境开放带动向西经济。")

heading1(doc, "十三、未来5—10年最值得观察的五条主线")
bullet(doc, "**主线1｜丝绸之路/口岸开放**：中欧/中亚班列、霍尔果斯/喀什。能否成向西开放核心。")
bullet(doc, "**主线2｜能源/新能源**：原煤5.53亿吨、新增新能源6482万千瓦。能否建国家能源基地。")
bullet(doc, "**主线3｜棉花/纺织**：棉花616.5万吨、纺织+26.1%。能否做深棉花产业链。")
bullet(doc, "**主线4｜旅游兴疆**：接待3.23亿人次/冰雪边境游。能否建世界旅游目的地。")
bullet(doc, "**主线5｜人口/民生**：+16万、农村收入首破2万。能否共同富裕。")
para(doc, "这五条，是新疆从\u201c能源+棉花+交通\u201d走向\u201c核心区+能源化工+旅游强\u201d的\u201c主赛道\u201d。")

heading1(doc, "十四、最终结论")
para(doc, "新疆2025年，本质上是\u201c**工业/外贸/能源+棉花、消费/PPI偏弱**\u201d的答卷：GDP21462.14亿、+5.5%、规上工业+7.7%、外贸727.22亿美元+19.2%、棉花616.5万吨、新能源+6482万千瓦、财政收入+10.5%。")
para(doc, "只要能源/外贸/棉花/旅游持续，新疆就站在\u201c枢纽+能源+开放\u201d增长极；若消费/价格受限，需承受\u201c生产强、内需弱\u201d结构。")
para(doc, "最稳观察信号：**一盯丝绸之路/口岸（开放）、二盯能源/新能源（基础）、三盯棉花/纺织（农业）、四盯旅游/消费（内需）、五盯财政/人口（质量）。**新疆，是\u201c丝路核心区+能源棉花\u201d新样本。")

heading1(doc, "附录A：主要资料来源与核验路径")
bullet(doc, "新疆维吾尔自治区2025年政府工作报告（2025年1月）——目标来源。")
bullet(doc, "《2025年新疆维吾尔自治区国民经济和社会发展统计公报》——GDP、工业、棉花、外贸实数。")
bullet(doc, "2026年新疆政府工作报告（2026年2月）——2025追认/能源/棉花/旅游。")
bullet(doc, "自治区财政厅、统计局、乌鲁木齐海关。")
heading2(doc, "核验说明")
para(doc, "本报告涉及数据以统计公报/官方口径为准；\u201c能源/棉花/丝绸之路\u201d等以官方口径为准。")

heading1(doc, "附录B：建议建立的年度跟踪仪表盘")
para(doc, "建议把下面10个\u201c测脉搏\u201d指标做成年度跟踪表：")
dash = [
    ["#", "指标", "2025基值", "为什么重要"],
    ["1", "GDP增速", "+5.5%（21462.14亿）", "总量与方向"],
    ["2", "规上工业增速", "+7.7%", "工业底盘"],
    ["3", "外贸/出口", "+19.2% / +24.2%", "开放"],
    ["4", "棉花产量", "616.5万吨（92.8%）", "农业"],
    ["5", "新能源新增装机", "6482万千瓦", "能源"],
    ["6", "固定资产投资", "+7.2%", "投资"],
    ["7", "社零增速", "+3.6%（4107.82亿）", "内需"],
    ["8", "常住人口/城镇化", "2639万 / 61.61%", "人口"],
    ["9", "财政收入", "+10.5%（2662.8亿）", "财政"],
    ["10", "CPI/PPI", "0% / -5.1%", "物价"],
]
table(doc, dash[0], dash[1:], widths=[0.8, 4.6, 4.4, 4.0])
para(doc, "把这10连起来看，外贸/能源/棉花（3/4/5）、人口/财政（8/9），都说明新疆在走向开放枢纽。")

# ============ 保存 ============
out = "/Users/x/Desktop/content-prod-lab/reports/新疆维吾尔自治区_2025年政府工作报告_深度研究_2026-08-13.docx"
doc.save(out)
print("SAVED:", out)
print("PARAGRAPHS:", len(doc.paragraphs))
print("TABLES:", len(doc.tables))
