# -*- coding: utf-8 -*-
"""Build 济宁市2025年政府工作报告 深度研究 DOCX, 参照地级市系列版式。"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DARK = RGBColor(0x1F, 0x2A, 0x44)
RED = RGBColor(0xB0, 0x1E, 0x2E)
GRAY = RGBColor(0x59, 0x60, 0x69)
BLUE = RGBColor(0x1F, 0x4E, 0x79)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
HEADER_FILL = "1F2A44"
K = "微软雅黑"


def set_run(run, size=11, bold=False, color=DARK, font=K):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    run.font.name = font
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), font)
    rFonts.set(qn("w:hAnsi"), font)
    rFonts.set(qn("w:eastAsia"), font)


def para(doc, text, size=11, bold=False, color=DARK, align=None,
         space_after=6, space_before=0, font=K):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        set_run(r, size=size, bold=bold or (i % 2 == 1), color=color, font=font)
    return p


def heading1(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(16)
    pf.space_after = Pt(8)
    r = p.add_run(text)
    set_run(r, size=16, bold=True, color=RED)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "B01E2E")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def heading2(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(10)
    pf.space_after = Pt(4)
    r = p.add_run(text)
    set_run(r, size=13, bold=True, color=BLUE)
    return p


def table(doc, headers, rows, widths=None, font_size=9.5):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_run(r, size=font_size, bold=True, color=WHITE)
        tcPr = hdr[i]._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), HEADER_FILL)
        tcPr.append(shd)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(val))
            set_run(r, size=font_size, color=DARK)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    return t


def quote_box(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(8)
    pf.left_indent = Pt(12)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), "B01E2E")
    pBdr.append(left)
    pPr.append(pBdr)
    r = p.add_run(text)
    set_run(r, size=11, color=DARK, font="楷体")
    return p


def bullet(doc, text, size=10.5):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.7)
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        set_run(r, size=size, bold=(i % 2 == 1), color=DARK)
    return p


# ================================================================= 文档主体
doc = Document()
sec = doc.sections[0]
sec.page_width = Cm(21)
sec.page_height = Cm(29.7)
sec.top_margin = Cm(2.4)
sec.bottom_margin = Cm(2.2)
sec.left_margin = Cm(2.5)
sec.right_margin = Cm(2.5)

st = doc.styles["Normal"]
st.font.name = K
st.font.size = Pt(11)
st.element.rPr.rFonts.set(qn("w:eastAsia"), K)


# ---- 封面 ----
para(doc, "遵义市2025年政府工作报告", size=24, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=60, space_after=4)
para(doc, "深度研究与\u201c容易被忽视的细节\u201d分析", size=16, bold=True, color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
para(doc, "从\u201c白酒之都、酱香茅台、装备制造、赤水河谷、长征红色旅游、辣椒/茶叶\u201d重新理解遵义", size=12, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)
para(doc, "\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501", color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
para(doc, "研究对象：2025年遵义市政府工作报告及同期财政、国民经济和社会发展统计资料、2026年官方复盘", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
para(doc, "标定日期：2026-08-14", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
doc.add_page_break()

# ---- 目录 ----
catalog = [
    "执行摘要",
    "一、研究口径与阅读方法",
    "二、先看遵义的特别底盘：白酒之都、装备制造、基础材料、红色旅游、省域副中心",
    "三、最关键的宏观错位：GDP+4.3%降档、白酒深度调整，但投资/农业/旅游稳",
    "四、容易被忽视、但信号很重的细节（15条）",
    "五、2025年GDP目标 vs 实际：对照表",
    "六、2025年增长是谁撑起来的？（结构归因）",
    "七、预算与财政的\u201c含金量\u201d",
    "八、民生底账：人口、收入与城乡",
    "九、城镇与农村：格局与均衡",
    "十、人口流入与流出",
    "十一、物价与货币环境",
    "十二、区域一体化：遵义在成渝贵、赤水河产区、西部陆海\u201c三圈\u201d里的位置",
    "十三、未来5—10年最值得观察的五条主线",
    "十四、最终结论：遵义在\u201c白酒+装备+红色\u201d里的增长逻辑",
    "附录A：主要资料来源与核验路径",
    "附录B：建议建立的年度跟踪仪表盘",
]
para(doc, "目　录", size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=10, space_after=10)
for item in catalog:
    para(doc, item, size=12, space_after=4)
doc.add_page_break()

# ---- 执行摘要 ----
heading1(doc, "执行摘要")
para(doc, "**核心判断**　2025年遵义最显著的是\u201cGDP 5206.12亿/+4.3%（2024年5027.61/+5.7%，大幅不及6%目标）、规上工业+2.7%、白酒产业深度调整\u201d、\u201c但固定资产投资892亿、民间投资+47%、进出口58亿、红色旅游+9%、民生71.8%\u201d。这说明遵义在\u201c白酒之都（茅台/酱香）+装备制造\u201d中，**白酒承压是最大拖累，工业/增收/投资走平**。")
para(doc, "把2025年目标（GDP+6%、规上+8%以上、固投+3%、社零+6%、财收+2%）、2025年实际（GDP+4.3%、规上+2.7%、财收+2.5%、城镇收入+3.7%/农村+5.9%）趋势看，遵义是\u201c白酒+白酒装备+农业\u201d路径：**白酒（1408亿）、基础材料（资源精深加工）、装备制造（航天）、能源、农业茶叶/辣椒\u201d是支柱；总量占全省22.1%（历史新高）。")
para(doc, "最容易记住的一句话：**遵义是\u201c世界酱香白酒之都（茅台）、长征文化圣地、成渝黔枢纽\u201d，靠\u201c白酒+装备制造+红色旅游\u201d驱动。**观察遵义，与其只看\u201cGDP 5206亿\u201d，不如看\u201c白酒1500亿目标、规上+2.7%（白酒深度调整）、红色旅游游客+9%、省域副中心（22.1%）\u201d。")
# ---- 一、研究口径与阅读方法 ----
heading1(doc, "一、研究口径与阅读方法")
para(doc, "本报告以\u201c遵义市政府工作报告（2025年2月）\u201d为起点，把\u201c2025年GDP目标（6%）\u201d与\u201c官方2025年（5206.12亿/+4.3%）\u201d并置对照，用\u201c2025年遵义市统计公报\u201d和\u201c2026年政府工作报告\u201d横向核验。")
para(doc, "口径提示：\u201cGDP、规上工业增速\u201d按可比价（实际）；\u201c投资、消费、进出口、财政\u201d按现价（名义）。\u201c常住人口\u201d用统计公报（约650万+）、城镇化率。")
para(doc, "指标体系（与研究口径一致）：**总量与增速、产业动能（白酒/装备/材料）、外贸、财政、民生与人口**。")
para(doc, "特别提示（不吃老本）：遵义2024年GDP 5027.61亿/+5.7%、2025年5206.12亿/+4.3%（白酒调整降档）；它不是\u201c只有茅台\u201d——**酱香白酒（千亿）、基础材料、航天装备、能源、茶叶辣椒、红色苏区\u201d才是真正底色；白酒首位产业。")
# ---- 二、先看遵义的特别底盘 ----
heading1(doc, "二、先看遵义的特别底盘：白酒之都、装备制造、基础材料、红色旅游、省域副中心")
para(doc, "遵义地处贵州北部、赤水河/乌江流域，是**世界酱香白酒之都（茅台镇/仁怀）、长征转折之城（遵义会议）、贵州省域副中心城市、全国辣椒/茶叶重要产区**；以\u201c茅台、习酒、珍酒、国台、董酒\u201d著称。2025年GDP 5206.12亿（贵州第2、占22.1%）、常住约650万、城镇化率约55%）。")
para(doc, "五个底盘名词，先立框架：")
bullet(doc, "**白酒（首位产业）**　茅台/酱香白酒千亿、2024产值1408亿（茅台镇/仁怀）、世界酱香核心区——\u201c白酒之都\u201d。")
bullet(doc, "**装备制造（航天）**　航天装备（江南航天）、先进装备480亿、特种装备——\u201c军工/制造\u201d。")
bullet(doc, "**基础材料/资源**　白酒配套、铝/材料、煤电、资源精深加工（400亿）——\u201c基础材料\u201d。")
bullet(doc, "**红色旅游/长征圣地**　遵义会议/娄山关/赤水河、红色旅游游客+9%、长征文化——\u201c红色文旅\u201d。")
bullet(doc, "**农业特色**　茶叶（湄潭/凤冈）、辣椒/蔬菜、林下经济、竹、高粱（酒用）——\u201c特色农业\u201d。")
para(doc, "这五根（白酒+装备+材料+红色+农业）构成遵义独特底盘：**左手白酒（茅台），右手军工装备+红色**。理解遵义，先理解\u201c白酒之都、遵义会议、省域副中心\u201d。")
# ---- 三、最关键的宏观错位 ----
heading1(doc, "三、最关键的宏观错位：GDP+4.3%降档、白酒深度调整，但投资/农业/旅游稳")
para(doc, "2025年遵义最需要辨析的一组\u201c错位\u201d：**GDP+4.3%（较5.7%降档）、规上工业+2.7%、白酒产业深度调整，但固定资产投资（892亿）、民间投资占比47%、农业总产值破1000亿、红色旅游+9%、进出口实绩企业+43家、财收+2.5%**。")
para(doc, "为什么\u201c白酒承压\u201d，投资农业旅游却稳？三解释：")
para(doc, "**其一，白酒深度调整、首位产业拖累**　白酒/酱香价格、茅台调整、产品结构；规上白酒承压——\u201c白酒周期\u201d。")
para(doc, "**其二，投资/工业/农业稳、多点支撑**　固投892亿（产业48.6%、工业+13.4%）、装备/材料、农业破1000亿——\u201c非酒产业补位\u201d。")
para(doc, "**其三，开放/旅游弱修复**　进出口58-88亿、红色旅游+9%，酒旅融合（酒+）——\u201c白酒+文旅\u201d。")
para(doc, "小结：遵义2025年是\u201c**白酒深度调整、投资农业旅游稳**\u201d：装备/材料/红色/农业稳，白酒/规上/进出口弱。")

# ---- 四、容易被忽视、但信号很重的细节（15条） ----
heading1(doc, "四、容易被忽视、但信号很重的细节（15条）")
bullet(doc, "**1.白酒产业深度调整（酱香）**\u201c遵义首位产业承压、茅台调整。\u201d")
bullet(doc, "**2.GDP+4.3%、占全省22.1%（历史新高）**\u201c总量稳、占比升。\u201d")
bullet(doc, "**3.规上工业+8.1%（2024）→+2.7%（2025）**\u201c白酒拖累工业降档。\u201d")
bullet(doc, "**4.固定资产投资892亿/工业投资+13.4%**\u201c产业投资强（48.6%）。\u201d")
bullet(doc, "**5.白酒产值1408亿（2024）、2025目标1500亿**\u201c白酒千亿集群。\u201d")
bullet(doc, "**6.装备制造（航天）480亿、基础材料400亿**\u201c白酒配套/军工/材料。\u201d")
bullet(doc, "**7.农业总产值破1000亿、茶叶/辣椒/高粱（酒用）**\u201c特色农业。\u201d")
bullet(doc, "**8.红色旅游游客/收入+13.2%/+14.3%（2024）**\u201c长征文旅。\u201d")
bullet(doc, "**9.进出口58亿/88亿、新增+43家**\u201c外贸稳、白酒/茶叶出口。\u201d")
bullet(doc, "**10.城镇+3.7%/农村+5.9%收入（农村快）**\u201c增收、城乡缩。\u201d")
bullet(doc, "**11.民生支出625亿（占71.8%）**\u201c民生重点。\u201d")
bullet(doc, "**12.省域副中心/遵义会议/世界酱香核心**\u201c定位。\u201d")
bullet(doc, "**13.白酒出口/贵州酱酒节、酒旅融合（酒+）**\u201c卖酒向生活方式。\u201d")
bullet(doc, "**14.仁怀经开区（国家级培育）、茅台机场口岸**\u201c开放平台。\u201d")
bullet(doc, "**15.2024 GDP 5027.61/+5.7%→2025 5206/+4.3%**\u201c亿元台阶、增速换挡。\u201d")
# ---- 五、2025年GDP目标 vs 实际：对照表 ----
heading1(doc, "五、2025年GDP目标 vs 实际：对照表")
table(doc,
    ["指标", "2025年目标", "2025年实际", "达成评价"],
    [
        ["地区生产总值(GDP)", "增长6%左右", "5206.12亿/4.3%", "差1.7pct"],
        ["规模以上工业", "增长8%以上", "+2.7%", "大幅不及"],
        ["固定资产投资", "增长3%左右", "892亿(投资稳)", "基本达标"],
        ["社会消费品零售总额", "增长6%左右", "——", "白酒调整"],
        ["一般公共预算收入", "增长2%以上", "+2.5%", "超额"],
        ["城镇/农村居民收入", "5%/7%左右", "+3.7%/+5.9%", "不及"],
        ["居民消费价格", "合理区间", "——", "稳"],
    ],
)
para(doc, "注：GDP、规上工业按可比价。**GDP（+4.3%）大幅不及6%、规上（+2.7%）低8%**（白酒深度调整）；**财收（+2.5%）超额、固投稳**。")
para(doc, "拆读：**白酒（第一位但调整）、装备/材料/农业、红色旅游、财收是支撑**；**规上（+2.7%）、白酒、进出口（58亿）**是短板——\u201c白酒周期、多点补位\u201d，是\u201c酱香之都\u201d样本。")

# ---- 六、2025年增长是谁撑起来的？（结构归因） ----
heading1(doc, "六、2025年增长是谁撑起来的？（结构归因）")
para(doc, "把遵义GDP的4.3%拆开：三次产业分别增4.4%、3.4%、5.1%（2024年一/二/三产 +3.9%/+6.5%/+5.4%，结构约14：40：46）。**第三产业+5.1%最快（红色文旅/酒旅）、第一产业+4.4%、第二产业（白酒+装备）+3.4%**。")
para(doc, "2026年遵义强调\u201c白酒稳底、重点产业突破、省域副中心、酒旅融合\u201d，聚焦**白酒（1512亿目标）、装备制造、基础材料、酱香/赤水河、红色文旅、茶叶辣椒**——核心是\u201c白酒稳-非酒求\u201d。")
para(doc, "**第二产业（工业/白酒/装备）**：规上+2.7%（白酒深度调整、航天装备/材料补位）——\u201c白酒调、配套强\u201d。")
para(doc, "**第三产业（服务业/文旅）**：+5.1%（红色旅游、酒旅、商贸、物流）——\u201c服务+红色/酒旅\u201d。")
para(doc, "**第一产业（农业）**：+4.4%（茶叶/辣椒/高粱/林下、总产值破1000亿）——\u201c特色农业稳\u201d。")
para(doc, "一句话归因：**2025年遵义增长\u201c靠第三产业（红色/酒旅）+农业+装备材料\u201d**，白酒拖累工业/GDP；\u201c白酒稳、产业突破\u201d是主线。")

# ---- 七、预算与财政的\u201c含金量\u201d ----
heading1(doc, "七、预算与财政的\u201c含金量\u201d")
para(doc, "2025年遵义**一般公共预算收入+2.5%**；民生支出625亿（占71.8%）；一般公共预算支出以民生为主。")
bullet(doc, "财收+2.5%（超额于GDP），白酒税收占比高——\u201c白酒税、含金量足\u201d。")
bullet(doc, "民生支出占71.8%（教育/社保/医疗）。")
bullet(doc, "金融：存贷款/撬动、争取资金1012亿（2024）——信贷+政策支持。")
para(doc, "**财政含金量小结**：财收+2.5%（稳/超）、民生71.8%、靠白酒税；财政对\u201c白酒稳、产业、红色、民生\u201d投入加大。")

# ---- 八、民生底账：人口、收入与城乡 ----
heading1(doc, "八、民生底账：人口、收入与城乡")
para(doc, "2025年遵义**城镇/农村居民人均可支配收入分别+3.7%/+5.9%**（农村快于城镇、城乡差距缩小）；2024年城镇+4.4%/农村+7.1%。就业：城镇新增就业约10.5万（2024）。")
para(doc, "人口画像：**常住约650万、城镇化率约58%**；红色/白酒/农业吸纳、劳动力转移。")
para(doc, "民生投入：民生支出625亿（71.8%）、教育医疗、就业、农村脱贫巩固——民生扎实、红色福祉。")

# ---- 九、城镇与农村：格局与均衡 ----
heading1(doc, "九、城镇与农村：格局与均衡")
para(doc, "遵义城镇化率约55-58%；县域经济强（仁怀白酒/赤水竹/湄潭茶/凤冈茶）；农村收入增速（+5.9%）>城镇（+3.7%），**城乡差距缩小**；乡村振兴、林下经济。")
para(doc, "农业底盘：**农业总产值破1000亿、茶叶/辣椒/高粱（酒用）/竹/林下、湄潭**——\u201c茶叶辣椒+酒用高粱\u201d。")
para(doc, "一句话：\u201c遵义是白酒+山地农业大市、县域特色经济强、乡村振兴\u201d。")

# ---- 十、人口流入与流出 ----
heading1(doc, "十、人口流入与流出")
para(doc, "遵义常住约650万（贵州人口大市）、城镇化约65%；\u201c白酒/农业\u201d吸纳，但部分青年外流（贵州/珠三角）；赤水河/遵义会议红色引客。")
para(doc, "结构观察：**山区农业县人口、外出务工、白酒从业（仁怀）**；仁怀/茅台吸附技术、营销。")
para(doc, "2026年目标：稳就业、育才——遵义靠\u201c白酒+红色+农业\u201d聚人。")

# ---- 十一、物价与货币环境 ----
heading1(doc, "十一、物价与货币环境")
para(doc, "2025年遵义**CPI在合理区间（低通胀）**、PPI/白酒价格调整——\u201c低通胀、需求弱\u201d。")
bullet(doc, "信贷扩张：存贷款/政策性资金（力争1012亿）——宽信用支持白酒/农业/产业。")
bullet(doc, "消费：白酒/酒旅/红色、以旧换新——消费弱修复。")
para(doc, "货币环境判断：**宽信用、CPI低**；遵义靠\u201c白酒+红色旅游\u201d稳需求（2026 CPI合理）。")

# ---- 十二、区域一体化：遵义的位置 ----
heading1(doc, "十二、区域一体化：遵义在成渝贵、赤水河产区、西部陆海\u201c三圈\u201d里的位置")
para(doc, "遵义是**贵州省域副中心城市、成渝贵经济带/连接成渝与贵阳、世界酱香白酒（赤水河谷产区）、西部陆海新通道北横节点**。")
bullet(doc, "**赤水河酱香产区**　仁怀/习水/赤水、茅台镇——世界酱香型白酒核心区。")
bullet(doc, "**成渝贵**　贵州副中心、接重庆/贵阳、渝贵铁路/赤水高速——成渝贵枢纽。")
bullet(doc, "**西部陆海新通道**　遵义综合交通、铁海联运、白酒/茶叶出口、综保区。")
para(doc, "一句话：**遵义在\u201c赤水河+成渝贵+西部陆海\u201d里，最核心是\u201c白酒产区+红色枢纽\u201d**；酱香白酒、区位、红色是最大优势。")

# ---- 十三、未来5—10年最值得观察的五条主线 ----
heading1(doc, "十三、未来5—10年最值得观察的五条主线")
bullet(doc, "**主线一：白酒/酱香（茅台/世界核心区）**\u201c白酒1512亿、茅台/酱香\u201d能否度过周期、稳中求进。")
bullet(doc, "**主线二：装备制造（航天/军工）**\u201c航天装备、白酒配套、特种装备\u201d补位。")
bullet(doc, "**主线三：红色文旅/酒旅融合**\u201c长征圣地、酒+生活\u201d文旅新生。")
bullet(doc, "**主线四：基础材料/能源/赤水河生态**\u201c白酒生态、资源精深加工\u201d绿色。")
bullet(doc, "**主线五：省域副中心/农业/人口**\u201cGDP 7000亿（2030）、茶叶辣椒、聚人\u201d。")

# ---- 十四、最终结论 ----
heading1(doc, "十四、最终结论：遵义在\u201c白酒+装备+红色\u201d里的增长逻辑")
para(doc, "把2025年闭环打穿：**遵义是\u201c世界酱香白酒之都、长征圣地、省域副中心\u201d**：GDP 5206.12亿/+4.3%、规上+2.7%、白酒1408亿、占全省22.1%。")
para(doc, "遵义不是\u201c只有茅台\u201d——它是**白酒+航天装备+基础材料+茶叶农业+红色文旅**的复合，靠\u201c白酒稳-非酒求\u201d驱动；但白酒/规上/进出口承压。")
para(doc, "一句话结论：**遵义是\u201c酱香之都、长征之城、成渝贵枢纽\u201d；观察它先看\u201c白酒产值、装备制造、红色文旅、酒旅融合\u201d，再看\u201c规上、进出口、白酒周期\u201d。**它是\u201c首位产业深度调整、其他多点支撑\u201d的贵州样本。")

# ---- 附录A ----
heading1(doc, "附录A：主要资料来源与核验路径")
bullet(doc, "《2025年遵义市政府工作报告》（2025年2月，2025年目标、2024年回顾+5.7%）")
bullet(doc, "《2025年遵义市国民经济和社会发展统计公报》（遵义市统计局，2026-05，2025年实际+4.3%）")
bullet(doc, "《2026年遵义市政府工作报告》（2026年2月，复盘+2026年目标）")
bullet(doc, "遵义市人民政府/统计局（zunyi.gov.cn）、遵义网")

# ---- 附录B：建议建立的年度跟踪仪表盘 ----
heading1(doc, "附录B：建议建立的年度跟踪仪表盘")
bullet(doc, "GDP总量/增速 vs 全年目标。")
bullet(doc, "规上工业（白酒/装备）增速。")
bullet(doc, "白酒产值/茅台/酱香周期。")
bullet(doc, "固定资产/工业/民间投资。")
bullet(doc, "红色旅游/酒旅/酒+。")
bullet(doc, "进出口/白酒出口。")
bullet(doc, "茶叶/辣椒/农业/高粱。")
bullet(doc, "财收/税收/民生%。")
bullet(doc, "常住/城镇化/劳动力转移。")
bullet(doc, "CPI/存贷款/综保区。")

# ---------------- 保存 ----------------
out = "/Users/x/Desktop/content-prod-lab/reports/遵义市_2025年政府工作报告_深度研究_2026-08-14.docx"
doc.save(out)
print("SAVED OK: 遵义市", out)
