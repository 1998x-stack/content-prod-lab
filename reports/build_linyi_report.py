# -*- coding: utf-8 -*-
"""Build 临沂市2025年政府工作报告 深度研究 DOCX, 参照地级市系列版式。"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DARK = RGBColor(0x1F, 0x2A, 0x44)
RED = RGBColor(0xB0, 0x1E, 0x2E)
GRAY = RGBColor(0x59, 0x60, 0x69)
BLUE = RGBColor(0x1F, 0x4E, 0x79)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
HEADER_FILL = "1F2A44"
K = "微软雅黑"


def set_run(run, size=11, bold=False, color=DARK, font=K):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    run.font.name = font
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), font)
    rFonts.set(qn("w:hAnsi"), font)
    rFonts.set(qn("w:eastAsia"), font)


def para(doc, text, size=11, bold=False, color=DARK, align=None,
         space_after=6, space_before=0, font=K):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        set_run(r, size=size, bold=bold or (i % 2 == 1), color=color, font=font)
    return p


def heading1(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(16)
    pf.space_after = Pt(8)
    r = p.add_run(text)
    set_run(r, size=16, bold=True, color=RED)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "B01E2E")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def heading2(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(10)
    pf.space_after = Pt(4)
    r = p.add_run(text)
    set_run(r, size=13, bold=True, color=BLUE)
    return p


def table(doc, headers, rows, widths=None, font_size=9.5):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_run(r, size=font_size, bold=True, color=WHITE)
        tcPr = hdr[i]._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), HEADER_FILL)
        tcPr.append(shd)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(val))
            set_run(r, size=font_size, color=DARK)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    return t


def quote_box(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(8)
    pf.left_indent = Pt(12)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), "B01E2E")
    pBdr.append(left)
    pPr.append(pBdr)
    r = p.add_run(text)
    set_run(r, size=11, color=DARK, font="楷体")
    return p


def bullet(doc, text, size=10.5):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.7)
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        set_run(r, size=size, bold=(i % 2 == 1), color=DARK)
    return p


# ================================================================= 文档主体
doc = Document()
sec = doc.sections[0]
sec.page_width = Cm(21)
sec.page_height = Cm(29.7)
sec.top_margin = Cm(2.4)
sec.bottom_margin = Cm(2.2)
sec.left_margin = Cm(2.5)
sec.right_margin = Cm(2.5)

st = doc.styles["Normal"]
st.font.name = K
st.font.size = Pt(11)
st.element.rPr.rFonts.set(qn("w:eastAsia"), K)


# ---- 封面 ----
para(doc, "临沂市2025年政府工作报告", size=24, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=60, space_after=4)
para(doc, "深度研究与\u201c容易被忽视的细节\u201d分析", size=16, bold=True, color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
para(doc, "从\u201c商贸物流名城、沂蒙老区、批发市场、五金与产业升级\u201d重新理解临沂", size=12, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)
para(doc, "\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501", color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
para(doc, "研究对象：2025年临沂市政府工作报告及同期财政、国民经济和社会发展统计资料、2026年官方复盘", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
para(doc, "标定日期：2026-08-13", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
doc.add_page_break()

# ---- 目录 ----
catalog = [
    "执行摘要",
    "一、研究口径与阅读方法",
    "二、先看临沂的特殊底盘：商贸物流城、沂蒙老区、八大传统产业与电子商贸",
    "三、最关键的宏观错位：GDP破6800亿但增速回落，工业/出口强，投资/地产弱",
    "四、容易被忽视、但信号很重的细节（15条）",
    "五、2025年GDP目标 vs 实际：对照表",
    "六、2025年增长是谁撑起来的？（结构归因）",
    "七、预算与财政的\u201c含金量\u201d",
    "八、民生底账：人口、收入与城乡",
    "九、城镇与农村：格局与均衡",
    "十、人口流入与流出",
    "十一、物价与货币环境",
    "十二、区域一体化：临沂在\u201c临沂都市圈+鲁南经济圈‘一带一路’物流\u201d里的位置",
    "十三、未来5—10年最值得观察的五条主线",
    "十四、最终结论：临沂在\u201c商贸物流+先进制造+沂蒙振兴\u201d里的增长逻辑",
    "附录A：主要资料来源与核验路径",
    "附录B：建议建立的年度跟踪仪表盘",
]
para(doc, "目　录", size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=10, space_after=10)
for item in catalog:
    para(doc, item, size=12, space_after=4)
doc.add_page_break()

# ---- 执行摘要 ----
heading1(doc, "执行摘要")
para(doc, "**核心判断**　2025年临沂最显著的是\u201cGDP 6862.2亿元、增长5.4%（低于5.5%目标/低于2024的5.7%）、三产占52.2%\u201d、\u201c规上工业+7.7%、进出口+4.9%\u201d、\u201c社会消费品零售3877.5亿/+5.2%\u201d、\u201c一般公共预算收入482.6亿/+3%\u201d。这说明临沂在\u201c商贸物流+先进制造+沂蒙老区\u201d下稳中向好，但**投资+0.4%、地产走弱、税收-1%**是短板。")
para(doc, "把2025年目标（GDP+5.5%/力争6%、规上+8%/力争10%、固投+6%、社零+6%、财政+3.5%）、2025年统计、2026年前瞻一起看，临沂是\u201c商贸物流名城+先进制造\u201d路径：**八大产业（13产业链6851亿）+商圈物流+电商/出口**是引擎。GDP居山东第5、全国人口大市。")
para(doc, "最容易记住的一句话：**临沂是\u201c商贸物流之都+沂蒙老区\u201d的山东大市，靠\u201c批发市场/物流（超万亿）+工业（装备/粮油/建材）+电商出口\u201d增长。**观察临沂，与其只看\u201cGDP 6862亿\u201d，不如看\u201c规上工业+7.7%、新能源商用车+27.9%、物流超万亿、市场采购出口914亿（全省第1）、物资网络863.9亿\u201d。")# ---- 一、研究口径与阅读方法 ----
heading1(doc, "一、研究口径与阅读方法")
para(doc, "本报告以\u201c临沂市2025年政府工作报告（2025年，张宝亮作）\u201d为起点，把\u201c2025年GDP目标（5.5%、力争6%）\u201d与\u201c2025年经济社会发展统计公报实际值（6862.2亿元/+5.4%）\u201d并置对照，再用2026年政府工作报告的复盘作事后验证。")
para(doc, "重点阅读三类证据：**一是指标目标是否达标**；**二是结构（谁贡献增长）是否匹配目标叙事**；**三是官方复盘如何自评、承认问题**。统计口径一般公共预算收入为地方级。")
para(doc, "统一采用\u201c同比增长%\u201d（GDP为不变价），财政与居民收入多为名义值，文中按原口径转录、不逐项换算。")

# ---- 二、特殊底盘 ----
heading1(doc, "二、先看临沂的特殊底盘：商贸物流城、沂蒙老区、八大传统产业与电子商贸")
para(doc, "**区位与身份**：临沂是山东省地级市、人口大市（常住约1100万），\u201c商贸物流名城\u201d\u201c沂蒙老区\u201d、鲁南经济圈中心城市；沂蒙精神发源地。")
para(doc, "**产业底盘**：一是商贸物流（物流总额破万亿、专业批发市场全国知名、TIR全国第1）；二是八大传统产业（板材/五金/陶瓷/粮油建材）+13产业链6851亿；三是先进制造（新能源商用车+27.9%/智能农机+11.2%/高端工程装备+10.1%）；四是电商/出口（市场采购出口914亿全省第1、中欧班列506列、网络863.9亿）。")
para(doc, "**人口底盘**：常住约1100万（七普1101.8万），山东人口大市、临沂都市圈核心。")
para(doc, "**市场与出口**：2025年社零3877.5亿/+5.2%；进出口1775.8亿/+4.9%（出口+5.2%、进口+2.2%）。")

# ---- 三、核心宏观错位 ----
heading1(doc, "三、最关键的宏观错位：GDP破6800亿但增速回落，工业/出口强，投资/地产弱")
para(doc, "**第一组错位**：2025年GDP目标5.5%、力争6%；实际6862.2亿/+5.4%**略低于目标**、也低于2024年的5.7%；但破6800亿、山东第5。")
para(doc, "**第二组错位**：规上工业+7.7%（目标8%略低）、进出口+4.9%；但**固投+0.4%（目标6%严重未达标）、地产-12.5%、税收-1%**。\u201c工业/出口强、投资弱\u201d。")
para(doc, "**第三组错位**：社零+5.2%（目标6%略低）内需尚可；但**乡村零售+7.3%、网络零售863.9亿**亮点。")
para(doc, "一句话：**临沂是\u201c商贸+制造强、投资/地产弱\u201d的沂蒙大市**——靠批发物流与工业，投资是短板。")

# ---- 四、15条细节 ----
heading1(doc, "四、容易被忽视、但信号很重的细节（15条）")
bullet(doc, "1. **GDP 6862.2亿/+5.4%**：破6800亿、山东第5、\u201c十四五\u201d收官。")
bullet(doc, "2. **规上工业+7.7%**：连续居全省前5，工业是稳定器。")
bullet(doc, "3. **新能源商用车+27.9%**：新动能放量。")
bullet(doc, "4. **智能农机+11.2%、高端工程装备+10.1%**：装备升级。")
bullet(doc, "5. **八大传统产业（板材/五金/建材/粮油）基本盘**：传统产业基本盘。")
bullet(doc, "6. **物流超万亿、TIR全国第1、中欧班列506列**：物流枢纽。")
bullet(doc, "7. **市场采购出口914.4亿（全省第1）**：外贸规模。")
bullet(doc, "8. **网络零售863.9亿/+7.4%、快递28.2亿件/+14.1%**：电商物流。")
bullet(doc, "9. **固投+0.4%（工业+19.3%）**：工业投强、整体弱。")
bullet(doc, "10. **地产-12.5%、销售面积-7.5%**：地产调整。")
bullet(doc, "11. **常住约1100万（七普口径）**：人口大市。")
bullet(doc, "12. **收入：城镇51513元/+4.5%、农村23127元/+5.7%**，城乡2.23。")
bullet(doc, "13. **财政482.6亿/+3%（税收-1%）**：达标但税收转弱。")
bullet(doc, "14. **规上企业4976家、新入库553家+112.7%**：企业活力。")
bullet(doc, "15. **新能源装机突破1000万千瓦**：绿色。")# ---- 五、目标 vs 实际对照表 ----
heading1(doc, "五、2025年GDP目标 vs 实际：对照表")
table(doc,
    ["指标", "2025年目标", "2025年实际", "是否达标"],
    [
        ["地区生产总值增速", "5.5%/力争6%", "6862.2亿元/+5.4%", "略低于目标"],
        ["规上工业增加值", "+8%/力争10%", "+7.7%", "略低于目标"],
        ["固定资产投资", "+6%/力争8%", "+0.4%", "严重未达标"],
        ["社会消费品零售总额", "+6%/力争8%", "3877.5亿元/+5.2%", "略低目标"],
        ["一般公共预算收入", "+3.5%左右", "482.6亿元/+3%", "略低目标"],
        ["居民人均可支配收入", "+5.5%/力争6%", "全体+5%", "略低目标"],
        ["进出口", "稳中提质", "1775.8亿元/+4.9%", "接近"],
    ],
    widths=[3.4, 2.5, 5.0, 3.1])
para(doc, "**简评**：临沂2025年**所有硬指标略低于目标**（GDP 5.4%/目标5.5、规上7.7%/8%、投资0.4%/6%、消费5.2%/6%）。工业/出口/商贸是亮点，**投资、税收**是明显短板。")

# ---- 六、结构归因 ----
heading1(doc, "六、2025年增长是谁撑起来的？（结构归因）")
para(doc, "**三大产业**：一产555.4亿/+4.2%、二产2726.4亿/+5.2%、三产3580.4亿/+5.6%；三产占52.2%。增长以三产（商贸/物流）+二产（工业）双支撑。")
para(doc, "**工业**：规上工业+7.7%、13产业链6851.3亿/+5.5%；**新能源商用车+27.9%、粮油加工+17.3%、智能农机+11.2%、绿色果蔬+10.7%、工程装备+10.1%**；高端不锈钢-2.7%。规上企业4976家。")
para(doc, "**服务业/商贸**：批发零售/物流（物流超万亿）、电商/快递（28.2亿件+14.1%）、市场采购。")
para(doc, "**增长归因**：临沂GDP主要靠**三产（商贸/物流/电商）+二产（工业/装备）+出口**；投资、地产贡献弱。")

# ---- 七、财政 ----
heading1(doc, "七、预算与财政的\u201c含金量\u201d")
para(doc, "2025年一般公共预算收入482.6亿元/+3%（税收313.9亿/-1%、占65.1%）；支出1011亿/+3.5%。")
para(doc, "**结构性**：收入\u201c达标（+3%）但税收-1%\u201d，靠非税支撑扩张；支出高增（+3.5%）保民生。")
para(doc, "**含金量**：临沂财政\u201c增收不增税\u201d，税收负增长反映制造业利润与地产走弱；民生支出与转移支付依赖度高。")

# ---------------- 八、民生 ----------------
heading1(doc, "八、民生底账：人口、收入与城乡")
para(doc, "常住人口约1100万（七普1101.8万）、山东人口大市。收入：**城镇51513元/+4.5%、农村23127元/+5.7%**，城乡比2.23。消费支出全体19383元/+4.5%（城镇22989/农村14690）。")
para(doc, "就业：城镇新增就业10.2万；粮食86.3亿斤（6连增）。")
para(doc, "**民生结论**：收入农村快于城镇、城乡差距2.23中等；人口大市、就业稳、粮食安全。")

# ---------------- 九、 ----------------
heading1(doc, "九、城镇与农村：格局与均衡")
para(doc, "临沂城镇化率约60%左右（山东非前列），中心城区（兰山/罗庄/河东）+商贸物流城；县域（沂南/沂水/兰陵/费县/莒南）人口多、乡镇经济（工业强镇/农业县）。")
para(doc, "城乡收入比2.23，农村收入增速快、边际收敛；乡村振兴、四雁人才支撑。")

# ---------------- 十、 ----------------
heading1(doc, "十、人口流入与流出")
para(doc, "临沂常住约1100万（人口输出型，外出务工多回流），临沂商城吸纳省内/鲁南人口；近年物流/电商从业回流。")
para(doc, "**流入**：物流/商贸就业、返乡创业；**流出**：青壮年外出（京津冀/长三角）。总体\u201c劳务大市、部分回流\u201d。")

# ---------------- 十一、 ----------------
heading1(doc, "十一、物价与货币环境")
para(doc, "2025年临沂CPI（公报未详列、据了解温和微增）；消费端通胀低。")
para(doc, "金融：存款13059.6亿（比年初+1252.1亿）、贷款11269.4亿（+359.3亿）；住户存款+978亿。\u201c宽中稳\u201d。")

# ---------------- 十二、 ----------------
heading1(doc, "十二、区域一体化：临沂在\u201c临沂都市圈+鲁南经济圈+‘一带一路’物流\u201d里的位置")
para(doc, "临沂是山东临沂都市圈核心、鲁南经济圈中心城市，对接长三角/京津冀；依托商贸物流（物流超万亿、TIR全国第1、中欧班列506列）成为\u201c一带一路\u201d物流枢纽。")
para(doc, "开放：市场采购出口914亿全省第1、跨境电商/海外仓、中欧班列。临沂是\u201c商贸+物流+开放\u201d枢纽城市。")

# ---------------- 十三、 ----------------
heading1(doc, "十三、未来5—10年最值得观察的五条主线")
bullet(doc, "1. **商贸物流升级**：物流超万亿、TIR、跨境电商、商城数字化。")
bullet(doc, "2. **先进制造/新能源**：新能源商用车+27.9%、高端装备、智能农机。")
bullet(doc, "3. **八大产业转型/板材五金**：板材/建材/家居升级、绿色制造。")
bullet(doc, "4. **沂蒙振兴/人口**：老区经济、乡村产业、人口回流。")
bullet(doc, "5. **投资与财政修复**：固投+0.4%、税收-1%后的投资/税收改善。")

# ---------------- 十四、 ----------------
heading1(doc, "十四、最终结论：临沂在\u201c商贸物流+先进制造+沂蒙振兴\u201d里的增长逻辑")
para(doc, "**结论**：临沂2025年的\u201c真相\u201d是——**GDP+5.4%（略低目标）、规上+7.7%、进出口+4.9%、固投+0.4%、财政+3%（税收-1%）**。它是\u201c商贸物流+工业+出口\u201d驱动、投资走弱的沂蒙大市。")
para(doc, "**对趋势判断**：物流/商贸/先进制造/新能源代表临沂的\u201c动能\u201d，投资/税收代表\u201c约束\u201d。**商贸升级+制造转型**决定潜力，**投资修复+人口回流**决定韧性。")
para(doc, "**若只看一个指标**：看**工业投资增速（+19.3%）与商贸物流总额（超万亿）**——临沂靠\u201c商城+物流\u201d撬动\u201c工业+电商\u201d，制造与商贸的协同（地产品率43%）是未来增长的关键。")

# ---------------- 附录A ----------------
heading1(doc, "附录A：主要资料来源与核验路径")
bullet(doc, "临沂市人民政府《2025年临沂市政府工作报告》（2025年1月，张宝亮）。")
bullet(doc, "临沂市统计局《2025年临沂市经济社会发展统计公报》（2026年3月）。")
bullet(doc, "《2026年临沂市政府工作报告》（2026年1月）。")
bullet(doc, "山东省2025年统计公报与临沂统计年鉴/七普交叉核验。")

# ---------------- 附录B ----------------
heading1(doc, "附录B：建议建立的年度跟踪仪表盘")
bullet(doc, "GDP总量/增速 vs 全年目标（+/-个百分点）。")
bullet(doc, "规上工业增加值及分行业（装备/粮油/建材/新能源商用车）增速。")
bullet(doc, "固定资产投资（总量/工业/基础设施/房地产）增速。")
bullet(doc, "社会消费品零售总额、网络零售、市场采购。")
bullet(doc, "进出口、出口、中欧班列、跨境电商。")
bullet(doc, "一般公共预算收入/税收/非税、财政支出、民生占比。")
bullet(doc, "常住人口、城镇化率、公共市场经营主体。")
bullet(doc, "CPI/核心CPI、规上工业利润。")
bullet(doc, "物流总额、TIR车次、快递量。")
bullet(doc, "临沂商城交易额、地产品率。")

# ---------------- 保存 ----------------
out = "/Users/x/Desktop/content-prod-lab/reports/临沂市_2025年政府工作报告_深度研究_2026-08-13.docx"
doc.save(out)
print("SAVED OK 临沂市_2025年政府工作报告_深度研究_2026-08-13.docx")