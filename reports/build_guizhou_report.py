# -*- coding: utf-8 -*-
"""Build 贵州省2025年政府工作报告 深度研究 DOCX, 参照省市系列版式。"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DARK = RGBColor(0x1F, 0x2A, 0x44)
RED = RGBColor(0xB0, 0x1E, 0x2E)
GRAY = RGBColor(0x59, 0x60, 0x69)
BLUE = RGBColor(0x1F, 0x4E, 0x79)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
HEADER_FILL = "1F2A44"
K = "微软雅黑"


def set_run(run, size=11, bold=False, color=DARK, font=K):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    run.font.name = font
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), font)
    rFonts.set(qn("w:hAnsi"), font)
    rFonts.set(qn("w:eastAsia"), font)


def para(doc, text, size=11, bold=False, color=DARK, align=None,
         space_after=6, space_before=0, font=K):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        set_run(r, size=size, bold=bold or (i % 2 == 1), color=color, font=font)
    return p


def heading1(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(16)
    pf.space_after = Pt(8)
    r = p.add_run(text)
    set_run(r, size=16, bold=True, color=RED)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "B01E2E")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def heading2(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(10)
    pf.space_after = Pt(4)
    r = p.add_run(text)
    set_run(r, size=13, bold=True, color=BLUE)
    return p


def table(doc, headers, rows, widths=None, font_size=9.5):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_run(r, size=font_size, bold=True, color=WHITE)
        tcPr = hdr[i]._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), HEADER_FILL)
        tcPr.append(shd)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(val))
            set_run(r, size=font_size, color=DARK)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    return t


def quote_box(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(8)
    pf.left_indent = Pt(12)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), "B01E2E")
    pBdr.append(left)
    pPr.append(pBdr)
    r = p.add_run(text)
    set_run(r, size=11, color=DARK, font="楷体")
    return p


def bullet(doc, text, size=10.5):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.7)
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        set_run(r, size=size, bold=(i % 2 == 1), color=DARK)
    return p


# ================================================================= 文档主体
doc = Document()
sec = doc.sections[0]
sec.page_width = Cm(21)
sec.page_height = Cm(29.7)
sec.top_margin = Cm(2.4)
sec.bottom_margin = Cm(2.2)
sec.left_margin = Cm(2.5)
sec.right_margin = Cm(2.5)

st = doc.styles["Normal"]
st.font.name = K
st.font.size = Pt(11)
st.element.rPr.rFonts.set(qn("w:eastAsia"), K)



# ---- 封面 ----
para(doc, "贵州省2025年政府工作报告", size=24, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=60, space_after=4)
para(doc, "深度研究与“容易被忽视的细节”分析", size=16, bold=True, color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
para(doc, "从“大数据/算力、白酒、旅游、扶贫成果与‘六大产业基地’新型工业化”重新理解贵州", size=12, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)
para(doc, "━━━━━━━━━━━━━━━━", color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
para(doc, "研究对象：2025年贵州省政府工作报告及同期财政、国民经济和社会发展统计资料、2026年官方复盘", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
para(doc, "标定日期：2026-08-13", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
doc.add_page_break()

# ---- 目录 ----
catalog = [
    "执行摘要",
    "一、研究口径与阅读方法",
    "二、先看贵州的特殊底盘：大数据/算力、白酒、旅游、扶贫成果与‘六大优势’",
    "三、最关键的宏观错位：GDP破2.35万亿、工业/出口偏强，但投资/消费/物价偏弱",
    "四、容易被忽视、但信号很重的细节（15条）",
    "五、2025年GDP目标 vs 实际：对照表",
    "六、2025年增长是谁撑起来的？（结构归因）",
    "七、预算与财政的“含金量”",
    "八、民生底账：人口、收入与城乡",
    "九、城镇与农村：格局与均衡",
    "十、人口流入与流出",
    "十一、物价与货币环境",
    "十二、区域一体化：贵州在“西部陆海通道+黔中城市群+大数据枢纽”里的位置",
    "十三、未来5—10年最值得观察的五条主线",
    "十四、最终结论：贵州在“大数据+新型工业化+旅游+白酒生态”里的增长逻辑",
    "附录A：主要资料来源与核验路径",
    "附录B：建议建立的年度跟踪仪表盘",
]
para(doc, "目　录", size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=10, space_after=10)
for item in catalog:
    para(doc, item, size=12, space_after=4)
doc.add_page_break()

# ---- 执行摘要 ----
heading1(doc, "执行摘要")
para(doc, "**核心判断**　如果只看新闻标题，2025年贵州最显眼的是“GDP破2.35万亿、增长4.9%”、“规上工业+7.0%”、“汽车产量+94.9%（新能源车+147.7%）”和“数字经济/算力全国领先”。但这份研究真正值得深读的，是这座“大数据枢纽+算力高地+白酒/旅游+新工业化”的山地省份，如何在固定资产投资（-3.7%）、房地产开发（-1.6%）、社会消费品零售（+1.7%）与CPI（-0.1%）偏弱的背景下，靠“汽车/电子信息制造+算力/大数据+白酒/旅游”稳住增长。")
para(doc, "把2025年初设定的目标（GDP增长5.5%左右）、2025年《国民经济和社会发展统计公报》、以及2026年报告对2025年的复盘放在一起看，贵州呈现清晰暗线：**从“白酒/资源/农业”的旧底盘，向“大数据/算力+新型工业化+旅游/绿色+互利共赢”转型**。旧引擎（白酒、传统能源、房地产、一般基建）在调整；新引擎（汽车、电子信息、算力/大数据、工业投资、旅游）被要求更快补位。这也是“在脱贫基础上、面向‘十五五’转型”的西部样本省。")
para(doc, "因此，本报告不按政府工作报告逐段复述，而采用“显性表述—同期数据—制度含义—长期影响”的方式，专门提取容易被忽视、但对判断贵州未来5—10年发展模式更有价值的信号。")
para(doc, "最容易记住的一句话：**贵州是“大数据/算力高地+‘六大优势’新型工业化+白酒/旅游/绿色能源”的西部样本，靠“工业+数字经济+旅游+白酒”撑起增长。**观察贵州，与其看“GDP 2.35万亿”，不如看“算力全国领先、汽车/电子信息制造、白酒、旅游、绿色能源与‘六大优势’”这几张名片。")
heading2(doc, "一页速览：2025年贵州经济的“表与里”")
table(doc,
    ["维度", "表面现象", "更深层信号"],
    [
        ["增长", "GDP 23562.17亿、+4.9%", "一产12.8%、二产31.0%、三产56.2%"],
        ["产业", "规上工业+7.0%", "汽车+42.3%、电子信息+35.9%"],
        ["外贸", "进出口848.36亿、-1.2%", "出口-10.4%、进口+17.7%"],
        ["投资", "固定资产投资-3.7%", "地产-1.6%、民间+2.1%"],
        ["财政", "一般公共预算收入2223.30亿、+2.5%", "税收+1.4%"],
        ["消费", "社零+1.7%", "餐饮-6.6%、通讯+20.3%"],
        ["人口", "常住3857万、城镇化率57.66%", "人口-3万、城镇新增就业62.6万"],
        ["开放", "数字经济全国领先/算力超160Eflops", "工业/算力双轮"],
    ],
    widths=[2.2, 5.2, 8.6])
para(doc, "", size=4, space_after=2)

# ---- 一、研究口径与阅读方法 ----
heading1(doc, "一、研究口径与阅读方法")
heading2(doc, "1.1 本报告的三份核心底稿")
bullet(doc, "**2025年《政府工作报告》**：2025年1月19日在省十四届人大三次会议上作，给出全年预期目标（GDP增长5.5%左右、规上工业7.5%以上、固投4%以上、社零5%以上、CPI涨幅2%左右）。")
bullet(doc, "**2025年《国民经济和社会发展统计公报》**：省统计局2026年4月15日发布，给出全年实际执行数与增速，是“事后验证”的锚点。")
bullet(doc, "**2026年贵州省政府工作报告及官方复盘**：对2025执行结果的追认，交叉验证“目标—执行—再评价”三段式。")
heading2(doc, "1.2 阅读方法：显性—数据—制度—长期")
para(doc, "每一章按“**显性表述 → 同期数据 → 制度含义 → 长期影响**”四层展开。其中“制度含义”回答“政府为什么这么排优先序”，“长期影响”回答“这对未来5—10年意味着什么”。")
para(doc, "关键判别：**当报告使用的“增长词”和统计公报里的实际数不一致时，数据优先。**例如2025年GDP目标5.5%、实际4.9%略低于目标；规上工业目标7.5%以上、实际7.0%基本达标；投资目标4%以上、实际-3.7%、未达标；社零目标5%以上、实际+1.7%、大幅未达标。差异反映：贵州“工业/数字经济强，投资/消费/物价弱”。")
heading2(doc, "1.3 指标取舍与单位说明")
para(doc, "本报告统一以“2025年统计公报+2026年报告复盘”为口径，增速均为可比价，绝对数为人民币，单位按要求标注。GDP 23562.17亿元为全省初步核算数。")

# ---- 二、底盘 ----
heading1(doc, "二、先看贵州的特殊底盘：大数据/算力、白酒、旅游、扶贫成果与‘六大优势’")
para(doc, "贵州的地盘，决定它的发展叙事。它不像沿海靠外贸，也不像资源型完全靠单一大宗，而是“**大数据/算力高地+‘六大优势’新型工业化+白酒/旅游+绿色能源+脱贫攻坚成果**”的组合。")
bullet(doc, "**大数据/算力**：数字经济连续多年全国领先，算力规模2025年突破160Eflops，是全国智算资源最多、能力最强地区之一，5G基站16.08万个、互联网出省带宽6.1万Gbps、移动互联网接入流量122.4亿GB（+15.7%）。")
bullet(doc, "**白酒**：茅台等酱香白酒是“名片产业”，但2025年酒、饮料和精制茶制造业增加值仅+2.4%、饮料酒产量-3.1%、烟酒类零售-12.7%——白酒在“挤泡沫/转换”。")
bullet(doc, "**旅游/山地绿色**：2025年接待游客+8.9%、旅游总花费+11.1%、来黔境外游客+50.1%，国家5A景区10个、4A景区159个。")
bullet(doc, "**扶贫成果与‘六大优势’**：在全面建成小康/消除绝对贫困基础上，贵州把自己重新定位为“具有资源、生态、能源、区位、产业、科技等比较优势”的省份，推动“六大产业基地”新型工业化（能源、白酒、轻纺、电子信息、健康医药、生态食品等）。")
para(doc, "**制度含义**：贵州不追求“大而全”，而是把“算力/大数据、白酒、旅游、绿色能源、新型工业化”当核心资产。这既是贵州特色，也是它面对“投资/消费弱”时的重点发力方向。")

# ---- 三、宏观错位 ----
heading1(doc, "三、最关键的宏观错位：GDP破2.35万亿、工业/出口偏强，但投资/消费/物价偏弱")
para(doc, "2025年贵州GDP 23562.17亿元、+4.9%（一产+4.4%、二产+5.1%、三产+4.9%）。表面看“稳中向好”，但拆开看是“**新引擎强、旧引擎/需求弱**”的错位：")
para(doc, "**强的部分**：规上工业+7.0%（采矿业+9.4%、制造业+7.2%）；汽车产量28.07万辆+94.9%、新能源汽车25.33万辆+147.7%；电子信息制造业+35.9%、电气机械+41.1%、有色冶炼+26.1%、化工+17.0%。")
para(doc, "**弱的部分**：固定资产投资-3.7%（三产-6.1%）；房地产开发-1.6%、商品房销售面积-1.4%；社零+1.7%（限额以上-0.9%、餐饮-6.6%）；进出口848.36亿、-1.2%（出口-10.4%）；CPI-0.1%、PPI-3.1%。")
para(doc, "**核心错位一句话**：贵州“工业/数字经济强（汽车+42.3%、电子信息+35.9%、算力/大数据全国领先），但投资/消费/物价/外贸出口弱”。2026年若内需/投资修复，增长有望从“工业单极”走向“工业+消费/投资”双轮。")
table(doc,
    ["强引擎", "数据", "弱项", "数据"],
    [
        ["汽车制造", "+42.3%", "固定资产投资", "-3.7%"],
        ["电子信息制造", "+35.9%", "房地产开发投资", "-1.6%"],
        ["电气机械制造", "+41.1%", "社会消费品零售", "+1.7%"],
        ["有色冶炼", "+26.1%", "进出口", "-1.2%"],
        ["算力/大数据", "全国领先/超160Eflops", "PPI", "-3.1%"],
    ],
    widths=[3.4, 3.4, 3.4, 3.4])
para(doc, "**错位结论**　贵州的增长“很强、但也很矛盾”。强的部分（工业/数字经济/投资结构）与弱的部分（投资/消费/物价/外贸）并存。2026年“稳工业/投资+扩内需/消费+续算力/白酒/旅游”是重点。")

# ---- 四、被忽视细节 ----
heading1(doc, "四、容易被忽视、但信号很重的细节（15条）")
details = [
    ("1", "规上工业+7.0%、制造业+7.2%", "制造升级。"),
    ("2", "汽车产量+94.9%、新能源车+147.7%", "汽车/新能源领跑。"),
    ("3", "电子信息制造+35.9%、电气机械+41.1%", "新质生产力/新兴制造。"),
    ("4", "有色冶炼+26.1%、化工+17.0%", "资源深加工。"),
    ("5", "食品制造+12.4%、采煤+8.8%", "基础产业稳。"),
    ("6", "数字经济连续全国领先、算力超160E", "大数据/算力。"),
    ("7", "互联网出省带宽6.1万Gbps、5G基站16.08万个", "数字基座。"),
    ("8", "旅游接待+8.9%、旅游花费+11.1%、境外+50.1%", "旅游强、入境回升。"),
    ("9", "白酒从“规模”转“卖生活方式”", "白酒调整/结构。"),
    ("10", "民间投资+2.1%、设备购置+12.8%", "民间/设备投资。"),
    ("11", "工业增加值占GDP 25.5%", "新型工业化。"),
    ("12", "常住3857万、-3万、城镇化率57.66%（+1.01pct）", "人口渐稳、城镇化。"),
    ("13", "居民人均可支配收入首破3万（30001元、+5.0%）", "收入破3万、城乡。"),
    ("14", "城镇新增就业62.6万人", "就业稳。"),
    ("15", "CPI-0.1%、PPI-3.1%", "通缩/物价弱。"),
]
for d in details:
    para(doc, "**%s**" % d[0], size=11)
    para(doc, d[1], size=10.5, color=GRAY)

    para(doc, "**备注**　这条“错位”在贵州尤其鲜明：增长靠工业/数字经济/投资结构，但投资/消费/物价弱。2026年若内需/消费修复，增长可能从“工业单极”走向“工业+消费/投资”多极。这条细节，正是贵州2026年最难也最值得盯的“变量”。", size=10.5)

# ---- 五、2025年目标 vs 实际 对照 ----
heading1(doc, "五、2025年GDP目标 vs 实际：对照表")
t5 = [
    ["指标", "2025目标", "2025实际", "达标判定"],
    ["GDP增速", "5.5%左右", "+4.9%", "略低于目标"],
    ["规上工业增加值增速", "7.5%以上", "+7.0%", "略低于目标"],
    ["固定资产投资增速", "4%以上", "-3.7%", "未达标"],
    ["社会消费品零售增速", "5%以上", "+1.7%", "大幅未达标"],
    ["一般公共预算收入增速", "2%左右", "+2.5%", "达标"],
    ["城镇/农村居民收入增速", "5%左右/7%左右", "+3.9%/+5.7%", "未达/低"],
    ["居民消费价格(CPI)", "涨幅2%左右", "-0.1%", "远低于"],
    ["常住人口城镇化率", "58%左右", "57.66%", "基本达标"],
    ["城镇新增就业", "60万人左右", "62.6万人", "达标"],
]
table(doc, t5[0], t5[1:], widths=[3.4, 3.0, 3.4, 3.6])
para(doc, "**对照结论**　目标“偏进攻”：GDP/规上工业/就业都定得不低，但**投资/消费/城镇收入是最大失分项**（固投-3.7%、社零+1.7%远低于5%、城镇收入+3.9%低于5%目标）。工业与就业接住了，靠“汽车/电子信息/算力+白酒/旅游”撑起4.9%的增长。")

# ---- 六、增速分项支撑 ----
heading1(doc, "六、2025年增长是谁撑起来的？（结构归因）")
para(doc, "GDP+4.9%背后，是“**新动能强（工业/数字经济/投资结构）、传统需求弱（消费/外贸/物价）**”的结构。把增长贡献拆开看：")
g6 = [
    ["引擎", "贡献/状态", "说明"],
    ["汽车制造", "+42.3%", "新能源车+147.7%、汽车+94.9%"],
    ["电子信息制造", "+35.9%", "大数据/算力、电子制造"],
    ["电气机械制造", "+41.1%", "新能源装备/电器"],
    ["有色冶炼", "+26.1%", "资源深加工"],
    ["化工制造", "+17.0%", "基础化工"],
    ["数字经济/算力", "全国领先", "大数据枢纽、算力160E"],
    ["民间投资/设备购置", "+2.1%/+12.8%", "民间/设备回暖"],
    ["房地产开发投资", "-1.6%", "地产回落、约束"],
    ["固定资产投资", "-3.7%", "三产/基建拖累"],
    ["CPI", "-0.1%", "物价偏弱、内需不足"],
]
table(doc, g6[0], g6[1:], widths=[3.4, 3.4, 6.6])
para(doc, "**一句话**　增长靠“工业+数字经济+投资结构”，但消费/外贸/物价是拖累。2026年考验贵州“能不能让内需/消费也跟着强起来”。")

# ---- 七、预算与财政 ----
heading1(doc, "七、预算与财政的“含金量”")
para(doc, "2025年贵州一般公共预算收入**2223.30亿元、+2.5%**，其中税收收入**+1.4%**。财政稳、但增速回落。")
bullet(doc, "一般公共预算收入+2.5%。")
bullet(doc, "税收+1.4%，随工业/数字经济企稳。")
bullet(doc, "财政“稳收+民生/工业/算力/白酒/旅游支出优先”，民生支出占比高。")

# ---- 八、民生底账 ----
heading1(doc, "八、民生底账：人口、收入与城乡")
para(doc, "2025年贵州常住人口**3857万人、-3万**，城镇化率**57.66%、+1.01pct**。")
para(doc, "居民人均可支配收入**30001元、+5.0%（首破3万）**，农村（+5.7%）快于城镇（+3.9%）。")
g8 = [
    ["指标", "数值", "信号"],
    ["常住人口", "3857万/-3万", "人口渐稳"],
    ["城镇化率", "57.66%/+1.01pct", "稳步城镇化"],
    ["居民人均可支配收入", "30001元/+5.0%", "收入破3万"],
    ["城镇新增就业", "62.6万人", "就业稳"],
]
table(doc, g8[0], g8[1:], widths=[4.2, 4.2, 4.6])
para(doc, "**民生观察**：居民收入破3万、人口渐稳，但物价偏弱、内需仍是短板。")

# ---- 九、城镇与农村 ----
heading1(doc, "九、城镇与农村：格局与均衡")
para(doc, "贵州城镇化率57.66%、+1.01pct，城乡协调推进；农村收入增速（+5.7%）快于城镇（+3.9%），城乡差仍待缩。")
bullet(doc, "农村居民收入+5.7%、快于城镇+3.9%。")
bullet(doc, "城乡收入比仍偏高，机会仍集中在贵阳/黔中都市圈。")
bullet(doc, "人口向贵阳/黔中城市群集中。")

# ---- 十、人口流入与流出 ----
heading1(doc, "十、人口流入与流出")
para(doc, "贵州2025年常住人口3857万人、-3万，人口净流出趋缓（但常住仍略降），外来流入趋缓。外来输入主要流向贵阳、黔中城市群、数字经济/制造/旅游区。")
para(doc, "未来看点：算力/大数据/汽车/旅游能否留下人口；若“数字经济+新型工业+旅游”成型，贵州有望从“外出劳务大省”走向“人口/就业托底”。")

# ---- 十一、物价与货币 ----
heading1(doc, "十一、物价与货币环境")
para(doc, "2025年贵州CPI**-0.1%**、PPI**-3.1%**，物价整体承压、需求偏弱。")
para(doc, "物价偏弱反映“工业/数字经济强、消费/投资弱”，与全国低通胀环境一致。2026年“把物价/需求稳住”是内需主线。")

# ---- 十二、区域一体化 ----
heading1(doc, "十二、区域一体化：贵州在“西部陆海通道+黔中城市群+大数据枢纽”里的位置")
para(doc, "贵州的核心战略坐标是“**数据枢纽+新型工业化+西部陆海新通道**”，承东启西、南下出海。")
bullet(doc, "大数据/算力：全国智算资源最多、能力最强地区之一。")
bullet(doc, "黔中城市群/贵阳：都市圈集聚。")
bullet(doc, "西部陆海新通道/南下出海：面向东盟/连接珠三角。")
bullet(doc, "绿色能源/新能源车：电力装机破1亿千瓦、新能源车产量破25万辆。")
para(doc, "若这一“算力/数据+新型工业+旅游/开放”闭环跑通，贵州将抢占新一轮“数字经济/西部陆海新通道”增长窗口。")

# ---- 十三、五条主线 ----
heading1(doc, "十三、未来5—10年最值得观察的五条主线")
lines5 = [
    ("① 大数据/算力/数字经济", "数字经济能否持续全国领先、产业化。"),
    ("② 新型工业化(汽车/电子信息)", "汽车、电子信息、新能源能否持续壮大。"),
    ("③ 白酒/旅游/绿色能源", "白酒企稳、旅游增长、绿色能源。"),
    ("④ 内需/消费/民间投资", "-3.7%/-1.6%，能否用新型工业/旅游补。"),
    ("⑤ 人口/就业/城乡", "人口净流出能否被数据/旅游托底。"),
]
for l in lines5:
    para(doc, "**%s**　%s" % l, space_after=6)

# ---- 十四、结论 ----
heading1(doc, "十四、最终结论：贵州在“大数据+新型工业化+旅游+白酒生态”里的增长逻辑")
para(doc, "贵州的2025年，本质上是“**大数据/算力/汽车/电子信息为核心，而投资/消费/外贸偏弱**”的答卷：GDP23562.17亿、+4.9%，规上工业+7.0%、数字经济全国领先、汽车+94.9%，但固投-3.7%、社零+1.7%、出口-10.4%、CPI-0.1%。")
para(doc, "只要算力/大数据、新型工业、旅游/绿能持续，贵州就站在“西部新兴增长极”的位；如果内需/消费持续偏弱、人口外流，贵州需承受“强工业、弱内需”的结构挑战。")
para(doc, "最稳观察信号：**一盯大数据/算力（数字经济）、二盯汽车/电子信息/新能源（新型工业）、三盯白酒/旅游（特色）、四盯投资/消费/民间（内需）、五盯人口/收入/城乡（民生）。**贵州，是“大数据+新型工业+旅游”的新样本。")

# ---- 附录A ----
heading1(doc, "附录A：主要资料来源与核验路径")
bullet(doc, "贵州省2025年《政府工作报告》——目标来源。")
bullet(doc, "《2025年贵州省国民经济和社会发展统计公报》（省统计局）—— GDP、工业、外贸、人口 实值。")
bullet(doc, "贵州省统计、贵阳海关、财政厅——外贸/财政。")
bullet(doc, "2026年贵州省政府工作报告——2025执行复盘。")
heading2(doc, "核验说明")
para(doc, "本报告所有增速、总量、占比均以统计公报/官方口径为基准，涉“大数据/算力/白酒/旅游/绿色能源”等以官方口径为准。")

# ---- 附录B ----
heading1(doc, "附录B：建议建立的年度跟踪仪表盘")
para(doc, "建议把下面10个“测脉搏”指标做成年度跟踪表：")
dash = [
    ["#", "指标", "2025基值", "为什么重要"],
    ["1", "GDP增速", "+4.9%", "总量与方向"],
    ["2", "规上工业增速", "+7.0%", "制造底盘"],
    ["3", "汽车/电子信息增速", "+42.3%/+35.9%", "新质生产力"],
    ["4", "数字经济/算力", "全国领先/超160E", "大数据/算力"],
    ["5", "固定资产投资/地产", "-3.7%/-1.6%", "投资结构"],
    ["6", "社零增速", "+1.7%", "内需消费"],
    ["7", "常住人口/城镇化率", "3857万/57.66%", "人口与城市"],
    ["8", "一般公共预算收入增速", "+2.5%", "财政质量"],
    ["9", "旅游/入境增速", "+8.9%/+50.1%", "旅游强"],
    ["10", "CPI/PPI", "-0.1%/-3.1%", "物价与需求"],
]
table(doc, dash[0], dash[1:], widths=[0.8, 4.6, 3.4, 4.4])
para(doc, "把这10个指标连起来看，任何一个“大数据/算力（4）、汽车/工业（3）、旅游（9）向上、投资/地产（5）修复”，都说明贵州在真正换挡。")

# ===================================== 保存
out = "/Users/x/Desktop/content-prod-lab/reports/贵州省_2025年政府工作报告_深度研究_2026-08-13.docx"
doc.save(out)
print("SAVED:", out)
print("PARAGRAPHS:", len(doc.paragraphs))
print("TABLES:", len(doc.tables))
