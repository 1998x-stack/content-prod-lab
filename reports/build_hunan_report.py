# -*- coding: utf-8 -*-
"""Build 湖南省2025年政府工作报告 深度研究 DOCX, 参照省市系列版式。"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DARK = RGBColor(0x1F, 0x2A, 0x44)
RED = RGBColor(0xB0, 0x1E, 0x2E)
GRAY = RGBColor(0x59, 0x60, 0x69)
BLUE = RGBColor(0x1F, 0x4E, 0x79)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
HEADER_FILL = "1F2A44"
K = "微软雅黑"


def set_run(run, size=11, bold=False, color=DARK, font=K):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    run.font.name = font
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), font)
    rFonts.set(qn("w:hAnsi"), font)
    rFonts.set(qn("w:eastAsia"), font)


def para(doc, text, size=11, bold=False, color=DARK, align=None,
         space_after=6, space_before=0, font=K):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        set_run(r, size=size, bold=bold or (i % 2 == 1), color=color, font=font)
    return p


def heading1(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(16)
    pf.space_after = Pt(8)
    r = p.add_run(text)
    set_run(r, size=16, bold=True, color=RED)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "B01E2E")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def heading2(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(10)
    pf.space_after = Pt(4)
    r = p.add_run(text)
    set_run(r, size=13, bold=True, color=BLUE)
    return p


def table(doc, headers, rows, widths=None, font_size=9.5):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_run(r, size=font_size, bold=True, color=WHITE)
        tcPr = hdr[i]._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), HEADER_FILL)
        tcPr.append(shd)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(val))
            set_run(r, size=font_size, color=DARK)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    return t


def quote_box(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(8)
    pf.left_indent = Pt(12)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), "B01E2E")
    pBdr.append(left)
    pPr.append(pBdr)
    r = p.add_run(text)
    set_run(r, size=11, color=DARK, font="楷体")
    return p


def bullet(doc, text, size=10.5):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.7)
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        set_run(r, size=size, bold=(i % 2 == 1), color=DARK)
    return p


# ================================================================= 文档主体
doc = Document()
sec = doc.sections[0]
sec.page_width = Cm(21)
sec.page_height = Cm(29.7)
sec.top_margin = Cm(2.4)
sec.bottom_margin = Cm(2.2)
sec.left_margin = Cm(2.5)
sec.right_margin = Cm(2.5)

st = doc.styles["Normal"]
st.font.name = K
st.font.size = Pt(11)
st.element.rPr.rFonts.set(qn("w:eastAsia"), K)



# ---- 封面 ----
para(doc, "湖南省2025年政府工作报告", size=24, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=60, space_after=4)
para(doc, "深度研究与“容易被忽视的细节”分析", size=16, bold=True, color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
para(doc, "从“工程机械、轨道交通、食品医药、基础设施建设与中部崛起”重新理解湖南", size=12, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)
para(doc, "━━━━━━━━━━━━━━━━", color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
para(doc, "研究对象：2025年湖南省政府工作报告及同期财政、国民经济和社会发展统计资料、2026年官方复盘", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
para(doc, "标定日期：2026-08-13", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
doc.add_page_break()

# ---- 目录 ----
catalog = [
    "执行摘要",
    "一、研究口径与阅读方法",
    "二、先看湖南的特殊底盘：工程机械/轨道交通、中部崛起、长株潭与基础设施",
    "三、最关键的宏观错位：GDP破5.53万亿、制造强，但投资/地产/外贸弱",
    "四、容易被忽视、但信号很重的细节（15条）",
    "五、2025年GDP目标 vs 实际：对照表",
    "六、2025年增长是谁撑起来的？（结构归因）",
    "七、预算与财政的“含金量”",
    "八、民生底账：人口、收入与城乡",
    "九、城镇与农村：格局与均衡",
    "十、人口流入与流出",
    "十一、物价与货币环境",
    "十二、区域一体化：湖南在“中部崛起+长株潭+粤港澳对接”里的位置",
    "十三、未来5—10年最值得观察的五条主线",
    "十四、最终结论：湖南在“工程机械+轨道交通+新型工业化+中部崛起”里的增长逻辑",
    "附录A：主要资料来源与核验路径",
    "附录B：建议建立的年度跟踪仪表盘",
]
para(doc, "目　录", size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=10, space_after=10)
for item in catalog:
    para(doc, item, size=12, space_after=4)
doc.add_page_break()

# ---- 执行摘要 ----
heading1(doc, "执行摘要")
para(doc, "**核心判断**　如果只看新闻标题，2025年湖南最显眼的是“GDP破5.53万亿、增长4.8%（稳居全国前十）”、“汽车产量162.6万辆（+30.8%）”、“新能源汽车超110万辆（全国第9）”和“混凝土机械（工程机械）+39.6%”。但这份研究真正值得深读的，是这座“中部崛起+长株潭+工程机械/轨道交通”的中部大省，如何在固定资产投资（-11.7%）、房地产开发（-28.6%）与人口自然增长率（-4.40‰）的背景下，靠“工程机械/汽车/高技术制造+消费/文旅”稳住增长。")
para(doc, "把2025年初设定的目标（GDP增长5.5%左右）、2025年《国民经济和社会发展统计公报》、以及2026年报告对2025年的复盘放在一起看，湖南呈现清晰暗线：**从“基建/传统制造”的旧底盘，向“工程机械/轨道交通+汽车/新能源+数字经济/文旅”转型**。旧引擎（地产/基建/传统制造）在调整；新引擎（工程机械出海、汽车/新能源、电子信息、食品医药、数字经济）被要求更快补位。")
para(doc, "因此，本报告不按政府工作报告逐段复述，而采用“显性表述—同期数据—制度含义—长期影响”的方式，专门提取容易被忽视、但对判断湖南未来5—10年发展模式更有价值的信号。")
para(doc, "最容易记住的一句话：**湖南是“中部崛起+长株潭+工程机械/轨道交通/食品医药”的中部样本，靠“工程机械出海+汽车/新能源+文旅/消费”撑起增长。**观察湖南，与其看“GDP 5.53万亿”，不如看“工程机械龙头、汽车产量、长株潭、中部崛起、新三样出口+88%”这几张名片。")
heading2(doc, "一页速览：2025年湖南经济的“表与里”")
table(doc,
    ["维度", "表面现象", "更深层信号"],
    [
        ["增长", "GDP 55308.7亿、+4.8%", "一产9.0%、二产36.0%、三产55.0%"],
        ["产业", "规上工业+6.2%", "汽车+30.8%、高技术制造+11.2%"],
        ["外贸", "进出口5414.1亿、-3.9%", "出口-4.1%、新三样+88%"],
        ["投资", "固定资产投资-11.7%", "地产-28.6%、基建-18.8%"],
        ["财政", "地方一般公共预算收入3507.6亿、+1.7%", "税收-4.9%、非税+12.6%"],
        ["消费", "社零21204.6亿、+3.6%", "智能手机+66.1%、通讯+49.9%"],
        ["人口", "常住6492万、城镇化63.0%", "自然增-4.40‰、老龄化24.72%"],
        ["开放", "中部崛起/长株潭/工程机械出海", "新三样出口+88%"],
    ],
    widths=[2.2, 5.4, 8.4])
para(doc, "", size=4, space_after=2)

# ---- 一、研究口径与阅读方法 ----
heading1(doc, "一、研究口径与阅读方法")
heading2(doc, "1.1 本报告的三份核心底稿")
bullet(doc, "**2025年《政府工作报告》**：2025年1月在省十四届人大三次会议上作，给出全年预期目标（GDP增长5.5%左右、规模工业增长6.5%以上、地方一般公共预算收入增长2.5%）。")
bullet(doc, "**2025年《国民经济和社会发展统计公报》**：湖南省统计局2026年3月25日发布，给出全年实际执行数与增速，是“事后验证”的锚点。")
bullet(doc, "**2026年湖南省政府工作报告及官方复盘**：对2025执行结果的追认，交叉验证“目标—执行—再评价”三段式。")
heading2(doc, "1.2 阅读方法：显性—数据—制度—长期")
para(doc, "每一章按“**显性表述 → 同期数据 → 制度含义 → 长期影响**”四层展开。其中“制度含义”回答“政府为什么这么排优先序”，“长期影响”回答“这对未来5—10年意味着什么”。")
para(doc, "关键判别：**当报告使用的“增长词”和统计公报里的实际数不一致时，数据优先。**例如2025年GDP目标5.5%左右、实际4.8%略低于目标；规上工业目标6.5%以上、实际6.2%基本达标；2024年GDP也达4.8%、投资/进出口未达年初预期。差异反映：湖南“工程机械/消费/文旅强，投资/地产/外贸弱”。")
heading2(doc, "1.3 指标取舍与单位说明")
para(doc, "本报告统一以“2025年统计公报+2026年报告复盘”为口径，增速均为可比价，绝对数为人民币。GDP 55308.7亿元为全省初步核算数（稳居全国前十强）。")

# ---- 二、底盘 ----
heading1(doc, "二、先看湖南的特殊底盘：工程机械/轨道交通、中部崛起、长株潭与基础设施")
para(doc, "湖南的地盘，取决于它作为“**中部崛起+长株潭都市圈+工程机械/轨道交通/食品医药强省**”的特殊定位。它是中部地区制造业与科教的高地。")
bullet(doc, "**工程机械**：混凝土机械产量4.4万台、+39.6%；三一重工、中联重科海外营收占比约60%，工程机械“出海”成效显著。")
bullet(doc, "**汽车/新能源**：汽车产量162.6万辆（+30.8%）、新能源汽车超110万辆（全国第9）；“新三样”出口+88%。")
bullet(doc, "**长株潭/中部崛起**：长株潭GDP 22878.5亿、约占全省41%，规上工业+7.7%；落实中部崛起“三基地一枢纽”，深度融入粤港澳/长三角对接。")
bullet(doc, "**轨道交通/基础设施**：高铁2501公里（全国第6）、高速公路8667公里；2026年推进6条高铁、17条高速扩容续建等“十大基础设施项目”。")
para(doc, "**制度含义**：湖南不追求总量冒进，而靠“工程机械/轨道交通等先进制造集群（中部第1）”+“长株潭都市圈”+“中部崛起战略”多重支撑。5个国家先进制造业集群居中部第1。")

# ---- 三、宏观错位 ----
heading1(doc, "三、最关键的宏观错位：GDP破5.53万亿、制造强，但投资/地产/外贸弱")
para(doc, "2025年湖南GDP 55308.7亿元、+4.8%（一产+4.2%、二产+4.1%、三产+5.3%）。表面看“稳中有进”，但拆开看是“**制造/消费强、投资/地产/外贸弱**”的错位：")
para(doc, "**强的部分**：规上工业+6.2%（汽车+30.8%、工程机械混凝土+39.6%、高技术制造+11.2%、装备制造+8.6%）；社零+3.6%（智能手机+66.1%、通讯+49.9%、家电+23.1%）；文旅/入境游客+19.1%。")
para(doc, "**弱的部分**：固定资产投资-11.7%（地产-28.6%、基建-18.8%、民间-10.0%）；进出口5414.1亿、-3.9%（出口-4.1%）；地方一般公共预算收入+1.7%（税收-4.9%）；人口自然增长率-4.40‰。")
para(doc, "**核心错位一句话**：湖南“制造/工程机械/消费/文旅强（汽车+30.8%、高技术+11.2%），但投资/地产/外贸弱、人口老龄化加深”。2026年“稳投资/地产、扩内需、续工程机械/新能源、推中部崛起”是重点。")
table(doc,
    ["强引擎", "数据", "弱项", "数据"],
    [
        ["汽车产量", "+30.8%", "固定资产投资", "-11.7%"],
        ["工程机械(混凝土)", "+39.6%", "房地产开发投资", "-28.6%"],
        ["高技术制造业", "+11.2%", "进出口", "-3.9%"],
        ["装备制造业", "+8.6%", "人口自然增长率", "-4.40‰"],
        ["社零/智能手机", "+3.6%/+66.1%", "地方税收", "-4.9%"],
    ],
    widths=[3.6, 3.0, 3.6, 3.0])
para(doc, "**错位结论**　湖南的增长“强制造、弱投资/地产/人口”。强的部分（工程机械/汽车/高技术/文旅）与弱的部分（投资/地产/人口/外贸）并存。2026年“稳制造业+补投资/地产/人口、续文旅消费”是重点。")

# ---- 四、被忽视细节 ----
heading1(doc, "四、容易被忽视、但信号很重的细节（15条）")
details = [
    ("1", "GDP 55308.7亿、+4.8%（全国前十强）", "总量/稳增。"),
    ("2", "规上工业+6.2%、汽车产量+30.8%", "工业/汽车强。"),
    ("3", "新能源汽车超110万辆（全国第9）", "新能源车。"),
    ("4", "混凝土机械(工程机械)+39.6%", "工程机械龙头。"),
    ("5", "高技术制造+11.2%、装备+8.6%", "新质生产力。"),
    ("6", "长株潭GDP 22878.5亿、规上工业+7.7%", "长株潭脊梁。"),
    ("7", "新三样出口+88%", "出口新动能。"),
    ("8", "民营增加值占规上工业64.5%、民间投资占比约65%", "民营底盘。"),
    ("9", "社零+3.6%、智能手机/通讯/家电高增", "消费升级。"),
    ("10", "文旅/入境游客+19.1%、张家界+32.5%", "文旅强。"),
    ("11", "入境游客破240万人次、工程机械出海60%", "开放/出海。"),
    ("12", "常住6492万、城镇化率63.0%（+0.93pct）", "人口/城镇化。"),
    ("13", "居民收入39545元、+5.0%、农村+5.8%", "收入稳、城乡扩。"),
    ("14", "CPI持平0.0%、PPI-1.0%", "物价/PPI弱。"),
    ("15", "粮食产量突破620亿斤（创历史新高）", "粮食安全。"),
]
for d in details:
    para(doc, "**%s**" % d[0], size=11)
    para(doc, d[1], size=10.5, color=GRAY)

    para(doc, "**备注**　这条“错位”在湖南尤其鲜明：增长靠工程机械/汽车/高技术/文旅，但投资/地产/外贸/人口弱。2026年若投资/地产修复、中部崛起发力，增长可能从“制造单极”走向“制造+内需/投资”多极。这条细节，正是湖南2026年最难也最值得盯的“变量”。", size=10.5)

# ---- 五、2025年目标 vs 实际 对照 ----
heading1(doc, "五、2025年GDP目标 vs 实际：对照表")
t5 = [
    ["指标", "2025目标", "2025实际", "达标判定"],
    ["GDP增速", "5.5%左右", "+4.8%", "略低于目标"],
    ["地方一般公共预算收入增速", "2.5%", "+1.7%", "未达标"],
    ["规模工业增加值增速", "6.5%以上", "+6.2%", "基本达标"],
    ["固定资产投资增速", "4.5%", "-11.7%", "大幅未达标"],
    ["社会消费品零售增速", "5.5%以上", "+3.6%", "未达标"],
    ["居民收入增速", "高于经济增速", "+5.0%", "达标"],
    ["城镇新增就业", "70万人", "82.7万人", "超额"],
    ["粮食产量", "616亿斤左右", "620亿斤", "超额"],
]
table(doc, t5[0], t5[1:], widths=[3.4, 3.0, 3.0, 3.8])
para(doc, "**对照结论**　目标“偏求稳”，但实际多处未达：GDP 4.8%（低于5.5%目标）、投资-11.7%（大幅失分）、社零3.6%（低于5.5%）。亮点在就业（82.7万）、粮食（620亿斤）、收入（5.0%高于GDP）。制造与消费撑住，投资/地产是最大短板。")

# ---- 六、增速分项支撑 ----
heading1(doc, "六、2025年增长是谁撑起来的？（结构归因）")
para(doc, "GDP+4.8%背后，是“**制造/消费/文旅强、投资/地产/人口弱**”的结构。把增长贡献拆开看：")
g6 = [
    ["引擎", "贡献/状态", "说明"],
    ["规上工业", "+6.2%", "汽车+30.8%、工程机械+39.6%"],
    ["高技术/装备制造", "+11.2%/+8.6%", "新质生产力"],
    ["长株潭都市圈", "+4.3%(GDP)", "规上工业+7.7%、全省脊梁"],
    ["消费/社零", "+3.6%", "智能手机+66.1%、通讯+49.9%"],
    ["文旅/入境", "+19.1%", "入境游客240万人次"],
    ["新三样出口", "+88%", "出口新动能"],
    ["房地产开发", "-28.6%", "地产调整"],
    ["固定资产投资", "-11.7%", "地产/基建拖累"],
    ["进出口", "-3.9%", "外贸-4.1%"],
    ["人口自然增长率", "-4.40‰", "人口/老龄化约束"],
]
table(doc, g6[0], g6[1:], widths=[3.6, 3.6, 6.0])
para(doc, "**一句话**　增长靠“工业（汽车/工程机械）+消费/文旅+长株潭”，但投资/地产/人口是最大拖累。2026年考验湖南“能不能让投资/地产/人口也跟着稳住”。")

# ---- 七、预算与财政 ----
heading1(doc, "七、预算与财政的“含金量”")
para(doc, "2025年湖南地方一般公共预算收入**3507.6亿元、+1.7%**（同口径+2.5%以上），其中税收收入**-4.9%**、非税收入**+12.6%**。")
bullet(doc, "财政收入+1.7%、税收-4.9%（地产/企业利润拖累）。")
bullet(doc, "财政“稳收+民生支出优先”，支持中部崛起/基建/产业。")
bullet(doc, "非税收入+12.6%，构成补充。")

# ---- 八、民生底账 ----
heading1(doc, "八、民生底账：人口、收入与城乡")
para(doc, "2025年湖南常住人口**6492万人**，城镇化率**63.0%、+0.93pct**。人口自然增长率**-4.40‰**（出生率4.83‰、死亡率9.24‰），60岁及以上占24.72%。")
para(doc, "居民人均可支配收入**39545元、+5.0%**，农村（+5.8%）快于城镇（+4.1%），城乡比2.27。")
g8 = [
    ["指标", "数值", "信号"],
    ["常住人口", "6492万", "人口总量大"],
    ["人口自然增长率", "-4.40‰", "自然负增"],
    ["城镇化率", "63.0%/+0.93pct", "稳步城镇化"],
    ["居民人均可支配收入", "39545元/+5.0%", "收入增速高于GDP"],
]
table(doc, g8[0], g8[1:], widths=[4.2, 4.0, 4.8])
para(doc, "**民生观察**：收入增速（5.0%）高于GDP（4.8%），但人口自然负增、老龄化加重（24.72%）是重要约束。")

# ---- 九、城镇与农村 ----
heading1(doc, "九、城镇与农村：格局与均衡")
para(doc, "湖南城镇化率63.0%、+0.93pct，城乡收入比2.27（缩小），城乡相对协调。")
bullet(doc, "农村居民收入+5.8%、快于城镇+4.1%。")
bullet(doc, "社零城镇+3.5%、乡村+4.0%，乡村消费引领。")
bullet(doc, "长株潭收入最高（58851元）、大湘西收入快（+6.0%）。")

# ---- 十、人口流入与流出 ----
heading1(doc, "十、人口流入与流出")
para(doc, "湖南常住6492万人，人口自然增长为负（-4.40‰），但总体人口规模仍大；长株潭、长沙等吸引人口/人才集聚。")
para(doc, "未来看点：中部崛起+长株潭都市圈+工程机械/汽车能否留住人口/吸引回流；若“制造业+科教高地”成势，湖南有望对冲人口自然负增。")

# ---- 十一、物价与货币 ----
heading1(doc, "十一、物价与货币环境")
para(doc, "2025年湖南CPI**持平（0.0%）**、PPI**-1.0%**，物价温和、工业品价格偏弱。")
para(doc, "物价基本平稳反映“供需平衡”，但PPI走弱、农产品价格-5.7%反映部分领域压力。2026年“扩内需、稳价格”是主线。")

# ---- 十二、区域一体化 ----
heading1(doc, "十二、区域一体化：湖南在“中部崛起+长株潭+粤港澳对接”里的位置")
para(doc, "湖南的核心战略坐标是“**中部崛起+长株潭都市圈+粤港澳/长三角对接**”，也是中部地区制造业与科教的高地。")
bullet(doc, "长株潭：GDP 22878.5亿（约41%）、规上工业+7.7%。")
bullet(doc, "中部崛起：“三基地一枢纽”，5个国家先进制造业集群居中部第1。")
bullet(doc, "粤港澳/长三角对接：深度融入，抢抓海南自贸港封关机遇。")
bullet(doc, "基础设施：高铁2501公里、高速公路8667公里。")
para(doc, "若“长株潭+中部崛起+制造业集群+对接粤港澳”成势，湖南将作为中部增长极实现更高能级。")

# ---- 十三、五条主线 ----
heading1(doc, "十三、未来5—10年最值得观察的五条主线")
lines5 = [
    ("① 工程机械/轨道交通", "工程机械出海、轨道交通产业能否持续。"),
    ("② 汽车/新能源/新一代信息技术", "新能源汽车、数字经济能否壮大。"),
    ("③ 长株潭都市圈/中部崛起", "长株潭协同、中部地位能否提升。"),
    ("④ 投资/地产/人口", "-11.7%/-28.6%，能否补投资/稳人口。"),
    ("⑤ 文旅/消费/食品医药", "文旅、食品医药、内需能否扩容。"),
]
for l in lines5:
    para(doc, "**%s**　%s" % l, space_after=6)

# ---- 十四、结论 ----
heading1(doc, "十四、最终结论：湖南在“工程机械+轨道交通+新型工业化+中部崛起”里的增长逻辑")
para(doc, "湖南的2025年，本质上是“**工程机械/汽车/高技术/文旅为核心，而投资/地产/人口弱**”的答卷：GDP55308.7亿、+4.8%（全国前十），规上工业+6.2%、汽车+30.8%、工程机械+39.6%、新三样出口+88%，但固投-11.7%、地产-28.6%、进出口-3.9%、人口自然增-4.40‰。")
para(doc, "只要工程机械/轨道交通出海、汽车/新能源、长株潭、中部崛起持续，湖南就站在“中部增长极”的位；如果投资/地产/人口持续偏弱、外贸不稳，湖南需承受“强制造、弱投资/人口”的结构挑战。")
para(doc, "最稳观察信号：**一盯工程机械/汽车（制造）、二盯中部崛起/长株潭（区域）、三盯文旅/消费（内需）、四盯投资/地产（约束）、五盯人口/收入（民生）。**湖南，是“工程机械+中部崛起+文旅”的新样本。")

# ---- 附录A ----
heading1(doc, "附录A：主要资料来源与核验路径")
bullet(doc, "湖南省2025年《政府工作报告》——目标来源。")
bullet(doc, "《2025年湖南省国民经济和社会发展统计公报》（省统计局，2026-03-25）——GDP、工业、外贸、人口实值。")
bullet(doc, "2026年湖南省政府工作报告——2025执行复盘+中部崛起/工程机械/长株潭。")
bullet(doc, "长沙海关、省财政厅——外贸/财政。")
heading2(doc, "核验说明")
para(doc, "本报告所有增速、总量、占比均以统计公报/官方口径为基准，涉“工程机械/轨道交通/中部崛起/长株潭”等以官方口径为准。")

# ---- 附录B ----
heading1(doc, "附录B：建议建立的年度跟踪仪表盘")
para(doc, "建议把下面10个“测脉搏”指标做成年度跟踪表：")
dash = [
    ["#", "指标", "2025基值", "为什么重要"],
    ["1", "GDP增速", "+4.8%", "总量与方向"],
    ["2", "规上工业增速", "+6.2%", "制造底盘"],
    ["3", "汽车/工程机械产量", "+30.8%/+39.6%", "制造龙头"],
    ["4", "进出口/新三样增速", "-3.9%/+88%", "外贸韧性"],
    ["5", "固定资产投资/地产", "-11.7%/-28.6%", "投资结构"],
    ["6", "社零增速", "+3.6%", "内需消费"],
    ["7", "常住人口/自然增长率", "6492万/-4.40‰", "人口与城市"],
    ["8", "地方一般公共预算收入增速", "+1.7%", "财政质量"],
    ["9", "长株潭GDP增速", "+4.3%", "区域脊梁"],
    ["10", "CPI/PPI", "0.0%/-1.0%", "物价与需求"],
]
table(doc, dash[0], dash[1:], widths=[0.8, 4.6, 3.4, 4.4])
para(doc, "把这10个指标连起来看，任何一个工程机械/汽车（3）、中部/长株潭（9）、文旅（内需）向上、投资/地产（5）修复，都说明湖南在真正换挡。")

# ===================================== 保存
out = "/Users/x/Desktop/content-prod-lab/reports/湖南省_2025年政府工作报告_深度研究_2026-08-13.docx"
doc.save(out)
print("SAVED:", out)
print("PARAGRAPHS:", len(doc.paragraphs))
print("TABLES:", len(doc.tables))
