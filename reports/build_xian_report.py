# -*- coding: utf-8 -*-
"""Build 西安市2025年政府工作报告 深度研究 DOCX, 参照省市系列版式。"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DARK = RGBColor(0x1F, 0x2A, 0x44)
RED = RGBColor(0xB0, 0x1E, 0x2E)
GRAY = RGBColor(0x59, 0x60, 0x69)
BLUE = RGBColor(0x1F, 0x4E, 0x79)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
HEADER_FILL = "1F2A44"
K = "微软雅黑"


def set_run(run, size=11, bold=False, color=DARK, font=K):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    run.font.name = font
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), font)
    rFonts.set(qn("w:hAnsi"), font)
    rFonts.set(qn("w:eastAsia"), font)


def para(doc, text, size=11, bold=False, color=DARK, align=None,
         space_after=6, space_before=0, font=K):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        set_run(r, size=size, bold=bold or (i % 2 == 1), color=color, font=font)
    return p


def heading1(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(16)
    pf.space_after = Pt(8)
    r = p.add_run(text)
    set_run(r, size=16, bold=True, color=RED)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "B01E2E")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def heading2(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(10)
    pf.space_after = Pt(4)
    r = p.add_run(text)
    set_run(r, size=13, bold=True, color=BLUE)
    return p


def table(doc, headers, rows, widths=None, font_size=9.5):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_run(r, size=font_size, bold=True, color=WHITE)
        tcPr = hdr[i]._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), HEADER_FILL)
        tcPr.append(shd)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(val))
            set_run(r, size=font_size, color=DARK)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    return t


def quote_box(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(8)
    pf.left_indent = Pt(12)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), "B01E2E")
    pBdr.append(left)
    pPr.append(pBdr)
    r = p.add_run(text)
    set_run(r, size=11, color=DARK, font="楷体")
    return p


def bullet(doc, text, size=10.5):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.7)
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        set_run(r, size=size, bold=(i % 2 == 1), color=DARK)
    return p


# ================================================================= 文档主体
doc = Document()
sec = doc.sections[0]
sec.page_width = Cm(21)
sec.page_height = Cm(29.7)
sec.top_margin = Cm(2.4)
sec.bottom_margin = Cm(2.2)
sec.left_margin = Cm(2.5)
sec.right_margin = Cm(2.5)

st = doc.styles["Normal"]
st.font.name = K
st.font.size = Pt(11)
st.element.rPr.rFonts.set(qn("w:eastAsia"), K)



# ---- 封面 ----
para(doc, "西安市2025年政府工作报告", size=24, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=60, space_after=4)
para(doc, "深度研究与\u201c容易被忽视的细节\u201d分析", size=16, bold=True, color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
para(doc, "从\u201c硬科技、新能源汽车、历史文化与副省级省会\u201d重新理解西安", size=12, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)
para(doc, "\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501", color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
para(doc, "研究对象：2025年西安市政府工作报告及同期财政、国民经济和社会发展统计资料、2026年官方复盘", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
para(doc, "标定日期：2026-08-13", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
doc.add_page_break()

# ---- 目录 ----
catalog = [
    "执行摘要",
    "一、研究口径与阅读方法",
    "二、先看西安的特殊底盘：硬科技、新能源汽车、历史文旅与副省级省会",
    "三、最关键的宏观错位：GDP破1.39万亿、进出口/社零强，但固投/财政/地产偏弱",
    "四、容易被忽视、但信号很重的细节（15条）",
    "五、2025年GDP目标 vs 实际：对照表",
    "六、2025年增长是谁撑起来的？（结构归因）",
    "七、预算与财政的\u201c含金量\u201d",
    "八、民生底账：人口、收入与城乡",
    "九、城镇与农村：格局与均衡",
    "十、人口流入与流出",
    "十一、物价与货币环境",
    "十二、区域一体化：西安在\u201c西安都市圈+关中平原城市群+一带一路枢纽\u201d里的位置",
    "十三、未来5—10年最值得观察的五条主线",
    "十四、最终结论：西安在\u201c硬科技+新能源汽车+文旅+科教\u201d里的增长逻辑",
    "附录A：主要资料来源与核验路径",
    "附录B：建议建立的年度跟踪仪表盘",
]
para(doc, "目　录", size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=10, space_after=10)
for item in catalog:
    para(doc, item, size=12, space_after=4)
doc.add_page_break()

# ---- 执行摘要 ----
heading1(doc, "执行摘要")
para(doc, "**核心判断**　如果只看新闻标题，2025年西安最显著的是\u201cGDP破1.39万亿、增长4.7%\u201d、\u201c进出口4987.9亿、+21.1%\u201d、\u201c规上工业总产值破万亿\u201d、\u201c接待游客3.27亿人次\u201d。但这份研究真正值得深读的，是这座\u201c硬科技+新能源汽车+历史文化\u201d的副省级省会，如何在GDP增速（4.7%）低于预期（5.5%）背景下，靠\u201c进出口（+21.1%）+消费（社零+5.3%）+新能源汽车/半导体+文旅\u201d支撑其结构性亮点。")
para(doc, "把2025年初设定的目标（GDP增长5.5%左右）、2025年《国民经济和社会发展统计公报》、2026年复盘放在一起看，西安呈现清晰暗线：**从\u201c钢铁石化/人口流出\u201d的旧底盘，向\u201c硬科技+新能源汽车/半导体+文旅+科教\u201d升级**。总量增速（4.7%）低于目标，但外贸、消费、硬科技产业是弹性亮点。")
para(doc, "因此，本报告不按政府工作报告逐段复述，而采用\u201c显性表述—同期数据—制度含义—长期影响\u201d的方式，专门提取容易被忽略、但对判断西安未来5—10年发展模式更有价值的信号。")
para(doc, "最容易记住的一句话：**西安是\u201c硬科技+新能源汽车+文旅+科教\u201d的西北龙头，靠\u201c进出口+产业制造+文旅+高校\u201d撑起增长。**观察西安，与其看\u201cGDP 1.39万亿\u201d，不如看\u201c硬科技、新能源汽车105万辆、进出口+21%、旅游3.27亿人次、研发强度5.5%\u201d这几张名片。")
heading2(doc, "一页速览：2025年西安经济的\u201c表与里\u201d")
table(doc,
    ["维度", "表面现象", "更深层信号"],
    [
        ["增长", "GDP 13902.67亿、+4.7%", "一产2.4%、二产28.6%、三产69.0%"],
        ["产业", "规上工业+5.7%", "铁路船舶+8.4%、专设+7.6%"],
        ["外贸", "进出口4987.9亿、+21.1%", "出口+25.8%、一带一路+11.6%"],
        ["投资", "固投-15.3%、基础设施-56.3%", "工业+2.4%、高技术+15.4%"],
        ["消费", "社零5721.21亿、+5.3%", "通讯+31.5%、新能源汽车+34.9%"],
        ["人口", "常住1323.63万、城镇化81.01%", "人口+6.87万、自然增长"],
        ["科教", "研发强度5.5%、硬科技", "半导体/光子/商业航天/文旅"],
    ],
    widths=[2.2, 5.6, 7.4])
para(doc, "", size=10, space_after=2)

# ---- 一、研究口径与阅读方法 ----
heading1(doc, "一、研究口径与阅读方法")
heading2(doc, "1.1 本报告的核心底稿")
bullet(doc, "**2025年《政府工作报告》**（2025年3月）——目标：GDP 5.5%左右、规上工业6.5%、固投1%、财政3%、进出口10%。")
bullet(doc, "**《2025年西安市统计公报》**（市统计局2026-05）——GDP、工业、外贸、人口实数。")
bullet(doc, "**2026年西安市政府工作报告/复盘**（2026年）——2025追认（进出口+21%等）与硬科技/新能源汽车展望。")
heading2(doc, "1.2 阅读方法：显性—数据—制度—长期")
para(doc, "**关键判别**：数据优先。例2025年GDP目标5.5%、实际4.7%（未达）；但进出口目标10%、实际21.1%超预期；规上工业目标6.5%、实际5.7%略欠。西安\u201c进出口/内需/硬科技好、GDP总量/固投/财政弱\u201d，穿透总量看产业升级。")

# ---- 二、底盘 ----
heading1(doc, "二、先看西安的特殊底盘：硬科技、新能源汽车、历史文旅与副省级省会")
para(doc, "西安的地盘取决于它作为\u201c**硬科技+科教资源密集+历史文化名城+副省级省会+一带一路枢纽**\u201d的特殊定位。它是西北龙头、科教大市（西工大、西安交大等）。")
bullet(doc, "**硬科技**：研发投入强度5.56%（全国顶尖）、技术合同4500亿；光子/半导体/新材料/超导/商业航天五大军团。")
bullet(doc, "**新能源汽车/电子**：汽车产量（比亚迪等）；新能源汽车105.14万辆；集成电路/SAM/三星/奕斯伟扩产。")
bullet(doc, "**历史文化/文旅**：旅游接待3.27亿人次、旅游总花费4026亿、国际游客大幅增长；“盛唐”大IP。")
bullet(doc, "**一带一路枢纽**：中欧班列（西安）全国第一、国际港；空/铁双港开放。")
para(doc, "这一底板决定了西安2025成绩单的\u201c底色\u201d：**只要硬科技/新能源车/文旅/一带一路持续，西安就站在\u201c硬科技+内陆枢纽\u201d增长极；若固投/财政/传统地产承压，西安需承受\u201c总量/外需强、投资/收入弱\u201d的结构挑战。**")

# ---- 三、宏观错位 ----
heading1(doc, "三、最关键的宏观错位：GDP破1.39万亿、进出口/社零强，但固投/财政/地产偏弱")
para(doc, "西安2025年最值得咀嚼的错位，是\u201c**进出口+消费+硬科技强、固投/财政/总量偏弱**\u201d。这种错位决定了对座西部省会城市的观察不能只看GDP增速。")
bullet(doc, "**GDP**：13902.67亿、+4.7%（目标5.5%未达）。一产332.80亿（+3.6%，占比2.4%）、二产3970.37亿（+4.8%，占比28.6%）、三产9599.50亿（+4.6%，占比69.0%）。")
bullet(doc, "**工业**：规上工业+5.7%（报告口径工业增加值+6.1%）；铁路/船舶/航空航天+8.4%、专用设备+7.6%、通用设备+6.7%。汽车产量148.27万辆、新能源汽车105.14万辆（-6.5%）。")
bullet(doc, "**外贸**：进出口4987.9亿、+21.1%（出口+25.8%）；一带一路+11.6%、机电产品出口+25.6%。")
bullet(doc, "**消费**：社零5721.21亿、+5.3%；通讯器材+31.5%、家电+30.8%、新能源汽车零售+34.9%。")
bullet(doc, "**固投/财政**：固投-15.3%、基础设施投资-56.3%；一般公共预算收入-2.3%。")
para(doc, "**为什么读这条**：西安作为\u201c西北副省级+硬科技\u201d，结构性矛盾是\u201c外循环/内需消费/硬科技强、固投/基础/财政弱\u201d。总量承压但外贸与新质制造弹性足。")

# ---- 四、15条细节 ----
heading1(doc, "四、容易被忽视、但信号很重的细节（15条）")
para(doc, "以下15条，都在2025年报告/公报里，但常被\u201cGDP 1.39万亿\u201d、\u201c4.7%\u201d等总量掩盖。它们是判断西安2025之后5—10年的关键小信号。")
bullet(doc, "**1. 进出口+21.1%、出口+25.8%**：外贸是西部省会罕见超高增长。")
bullet(doc, "**2. 一带一路出口+11.6%、电动载人汽车+66.1%**：新三样/一带一路逻辑。")
bullet(doc, "**3. 研发强度5.56%（全国领先）**：硬科技/科教资源密集。")
bullet(doc, "**4. 新能源汽车产量105.14万辆**：比亚迪等新能车制造基地。")
bullet(doc, "**5. 光子/半导体/商业航天**：6G光子中试、商业航天、低空、无人机（民用无人机+78.5%）。")
bullet(doc, "**6. 中欧班列（西安）全国第一**：内陆开放/国际陆港枢纽。")
bullet(doc, "**7. 旅游3.27亿人次、总花费4026亿**：历史文化大IP（文旅强）。")
bullet(doc, "**8. 常住1323.63万、+6.87万、城镇化81.01%**：人口持续净流入。")
bullet(doc, "**9. 民营经济占GDP 50.5%**：民企活力较好。")
bullet(doc, "**10. 高技术制造投资+15.4%、科服+20.3%**：有效投资的含科量。")
bullet(doc, "**11. 战略新兴产业+6.3%**：硬科技落地。")
bullet(doc, "**12. 全国重点实验室30家、高企1.5万家**：科教资源、国家创新体系。")
bullet(doc, "**13. 中欧班列/一带一路**：西安港集装箱/空/铁双枢纽。")
bullet(doc, "**14. 固投-15.3%、基础设施-56.3%**：地产/基建调整拖累。")
bullet(doc, "**15. 财政-2.3%**：收入下滑、政策空间受束（依赖硬科技/文旅税源）。")
para(doc, "", size=10, space_after=2)

# ---- 五、对照表 ----
heading1(doc, "五、2025年GDP目标 vs 实际：对照表")
para(doc, "2025年报告中列出的主要预期目标，与2025年实际完成：")
tb = [
    ["指标", "2025年目标", "2025年实际（公报/报告）", "达标判定"],
    ["GDP增速", "5.5%左右", "+4.7%（13902.67亿）", "未达标"],
    ["规上工业增加值", "6.5%", "+5.7%（公报）/6.1%（报告）", "略欠"],
    ["固定资产投资", "1%", "-15.3%", "大幅回落"],
    ["社会消费品零售总额", "5%左右", "+5.3%（5721.21亿）", "达标"],
    ["进出口总值", "10%左右", "+21.1%（4987.9亿）", "超预期"],
    ["一般公共预算收入", "3%以上", "-2.3%（979.35亿）", "未达标"],
    ["城镇新增就业", "18万人", "18.1万人", "达标"],
    ["CPI", "2%左右", "+0.5%", "温和"],
]
table(doc, tb[0], tb[1:], widths=[3.2, 3.2, 4.8, 4.0])
para(doc, "", size=10, space_after=2)
para(doc, "**要点**：社零/就业/外贸达标或超预期；GDP/财政/固投未达标——西安\u201c消费外贸强、总量/投资/财政弱\u201d。")

# ---- 六、结构归因 ----
heading1(doc, "六、2025年增长是谁撑起来的？（结构归因）")
heading2(doc, "6.1 三次产业：三产主导、二产稳")
para(doc, "**三产（+4.6%、占比69.0%）主导**：文化旅游/商贸/金融/科技服务。二产+4.8%（占比28.6%）、一产+3.6%（2.4%）。")
heading2(doc, "6.2 工业：自动化/设备制造稳")
para(doc, "规上工业+6.1%；铁路/船舶/航天+8.4%、专用设备+7.6%。装备制造+6.6%。新能源汽车/半导体结构优化。")
heading2(doc, "6.3 消费企稳")
para(doc, "社零+5.3%；家电/通讯/新能源汽车零售高增、通讯器材+31.5%、家电+30.8%、新能源车零售+34.9%。")
heading2(doc, "6.4 外贸强")
para(doc, "进出口+21.1%、出口+25.8%；新三样出口高（电动+66%）。西部开放亮点。")
heading2(doc, "6.5 投资弱")
para(doc, "固投-15.3%、基础设施-56.3%；工业+2.4%、高技术+15.4%——总量下滑但结构含新。")
para(doc, "**一句话归因**：西安2025年\u201c**进出口/消费/硬科技制造**\u201d为主引擎，\u201c固投/地产/财政\u201d走弱——靠外贸+文旅+硬科技驱动、投资收缩的省会。")

# ---- 七、财政 ----
heading1(doc, "七、预算与财政的\u201c含金量\u201d")
para(doc, "**地方一般公共预算收入979.35亿元、-2.3%**；税收入701.77亿、-1.8%；一般公共预算支出1513.02亿元、-3.1%。")
bullet(doc, "**收入承压下滑**：-2.3%，反映地产/土地/企业利润与一般性支出收缩；财政质量待修。")
bullet(doc, "**民生/科创投入**：民生占比较高、科技支出（硬科技）部分支撑。")
bullet(doc, "**债务防范**：化解隐性债务（2024年1040亿）、土地收入走弱。")
para(doc, "**财政含义**：西安\u201c收入下滑、支出收缩、政策空间受限\u201d，需靠硬科技/文旅/新制造创造税源。")

# ---- 八、民生 ----
heading1(doc, "八、民生底账：人口、收入与城乡")
para(doc, "西安\u201c十四五\u201d人口稳增：**常住人口1323.63万人、比上年+6.87万；城镇化率81.01%（+0.58pct）**；居民人均可支配收入47496元、+5.4%。")
bullet(doc, "**收入**：人均可支配收入47496元、+5.4%（快于GDP）；城镇56324元、+4.9%，农村22425元、+5.9%（城乡收入比2.51、缩小）。")
bullet(doc, "**就业**：城镇新增就业18.1万（目标18万）；就业总体稳。")
bullet(doc, "**社保**：养老/医保完善；低保/医保健全。")
bullet(doc, "**人口**：常住+6.87万、城镇化81.01%；高教/硬科技吸引青年。")
para(doc, "**民生含义**：西安\u201c人口/收入/就业\u201d稳，城镇化持续提升，科教/文旅强（回流本地人口）。")

# ---- 九、城乡 ----
heading1(doc, "九、城镇与农村：格局与均衡")
para(doc, "**城镇化率81.01%**，西安高度城镇化；城乡居民收入比2.51（-0.02），县域追赶。")
bullet(doc, "**城市**：副省级/都市圈核心，文化/科技/交通枢纽；轨道交通/西安东站扩建。")
bullet(doc, "**农村**：农业生产稳（粮食141.5万吨）；县域（周至/蓝田等）文旅农业结合。")
para(doc, "**城乡均衡**：西安\u201c城市强、县域追赶\u201d，农文旅+城郊融合。")

# ---- 十、人口流 ----
heading1(doc, "十、人口流入与流出")
para(doc, "**西安净流入**：2025年常住+6.87万、自然增长平稳。在全国人口总体收缩背景下，靠科教/人才/产业吸引。")
bullet(doc, "**流入**：高校（西交大/西工大等）+硬科技/新能源汽车/文旅岗位。")
bullet(doc, "**竞争**：与其他副省级省会（郑州/成都等）争夺；靠科教、文化、气候差异化。")
para(doc, "人口方向决定中长期需求与增长；西安的\u201c科教+文旅+硬科技\u201d是其长逻辑。")

# ---- 十一、物价 ----
heading1(doc, "十一、物价与货币环境")
para(doc, "**2025年西安CPI上涨0.5%**：食品-0.7%、交通通信-1.6%；教育文化娱乐+3.7%、其他+9.0%。")
bullet(doc, "**物价**：温和上调但整体低位，食品/交通负、服务价格上行。")
bullet(doc, "**货币/流动性**：金融多元、贷款/储蓄稳；中欧班列、国际港物流活跃。")
para(doc, "**物价含义**：西安\u201c通胀温和、结构分化\u201d，消费修复由服务/文旅拉；总需求压力可控。")

# ---- 十二、区域一体化 ----
heading1(doc, "十二、区域一体化：西安在\u201c西安都市圈+关中平原城市群+一带一路枢纽\u201d里的位置")
para(doc, "西安处于**西安都市圈+关中平原城市群+国家中心城市**核心：既是西北龙头，也是\u201c一带一路\u201d内陆开放枢纽（中欧班列集结中心）。")
bullet(doc, "**西安都市圈**：西咸新区一体化；市域铁路、关中协同。")
bullet(doc, "**关中平原城市群**：作为主轴，带动宝鸡/咸阳/西延等。")
bullet(doc, "**一带一路/枢纽**：中欧班列（西安）全国第一、西安国际港；空铁双港、跨境电商。")
bullet(doc, "**关中-天水**：沿丝绸之路经济带节点城市。")
para(doc, "**区域含义**：西安作为\u201c西部开放枢纽+都市圈+一带一路\u201d，靠硬科技+文旅+走廊带带动关中/丝路。")

# ---- 十三、五条主线 ----
heading1(doc, "十三、未来5—10年最值得观察的五条主线")
bullet(doc, "**主线1｜硬科技/科教**：研发强度5.56%、半导体/光子/超导/商业航天。能否把科教优势变成产业/税收。")
bullet(doc, "**主线2｜新能源汽车/新质制造**：比亚迪、新能车105万辆；能否在新能源/半导体链站稳。")
bullet(doc, "**主线3｜一带一路/外贸**：进出口+21%、中欧班列全国第一。能否持续扮演西部开放枢纽。")
bullet(doc, "**主线4｜文旅/历史文化IP**：旅游3.27亿人次、国际游客。能否依托大IP做大文旅消费。")
bullet(doc, "**主线5｜都市圈/人口**：常住+6.87万、城镇化约81%。能否把\u201c科教+文旅\u201d变成长效人口/内需引擎。")
para(doc, "这五条，是西安从\u201c历史文化+科教城市\u201d走向\u201c硬科技+开放枢纽\u201d的\u201c主赛道\u201d。")

# ---- 十四、结论 ----
heading1(doc, "十四、最终结论：西安在\u201c硬科技+新能源汽车+文旅+科教\u201d里的增长逻辑")
para(doc, "西安2025年，本质上是\u201c**进出口/消费/硬科技强、固投/财政/总量弱**\u201d的答卷：GDP13902.67亿、+4.7%、进出口4987.9亿+21.1%、社零+5.3%、规上工业+5.7%、研发强度5.56%、新能源汽车105万辆、常住1323.63万。")
para(doc, "只要硬科技/新能源汽车/文旅/一带一路持续，西安就站在\u201c硬科技+开放枢纽\u201d增长极；若固投/财政/基础设施承压，西安需承受\u201c外需消费强、总量投资弱\u201d的结构挑战。")
para(doc, "最稳观察信号：**一盯硬科技/科教（引擎）、二盯进出口/一带一路（开放）、三盯新能源汽车/产业（制造）、四盯文旅/消费（内需）、五盯人口/城区（长期）。**西安，是\u201c历史文化+硬科技\u201d的新样本。")

# ---- 附录A ----
heading1(doc, "附录A：主要资料来源与核验路径")
bullet(doc, "西安市2025年政府工作报告（2025年3月）——目标来源。")
bullet(doc, "《2025年西安市国民经济和社会发展统计公报》（市统计局）——GDP、工业、外贸、人口实值。")
bullet(doc, "2026年西安市政府工作报告（2026年）——2025追认/硬科技/新能源汽车/文旅。")
bullet(doc, "西安海关、市财政局（外贸/财政）。")
heading2(doc, "核验说明")
para(doc, "本报告涉及数据以统计公报/官方口径为准；\u201c硬科技/新能源汽车/文旅/一带一路\u201d等以官方口径为准。")

# ---- 附录B ----
heading1(doc, "附录B：建议建立的年度跟踪仪表盘")
para(doc, "建议把下面10个\u201c测脉搏\u201d指标做成年度跟踪表：")
dash = [
    ["#", "指标", "2025基值", "为什么重要"],
    ["1", "GDP增速", "+4.7%（13902.67亿）", "总量与方向"],
    ["2", "规模以上工业增速", "+5.7%", "制造底盘"],
    ["3", "进出口增速", "+21.1%（4987.9亿）", "开放/外贸"],
    ["4", "新能源汽车产量", "105.14万辆", "新能源车产业"],
    ["5", "硬科技/研发强度", "技术4500亿/5.56%", "硬科技"],
    ["6", "固定资产投资/基建", "-15.3% / -56.3%", "投资结构"],
    ["7", "社零增速", "+5.3%（5721.21亿）", "内需消费"],
    ["8", "常住人口/城镇化", "1323.63万 / 81.01%", "人口与城市"],
    ["9", "地方财政收入", "-2.3%（979.35亿）", "财政质量"],
    ["10", "CPI", "+0.5%", "物价"],
]
table(doc, dash[0], dash[1:], widths=[0.8, 4.6, 4.4, 4.0])
para(doc, "把这10个指标连起来看，进出口/硬科技/新能源车（3/4/5）、固投/财政（6/9），都说明西安在真正换挡。")

# ============ 保存 ============
out = "/Users/x/Desktop/content-prod-lab/reports/西安市_2025年政府工作报告_深度研究_2026-08-13.docx"
doc.save(out)
print("SAVED:", out)
print("PARAGRAPHS:", len(doc.paragraphs))
print("TABLES:", len(doc.tables))
