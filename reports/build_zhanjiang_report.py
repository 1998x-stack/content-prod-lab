# -*- coding: utf-8 -*-
"""Build 济宁市2025年政府工作报告 深度研究 DOCX, 参照地级市系列版式。"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DARK = RGBColor(0x1F, 0x2A, 0x44)
RED = RGBColor(0xB0, 0x1E, 0x2E)
GRAY = RGBColor(0x59, 0x60, 0x69)
BLUE = RGBColor(0x1F, 0x4E, 0x79)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
HEADER_FILL = "1F2A44"
K = "微软雅黑"


def set_run(run, size=11, bold=False, color=DARK, font=K):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    run.font.name = font
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), font)
    rFonts.set(qn("w:hAnsi"), font)
    rFonts.set(qn("w:eastAsia"), font)


def para(doc, text, size=11, bold=False, color=DARK, align=None,
         space_after=6, space_before=0, font=K):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        set_run(r, size=size, bold=bold or (i % 2 == 1), color=color, font=font)
    return p


def heading1(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(16)
    pf.space_after = Pt(8)
    r = p.add_run(text)
    set_run(r, size=16, bold=True, color=RED)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "B01E2E")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def heading2(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(10)
    pf.space_after = Pt(4)
    r = p.add_run(text)
    set_run(r, size=13, bold=True, color=BLUE)
    return p


def table(doc, headers, rows, widths=None, font_size=9.5):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_run(r, size=font_size, bold=True, color=WHITE)
        tcPr = hdr[i]._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), HEADER_FILL)
        tcPr.append(shd)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(val))
            set_run(r, size=font_size, color=DARK)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    return t


def quote_box(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(8)
    pf.left_indent = Pt(12)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), "B01E2E")
    pBdr.append(left)
    pPr.append(pBdr)
    r = p.add_run(text)
    set_run(r, size=11, color=DARK, font="楷体")
    return p


def bullet(doc, text, size=10.5):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.7)
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        set_run(r, size=size, bold=(i % 2 == 1), color=DARK)
    return p


# ================================================================= 文档主体
doc = Document()
sec = doc.sections[0]
sec.page_width = Cm(21)
sec.page_height = Cm(29.7)
sec.top_margin = Cm(2.4)
sec.bottom_margin = Cm(2.2)
sec.left_margin = Cm(2.5)
sec.right_margin = Cm(2.5)

st = doc.styles["Normal"]
st.font.name = K
st.font.size = Pt(11)
st.element.rPr.rFonts.set(qn("w:eastAsia"), K)


# ---- 封面 ----
para(doc, "湛江市2025年政府工作报告", size=24, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=60, space_after=4)
para(doc, "深度研究与\u201c容易被忽视的细节\u201d分析", size=16, bold=True, color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
para(doc, "从\u201c钢铁石化临港、四绿一蓝、省域副中心、粤西枢纽\u201d重新理解湛江", size=12, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)
para(doc, "\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501", color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
para(doc, "研究对象：2025年湛江市政府工作报告及同期财政、国民经济和社会发展统计资料、2026年官方复盘", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
para(doc, "标定日期：2026-08-14", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
doc.add_page_break()

# ---- 目录 ----
catalog = [
    "执行摘要",
    "一、研究口径与阅读方法",
    "二、先看湛江的特殊底盘：钢铁石化、四绿五蓝海洋产业、省域副中心",
    "三、最关键的宏观错位：工业全省第1强，但固投-13%、外资-96%、消费投资双弱",
    "四、容易被忽视、但信号很重的细节（15条）",
    "五、2025年GDP目标 vs 实际：对照表",
    "六、2025年增长是谁撑起来的？（结构归因）",
    "七、预算与财政的\u201c含金量\u201d",
    "八、民生底账：人口、收入与城乡",
    "九、城镇与农村：格局与均衡",
    "十、人口流入与流出",
    "十一、物价与货币环境",
    "十二、区域一体化：湛江在大湾区、北部湾、海南自贸港、东盟\u201c四面\u201d里的位置",
    "十三、未来5—10年最值得观察的五条主线",
    "十四、最终结论：湛江在\u201c钢铁石化+海洋经济+省域副中心\u201d里的增长逻辑",
    "附录A：主要资料来源与核验路径",
    "附录B：建议建立的年度跟踪仪表盘",
]
para(doc, "目　录", size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=10, space_after=10)
for item in catalog:
    para(doc, item, size=12, space_after=4)
doc.add_page_break()

# ---- 执行摘要 ----
heading1(doc, "执行摘要")
para(doc, "**核心判断**　2025年湛江最亮的是\u201c规上工业+10.7%（全省第1）、GDP 3952.94亿/+4.5%（增速全省第4）、税收+11.7%（全省第1）\u201d、\u201c但固定投资-13%、房地产-26%、实际利用外资-96%\u201d、\u201c社零+2.8%、进出口+1.9%、CPI-0.2%\u201d。这说明湛江在\u201c亿吨经济+海洋强市\u201d战略中，**工业投资强、但民企投资/外资/地产洞**。")
para(doc, "把2025年目标（GDP+5%、规上+6%、固投+5%、社零+5%、进出口+5%、财收+3%）、2025年实际（GDP+4.5%-规上+10.7%、固投-13%、社零+2.8%、进出口+1.9%、财收+6.9%）趋势看，湛江是\u201c港口临港+钢铁石化+家电+农业水产\u201d路径：**钢铁、石化、能源、海洋（四绿一蓝）**是支柱（工业对GDP贡献47%）。")
para(doc, "最容易记住的一句话：**湛江是\u201c海洋城市、重化工业、省域副中心\u201d，靠\u201c钢铁石化家电出口+农林水\u201d驱动，投资增长主要靠工业（占投资57.1%）。**观察湛江，与其只看\u201cGDP 3953亿\u201d，不如看\u201c规上+10.7%全省第1、化学制造+24.9%、钢铁/家电出口、重点港泊159个\u201d。")
# ---- 一、研究口径与阅读方法 ----
heading1(doc, "一、研究口径与阅读方法")
para(doc, "本报告以\u201c湛江市政府工作报告（2025年，湛江市市长作）\u201d为起点，把\u201c2025年GDP目标（5%）\u201d与\u201c官方2025年（3952.94亿/+4.5%）\u201d并置对照，用\u201c2025年湛江市统计公报\u201d和\u201c2026年政府工作报告\u201d横向核验。")
para(doc, "口径提示：\u201cGDP、规上工业增速\u201d按可比价（实际）；\u201c投资、消费、进出口、财政\u201d按现价（名义）。\u201c常住人口\u201d用统计公报（712.29万）、城镇化率49.63%。")
para(doc, "指标体系（与研究口径一致）：**总量与增速、产业动能（工业/临港石化海洋）、外贸（靠工业出口）、财政、民生与人口**。")
para(doc, "特别提示（不吃老本）：湛江2024年GDP最终核实3839.93亿、+1.2%；2025年+4.5%、突破3900亿；它不是\u201c只有旅游\u201d——**钢铁、石化（巴斯夫）、能源（核电）、家电、海洋水产\u201d才是真正底色，工业是拉动主力。")
# ---- 二、先看湛江的特殊底盘 ----
heading1(doc, "二、先看湛江的特殊底盘：钢铁石化、四绿五蓝海洋产业、省域副中心")
para(doc, "湛江地处粤西雷州半岛，是**中国大陆最南端沿海城市、广东省域副中心城市、海洋强市**；是国家战略\u201c临港重化工业基地\u201d（宝钢湛江、中科炼化/巴斯夫一体化、廉江核电）。2025年GDP 3952.94亿（逼近4000亿）、常住712.29万、城镇化率49.63%。")
para(doc, "五个底盘名词，先立框架：")
bullet(doc, "**钢铁**　宝钢湛江钢铁基地、千万吨级——\u201c临港重化\u201d。")
bullet(doc, "**石化**　中科炼化、巴斯夫一体化基地（东海岛），化学制造+24.9%。")
bullet(doc, "**绿色能源（四绿一蓝）**　核电（廉江核电站）、风电、光伏、生物质\u2014\u2014\u201c四绿\u201d。")
bullet(doc, "**蓝色海洋经济**　海洋渔业（水产品137.42万吨/全国水产大市）、港口（159个泊位）\u2014\u2014\u201c一蓝\u201d。")
bullet(doc, "**农业/临港**　蔗糖（广东最大）、热带水果、对虾贝类；雷州半岛、徐闻菠萝等。")
para(doc, "这五根（钢铁+石化+能源+海洋+农业临港）构成湛江独特底盘：**左手重化工（钢铁石化），右手海洋+绿色能源**。理解湛江，先理解\u201c港口+临港重化+海洋强市、省域副中心\u201d。")
# ---- 三、最关键的宏观错位 ----
heading1(doc, "三、最关键的宏观错位：工业全省第1强，但固投-13%、外资-96%、消费投资双弱")
para(doc, "2025年湛江最需要辨析的一组\u201c错位\u201d：**规上工业+10.7%（全省第1）、化学制造+24.9%、税收+11.7%（全省第1）强，但固投-13%、房地产开发-26%、实际利用外资-96%、社零+2.8%、CPI-0.2%**。")
para(doc, "为什么\u201c工业这么强\u201d，投资与外资却塌？三解释：")
para(doc, "**其一，工业（尤其是钢石）有一轮大扩张**　规上+10.7%（全省第1）、化学制造+24.9%、工业对GDP贡献47%、承接产业转移主平台+37.1%——\u201c临港重化+转移\u201d强。")
para(doc, "**其二，投资、地产是所有断崖**　固投-13%（工业投资-10%、基建+6%）、房地产-26%、商品房销售-4.5%；\u201c投资靠工业、地产空转\u201d。")
para(doc, "**其三，外资、消费、外贸弱**　实际利用外资-96%（1.51亿）、外资断崖；社零+2.8%、进出口+1.9%、CPI-0.2%、PPI弱——\u201c需求弱、外资出逃\u201d。")
para(doc, "小结：湛江2025年是\u201c**强工业、弱地产外资、内需缓**\u201d：钢化/能源/农业强，房地产投资、外资、消费、物价弱。")

# ---- 四、容易被忽视、但信号很重的细节（15条） ----
heading1(doc, "四、容易被忽视、但信号很重的细节（15条）")
bullet(doc, "**1.规上工业+10.7%全省第1/工业对GDP贡献47%**\u201c钢铁、石化、能源拉动（化学制造+24.9%）。\u201d")
bullet(doc, "**2.国内税收476.1亿/+11.7%（全省第1）**\u201c税收创高、含金量足（营收-税收差）。\u201d")
bullet(doc, "**3.园区规上+13.1%、产业转移主平台+37.1%**\u201c承接大湾区产业转移（对接47个项目）。\u201d")
bullet(doc, "**4.固定资产投资-12%（工业占57.1%）**\u201c投资靠工业（钢铁/石化/能源）。\u201d")
bullet(doc, "**5.房地产开发投资-26%、商品房销售-4.5%**\u201c地产深度调整、空置。\u201d")
bullet(doc, "**6.实际利用外资-96%（仅1.51亿）**\u201c外资断崖、开放短板。\u201d")
bullet(doc, "**7.进出口628.57亿/+1.9%**\u201c钢材+25.7%（汽车/家电出口弱）。\u201d")
bullet(doc, "**8.社零1620.06亿/+2.8%（社零全省第6）**\u201c消费弱修复、餐饮+0.9%。\u201d")
bullet(doc, "**9.水产（对虾/网箱/大型养殖平台）已发**\u201c\u2018蓝色粮仓\u2019、向海洋要粮食。\u201d")
bullet(doc, "**10.廉江核电/海上风电/氢能——\u2018四绿一蓝\u2019**\u201c能源、低碳转型。\u201d")
bullet(doc, "**11.巴斯夫一体化（绿色石化）一期投产**\u201c外资大项目龙头（但整体外资-96%）。\u201d")
bullet(doc, "**12.新引进项目-476个/+48.8%、专精特新170家**\u201c招商、民营生态、新质。\u201d")
bullet(doc, "**13.居民收入32336元/+9.4%**\u201c农村+5.9%>城镇+3.9%——消费力偏弱。\u201d")
bullet(doc, "**14.CPI-0.2%（交通通信-3.1%）**\u201c低通胀、需求疲。\u201d")
bullet(doc, "**15.常住712.29万/+0.21万**\u201c农业大市但人口增量微弱、城镇化49.63%。\u201d")
# ---- 五、2025年GDP目标 vs 实际：对照表 ----
heading1(doc, "五、2025年GDP目标 vs 实际：对照表")
table(doc,
    ["指标", "2025年目标", "2025年实际", "达成评价"],
    [
        ["地区生产总值(GDP)", "增长5%", "3952.94亿/4.5%", "差0.5pct"],
        ["规上工业增加值", "增长6%", "+10.7%(全省第1)", "大幅超额"],
        ["固定资产投资", "增长5%", "-13%", "大幅下行"],
        ["社会消费品零售总额", "增长5%", "1620.06亿/+2.8%", "不及目标"],
        ["进出口总额", "增长5%", "628.57亿/+1.9%", "不及目标"],
        ["一般公共预算收入", "增长3%", "215.10亿/+6.9%", "超额"],
        ["居民收入", "与GDP同步", "32336元/+5.0%", "略高于GDP"],
    ],
)
para(doc, "注：GDP、规上工业按可比价。**规上工业（+10.7%）、财收（+6.9%）超额**；**GDP（+4.5%差0.5pct）、固投（-13%）、社零（+2.8%）、进出口（+1.9%）不及目标**。")
para(doc, "拆读：**工业（+10.7%全省第1）、税收、承接转移、化料是亮色**；**固投（-13%）、外资（-96%）、消费（+2.8%）、CPI（-0.2%）**是短板——\u201c工业强、需求弱\u201d，是\u201c临港重化+海洋\u201d样本。")

# ---- 六、2025年增长是谁撑起来的？（结构归因） ----
heading1(doc, "六、2025年增长是谁撑起来的？（结构归因）")
para(doc, "把湛江GDP的4.5%拆开：三次产业分别增3.5%、4.7%、4.8%（结构18.6：32.6：48.8）。第三产业（服务业）贡献率51.6%、第二产业贡献33.0%、第一产业15.4%——**三产带动但工业最猛**。")
para(doc, "2026年湛江强调\u201c十五五、省域副中心、海洋强市、新能源\u201d，聚焦**钢铁石化升级、新能源（四绿一蓝）、水产海洋经济、临港、承接大湾区转移**——核心是\u201c重工+海洋强市\u201d。")
para(doc, "**第二产业（工业）**：规上+10.7%（全省第1）、化学制造+24.9%、工业对GDP贡献47%、园区+13.1%、主平台+37.1%——\u201c临港重化强\u201d。")
para(doc, "**第三产业（服务业）**：+4.8%（商贸、物流、港口、旅游总收入+11.1%）——\u201c服务+临港港口\u201d。")
para(doc, "**第一产业（农业/海洋水产）**：+3.5%（粮食159.61万吨、糖蔗、水产品137.42万吨/全国水产大市）——\u201c农业海洋稳\u201d。")
para(doc, "一句话归因：**2025年湛江增长\u201c靠规上工业（钢铁石化）+服务业+农林水产\u201d**，房地、外资、消费弱；\u201c工业/海洋\u201d是核心。")

# ---- 七、预算与财政的\u201c含金量\u201d ----
heading1(doc, "七、预算与财政的\u201c含金量\u201d")
para(doc, "2025年湛江**一般公共预算收入215.10亿元（+6.9%）**；税收收入131.26亿（+6.1%）、国内税收476.1亿（+11.7%，全省第1）；一般公共预算支出568.16亿（+6.6%）、民生占81.7%。")
bullet(doc, "税收+11.7%国内、地方收入+6.9%——\u201c工业税收旺\u201d。")
bullet(doc, "民生支出占81.7%（教育/社保/医疗）。")
bullet(doc, "金融支撑：存款+5.4%、贷款+8.9%、存贷比87.49%（创历史新高）——信贷密集支持临港/工业/农业。")
para(doc, "**财政含金量小结**：财收+6.9%（高GDP增速）、靠钢铁/石化/家电工业税、民生81.7%；财政对\u201c钢铁石化、海洋、承接转移、民生\u201d投入加大。")
# ---- 八、民生底账：人口、收入与城乡 ----
heading1(doc, "八、民生底账：人口、收入与城乡")
para(doc, "2025年湛江**居民人均可支配收入32336元（+5.0%）**，其中城镇39943元（+3.9%）、农村25145元（+5.9%），城乡比1.59（小幅收窄）。就业：城镇新增就业6.68万人。")
para(doc, "人口画像：**常住712.29万/+0.21万、城镇化率49.63%（+1.04pct）**；农业人口仍约一半，但城镇化率提升快（+1.04pct）。")
para(doc, "民生投入：新增公办学位、居民医保住院待遇提高近10pct、保障房2326套——民生投入扎实但收入/消费偏低。")

# ---- 九、城镇与农村：格局与均衡 ----
heading1(doc, "九、城镇与农村：格局与均衡")
para(doc, "湛江城镇化率49.63%，仍为粤西农业大市；农村收入增速（+5.9%）高于城镇（+3.9%），**城乡比缩至1.59**，县域为蔗/水/果产业带；乡村振兴与海洋牧场并进。")
para(doc, "农业底盘：**粮食159.61万吨、糖蔗914.16万吨（广东最大）、蔬菜503.5万吨、园林水果361.1万吨、猪牛羊禽肉56.48万吨、水产品137.42万吨**——\u201c两高一优、向海洋要粮\u201d。")
para(doc, "一句话：\u201c湛江是果蔬蔗糖/水产大市、城镇化率刚过半\u201d，但\u201c农业海洋强、县域经\u201d受转移带动提升。")

# ---- 十、人口流入与流出 ----
heading1(doc, "十、人口流入与流出")
para(doc, "湛江常住712.29万（净增+0.21万）、城镇化49.63%；\u201c农业大市外流、但主城区（霞山/赤坎/经开区）+临港\u201d吸引人才/工人，本地人口较多流出赴珠三角。")
para(doc, "结构观察：**出生率10.50‰、自然增长+3.93‰**（粤西人口自然增长为正）；但青壮年外出打工多。")
para(doc, "2026年目标：城镇新增就业4.5万、引进博士275名、高技能人才48.5万等；湛江靠\u201c临港重化+海洋+大学城\u201d聚人。")

# ---- 十一、物价与货币环境 ----
heading1(doc, "十一、物价与货币环境")
para(doc, "2025年湛江**CPI-0.2%**（交通通信-3.1%、居住-0.6%；衣着+2.6%）——\u201c低通胀、需求弱\u201d。")
bullet(doc, "信贷扩张：存款+5.4%、贷款+8.9%、存贷比87.49%——宽信用支持临港/工业/农业。")
bullet(doc, "消费：社零+2.8%、餐饮+0.9%、商品零售+3.1%——缓慢修复、温。")
para(doc, "货币环境判断：**宽信用、CPI-0.2%**；湛江靠\u201c工业+基建+海洋经济\u201d稳需求（2026 CPI必温和回升）。")

# ---- 十二、区域一体化：湛江的位置 ----
heading1(doc, "十二、区域一体化：湛江在大湾区、北部湾、海南自贸港、东盟\u201c四面\u201d里的位置")
para(doc, "湛江是**国家战略\u201c广东省域副中心城市\u201d、粤港澳大湾区西向门户、对接海南自贸港的前沿、北部湾城市群节点、面向东盟（RCEP）的海上门户**。")
bullet(doc, "**对接大湾区**　承接珠三角产业转移（产业转移主平台+37.1%、对接47个项目）。")
bullet(doc, "**连接海南**　琼州海峡、湛江海港、面向海南自贸港（货物中转/水产品）。")
bullet(doc, "**面向东盟/北部湾**　RCEP（对RCEP+15.3%）、港口159个泊位、并行港（湛江-洋浦）。")
para(doc, "一句话：**湛江在\u201c大湾区、海南、北部湾、东盟\u201d四面交汇的中枢，最核心是\u201c海洋+临港重化\u201d**；区位是最大优势（联通西南/华南/东盟）。")

# ---- 十三、未来5—10年最值得观察的五条主线 ----
heading1(doc, "十三、未来5—10年最值得观察的五条主线")
bullet(doc, "**主线一：临港重化升级（钢铁/石化/新能源）**\u201c钢铁+石化（巴斯夫）+核电/风电\u201d能否成世界级绿色石化/临港基地。")
bullet(doc, "**主线二：海洋经济（蓝色粮仓）**\u201c对虾/水产/网箱、海上牧场\u201d、深海装备。")
bullet(doc, "**主线三：承接大湾区产业转移**\u201c产业主平台+37%\u201d能否再造新工业（装备/家电）。")
bullet(doc, "**主线四：省域副中心+枢纽（海口/铁路/港）**\u201c418.7公里铁路、机场吞吐312万、港口159泊位\u201d大数据联通。")
bullet(doc, "**主线五：绿色能源（碳中和）+外资回暖**\u201c巴斯夫、核电、风电\u201d、外资-27%如何止血回升。")

# ---- 十四、最终结论 ----
heading1(doc, "十四、最终结论：湛江在\u201c钢铁石化+海洋经济+省域副中心\u201d里的增长逻辑")
para(doc, "把2025年闭环打穿：**湛江是\u201c临港重化、海洋强市、省域副中心\u201d**：GDP 3952.94亿/+4.5%、规上+10.7%（全省第1）、税收+11.7%、水产/海洋/四绿一蓝。")
para(doc, "湛江不是\u201c只有钢铁\u201d——它是**钢铁+石化+能源（核电）+海洋+家电**的复合，靠\u201c工业+海洋\u201d驱动；但房地产投资、外资、消费、物价弱。")
para(doc, "一句话结论：**湛江是\u201c亿吨大港、临港重化、海洋之都\u201d；观察它先看\u201c工业/钢铁石化的出口、外资、港口吞吐、向海洋\u201d，再看\u201c地产、消费、物价\u201d。**它是\u201c工业强、需求弱、开放待活\u201d的粤西样本。")

# ---- 附录A ----
heading1(doc, "附录A：主要资料来源与核验路径")
bullet(doc, "《2025年湛江市政府工作报告》（2025年3月，2025年目标、2024年回顾3839.93亿）")
bullet(doc, "《2025年湛江市国民经济和社会发展统计公报》（湛江市统计局,国家统计局湛江调查队，2026-04-23，2025年实际）")
bullet(doc, "《2026年湛江市政府工作报告》（2026年3月，复盘+2026年目标）")
bullet(doc, "湛江市人民政府/统计局、湛江日报等")

# ---- 附录B ----
heading1(doc, "附录B：建议建立的年度跟踪仪表盘")
bullet(doc, "GDP总量/增速 vs 全年目标。")
bullet(doc, "规上工业（钢铁/石化/能源）增速、四大类。")
bullet(doc, "承接转移、主平台。")
bullet(doc, "固定资产/工业/房地产投资。")
bullet(doc, "港口吞吐、铁路、机场。")
bullet(doc, "社零、餐饮、以旧换新。")
bullet(doc, "进出口、钢材/家电出口、外资。")
bullet(doc, "地方一般公共预算/税收/民生%。")
bullet(doc, "常住/城镇化、农业人口（户籍-常住）。")
bullet(doc, "CPI、存贷款、化石/绿色。")

# ---------------- 保存 ----------------
out = "/Users/x/Desktop/content-prod-lab/reports/湛江市_2025年政府工作报告_深度研究_2026-08-14.docx"
doc.save(out)
print("SAVED OK: 湛江市", out)
