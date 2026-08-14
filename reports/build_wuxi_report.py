# -*- coding: utf-8 -*-
"""Build 无锡市2025年政府工作报告 深度研究 DOCX, 参照地级市系列版式。"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DARK = RGBColor(0x1F, 0x2A, 0x44)
RED = RGBColor(0xB0, 0x1E, 0x2E)
GRAY = RGBColor(0x59, 0x60, 0x69)
BLUE = RGBColor(0x1F, 0x4E, 0x79)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
HEADER_FILL = "1F2A44"
K = "微软雅黑"


def set_run(run, size=11, bold=False, color=DARK, font=K):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    run.font.name = font
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), font)
    rFonts.set(qn("w:hAnsi"), font)
    rFonts.set(qn("w:eastAsia"), font)


def para(doc, text, size=11, bold=False, color=DARK, align=None,
         space_after=6, space_before=0, font=K):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        set_run(r, size=size, bold=bold or (i % 2 == 1), color=color, font=font)
    return p


def heading1(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(16)
    pf.space_after = Pt(8)
    r = p.add_run(text)
    set_run(r, size=16, bold=True, color=RED)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "B01E2E")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def heading2(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(10)
    pf.space_after = Pt(4)
    r = p.add_run(text)
    set_run(r, size=13, bold=True, color=BLUE)
    return p


def table(doc, headers, rows, widths=None, font_size=9.5):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_run(r, size=font_size, bold=True, color=WHITE)
        tcPr = hdr[i]._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), HEADER_FILL)
        tcPr.append(shd)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(val))
            set_run(r, size=font_size, color=DARK)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    return t


def quote_box(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(8)
    pf.left_indent = Pt(12)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), "B01E2E")
    pBdr.append(left)
    pPr.append(pBdr)
    r = p.add_run(text)
    set_run(r, size=11, color=DARK, font="楷体")
    return p


def bullet(doc, text, size=10.5):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.7)
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        set_run(r, size=size, bold=(i % 2 == 1), color=DARK)
    return p


# ================================================================= 文档主体
doc = Document()
sec = doc.sections[0]
sec.page_width = Cm(21)
sec.page_height = Cm(29.7)
sec.top_margin = Cm(2.4)
sec.bottom_margin = Cm(2.2)
sec.left_margin = Cm(2.5)
sec.right_margin = Cm(2.5)

st = doc.styles["Normal"]
st.font.name = K
st.font.size = Pt(11)
st.element.rPr.rFonts.set(qn("w:eastAsia"), K)


# ---- 封面 ----
para(doc, "无锡市2025年政府工作报告", size=24, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=60, space_after=4)
para(doc, "深度研究与\u201c容易被忽视的细节\u201d分析", size=16, bold=True, color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
para(doc, "从\u201c制造强市、物联网之都、集成电路、民营经济、长三角\u201d重新理解无锡", size=12, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)
para(doc, "\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501", color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
para(doc, "研究对象：2025年无锡市政府工作报告及同期财政、国民经济和社会发展统计资料、2026年官方复盘", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
para(doc, "标定日期：2026-08-14", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
doc.add_page_break()

# ---- 目录 ----
catalog = [
    "执行摘要",
    "一、研究口径与阅读方法",
    "二、先看无锡的特殊底盘：制造强市、物联网、集成电路、装备、民营经济与长三角",
    "三、最关键的宏观错位：GDP 16773.94亿/5.1%，工业外贸强但固投地产大跌、消费偏弱",
    "四、容易被忽视、但信号很重的细节（15条）",
    "五、2025年GDP目标 vs 实际：对照表",
    "六、2025年增长是谁撑起来的？（结构归因）",
    "七、预算与财政的\u201c含金量\u201d",
    "八、民生底账：人口、收入与城乡",
    "九、城镇与农村：格局与均衡",
    "十、人口流入与流出",
    "十一、物价与货币环境",
    "十二、区域一体化：无锡在\u201c长三角+上海都市圈+太湖湾科创带\u201d里的位置",
    "十三、未来5—10年最值得观察的五条主线",
    "十四、最终结论：无锡在\u201c智造强市+物联网+民营经济\u201d里的增长逻辑",
    "附录A：主要资料来源与核验路径",
    "附录B：建议建立的年度跟踪仪表盘",
]
para(doc, "目　录", size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=10, space_after=10)
for item in catalog:
    para(doc, item, size=12, space_after=4)
doc.add_page_break()

# ---- 执行摘要 ----
heading1(doc, "执行摘要")
para(doc, "**核心判断**　2025年无锡最显著的是\u201cGDP 16773.94亿元、增长5.1%（达到5%目标）、人均22.3万、江苏第3\u201d、\u201c规上工业+5.9%、进出口8292.76亿/+7.6%（出口+14.2%）、高新技术产业产值+4.3%\u201d、\u201c民营经济增加值11182.18亿/占66.7%\u201d、\u201c但固投3979.11亿/-13.3%、房地产-20.1%、社零+3.1%\u201d、\u201c财收1225.39亿/+2.0%、常住753.74万\u201d。这说明无锡在\u201c制造强市+物联网\u201d的升级中，**外贸、民营、制造强但固投地产消费偏弱**。")
para(doc, "把2025年目标（GDP+5%以上/社零+5.5%左右）、2025年统计（GDP+5.1%达成、规上+5.9%、固投-13.3%、社零+3.1%低于目标、财收+2.0%）、趋势一起看，无锡是\u201c制造+民营+外贸\u201d路径：**集成电路、物联网、高端装备、新能源、民营经济**是支柱；2025年总量16774亿居江苏第2（次于苏州）。")
para(doc, "最容易记住的一句话：**无锡是\u201c长三角制造强市、物联网之都、民营经济标杆\u201d，靠\u201c集成电路+物联网+装备+外贸\u201d增长。**观察无锡，与其只看\u201cGDP 16774亿\u201d，不如看\u201c高新技术产业产值14629.84亿/占54%、进出口+7.6%、民营占66.7%、上市公司218家\u201d。")

# ---- 一、研究口径与阅读方法 ----
heading1(doc, "一、研究口径与阅读方法")
para(doc, "本报告以\u201c无锡市政府工作报告（2025年，赵建军作）\u201d为起点，把\u201c2025年GDP目标（5%以上）\u201d与\u201c官方2025年GDP（16773.94亿元/+5.1%）\u201d并置对照，并用\u201c2025年无锡市统计公报\u201d和\u201c2026年政府工作报告复盘\u201d横向核验。")
para(doc, "口径提示：\u201cGDP、规上工业增加值增速\u201d按可比价（实际增速）；\u201c投资、消费、进出口、财政收入\u201d按现价（名义增速）。涉及\u201c常住人口\u201d用统计公报常住口径（753.74万），城镇化率用官方公布值（83.80%）。")
para(doc, "指标体系（与研究口径一致）：五大象限——**总量与增速（GDP）、产业动能（集成电路/物联网/装备/民营）、外贸、财政质量、民生与人口**。")
para(doc, "特别提示（不吃老本）：无锡2024年GDP约1.62万亿/+5.8%，2025年放缓到+5.1%、固投-13.3%；它不只是\u201c长三角制造\u201d——**物联网（世界物联网博览会）、集成电路（SK海力士/华虹）、民营经济（占2/3）**是真正的底色，要看\u201c制造强、投资弱\u201d的转型矛盾。")
# ---- 二、先看无锡的特殊底盘 ----
heading1(doc, "二、先看无锡的特殊底盘：制造强市、物联网、集成电路、装备、民营经济与长三角")
para(doc, "无锡地处太湖之滨、长三角苏南，是**中国制造业强市、物联网之都（世界物联网博览会）、集成电路重镇、民营经济标杆**。2025年GDP 16773.94亿元、常住人口753.74万，人均22.3万，江苏第2。")
para(doc, "五个底盘名词，先立框架：")
bullet(doc, "**制造强市/集成电路**　无锡是全国集成电路产业高地（SK海力士、华虹半导体、长电科技），高新技术产业产值14629.84亿；物联网之都（国家传感网创新示范区）。")
bullet(doc, "**物联网之都**　物联网产业世界级，世界物联网博览会引起全球关注，\u201c一感两网\u201d、智能传感。")
bullet(doc, "**高端装备/新能源**　装备制造、新能源（光伏）、汽车、高端材料（特钢/纺织）——制造集群广。")
bullet(doc, "**民营经济标杆**　民营经济增加值11182.18亿、占经济总量66.7%，太湖湾\u201c亲清\u201d、上市公司218家。")
bullet(doc, "**外贸强市**　进出口8292.76亿/+7.6%、出口+14.2%，\u201c外贸百强\u201d、高技术产品出口+24.8%。")
para(doc, "这五根支柱（集成电路+物联网+装备+民营+长三角）构成无锡独特底盘：**左手科技制造（集成电路），右手民营+外贸，腹地是长三角**。理解无锡，先理解\u201c苏南民营制造全能冠军\u201d。")

# ---- 三、最关键的宏观错位 ----
heading1(doc, "三、最关键的宏观错位：GDP 16773.94亿/5.1%，工业外贸强但固投地产大跌、消费偏弱")
para(doc, "2025年无锡最需要辨析的一组\u201c错位\u201d：**GDP 5.1%达成5%目标、规上工业+5.9%、进出口+7.6%强，但固投-13.3%、房地产-20.1%、社零+3.1%、CPI持平**。")
para(doc, "为什么\u201c外贸强、制造强、经济却在放缓\u201d？三个解释：")
para(doc, "**其一，工业提质但增速放缓**　规上+5.9%、高新技术产业+4.3%（占54%），但集成电路、光伏、汽车面临价格战/回调——\u201c总量大、结构优、增速常态化\u201d。")
para(doc, "**其二，投资断崖**　固投-13.3%（制造业-11.2%、三产-15.9%、房地产-20.1%、民间-13.3%）——地产与制造业投资双降，是2025年最大短板。")
para(doc, "**其三，消费偏弱**　社零+3.1%（低于5.5%目标），汽车-1.2%（新能车+37.5%分流）；\u201c消费需求温和\u201d。")
para(doc, "小结：无锡2025年是\u201c**稳总量、强外贸科技&民营、弱投资地产消费**\u201d的一年：集成电路、物联网、外贸撑增量，但投资、地产、消费偏软。")

# ---- 四、容易被忽视、但信号很重的细节（15条） ----
heading1(doc, "四、容易被忽视、但信号很重的细节（15条）")
bullet(doc, "**1.民营经济增加值11182.18亿/占66.7%**　民营经济是绝对主力（增速+4.8%）。")
bullet(doc, "**2.进出口8292.76亿/+7.6%、出口+14.2%，高新技术产品出口+24.8%**　外贸表现亮眼。")
bullet(doc, "**3.高新技术产业产值14629.84亿/占规上工业54%**　科技创新成色十足。")
bullet(doc, "**4.国家级专精特新、\u2018灯塔工厂\u2019、供应链大市**（\u201c四上\u201d2977家、\u201c独角兽39家\u201d）。")
bullet(doc, "**5.集成电路/物联网（SK海力士、华虹、国家传感网）**\u201c制造强市的科技底座\u201d。")
bullet(doc, "**6.上市公司218家（全国前列）、直接融资/并购383亿**\u201c资本活水\u201d。")
bullet(doc, "**7.社零4418.47亿/+3.1%、线上零售+37.3%**\u201c电商新消费爆发（新能源汽车+37.5%）。\u201d")
bullet(doc, "**8.实际使用外资30.3亿美元/-11.9%**\u201c外资逆流略降，高技术占比35.9%。\u201d")
bullet(doc, "**9.财收1225.39亿/+2.0%、税收+2.5%、民生占80%**\u201c财政质稳、民生投入大。\u201d")
bullet(doc, "**10.居民收入75200元/+4.1%、城乡比1.69**\u201c全国领先；农村+4.5%快于城镇+3.9%。\u201d")
bullet(doc, "**11.常住753.74万/城镇化83.8%（+0.29pct）**\u201c苏南人口大市、城镇化率高。\u201d")
bullet(doc, "**12.CPI持平（交通通信-1.9%、食品烟酒-0.5%）**\u201c物价稳、需求温和。\u201d")
bullet(doc, "**13.地铁运营141公里/5条线、内环太湖**\u201c城际+地铁1小时都市圈。\u201d")
bullet(doc, "**14.太湖治理Ⅲ类水质、PM2.5 26μg**\u201c生态绿色（太湖水科）\u201d。")
bullet(doc, "**15.旅游总花费1916.34亿/+9.4%、离境退税+5倍**\u201c入境旅游/过境购物恢复。\u201d")
# ---- 五、2025年GDP目标 vs 实际：对照表 ----
heading1(doc, "五、2025年GDP目标 vs 实际：对照表")
table(doc,
    ["指标", "2025年目标", "2025年实际", "达成评价"],
    [
        ["地区生产总值(GDP)", "增长5%以上", "16773.94亿/5.1%", "达成"],
        ["社会消费品零售总额", "增长5.5%左右", "4418.47亿/3.1%", "未达成，差2.4pct"],
        ["一般公共预算收入", "稳定增长", "1225.39亿/2.0%", "达成"],
        ["规上工业增加值", "——", "+5.9%", "稳健"],
        ["进出口总额", "稳中提质", "8292.76亿/+7.6%", "超额"],
        ["固定资产投资", "——", "-13.3%", "大幅下行"],
        ["居民收入", "与经济增长基本同步", "75200元/+4.1%", "基本同步"],
    ],
)
para(doc, "注：GDP、规上工业按可比价；投资、消费、进出口、财收按现价。**GDP（5.1%）、财收（+2.0%）、进出口（+7.6%）达成/超额**，**社零（+3.1%）低于5.5%目标、固投（-13.3%）下行**。")
para(doc, "拆读：**进出口（+7.6%）、民营（66.7%）、高新技术产业（54%）是亮色**，**固投（-13.3%）、地产（-20.1%）、社零（+3.1%）、CPI（0%）偏弱**；\u201cGDP目标5%以上\u201d实际5.1%——\u201c达成但内需弱\u201d，是\u201c制造强市\u201d的典型样本。")

# ---- 六、2025年增长是谁撑起来的？（结构归因） ----
heading1(doc, "六、2025年增长是谁撑起来的？（结构归因）")
para(doc, "把无锡GDP的5.1%拆开：三次产业分别增2.6%、5.0%、5.1%（结构0.9：46.9：52.2）。**第二产业（工业）与第三产业（服务业）双撑，第一产业（渔业/农业）很弱**；民营占66.7%。")
para(doc, "2026年无锡强调\u201c翻万亿2.0\u201d、\u201c十五五\u201d突破2万亿、未来产业（AI核心+20%），聚焦\u201c物联网、集成电路、装备、新能源、未来产业\u201d——核心是\u201c科技驱动、制造强市\u201d。")
para(doc, "**第二产业（工业制造）**：规上工业+5.9%、高新技术产业+4.3%（14629.84亿/占54%）；集成电路、物联网、装备、新能源——\u201c制造强\u201d。")
para(doc, "**第三产业（服务业）**：+5.1%（信息服务/物流/贸易）；金融（存31万亿+/贷款+8.4%）、证券交易额+47.3%、旅游+9.4%——服务业稳。")
para(doc, "**外贸（开放型）**：进出口+7.6%（出口+14.2%）、高技术产品出口+24.8%——外需对冲内需。")
para(doc, "一句话归因：**2025年无锡增长\u201c靠工业（集成电路/物联网）+外贸+民营\u201d**，投资、地产、消费内需偏弱；\u201c制造强、投资弱\u201d是核心特征。")

# ---- 七、预算与财政的\u201c含金量\u201d ----
heading1(doc, "七、预算与财政的\u201c含金量\u201d")
para(doc, "2025年无锡**一般公共预算收入1225.39亿元（+2.0%）**，其中税收960.83亿元（+2.5%），税收占比约78.4%；一般公共预算支出1274.84亿元，**民生领域支出占80%**。")
bullet(doc, "税收结构：税收+2.5%快于财收，增值税/所得税主体税种稳，\u201c增收靠科技制造税\u201d。")
bullet(doc, "民生支出占80%，投向教育、社保、医疗、养老；城乡居民低保提至1167元/月。")
bullet(doc, "金融支撑：存款余额30831.8亿（+6.6%，\u201c存款余额破5.8万亿\u201d\u201c十五五\u201d）、贷款27251.13亿（+8.4%）；上市公司218家、证券交易17.93万亿（+47.3%）。")
para(doc, "**财政含金量小结**：财收+2.0%稳增、税收+2.5%、民生占80%，\u201c稳收入、强民生、活金融\u201d；财政对科创、集成电路、太湖治理投入大。")

# ---- 八、民生底账：人口、收入与城乡 ----
heading1(doc, "八、民生底账：人口、收入与城乡")
para(doc, "2025年无锡**居民人均可支配收入75200元（+4.1%）**，其中城镇83158元（+3.9%）、农村49181元（+4.5%），**城乡比1.69（缩小）**。就业：城镇新增就业14.59万人。")
para(doc, "人口画像：**常住753.74万、户籍523.07万、城镇化率83.8%（+0.29pct）**，是\u201c常住>户籍\u201d的人口净流入城市、苏南富裕高城镇化。")
para(doc, "民生投入：社保普及（养老428.16万/多险种）、医疗救助60.7万人次、保障房11847套——民生保障扎实。")

# ---- 九、城镇与农村：格局与均衡 ----
heading1(doc, "九、城镇与农村：格局与均衡")
para(doc, "无锡常住城镇化率83.8%（江苏前列），城乡高度融合；农村收入增速（+4.5%）高于城镇（+3.9%），**城乡比缩小1.69**，全国领先。")
para(doc, "农业底盘：粮食产量57万吨（+0.4%）、蔬菜127.18万吨、水产9.5万吨、茶叶/瓜果——苏南鱼米（大闸蟹/水蜜桃——阳山、太湖蟹）。")
para(doc, "一句话：\u201c无锡城镇化高、农村收入高、城乡比优秀\u201d，农业体量小但精致（水蜜桃、水产），\u201c以工带农、城乡一体\u201d。")

# ---- 十、人口流入与流出 ----
heading1(doc, "十、人口流入与流出")
para(doc, "无锡常住753.74万（高于户籍523万），是江苏净流入型、产业工人与人才集聚城市；\u201c长三角制造带\u201d吸纳周边与外地人口。")
para(doc, "结构观察：**人口正增长（+0.43%）、青年/科创人才导入**；老龄化低于全省均值、城镇化83.8%高位稳定。")
para(doc, "2026年目标：城镇新增就业13万、集聚青年、引智——无锡把\u201c人口+人才\u201d作为\u201c十四五\u201d续力。")

# ---- 十一、物价与货币环境 ----
heading1(doc, "十一、物价与货币环境")
para(doc, "2025年无锡**CPI与上年持平（0%）**（交通通信-1.9%、食品烟酒-0.5%、居住-0.7%；衣着+1.8%）——**物价稳、需求温和**。")
bullet(doc, "信贷扩张：存款30831.8亿（+6.6%）、贷款27251.13亿（+8.4%）、证交17.93万亿（+47.3%）——流动性充裕。")
bullet(doc, "消费：新能车+37.5%、线上零售+37.3%——\u201c新消费\u201d带动（CPI平）。")
para(doc, "货币环境判断：**宽信用、CPI持平**（近0）；\u201c资金活、物价稳\u201d，无锡靠\u201c设备更新、以旧换新、刺激新消费\u201d稳内需（2026扩内需）。")

# ---- 十二、区域一体化：无锡的位置 ----
heading1(doc, "十二、区域一体化:无锡在\u201c长三角+上海都市圈+太湖湾科创带\u201d里的位置")
para(doc, "无锡地处苏南、长三角核心、太湖之滨，是**上海大都市圈\u201c1+8\u201d成员、太湖湾科创带核心、苏锡常一体化\u201d节点**。")
bullet(doc, "**长三角一体**　无锡与苏州、常州构成\u201c苏锡常\u201d制造带，承接上海科创外溢（\u201c锡沪协同\u201d）。")
bullet(doc, "**太湖湾科创带**　太湖沿线科创走廊（集成电路、物联网）、合肥-南京协调。")
bullet(doc, "**交通枢纽**　京沪高铁/长江经济带、无锡硕放机场（旅客1096万）、地铁5线141公里。")
para(doc, "一句话：**无锡在\u201c长三角+上海都市圈+太湖湾科创带\u201d里，最核心的定位是\u201c苏南先进制造+科创中心\u201d**——区位、制造、科技是最大优势。")

# ---- 十三、未来5—10年最值得观察的五条主线 ----
heading1(doc, "十三、未来5—10年最值得观察的五条主线")
bullet(doc, "**主线一：集成电路/半导体国产化**\u201cSK海力士、华虹、长电\u201d能否顶住产能/价格，突破\u201c卡脖子\u201d。")
bullet(doc, "**主线二：AI+制造/未来产业**\u201c未来产业目标1450亿、AI+20%\u201d能否成新增长极。")
bullet(doc, "**主线三：物联网产业化（应用落地）**\u201c世界物联网\u201d赛后能否把\u201c技术\u201d做成\u201c产业\u201d。")
bullet(doc, "**主线四：民营经济与资本**\u201c民营66.7%、上市公司218\u201d能否\u201c活水\u201d支撑新质。")
bullet(doc, "**主线五：投资与内需**\u201c固投-13.3%、地产-20.1%\u201d能否在\u201c两新/设备更新\u201d下重启。")

# ---- 十四、最终结论 ----
heading1(doc, "十四、最终结论：无锡在\u201c智造强市+物联网+民营经济\u201d里的增长逻辑")
para(doc, "把2025年闭环打穿：**无锡是\u201c长三角制造强市、物联网之都\u201d**：GDP 16773.94亿/+5.1%、规上工业+5.9%、进出口+7.6%、民营占66.7%、高新技术产业产值14629.84亿/占54%。")
para(doc, "无锡不是\u201c只靠加工\u201d——它是**集成电路+物联网+装备+民营+外贸**的复合制造，靠\u201c科技+民营+外贸\u201d驱动；但投资、地产、消费内需偏弱，\u201c制造强、投资弱\u201d。")
para(doc, "一句话结论：**无锡是\u201c长三角制造强市、物联网之都\u201d；观察它先看\u201c集成电路、物联网、民营经济、外贸\u201d，再看\u201c固投、地产、消费\u201d。**它是\u201c制造优、科技强、内需弱\u201d的苏南样本。")

# ---- 附录A ----
heading1(doc, "附录A：主要资料来源与核验路径")
bullet(doc, "《2025年无锡市政府工作报告》（2025年1月11日，赵建军作，2025年目标、2024年回顾）")
bullet(doc, "《2025年无锡市国民经济和社会发展统计公报》（无锡市统计局，2026-05-09，2025年实际数据）")
bullet(doc, "《2026年无锡市政府工作报告》（无锡市发改委，2026年1月，2025年复盘+2026年目标）")
bullet(doc, "无锡市政府官网、无锡市统计局（wuxi.gov.cn/tj）")

# ---- 附录B ----
heading1(doc, "附录B：建议建立的年度跟踪仪表盘")
bullet(doc, "GDP总量/增速 vs 全年目标（可比价）。")
bullet(doc, "规上工业增加值及分行业（集成电路/物联网/装备/新能源）增速。")
bullet(doc, "高新技术产业产值/占比、民营企业产值/占比。")
bullet(doc, "固定资产投资/工业/房地产/民间投资增速。")
bullet(doc, "社会消费品零售总额、线上零售、新能源汽车。")
bullet(doc, "进出口、出口/进口、高技术产品出口、实际使用外资。")
bullet(doc, "一般公共预算收入、税收/非税、民生支出占比。")
bullet(doc, "常住人口、城镇化率、城镇新增就业。")
bullet(doc, "CPI、金融存贷款、证券交易。")
bullet(doc, "上市公司数、物联网/集成电路订单、太湖水质。")

# ---------------- 保存 ----------------
out = "/Users/x/Desktop/content-prod-lab/reports/无锡市_2025年政府工作报告_深度研究_2026-08-14.docx"
doc.save(out)
print("SAVED OK: 无锡市", out)
