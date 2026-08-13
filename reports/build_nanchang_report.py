# -*- coding: utf-8 -*-
"""Build 南昌市2025年政府工作报告 深度研究 DOCX, 参照省市系列版式。"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DARK = RGBColor(0x1F, 0x2A, 0x44)
RED = RGBColor(0xB0, 0x1E, 0x2E)
GRAY = RGBColor(0x59, 0x60, 0x69)
BLUE = RGBColor(0x1F, 0x4E, 0x79)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
HEADER_FILL = "1F2A44"
K = "微软雅黑"


def set_run(run, size=11, bold=False, color=DARK, font=K):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    run.font.name = font
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), font)
    rFonts.set(qn("w:hAnsi"), font)
    rFonts.set(qn("w:eastAsia"), font)


def para(doc, text, size=11, bold=False, color=DARK, align=None,
         space_after=6, space_before=0, font=K):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        set_run(r, size=size, bold=bold or (i % 2 == 1), color=color, font=font)
    return p


def heading1(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(16)
    pf.space_after = Pt(8)
    r = p.add_run(text)
    set_run(r, size=16, bold=True, color=RED)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "B01E2E")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def heading2(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(10)
    pf.space_after = Pt(4)
    r = p.add_run(text)
    set_run(r, size=13, bold=True, color=BLUE)
    return p


def table(doc, headers, rows, widths=None, font_size=9.5):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_run(r, size=font_size, bold=True, color=WHITE)
        tcPr = hdr[i]._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), HEADER_FILL)
        tcPr.append(shd)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(val))
            set_run(r, size=font_size, color=DARK)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    return t


def quote_box(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(8)
    pf.left_indent = Pt(12)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), "B01E2E")
    pBdr.append(left)
    pPr.append(pBdr)
    r = p.add_run(text)
    set_run(r, size=11, color=DARK, font="楷体")
    return p


def bullet(doc, text, size=10.5):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.7)
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        set_run(r, size=size, bold=(i % 2 == 1), color=DARK)
    return p


# ================================================================= 文档主体
doc = Document()
sec = doc.sections[0]
sec.page_width = Cm(21)
sec.page_height = Cm(29.7)
sec.top_margin = Cm(2.4)
sec.bottom_margin = Cm(2.2)
sec.left_margin = Cm(2.5)
sec.right_margin = Cm(2.5)

st = doc.styles["Normal"]
st.font.name = K
st.font.size = Pt(11)
st.element.rPr.rFonts.set(qn("w:eastAsia"), K)



# ---- 封面 ----
para(doc, "南昌市2025年政府工作报告", size=24, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=60, space_after=4)
para(doc, "深度研究与\u201c容易被忽视的细节\u201d分析", size=16, bold=True, color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
para(doc, "从\u201c虚拟现实、数字经济、现代制造与中部省会\u201d重新理解南昌", size=12, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)
para(doc, "\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501", color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
para(doc, "研究对象：2025年南昌市政府工作报告及同期财政、国民经济和社会发展统计资料、2026年官方复盘", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
para(doc, "标定日期：2026-08-13", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
doc.add_page_break()

# ---- 目录 ----
catalog = [
    "执行摘要",
    "一、研究口径与阅读方法",
    "二、先看南昌的特殊底盘：虚拟现实、数字经济、中部制造与省会引领",
    "三、最关键的宏观错位：GDP破8100亿、工业/数字强，但消费/外贸/地产偏弱",
    "四、容易被忽视、但信号很重的细节（15条）",
    "五、2025年GDP目标 vs 实际：对照表",
    "六、2025年增长是谁撑起来的？（结构归因）",
    "七、预算与财政的\u201c含金量\u201d",
    "八、民生底账：人口、收入与城乡",
    "九、城镇与农村：格局与均衡",
    "十、人口流入与流出",
    "十一、物价与货币环境",
    "十二、区域一体化：南昌在\u201c南昌都市圈+长江中游城市群+中部崛起\u201d里的位置",
    "十三、未来5—10年最值得观察的五条主线",
    "十四、最终结论：南昌在\u201c虚拟/数字+现代制造+省会引领\u201d里的增长逻辑",
    "附录A：主要资料来源与核验路径",
    "附录B：建议建立的年度跟踪仪表盘",
]
para(doc, "目　录", size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=10, space_after=10)
for item in catalog:
    para(doc, item, size=12, space_after=4)
doc.add_page_break()

# ---- 执行摘要 ----
heading1(doc, "执行摘要")
para(doc, "**核心判断**　如果只看新闻标题，2025年南昌最显著的是\u201cGDP破8100亿、增长5.0%左右\u201d、\u201c规模以上工业+7.0%\u201d、\u201c电子信息3000亿、汽车装备2500亿\u201d、\u201c数字经济核心产业营收2900亿\u201d。但这份研究真正值得深读的，是这座\u201c虚拟/数字经济+中部制造+省会引领\u201d的中部城市，如何在消费偏弱（社零+2.6%）、外贸依存不高（进出口1140亿）、房地产持续调整的背景下，靠\u201c高技术制造（+17%）+装备制造（+15.5%）+数字经济\u201d保住4.7%的增长（公报核算口径）。")
para(doc, "把2025年初设定的目标（GDP增长5.0%左右）、2025年《国民经济和社会发展统计公报》、2026年复盘放在一起看，南昌呈现清晰暗线：**从\u201c传统制造/房地产\u201d的旧底盘，向\u201c虚拟现实+数字经济+先进制造+省会引领\u201d升级换代**。旧引擎（房地产、部分传统产能）在调整；新引擎（电子信息/汽车装备、数字经济/VR、航空、现代服务）被要求更快补位。")
para(doc, "因此，本报告不按政府工作报告逐段复述，而采用\u201c显性表述—同期数据—制度含义—长期影响\u201d的方式，专门提取容易被忽略、但对判断南昌未来5—10年发展模式更有价值的信号。")
para(doc, "最容易记住的一句话：**南昌是\u201c中部军工与虚拟/数字经济+现代制造\u201d的重镇，靠\u201c虚拟/数字+电子信息/汽车装备+省会引领+现代服务业\u201d撑起增长。**观察南昌，与其看\u201cGDP 8100亿\u201d，不如看\u201c数字经济2900亿、电子信息3000亿、汽车装备2500亿、城镇率80%\u201d这几张名片。")
heading2(doc, "一页速览：2025年南昌经济的\u201c表与里\u201d")
table(doc,
    ["维度", "表面现象", "更深层信号"],
    [
        ["增长", "GDP 8141.69亿、+4.7%（公报）/约5.0%左右（报告）", "一产3.2%、二产41.5%、三产55.3%"],
        ["产业", "规上工业+7.0%", "高技术+17%、装备制造+15.5%、电子信息3000亿"],
        ["外贸", "进出口1140亿、+4.6%", "外贸占比较低、依赖内需"],
        ["投资", "固投目标5.5%、实际放缓；工业投资两位数增长", "民间投资比重+3.2pct、房地产调整"],
        ["消费", "社零2921.78亿、+2.6%", "首店100家、以旧换新带动360亿"],
        ["人口", "常住约670多万、城镇化率约80%（跃升）", "\u201c十四五\u201d人口年均净流入"],
        ["数字", "规上数字经济核心产业营收约2900亿", "虚拟现实创新中心、数据要素"],
    ],
    widths=[2.2, 5.6, 8.2])
para(doc, "", size=10, space_after=2)

# ---- 一、研究口径与阅读方法 ----
heading1(doc, "一、研究口径与阅读方法")
heading2(doc, "1.1 本报告的核心底稿")
bullet(doc, "**2025年《政府工作报告》**（2025年1月）——目标：GDP 5.0%左右、规上工业8.0%、固投5.5%、社零6%以上。")
bullet(doc, "**《2025年南昌市统计公报》**（市统计局2026-04）——GDP、工业、外贸、人口实数。")
bullet(doc, "**2026年南昌市政府工作报告/复盘**——对2025执行追认与房产/产业展望。")
heading2(doc, "1.2 阅读方法：显性—数据—制度—长期")
para(doc, "每章按\u201c**显性表述→同期数据→制度含义→长期影响**\u201d四层展开。")
para(doc, "**关键判别**：数据优先。例2025年GDP目标5.0%、实际5.0%左右（报告初估）/核算4.7%（公报口径）；工业目标8.0%、实际7.0%略低于目标，但高技术制造+17%。南昌\u201c工业/数字偏强、消费/外贸/地产偏弱、结构在换\u201d，需穿透GDP总量看结构。")

# ---- 二、底盘 ----
heading1(doc, "二、先看南昌的特殊底盘：虚拟现实、数字经济、中部制造与省会引领")
para(doc, "南昌的底盘，取决于它作为\u201c**虚拟现实/数字经济重镇+中部现代制造基地+江西省会+省会引领战略**\u201d的特殊定位。它是江西政治经济文化中心、\u201c豫章·洪都\u201d。")
bullet(doc, "**数字经济**：国家虚拟现实创新中心能力建设通过验收，创新能力指数列国家创新型城市第30位，数字科研百强第70位；规上数字经济核心产业营收约2900亿。")
bullet(doc, "**电子信息/汽车装备**：全产业链年营收分别有望突破3000亿/2500亿；商飞等先进制造（C909/C919/C929谱系，C909交付占全国1/3）。")
bullet(doc, "**新材料/医药健康**：与电子信息、汽车装备并列为4大千亿级产业链（新材料、医药健康另两块）。")
bullet(doc, "**省会引领**：\u201c一枢纽四中心\u201d建设，深度参与中部城市群/长江中游城市群/都市圈协同。")
para(doc, "这一底板几乎决定了南昌2025成绩单的：**只要虚拟/数字经济、电子信息/汽车装备、现代制造持续，南昌就站在\u201c中部制造+数字\u201d增长极；若传统产能/地产收缩过快，制造空间\u201c进退两难\u201d。")

# ---- 三、宏观错位 ----
heading1(doc, "三、最关键的宏观错位：GDP破8100亿、工业/数字强，但消费/外贸/地产偏弱")
para(doc, "南昌2025年最值得咀嚼的错位，是\u201c**增长靠工业+数字+省会服务，消费/外贸/地产却偏弱**\u201d。这种错位决定了对这座中部城市的观察不能只看GDP总量。")
bullet(doc, "**GDP**：8141.69亿、+4.7%（公报核算口径）；报告口径\u201c预计增长5.0%左右\u201d。一产261.68亿（+3.2%）、二产3381.13亿（+5.5%）、三产4498.88亿（+4.2%，占比55.3%）。")
bullet(doc, "**工业**：规上工业增加值+7.0%；高技术制造+17%、装备制造+15.5%；规上服务业营收+7.5%左右。")
bullet(doc, "**消费**：社零2921.78亿、+2.6%（前一年约+4%，明显回落）；首店100家、以旧换新带动360亿、\u201c乐购洪城\u201d500场。")
bullet(doc, "**外贸**：进出口1140亿、+4.6%；外贸依存度约14%，明显低于沿海外贸型城市。")
bullet(doc, "**投资/地产**：固投目标5.5%、实际整体放缓（工业投资两位数增长、民间投资比重+3.2pct）；房地产持续调整。")
para(doc, "**为什么读这条**：南昌作为\u201c中部管会+制造+数字\u201d城市，现阶段结构性矛盾是\u201c工业/数字强、消费/外贸偏弱、地产调整\u201d。经济总量稳健，但消费/外贸的拉动仍需政策托底。")

# ---- 四、15条细节 ----
heading1(doc, "四、容易被忽视、但信号很重的细节（15条）")
para(doc, "以下15条，都在2025年报告/公报里存在，但常被\u201cGDP 8100亿\u201d、\u201c工业7.0%\u201d等总量叙事掩盖。它们是判断南昌2025之后5—10年的关键小信号。")
bullet(doc, "**1. 规上工业+7.0%、高技术制造+17%**：工业8.0%目标未完全达到（7.0%），但高技术制造17%说明\u201c高端化、新兴化\u201d方向明确，只是整体盘子被传统/房地产扰动。")
bullet(doc, "**2. 装备制造+15.5%**：装备是\u201c现代制造\u201d的硬件，南昌在中部装备圈的相对优势强化。")
bullet(doc, "**3. 4大千亿级产业链成型**：电子信息（约3000亿）、汽车及装备（约2500亿）、新材料、医药健康——形成\u201c千亿级\u201d产业集群。")
bullet(doc, "**4. 商飞大飞机全谱系（C909/C919/C929）**：C909交付量占全国1/3，是南昌在\u201c国家大飞机\u201d里的独特位。")
bullet(doc, "**5. 虚拟现实/数字经济**：国家虚拟现实创新中心通过验收；规上数字经济核心产业营收约2900亿。")
bullet(doc, "**6. 数字经济\u201c数据要素\u201d示范**：\u201c数安通\u201d平台全国率先上线，入选国家数据基础设施建设典型案例；数字经济核心产业营收2900亿。")
bullet(doc, "**7. 全社会研发投入强度约2.2%**：比\u201c十四五\u201d初期明显提升，目标2026年2.28%。")
bullet(doc, "**8. 人才吸引力**：\u201c十四五\u201d累计吸引超58万名大学生与技能人才留昌，2025年新增约15万。常住人口平均年龄33岁（年轻）。")
bullet(doc, "**9. 科技创新动能**：科技型中小企业突破4000家、高新技术企业有望突破2000家、国家级专精特新\u201c小巨人\u201d40家；每万人口发明专利从12.5件提升至41.5件。")
bullet(doc, "**10. 城镇化率约80%（跃升）**：\u201c十四五\u201d期间城镇化率从约70%提升至80%左右，都市化/省会集聚明显。")
bullet(doc, "**11. 会展经济**：规模以上展会140场、带动综合消费超230亿，获评全国最具竞争力会展城市。")
bullet(doc, "**12. 文旅**：年均游客近2亿人次；\u201c文旅+体育+演艺\u201d融合，全球旅游目的地目标。")
bullet(doc, "**13. 财政质地**：地方一般公共预算收入537.77亿、+2.2%，税收占比约60.5%（可观）——财政质量好于多数同类城市。")
bullet(doc, "**14. 社会治理/安全**：生产安全事故起数、死亡人数分别下降12%、8.3%；融资平台\u201c三下降\u201d（数量/债务/成本）。")
bullet(doc, "**15. 一枢纽四中心**：作为\u201c中部中心城市\u201d，深度参与中部崛起/长江中游城市群/南昌都市圈协同。")
para(doc, "", size=10, space_after=2)

# ---- 五、对照表 ----
heading1(doc, "五、2025年GDP目标 vs 实际：对照表")
para(doc, "2025年报告中列出的主要预期目标（2025年初设定），和2025年《统计公报》实际完成情况：")
tb = [
    ["指标", "2025年目标", "2025年实际（公报/报告）", "达标判定"],
    ["GDP增速", "5.0%左右", "约5.0%左右（报告初估）/核算4.7%（公报）", "基本达标"],
    ["规上工业增加值", "8.0%左右", "+7.0%", "略低于目标"],
    ["固定资产投资", "5.5%左右", "整体放缓；工业投资两位数", "未完全达标"],
    ["社会消费品零售总额", "增长6%以上", "+2.6%（2921.78亿）", "明显回落"],
    ["进出口", "促稳提质", "1140亿、+4.6%", "平稳略增"],
    ["地方一般公共预算收入", "——", "537.77亿、+2.2%", "稳健"],
    ["城镇新增就业", "——", "新增城镇就业9万人+", "达标"],
    ["居民人均可支配收入", "与经济增长同步", "（公报未公布绝对值）", "——"],
]
table(doc, tb[0], tb[1:], widths=[3.2, 3.2, 4.4, 4.8])
para(doc, "", size=10, space_after=2)
para(doc, "**要点**：GDP基本达标、工业/固投略弱、消费回落明显，是南昌\u201c总量稳、结构待换、内需待振\u201d的直接缩影。")

# ---- 六、结构归因 ----
heading1(doc, "六、2025年增长是谁撑起来的？（结构归因）")
heading2(doc, "6.1 三次产业：二产最利、三产底盘、一产稳")
para(doc, "**二产（第二产业）是最大引擎**：二产增加值3381.13亿、+5.5%，高于GDP增速，说明\u201c制造/基建\u201d在撑着增长。三产4498.88亿、+4.2%（占比55.3%），现代服务是\u201c省会\u201d底盘但拉动温和。一产261.68亿、+3.2%稳定。")
heading2(doc, "6.2 工业：高技术驱动、装备凸显")
para(doc, "规上工业+7.0%，其中**高技术制造业+17%、装备制造业+15.5%**，电子信息/汽车装备/新材料/医药健康4大千亿级产业链。工业营收预计达8000亿。")
heading2(doc, "6.3 消费偏弱")
para(doc, "社零2921.78亿、+2.6%，远低于2024年约3%、更低于\u201c6%±\u201d目标；需求/收入成为拖累。")
heading2(doc, "6.4 外贸若干、不主导")
para(doc, "进出口1140亿、+4.6%，体量小、依存度低，更多是\u201c稳外贸\u201d，对增长贡献有限。")
heading2(doc, "6.5 投资：工业强、地产弱")
para(doc, "工业投资两位数增长、民间投资比重+3.2pct；但地产调整拖累整体固投。")
para(doc, "**一句话归因**：2025年南昌增长\u201c**工业/技术制造+省会服务+数字经济**\u201d是主引擎，\u201c消费/内需/地产\u201d偏弱——典型的\u201c生产型、投资驱动、消费待补\u201d中部省会。")

# ---- 七、财政 ----
heading1(doc, "七、预算与财政的\u201c含金量\u201d")
para(doc, "**地方一般公共预算收入537.77亿、+2.2%**；其中税收325.35亿、占比约60.5%。一般公共预算支出914.44亿。")
bullet(doc, "**收入/纯税较好**：增速+2.2%高于部分中部城市，税收占比约六成、靠产业/数字税基。")
bullet(doc, "**民生硬度**：新增就业9万+、创业担保贷款25.5亿带动3.3万人次；社保体系、育儿补贴惠及婴幼儿15万人（江西率先推生育津贴直发至个人）。")
bullet(doc, "**债务防控**：融资平台数量/债务规模/融资成本\u201c三下降\u201d，风险总体可控。")
para(doc, "**财政含义**：南昌财政\u201c质地尚可、民生偏强\u201d，但消费/收入端待提振——政策空间有限、优先保民生与债务防风险。")

# ---- 八、民生 ----
heading1(doc, "八、民生底账：人口、收入与城乡")
para(doc, "南昌\u201c十四五\u201d人口表现亮眼：**常住人口约667万+、城镇化率跃升至约80%**（2024年末667.04万、城镇化率79.9%；2025年继续小幅跃升），人口结构年轻（平均年龄33岁），累计吸引超58万名大学生/技能人才。")
bullet(doc, "**人口**：常住约670万级别，2024年末667.04万、城镇化率79.9%；2025年小幅提升至约80%。前后增加了数万级别增长（江西全省2025年减少28万，南昌逆势净流入，省会虹吸明显）。")
bullet(doc, "**就业**：2025年新增城镇就业超9万人（累计比上年+21.9%）；创业担保贷款25.5亿带动3.3万人次。")
bullet(doc, "**收入**：居民人均可支配收入\u201c与经济增长同步\u201d；\u201c十四五\u201d年均+4.7%（但2025绝对值未在公报公布）。")
bullet(doc, "**社保**：育儿补贴惠及婴幼儿15万人；生育津贴直发至个人（江西率先）；企业职工养老保险全国统筹全面落实。")
para(doc, "**民生含义**：南昌在\u201c人口/就业/城镇化\u201d上有亮点（年轻、净流入），但在\u201c收入/消费\u201d上是隐忧（增长放缓）。数据优先看民生总量。")

# ---- 九、城乡 ----
heading1(doc, "九、城镇与农村：格局与均衡")
para(doc, "**城镇化率约80%**，南昌的城镇/省会主导特征明显。\u201c十四五\u201d城镇化率从约70%提升至80%左右，都市化/集聚明显。")
bullet(doc, "**城市**：轨道1、2号线延长线、洪腾高速、西二绕城高速通车；隆兴大桥公路段开通。\u201c东进南延西拓北融中兴\u201d战略，高标准打造20个\u201c席地而坐\u201d城市客厅、32个城中村改造。")
bullet(doc, "**农村**：粮食播面512万亩、总产量43亿斤（2026目标42亿斤）；蔬菜播面67.3万亩、高标准农田5.3万亩；国家级龙头企业18家；行政村集体收入14.3亿元；2.13万名脱贫人口稳定就业。")
para(doc, "**城乡均衡**：南昌\u201c城市更新+乡村振兴\u201d并进，但总体\u201c城市强、县域追赶\u201d。")

# ---- 十、人口流 ----
heading1(doc, "十、人口流入与流出")
para(doc, "**南昌呈明确净流入**：\u201c十四五\u201d累计吸引超58万名大学生/技能人才；2025年新增约15万。全省（江西）2025年约净减28万，但南昌逆势净流入，凸显\u201c省会虹吸\u201d。")
bullet(doc, "**流入**：高校/产业/VR/航天等吸引青年；平均年龄33岁。")
bullet(doc, "**竞争**：与长沙、武汉、合肥等中部省会抢人；南昌需靠\u201c制造+数字\u201d岗位与城市品质留人。")
para(doc, "人口方向决定中长期需求与增长；南昌的\u201c年轻+省会集聚\u201d是其最硬的长逻辑之一。")

# ---- 十一、物价 ----
heading1(doc, "十一、物价与货币环境")
para(doc, "公报未单列CPI细项，但2025年整体全国物价低位。以CPI为观察、南昌呈\u201c温和\u201d特征。")
bullet(doc, "**物价**：全国CPI低位、需求弱，南昌同样面临\u201c消费偏弱、通缩压力可控\u201d。")
bullet(doc, "**货币/流动性**：民间投资比重+3.2pct、创业担保贷款25.5亿、\u201c数安通\u201d数据要素。")
para(doc, "**物价含义**：南昌消费端偏弱与全国通缩同向，对\u201c刺激消费\u201d构成内需政策空间。")

# ---- 十二、区域一体化 ----
heading1(doc, "十二、区域一体化：南昌在\u201c南昌都市圈+长江中游城市群+中部崛起\u201d里的位置")
para(doc, "南昌处于**中部崛起+长江中游城市群（南昌都市圈+武汉都市圈+长株潭都市圈）**交汇：既是江西省会，也是中部\u201c一枢纽四中心\u201d中心城市。")
bullet(doc, "**都市圈**：推动南昌都市圈与武汉、长株潭都市圈协同，深度参与中部崛起/长江中游城市群建设。")
bullet(doc, "**枢纽**：\u201c一枢纽四中心\u201d——交通枢纽+制造/科创/金融/开放中心；机场三期、昌九高铁、昌抚高铁、轨道交通三期推进。")
bullet(doc, "**开放**：开通外贸货运班列24条，南昌机场通航79城、客运吞吐量增速居中部省会机场首位、组织700+企业参加展会。")
bullet(doc, "**服务全省**：2460个事项\u201c省内通办\u201d、省内地市1万多个事项南昌可代收代办；唯一\u201c强省会\u201d带动江西。")
para(doc, "**区域含义**：南昌作为江西唯一强省会、中部交通枢纽，与武汉/长沙/合肥争夺中部资源——需靠数字/VR/制造/旅游差异化突出。")

# ---- 十三、五条主线 ----
heading1(doc, "十三、未来5—10年最值得观察的五条主线")
bullet(doc, "**主线1｜虚拟现实/数字经济**：国家虚拟现实创新中心+数据要素+数字经济核心产业营收2900亿。全球VR/数字产业地位是否稳固。")
bullet(doc, "**主线2｜电子信息与汽车装备**：电子信息3000亿/汽车装备2500亿+商飞大飞机全谱系。能否守住/突破\u201c千亿\u201d并升级为\u201c两千亿+龙头集群\u201d。")
bullet(doc, "**主线3｜省会引领与都市圈**：\u201c一枢纽四中心\u201d+都市圈协同、长江中游城市群。能否当好\u201c中部中心城市\u201d并带动江西。")
bullet(doc, "**主线4｜未来产业（低空经济/具身智能/未来显示）**：三个优先领域、未来产业三年行动。能否孵化新成长极。")
bullet(doc, "**主线5｜人口与内需**：青年净流入+城镇化80%+消费/收入修复。能否把\u201c人口集聚\u201d变成\u201c内需/消费\u201d的长期引擎。")
para(doc, "这五条，是南昌从\u201c中部制造/数字城市\u201d走向\u201c现代制造+数字经济全国的强省会\u201d的\u201c主赛道\u201d。")

# ---- 十四、结论 ----
heading1(doc, "十四、最终结论：南昌在\u201c虚拟/数字+现代制造+省会引领\u201d里的增长逻辑")
para(doc, "南昌2025年，本质上是\u201c**工业/技术制造+数字经济核心、消费/外贸/地产偏弱**\u201d的答卷：GDP8141.69亿、+4.7%（报告约5.0%左右）、规上工业+7.0%、高技术制造+17%、规模数字经济核心产业营收2900亿、第一/二/三产结构3.2:41.5:55.3、社零+2.6%、出口+4.6%。")
para(doc, "只要虚拟/数字、电子信息/汽车装备、现代制造持续，南昌就站在\u201c中部制造+数字\u201d增长极；如果传统产能/地产收缩过快，南昌需承受\u201c制造/数字强、消费/外贸/地产弱\u201d的结构挑战。")
para(doc, "最稳观察信号：**一盯数字经济/VR（引擎）、二盯电子/汽车制造（制造）、三盯商飞大飞机/现代产业（技术）、四盯消费/收入（内需）、五盯人口/都市圈（长期）。**南昌，是\u201c中部制造+虚拟/数字经济\u201d的新样本。")

# ---- 附录A ----
heading1(doc, "附录A：主要资料来源与核验路径")
bullet(doc, "南昌市2025年政府工作报告（2025年1月）——目标来源。")
bullet(doc, "《2025年南昌市国民经济和社会发展统计公报》（市统计局）——GDP、工业、外贸、人口实值。")
bullet(doc, "2026年南昌市政府工作报告（2026年1月）——2025复盘/虚拟现实/数字经济/省会引领。")
bullet(doc, "南昌海关（外贸）、市财政。")
heading2(doc, "核验说明")
para(doc, "本报告涉及数据以统计公报/官方口径为准；\u201c虚拟/数字经济/制造/都市圈\u201d等以官方口径为准。")

# ---- 附录B ----
heading1(doc, "附录B：建议建立的年度跟踪仪表盘")
para(doc, "建议把下面10个\u201c测脉搏\u201d指标做成年度跟踪表：")
dash = [
    ["#", "指标", "2025基值", "为什么重要"],
    ["1", "GDP增速", "+4.7%（核算）", "总量与方向"],
    ["2", "规模以上工业增速", "+7.0%", "制造底盘"],
    ["3", "高技术/装备制造", "+17%/+15.5%", "产业升级"],
    ["4", "数字经济核心产业营收", "约2900亿", "数字经济"],
    ["5", "进出口增速", "+4.6%（1140亿）", "外贸韧性"],
    ["6", "固定资产投资/民间", "工业两位数/比重+3.2pct", "投资结构"],
    ["7", "社零增速", "+2.6%（2921.78亿）", "内需消费"],
    ["8", "常住人口/城镇化", "约670万+/约80%", "人口与城市"],
    ["9", "地方财政收入", "+2.2%（537.77亿）", "财政质量"],
    ["10", "CPI/就业", "低位 / 新增9万+", "物价/就业"],
]
table(doc, dash[0], dash[1:], widths=[0.8, 4.6, 3.8, 4.0])
para(doc, "把这10个指标连起来看，数字经济/制造/VR（3/4）、产业升级（3）、消费（7）、人口（8），都说明南昌在真正换挡。")

# ============ 保存 ============
out = "/Users/x/Desktop/content-prod-lab/reports/南昌市_2025年政府工作报告_深度研究_2026-08-13.docx"
doc.save(out)
print("SAVED:", out)
print("PARAGRAPHS:", len(doc.paragraphs))
print("TABLES:", len(doc.tables))
