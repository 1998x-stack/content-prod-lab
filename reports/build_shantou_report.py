# -*- coding: utf-8 -*-
"""Build 汕头市2025年政府工作报告 深度研究 DOCX, 参照地级市系列版式。"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DARK = RGBColor(0x1F, 0x2A, 0x44)
RED = RGBColor(0xB0, 0x1E, 0x2E)
GRAY = RGBColor(0x59, 0x60, 0x69)
BLUE = RGBColor(0x1F, 0x4E, 0x79)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
HEADER_FILL = "1F2A44"
K = "微软雅黑"


def set_run(run, size=11, bold=False, color=DARK, font=K):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    run.font.name = font
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), font)
    rFonts.set(qn("w:hAnsi"), font)
    rFonts.set(qn("w:eastAsia"), font)


def para(doc, text, size=11, bold=False, color=DARK, align=None,
         space_after=6, space_before=0, font=K):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        set_run(r, size=size, bold=bold or (i % 2 == 1), color=color, font=font)
    return p


def heading1(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(16)
    pf.space_after = Pt(8)
    r = p.add_run(text)
    set_run(r, size=16, bold=True, color=RED)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "B01E2E")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def heading2(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(10)
    pf.space_after = Pt(4)
    r = p.add_run(text)
    set_run(r, size=13, bold=True, color=BLUE)
    return p


def table(doc, headers, rows, widths=None, font_size=9.5):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_run(r, size=font_size, bold=True, color=WHITE)
        tcPr = hdr[i]._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), HEADER_FILL)
        tcPr.append(shd)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(val))
            set_run(r, size=font_size, color=DARK)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    return t


def quote_box(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(8)
    pf.left_indent = Pt(12)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), "B01E2E")
    pBdr.append(left)
    pPr.append(pBdr)
    r = p.add_run(text)
    set_run(r, size=11, color=DARK, font="楷体")
    return p


def bullet(doc, text, size=10.5):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.7)
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        set_run(r, size=size, bold=(i % 2 == 1), color=DARK)
    return p


# ================================================================= 文档主体
doc = Document()
sec = doc.sections[0]
sec.page_width = Cm(21)
sec.page_height = Cm(29.7)
sec.top_margin = Cm(2.4)
sec.bottom_margin = Cm(2.2)
sec.left_margin = Cm(2.5)
sec.right_margin = Cm(2.5)

st = doc.styles["Normal"]
st.font.name = K
st.font.size = Pt(11)
st.element.rPr.rFonts.set(qn("w:eastAsia"), K)


# ---- 封面 ----
para(doc, "汕头市2025年政府工作报告", size=24, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=60, space_after=4)
para(doc, "深度研究与\u201c容易被忽视的细节\u201d分析", size=16, bold=True, color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
para(doc, "从\u201c粤东中心城市、潮汕侨乡、玩具之都（澄海）、纺织服装、港口\u201d重新理解汕头", size=12, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)
para(doc, "\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501", color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
para(doc, "研究对象：2025年汕头市政府工作报告及同期财政、国民经济和社会发展统计资料、2026年官方复盘", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
para(doc, "标定日期：2026-08-14", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
doc.add_page_break()

# ---- 目录 ----
catalog = [
    "执行摘要",
    "一、研究口径与阅读方法",
    "二、先看汕头的特殊底盘：粤东中心、潮汕侨乡、玩具之都、纺织服装、港口",
    "三、最关键的宏观错位：GDP 3023.83亿/0.1%，三产强但二产、固投-58.2%、外贸崩塌",
    "四、容易被忽视、但信号很重的细节（15条）",
    "五、2025年GDP目标 vs 实际：对照表",
    "六、2025年增长是谁撑起来的？（结构归因）",
    "七、预算与财政的\u201c含金量\u201d",
    "八、民生底账：人口、收入与城乡",
    "九、城镇与农村：格局与均衡",
    "十、人口流入与流出",
    "十一、物价与货币环境",
    "十二、区域一体化：汕头在\u201c粤东+粤港澳大湾区+侨经济\u201d里的位置",
    "十三、未来5—10年最值得观察的五条主线",
    "十四、最终结论：汕头在\u201c玩具+纺织+侨乡开放\u201d里的增长逻辑",
    "附录A：主要资料来源与核验路径",
    "附录B：建议建立的年度跟踪仪表盘",
]
para(doc, "目　录", size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=10, space_after=10)
for item in catalog:
    para(doc, item, size=12, space_after=4)
doc.add_page_break()

# ---- 执行摘要 ----
heading1(doc, "执行摘要")
para(doc, "**核心判断**　2025年汕头最显著的是\u201cGDP 3023.83亿元、增长0.1%（近乎停滞、远低于5%目标）\u201d、\u201c二产-7.3%、规上工业低位、三产+5.4%\u201d、\u201c固投-58.2%（历史罕见）、进出口-8.1%、财收-2.67%\u201d、\u201c但三产占57.7%、常住557.69万、玩具出口110.6亿\u201d。这说明汕头在\u201c潮汕侨乡+轻工\u201d中，**产业与需求双承压、经济接近停滞**。")
para(doc, "把2025年目标（GDP+5%/规上+5%/固投+3%/社零+4%）、2025年统计（GDP+0.1%远低于目标、规上低位、固投-58.2%、社零+0.2%、财收-2.67%）、趋势一起看，汕头是\u201c玩具+纺织+侨乡\u201d路径：**玩具（澄海）、纺织服装、轻工、港口、侨经济、跨境电商**是支柱；2025年总量3023.8亿居广东第10（粤东老牌，但被周边追赶）。")
para(doc, "最容易记住的一句话：**汕头是\u201c粤东中心城市、潮汕侨乡、中国玩具之都\u201d，以\u201c玩具出口+纺织服装+侨乡开放\u201d为底色，2025年在深度调整。**观察汕头，与其只看\u201cGDP 3023亿\u201d，不如看\u201c玩具出口110.6亿（占出口26.1%）、纺织服装（弹力面料全国第一）、港口集装箱170万标箱、三产+5.4%\u201d。")

# ---- 一、研究口径与阅读方法 ----
heading1(doc, "一、研究口径与阅读方法")
para(doc, "本报告以\u201c汕头市政府工作报告（2025年）\u201d为起点，把\u201c2025年GDP目标（5%）\u201d与\u201c官方2025年GDP（3023.83亿元/+0.1%）\u201d并置对照，并用\u201c2025年汕头市统计公报\u201d和\u201c2026年政府工作报告\u201d横向核验。")
para(doc, "口径提示：\u201cGDP、规上工业增加值增速\u201d按可比价（实际增速）；\u201c投资、消费、进出口、财政收入\u201d按现价（名义增速）。涉及\u201c常住人口\u201d用统计公报常住口径（557.69万），城镇化率用公布值（71.92%）。")
para(doc, "指标体系（与研究口径一致）：五大象限——**总量与增速（GDP）、产业动能（玩具/纺织/港口/侨商）、投资、财政、民生与人口**。")
para(doc, "特别提示（不吃老本）：汕头2025年GDP仅+0.1%、\u201c-7.3%的二产+0.1%总量\u201d、\u201c-58.2%的固投\u201d是历史级深度调整（广东经济增速垫底地市之一）；它靠**玩具、纺织、侨乡、港**立市——真正要看的是\u201c传统轻工转型升级、外向外贸企稳、侨资侨智回流\u201d。")
# ---- 二、先看汕头的特殊底盘 ----
heading1(doc, "二、先看汕头的特殊底盘：粤东中心、潮汕侨乡、玩具之都、纺织服装、港口")
para(doc, "汕头地处粤东、濒海临潮汕平原，是**粤东中心城市、潮汕侨乡、中国玩具之都（澄海）、纺织服装基地、经济特区（首批）**。2025年GDP 3023.83亿元、常住557.69万，粤东第1、广东第10。")
para(doc, "五个底盘名词，先立框架：")
bullet(doc, "**玩具之都（澄海）**　玩具品牌拥有量、专利授权量全国第一，玩具出口110.6亿（占出口26.1%）；\u201c澄海玩具创意产业带\u201d（跨境电商试点）。")
bullet(doc, "**纺织服装**　纺织服装弹力面料年织布/印染量全国第一、中国家居服采购基地；汕头国际纺织城入园超1000家。")
bullet(doc, "**潮汕侨乡/经济特区**　首批经济特区、2000万潮汕侨胞、侨资/侨智，\u201c侨经济\u201d、招商引资。")
bullet(doc, "**港口/轻工**　汕头港保税、集装箱170.21万标箱、货物吞吐量3656万吨；食品、造纸、化工等轻工。")
bullet(doc, "**三新两特一大**　新能源（海上风电）、新材料（化工）、新一代电子信息+玩具、纺织服装+大健康。")
para(doc, "这五根（玩具+纺织+侨乡+港口+三新两特）构成汕头独特底盘：**左手玩具/纺织（传统产业），右手侨乡/港口（开放禀赋）**。理解汕头，先理解\u201c潮汕商业文化+侨乡开放\u201d。")

# ---- 三、最关键的宏观错位 ----
heading1(doc, "三、最关键的宏观错位：GDP 3023.83亿/0.1%，三产强但二产、固投-58.2%、外贸崩塌")
para(doc, "2025年汕头最需要辨析的一组\u201c错位\u201d：**GDP 0.1%（远低5%目标）、二产-7.3%、固投-58.2%、进出口-8.1%、财收-2.67%，但三产+5.4%（占57.7%）、税收+3.02%**。")
para(doc, "为什么\u201c三产强、税收正\u201d，经济总量却接近停滞？三个解释：")
para(doc, "**其一，二产断崖**　第二产业-7.3%（全部工业增加值1004.28亿，规上低位）、固投-58.2%（房地产-38.8%、工业投资大降）——\u201c制造业+地产双杀\u201d拖累总量。")
para(doc, "**其二，外贸崩塌**　进出口-8.1%（出口-9.1%），对美-15.8%、对东盟-27.6%——潮汕轻工/玩具出口受贸易摩擦与外部需求压制。")
para(doc, "**其三，三产独强**　老三产+5.4%（金融/物流/旅游/商务/到港），旅游收入+23.9%、港口；但不足以托底\u201c二产崩塌\u201d。")
para(doc, "小结：汕头2025年是\u201c**三产强、二产投资外贸崩塌、深度调整**\u201d的一年：旅游、港口、侨经济有亮点，但**工业、固投、财收、外贸**全面承压。")

# ---- 四、容易被忽视、但信号很重的细节（15条） ----
heading1(doc, "四、容易被忽视、但信号很重的细节（15条）")
bullet(doc, "**1.玩具出口110.6亿（占26.1%）**　澄海玩具（专利全国第一）+跨境电商是\u201c隐藏支点\u201d。")
bullet(doc, "**2.纺织服装弹力面料/印染量全国第一**\u201c汕货\u201d纺织集群、家居服基地。\u201d")
bullet(doc, "**3.实际利用外资12.98亿/+136.8%**\u201c侨商/外资回流（招商引资+）。\u201d")
bullet(doc, "**4.旅游收入229.6亿/+23.9%、过夜游客1016万/+12.8%**\u201c潮汕文旅爆发（美食/海滩）。\u201d")
bullet(doc, "**5.港口集装箱170.21万标箱（粤东枢纽）**\u201c汕头+潮州同港、跨境物流。\u201d")
bullet(doc, "**6.海上风电（汕头风电35.84亿千瓦时/+17.9%）**\u201c三新两特一大\u201d新能源（华能/汕头）。\u201d")
bullet(doc, "**7.印制电路板+11.8%、新一代电子+4.2%**\u201c制造业（电子）亮点。\u201d")
bullet(doc, "**8.高技术制造+2.3%、先进制造占35.9%**\u201c结构转型中。\u201d")
bullet(doc, "**9.民间投资占47.4%、工业投资占39.7%**\u201c民营企业基础（潮商）。\u201d")
bullet(doc, "**10.居民收入36216元/+3.9%、城乡比1.60**\u201c农村+5.8%快于城镇+3.3%（收窄）。\u201d")
bullet(doc, "**11.常住557.69万/+0.14万、城镇化71.92%**\u201c人口稳定（侨乡/来汕人口）、出生>死亡。\u201d")
bullet(doc, "**12.CPI+0.4%**\u201c温和通胀（教育/医疗/其他用品+）。\u201d")
bullet(doc, "**13.粮食46.34万吨、水产48.63万吨**\u201c农业稳（水产/潮菜）。\u201d")
bullet(doc, "**14.市场主体超57万户（粤东西北第一）**\u201c民营经济小微活跃。\u201d")
bullet(doc, "**15.恩格尔系数41.5%**\u201c消费结构偏生存性、收入偏低。\u201d")
# ---- 五、2025年GDP目标 vs 实际：对照表 ----
heading1(doc, "五、2025年GDP目标 vs 实际：对照表")
table(doc,
    ["指标", "2025年目标", "2025年实际", "达成评价"],
    [
        ["地区生产总值(GDP)", "增长5%", "3023.83亿/0.1%", "大幅未达成"],
        ["规上工业增加值", "增长5%", "低位(高技术+2.3%)", "未达成"],
        ["固定资产投资", "增长3%", "-58.2%", "大幅未达成"],
        ["社会消费品零售总额", "增长4%左右", "1483.23亿/+0.2%", "未达成"],
        ["进出口总额", "——", "518.20亿/-8.1%", "负增长"],
        ["一般公共预算收入", "——", "170.36亿/-2.67%", "转负"],
        ["居民收入", "与经济增长同步", "36216元/+3.9%", "快于GDP"],
    ],
)
para(doc, "注：GDP、规上工业按可比价。**2025年汕头全面未达目标**：GDP仅+0.1%、固投-58.2%、进出口-8.1%、财收-2.8%——属广东增速垫底地市；居民收入（+3.9%）是唯一优于GDP的指标。")
para(doc, "拆读：**旅游（+23.9%）、外资（+136.8%）、玩具、纺织**是结构亮点；**固投（-58.2%）、二产（-7.3%）、外贸（-8.1%）**是深度拖累——\u201c深度调整\u201d是2025年核心词。")

# ---- 六、2025年增长是谁撑起来的？（结构归因） ----
heading1(doc, "六、2025年增长是谁撑起来的？（结构归因）")
para(doc, "把汕头GDP的0.1%拆开：三次产业分别增3.0%、-7.3%、5.4%（结构4.7：37.6：57.7）。**第三产业（服务业）是唯一正增长且占比最高，第二产业（工业+建筑）垮塌，第一产业（农业）稳但体量小**。")
para(doc, "2026年汕头强调\u201c5%增速、制造业当家\u201d，聚焦**三新两特一大、玩具/纺织升级、海上风电、侨招商引资、免税/开放**——核心是\u201c稳二产、强三产、抓文旅基建\u201d。")
para(doc, "**第二产业（工业）**：二产-7.3%（玩具/纺织产能下行）、高技术+2.3%、电子+11.8%（PCB），风电/新材料有布局——\u201c传统强、景气弱\u201d。")
para(doc, "**第三产业（服务业）**：+5.4%（旅游+23.9%、港口、金融、侨商），占GDP 57.7%——\u201c三产是唯一引擎\u201d。")
para(doc, "**外贸（开放）**：-8.1%（出口-9.1%），但\u201c新业态跨境电商、侨商投资\u201d+——\u201c外部对冲\u201d。")
para(doc, "一句话归因：**2025年汕头增长\u201c靠服务业（旅游/港口）+侨经济\u201d托底**，但工业、固投、外贸深度调整；\u201c三产独强、结构失衡\u201d是核心特征。")

# ---- 七、预算与财政的\u201c含金量\u201d ----
heading1(doc, "七、预算与财政的\u201c含金量\u201d")
para(doc, "2025年汕头**一般公共预算收入170.36亿元（-2.67%）**，其中税收117.43亿元（+3.02%）；一般公共预算支出385.90亿元（+0.2%）、民生支出占78.3%。")
bullet(doc, "税收结构：税收+3.02%正、财收-2.67%转负（非税/土地承压）——\u201c税收尚好、总量受土地拖累\u201d。")
bullet(doc, "民生支出：占78.3%（+1.5%），投向教育/社保/医疗/社保。")
bullet(doc, "金融支撑：存款6011.61亿（+6.4%）、贷款3314.88亿（+5.9%）——稳健、支持民营小微。")
para(doc, "**财政含金量小结**：财收-2.67%转负、税收+3.02%；\u201c税收稳、总量弱\u201d；财政对公投基建、民生、三新两特投入加大。")

# ---- 八、民生底账：人口、收入与城乡 ----
heading1(doc, "八、民生底账：人口、收入与城乡")
para(doc, "2025年汕头**居民人均可支配收入36216元（+3.9%）**，其中城镇40569元（+3.3%）、农村25285元（+5.8%），城乡比1.60（缩小）。")
para(doc, "人口画像：**常住557.69万/+0.14万、城镇化71.92%**，出生率7.89‰>死亡率7.32‰（人口正增长）；潮汕是\u201c多子多福\u201d文化区、人口年轻。")
para(doc, "民生投入：支出占78.3%保民生、医疗/教育倾斜；恩格尔系数41.5%——收入中等、生活成本高（伙食占比高）。")

# ---- 九、城镇与农村：格局与均衡 ----
heading1(doc, "九、城镇与农村：格局与均衡")
para(doc, "汕头常住城镇化率71.92%（广东地市偏中），城乡较均衡；农村收入增速（+5.8%）高于城镇（+3.3%），**城乡比缩小（1.60）**。")
para(doc, "农业底盘：**粮食46.34万吨**、园林水果38.55万吨、水产48.63万吨、蔬菜183.21万吨——\u201c潮汕菜农/水产、广东鱼米\u201d。")
para(doc, "一句话：\u201c汕头城镇化中上、农村收入快、城乡比好\u201d；但\u201c产业工业弱、增收靠外出/侨汇\u201d，\u201c乡村振兴\u201d需产业支撑。")

# ---- 十、人口流入与流出 ----
heading1(doc, "十、人口流入与流出")
para(doc, "汕头常住557.69万（+0.14万）、城镇化71.92%，是常住净增型（出生率高、来汕人口），与粤东\u201c潮汕人口外流至深圳/珠三角/海上\u201d的传统不同——\u201c06年出生潮+打工人回流/来汕\u201d在改善。")
para(doc, "结构观察：**出生率（7.89‰）高于死亡率（7.32‰）、全国少有的正自然增长**；潮商/大学生回流，\u201c侨乡\u201d人口吸引力回升。")
para(doc, "2026年目标：引才/留人（粤东中心）、潮汕青年返乡创业——汕头把\u201c人口+侨智\u201d作强市抓手。")

# ---- 十一、物价与货币环境 ----
heading1(doc, "十一、物价与货币环境")
para(doc, "2025年汕头**CPI+0.4%**（教育文化娱乐+0.7%、医疗保险+1.1%、居住-1.2%、交通通信-1.6%）——\u201c温和通胀（服务+）、地产交通-。\u201d")
bullet(doc, "信贷：存款+6.4%、贷款+5.9%，宽信用支撑民营小微/消费。")
bullet(doc, "消费：家电+292.9%、通讯器材+200.8%（以旧换新/电商）、日用品+11%。")
para(doc, "货币环境判断：**宽信用、CPI+0.4%温和**；\u201c资金稳、物价温和\u201d，汕头靠\u201c以旧换新、文旅、侨商\u201d促消费（2026目标+4%）。")

# ---- 十二、区域一体化：汕头的位置 ----
heading1(doc, "十二、区域一体化：汕头在\u201c粤东+大湾区+侨经济\u201d里的位置")
para(doc, "汕头是\u201c粤东中心城市、广东省域副中心（拟）、经济特区\u201d，地处\u201c潮汕都市圈+粤港澳大湾区辐射+海上丝绸之路\u201d节点。")
bullet(doc, "**粤东中心**　汕头+潮州+揭阳\u201c汕潮揭都市圈\u201d，汕头为龙头、带动轻工玩具。")
bullet(doc, "**大湾区联动**　高铁（汕汕/厦深）、汕头港对接大湾区产业外溢。")
bullet(doc, "**侨乡/海上丝路**　2000万潮汕侨胞、侨资/跨境电商、一带一路港口枢纽。")
para(doc, "一句话：**汕头在\u201c粤东+侨乡\u201d里，最核心的定位是\u201c粤东中心、侨乡开放桥头堡\u201d**——侨、湾区、港口是最大优势。")

# ---- 十三、未来5—10年最值得观察的五条主线 ----
heading1(doc, "十三、未来5—10年最值得观察的五条主线")
bullet(doc, "**主线一：玩具/纺织转型升级**\u201c澄海玩具出口110亿、纺织集群\u201d能否+跨境电商、+一带一路\u201d。")
bullet(doc, "**主线二：海上风电/三新两特**\u201c风电35.84亿/、新材料\u201d能否撑\u201c新增长极\u201d。")
bullet(doc, "**主线三：侨经济/回归**\u201c外资+136.8%、归侨\u201d能否\u201c侨资/侨智\u201d带动。")
bullet(doc, "**主线四：投资与基建**\u201c固投-58.2%\u201d能否靠\u201c出海基建\u201d重回。")
bullet(doc, "**主线五：外贸企稳（玩具/纺织）**\u201c出口-9.1%、对美-15.8%\u201d能否在\u201c新兴市场/东南亚\u201d企稳。")

# ---- 十四、最终结论 ----
heading1(doc, "十四、最终结论：汕头在\u201c玩具+纺织+侨乡开放\u201d里的增长逻辑")
para(doc, "把2025年闭环打穿：**汕头是\u201c粤东中心、潮汕侨乡、玩具之都\u201d**：GDP 3023.83亿/+0.1%、规上低位、固投-58.2%、进出口-8.1%、旅游+23.9%、三产占57.7%。")
para(doc, "汕头不是\u201c只有玩具\u201d——它是**玩具+纺织+侨乡+港+旅游**的复合，靠\u201c服务业+侨乡\u201d维持；但工业、投资、外贸深调，\u201c2025经济接近停滞\u201d要破局。")
para(doc, "一句话结论：**汕头是\u201c粤东中心、潮汕侨乡、玩具之都\u201d；观察它先看\u201c玩具出口、纺织、旅游、侨乡\u201d，再看\u201c二产、固投、外贸、财政\u201d。**它是2025\u201c三产稳、二产崩、待复苏\u201d的\u201c侨乡特区\u201d样本。")

# ---- 附录A ----
heading1(doc, "附录A：主要资料来源与核验路径")
bullet(doc, "《2025年汕头市政府工作报告》（2025年2月，陈涛作，2025年目标、2024年回顾）")
bullet(doc, "《2025年汕头市国民经济和社会发展统计公报》（汕头市统计局，2026-05-07，2025年实际数据）")
bullet(doc, "《2026年汕头市政府工作报告》（2026年2月，2025年复盘+2026年目标）")
bullet(doc, "汕头市政府、汕头市统计局（shantou.gov.cn/tjj）")

# ---- 附录B ----
heading1(doc, "附录B：建议建立的年度跟踪仪表盘")
bullet(doc, "GDP总量/增速 vs 全年目标（可比价）。")
bullet(doc, "规上工业增加值及分行业（玩具/纺织/电子/风电）增速。")
bullet(doc, "玩具出口额/占比、跨境电商。")
bullet(doc, "纺织服装/家居服、印染量。")
bullet(doc, "固定资产投资/工业/房地产/民间投资增速。")
bullet(doc, "旅游收入/人数、港口吞吐/集装箱。")
bullet(doc, "进出口、出口/进口、对美/东盟。")
bullet(doc, "一般公共预算收入、税收/非税、民生支出%。")
bullet(doc, "常住人口、自然增长、城镇化率。")
bullet(doc, "CPI、金融存贷款、华侨/外资利用。")

# ---------------- 保存 ----------------
out = "/Users/x/Desktop/content-prod-lab/reports/汕头市_2025年政府工作报告_深度研究_2026-08-14.docx"
doc.save(out)
print("SAVED OK: 汕头市", out)
