# -*- coding: utf-8 -*-
"""Build 兰州市2025年政府工作报告 深度研究 DOCX, 参照省会系列版式。"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DARK = RGBColor(0x1F, 0x2A, 0x44)
RED = RGBColor(0xB0, 0x1E, 0x2E)
GRAY = RGBColor(0x59, 0x60, 0x69)
BLUE = RGBColor(0x1F, 0x4E, 0x79)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
HEADER_FILL = "1F2A44"
K = "微软雅黑"


def set_run(run, size=11, bold=False, color=DARK, font=K):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    run.font.name = font
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), font)
    rFonts.set(qn("w:hAnsi"), font)
    rFonts.set(qn("w:eastAsia"), font)


def para(doc, text, size=11, bold=False, color=DARK, align=None,
         space_after=6, space_before=0, font=K):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        set_run(r, size=size, bold=bold or (i % 2 == 1), color=color, font=font)
    return p


def heading1(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(16)
    pf.space_after = Pt(8)
    r = p.add_run(text)
    set_run(r, size=16, bold=True, color=RED)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "B01E2E")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def heading2(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(10)
    pf.space_after = Pt(4)
    r = p.add_run(text)
    set_run(r, size=13, bold=True, color=BLUE)
    return p


def table(doc, headers, rows, widths=None, font_size=9.5):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_run(r, size=font_size, bold=True, color=WHITE)
        tcPr = hdr[i]._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), HEADER_FILL)
        tcPr.append(shd)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(val))
            set_run(r, size=font_size, color=DARK)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    return t


def quote_box(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(8)
    pf.left_indent = Pt(12)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), "B01E2E")
    pBdr.append(left)
    pPr.append(pBdr)
    r = p.add_run(text)
    set_run(r, size=11, color=DARK, font="楷体")
    return p


def bullet(doc, text, size=10.5):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.7)
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        set_run(r, size=size, bold=(i % 2 == 1), color=DARK)
    return p


# ================================================================= 文档主体
doc = Document()
sec = doc.sections[0]
sec.page_width = Cm(21)
sec.page_height = Cm(29.7)
sec.top_margin = Cm(2.4)
sec.bottom_margin = Cm(2.2)
sec.left_margin = Cm(2.5)
sec.right_margin = Cm(2.5)

st = doc.styles["Normal"]
st.font.name = K
st.font.size = Pt(11)
st.element.rPr.rFonts.set(qn("w:eastAsia"), K)


# ---- 封面 ----
para(doc, "兰州市2025年政府工作报告", size=24, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=60, space_after=4)
para(doc, "深度研究与\u201c容易被忽视的细节\u201d分析", size=16, bold=True, color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
para(doc, "从\u201c西北交通枢纽、有色冶金、新能源装备、石化产业与黄河治理\u201d重新理解兰州", size=12, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)
para(doc, "\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501", color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
para(doc, "研究对象：2025年兰州市政府工作报告及同期财政、国民经济和社会发展统计资料、2026年官方复盘", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
para(doc, "标定日期：2026-08-13", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
doc.add_page_break()

# ---- 目录 ----
catalog = [
    "执行摘要",
    "一、研究口径与阅读方法",
    "二、先看兰州的特殊底盘：西北枢纽、石化/有色冶金、生物医药与黄河之都",
    "三、最关键的宏观错位：GDP破3900亿达标，工业/出口爆发，但投资/消费/固投转负",
    "四、容易被忽视、但信号很重的细节（15条）",
    "五、2025年GDP目标 vs 实际：对照表",
    "六、2025年增长是谁撑起来的？（结构归因）",
    "七、预算与财政的\u201c含金量\u201d",
    "八、民生底账：人口、收入与城乡",
    "九、城镇与农村：格局与均衡",
    "十、人口流入与流出",
    "十一、物价与货币环境",
    "十二、区域一体化：兰州在\u201c强省会+兰州都市圈+面向中亚\u2019一带一路\u2019\u201d里的位置",
    "十三、未来5—10年最值得观察的五条主线",
    "十四、最终结论：兰州在\u201c石化有色+新能源装备+枢纽物流\u201d里的增长逻辑",
    "附录A：主要资料来源与核验路径",
    "附录B：建议建立的年度跟踪仪表盘",
]
para(doc, "目　录", size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=10, space_after=10)
for item in catalog:
    para(doc, item, size=12, space_after=4)
doc.add_page_break()

# ---- 执行摘要 ----
heading1(doc, "执行摘要")
para(doc, "**核心判断**　2025年兰州最显著的是\u201cGDP 3903.61亿元、增长5.5%（达5.5%目标）、三产占65.6%\u201d、\u201c规上工业+8.7%但电子工业+64.1%、有色工业+50.6%\u201d、\u201c外贸进出口143亿元/+49.1%、出口+128.1%\u201d、\u201c常住人口445.14万、自然增长率-4.04\u2030\u201d。这说明兰州在\u201c强省会行动\u201d下实现力争上游，但**投资-2.8%、消费+0.7%、地产走弱**是明显短板。")
para(doc, "把2025年目标（GDP 5.5%/规上工业+7%/固投+10%/社零+5%/财政+3%）、2025年统计、2026年前瞻一起看，兰州呈现\u201c西北工业省会\u201d的典型路径：**以石化/有色/新材料/生物医药为引擎、出口爆发、但内需偏弱**。总量3903亿居甘肃首位、对全省增长贡献率28%。")
para(doc, "最容易记住的一句话：**兰州是\u201c西北交通枢纽+石化/有色/生物医药之城\u201d，靠\u201c工业（石化/有色/新能源装备）+出口+物流枢纽+黄河文旅\u201d增长。**观察兰州，与其只看\u201cGDP 3903亿\u201d，不如看\u201c有色+50.6%、电子+64.1%、出口+128.1%、工业增加值1697亿（规上）\u201d。")# ---- 一、研究口径与阅读方法 ----
heading1(doc, "一、研究口径与阅读方法")
para(doc, "本报告以\u201c兰州市2025年政府工作报告（2025年1月，刘建勋作）\u201d为起点，把\u201c2025年GDP目标（5.5%左右）\u201d与\u201c2025年国民经济和社会发展统计公报实际值（3903.61亿元/+5.5%）\u201d并置对照，再用2026年政府工作报告的复盘作事后验证。")
para(doc, "重点阅读三类证据：**一是指标目标是否达标**；**二是结构（谁贡献增长）是否匹配目标叙事**；**三是官方复盘如何自评、承认问题**。统计口径一般公共预算收入为地方级。")
para(doc, "统一采用\u201c同比增长%\u201d（GDP为不变价、多数指标按可比价），财政与居民收入多为名义值，文中按原口径转录、不逐项换算。")

# ---- 二、特殊底盘 ----
heading1(doc, "二、先看兰州的特殊底盘：西北枢纽、石化/有色冶金、生物医药与黄河之都")
para(doc, "**区位与身份**：兰州是甘肃省会、中国西北地区重要工业基地和交通/物流枢纽（亚欧大陆桥节点、中川国际机场三期投运破1700万人次）、黄河穿城而过的\u201c黄河之都\u201d。实施\u201c强省会\u201d行动、兰州-兰州新区一体化，是兰西城市群核心。")
para(doc, "**产业底盘**：一是石化（兰州石化百万吨乙烯、产值破千亿）；二是有色冶金/新材料（产值均超500亿、原铝/电解锌）；三是新能源装备/电力（发电装机1043万千瓦/+25.7%、太阳能发电+37.5%）；四是生物医药（兰州生物制品、500亿营收）；五是数据信息/数字经济。")
para(doc, "**人口底盘**：2025年末常住人口445.14万/+0.34%（净增1.49万），城镇化率86.37%；高校（兰州大学、兰理工等）提供人才。")
para(doc, "**市场与出口**：2025年社零1474.3亿元/+0.7%；外贸进出口143亿元/+49.1%、出口+128.1%（\u201c一带一路\u201d占59.4%）。出口爆发是2025年最大亮点。")

# ---- 三、核心宏观错位 ----
heading1(doc, "三、最关键的宏观错位：GDP破3900亿达标，工业/出口爆发，但投资转负、消费弱")
para(doc, "**第一组错位**：2025年GDP目标5.5%左右，实际3903.61亿元/+5.5%**精准达标**，且总量从3662亿跃升至3900亿、\u201c十四五\u201d年均+4.3%。但**投资-2.8%（目标+10%）与消费+0.7%（目标+5%）双双低于目标**。")
para(doc, "**第二组错位**：规上工业+8.7%（目标+7%达标），其中**有色工业+50.6%、电子工业+64.1%爆发**；但建材-24.6%、纺织-3.3%、平板玻璃归零——新旧动能分化。出口+128.1%爆炸，进口-13.3%。")
para(doc, "**第三组错位**：常住人口+0.34%微增，但**自然增长率-4.04\u2030（出生4.34\u2030<死亡8.38\u2030）**严重负自然增长，人口结构性衰退明显。")
para(doc, "一句话：**兰州是\u201c工业/出口强、投资/消费弱、人口自然负增长\u201d的西北省会**——增长靠工业与出口，内需与人口是长期短板。")

# ---- 四、15条细节 ----
heading1(doc, "四、容易被忽视、但信号很重的细节（15条）")
bullet(doc, "1. **GDP精准达标**：3903.61亿/+5.5%恰为预算目标，\u201c十四五\u201d年均+4.3%。")
bullet(doc, "2. **有色工业+50.6%**：黄金产量+224.9%、原铝+2.4%，有色是2025最大动能。")
bullet(doc, "3. **电子工业+64.1%**：电子信息增速爆炸，但占比仍低。")
bullet(doc, "4. **规上工业总产值破3000亿**：石化、有色、生物医药三大集群支撑。")
bullet(doc, "5. **出口+128.1%**：出口96.5亿元、进口46.5亿元，贸易从逆转顺。")
bullet(doc, "6. **外贸+49.1%、一带一路占59.4%**：开放的“一带一路”通道。")
bullet(doc, "7. **固投-2.8%**：目标+10%未达成，工业化投资（第二产+10.4%）为正但整体转负。")
bullet(doc, "8. **汽车类零售-9.3%、石油-15.8%**：地产依附的消费链条降温。")
bullet(doc, "9. **社零+0.7%**：比目标+5%低约4个百分点，网络零售+29.3%是亮点。")
bullet(doc, "10. **黄金产量+224.9%**：兰州及周边黄金产业爆发，2026年拟继续。")
bullet(doc, "11. **常住人口445.14万/城镇化86.37%**：但自然增长率-4.04‰。")
bullet(doc, "12. **收入：城镇52660元/+4.3%、农村21319元/+6.5%**，城乡比2.47（缩小）。")
bullet(doc, "13. **CPI-0.1%**：通缩压力持续。")
bullet(doc, "14. **一般公共预算收入260.91亿/+3.57%**：达标（目标+3%），非税+10.8%。")
bullet(doc, "15. **旅游国内1.29亿人次/+17.3%、收入1000.12亿/+22.5%**：黄河文旅+夜游黄河。")# ---- 五、目标 vs 实际对照表 ----
heading1(doc, "五、2025年GDP目标 vs 实际：对照表")
table(doc,
    ["指标", "2025年目标", "2025年实际", "是否达标"],
    [
        ["地区生产总值增速", "5.5%左右", "3903.61亿元/+5.5%", "精准达标"],
        ["规上工业增加值", "+7%左右", "+8.7%", "超目标"],
        ["固定资产投资", "+10%左右", "-2.8%（转负）", "严重未达标"],
        ["社会消费品零售总额", "+5%", "1474.3亿元/+0.7%", "未达标"],
        ["一般公共预算收入", "+3%", "260.91亿元/+3.57%", "略超目标"],
        ["城镇/农村人均可支配收入", "+6%/+8%", "+4.3%/+6.5%", "均低于目标"],
        ["外贸进出口", "约150亿元", "143亿元/+49.1%", "总额略低于、增速高"],
    ],
    widths=[3.4, 2.5, 5.0, 3.1])
para(doc, "**简评**：**GDP、规上工业、财政\u201c达标\u201d**，但**投资、消费、居民收入均未达目标**；外贸虽总量略低150亿但+49.1%高增、出口+128.1%爆炸。这是\u201c工业与出口强、内需弱\u201d的典型错位年。")

# ---- 六、结构归因 ----
heading1(doc, "六、2025年增长是谁撑起来的？（结构归因）")
para(doc, "**三大产业**：一产75.29亿/+5.3%、二产1269.24亿/+5.9%、三产2559.08亿/+5.2%；三产占GDP 65.6%。增长由\u201c三产+二产\u201d双轮，二产（工业+8.7%）快于三产。")
para(doc, "**工业内部**：工业增加值934.53亿/+8.3%，规上+8.7%；**有色工业+50.6%、电子工业+64.1%、电力+17.6%、其他工业+14.4%**最亮眼；建材-24.6%、纺织-3.3%、石化-1.1%相对拖累。规上工业利润105.9亿/+59.2%。")
para(doc, "**服务业**：批发零售+13.5%、租赁商务+20.9%、信息传输软件+15.2%快；金融+0.3%、房地产+1.8%低迷。")
para(doc, "**增长归因**：兰州增长主要靠**工业（有色/电子/电力）+服务业（批发零售/信息/物流）+出口**；投资、地产、部分零售为负贡献。")

# ---- 七、财政 ----
heading1(doc, "七、预算与财政的\u201c含金量\u201d")
para(doc, "2025年一般公共预算收入260.91亿元/+3.57%（恰达标目标+3%），其中税收175.88亿/+0.39%、非税85.03亿/+10.84%；一般公共预算支出588.46亿/+7.45%，民生支出474.9亿/+6.24%（占80.7%）。")
para(doc, "**结构性**：非税高增（+10.8%）支撑收入，但增值税-2.37%、企业所得税-4.58%等主体税负回；支出高增（+7.45%）支撑基建/民生，财政\u201c靠非税+民生支出\u201d特征明显。")
para(doc, "**含金量**：收入\u201c量增质平\u201d（税弱、非税强），支出\u201c扩民生、扩基建\u201d；财政更多依托转移支付与国资，是\u201c强省会\u201d但自主财力弱\u201d。")

# ---------------- 八、民生 ----------------
heading1(doc, "八、民生底账：人口、收入与城乡")
para(doc, "常住人口445.14万/+0.34%（净增1.49万）、城镇化率86.37%；出生率4.34\u2030/死亡率8.38\u2030，**自然增长率-4.04\u2030**。收入：**城镇52660元/+4.3%、农村21319元/+6.5%**，城乡比2.47（缩小0.05）。")
para(doc, "**民生结论**：收入缓增但农村快于城镇、差距缩小；人口已是\u201c低出生+高老+机械微流入\u201d结构，老龄化（60+占23.3%）是长期压力。")

# =========== 九、城镇与农村 =============
heading1(doc, "九、城镇与农村：格局与均衡")
para(doc, "兰州城镇化率86.37%（全国前列），主城区承载石化/有色/服务/枢纽；县域（兰州新区、红古、永登等）承载有色冶金、新材料、农业。")
para(doc, "城乡收入比2.47，高于全国、逼近西部常态；农村收入增速（+6.5%）快于城镇（+4.3%）在收敛，但农村县基础仍薄弱。")

# =========== 十、人口流入流出 =============
heading1(doc, "十、人口流入与流出")
para(doc, "兰州常住人口445.14万、微增1.49万，但**自然负增长（出生4.34\u2030vs死亡8.38\u2030，-4.04\u2030）**，增长依赖机械流入（高校+产业劳动力+新市民）。")
para(doc, "**流入**：高校、石化/有色就业、新区、\u201c留兰人才\u201d；**流出**：部分低收入/年轻劳动力外迁至东部。净流入强度有限、人口总量趋缓，是兰州最需警惕的长期变量。")

# =========== 十一、物价 =============
heading1(doc, "十一、物价与货币环境")
para(doc, "2025年兰州CPI**累计下降0.1%**（通缩），食品烟酒-0.2%、交通通信-2.2%、衣着-1.2%；服务指数99.8。为低通胀/微通缩。")
para(doc, "金融：全市存款余额12054亿/+7.4%、住户存款+8.6%；贷款16551亿/+1.1%（住户贷款-0.68%）。\u201c宽货币、弱需求\u201d，资产市场（股票市值+29.2%）回暖。")

# =========== 十二、区域一体化 =============
heading1(doc, "十二、区域一体化：兰州在\u201c强省会+兰州都市圈+面向中亚\u2019一带一路\u2019\u201d里的位置")
para(doc, "兰州是甘肃省会、\u201c强省会\u201d核心，\u201c兰州-兰州新区\u201d一体化、兰西城市群、兰州都市圈；对全省经济增长贡献率28%、首位度30%。交通枢纽（兰州新区+中川国际机场三期1700万人次、兰张三四线铁路、G30等）支撑\u201c中心城市-成渝/川渝\u201d廊道。")
para(doc, "对外开放上，兰州面向中亚、欧洲，国际货运班列1118列、货值29亿美元，外贸+49.1%尤其出口+128%（对\u201c一带一路\u201d国家占59.4%）。兰州是\u201c西北物流枢纽+开放口岸\u201d型省会。")

# =========== 十三、五条主线 =============
heading1(doc, "十三、未来5—10年最值得观察的五条主线")
bullet(doc, "1. **石化/有色/新材料集群**：兰州石化百万吨乙烯、有色冶金破千亿、黄金+224.9%，工业升级主升浪。")
bullet(doc, "2. **新能源装备+电力**：发电装机1043万千瓦、太阳能+37.5%、光储/零碳经济。")
bullet(doc, "3. **出口/开放与物流枢纽**：出口+128%、跨境班列、航空/铁路枢纽，\u201c引进来\u201d+走出去。")
bullet(doc, "4. **生物医药+数据信息**：兰州生物、数据信息550亿，战略性新兴产业占比升至11.5%。")
bullet(doc, "5. **强省会+人口/消费**：内需+0.7%弱、人口自然负增长，能否靠黄河文旅+美丽乡村+金融拉动消费与人口。")

# ============ 十四、最终结论 =============
heading1(doc, "十四、最终结论：兰州在\u201c石化有色+新能源装备+枢纽物流\u201d里的增长逻辑")
para(doc, "**结论**：兰州2025年的\u201c真相\u201d是——**GDP+5.5%（达标）、工业+8.7%（有色/电子爆发）、出口+128%、外贸+49%，但投资-2.8%、消费+0.7%、人口自然-4.04\u2030**。它是\u201c工业强、出口强、但内需弱\u201d的西北省会，正在\u201c强省会+新型工业化+开放\u201d路径。")
para(doc, "**对趋势判断**：工业与出口代表兰州的\u201c硬实力\u201d，投资/消费/人口代表\u201c软约束\u201d。**有色/电子/电力/生物医药+出口**决定中期潜力，**投资盘活+消费/人口**决定长期韧性。")
para(doc, "**若只看一个指标**：看**规上工业中的有色工业增速（+50.6%）/电子（+64.1%）+出口增速（+128%）**——它们比GDP更能说明兰州\u201c工业+开放\u201d成色。兰州靠\u201c一个强省会、一条铁路、一条黄河、一堆矿\u201d，但2026冲刺五千亿要靠工业和开放。")

# ---------------- 附录A ----------------
heading1(doc, "附录A：主要资料来源与核验路径")
bullet(doc, "兰州市人民政府《兰州市政府工作报告（2025）》（2025年1月）。")
bullet(doc, "兰州市统计局《2025年兰州市国民经济和社会发展统计公报》（2026年5月）。")
bullet(doc, "《2026年兰州市政府工作报告》（2026年1月）及\u201c数读\u201d解读。")
bullet(doc, "西北各省统计公报与兰州统计年鉴交叉核验。")

# ---------------- 附录B ----------------
heading1(doc, "附录B：建议建立的年度跟踪仪表盘")
bullet(doc, "GDP总量/增速 vs 全年目标（+/-个百分点）。")
bullet(doc, "规上工业增加值及分行业（有色/电子/电力/石化/建材）增速。")
bullet(doc, "固定资产投资（总量/工业/基建/民间/房地产）增速。")
bullet(doc, "社会消费品零售总额与2026年目标（市区限额以上）。")
bullet(doc, "外贸进出口额（美元/人民币）、出口、\u201c一带一路\u201d占比。")
bullet(doc, "一般公共预算收入/税收/非税、财政支出、民生占比。")
bullet(doc, "常住人口、自然增长率、城镇化率、高校留兰。")
bullet(doc, "CPI/核心CPI、规上工业利润总额。")
bullet(doc, "旅游人数/旅游总收入（黄河文旅）、夜游黄河。")
bullet(doc, "发电装机/光伏、有色/黄金产量、兰州牛肉拉面。")

# ---------------- 保存 ----------------
out = "/Users/x/Desktop/content-prod-lab/reports/兰州市_2025年政府工作报告_深度研究_2026-08-13.docx"
doc.save(out)
print("SAVED OK 兰州市_2025年政府工作报告_深度研究_2026-08-13.docx")