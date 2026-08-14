# -*- coding: utf-8 -*-
"""Build 济宁市2025年政府工作报告 深度研究 DOCX, 参照地级市系列版式。"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DARK = RGBColor(0x1F, 0x2A, 0x44)
RED = RGBColor(0xB0, 0x1E, 0x2E)
GRAY = RGBColor(0x59, 0x60, 0x69)
BLUE = RGBColor(0x1F, 0x4E, 0x79)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
HEADER_FILL = "1F2A44"
K = "微软雅黑"


def set_run(run, size=11, bold=False, color=DARK, font=K):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    run.font.name = font
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), font)
    rFonts.set(qn("w:hAnsi"), font)
    rFonts.set(qn("w:eastAsia"), font)


def para(doc, text, size=11, bold=False, color=DARK, align=None,
         space_after=6, space_before=0, font=K):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        set_run(r, size=size, bold=bold or (i % 2 == 1), color=color, font=font)
    return p


def heading1(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(16)
    pf.space_after = Pt(8)
    r = p.add_run(text)
    set_run(r, size=16, bold=True, color=RED)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "B01E2E")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def heading2(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(10)
    pf.space_after = Pt(4)
    r = p.add_run(text)
    set_run(r, size=13, bold=True, color=BLUE)
    return p


def table(doc, headers, rows, widths=None, font_size=9.5):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_run(r, size=font_size, bold=True, color=WHITE)
        tcPr = hdr[i]._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), HEADER_FILL)
        tcPr.append(shd)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(val))
            set_run(r, size=font_size, color=DARK)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    return t


def quote_box(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(8)
    pf.left_indent = Pt(12)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), "B01E2E")
    pBdr.append(left)
    pPr.append(pBdr)
    r = p.add_run(text)
    set_run(r, size=11, color=DARK, font="楷体")
    return p


def bullet(doc, text, size=10.5):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.7)
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        set_run(r, size=size, bold=(i % 2 == 1), color=DARK)
    return p


# ================================================================= 文档主体
doc = Document()
sec = doc.sections[0]
sec.page_width = Cm(21)
sec.page_height = Cm(29.7)
sec.top_margin = Cm(2.4)
sec.bottom_margin = Cm(2.2)
sec.left_margin = Cm(2.5)
sec.right_margin = Cm(2.5)

st = doc.styles["Normal"]
st.font.name = K
st.font.size = Pt(11)
st.element.rPr.rFonts.set(qn("w:eastAsia"), K)


# ---- 封面 ----
para(doc, "菏泽市2025年政府工作报告", size=24, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=60, space_after=4)
para(doc, "深度研究与\u201c容易被忽视的细节\u201d分析", size=16, bold=True, color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
para(doc, "从\u201c突破菏泽、生物医药、化工新材料、牡丹文化、鲁苏豫皖交界\u201d重新理解菏泽", size=12, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)
para(doc, "\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501", color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
para(doc, "研究对象：2025年菏泽市政府工作报告及同期财政、国民经济和社会发展统计资料、2026年官方复盘", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
para(doc, "标定日期：2026-08-14", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
doc.add_page_break()

# ---- 目录 ----
catalog = [
    "执行摘要",
    "一、研究口径与阅读方法",
    "二、先看菏泽的特殊底盘：突破菏泽、生物医药、化工、牡丹文旅、农业大市",
    "三、最关键的宏观错位：突破菏泽、工业外贸稳，但GDP+5.5%、财收+1.3%、地产弱",
    "四、容易被忽视、但信号很重的细节（15条）",
    "五、2025年GDP目标 vs 实际：对照表",
    "六、2025年增长是谁撑起来的？（结构归因）",
    "七、预算与财政的\u201c含金量\u201d",
    "八、民生底账：人口、收入与城乡",
    "九、城镇与农村：格局与均衡",
    "十、人口流入与流出",
    "十一、物价与货币环境",
    "十二、区域一体化：菏泽在鲁苏豫皖、中原经济区、黄河流域\u201c三圈\u201d里的位置",
    "十三、未来5—10年最值得观察的五条主线",
    "十四、最终结论：菏泽在\u201c突破菏泽+生物医药+化工农业\u201d里的增长逻辑",
    "附录A：主要资料来源与核验路径",
    "附录B：建议建立的年度跟踪仪表盘",
]
para(doc, "目　录", size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=10, space_after=10)
for item in catalog:
    para(doc, item, size=12, space_after=4)
doc.add_page_break()

# ---- 执行摘要 ----
heading1(doc, "执行摘要")
para(doc, "**核心判断**　2025年菏泽最显著的是\u201cGDP 约4920-4937亿（+5.5%左右、名义+134.82亿）、规上工业+6.5%-7%、进出口646亿/+2.6%、固定资产投资增速好于全省、存贷款+11%/+9%\u201d、\u201c但一般公共预算收入+1.3%、社零+5.5%、地产偏弱、物价低位\u201d。这说明菏泽在\u201c突破菏泽+鲁西南增长极\u201d中，**工业投资稳、但财政消费地产弱**。")
para(doc, "把2025年目标（GDP+5.5%、规上+7%、固投+6%、进出口稳定）、2025年实际（GDP约4937亿/+5.5%左右、规上+6.5%、财收+1.3%、进出口+2.6%）趋势看，菏泽是\u201c突破菏泽+转型\u201d路径：**生物医药、化工新材料、机电装备、牡丹文化、农业产业**是支柱。")
para(doc, "最容易记住的一句话：**菏泽是\u201c牡丹之都、生物医药之都、鲁西南增长极\u201d，靠\u201c突破菏泽+医药化工+农业\u201d驱动。**观察菏泽，与其只看\u201cGDP 4937亿\u201d，不如看\u201c规上+6.5%、进出口646亿、生物医药、牡丹文旅、存贷款+11%/+9%\u201d。")
# ---- 一、研究口径与阅读方法 ----
heading1(doc, "一、研究口径与阅读方法")
para(doc, "本报告以\u201c菏泽市政府工作报告（2025年1月）\u201d为起点，把\u201c2025年GDP目标（5.5%）\u201d与\u201c官方2025年（约4937亿/+5.5%左右）\u201d并置对照，用\u201c2025年菏泽市统计公报\u201d和\u201c2026年计划执行与政府工作报告\u201d横向核验。")
para(doc, "口径提示：\u201cGDP、规上工业增速\u201d按可比价（实际）；\u201c投资、消费、进出口、财政\u201d按现价（名义）。\u201c常住人口\u201d用统计公报（约880万+）、城镇化率。")
para(doc, "指标体系（与研究口径一致）：**总量与增速、产业动能（医药/化工/农业）、外贸、财政、民生与人口**。")
para(doc, "特别提示（不吃老本）：菏泽2024年GDP 4802.58亿/+6.1%、2025年约4937亿/+5.5%左右（增速降档）；它不是\u201c只有牡丹\u201d——**生物医药、化工新材料、牡丹文旅、农业大市、鲁西南都市\u201d才是真正底色；突破菏泽战略。")
# ---- 二、先看菏泽的特殊底盘 ----
heading1(doc, "二、先看菏泽的特殊底盘：突破菏泽、生物医药、化工、牡丹文旅、农业大市")
para(doc, "菏泽地处山东省西南部、鲁苏豫皖四省交界、黄河南岸，是**\u201c突破菏泽\u201d战略实施地、牡丹之都（中国牡丹之都）、生物医药之都、农业大市**；以\u201c牡丹、中药、化工、农产品、鲁西南\u201d著称。2025年GDP约4937亿（山东中西部）、常住约880万（山东人口大市）、城镇化率约55%。")
para(doc, "五个底盘名词，先立框架：")
bullet(doc, "**生物医药**　菏泽生物医药产业园/头药、牡丹医药——\u201c医药之都\u201d。")
bullet(doc, "**化工新材料**　化工产业园、石油化工/煤化工、新材料——\u201c化工\u201d。")
bullet(doc, "**牡丹/文化文旅**　牡丹（国花/牡丹节）、水浒文化——\u201c牡丹之都、文旅\u201d。")
bullet(doc, "**农业大市**　粮食/牛羊/特色农业（曹县）、全国最大牡丹/蒜\u2014\u2014\u201c农业强市\u201d。")
bullet(doc, "**装备机电/鲁西南都市**　机电装备、商贸物流、鲁西南中心城市——\u201c突破菏泽\u201d。")
para(doc, "这五根（医药+化工+牡丹+农业+鲁西南都市）构成菏泽独特底盘：**左手医药化工，右手牡丹农业（鲁西南）**。理解菏泽，先理解\u201c突破菏泽、牡丹之都、农业大市\u201d。")
# ---- 三、最关键的宏观错位 ----
heading1(doc, "三、最关键的宏观错位：突破菏泽、工业外贸稳，但GDP+5.5%降档、财收+1.3%、地产弱")
para(doc, "2025年菏泽最需要辨析的一组\u201c错位\u201d：**GDP约4937亿/+5.5%左右（较2024的6.1%降档）、规上工业+6.5%、固定资产投资增速好于全省、进出口646亿/+2.6%、存贷款+11%/+9%强，但一般公共预算收入+1.3%、社零+5.5%、地产/物价偏弱**。")
para(doc, "为什么\u201c工业投资外贸\u201d稳，财政与消费却弱？三解释：")
para(doc, "**其一，突破菏泽/工业/投资稳、体量大**　规上+6.5%、固投好于全省、进出口+2.6%、存贷款高增——\u201c产业投资稳\u201d。")
para(doc, "**其二，财收、消费、地产弱**　财收+1.3%（政策性/减税）、社零+5.5%、房地产调整——\u201c财政消费弱\u201d。")
para(doc, "**其三，物价/需求温**　CPI低位、以旧换新——\u201c内需弱、量价\u201d。")
para(doc, "小结：菏泽2025年是\u201c**突破工业外贸稳、财政消费地产弱**\u201d：医药/化工/外贸存贷强，财收、地产、通胀弱。")

# ---- 四、容易被忽视、但信号很重的细节（15条） ----
heading1(doc, "四、容易被忽视、但信号很重的细节（15条）")
bullet(doc, "**1.GDP约4937亿/+5.5%左右（名义+134.82亿）**\u201c总量近5000亿、增速降档。\u201d")
bullet(doc, "**2.规模以上工业+6.5%-7%**\u201c医药/化工/制造拉动。\u201d")
bullet(doc, "**3.进出口646亿/+2.6%（外贸稳）**\u201c开放稳。\u201d")
bullet(doc, "**4.固定资产投资增速好于全省**\u201c突破菏泽/项目投资旺。\u201d")
bullet(doc, "**5.存贷款+11%/+9%（金融强）**\u201c信贷支持、宽信用。\u201d")
bullet(doc, "**6.生物医药（园区/头药）**\u201c医药之都。\u201d")
bullet(doc, "**7.牡丹/牡丹节/水浒文化**\u201c牡丹文旅。\u201d")
bullet(doc, "**8.一般公共预算收入333.4亿/+1.3%**\u201c财政稳、偏慢。\u201d")
bullet(doc, "**9.社零+5.5%、居民收入+5%**\u201c收入与增长同步。\u201d")
bullet(doc, "**10.化工新材料、装备机电**\u201c化工/机电。\u201d")
bullet(doc, "**11.农业大市（粮食/牛羊/曹县）**\u201c农业强、鲁西南粮仓。\u201d")
bullet(doc, "**12.突破菏泽战略（省重点）**\u201c政策红利/平台。\u201d")
bullet(doc, "**13.研发投入两位数增长、连续2年全省前列**\u201c创新/新质。\u201d")
bullet(doc, "**14.牡丹唯一定都/中国牡丹之都**\u201c文化IP。\u201d")
bullet(doc, "**15.常住880万+/鲁西南人口大市**\u201c人口/都市。\u201d")
# ---- 五、2025年GDP目标 vs 实际：对照表 ----
heading1(doc, "五、2025年GDP目标 vs 实际：对照表")
table(doc,
    ["指标", "2025年目标", "2025年实际", "达成评价"],
    [
        ["地区生产总值(GDP)", "增长5.5%左右", "约4937亿/+5.5%左右", "基本达成"],
        ["规模以上工业", "增长7%左右", "+6.5%左右", "略低"],
        ["固定资产投资", "增长6%左右", "增速好于全省", "好"],
        ["社会消费品零售总额", "——", "+5.5%", "稳增"],
        ["进出口总额", "稳定", "646亿/+2.6%", "稳增"],
        ["一般公共预算收入", "——", "333.4亿/+1.3%", "偏慢"],
        ["居民收入", "与增长同步", "+5%", "同步"],
    ],
)
para(doc, "注：GDP、规上工业按可比价。**GDP（+5.5%左右）基本达成、固投（好于全省）、进出口（+2.6%）稳**；**规上工业（+6.5%略低7%）、财收（+1.3%）偏慢**。")
para(doc, "拆读：**突破菏泽工业投资、外贸、生物医药、存贷款是亮色**；**财收（+1.3%）、地产、物价（低位）**是短板——\u201c产业投资稳、财政消费弱\u201d，是\u201c突破菏泽转型\u201d样本。")

# ---- 六、2025年增长是谁撑起来的？（结构归因） ----
heading1(doc, "六、2025年增长是谁撑起来的？（结构归因）")
para(doc, "菏泽以工业（医药化工）为主导、农业（第一）为底盘、加上牡丹等服务，\u201c第二+第三产业\u201d拉动。服务业增加值可比+6.5%、规上工业+6.5%。")
para(doc, "2026年菏泽强调\u201c突破菏泽、鲁西南都市、黄河流域生态\u201d，聚焦**生物医药、化工新材料、牡丹文化、装备机电、农业产业化、鲁苏豫皖枢纽**——核心是\u201c产业+突破\u201d。")
para(doc, "**第二产业（工业/制造）**：规上+6.5%、生物医药、化工、机电——\u201c医药化工强\u201d。")
para(doc, "**第三产业（服务业）**：可比+6.5%（商贸、物流、文旅/牡丹、医养）——\u201c服务+牡丹\u201d。")
para(doc, "**第一产业（农业）**：粮食/牛羊/曹县、特色农业——\u201c农业大市稳\u201d。")
para(doc, "一句话归因：**2025年菏泽增长\u201c靠工业（医药化工）+服务业（商贸/文旅）+农业\u201d**，财政消费弱；\u201c突破+产业\u201d是核心。")

# ---- 七、预算与财政的\u201c含金量\u201d ----
heading1(doc, "七、预算与财政的\u201c含金量\u201d")
para(doc, "2025年菏泽**一般公共预算收入333.4亿元（+1.3%）**；民生支出占比约78.5%；财政支出投向教育/社保/医疗。")
bullet(doc, "财收+1.3%（增速偏慢、工业减税）、民生78.5%——\u201c财政稳、含金量待升\u201d。")
bullet(doc, "金融支撑：存款+11%、贷款+9%——宽信用支持产业/农业。")
bullet(doc, "突破菏泽政策：省重点支持资金/专项债——政策红利。")
para(doc, "**财政含金量小结**：财收+1.3%（低GDP）、民生78.5%、金融活；财政对\u201c医药、牡丹、民生、突破\u201d投入加大。")

# ---- 八、民生底账：人口、收入与城乡 ----
heading1(doc, "八、民生底账：人口、收入与城乡")
para(doc, "2025年菏泽**全体居民人均可支配收入+5%**（与经济增长同步）；城镇新增就业6.5万+；收入水平全省偏低、城乡差距待缩。")
para(doc, "人口画像：**常住约880万（山东人口第2大市）、城镇化率约50%**；农业县多、人口净流出→中心城区、外出务工。")
para(doc, "民生投入：教育/社保/医保、民生支出78.5%、医疗资源——民生投入扎实。")

# ---- 九、城镇与农村：格局与均衡 ----
heading1(doc, "九、城镇与农村：格局与均衡")
para(doc, "菏泽城镇化率约50%（山东偏低）；县域经济强（牡丹区1185亿/曹县654亿/郓城604亿/单县518亿）；农村收入增速（快于城镇），**城乡差距逐步缩小**；乡村振兴、农业产业化。")
para(doc, "农业底盘：**粮食（全国产粮大市）、牛羊畜牧、牡丹、蒜/蔬菜、曹县电商（汉服/演出服）**——\u201c农业大市+电商\u201d。")
para(doc, "一句话：\u201c菏泽是农业大市+县域强、农村电商（曹县）、城乡融合推进\u201d。")

# ---- 十、人口流入与流出 ----
heading1(doc, "十、人口流入与流出")
para(doc, "菏泽常住约880万（山东人口大市）、城镇化率约50%；\u201c农业县\u201d人口外流（青岛/济南/长三角）、主城区+曹县吸引电商/劳务。")
para(doc, "结构观察：**自然增长/外出务工、回流（新型城镇）**；曹县汉服电商+返乡创业。")
para(doc, "2026年目标：城镇新增就业6.5万人、育才——菏泽靠\u201c突破产业+牡丹+电商\u201d聚人。")

# ---- 十一、物价与货币环境 ----
heading1(doc, "十一、物价与货币环境")
para(doc, "2025年菏泽**CPI低位（温和）**、居民消费价格在合理区间——\u201c低通胀、需求温\u201d。")
bullet(doc, "信贷扩张：存款+11%、贷款+9%——宽信用支持产业/农业/电商。")
bullet(doc, "消费：社零+5.5%、以旧换新——消费稳。")
para(doc, "货币环境判断：**宽信用、CPI低位**；菏泽靠\u201c产业+农业+电商\u201d稳需求（2026 CPI合理）。")

# ---- 十二、区域一体化：菏泽的位置 ----
heading1(doc, "十二、区域一体化：菏泽在鲁苏豫皖、中原经济区、黄河流域\u201c三圈\u201d里的位置")
para(doc, "菏泽是**鲁苏豫皖四省交界中心城市、中原经济区与黄河流域生态保护重要节点、山东省\u201c突破菏泽\u201d战略实施地**。")
bullet(doc, "**鲁苏豫皖交界**　四省交界、鲁西南/豫东/苏北/皖北枢纽、交通十字。")
bullet(doc, "**中原经济区/黄河流域**　黄河曹城、中原城市群、绿色低碳先行区。")
bullet(doc, "**突破菏泽**　省重点战略、产业转移/平台/政策红利。")
para(doc, "一句话：**菏泽在\u201c鲁苏豫皖+黄河+突破菏泽\u201d里，最核心是\u201c医药化工+牡丹+鲁西南枢纽\u201d**；区位、人口、政策是优势。")

# ---- 十三、未来5—10年最值得观察的五条主线 ----
heading1(doc, "十三、未来5—10年最值得观察的五条主线")
bullet(doc, "**主线一：生物医药（之都）**\u201c医药园区/头药、中药\u201d能否成大产业。")
bullet(doc, "**主线二：化工新材料/装备机电**\u201c化工升级、突破菏泽承接转移\u201d。")
bullet(doc, "**主线三：牡丹文旅/文化IP**\u201c牡丹节、水浒、黄河流域文旅\u201d。")
bullet(doc, "**主线四：农业/农村电商（曹县）**\u201c粮仓、汉服/演出电商、产业化\u201d。")
bullet(doc, "**主线五：人口/财收/枢纽**\u201c城镇化、财收提质、鲁西南大城市\u201d。")

# ---- 十四、最终结论 ----
heading1(doc, "十四、最终结论：菏泽在\u201c突破菏泽+生物医药+化工农业\u201d里的增长逻辑")
para(doc, "把2025年闭环打穿：**菏泽是\u201c牡丹之都、生物医药之都、鲁西南增长极\u201d**：GDP约4937亿/+5.5%左右、规上+6.5%、进出口646亿、生物医药/牡丹/农业大市。")
para(doc, "菏泽不是\u201c只有牡丹\u201d——它是**医药+化工+牡丹文旅+农业电商+鲁西南枢纽**的复合，靠\u201c突破+产业\u201d驱动；但财收、地产、消费弱。")
para(doc, "一句话结论：**菏泽是\u201c牡丹之都、医药之城、鲁西南增长极\u201d；观察它先看\u201c生物医药、化工、突破项目、外贸、存贷款\u201d，再看\u201c财收、地产、消费\u201d。**它是\u201c产业突破中、财政消费待强\u201d的鲁西南样本。")

# ---- 附录A ----
heading1(doc, "附录A：主要资料来源与核验路径")
bullet(doc, "《2025年菏泽市政府工作报告》（2025年1月，2025年目标、2024年回顾+6.1%）")
bullet(doc, "《菏泽市2025年国民经济和社会发展计划执行情况》（2026年1月，2025年执行情况）")
bullet(doc, "《菏泽市2026年政府工作报告》（2026年1月，复盘+2026年目标）")
bullet(doc, "菏泽市人民政府/统计局（heze.gov.cn）、菏泽日报")

# ---- 附录B ----
heading1(doc, "附录B：建议建立的年度跟踪仪表盘")
bullet(doc, "GDP总量/增速 vs 全年目标。")
bullet(doc, "规模以上工业（医药/化工）增速。")
bullet(doc, "突破菏泽/重点项目。")
bullet(doc, "固定资产/工业/房地产投资。")
bullet(doc, "进出口、存贷款。")
bullet(doc, "社零/居民收入。")
bullet(doc, "生物医药/牡丹文旅/电商。")
bullet(doc, "财收/税收/民生%。")
bullet(doc, "常住/城镇化/回流。")
bullet(doc, "CPI、农民收入/农业。")

# ---------------- 保存 ----------------
out = "/Users/x/Desktop/content-prod-lab/reports/菏泽市_2025年政府工作报告_深度研究_2026-08-14.docx"
doc.save(out)
print("SAVED OK: 菏泽市", out)
