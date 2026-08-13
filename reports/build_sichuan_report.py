# -*- coding: utf-8 -*-
"""Build 重庆市2025年政府工作报告 深度研究 DOCX, 参照北京/上海/杭州/天津系列版式。"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DARK = RGBColor(0x1F, 0x2A, 0x44)
RED = RGBColor(0xB0, 0x1E, 0x2E)
GRAY = RGBColor(0x59, 0x60, 0x69)
BLUE = RGBColor(0x1F, 0x4E, 0x79)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
HEADER_FILL = "1F2A44"
K = "微软雅黑"


def set_run(run, size=11, bold=False, color=DARK, font=K):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    run.font.name = font
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), font)
    rFonts.set(qn("w:hAnsi"), font)
    rFonts.set(qn("w:eastAsia"), font)


def para(doc, text, size=11, bold=False, color=DARK, align=None,
         space_after=6, space_before=0, font=K):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        set_run(r, size=size, bold=bold or (i % 2 == 1), color=color, font=font)
    return p


def heading1(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(16)
    pf.space_after = Pt(8)
    r = p.add_run(text)
    set_run(r, size=16, bold=True, color=RED)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "B01E2E")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def heading2(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(10)
    pf.space_after = Pt(4)
    r = p.add_run(text)
    set_run(r, size=13, bold=True, color=BLUE)
    return p


def table(doc, headers, rows, widths=None, font_size=9.5):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_run(r, size=font_size, bold=True, color=WHITE)
        tcPr = hdr[i]._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), HEADER_FILL)
        tcPr.append(shd)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(val))
            set_run(r, size=font_size, color=DARK)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    return t


def quote_box(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(8)
    pf.left_indent = Pt(12)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), "B01E2E")
    pBdr.append(left)
    pPr.append(pBdr)
    r = p.add_run(text)
    set_run(r, size=11, color=DARK, font="楷体")
    return p


def bullet(doc, text, size=10.5):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.7)
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        set_run(r, size=size, bold=(i % 2 == 1), color=DARK)
    return p


# ================================================================= 文档主体
doc = Document()
sec = doc.sections[0]
sec.page_width = Cm(21)
sec.page_height = Cm(29.7)
sec.top_margin = Cm(2.4)
sec.bottom_margin = Cm(2.2)
sec.left_margin = Cm(2.5)
sec.right_margin = Cm(2.5)

st = doc.styles["Normal"]
st.font.name = K
st.font.size = Pt(11)
st.element.rPr.rFonts.set(qn("w:eastAsia"), K)


# ---- 封面 ----
para(doc, "四川省2025年政府工作报告", size=24, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=60, space_after=4)
para(doc, "深度研究与“容易被忽视的细节”分析", size=16, bold=True, color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
para(doc, "从“成渝经济圈、电子信息、装备制造、白酒与人口大省转型”重新理解四川", size=12, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)
para(doc, "━━━━━━━━━━━━━━━━", color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
para(doc, "研究对象：2025年四川省政府工作报告及同期财政、国民经济和社会发展统计资料、2026年官方复盘", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
para(doc, "整理时间：2026年8月", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
para(doc, "分析性质：分析性研究 ｜ 非官方解读", size=11, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
doc.add_page_break()

# ---- 目录 ----
catalog = [
    "执行摘要",
    "一、研究口径与阅读方法",
    "二、先看四川的特殊底盘：西部大省、成渝双城与人口大省",
    "三、最关键的宏观错位：GDP破6.77万亿、工业稳，但外贸与物价偏弱",
    "四、容易被忽视、但信号很重的细节（15条）",
    "五、把所有细节连起来：四川正在经历的“六个换挡”",
    "六、增长暗线：电子信息、锂电与新能源、西部枢纽",
    "七、财政暗线：收入稳、税收质量尚可，民生支出大",
    "八、产业暗线：从“六大优势产业”到“新质生产力”",
    "九、区域格局：成渝双城经济圈与省内“一干多支”",
    "十、人口与城市：人口大省、外流与城镇化",
    "十一、民营经济：占近六成、市场主体活跃",
    "十二、事后验证：用2025年实际执行结果检验报告判断",
    "十三、未来5—10年最值得观察的五条主线",
    "十四、最终结论：四川在“西部大开发+成渝双城”里的增长逻辑",
    "附录A：主要资料来源与核验路径",
    "附录B：建议建立的年度跟踪仪表盘",
]
para(doc, "目　录", size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=10, space_after=10)
for item in catalog:
    para(doc, item, size=12, space_after=4)
doc.add_page_break()

# ---- 执行摘要 ----
heading1(doc, "执行摘要")
para(doc, "**核心判断**　如果只看新闻标题，2025年四川最显眼的是“GDP突破6.77万亿、增长5.5%”、“规上工业+6.5%”和“汽车产量115万辆、+29.6%”。但政府工作报告里真正值得深读的，是这一座“西部大省+成渝双城+人口大省”如何在进出口转负（-1.3%）、工业品价格通缩（PPI-2.8%）、人口减少（-46万）的背景下，靠“电子信息、锂电、装备制造、西部枢纽与新质生产力”稳住增长。")
para(doc, "把2025年初《政府工作报告》设定的目标、2025年《国民经济和社会发展统计公报》、以及2026年报告对2025年的复盘放在一起看，四川呈现清晰暗线：**从“人口+基建+传统制造”的旧依赖，向“成渝双城+新质生产力+西部开放枢纽+绿色低碳（锂电/水电）”转型**。旧引擎（一般基建、地产、低端加工）在调整；新引擎（电子信息、锂电、储能、装备、高端服务）被要求更快补位。")
para(doc, "因此，本报告不按政府工作报告逐段复述，而采用“显性表述—同期数据—制度含义—长期影响”的方式，专门提取容易被忽视、但对判断四川未来5—10年发展模式更有价值的信号。")
para(doc, "最容易记住的一句话：**四川是西部第一大经济体、成渝双城经济圈的核心，也是“水电+锂电+电子信息”的产业重镇——它既要用深耕内陆的枢纽与绿色低碳转型挑起“西部大开发”重担，也要在人口外流、外贸转负与通缩中，靠“双圈协同+新质生产力+消费升级”稳住基本盘。**观察四川，与其看“总量”，不如看“新质生产力的密度、成渝协同、锂电与半导体、西部通道与消费”这几张名片。")

heading2(doc, "一页速览：2025年四川经济的“表与里”")
table(doc,
    ["维度", "表面现象", "更深层信号"],
    [
        ["增长", "GDP 67665.3亿、+5.5%", "西部第一、总量全国第5"],
        ["产业", "规上工业+6.5%、高技术制造+12.3%", "锂电/汽车/机器人高增"],
        ["外贸", "进出口10318.1亿、-1.3%", "加工贸易+3.1%、一般贸易转弱"],
        ["投资", "固定资产投资-2.4%", "工业投资+7.3%、地产-8.5%"],
        ["财政", "一般公共预算收入+3.9%（税收+3.6%）", "收入较稳、税收尚可"],
        ["消费", "社零29135.4亿、+5.1%", "乡村+6.1%、线上+19.5%"],
        ["人口", "常住8318万、城镇化率61.38%", "人口-46万、自然增-4.22‰"],
        ["物价", "CPI-0.3%、PPI-2.8%", "通缩压力"],
    ],
    widths=[2.2, 5.2, 8.6])
para(doc, "", size=4, space_after=2)

# ---- 一、研究口径与阅读方法 ----
heading1(doc, "一、研究口径与阅读方法")
heading2(doc, "1.1 本报告的三份核心底稿")
bullet(doc, "**2025年《政府工作报告》**：2025年1月在省十四届人大三次会议上作，给出全年预期目标（GDP增长5.5%以上、城镇新增就业85万人、城镇调查失业率5.5%左右、CPI涨幅2%左右等）。固定资产投资/进出口等未设具体速率目标。")
bullet(doc, "**2025年《国民经济和社会发展统计公报》**：四川省统计局2026年3月17日发布，给出全年实际执行数与增速，是“事后验证”的锚点。")
bullet(doc, "**2026年四川省政府工作报告及官方复盘**：对2025执行结果的追认，交叉验证“目标—执行—再评价”三段式。")
heading2(doc, "1.2 阅读方法：显性—数据—制度—长期")
para(doc, "每一章按“**显性表述 → 同期数据 → 制度含义 → 长期影响**”四层展开。其中“制度含义”回答“政府为什么这么排优先序”，“长期影响”回答“这对未来5—10年意味着什么”。")
para(doc, "关键判别：**当报告使用的“增长词”和统计公报里的实际数不一致时，数据优先。**例如2025年初定“GDP增长5.5%以上”、实际+5.5%达标；固定资产投资则约-2.4%；进出口转负-1.3%（因电子外需波动）。差异反映：四川在工业/消费增长同时，外贸与价格端承压。")

# ---- 二、特殊底盘 ----
heading1(doc, "二、先看四川的特殊底盘：西部大省、成渝双城与人口大省")
para(doc, "在所有省份里，四川的“底盘”独特：**西部第一大经济体+成渝双城经济圈核心+人口近1亿（户籍约9017万）**三合一。面积、人口、资源都居西部前列，是“西部大开发”与“长江经济带”与“一带一路”交汇的重镇。")
para(doc, "这决定四川的多重身份并存：**产业重镇**（电子信息、装备、白酒、锂电）、**绿色能源**（水电、锂电、光伏）、**西部枢纽**（陆海新通道、中欧班列、航空枢纽）、**人口大省**（劳动输出与回流并存）。")
heading2(doc, "2.1 电子信息与锂电")
para(doc, "四川是全国电子信息、光伏、锂电（锂盐）重要基地，宁德时代等锂电龙头在宜宾/遂宁等地布局。2025年锂离子电池产量+45.1%、汽车产量+29.6%，说明“锂电+新能源”正在放量。")
heading2(doc, "2.2 成渝双城与西部枢纽")
para(doc, "成渝地区双城经济圈是西部首个“双城”国家战略，成都与重庆联动，共同承接投资、产业与开放。四川也是西部陆海新通道、中欧班列、连接南亚东南亚的门户。")
heading2(doc, "2.3 白酒与食品")
para(doc, "四川是“白酒第一省”（五粮液、泸州老窖等），白酒产量（商品量）125.2万千升、同比-7.0%（需求走弱）。消费与品牌是四川特色。")

# ---- 三、宏观错位 ----
heading1(doc, "三、最关键的宏观错位：GDP破6.77万亿、工业稳，但外贸与物价偏弱")
para(doc, "把2025年四川的宏观面放进一张表，会出现令人惊讶的“错位”：表观增长来自工业/消费，而外贸与物价/人口却在承压。这个错位，正是读懂四川的关键。")
heading2(doc, "3.1 “强的一面”")
table(doc,
    ["指标", "2025实绩", "信号"],
    [
        ["GDP", "67665.3亿、+5.5%", "西部第一、超5.5%目标"],
        ["规上工业", "+6.5%", "制造业强、高技术+12.3%"],
        ["汽车产量", "115.1万辆、+29.6%", "汽车/新能源放量"],
        ["工业投资", "+7.3%", "工业投资强"],
        ["社零", "+5.1%", "内需消费稳"],
        ["民营经济", "占56.4%", "民营底盘厚"],
    ],
    widths=[3.2, 5.2, 6.0])
heading2(doc, "3.2 “弱的一面”")
table(doc,
    ["指标", "2025实值", "信号"],
    [
        ["固定资产投资", "-2.4%", "投资偏弱"],
        ["进出口", "-1.3%（出口-1.5%）", "外贸转负"],
        ["房地产开发投资", "-8.5%", "地产收缩"],
        ["CPI", "-0.3%", "物价走低"],
        ["PPI", "-2.8%", "工业通缩"],
        ["常住人口", "-46万", "人口减少"],
    ],
    widths=[3.2, 5.2, 6.0])
para(doc, "**错位结论**　四川的增长“很强、但也很矛盾”。强的部分（工业/高技术/消费）与弱的部分（外贸/投资/价格/人口）并存。**真正的焦点不是“有没有增长”，而是“外部（外贸）与价格的压力”**：工业内部高增、但进出口转负、PPI通缩，说明“量强价弱+外需波动”。2026年“稳定外贸+扩大内需+新质生产力”是四川的着力点。")

# ---- 四、被忽视细节 ----
heading1(doc, "四、容易被忽视、但信号很重的细节（15条）")
details = [
    ("1", "规上工业+6.5%，其中高技术制造+12.3%", "新质生产力、高技术成为主要增量。"),
    ("2", "汽车产量115.1万辆、+29.6%", "汽车/新能源车放量。"),
    ("3", "锂离子电池产量+45.1%", "锂电/储能景气（四川是锂盐基地）。"),
    ("4", "工业机器人产量+45.9%", "智能制造、机器人新赛道。"),
    ("5", "六大优势产业+6.8%、绿色低碳优势产业+7.8%", "绿色低碳、优势产业领先。"),
    ("6", "进出口10318亿、-1.3%，加工贸易+3.1%", "外需波动、加工贸易仍稳。"),
    ("7", "白酒产量125.2万千升、-7.0%", "白酒走弱、消费品牌承压。"),
    ("8", "工业投资+7.3%、高技术制造投资+5.3%", "工业/高技投资强。"),
    ("9", "社零+5.1%、乡村+6.1%、线上+19.5%", "下沉+线上消费活跃。"),
    ("10", "民营经济占56.4%、+5.5%", "民营底盘厚、市场活跃。"),
    ("11", "一般公共预算收入+3.9%、税收+3.6%", "财政稳、税收尚可。"),
    ("12", "常住8318万、城镇化率61.38%（+1.28pct）", "人口-46万、城镇化提高。"),
    ("13", "人口自然增-4.22‰、出生率5.20‰", "自然负增、人口减少。"),
    ("14", "居民人均可支配收入36120元、+5.2%", "收入与增长同步。"),
    ("15", "CPI-0.3%、PPI-2.8%", "通缩压力、量强价弱。"),
]
for d in details:
    para(doc, "**%s**" % d[0], size=11)
    para(doc, d[1], size=10.5, color=GRAY)

# ---- 五、六个换挡 ----
heading1(doc, "五、把所有细节连起来：四川正在经历的“六个换挡”")
gear = [
    ("1．动能换挡：从“人口/基建驱动”到“新质生产力+绿色低碳”", "高技术+12.3%、锂电+45.1%、绿色低碳+7.8%。"),
    ("2．产业换挡：从“传统制造/白酒”到“电子信息+锂电+装备”", "汽车+29.6%、机器人+45.9%。"),
    ("3．投资换挡：从“基建/地产”到“制造业/高技术”", "工业投资+7.3%、高技术+5.3%、地产-8.5%。"),
    ("4．开放换挡：从“依赖电子代工”到“西部枢纽+多元化”", "加工贸易+3.1%、陆海新通道/中欧班列深化。"),
    ("5．人口换挡：从“净流出”到“成都吸附+人口存量”", "常住-46万、成都/都市圈集聚。"),
    ("6．动能换挡：从“工业量强”到“量质提升+价格企稳”", "工业高增但PPI为负，需提升价格。"),
]
for g in gear:
    para(doc, "**%s**　%s" % g, space_after=6)

# ---- 六、增长暗线 ----
heading1(doc, "六、增长暗线：电子信息、锂电与新能源、西部枢纽")
heading2(doc, "6.1 电子信息+锂电")
para(doc, "四川是电子信息（成都）、锂电（宜宾/遂宁）、光伏重镇。高技术制造+12.3%、锂电产量+45.1%，说明“新能源+高端制造”正放量，是四川最重要的新动能。")
heading2(doc, "6.2 绿色低碳+水电")
para(doc, "四川水电资源丰富，绿色低碳优势产业+7.8%。在“双碳”与锂电/储能需求下，四川具备“绿电+锂电+储能”协同。")
para(doc, "**这条暗线意味着**：四川的增长叙事正从“人口+基建”转向“新质生产力+绿色低碳+西部枢纽”。看四川，盯住“高技术制造、锂电、绿色低碳”这几组数。")

# ---- 七、财政暗线 ----
heading1(doc, "七、财政暗线：收入稳、税收尚可，民生支出大")
para(doc, "2025年四川地方一般公共预算收入5853.9亿元、+3.9%，税收3732.1亿元、+3.6%。收入增速还不错，税收尚可，主要来自制造/消费/白酒等税源。")
para(doc, "支出向民生、交通、转型倾斜。四川在“稳增长—稳民生—促转型”间平衡，税收质量相对稳健。")
para(doc, "**制度含义**　四川财政“收入稳、税收尚可”，为西部省份提供了样本。重点是保持“锂电+电子信息+新质生产力”的真实税源，摆脱对地产/一次性收入的依赖。")

# ---- 八、产业暗线 ----
heading1(doc, "八、产业暗线：从“六大优势产业”到“新质生产力”")
heading2(doc, "8.1 四川产业的“表”")
table(doc,
    ["指标", "2025增速/信号", "解读"],
    [
        ["规上工业", "+6.5%", "制造强省"],
        ["高技术制造", "+12.3%", "新质生产力"],
        ["汽车产量", "115.1万辆、+29.6%", "汽车/新能源"],
        ["锂离子电池", "+45.1%", "锂电储能"],
        ["工业机器人", "+45.9%", "智能制造"],
        ["六大优势产业", "+6.8%", "优势产业"],
        ["绿色低碳", "+7.8%", "绿色转型"],
    ],
    widths=[5.2, 4.0, 5.0])
heading2(doc, "8.2 从“白酒/低端加工”到“电子信息/锂电/高端制造”")
para(doc, "四川过去以白酒、食品、装备、电子信息为主，2025年显示“锂电+汽车+机器人+绿色低碳”正加速成为新增长。“六大优势产业”是工业底盘，而“新质生产力”是未来弹性。")

# ---- 九、区域 ----
heading1(doc, "九、区域格局：成渝双城经济圈与省内“一干多支”")
para(doc, "成渝双城经济圈是西部重要增长极，成都与重庆双核驱动，川渝共同承接先进制造、电子信息、物流与文旅。四川以成都为“主干”，带动德阳、绵阳、乐山、宜宾、泸州等“多支点”。")
para(doc, "川西北、川东北等民族/山区地区，承担生态、文旅、振兴功能。四川“一干多支、五区协同”，是国家西部大开发与区域协调的实践。")

# ---- 十、人口与城市 ----
heading1(doc, "十、人口与城市：人口大省、外流与城镇化")
heading2(doc, "10.1 人口总量")
para(doc, "2025年末四川常住8318万人、城镇化率61.38%（较上年提高1.28个百分点），但比上年减少46万人，人口自然增率-4.22‰。四川是人口大省，正经历“净流出+自然负增”。")
heading2(doc, "10.2 城市与就业")
para(doc, "成都吸附力强、是西部人口/人才集聚地；其余市州仍有人口流出。如何在“成都强+区域协同”下拉动产业与就业、留住人口，是四川长期的课题。")

# ---- 十一、民营经济 ----
heading1(doc, "十一、民营经济：占近六成、市场主体活跃")
heading2(doc, "11.1 民企地位")
para(doc, "四川民营经济增加值38172亿元、占GDP 56.4%、+5.5%；民营经营主体922.5万户（占97%）。民营是四川经济最重要的“底盘”。")
heading2(doc, "11.2 政策与服务")
para(doc, "四川持续优化营商环境、支持民间投资与专精特新。民企/个体工商户活跃，是四川未来10年高质量发展的关键。")

# ---- 十二、事后验证 ----
heading1(doc, "十二、事后验证：用2025年实际执行结果检验报告判断")
heading2(doc, "12.1 年初目标与年末执行对比")
table(doc,
    ["指标", "2025年初目标", "2025实际", "偏差"],
    [
        ["GDP增速", "5.5%以上", "+5.5%", "达标"],
        ["城镇新增就业", "85万以上", "达标", "达标"],
        ["CPI", "2%左右", "-0.3%", "明显低"],
        ["固定资产投资", "未设", "约-2.4%", "偏弱"],
        ["进出口", "未设", "-1.3%", "转负"],
        ["居民收入", "与增长同步", "+5.2%", "同步"],
    ],
    widths=[3.0, 3.4, 3.0, 4.0])
heading2(doc, "12.2 判断验证")
para(doc, "三条核心判断得到支持：**（1）工业/高技术/消费强（规上+6.5%、高技术+12.3%），验证；（2）外贸/投资/价格偏弱（进出口-1.3%、PPI-2.8%、固投-2.4%），验证；（3）民营/绿色/新质强，验证。**")
para(doc, "核心观察：四川靠“工业+新质生产力+消费+民营”守住5.5%增长，但外贸转负、人口减少、价格通缩是隐忧。2026年“稳定外贸+扩大内需+新质生产力+成渝协同”是重点。")

# ---- 十三、五条主线 ----
heading1(doc, "十三、未来5—10年最值得观察的五条主线")
lines5 = [
    ("① 电子信息/锂电/高端制造（新质生产力）", "高技术+12.3%、锂电+45.1%，能否持续放大。"),
    ("② 成渝双城经济圈协同", "双城共建能否把成都/重庆做强、产业联动。"),
    ("③ 西部枢纽与外贸", "在中欧班列/陆海新通道下，能否稳定电子外贸与多元化。"),
    ("④ 白酒/消费修复与内需", "白酒-7%、消费能否在刺激政策下修复。"),
    ("⑤ 人口/城乡/西部开发", "在人口减少下，如何用“回流+新质生产力+山区振兴”稳住。"),
]
for l in lines5:
    para(doc, "**%s**　%s" % l, space_after=6)

# ---- 十四、结论 ----
heading1(doc, "十四、最终结论：四川在“西部大开发+成渝双城”里的增长逻辑")
para(doc, "四川的2025年，本质上是“**工业+新质生产力+消费为核心，而对外贸/价格/人口存压力**”的答卷：GDP 6.77万亿、西部第一、高技术+12.3%，代价是进出口转负、PPI通缩、人口减少。")
para(doc, "只要电子信息/锂电/装备、成渝协同、新质生产力、消费能接住，四川就站在“西部大开发”头排；如果外贸、价格、人口持续偏弱，四川需承受“转型与外部”双重压力。")
para(doc, "最稳妥的观察信号：**一盯高技术/锂电/新质生产力（动能）、二盯成渝协同（区域）、三盯外贸/投资（开放）、四盯消费/白酒（内需）、五盯人口与财政（社会/财政）。**四川，是西部大开发的“火车头”。")

# ---- 附录A ----
heading1(doc, "附录A：主要资料来源与核验路径")
bullet(doc, "四川省2025年《政府工作报告》（2025年1月）——目标来源。")
bullet(doc, "《2025年四川省国民经济和社会发展统计公报》（四川省统计局，2026-03-17）——GDP、工业、外贸、人口实值。")
bullet(doc, "四川省经济和信息化厅/统计专题——电子信息、锂电、优势产业。")
bullet(doc, "2026年四川省政府工作报告——2025执行复盘。")
bullet(doc, "成都海关、商务厅、财政厅——外贸/财政实况。")
heading2(doc, "核验说明")
para(doc, "本报告所有增速、总量、占比均以统计公报/官方发布为基准。涉“电子信息产业增加值”“白酒增加值”“新能源汽车”等未单列披露项，以官方口径为准。")

# ---- 附录B ----
heading1(doc, "附录B：建议建立的年度跟踪仪表盘")
para(doc, "建议把下面10个“测脉搏”指标做成年度跟踪表：")
dash = [
    ["#", "指标", "2025基值", "为什么重要"],
    ["1", "GDP增速", "+5.5%", "总量与方向"],
    ["2", "规上工业增速", "+6.5%", "制造底盘"],
    ["3", "高技术制造增速", "+12.3%", "新质生产力"],
    ["4", "锂电产量增速", "+45.1%", "新能源/储能"],
    ["5", "进出口增速", "-1.3%", "外贸韧性"],
    ["6", "社零增速", "+5.1%", "内需消费"],
    ["7", "一般公共预算收入/税收", "+3.9%/+3.6%", "财政质量"],
    ["8", "常住人口/城镇化率", "8318万/61.46%", "人口与城镇化"],
    ["9", "民营经济占比", "56.4%", "民营活力"],
    ["10", "居民人均可支配收入增速", "+5.2%", "民生获得感"],
]
table(doc, dash[0], dash[1:], widths=[0.8, 4.5, 3.2, 5.0])
para(doc, "把这10个指标连起来看，任何一个“新引擎（2/3/4）向上、旧引擎（5）修复”，都说明四川在真正换挡；反之则是旧路径的反复。")

# ========================================= 保存
out = "/Users/x/Desktop/content-prod-lab/reports/四川省_2025年政府工作报告_深度研究_2026-08-13.docx"
doc.save(out)
print("SAVED:", out)
print("PARAGRAPHS:", len(doc.paragraphs))
print("TABLES:", len(doc.tables))
