# -*- coding: utf-8 -*-
"""Build 黑龙江省2025年政府工作报告 深度研究 DOCX, 参照省市系列版式。"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DARK = RGBColor(0x1F, 0x2A, 0x44)
RED = RGBColor(0xB0, 0x1E, 0x2E)
GRAY = RGBColor(0x59, 0x60, 0x69)
BLUE = RGBColor(0x1F, 0x4E, 0x79)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
HEADER_FILL = "1F2A44"
K = "微软雅黑"


def set_run(run, size=11, bold=False, color=DARK, font=K):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    run.font.name = font
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), font)
    rFonts.set(qn("w:hAnsi"), font)
    rFonts.set(qn("w:eastAsia"), font)


def para(doc, text, size=11, bold=False, color=DARK, align=None,
         space_after=6, space_before=0, font=K):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        set_run(r, size=size, bold=bold or (i % 2 == 1), color=color, font=font)
    return p


def heading1(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(16)
    pf.space_after = Pt(8)
    r = p.add_run(text)
    set_run(r, size=16, bold=True, color=RED)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "B01E2E")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def heading2(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(10)
    pf.space_after = Pt(4)
    r = p.add_run(text)
    set_run(r, size=13, bold=True, color=BLUE)
    return p


def table(doc, headers, rows, widths=None, font_size=9.5):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_run(r, size=font_size, bold=True, color=WHITE)
        tcPr = hdr[i]._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), HEADER_FILL)
        tcPr.append(shd)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(val))
            set_run(r, size=font_size, color=DARK)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    return t


def quote_box(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(8)
    pf.left_indent = Pt(12)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), "B01E2E")
    pBdr.append(left)
    pPr.append(pBdr)
    r = p.add_run(text)
    set_run(r, size=11, color=DARK, font="楷体")
    return p


def bullet(doc, text, size=10.5):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.7)
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        set_run(r, size=size, bold=(i % 2 == 1), color=DARK)
    return p


# ================================================================= 文档主体
doc = Document()
sec = doc.sections[0]
sec.page_width = Cm(21)
sec.page_height = Cm(29.7)
sec.top_margin = Cm(2.4)
sec.bottom_margin = Cm(2.2)
sec.left_margin = Cm(2.5)
sec.right_margin = Cm(2.5)

st = doc.styles["Normal"]
st.font.name = K
st.font.size = Pt(11)
st.element.rPr.rFonts.set(qn("w:eastAsia"), K)



# ---- 封面 ----
para(doc, "黑龙江省2025年政府工作报告", size=24, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=60, space_after=4)
para(doc, "深度研究与“容易被忽视的细节”分析", size=16, bold=True, color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
para(doc, "从“冰雪经济、装备制造、对俄贸易、粮食基地与人口外流”重新理解黑龙江", size=12, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)
para(doc, "━━━━━━━━━━━━━━━━", color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
para(doc, "研究对象：2025年黑龙江省政府工作报告及同期财政、国民经济和社会发展统计资料、2026年官方复盘", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
para(doc, "标定日期：2026-08-13", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
doc.add_page_break()

# ---- 目录 ----
catalog = [
    "执行摘要",
    "一、研究口径与阅读方法",
    "二、先看黑龙江的特殊底盘：冰雪经济、装备制造、对俄贸易、粮食基地与人口外流",
    "三、最关键的宏观错位：GDP近1.69万亿、装备/冰雪/财政强，但工业/投资/人口弱",
    "四、容易被忽视、但信号很重的细节（15条）",
    "五、2025年GDP目标 vs 实际：对照表",
    "六、2025年增长是谁撑起来的？（结构归因）",
    "七、预算与财政的“含金量”",
    "八、民生底账：人口、收入与城乡",
    "九、城镇与农村：格局与均衡",
    "十、人口流入与流出",
    "十一、物价与货币环境",
    "十二、区域一体化：黑龙江在“对俄开放+中俄合作+东北振兴”里的位置",
    "十三、未来5—10年最值得观察的五条主线",
    "十四、最终结论：黑龙江在“粮食+冰雪+装备+对俄”里的增长逻辑",
    "附录A：主要资料来源与核验路径",
    "附录B：建议建立的年度跟踪仪表盘",
]
para(doc, "目　录", size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=10, space_after=10)
for item in catalog:
    para(doc, item, size=12, space_after=4)
doc.add_page_break()

# ---- 执行摘要 ----
heading1(doc, "执行摘要")
para(doc, "**核心判断**　如果只看新闻标题，2025年黑龙江最显眼的是“GDP近1.69万亿、增长4.2%”、“粮食产量8200.3万吨（连续16年全国第1）”、“装备工业+8.4%、电气机械+23.7%”和“游客+10.3%/入境+47.8%”。但这份研究真正值得深读的，是这座“粮食基地+冰雪经济+装备制造+对俄窗口”的东北大省，如何在工业偏弱（规上工业+2.8%）、固定资产投资（-10.3%）与人口持续外流的背景下，靠“装备制造+冰雪文旅+对俄贸易+粮食”实现4.2%的增长。")
para(doc, "把2025年初设定的目标（GDP增长5%左右）、2025年《国民经济和社会发展统计公报》、以及2026年报告对2025年的复盘放在一起看，黑龙江呈现清晰暗线：**从“粮食/能源/传统重工业”的旧底盘，向“装备制造+冰雪/旅游+对俄开放+生物经济/数字经济”转型**。旧引擎（能源/石化/传统重工）在调整；新引擎（装备/电气、冰雪旅游、对俄、粮食/食材加工）被要求更快补位。")
para(doc, "因此，本报告不按政府工作报告逐段复述，而采用“显性表述—同期数据—制度含义—长期影响”的方式，专门提取容易被忽视、但对判断黑龙江未来5—10年发展模式更有价值的信号。")
para(doc, "最容易记住的一句话：**黑龙江是“粮食安全（连续16年全国第1）+冰雪经济+装备制造+对俄开放”的边疆样本，靠“粮食+装备+冰雪+对俄”撑起增长。**观察黑龙江，与其看“GDP 1.69万亿”，不如看“粮食8200万吨、冰雪旅游、装备电气、对俄贸易、亚冬会”这几张名片。")
heading2(doc, "一页速览：2025年黑龙江经济的“表与里”")
table(doc,
    ["维度", "表面现象", "更深层信号"],
    [
        ["增长", "GDP 16878.0亿、+4.2%", "一产19.9%、二产23.8%、三产56.4%"],
        ["产业", "规上工业+2.8%", "装备+8.4%、电气机械+23.7%"],
        ["外贸", "进出口3125.9亿", "出口+5.0%、民企+6.1%"],
        ["投资", "固定资产投资-10.3%", "铁路+49.9%、技改+11.3%"],
        ["财政", "一般公共预算收入1535.0亿、+5.7%", "支出+9.5%"],
        ["消费", "社零5622.1亿、+3.2%", "文化办公+1.5倍、穿戴+1.9倍"],
        ["人口", "常住3001万、城镇化68.98%", "人口外流压力"],
        ["开放", "对俄贸易/中俄合作/东北亚", "粮食/冰雪/装备"],
    ],
    widths=[2.2, 5.4, 8.4])
para(doc, "", size=10, space_after=2)

# ---- 一、研究口径与阅读方法 ----
heading1(doc, "一、研究口径与阅读方法")
heading2(doc, "1.1 本报告的三份核心底稿")
bullet(doc, "**2025年《政府工作报告》**：2025年1月在省十四届人大三次会议上作，给出全年预期目标（GDP增长5%左右、规上工业5%左右、固投6%左右、社零5%左右、粮食1600亿斤以上）。")
bullet(doc, "**2025年《国民经济和社会发展统计公报》**：省统计局2026年3月31日发布，给出全年实际执行数与增速，是“事后验证”的锚点。")
bullet(doc, "**2026年黑龙江省政府工作报告及官方复盘**：对2025执行结果的追认，交叉验证“目标—执行—再评价”三段式。")
heading2(doc, "1.2 阅读方法：显性—数据—制度—长期")
para(doc, "每一章按“**显性表述 → 同期数据 → 制度含义 → 长期影响**”四层展开。其中“制度含义”回答“政府为什么这么排优先序”，“长期影响”回答“这对未来5—10年意味着什么”。")
para(doc, "关键判别：**当报告使用的“增长词”和统计公报里的实际数不一致时，数据优先。**例如2025年GDP目标5%左右、实际4.2%略低于目标；规上工业目标5%左右、实际2.8%；社零目标5%左右、实际3.2%。差异反映：黑龙江“粮食/冰雪/装备/财政强，工业/投资/人口弱”。")
heading2(doc, "1.3 指标取舍与单位说明")
para(doc, "本报告统一以“2025年统计公报+2026年报告复盘”为口径，增速均为可比价，绝对数为人民币，粮食以亿斤/万吨计。GDP以初步核算为准。")

# ---- 二、底盘 ----
heading1(doc, "二、先看黑龙江的特殊底盘：冰雪经济、装备制造、对俄贸易、粮食基地与人口外流")
para(doc, "黑龙江的地盘，取决于它作为“**国家粮食安全基地+冰雪经济强省+装备制造+对俄开放前沿**”的特殊定位。它是中国粮食安全的“压舱石”。")
bullet(doc, "**粮食基地**：粮食产量8200.3万吨、连续16年全国第1（水稻2514.2、玉米4735.7、大豆903.9万吨）；高标准农田近1.2亿亩；绿色有机食品基地超1亿亩；玉米深加工全国第1。")
bullet(doc, "**冰雪经济**：游客+10.3%、旅游花费+16.5%、入境+47.8%；哈尔滨冰雪大世界（全球最大主题乐园）、亚冬会（第九届）、冬运会十四连冠；冰雪旅游“顶流”。")
bullet(doc, "**装备制造**：装备工业+8.4%、电气机械+23.7%、通用设备+11.3%；发电机组3796万千瓦（+27.5%）、汽车发动机3218.7万千瓦（+34.0%）、汽轮机+26.1%、集成电路+11.1%。")
bullet(doc, "**对俄开放**：进出口3125.9亿、出口+5.0%、民企+6.1%；2024年对俄进出口增速高于全国8.3个百分点；边境贸易/中俄合作。")
para(doc, "**制度含义**：黑龙江把“粮食安全、冰雪经济、装备制造、对俄开放”当核心资产，深入推进“东北振兴/产业转型”，生物经济、数字经济、旅游是新增量。")

# ---- 三、宏观错位 ----
heading1(doc, "三、最关键的宏观错位：GDP近1.69万亿、装备/冰雪/财政强，但工业/投资/人口弱")
para(doc, "2025年黑龙江GDP 16878.0亿元、+4.2%（一产+3.5%、二产+1.5%、三产+5.5%）。表面看“稳中有进”，但拆开看是“**装备/冰雪/财政强、工业/投资/人口弱**”的错位：")
para(doc, "**强的部分**：规上工业+2.8%（装备+8.4%、电气机械+23.7%、原煤+16.6%）；财政收入+5.7%；旅游花费+16.5%、入境+47.8%；粮食8200.3万吨（连续16年全国第1）。")
para(doc, "**弱的部分**：固定资产投资-10.3%（基建-11.4%）；房地产开发-7.3%；三产/传统工业偏弱（规上工业+2.8%）；人口持续外流（常住3001万）。")
para(doc, "**核心错位一句话**：黑龙江“粮食/冰雪/装备/财政强（装备+8.4%、入境+47.8%），但工业/投资强具空间、人口外流”。2026年“稳工业/投资、强冰雪/对俄、保粮食安全”是重点。")
table(doc,
    ["强引擎", "数据", "弱项", "数据"],
    [
        ["装备工业", "+8.4%", "固定资产投资", "-10.3%"],
        ["电气机械制造", "+23.7%", "房地产开发投资", "-7.3%"],
        ["旅游花费/入境", "+16.5%/+47.8%", "规上工业增速", "+2.8%"],
        ["财政收入", "+5.7%", "社会消费品零售", "+3.2%"],
        ["粮食产量", "8200.3万吨(全国第1)", "人口外流", "常住3001万"],
    ],
    widths=[3.8, 3.4, 3.4, 3.0])
para(doc, "**错位结论**　黑龙江的增长“强粮食/冰雪/装备/财政、弱工业/投资”。2026年“稳工业/投资、强冰雪/对俄、补人口”是重点。")

# ---- 四、被忽视细节 ----
heading1(doc, "四、容易被忽视、但信号很重的细节（15条）")
details = [
    ("1", "GDP 16878.0亿、+4.2%", "总量/增长回升。"),
    ("2", "粮食8200.3万吨、连续16年全国第1", "粮食安全。"),
    ("3", "装备工业+8.4%（电气+23.7%）", "装备/电气领跑。"),
    ("4", "发电机组+27.5%、汽车发动机+34.0%", "装备产品。"),
    ("5", "冰雪旅游/亚冬会/冰雪大世界", "冰雪经济强。"),
    ("6", "游客+10.3%、旅游花费+10.5%", "旅游/冰雪。"),
    ("7", "入境游客+47.8%、花费+42.8%", "入境回暖。"),
    ("8", "进出口3125.9亿、对俄、民企+6.1%", "对俄/开放。"),
    ("9", "规上工业+2.8%、原煤+16.6%", "工业/能源。"),
    ("10", "财政收入+5.7%、支出+9.5%", "财政稳增。"),
    ("11", "文化办公+1.5倍、可穿戴+1.9倍", "消费升级/新品。"),
    ("12", "常住3001万、城镇化率68.98%+0.93pct", "人口/城镇化。"),
    ("13", "居民收入32851元、+5.1%、农村+5.8%", "收入稳、农村快。"),
    ("14", "城镇新增就业35.8万、超额完成", "就业稳。"),
    ("15", "CPI+0.2%、PPI-5.3%", "通缩/工业品弱。"),
]
for d in details:
    para(doc, "**%s**" % d[0], size=11)
    para(doc, d[1], size=10.5, color=GRAY)

    para(doc, "**备注**　这条“错位”在黑龙江尤其鲜明：增长靠粮食/冰雪/装备/财政，但工业/投资/人口弱。2026年若工业/投资修复、对俄深化，增长可能从“装备/冰雪单极”走向“装备/冰雪+产业/投资”多极。这条细节，正是黑龙江2026年最难也最值得盯的“变量”。", size=10.5)

# ---- 五、2025年目标 vs 实际 对照 ----
heading1(doc, "五、2025年GDP目标 vs 实际：对照表")
t5 = [
    ["指标", "2025目标", "2025实际", "达标判定"],
    ["GDP增速", "5%左右", "+4.2%", "略低于目标"],
    ["规上工业增加值增速", "5%左右", "+2.8%", "未达标"],
    ["固定资产投资增速", "6%左右", "-10.3%", "大幅未达标"],
    ["社会消费品零售增速", "5%左右", "+3.2%", "未达标"],
    ["粮食产量", "1600亿斤以上", "8200.3万吨(约1640亿斤)", "超额"],
    ["居民消费价格(CPI)", "涨幅2%左右", "+0.2%", "远低于"],
    ["城镇调查失业率", "5.8%左右", "约5.3%", "达标"],
]
table(doc, t5[0], t5[1:], widths=[3.4, 3.0, 3.0, 3.6])
para(doc, "**对照结论**　目标“稳中求进”，实际“有缺口”：GDP 4.2%、规上工业2.8%、社零3.2%、固投-10.3%（均未达目标）；但粮食超额（1600亿+）、CPI 0.2%、就业达标。粮食/冰雪/装备/财政接住，工业/投资是短板。")

# ========= 六、增速分项支撑 =========
heading1(doc, "六、2025年增长是谁撑起来的？（结构归因）")
para(doc, "GDP+4.2%背后，是“**装备/冰雪/粮食/财政强、工业/投资弱**”的结构。把增长贡献拆开看：")
g6 = [
    ["引擎", "贡献/状态", "说明"],
    ["装备/装备工业", "+8.4%", "电气机械+23.7%、发电机组+27.5%"],
    ["冰雪/旅游", "花费+16.5%", "游客+10.3%、入境+47.8%"],
    ["粮食/农业", "8200.3万吨", "连续16年全国第1"],
    ["财政收入", "+5.7%", "财政稳增"],
    ["对俄/出口", "出口+5.0%", "民企+6.1%"],
    ["规上工业", "+2.8%", "工业偏弱"],
    ["固定资产投资", "-10.3%", "投资拖累"],
    ["房地产开发", "-7.3%", "地产调整"],
    ["人口外流", "常住3001万", "人口约束"],
    ["CPI/PPI", "0.2%/-5.3%", "物价通缩"],
]
table(doc, g6[0], g6[1:], widths=[3.6, 3.6, 6.0])
para(doc, "**一句话**　增长靠“装备+冰雪+粮食+财政”，工业/投资/人口是短板。2026年考验黑龙江“装备/冰雪能否带动工业投资修复”。")

# ========== 七、预算与财政 ==========
heading1(doc, "七、预算与财政的“含金量”")
para(doc, "2025年黑龙江一般公共预算收入**1535.0亿元、+5.7%**（支出+9.5%）。财政收入稳步增长。")
bullet(doc, "财政收入+5.7%、支出（可比）+9.5%，稳增。")
bullet(doc, "财政“稳收+民生（社保+9.3%、教育等）+粮食/冰雪/对俄支持”。")
bullet(doc, "支持东北振兴、农业科技、装备制造。")

# ---- 八、民生底账 ----
heading1(doc, "八、民生底账：人口、收入与城乡")
para(doc, "2025年黑龙江常住人口**3001万人**，城镇化率**68.98%、+0.93pct**。人口持续外流。")
para(doc, "居民人均可支配收入**32851元、+5.1%**，农村（+5.8%）快于城镇（+4.4%）。")
g8 = [
    ["指标", "数值", "信号"],
    ["常住人口", "3001万", "人口外流压力"],
    ["城镇化率", "68.98%/+0.93pct", "较高城镇化"],
    ["居民人均可支配收入", "32851元/+5.1%", "收入稳、农村快"],
    ["城镇新增就业", "35.8万人", "就业稳"],
]
table(doc, g8[0], g8[1:], widths=[4.2, 4.0, 4.8])
para(doc, "**民生观察**：收入增速+5.1%不错、城镇化较高，但人口持续外流是重要约束。")

# ---- 九、城镇与农村 ----
heading1(doc, "九、城镇与农村：格局与均衡")
para(doc, "黑龙江城镇化率68.98%、+0.93pct，农村居民收入快于城镇。")
bullet(doc, "农村居民收入+5.8%、快于城镇+4.4%。")
bullet(doc, "粮食大省农业资源强、畜牧（猪牛羊禽+2.9%）。")
bullet(doc, "人口向哈尔滨/都市圈集中。")

# ---- 十、人口流入与流出 ----
heading1(doc, "十、人口流入与流出")
para(doc, "黑龙江常住约3001万人、人口持续外流（全国人口流出最重省份之一）。")
para(doc, "未来看点：冰雪/旅游/装备/对俄能否留住人口；若“冰雪产业+对俄开放+粮食经济”成势，黑龙江有望减缓人口流出。")

# ---- 十一、物价与货币环境 ----
heading1(doc, "十一、物价与货币环境")
para(doc, "2025年黑龙江CPI**+0.2%**、PPI**-5.3%**，物价温和、工业品价格深度走弱（石油石化）。")
para(doc, "物价偏弱反映“上游工业品价格下行”，与全国低通胀一致。2026年“扩内需、稳价格”是主线。")

# ---- 十二、区域一体化：对俄开放 ----
heading1(doc, "十二、区域一体化：黑龙江在“对俄开放+中俄合作+东北振兴”里的位置")
para(doc, "黑龙江的核心战略坐标是“**对俄开放前沿+东北振兴+中俄合作**”，是中俄/东北亚开放的窗口。")
bullet(doc, "对俄开放：边境贸易、黑河/绥芬河口岸、跨江大桥。")
bullet(doc, "中俄合作：对俄进出口增速高于全国。")
bullet(doc, "东北振兴：装备制造、冰雪经济、粮食基地。")
bullet(doc, "生态/冰雪：大小兴安岭、冰雪旅游大省。")
para(doc, "若“对俄开放+冰雪经济+装备制造+粮食安全”成势，黑龙江将作为东北亚开放与冰雪经济强省。")

# ============ 十三、五条主线 ============
heading1(doc, "十三、未来5—10年最值得观察的五条主线")
lines5 = [
    ("① 粮食安全/食品加工", "粮食基地、食材加工能否升级。"),
    ("② 冰雪经济/旅游", "冰雪产业、亚冬会能否持续壮大。"),
    ("③ 装备制造/电气", "发电机组、发动机、电气高端制造。"),
    ("④ 对俄开放/贸易", "对俄贸易、中俄合作能否深化。"),
    ("⑤ 人口/工业投资", "人口外流能否减缓、工业投资修复。"),
]
for l in lines5:
    para(doc, "**%s**　%s" % l, space_after=6)

# ---- 十四、结论 ----
heading1(doc, "十四、最终结论：黑龙江在“粮食+冰雪+装备+对俄”里的增长逻辑")
para(doc, "黑龙江的2025年，本质上是“**粮食/冰雪/装备/对俄为核心，而工业/投资/人口弱**”的答卷：GDP16878.0亿、+4.2%、粮食8200.3万吨（16年全国第1）、装备+8.4%、旅游花费+16.5%、入境+47.8%，但固投-10.3%、地产-7.3%、规上工业+2.8%、人口持续外流。")
para(doc, "只要粮食安全、冰雪经济、装备制造、对俄开放持续，黑龙江就站在“东北振兴/冰雪强省”的位置；如果工业/投资/人口持续偏弱，黑龙江需承受“农业/旅游强、工业/人口弱”的结构挑战。")
para(doc, "最稳观察信号：**一盯粮食（安全）、二盯冰雪（消费）、三盯装备（制造）、四盯对俄（开放）、五盯工业投资/人口（约束）。**黑龙江，是“粮食+冰雪+对俄”的样本。")

# ---- 附录A ----
heading1(doc, "附录A：主要资料来源与核验路径")
bullet(doc, "黑龙江2025年《政府工作报告》——目标来源。")
bullet(doc, "《2025年黑龙江省国民经济和社会发展统计公报》——GDP、工业、外贸、粮食、产业。")
bullet(doc, "黑龙江省统计（GDP）、哈尔滨海关（对俄）、省财政厅（财政）。")
bullet(doc, "2026年黑龙江省政府工作报告——2025复盘。")
heading2(doc, "核验说明")
para(doc, "本报告所有增速、总量、占比均以统计公报/官方口径为基准，涉“粮食/冰雪/装备/对俄”等以官方为口径。")

# ---- 附录B ----
heading1(doc, "附录B：建议建立的年度跟踪仪表盘")
para(doc, "建议把下面10个“测脉搏”指标做成年度跟踪表：")
dash = [
    ["#", "指标", "2025基值", "为什么重要"],
    ["1", "GDP增速", "+4.2%", "总量与方向"],
    ["2", "规上工业增速", "+2.8%", "制造底盘"],
    ["3", "粮食产量", "8200.3万吨", "粮食安全"],
    ["4", "装备/发电机组增长", "+8.4%/+27.5%", "高端制造"],
    ["5", "固定资产投资/地产", "-10.3%/-7.3%", "投资结构"],
    ["6", "社零增速", "+3.2%", "内需消费"],
    ["7", "常住人口/城镇化率", "3001万/68.98%", "人口与城市"],
    ["8", "一般公共预算收入", "+5.7%", "财政质量"],
    ["9", "对俄/出口增速", "+5.0%", "对俄开放"],
    ["10", "CPI/PPI", "+0.2%/-5.3%", "物价与需求"],
]
table(doc, dash[0], dash[1:], widths=[0.8, 4.6, 3.4, 4.4])
para(doc, "把这10个指标连起来看，装备（4）、对俄（9）等是否真在换挡。")

# ============ 保存 ============
out = "/Users/x/Desktop/content-prod-lab/reports/黑龙江省_2025年政府工作报告_深度研究_2026-08-13.docx"
doc.save(out)
print("SAVED:", out)
print("PARAGRAPHS:", len(doc.paragraphs))
print("TABLES:", len(doc.tables))
