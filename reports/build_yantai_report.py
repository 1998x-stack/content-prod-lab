# -*- coding: utf-8 -*-
"""Build 济宁市2025年政府工作报告 深度研究 DOCX, 参照地级市系列版式。"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DARK = RGBColor(0x1F, 0x2A, 0x44)
RED = RGBColor(0xB0, 0x1E, 0x2E)
GRAY = RGBColor(0x59, 0x60, 0x69)
BLUE = RGBColor(0x1F, 0x4E, 0x79)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
HEADER_FILL = "1F2A44"
K = "微软雅黑"


def set_run(run, size=11, bold=False, color=DARK, font=K):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    run.font.name = font
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), font)
    rFonts.set(qn("w:hAnsi"), font)
    rFonts.set(qn("w:eastAsia"), font)


def para(doc, text, size=11, bold=False, color=DARK, align=None,
         space_after=6, space_before=0, font=K):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        set_run(r, size=size, bold=bold or (i % 2 == 1), color=color, font=font)
    return p


def heading1(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(16)
    pf.space_after = Pt(8)
    r = p.add_run(text)
    set_run(r, size=16, bold=True, color=RED)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "B01E2E")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def heading2(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(10)
    pf.space_after = Pt(4)
    r = p.add_run(text)
    set_run(r, size=13, bold=True, color=BLUE)
    return p


def table(doc, headers, rows, widths=None, font_size=9.5):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_run(r, size=font_size, bold=True, color=WHITE)
        tcPr = hdr[i]._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), HEADER_FILL)
        tcPr.append(shd)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(val))
            set_run(r, size=font_size, color=DARK)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    return t


def quote_box(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(8)
    pf.left_indent = Pt(12)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), "B01E2E")
    pBdr.append(left)
    pPr.append(pBdr)
    r = p.add_run(text)
    set_run(r, size=11, color=DARK, font="楷体")
    return p


def bullet(doc, text, size=10.5):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.7)
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        set_run(r, size=size, bold=(i % 2 == 1), color=DARK)
    return p


# ================================================================= 文档主体
doc = Document()
sec = doc.sections[0]
sec.page_width = Cm(21)
sec.page_height = Cm(29.7)
sec.top_margin = Cm(2.4)
sec.bottom_margin = Cm(2.2)
sec.left_margin = Cm(2.5)
sec.right_margin = Cm(2.5)

st = doc.styles["Normal"]
st.font.name = K
st.font.size = Pt(11)
st.element.rPr.rFonts.set(qn("w:eastAsia"), K)


# ---- 封面 ----
para(doc, "烟台市2025年政府工作报告", size=24, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=60, space_after=4)
para(doc, "深度研究与\u201c容易被忽视的细节\u201d分析", size=16, bold=True, color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
para(doc, "从\u201c万亿制造强市、海洋经济、绿色石化、东方航天港、清洁能源\u201d重新理解烟台", size=12, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)
para(doc, "\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501", color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
para(doc, "研究对象：2025年烟台市政府工作报告及同期财政、国民经济和社会发展统计资料、2026年官方复盘", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
para(doc, "标定日期：2026-08-14", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
doc.add_page_break()

# ---- 目录 ----
catalog = [
    "执行摘要",
    "一、研究口径与阅读方法",
    "二、先看烟台的特别底盘：万亿制造、海洋经济、裕龙万华、东方航天港",
    "三、最关键的宏观错位：工业出口海洋强，但固投-18.8%、财收+0.2%、CPI弱",
    "四、容易被忽视、但信号很重的细节（15条）",
    "五、2025年GDP目标 vs 实际：对照表",
    "六、2025年增长是谁撑起来的？（结构归因）",
    "七、预算与财政的\u201c含金量\u201d",
    "八、民生底账：人口、收入与城乡",
    "九、城镇与农村：格局与均衡",
    "十、人口流入与流出",
    "十一、物价与货币环境",
    "十二、区域一体化：烟台在胶东经济圈、山东半岛沿海经济带、京津冀\u201c蓝\u201d里的位置",
    "十三、未来5—10年最值得观察的五条主线",
    "十四、最终结论：烟台在\u201c工业临港+海洋经济+绿色低碳\u201d里的增长逻辑",
    "附录A：主要资料来源与核验路径",
    "附录B：建议建立的年度跟踪仪表盘",
]
para(doc, "目　录", size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=10, space_after=10)
for item in catalog:
    para(doc, item, size=12, space_after=4)
doc.add_page_break()

# ---- 执行摘要 ----
heading1(doc, "执行摘要")
para(doc, "**核心判断**　2025年烟台最显著的是\u201cGDP 11350.98亿/+6.1%（突破万亿、高于全国全省、山东第3）、规上工业+13.5%（制造业+16.1%）、进出口5379.32亿/+13.7%、海洋生产总值破3000亿\u201d、\u201c但固定资产投资-18.8%、民间投资-26%、房地产开发-6.1%、一般公共预算收入+0.2%（差3.5%目标）、CPI-0.2%\u201d。这说明烟台在\u201c万亿工业强市+海洋经济\u201d中，**工业海洋外贸强、投资地产财政弱**。")
para(doc, "把2025年目标（GDP+5.5%、规上+7.5%、财收+3.5%、社零高于全省）、2025年实际（GDP+6.1%、规上+13.5%、社零+5.8%、进出口+13.7%、财收+0.2%）趋势看，烟台是\u201c临港重工业+海洋\u201d路径：**石化（裕龙/万华）、有色、机械、核电（绿色）、海洋\u201d是支柱；规上工业总产值1.3万亿。")
para(doc, "最容易记住的一句话：**烟台是\u201c万亿工业强市、海洋经济大市\u201d，靠\u201c石化新能源+有色装备+海洋\u201d增长。**观察烟台，与其只看\u201cGDP 11351亿\u201d，不如看\u201c规上+13.5%、进出口破5000亿（+13.7%）、海洋3000亿、裕龙万华投产、清洁能源2000万千瓦\u201d。")
# ---- 一、研究口径与阅读方法 ----
heading1(doc, "一、研究口径与阅读方法")
para(doc, "本报告以\u201c烟台市政府工作报告（2025年1月，郑德雁作）\u201d为起点，把\u201c2025年GDP目标（5.5%）\u201d与\u201c官方2025年（11350.98亿/+6.1%）\u201d并置对照，用\u201c2025年烟台市统计公报\u201d和\u201c2026年政府工作报告\u201d横向核验。")
para(doc, "口径提示：\u201cGDP、规上工业增速\u201d按可比价（实际）；\u201c投资、消费、进出口、财政\u201d按现价（名义）。\u201c常住人口\u201d用统计公报（700.05万）、城镇化率70.92%。")
para(doc, "指标体系（与研究口径一致）：**总量与增速、产业动能（石化/有色/海洋/核电）、外贸、财政、民生与人口**。")
para(doc, "特别提示（不吃老本）：烟台2024年GDP迈入万亿（+5.8%）、2025年11350.98亿/+6.1%（北方首个万亿级地级市）；它不是\u201c只有葡萄/海鲜\u201d——**裕龙石化、万华化工、有色、核电、海洋经济、东方航天港\u201d才是真正底色。")
# ---- 二、先看烟台的特别底盘 ----
heading1(doc, "二、先看烟台的特别底盘：万亿制造、海洋经济、裕龙万华、东方航天港")
para(doc, "烟台地处山东半岛东北部、渤海与黄海之滨，是**国家万亿级工业强市、海洋经济大市、绿色低碳发展示范区**；是\u201c裕龙石化、万华化学、东方航天港、核能之城（海阳）\u201d的代名词。2025年GDP 11350.98亿（北方首个万亿级地级市）、常住700.05万、城镇化率70.92%、山东第3（青岛/济南之后）。")
para(doc, "五个底盘名词，先立框架：")
bullet(doc, "**绿色石化（裕龙/万华）**　裕龙岛炼化一体化一期（1188亿）、万华新材料低碳产业园（1260亿）、石油加工+311.7%/化工+46.2%——\u201c石化龙头\u201d。")
bullet(doc, "**有色与机械制造**　有色/黄金（亚洲最大海底金矿）、机械装备（潍柴弗迪新能源动力产业园）——\u201c重工业\u201d。")
bullet(doc, "**绿色能源（核电/风光）**　清洁能源装机2000万千瓦（全省第1）、海阳零碳供暖（全国首个）——\u201c绿色低碳\u201d。")
bullet(doc, "**海洋经济**　海洋生产总值破3000亿、港口吞吐5.4亿吨/集箱556万标箱、海上发射母港（东方航天港、发射22次137颗卫星）——\u201c海洋强市\u201d。")
bullet(doc, "**航空航天/高端装备**　东方航天港、潍柴弗迪新能源动力——\u201c新质生产\u201d。")
para(doc, "这五根（石化+有色机械+绿色能源+海洋+航天）构成烟台独特底盘：**左手万亿工业（石化/有色），右手海洋+太空**。理解烟台，先理解\u201c工业强市、海洋大市、绿色低碳\u201d。")
# ---- 三、最关键的宏观错位 ----
heading1(doc, "三、最关键的宏观错位：工业海洋外贸强，但固投-18.8%、财收+0.2%、CPI弱")
para(doc, "2025年烟台最需要辨析的一组\u201c错位\u201d：**规上工业+13.5%（制造业+16.1%）强、进出口+13.7%、海洋3000亿，但固定资产投资-18.8%、民间投资-26%、房地产开发-6.1%、一般公共预算收入+0.2%（差3.5%目标）、CPI-0.2%**。")
para(doc, "为什么\u201c工业/海洋/外贸这么强\u201d，投资与财政却弱？三解释：")
para(doc, "**其一，石化/有色/海洋扩产、体量巨大**　规上+13.5%（石油加工+311.7%、化工+46.2%、有色+15.7%）、裕龙/万华投产——\u201c产能释放大\u201d。")
para(doc, "**其二，投资/地产/财政偏弱**　固投-18.8%（制造业-34.2%、民间-26%）、地产-6.1%/销售-15.2%、财收+0.2%（税收-4.3%）——\u201c投资财政弱\u201d。")
para(doc, "**其三，外贸强、消费/物价温**　进出口+13.7%破5000亿、社零+5.8%（线上+32.3%）；但CPI-0.2%、PPI弱——\u201c量好价弱\u201d。")
para(doc, "小结：烟台2025年是\u201c**工业海洋外贸强、投资财政地产弱**\u201d：石化/有色/海洋强，固投、财收、物价弱。")

# ---- 四、容易被忽视、但信号很重的细节（15条） ----
heading1(doc, "四、容易被忽视、但信号很重的细节（15条）")
bullet(doc, "**1.规上工业+13.5%（制造业+16.1%）**\u201c裕龙/万华/石化有色拉动。\u201d")
bullet(doc, "**2.石油加工+311.7%、化学制造+46.1%、有色+15.7%**\u201c石化有色爆发、绿色石化龙头。\u201d")
bullet(doc, "**3.裕龙岛炼化一期/万华系列投产（总投资1188亿/1260亿）**\u201c万亿级重工业落地。\u201d")
bullet(doc, "**4.进出口5379.32亿/+13.7%、进口+23.2%**\u201c对欧盟+26.7%、一带一路+13.1%——外贸强。\u201d")
bullet(doc, "**5.海洋生产总值破3000亿、港口吞吐5.4亿吨/集箱556万**\u201c海洋经济大市。\u201d")
bullet(doc, "**6.东方航天港：发射22次/137颗卫星（全国唯一海上发射母港）**\u201c航天+海洋。\u201d")
bullet(doc, "**7.清洁能源装机2000万千瓦（全省第1）、海阳零碳供暖**\u201c绿色低碳、核能之城。\u201d")
bullet(doc, "**8.固定资产投资-18.8%（制造业-34.2%、民间-26%）**\u201c投资大降、扩产基数。\u201d")
bullet(doc, "**9.房地产开发-6.1%、商品房销售-15.2%**\u201c地产调整。\u201d")
bullet(doc, "**10.一般公共预算收入700.08亿/+0.2%（税收-4.3%）**\u201c财政弱、税基缺。\u201d")
bullet(doc, "**11.社零4068.85亿/+5.8%、线上+32.3%、餐饮+9.9%**\u201c消费/文旅热（旅游游客1亿人次/收入1265亿）。\u201d")
bullet(doc, "**12.潍柴弗迪新能源动力产业园投产后（首期280亿）**\u201c新能源汽车+100.9%、新能源。\u201d")
bullet(doc, "**13.居民收入53562元/+5.3%、城乡比2.03**\u201c农村+5.7%>城镇+5.8%。\u201d")
bullet(doc, "**14.CPI-0.2%（食品烟酒-0.6%、蛋-8.6%）**\u201c低通胀。\u201d")
bullet(doc, "**15.常住700.05万、城镇化率70.92%**\u201c山东次人口城市、城镇化稳升。\u201d")
# ---- 五、2025年GDP目标 vs 实际：对照表 ----
heading1(doc, "五、2025年GDP目标 vs 实际：对照表")
table(doc,
    ["指标", "2025年目标", "2025年实际", "达成评价"],
    [
        ["地区生产总值(GDP)", "增长5.5%左右", "11350.98亿/6.1%", "大幅超额"],
        ["规模以上工业", "增长7.5%以上", "+13.5%(制造+16.1%)", "大幅超额"],
        ["一般公共预算收入", "增长3.5%左右", "700.08亿/+0.2%", "差3.3pct"],
        ["固定资产投资", "合理增长", "-18.8%", "大幅下行"],
        ["社会消费品零售总额", "高于全省平均", "4068.85亿/+5.8%", "超额"],
        ["进出口总额", "稳量提质", "5379.32亿/+13.7%", "大幅超额"],
        ["居民收入", "与经济增长同步", "53562元/+5.3%", "略低GDP"],
    ],
)
para(doc, "注：GDP、规上工业按可比价。**GDP（+6.1%）、规上（+13.5%）、进出口（+13.7%）、社零（+5.8%）超额**；**固投（-18.8%）、财收（+0.2%差3.5%）、房地产（-6.1%）**偏弱。")
para(doc, "拆读：**石化/有色/海洋/出口是亮色**；**固投（-18.8%）、财政（+0.2%）、税收（-4.3%）、物价（-0.2%）**是短板——\u201c工业外贸强、投资财政弱\u201d，是\u201c万亿工业强市\u201d样本。")

# ---- 六、2025年增长是谁撑起来的？（结构归因） ----
heading1(doc, "六、2025年增长是谁撑起来的？（结构归因）")
para(doc, "把烟台GDP的6.1%拆开：三次产业分别增4.1%、8.0%、4.8%（结构6.2：42.4：51.4）。**第二产业（工业）增速8.0%、贡献最大**，第三产业（服务业）占比51.4%，第一产业（农业/海洋）稳\u2014\u2014\u201c二产领跑\u201d。")
para(doc, "2026年烟台强调\u201c万亿新征程、海洋强市、绿色低碳\u201d，聚焦**石化/绿色能源、海洋经济、航空、新材料、县域千亿（龙口破2000亿/招远莱州千亿）**——核心是\u201c工业+海洋+绿色\u201d。")
para(doc, "**第二产业（工业）**：规上+13.5%（石油加工+311.7%、化工+46.2%、有色+15.7%）、裕龙/万华/潍柴弗迪——\u201c重化+绿色制造\u201d强。")
para(doc, "**第三产业（服务业）**：+4.9%（商贸、物流、港口/集装箱、旅游收入1265亿）——\u201c服务+开放\u201d。")
para(doc, "**第一产业（农业/海洋）**：+4.1%（果品大市：苹果702.4万吨+3.1%、海洋渔业）——\u201c农业海洋稳\u201d。")
para(doc, "一句话归因：**2025年烟台增长\u201c靠第二产业（石化有色制造）+海洋+外贸\u201d**，固定投资、财政弱；\u201c工业海洋\u201d是核心。")

# ---- 七、预算与财政的\u201c含金量\u201d ----
heading1(doc, "七、预算与财政的\u201c含金量\u201d")
para(doc, "2025年烟台**一般公共预算收入700.08亿元（+0.2%）**；税收409.79亿（-4.3%、占比58.5%）；一般公共预算支出1051.93亿（+4.7%）、民生占79.5%。")
bullet(doc, "税收-4.3%、占比58.5%（财政质量受规上制造业基-拖累）——\u201c税基承压\u201d。")
bullet(doc, "民生支出占79.5%（教育/社保/医疗）。")
bullet(doc, "金融支撑：存款+9.0%、贷款+10.4%——信贷充裕支持万亿工业/海洋/绿色。")
para(doc, "**财政含金量小结**：财收+0.2%（低GDP增速）、税收-4.3%（含金量待升），民生79.5%；财政对\u201c绿色低碳、海洋、新质、民生\u201d投入加大。")

# ---- 八、民生底账：人口、收入与城乡 ----
heading1(doc, "八、民生底账：人口、收入与城乡")
para(doc, "2025年烟台**居民人均可支配收入53562元（+5.3%）**，其中城镇65009元（+4.8%）、农村31965元（+5.7%），城乡比2.03（缩小0.02）。消费：人均消费支出32904元（+3.9%）。就业：城镇新增就业11.1万人。")
para(doc, "人口画像：**常住700.05万、城镇化率70.92%（+0.81pct）**；烟台工业/海洋/县域吸纳就业、出生率偏低。")
para(doc, "民生投入：民生支出3584亿（十四五75.7%）、城镇新增就业十四五54万、改造老旧小区1662个、公共医疗教育——民生扎实。")

# ---- 九、城镇与农村：格局与均衡 ----
heading1(doc, "九、城镇与农村：格局与均衡")
para(doc, "烟台城镇化率70.92%；县域经济极强（龙口冲刺2000亿/招远千亿/莱州千亿县、福山/芝罘制造业），苹果大市（702.4万吨）；农村收入增速（+5.7%）>城镇（+4.8%），**城乡比2.03缩小**；高标准农田、乡村振兴。")
para(doc, "农业底盘：**粮食191.29万吨、油料41.14万吨、水果898.97万吨（苹果702.4万吨+3.1%）、海洋渔业**——\u201c果品大市+海洋牧场\u201d。")
para(doc, "一句话：\u201c烟台是万亿工业+果品海洋大市、县域（龙口）发达\u201d，城乡均衡全国前列。")

# ---- 十、人口流入与流出 ----
heading1(doc, "十、人口流入与流出")
para(doc, "烟台常住700.05万（山东第6人口城市）、城镇化70.92%；\u201c制造业/海洋/核电\u201d吸纳就业，但人口总量偏稳、部分青年外流至北京/青岛；主城区/开发区吸引人才。")
para(doc, "结构观察：**出生率偏低（老龄化）、海洋人才（海洋大学）**；引进人才十四五31.5万。")
para(doc, "2026年目标：城镇新增就业10万人以上、引进人才——烟台靠\u201c工业+海洋+绿色\u201d聚人。")

# ---- 十一、物价与货币环境 ----
heading1(doc, "十一、物价与货币环境")
para(doc, "2025年烟台**CPI-0.2%**（食品烟酒-0.6%、交通通信-2.9%、蛋-8.6%、畜肉-5.7%；其他用品+13.1%）——\u201c低通胀、需求弱\u201d。")
bullet(doc, "信贷扩张：存款+9.0%、贷款+10.4%——宽信用支持万亿工业/海洋。")
bullet(doc, "消费：社零+5.8%、线上+32.3%、餐饮+9.9%、文旅（游客1亿/收入1265亿）——消费热。")
para(doc, "货币环境判断：**宽信用、CPI-0.2%**；烟台靠\u201c工业+海洋+消费\u201d稳需求（2026 CPI合理）。")

# ---- 十二、区域一体化：烟台的位置 ----
heading1(doc, "十二、区域一体化：烟台在胶东经济圈、山东半岛沿海经济带、京津冀\u201c蓝\u201d里的位置")
para(doc, "烟台是**胶东经济圈核心、山东半岛沿海经济带龙头、环渤海经济圈、对接日韩/京津冀（京津冀协同）重要枢纽**。")
bullet(doc, "**胶东经济圈**　青岛/烟台/威海一体、海洋强省、胶东一体化。")
bullet(doc, "**对日韩/开放**　对欧盟+26.7%、一带一路+13.1%、RCEP 1817亿——开放强市、日韩近樑。")
bullet(doc, "**向海向洋**　东方航天港、海洋牧场、环渤海、海上发射。")
para(doc, "一句话：**烟台在\u201c胶东+渤海+对日韩\u201d里，最核心是\u201c工业+海洋+开放\u201d**；区位、万亿制造、港口是最大优势。")

# ---- 十三、未来5—10年最值得观察的五条主线 ----
heading1(doc, "十三、未来5—10年最值得观察的五条主线")
bullet(doc, "**主线一：绿色石化与新材料（裕龙/万华）**\u201c石化炼化、新材料\u201d能否冲击万亿产业集群。")
bullet(doc, "**主线二：海洋经济（蓝色粮仓/海上发射）**\u201c海洋3000亿、航天港、深海\u201d。")
bullet(doc, "**主线三：绿色能源（核电/风光）**\u201c清洁能源2000万→3500万、零碳园区\u201d。")
bullet(doc, "**主线四：高端制造/新能源车（潍柴弗迪）**\u201c动力电池、航空航天\u201d新质。")
bullet(doc, "**主线五：人口/县域/财政质量**\u201c龙口2000亿、千亿县、财收\u201d如何\u201c提优、聚人\u201d。")

# ---- 十四、最终结论 ----
heading1(doc, "十四、最终结论：烟台在\u201c工业临港+海洋经济+绿色低碳\u201d里的增长逻辑")
para(doc, "把2025年闭环打穿：**烟台是\u201c万亿工业强市、海洋经济大市、绿色低碳\u201d**：GDP 11350.98亿/+6.1%（北方首个万亿地级市）、规上+13.5%、进出口5379.32亿、海洋3000亿、清洁能源2000万千瓦。")
para(doc, "烟台不是\u201c只有海鲜/葡萄酒\u201d——它是**绿色石化+有色+核电绿色能源+海洋+航空航天**的复合，靠\u201c万亿工业+海洋\u201d驱动；但固定投资、财政、物价弱。")
para(doc, "一句话结论：**烟台是\u201c万亿工业大市、海洋强市、绿色之城\u201d；观察它先看\u201c石化/新能源、海洋、进出口、港口\u201d，再看\u201c固投、财政、税收\u201d。**它是\u201c工业海洋强、投资财政待优\u201d的山东样本。")

# ---- 附录A ----
heading1(doc, "附录A：主要资料来源与核验路径")
bullet(doc, "《2025年烟台市政府工作报告》（2025年1月，郑德雁作，2025年目标、2024年回顾万亿）")
bullet(doc, "《2025年烟台市国民经济和社会发展统计公报》（烟台市统计局，2026-03-31，2025年实际）")
bullet(doc, "《2026年烟台市政府工作报告》（2026年1月，复盘+2026年目标）")
bullet(doc, "烟台市人民政府/统计局、中国经济网等")

# ---- 附录B ----
heading1(doc, "附录B：建议建立的年度跟踪仪表盘")
bullet(doc, "GDP总量/增速 vs 全年目标。")
bullet(doc, "规上工业（石化/有色/制造）增速。")
bullet(doc, "裕龙/万华/潍柴投产与产能。")
bullet(doc, "固定资产/工业/房地产投资。")
bullet(doc, "港口吞吐、集箱、海洋经济。")
bullet(doc, "社零、线上、旅游。")
bullet(doc, "进出口、对欧盟/日韩、外资。")
bullet(doc, "一般公共预算/税收占比/民生%。")
bullet(doc, "常住/城镇化、青年、县域千亿。")
bullet(doc, "CPI、存贷款、清洁能源/发射。")

# ---------------- 保存 ----------------
out = "/Users/x/Desktop/content-prod-lab/reports/烟台市_2025年政府工作报告_深度研究_2026-08-14.docx"
doc.save(out)
print("SAVED OK: 烟台市", out)
