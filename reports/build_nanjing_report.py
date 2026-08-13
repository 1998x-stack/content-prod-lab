# -*- coding: utf-8 -*-
"""Build 南京市2025年政府工作报告 深度研究 DOCX, 参照省市系列版式。"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DARK = RGBColor(0x1F, 0x2A, 0x44)
RED = RGBColor(0xB0, 0x1E, 0x2E)
GRAY = RGBColor(0x59, 0x60, 0x69)
BLUE = RGBColor(0x1F, 0x4E, 0x79)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
HEADER_FILL = "1F2A44"
K = "微软雅黑"


def set_run(run, size=11, bold=False, color=DARK, font=K):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    run.font.name = font
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), font)
    rFonts.set(qn("w:hAnsi"), font)
    rFonts.set(qn("w:eastAsia"), font)


def para(doc, text, size=11, bold=False, color=DARK, align=None,
         space_after=6, space_before=0, font=K):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        set_run(r, size=size, bold=bold or (i % 2 == 1), color=color, font=font)
    return p


def heading1(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(16)
    pf.space_after = Pt(8)
    r = p.add_run(text)
    set_run(r, size=16, bold=True, color=RED)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "B01E2E")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def heading2(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(10)
    pf.space_after = Pt(4)
    r = p.add_run(text)
    set_run(r, size=13, bold=True, color=BLUE)
    return p


def table(doc, headers, rows, widths=None, font_size=9.5):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_run(r, size=font_size, bold=True, color=WHITE)
        tcPr = hdr[i]._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), HEADER_FILL)
        tcPr.append(shd)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(val))
            set_run(r, size=font_size, color=DARK)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    return t


def quote_box(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(8)
    pf.left_indent = Pt(12)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), "B01E2E")
    pBdr.append(left)
    pPr.append(pBdr)
    r = p.add_run(text)
    set_run(r, size=11, color=DARK, font="楷体")
    return p


def bullet(doc, text, size=10.5):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.7)
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        set_run(r, size=size, bold=(i % 2 == 1), color=DARK)
    return p


# ================================================================= 文档主体
doc = Document()
sec = doc.sections[0]
sec.page_width = Cm(21)
sec.page_height = Cm(29.7)
sec.top_margin = Cm(2.4)
sec.bottom_margin = Cm(2.2)
sec.left_margin = Cm(2.5)
sec.right_margin = Cm(2.5)

st = doc.styles["Normal"]
st.font.name = K
st.font.size = Pt(11)
st.element.rPr.rFonts.set(qn("w:eastAsia"), K)



# ---- 封面 ----
para(doc, "南京市2025年政府工作报告", size=24, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=60, space_after=4)
para(doc, "深度研究与\u201c容易被忽视的细节\u201d分析", size=16, bold=True, color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
para(doc, "从\u201c软件名城、智能电网、科教产学研与长三角\u201d重新理解南京", size=12, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)
para(doc, "\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501", color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
para(doc, "研究对象：2025年南京市政府工作报告及同期财政、国民经济和社会发展统计资料、2026年官方复盘", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
para(doc, "标定日期：2026-08-13", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
doc.add_page_break()

# ---- 目录 ----
catalog = [
    "执行摘要",
    "一、研究口径与阅读方法",
    "二、先看南京的特殊底盘：软件名城、智能电网、产学研与长三角城市群",
    "三、最关键的宏观错位：GDP破1.94万亿、三产强，但固投/外贸/地产偏弱",
    "四、容易被忽视、但信号很重的细节（15条）",
    "五、2025年 GDP/目标 vs 实际：对照表",
    "六、2025年增长是谁撑起来的？（结构归因）",
    "七、预算与财政的\u201c含金量\u201d",
    "八、民生底账：人口、收入与城乡",
    "九、城镇与农村：格局与均衡",
    "十、人口流入与流出",
    "十一、物价与货币环境",
    "十二、区域一体化：南京在\u201c南京都市圈+长三角一体化+宁镇扬\u201d里的位置",
    "十三、未来5—10年最值得观察的五条主线",
    "十四、最终结论：南京在\u201c软件/智能电网+产学研+省会引领\u201d里的增长逻辑",
    "附录A：主要资料来源与核验路径",
    "附录B：建议建立的年度跟踪仪表盘",
]
para(doc, "目　录", size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=10, space_after=10)
for item in catalog:
    para(doc, item, size=12, space_after=4)
doc.add_page_break()

# ---- 执行摘要 ----
heading1(doc, "执行摘要")
para(doc, "**核心判断**　如果只看新闻标题，2025年南京最显著的是\u201cGDP破1.94万亿、增长5.2%\u201d、\u201c软件与信息服务成为首个万亿级产业\u201d、\u201c智能电网产业规模超五千亿\u201d、\u201c常住人口963.85万\u201d。但这份研究真正值得深读的，是这座\u201c软件名城+科教产学研+长三角综合枢纽\u201d的省会，如何在消费回升（+3.5%）、外贸小幅回落（-1.4%）、房地产调整（固投-4.9%、房投-13.1%）的背景下，靠\u201c三产(+6.0%)+软件/智能电网/制造业升级+科教人才\u201d实现5.2%的稳健增长。")
para(doc, "把2025年初设定的目标（GDP增长5%左右）、2025年《国民经济和社会发展统计公报》、2026年复盘放在一起看，南京呈现清晰暗线：**从\u201c重工业/房地产\u201d的旧底盘，向\u201c数字经济/软件+智能电网+先进制造+科教产学研\u201d的高能级升级**。旧引擎（钢铁石化、部分地产）在调整；新引擎（软件万亿、智能电网五千亿、智能制造、科教人才）被要求更快补位。")
para(doc, "因此，本报告不按政府工作报告逐段复述，而采用\u201c显性表述—同期数据—制度含义—长期影响\u201d的方式，专门提取容易被忽略、但对判断南京未来5—10年发展模式更有价值的信号。")
para(doc, "最容易记住的一句话：**南京是\u201c软件名城+科教产学研+长三角综合枢纽\u201d的高能级省会，靠\u201c软件/智能电网+三产+科教人才+都市圈\u201d撑起增长。**观察南京，与其看\u201cGDP 1.94万亿\u201d，不如看\u201c软件万亿、智能电网五千亿、研发强度4%、城镇化87.5%、万人发明专利全国前三\u201d这几张名片。")
heading2(doc, "一页速览：2025年南京经济的\u201c表与里\u201d")
table(doc,
    ["维度", "表面现象", "更深层信号"],
    [
        ["增长", "GDP 19428.78亿、+5.2%", "一产1.7%、二产30.2%、三产68.0%"],
        ["产业", "规上工业+5.8%", "电气+5.2%、医药+10.8%、集成电路+28.2%"],
        ["外贸", "进出口5380.9亿、-1.4%", "出口+0.7%、外资-53.7%"],
        ["投资", "固投-4.9%、扣除地产+2.4%", "工业+12.1%、房建-13.9%"],
        ["消费", "社零8135.85亿、+3.5%", "新能源汽车+33.1%、金银珠宝+15.4%"],
        ["人口", "常住963.85万、城镇化87.5%", "人口+6.15万、自然增0.35‰"],
        ["数字", "软件万亿级、智能电网超五千亿", "科教产学研、研发强度4%"],
    ],
    widths=[2.2, 5.6, 7.4])
para(doc, "", size=10, space_after=2)

# ---- 一、研究口径与阅读方法 ----
heading1(doc, "一、研究口径与阅读方法")
heading2(doc, "1.1 本报告的核心底稿")
bullet(doc, "**2025年《政府工作报告》**（2025年1月）——目标：GDP 5%左右、一般公共预算收入正增长、社零5%左右、城镇就业21万。")
bullet(doc, "**《2025年南京市统计公报》**（市统计局2026-03）——GDP、工业、外贸、人口实数。")
bullet(doc, "**2026年南京市政府工作报告/复盘**——对2025执行追认与软件/智能电网展望。")
heading2(doc, "1.2 阅读方法：显性—数据—制度—长期")
para(doc, "**关键判别**：数据优先。例2025年GDP目标5%、实际5.2%达标；社零目标5%、实际3.5%偏弱。南京\u201c服务业/软件/科教强、消费/外贸/地产偏弱、结构在换\u201d，需穿透GDP总量看\u201c万亿软件+5000亿智能电网+科研4%\u201d。")

# ---- 二、底盘 ----
heading1(doc, "二、先看南京的特殊底盘：软件名城、智能电网、产学研与长三角城市群")
para(doc, "南京的地盘取决于它作为\u201c**软件名城+智能电网+科教产学研中心+长三角综合枢纽**\u201d的特殊定位。它是江苏省会、\u201c六朝古都·金陵\u201d，位于长三角城市群核心、长江经济带。")
bullet(doc, "**软件名城**：软件和信息服务业成为全国首批万亿级产业；在工信部软件名城评估中位居第四；重点新兴平台软件、工业软件、信创、开源生态。")
bullet(doc, "**智能电网/新型电力**：智能电网产业规模超五千亿元；中国电科院南京科研基地、南瑞继保智能化电气装备等支撑\u201c电力智能化\u201d名片。")
bullet(doc, "**科教产学研**：研发投入强度达4%（全国前列）；紫金山实验室发布全球首个6G通智感融合外场试验网；全国重点实验室34家；万人有效发明专利187.21件。")
bullet(doc, "**综合枢纽/都市圈**：\u201c南京都市圈+长三角一体化+宁镇扬\u201d；高铁/轨道交通发达；制造与服务业双轮驱动。")
para(doc, "这一底板决定了南京2025成绩单的\u201c底色\u201d：**只要软件/智能电网/科教/长三角服务持续，南京就站在\u201c数字经济+高端制造\u201d增长极；若重工业/外贸/地产承压，南京需承受\u201c服务业强、传统制造与外贸偏弱\u201d的结构挑战。**")

# ---- 三、宏观错位 ----
heading1(doc, "三、最关键的宏观错位：GDP破1.94万亿、三产强，但固投/外贸/地产偏弱")
para(doc, "南京2025年最值得咀嚼的错位，是\u201c**增长靠三产+软件/科教+高端制造，固投/外贸/地产却偏弱**\u201d。这种错位决定了对座上这座长三角省会城市的观察不能只看GDP总量。")
bullet(doc, "**GDP**：19428.78亿、+5.2%。一产338.50亿（+3.4%，占比1.7%）、二产5873.07亿（+3.7%，占比30.2%）、三产13217.21亿（+6.0%，占比68.0%）。")
bullet(doc, "**工业**：规上工业增加值+5.8%；电气机械+5.2%、医药+10.8%、黑色金属+11.0%、仪器仪表+12.1%；计算机通信电子+7.0%。新能源汽车产量+99.0%、工业机器人+35.4%、集成电路+28.2%。")
bullet(doc, "**消费**：社零8135.85亿、+3.5%；新能源汽车零售+33.1%、金银珠宝+15.4%。")
bullet(doc, "**外贸**：进出口5380.9亿、-1.4%（出口+0.7%、进口-5.1%）；实际使用外资13.12亿美元、-53.7%。")
bullet(doc, "**固投/地产**：固投-4.9%、扣除房地产后+2.4%（工业投资+12.1%）；房地产投资-13.1%、商品房销售-8.4%。")
para(doc, "**为什么读这条**：南京作为\u201c长三角服务+科教型省会\u201d，现阶段结构性矛盾是\u201c服务业强、外贸/外资弱、地产调整、传统制造承压\u201d。经济靠三产和科技制造业支撑，但外部开放与房地产开发仍偏弱。")

# ---- 四、15条细节 ----
heading1(doc, "四、容易被忽视、但信号很重的细节（15条）")
para(doc, "以下15条，都在2025年报告/公报里存在，但常被\u201cGDP 1.94万亿\u201d、\u201c5.2%\u201d等总量叙事掩盖。它们是判断南京2025之后5—10年的关键小信号。")
bullet(doc, "**1. 软件成为首个万亿级产业**：软件与信息服务业跨过万亿，是南京\u201c数字经济\u201d的里程碑。")
bullet(doc, "**2. 智能电网产业规模超五千亿**：\u201c新型电力\u201d是南京特色名片，全国领先。")
bullet(doc, "**3. 规上工业+5.8%**：略低于GDP（5.2%），但先进制造（集成电路+28.2%、汽车+99.9%、工业机器人+35.4%）拉动强劲。")
bullet(doc, "**4. 民营工业企业+11.1%**：高于国有（+8.2%）、股份制（+9.3%），凸显民营制造与数字经济活力。")
bullet(doc, "**5. 三产+6.0%、占GDP 68%**：服务业主导，现代服务业/科教是南京最硬的底盘。")
bullet(doc, "**6. 科研强度：全社会研发强度约4%（目标）**：投入强度全国前列，科技教育高地。")
bullet(doc, "**7. 万人有效发明专利187.21件**：与北京、深圳并列全国前列；有效发明专利17.9万件。")
bullet(doc, "**8. 人才+415万**：人才资源总量同比增长8.9%，科教/创新人才聚集。")
bullet(doc, "**9. 独角兽企业12家、新增上市公司5家（累计167家）**：科创资本生态。")
bullet(doc, "**10. 未来产业**：未来网络、虚拟现实等6个省级试点；未来产业业务收入增长20%。")
bullet(doc, "**11. 城市更新/品质**：颐和路入选全国城市更新典型，62个城市更新项目、138个老旧小区改造。")
bullet(doc, "**12. 国际化/入境游**：入境过夜游客68.2万人次、+17.6%；旅游花费2390.2亿、+9.3%。")
bullet(doc, "**13. 财政质量**：地方一般公共预算收入1620.91亿、+1.6%，税收占比约81.7%（高）。")
bullet(doc, "**14. 绿色低碳**：PM2.5年均27.1微克/立方米、优良率87.4%；全社会用电量823.78亿千瓦时、+2.6%。")
bullet(doc, "**15. 都市圈互认互贷**：南京都市圈公积金、地铁450km、长三角大飞机集群华东办。")
para(doc, "", size=10, space_after=2)

# ---- 五、对照表 ----
heading1(doc, "五、2025年 GDP/目标 vs 实际：对照表")
para(doc, "2025年报告中列出的主要预期目标，与2025年《统计公报》实际完成情况：")
tb = [
    ["指标", "2025年目标", "2025年实际（公报/报告）", "达标判定"],
    ["GDP增速", "5%左右", "+5.1%（19428.78亿）", "基本达标"],
    ["一般公共预算收入", "正增长", "+1.6%（1620.91亿）", "达标"],
    ["社会消费品零售总额", "5%左右", "+3.5%（8135.85亿）", "偏弱"],
    ["城镇新增就业", "21万人", "21.91万人（公报口径）", "达标"],
    ["外贸进出口", "稳中提质", "5380.9亿、-1.4%", "回落"],
    ["全社会研发强度", "4%", "约3.9%（公报）", "接近"],
    ["居民人均可支配收入", "与经济增长同步", "+4.1%（78243元）", "基本同步"],
]
table(doc, tb[0], tb[1:], widths=[3.2, 3.2, 4.6, 4.0])
para(doc, "", size=10, space_after=2)
para(doc, "**要点**：GDP/就业/财政达标，社零/进出口偏弱——南京\u201c收入端稳定、消费/外贸待修\u201d。")

# ---- 六、结构归因 ----
heading1(doc, "六、2025年增长是谁撑起来的？（结构归因）")
heading2(doc, "6.1 三次产业：三产绝对主导")
para(doc, "**三产（第三产业）是绝对引擎**：13217.21亿、+6.0%、占比68%。一、二产基数小、增速屈居（一产+3.4%、二产+3.7%、占比30.2%）。三产靠软信/金融/科技服务拉动，南京“服务业+科教”模式鲜明。")
heading2(doc, "6.2 工业：高技术/先进制造驱动")
para(doc, "规上工业+5.8%；新能源汽车+99.9%、集成电路+28.2%、机器人+35.4%。软件万亿、智能电网五千亿——新兴产业是工业亮点，全部产业盘靠它。")
heading2(doc, "6.3 消费偏弱")
para(doc, "社零+3.5%、低于目标5%；高档/耐用品/汽电部分亮眼（新能源汽车、金银珠宝），但总量回落与地产、收入预期相关。")
heading2(doc, "6.4 外贸回落、外资下滑")
para(doc, "进出口-1.4%（出口+0.7%）；实际使用外资-53.7%——外部资本/外贸链在调整。")
heading2(doc, "6.5 投资：工业强、地产弱")
para(doc, "固投-4.9%、扣除地产+2.4%；工业投资+12.1%；房地产投资-13.1%、销售-8.4%——地产拖累、工业/基建补位。")
para(doc, "**一句话归因**：南京2025年\u201c**服务(三产)+科教/软件+先进制造/新能源**\u201d为主引擎，\u201c消费/外贸/地产\u201d偏弱——高能级服务型省会、投资/地产调整期。")

# ---- 七、财政 ----
heading1(doc, "七、预算与财政的\u201c含金量\u201d")
para(doc, "**地方一般公共预算收入1620.91亿元、+1.6%**；税收1324.12亿（税占比81.7%）；一般公共预算支出1704.89亿元、与上年持平。")
bullet(doc, "**收入质地好**：税占比81.7%，财政质量居全国城市前列、税收靠现代服务业/软件。")
bullet(doc, "**民生硬度**：居民人均可支配收入78243元、+4.1%；低保标准统一每月1115元；新增城镇就业近22万。")
bullet(doc, "**风险防控**：新增5家上市、金融存款增长15.3%，金融稳健；政府债务可控。")
para(doc, "**财政含义**：南京\u201c量少质高、服务驱动、民生稳\u201d，但增长温和——政策空间有限、需靠创新/科教驱动长期增长。")

# ---- 八、民生 ----
heading1(doc, "八、民生底账：人口、收入与城乡")
para(doc, "南京\u201c十四五\u201d人口稳增：**常住人口963.85万人、比上年增加6.15万人（+0.64%）；城镇化率87.5%（+0.2pct）**。人口自然增长率0.35‰（出生5.39‰、死亡5.04‰），保持净流入+自然增长共存。")
bullet(doc, "**收入**：居民人均可支配收入78243元、+4.1%；城镇86320元、+3.9%，农村40471元、+4.6%（城乡比缩小）。")
bullet(doc, "**就业**：城镇新增就业21.91万人、提供就业岗位57.2万次；就业总体稳定。")
bullet(doc, "**社保**：职工养老/失业/工伤参保分别400.53/344.29/388.16万人；低保3.97万人，低保标准统一每月1115元。")
bullet(doc, "**预期寿命**：户籍人口人均预期寿命首次突破85岁，健康城市领先。")
para(doc, "**民生含义**：南京\u201c收入/就业/康寿\u201d全面领先，人口净流入+城镇化高，是长三角高质量民生的样本。")

# ---- 九、城乡 ----
heading1(doc, "九、城镇与农村：格局与均衡")
para(doc, "**城镇化率87.5%**，南京高度城镇化/都市化；城乡收入比约2.13:1（86320/40471）、居全国较优水平。")
bullet(doc, "**城市**：轨道交通15条、运营里程524.4公里；颐和路历史文化街区入选全国城市更新典型案例；18个更新实验区、138个老旧小区改造。")
bullet(doc, "**农村**：粮食总产100.8万吨、+0.5%；高效粮油、休闲农业；市级以上龙头企业286家；城乡收入比缩小。")
para(doc, "**城乡均衡**：南京\u201c城市极强、县域补齐\u201d，以都市圈/乡村振兴推动城乡融合。")

# ---- 十、人口流 ----
heading1(doc, "十、人口流入与流出")
para(doc, "**南京呈稳定净流入**：2025年常住人口+6.16万、自然增长0.35‰（出生5.39‰、死亡5.04‰），叠加机械增长净流入，在全国人口收缩背景下依然吸人。")
bullet(doc, "**流入**：科教人才+留学归国；人才资源总量415万、+8.9%；每年新增大学生/研发人才。")
bullet(doc, "**竞争**：与上海/杭州/苏州/合肥等长三角强邻争夺高端人才；靠科教、软件、都市圈留人。")
para(doc, "人口方向决定中长期需求与增长；南京的\u201c高收入、科创、都市圈\u201d是其最硬的长逻辑之一。")

# ---- 十一、物价 ----
heading1(doc, "十一、物价与货币环境")
para(doc, "**2025年南京CPI下降0.2%**（2024年+0.4%转负）：食品烟酒-1.1%、交通通信-3.5%、居住-0.5%；衣着+1.9%、其他+10.8%。PPI出厂-2.5%、购进-3.0%。")
bullet(doc, "**物价**：CPI转跌、PPI持续负，反映终端需求偏弱、上游工业品价格压力，与微/内需相符。")
bullet(doc, "**货币/流动性**：本外币存款65688.44亿、+15.3%，贷款63201.98亿、+7.5%；资本市场活跃（新增5家上市）。")
para(doc, "**物价含义**：南京\u201c通缩压力\u201d与全国同步，居民购买力/消费有待稳定。可关注收入与服务的再通胀。")

# ---- 十二、区域一体化 ----
heading1(doc, "十二、区域一体化：南京在\u201c南京都市圈+长三角一体化+宁镇扬\u201d里的位置")
para(doc, "南京处于**长三角城市群+南京都市圈+宁镇扬一体化**核心：既是江苏省会，也是跨省域\u201c国家中心城市\u201d式核心枢纽。")
bullet(doc, "**南京都市圈**：公积金互认互贷、综合交通运输规划、低空应用协同一体化；\u201c苏皖交会\u201d辐射安徽。")
bullet(doc, "**长三角一体化**：与长三角共建大飞机国家先进制造业集群（华东办落户南京）；面向沪杭苏、长三角科创共同体。")
bullet(doc, "**宁镇扬**：落实宁镇扬一体化三年行动；与镇江、扬州协同。")
bullet(doc, "**交通枢纽**：机场旅客吞吐量3137.82万人次、港口集装箱401.33万标箱；地铁524公里。")
para(doc, "**区域含义**：南京作为长三角副中心北方门户、跨省省会，靠\u201c都市圈+大飞机/软件/科创\u201d承东启西、服务长三角与安徽。")

# ---- 十三、五条主线 ----
heading1(doc, "十三、未来5—10年最值得观察的五条主线")
bullet(doc, "**主线1｜软件/数字经济**：软件万亿+智能电网5000亿+信创。能否推动数字经济持续做大做强。")
bullet(doc, "**主线2｜科教产学研**：研发强度4%、紫金山实验室、34家全国重点实验室。能否把科教优势变成产业优势。")
bullet(doc, "**主线3｜智能制造与新能源**：集成电路+28.2%、新能源汽车+99.9%、机器人+35.4%。能否在新能源/半导体链站稳。")
bullet(doc, "**主线4｜都市圈与长三角协同**：南京都市圈+大飞机华东集群+宁镇扬。能否成为长三角西北增长极。")
bullet(doc, "**主线5｜人口与消费**：常住+6万、收入+4.1%、城镇化87.5%。能否把\u201c高收入+人口\u201d转化为\u201c内需/消费\u201d的长期引擎。")
para(doc, "这五条，是南京从\u201c服务+科教省会\u201d走向\u201c数字经济科创强市+都市圈核心\u201d的\u201c主赛道\u201d。")

# ---- 十四、结论 ----
heading1(doc, "十四、最终结论：南京在\u201c软件/智能电网+产学研+省会引领\u201d里的增长逻辑")
para(doc, "南京2025年，本质上是\u201c**软件/智能电网+科教/三产驱动、消费/外贸/地产偏弱**\u201d的答卷：GDP19428.78亿、+5.2%、三产占比68%、软件万亿、智能电网5000亿、规上工业+5.8%、社零+3.5%、进出口-1.4%、财政+1.6%（税占81.7%）。")
para(doc, "只要软件/智能电网、科教产学研、长三角协同持续，南京就站在\u201c数字经济+高端制造\u201d增长极；若外贸/外资、地产承压，南京需承受\u201c服务业强、传统制造与外需弱\u201d的结构挑战。")
para(doc, "最稳观察信号：**一盯软件/数字经济（引擎）、二盯智能电网/高端制造（制造）、三盯科教产学研（创新）、四盯消费/收入（内需）、五盯都市圈/人口（长期）。**南京，是\u201c软件名城+科教/智能电网\u201d的新样本。")

# ---- 附录A ----
heading1(doc, "附录A：主要资料来源与核验路径")
bullet(doc, "南京市2025年政府工作报告（2025年1月）——目标来源。")
bullet(doc, "《2025年南京市国民经济和社会发展统计公报》（市统计局）——GDP、工业、外贸、人口实值。")
bullet(doc, "2026年南京市政府工作报告（2026年1月）——2025复盘/软件/智能电网/长三角。")
bullet(doc, "金陵海关（外贸、口岸）、市财政局。")
heading2(doc, "核验说明")
para(doc, "本报告涉及数据以统计公报/官方口径为准；\u201c软件/智能电网/科教产学研/都市圈\u201d等以官方口径为准。")

# ---- 附录B ----
heading1(doc, "附录B：建议建立的年度跟踪仪表盘")
para(doc, "建议把下面10个\u201c测脉搏\u201d指标做成年度跟踪表：")
dash = [
    ["#", "指标", "2025基值", "为什么重要"],
    ["1", "GDP增速", "+5.2%（19428.78亿）", "总量与方向"],
    ["2", "规模以上工业增速", "+5.8%", "制造底盘"],
    ["3", "软件/智能电网产值", "超万亿/超5000亿", "产业新边界"],
    ["4", "三产占比/增速", "68% / +6.0%", "服务业底盘"],
    ["5", "进出口增速", "-1.4%（5380.9亿）", "外贸韧性"],
    ["6", "固定资产投资/工业", "-4.9% / 工业+12.1%", "投资结构"],
    ["7", "社零增速", "+3.5%（8135.85亿）", "内需消费"],
    ["8", "常住人口/城镇化", "963.85万 / 87.5%", "人口与城市"],
    ["9", "地方财政收入/税占", "+1.6%(1620.91亿) / 81.7%", "财政质量"],
    ["10", "CPI", "-0.2%", "物价"],
]
table(doc, dash[0], dash[1:], widths=[0.8, 4.6, 4.0, 3.8])
para(doc, "把这10个指标连起来看，软件/智能电网/科教（3/5）、消费与外贸（5/7）、人口（8），都说明南京在真正换挡。")

# ============ 保存 ============
out = "/Users/x/Desktop/content-prod-lab/reports/南京市_2025年政府工作报告_深度研究_2026-08-13.docx"
doc.save(out)
print("SAVED:", out)
print("PARAGRAPHS:", len(doc.paragraphs))
print("TABLES:", len(doc.tables))
