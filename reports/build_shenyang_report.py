# -*- coding: utf-8 -*-
"""Build 沈阳市2025年政府工作报告 深度研究 DOCX, 参照省会系列版式。"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DARK = RGBColor(0x1F, 0x2A, 0x44)
RED = RGBColor(0xB0, 0x1E, 0x2E)
GRAY = RGBColor(0x59, 0x60, 0x69)
BLUE = RGBColor(0x1F, 0x4E, 0x79)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
HEADER_FILL = "1F2A44"
K = "微软雅黑"


def set_run(run, size=11, bold=False, color=DARK, font=K):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    run.font.name = font
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), font)
    rFonts.set(qn("w:hAnsi"), font)
    rFonts.set(qn("w:eastAsia"), font)


def para(doc, text, size=11, bold=False, color=DARK, align=None,
         space_after=6, space_before=0, font=K):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        set_run(r, size=size, bold=bold or (i % 2 == 1), color=color, font=font)
    return p


def heading1(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(16)
    pf.space_after = Pt(8)
    r = p.add_run(text)
    set_run(r, size=16, bold=True, color=RED)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "B01E2E")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def heading2(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(10)
    pf.space_after = Pt(4)
    r = p.add_run(text)
    set_run(r, size=13, bold=True, color=BLUE)
    return p


def table(doc, headers, rows, widths=None, font_size=9.5):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_run(r, size=font_size, bold=True, color=WHITE)
        tcPr = hdr[i]._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), HEADER_FILL)
        tcPr.append(shd)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(val))
            set_run(r, size=font_size, color=DARK)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    return t


def quote_box(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(8)
    pf.left_indent = Pt(12)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), "B01E2E")
    pBdr.append(left)
    pPr.append(pBdr)
    r = p.add_run(text)
    set_run(r, size=11, color=DARK, font="楷体")
    return p


def bullet(doc, text, size=10.5):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.7)
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        set_run(r, size=size, bold=(i % 2 == 1), color=DARK)
    return p


# ================================================================= 文档主体
doc = Document()
sec = doc.sections[0]
sec.page_width = Cm(21)
sec.page_height = Cm(29.7)
sec.top_margin = Cm(2.4)
sec.bottom_margin = Cm(2.2)
sec.left_margin = Cm(2.5)
sec.right_margin = Cm(2.5)

st = doc.styles["Normal"]
st.font.name = K
st.font.size = Pt(11)
st.element.rPr.rFonts.set(qn("w:eastAsia"), K)


# ---- 封面 ----
para(doc, "沈阳市2025年政府工作报告", size=24, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=60, space_after=4)
para(doc, "深度研究与\u201c容易被忽视的细节\u201d分析", size=16, bold=True, color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
para(doc, "从\u201c老工业基地转型、装备制造、工业机器人、东北振兴与汽车\u201d重新理解沈阳", size=12, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)
para(doc, "\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501", color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
para(doc, "研究对象：2025年沈阳市政府工作报告及同期财政、国民经济和社会发展统计资料、2026年官方复盘", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
para(doc, "标定日期：2026-08-13", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
doc.add_page_break()

# ---- 目录 ----
catalog = [
    "执行摘要",
    "一、研究口径与阅读方法",
    "二、先看沈阳的特殊底盘：老工业基地、装备制造、工业机器人、汽车与东北振兴",
    "三、最关键的宏观错位：GDP破9000亿但增速骤降至2%，工业/装备/投资全线承压",
    "四、容易被忽视、但信号很重的细节（15条）",
    "五、2025年GDP目标 vs 实际：对照表",
    "六、2025年增长是谁撑起来的？（结构归因）",
    "七、预算与财政的\u201c含金量\u201d",
    "八、民生底账：人口、收入与城乡",
    "九、城镇与农村：格局与均衡",
    "十、人口流入与流出",
    "十一、物价与货币环境",
    "十二、区域一体化：沈阳在\u201c沈阳都市圈+东北振兴+面向东北亚\u201d里的位置",
    "十三、未来5—10年最值得观察的五条主线",
    "十四、最终结论：沈阳在\u201c装备+机器人+汽车+科技\u201d里的增长逻辑",
    "附录A：主要资料来源与核验路径",
    "附录B：建议建立的年度跟踪仪表盘",
]
para(doc, "目　录", size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=10, space_after=10)
for item in catalog:
    para(doc, item, size=12, space_after=4)
doc.add_page_break()

# ---- 执行摘要 ----
heading1(doc, "执行摘要")
para(doc, "**核心判断**　2025年沈阳最显著的是\u201cGDP 9100.3亿元、增长2.0%（远低于5.5%目标）、三产占63.1%\u201d、\u201c规上工业-6.0%、装备制造-7.4%\u201d、\u201c固投-24.9%、财政-3.8%\u201d、\u201c常住人口927.6万/城镇化85.87%\u201d。这说明沈阳在\u201c顶住汽车/地产压力\u201d下艰难增长，**工业、投资、财政全线负增长**是标志性困难年。")
para(doc, "把2025年目标（GDP+5.5%/规上+6%/固投+6%）、2025年统计、2026年前瞻一起看，沈阳呈现\u201c东北振兴+装备制造\u201d路径：**工业机器人（5170套+4.2%）、变压器（+57.8%）、锂电（+57.5%）、出口（+8.8%）**是少数据点，但汽车-10.3%、装备-7.4%拖累。总量9100亿居东北首位。")
para(doc, "最容易记住的一句话：**沈阳是\u201c东北老工业基地+装备/机器人之城\u201d的省会（辽宁省），靠\u201c装备、工业机器人、汽车、科技（沈阳材料科学国家研究中心）\u201d转型。**2025年当地经济逆风、增速仅+2%，是\u201c新旧动能换挡阵痛\u201d的一年。观察沈阳，与其只看\u201cGDP 9100亿\u201d，不如看\u201c工业机器人+4.2%、变压器+57.8%、锂电+57.5%、出口+8.8%、战略性新兴占33.5%\u201d。")# ---- 一、研究口径与阅读方法 ----
heading1(doc, "一、研究口径与阅读方法")
para(doc, "本报告以\u201c沈阳市2025年政府工作报告（2025年）\u201d为起点，把\u201c2025年GDP目标（5.5%以上）\u201d与\u201c2025年国民经济和社会发展统计公报实际值（9100.3亿元/+2.0%）\u201d并置对照，再用2026年政府工作报告的复盘作事后验证。")
para(doc, "重点阅读三类证据：**一是指标目标是否达标**；**二是结构（谁贡献增长）是否匹配目标叙事**；**三是官方复盘如何自评、承认问题**。统计口径一般公共预算收入为地方级。")
para(doc, "统一采用\u201c同比增长%\u201d（GDP为不变价），财政与居民收入多为名义值，文中按原口径转录、不逐项换算。")

# ---- 二、特殊底盘 ----
heading1(doc, "二、先看沈阳的特殊底盘：老工业基地、装备制造、工业机器人、汽车与东北振兴")
para(doc, "**区位与身份**：沈阳是辽宁省会、副省级城市，\u201c共和国装备部\u201d\u201c工业机器人之城\u201d，国家全面振兴东北战略核心、沈阳都市圈获批；沈阳材料科学国家研究中心、沈阳自贸区。")
para(doc, "**产业底盘**：一是装备制造（占规上68%、变压器/矿山装备/气体压缩机）；二是工业机器人（产量5170套/+4.2%）；三是汽车（宝马/华晨、汽车制造占大、2025-10.3%）；四是科技/新兴（战略性新兴占33.5%、锂电+57.5%、数据标注、AI算力）；五是航空/轨道装备。")
para(doc, "**人口底盘**：2025年末常住人口927.6万/+3.3万、城镇化率85.87%；是东北人口大市。")
para(doc, "**市场与出口**：2025年社零4195.4亿/+1.0%；进出口1353.7亿/-7.6%（出口+8.8%）。")

# ---- 三、核心宏观错位 ----
heading1(doc, "三、最关键的宏观错位：GDP破9000亿但增速骤降至2%，工业/装备/投资全线承压")
para(doc, "**第一组错位**：2025年GDP目标5.5%以上，实际9100.3亿/+2.0%**巨大落差**（低3.5pct），增速明显低于2024的5.2%；\u201c十四五\u201d连跨7000/8000/9000亿但尾部失速。")
para(doc, "**第二组错位**：规上工业-6.0%、装备-7.4%、汽车-10.3%、固投-24.9%全线走弱；但**出口+8.8%、洗衣机/变压器/锂电增长、战略性新兴占33.5%**，体现结构性分化。")
para(doc, "**第三组错位**：常住927.6万/城镇化85.87%但**自然-4.58\u2030**（出生4.40\u2030<死亡8.98\u2030）严重负自然增长；财政-3.8%、支出-11.3%双降。")
para(doc, "一句话：**沈阳是\u201c装备/机器人强、但汽车/投资/财政全面承压\u201d的东北工业省会**——2025年是\u201c被动换挡、逆风\u201d的一年。")

# ---- 四、15条细节 ----
heading1(doc, "四、容易被忽视、但信号很重的细节（15条）")
bullet(doc, "1. **GDP 9100.3亿/+2.0%**：远低于5.5%目标、东北首个破9000亿但失速。")
bullet(doc, "2. **规上工业-6.0%**：制造业-5.9%，工业压力大。")
bullet(doc, "3. **装备制造-7.4%**：占规上68%的支柱负增长。")
bullet(doc, "4. **汽车制造业-10.3%**：宝马/汽车资源调整拖累。")
bullet(doc, "5. **工业机器人5170套/+4.2%**：\u201c机器人之城\u201d亮点。")
bullet(doc, "6. **变压器+57.8%、锂电+57.5%**：新动能/绿电装备。")
bullet(doc, "7. **固投-24.9%**：制造业-18.2%塌陷，是最大隐患。")
bullet(doc, "8. **出口+8.8%、入境游客+43.5%**：外向/文旅亮点。")
bullet(doc, "9. **社零+1.0%**：消费微增。")
bullet(doc, "10. **常住927.6万/城镇化85.87%**：东北最大人口之一、净流入3.3万。")
bullet(doc, "11. **收入：城镇58297元/+3.9%、农村27196元/+5.5%**，城乡比2.14。")
bullet(doc, "12. **CPI+0.6%**：温和、比多数城市高。")
bullet(doc, "13. **财政794.2亿/-3.8%、支出-11.3%**：财政收缩。")
bullet(doc, "14. **战略性新兴占规上33.5%、技改+12.2%**：转型实质推进。")
bullet(doc, "15. **粮食420万吨/+4.9%、农业稳**：粮仓底盘。")# ---- 五、目标 vs 实际对照表 ----
heading1(doc, "五、2025年GDP目标 vs 实际：对照表")
table(doc,
    ["指标", "2025年目标", "2025年实际", "是否达标"],
    [
        ["地区生产总值增速", "5.5%以上", "9100.3亿元/+2.0%", "严重未达标"],
        ["规上工业增加值", "（与GDP协调）", "-6.0%", "负增长"],
        ["固定资产投资", "+6%", "-24.9%（转负）", "严重未达标"],
        ["社会消费品零售总额", "（促消费）", "4195.4亿元/+1.0%", "温和"],
        ["一般公共预算收入", "（3%左右）", "794.2亿元/-3.8%", "负增长未达标"],
        ["城镇/农村人均可支配收入", "与GDP同步", "+3.9%/+5.5%", "高于GDP"],
        ["出口总额", "（设目标）", "+8.8%", "亮点"],
    ],
    widths=[3.4, 2.5, 5.0, 3.1])
para(doc, "**简评**：2025年沈阳**几乎所有硬指标未达目标**，GDP骤降至+2.0%、工业/投资/财政负增长。仅**出口、机器人、变压器、新能储、居民收入**为亮点，是\u201c新一轮换挡阵痛\u201d之年。")

# ---- 六、结构归因 ----
heading1(doc, "六、2025年增长是谁撑起来的？（结构归因）")
para(doc, "**三大产业**：一产376.1亿/+3.6%、二产2981.0亿/-4.5%、三产5743.2亿/+5.0%；三产占63.1%。增长几乎全靠三产（+5%），二产（工业/建筑）拖累为负。")
para(doc, "**工业**：规上工业-6.0%；**电气机械+15.4%、铁路船舶航空航天+3.1%、电子+2.0%、农副食品+3.1%**；**汽车-10.3%、专用设备-10.5%、通用设备-8.1%、医药-5.2%**拖累。装备-7.4%占规上68%。工业机器人+4.2%、变压器+57.8%、锂电+57.5%。")
para(doc, "**服务业**：三产+5%（批发零售/金融/文旅/科技服务），旅游/赛事/消费支撑。")
para(doc, "**增长归因**：沈阳GDP增长主要靠**三产（消费/金融/文旅/科技）+出口+部分新兴制造**；汽车、装备、固投、财政为负。")

# ---- 七、财政 ----
heading1(doc, "七、预算与财政的\u201c含金量\u201d")
para(doc, "2025年一般公共预算收入794.2亿元/-3.8%，支出1031.9亿/-11.3%（卫生健康+10.7%、教育+0.1%、社保-3.2%、科技支出-9.3%、农林水-19.4%）。")
para(doc, "**结构性**：收入、支出**双下降**，反映\u201c经济下行减收+化债/压缩支出\u201d；民生（卫生+10.7%）仍有增。")
para(doc, "**含金量**：沈阳财政\u201c收缩型\u201d，2025年是东北老工业基地财政压力的缩影，支出让位于化债、民生（卫生）优先。")

# ---------------- 八、民生 ----------------
heading1(doc, "八、民生底账：人口、收入与城乡")
para(doc, "常住人口927.6万/+3.3万、城镇化85.87%；出生4.40\u2030/死亡8.98\u2030，**自然-4.58\u2030**。收入：**城镇58297元/+3.9%、农村27196元/+5.5%**，城乡比2.14（缩小0.04）。")
para(doc, "就业：城镇新增就业15万；消费支出：城镇44505元、农村19060元。")
para(doc, "**民生结论**：收入农村快于城镇、差距2.14偏小（东北相对均衡）；人口自然严重负增长、靠机械流入；就业稳兜底。")

# ---------------- 九 ----------------
heading1(doc, "九、城镇与农村：格局与均衡")
para(doc, "沈阳城镇化率85.87%（很高），中心城区（和平/沈河/铁西/浑南）承载装备/汽车/金融；县域（辽中/法库/康平）发展农业与加工。")
para(doc, "城乡收入比2.14（东北较小、较均衡），农村收入+5.5%快于城镇、差距缩小；粮仓底盘支撑。")

# =========== 十 ============
heading1(doc, "十、人口流入与流出")
para(doc, "沈阳常住927.6万、净流入3.3万，但自然-4.58\u2030，增长靠机械流入（高校：东北大学/辽宁大学+装备/汽车+新市民）。")
para(doc, "**流入**：大学生留沈、东北人口回流（省内）；**流出**：青年外迁华北/东部。整体\u201c微弱净流入、结构老化\u201d。")

# =========== 十一 ============
heading1(doc, "十一、物价与货币环境")
para(doc, "2025年沈阳CPI**上涨0.6%**（温和），食品烟酒+0.5%；较多数城市偏高一点。")
para(doc, "金融：存款24635.3亿、贷款（公积金/房贷等）慢增。\u201c温和通胀、流动性稳\u201d。")

# =========== 十二 ============
heading1(doc, "十二、区域一体化：沈阳在\u201c沈阳都市圈+东北振兴+面向东北亚\u201d里的位置")
para(doc, "沈阳是沈阳都市圈核心、辽宁省乃至东北政治经济中心，国家全面振兴东北承载地；与\u201c长春都市圈\u201d\u201c大连\u201d联动，面向俄远东/东北亚（沈阳-海参崴/莫斯科航线）。")
para(doc, "产业协同：装备（沈鼓/沈变）、机器人（新松）、汽车（宝马）、航空/轨道；自贸区+综保区、桃仙机场。沈阳是\u201c东北振兴头部\u201d城市。")

# ============ 十三 =============
heading1(doc, "十三、未来5—10年最值得观察的五条主线")
bullet(doc, "1. **装备崛起与工业机器人**：变压器/机器人/重大装备突破，\u201c新装备\u201d。")
bullet(doc, "2. **汽车（新能源/宝马）**：新能源汽车与燃油车切换，2025-10.3%能否修复。")
bullet(doc, "3. **科技/新质生产力**：材料国家中心、AI算力1830P、低空/具身。")
bullet(doc, "4. **消费/文旅**：游客2.3亿、演唱会/赛事、入境+43.5%。")
bullet(doc, "5. **投资与财政修复**：固投-24.9%、财政-3.8%后，靠项目/央地合作修复。")

# ============ 十四 =============
heading1(doc, "十四、最终结论：沈阳在\u201c装备+机器人+汽车+科技\u201d里的增长逻辑")
para(doc, "**结论**：沈阳2025年的\u201c真相\u201d是——**GDP+2.0%（远低于目标）、规上-6.0%、装备-7.4%、固投-24.9%、财政-3.8%**。它是\u201c老工业基地换挡阵痛\u201d的东北省会，靠**机器人/变压器/锂电/出口/科技**对冲下行。")
para(doc, "**对趋势判断**：装备+机器人+汽车+科技代表沈阳\u201c长逻辑\u201d，2025年是\u201c旧动能（汽车/传统装备）失速、新动能（新能源/机器人）尚小\u201d的青黄不接。**工业转型+投资修复**决定未来，**出口/消费/人口流失**是长期压力。")
para(doc, "**若只看一个指标**：看**规上工业中的工业机器人产量增速（+4.2%）与汽车制造增速（-10.3%）**——沈阳处在\u201c旧汽车换挡、新装备爬坡\u201d的关键期，装备/机器人能否接替汽车决定了沈阳未来的增长天花板。")

# ------------- 附录A -------------
heading1(doc, "附录A：主要资料来源与核验路径")
bullet(doc, "沈阳市人民政府《2025年沈阳市政府工作报告》（2025年）。")
bullet(doc, "沈阳市统计局《2025年沈阳市国民经济和社会发展统计公报》（2026年5月）。")
bullet(doc, "《2026年沈阳市政府工作报告》（2026年1月）及中国县域摘要。")
bullet(doc, "辽宁省2025年统计公报、沈阳统计年鉴交叉核验。")

# ------------- 附录B -------------
heading1(doc, "附录B：建议建立的年度跟踪仪表盘")
bullet(doc, "GDP总量/增速 vs 全年目标（+/-百分点）。")
bullet(doc, "规上工业增加值及分行业（装备/汽车/机器人/电气/锂电）增速。")
bullet(doc, "固定资产投资（总量/制造业/基建/民间/房地产）增速。")
bullet(doc, "社会消费品零售总额、游客/入境。")
bullet(doc, "进出口（人民币）、出口、入境游客。")
bullet(doc, "一般公共预算收入/税收/非税、财政支出、民生占比。")
bullet(doc, "常住人口、自然增长率、城镇化率、高校留沈。")
bullet(doc, "CPI/核心CPI、规上工业利润与营收。")
bullet(doc, "工业机器人产量、变压器/锂电、战略性新兴占比。")
bullet(doc, "沈阳（宝马）汽车产量、技改投资。")

# ------------- 保存 -------------
out = "/Users/x/Desktop/content-prod-lab/reports/沈阳市_2025年政府工作报告_深度研究_2026-08-13.docx"
doc.save(out)
print("SAVED OK 沈阳市_2025年政府工作报告_深度研究_2026-08-13.docx")