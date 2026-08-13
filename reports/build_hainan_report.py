# -*- coding: utf-8 -*-
"""Build 海南省2025年政府工作报告 深度研究 DOCX, 参照省市系列版式。"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DARK = RGBColor(0x1F, 0x2A, 0x44)
RED = RGBColor(0xB0, 0x1E, 0x2E)
GRAY = RGBColor(0x59, 0x60, 0x69)
BLUE = RGBColor(0x1F, 0x4E, 0x79)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
HEADER_FILL = "1F2A44"
K = "微软雅黑"


def set_run(run, size=11, bold=False, color=DARK, font=K):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    run.font.name = font
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), font)
    rFonts.set(qn("w:hAnsi"), font)
    rFonts.set(qn("w:eastAsia"), font)


def para(doc, text, size=11, bold=False, color=DARK, align=None,
         space_after=6, space_before=0, font=K):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        set_run(r, size=size, bold=bold or (i % 2 == 1), color=color, font=font)
    return p


def heading1(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(16)
    pf.space_after = Pt(8)
    r = p.add_run(text)
    set_run(r, size=16, bold=True, color=RED)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "B01E2E")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def heading2(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(10)
    pf.space_after = Pt(4)
    r = p.add_run(text)
    set_run(r, size=13, bold=True, color=BLUE)
    return p


def table(doc, headers, rows, widths=None, font_size=9.5):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_run(r, size=font_size, bold=True, color=WHITE)
        tcPr = hdr[i]._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), HEADER_FILL)
        tcPr.append(shd)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(val))
            set_run(r, size=font_size, color=DARK)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    return t


def quote_box(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(8)
    pf.left_indent = Pt(12)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), "B01E2E")
    pBdr.append(left)
    pPr.append(pBdr)
    r = p.add_run(text)
    set_run(r, size=11, color=DARK, font="楷体")
    return p


def bullet(doc, text, size=10.5):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.7)
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        set_run(r, size=size, bold=(i % 2 == 1), color=DARK)
    return p


# ================================================================= 文档主体
doc = Document()
sec = doc.sections[0]
sec.page_width = Cm(21)
sec.page_height = Cm(29.7)
sec.top_margin = Cm(2.4)
sec.bottom_margin = Cm(2.2)
sec.left_margin = Cm(2.5)
sec.right_margin = Cm(2.5)

st = doc.styles["Normal"]
st.font.name = K
st.font.size = Pt(11)
st.element.rPr.rFonts.set(qn("w:eastAsia"), K)



# ---- 封面 ----
para(doc, "海南省2025年政府工作报告", size=24, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=60, space_after=4)
para(doc, "深度研究与“容易被忽视的细节”分析", size=16, bold=True, color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
para(doc, "从“自贸港封关运作、免税消费、热带农业、海洋经济与旅游升级”重新理解海南", size=12, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)
para(doc, "━━━━━━━━━━━━━━━━", color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
para(doc, "研究对象：2025年海南省政府工作报告及同期财政、国民经济和社会发展统计资料、2026年官方复盘", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
para(doc, "标定日期：2026-08-13", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
doc.add_page_break()

# ---- 目录 ----
catalog = [
    "执行摘要",
    "一、研究口径与阅读方法",
    "二、先看海南的特殊底盘：自贸港封关、免税消费、热带农业、海洋经济与旅游",    
    "三、最关键的宏观错位：GDP破8000亿、消费/工业强，但投资/地产弱",
    "四、容易被忽视、但信号很重的细节（15条）",
    "五、2025年GDP目标 vs 实际：对照表",
    "六、2025年增长是谁撑起来的？（结构归因）",
    "七、预算与财政的“含金量”",
    "八、民生底账：人口、收入与城乡",
    "九、城镇与农村：格局与均衡",
    "十、人口流入与流出",
    "十一、物价与货币环境",
    "十二、区域一体化：海南在“自贸港+封关+南海/海洋”里的位置",
    "十三、未来5—10年最值得观察的五条主线",
    "十四、最终结论：海南在“自贸港封关+免税+海洋+旅游”里的增长逻辑",
    "附录A：主要资料来源与核验路径",
    "附录B：建议建立的年度跟踪仪表盘",
]
para(doc, "目　录", size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=10, space_after=10)
for item in catalog:
    para(doc, item, size=12, space_after=4)
doc.add_page_break()

# ---- 执行摘要 ----
heading1(doc, "执行摘要")
para(doc, "**核心判断**　如果只看新闻标题，2025年海南最显眼的是“GDP破8000亿、增长4.0%”、“规上工业+9.9%（全国第2）”、“游客1.06亿人次（+9.1%）”和“封关运作启动（12月18日）”。但这份研究真正值得深读的，是这座“自贸港封关+免税+海洋+热带农业+旅游”的海岛省份，如何在固定资产投资（-16.5%）、房地产开发投资（-23.0%）与CPI（-0.1%）偏弱的背景下，靠“规上工业+免税消费+旅游/海洋+跨境”稳住增长。")
para(doc, "把2025年初设定的目标（GDP增长6%以上）、2025年《国民经济和社会发展统计公报》、以及2026年报告对2025年的复盘放在一起看，海南呈现清晰暗线：**从“房地产依赖”的旧底盘，向“自贸港封关+免税消费+海洋经济+旅游业/高新技术”转型**。旧引擎（房地产、一般基建）在调整；新引擎（规上工业、免税、旅游/海洋、深海科技、种业）被要求更快补位。这也是“全国唯一全域自贸港/全岛封关运作年”的特殊样本。")
para(doc, "因此，本报告不按政府工作报告逐段复述，而采用“显性表述—同期数据—制度含义—长期影响”的方式，专门提取容易被忽视、但对判断海南未来5—10年发展模式更有价值的信号。")
para(doc, "最容易记住的一句话：**海南是“全岛自贸港+封关运作+免税消费+热带农业/海洋经济/旅游”的全国唯一样本，靠“免税+工业+旅游+海洋”撑起增长。**观察海南，与其看“GDP 8108.85亿”，不如看“封关运作、免税/零关税、海洋经济36%、热带农业、旅游1.06亿人次”这几张名片。")
heading2(doc, "一页速览：2025年海南经济的“表与里”")
table(doc,
    ["维度", "表面现象", "更深层信号"],
    [
        ["增长", "GDP 8108.85亿、+4.0%", "一产20.3%、二产18.0%、三产61.7%"],
        ["产业", "规上工业+9.9%", "高技术制造+28.1%、单体17.78万就业"],
        ["外贸", "货物进出口2760.03亿、-0.7%", "出口-10.6%、服务进出口+22.1%"],
        ["投资", "固定资产投资-16.5%", "地产-23.0%、二产-30.2%"],
        ["财政", "地方一般公共预算收入903.74亿、+1.5%", "税收+4.5%"],
        ["消费", "社零2673.58亿、+5.0%", "乡村+8.7%、最高增幅"],
        ["人口", "常住1055万、+7万、城镇化64.08%", "三产61.7%、都市圈"],
        ["开放", "封关运作(12月18日启动)", "自贸港/零关税/免税"],
    ],
    widths=[2.2, 5.4, 8.4])
para(doc, "", size=4, space_after=2)

# ---- 一、研究口径与阅读方法 ----
heading1(doc, "一、研究口径与阅读方法")
heading2(doc, "1.1 本报告的三份核心底稿")
bullet(doc, "**2025年《政府工作报告》**：2025年1月在省七届人大三次会议上作，给出全年预期目标（GDP增长6%以上、地方一般公共预算收入增长3%左右、固投8%以上、社零10%以上）。")
bullet(doc, "**2025年《国民经济和社会发展统计公报》**：海南省统计局2026年2月发布，给出全年实际执行数与增速，是“事后验证”的锚点。")
bullet(doc, "**2026年海南省政府工作报告及官方复盘**：对2025执行结果的追认，交叉验证“目标—执行—再评价”三段式。")
heading2(doc, "1.2 阅读方法：显性—数据—制度—长期")
para(doc, "每一章按“**显性表述 → 同期数据 → 制度含义 → 长期影响**”四层展开。其中“制度含义”回答“政府为什么这么排优先序”，“长期影响”回答“这对未来5—10年意味着什么”。")
para(doc, "关键判别：**当报告使用的“增长词”和统计公报里的实际数不一致时，数据优先。**例如2025年GDP目标6%以上、实际4.0%未达标；固投目标8%以上、实际-16.5%、大幅未达标；社零目标10%以上、实际+5.0%、未达标；财政收入目标3%左右、实际+1.5%、未达标。差异反映：海南“消费/工业强、投资/地产弱”。")
heading2(doc, "1.3 指标取舍与单位说明")
para(doc, "本报告统一以“2025年统计公报+2026年报告复盘”为口径，增速均为可比价，绝对数为人民币。GDP 8108.85亿元为全省初步核算数。")

# ---- 二、底盘 ----
heading1(doc, "二、先看海南的特殊底盘：自贸港封关、免税消费、热带农业、海洋经济与旅游")
para(doc, "海南的地盘，取决于它作为“**全国唯一全域自贸港+全岛封关运作+热带岛屿**”的特殊定位。它不像内地靠投资/工业总量，而是“**离岛自贸+免税消费+热带农业+海洋经济+旅游**”的组合。")
bullet(doc, "**自贸港/封关**：2025年12月18日全岛封关运作正式启动，标志自贸港从政策试点转向制度安排。零关税税目比例从21%增至74%、扩大至6637个税目；离岛免税商品扩至47大类；封关后一个月新增经营主体2.68万户、离岛免税销售+46.8%。")
bullet(doc, "**免税消费**：2025年社零+5.0%、增速全国第6；“机票当钱花”等促销超70场、带动离岛免税消费近40亿元；封关后离岛免税销售显著放量。")
bullet(doc, "**热带农业/海洋**：海洋生产总值占GDP比重提升至36%以上；农林牧渔总产值2631.86亿（+4.7%）；深水网箱养殖+7.2%、工厂化养殖+21.1%；热带特色高效农业增加值占农业约65%。")
bullet(doc, "**旅游**：2025年接待游客1.06亿人次（+9.1%）、游客总花费2254.32亿（+10.5%）、入境游客150.05万人次（+35.2%）；医疗旅游86.5万人次（+1.1倍）。")
para(doc, "**制度含义**：海南不追求“大而全”，而是把“自贸港制度红利、免税/零关税、海洋、热带农业、旅游”当核心资产。这既是海南最大特色，也是它面对“投资/地产弱”时的最独特筹码。")

# ---- 三、宏观错位 ----
heading1(doc, "三、最关键的宏观错位：GDP破8000亿、消费/工业强，但投资/地产弱")
para(doc, "2025年海南GDP 8108.85亿元、+4.0%（一产+4.4%、二产+1.0%、三产+4.6%）。表面看“温和增长”，但拆开看是“**消费/工业强、投资/地产弱**”的错位：")
para(doc, "**强的部分**：规上工业+9.9%（重工业+15.4%、石油天然气开采+48.7%、高技术制造+28.1%、燃料加工+24.7%）；社零2673.58亿、+5.0%（乡村+8.7%、汽车+33.0%、通讯+56.6%）；服务进出口+22.1%；旅游1.06亿人次（+9.1%）、总花费+10.5%。")
para(doc, "**弱的部分**：固定资产投资-16.5%（二产-30.2%、三产-13.0%）；房地产开发投资-23.0%、新开工面积-38.0%；货物进出口-0.7%（出口-10.7%）；CPI-0.1%、PPI-4.5%。")
para(doc, "**核心错位一句话**：海南“消费/工业/免税/旅游强（社零+5.0%、规上工业+9.9%），但投资/地产/出口弱”。2026年封关红利释放后，若“免税+海洋+跨境”继续放量、投资须修复，增长有望从“消费单极”走向“消费+投资/工业”双轮。")
table(doc,
    ["强引擎", "数据", "弱项", "数据"],
    [
        ["社会消费品零售", "+5.0%", "固定资产投资", "-16.5%"],
        ["规上工业", "+9.9%", "房地产开发投资", "-23.0%"],
        ["高技术制造业", "+28.1%", "二产投资", "-30.2%"],
        ["旅游总花费", "+10.5%", "货物进出口", "-0.7%"],
        ["服务进出口", "+22.1%", "CPI", "-0.1%"],
    ],
    widths=[3.6, 3.2, 3.6, 3.2])
para(doc, "**错位结论**　海南的增长“很特殊、也很矛盾”。强的部分（消费/工业/免税/旅游/海洋）与弱的部分（投资/地产/出口/物价）并存。2026年“稳投资/地产+续免税/海洋/旅游”是重点。")

# ---- 四、被忽视细节 ----
heading1(doc, "四、容易被忽视、但信号很重的细节（15条）")
details = [
    ("1", "全岛封关运作2025-12-18启动", "自贸港/制度安排。"),
    ("2", "零关税税目比例21%→74%、6637个税目", "零关税/享惠扩围。"),
    ("3", "离岛免税商品扩至47大类、封关后+46.8%", "免税消费强。"),
    ("4", "规上工业+9.9%（全国第2）、高技术+28.1%", "制造/技改强。"),
    ("5", "石油天然气开采+48.7%、燃料加工+24.7%", "能源/石化。"),
    ("6", "四大主导产业占GDP 65.9%（+2.2pct）", "产业集中。"),
    ("7", "社零+5.0%（居全国第6）、乡村+8.7%", "内需/下沉强。"),
    ("8", "旅游1.06亿人次+9.1%、总花费+10.5%", "旅游强、入境+35.2%。"),
    ("9", "医疗旅游86.5万人次、+1.1倍", "乐城/医疗旅游。"),
    ("10", "海洋生产总值占GDP超36%", "海洋经济。"),
    ("11", "固定资产-16.5%、地产-23.0%", "投资/地产调整。"),
    ("12", "服务进出口+22.1%、实际使用外资+19.9%", "自贸港/跨境强。"),
    ("13", "常住1055万、+7万、城镇化率64.08%（+1.00pct）", "人口流入、城镇化。"),
    ("14", "居民收入36306元、+4.2%、农村+5.2%", "民生收入回暖。"),
    ("15", "CPI-0.1%、PPI-4.5%", "物价/通缩压力。"),
]
for d in details:
    para(doc, "**%s**" % d[0], size=11)
    para(doc, d[1], size=10.5, color=GRAY)

    para(doc, "**备注**　这条“错位”在海南尤其鲜明：增长靠消费/工业/免税/旅游/海洋，但投资/地产弱。2026年若投资/地产修复、封关红利持续释放，增长可能从“消费单极”走向“消费+投资/工业”多极。这条细节，正是海南2026年最难也最值得盯的“变量”。", size=10.5)

# ---- 五、2025年目标 vs 实际 对照 ----
heading1(doc, "五、2025年GDP目标 vs 实际：对照表")
t5 = [
    ["指标", "2025目标", "2025实际", "达标判定"],
    ["GDP增速", "6%以上", "+4.0%", "未达标"],
    ["地方一般公共预算收入增速", "3%左右", "+1.5%", "未达标"],
    ["固定资产投资增速", "8%以上", "-16.5%", "大幅未达标"],
    ["社会消费品零售增速", "10%以上", "+5.0%", "未达标"],
    ["城镇/农村居民收入增速", "6.5%左右/7.5%左右", "+3.4%/+5.2%", "未达"],
    ["居民消费价格(CPI)", "涨幅2%左右", "-0.1%", "远低于"],
    ["城镇调查失业率", "控制在5.5%左右", "5.2%", "达标"],
]
table(doc, t5[0], t5[1:], widths=[3.4, 3.0, 3.0, 3.8])
para(doc, "**对照结论**　目标“偏进攻”：GDP/固投/社零都定得不低（尤其社零10%+），但**投资/社零/GDP是最大失分项**（固投-16.5%、社零+5.0%未达10%、GDP+4.0%低于6%）。规上工业、就业（城镇调查失业率5.2%）接住了，靠“消费/工业/免税/旅游”撑起增长。")

# ---- 六、增速分项支撑 ----
heading1(doc, "六、2025年增长是谁撑起来的？（结构归因）")
para(doc, "GDP+4.0%背后，是“**三产/消费/免税/旅游强、投资/地产弱**”的结构。把增长贡献拆开看：")
g6 = [
    ["引擎", "贡献/状态", "说明"],
    ["消费(社零)", "+5.0%", "社零2673.58亿、乡村+8.7%"],
    ["规上工业", "+9.9%", "高技术+28.1%、石化/能源"],
    ["旅游/文旅", "+10.5%(花费)", "游客1.06亿人次、入境+35.2%"],
    ["免税/零关税", "封关后+46.8%", "离岛免税/零关税扩围"],
    ["海洋经济", "占GDP超36%", "海洋/深海科技/南繁"],
    ["服务进出口", "+22.1%", "自贸港跨境"],
    ["房地产投资", "-23.0%", "地产调整、约束"],
    ["固定资产投资", "-16.5%", "地产/二产拖累"],
    ["货物进出口", "-0.7%", "出口-10.7%"],
    ["CPI", "-0.1%", "物价偏弱"],
]
table(doc, g6[0], g6[1:], widths=[3.4, 4.4, 5.6])
para(doc, "**一句话**　增长靠“消费/文旅+工业+免税+海洋”，但房地产/基建投资是最大拖累。2026年考验海南“能不能在释放封关红利的同时、稳住投资”。")

# ---- 七、预算与财政 ----
heading1(doc, "七、预算与财政的“含金量”")
para(doc, "2025年海南全口径一般公共预算收入**1639.34亿元、+1.1%**，其中地方一般公共预算收入**903.74亿元、+1.5%**，税收收入**690.13亿元、+4.5%**。")
bullet(doc, "地方一般公共预算收入+1.5%、税收+4.5%（快于总量）。")
bullet(doc, "财政“稳收+自贸港/民生/基建支出优先”，支撑封关、海洋、旅游。")
bullet(doc, "免税/零关税带动涉税放量：零关税享惠主体封关后新增超1万家。")

# ---- 八、民生底账 ----
heading1(doc, "八、民生底账：人口、收入与城乡")
para(doc, "2025年海南常住人口**1055万人、+7万**，城镇化率**64.08%、+1.00pct**。")
para(doc, "全省居民人均可支配收入**36306元、+4.2%**，农村（+5.2%）连续15年快于城镇（+3.4%）。")
g8 = [
    ["指标", "数值", "信号"],
    ["常住人口", "1055万/+7万", "人口流入"],
    ["城镇化率", "64.08%/+1.00pct", "稳步城镇化"],
    ["居民人均可支配收入", "36306元/+4.2%", "收入/民生稳"],
    ["城镇新增就业", "17.78万人", "就业稳"],
]
table(doc, g8[0], g8[1:], widths=[4.2, 4.2, 4.6])
para(doc, "**民生观察**：人口净流入、收入回暖、农村快于城镇，但物价偏弱仍是短板。")

# ---- 九、城镇与农村 ----
heading1(doc, "九、城镇与农村：格局与均衡")
para(doc, "海南城镇化率64.08%、+1.00pct，三产占比61.7%，城乡协调推进。")
bullet(doc, "农村居民收入+5.2%、快于城镇+3.4%（连续15年）。")
bullet(doc, "乡村社零+8.7%、快于城镇+4.1%，下沉明显。")
bullet(doc, "机会集中在海口经济圈、三亚经济圈、儋洋经济圈。")

# ---- 十、人口流入与流出 ----
heading1(doc, "十、人口流入与流出")
para(doc, "海南2025年常住人口1055万、净增7万，是少有的人口流入省份，主要流向海口/三亚/儋洋等自贸港/旅游/产业区。")
para(doc, "未来看点：封关运作+免税/海洋/旅游能否持续吸引人口与人才；若“自贸港+免税+海洋科技”成型，海南有望持续人口流入/人才汇聚。")

# ---- 十一、物价与货币 ----
heading1(doc, "十一、物价与货币环境")
para(doc, "2025年海南CPI**-0.1%**、PPI**-4.5%**，物价整体承压、工业品价格走弱。")
para(doc, "物价偏弱反映“消费/服务强、工业品价格弱”，与全国低通胀环境一致。2026年“稳价格/扩内需/虑就业”是主线。")

# ---- 十二、区域一体化 ----
heading1(doc, "十二、区域一体化：海南在“自贸港+封关+南海/海洋”里的位置")
para(doc, "海南的核心战略坐标是“**全国唯一全域自贸港+全岛封关运作+南海/海洋+免税**”，也是中国高水平开放的前沿。")
bullet(doc, "自贸港/封关：12月18日启动，制度红利扩展。")
bullet(doc, "海洋经济：占GDP超36%，深海科技/南繁/大养殖。")
bullet(doc, "免税/跨境：离岛免税、服务进出口+22.1%、实际使用外资+19.9%。")
bullet(doc, "三大经济圈：海口/三亚/儋洋，各具分工。")
para(doc, "若“封关+免税+海洋+旅游”闭环跑通，海南将作为全国开放前沿抢占新一轮“封关/跨境/海洋”窗口。")

# ---- 十三、五条主线 ----
heading1(doc, "十三、未来5—10年最值得观察的五条主线")
lines5 = [
    ("① 封关运作/免税消费", "封关红利能否持续放量、免税扩大。"),
    ("② 海洋经济/深海科技", "海洋生产总值、深海/南繁能否壮大。"),
    ("③ 旅游/医疗/免税", "旅游1亿+人次能否升级为高端消费。"),
    ("④ 投资/地产再平衡", "-16.5%/-23.0%，能否用工业/海洋/免税补。"),
    ("⑤ 人口/就业/民生", "人口流入能否被自贸港/旅游托底。"),
]
for l in lines5:
    para(doc, "**%s**　%s" % l, space_after=6)

# ---- 十四、结论 ----
heading1(doc, "十四、最终结论：海南在“自贸港封关+免税+海洋+旅游”里的增长逻辑")
para(doc, "海南的2025年，本质上是“**免税/消费/旅游/海洋/工业为核心，而投资/地产偏弱**”的答卷：GDP8108.85亿、+4.0%，规上工业+9.9%、社零+5.0%、游客1.06亿人次、海洋占GDP超36%、封关启动，但固投-16.5%、地产-23.0%、出口-10.7%、CPI-0.1%。")
para(doc, "只要封关红利、免税/海洋、旅游/医疗、工业持续，海南就站在“全国自贸港+南海海洋”的增长位；如果房地产/投资持续偏弱、出口/工业不稳，海南需承受“消费强、投资弱”的结构挑战。")
para(doc, "最稳观察信号：**一盯封关/免税/零关税（制度红利）、二盯海洋/免税/旅游（开放）、三盯规上工业/制造（动能）、四盯投资/地产（约束）、五盯消费/人口/就业（内需）。**海南，是“自贸港+海洋+旅游”的新样本。")

# ---- 附录A ----
heading1(doc, "附录A：主要资料来源与核验路径")
bullet(doc, "海南省2025年《政府工作报告》——目标来源。")
bullet(doc, "《2025年海南省国民经济和社会发展统计公报》（省统计局）——GDP、工业、外贸、人口实值。")
bullet(doc, "2026年海南省政府工作报告及2025年计划执行情况——2025执行复盘+封关运作、免税、海洋等。")
bullet(doc, "海口海关、省财政厅、省商务厅——对外贸易/财政/服务贸易。")
heading2(doc, "核验说明")
para(doc, "本报告所有增速、总量、占比均以统计公报/官方口径为基准，涉“自贸港/封关/免税/海洋经济/热带农业”等以官方口径为准。")

# ---- 附录B ----
heading1(doc, "附录B：建议建立的年度跟踪仪表盘")
para(doc, "建议把下面10个“测脉搏”指标做成年度跟踪表：")
dash = [
    ["#", "指标", "2025基值", "为什么重要"],
    ["1", "GDP增速", "+4.0%", "总量与方向"],
    ["2", "规上工业增速", "+9.9%", "制造/技改"],
    ["3", "社会消费品零售增速", "+5.0%", "免税/内需"],
    ["4", "旅游人次/花费", "1.06亿/+10.5%", "旅游/消费"],
    ["5", "固定资产投资/地产", "-16.5%/-23.0%", "投资结构"],
    ["6", "海洋生产总值占GDP", "超36%", "海洋经济"],
    ["7", "常住人口/城镇化率", "1055万/64.08%", "人口与城市"],
    ["8", "地方一般公共预算收入增速", "+1.5%", "财政质量"],
    ["9", "货物进出口增速", "-0.7%", "自贸港双循环"],
    ["10", "CPI/PPI", "-0.1%/-4.5%", "物价与需求"],
]
table(doc, dash[0], dash[1:], widths=[0.8, 4.6, 3.4, 4.4])
para(doc, "把这10个指标连起来看，任何一个封关/免税（3）、旅游（4）、工业（2）、海洋（6）向上、投资/地产（5）修复，都说明海南在真正换挡。")

# ===================================== 保存
out = "/Users/x/Desktop/content-prod-lab/reports/海南省_2025年政府工作报告_深度研究_2026-08-13.docx"
doc.save(out)
print("SAVED:", out)
print("PARAGRAPHS:", len(doc.paragraphs))
print("TABLES:", len(doc.tables))
