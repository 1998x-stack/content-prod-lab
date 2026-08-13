# -*- coding: utf-8 -*-
"""Build 重庆市2025年政府工作报告 深度研究 DOCX, 参照北京/上海/杭州/天津系列版式。"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DARK = RGBColor(0x1F, 0x2A, 0x44)
RED = RGBColor(0xB0, 0x1E, 0x2E)
GRAY = RGBColor(0x59, 0x60, 0x69)
BLUE = RGBColor(0x1F, 0x4E, 0x79)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
HEADER_FILL = "1F2A44"
K = "微软雅黑"


def set_run(run, size=11, bold=False, color=DARK, font=K):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    run.font.name = font
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), font)
    rFonts.set(qn("w:hAnsi"), font)
    rFonts.set(qn("w:eastAsia"), font)


def para(doc, text, size=11, bold=False, color=DARK, align=None,
         space_after=6, space_before=0, font=K):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        set_run(r, size=size, bold=bold or (i % 2 == 1), color=color, font=font)
    return p


def heading1(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(16)
    pf.space_after = Pt(8)
    r = p.add_run(text)
    set_run(r, size=16, bold=True, color=RED)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "B01E2E")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def heading2(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(10)
    pf.space_after = Pt(4)
    r = p.add_run(text)
    set_run(r, size=13, bold=True, color=BLUE)
    return p


def table(doc, headers, rows, widths=None, font_size=9.5):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_run(r, size=font_size, bold=True, color=WHITE)
        tcPr = hdr[i]._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), HEADER_FILL)
        tcPr.append(shd)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(val))
            set_run(r, size=font_size, color=DARK)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    return t


def quote_box(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(8)
    pf.left_indent = Pt(12)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), "B01E2E")
    pBdr.append(left)
    pPr.append(pBdr)
    r = p.add_run(text)
    set_run(r, size=11, color=DARK, font="楷体")
    return p


def bullet(doc, text, size=10.5):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.7)
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        set_run(r, size=size, bold=(i % 2 == 1), color=DARK)
    return p


# ================================================================= 文档主体
doc = Document()
sec = doc.sections[0]
sec.page_width = Cm(21)
sec.page_height = Cm(29.7)
sec.top_margin = Cm(2.4)
sec.bottom_margin = Cm(2.2)
sec.left_margin = Cm(2.5)
sec.right_margin = Cm(2.5)

st = doc.styles["Normal"]
st.font.name = K
st.font.size = Pt(11)
st.element.rPr.rFonts.set(qn("w:eastAsia"), K)


# ---- 封面 ----
para(doc, "深圳市2025年政府工作报告", size=24, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=60, space_after=4)
para(doc, "深度研究与“容易被忽视的细节”分析", size=16, bold=True, color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
para(doc, "从“20+8产业集群、人工智能+机器人、科技金融、前海与外贸强外资弱”重新理解深圳", size=12, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)
para(doc, "━━━━━━━━━━━━━━━━", color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
para(doc, "研究对象：2025年深圳市政府工作报告及同期财政、国民经济和社会发展统计资料、2026年官方复盘", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
para(doc, "整理时间：2026年8月", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
para(doc, "分析性质：分析性研究 ｜ 非官方解读", size=11, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
doc.add_page_break()

# ---- 目录 ----
catalog = [
    "执行摘要",
    "一、研究口径与阅读方法",
    "二、先看深圳的特殊底盘：创新之都、外贸强市与改革开放窗口",
    "三、最关键的宏观错位：GDP破3.87万亿、战新与科技强，但投资与地产深度调整",
    "四、容易被忽视、但信号很重的细节（15条）",
    "五、把所有细节连起来：深圳正在经历的“六个换挡”",
    "六、增长暗线：20+8产业集群、人工智能/机器人与科技金融（新质生产力）",
    "七、财政暗线：收入稳、民营与科技税源，土地财政退潮",
    "八、产业暗线：从“代工/地产”到“高端制造+战略性新兴+科技”",
    "九、区域格局：前海、大湾区与深港协同",
    "十、人口与城市：超1800万、全国最年轻的一线城市",
    "十一、民营经济：占六成进出口、科创主体",
    "十二、事后验证：用2025年实际执行结果检验报告判断",
    "十三、未来5—10年最值得观察的五条主线",
    "十四、最终结论：深圳在“创新驱动+对外开放”里的增长逻辑",
    "附录A：主要资料来源与核验路径",
    "附录B：建议建立的年度跟踪仪表盘",
]
para(doc, "目　录", size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=10, space_after=10)
for item in catalog:
    para(doc, item, size=12, space_after=4)
doc.add_page_break()

# ---- 执行摘要 ----
heading1(doc, "执行摘要")
para(doc, "**核心判断**　如果只看新闻标题，2025年深圳最显眼的是“GDP突破3.87万亿、增长5.5%、居一线城市首位”、“高技术制造占规上59.9%”和“战略性新兴产业占GDP43%”。但这份研究真正值得深读的，是这一座“创新之都+商贸强市+开放的窗口”如何在固定资产投资（-21.7%）、房地产开发投资（-31.0%）与出口（-2.6%）调整的背景下，靠“科技+战新+金融+产业投资+机器人/无人机”守住增长。")
para(doc, "把2025年初设定的目标（GDP增长5.5%）、2025年《国民经济和社会发展统计公报》、以及2026年报告对2025年的复盘放在一起看，深圳呈现清晰暗线：**从“外贸+地产/土地+传统制造”的旧依赖，向“科技创新+战略性新兴产业+人才+对外开放”转型**。旧引擎（房地产、土地财政、一般代工制造、部分出口）在调整；新引擎（20+8产业集群、人工智能/机器人、生物医药、新能源车、软件/金融、前海开放）被要求更快补位。")
para(doc, "因此，本报告不按政府工作报告逐段复述，而采用“显性表述—同期数据—制度含义—长期影响”的方式，专门提取容易被忽视、但对判断深圳未来5—10年发展模式更有价值的信号。")
para(doc, "最容易记住的一句话：**深圳是“中国改革开放+科技创新”的窗口城市，以“战新（占GDP43%）+制造（高技术59%）+科技金融+人才”为底盘，靠“20+8产业集群”“人工智能+机器人”和对外开放守全国之首。**观察深圳，与其看“GDP 3.87万亿”，不如看“战新/高技术占比、机器人/无人机、科技金融、研发与人口吸引力”这几张名片。")
heading2(doc, "一页速览：2025年深圳经济的“表与里”")
table(doc,
    ["维度", "表面现象", "更深层信号"],
    [
        ["增长", "GDP 38731.80亿、+5.5%", "一线城市首位、服务+6.3%"],
        ["产业", "规上工业+5.4%、高技术+5.6%", "战新占GDP43%、装备+8.2%"],
        ["外贸", "进出口45533.89亿、+1.4%（出口-2.6%）", "高科技出口+10.1%、占46.6%"],
        ["投资", "固定资产投资-21.7%", "地产-31.0%、工业-27.7%"],
        ["财政", "一般公共预算收入4163.77亿、+6.4%", "财政强、土地退潮"],
        ["消费", "社零10259.93亿、+2.3%", "线上+10.5%、以旧换新"],
        ["人口", "常住1824.85万、城镇化率99.79%", "最年轻一线、人口+26万"],
        ["科技", "战新占GDP43%、工业机器人+43.1%", "人工智能/机器人新赛道"],
    ],
    widths=[2.2, 5.2, 8.6])
para(doc, "", size=4, space_after=2)

# ---- 一、研究口径与阅读方法 ----
heading1(doc, "一、研究口径与阅读方法")
heading2(doc, "1.1 本报告的三份核心底稿")
bullet(doc, "**2025年《政府工作报告》**：2025年2月25日在市七届人大六次会议上作，给出全年预期目标（GDP增长5.5%左右等）。固定资产投资/社零/进出口等未单独设目标。")
bullet(doc, "**2025年《国民经济和社会发展统计公报》**：深圳市统计局2026年5月25日发布，给出全年实际执行数与增速，是“事后验证”的锚点。")
bullet(doc, "**2026年深圳市政府工作报告及官方复盘**：对2025执行结果的追认，交叉验证“目标—执行—再评价”三段式。")
heading2(doc, "1.2 阅读方法：显性—数据—制度—长期")
para(doc, "每一章按“**显性表述 → 同期数据 → 制度含义 → 长期影响**”四层展开。其中“制度含义”回答“政府为什么这么排优先序”，“长期影响”回答“这对未来5—10年意味着什么”。")
para(doc, "关键判别：**当报告使用的“增长词”和统计公报里的实际数不一致时，数据优先。**例如2025年GDP目标5.5%、实际5.5%达标；但固定资产投资-21.7%、地产-31%，说明深圳“科技/金融强、投资/地产深度调整”。以实际执行为准。")

# ---- 二、特殊底盘 ----
heading1(doc, "二、先看深圳的特殊底盘：创新之都、外贸强市与改革开放窗口")
para(doc, "在所有一线城市里，深圳的“底盘”独特：**创新之城+外贸强市+改革开放的窗口**三合一。常住1824.85万、城镇化率99.79%，是全国经济密度与创新浓度最高的城市之一。")
para(doc, "这决定深圳的多重身份并存：**战新与制造**（战新占GDP43%、高技术占规上59.9%）、**外贸强市**（进出口4.55万亿、居内地城市首位）、**科技金融**（金融机构、创投）、**对外开放**（前海、深港、都市紧密）、**人才高地**（人口年轻、286+万高校生）。")
heading2(doc, "2.1 战新与新质生产力")
para(doc, "深圳以“20+8产业集群”为骨架，战略性新兴产业占GDP43%、高技术制造占59.9%。工业机器人+43%、无人机+40%、3D打印+45%，人工智能/机器人/新能源汽车是新赛道的发动机。")
heading2(doc, "2.2 外贸/金融/开放")
para(doc, "深圳出口连续33年内地城市第一、金融业增加值+12.1%（占GDP）、前海/深港深化开放。深圳是“制造+金融+外贸+开放”四位一体。")
heading2(doc, "2.3 人口/人才")
para(doc, "深圳常住1824.85万、城镇化99.79%，最年轻的一线城市（外来人口/年轻人多）。科技创新与人才相生相成。")

# ---- 三、宏观错位 ----
heading1(doc, "三、最关键的宏观错位：GDP破3.87万亿、战新与科技强，但投资与地产深度调整")
para(doc, "把2025年深圳的宏观面放进一张表，会出现令人惊讶的“错位”：表观增长来自战新/制造/金融/科技，而投资与地产深度调整。这个错位，正是读懂深圳的关键。")
heading2(doc, "3.1 “强的一面”")
table(doc,
    ["指标", "2025实绩", "信号"],
    [
        ["GDP", "38731.80亿、+5.5%", "一线城市首位"],
        ["高技术制造", "+5.6%、占规上59.9%", "新质生产力"],
        ["战略性新兴", "占GDP 43.0%", "战新主导"],
        ["装备制造", "+8.2%、占72.6%", "高端制造"],
        ["金融业", "增加值+12.1%", "科技金融"],
        ["一般公共预算收入", "+6.4%", "财政强"],
    ],
    widths=[3.2, 5.4, 6.0])
heading2(doc, "3.2 “调整面大的一面”")
table(doc,
    ["指标", "2025实值", "信号"],
    [
        ["固定资产投资", "-21.7%", "大幅调整"],
        ["房地产开发投资", "-31.0%", "地产深度调整"],
        ["工业投资", "-27.7%", "投资回落"],
        ["出口", "-2.6%", "出口回落"],
        ["社零", "+2.3%", "消费偏弱"],
    ],
    widths=[3.2, 5.4, 6.0])
para(doc, "**错位结论**　深圳的增长“很强、但也在结构性调整”。强的部分（GDP/战新/科技/金融/财政）与弱的部分（投资/地产/出口/消费）并存。**真正的焦点是“科技/金融强，投资/地产/出口调整”**：GDP/战新/财政都强，但固定资产投资-21.7%、地产-31%大幅收缩。2026年深圳在“战新+科技，稳投资/地产/出口”间平衡。")

# ---- 四、被忽视细节 ----
heading1(doc, "四、容易被忽视、但信号很重的细节（15条）")
details = [
    ("1", "规上工业+5.4%、高技术制造+5.6%（占59.9%）", "结构高度先进。"),
    ("2", "战新产业占GDP 43.0%", "战新全国最高之一。"),
    ("3", "装备制造+8.2%、占72.6%", "装备/智能制造。"),
    ("4", "工业机器人产量+43.1%、无人机+40.1%", "机器人与无人机放量。"),
    ("5", "3D打印、可穿戴/智能硬件高增", "智能硬件/新赛道。"),
    ("6", "金融业增加值5261.56亿、+12.1%（占GDP比重高）", "科技金融强。"),
    ("7", "信息软件/服务+10.3%", "软件/信息产业。"),
    ("8", "高新技术产品出口+10.1%、占46.6%", "出口结构向高技术。"),
    ("9", "固定资产投资-21.7%、地产-31.0%", "投资/地产深度调整。"),
    ("10", "民间投资占34.8%、技改投资+19.2%", "民企/技改强。"),
    ("11", "社零+2.3%、线上+10.5%", "消费偏弱、线上强。"),
    ("12", "常住1824.85万、（净增25.9万）城镇化99.79%", "人口强、最年轻一线。"),
    ("13", "居民人均可支配收入84945元、+4.7%", "收入高。"),
    ("14", "出口连续33年居内地大中城市首位、港口集装箱吞吐量+6.0%", "外贸/开放双强。"),
    ("15", "CPI+0.2%、PPI+0.3%（全国罕见工业价正）", "深圳物价较稳、工业品价正。"),
]
for d in details:
    para(doc, "**%s**" % d[0], size=11)
    para(doc, d[1], size=10.5, color=GRAY)

# ---- 五、六个换挡 ----
heading1(doc, "五、把所有细节连起来：深圳正在经历的“六个换挡”")
gear = [
    ("1．动能换挡：从“代工/地产驱动”到“战新+科技+新质生产力”", "战新占GDP43%、高技术59.9%、机器人/无人机。"),
    ("2．产业换挡：从“电子代工/地产”到“20+8产业集群+AI/机器人”", "人工智能、半导体、新能源、生物医药、无人机。"),
    ("3．投资换挡：从“地产/土地”到“科创/产业投资+技改”", "地产-31%、技改+19.2%、民间34.8%。"),
    ("4．开放换挡：从“传统出口”到“高新技术+前海/深港”", "高科技出口+10.1%、前海、数字/服务。"),
    ("5．人口换挡：从“人口流入”到“人才/科创集聚”", "常住1824万、最年轻一线、人才洼地。"),
    ("6．动能换挡：从“土地财政”到“科技金融+民营税源”", "金融+12.1%、财政收入稳、土地退潮。"),
]
for g in gear:
    para(doc, "**%s**　%s" % g, space_after=6)

# ---- 六、增长暗线 ----
heading1(doc, "六、增长暗线：20+8产业集群、人工智能/机器人与科技金融（新质生产力）")
heading2(doc, "6.1 20+8产业集群+新质")
para(doc, "深圳以“20+8产业集群”为骨架，战新占GDP43%、高技术占59.9%。工业机器人+43%、无人机+40%、3D打印+45%，人工智能/机器人/新能源车是新赛道发动机。")
heading2(doc, "6.2 科技金融")
para(doc, "金融业增加值+12.1%、占GDP重，信息软件+10.3%。深圳“创投+科技+金融”生态，是科技转化与独角兽的土壤。")
para(doc, "**这条暗线意味着**：深圳的增长叙事正从“外贸/地产”转向“战新+科技+金融+开放”。看深圳，盯住“战新/高技术占比、机器人/无人机、科技金融、研发与人口”这几组数。")

# ---- 七、财政暗线 ----
heading1(doc, "七、财政暗线：收入稳、民营与科技税源，土地财政退潮")
para(doc, "2025年深圳一般公共预算收入4163.77亿、+6.4%，财政来源已从“土地财政”转向“科技/金融/民营”税源。支出向科创、民生、前海倾斜。")
para(doc, "**制度含义**　深圳“土地财政退潮、科技/民营税源为主”是财政健康化。关键是保持“战新+科技+金融”税源，支撑创新与改革。")

# ---- 八、产业暗线 ----
heading1(doc, "八、产业暗线：从“代工/地产”到“高端制造+战略性新兴+科技”")
heading2(doc, "8.1 深圳产业的“表”")
table(doc,
    ["指标", "2025增速/信号", "解读"],
    [
        ["规上工业", "+5.4%", "制造强市"],
        ["高技术制造", "+5.6%、占59.9%", "高技术主导"],
        ["战略性新兴", "占GDP43%", "战新最高之一"],
        ["装备制造", "+8.2%、占72.6%", "高端装备"],
        ["工业机器人", "+43.1%", "机器人赛道"],
        ["无人机", "+40.1%", "低空/无人机"],
        ["出口/高科技", "高科技出口+10.1%", "出口向高技术"],
    ],
    widths=[4.6, 3.6, 5.0])
heading2(doc, "8.2 从“代工/地产”到“科技/战新”")
para(doc, "深圳从代工/地产，升级为“20+8产业集群+高端装备+AI/机器人+科技金融”。战新占GDP 43%、高技术占60%，是深圳“创新驱动”的底盘。")

# ---- 九、区域 ----
heading1(doc, "九、区域格局：前海、大湾区与深港协同")
para(doc, "深圳是粤港澳大湾区核心，前海深港现代服务业合作区/CBD等推进制度型开放。深港协同、广州-深圳双城联动，是广东都市圈/大湾区的重要引擎。")
para(doc, "以前海、河套、光明科学城等平台，深圳开展面向港澳与国际的制度型开放。大湾区一体化+前海，是深圳开放与创新的桥头堡。")

# ---- 十、人口与城市 ----
heading1(doc, "十、人口与城市：超1800万、全国最年轻的一线城市")
heading2(doc, "10.1 人口总量")
para(doc, "2025年末深圳常住1824.85万、城镇化率99.79%、净增25.9万。深圳是全国最年轻、人口流入最强的一线城市之一。")
heading2(doc, "10.2 城市与人才")
para(doc, "深圳人口年轻、人才集聚，科技创新+城市治理是特色。城市向东/维搬迁、土地高效利用，支撑超大城市宜居与创新。")

# ---- 十一、民营经济 ----
heading1(doc, "十一、民营经济：占六成进出口、科创主体")
heading2(doc, "11.1 民企地位")
para(doc, "深圳民营企业发达，民营进出口占六成多（民企贸易强），市场主体活跃。民营+科技企业是深圳经济核心底盘。")
heading2(doc, "11.2 政策/金融")
para(doc, "深圳优化营商环境、创投+科技金融+持续改革。民营+专精特新+科创，是未来10年深圳的核心。")

# ---- 十二、事后验证 ----
heading1(doc, "十二、事后验证：用2025年实际执行结果检验报告判断")
heading2(doc, "12.1 年初目标与年末执行对比")
table(doc,
    ["指标", "2025年初目标", "2025实际", "偏差"],
    [
        ["GDP增速", "5.5%左右", "+5.5%", "达标"],
        ["高技术制造", "保持领先", "+5.6%、占59.9%", "强"],
        ["战略性新兴", "占GDP比重提升", "43.0%", "强"],
        ["固定资产投资", "未设", "-21.7%", "大幅调整"],
        ["进出口", "稳量提质", "+1.4%", "基本稳"],
        ["居民收入", "与增长同步", "+4.7%", "同步"],
    ],
    widths=[3.0, 3.2, 3.0, 4.4])
heading2(doc, "12.2 判断验证")
para(doc, "三条核心判断得到支持：**（1）战新/科技/金融强（占GDP43%、金融+12.1%），验证；（2）投资/地产深度调整（固投-21.7%、地产-31%），验证；（3）外贸/开放/科创韧性强，验证。**")
para(doc, "核心观察：深圳靠“战新+科技+金融+开放”守住一线首位，但投资/地产深度调整是最大约束。2026年“稳投资/地产、强化科技与开放”是重点。")

# ---- 十三、五条主线 ----
heading1(doc, "十三、未来5—10年最值得观察的五条主线")
lines5 = [
    ("① 20+8集群/人工智能/机器人/芯片", "从“制造”到“智能/自主”能否兑现。"),
    ("② 科技金融与创投", "创投、前海、深交所、跨境，能否孵化更多一流企。"),
    ("③ 出口/开放/前海/深港", "高科技出口、前海、开放，能否对冲传统出口波动。"),
    ("④ 地产/投资/土地财政退潮", "-31%、-21%，能否用科技与产业补投资。"),
    ("⑤ 人口/人才、保障性住房", "年轻人口、人才、住房，能否维持年轻/宜居。"),
]
for l in lines5:
    para(doc, "**%s**　%s" % l, space_after=6)

# ---- 十四、结论 ----
heading1(doc, "十四、最终结论：深圳在“创新驱动+对外开放”里的增长逻辑")
para(doc, "深圳的2025年，本质上是“**战新+科技+金融+开放为核心，而地产/投资深度调整**”的答卷：GDP3.87万亿、一线首位、战新43%，代价是地产-31%、投资-21.7%。")
para(doc, "只要“20+8、AI/机器人、金融科技、开放”能接住，深圳就仍是创新之都；如果地产/投资继续深调、传统出口疲弱，深圳需承受“科技强、投资/地产弱”的转型。")
para(doc, "最稳妥的观察信号：**一盯战新/AI/机器人（动能）、二盯科技金融（金融）、三盯出口/前海（开放）、四盯地产/投资（约束）、五盯人口/住房（底座）。**深圳，是“创新+开放”的样板城市。")

# ---- 附录A ----
heading1(doc, "附录A：主要资料来源与核验路径")
bullet(doc, "深圳市2025年《政府工作报告》（2025年2月）——目标来源。")
bullet(doc, "《2025年深圳市国民经济和社会发展统计公报》（深圳市统计局，2026-05-25）——GDP、工业、外贸、人口实值。")
bullet(doc, "深圳市统计局/科创专题、20+8集群专报——战新与科创。")
bullet(doc, "2026年深圳市政府工作报告——2025执行复盘。")
bullet(doc, "深圳海关、市财政厅（局）、前海——外贸/财政/开放。")
heading2(doc, "核验说明")
para(doc, "本报告所有增速、总量、占比均以统计公报/官方发布为基准。涉“前海专项”“税收收入”等未披露项以官方口径为准。")

# ---- 附录B ----
heading1(doc, "附录B：建议建立的年度跟踪仪表盘")
para(doc, "建议把下面10个“测脉搏”指标做成年度跟踪表：")
dash = [
    ["#", "指标", "2025基值", "为什么重要"],
    ["1", "GDP增速", "+5.5%", "总量与方向"],
    ["2", "高技术制造占比", "59.9%", "新质生产力"],
    ["3", "战略性新兴占GDP", "43.0%", "战新驱动"],
    ["4", "工业机器人/无人机", "+43.1%/+40.1%", "AI/机器人"],
    ["5", "进出口/高科技出口", "+1.4%/+10.1%", "外贸韧性"],
    ["6", "固定资产投资/地产", "-21.7%/-31.0%", "投资/地产调整"],
    ["7", "金融业增速", "+12.1%", "科技金融"],
    ["8", "常住人口/城镇化率", "1824.85万/99.79%", "人口与城市"],
    ["9", "一般公共预算收入", "+6.4%", "财政质量"],
    ["10", "居民人均可支配收入增速", "+4.7%", "民生获得感"],
]
table(doc, dash[0], dash[1:], widths=[0.8, 4.5, 3.2, 4.6])
para(doc, "把这10个指标连起来看，任何一个“新引擎（2/3/4）向上、旧引擎（6）调整”，都说明深圳在真正换挡；反之则是旧路径的反复。")

# ========================================= 保存
out = "/Users/x/Desktop/content-prod-lab/reports/深圳市_2025年政府工作报告_深度研究_2026-08-13.docx"
doc.save(out)
print("SAVED:", out)
print("PARAGRAPHS:", len(doc.paragraphs))
print("TABLES:", len(doc.tables))
