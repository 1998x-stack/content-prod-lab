# -*- coding: utf-8 -*-
"""Build 泉州市2025年政府工作报告 深度研究 DOCX, 参照地级市系列版式。"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DARK = RGBColor(0x1F, 0x2A, 0x44)
RED = RGBColor(0xB0, 0x1E, 0x2E)
GRAY = RGBColor(0x59, 0x60, 0x69)
BLUE = RGBColor(0x1F, 0x4E, 0x79)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
HEADER_FILL = "1F2A44"
K = "微软雅黑"


def set_run(run, size=11, bold=False, color=DARK, font=K):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    run.font.name = font
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), font)
    rFonts.set(qn("w:hAnsi"), font)
    rFonts.set(qn("w:eastAsia"), font)


def para(doc, text, size=11, bold=False, color=DARK, align=None,
         space_after=6, space_before=0, font=K):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        set_run(r, size=size, bold=bold or (i % 2 == 1), color=color, font=font)
    return p


def heading1(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(16)
    pf.space_after = Pt(8)
    r = p.add_run(text)
    set_run(r, size=16, bold=True, color=RED)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "B01E2E")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def heading2(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(10)
    pf.space_after = Pt(4)
    r = p.add_run(text)
    set_run(r, size=13, bold=True, color=BLUE)
    return p


def table(doc, headers, rows, widths=None, font_size=9.5):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_run(r, size=font_size, bold=True, color=WHITE)
        tcPr = hdr[i]._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), HEADER_FILL)
        tcPr.append(shd)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(val))
            set_run(r, size=font_size, color=DARK)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    return t


def quote_box(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(8)
    pf.left_indent = Pt(12)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), "B01E2E")
    pBdr.append(left)
    pPr.append(pBdr)
    r = p.add_run(text)
    set_run(r, size=11, color=DARK, font="楷体")
    return p


def bullet(doc, text, size=10.5):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.7)
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        set_run(r, size=size, bold=(i % 2 == 1), color=DARK)
    return p


# ================================================================= 文档主体
doc = Document()
sec = doc.sections[0]
sec.page_width = Cm(21)
sec.page_height = Cm(29.7)
sec.top_margin = Cm(2.4)
sec.bottom_margin = Cm(2.2)
sec.left_margin = Cm(2.5)
sec.right_margin = Cm(2.5)

st = doc.styles["Normal"]
st.font.name = K
st.font.size = Pt(11)
st.element.rPr.rFonts.set(qn("w:eastAsia"), K)


# ---- 封面 ----
para(doc, "泉州市2025年政府工作报告", size=24, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=60, space_after=4)
para(doc, "深度研究与\u201c容易被忽视的细节\u201d分析", size=16, bold=True, color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
para(doc, "从\u201c海丝起点、泉企民营、鞋服（晋江）、石材、石化、侨乡\u201d重新理解泉州", size=12, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)
para(doc, "\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501", color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
para(doc, "研究对象：2025年泉州市政府工作报告及同期财政、国民经济和社会发展统计资料、2026年官方复盘", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
para(doc, "标定日期：2026-08-14", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
doc.add_page_break()

# ---- 目录 ----
catalog = [
    "执行摘要",
    "一、研究口径与阅读方法",
    "二、先看泉州的特殊底盘：海丝起点、泉服、鞋Jin江、纺织石材、石化与民营经济",
    "三、最关键的宏观错位：GDP 13778.34亿/5.3%，制造强但地产负、外贸下滑、房地产-32.3%",
    "四、容易被忽视、但信号很重的细节（15条）",
    "五、2025年GDP目标 vs 实际：对照表",
    "六、2025年增长是谁撑起来的？（结构归因）",
    "七、预算与财政的\u201c含金量\u201d",
    "八、民生底账：人口、收入与城乡",
    "九、城镇与农村：格局与均衡",
    "十、人口流入与流出",
    "十一、物价与货币环境",
    "十二、区域一体化：泉州在\u201c闽西南+海丝核心区+两岸融合\u201d里的位置",
    "十三、未来5—10年最值得观察的五条主线",
    "十四、最终结论：泉州在\u201c民营制造+海丝侨乡\u201d里的增长逻辑",
    "附录A：主要资料来源与核验路径",
    "附录B：建议建立的年度跟踪仪表盘",
]
para(doc, "目　录", size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=10, space_after=10)
for item in catalog:
    para(doc, item, size=12, space_after=4)
doc.add_page_break()

# ---- 执行摘要 ----
heading1(doc, "执行摘要")
para(doc, "**核心判断**　2025年泉州最显著的是\u201cGDP 13778.34亿元、增长5.3%、福建第1、万亿元级\u201d、\u201c规上工业+7.0%（制造业+7.4%、纺织+19.3%、文教工美+33.8%）\u201d、\u201c民营经济主体171.12万户、\u2019泉服\u2019鞋开\u201d、\u201c但房地产-32.3%、进出口-12.9%、外贸承压\u201d、\u201c地方一般公共预算收入592.07亿/+3.4%\u201d。这说明泉州在\u201c泉企民营+海丝\u201d中，**工业、民营强但地产、外贸调整**。")
para(doc, "把2025年目标（GDP+5%左右）、2025年统计（GDP+5.3%达成、规上+7.0%、社零+4.1%、进出口-12.9%、财收+3.4%）、趋势一起看，泉州是\u201c民营企业+纺织/鞋/石材/石化\u201d路径：**纺织服装（晋江安踏）、石油化工（湄洲湾）、石材（南安）、茶/文旅、海丝**是支柱；2025年总量13778亿居福建第1。")
para(doc, "最容易记住的一句话：**泉州是\u201c海丝起点、民营经济第一城（鞋/服装/石材/安踏）、福建GDP之王\u201d，靠\u201c民营制造+侨乡海丝\u201d增长。**观察泉州，与其只看\u201cGDP 13778亿\u201d，不如看\u201c纺织+19.3%、文教工美+33.8%、民营171万户、港口集装箱206万标箱、海丝\u201d。")

# ---- 一、研究口径与阅读方法 ----
heading1(doc, "一、研究口径与阅读方法")
para(doc, "本报告以\u201c泉州市政府工作报告（2025年，蔡战胜作）\u201d为起点，把\u201c2025年GDP目标（5%左右）\u201d与\u201c官方2025年GDP（13778.34亿元/+5.3%）\u201d并置对照，并用\u201c2025年泉州市统计公报\u201d和\u201c2026年政府工作报告\u201d横向核验。")
para(doc, "口径提示：\u201cGDP、规上工业增加值增速\u201d按可比价（实际增速）；\u201c投资、消费、进出口、财政收入\u201d按现价（名义增速）。涉及\u201c常住人口\u201d公报未披露，本报告引用户籍773.49万、常住约890万（结合公开数据）；城镇化率约70%+。")
para(doc, "指标体系（与研究口径一致）：五大象限——**总量与增速（GDP）、产业动能（鞋服/石化/石材/纺织/民营）、投资、财政、民生与海丝**。")
para(doc, "特别提示（不吃老本）：泉州是**福建GDP第一城、民营经济标杆（安踏、恒安、达利）、海丝起点**；2025年工业强（规上+7.0%）但地产、外贸调整——真正要看的是\u201c民营制造转型、地产去库存、外向海丝\u201d。")
# ---- 二、先看泉州的特殊底盘 ----
heading1(doc, "二、先看泉州的特殊底盘：海丝起点、泉企民营、鞋服（晋江）、纺织石材、石化与港口")
para(doc, "泉州地处福建东南沿海、闽南，是**海丝起点、福建GDP第一城、中国民营经济标杆（泉州模式）、侨乡**。2025年GDP 13778.34亿元、户籍773.49万（常住约890万），人均突破15万，居福建第1。")
para(doc, "五个底盘名词，先立框架：")
bullet(doc, "**民营经济看泉**　民营经济主体171.12万户（全省第一）、\u201c泉州模式\u201d；安踏、恒安、达利、七匹狼等民企。")
bullet(doc, "**纺织服装（晋江）**　晋江是\u201c中国鞋都\u201d，安踏/特步/361°、运动鞋服，纺织业+19.3%。")
bullet(doc, "**石油化工（湄洲湾）**　泉港/湄洲湾石化基地，石油化工是GDP第一支柱；石化+化工材料。")
bullet(doc, "**石材（南安）**\u201c中国石材之都\u201d，建材、水暖，传统强项。")
bullet(doc, "**海丝/港口**　泉州港（海丝起点、全球贸易史重镇），集装箱206.46万标箱、吞吐量1.16亿吨；侨/海丝开放、跨境电商。")
para(doc, "这五根（民营+纺织+石化+石材+海丝）构成泉州独特底盘：**左手民营制造（鞋/服/建材），右手海丝港口，根子是晋江精神（爱拼敢赢）**。理解泉州，先理解\u201c泉州民企+侨乡\u201d。")

# ---- 三、最关键的宏观错位 ----
heading1(doc, "三、最关键的宏观错位：GDP 13778.34亿/5.3%，制造强但地产负、外贸下滑、房地产-32.3%")
para(doc, "2025年泉州最需要辨析的一组\u201c错位\u201d：**GDP 5.3%（达成5%目标）、规上工业+7.0%（制造强）但房地产-32.3%、进出口-12.9%、固投+1.1%低**。")
para(doc, "为什么\u201c工业这么强、经济却在放缓并靠地产外贸拖累\u201d？三个解释：")
para(doc, "**其一，工业强、结构优**　规上+7.0%（纺织+19.3%、文教工美+33.8%）、制造业+7.4%——但增速受制于\u201c高基数+传统（纺织/建材）转型\u201d。")
para(doc, "**其二，地产深度调整**　房地产-32.3%、商品房销售-24.1%、国内贷款-58.3%——地产拖累固投（+1.1%）与财政（土地）。")
para(doc, "**其三，外贸转负**　进出口-12.9%（出口-16.9%）、对美-21.3%、对东盟-19.6%——泉服/纺织/石材出口受贸易摩擦压制。")
para(doc, "小结：泉州2025年是\u201c**强工业民营、弱地产外贸、结构性调整\u201d**的一年：制造/纺织强、民间投资稳，但地产、外贸、固投偏弱。")

# ---- 四、容易被忽视、但信号很重的细节（15条） ----
heading1(doc, "四、容易被忽视、但信号很重的细节（15条）")
bullet(doc, "**1.规上工业+7.0%、制造业+7.4%**\u201c纺织+19.3%、文教工美+33.8%强（泉州制造）。\u201d")
bullet(doc, "**2.民营经济主体171.12万户（全省第一）**\u201c晋江经验：爱拼才会赢（安踏。恒安。达利）。\u201d")
bullet(doc, "**3.社零6416.07亿/+4.1%（全省第二）**\u201c万亿消费、新能车+23.4%、家电+47.8%。\u201d")
bullet(doc, "**4.进出口2363.79亿/-12.9%**\u201c对美-21.3%、对东盟-19.6%，海丝出口承压。\u201d")
bullet(doc, "**5.房地产-32.3%、商品房销售-24.1%**\u201c地产深调，是2025年最大拖累。\u201d")
bullet(doc, "**6.地方一般公共预算收入592.07亿/+3.4%**\u201c财政稳（税收+、土地-）。\u201d")
bullet(doc, "**7.港口集装箱206.46万标箱/+1.1%**\u201c海丝、面向东南亚。\u201d")
bullet(doc, "**8.工业投资+14.8%、电力热力投资+57.2%**\u201c能源/设备升级投资热情。\u201d")
bullet(doc, "**9.文教工美+33.8%、纺织+19.3%、通用设备+23.1%**\u201c泉州制造升级（专精特新62家）。\u201d")
bullet(doc, "**10.居民收入54858元/+5.1%、城乡比2.00**\u201c农村+5.7%快于城镇+4.6%（持平/略缩小）。\u201d")
bullet(doc, "**11.常住约800万、户籍773.49万**\u201c闽南最大城市、侨乡（海外侨胞800万+）。\u201d")
bullet(doc, "**12.高新技术企业2730家**\u201c科创成色提升（研发+11.2%）。\u201d")
bullet(doc, "**13.CPI+1.0%**\u201c温和通胀（衣着+5.7%、其他用品+10.7%）。\u201d")
bullet(doc, "**14.粮食52.91万吨/+2.5%**\u201c农业稳（茶11.54万吨、水产品120万吨）。\u201d")
bullet(doc, "**15.旅游总花费1348.30亿/+11.3%**\u201c世遗泉州文旅爆发（入境游客+34.2%）。\u201d")
# ---- 五、2025年GDP目标 vs 实际：对照表 ----
heading1(doc, "五、2025年GDP目标 vs 实际：对照表")
table(doc,
    ["指标", "2025年目标", "2025年实际", "达成评价"],
    [
        ["地区生产总值(GDP)", "增长5%左右", "13778.34亿/5.3%", "达成"],
        ["规上工业增加值", "——", "+7.0%", "工业强"],
        ["固定资产投资", "——", "+1.1%", "偏低"],
        ["社会消费品零售总额", "——", "6416.07亿/+4.1%", "稳健"],
        ["进出口总额", "——", "2363.79亿/-12.9%", "负增长"],
        ["地方一般公共预算收入", "——", "592.07亿/+3.4%", "稳增"],
        ["居民收入", "与经济增长同步", "54858元/+5.1%", "同步"],
    ],
)
para(doc, "注：GDP、规上工业按可比价。**GDP（+5.3%）达成、规上（+7.0%）、财收（+3.4%）稳健**；**进出口（-12.9%）、房地产（-32.3%）转负**。")
para(doc, "拆读：**工业（纺织+19.3%/文教工美+33.8%）、民营（171万户）、社零（+4.1%）是亮色**；**地产（-32.3%）、外贸（-12.9%）是短板**——\u201c制造强、地产外贸弱\u201d，是闽南万亿城样本。")

# ---- 六、2025年增长是谁撑起来的？（结构归因） ----
heading1(doc, "六、2025年增长是谁撑起来的？（结构归因）")
para(doc, "把泉州GDP的5.3%拆开：三次产业分别增3.6%、6.0%、4.7%（结构2.0：50.5：47.5）。**第二产业（工业）是主引擎（+6.0%），第三产业（服务业）稳，第一产业（农业）弱体量小**。")
para(doc, "2026年泉州强调\u201c民营经济、海丝、世界遗产城市\u201d，聚焦**纺织服装升级、石化、石材、新材料、海洋经济、文旅会展**——核心是\u201c实业民营+海丝链接\u201d。")
para(doc, "**第二产业（工业）**：规上+7.0%（纺织+19.3%、文教工美+33.8%、石化）、装备/通用设备+23.1%——\u201c制造业强\u201d。")
para(doc, "**第三产业（服务业）**：批发零售、文旅（旅游+11.3%）、港口物流、数字经济——\u201c服务业稳\u201d。")
para(doc, "**外贸（开放）**：进出口-12.9%转负、对美/东盟下滑，但海丝/跨境电商、新外贸——\u201c外向调整\u201d。")
para(doc, "一句话归因：**2025年泉州增长\u201c靠工业（纺织/制造）+民营+消费\u201d**，地产、外贸拖累；\u201c实业强、地产外贸弱\u201d是泉州核心特征。")

# ---- 七、预算与财政的\u201c含金量\u201d ----
heading1(doc, "七、预算与财政的\u201c含金量\u201d")
para(doc, "2025年泉州**地方一般公共预算收入592.07亿元（+3.4%）**、一般公共预算总收入991.6亿（+1.7%）；一般公共预算支出880.29亿元（+4.6%）。")
bullet(doc, "税收结构：税务组织收入+0.5%、税收814.34亿（-1.2%）——\u201c财收+3.4%靠非税/国企，税收微降\u201d。")
bullet(doc, "民生支出：教育230.04亿（+3.7%）、公积/社保——民生优先。")
bullet(doc, "金融支撑：存款13315.22亿（+6.0%）、贷款13050.82亿（+6.4%，普惠小微3277亿）——信贷宽对民企/小微。")
para(doc, "**财政含金量小结**：财收+3.4%稳增、税收微降；\u201c稳收入、宽信贷（对民营）\u201d；财政对制造、海丝、民生投入加大。")

# ---- 八、民生底账：人口、收入与城乡 ----
heading1(doc, "八、民生底账：人口、收入与城乡")
para(doc, "2025年泉州**居民人均可支配收入54858元（+5.1%）**，其中城镇66511元（+4.6%）、农村33321元（+5.7%），城乡比2.00（略缩小）。就业：城镇新增就业9.49万人。")
para(doc, "人口画像：**户籍773.49万（常住约890万）**，闽南人口大市、超大侨乡（海外侨胞数百/800万）；人口结构相对年轻、劳动密集（纺织/鞋业）吸纳劳动力。")
para(doc, "民生投入：中学/中小学学位、医疗床位、五险新增14万人次；职业技能培训5.2万人次——民生保障扎实。")

# ---- 九、城镇与农村：格局与均衡 ----
heading1(doc, "九、城镇与农村：格局与均衡")
para(doc, "泉州常住城镇化率约70%（闽南发达），城乡格局较均衡；农村收入增速（+5.7%）高于城镇（+4.6%），**城乡比缩小（2.00）**。")
para(doc, "农业底盘：**粮食52.91万吨（+2.5%）**、茶叶11.54万吨、园林水果24.51万吨、水产品120万吨——\u201c福建茶香+渔+果蔬\u201d。")
para(doc, "一句话：\u201c泉州城镇制造强、农村产业兴（乡镇厂/晋江/南安）\u201d、城乡比尚可，但\u201c纺织劳动密集、县域经济靠乡镇\u201d。")

# ---- 十、人口流入与流出 ----
heading1(doc, "十、人口流入与流出")
para(doc, "泉州常住约890万（户籍773万），是\u201c常住>户籍\u201d的人口净流入型；闽南制造业吸引省内/中西部劳动力（晋江/石狮工厂），\u201c人口+劳动力\u201d充足。")
para(doc, "结构观察：**全国少有的出生率高于死亡率**、人口老龄化相对较轻；侨乡（海外800万惠安等）双向流动、跨境人口。")
para(doc, "2026年目标：引才11.2万（2025已）、大学生/人才南下——泉州以\u201c民营制造+海丝\u201d聚人。")

# ---- 十一、物价与货币环境 ----
heading1(doc, "十一、物价与货币环境")
para(doc, "2025年泉州**CPI同比+1.0%**（衣着+5.7%、其他用品+10.7%；食品烟酒+0.4%、交通通信-2.1%）——**温和通胀（消费品涨价）**。")
bullet(doc, "信贷扩张：贷款+6.4%、存款+6.0%，普惠小微/绿色贷款增长——宽信用对民营小微。")
bullet(doc, "消费：新能源车+23.4%、家电+47.8%、线上+17.8%——\u201c换新/新消费\u201d。\u201d")
para(doc, "货币环境判断：**宽信用、CPI+1.0%温和**；泉州靠\u201c以旧换新、世遗文旅、民营制造\u201d稳内需（2026扩内需）。")

# ---- 十二、区域一体化：泉州的位置 ----
heading1(doc, "十二、区域一体化：泉州在\u201c闽西南+海丝核心区+两岸融合\u201d里的位置")
para(doc, "泉州是\u201c海上丝绸之路核心区\u201d、闽西南城市群、福建GDP第一城，地处\u201c闽东南沿海+对台前沿\u201d。")
bullet(doc, "**海丝核心区**　泉州海丝起点（开元寺/洛阳桥世界遗产）、面向东南亚/阿拉伯——\u201c21世纪海丝\u201d枢纽。")
bullet(doc, "**闽西南协同**　泉州-厦门-漳州\u201c闽西南都市圈\u201d，泉为制造、厦为航运自贸、漳为农业。")
bullet(doc, "**两岸融合**　晋江/泉州对台（台湾80%祖籍闽南）、金门——两岸经贸人文纽带。")
para(doc, "一句话：**泉州在\u201c闽西南+海丝\u201d里，最核心的定位是\u201c福建制造、海丝起点\u201d**——民营、侨、海丝是最大优势。")

# ---- 十三、未来5—10年最值得观察的五条主线 ----
heading1(doc, "十三、未来5—10年最值得观察的五条主线")
bullet(doc, "**主线一：民营制造升级（纺织/鞋/建材）**\u201c安踏、恒安、纺织+19.3%\u201d能否由\u201c鞋服\u201d到\u201c智造/品牌\u201d。")
bullet(doc, "**主线二：石化/石化新材料**\u201c湄洲湾石化\u201d能否由\u201c炼油\u201d到\u201c化工新城\u201d。")
bullet(doc, "**主线三：海丝/跨境出海**\u201c海丝+华侨+跨境电商\u201d能否对冲\u201c对美外贸转负\u201d。")
bullet(doc, "**主线四：世遗文旅/城市更新**\u201c世界遗产（泉州）旅游+11.3%、泉州古城\u201d新动能。")
bullet(doc, "**主线五：地产与投资**\u201c房地产-32.3%\u201d能否靠\u201c保障房、城市更新\u201d软着陆。")

# ---- 十四、最终结论 ----
heading1(doc, "十四、最终结论：泉州在\u201c民营制造+海丝侨乡\u201d里的增长逻辑")
para(doc, "把2025年闭环打穿：**泉州是\u201c福建GDP第一城、海丝起点、民营之都\u201d**：GDP 13778.34亿/+5.3%、规上+7.0%（纺织+19.3%）、社零+4.1%、财收+3.4%。")
para(doc, "泉州不是\u201c只有鞋和石材\u201d——它是**纺织服装+石油化工+建材+海丝+民营**的复合制造之都，靠\u201c实业民营+侨乡宽容开放\u201d驱动；但地产、外贸调整，\u201c制造强、地产弱\u201d。")
para(doc, "一句话结论：**泉州是\u201c海丝起点、民营经济第一城、福建制造\u201d；观察它先看\u201c纺织服装（晋江）、石化、石材、民营、海丝\u201d，再看\u201c地产、外贸、固投\u201d。**它是\u201c制造强、民营活、外贸弱\u201d的闽南样本。")

# ---- 附录A ----
heading1(doc, "附录A：主要资料来源与核验路径")
bullet(doc, "《2025年泉州市政府工作报告》（2025年1月，蔡战胜作，2025年目标、2024年回顾）")
bullet(doc, "《2025年泉州市国民经济和社会发展统计公报》（泉州市统计局，2026-03-31，2025年实际数据）")
bullet(doc, "《2026年泉州市政府工作报告》（2026年1月，2025年复盘+2026年目标）")
bullet(doc, "泉州市政府、泉州市统计局（quanzhou.gov.cn/tjj）")

# ---- 附录B ----
heading1(doc, "附录B：建议建立的年度跟踪仪表盘")
bullet(doc, "GDP总量/增速 vs 全年目标（可比价）。")
bullet(doc, "规上工业增加值及分行业（纺织/石化/建材/机电/高技术）增速。")
bullet(doc, "民营经济主体数、民间投资、小微企业贷款。")
bullet(doc, "固定资产投资/工业/房地产投资增速。")
bullet(doc, "社会消费品零售总额、线上、以旧换新。")
bullet(doc, "进出口、出口、海丝/东盟/美国。")
bullet(doc, "地方一般公共预算收入、税收/非税、支出。")
bullet(doc, "常住人口、城镇新增就业、侨/人才。")
bullet(doc, "CPI、金融存贷款、普惠微贷。")
bullet(doc, "港口集装箱、旅游、海丝项目。")

# ---------------- 保存 ----------------
out = "/Users/x/Desktop/content-prod-lab/reports/泉州市_2025年政府工作报告_深度研究_2026-08-14.docx"
doc.save(out)
print("SAVED OK: 泉州市", out)
