# -*- coding: utf-8 -*-
"""Build 常州市2025年政府工作报告 深度研究 DOCX, 参照地级市系列版式。"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DARK = RGBColor(0x1F, 0x2A, 0x44)
RED = RGBColor(0xB0, 0x1E, 0x2E)
GRAY = RGBColor(0x59, 0x60, 0x69)
BLUE = RGBColor(0x1F, 0x4E, 0x79)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
HEADER_FILL = "1F2A44"
K = "微软雅黑"


def set_run(run, size=11, bold=False, color=DARK, font=K):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    run.font.name = font
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), font)
    rFonts.set(qn("w:hAnsi"), font)
    rFonts.set(qn("w:eastAsia"), font)


def para(doc, text, size=11, bold=False, color=DARK, align=None,
         space_after=6, space_before=0, font=K):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        set_run(r, size=size, bold=bold or (i % 2 == 1), color=color, font=font)
    return p


def heading1(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(16)
    pf.space_after = Pt(8)
    r = p.add_run(text)
    set_run(r, size=16, bold=True, color=RED)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "B01E2E")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def heading2(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(10)
    pf.space_after = Pt(4)
    r = p.add_run(text)
    set_run(r, size=13, bold=True, color=BLUE)
    return p


def table(doc, headers, rows, widths=None, font_size=9.5):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_run(r, size=font_size, bold=True, color=WHITE)
        tcPr = hdr[i]._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), HEADER_FILL)
        tcPr.append(shd)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(val))
            set_run(r, size=font_size, color=DARK)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    return t


def quote_box(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(8)
    pf.left_indent = Pt(12)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), "B01E2E")
    pBdr.append(left)
    pPr.append(pBdr)
    r = p.add_run(text)
    set_run(r, size=11, color=DARK, font="楷体")
    return p


def bullet(doc, text, size=10.5):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.7)
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        set_run(r, size=size, bold=(i % 2 == 1), color=DARK)
    return p


# ================================================================= 文档主体
doc = Document()
sec = doc.sections[0]
sec.page_width = Cm(21)
sec.page_height = Cm(29.7)
sec.top_margin = Cm(2.4)
sec.bottom_margin = Cm(2.2)
sec.left_margin = Cm(2.5)
sec.right_margin = Cm(2.5)

st = doc.styles["Normal"]
st.font.name = K
st.font.size = Pt(11)
st.element.rPr.rFonts.set(qn("w:eastAsia"), K)


# ---- 封面 ----
para(doc, "常州市2025年政府工作报告", size=24, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=60, space_after=4)
para(doc, "深度研究与\u201c容易被忽视的细节\u201d分析", size=16, bold=True, color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
para(doc, "从\u201c智能制造、新能源之都（动力电池/光伏）、科教城、新质生产力\u201d重新理解常州", size=12, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)
para(doc, "\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501", color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
para(doc, "研究对象：2025年常州市政府工作报告及同期财政、国民经济和社会发展统计资料、2026年官方复盘", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
para(doc, "标定日期：2026-08-14", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
doc.add_page_break()

# ---- 目录 ----
catalog = [
    "执行摘要",
    "一、研究口径与阅读方法",
    "二、先看常州的特殊底盘：智能制造、新能源之都、动力电池、科教城与长三角区位",
    "三、最关键的宏观错位：GDP 11158.7亿/5.2%，工业缓、外贸逆势大涨，但固投地产大幅下行",
    "四、容易被忽视、但信号很重的细节（15条）",
    "五、2025年GDP目标 vs 实际：对照表",
    "六、2025年增长是谁撑起来的？（结构归因）",
    "七、预算与财政的\u201c含金量\u201d",
    "八、民生底账：人口、收入与城乡",
    "九、城镇与农村：格局与均衡",
    "十、人口流入与流出",
    "十一、物价与货币环境",
    "十二、区域一体化：常州在\u201c长三角+苏南+上海都市圈\u201d里的位置",
    "十三、未来5—10年最值得观察的五条主线",
    "十四、最终结论：常州在\u201c新能源汽车+智能制造+二代\u201d里的增长逻辑",
    "附录A：主要资料来源与核验路径",
    "附录B：建议建立的年度跟踪仪表盘",
]
para(doc, "目　录", size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=10, space_after=10)
for item in catalog:
    para(doc, item, size=12, space_after=4)
doc.add_page_break()

# ---- 执行摘要 ----
heading1(doc, "执行摘要")
para(doc, "**核心判断**　2025年常州最显著的是\u201cGDP 11158.7亿元、增长5.2%、人均破20万、江苏第5座万亿之城\u201d、\u201c规上工业+4.6%（高技术制造+12.3%、动力储能电池+14.1%）\u201d、\u201c进出口3728.1亿/+13.9%全省第一、实际使用外资13.96亿美元逆势\u201d、\u201c但固投-19.0%、房地产-41.1%、CPI-0.2%\u201d、\u201c财收715.5亿/+2.5%、税收占比85.5%全省第一\u201d。这说明常州在\u201c智能制造+新能源\u201d的转型中，**外贸强、产业优，但投资地产与消费失衡**。")
para(doc, "把2025年目标（GDP+5%以上/社零+5%左右）、2025年统计（GDP+5.2%达成、规上+4.6%、固投-19.0%、社零+2.6%、财收+2.5%）、趋势一起看，常州是\u201c存量升级+增量新能源\u201d路径：**整车/动力电池、电子、光伏、高端装备、科教城**是支柱；2025年总量11159亿居江苏第5。")
para(doc, "最容易记住的一句话：**常州是\u201c江苏第5座万亿之城、新能源之都（理想/比亚迪/新能源电池）\u201d，靠\u201c新能源汽车+智能制造+外贸\u201d增长。**观察常州，与其只看\u201cGDP 11158亿\u201d，不如看\u201c新能源整车80万辆、动力储能电池+19.1%、进出口+13.9%、高技术制造+12.3%\u201d。")

# ---- 一、研究口径与阅读方法 ----
heading1(doc, "一、研究口径与阅读方法")
para(doc, "本报告以\u201c常州市政府工作报告（2025年）\u201d为起点，把\u201c2025年GDP目标（5%以上）\u201d与\u201c官方2025年GDP（11158.7亿元/+5.2%）\u201d并置对照，并用\u201c2025年常州市统计公报\u201d和\u201c2026年政府工作报告复盘\u201d作为横向核验。")
para(doc, "口径提示：\u201cGDP、规上工业增加值增速\u201d按可比价（实际增速）；\u201c投资、消费、进出口、财政收入\u201d按现价（名义增速）。涉及\u201c常住人口\u201d用统计公报常住口径（约541.5万），城镇化率用官方公布值（约79.9%）。")
para(doc, "指标体系（与研究口径一致）：五大象限——**总量与增速（GDP）、产业动能（整车/电池/电子/光伏）、外贸、财政质量、民生与人口**。")
para(doc, "特别提示（不吃老本）：常州在2024年突破万亿（10789.7亿、+6.8%左右），2025年放缓到+5.2%，但**第二产业+3.3%偏缓、第三产业+7.0%快**，\u201c投资-19%但外贸+13.9%\u201d——常州不是\u201c只有电池\u201d，要看它\u201c加工贸易+高贸新+智能装备\u201d的复合底色。")
# ---- 二、先看常州的特殊底盘 ----
heading1(doc, "二、先看常州的特殊底盘：智能制造、新能源之都、动力电池、科教城与长三角区位")
para(doc, "常州地处江苏南部、长三角核心区，是**江苏第5座万亿之城、全国制造业重镇、新能源之都**。2025年GDP 11158.7亿元、常住人口约541.5万，人均GDP破20万，城镇化率约79.9%。")
para(doc, "五个底盘名词，先立框架：")
bullet(doc, "**智能制造/装备**　高端装备、工业机器人、智能装备制造全国领先，国家制造业单项冠军6家、\u201c灯塔工厂\u201d2家、智能工厂6家。")
bullet(doc, "**新能源之都（整车+电池）**　2025年新能源整车产量超80万辆（理想/比亚迪等），动力及储能电池开票销售+19.1%，产业产值向万亿迈进。")
bullet(doc, "**动力电池/储能**　动力及储能电池产值+14.1%，与宁德/比亚迪在溧阳/金坛布局，全产业链生态。")
bullet(doc, "**科教城**　常州科教城、江苏理工等，全社会研发经费3.62%，\u201c专精特新\u201d小巨人87家、高新技术产值占规上工业60%以上。")
bullet(doc, "**外贸强市**　进出口3728.1亿/+13.9%全省第一；跨境电商+53.4%、\u201c新三样\u201d出口翻番。")
para(doc, "这五根支柱（制造+新能源+科教+外贸+区位）构成常州独特底盘：**左手智能装备（存量），右手新能源汽车（增量），底子是科教城与外贸**。理解常州，先理解\u201c江苏制造第三城\u201d的底蕴。")

# ---- 三、最关键的宏观错位 ----
heading1(doc, "三、最关键的宏观错位：GDP 11158.7亿/5.2%，工业缓、外贸逆势大涨，但固投地产大幅下行")
para(doc, "2025年常州最需要辨析的一组\u201c错位\u201d：**GDP 5.2%达成5%目标、规上工业+4.6%、进出口+13.9%强势，但固投-19.0%、房地产-41.1%、社零+2.6%、CPI-0.2%**。")
para(doc, "为什么\u201c外贸这么强、工业却偏缓、投资大跌\u201d？三个解释：")
para(doc, "**其一，工业在\u201c高基数\u201d上放缓**　规上工业+4.6%（2024年+6.8%），电子+12.3%、生物医药+5.1%高增，但动力电池价格战、纺织/机械个位——\u201c总量大、增速常态化\u201d。")
para(doc, "**其二，外贸是最大亮点**　进出口+13.9%全省第一、\u201c新三样\u201d（新能源车/锂电/光伏）出口翻番、跨境电商+53.4%——外需与高附加值对冲了内需。")
para(doc, "**其三，投资地产断崖**　固投-19.0%（工业-5.1%、服务业-32.4%）、房地产-41.1%、商品房销售承压——\u201c地产调整拖累投资\u201d是常州2025年最大痛点。")
para(doc, "小结：常州2025年是\u201c**稳总量、强外贸科技、弱投资地产**\u201d的一年：产业（新能源/电子）+外贸撑增量，但**投资、地产、部分消费**承压，\u201c外向强、内需弱\u201d是核心矛盾。")

# ---- 四、容易被忽视、但信号很重的细节（15条） ----
heading1(doc, "四、容易被忽视、但信号很重的细节（15条）")
bullet(doc, "**1.进出口3728.1亿/+13.9%，全省第一**　出口+17%、\u201c新三样\u201d翻番、跨境电商+53.4%——外贸是常州最大增量。")
bullet(doc, "**2.实际使用外资13.96亿美元、逆势增长**　在全球外资收缩背景下仍为正，新能源/制造吸引外资。")
bullet(doc, "**3.高技术制造业+12.3%（电子设备+28.4%）**　高技术/电子是工业\u201c新引擎\u201d，占规上工业60%以上产值。")
bullet(doc, "**4.新能源整车超80万辆**　理想/比亚迪在常产线满产，新能源产业向万亿迈进。")
bullet(doc, "**5.动力储能电池产值+14.1%、开票+19.1%**　宁德/比亚迪/溧阳动力电池产业集群。")
bullet(doc, "**6.税收收入+4.2%、税收占比85.5%全省第一**　财政含金量冠江苏。")
bullet(doc, "**7.社零+2.6%（2918.1亿）**　消费平稳但偏弱；通讯器材+32.9%、智能手机+40.1%、新能源汽车+17.4%——\u201c新消费\u201d亮眼。")
bullet(doc, "**8.固投-19.0%、房地产-41.1%**　投资/地产深调，是2025年最大拖累。")
bullet(doc, "**9.设备购置+7.5%、工业设备+10.3%**　\u201c设备更新\u201d投资领跑，制造业长期看好。")
bullet(doc, "**10.居民收入68295元/+4.0%、城乡比1.742**　农村+4.6%快于城镇+3.7%，全国最优水平。")
bullet(doc, "**11.常住人口约541.5万/城镇化率79.9%**　江苏人口大市、城镇化率高、科创人才导入12万+。")
bullet(doc, "**12.CPI-0.2%、但12月+0.9%**　物价温和、年底转正；PPI-3.3%（工业品价跌）需关注。")
bullet(doc, "**13.研发经费3.62%占GDP**　高于全省平均，\u201c科教城+高企2000+\u201d支撑新质生产力。")
bullet(doc, "**14.\u201c专精特新\u201d小巨人87家、单项冠军6家**　隐形冠军集群，制造\u201c质地\u201d优。")
bullet(doc, "**15.城镇新增就业超11万/引进人才超12万**　人口、人才双增长，制造业强市留人。")
# ---- 五、2025年GDP目标 vs 实际：对照表 ----
heading1(doc, "五、2025年GDP目标 vs 实际：对照表")
table(doc,
    ["指标", "2025年目标", "2025年实际", "达成评价"],
    [
        ["地区生产总值(GDP)", "增长5%以上", "11158.7亿/5.2%", "达成"],
        ["一般公共预算收入", "保持稳定增长", "715.5亿/2.5%", "达成(稳增)"],
        ["社会消费品零售总额", "增长5%左右", "2918.1亿/2.6%", "未达成，差2.4pct"],
        ["规上工业增加值", "——", "+4.6%", "放缓"],
        ["进出口总额", "稳中提质", "3728.1亿/+13.9%", "大幅超额(全省第一)"],
        ["固定资产投入", "——", "-19.0%", "大幅下行"],
        ["居民收入", "与经济增长基本同步", "68295元/+4.0%", "略低于GDP"],
    ],
)
para(doc, "注：GDP、规上工业按可比价；投资、消费、进出口、财收按现价。**GDP（5.2%）、进出口（+13.9%）、财收（+2.5%）达成**，**社零、固投偏弱**。")
para(doc, "拆读：**外贸（+13.9%全省第一）、高科技（+12.3%）、税收占比（85.5%全省第一）是亮色**，**固投（-19%）、房地产（-41.1%）、CPI（-0.2%）是短板**；\u201cGDP目标5%以上\u201d实际5.2%——\u201c稳健达标、内外分化\u201d，是\u201c万亿之城\u201d的典型样本。")

# ---- 六、2025年增长是谁撑起来的？（结构归因） ----
heading1(doc, "六、2025年增长是谁撑起来的？（结构归因）")
para(doc, "把常州GDP的5.2%拆开：三次产业分别增2.1%、3.3%、7.0%（2024年结构约1.6：46.2：52.2）。**第三产业（服务业）是主引擎（+7.0%），第二产业（工业）稳，第一产业（农业）弱**。")
para(doc, "2026年常州强调\u201c稳中求进、提质增效、做优增量\u201d，聚焦**新质生产力、新能源产业、数字经济（AI核心200亿）、人才科教**——核心是\u201c制造强市+外向型\u201d。")
para(doc, "**第二产业（制造业）**：规上工业+4.6%、高技术制造+12.3%；新能源整车80万辆、动力储能电池+14.1%、装备/电子强——\u201c存量（制造）+增量（新能源）\u201d并进。")
para(doc, "**第三产业（服务业）**：+7.0%（信息服务/软件/现代物流）；跨境电商+53.4%、数字贸易、展会（足球联赛拉动15亿）。")
para(doc, "**外贸（开放型）**：进出口+13.9%全省第一、出口+17%、实际使用外资+逆势增长（13.96亿美元）——外部需求是2025年最大对冲。")
para(doc, "一句话归因：**2025年常州增长\u201c靠外贸+服务业+高技术制造\u201d**，**投资、地产、消费偏弱**；\u201c外向强（外贸）、内生弱（固投地产）\u201d是常州核心结构性特征。")

# ---- 七、预算与财政的\u201c含金量\u201d ----
heading1(doc, "七、预算与财政的\u201c含金量\u201d")
para(doc, "2025年常州**一般公共预算收入715.5亿元（+2.5%）**，其中税收收入611.4亿元（+4.2%）、**税收占比85.5%位列全省第一**（最高质量）。财收/GDP匹配、税收质量冠江苏。")
bullet(doc, "税收结构：增值税/所得税正增长，税收+4.2%快于财收总增速——财政\u201c真金白银\u201d质量高。")
bullet(doc, "民生与产业：logcom支出超80%（民生类），研发经费370\u2018三年\u2019+；先进制造、新能源、教育投入大。")
bullet(doc, "金融支撑：存款19940.9亿（+5.0%）、贷款19566.9亿（+5.9%）；直接融资/上市4家（累计60+）、增发支持制造。")
para(doc, "**财政含金量小结**：财收+2.5%、税收占比85.5%全省第一，**\u201c税收质量冠江苏\u201d**；财政对\u201c新能源+科技城\u201d\u201c人才\u201d投入，与其\u201c万亿新城\u201d地位匹配。")

# ---- 八、民生底账：人口、收入与城乡 ----
heading1(doc, "八、民生底账：人口、收入与城乡")
para(doc, "2025年常州**居民人均可支配收入68295元（+4.0%）**，其中城镇77763元（+3.7%）、农村44647元（+4.6%），**城乡比1.742（领先全国）**。就业：城镇新增就业超11万人、引进各类人才超12万人。")
para(doc, "人口画像：**常住人口约541.5万、城镇化率79.9%**，是江苏人口大市且城镇化率高；\u201c科教城+高企\u201d带动人才导入，制造业留人。")
para(doc, "民生投入：低保标准提至1080元/月、社保参保率高、教育/养老投入；+5.0%参保率、普惠托育——民生保障扎实。")

# ---- 九、城镇与农村：格局与均衡 ----
heading1(doc, "九、城镇与农村：格局与均衡")
para(doc, "常州城镇化率约79.9%（江苏前列），城乡格局均衡；农村收入增速（+4.6%）高于城镇（+3.7%），**城乡收入比缩小（1.742）**，全国领先。")
para(doc, "农业底盘：粮食产量71.3万吨（+0.4%）、蔬菜、生猪；常州辖溧阳（天目湖/茶）、金坛（储能），\u201c农业+旅游+制造业\u201d乡镇经济发达。")
para(doc, "一句话：\u201c常州城镇强、农村收入快、城乡比全国最优\u201d，但\u201c粮食体量小、农业占GDP低\u201d，更多靠\u201c县城/产业向乡镇传导\u201d。")

# ---- 十、人口流入与流出 ----
heading1(doc, "十、人口流入与流出")
para(doc, "常州市**常住人口约541.5万、城镇化率79.9%**，是中部以\u201c外来务工+人才导入\u201d为主的净流入型城市；江苏苏南城市人口总体平稳，常州靠\u201c新能源+制造\u201d吸纳。")
para(doc, "结构观察：**人口总量稳、户籍少（常住多于户籍）**，农民工与技术工人、人才（+12万人）净流入；老龄化略低于或全省持平。")
para(doc, "2026年目标：引进人才10万以上、新增就业10万——常州把\u201c人口+人才\u201d作为\u201c万亿之城\u201d持续动能。")

# ---- 十一、物价与货币环境 ----
heading1(doc, "十一、物价与货币环境")
para(doc, "2025年常州**CPI全年累计-0.2%**（食品烟酒-1.3%、交通通信-3.5%；衣着+0.8%、居住+0.3%），**物价温和偏弱、年底12月同比转正+0.9%**；PPI-3.3%（工业品价格走弱）。")
bullet(doc, "信贷：存款19940.9亿（+5.0%）、贷款19566.9亿（+5.9%），信贷稳健充裕。")
bullet(doc, "以旧换新：通讯器材+32.9%、家具+27.5%、新能源车+17.4%——政策贴补拉动\u201c新消费\u201d。")
para(doc, "货币环境判断：**宽信用、低物价（CPI为负）**；\u201c资金充裕、价格走弱\u201d，常州通过\u201c以旧换新、刺激内需\u201d稳物价、促消费（2026目标+2%）。")

# ---- 十二、区域一体化：常州的位置 ----
heading1(doc, "十二、区域一体化：常州在\u201c长三角+苏南+上海都市圈\u201d里的位置")
para(doc, "常州地处长三角苏南板块、上海大都市圈外围，是\u201c沪宁线\u201d制造业带区段重要节点，与无锡/镇江/南京/上海联动。")
bullet(doc, "**苏南协同**　常州与无锡、镇江、苏州构成\u201c苏锡常\u201d制造带，承接上海科创外溢。")
bullet(doc, "**上海枢纽**　高铁京沪、沿江高铁+沪蓉高速，\u201c1小时都市圈\u201d，承接上海产业、制造与科创辐射。")
bullet(doc, "**科教城/园区**　常州科教城、中德创新园，产教融合；高铁、水运、空港物流。")
para(doc, "一句话：**常州在\u201c长三角+苏南城市群\u201d里，最核心的定位是\u201c智能制造+新能源之都\u201d**——苏南区位、制造业底蕴是常州的最大优势。")

# ---- 十三、未来5—10年最值得观察的五条主线 ----
heading1(doc, "十三、未来5—10年最值得观察的五条主线")
bullet(doc, "**主线一：新能源跨界（车+电池+储能）**\u201c新能源整车80万辆、储能电池+19.1%\u201d能否突破\u201c电池价格战\u201d、走向全球。")
bullet(doc, "**主线二：智能制造新质（AI+装备）**\u201cAI核心产业200亿、智能工厂\u201d能否接续\u201c制造强市\u201d。")
bullet(doc, "**主线三：外贸高附加值**\u201c新三样\u201d出口翻番、跨境电商+53.4%能否稳定\u201c外向\u201d。")
bullet(doc, "**主线四：地产与投资缺口**\u201c恒大/固投-19%\u201d能否止跌、重拾投资。")
bullet(doc, "**主线五：人口与人才**\u201c24万人才/引职\u201d、城镇化79.9%能否持续\u201c强城、聚人\u201d。")

# ---- 十四、最终结论 ----
heading1(doc, "十四、最终结论：常州在\u201c新能源汽车+智能制造+外向型\u201d里的增长逻辑")
para(doc, "把2025年闭环打穿：**常州是\u201c江苏万亿制造强市\u201d**：GDP 11158.7亿/+5.2%、规上工业+4.6%、进出口+13.9%、新能源整车80万辆、财收+2.5%（税收占85.5%）。")
para(doc, "常州不是\u201c只靠电池\u201d——它把**智能制造（存量）+新能源（增量）+外贸（外向）+科教（人才）**组合成\u201c复合新质\u201d；但投资、地产、消费内需偏弱，\u201c外向强、内生弱\u201d。")
para(doc, "一句话结论：**常州是\u201c智能制造+新能源之都\u201d的万亿强市；观察它先看\u201c新能源整车/电池、外贸、高技术制造\u201d，再看\u201c固投、地产、消费\u201d。**它是\u201c制造优、外贸强、内需弱\u201d的江苏\u201c新能源\u201d样本。")

# ---- 附录A ----
heading1(doc, "附录A：主要资料来源与核验路径")
bullet(doc, "《2025年常州市政府工作报告》（2025年1月，盛蕾作，2025年目标、2024年回顾）")
bullet(doc, "《2025年常州市国民经济运行情况》（常州市统计局，2026-01-30，2025年实际数据）")
bullet(doc, "《2026年常州市政府工作报告》（常州市政府，2026年1月，2025年复盘+2026年目标）")
bullet(doc, "常州市政府官网、常州市统计局（changzhou.gov.cn/tjj）")
bullet(doc, "聚汇数据·常州人口（常住约541.5万、城镇化率79.9%）")

# ---- 附录B ----
heading1(doc, "附录B：建议建立的年度跟踪仪表盘")
bullet(doc, "GDP总量/增速 vs 全年目标（可比价）。")
bullet(doc, "规上工业增加值及分行业（电子/新能源/电池/装备）增速。")
bullet(doc, "新能源整车产量、动力储能电池产值/开票。")
bullet(doc, "固定资产投资/工业/服务业/房地产投资增速。")
bullet(doc, "社会消费品零售总额、\u201c新三样\u201d、线上。")
bullet(doc, "外贸、\u201c新三样\u201d出口、跨境电商、实际使用外资。")
bullet(doc, "一般公共预算收入、税收/非税、税收占比。")
bullet(doc, "常住人口、城镇化率、城镇新增就业、引进人才。")
bullet(doc, "CPI/PPI、金融存贷款。")
bullet(doc, "科技城研发、高企、专精特新、灯塔工厂。")

# ---------------- 保存 ----------------
out = "/Users/x/Desktop/content-prod-lab/reports/常州市_2025年政府工作报告_深度研究_2026-08-14.docx"
doc.save(out)
print("SAVED OK: 常州市", out)
