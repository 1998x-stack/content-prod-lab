# -*- coding: utf-8 -*-
"""Build 济宁市2025年政府工作报告 深度研究 DOCX, 参照地级市系列版式。"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DARK = RGBColor(0x1F, 0x2A, 0x44)
RED = RGBColor(0xB0, 0x1E, 0x2E)
GRAY = RGBColor(0x59, 0x60, 0x69)
BLUE = RGBColor(0x1F, 0x4E, 0x79)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
HEADER_FILL = "1F2A44"
K = "微软雅黑"


def set_run(run, size=11, bold=False, color=DARK, font=K):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    run.font.name = font
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), font)
    rFonts.set(qn("w:hAnsi"), font)
    rFonts.set(qn("w:eastAsia"), font)


def para(doc, text, size=11, bold=False, color=DARK, align=None,
         space_after=6, space_before=0, font=K):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        set_run(r, size=size, bold=bold or (i % 2 == 1), color=color, font=font)
    return p


def heading1(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(16)
    pf.space_after = Pt(8)
    r = p.add_run(text)
    set_run(r, size=16, bold=True, color=RED)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "B01E2E")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def heading2(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(10)
    pf.space_after = Pt(4)
    r = p.add_run(text)
    set_run(r, size=13, bold=True, color=BLUE)
    return p


def table(doc, headers, rows, widths=None, font_size=9.5):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_run(r, size=font_size, bold=True, color=WHITE)
        tcPr = hdr[i]._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), HEADER_FILL)
        tcPr.append(shd)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(val))
            set_run(r, size=font_size, color=DARK)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    return t


def quote_box(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(8)
    pf.left_indent = Pt(12)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), "B01E2E")
    pBdr.append(left)
    pPr.append(pBdr)
    r = p.add_run(text)
    set_run(r, size=11, color=DARK, font="楷体")
    return p


def bullet(doc, text, size=10.5):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.7)
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        set_run(r, size=size, bold=(i % 2 == 1), color=DARK)
    return p


# ================================================================= 文档主体
doc = Document()
sec = doc.sections[0]
sec.page_width = Cm(21)
sec.page_height = Cm(29.7)
sec.top_margin = Cm(2.4)
sec.bottom_margin = Cm(2.2)
sec.left_margin = Cm(2.5)
sec.right_margin = Cm(2.5)

st = doc.styles["Normal"]
st.font.name = K
st.font.size = Pt(11)
st.element.rPr.rFonts.set(qn("w:eastAsia"), K)


# ---- 封面 ----
para(doc, "襄阳市2025年政府工作报告", size=24, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=60, space_after=4)
para(doc, "深度研究与\u201c容易被忽视的细节\u201d分析", size=16, bold=True, color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
para(doc, "从\u201c汽车之都、新能源与智能网联商用车、光电、高端装备、汉江枢纽\u201d重新理解襄阳", size=12, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)
para(doc, "\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501", color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
para(doc, "研究对象：2025年襄阳市政府工作报告及同期财政、国民经济和社会发展统计资料、2026年官方复盘", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
para(doc, "标定日期：2026-08-14", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
doc.add_page_break()

# ---- 目录 ----
catalog = [
    "执行摘要",
    "一、研究口径与阅读方法",
    "二、先看襄阳的特殊底盘：汽车之都、智能网联、光电、汉江枢纽、区域性中心城市",
    "三、最关键的宏观错位：GDP+2.1%显著降档、社零-6.9%，但高技术制造+29.7%、财收+6.4%",
    "四、容易被忽视、但信号很重的细节（15条）",
    "五、2025年GDP目标 vs 实际：对照表",
    "六、2025年增长是谁撑起来的？（结构归因）",
    "七、预算与财政的\u201c含金量\u201d",
    "八、民生底账：人口、收入与城乡",
    "九、城镇与农村：格局与均衡",
    "十、人口流入与流出",
    "十一、物价与货币环境",
    "十二、区域一体化：襄阳在湖北中部、汉江生态经济带、成渝鄂\u201c三圈\u201d里的位置",
    "十三、未来5—10年最值得观察的五条主线",
    "十四、最终结论：襄阳在\u201c汽车+智能网联+汉江枢纽\u201d里的增长逻辑",
    "附录A：主要资料来源与核验路径",
    "附录B：建议建立的年度跟踪仪表盘",
]
para(doc, "目　录", size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=10, space_after=10)
for item in catalog:
    para(doc, item, size=12, space_after=4)
doc.add_page_break()

# ---- 执行摘要 ----
heading1(doc, "执行摘要")
para(doc, "**核心判断**　2025年襄阳最显著的是\u201cGDP 6113.85亿/+2.1%（大幅不及6.5%、汽车-4.7%拖累）、规模以上工业+0.7%（制造业+10.2%）、固定资产投资-4.7%、房地产-21.3%\u201d、\u201c但高技术制造+29.7%、新能源汽车8万辆、地方财收+6.4%、招商引资亿元项目456个（历史最好）\u201d、\u201c社零-6.9%、进出口-13%\u201d。这说明襄阳在\u201c汽车之都+区域性中心城市\u201d中，**产业升级（高科制造）强、传统汽车/投资消费弱**。")
para(doc, "把2025年目标（GDP+6.5%、规上+8%、固投+8%、社零+7%、进出口+5%）、2025年实际（GDP+2.1%、规上+0.7%、固投-4.7%、社零-6.9%、进出口-13%）趋势看，襄阳是\u201c汽车+转型\u201d路径：**新能源与智能网联商用车、新能源新材料、高端装备、光电**是支柱；高技术制造+29.7%显亮点。")
para(doc, "最容易记住的一句话：**襄阳是\u201c湖北第二城、中国汽车城、汉江经济带中心城市\u201d，靠\u201c智能网联汽车+装备制造+光电\u201d驱动。**观察襄阳，与其只看\u201cGDP 6114亿\u201d，不如看\u201c高技术制造+29.7%、新能源汽车8万辆、全国唯一国家级车联网先导区、电子+38.8%、引入百亿项目11个\u201d。")
# ---- 一、研究口径与阅读方法 ----
heading1(doc, "一、研究口径与阅读方法")
para(doc, "本报告以\u201c襄阳市政府工作报告（2025年2月，杜海洋作）\u201d为起点，把\u201c2025年GDP目标（6.5%）\u201d与\u201c官方2025年（6113.85亿/+2.1%）\u201d并置对照，用\u201c2025年襄阳市统计公报\u201d和\u201c2026年政府工作报告\u201d横向核验。")
para(doc, "口径提示：\u201cGDP、规上工业增速\u201d按可比价（实际）；\u201c投资、消费、进出口、财政\u201d按现价（名义）。\u201c常住人口\u201d用统计公报（525.12万）、城镇化率66.21%。")
para(doc, "指标体系（与研究口径一致）：**总量与增速、产业动能（汽车/光电/高端制造）、外贸、财政、民生与人口**。")
para(doc, "特别提示（不吃老本）：襄阳2024年GDP最终核实约5985亿、2025年6113.85亿/+2.1%（增速大幅降档）；它不是\u201c只有车\u201d——**新能源智能网联商用车、光电（乔翔/冠捷）、高端装备、新材料\u201d才是新底色；高技术制造+29.7%。")
# ---- 二、先看襄城的特殊底盘 ----
heading1(doc, "二、先看襄城的特殊底盘：汽车城、智能网联、光电、汉江枢纽、区域性中心城市")
para(doc, "襄阳地处湖北省西北部、汉江中游，是**湖北省域副中心城市（湖北第二城）、中国汽车之都（东风日产襄樊基地）、汉江经济带核心城市**；以\u201c东风新能源智能越野车、全国唯一国家级车联网先导区\u201d著称。2025年GDP 6113.85亿（湖北第2、武汉之后）、常住525.12万、城镇化率66.21%。")
para(doc, "五个底盘名词，先立框架：")
bullet(doc, "**汽车/智能网联**　东风日产、风神、新能源智能越野车（8万辆）、全国唯一国家车联网先导区（智能网联商用车千亿）——\u201c汽车之都\u201d。")
bullet(doc, "**光电产业**　光电（乔翔/冠捷）、电子+38.8%、手机512万台、锂电+113.1%——\u201c光电都市\u201d。")
bullet(doc, "**高端装备/军工**　航空航天（航宇嘉泰）、轨道交通、军转民——\u201c装备制造\u201d。")
bullet(doc, "**汉江枢纽/物流**　襄阳港（全国内河主要港口）、郑渝/襄荆高铁、自贸片区（内陆第一）——\u201c汉江枢纽\u201d。")
bullet(doc, "**区域性中心城市**　湖北副中心、综合保税区、跨境电商综试区——\u201c区域性中心城市\u201d。")
para(doc, "这五根（汽车+光电+装备+汉江枢纽+副中心）构成襄阳独特底盘：**左手汽车智能网联，右手光电汉江枢纽**。理解襄阳，先理解\u201c汽车城、区域中心、汉江\u201d。")
# ---- 三、最关键的宏观错位 ----
heading1(doc, "三、最关键的宏观错位：GDP+2.1%显著降档、社零-6.9%，但高技术制造+29.7%、财收+6.4%")
para(doc, "2025年襄阳最需要辨析的一组\u201c错位\u201d：**GDP+2.1%（大幅不及6.5%）、规上工业+0.7%（汽车-4.7%）、固定资产投资-4.7%、房地产开发-21.3%、社零-6.9%、进出口-13%，但高技术制造+29.7%、电子+38.8%、地方财收+6.4%、招商亿元项目456个**。")
para(doc, "为什么\u201c汽车/投资/消费\u201d全面承压，高技术制造与财政却好？三解释：")
para(doc, "**其一，汽车产业深度调整、基盘拖累**　汽车产量-4.7%、传统燃油车承压（新能源车8万辆）、二产仅+0.2%——\u201c汽车转型期\u201d。")
para(doc, "**其二，高技术/新质制造强**　高技术制造+29.7%（电子+38.8%、锂电+113.1%）、智能网联/光电——\u201c新质亮点\u201d。")
para(doc, "**其三，投资消费外贸弱、财政招商好**　固投-4.7%、地产-21.3%、社零-6.9%、进出口-13%；但财收+6.4%、招商456项（百亿11个）——\u201c旧引擎回调、新增长蓄力\u201d。")
para(doc, "小结：襄阳2025年是\u201c**产业转档（汽车回调）、投资消费外贸弱，但高技术/财收/招商强**\u201d：智能网联/光电/装备蓄力，汽车/地产/消费弱。")

# ---- 四、容易被忽视、但需重点关注的细节（考虑局部忽视信号 —— 部署15条） ----
heading1(doc, "四、容易被忽视、但信号很重的细节（15条）")
bullet(doc, "**1.高技术制造+29.7%（电子通信+38.8%）**\u201c新质制造强。\u201d")
bullet(doc, "**2.新能源汽车8万辆（+1.1%）、锂电+113.1%**\u201c新能源车/电池。\u201d")
bullet(doc, "**3.全国唯一国家级车联网先导区（智能网联商用车千亿）**\u201c智能网联。\u201d")
bullet(doc, "**4.招商引资456个亿/3504亿、百亿11个（历史最好）**\u201c项目招商逆势。\u201d")
bullet(doc, "**5.地方一般公共预算收入+6.4%（财收+3.9%）**\u201c财政稳。\u201d")
bullet(doc, "**6.汽车产量28.4万辆/-4.7%（产能调整）**\u201c传统车拖累。\u201d")
bullet(doc, "**7.固定资产投资-4.7%、房地产-21.3%**\u201c投资地产弱。\u201d")
bullet(doc, "**8.社零-6.9%、进出口-13%（出口-11.6%）**\u201c内需外贸弱。\u201d")
bullet(doc, "**9.CPI-0.2%、居住/食品降、衣着+4.3%**\u201c低通胀。\u201d")
bullet(doc, "**10.汉江枢纽（襄阳港内河主要港口、郑渝/襄荆高铁）**\u201c枢纽+物流，交通大通道。\u201d")
bullet(doc, "**11.襄州/枣阳/宜城县域经济（900/550亿目标）**\u201c县域成增长仪。\u201d")
bullet(doc, "**12.光电/显示（乔翔/冠捷）、科技项目**\u201c新质+光电。\u201d")
bullet(doc, "**13.区域中心城市-综合保税/跨境/自贸片区（内陆第一）**\u201c开放平台。\u201d")
bullet(doc, "**14.居民收入41785元/+5.1%（农村+5.8%）**\u201c增收、城乡比缩。\u201d")
bullet(doc, "**15.常住525.12万、城镇化66.21%（湖北副中心）**\u201c人口/城市。\u201d")
# ---- 五、2025年GDP目标 vs 实际：对照表 ----
heading1(doc, "五、2025年GDP目标 vs 实际：对照表")
table(doc,
    ["指标", "2025年目标", "2025年实际", "达成评价"],
    [
        ["地区生产总值(GDP)", "增长6.5%左右", "6113.85亿/2.1%", "大幅不及"],
        ["规模以上工业", "增长8%左右", "+0.7%", "大幅不及"],
        ["固定资产投资", "增长8%左右", "-4.7%", "大幅不及"],
        ["社会消费品零售总额", "增长7%左右", "-6.9%", "负增长"],
        ["进出口总额", "增长5%左右", "-13%", "大幅负增长"],
        ["地方一般公共预算收入", "——", "303.7亿/+6.4%", "超额"],
        ["居民收入", "与经济增长同步", "41785元/+5.1%", "高于GDP"],
    ],
)
para(doc, "注：GDP、规模工业按可比价。**GDP（+2.1%）、规上（+0.7%）、固投（-4.7%）、社零（-6.9%）、进出口（-13%）大幅不及目标**；**高技术制造（+29.7%）、地方财收（+6.4%）、招商引资**是亮点。")
para(doc, "拆读：**高技术制造、智能网联/光电、财政、招商、文旅是亮色**；**汽车（-4.7%）、地产（-21.3%）、社零（-6.9%）、进出口（-13%）**是短板——\u201c产业转档、汽车回调、新增长蓄力\u201d，是\u201c副中心转型\u201d样本。")

# ---- 六、2025年增长是谁撑起来的？（结构归因） ----
heading1(doc, "六、2025年增长是谁撑起来的？（结构归因）")
para(doc, "把襄阳GDP的2.1%拆开：三次产业分别增3.4%、0.2%、3.6%（结构8.8：38.8：52.4）。**第三产业（服务业）+3.6%是最强拉动、第二产业（工业）仅+0.2%（汽车拖累）、第一产业（农业）稳**。")
para(doc, "2026年襄阳强调\u201c区域性中心城市、汉江枢纽\u201d，聚焦**智能网联新能源汽车（东风百亿）、光电、高端装备、汉江新区、低空经济、县域**——核心是\u201c汽车转型+新质\u201d。")
para(doc, "**第二产业（工业/制造）**：规上+0.7%（制造+10.2%、采掘-10.8%、电力-16.7%）、高技术+29.7%、汽车-4.7%——\u201c汽车回调、高技术强\u201d。")
para(doc, "**第三产业（服务业）**：+3.6%（商贸、物流/港口、文旅900亿+11.4%、会展）——\u201c服务+文旅\u201d。")
para(doc, "**第一产业（农业）**：+3.4%（粮食474.3万吨、油料、畜禽、生科）——\u201c农业稳\u201d。")
para(doc, "一句话归因：**2025年襄阳增长\u201c靠第三产业（文旅/商贸）+高技术制造\u201d**，汽车/投资/消费弱；**新质与旧基底背离**——\u201c二次曲线\u201d切换中。")

# ---- 七、预算与财政的\u201c含量\u201d ----
heading1(doc, "七、预算与财政的\u201c含金量\u201d")
para(doc, "2025年襄阳**地方一般公共预算收入303.7亿元（+6.4%）**；财政总收入501.2亿（+3.9%）；支出758.3亿（-5.5%）。")
bullet(doc, "地方财收+6.4%（超额于GDP）——\u201c财政稳、含金量好\u201d。")
bullet(doc, "支出-5.5%（土地财政/基建收窄）——\u201c民生还可\u201d。")
bullet(doc, "金融支撑：存款+9.8%（住户+532.99亿）、贷款+9.0%（企事业+406.4亿）——宽信用支持产业/民生。")
para(doc, "**财政含金量小结**：地方财收+6.4%（逆势好、含金量足）、存贷高增、财政对\u201c智能网联/汽车转型/民生\u201d投入加大。")
# ---- 八、民生底账：人口、收入与城乡 ----
heading1(doc, "八、民生底账：人口、收入与城乡")
para(doc, "2025年襄阳**全体居民人均可支配收入41785元（+5.1%）**，其中城镇50412元（+4.6%）、农村27078元（+5.8%），城乡比1.86（缩小）。消费：人均消费支出26573元（+3.2%）。就业：城镇新增就业10.77万人（超额）。")
para(doc, "人口画像：**常住525.12万、城镇化率66.21%**；湖北副中心、青年外流武汉、主城（襄城/樊城/东津）+县域。")
para(doc, "民生投入：低保/养老5.01万张、医保互认、托育4.8个/千人——民生扎实、健康城市。")

# ---- 九、城镇与农村：格局与均衡 ----
heading1(doc, "九、城镇与农村：格局与均衡")
para(doc, "襄阳城镇化率66.21%；县域经济强（襄州/枣阳目标900亿、宜城/谷城550亿、老河口500亿）；农村收入增速（+5.8%）>城镇（+4.6%），**城乡比缩小**；汉江流域农业。")
para(doc, "农业底盘：**粮食474.3万吨、油料、畜禽（生猪678万头）、水产、水果**——\u201c江汉粮仓\u201d。")
para(doc, "一句话：\u201c襄阳是汽车+农业大市、县域经济强、区域中心\u201d。")

# ---- 十、人口流入与流出 ----
heading1(doc, "十、人口流入与流出")
para(doc, "襄阳常住525.12万（湖北副中心）、城镇化66.21%；\u201c汽车/光电/高校（6所）\u201d吸纳，但部分青年外流武汉；主城区+东津新区吸引，\u201c十五五\u201d目标常住城镇化率70%。")
para(doc, "结构观察：**户籍569.22万>常住（净流出）、老龄化**；汽车/智造吸附技术人才。")
para(doc, "2026年目标：新增就业10万、育才——襄阳靠\u201c汽车+光电+汉江枢纽\u201d聚人。")

# ---- 十一、物价与货币环境 ----
heading1(doc, "十一、物价与货币环境")
para(doc, "2025年襄阳**CPI-0.2%**（食品烟酒-1.5%、交通通信-2.2%、医疗-2.0%；衣着+4.3%）——\u201c低通胀、需求弱\u201d。")
bullet(doc, "信贷扩张：存款+9.8%、贷款+9.0%——宽信用支持产业/汽车/文旅。")
bullet(doc, "消费：社零-6.9%（负）、以旧换新58亿、文旅900亿+11.4%——\u201c商品弱、文旅强\u201d。")
para(doc, "货币环境判断：**宽信用、CPI-0.2%**；襄阳靠\u201c文旅+新质+汉江\u201d稳需求（2026 CPI 2%）。")

# ---- 十二、区域一体化：襄阳的位置 ----
heading1(doc, "十二、区域一体化：襄阳在湖北中部、汉江生态经济带、成渝鄂\u201c三圈\u201d里的位置")
para(doc, "襄阳是**湖北省域副中心城市、汉江生态经济带核心城市、中部地区重要枢纽（郑渝/襄荆高铁、襄阳港内河主要港口）**。")
bullet(doc, "**湖北中部**　武汉之后湖北第二极、武襄十城市群、承接武汉辐射。")
bullet(doc, "**汉江生态经济带**　汉江大通道、襄阳港（内河主要港口）、南水北调/生态。")
bullet(doc, "**中部/成渝联通**　郑渝高铁（襄阳-重庆/郑州）、中部崛起、长江经济带。")
para(doc, "一句话：**襄阳在\u201c湖北中部+汉江+成渝\u201d里，最核心是\u201c汽车+枢纽+区域中心\u201d**；区位、汽车、汉江是优势。")

# ---- 十三、未来5—10年最值得观察的五条主线 ----
heading1(doc, "十三、未来5—10年最值得观察的五条主线")
bullet(doc, "**主线一：新能源智能网联汽车（东风百亿）**\u201c车联网/智能越野、千亿商用车\u201d能否接传统车。")
bullet(doc, "**主线二：光电显示/电子（乔翔/冠捷）**\u201c电子+38.8%、光电\u201d新高地。")
bullet(doc, "**主线三：高端装备/低空/新材料**\u201c装备、航空航天、材料\u201d新质。")
bullet(doc, "**主线四：汉江枢纽/物流/开放**\u201c襄阳港、综保区、自贸片区、跨境电商\u201d。")
bullet(doc, "**主线五：县域/人口/消费止跌**\u201c枣阳宜城、城镇化70%、汽车消费回\u201d。")

# ---- 十四、最终结论 ----
heading1(doc, "十四、最终结论：襄阳在\u201c汽车+智能网联+光电子+汉江枢纽\u201d里的增长逻辑")
para(doc, "把2025年闭环打穿：**襄阳是\u201c湖北副中心、汽车之都、汉江枢纽\u201d**：GDP 6113.85亿/+2.1%、高技术制造+29.7%、新能源车8万辆、财收+6.4%。")
para(doc, "襄阳不是\u201c只有车\u201d——它是**智能网联汽车+光电+高端装备+汉江枢纽+文旅**的复合，靠\u201c新质+区位\u201d蓄势；但汽车/投资/消费/外贸（传统基底）回调。")
para(doc, "一句话结论：**襄阳是\u201c汽车城、区域中心、汉江要冲\u201d；观察它先看\u201c智能网联、光电、高端制造、招商、财政\u201d，再看\u201c汽车产量、地产、消费\u201d。**它是\u201c新质蓄力、旧基底重塑\u201d的中部转型样本。")

# ---- 附录A ----
heading1(doc, "附录A：主要资料来源与核验路径")
bullet(doc, "《2025年襄阳市政府工作报告》（2025年2月，杜海洋作，2025年目标、2024年回顾6000+）")
bullet(doc, "《2025年襄阳市国民经济和社会发展统计公报》（襄阳市统计局，2026-06-23，2025年实际）")
bullet(doc, "《2026年襄阳市政府工作报告》（2026年1月，杜海洋，复盘+2026目标）")
bullet(doc, "襄阳市政府/统计局（xiangyang.gov.cn）")

# ---- 附录B ----
heading1(doc, "附录B：建议建立的年度跟踪仪表盘")
bullet(doc, "GDP总量/增速 vs 全年目标。")
bullet(doc, "规上工业（高技术/汽车）增速。")
bullet(doc, "智能网联/新能源车、光电。")
bullet(doc, "固定资产/工业/房地产投资。")
bullet(doc, "招商项目、百亿项目。")
bullet(doc, "社零/文旅/进出口。")
bullet(doc, "财收/民生/汉江。")
bullet(doc, "常住/城镇化/县域。")
bullet(doc, "CPI/存贷款/交通枢纽。")

# ---------------- 保存 ----------------
out = "/Users/x/Desktop/content-prod-lab/reports/襄阳市_2025年政府工作报告_深度研究_2026-08-14.docx"
doc.save(out)
print("SAVED OK: 襄阳市", out)
