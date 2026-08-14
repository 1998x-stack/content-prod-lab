# -*- coding: utf-8 -*-
"""Build 济宁市2025年政府工作报告 深度研究 DOCX, 参照地级市系列版式。"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DARK = RGBColor(0x1F, 0x2A, 0x44)
RED = RGBColor(0xB0, 0x1E, 0x2E)
GRAY = RGBColor(0x59, 0x60, 0x69)
BLUE = RGBColor(0x1F, 0x4E, 0x79)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
HEADER_FILL = "1F2A44"
K = "微软雅黑"


def set_run(run, size=11, bold=False, color=DARK, font=K):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    run.font.name = font
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), font)
    rFonts.set(qn("w:hAnsi"), font)
    rFonts.set(qn("w:eastAsia"), font)


def para(doc, text, size=11, bold=False, color=DARK, align=None,
         space_after=6, space_before=0, font=K):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        set_run(r, size=size, bold=bold or (i % 2 == 1), color=color, font=font)
    return p


def heading1(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(16)
    pf.space_after = Pt(8)
    r = p.add_run(text)
    set_run(r, size=16, bold=True, color=RED)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "B01E2E")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def heading2(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(10)
    pf.space_after = Pt(4)
    r = p.add_run(text)
    set_run(r, size=13, bold=True, color=BLUE)
    return p


def table(doc, headers, rows, widths=None, font_size=9.5):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_run(r, size=font_size, bold=True, color=WHITE)
        tcPr = hdr[i]._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), HEADER_FILL)
        tcPr.append(shd)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(val))
            set_run(r, size=font_size, color=DARK)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    return t


def quote_box(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(8)
    pf.left_indent = Pt(12)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), "B01E2E")
    pBdr.append(left)
    pPr.append(pBdr)
    r = p.add_run(text)
    set_run(r, size=11, color=DARK, font="楷体")
    return p


def bullet(doc, text, size=10.5):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.7)
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        set_run(r, size=size, bold=(i % 2 == 1), color=DARK)
    return p


# ================================================================= 文档主体
doc = Document()
sec = doc.sections[0]
sec.page_width = Cm(21)
sec.page_height = Cm(29.7)
sec.top_margin = Cm(2.4)
sec.bottom_margin = Cm(2.2)
sec.left_margin = Cm(2.5)
sec.right_margin = Cm(2.5)

st = doc.styles["Normal"]
st.font.name = K
st.font.size = Pt(11)
st.element.rPr.rFonts.set(qn("w:eastAsia"), K)


# ---- 封面 ----
para(doc, "丽江市2025年政府工作报告", size=24, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=60, space_after=4)
para(doc, "深度研究与\u201c容易被忽视的细节\u201d分析", size=16, bold=True, color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
para(doc, "从\u201c世界文化旅游名城、玉龙雪山、丽江古城、三江并流、纳西东巴\u201d重新理解丽江", size=12, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)
para(doc, "\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501", color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
para(doc, "研究对象：2025年丽江市政府工作报告及同期财政、国民经济和社会发展统计资料、2026年官方复盘", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
para(doc, "标定日期：2026-08-14", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
doc.add_page_break()

# ---- 目录 ----
catalog = [
    "执行摘要",
    "一、研究口径与阅读方法",
    "二、先看丽江的特别底盘：文化名城、雪山、东巴纳西、绿色能源、山城",
    "三、最关键的宏观错位：规上+8.8%、GDP+4.5%、但CPI-0.5%、社零+3.7%、旅游产业化",
    "四、容易被忽视、但信号很重的细节（15条）",
    "五、2025年GDP目标 vs 实际：对照表",
    "六、2025年增长是谁撑起来的？（结构归因）",
    "七、预算与财政的\u201c含金量\u201d",
    "八、民生底账：人口、收入与城乡",
    "九、城镇与农村：格局与均衡",
    "十、人口流入与流出",
    "十一、物价与货币环境",
    "十二、区域一体化：丽江在丽江—大理度假带、滇西北、金沙江流域里的位置",
    "十三、未来5\u201310年最值得观察的五条主线",
    "十四、最终结论：丽江在\u201c旅游+绿色\u201d里的增长逻辑",
    "附录A：主要资料来源与核验路径",
    "附录B：建议建立的年度跟踪仪表盘",
]
para(doc, "目　录", size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=10, space_after=10)
for item in catalog:
    para(doc, item, size=12, space_after=4)
doc.add_page_break()

# ---- 执行摘要 ----
heading1(doc, "执行摘要")
para(doc, "**核心判断**　2025年丽江最显著的是\u201cGDP 744.27亿/+4.5%、三产占54%、规上工业+8.8%（制造+15.1%、电力+6.2%）、境外游客90万人次+44.12%、民间投资+9.5%、房地产+8.4%转正\u201d、\u201c但社零+3.7%、CPI-0.5%、旅游主导下内需偏弱\u201d。这说明丽江在\u201c世界文化旅游名城+绿色能源\u201d中，**旅游与绿色能源稳、工业制造在回升，但消费与价格偏弱**。")
para(doc, "把丽江2025年目标（GDP+4.5%以上、规上+7.5%、一产+3.5%、二产+5.7%、三产+3.8%）与实际（GDP+4.5%、规上+8.8%、一产+3.1%、二产+6.2%、三产+3.8%）对照，丽江是\u201c文旅+绿电\u201d路径：**丽江古城、玉龙雪山、束河、泸沽湖旅游；水电清洁能源**是支柱。")
para(doc, "最容易记住的一句话：**丽江是\u201c世界文化旅游名城、中国最美雪山玉龙、茶马古道、纳西东巴文化\u201d，靠\u201c旅游+绿色电力\u201d驱动。**观察丽江，与其只看\u201cGDP 744亿/+4.5%\u201d，不如看\u201c接待游客千万级、境外90万人次、规上工业+8.8%、清洁能源、民营经济占60%\u201d。")

# ---- 一、研究口径与阅读方法 ----
heading1(doc, "一、研究口径与阅读方法")
para(doc, "**本文的最终目标**，不是复述丽江市2025年政府工作报告的原文，而是把它放到\u201c目标→实际\u201d和\u201c长期底盘→年度对撞\u201d的坐标里重新读一遍，找出丽江2025年到底靠什么增长、哪里在换挡、哪里是暗线。")
para(doc, "**三组资料互证**：（1）《2025年丽江市国民经济和社会发展统计公报》与《2025年丽江经济稳中有进》（市统计局）提供2025年**实际完成**数据；（2）《丽江市2026年政府工作报告》（2026-02-09市五届人大六次会议）既复盘2025年、又给出2026年目标与主线；（3）《丽江市2025年政府工作报告》提供**目标设定**。三份材料交叉核验，避免单源失真。")
para(doc, "**核心框架**：先交代丽江的\u201c底盘\u201d（文化名城+绿色能源+雪山景区的产业与地理禀赋），再用\u201c目标vs实际\u201d对照表定位最关键的\u201c宏观错位\u201d，接着用15条细节捕捉暗线，最后从财政、民生、城乡、人口、物价五个侧面把错位坐实，落到2026年及\u201c十五五\u201d主线与结论。")
para(doc, "**口径说明**：GDP 744.27亿元+4.5%、2024年经最终核实为708.51亿+4.6%。三次产业增加值88.62/252.22/403.44、占比11.9:33.9:54.2。固投+0.6%、房地产+8.4%（转正）。社零334.05亿+3.7%。CPI累计-0.5%。常住人口城镇化率51.44%。民营经济占GDP 60%。")
para(doc, "**为什么值得单独研究丽江**：GDP+4.5%（低中速）、规上+8.8%（较好），是\u201c世界级文旅+绿色能源\u201d的旅游型城市；但社零+3.7%、CPI-0.5%、旅游对税费贡献受门票/市场整顿影响。\u201c文旅强、内需消费弱、工业制造回升\u201d的组合，是观察\u201c旅游型地级市在转型\u201d的上佳切片。")
# ---- 二、先看丽江的特别底盘 ----
heading1(doc, "二、先看丽江的特别底盘：文化名城、雪山、东巴纳西、绿色能源、山城")
para(doc, "丽江的成长逻辑，不能只看GDP数字，而要先把\u201c底盘\u201d摊开。底盘决定了几十年、上百年的禀赋，也决定了2025年增长来源和约束。")
bullet(doc, "**世界文化旅游名城**：丽江古城（世界遗产）、玉龙雪山、束河古镇、泸沽湖（摩梭）、茶马古道，是云南/中国最知名的旅游目的地之一。A级景区43个（+24个）。")
bullet(doc, "**三江并流·雪山高原生态**：地处青藏高原南缘的滇西北、横断山脉，玉龙雪山、老君山、金沙江，生态与生物多样性极丰富。")
bullet(doc, "**纳西东巴文化**：纳西族东巴文化、东巴文字（世界记忆遗产）、白沙壁画，是丽江的独特文化名片。")
bullet(doc, "**绿色能源/水电**：金沙江流域的水电、风电、光伏（新增新能源装机103.5万千瓦），清洁能源是工业与电力支柱。")
bullet(doc, "**山地特色农业**：高原特色农业（中药材+6.4%、园林水果+6.5%、花卉+6.5%），一产3.1%稳增。")
heading2(doc, "底盘的产业钟摆")
para(doc, "把底盘归纳成一句话：**丽江是\u201c旅游+绿色能源\u201d的文旅与生态城市**。文化名城定流量、玉龙雪山供名片、水电风电供能源、纳西东巴供底蕴。城市资产在\u201c旅游+绿电\u201d链条上。2025年规上+8.8%（工业制造回升）是亮点，但社零+3.7%、CPI-0.5%、旅游的财税/消费转化有待提升。")

# ---- 三、最关键的宏观错位 ----
heading1(doc, "三、最关键的宏观错位：规上+8.8%、工业回升、但CPI-0.5%、社零+3.7%、旅游内需偏弱")
para(doc, "把2025年数据摊开，丽江最醒目的不是\u201c+4.5%\u201d，而是**工业制造回升、旅游内需与价格偏弱**的三重错位。")
quote_box(doc, "宏观错位一号：**工业回升 vs 旅游消费偏弱**。规上工业+8.8%（制造+15.1%、电力+6.2%、通信电子+16.3%），制造在好转；但社零+3.7%、CPI-0.5%、旅游为主导的内需与价格偏冷——丽江经济高度依赖旅游，而旅游消费转化内需有限。")
quote_box(doc, "宏观错位二号：**投资回稳 vs 结构**。固投+0.6%（房地产转正+8.4%、民间+9.5%、产业投资+3.1%），投资在企稳但整体增速低；产业/民间投资占比约32~48%。")
quote_box(doc, "宏观错位三号：**三产主导 vs 一产弱、就业结构偏旅**。三产占54.2%、旅游接待大幅增长（境外90万人次+44.1%），但一产+3.1%、居民与社会就业偏旅游波动。")
para(doc, "三条错位加在一起，指向同一个结论：丽江2025年是\u201c旅游+绿色+工业回升\u201d的文旅型增长。工业是2025年的惊喜（+8.8%），但消费（社零+3.7%）、价格（CPI-0.5%）、财政对旅游的转化偏弱。丽江要提升增长质量，须把旅游的流量（游客人次）转化为消费收入与税基、把绿色电力转化为工业制造。")

# ---- 四、容易被忽视、但信号很重的细节（15条） ----
heading1(doc, "四、容易被忽视、但信号很重的细节（15条）")
para(doc, "下面这15条，散落在统计公报的角落里，单独看无关紧要，合在一起却拼出丽江2025年真实的结构肌理。")
bullet(doc, "**1. 规上工业+8.8%、计算机通信电子+16.3%（贡献率39.7%）**：丽江的电子/通信制造在快速发育，是工业新苗。")
bullet(doc, "**2. 制造业贡献率48.9%、电力燃气水50.2%**：制造与电力几乎各占规上工业一半，结构更均衡。")
bullet(doc, "**3. 采矿业+30.1%**：矿产资源（石灰石/砂石/建材）在涨。")
bullet(doc, "**4. 境外游客90万人次、+44.12%**：入境游放量、是丽江国际化亮点（从5万→90万）。")
bullet(doc, "**5. 房地产投资+8.4%转正**：在多数城市地产下滑中，丽江地产投资罕见转正（文旅地产）。")
bullet(doc, "**6. 民间投资+9.5%**：民间投资高于总量，文旅/民宿/民宿资本活跃。")
bullet(doc, "**7. 基础设施投资+4.2%占68.6%**：基建主导投资（交通/电力）。")
bullet(doc, "**8. 民营经济增加值占GDP 60%**：丽江经济以民宿/旅游/私企为主。")
bullet(doc, "**9. 粮食54.18万吨+1.53%**：粮食稳产。")
bullet(doc, "**10. 各项贷款破980亿**：贷款破千亿、资金对文旅/实体支撑。")
bullet(doc, "**11. 外贸进出口3636.9万美元+93.9%**：外贸放量（虽总量小）。")
bullet(doc, "**12. 城乡居民收入增速3.2%/6.4%、城乡比缩至2.53**：农村快于城镇、城乡差距缩小。")
bullet(doc, "**13. 人均GDP（约6.5万）、全省排名升2位**：人均提升。")
bullet(doc, "**14. CPI-0.5%（食品-1.3%、交通通信-2.1%）**：通缩在食品/交通、对游客/居民购买偏弱。")
bullet(doc, "**15. 城镇化率51.44%**：城镇化中等、山区人口向城/旅游区集中。")
# ---- 五、2025年GDP目标 vs 实际：对照表 ----
heading1(doc, "五、2025年GDP目标 vs 实际：对照表")
para(doc, "把丽江市2025年政府工作报告设定的目标，与实际值逐项对照：")
table(doc,
    ["指标", "2025年目标", "2025年实际", "判定"],
    [
        ["地区生产总值增速", "+4.5%以上", "744.27亿，+4.5%", "达成"],
        ["第一产业增速", "+3.5%", "+3.1%", "略低"],
        ["第二产业增速", "+5.7%", "+6.2%", "达成"],
        ["规上工业增加值增速", "+7.5%", "+8.8%", "达成，超额1.3pct"],
        ["第三产业增速", "+3.8%", "+3.8%", "达成"],
        ["固定资产投资增速", "（转正）", "+0.6%", "达成"],
        ["社会消费品零售总额", "（稳）", "+3.7%", "稳"],
    ],
    widths=[4.6, 3.0, 4.6, 3.8],
)
para(doc, "这张表透露丽江2025年的\u201c成色\u201d：**GDP、规上、二产、投资均达成**；一产略低、社零+3.7%偏弱。\u201c工业与投资达标、消费价格偏弱\u201d是对2025年最简练的总结——丽江靠工业制造（+8.8%）和固投企稳托住总量，但内需消费仍待提振。")

# ---- 六、2025年增长是谁撑起来的？（结构归因） ----
heading1(doc, "六、2025年增长是谁撑起来的？（结构归因）")
para(doc, "丽江2025年GDP+4.5%的功劳簿，大致可以拆成几块：")
bullet(doc, "**第一功：规上工业+8.8%（工业引擎）**。制造+15.1%、电力燃气水+6.2%、采矿业+30.1%、通信电子+16.3%；绿色能源（水电/光伏）与制造是工业双引擎。")
bullet(doc, "**第二功：三产+3.8%（旅游/服务）**。第三产业403.44亿、占比54.2%，旅游（境外90万人次+44%）、酒店民宿、交通在支撑。")
bullet(doc, "**第三功：投资企稳（固投+0.6%）**。房地产+8.4%转正、民间+9.5%、产业投资+3.1%，投资在文旅地产与绿电上企稳。")
bullet(doc, "**低速项：一产+3.1%、社零+3.7%、CPI-0.5%**。农业稳、但内需消费与价格偏弱。")
para(doc, "**结构归因结论**：2025年的丽江增长是\u201c规上工业+三产旅游\u201d双轮，工业最亮（+8.8%）；投资因文旅地产转正企稳；但社零+3.7%、CPI-0.5%显示旅游内需消费转化弱。丽江的增长多靠生产与投资，消费与价格待激活。")

# ---- 七、预算与财政的\u201c含金量\u201d ----
heading1(doc, "七、预算与财政的\u201c含金量\u201d")
para(doc, "看丽江财政，需正视\u201c旅游大市但财税转化有限\u201d。")
table(doc, ["财政指标", "2025年情况", "说明"],
    [
        ["一般公共预算收入", "（增长，规模较小）", "低于GDP万亿级"],
        ["税收依赖", "主要靠旅游/基建/制造业", "税基偏薄"],
        ["财政支出", "民生为主", "旅游大市保障投入"],
        ["转移支付", "云南生态/边区依赖", "财政自给弱"],
    ],
    widths=[4.6, 5.6, 4.8],
)
para(doc, "财政核心判断：**旅游富民、财政转化有限**。丽江经济以民营旅游（占60%）为主、企业规模小，税收/收入规模与GDP不匹配，财政靠一般转移支付与旅游消费（门票/酒店）支撑、自给率不高。这是旅游型城市的共性——富旅游、穷财政。")
quote_box(doc, "财政与宏观的勾连：**旅游富民不富财政**。游客人次大增（境外90万+44%）、但财政收入/人均收入并未同步高增，主因旅游收入分散于民宿/个体、且票务/税收口径有限。丽江要\u201c流量变税源\u201d，需在旅游增加值、酒店住宿、文创产业链上做增量抽取。")

# ================= 八、民生底账 =================
heading1(doc, "八、民生底账：人口、收入与城乡")
bullet(doc, "**居民收入**：城镇居民人均可支配收入+3.2%、农村+6.4%；城乡比缩至2.53（“十四五”末、从2.99缩小）。")
bullet(doc, "**就业**：城镇化率51.44%、旅游/服务业就业主导。")
bullet(doc, "**收入与消费**：社零+3.7%、CPI-0.5%、居民购买力在价格下行下尚可。")
bullet(doc, "**民生**：财政民生支出、保障旅游与基本公共服务。")
para(doc, "民生综合评价：**城镇收入低速、农村快、城乡差距缩小**；旅游就业灵活但收入波动。丽江民生受旅游景气影响大。")

# ================= 九、城镇与农村 =================
heading1(doc, "九、城镇与农村：格局与均衡")
bullet(doc, "**城镇化率51.44%**：常住约125万人、城镇约64万，中低城镇化、山区分散。")
bullet(doc, "**农村收入+6.4%、快于城镇+3.2%**：乡村依托旅游（客栈/特产/农业）增收。")
bullet(doc, "**农业**：中药材+6.4%、果品+6.5%、花卉+6.5%、粮食54.18万吨。")
bullet(doc, "**城乡差缩小**：城乡比缩至2.53（“十四五”从2.99降至2.53）。")
para(doc, "丽江城乡均衡靠旅游普惠民宿与农村经济，城乡差距缩小，是山区文旅共富样本。")

# ============ 十、人口流入与流出 ============
heading1(doc, "十、人口流入与流出")
bullet(doc, "**常住约125万**：城镇化率51.44%、云南山区市、人口规模小、相对稳定。")
bullet(doc, "**旅游人口/新丽江人**：旅游带来流动人口、外地创业者（新丽江人）、康养人群。")
bullet(doc, "**自然增长**：民族地区出生率较高、但总体低增。")
quote_box(doc, "人口：常住125万左右、旅游人口/新丽江人带来流动性，是滇西北中人口相对稳的市；靠文旅与宜居留住青壮年。")

# ============ 十一、物价与货币 ============
heading1(doc, "十一、物价与货币环境")
bullet(doc, "**CPI-0.5%**：食品-1.3%、交通通信-2.1%、医疗-3.6%，通缩偏冷、消费疲弱。")
bullet(doc, "**金融**：各项贷款破980亿，资金对文旅/基建支撑。")
para(doc, "CPI通缩+贷款稳增：反映丽江消费需求偏弱但资金充裕，文旅/基建受支持。")

# ========== 十二、区域一体化 ==========
heading1(doc, "十二、区域一体化：丽江在丽江—大理度假带、滇西北、金沙江流域里的位置")
bullet(doc, "**滇西北文旅经济带**：丽江—大理—香格里拉（迪庆）组成的滇西北旅游圈，是大香格里拉旅游环线核心。")
bullet(doc, "**金沙江流域绿色能源**：丽江在金沙江中游，水电/风光基地是西南能源廊道。")
bullet(doc, "**丽江—大理/昆明联动**：高铁连接大理、昆明，旅游客源与经济联动。")
bullet(doc, "**县域**：古城（主城/文遗）、玉龙（雪山/文旅）、永胜（芒果/农旅）、华坪（芒果）、宁蒗（泸沽湖/摩梭），构成文旅+农业骨架。")
para(doc, "丽江的区域坐标：**在\u201c大滇西文旅带+金沙江绿色能源\u201d双轮下，丽江是云南文旅门户与绿色能源节点**。若把丽江—大理—香格里拉联动做实，文旅与绿电能级更大。")

# ============ 十三、未来5-10年主线 ============
heading1(doc, "十三、未来5\u201310年最值得观察的五条主线")
bullet(doc, "**主线一：旅游从\u201c人次\u201d到\u201c收入/税源\u201d**。若把游客量（千万级）转化为旅游收入、酒店/文创/度假盈利与财税、消费，丽江富流量的成色质变。")
bullet(doc, "**主线二：绿色能源与工业制造（水电→绿色制造）**。规上+8.8%、通信电子+16.3%，若把清洁能源转化为光伏、新材料等制造，工业新引擎。")
bullet(doc, "**主线三：康养/度假/文创产业**。候鸟/康养/度假取代观光，丽江人均逗留与消费提升。")
bullet(doc, "**主线四：国际旅游市场**。境外游客90万人次（+44%），若重启国际航班/签证，国际化成长期变量。")
bullet(doc, "**主线五：生态与乡村**。三江并流生态、山地康养、乡村文旅，生态价值向经济转化。")
para(doc, "五条主线里，**最值得盯的是主线一（旅游→收入税级）与主线二（绿电→制造）**——丽江的赛道在\u201c旅游+绿色\u201d，能否把流量和电力炼成收入、制造与财税，是下一个十年的分水岭。")

# ============ 十四、最终结论 ============
heading1(doc, "十四、最终结论：丽江在\u201c旅游+绿色\u201d里的增长逻辑")
para(doc, "综合2025年的开局与结构，给丽江一个平衡的结论：")
para(doc, "**一、增长平稳、工业是惊喜。**GDP+4.5%、规上工业+8.8%（制造+15.1%、电力通信+16.3%）、房地产+8.4%转正、境外游客+44%，是旅游名城在工业与开放上难得的正增量。")
para(doc, "**二、但消费/价格/财政偏弱。**社零+3.7%、CPI-0.5%、财政自给不足，旅游大市的\u201c人次富、税源薄\u201d是核心症结。")
para(doc, "**三、绿电与文旅是变量。**金沙江水电风光伏、丽江大理香格里拉旅游圈，产业能级有待放大。")
para(doc, "**四、2026年前瞻。**在\u201c稳文旅、强工业和绿电、激活消费与税源\u201d约束下，丽江2026年边际变量是**旅游收入转化、工业绿电链条、国际游客、康养文创、财政税收**。若把旅游的流量、绿电的能量转为本地收入与制造，丽江能走出\u201c文旅富市、财政偏紧\u201d的宿命。")
para(doc, "**一句话总结：丽江2025年交出了\u201c+4.5%、规上+8.8%、境外游客+44.12%\u201d的答卷，旅游名城在工业与绿电的回升中稳住基本盘；但更该问的是——在这个以游客流量著称的城市，这座城市能否把\u201c雪山+古城+东巴\u201d的招牌，炼成\u201c文旅营收+绿色制造+财税\u201d的均衡增长，让丽江不只出名、更富民强市。**当下看，稳中有进，但消费与财政的隐忧提醒：把流量变成质量，才有可持续的未来。")

# ============ 附录A ============
heading1(doc, "附录A：主要资料来源与核验路径")
bullet(doc, "《2025年丽江市国民经济和社会发展统计公报》（丽江市统计局，2026-02）：GDP、三次产业、工业、投资、收入等。")
bullet(doc, "《2025年丽江经济稳中有进》（丽江市统计局/云南，2026-01）：GDP 744.27亿、规上+8.8%、社零+3.7%、CPI-0.5%等。")
bullet(doc, "《丽江市2026年政府工作报告》（2026-02-09聂金辉）：2025年回顾、2026年目标。")
bullet(doc, "核验方式：以统计公报/经济运行情况为准、以政府工作报告为框架，交叉比对目标-实际。")

# ============ 附录B ============
heading1(doc, "附录B：建议建立年度跟踪仪表盘")
table(doc, ["维度", "跟踪指标", "用途/预警"],
    [
        ["总量", "GDP、人均GDP", "守住+4.5%"],
        ["旅游", "游客人次、旅游收入、境外游客", "文旅景气"],
        ["工业", "规上工业、通信电子、电力", "制造回升"],
        ["消费", "社零、CPI", "内需/价格"],
        ["投资", "固投、房地产、民间投资", "投资质量"],
        ["财政", "一般公共预算收入、税收", "旅游税源化"],
        ["人口", "常住、城镇化率", "吸引流动性"],
    ],
    widths=[2.6, 6.2, 6.0],
)

# ========== 保存
doc.save("/Users/x/Desktop/content-prod-lab/reports/丽江市_2025年政府工作报告_深度研究_2026-08-14.docx")
print("SAVED OK: 丽江市 /Users/x/Desktop/content-prod-lab/reports/丽江市_2025年政府工作报告_深度研究_2026-08-14.docx")
