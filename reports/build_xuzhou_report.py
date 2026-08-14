# -*- coding: utf-8 -*-
"""Build 徐州市2025年政府工作报告 深度研究 DOCX, 参照地级市系列版式。"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DARK = RGBColor(0x1F, 0x2A, 0x44)
RED = RGBColor(0xB0, 0x1E, 0x2E)
GRAY = RGBColor(0x59, 0x60, 0x69)
BLUE = RGBColor(0x1F, 0x4E, 0x79)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
HEADER_FILL = "1F2A44"
K = "微软雅黑"


def set_run(run, size=11, bold=False, color=DARK, font=K):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    run.font.name = font
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), font)
    rFonts.set(qn("w:hAnsi"), font)
    rFonts.set(qn("w:eastAsia"), font)


def para(doc, text, size=11, bold=False, color=DARK, align=None,
         space_after=6, space_before=0, font=K):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        set_run(r, size=size, bold=bold or (i % 2 == 1), color=color, font=font)
    return p


def heading1(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(16)
    pf.space_after = Pt(8)
    r = p.add_run(text)
    set_run(r, size=16, bold=True, color=RED)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "B01E2E")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def heading2(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(10)
    pf.space_after = Pt(4)
    r = p.add_run(text)
    set_run(r, size=13, bold=True, color=BLUE)
    return p


def table(doc, headers, rows, widths=None, font_size=9.5):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_run(r, size=font_size, bold=True, color=WHITE)
        tcPr = hdr[i]._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), HEADER_FILL)
        tcPr.append(shd)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(val))
            set_run(r, size=font_size, color=DARK)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    return t


def quote_box(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(8)
    pf.left_indent = Pt(12)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), "B01E2E")
    pBdr.append(left)
    pPr.append(pBdr)
    r = p.add_run(text)
    set_run(r, size=11, color=DARK, font="楷体")
    return p


def bullet(doc, text, size=10.5):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.7)
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        set_run(r, size=size, bold=(i % 2 == 1), color=DARK)
    return p


# ================================================================= 文档主体
doc = Document()
sec = doc.sections[0]
sec.page_width = Cm(21)
sec.page_height = Cm(29.7)
sec.top_margin = Cm(2.4)
sec.bottom_margin = Cm(2.2)
sec.left_margin = Cm(2.5)
sec.right_margin = Cm(2.5)

st = doc.styles["Normal"]
st.font.name = K
st.font.size = Pt(11)
st.element.rPr.rFonts.set(qn("w:eastAsia"), K)


# ---- 封面 ----
para(doc, "徐州市2025年政府工作报告", size=24, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=60, space_after=4)
para(doc, "深度研究与\u201c容易被忽视的细节\u201d分析", size=16, bold=True, color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
para(doc, "从\u201c老工业基地、工程机械之都、淮海经济区中心、绿色低碳能源（光伏）\u201d重新理解徐州", size=12, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)
para(doc, "\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501", color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
para(doc, "研究对象：2025年徐州市政府工作报告及同期财政、国民经济和社会发展统计资料、2026年官方复盘", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
para(doc, "标定日期：2026-08-14", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
doc.add_page_break()

# ---- 目录 ----
catalog = [
    "执行摘要",
    "一、研究口径与阅读方法",
    "二、先看徐州的特殊底盘：老工业基地、工程机械、淮海经济区中心、绿色低碳能源",
    "三、最关键的宏观错位：GDP 9957.22亿/5.8%低于6%目标，工业强但固投负增、楼市外贸下行",
    "四、容易被忽视、但信号很重的细节（15条）",
    "五、2025年GDP目标 vs 实际：对照表",
    "六、2025年增长是谁撑起来的？（结构归因）",
    "七、预算与财政的\u201c含金量\u201d",
    "八、民生底账：人口、收入与城乡",
    "九、城镇与农村：格局与均衡",
    "十、人口流入与流出",
    "十一、物价与货币环境",
    "十二、区域一体化：徐州在\u201c淮海经济区+苏皖鲁豫区域性中心城市\u201d里的位置",
    "十三、未来5—10年最值得观察的五条主线",
    "十四、最终结论：徐州在\u201c工程机械+高端装备+淮海中心\u201d里的增长逻辑",
    "附录A：主要资料来源与核验路径",
    "附录B：建议建立的年度跟踪仪表盘",
]
para(doc, "目　录", size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=10, space_after=10)
for item in catalog:
    para(doc, item, size=12, space_after=4)
doc.add_page_break()

# ---- 执行摘要 ----
heading1(doc, "执行摘要")
para(doc, "**核心判断**　2025年徐州最显著的是\u201cGDP 9957.22亿元、增长5.8%但低于6%目标、冲万亿（差约43亿）、人均11.08万\u201d、\u201c规上工业+6.8%（工程机械\u2018一号产业\u2019+19%、高技术+9%）\u201d、\u201c社零+5.5%居全省第一\u201d、\u201c但固投-11.9%、房地产-21.7%、进出口-2.5%\u201d、\u201c财收575.33亿/+2.7%、常住896.71万\u201d。这说明徐州在\u201c老工业基地转型+工程机械+淮海中心\u201d中，**工业、消费强但投资出口地产偏弱**。")
para(doc, "把2025年目标（GDP+6%左右/固投+7%/社零+6.5%/进出口+6%以上）、2025年统计（GDP+5.8%低于目标、规上+6.8%、固投-11.9%、社零+5.5%、财收+2.7%）、趋势一起看，徐州是\u201c淮海经济区中心+工程机械\u201d路径：**工程机械（徐工世界级）、绿色低碳能源（光伏/储能）、半导体、装备制造、淮海商贸**是支柱；2025年总量9957亿居江苏第8，冲刺万亿。")
para(doc, "最容易记住的一句话：**徐州是\u201c工程机械之都（徐工）、淮海经济区中心、老工业城市转型样板\u201d，靠\u201c工程机械+绿色能源+区域中心\u201d增长。**观察徐州，与其只看\u201cGDP 9957亿\u201d，不如看\u201c工程机械+19%、高技术制造+9%、\u2019343\u2019集群8000亿、徐州港集箱+20.4%\u201d。")

# ---- 一、研究口径与阅读方法 ----
heading1(doc, "一、研究口径与阅读方法")
para(doc, "本报告以\u201c徐州市政府工作报告（2025年）\u201d为起点，把\u201c2025年GDP目标（6%左右）\u201d与\u201c官方2025年GDP（9957.22亿元/+5.8%）\u201d并置对照，并用\u201c2025年徐州市统计公报\u201d和\u201c2026年政府工作报告复盘\u201d作为横向核验。")
para(doc, "口径提示：\u201cGDP、规上工业增加值增速\u201d按可比价（实际增速）；\u201c投资、消费、进出口、财政收入\u201d按现价（名义增速）。涉及\u201c常住人口\u201d用统计公报常住口径（896.71万），城镇化率用官方公布值（69.08%）。")
para(doc, "指标体系（与研究口径一致）：五大象限——**总量与增速（GDP）、产业动能（工程机械/绿色能源/半导体）、贸销、财政质量、民生与淮海中心**。")
para(doc, "特别提示（不吃老本）：徐州2024年GDP 9537.1亿/+6.4%，2025年+5.8%略放缓且跌破6%目标，\u201c距万亿仅差约42亿\u201d；同时它是**人口大市（897万）**的非沿海万亿冲刺城市——真正要看的不是\u201cGDP数字\u201d，而是**工程机械全球竞争力、绿色能源、区域中心能级、从煤矿到智造的转型成色**。")
# ---- 二、先看徐州的特殊底盘 ----
heading1(doc, "二、先看徐州的特殊底盘：老工业基地、工程机械、淮海经济区中心、绿色低碳能源")
para(doc, "徐州地处苏鲁豫皖四省交界、淮海经济区中心，是**工程机械之都（徐工）、江苏省老工业城市转型典型、铁公水航综合性枢纽**。2025年GDP 9957.22亿、常住人口896.71万（江苏人口第3），人均11.08万。")
para(doc, "五个底盘名词，先立框架：")
bullet(doc, "**工程机械之都**　徐工集团世界级，工程机械\u201c一号产业\u201d2025年增长19%，起重机/挖掘机/装载机全球领先。")
bullet(doc, "**老工业基地转型**　从\u201c煤炭之城（煤矿开采）\u201d转型\u201c智造强市\u201d；新能源（光伏/储能）装机占比47.4%、绿色低碳能源崛起。")
bullet(doc, "**淮海经济区中心**　“四中心一枢纽”（区域经济/科创/消费/开放），“343”创新集群总规模突破8000亿。")
bullet(doc, "**半导体/高端制造**　半导体、电子、智能家电集聚；高技术制造+9%，徐州作为内陆\u201c制造洼地\u201d承接东部产业。")
bullet(doc, "**综合枢纽**　徐州观音机场旅客408万、港口吞吐量6948万吨/集箱34.5万TEU+20.4%、\u2018中欧班列集结中心\u2019、高铁枢纽。")
para(doc, "这三根支柱（工程机械+老工业转型+淮海中心）构成徐州独特底盘：**左手工程机械（徐工全球），右手绿色能源（光伏/储能）+淮海辐射**。理解徐州，先理解它\u201c四省交界、交通枢纽、制造民企\u201d的底色。")

# ---- 三、最关键的宏观错位 ----
heading1(doc, "三、最关键的宏观错位：GDP 9957.22亿/5.8%低于6%目标，工业强但固投负增、楼市外贸下行")
para(doc, "2025年徐州最需要辨析的一组\u201c错位\u201d：**GDP 5.8%未达6%目标、规上工业+6.8%、高技术+9%、社零+5.5%全省第一，但固投-11.9%、房地产-21.7%、进出口-2.5%**。")
para(doc, "为什么\u201c工业消费强、经济却在放缓\u201d？三个解释：")
para(doc, "**其一，投资断崖**　固投-11.9%（工业-11.4%、地产-21.7%、民间-10.5%）；地产/工业投资双降，是GDP冲万亿\u201c临门一脚\u201d的主要拖累。")
para(doc, "**其二，外贸走弱**　进出口1121.21亿/-2.5%（出口-3.5%），对美-41.6%、对欧盟-20.9%——贸易摩擦与外部需求压制外向。")
para(doc, "**其三，工业强但基建地产弱**　工程机械+19%、装备/半导体强；但\u201c固投、地产、部分传统制造\u201d承压，\u201c强制造、弱投资\u201d并存。")
para(doc, "小结：徐州2025年是\u201c**稳工业、强消费、弱投资地产外贸**\u201d的一年：工程机械、绿色能源、淮海商贸撑增量，但**投资、地产、出口**拖累，\u201c距万亿差临门一脚\u201d。")

# ---- 四、容易被忽视、但信号很重的细节（15条） ----
heading1(doc, "四、容易被忽视、但信号很重的细节（15条）")
bullet(doc, "**1.工程机械\u201c一号产业\u201d增长+19%**　徐工世界级、挖掘机/起重机全球领先，装备制造强。")
bullet(doc, "**2.高技术制造业+9%、\u2019343\u2019集群总规模8000亿**　半导体/智能材料/绿色能源撑\u201c第二曲线\u201d。")
bullet(doc, "**3.可再生能源发电装机占比47.4%**　光伏/储能为代表的绿色低碳能源快速崛起。")
bullet(doc, "**4.社零+5.5%居全省第一**　消费内需强（中国中部），以旧换新带动210亿。")
bullet(doc, "**5.服务业营业收入+14.9%、徐州港集箱+20.4%**　现代物流、商贸崛起（淮海中心\u201c四中心\u201d）。")
bullet(doc, "**6.民间投资占比81%**　民营企业是绝对主力（活力强、占比高）。")
bullet(doc, "**7.固投-11.9%、地产-21.7%、第一产业投资-27.4%**　投资/地产深调，是2025年最大拖累。")
bullet(doc, "**8.进出口-2.5%、出口-3.5%**　对美-41.6%、对欧盟-20.9%；对\u201c一带一路\u201d+6.6%（占68%）优于整体。")
bullet(doc, "**9.财收575.33亿/+2.7%、税收+2.1%**　个税+20.8%、增值税+8.6%，财政稳、收入质量尚可。")
bullet(doc, "**10.居民收入42230元/+4.9%、城乡比1.60**　全国领先；农村+5.6%快于城镇+4.3%。")
bullet(doc, "**11.常住896.71万/城镇化69.08%**　江苏人口大市、户籍超1000万、城镇化提升+0.68pct。")
bullet(doc, "**12.CPI+0.1%（四升四降）**　物价温和、居住/医疗保健上涨，食品烟酒/交通-。")
bullet(doc, "**13.高技术制造/半导体**　国家重要半导体材料基地、智能装备。")
bullet(doc, "**14.旅游9406.78万人次/988亿**　淮海旅游客源、旅游收入+10.3%。")
bullet(doc, "**15.中欧班列/跨境枢纽**\u201c连徐淮\u201d国家综合货运枢纽、徐州空港跨境监管中心——内陆开放纵深。")
# ---- 五、2025年GDP目标 vs 实际：对照表 ----
heading1(doc, "五、2025年GDP目标 vs 实际：对照表")
table(doc,
    ["指标", "2025年目标", "2025年实际", "达成评价"],
    [
        ["地区生产总值(GDP)", "增长6%左右", "9957.22亿/5.8%", "未达成，差0.2pct"],
        ["固定资产投资", "增长7%左右", "-11.9%", "大幅未达成"],
        ["社会消费品零售总额", "增长6.5%左右", "+5.5%(全省第一)", "未达成，差1pct"],
        ["进出口总额", "增长6%以上", "1121.21亿/-2.5%", "未达成"],
        ["规上工业增加值", "——", "+6.8%", "工业强"],
        ["一般公共预算收入", "——", "575.33亿/+2.7%", "稳增"],
        ["居民收入", "与经济增长基本同步", "42230元/+4.9%", "总体同步"],
    ],
)
para(doc, "注：GDP、规上工业按可比价；投资、消费、进出口、财收按现价。**GDP（5.8%）略低于6%目标、固投（-11.9%）、进出口（-2.5%）未达**；**规上工业（+6.8%）、社零（全省第一）、财收（+2.7%）稳健**。")
para(doc, "拆读：**工程机械（+19%）、高技术（+9%）、社零（+5.5%全省第一）是亮色**，**固投（-11.9%）、地产（-21.7%）、外贸（-2.5%）是短板**；\u201cGDP目标6%\u201d实际5.8%、\u201c距万亿差约43亿\u201d——\u201c临门一脚、稳健收官\u201d，是\u201c淮海中心\u201d冲刺万亿的样本。")

# ---- 六、2025年增长是谁撑起来的？（结构归因） ----
heading1(doc, "六、2025年增长是谁撑起来的？（结构归因）")
para(doc, "把徐州GDP的5.8%拆开：三次产业分别增3.0%、3.5%、8.0%（结构8.0：36.7：55.3）。**第三产业（服务业）是主引擎（+8.0%），第二产业（工业）是压舱石，第一产业（农业）稳**。")
para(doc, "2026年徐州强调\u201c淮海经济区中心城市\u201d建设、\u201c四中心一枢纽\u201d，聚焦**工程机械、绿色低碳能源、半导体、商贸物流、区域中心**——核心是\u201c智造强市+区域辐射\u201d。")
para(doc, "**第二产业（工业）**：规上工业+6.8%（重工业+7.7%）、工程机械+19%、高技术制造+9%、半导体/新能源装机47.4%——\u201c一号产业\u201d+绿色能源双轮。")
para(doc, "**第三产业（服务业）**：+8.0%（服务业营收+14.9%）；物流（徐州港集箱+20.4%）、商贸、文旅（旅游+10.3%）、区域消费中心。")
para(doc, "**外贸（开放型）**：进出口-2.5%转弱，但\u201c一带一路\u201d出口+6.6%（占68%）、加工贸易+66%——\u201c内需补外需\u201d。")
para(doc, "一句话归因：**2025年徐州增长\u201c靠服务业+工程机械+消费\u201d**，投资、地产、外贸转弱；\u201c强制造、强消费、弱投资外贸\u201d是徐州核心特征。")

# ---- 七、预算与财政的\u201c含金量\u201d ----
heading1(doc, "七、预算与财政的\u201c含金量\u201d")
para(doc, "2025年徐州**一般公共预算收入575.33亿元（+2.7%）**，其中税收416.58亿元（+2.1%），税收占比约72%；支出1053.5亿元（+0.1%）。")
bullet(doc, "税收结构：增值税+8.6%、个税+20.8%、所得税+0.4%——主体税种回正、个税大涨。")
bullet(doc, "支出民生：社保就业支出+12.8%、卫生健康+8.9%、教育+4.4%——民生优先。")
bullet(doc, "金融支撑：存款13023.97亿（+8.9%）、贷款13515.18亿（+11.6%）；上市公司17家、市值+28.8%；制造业贷款+13.8%。")
para(doc, "**财政含金量小结**：财收+2.7%与GDP匹配、税收质量尚可；\u201c稳收入、强民生、宽信贷\u201d；财政对工程机械、绿色能源、中心城市建设投入在加大。")

# ---- 八、民生底账：人口、收入与城乡 ----
heading1(doc, "八、民生底账：人口、收入与城乡")
para(doc, "2025年徐州**居民人均可支配收入42230元（+4.9%）**，其中城镇48978元（+4.3%）、农村30639元（+5.6%），**城乡比1.60（全国领先）**。就业：城镇新增就业9.47万人。")
para(doc, "人口画像：**常住人口896.71万、城镇化率69.08%（+0.68pct）**，是江苏人口第3大市（户籍超1000万）；老龄化+教育大市，高校在校25.65万。")
para(doc, "民生投入：低保标准提至813元/月、养老床位3.64万张、城中村改造2886万㎡——民生保底扎实。")

# ---- 九、城镇与农村：格局与均衡 ----
heading1(doc, "九、城镇与农村：格局与均衡")
para(doc, "徐州常住城镇化率69.08%（江苏地市平均偏下但快速提升），城乡格局相对均衡；农村收入增速（+5.6%）高于城镇（+4.3%），**城乡比缩小（1.60）**，全国领先。")
para(doc, "农业底盘：**粮食总产量504.15万吨（连续7年百亿斤以上）**、蔬菜1523.12万吨、猪出栏400.92万头、高标准农田520千公顷——正是淮海\u201c大粮仓\u201d。")
para(doc, "一句话：\u201c徐州农业粮仓大、农村收入快、城乡比优秀\u201d，但\u201c农业人口多、要从县域/乡镇转城镇化\u201d。")

# ---- 十、人口流入与流出 ----
heading1(doc, "十、人口流入与流出")
para(doc, "徐州常住896.71万、城镇化率69.08%（户籍人口超1050万），是\u201c户籍多于常住\u201d的净流出型人口大市（外出务工），但**作为淮海中心正反向吸引周边县市人口就业**。")
para(doc, "结构观察：**城镇化率仍低于江苏平均（约75%）**，中心城区集聚力在增强（轨道交通\u201c四线成网\u201d）；常住出生率4.72‰、死亡率8.22‰（自然增长负）。")
para(doc, "2026年目标：城镇新增就业6万以上、引进高校毕业生5万——徐州以\u201c中心+产业\u201d把\u201c返乡、周边\u201d人口接住。")

# ---- 十一、物价与货币环境 ----
heading1(doc, "十一、物价与货币环境")
para(doc, "2025年徐州**CPI同比+0.1%**（衣着+2.0%、居住+0.9%、医疗保健+6.1%、其他+9.6%；食品烟酒-1.9%、交通通信-3.3%）——整体温和、服务价格上涨。")
bullet(doc, "信贷扩张：贷款+11.6%（制造业贷款+13.8%）、存款+8.9%（住户+11%），充裕支撑制造/基建。")
bullet(doc, "工业品价格：PPI受工程机械/能源价影响走弱（公报提\u201c出厂价、购进\u201d均有降幅）。")
para(doc, "货币环境判断：**宽信用、CPI基本平稳**，\u201c资金充裕、价格温和\u201d；徐州依托工程机械、绿色能源、区域消费，2026年扩内需（社零目标6%）。")

# ---- 十二、区域一体化：徐州的位置 ----
heading1(doc, "十二、区域一体化：徐州在\u201c淮海经济区+苏皖鲁豫区域性中心城市\u201d里的位置")
para(doc, "徐州是**淮海经济区中心城市、苏皖鲁豫四省交界区域性中心城市**，\u201c四中心一枢纽\u201d（区域经济/科创/消费/开放中心+综合枢纽）。")
bullet(doc, "**淮海协同**　苏皖鲁豫四省交界地区协同推进（中央区域办定位），\u201c连徐淮\u201d国家综合货运枢纽补链强链。")
bullet(doc, "**交通枢纽**　观音机场旅客408万/+19%、港口吞吐量6949万吨/集箱+20.4%、中欧班列集结中心、高铁\u2014铁公水航俱全。")
bullet(doc, "**区域消费/经济中心**　工业/服务业\u201c两个1.3万亿\u201d目标、社零5800亿目标、区域中心消费、商贸辐射。")
para(doc, "一句话：**徐州在\u201c淮海经济区+苏皖鲁豫\u201d里，最核心的定位是\u201c国家中心城市（区域经济/消费/科创枢纽）\u201d**——四省交界、综合枢纽是大前提。")

# ---- 十三、未来5—10年最值得观察的五条主线 ----
heading1(doc, "十三、未来5—10年最值得观察的五条主线")
bullet(doc, "**主线一：工程机械国际化**\u201c徐工+工程机械+19%\u201d能否在全球/智能化（无人港、新能源装备）上持续领先。")
bullet(doc, "**主线二：绿色低碳能源（光伏/储能）**\u201c可再生能源装机47.4%\u201d能否形成\u201c光伏+储能+电池\u201d新增长极。")
bullet(doc, "**主线三：半导体/高端制造**\u201c高技术+9%、半导体材料\u201d承接东部产业能否成\u201c内陆智造\u201d。")
bullet(doc, "**主线四：区域中心能级**\u201c四中心一枢纽、两个1.3万亿\u201d能否做实淮海中心城市。")
bullet(doc, "**主线五：人口与城市**\u201c897万人口、城镇化69%\u201d能否靠中心+产业逆转\u201c外出回流\u201d。")

# ---- 十四、最终结论 ----
heading1(doc, "十四、最终结论：徐州在\u201c工程机械+高端装备+淮海中心\u201d里的增长逻辑")
para(doc, "把2025年闭环打穿：**徐州是\u201c淮海经济区中心、工程机械之都\u201d**：GDP 9957.22亿/+5.8%、规上工业+6.8%、社零+5.5%全省第一、工程机械+19%、财收+2.7%。")
para(doc, "徐州不是\u201c只有煤矿\u201d——它已成为\u201c工程机械全球名企（徐工）+绿色能源+淮海消费中心\u201d的复合经济；但投资、地产、外贸偏弱，\u201c临门一脚冲万亿\u201d。")
para(doc, "一句话结论：**徐州是\u201c淮海经济区中心、工程机械之都\u201d；观察它先看\u201c工程机械、绿色能源、淮海中心、消费\u201d，再看\u201c固投、地产、外贸\u201d。**它是\u201c制造业强、消费强、投资弱\u201d的\u201c淮海中心冲刺万亿\u201d样本。")

# ---- 附录A ----
heading1(doc, "附录A：主要资料来源与核验路径")
bullet(doc, "《2025年徐州市政府工作报告》（2025年1月14日，沈峻峰作，2025年目标、2024年回顾）")
bullet(doc, "《2025年徐州市国民经济和社会发展统计公报》（徐州市统计局，2026-06-17，2025年实际数据）")
bullet(doc, "《2026年徐州市政府工作报告》（徐州市政府，2026年2月，2025年复盘+2026年目标）")
bullet(doc, "徐州市政府官网、徐州市统计局（xz.gov.cn/tj）")
bullet(doc, "徐州市统计公报（2026-06发布）")

# ---- 附录B ----
heading1(doc, "附录B：建议建立的年度跟踪仪表盘")
bullet(doc, "GDP总量/增速 vs 全年目标（\u201c万亿冲刺\u201d进度）。")
bullet(doc, "规上工业增加值及分行业（工程机械/绿色能源/半导体/装备）增速。")
bullet(doc, "工程机械\u201c一号产业\u201d产值/出口、徐工订单。")
bullet(doc, "可再生能源/光伏/储能装机。")
bullet(doc, "固定资产投资/工业/服务业/房地产投资增速。")
bullet(doc, "社会消费品零售总额（\u201c四中心\u201d消费/L第一）、以旧换新。")
bullet(doc, "进出口、\u201c一带一路\u201d、中欧班列。")
bullet(doc, "一般公共预算收入、税收/非税、支出结构。")
bullet(doc, "常住人口、城镇化率、城镇新增就业、引进大学生。")
bullet(doc, "CPI、金融存贷款、制造业贷款。")

# ---------------- 保存 ----------------
out = "/Users/x/Desktop/content-prod-lab/reports/徐州市_2025年政府工作报告_深度研究_2026-08-14.docx"
doc.save(out)
print("SAVED OK: 徐州市", out)
