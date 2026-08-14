# -*- coding: utf-8 -*-
"""Build 济宁市2025年政府工作报告 深度研究 DOCX, 参照地级市系列版式。"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DARK = RGBColor(0x1F, 0x2A, 0x44)
RED = RGBColor(0xB0, 0x1E, 0x2E)
GRAY = RGBColor(0x59, 0x60, 0x69)
BLUE = RGBColor(0x1F, 0x4E, 0x79)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
HEADER_FILL = "1F2A44"
K = "微软雅黑"


def set_run(run, size=11, bold=False, color=DARK, font=K):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    run.font.name = font
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), font)
    rFonts.set(qn("w:hAnsi"), font)
    rFonts.set(qn("w:eastAsia"), font)


def para(doc, text, size=11, bold=False, color=DARK, align=None,
         space_after=6, space_before=0, font=K):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        set_run(r, size=size, bold=bold or (i % 2 == 1), color=color, font=font)
    return p


def heading1(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(16)
    pf.space_after = Pt(8)
    r = p.add_run(text)
    set_run(r, size=16, bold=True, color=RED)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "B01E2E")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def heading2(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(10)
    pf.space_after = Pt(4)
    r = p.add_run(text)
    set_run(r, size=13, bold=True, color=BLUE)
    return p


def table(doc, headers, rows, widths=None, font_size=9.5):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_run(r, size=font_size, bold=True, color=WHITE)
        tcPr = hdr[i]._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), HEADER_FILL)
        tcPr.append(shd)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(val))
            set_run(r, size=font_size, color=DARK)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    return t


def quote_box(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(8)
    pf.left_indent = Pt(12)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), "B01E2E")
    pBdr.append(left)
    pPr.append(pBdr)
    r = p.add_run(text)
    set_run(r, size=11, color=DARK, font="楷体")
    return p


def bullet(doc, text, size=10.5):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.7)
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        set_run(r, size=size, bold=(i % 2 == 1), color=DARK)
    return p


# ================================================================= 文档主体
doc = Document()
sec = doc.sections[0]
sec.page_width = Cm(21)
sec.page_height = Cm(29.7)
sec.top_margin = Cm(2.4)
sec.bottom_margin = Cm(2.2)
sec.left_margin = Cm(2.5)
sec.right_margin = Cm(2.5)

st = doc.styles["Normal"]
st.font.name = K
st.font.size = Pt(11)
st.element.rPr.rFonts.set(qn("w:eastAsia"), K)


# ---- 封面 ----
para(doc, "绍兴市2025年政府工作报告", size=24, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=60, space_after=4)
para(doc, "深度研究与\u201c容易被忽视的细节\u201d分析", size=16, bold=True, color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
para(doc, "从\u201c民营经济强市、黄酒/水乡文化、集成电路、先进高分子、数字经济\u201d重新理解绍兴", size=12, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)
para(doc, "\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501", color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
para(doc, "研究对象：2025年绍兴市政府工作报告及同期财政、国民经济和社会发展统计资料、2026年官方复盘", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
para(doc, "标定日期：2026-08-14", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
doc.add_page_break()

# ---- 目录 ----
catalog = [
    "执行摘要",
    "一、研究口径与阅读方法",
    "二、先看绍兴的特殊底盘：江南水乡、黄酒文化、制造业强市、集成电路",
    "三、最关键的宏观错位：GDP+6.5%全省第2、工业强，但地产-28.4%、进出口-6.6%",
    "四、容易被忽视、但信号很重的细节（15条）",
    "五、2025年GDP目标 vs 实际：对照表",
    "六、2025年增长是谁撑起来的？（结构归因）",
    "七、预算与财政的\u201c含金量\u201d",
    "八、民生底账：人口、收入与城乡",
    "九、城镇与农村：格局与均衡",
    "十、人口流入与流出",
    "十一、物价与货币环境",
    "十二、区域一体化：绍兴在杭州都市圈、G60科创走廊、长三角\u201c三圈\u201d里的位置",
    "十三、未来5—10年最值得观察的五条主线",
    "十四、最终结论：绍兴在\u201c民营制造+集成电路+水乡文化\u201d里的增长逻辑",
    "附录A：主要资料来源与核验路径",
    "附录B：建议建立的年度跟踪仪表盘",
]
para(doc, "目　录", size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=10, space_after=10)
for item in catalog:
    para(doc, item, size=12, space_after=4)
doc.add_page_break()

# ---- 执行摘要 ----
heading1(doc, "执行摘要")
para(doc, "**核心判断**　2025年绍兴最显著的是\u201cGDP 8932亿/+6.5%（达成6.5%目标、增速全省第2）、规上工业+9.6%（全省第2）、项目投资+8.2%、民营经济占76%GDP\u201d、\u201c但房地产开发-28.4%、进出口-6.6%、制造业内需强但地产外贸弱\u201d、\u201cCPI+0.1%\u201d。这说明绍兴在\u201c江南水乡+民营制造\u201d中，**制造业/民营/内需强，但地产外贸弱**。")
para(doc, "把2025年目标（GDP+6.5%、规上+8%、固投+3%、社零+6.5%、财收+2.8%、研发3.3%）、2025年实际（GDP+6.5%达成、规上+9.6%、社零+4.8%、财收+2.5%、进出口-6.6%、项目投资+8.2%）趋势看，绍兴是\u201c制造+民营\u201d路径：**集成电路、先进高分子（袜/印染/化工）、高端装备、黄酒\u201d是支柱；GDP冲万亿（2027年）。")
para(doc, "最容易记住的一句话：**绍兴是\u201c黄酒之乡、水乡文化、民营制造强市\u201d，靠\u201c产业+民营+数字经济\u201d增长。**观察绍兴，与其只看\u201cGDP 8932亿\u201d，不如看\u201c民营占76%GDP、集成电路900亿、先进高分子1100亿、规上+9.6%、进出口-6.6%（外贸承压）\u201d。")
# ---- 一、研究口径与阅读方法 ----
heading1(doc, "一、研究口径与阅读方法")
para(doc, "本报告以\u201c绍兴市政府工作报告（2025年2月，吴登芬作）\u201d为起点，把\u201c2025年GDP目标（6.5%）\u201d与\u201c官方2025年（8932亿/+6.5%）\u201d并置对照，用\u201c2025年绍兴市统计公报\u201d和\u201c2026年政府工作报告\u201d横向核验。")
para(doc, "口径提示：\u201cGDP、规上工业增速\u201d按可比价（实际）；\u201c投资、消费、进出口、财政\u201d按现价（名义）。\u201c常住人口\u201d用统计公报（544.3万）、城镇化率75.7%。")
para(doc, "指标体系（与研究口径一致）：**总量与增速、产业动能（制造/集成电路/民营）、外贸、财政、民生与人口**。")
para(doc, "特别提示（不吃老本）：绍兴2024年GDP~8300亿（+6.5%）、2025年8932亿/+6.5%；它不是\u201c只有黄酒/鲁迅\u201d——**集成电路、先进高分子、高端装备、民营制造、数字经济\u201d才是真正底色；民营占76% GDP。")
# ---- 二、先看绍兴的特殊底盘 ----
heading1(doc, "二、先看绍兴的特殊底盘：江南水乡、黄酒文化、制造业强市、集成电路")
para(doc, "绍兴地处浙江中北部、杭州湾南岸、环杭州湾大湾区，是**江南水乡、历史文化名城（鲁迅故里、黄酒之乡）、民营经济强市**；以\u201c中国轻纺城、中国印染、芯片（集成电路）、高端制造\u201d著称。2025年GDP 8932亿（冲万亿）、常住544.3万、城镇化率75.7%、浙江第5。")
para(doc, "五个底盘名词，先立框架：")
bullet(doc, "**民营经济**　民营占76%GDP/84%税收/89%就业；民企500强11家、上市公司85家——\u201c民营大市\u201d。")
bullet(doc, "**集成电路（绍芯）**　集成电路产业900亿、越城/柯桥芯片——\u201c绍兴芯片、数字经济\u201d。")
bullet(doc, "**先进高分子/传统制造**　先进高分子1100亿、印染/袜业/纺机、高端装备、化工——\u201c制造业强\u201d。")
bullet(doc, "**黄酒/文化产业**　绍兴黄酒、鲁迅故里、水乡古镇（安昌/仓桥）——\u201c文化+旅游\u201d。")
bullet(doc, "**水乡/生态/科创**　环杭州湾大湾区、G60科创走廊、教科人一体化试验区、低空经济——\u201c新质+生态\u201d。")
para(doc, "这五根（民营+芯片+先进制造+黄酒文化+水乡科创）构成绍兴独特底盘：**左手民营制造（集成电路），右手黄酒水乡文化**。理解绍兴，先理解\u201c民营、芯片、水乡\u201d。")
# ---- 三、最关键的宏观错位 ----
heading1(doc, "三、最关键的宏观错位：GDP+6.5%全省第2、工业强，但地产-28.4%、进出口-6.6%")
para(doc, "2025年绍兴最需要辨析的一组\u201c错位\u201d：**GDP+6.5%（全省第2）、规上工业+9.6%（全省第2）、项目投资+8.2%、民营强，但房地产开发-28.4%、进出口-6.6%（出口-4.2%）、CPI仅+0.1%**。")
para(doc, "为什么\u201c工业/投资/民营这么强\u201d，地产与外贸却弱？三解释：")
para(doc, "**其一，制造业/民营/内需强、体量大**　规上+9.6%（电子+20.9%、化工+19.7%、电气+14.5%）、项目投资+8.2%、民营占76%GDP、数字经济+6.2%——\u201c制造内需强\u201d。")
para(doc, "**其二，地产、外贸、外部弱**　房地产开发-28.4%、进口-30.2%、进出口-6.6%（出口-4.2%）——\u201c地产外贸弱\u201d。")
para(doc, "**其三，消费/物价温**　社零+4.8%（网络+8.9%）、但CPI+0.1%、汽车/楼市弱——\u201c内需修复中、量价\u201d。")
para(doc, "小结：绍兴2025年是\u201c**制造民营内需强、地产外贸弱**\u201d：集成电路/先进高分子/制造强，地产、进出口、CPI弱。")

# ---- 四、容易被忽视、但信号很重的细节（15条） ----
heading1(doc, "四、容易被忽视、但信号很重的细节（15条）")
bullet(doc, "**1.规上工业+9.6%全省第2（电子+20.9%/化工+19.7%）**\u201c绍兴制造强。\u201d")
bullet(doc, "**2.民营经济占76%GDP/84%税收/89%就业**\u201c民营绝对主力。\u201d")
bullet(doc, "**3.集成电路产业900亿、先进高分子1100亿**\u201c绍芯、先进制造。\u201d")
bullet(doc, "**4.项目投资+8.2%（扣除地产后）**\u201c工业/项目投资旺。\u201d")
bullet(doc, "**5.房地产开发685亿/-28.4%、房屋销售611万㎡/+0.1%**\u201c地产深度调整。\u201d")
bullet(doc, "**6.进出口3864亿/-6.6%（出口-4.2%/进口-30.2%）**\u201c外贸承压、高新产品出口+16.3%。\u201d")
bullet(doc, "**7.社零3255亿/+4.8%、网络零售668亿/+8.6%**\u201c消费/电商旺。\u201d")
bullet(doc, "**8.一般公共预算收入603亿/+2.5%（税收459亿/+2.4%）**\u201c财政稳增。\u201d")
bullet(doc, "**9.数字经济核心产业550亿（占6.2%）**\u201c数字经济、AI、低空。\u201d")
bullet(doc, "**10.新增专精特新小巨人15家/高企200家**\u201c新质生产力、民营创新。\u201d")
bullet(doc, "**11.出口高新+16.3%、实到外资7亿美元/+7.4%**\u201c外贸结构优、外资稳。\u201d")
bullet(doc, "**12.长三角/G60科创走廊、教科人一体化试验区**\u201c科创平台。\u201d")
bullet(doc, "**13.居民收入76732元/+5.3%、城乡比1.61**\u201c农村+5.8%>城镇+4.8%。\u201d")
bullet(doc, "**14.CPI+0.1%（教育文娱+3.1%、交通通信-3.2%）**\u201c物价低、量价。\u201d")
bullet(doc, "**15.常住544.3万/+1.4万、城镇化75.7%**\u201c人口稳、城镇化提升。\u201d")
# ---- 五、2025年GDP目标 vs 实际：对照表 ----
heading1(doc, "五、2025年GDP目标 vs 实际：对照表")
table(doc,
    ["指标", "2025年目标", "2025年实际", "达成评价"],
    [
        ["地区生产总值(GDP)", "增长6.5%左右", "8932亿/6.5%", "达成(全省第2)"],
        ["规模以上工业", "增长8%以上", "+9.6%(全省第2)", "超额"],
        ["一般公共预算收入", "增长2.8%", "603亿/+2.5%", "略低"],
        ["固定资产投资", "增长3%以上", "-3.2%(项目+8.2%)", "固投下行"],
        ["社会消费品零售总额", "增长6.5%", "3255亿/+4.8%", "差1.7pct"],
        ["进出口总额", "促稳提质", "3864亿/-6.6%", "外贸承压"],
        ["居民收入", "与经济增长同步", "76732元/+5.3%", "略低GDP"],
    ],
)
para(doc, "注：GDP、规上工业按可比价。**GDP（+6.5%）、规上（+9.6%）达成/超额**；**固投（含地产-3.2%）、社零（+4.8%）、进出口（-6.6%）不及目标**；项目投资（+8.2%）超。")
para(doc, "拆读：**民营、集成电路/先进高分子、制造、项目投资、数字经济是亮色**；**地产（-28.4%）、进出口（-6.6%）、CPI（+0.1%）、社零**是短板——\u201c制造民营强、地产外贸弱\u201d，是\u201c浙江制造强市\u201d样本。")

# ---- 六、2025年增长是谁撑起来的？（结构归因） ----
heading1(doc, "六、2025年增长是谁撑起来的？（结构归因）")
para(doc, "把绍兴GDP的6.5%拆开：三次产业分别增4.1%、6.1%、6.9%（结构3.0：47.0：50.0）。**第三产业（服务业）+6.9%最快、第二产业（制造业）+6.1%**、第一产业（农业）稳\u2014\u2014\u201c二产+三产双轮\u201d。")
para(doc, "2026年绍兴强调\u201c冲万亿（2027）、新质生产力、长三角一体化\u201d，聚焦**集成电路、先进高分子新材料、高端装备、数字经济/AI/低空、黄酒文化旅游**——核心是\u201c民营制造+水乡\u201d。")
para(doc, "**第二产业（工业）**：规上+9.6%（电子+20.9%、化工+19.7%、电气+14.5%）、集成电路900亿、先进高分子1100亿——\u201c制造业强\u201d。")
para(doc, "**第三产业（服务业）**：+6.9%（商贸、物流、网络零售、文旅）——\u201c服务业+电商\u201d。")
para(doc, "**第一产业（农业）**：+4.1%（粮食、水产、黄酒原料糯米、乡村旅游）——\u201c农业稳\u201d。")
para(doc, "一句话归因：**2025年绍兴增长\u201c靠制造业（集成电路/先进制造）+服务业（网络/文旅）+民营\u201d**，地产、外贸弱；\u201c产业+民营\u201d是核心。")

# ---- 七、预算与财政的\u201c含金量\u201d ----
heading1(doc, "七、预算与财政的\u201c含金量\u201d")
para(doc, "2025年绍兴**一般公共预算收入603亿元（+2.5%）**；税收收入459亿（+2.4%）；一般公共预算支出812亿（-2.2%）。")
bullet(doc, "税收+2.4%（近80%），税收占比高（76%）——\u201c财税实业强、含金量足\u201d。")
bullet(doc, "民营税收占比84%——\u201c税收靠民营\u201d。")
bullet(doc, "金融支撑：贷款1.79万亿（+8.4%）——信贷支持制造/民营/涉水。")
para(doc, "**财政含金量小结**：财收+2.5%（与GDP匹配）、税收占比高、民营贡献84%；财政对\u201c集成电路、民生、新质\u201d投入加大。")

# ---- 八、民生底账：人口、收入与城乡 ----
heading1(doc, "八、民生底账：人口、收入与城乡")
para(doc, "2025年绍兴**居民人均可支配收入76732元（+5.3%）**，其中城镇87605元（+4.8%）、农村54535元（+5.8%），城乡比1.61（收窄）。就业：城镇新增就业9.9万人、大学生15.1万。")
para(doc, "人口画像：**常住544.3万/+1.4万、城镇化率75.7%（+1.2pct）**；人口稳增、自然增长率-2.6‰（低生育）。")
para(doc, "民生投入：低保人均月1190元、适老化改造9000多户、养老/医疗——民生扎实（城乡融合共同富裕）。")

# ---- 九、城镇与农村：格局与均衡 ----
heading1(doc, "九、城镇与农村：格局与均衡")
para(doc, "绍兴城镇化率75.7%；县域经济极强（柯桥轻纺/上虞化工/诸暨袜业、嵊州领带、新昌轴承/低空）；农村收入增速（+5.8%）>城镇（+4.8%），**城乡比1.61缩窄**；和美乡村覆盖70%。")
para(doc, "农业底盘：**粮食、水产、黄酒原料（糯米/鉴湖）、乡村旅游、低效农户收入+10.8%**——\u201c江南鱼米之乡\u201d。")
para(doc, "一句话：\u201c绍兴是制造+县域经济强、城乡融合全国前列、水乡古镇\u201d。")

# ---- 十、人口流入与流出 ----
heading1(doc, "十、人口流入与流出")
para(doc, "绍兴常住544.3万（净增+1.4万）、城镇化75.7%；\u201c民营制造/芯片/新昌\u201d吸纳就业，部分青年流杭甬，主城（越城/柯桥）+湾区吸引。")
para(doc, "结构观察：**自然增长-2.6‰（老龄化、低生育）**；外来/大学生（15万）补充。")
para(doc, "2026年目标：新增城镇就业9.5万、引高层人才150名——绍兴靠\u201c产业+民营+湾区\u201d聚人。")

# ---- 十一、物价与货币环境 ----
heading1(doc, "十一、物价与货币环境")
para(doc, "2025年绍兴**CPI+0.1%**（教育文娱+3.1%、衣着+1.3%；食品烟酒-0.7%、交通通信-3.2%）——\u201c物价低、需求温\u201d。")
bullet(doc, "信贷扩张：贷款+8.4%——宽信用支持制造/民营。")
bullet(doc, "消费：社零+4.8%、网络零售+8.6%——消费/电商旺。")
para(doc, "货币环境判断：**宽信用、CPI+0.1%**；绍兴靠\u201c制造+消费+数字经济\u201d稳需求（2026 CPI 2%）。")

# ---- 十二、区域一体化：绍兴的位置 ----
heading1(doc, "十二、区域一体化：绍兴在杭州都市圈、G60科创走廊、长三角\u201c三圈\u201d里的位置")
para(doc, "绍兴是**杭州都市圈重要成员、G60科创走廊节点、环杭州湾大湾区核心、长三角一体化（融杭联甬接沪）重要城市**。")
bullet(doc, "**杭州都市圈/杭绍同城**　杭绍地铁/城际、G60科创走廊、承接杭甬辐射。")
bullet(doc, "**环杭州湾大湾区**　绍兴滨海新区、湾区制造业、绍兴科创走廊。")
bullet(doc, "**长三角/开放**　上海/宁波港联动、跨境电商、民营外向。")
para(doc, "一句话：**绍兴在\u201c杭绍甬+环湾区+长三角\u201d里，最核心是\u201c民营制造+科创走廊\u201d**；区位、产业、民营是最大优势。")

# ---- 十三、未来5—10年最值得观察的五条主线 ----
heading1(doc, "十三、未来5—10年最值得观察的五条主线")
bullet(doc, "**主线一：集成电路（绍芯/越城芯片）**\u201c900亿\u201d能否冲千亿/万亿、强链。")
bullet(doc, "**主线二：先进高分子/新材料**\u201c1100亿、印染化工升级\u201d。")
bullet(doc, "**主线三：民营经济/上市企业**\u201c85家上市、民企500强\u201d强链、出海。")
bullet(doc, "**主线四：数字经济/AI/低空**\u201c数字经济6.2%、AI核心350亿（2026）、低空200亿\u201d。")
bullet(doc, "**主线五：黄酒文旅/人口/冲万亿**\u201c水乡文旅、老龄化、2027万亿\u201d如何\u201c稳人口、提消费\u201d。")

# ---- 十四、最终结论 ----
heading1(doc, "十四、最终结论：绍兴在\u201c民营制造+集成电路+水乡文化\u201d里的增长逻辑")
para(doc, "把2025年闭环打穿：**绍兴是\u201c江南水乡、民营制造强市\u201d**：GDP 8932亿/+6.5%（全省第2）、规上+9.6%、民营占76%GDP、集成电路900亿、先进高分子1100亿。")
para(doc, "绍兴不是\u201c只有黄酒\u201d——它是**集成电路+先进高分子+高端装备+民营+数字经济**的复合，靠\u201c制造+民营\u201d驱动；但地产、进出口、CPI弱。")
para(doc, "一句话结论：**绍兴是\u201c黄酒水乡、民营大市、制造强\u201d；观察它先看\u201c民营、集成电路、先进高分子、数字经济、项目投资\u201d，再看\u201c地产、进出口、社零\u201d。**它是\u201c制造民营强、地产外贸弱\u201d的浙江样本。")

# ---- 附录A ----
heading1(doc, "附录A：主要资料来源与核验路径")
bullet(doc, "《2025年绍兴市政府工作报告》（2025年2月，吴登芬作，2025年目标、2024年回顾+6.5%）")
bullet(doc, "《2025年绍兴市国民经济和社会发展统计公报》（绍兴市统计局，2026-05-08，2025年实际）")
bullet(doc, "《2026年绍兴市政府工作报告》（2026年2月，复盘+2026年目标）")
bullet(doc, "绍兴市人民政府/统计局（sx.gov.cn）")

# ---- 附录B ----
heading1(doc, "附录B：建议建立的年度跟踪仪表盘")
bullet(doc, "GDP总量/增速 vs 全年目标。")
bullet(doc, "规模以上工业（集成电路/先进高分子）增速。")
bullet(doc, "民营经济、营收/税收占比。")
bullet(doc, "固定资产/工业/房地产投资。")
bullet(doc, "数字经济/AI/低空。")
bullet(doc, "社零、网络零售、文旅。")
bullet(doc, "进出口、外资、高新产品出口。")
bullet(doc, "一般公共预算/税收/民生%。")
bullet(doc, "常住/城镇化、青年、人口自然。")
bullet(doc, "CPI、存贷款、数据/湿地。")

# ---------------- 保存 ----------------
out = "/Users/x/Desktop/content-prod-lab/reports/绍兴市_2025年政府工作报告_深度研究_2026-08-14.docx"
doc.save(out)
print("SAVED OK: 绍兴市", out)
