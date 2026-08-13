# -*- coding: utf-8 -*-
"""Build 重庆市2025年政府工作报告 深度研究 DOCX, 参照北京/上海/杭州/天津系列版式。"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DARK = RGBColor(0x1F, 0x2A, 0x44)
RED = RGBColor(0xB0, 0x1E, 0x2E)
GRAY = RGBColor(0x59, 0x60, 0x69)
BLUE = RGBColor(0x1F, 0x4E, 0x79)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
HEADER_FILL = "1F2A44"
K = "微软雅黑"


def set_run(run, size=11, bold=False, color=DARK, font=K):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    run.font.name = font
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), font)
    rFonts.set(qn("w:hAnsi"), font)
    rFonts.set(qn("w:eastAsia"), font)


def para(doc, text, size=11, bold=False, color=DARK, align=None,
         space_after=6, space_before=0, font=K):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        set_run(r, size=size, bold=bold or (i % 2 == 1), color=color, font=font)
    return p


def heading1(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(16)
    pf.space_after = Pt(8)
    r = p.add_run(text)
    set_run(r, size=16, bold=True, color=RED)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "B01E2E")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def heading2(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(10)
    pf.space_after = Pt(4)
    r = p.add_run(text)
    set_run(r, size=13, bold=True, color=BLUE)
    return p


def table(doc, headers, rows, widths=None, font_size=9.5):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_run(r, size=font_size, bold=True, color=WHITE)
        tcPr = hdr[i]._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), HEADER_FILL)
        tcPr.append(shd)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(val))
            set_run(r, size=font_size, color=DARK)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    return t


def quote_box(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(8)
    pf.left_indent = Pt(12)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), "B01E2E")
    pBdr.append(left)
    pPr.append(pBdr)
    r = p.add_run(text)
    set_run(r, size=11, color=DARK, font="楷体")
    return p


def bullet(doc, text, size=10.5):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.7)
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        set_run(r, size=size, bold=(i % 2 == 1), color=DARK)
    return p


# ================================================================= 文档主体
doc = Document()
sec = doc.sections[0]
sec.page_width = Cm(21)
sec.page_height = Cm(29.7)
sec.top_margin = Cm(2.4)
sec.bottom_margin = Cm(2.2)
sec.left_margin = Cm(2.5)
sec.right_margin = Cm(2.5)

st = doc.styles["Normal"]
st.font.name = K
st.font.size = Pt(11)
st.element.rPr.rFonts.set(qn("w:eastAsia"), K)


# ---- 封面 ----
para(doc, "江苏省2025年政府工作报告", size=24, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=60, space_after=4)
para(doc, "深度研究与“容易被忽视的细节”分析", size=16, bold=True, color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
para(doc, "从“制造业大省、园区经济、外贸外资与共同富裕”重新理解江苏的增长逻辑", size=12, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)
para(doc, "━━━━━━━━━━━━━━━━", color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
para(doc, "研究对象：2025年江苏省政府工作报告及同期财政、国民经济和社会发展统计资料、2026年官方复盘", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
para(doc, "整理时间：2026年8月", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
para(doc, "分析性质：分析性研究 ｜ 非官方解读", size=11, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
doc.add_page_break()

# ---- 目录 ----
catalog = [
    "执行摘要",
    "一、研究口径与阅读方法",
    "二、先看江苏的特殊底盘：制造业强省、园区经济与万亿城市矩阵",
    "三、最关键的宏观错位：GDP破14.2万亿、先进制造强，但投资与物价偏弱",
    "四、容易被忽视、但信号很重的细节（15条）",
    "五、把所有细节连起来：江苏正在经历的“六个换挡”",
    "六、增长暗线：制造业当家、园区经济与新质生产力",
    "七、财政暗线：收入稳、税收质量高、民生支出大",
    "八、产业暗线：从“世界工厂”到“先进制造集群+智造强省”",
    "九、区域格局：苏南苏中苏北、长三角与园区城市群",
    "十、人口与城市：长三角人口重镇、老龄化",
    "十一、民营经济：占近六成、工业民企贡献大",
    "十二、事后验证：用2025年实际执行结果检验报告判断",
    "十三、未来5—10年最值得观察的五条主线",
    "十四、最终结论：江苏在“高质量发展+共同富裕”里的增长逻辑",
    "附录A：主要资料来源与核验路径",
    "附录B：建议建立的年度跟踪仪表盘",
]
para(doc, "目　录", size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=10, space_after=10)
for item in catalog:
    para(doc, item, size=12, space_after=4)
doc.add_page_break()

# ---- 执行摘要 ----
heading1(doc, "执行摘要")
para(doc, "**核心判断**　如果只看新闻标题，2025年江苏最显眼的是“GDP突破14.24万亿、增长5.3%”、“实际使用外资规模全国第一”和“新能源汽车产量157万辆、+32.4%”。但政府工作报告里真正值得深读的，是这一座“制造业强省+园区经济+外贸外资重镇”如何在固定资产投资（-12.7%）与房地产（-21.6%）深度调整的背景下，靠“先进制造、园区经济与新质生产力”稳住增长。")
para(doc, "把2025年初《政府工作报告》设定的目标、2025年《国民经济和社会发展统计公报》、以及2026年报告对2025年的复盘放在一起看，江苏呈现清晰暗线：**从“外贸/基建/传统制造”的旧依赖，向“先进制造+创新+园区经济+高质量发展”转型**。旧引擎（房地产、一般基建）在收缩；新引擎（高新技术、装备制造、新能源汽车、集成电路、民营）被要求更快补位。")
para(doc, "因此，本报告不按政府工作报告逐段复述，而采用“显性表述—同期数据—制度含义—长期影响”的方式，专门提取容易被忽视、但对判断江苏未来5—10年发展模式更有价值的信号。")
para(doc, "最容易记住的一句话：**江苏是全国“制造业+园区+高质量”的样板——既要做“挑起大梁”的强省，也要在房地产、投资调整中靠“实业+创新+公平”守住基本盘。**观察江苏，与其看“总量”，不如看“先进制造占比、园区/民营活力、外资与外贸韧性、公共财政收入与城乡差距”这五张名片。")
heading2(doc, "一页速览：2025年江苏经济的“表与里”")
table(doc,
    ["维度", "表面现象", "更深层信号"],
    [
        ["增长", "GDP 142351.5亿、+5.3%", "长三角第二、制造业强"],
        ["产业", "规上工业+6.5%、高技术+11.9%", "高新技术占52.1%、装备占56%"],
        ["外贸", "进出口5.95万亿、+6.0%（出口+8.4%）", "外资规模全国第一、机电出口强"],
        ["投资", "固定资产投资-12.7%", "房地产-21.6%、制造业-9.8%"],
        ["财政", "一般公共预算收入+2.1%（税收+3.8%）", "税收质量高、占77.4%"],
        ["消费", "社零46394.2亿、+3.3%", "新能源车、家电获提升"],
        ["人口", "常住8518万、城镇化率76.2%", "人口负增-3.6‰"],
        ["物价", "CPI-0.2%、PPI-2.6%", "工业通缩压力"],
    ],
    widths=[2.2, 5.2, 8.6])
para(doc, "", size=4, space_after=2)

# ---- 一、研究口径与阅读方法 ----
heading1(doc, "一、研究口径与阅读方法")
heading2(doc, "1.1 本报告的三份核心底稿")
bullet(doc, "**2025年《政府工作报告》**：2025年1月19日在省十四届人大三次会议上作，给出全年预期目标（GDP增长5%以上、社零增长5.5%左右、城镇新增就业120万人以上、CPI涨幅2%左右、居民收入与增长基本同步、研发强度3.35%左右等）。固定资产投资/进出口/财政收入未设具体百分数，以方向性要求表述。")
bullet(doc, "**2025年《国民经济和社会发展统计公报》**：江苏省统计局2026年3月19日发布，给出全年实际执行数与增速，是“事后验证”的锚点。")
bullet(doc, "**2026年江苏省政府工作报告及官方复盘**：对2025执行结果的追认，交叉验证“目标—执行—再评价”三段式。")
heading2(doc, "1.2 阅读方法：显性—数据—制度—长期")
para(doc, "每一章按“**显性表述 → 同期数据 → 制度含义 → 长期影响**”四层展开。其中“制度含义”回答“政府为什么这么排优先序”，“长期影响”回答“这对未来5—10年意味着什么”。")
para(doc, "关键判别：**当报告使用的“增长词”和统计公报里的实际数不一致时，数据优先。**例如2025年初定“GDP增长5%以上”，实际+5.3%达标；“社零5.5%左右”，实际+3.3%；研发投入强度3.35%，与实况接近。差异清晰反映：江苏在房地产/投资调整下仍靠“制造业+外贸+高质量”守住增长。")

# ---- 二、特殊底盘 ----
heading1(doc, "二、先看江苏的特殊底盘：制造业强省、园区经济与万亿城市矩阵")
para(doc, "在所有省份里，江苏的“底盘”独特：**经济大省+制造业强省+园区经济+万亿城市片**四合一。苏州、南京、无锡、南通等多城进入万亿GDP行列，是全国最“均衡”的发达省份之一。")
para(doc, "这决定江苏的五重身份并存：**制造引擎**（工业增加值全国前列）、**外贸外资大户**（外资规模全国第一、出口机电/集成电路强）、**园区经济高地**（苏州工业园等全国样板）、**民营与外资并存**（民营占58.5%）、**科教重镇**（南京高校资源丰富）。")
heading2(doc, "2.1 制造业为核")
para(doc, "江苏规上工业增加值+6.5%，高新技术产业占52.1%、装备制造占56%。新能源汽车产量157万辆、+32.4%，集成电路+12.3%。制造业是江苏最坚实的“底盘”。")
heading2(doc, "2.2 园区经济与开放")
para(doc, "江苏以“园区经济”闻名，苏州工业园区等是全国开放与制造样板。实际使用外资规模全国第一（余额168亿美元，占全国16.1%），机电/高技术出口强，是“外资+外贸”双强省份。")
heading2(doc, "2.3 城市矩阵与区域")
para(doc, "苏南（苏州/无锡/常州）、苏中（南通/扬州/泰州）、苏北（徐州/盐城/连云港等）差异仍存，但整体是全国城乡发展较均衡的省份之一。")

# ---- 三、宏观错位 ----
heading1(doc, "三、最关键的宏观错位：GDP破14.2万亿、先进制造强，但投资与物价偏弱")
para(doc, "把2025年江苏的宏观面放进一张表，会出现令人惊讶的“错位”：表观增长来自工业与出口，而投资的房地产/基建与前端的PPI却在收缩。这个错位，正是读懂江苏的关键。")
heading2(doc, "3.1 “强的一面”")
table(doc,
    ["指标", "2025实绩", "信号"],
    [
        ["GDP", "142351.5亿、+5.3%", "长三角第二、超5%目标"],
        ["规上工业", "+6.5%", "制造业强、行业面75%"],
        ["高技术制造", "+11.9%", "对规上贡献43%"],
        ["装备制造", "+8.8%、占56%", "高端制造主导"],
        ["进出口", "5.95万亿、+6.0%（出口+8.4%）", "机电/高技术出口强"],
        ["实际使用外资", "全国第一", "外资信心强"],
    ],
    widths=[3.2, 5.2, 6.0])
heading2(doc, "3.2 “弱的一面”")
table(doc,
    ["指标", "2025实值", "信号"],
    [
        ["固定资产投资", "-12.7%", "投资深度调整"],
        ["房地产开发投资", "-21.6%", "地产收缩"],
        ["制造业投资", "-9.8%", "制造投资也降"],
        ["一般公共预算收入", "+2.1%", "收入低增"],
        ["CPI", "-0.2%", "消费价走低"],
        ["PPI", "-2.6%", "工业通缩"],
    ],
    widths=[3.2, 5.2, 6.0])
para(doc, "**错位结论**　江苏的增长“很强、但也很矛盾”。强的部分（规上工业/高技术/装备/出口/外资）与弱的部分（固投/房地产/CPI/PPI）并存。**真正的焦点不是“有没有增长”，而是“量的质量与结构呼应”**：制造业与出口强撑起5.3%增长，但房地产与投资拖累整体。2026年稳住投资、补内需、让先进制造与民营释放更多活力，是江苏的关键。")

# ---- 四、被忽视细节 ----
heading1(doc, "四、容易被忽视、但信号很重的细节（15条）")
details = [
    ("1", "高技术制造业+11.9%、对规上工业增长贡献率43%", "新质生产力是主要增量。"),
    ("2", "装备制造业增加值+8.8%、占56%", "高端装备主导制造。"),
    ("3", "新能源汽车产量157.48万辆、+32.4%", "新能源车量大增长。"),
    ("4", "集成电路产量1734亿块、+12.3%", "半导体强省。"),
    ("5", "工业机器人产量13.14万台套、+31.5%", "智能制造、机器人赛道。"),
    ("6", "进出口5.95万亿、+6.0%、出口机电+11.6%", "机电产品出口强。"),
    ("7", "电动汽车出口+140.9%、船舶+37%", "新能源车/船舶出海强劲。"),
    ("8", "实际使用外资全国第一、占全国16.1%", "外资高地、开放强。"),
    ("9", "固定资产投资-12.7%、房地产-21.6%", "房地产/投资深度调整。"),
    ("10", "民营企业增加值占GDP 58.5%、工业民企贡献64.2%", "民营底盘厚、贡献大。"),
    ("11", "税收占一般公共预算收入77.4%", "收入质量高、税收主导。"),
    ("12", "常住8518万、城镇化率76.2%、人口自然增-3.6‰", "人口负增、城镇化高。"),
    ("13", "CPI-0.2%、PPI-2.6%", "通缩压力。"),
    ("14", "居民人均可支配收入57971元、+4.6%、城乡比2.02", "收入高、城乡差距缩小。"),
    ("15", "研发投入强度约3.35%、全产业高企林立", "创新驱动与高质量底色。"),
]
for d in details:
    para(doc, "**%s**" % d[0], size=11)
    para(doc, d[1], size=10.5, color=GRAY)

# ---- 五、六个换挡 ----
heading1(doc, "五、把所有细节连起来：江苏正在经历的“六个换挡”")
gear = [
    ("1．动能换挡：从“房产/基建驱动”到“先进制造+新质生产力”", "高技术+11.9%、装备占56%成为新引擎。"),
    ("2．产业换挡：从“传统制造/代工”到“先进制造集群+智造强省”", "电子信息、新能源车、集成电路放量。"),
    ("3．投资换挡：从“基建/地产”到“设备更新/高端制造”", "设备投资+9.1%、航空/半导体投资强，地产-21.6%。"),
    ("4．开放换挡：从“传统外贸”到“机电/高技术+外资”", "机电出口+11.6%、电动汽车+140.9%、外资全国第一。"),
    ("5．人口换挡：从“净流入”到“高质量流入+人口存量”", "常住8518万、城镇化76.2%、人口负增-3.6‰。"),
    ("6．社会换挡：从“追求规模”到“共同富裕+高质量”", "民营占58.5%、城乡比2.02、公共优质。"),
]
for g in gear:
    para(doc, "**%s**　%s" % g, space_after=6)

# ---- 六、增长暗线 ----
heading1(doc, "六、增长暗线：制造业当家、园区经济与新质生产力")
heading2(doc, "6.1 制造业+新质生产力")
para(doc, "江苏以先进制造为核，高技术制造+11.9%、对规上工业贡献43%；装备制造占56%。新能源汽车、集成电路、工业机器人、3D打印等高增，说明江苏在“制造业当家+新质生产力”上走在全国前列。")
heading2(doc, "6.2 园区经济与开放")
para(doc, "苏州工业园等园区承载外资/制造的高效模式，实际使用外资全国第一。园区经济推动江苏制造与开放深度结合，机电/高技术出口强。")
para(doc, "**这条暗线意味着**：江苏的增长叙事正从“外贸+基建”转向“先进制造+创新+园区经济”。看江苏，盯住“高新技术占比”与“外资、机电出口”两组数。")

# ---- 七、财政暗线 ----
heading1(doc, "七、财政暗线：收入稳、税收质量高、民生支出大")
para(doc, "2025年江苏一般公共预算收入10245.7亿元、+2.1%，税收收入7929亿元、+3.8%，税收占77.4%——收入质量较高，主要来自制造业/服务的真实税源。支出高、民生投入大。")
para(doc, "**制度含义**　江苏财政“税收质量高、收入稳健”为全国提供了平衡地产下行的样本。重点是保持制造业/创新税源、优化支出结构，避免地产依赖。")

# ---- 八、产业暗线 ----
heading1(doc, "八、产业暗线：从“世界工厂”到“先进制造集群+智造强省”")
heading2(doc, "8.1 江苏产业的“表”")
table(doc,
    ["指标", "2025增速/占比", "解读"],
    [
        ["规上工业", "+6.5%", "制造强省"],
        ["高技术制造", "+11.9%", "新质生产力"],
        ["装备制造", "+8.8%、占56%", "高端装备主导"],
        ["高新技术产业", "占52.1%", "高新技术过半"],
        ["新能源汽车", "157万辆、+32.4%", "新能源放量"],
        ["集成电路", "+12.3%", "半导体强省"],
        ["工业机器人", "+31.5%", "智能制造"],
    ],
    widths=[5.4, 3.0, 5.2])
heading2(doc, "8.2 从“代工”到“智造/集群”")
para(doc, "江苏从全球“世界工厂”升级为“先进制造集群+智造强省”。高端装备、半导体、新能源车、机器人等是新增长，园区+创新+民营构成“制造业当家”的新底盘。")

# ---- 九、区域 ----
heading1(doc, "九、区域格局：苏南苏中苏北、长三角与园区城市群")
para(doc, "江苏“苏南+苏中+苏北”格局：苏南（苏州/无锡/常州）是制造与外资高地；苏中（南通/扬州）是先进制造与服务在长三角的承接地；苏北（徐州/盐城/连云港）相对滞后但基础设施与产业承接加快。")
para(doc, "长三角一体化下，南京、苏州、无锡、常州等多城进入万亿GDP，苏州工业园等园区承接外资与制造。区域均衡、园区经济，是江苏优势。")

# ---- 十、人口与城市 ----
heading1(doc, "十、人口与城市：长三角人口重镇、老龄化")
heading2(doc, "10.1 人口总量")
para(doc, "2025年江苏常住8518万人、城镇化率76.2%，高于全国。但人口自然增长率-3.6‰、出生率4.2‰，说明“人口存量竞争+老龄化”加深。")
heading2(doc, "10.2 城市与老龄")
para(doc, "南京、苏州等城市是长三角人口与人才集聚地，但整体老龄化（65岁以上约19%估计）带来养老/医疗压力。如何以“创新吸引力”对冲“自然负增”，是江苏长期的课题。")

# ---- 十一、民营经济 ----
heading1(doc, "十一、民营经济：占近六成、工业民企贡献大")
heading2(doc, "11.1 民企地位")
para(doc, "江苏民营经济增加值占GDP 58.5%，规上工业民企增加值+7.6%（贡献64.2%）、规上服务业民营+9.1%（贡献78.3%）。民营是江苏最重要的增长发动机。")
heading2(doc, "11.2 政策与创新")
para(doc, "江苏持续优化营商环境、支持民间投资与专精特新。民企/个体工商户的活跃，是江苏未来10年高质量发展的关键。")

# ---- 十二、事后验证 ----
heading1(doc, "十二、事后验证：用2025年实际执行结果检验报告判断")
heading2(doc, "12.1 年初目标与年末执行对比")
table(doc,
    ["指标", "2025年初目标", "2025实际", "偏差"],
    [
        ["GDP增速", "5%以上", "+5.3%", "达标"],
        ["社零", "5.5%左右", "+3.3%", "低于目标"],
        ["城镇新增就业", "120万以上", "141.14万", "超额"],
        ["研发强度", "3.35%左右", "接近达标", "达标"],
        ["居民收入", "与增长基本同步", "+4.6%", "同步"],
        ["CPI", "2%左右", "-0.2%", "明显低"],
    ],
    widths=[3.0, 3.6, 3.0, 4.0])
heading2(doc, "12.2 判断验证")
para(doc, "三条核心判断得到支持：**（1）制造业/高技术/装备强（规上+6.5%、高技术+11.9%），验证；（2）房地产与投资调整（-21.6%/-12.7%）拖累总量，验证；（3）外资/外贸/民营强（外资全国第一、民企贡献极高），验证。**")
para(doc, "核心观察：江苏在房地产/投资调整下，仍依靠先进制造、外资外贸与民营守住5.3%增长，且收入质量高、民生投入大。2026年看“先进制造+内需消费+投资再平衡”。")

# ---- 十三、五条主线 ----
heading1(doc, "十三、未来5—10年最值得观察的五条主线")
lines5 = [
    ("① 先进制造与新质生产力", "高新技术占比、集成电路/机器人/新能源，能否持续放大。"),
    ("② 园区经济与开放", "外资/机电出口能否持续、园区可否孵化更多独角兽。"),
    ("③ 民营科技创新", "民企占58.5%、能否继续保持制造业活力。"),
    ("④ 房地产/投资再平衡", "-21.6%/-12.7%，能否在制造业与设备投资中找回增长。"),
    ("⑤ 人口与共同富裕", "在人口负增下，如何用创新与人才战略留住人口。"),
]
for l in lines5:
    para(doc, "**%s**　%s" % l, space_after=6)

# ---- 十四、结论 ----
heading1(doc, "十四、最终结论：江苏在“高质量发展+共同富裕”里的增长逻辑")
para(doc, "江苏的2025年，本质上是“**制造业/高质量为核心，而对房地产/投资依赖下降**”的答卷：GDP 14.24万亿、高技术+11.9%、外资总量第一，代价是固投-12.7%、地产-21.6、通胀低。")
para(doc, "只要先进制造、园区、民营/外资、消费能接住，江苏就仍是全国“高质量”头排；如果房地产、投资与通胀持续拖累，江苏需承受“换挡”阵痛。")
para(doc, "最稳妥的观察信号：**一盯制造业/高技术（动能）、二盯外资/外贸/民营（开放底座）、三盯房地产/投资（出清）、四盯城乡收入与民生（共富）、五盯科创/研发（长期）。**江苏，是中国经济“制造+高质量”的压舱石。")

# ---- 附录A ----
heading1(doc, "附录A：主要资料来源与核验路径")
bullet(doc, "江苏省2025年《政府工作报告》（2025年1月）——目标来源。")
bullet(doc, "《2025年江苏省国民经济和社会发展统计公报》（江苏省统计局，2026-03-19）——GDP、工业、外贸、人口等实值。")
bullet(doc, "江苏省统计/经济运行、外贸专篇、外资专报——园区经济与外资。")
bullet(doc, "2026年江苏省政府工作报告——2025执行复盘。")
bullet(doc, "南京海关、商务厅、财政厅——外贸与财政实况。")
heading2(doc, "核验说明")
para(doc, "本报告所有增速、总量、占比均以统计公报/官方发布为基准。涉“苏南/苏北分区”“战略性新兴产业”等以官方口径为准。")

# ---- 附录B ----
heading1(doc, "附录B：建议建立的年度跟踪仪表盘")
para(doc, "建议把下面10个“测脉搏”指标做成年度跟踪表：")
dash = [
    ["#", "指标", "2025基值", "为什么重要"],
    ["1", "GDP增速", "+5.3%", "总量与方向"],
    ["2", "规上工业增速", "+6.5%", "制造底盘"],
    ["3", "高新技术产业占比", "52.1%", "新质生产力"],
    ["4", "进出口/出口增速", "+6.0%/+8.4%", "外贸韧性"],
    ["5", "实际使用外资", "全国第一", "外资信心"],
    ["6", "固定资产投资增速", "-12.7%", "投资动能"],
    ["7", "税收收入/占比", "+3.8%/77.4%", "财政质量"],
    ["8", "常住人口/城镇化率", "8518万/76.2%", "人口与城市"],
    ["9", "社零增速", "+3.3%", "内需消费"],
    ["10", "居民人均可支配收入增速", "+4.6%", "民生获得感"],
]
table(doc, dash[0], dash[1:], widths=[0.8, 4.5, 3.4, 4.8])
para(doc, "把这10个指标连起来看，任何一个“新引擎（2/3/4/5）向上、旧引擎（6）出清”，都说明江苏在真正换挡；反之则是结构阵痛的持续。")

# ========================================= 保存
out = "/Users/x/Desktop/content-prod-lab/reports/江苏省_2025年政府工作报告_深度研究_2026-08-13.docx"
doc.save(out)
print("SAVED:", out)
print("PARAGRAPHS:", len(doc.paragraphs))
print("TABLES:", len(doc.tables))
