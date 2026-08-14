# -*- coding: utf-8 -*-
"""Build 佛山市2025年政府工作报告 深度研究 DOCX, 参照地级市系列版式。"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DARK = RGBColor(0x1F, 0x2A, 0x44)
RED = RGBColor(0xB0, 0x1E, 0x2E)
GRAY = RGBColor(0x59, 0x60, 0x69)
BLUE = RGBColor(0x1F, 0x4E, 0x79)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
HEADER_FILL = "1F2A44"
K = "微软雅黑"


def set_run(run, size=11, bold=False, color=DARK, font=K):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    run.font.name = font
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), font)
    rFonts.set(qn("w:hAnsi"), font)
    rFonts.set(qn("w:eastAsia"), font)


def para(doc, text, size=11, bold=False, color=DARK, align=None,
         space_after=6, space_before=0, font=K):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        set_run(r, size=size, bold=bold or (i % 2 == 1), color=color, font=font)
    return p


def heading1(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(16)
    pf.space_after = Pt(8)
    r = p.add_run(text)
    set_run(r, size=16, bold=True, color=RED)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "B01E2E")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def heading2(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(10)
    pf.space_after = Pt(4)
    r = p.add_run(text)
    set_run(r, size=13, bold=True, color=BLUE)
    return p


def table(doc, headers, rows, widths=None, font_size=9.5):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_run(r, size=font_size, bold=True, color=WHITE)
        tcPr = hdr[i]._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), HEADER_FILL)
        tcPr.append(shd)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(val))
            set_run(r, size=font_size, color=DARK)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    return t


def quote_box(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(8)
    pf.left_indent = Pt(12)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), "B01E2E")
    pBdr.append(left)
    pPr.append(pBdr)
    r = p.add_run(text)
    set_run(r, size=11, color=DARK, font="楷体")
    return p


def bullet(doc, text, size=10.5):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.7)
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        set_run(r, size=size, bold=(i % 2 == 1), color=DARK)
    return p


# ================================================================= 文档主体
doc = Document()
sec = doc.sections[0]
sec.page_width = Cm(21)
sec.page_height = Cm(29.7)
sec.top_margin = Cm(2.4)
sec.bottom_margin = Cm(2.2)
sec.left_margin = Cm(2.5)
sec.right_margin = Cm(2.5)

st = doc.styles["Normal"]
st.font.name = K
st.font.size = Pt(11)
st.element.rPr.rFonts.set(qn("w:eastAsia"), K)


# ---- 封面 ----
para(doc, "佛山市2025年政府工作报告", size=24, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=60, space_after=4)
para(doc, "深度研究与\u201c容易被忽视的细节\u201d分析", size=16, bold=True, color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
para(doc, "从\u201c佛山制造、家电、建材、顺德与民营经济\u201d重新理解佛山", size=12, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)
para(doc, "\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501", color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
para(doc, "研究对象：2025年佛山市政府工作报告及同期财政、国民经济和社会发展统计资料、2026年官方复盘", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
para(doc, "标定日期：2026-08-13", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
doc.add_page_break()

# ---- 目录 ----
catalog = [
    "执行摘要",
    "一、研究口径与阅读方法",
    "二、先看佛山的特殊底盘：佛山制造、家电/建材、顺德民营经济与粤港澳大湾区",
    "三、最关键的宏观错位：GDP几乎滞涨、规上工业/固投/出口全线负增长",
    "四、容易被忽视、但信号很重的细节（15条）",
    "五、2025年GDP目标 vs 实际：对照表",
    "六、2025年增长是谁撑起来的？（结构归因）",
    "七、预算与财政的\u201c含金量\u201d",
    "八、民生底账：人口、收入与城乡",
    "九、城镇与农村：格局与均衡",
    "十、人口流入与流出",
    "十一、物价与货币环境",
    "十二、区域一体化：佛山在\u201c粤港澳大湾区+广佛同城+制造业\u201d里的位置",
    "十三、未来5—10年最值得观察的五条主线",
    "十四、最终结论：佛山在\u201c家电+制造+民营\u201d里的增长逻辑",
    "附录A：主要资料来源与核验路径",
    "附录B：建议建立的年度跟踪仪表盘",
]
para(doc, "目　录", size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=10, space_after=10)
for item in catalog:
    para(doc, item, size=12, space_after=4)
doc.add_page_break()

# ---- 执行摘要 ----
heading1(doc, "执行摘要")
para(doc, "**核心判断**　2025年佛山最显著的是\u201cGDP 13157.35亿元、增长0.2%（远低于5%目标）、二产占49.6%\u201d、\u201c规上工业-4.3%、固投-16.0%、进出口-3.7%\u201d、\u201c常住人口979.19万/城镇化95.80%\u201d。这说明佛山在\u201c制造业深度调整、外部承压\u201d下遇到**罕见的全面下行**，GDP几乎零增长。")
para(doc, "把2025年目标（GDP+5%/规上+6%/财政+2%/进出口+3%）、2025年统计、趋势一起看，佛山是\u201c万亿制造强市+民营经济重镇\u201d，但2025年**制造业-4.3%、固投-16%、地产-27.6%**矩阵回落。总量1.32万亿居广东前列。")
para(doc, "最容易记住的一句话：**佛山是\u201c佛山制造+家电/建材+顺德-民营\u201d的万亿制造地级市**，2025年遭遇**工业、投资、出口、地产全线下行**的一年。观察佛山，与其只看\u201cGDP 1.3万亿\u201d，不如看\u201c家电制造+6.9%、高技术+3.6%、先进制造占55.7%、出口机电+5.5%、民营占比约6成\u201d。")# ---- 一、研究口径与阅读方法 ----
heading1(doc, "一、研究口径与阅读方法")
para(doc, "本报告以\u201c佛山市政府工作报告（2025年，白涛作）\u201d为起点，把\u201c2025年GDP目标（5%以上）\u201d与\u201c2025年国民经济和社会发展统计公报实际值（13157.35亿元/+0.2%）\u201d并置对照，并回看2024年实际（13361.9亿/+1.3%）。")
para(doc, "重点阅读三类证据：**一是指标目标是否达标**；**二是结构（谁贡献增长或下拖累）**；**三是官方在\u201c再创新佛山\u201d中的攻坚方向**。统计口径一般公共预算收入为地方级。")
para(doc, "统一采用\u201c同比增长%\u201d（GDP为不变价），多数指标按可比口径，文中按原口径转录、不逐项换算。")

# ---- 二、特殊底盘 ----
heading1(doc, "二、先看佛山的特殊底盘：佛山制造、家电/建材、顺德民营与大湾区")
para(doc, "**区位与身份**：佛山是广东省地级市、万亿制造强市，\u201c广东制造重镇\u201d，粤港澳大湾区核心（广佛同城），顺德/南海民营经济最发达。")
para(doc, "**产业底盘**：一是家电家用电器（美的等，家电制造+6.8%）；二是建材/陶瓷/家居（佛山陶瓷之都）；三是先进制造/装备（先进制造占规上55.7%、机电出口）；四是民营经济（民投占固投65.5%、市场主体171万户）。")
para(doc, "**人口底盘**：2025年末常住人口979.19万、城镇化率95.80%；特大人口城市。")
para(doc, "**市场与出口**：2025年社零3945.94亿/+0.1%；进出口4813.6亿/-3.7%（出口-3.7%、进口-3.6%）。")

# ---- 三、核心宏观错位 ----
heading1(doc, "三、最关键的宏观错位：GDP几乎停止、工业/固投/出口全线下行（困难年）")
para(doc, "**第一组错位**：2025年GDP目标5%以上，实际13157.35亿/+0.2%，几乎零增长（2024高+1.3%回落）；制造业（二产-2.1%）拖累是主因。")
para(doc, "**第二组错位**：**规上工业-4.3%、固投-16.0%、房地产-27.6%、进出口-3.7%**全线下行；仅高技术+3.6%、家电制造+6.8%、出口机电+5.5%相对稳。\u201c制造调整+地产拖累+外贸承压\u201d。")
para(doc, "**第三组错位**：常住979万/城镇化95.8%、户籍自然+5.17\u2030（人口正增）；但**社零+0.1%、CPI-0.2%**通缩、财政-1.8%。")
para(doc, "一句话：**佛山2025年是\u201c制造/工业/固投/出口全线下行\u201d的深度调整年**——靠家电、高技术、出口机电对冲，但总量承压。")

# ---- 四、15条细节 ----
heading1(doc, "四、容易被忽视、但信号很重的细节（15条）")
bullet(doc, "1. **GDP +0.2%**：制造业-2.1%拖累，二产占49.6%。")
bullet(doc, "2. **规上工业-4.3%**：外商及港澳台-6.7%、中型-13.1%、股份制-3%。")
bullet(doc, "3. **家电制造+6.8%**：美的等家电逆势。")
bullet(doc, "4. **高技术制造+3.6%（医药+9.9%）**：新质生产力。")
bullet(doc, "5. **先进制造占55.7%（+3.2pct）**：结构升级。")
bullet(doc, "6. **固投-16.0%、房地产-27.4%**：地产深度向下。")
bullet(doc, "7. **工业/技改-17.5%**：技改周期调整。")
bullet(doc, "8. **进出口-3.7%、对美出口-21.2%**：外贸（美国）承压。")
bullet(doc, "9. **出口机电+5.5%、对欧-0.7%**：机电产品分化。")
bullet(doc, "10. **常住979.19万/城镇化95.65%**：人口稳。")
bullet(doc, "11. **收入：城镇72298元/+3.1%、农村45658元/+5.0%**，城乡1.58。")
bullet(doc, "12. **CPI-0.2%（通缩）**：价格低迷。")
bullet(doc, "13. **财政800.27亿/-1.8%（税收+1.2%）**：财政微降、税稳。")
bullet(doc, "14. **民营占固投65.5%**：民营韧性。")
bullet(doc, "15. **实际使用外资38.57亿/+50.4%、高企8290家**：外资/科创加码。")# ---- 五、目标 vs 实际对照表 ----
heading1(doc, "五、2025年GDP目标 vs 实际：对照表")
table(doc,
    ["指标", "2025年目标", "2025年实际", "是否达标"],
    [
        ["地区生产总值增速", "5%以上", "13157.35亿元/+0.2%", "严重未达标"],
        ["规上工业增加值", "+6%", "-4.3%（转负）", "严重未达标"],
        ["固定资产投资", "提质提效", "-16.0%（转负）", "未达标"],
        ["社会消费品零售总额", "+5%左右", "3945.94亿元/+0.1%", "未达标"],
        ["进出口总额", "+3%", "4813.6亿元/-3.7%", "未达标(转负)"],
        ["地方一般公共预算收入", "+2%", "800.27亿元/-1.8%", "未达标"],
        ["居民人均可支配收入", "与GDP同步", "+3.3%（全体）", "优于GDP"],
    ],
    widths=[3.4, 2.5, 5.0, 3.1])
para(doc, "**简评**：2025年佛山**所有硬指标均未达标**，GDP +0.2%、工业/固投/进出口转负，是**罕见困难年**。**家电制造、高技术、出口机电、居民收入**相对稳，是零增长下的分化亮点。")

# ---- 六、结构归因 ----
heading1(doc, "六、2025年增长是谁撑起来的？（结构归因）")
para(doc, "**三大产业**：一产251.25亿/+4.1%、二产6520.37亿/-2.1%、三产6385.72亿/+2.5%；二产占49.6%。增长由**三产+一产**支撑，二产（工业/建筑）转负拖累。")
para(doc, "**工业**：规上-4.3%；**家电+6.8%、高技术+3.6%、先进制造占55.7%**；但优势传统-4.0%（建材-14.3%、家具-6.7%、金属-13%）。规上利润总额-20.7%。")
para(doc, "**服务业**：金融/批发零售/文旅（旅游总收入+8.6%）稳；但消费低增。")
para(doc, "**增长归因**：佛山GDP靠**家电/高技术/出口机电/三产**对冲；工业整体、地产、外贸（对美）显著拖累。")

# ---- 七、财政 ----
heading1(doc, "七、预算与财政的\u201c含金量\u201d")
para(doc, "2025年地方一般公共预算收入800.27亿元/-1.7%（可比，税收493.01亿/+1.2%）；支出920.81亿/+0.1%。")
para(doc, "**结构性**：收入\u201c微降（-1.8%）但税收+1.2%\u201d，靠税基（工业利润-20.7%承压）但有韧性；支出基本持平。")
para(doc, "**含金量**：佛山财政\u201c税收稳、总量微降\u201d；在地产-27%与工业利润下行的背景下，税收+1.2%体现民营与家电韧性。")

# ---------------- 八、民生 ----------------
heading1(doc, "八、民生底账：人口、收入与城乡")
para(doc, "常住人口979.19万/城镇化95.80%、户籍人口541.43万（出生10.02\u2030/自然+5.17\u2030）；收入：**居民71130元/+3.3%、城镇72298元/+3.1%、农村45658元/+5.0%**，城乡比1.58。")
para(doc, "就业：城镇新增就业10.67万。")
para(doc, "**民生结论**：收入农村快于城镇、城乡差距1.58小；人口正自然增长+高城镇化；就业稳，民生普惠好。")

# ---------------- 九 ----------------
heading1(doc, "九、城镇与农村：格局与均衡")
para(doc, "佛山城镇化率95.8%（近直辖市水平），中心城区（禅城）与制造强镇（顺德/南海）融合；农村（顺德村居）承接制造与村级工业园。")
para(doc, "城乡收入比1.58（广东/全国低位），农村收入+5.4%快于城镇+3.1%，极均衡。")

# ---------------- 十 ----------------
heading1(doc, "十、人口流入与流出")
para(doc, "佛山常住979万、户籍541万，年均自然+正、大量非户籍产业人口；多靠制造业吸引（常驻/流动）。")
para(doc, "**流入**：制造业、家电、供应链人口+大专/蓝领；**流出**：赴深圳/广州、部分返乡。总体\u201c强制造人口、流动大\u201d。")

# ---------------- 十一 ----------------
heading1(doc, "十一、物价与货币环境")
para(doc, "2025年佛山CPI**下降0.2%**（通缩），交通通信-3.7%、居住-0.2%；衣着+0.4%。")
para(doc, "金融：存款29019.20亿（住户+7.5%）、贷款22004.21亿/+3.8%。\u201c宽中稳、实体融资\u201d。")

# ---------------- 十二 ----------------
heading1(doc, "十二、区域一体化：佛山在\u201c粤港澳大湾区+广佛同城+制造业\u201d里的位置")
para(doc, "佛山是粤港澳大湾区核心、**广佛同城**（与广州共建广佛都市圈），制造业与大湾区/东盟链接；佛山是“万亿制造+大湾区”的重要承载。")
para(doc, "制造业协同：顺德（家电）、乐平（机器人）、大湾区协同；对香港/欧盟/机电贸易。")

# ---------------- 十三 ----------------
heading1(doc, "十三、未来5—10年最值得观察的五条主线")
bullet(doc, "1. **家电/制造升级与智造**：美的·家电数字化、新质生产力。")
bullet(doc, "2. **高端装备/机器人**：先进制造55.7%、机器人产业。")
bullet(doc, "3. **民营经济/万亿市场**：顺德的民营活力。")
bullet(doc, "4. **地产与内需修复**：地产-27%、固投-16%后的修复。")
bullet(doc, "5. **科创/转型**：高企8290首家、广东制造业500强102家。")

# ---------------- 十四 ----------------
heading1(doc, "十四、最终结论：佛山在\u201c家电+制造+民营\u201d里的增长逻辑")
para(doc, "**结论**：佛山2025年的\u201c真相\u201d是——**GDP +0.2%（严重低于5%目标）、规上-4.3%、固投-16%、进出口-3.7%、财政-1.8%**；是\u201c制造/地产/出口全线下行\u201d的调整年。")
para(doc, "**对趋势**：家电/高技术/先进制造代表\u201c结构升级\u201d，地产/固投/传统制造代表\u201c调整\u201d。**工业升级+民营信心+出口转结构**决定复苏，但2025年是难得一遇的下行压力年。")
para(doc, "**若只看一个指标**：看**制造业投资增速（工业技改-17.5%）与家电制造增速（+6.8%）**——佛山在\u201c地产+出口\u201d拖累下，靠**家电/高技术/民营**寻找\u201c再创新佛山\u201d的新增长极，投资与地产的修复是其能否跳出滞涨的关键。")

# ---------------- 附录A ----------------
heading1(doc, "附录A：主要资料来源与核验路径")
bullet(doc, "佛山市人民政府《2025年佛山市政府工作报告》（2025年，白涛）。")
bullet(doc, "佛山市统计局《2025年佛山市国民经济和社会发展统计公报》（2026年）。")
bullet(doc, "佛山日报/佛山+客户端、广东省2025年统计公报交叉核验。")

# ---------------- 附录B ----------------
heading1(doc, "附录B：建议建立的年度跟踪仪表盘")
bullet(doc, "GDP总量/增速 vs 全年目标（+/-百分点）。")
bullet(doc, "规上工业增加值及分行业（家电/建材/高技术/装备）增速。")
bullet(doc, "固定资产投资（总量/工业/房地产/民间）增速。")
bullet(doc, "社会消费品零售总额、家电/新能源汽车。")
bullet(doc, "进出口总额、出口、对美/欧增速。")
bullet(doc, "一般公共预算收入/税收/非税、财政支出。")
bullet(doc, "常住人口、城镇化率、户籍自然。")
bullet(doc, "CPI/核心CPI、规上工业利润。")
bullet(doc, "居民收入（城镇/农村）、城乡比。")
bullet(doc, "高新技术企业、PCT专利、民营主体。")

# ---------------- 保存 ----------------
out = "/Users/x/Desktop/content-prod-lab/reports/佛山市_2025年政府工作报告_深度研究_2026-08-13.docx"
doc.save(out)
print("SAVED OK 佛山市_2025年政府工作报告_深度研究_2026-08-13.docx")