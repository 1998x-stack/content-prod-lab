# -*- coding: utf-8 -*-
"""Build 济宁市2025年政府工作报告 深度研究 DOCX, 参照地级市系列版式。"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DARK = RGBColor(0x1F, 0x2A, 0x44)
RED = RGBColor(0xB0, 0x1E, 0x2E)
GRAY = RGBColor(0x59, 0x60, 0x69)
BLUE = RGBColor(0x1F, 0x4E, 0x79)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
HEADER_FILL = "1F2A44"
K = "微软雅黑"


def set_run(run, size=11, bold=False, color=DARK, font=K):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    run.font.name = font
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), font)
    rFonts.set(qn("w:hAnsi"), font)
    rFonts.set(qn("w:eastAsia"), font)


def para(doc, text, size=11, bold=False, color=DARK, align=None,
         space_after=6, space_before=0, font=K):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        set_run(r, size=size, bold=bold or (i % 2 == 1), color=color, font=font)
    return p


def heading1(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(16)
    pf.space_after = Pt(8)
    r = p.add_run(text)
    set_run(r, size=16, bold=True, color=RED)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "B01E2E")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def heading2(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(10)
    pf.space_after = Pt(4)
    r = p.add_run(text)
    set_run(r, size=13, bold=True, color=BLUE)
    return p


def table(doc, headers, rows, widths=None, font_size=9.5):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_run(r, size=font_size, bold=True, color=WHITE)
        tcPr = hdr[i]._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), HEADER_FILL)
        tcPr.append(shd)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(val))
            set_run(r, size=font_size, color=DARK)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    return t


def quote_box(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(8)
    pf.left_indent = Pt(12)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), "B01E2E")
    pBdr.append(left)
    pPr.append(pBdr)
    r = p.add_run(text)
    set_run(r, size=11, color=DARK, font="楷体")
    return p


def bullet(doc, text, size=10.5):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.7)
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        set_run(r, size=size, bold=(i % 2 == 1), color=DARK)
    return p


# ================================================================= 文档主体
doc = Document()
sec = doc.sections[0]
sec.page_width = Cm(21)
sec.page_height = Cm(29.7)
sec.top_margin = Cm(2.4)
sec.bottom_margin = Cm(2.2)
sec.left_margin = Cm(2.5)
sec.right_margin = Cm(2.5)

st = doc.styles["Normal"]
st.font.name = K
st.font.size = Pt(11)
st.element.rPr.rFonts.set(qn("w:eastAsia"), K)


# ---- 封面 ----
para(doc, "乐山市2025年政府工作报告", size=24, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=60, space_after=4)
para(doc, "深度研究与\u201c容易被忽视的细节\u201d分析", size=16, bold=True, color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
para(doc, "从\u201c乐山大佛、峨眉山、世界双遗产、中国多晶硅之都、三江汇流\u201d重新理解乐山", size=12, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)
para(doc, "\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501", color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
para(doc, "研究对象：2025年乐山市政府工作报告及同期财政、国民经济和社会发展统计资料、2026年官方复盘", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
para(doc, "标定日期：2026-08-14", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
doc.add_page_break()

# ---- 目录 ----
catalog = [
    "执行摘要",
    "一、研究口径与阅读方法",
    "二、先看乐山的特别底盘：大佛峨眉、多晶硅、文旅、三江、晶硅之都",
    "三、最关键的宏观错位：三产+7.3%、社零+6.5%、但规上-9.1%、多晶硅-66.4%、二产-1.5%",
    "四、容易被忽视、但信号很重的细节（15条）",
    "五、2025年GDP目标 vs 实际：对照表",
    "六、2025年增长是谁撑起来的？（结构归因）",
    "七、预算与财政的\u201c含金量\u201d",
    "八、民生底账：人口、收入与城乡",
    "九、城镇与农村：格局与均衡",
    "十、人口流入与流出",
    "十一、物价与货币环境",
    "十二、区域一体化：乐山在川南、成渝双城、成都都市圈、峨眉乐山文旅带里的位置",
    "十三、未来5\u201310年最值得观察的五条主线",
    "十四、最终结论：乐山在\u201c旅游+多晶硅+三产\u201d里的增长逻辑",
    "附录A：主要资料来源与核验路径",
    "附录B：建议建立的年度跟踪仪表盘",
]
para(doc, "目　录", size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=10, space_after=10)
for item in catalog:
    para(doc, item, size=12, space_after=4)
doc.add_page_break()

# ---- 执行摘要 ----
heading1(doc, "执行摘要")
para(doc, "**核心判断**　2025年乐山最显著的是\u201cGDP 2501.54亿/+3.1%、三产+7.3%（占比升至50.6%）、社零+6.5%、旅游（峨眉674.67万人次/乐山大佛810.66万人次）、新能源储能+70%\u201d、\u201c但二产-1.5%、规上-9.1%（重工-11.7%）、多晶硅-66.4%、进出口-4.1%\u201d。这说明乐山在\u201c世界双遗产+晶硅光伏\u201d中，**三产与旅游消费强、但多晶硅/工业在下、总量仅+3.1%**。")
para(doc, "把乐山2025年目标与实际对照，乐山是\u201c文旅+晶硅\u201d路径：**峨眉山/乐山大佛、多晶硅光伏、新储能、文旅会展、食品饮料**是支柱。三产+7.3%支撑了总量的低增长，二产/晶硅在下行。")
para(doc, "最容易记住的一句话：**乐山是\u201c世界双遗产地（大佛+峨眉山）、中国多晶硅之都\u201d，靠\u201c旅游+晶硅+新储能\u201d驱动。**观察乐山，与其只看\u201cGDP 2502亿/+3.1%\u201d，不如看\u201c多晶硅11.36万吨-66.4%、新型储能+70%、峨眉乐山大佛游客、三产占50.6%\u201d。")

# ---- 一、研究口径与阅读方法 ----
heading1(doc, "一、研究口径与阅读方法")
para(doc, "**本文的最终目标**，不是复述乐山市2025年政府工作报告的原文，而是把它放到\u201c目标→实际\u201d和\u201c长期底盘→年度对撞\u201d的坐标里重新读一遍，找出乐山2025年到底靠什么增长、哪里在换挡、哪里是暗线。")
para(doc, "**三组资料互证**：（1）《2025年乐山市国民经济和社会发展统计公报》（乐山市统计局，2026-04-02发布）提供全部2025年**实际完成**数据；（2）《乐山市2026年政府工作报告》（2026-01-21市八届人大七次会议）复盘2025年；（3）《乐山市2025年政府工作报告》与计划草案提供**目标设定**。三份材料交叉核验，避免单源失真。")
para(doc, "**核心框架**：先交代乐山的\u201c底盘\u201d（大佛峨眉+多晶硅+三江文旅），再用\u201c目标vs实际\u201d对照表定位宏观错位，接着用15条细节捕捉暗线，最后从财政、民生、城乡、人口、物价坐实，落到2026/十五五主线与结论。")
para(doc, "**口径说明**：GDP 2501.54亿元+3.1%、2024年经最终核实为2427.64亿+4.3%。三次产业增加值268.28/966.64/1266.62、占比10.7:38.7:50.6。固投+5.3%、房地产-2.0%。社零1004.10亿+6.5%。一般公共预算总收入282.02亿+3.2%、地方172.16亿+5.6%。进出口77.28亿-4.1%。居民收入35999元+4.6%。常住约341.2万（户籍）、自然减少。")
para(doc, "**为什么值得单独研究乐山**：GDP+3.1%、二产-1.5%、规上-9.1%（多晶硅-66.4%），是\u201c晶硅光伏下行+文旅韧性\u201d的一年；但三产+7.3%、社零+6.5%、旅游强、新储能+70%。\u201c旧工业（晶硅）去产能、新三产（文旅）上\u201d，是观察\u201c资源/制造型城市转型\u201d的上佳切片。")
# ---- 二、先看乐山的特别底盘 ----
heading1(doc, "二、先看乐山的特别底盘：大佛峨眉、多晶硅、文旅、三江、晶硅之都")
para(doc, "乐山的成长逻辑，不能只看GDP数字，而要先把\u201c底盘\u201d摊开。底盘决定了几十年、上百年的禀赋，也决定了2025年增长的来源和约束。")
bullet(doc, "**世界双遗产·旅游之都**：乐山大佛（世界文化与自然双遗产）与峨眉山，是四川最知名的目的地之一；2025年峨眉山接待674.67万人次、乐山大佛810.66万人次。")
bullet(doc, "**中国多晶硅/晶硅光伏之都**：乐山是全国最重要的多晶硅/晶硅光伏基地（永祥/通威等），2025年多晶硅产量11.36万吨（-66.4%）、单晶硅7.89万吨（-48.7%），晶硅光伏是工业与转型的支柱。")
bullet(doc, "**新型储能/新材料新赛道**：新型储能产值+70%、低空经济+28.3%，是晶硅之外的新兴赛道。")
bullet(doc, "**冶金/建材/化工（黑色冶炼+24.6%、造纸+26.5%）**：黑色金属冶炼、非金属矿、建材陶瓷等传统工业。")
bullet(doc, "**三江汇流·文旅峨眉乐山文旅带**：岷江/大渡河/青衣江三江汇流，乐山-峨眉山文旅带、会展/美食（跷脚牛肉/钵钵鸡）名城。")
heading2(doc, "底盘的产业钟摆")
para(doc, "把底盘归纳成一句话：**乐山是\u201c晶硅光伏+文旅\u201d的双轮城市**。多晶硅定工业、峨眉乐山大佛供旅游、新储能/低空添转型、三江美食供消费。城市资产在\u201c晶硅→储能\u201d与\u201c旅游→三产\u201d两条链上。2025年晶硅-66%去产能拖累工业（规上-9.1%），但三产+7.3%、旅游消费强，是典型的\u201c旧产能去化、新动能（文旅/储能）补位\u201d。")

# ---- 三、最关键的宏观错位 ----
heading1(doc, "三、最关键的宏观错位：三产+7.3%、社零+6.5%、但规上-9.1%、多晶硅-66.4%、二产-1.5%")
para(doc, "把2025年数据摊开，乐山最醒目的不是\u201c+3.1%\u201d，而是**三产文旅强、二产晶硅弱**的换挡错位。")
quote_box(doc, "宏观错位一号：**三产+7.3% vs 二产-1.5%**。第三产业1266.62亿、占比50.6%（首超50%）、批发零售+9.4%、营利性服务+11.2%、住宿餐饮+7.6%，文旅/服务旺；但二产-1.5%、规上工业-9.1%（重工-11.7%），工业被晶硅拖累。")
quote_box(doc, "宏观错位二号：**旅游消费强 vs 晶硅外贸弱**。社零1004.1亿+6.5%、峨眉大佛游客旺；但多晶硅产量-66.4%、晶硅光伏产业-65.6%、进出口-4.1%——外贸被光伏价格战拖累、内需消费在撑。")
quote_box(doc, "宏观错位三号：**新产能上（储能+70%） vs 旧工业下**。新型储能+70%、低空经济+28.3%、半导体分立器件+17.6%；但晶硅光伏-65.6%、绿色化工-15.1%、装备-10.6%，新旧动能青黄不接。")
para(doc, "三条错位加在一起，指向同一个结论：乐山2025年是\u201c三产/文旅+社零/储能强、晶硅/工业弱\u201d的换挡年。总量靠三产与消费托底（+3.1%），但晶硅光伏深度去产能是最大拖累。乐山的下一程，取决于把\u201c晶硅→储能/新型材料\u201d的转型与新三产放大。")

# ---- 四、容易被忽视、但信号很重的细节（15条） ----
heading1(doc, "四、容易被忽视、但信号很重的细节（15条）")
para(doc, "下面这15条，散落在统计公报的角落里，单独看无关紧要，合在一起却拼出乐山2025年真实的结构肌理。")
bullet(doc, "**1. 规上工业-9.1%（但计算机通信电子+38.1%、造纸+26.5%、黑色冶金+24.6%、电力+17.7%）**：光电/造纸冶金在逆势增长。")
bullet(doc, "**2. 多晶硅11.36万吨-66.4%、单晶硅-48.7%**：晶硅在价格战下大幅减产。")
bullet(doc, "**3. 新型储能+70%、低空经济+28.3%**：新赛道在放量。")
bullet(doc, "**4. 规上工业营收1415.3亿-10.7%、利润102.6亿+211.9%**：营收减、利润大增（晶硅成本降/其他溢）。")
bullet(doc, "**5. 工业用电量211亿千瓦时-34.7%**：晶硅/建材高耗能去产能、用电大降。")
bullet(doc, "**6. 三产占GDP50.6%（首次过半）**：乐山结构向服务主导转型。")
bullet(doc, "**7. 民营经济增加值1518.93亿占60.7%（三产+11.9%）**：民营三产（旅游/餐饮）活跃。")
bullet(doc, "**8. 社零1004.1亿+6.5%、限上+8.1%（通讯+59%、家电+12.1%）**：消费升级、通讯家电强。")
bullet(doc, "**9. 旅游：峨眉山674.67万人次/乐山大佛810.66万人次**：世界双遗产流量。")
bullet(doc, "**10. 一产+4.0%（茶叶+6.0%、水果+5.6%、蔬菜+4.1%）**：农业特色稳。")
bullet(doc, "**11. 货币：存款4759.54亿+8.8%、贷款3413.42亿+11.7%**：贷款高增、金融活。")
bullet(doc, "**12. 房租/地产：房地产投资-2.0%、销售**：地产缓。")
bullet(doc, "**13. 高新企业153家、专利授权1526件**：创新在积累。")
bullet(doc, "**14. 住户存款3581.56亿+9.8%**：储蓄高、消费潜力大。")
bullet(doc, "**15. 出生13908、死亡28388、自然负增**：人口自然减少-2.8万（户籍）。")
# ---- 五、2025年GDP目标 vs 实际：对照表 ----
heading1(doc, "五、2025年GDP目标 vs 实际：对照表")
para(doc, "把乐山市2025年预期目标与实际值逐项对照：")
table(doc,
    ["指标", "2025年目标", "2025年实际", "判定"],
    [
        ["地区生产总值增速", "+5.5%左右", "2501.54亿，+3.1%", "明显低于目标"],
        ["规上工业", "+8%左右", "-9.1%", "大幅未达"],
        ["固定资产投资", "+6%以上", "+5.3%", "略低"],
        ["社会消费品零售总额", "+6%左右", "1004.1亿，+6.5%", "超额"],
        ["一般公共预算收入", "+5%左右", "地方172.16亿，+5.6%", "达成"],
        ["进出口总额", "（稳外贸）", "77.28亿，-4.1%", "未达"],
    ],
    widths=[4.6, 3.0, 4.6, 3.8],
)
para(doc, "这张表透露乐山2025年的\u201c成色\u201d：**社零、财收达成/超额**；但**GDP+3.1%（差5.5%）、规上-9.1%、进出口-4.1%远低目标**。\u201c三产消费强、工业/晶硅拖累\u201d是对2025年最简练的总结——总量被晶硅下行压制，靠三产/社零托底。")

# ---- 六、2025年增长是谁撑起来的？（结构归因） ----
heading1(doc, "六、2025年增长是谁撑起来的？（结构归因）")
para(doc, "乐山2025年GDP+3.1%的功劳簿，大致可以拆成几块：")
bullet(doc, "**第一功：三产+7.3%（绝对引擎）**。第三产业1266.62亿、占比50.6%；批发零售+9.4%、营利性服务+11.2%、住宿餐饮+7.6%、金融+8.8%、交通+6.4%。文旅/会展/金融撑起服务。")
bullet(doc, "**第二功：消费/社零+6.5%**。社零1004.1亿、限上+8.1%（通讯+59%、家电+12.1%），以旧换新+文旅消费拉动。")
bullet(doc, "**第三功：一产+4.0%**。茶叶+6%、水果+5.6%、蔬菜+4.1%，农业稳增。")
bullet(doc, "**拖累：二产-1.5%**。规上-9.1%、晶硅光伏-65.6%去产能，拖累工业与总量。")
para(doc, "**结构归因结论**：2025年的乐山增长是\u201c三产+消费+农业\u201d驱动，总量靠文旅/服务/社零（三产50.6%）托住；二产/晶硅严重拖累。乐山已从\u201c工业多晶硅驱动\u201d转向\u201c三产文旅+新储能\u201d，是结构性转轨的一年。")

# ---- 七、预算与财政的\u201c含金量\u201d ----
heading1(doc, "七、预算与财政的\u201c含金量\u201d")
table(doc, ["财政指标", "2025年数值", "同比", "点评"],
    [
        ["一般公共预算总收入", "282.02亿元", "+3.2%", "稳"],
        ["地方一般公共预算", "172.16亿元", "+5.6%", "达成目标"],
        ["＃税收收入", "79.81亿元", "+6.8%", "增"],
        ["一般公共预算支出", "411.33亿元", "+8.8%", "支出扩张"],
        ["存款/贷款", "4759.54/3413.42亿", "+8.8%/+11.7%", "金融活跃"],
    ],
    widths=[4.6, 3.4, 3.0, 4.6],
)
para(doc, "财政核心判断：**收入稳增、支出扩张、质量尚可**。地方一般公共预算+5.6%、税收+6.8%，在规上-9.1%下财税仍正增（靠三产/消费税基）；支出+8.8%加大民生与基建；存贷款双位数增长、金融对经济支撑好。")
quote_box(doc, "财政与宏观的勾连：**工业负、财税稳**。乐山工业（晶硅）-9.1%但财税仍+5.6%，因为财政税基已从\u201c重工业\u201d向\u201c三产消费\u201d转移；这验证乐山经济结构已实质服务化，抗晶硅周期能力强。")

# ---- 八、民生底账 ----
heading1(doc, "八、民生底账：人口、收入与城乡")
bullet(doc, "**居民收入**：全体居民35999元+4.6%；城镇49195元+4.1%、农村22801元+5.3%；城乡比2.16。")
bullet(doc, "**就业**：旅游/三产就业、民营（占GDP60.7%）吸纳。")
bullet(doc, "**收入与消费**：住户存款3581.56亿+9.8%、储蓄高。")
bullet(doc, "**社保**：城/乡养老123.44+117.92万、医保304万（职工+居民）。")
para(doc, "民生综合评价：**收入中高速、农村快于城镇、储蓄充裕**；就业靠三产旅游、人口自然减少（-2.8万）是隐忧。")

# ================= 九、城镇与农村 =================
heading1(doc, "九、城镇与农村：格局与均衡")
bullet(doc, "**城镇化中高**：乐山户籍341.2万、城镇化率较高（四川中西部）。")
bullet(doc, "**收入城乡**：城49195/农22801、比2.16，相对均衡。")
bullet(doc, "**农业**：粮食127.39万吨+0.7%、茶叶6.5万吨+6%、水果30.9万吨+5.6%、农一产+4.0%。")
para(doc, "乐山城乡相对均衡（2.16）、农业特色（茶叶/果蔬）较强，是成都平原南缘农业旅游一体。")

# ============ 十、人口流入与流出 ============
heading1(doc, "十、人口流入与流出")
bullet(doc, "**总量减**：户籍341.2万、-2.8万；出生13908、死亡28388，自然负增。")
bullet(doc, "**人口外流**：四川地级市普遍、年轻人口流向成都/长三角。")
bullet(doc, "**对经济含义**：户籍341万、常住328左右，人口总量较大但呈自然减少。")
quote_box(doc, "人口：户籍341.2万、自然-2.8万，规模较大但自然下降。乐山要留人，靠文旅/储能/三产就业。")

# ============ 十一、物价与货币 ============
heading1(doc, "十一、物价与货币环境")
bullet(doc, "**CPI-0.7%**：食品-2.0%、衣着-3.9%、交通-3.5%、消费品-1.9%，通缩、消费价格偏冷（服务+1.2%）。")
bullet(doc, "**金融**：存款4759.54亿+8.8%、贷款3413.42亿+11.7%，贷款高增（中长期+9.0%、短期+22.5%）。")
para(doc, "温和通缩+信贷高增：乐山资金充裕、贷款活跃，对基建/服务/储能支撑好。")

# ========== 十二、区域一体化 ==========
heading1(doc, "十二、区域一体化：乐山在成渝双城、川南、成都都市圈、峨眉乐山文旅带里的位置")
bullet(doc, "**成都都市圈/成渝双城经济圈**：乐山是成渝中间、川南门户，承接成都产业外溢与旅游客源。")
bullet(doc, "**峨眉乐山文旅带**：乐山-峨眉山、乐山-宜宾/泸州川南城市群，文旅/交通联动。")
bullet(doc, "**多晶硅/晶硅基地**：全国晶硅光伏产业链重地（永祥/通威），锂电、新储能新园区。")
bullet(doc, "**县域**：市中区/五通桥（工业/晶硅）、峨眉山（文旅）、夹江/犍为（建材/茶叶）、沐川/峨边（生态），构成\u201c文旅+晶硅+农业\u201d。")
para(doc, "乐山坐标：**在\u201c成都都市圈+川南\u201d下，乐山是川南文旅门户与晶硅光伏基地**。若文旅会展、晶硅转型储能、成渝联动做实，能级更大。")

# ============ 十三、未来5-10年主线 ============
heading1(doc, "十三、未来5\u201310年最值得观察的五条主线")
bullet(doc, "**主线一：晶硅→储能/新材料转型**。多晶硅-66.4%去产能，若向新型储能（+70%）、固态电池、N型电池转型，工业再起。")
bullet(doc, "**主线二：世界双遗产文旅的提质（峨眉乐山大佛）**。把\u201c人次\u201d做\u201c收入\"、酒店/度假/会展/美食，三产再升级。")
bullet(doc, "**主线三：新质制造（通信电子/半导体）**。计算机通信+38.1%、半导体分立器件+17.6%，可成第二工业。")
bullet(doc, "**主线四：晶硅价格周期与新能源产业卷**。光伏价格战、产能退出，乐山能否靠成本/新电池（N型）逆袭。")
bullet(doc, "**主线五：人口与消费**。户籍-2.8万，靠三产/储能就业与文旅留人。")
para(doc, "五条主线里，**最值得盯的是主线一（晶硅→储能）与主线二（文旅升级）**——乐山从\u201c多晶硅（+文旅）\u201d转向\u201c储能+新材料+文旅\u201d的复合，是未来十年关键。")

# ============ 十四、最终结论 ============
heading1(doc, "十四、最终结论：乐山在\u201c旅游+多晶硅+三产\u201d里的增长逻辑")
para(doc, "综合2025年的开局与结构，给乐山一个平衡的结论：")
para(doc, "**一、三产成主引擎、消费强。**GDP+3.1%、三产+7.3%（占比50.6%）、社零+6.5%，旅游/服务/消费托住总量，乐山已从\u201c工业主导\u201d转向\u201c三产主导\u201d。")
para(doc, "**二、但晶硅/工业深度去产能。**规上-9.1%、多晶硅-66.4%、工业用电-34.7%，晶硅价格战是最大拖累。")
para(doc, "**三、财政/金融稳、结构转。**财税+5.6%（三产/消费税基）、贷款+11.7%、新储能+70%——乐山被动在向新质转。")
para(doc, "**四、2026年前瞻。**在\u201c稳三产、促储能、纾晶硅、促消费、稳财政\u201d约束下，乐山2026年边际变量是**晶硅价格/储能放量、文旅收入/人次、半导体/计算机、社零、新产业**。若能靠储能/新材料/文旅对冲晶硅，乐山有望企稳回升。")
para(doc, "**一句话总结：乐山2025年交出了\u201c+3.1%、三产+7.3%占半、社零+6.5%\u201d的答卷，世界双遗产的文旅与消费托住了总量，但成也晶硅、阻也晶硅——多晶硅-66%的阵痛仍在；更该问的是，在晶硅价格周期与文旅红利的交汇处，这座城市能否把\u201c大佛峨眉+储能\u201d炼成\u201c旅游+储能+高端制造\u201d的新竞争力，让乐山从\u201c靠晶硅\u201d真正走向\u201c靠文旅+新质\u201d的更强韧性。**当下看，结构在变好（三产过半），但晶硅拖累与人口下降提醒——转型升级的乐山，路正长。")

# ============ 附录A ============
heading1(doc, "附录A：主要资料来源与核验路径")
bullet(doc, "《2025年乐山市国民经济和社会发展统计公报》（乐山市统计局，2026-04）：GDP、三次产业、工业、投资、消费、财政、收入、人口等。")
bullet(doc, "《乐山市2026年政府工作报告》（2026-01-21赵迎春）：2025年回顾、2026年目标。")
bullet(doc, "《乐山市2025年政府工作报告》：2025年预期目标。")
bullet(doc, "核验方式：以统计公报为准、以政府工作报告为框架，交叉比对目标-实际、剔除五经普修订差异。")

# ============ 附录B ============
heading1(doc, "附录B：建议建立年度跟踪仪表盘")
table(doc, ["维度", "跟踪指标", "用途/预警"],
    [
        ["总量", "GDP、人均、三产占比", "三产/韧性"],
        ["工业", "规上工业、多晶硅/晶硅", "工业合规"],
        ["储能", "新型储能、低空、半导体", "新赛道"],
        ["消费", "社零、旅游人次/收入", "内需/文旅"],
        ["投资", "固投、房地产、文旅/制造", "投资质量"],
        ["财政", "收入、税收、支出", "财政韧性"],
        ["人口", "常住/户籍、自然增长率", "人口"],
    ],
    widths=[2.6, 6.2, 6.0],
)

# ========== 保存
doc.save("/Users/x/Desktop/content-prod-lab/reports/乐山市_2025年政府工作报告_深度研究_2026-08-14.docx")
print("SAVED OK: 乐山市 /Users/x/Desktop/content-prod-lab/reports/乐山市_2025年政府工作报告_深度研究_2026-08-14.docx")
