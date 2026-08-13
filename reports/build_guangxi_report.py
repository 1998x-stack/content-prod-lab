# -*- coding: utf-8 -*-
"""Build 广西壮族自治区2025年政府工作报告 深度研究 DOCX, 参照省市系列版式。"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DARK = RGBColor(0x1F, 0x2A, 0x44)
RED = RGBColor(0xB0, 0x1E, 0x2E)
GRAY = RGBColor(0x59, 0x60, 0x69)
BLUE = RGBColor(0x1F, 0x4E, 0x79)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
HEADER_FILL = "1F2A44"
K = "微软雅黑"


def set_run(run, size=11, bold=False, color=DARK, font=K):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    run.font.name = font
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), font)
    rFonts.set(qn("w:hAnsi"), font)
    rFonts.set(qn("w:eastAsia"), font)


def para(doc, text, size=11, bold=False, color=DARK, align=None,
         space_after=6, space_before=0, font=K):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        set_run(r, size=size, bold=bold or (i % 2 == 1), color=color, font=font)
    return p


def heading1(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(16)
    pf.space_after = Pt(8)
    r = p.add_run(text)
    set_run(r, size=16, bold=True, color=RED)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "B01E2E")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def heading2(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(10)
    pf.space_after = Pt(4)
    r = p.add_run(text)
    set_run(r, size=13, bold=True, color=BLUE)
    return p


def table(doc, headers, rows, widths=None, font_size=9.5):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_run(r, size=font_size, bold=True, color=WHITE)
        tcPr = hdr[i]._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), HEADER_FILL)
        tcPr.append(shd)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(val))
            set_run(r, size=font_size, color=DARK)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    return t


def quote_box(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(8)
    pf.left_indent = Pt(12)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), "B01E2E")
    pBdr.append(left)
    pPr.append(pBdr)
    r = p.add_run(text)
    set_run(r, size=11, color=DARK, font="楷体")
    return p


def bullet(doc, text, size=10.5):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.7)
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        set_run(r, size=size, bold=(i % 2 == 1), color=DARK)
    return p


# ================================================================= 文档主体
doc = Document()
sec = doc.sections[0]
sec.page_width = Cm(21)
sec.page_height = Cm(29.7)
sec.top_margin = Cm(2.4)
sec.bottom_margin = Cm(2.2)
sec.left_margin = Cm(2.5)
sec.right_margin = Cm(2.5)

st = doc.styles["Normal"]
st.font.name = K
st.font.size = Pt(11)
st.element.rPr.rFonts.set(qn("w:eastAsia"), K)


# ---- 封面 ----
para(doc, "广西壮族自治区2025年政府工作报告", size=24, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=60, space_after=4)
para(doc, "深度研究与“容易被忽视的细节”分析", size=16, bold=True, color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
para(doc, "从“面向东盟、港口、西部陆海新通道、平陆运河与特色农业”重新理解广西", size=12, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)
para(doc, "━━━━━━━━━━━━━━━━", color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
para(doc, "研究对象：2025年广西壮族自治区政府工作报告及同期财政、国民经济和社会发展统计资料、2026年官方复盘", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
para(doc, "整理时间：2026年8月", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
para(doc, "分析性质：分析性研究 ｜ 非官方解读", size=11, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
doc.add_page_break()

# ---- 目录 ----
catalog = [
    "执行摘要",
    "一、研究口径与阅读方法",
    "二、先看广西的特殊底盘：面向东盟、港口与通道、特色农业/少数民族",
    "三、最关键的宏观错位：GDP破2.97万亿、制造/出口强，但投资/地产与物价偏弱",
    "四、容易被忽视、但信号很重的细节（15条）",
    "五、2025年GDP目标 vs 实际：对照表",
    "六、2025年增长是谁撑起来的？（结构归因）",
    "七、预算与财政的“含金量”",
    "八、民生底账：人口、收入与城乡",
    "九、城镇与农村：格局与均衡",
    "十、人口流入与流出",
    "十一、物价与货币环境",
    "十二、区域一体化：广西在“西部陆海新通道+北部湾+东盟”里的位置",
    "十三、未来5—10年最值得观察的五条主线",
    "十四、最终结论：广西在“面向东盟+通道+制造”里的增长逻辑",
    "附录A：主要资料来源与核验路径",
    "附录B：建议建立的年度跟踪仪表盘",
]
para(doc, "目　录", size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=10, space_after=10)
for item in catalog:
    para(doc, item, size=12, space_after=4)
doc.add_page_break()

# ---- 执行摘要 ----
heading1(doc, "执行摘要")
para(doc, "**核心判断**　如果只看新闻标题，2025年广西最显眼的是“GDP破2.97万亿、增长5.1%”、“出口+10.6%、对一带一路+11.0%”和“高技术制造+23.7%、装备+16.5%”。但这份研究真正值得深读的，是这座“面向东盟+港口+西部陆海新通道+平陆运河”的边疆省份，如何在固定资产投资（-8.2%）、房地产开发投资（-16.2%）与CPI（-0.3%）的背景下，靠“装备/高技术制造、出口、通道（北部湾/平陆运河）+特色农业”稳住增长。")
para(doc, "把2025年初设定的目标（GDP增长5%左右）、2025年《国民经济和社会发展统计公报》、以及2026年报告对2025年的复盘放在一起看，广西呈现清晰暗线：**从“轻工/有色金属/农业”的旧底盘，向“面向东盟开放+西部通道（平陆运河/北部湾）+先进制造+特色农业”转型**。旧引擎（传统轻工、有色金属、房地产、一般基建）在调整；新引擎（装备制造、高技术制造、新能源车、出口、通道经济）被要求更快补位。这也是人口近5000万、少数民族聚居、处在“兴边富民+通道门户”双重定位的省份。")
para(doc, "因此，本报告不按政府工作报告逐段复述，而采用“显性表述—同期数据—制度含义—长期影响”的方式，专门提取容易被忽视、但对判断广西未来5—10年发展模式更有价值的信号。")
para(doc, "最容易记住的一句话：**广西是“面向东盟开放门户+西部陆海新通道枢纽+特色农业/生态”的边疆省份，靠“平陆运河/北部湾港+先进制造+出口”把内陆与大海连通。**观察广西，与其看“GDP 2.97万亿”，不如看“西部通道/平陆运河、东盟贸易、装备/高技术制造、港口与农业特色”这几张名片。")
heading2(doc, "一页速览：2025年广西经济的“表与里”")
table(doc,
    ["维度", "表面现象", "更深层信号"],
    [
        ["增长", "GDP 29727.45亿、+5.1%", "二产+5.0%、三产+5.4%"],
        ["产业", "规上工业+7.7%、高技术+23.7%", "装备+16.5%、先进制造提升"],
        ["外贸", "进出口8192.62亿、+8.4%（出口+10.6%）", "一带一路+11.0%、面向东盟"],
        ["投资", "固定资产投资-8.2%", "工业占43.5%、地产-16.2%"],
        ["财政", "一般公共预算收入1922.05亿、+4.6%", "税收+3.9%"],
        ["消费", "社零8396.30亿、+3.0%", "乡村+7.0%、线上+10.4%"],
        ["人口", "常住4989万、城镇化率58.09%", "人口-24万、自然增-0.90‰"],
        ["开放", "西部陆海新通道/平陆运河/北部湾港", "东盟门户"],
    ],
    widths=[2.2, 5.2, 8.6])
para(doc, "", size=4, space_after=2)

# ---- 一、研究口径与阅读方法 ----
heading1(doc, "一、研究口径与阅读方法")
heading2(doc, "1.1 本报告的三份核心底稿")
bullet(doc, "**2025年《政府工作报告》**：2025年1月在自治区十四届人大三次会议上作，给出全年预期目标（GDP增长5%左右、部署重大项目建设等）。固定资产投资/进出口等未全部设具体速率。")
bullet(doc, "**2025年《国民经济和社会发展统计公报》**：广西统计局2026年4月9日发布，给出全年实际执行数与增速，是“事后验证”的锚点。")
bullet(doc, "**2026年广西政府工作报告及官方复盘**：对2025执行结果的追认，交叉验证“目标—执行—再评价”三段式。")
heading2(doc, "1.2 阅读方法：显性—数据—制度—长期")
para(doc, "每一章按“**显性表述 → 同期数据 → 制度含义 → 长期影响**”四层展开。其中“制度含义”回答“政府为什么这么排优先序”，“长期影响”回答“这对未来5—10年意味着什么”。")
para(doc, "关键判别：**当报告使用的“增长词”和统计公报里的实际数不一致时，数据优先。**例如2025年GDP目标5%、实际5.1%达标；但固定资产投资-8.2%、地产-16.2%。差异反映：广西“装备/出口/通道强，投资/地产/物价弱”。")

# ---- 二、特殊底盘 ----
heading1(doc, "二、先看广西的特殊底盘：面向东盟、港口与通道、特色农业/少数民族")
para(doc, "在所有省份里，广西的“底盘”独特：**面向东盟+西部陆海新通道门户+港口（北部湾）+特色农业/少数民族**四合一。常住4989万、城镇化率58.09%，是内陆连通大海的“通道省”。")
para(doc, "这决定广西的多重身份并存：**面向东盟/开放门户**（东盟第一贸易、北部湾港）、**西部通道/平陆运河**、**港口/海港**（北部湾港）、**特色农业**（甘蔗/水果/糖）、**民族地区**（壮族等、乡村振兴）。")
heading2(doc, "2.1 面向东盟+通道")
para(doc, "广西是中国面向东盟的门户，对一带一路+11%、东盟/RCEP贸易强。平陆运河+北部湾港+西部陆海新通道，让广西“向海+通陆”。")
heading2(doc, "2.2 先进制造+出口")
para(doc, "装备制造+16.5%、高技术+23.7%，出口+10.6%。新能源汽车、锂电池、风电装备等新赛道放量。制造升级+出口，是广西从“农业/有色金属”到制造强省的努力。")
heading2(doc, "2.3 特色农业/生态")
para(doc, "广西甘蔗/糖、果、铝、药材等特色农业与生态资源丰富；同时是少数民族聚居、兴边富民的重点区。")

# ---- 三、宏观错位 ----
heading1(doc, "三、最关键的宏观错位：GDP破2.97万亿、制造/出口强，但投资/地产与物价偏弱")
para(doc, "把2025年广西的宏观面放进一张表，会出现令人意外的“错位”：表观增长来自出口/制造/通道，而投资/地产/物价偏弱。这个错位，正是读懂广西的关键。")
heading2(doc, "3.1 “强的一面”")
table(doc,
    ["指标", "2025实绩", "信号"],
    [
        ["GDP", "29727.45亿、+5.1%", "破3万亿"],
        ["规上工业", "+7.7%", "制造强"],
        ["高技术制造", "+23.7%", "高技术领跑"],
        ["装备制造", "+16.5%", "高端装备"],
        ["出口/一带一路", "+10.6%/+11.0%", "东盟门户"],
        ["一般公共预算收入", "+4.6%", "财政稳"],
    ],
    widths=[3.2, 5.4, 6.0])
heading2(doc, "3.2 “弱/调整的一面”")
table(doc,
    ["指标", "2025实值", "信号"],
    [
        ["固定资产投资", "-8.2%", "投资回落"],
        ["房地产开发投资", "-16.2%", "地产调整"],
        ["基础设施投资", "-11.1%", "基建放缓"],
        ["CPI", "-0.3%", "物价走弱"],
        ["社会消费品零售", "+3.0%", "消费偏弱"],
    ],
    widths=[3.2, 5.4, 6.0])
para(doc, "**错位结论**　广西的增长“很强、但也很矛盾”。强的部分（制造/出口/通道/财政）与弱的部分（投资/地产/物价/消费）并存。**真正的焦点是“制造/出口强，投资/物价弱”**：高技术+23.3%拉动增长，但固投-8.2%、地产-16.2%。2026年“面向东盟+通道+稳内需/地产”是重点。")

# ---- 四、被忽视细节 ----
heading1(doc, "四、容易被忽视、但信号很重的细节（15条）")
details = [
    ("1", "规上工业+7.7%、制造业占GDP上升", "制造升级。"),
    ("2", "装备制造+16.5%、高技术+23.7%", "装备/高技术领跑。"),
    ("3", "光缆+59.5%、服务机器人+23.0%", "新兴/智能产品。"),
    ("4", "太阳能超白玻璃+60.3%、锂电+54.5%", "新能源链。"),
    ("5", "风力发电机组+30.6%、电子元件+19.1%", "能源/电子。"),
    ("6", "出口8192.62亿、+8.4%、出口+10.6%", "出口强、东盟。"),
    ("7", "对一带一路+11.0%、RCEP+6.1%", "东盟/一带一路。"),
    ("8", "实际利用外资-2.6%、但对外投资+69.1%", "外资入、出海强。"),
    ("9", "固定资产投资-8.2%、工业占43.5%", "工业投资高、占比升。"),
    ("10", "民间投资+2.4%、占36.9%", "民间投资增。"),
    ("11", "社零/乡村+7.0%、线上+10.4%", "内需/线上强。"),
    ("12", "常住4989万、-24万、城镇化58.09%（+0.70pct）", "人口负增、城镇化。"),
    ("13", "居民人均可支配收入32721元、+5.1%", "收入稳、城乡扩。"),
    ("14", "西部陆海新通道/平陆运河/农民市", "通道/枢纽。"),
    ("15", "CPI-0.3%、PPI-2.9%", "通缩/物价弱。"),
]
for d in details:
    para(doc, "**%s**" % d[0], size=11)
    para(doc, d[1], size=10.5, color=GRAY)

    para(doc, "**备注**　这条“错位”在广西尤其鲜明：增长靠工业/出口/通道，但投资/物价/消费偏弱。2026年若内需/投资修复，增长可能从工业单极走向“制造+内需+投资”多极。这条细节，正是广西2026年最难也最值得盯的“变量”。", size=10.5)

# ---- 五、2025年目标 vs 实际 对照 ---- 
heading1(doc, "五、2025年GDP目标 vs 实际：对照表")
t5 = [
    ["指标", "2025目标", "2025实际", "达标判定"],
    ["GDP增速", "5.0%左右", "+5.1%", "达标(略超)"],
    ["规上工业增加值增速", "7.0%以上", "+7.7%", "达标"],
    ["固定资产投资增速", "3.0%左右", "-8.2%", "大幅未达标"],
    ["社会消费品零售增速", "5.0%左右", "+3.0%", "未达标"],
    ["一般公共预算收入增速", "3.0%左右", "+4.6%", "达标"],
    ["居民人均可支配收入增速", "与GDP同步", "+5.1%", "基本同步"],
    ["进出口总额/增速", "稳中有升", "8192.62亿/+8.4%", "达标"],
    ["常住人口城镇化率", "提高", "58.09%", "达标"],
    ["CPI", "3.0%左右", "-0.3%", "远低于/偏弱"],
]
table(doc, t5[0], t5[1:], widths=[3.4, 3.0, 3.6, 3.4])
para(doc, "**对照结论**　目标“偏进攻”：工业/财政/民生都定得不低，但**投资/消费/物价是最大失分项**（固投-8.2%、社零+3.0%未达5.0%、CPI负值）。工业与出口接住了，靠“面向东盟+通道”撑起5.1%的增长。")

# ---- 六、增速分项支撑 ---- 
heading1(doc, "六、2025年增长是谁撑起来的？（结构归因）")
para(doc, "GDP+5.1%背后，是“**新动能强、传统/投资弱**”的结构。把增长贡献拆开看：")
g6 = [
    ["引擎", "贡献/状态", "说明"],
    ["高技术制造业", "+23.7%", "最亮引擎：材料/锂电/装备"],
    ["装备制造", "+16.5%", "高端装备、出海支撑"],
    ["新能源(锂电/风/光伏)", "放量", "新能源汽车/锂电池/风电"],
    ["出口(面向东盟)", "+10.6%", "东盟/一带一路外贸"],
    ["财政/税收", "+4.6%/+3.9%", "积极财政、税收转正"],
    ["民间投资", "+2.4%", "民间回暖"],
    ["地产投资", "-16.2%", "地产回落、约束"],
    ["固定资产投资", "-8.2%", "建/地产/基建拖累"],
    ["CPI", "-0.3%", "物价偏弱、内需不足"],
]
table(doc, g6[0], g6[1:], widths=[3.4, 3.4, 6.6])
para(doc, "**一句话**　增长靠“工业+出口+通道”，但投资/地产/物价是拖累。2026年考验广西“能不能让内需/投资/地产也跟着强起来”。")

# ---- 七、预算与财政 ---- 
heading1(doc, "七、预算与财政的“含金量”")
para(doc, "2025年广西一般公共预算收入**1922.05亿元、+4.6%**，其中税收收入**+3.9%**。财政平稳、质量上移形态。")
para(doc, "**要点**：")
bullet(doc, "一般公共预算收入+4.6%，与GDP同步。")
bullet(doc, "税收+3.9%，随工业/出口企稳。")
bullet(doc, "财政“稳收+工业/通道支出优先”，支撑平陆运河、西部陆海新通道等重大工程。")

# ---- 八、民生与居民 ---- 
heading1(doc, "八、民生底账：人口、收入与城乡")
para(doc, "2025年广西常住人口**4989万人、-24万**，城镇化率**58.09%、+0.70pct**。")
para(doc, "居民人均可支配收入**32721元、+5.1%**，城乡收入差仍待缩小。")
g8 = [
    ["指标", "数值", "信号"],
    ["常住人口", "4989万/-24万", "人口负增、流出"],
    ["城镇化率", "58.09%/+0.70pct", "稳步城镇化"],
    ["居民人均可支配收入", "32721元/+5.1%", "收入/民生稳"],
    ["CPI", "-0.3%", "物价偏弱"],
]
table(doc, g8[0], g8[1:], widths=[4.2, 4.2, 4.6])
para(doc, "**民生观察**：收入增长与GDP同步，但人口仍在净流出、物价偏弱，内需仍是最大短板。")

# ---- 九、城镇与农村 ---- 
heading1(doc, "九、城镇与农村：格局与均衡")
para(doc, "广西城镇化率58.09%、+0.70pct，城乡协调持续推进；乡村/线上消费好于总体，城乡收入差仍待缩。")
bullet(doc, "乡村社零+7.0%、线上+10.4%，消费下沉明显。")
bullet(doc, "城乡收入比仍偏高，机会仍集中在南宁/北部湾/沿边。")
bullet(doc, "人口向都市圈/通道枢纽集中。")

# ---- 十、外来人口 ---- 
heading1(doc, "十、人口流入与流出")
para(doc, "广西2025年人口净流出压力仍在（常住-24万），外来输入主要流向南宁、北部湾、沿边开放和制造/通道地区；外出务工与“回流”并存。")
para(doc, "未来看点：面向东盟/平陆运河/沿海制造能否留住人口；若“西部陆海新通道+制造/园区”成型，广西有望从“流出大省”走向“人口/就业托底”。")

# ---- 十一、物价与货币 ---- 
heading1(doc, "十一、物价与货币环境")
para(doc, "2025年广西CPI**-0.3%**、PPI**-2.9%**，物价整体承压、需求偏弱。")
para(doc, "物价偏弱反映“制造强、消费/投资弱”，与全国低通胀环境一致。2026年“把物价/需求稳住”是内需主线。")

# ---- 十二、区域一体化 ---- 
heading1(doc, "十二、区域一体化：广西在“西部陆海新通道+北部湾+东盟”里的位置")
para(doc, "广西的核心战略坐标是“**面向东盟+西部陆海新通道+北部湾**”，也是中国-东盟合作的门户。")
bullet(doc, "西部陆海新通道：平陆运河、北部湾港扩容，出海枢纽。")
bullet(doc, "面向东盟：进出口+8.4%、对外投资+69.1%，通道经济。")
bullet(doc, "沿边/沿海开放：RCEP+6.1%，协定红利。")
para(doc, "若这一“通道+制造+开放”闭环跑通，广西将抢占新一轮“东盟/蓝色”增长窗口。")

# ---- 十三、五条主线 ---- 
heading1(doc, "十三、未来5—10年最值得观察的五条主线")
lines5 = [
    ("① 面向东盟/跨境出海", "东盟、一带一路能否放大广西开放红利。"),
    ("② 通道/枢纽(平陆运河/北部湾)", "西部陆海新通道能否成为增长极。"),
    ("③ 高技术/装备/新能源升级", "高技术、锂电、风电能否持续壮大。"),
    ("④ 投资/地产再平衡", "-8.2%/-16.2%，能否用通道/制造补。"),
    ("⑤ 人口/城乡/民生", "人口净流出能否被制造/通道托底。"),
]
for l in lines5:
    para(doc, "**%s**　%s" % l, space_after=6)

# ---- 十四、结论 ---- 
heading1(doc, "十四、最终结论：广西在“面向东盟+通道+制造”里的增长逻辑")
para(doc, "广西的2025年，本质上是“**制造/出口/东盟/通道为核心，而投资/地产/物价偏弱**”的答卷：GDP29727.45亿、+5.1%，高技术+23.7%、外贸+8.4%，但固投-8.2%、地产-16.2%、CPI-0.3%。")
para(doc, "只要面向东盟开放、通道建设、高技术制造能持续，广西就站在“西部陆海新通道+东盟门户”的增长位；如果投资/地产/内需持续偏弱、人口继续流出，广西需承受“外向强、内需弱”的结构挑战。")
para(doc, "最稳观察信号：**一盯东盟/出海（开放）、二盯通道/平陆（枢纽）、三盯高技术/装备/新能源（动能）、四盯投资/地产（约束）、五盯消费/人口/民生（内需）。**广西，是“面向东盟+通道+制造”的新样本。")

# ---- 附录A ---- 
heading1(doc, "附录A：主要资料来源与核验路径")
bullet(doc, "广西壮族自治区2025年《政府工作报告》——目标来源。")
bullet(doc, "《2025年广西壮族自治区国民经济和社会发展统计公报》—— GDP、工业、外贸、人口 实值。")
bullet(doc, "广西统计、南宁海关、财政厅——外贸/财政。")
bullet(doc, "2026年广西区政府工作报告——2025执行复盘。")
heading2(doc, "核验说明")
para(doc, "本报告所有增速、总量、占比均以统计公报/官方口径为基准，涉“西部陆海新通道/平陆运河/北部湾”等以官方口径为准。")

# ---- 附录B ---- 
heading1(doc, "附录B：建议建立的年度跟踪仪表盘")
para(doc, "建议把下面10个“测脉搏”指标做成年度跟踪表：")
dash = [
    ["#", "指标", "2025基值", "为什么重要"],
    ["1", "GDP增速", "+5.1%", "总量与方向"],
    ["2", "规上工业增速", "+7.7%", "制造底盘"],
    ["3", "高技术制造增速", "+23.7%", "新质生产力"],
    ["4", "进出口/出口增速", "+8.4%/+10.6%", "东盟/通道"],
    ["5", "固定资产投资/地产", "-8.2%/-16.2%", "投资结构"],
    ["6", "社零增速", "+3.0%", "内需消费"],
    ["7", "常住人口/城镇化率", "4989万/58.09%", "人口与城市"],
    ["8", "一般公共预算收入增速", "+4.6%", "财政质量"],
    ["9", "居民人均可支配收入增速", "+5.1%", "民生获得感"],
    ["10", "CPI/PPI", "-0.3%/-2.9%", "物价与需求"],
]
table(doc, dash[0], dash[1:], widths=[0.8, 4.6, 3.4, 4.4])
para(doc, "把这10个指标连起来看，任何一个“东盟/通道（4）、高技术（3）、制造（2）向上、投资/地产（5）修复”，都说明广西在真正换挡。")

# ========================================= 保存
out = "/Users/x/Desktop/content-prod-lab/reports/广西壮族自治区_2025年政府工作报告_深度研究_2026-08-13.docx"
doc.save(out)
print("SAVED:", out)
print("PARAGRAPHS:", len(doc.paragraphs))
print("TABLES:", len(doc.tables))
