# -*- coding: utf-8 -*-
"""Build 长春市2025年政府工作报告 深度研究 DOCX, 参照省会系列版式。"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DARK = RGBColor(0x1F, 0x2A, 0x44)
RED = RGBColor(0xB0, 0x1E, 0x2E)
GRAY = RGBColor(0x59, 0x60, 0x69)
BLUE = RGBColor(0x1F, 0x4E, 0x79)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
HEADER_FILL = "1F2A44"
K = "微软雅黑"


def set_run(run, size=11, bold=False, color=DARK, font=K):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    run.font.name = font
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), font)
    rFonts.set(qn("w:hAnsi"), font)
    rFonts.set(qn("w:eastAsia"), font)


def para(doc, text, size=11, bold=False, color=DARK, align=None,
         space_after=6, space_before=0, font=K):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        set_run(r, size=size, bold=bold or (i % 2 == 1), color=color, font=font)
    return p


def heading1(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(16)
    pf.space_after = Pt(8)
    r = p.add_run(text)
    set_run(r, size=16, bold=True, color=RED)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "B01E2E")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def heading2(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(10)
    pf.space_after = Pt(4)
    r = p.add_run(text)
    set_run(r, size=13, bold=True, color=BLUE)
    return p


def table(doc, headers, rows, widths=None, font_size=9.5):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_run(r, size=font_size, bold=True, color=WHITE)
        tcPr = hdr[i]._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), HEADER_FILL)
        tcPr.append(shd)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(val))
            set_run(r, size=font_size, color=DARK)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    return t


def quote_box(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(8)
    pf.left_indent = Pt(12)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), "B01E2E")
    pBdr.append(left)
    pPr.append(pBdr)
    r = p.add_run(text)
    set_run(r, size=11, color=DARK, font="楷体")
    return p


def bullet(doc, text, size=10.5):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.7)
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        set_run(r, size=size, bold=(i % 2 == 1), color=DARK)
    return p


# ================================================================= 文档主体
doc = Document()
sec = doc.sections[0]
sec.page_width = Cm(21)
sec.page_height = Cm(29.7)
sec.top_margin = Cm(2.4)
sec.bottom_margin = Cm(2.2)
sec.left_margin = Cm(2.5)
sec.right_margin = Cm(2.5)

st = doc.styles["Normal"]
st.font.name = K
st.font.size = Pt(11)
st.element.rPr.rFonts.set(qn("w:eastAsia"), K)


# ---- 封面 ----
para(doc, "长春市2025年政府工作报告", size=24, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=60, space_after=4)
para(doc, "深度研究与\u201c容易被忽视的细节\u201d分析", size=16, bold=True, color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
para(doc, "从\u201c汽车城、全球整车龙头、高端装备、冰雪经济与光电信息\u201d重新理解长春", size=12, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)
para(doc, "\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501", color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
para(doc, "研究对象：2025年长春市政府工作报告及同期财政、国民经济和社会发展统计资料、2026年官方复盘", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
para(doc, "标定日期：2026-08-13", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
doc.add_page_break()

# ---- 目录 ----
catalog = [
    "执行摘要",
    "一、研究口径与阅读方法",
    "二、先看长春的特殊底盘：汽车城、装备制造、粮仓与冰雪文旅",
    "三、最关键的宏观错位：GDP破8000亿但弱于目标，工业强、投资与进出口双降",
    "四、容易被忽视、但信号很重的细节（15条）",
    "五、2025年GDP目标 vs 实际：对照表",
    "六、2025年增长是谁撑起来的？（结构归因）",
    "七、预算与财政的\u201c含金量\u201d",
    "八、民生底账：人口、收入与城乡",
    "九、城镇与农村：格局与均衡",
    "十、人口流入与流出",
    "十一、物价与货币环境",
    "十二、区域一体化：长春在\u201c长春都市圈+中韩(长春)示范+东北亚\u201d里的位置",
    "十三、未来5—10年最值得观察的五条主线",
    "十四、最终结论：长春在\u201c汽车+装备+光电+冰雪\u201d里的增长逻辑",
    "附录A：主要资料来源与核验路径",
    "附录B：建议建立的年度跟踪仪表盘",
]
para(doc, "目　录", size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=10, space_after=10)
for item in catalog:
    para(doc, item, size=12, space_after=4)
doc.add_page_break()

# ---- 执行摘要 ----
heading1(doc, "执行摘要")
para(doc, "**核心判断**　2025年长春最显著的是\u201cGDP 8005.59亿元、增长4.9%（低于5.5%目标）、三产占58.8%\u201d、\u201c规上工业增加值+6.2%（汽车制造业+3.4%）\u201d、\u201c进出口1057.92亿元/-17.0%\u201d、\u201c常住人口910.50万、城镇化69.2%\u201d。这说明长春在\u201c汽车+装备+光电+冰雪\u201d驱动下保持正增长，但**投资-16.5%、进出口-17%双降**是显著压力。")
para(doc, "把2025年目标（GDP+5.5%以上/规上+7%/固投+5%/社零+6%）、2025年统计、2026年前瞻一起看，长春是\u201c东北振兴+汽车强市\u201d路径：**汽车（红旗/新能源）、冰雪经济、光电**是亮点，但投资与开放偏弱。总量8005亿居吉林第一、占全省约一半。")
para(doc, "最容易记住的一句话：**长春是\u201c中国汽车城+高端装备+冰雪文旅\u201d的吉林省会，靠\u201c整车（红旗/大众）、新能源汽车、光电信息、冰雪旅游\u201d增长。**观察长春，与其只看\u201cGDP 8005亿\u201d，不如看\u201c新能源汽车产量+25.6%、红旗销量破45万、冰雪游客2.3亿/+33%、光电产业破1000亿\u201d。")# ---- 一、研究口径与阅读方法 ----
heading1(doc, "一、研究口径与阅读方法")
para(doc, "本报告以\u201c长春市2025年政府工作报告（2025年1月，王子联作）\u201d为起点，把\u201c2025年GDP目标（5.5%以上）\u201d与\u201c2025年国民经济和社会发展统计公报实际值（8005.59亿元/+4.9%）\u201d并置对照，再用2026年政府工作报告的复盘作事后验证。")
para(doc, "重点阅读三类证据：**一是指标目标是否达标**；**二是结构（谁贡献增长）是否匹配目标叙事**；**三是官方复盘如何自评、承认问题**。统计口径以统计公报与市财政局科为准。")
para(doc, "统一采用\u201c同比增长%\u201d（GDP为不变价），财政与居民收入多为名义值，文中按原口径转录、不逐项换算。")

# ---- 二、特殊底盘 ----
heading1(doc, "二、先看长春的特殊底盘：汽车城、装备制造、粮仓与冰雪文旅")
para(doc, "**区位与身份**：长春是吉林省会、副省级城市、中国\u201c汽车城\u201d（一汽所在地），老工业基地+国家全面振兴东北核心，是长春都市圈、中韩(长春)国际合作示范、长吉图先导区承载地。")
para(doc, "**产业底盘**：一是汽车（红旗/智能网联新能源、新能源汽车产量+25.6%、红旗销量破45万辆）；二是高端装备（氢能动车组、\u201c吉林一号\u201d卫星144颗）；三是光电信息（产业规模破1000亿）+医药健康（+10.2%）；四是冰雪经济/文旅（游客2.3亿/+33%）。")
para(doc, "**农业底盘**：全国著名粮仓（玉米/水稻），粮食产量1329.15万吨。")
para(doc, "**人口底盘**：2025年末常住人口910.50万/+（城镇化69.15%、提高0.72pct），是东北最大人口城市之一。")
para(doc, "**市场与出口**：2025年社零2551.17亿/+3.9%；进出口1057.92亿/-17.0%（出口+1.4%、进口-23.8%）。")

# ---- 三、核心宏观错位 ----
heading1(doc, "三、最关键的宏观错位：GDP破8000亿但弱于目标，工业强、投资与进出口双降")
para(doc, "**第一组错位**：2025年GDP目标5.5%以上，实际8005.59亿/+4.9%**低于目标0.6pct**，但\u201c十四五\u201d连破两个千亿台阶、总量破8000亿；增速高于全国平均。")
para(doc, "**第二组错位**：规上工业增加值+6.2%（目标+7%略低）、工业增加值+7.6%；但**投资-16.5%、进出口-17.0%、房地产-23.5%**双降。\u201c工业强、投资弱、外贸弱\u201d。")
para(doc, "**第三组错位**：常住910.5万、跨市净流入6.42万（2025），但**自然增长率-4.87\u2030**（出生3.66\u2030<死亡8.53\u2030）严重负自然增长，人口靠机械流入维持。")
para(doc, "一句话：**长春是\u201c汽车强、投资与外贸弱、人口自然负增长\u201d的东北省会**——靠汽车/装备与冰雪维持增长。")

# ---- 四、15条细节 ----
heading1(doc, "四、容易被忽视、但信号很重的细节（15条）")
bullet(doc, "1. **GDP破8000亿**：8005.59亿/+4.9%、\u201c十四五\u201d跨两个千亿台阶。")
bullet(doc, "2. **汽车制造业+3.4%**：慢于全行业，整车资源整合/新能源转型承压。")
bullet(doc, "3. **新能源汽车产量+25.6%、红旗45万辆**：智能网联新能源主线。")
bullet(doc, "4. **中成药+0.8%、医药健康+10.2%**：医药产业回暖。")
bullet(doc, "5. **光电信息破1000亿、\u201c吉林一号\u201d144颗**：新质生产力。")
bullet(doc, "6. **固定资产-16.5%**：工业投资-9.2%、房地产-23.5%塌陷。")
bullet(doc, "7. **进出口-17.0%（出口+1.4%/进口-23.8%）**：开放承压、进口回落。")
bullet(doc, "8. **冰雪游客2.3亿/+33%、总花费3609.9亿/+19.3%**：冰雪经济爆发。")
bullet(doc, "9. **工业利润401.11亿/-20.1%**：盈利承压但总盘大。")
bullet(doc, "10. **常住910.5万、城镇化69.2%**：东北人口最大城市之一、净流入6.42万。")
bullet(doc, "11. **收入：城镇44250元/+4.5%、农村23042元/+5.8%**，城乡1.92。")
bullet(doc, "12. **CPI -0.1%**：低通胀。")
bullet(doc, "13. **就业10.47万、农村转移100万**：就业稳兜底。")
bullet(doc, "14. **地方财政收入424.17亿/+3.1%**：唯一\u201c达标\u201d的硬指标（+3%）。")
bullet(doc, "15. **技术合同268亿、高新技术1361户**：科创但转化偏慢。")# ---- 五、目标 vs 实际对照表 ----
heading1(doc, "五、2025年GDP目标 vs 实际：对照表")
table(doc,
    ["指标", "2025年目标", "2025年实际", "是否达标"],
    [
        ["地区生产总值增速", "5.5%以上", "8005.59亿元/+4.9%", "未达标"],
        ["规上工业增加值", "+7%以上", "+6.2%", "略低于目标"],
        ["固定资产投资", "+5%以上", "-16.5%（转负）", "严重未达标"],
        ["社会消费品零售总额", "+6%以上", "2551.17亿元/+3.9%", "未达标"],
        ["一般公共预算收入", "+3%", "424.17亿元/+3.1%", "达标"],
        ["城镇/农村人均可支配收入", "高于经济增速", "+4.5%/+5.8%", "均低于GDP但接近"],
        ["CPI涨幅", "2%左右", "-0.1%", "偏低"],
    ],
    widths=[3.4, 2.5, 5.0, 3.1])
para(doc, "**简评**：2025年长春除\u201c财政\u201d外**几乎所有硬指标未达目标**；**投资-16.5%最惨**。汽车+冰雪是韧性来源，工业+6.2%接近目标。")

# ---- 六、结构归因 ----
heading1(doc, "六、2025年增长是谁撑起来的？（结构归因）")
para(doc, "**三大产业**：一产529.04亿/+4.5%、二产2766.17亿/+5.1%、三产4710.37亿/+4.7%；三产占GDP 58.8%。增长以三产+二产双轮。")
para(doc, "**工业**：工业增加值2315.74亿/+7.6%；规上+6.2%、制造业盈利394.83亿。**汽车制造+3.4%、农副食品+3.8%、电力热力+3.8%、电子+3.2%、通用设备+14.1%**；电气机械-29.5%、专用设备-4.4%、黑色金属-14.4%。工业利润总额401.11亿/-20.1%。")
para(doc, "**服务业**：批发零售+3.3%、其他服务+8.1%、金融-1.5%、交通+1.4%、房地产+1.7%。冰雪/文旅拉动其他服务业。")
para(doc, "**增长归因**：长春增长主要靠**工业（汽车/装备/光电/医药）+三产（其他服务/文旅/冰雪）+消费**；投资、进出口、地产为负。")

# ---- 七、财政 ----
heading1(doc, "七、预算与财政的\u201c含金量\u201d")
para(doc, "2025年地方财政收入424.17亿元/+3.1%（达标目标+3%），税收264.87亿/+4.6%、非税159.30亿/+0.7%；地方财政支出1145.19亿/+8.4%。")
para(doc, "**结构性**：收入靠税收（+4.6%）、财政支出高增（+8.4%）、民生占大头；地方财政收入增速在副省级城市靠前。")
para(doc, "**含金量**：财政\u201c收入稳增、税收强、民生支出高\u201d，是长春少数\u201c达标\u201d亮点；但靠传统车税与转移支付支撑，自主财源仍依赖汽车为主。")

# ---------------- 八、民生 ----------------
heading1(doc, "八、民生底账：人口、收入与城乡")
para(doc, "常住人口910.50万/+（跨市净流入6.42万）、城镇化率69.15%（提高0.72pct）；出生率3.66\u2030/死亡率8.53\u2030，**自然增长率-4.87\u2030**。收入：**城镇44250元/+4.5%、农村23042元/+5.8%**，城乡比1.92。")
para(doc, "就业：城镇新增就业10.47万、农村转移100.06万。消费支出、医保覆盖广。")
para(doc, "**民生结论**：收入农村快于城镇、城乡差距约1.92；人口自然严重负增长、靠机械流入维持总量；就业兜底实。")

# ---------------- 九 ----------------
heading1(doc, "九、城镇与农村：格局与均衡")
para(doc, "长春城镇化率69.2%（较全国低），中心城区（朝阳/南关/高新/新区）承载汽车/金融/光电；县域（榆树、农安、德惠）是重要粮仓（玉米/水稻）。")
para(doc, "城乡收入比1.92（差距适中、较均衡），农村收入快于城镇、依托粮仓与劳务；乡村振兴与\u201c黑土地\u201d。")

# =========== 十 ============
heading1(doc, "十、人口流入与流出")
para(doc, "长春常住910.5万、跨市净流入6.42万（2025），但**自然-4.87\u2030**，增长依赖机械流入（高校+汽车/装备/光电就业+新市民）。")
para(doc, "**流入**：高校（吉林大学等）+汽车/光电/冰雪+留学生留驻；**流出**：东北青壮年/轻劳动外迁。净流入为正，但人口结构老龄（60+比重高）是长期压力。")

# =========== 十一 ============
heading1(doc, "十一、物价与货币环境")
para(doc, "2025年长春CPI**下降0.1%**（通缩），\u201c五涨三降\u201d（衣着+2%、食品-1.1%、交通通信-3.6%）。低通胀/微通缩。")
para(doc, "金融：存款余额20903.53亿/比年初+757.70亿、贷款19191.25亿/+464.70亿。\u201c宽中稳\u201d。")

# =========== 十二 ============
heading1(doc, "十二、区域一体化：长春在\u201c长春都市圈+中韩示范+长吉图\u201d里的位置")
para(doc, "长春是长春都市圈核心（已获批规划）、长吉图开发开放先导区、中韩(长春)国际合作示范区+长春新区。与\u201c沈阳都市圈/京津冀\u201d联动，面向\u201c一带一路\u201d、东北亚。")
para(doc, "产业协同：汽车（一汽-京津冀-长春）、光电、绿色能源、现代农机等集群布局。长春是东三省\u201c外向+都市圈\u201d的核心城市。")

# ============ 十三 =============
heading1(doc, "十三、未来5—10年最值得观察的五条主线")
bullet(doc, "1. **新能源汽车+智能网联**：红旗/大众新车型+新能源汽车产量+25.6%，2026目标+30%。")
bullet(doc, "2. **高端装备+光电信息**：氢能动车组、\u201c吉林一号\u201d、光电破1000亿。")
bullet(doc, "3. **冰雪经济+文旅**：瓦萨滑雪节、亚洲杯等，冰雪游客+33%。")
bullet(doc, "4. **科创与\u201c3转4强7新\u201d**：科创基金300亿、医药健康、低空/具身智能。")
bullet(doc, "5. **投资与开放修复**：固投-16.5%后能否靠都市圈/招商引资修复（2026目标1300项目）。")

# ============ 十四 =============
heading1(doc, "十四、最终结论：长春在\u201c汽车+装备+光电+冰雪\u201d里的增长逻辑")
para(doc, "**结论**：长春2025年的\u201c真相\u201d是——**GDP+4.9%（低于5.5%）、规上工业+6.2%、固投-16.5%、进出口-17%、人口自然-4.87\u2030、财政+3.1%**。它是\u201c汽车强、冰雪火、但投资/外贸弱\u201d的东北省会。")
para(doc, "**对趋势判断**：汽车+装备+光电+冰雪代表长春的\u201c增长点\u201d，投资/开放/人口代表\u201c约束\u201d。**新能源汽车/光电/冰雪**决定成色，**投资修复+创新转化**决定韧性。")
para(doc, "**若只看一个指标**：看**规上工业中的新能源汽车产量增速（+25.6%）+固定资产投资增速（-16.5%）**——长春是\u201c汽车强、但投资缺口大\u201d的老工业基地，能否依托汽车+光电+冰雪修复内需是其转型关键。")

# ------------- 附录A -------------
heading1(doc, "附录A：主要资料来源与核验路径")
bullet(doc, "长春市人民政府《长春市人民政府2025年工作报告》（2025年1月）。")
bullet(doc, "长春市统计局《2025年长春市国民经济和社会发展统计公报》（2026年5月）。")
bullet(doc, "《2026年长春市政府工作报告》（2026年2月）及一图读懂。")
bullet(doc, "吉林省2025年统计公报、长春统计年鉴交叉核验。")

# ------------- 附录B -------------
heading1(doc, "附录B：建议建立的年度跟踪仪表盘")
bullet(doc, "GDP总量/增速 vs 全年目标（+/-百分点）。")
bullet(doc, "规上工业增加值及分行业（汽车/新能源/光电/医药/装备）增速。")
bullet(doc, "固定资产投资（总量/工业/基础设施/房地产）增速。")
bullet(doc, "社会消费品零售总额、网络零售。")
bullet(doc, "货物进出口（人民币）、出口、一汽出口。")
bullet(doc, "地方财政收入/税收/非税、财政支出、民生占比。")
bullet(doc, "常住人口、自然增长率、城镇化率、高校留长。")
bullet(doc, "CPI/核心CPI、规上工业利润总额。")
bullet(doc, "冰雪旅游游客/总花费、冰雪项目、入境游客。")
bullet(doc, "新能源汽车产量、红旗销量、光电产业、\u201c吉林一号\u201d卫星。")

# ------------- 保存 -------------
out = "/Users/x/Desktop/content-prod-lab/reports/长春市_2025年政府工作报告_深度研究_2026-08-13.docx"
doc.save(out)
print("SAVED OK 长春市_2025年政府工作报告_深度研究_2026-08-13.docx")