# -*- coding: utf-8 -*-
"""Build 长沙市2025年政府工作报告 深度研究 DOCX, 参照省市系列版式。"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DARK = RGBColor(0x1F, 0x2A, 0x44)
RED = RGBColor(0xB0, 0x1E, 0x2E)
GRAY = RGBColor(0x59, 0x60, 0x69)
BLUE = RGBColor(0x1F, 0x4E, 0x79)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
HEADER_FILL = "1F2A44"
K = "微软雅黑"


def set_run(run, size=11, bold=False, color=DARK, font=K):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    run.font.name = font
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), font)
    rFonts.set(qn("w:hAnsi"), font)
    rFonts.set(qn("w:eastAsia"), font)


def para(doc, text, size=11, bold=False, color=DARK, align=None,
         space_after=6, space_before=0, font=K):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        set_run(r, size=size, bold=bold or (i % 2 == 1), color=color, font=font)
    return p


def heading1(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(16)
    pf.space_after = Pt(8)
    r = p.add_run(text)
    set_run(r, size=16, bold=True, color=RED)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "B01E2E")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def heading2(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(10)
    pf.space_after = Pt(4)
    r = p.add_run(text)
    set_run(r, size=13, bold=True, color=BLUE)
    return p


def table(doc, headers, rows, widths=None, font_size=9.5):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_run(r, size=font_size, bold=True, color=WHITE)
        tcPr = hdr[i]._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), HEADER_FILL)
        tcPr.append(shd)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(val))
            set_run(r, size=font_size, color=DARK)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    return t


def quote_box(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(8)
    pf.left_indent = Pt(12)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), "B01E2E")
    pBdr.append(left)
    pPr.append(pBdr)
    r = p.add_run(text)
    set_run(r, size=11, color=DARK, font="楷体")
    return p


def bullet(doc, text, size=10.5):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.7)
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        set_run(r, size=size, bold=(i % 2 == 1), color=DARK)
    return p


# ================================================================= 文档主体
doc = Document()
sec = doc.sections[0]
sec.page_width = Cm(21)
sec.page_height = Cm(29.7)
sec.top_margin = Cm(2.4)
sec.bottom_margin = Cm(2.2)
sec.left_margin = Cm(2.5)
sec.right_margin = Cm(2.5)

st = doc.styles["Normal"]
st.font.name = K
st.font.size = Pt(11)
st.element.rPr.rFonts.set(qn("w:eastAsia"), K)



# ---- 封面 ----
para(doc, "长沙市2025年政府工作报告", size=24, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=60, space_after=4)
para(doc, "深度研究与“容易被忽视的细节”分析", size=16, bold=True, color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
para(doc, "从“工程机械之都、文创新、网红消费、民生与财政”重新理解长沙", size=12, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)
para(doc, "━━━━━━━━━━━━━━━━", color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
para(doc, "研究对象：2025年长沙市政府工作报告及同期财政、国民经济和社会发展统计资料、2026年官方复盘", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
para(doc, "标定日期：2026-08-13", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
doc.add_page_break()

# ---- 目录 ----
catalog = [
    "执行摘要",
    "一、研究口径与阅读方法",
    "二、先看长沙的特殊底盘：工程机械之都、文创/网红消费、科创与民生",
    "三、最关键的宏观错位：GDP破1.57万亿、规上工业强，但投资/地产/财税弱",
    "四、容易被忽视、但信号很重的细节（15条）",
    "五、2025年GDP目标 vs 实际：对照表",
    "六、2025年增长是谁撑起来的？（结构归因）",
    "七、预算与财政的“含金量”",
    "八、民生底账：人口、收入与城乡",
    "九、城镇与农村：格局与均衡",
    "十、人口流入与流出",
    "十一、物价与货币环境",
    "十二、区域一体化：长沙在“长株潭一体化+中部崛起+长沙都市圈”里的位置",
    "十三、未来5—10年最值得观察的五条主线",
    "十四、最终结论：长沙在“工程机械+科创/文创+人口/民生”里的增长逻辑",
    "附录A：主要资料来源与核验路径",
    "附录B：建议建立的年度跟踪仪表盘",
]
para(doc, "目　录", size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=10, space_after=10)
for item in catalog:
    para(doc, item, size=12, space_after=4)
doc.add_page_break()

# ---- 执行摘要 ----
heading1(doc, "执行摘要")
para(doc, "**核心判断**　如果只看新闻标题，2025年长沙最显眼的是“GDP破1.57万亿（中部城市前列）、增长4.0%”、“规上工业+7.3%”、“高技术制造+14.3%”和“旅游23334万人次（+8.8%）”。但这份研究真正值得深读的，是这座“工程机械之都+文创/网红消费”的中部省会，如何在固定资产投资（-19.3%）、商品房销售面积（-10.2%）与财税偏弱的背景下，靠“工程机械/电子信息/高技术制造+消费/文旅”实现4.0%的增长。")
para(doc, "把2025年初设定的目标（GDP增长5.5%以上）、2025年《国民经济和社会发展统计公报》、以及2026年报告对2025年的复盘放在一起看，长沙呈现清晰暗线：**从“工程机械/制造业”的强底盘，向“高端装备+电子信息/高技术+文创新消费”转型**。旧引擎（地产、传统基建）在调整；新引擎（电子信息、高技术制造、文创/网红消费）被要求更快补位。")
para(doc, "因此，本报告不按政府工作报告逐段复述，而采用“显性表述—同期数据—制度含义—长期影响”的方式，专门提取容易被忽视、但对判断长沙未来5—10年发展模式更有价值的信号。")
para(doc, "最容易记住的一句话：**长沙是“工程机械之都+创新高地+网红消费之都”的中点城市样本，靠“高端制造+文创/网红消费+科创”撑起增长。**观察长沙，与其看“GDP 1.57万亿”，不如看“工程机械、电子信息、高技术制造、文创、网红消费、长沙都市圈”这几张名片。")
heading2(doc, "一页速览：2025年长沙经济的“表与里”")
table(doc,
    ["维度", "表面现象", "更深层信号"],
    [
        ["增长", "GDP 15737.82亿、+4.0%", "一产3.1%、二产35.0%、三产61.9%"],
        ["产业", "规上工业+7.3%", "高技术+14.3%、装备+8.4%"],
        ["外贸", "进出口2875.4亿、+3.5%", "出口+4.2%、机电占68.8%"],
        ["投资", "固定资产投资-19.3%", "商品房面积-10.2%、三产-22.4%"],
        ["财政", "地方一般公共预算收入1296.87亿、+2.6%", "支出-2.3%"],
        ["消费", "社零5738.93亿、+3.9%", "家电+39.0%、通讯+56.9%"],
        ["人口", "常住1072.14万、城镇化84.51%", "自然增-0.67‰、增长1.0%"],
        ["文旅", "旅游23334.63万人次、+8.8%", "工程机械+文创+网红"],
    ],
    widths=[2.2, 5.4, 8.4])
para(doc, "", size=10, space_after=2)

# ---- 一、研究口径与阅读方法 ----
heading1(doc, "一、研究口径与阅读方法")
heading2(doc, "1.1 本报告的核心底稿")
bullet(doc, "**2025年《政府工作报告》**：2025年1月在省十四届人大三次会议上作（目标：GDP 5.5%以上、规上工业7%+、固投4.5%、社零6%以上）。")
bullet(doc, "**2025年《国民经济统计公报》**（市统计局2026-04发布）——GDP、工业、贸易、人口实数。")
bullet(doc, "**2026年长沙市政府工作报告/复盘**——对2025执行结果追认。")
heading2(doc, "1.2 阅读方法：显性—数据—制度—长期")
para(doc, "每章按“**显性表述→同期数据→制度含义→长期影响**”四层展开。")
para(doc, "**关键判别**：数据优先。例2025年GDP目标5.5%、实际4.0%；规上工业目标7%、实际7.3%达标。长沙“高技术/消费/网红强，投资/地产弱”。")

# ---- 二、底盘 ----
heading1(doc, "二、先看长沙的特殊底盘：工程机械之都、文创/网红消费、科创与民生")
para(doc, "长沙的地盘，取决于它作为“**工程机械之都+科创芯片+文创新消费+网红消费之都**”的特殊定位。它是中部最具活力的省会之一。")
bullet(doc, "**工程机械**：三一重工、中联重科等全球龙头，产业集群“全球工程机械之都”；工程机械投向“电动化/智能化/国际化”。")
bullet(doc, "**文创/网红消费**：文和友、茶颜悦色、黑色经典等网红消费；马栏山视频文创产业园、指尖造物；旅游23334万人次（+8.8%）。")
bullet(doc, "**科创**：岳麓山实验室、湘江科学城、算力、机器人；高技术制造+14.3%；长沙是“全国创新城市”代表。")
bullet(doc, "**民生**：长沙房价低、宜居，“人才引进”吸纳人口（常住+1.0%）、网红“young”活力。")
para(doc, "**制度含义**：长沙把“工程机械、科创、文创/网红消费、宜居民生”当核心资产，面向长株潭一体化与中部崛起。")

# ---- 三、宏观错位 ----
heading1(doc, "三、最关键的宏观错位：GDP破1.57万亿、规上工业强，但投资/地产/财税弱")
para(doc, "2025年长沙GDP 15737.82亿元、+4.0%（一产+4.4%、二产+3.1%、三产+4.5%）。表面看“稳中向好”，但拆开看是“**规上工业/高技术/消费强、投资/地产/财税弱**”的错位：")
para(doc, "**强的部分**：规上工业+7.3%（高技术+14.3%、装备+8.4%、电子信息+18.7%、化工+16.4%）；进出口+3.5%（出口+4.2%）；社零+3.9%（家电+39.0%、通讯+56.9%）；旅游+8.8%。")
para(doc, "**弱的部分**：固定资产投资-19.3%（三产-22.4%、二产-14.3%）；商品房销售面积-10.2%；地方财政收入+2.6%（但支出-2.3%）；CPI +0.6%。")
para(doc, "**核心错位一句话**：长沙“高技术/工业/消费/文创强（高技术+14.3%、社零+3.9%），但投资/地产/财税弱”。2026年“稳投资/地产、强科创/文创、扩大内需/网红”是重点。")
table(doc,
    ["强引擎", "数据", "弱项", "数据"],
    [
        ["规上工业", "+7.3%", "固定资产投资", "-19.3%"],
        ["高技术制造业", "+14.3%", "商品房销售面积", "-10.2%"],
        ["电子信息", "+18.7%", "社会消费品零售总额", "+3.9%"],
        ["旅游/接待", "+8.8%", "地方财政收入", "+2.6%"],
        ["进出口", "+3.5%", "CPI", "+0.6%"],
    ],
    widths=[3.6, 3.0, 4.0, 3.0])
para(doc, "**错位结论**　长沙的增长“强制造/消费/文创、弱投资/地产/财税”。2026年“强科创/文创、稳投资/地产、扩内需”是重点。")

# ---- 四、被忽视细节 ----
heading1(doc, "四、容易被忽视、但信号很重的细节（15条）")
details = [
    ("1", "GDP 15737.82亿、+4.0%（中部前列）", "总量/活力。"),
    ("2", "规上工业+7.0%（高技术+14.3%）", "工业/高技术强。"),
    ("3", "装备制造+8.4%（工程机械）", "工程机械之都。"),
    ("4", "电子信息制造+18.7%、通用设备+8.6%", "电子信息/通用。"),
    ("5", "化工+16.4%、有色冶炼+14.6%", "材料/有色。"),
    ("6", "旅游23334万人次、+8.8%、花费+6.8%", "文创/网红消费。"),
    ("7", "机电出口占68.8%、高技术出口22.1%", "出口结构好。"),
    ("8", "进出口2875.4亿、+3.5%", "外贸稳。"),
    ("9", "社零5738.93亿、+3.9%，家电/通讯旺", "消费升级/网红。"),
    ("10", "长沙都市圈/长株潭一体化", "区域/都市圈。"),
    ("11", "常住+1.0%、城镇化84.51%、人才流入", "人口/人才。"),
    ("12", "常住1072.14万、自然增-0.67‰", "人口/老龄化低。"),
    ("13", "城镇收入72117元、+3.5%、农村+5.3%", "收入/农村快。"),
    ("14", "科技园区/算力/科创，技术合同1520亿", "科创高地。"),
    ("15", "CPI+0.6%（温和）", "物价温和。"),
]
for d in details:
    para(doc, "**%s**" % d[0], size=11)
    para(doc, d[1], size=10.5, color=GRAY)

    para(doc, "**备注**　这条“错位”在长沙尤其鲜明：增长靠规上工业/高技术/消费/文创，但投资/地产/财税弱。2026年若地产/投资修复，增长可能从“工业+消费”走向“工业+投资/地产”多极。这条细节，正是长沙2026年最难也最值得盯的“变量”。", size=10.5)

# ---- 五、2025年目标 vs 实际 对照 ----
heading1(doc, "五、2025年GDP目标 vs 实际：对照表")
t5 = [
    ["指标", "2025目标", "2025实际", "达标判定"],
    ["GDP增速", "5.5%以上", "+4.0%", "未达标"],
    ["规上工业增加值增速", "7%以上", "+7.3%", "达标"],
    ["固定资产投资增速", "4.5%", "-19.3%", "大幅未达标"],
    ["社会消费品零售增速", "6%以上", "+3.9%", "未达标"],
    ["地方一般公共预算收入增速", "2.5%", "+2.6%", "达标"],
    ["居民消费价格(CPI)", "涨幅3%左右", "+0.6%", "远低于"],
]
table(doc, t5[0], t5[1:], widths=[3.4, 3.0, 3.0, 3.6])
para(doc, "**对照结论**　目标“偏进攻”，实际“分化”：规上工业+7.3%、财政+2.6%达标；但GDP 4.0%、固投-19.3%、社零3.9%（均未达）。工业/文创/消费接住，投资/地产是短板。")

# ---- 六、增速分项支撑 ----
heading1(doc, "六、2025年增长是谁撑起来的？（结构归因）")
para(doc, "GDP+4.0%背后，是“**规上工业/高技术/消费/文创强、投资/地产弱**”的结构。把增长贡献拆开看：")
g6 = [
    ["引擎", "贡献/状态", "说明"],
    ["规上工业", "+7.3%", "高技术+14.3%、装备+8.4%"],
    ["电子信息/化工", "+18.7%/+16.4%", "电子信息制造强"],
    ["消费/社零", "+3.9%", "家电+39.0%、通讯+56.9%"],
    ["文创/网红消费", "旅游+8.8%", "文和友/马栏山/网红"],
    ["进出口", "+3.5%", "出口+4.2%、机电占68.8%"],
    ["房地产投资", "-11.8%", "地产调整"],
    ["固定资产投资", "-19.3%", "三产/地产拖累"],
    ["地方财政收入", "+2.6%", "财税稳"],
    ["CPI", "+0.6%", "物价温和"],
]
table(doc, g6[0], g6[1:], widths=[3.6, 3.6, 6.0])
para(doc, "**一句话**　增长靠“高技术制造+消费/文创+出口”，地产/投资是拖累。2026年考验长沙“能否让投资/地产也稳”。")

# ---- 七、预算与财政 ----
heading1(doc, "七、预算与财政的“含金量”")
para(doc, "2025年长沙地方一般公共预算收入**1296.87亿元、+2.6%**（支出-2.3%）。")
bullet(doc, "财政收入+2.6%（较上年回落），但仍正增长。")
bullet(doc, "财政“稳收+民生（教育/卫生）+科创/文创支出”。")
bullet(doc, "支持工程机械、科创、网红消费发展。")

# ---- 八、民生底账 ----
heading1(doc, "八、民生底账：人口、收入与城乡")
para(doc, "2025年长沙常住人口**1072.14万人、+1.0%**，城镇化率**84.51%+0.52pct**。自然增**-0.67‰**。")
para(doc, "城镇居民可支配收入**72117元、+3.5%**，农村**48069元、+5.3%**。")
g8 = [
    ["指标", "数值", "信号"],
    ["常住人口", "1072.14万/+1.0%", "人口流入"],
    ["城镇化率", "84.51%/+0.52pct", "高城镇化"],
    ["城镇居民收入", "72117元/+3.5%", "收入稳"],
    ["农村居民收入", "48069元/+5.3%", "农村快于城镇"],
]
table(doc, g8[0], g8[1:], widths=[4.2, 4.0, 4.8])
para(doc, "**民生观察**：人口持续净流入、农村收入快于城镇，城镇化率84.51%较高；物价温和（+0.6%）。长沙是“人口/宜居”强市。")

# ---- 九、城镇与农村 ----
heading1(doc, "九、城镇与农村：格局与均衡")
para(doc, "长沙城镇化率84.51%、人口持续流入，城乡收入差距小。")
bullet(doc, "农村居民收入+5.3%、快于城镇+3.5%。")
bullet(doc, "乡村社零+4.9%、快于城镇+3.7%，下沉。")
bullet(doc, "湘江新区、长沙经开区、星沙等制造业/创新基地。")

# ---- 十、人口流入与流出 ----
heading1(doc, "十、人口流入与流出")
para(doc, "长沙2025年常住+1.0%（净流入），是高吸引力省会（房价低/校园多/劳动）。")
para(doc, "未来看点：产业/科创/宜居能否持续吸引“新长沙人”/回流。")

# ---- 十一、物价与货币 ----
heading1(doc, "十一、物价与货币环境")
para(doc, "2025年长沙CPI**+0.6%**，物价温和上涨。")
para(doc, "物价温和反映“供需平衡”，长沙消费/文旅强。2026年“扩内需、促消费”是主线。")

# ---- 十二、区域一体化 ----
heading1(doc, "十二、区域一体化：长沙在“长株潭一体化+中部崛起+长沙都市圈”里的位置")
para(doc, "长沙的核心战略坐标是“**长株潭一体化+长沙都市圈+中部崛起”，是中部增长极核心城市。")
bullet(doc, "长株潭：长沙+株洲+湘潭一体化，都市圈。")
bullet(doc, "工程机械：三一/中联，全球工程机械之都。")
bullet(doc, "科创：岳麓山、湘江科学城、马栏山文创。")
bullet(doc, "网红消费：长沙都市圈“夜经济/网红”。")
para(doc, "若“长株潭+都市圈+科创+文创”成势，长沙将是中部最具活力的“网红+制造”强市。")

# ---- 十三、五条主线 ----
heading1(doc, "十三、未来5—10年最值得观察的五条主线")
lines5 = [
    ("① 工程机械/高端装备", "电动化/智能化/国际化能否持续。"),
    ("② 电子信息/高技术制造", "电子信息、半导体/机器人能否成极。"),
    ("③ 文创/网红消费", "娱乐经济、网红消费、有时尚之都。"),
    ("④ 科创——创新高地", "湘江科创、计算/大模型。"),
    ("⑤ 人口/民生/地产", "人才流入、房价低、地产修复。"),
]
for l in lines5:
    para(doc, "**%s**　%s" % l, space_after=6)

# ---- 十四、结论 ----
heading1(doc, "十四、最终结论：长沙在“工程机械+科创/文创+人口/民生”里的增长逻辑")
para(doc, "长沙的2025年，本质上是“**规上工业/高技术/消费/文创核心，投资/地产/财税弱**”的答卷：GDP15737.82亿、+4.0%、规上工业+7.3%、高技术+14.3%、旅游+8.8%、常住+1.0%，但固投-19.3%、商品房面积-10.2%、财政+2.6%。")
para(doc, "只要工程机械、科创、文创/网红、人口流入持续，长沙就站在“中部都市增长极”之位；如果房地产/投资偏弱，长沙需承受“消费/制造强、投资/地产弱”的结构挑战。")
para(doc, "最稳观察信号：**一盯工程机械/装备（制造）、二盯科创/计算（创新）、三盯文创/网红（消费）、四盯地产/投资（约束）、五盯人口/民生（支撑）。**长沙，是“工程机械+科创+网红”的新样本。")

# ---- 附录A ----
heading1(doc, "附录A：主要资料来源与核验路径")
bullet(doc, "长沙市2025年《政府工作报告》（2025年1月）——目标来源。")
bullet(doc, "《长沙2025年国民经济和社会发展统计公报》（市统计局）——GDP、工业、外贸、人口实值。")
bullet(doc, "2026年长沙市政府工作报告——2025复盘/长株潭/科创。")
bullet(doc, "长沙海关（外贸）、市财政（财政）。")
heading2(doc, "核验说明")
para(doc, "本报告所有增速、总量、占比均以统计公报/官方口径为基准，涉“工程机械/文创/科创”等以官方为口径。")

# ---- 附录B ----
heading1(doc, "附录B：建议建立的年度跟踪仪表盘")
para(doc, "建议把下面10个“测脉搏”指标做成年度跟踪表：")
dash = [
    ["#", "指标", "2025基值", "为什么重要"],
    ["1", "GDP增速", "+4.0%", "总量与方向"],
    ["2", "规上工业增速", "+7.3%", "制造底盘"],
    ["3", "高技术制造增速", "+14.3%", "新动能"],
    ["4", "进出口/出口增速", "+3.5%/+4.2%", "外贸"],
    ["5", "固定资产投资/地产", "-19.3%/-11.8%", "投资"],
    ["6", "社零增速", "+3.9%", "内需消费"],
    ["7", "常住人口/城镇化率", "1072万/84.51%", "人口与城市"],
    ["8", "地方财政收入", "+2.6%", "财政质量"],
    ["9", "旅游/文创", "23334万人次", "文创消费"],
    ["10", "CPI", "+0.6%", "物价"],
]
table(doc, dash[0], dash[1:], widths=[0.8, 4.6, 3.4, 4.4])
para(doc, "把这10个指标连起来看，装备/高技术/文创/人口（2/3/7）向上、经开/居住（5）修复，都说明长沙在真正换挡。")

# ============ 保存 ============
out = "/Users/x/Desktop/content-prod-lab/reports/长沙市_2025年政府工作报告_深度研究_2026-08-13.docx"
doc.save(out)
print("SAVED:", out)
print("PARAGRAPHS:", len(doc.paragraphs))
print("TABLES:", len(doc.tables))
