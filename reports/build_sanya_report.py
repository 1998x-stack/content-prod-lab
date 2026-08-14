# -*- coding: utf-8 -*-
"""Build 济宁市2025年政府工作报告 深度研究 DOCX, 参照地级市系列版式。"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DARK = RGBColor(0x1F, 0x2A, 0x44)
RED = RGBColor(0xB0, 0x1E, 0x2E)
GRAY = RGBColor(0x59, 0x60, 0x69)
BLUE = RGBColor(0x1F, 0x4E, 0x79)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
HEADER_FILL = "1F2A44"
K = "微软雅黑"


def set_run(run, size=11, bold=False, color=DARK, font=K):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    run.font.name = font
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), font)
    rFonts.set(qn("w:hAnsi"), font)
    rFonts.set(qn("w:eastAsia"), font)


def para(doc, text, size=11, bold=False, color=DARK, align=None,
         space_after=6, space_before=0, font=K):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        set_run(r, size=size, bold=bold or (i % 2 == 1), color=color, font=font)
    return p


def heading1(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(16)
    pf.space_after = Pt(8)
    r = p.add_run(text)
    set_run(r, size=16, bold=True, color=RED)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "B01E2E")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def heading2(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(10)
    pf.space_after = Pt(4)
    r = p.add_run(text)
    set_run(r, size=13, bold=True, color=BLUE)
    return p


def table(doc, headers, rows, widths=None, font_size=9.5):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_run(r, size=font_size, bold=True, color=WHITE)
        tcPr = hdr[i]._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), HEADER_FILL)
        tcPr.append(shd)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(val))
            set_run(r, size=font_size, color=DARK)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    return t


def quote_box(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(8)
    pf.left_indent = Pt(12)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), "B01E2E")
    pBdr.append(left)
    pPr.append(pBdr)
    r = p.add_run(text)
    set_run(r, size=11, color=DARK, font="楷体")
    return p


def bullet(doc, text, size=10.5):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.7)
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        set_run(r, size=size, bold=(i % 2 == 1), color=DARK)
    return p


# ================================================================= 文档主体
doc = Document()
sec = doc.sections[0]
sec.page_width = Cm(21)
sec.page_height = Cm(29.7)
sec.top_margin = Cm(2.4)
sec.bottom_margin = Cm(2.2)
sec.left_margin = Cm(2.5)
sec.right_margin = Cm(2.5)

st = doc.styles["Normal"]
st.font.name = K
st.font.size = Pt(11)
st.element.rPr.rFonts.set(qn("w:eastAsia"), K)


# ---- 封面 ----
para(doc, "三亚市2025年政府工作报告", size=24, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=60, space_after=4)
para(doc, "深度研究与\u201c容易被忽视的细节\u201d分析", size=16, bold=True, color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
para(doc, "从\u201c海南自贸港、热带旅游、离岛免税、崖州湾科技城、候鸟康养、深海南繁\u201d重新理解三亚", size=12, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)
para(doc, "\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501", color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
para(doc, "研究对象：2025年三亚市政府工作报告及同期财政、国民经济和社会发展统计资料、2026年官方复盘", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
para(doc, "标定日期：2026-08-14", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
doc.add_page_break()

# ---- 目录 ----
catalog = [
    "执行摘要",
    "一、研究口径与阅读方法",
    "二、先看三亚的特殊底盘：自贸港、旅游免税、崖州湾、候鸟、深海",
    "三、最关键的宏观错位：GDP破千亿、社零/旅游/免税强，但二产-3.1%、GDP+4.8%低6.5%、规上弱",
    "四、容易被忽视、但信号很重的细节（15条）",
    "五、2025年GDP目标 vs 实际：对照表",
    "六、2025年增长是谁撑起来的？（结构归因）",
    "七、预算与财政的\u201c含金量\u201d",
    "八、民生底账：人口、收入与城乡",
    "九、城镇与农村：格局与均衡",
    "十、人口流入与流出",
    "十一、物价与货币环境",
    "十二、区域一体化：三亚在海南自贸港、三亚经济圈、国际旅游消费中心\u201c三圈\u201d里的位置",
    "十三、未来5—10年最值得观察的五条主线",
    "十四、最终结论：三亚在\u201c自贸港+旅游免税+崖州湾\u201d里的增长逻辑",
    "附录A：主要资料来源与核验路径",
    "附录B：建议建立的年度跟踪仪表盘",
]
para(doc, "目　录", size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=10, space_after=10)
for item in catalog:
    para(doc, item, size=12, space_after=4)
doc.add_page_break()

# ---- 执行摘要 ----
heading1(doc, "执行摘要")
para(doc, "**核心判断**　2025年三亚最显著的是\u201cGDP 1033.88亿/+4.8%（突破千亿、第三产业占75.9%）、社会消费品零售总额573亿/+12.7%（汽车+125.3%）、旅游1032亿/+9.5%（入境+41%）、离岛免税195.5亿（占全省62.5%）\u201d、\u201c但第二产业-3.1%、规上工业-9.2%、GDP+4.8%低6.5%目标\u201d。这说明三亚在\u201c海南自贸港+热带旅游\u201d中，**旅游消费免税强、二产/工业弱**。")
para(doc, "把2025年目标（GDP+6.5%、固投+8.5%、社零+10%、财收+5%、货物贸易+15%）、2025年实际（GDP+4.8%、社零+12.7%、固投+0.7%、货物+7.8%、服务+50.3%）趋势看，三亚是\u201c旅游+免税+自贸港\u201d路径：**旅游民宿、免税消费、崖州湾（南繁/深海）、海洋、游艇、会展体育\u201d是支柱；三产76%。")
para(doc, "最容易记住的一句话：**三亚是\u201c热带滨海旅游城市、海南自贸港重要门面、国际旅游消费中心\u201d，靠\u201c旅游+免税+崖州湾科创\u201d增长。**观察三亚，与其只看\u201cGDP 1034亿\u201d，不如看\u201c游客3615万（入境106万+41%）、旅游花费1032亿、免税195.5亿、汽车消费+125%、崖州湾南繁深海\u201d。")
# ---- 一、研究口径与阅读方法 ----
heading1(doc, "一、研究口径与阅读方法")
para(doc, "本报告以\u201c三亚市政府工作报告（2025年1月）\u201d为起点，把\u201c2025年GDP目标（6.5%）\u201d与\u201c官方2025年（1033.88亿/+4.8%）\u201d并置对照，用\u201c2025年三亚市统计公报\u201d和\u201c2026年政府工作报告\u201d横向核验。")
para(doc, "口径提示：\u201cGDP、规上工业增速\u201d按可比价（实际）；\u201c投资、消费、进出口、财政\u201d按现价（名义）。\u201c常住人口\u201d用统计公报（约110万+）、城镇化率高。")
para(doc, "指标体系（与研究口径一致）：**总量与增速、产业动能（旅游/免税/崖州湾）、外贸、消费、财政、民生与人口**。")
para(doc, "特别提示（不吃老本）：三亚2024年GDP首次破千亿（+3.4%）、2025年1033.88亿/+4.8%；它不是\u201c只有沙滩\u201d——**海南自贸港、离岛免税、崖州湾国家实验室（南繁/深海）、游艇邮轮、会展体育（亚沙会）\u201d才是真正底色。")
# ---- 二、先看三亚的特殊底盘 ----
heading1(doc, "二、先看三亚的特殊底盘：自贸港、旅游免税、崖州湾、候鸟、深海")
para(doc, "三亚地处海南岛最南端、南海之滨，是**海南自由贸易港核心城市、国际旅游消费中心（旅游免税）、崖州湾科技城（南繁育种/深海）、候鸟康养目的地、热带滨海旅游（亚龙湾/蜈支洲/天涯海角/南山）**。2025年GDP 1033.88亿（破千亿）、常住约110万、城镇化率高、海南省第2（海口之后）。")
para(doc, "五个底盘名词，先立框架：")
bullet(doc, "**旅游/游艇邮轮**　亚龙湾/海棠湾/南山/蜈支洲、游艇1400艘（全国第2）、国际邮轮（全国第1）——\u201c热带旅游\u201d。")
bullet(doc, "**离岛免税/消费**　离岛免税（195.5亿、占全省62.5%）、国际免税城4A、社零+12.7%——\u201c免税消费\u201d。")
bullet(doc, "**崖州湾科技城**　南繁硅谷（野生稻种质圃全球最大）、深海技术全国重点实验室、国家实验室——\u201c南繁+深海\u201d。")
bullet(doc, "**自贸港/开放**　琼港合作、免税新政、免税政策、加工增值、封关运作、境外航线42条——\u201c自贸港窗口\u201d。")
bullet(doc, "**会展/体育/候鸟康养**　海南岛电影节、亚沙会（第六届）、FE电动方程式、环岛自行车、候鸟/医养——\u201c会展康养\u201d。")
para(doc, "这五根（旅游+免税+崖州湾+自贸港+会展）构成三亚独特底盘：**左手旅游免税（消费），右手崖州湾（科创/藏育）**。理解三亚，先理解\u201c旅游、免税、崖州湾\u201d。")
# ---- 三、最关键的宏观错位 ----
heading1(doc, "三、最关键的宏观错位：GDP破千亿、社零/旅游/免税强，但二产-3.1%、GDP+4.8%低6.5%、规上弱")
para(doc, "2025年三亚最需要辨析的一组\u201c错位\u201d：**GDP+4.8%（低6.5%）、社会消费品零售总额+12.7%、旅游总花费1032亿/+9.5%、入境游客+41%、离岛免税195.5亿（占62.5%）强，但第二产业-3.1%、规上工业-9.2%、GDP增速低目标**。")
para(doc, "为什么\u201c消费旅游免税\u201d强，GDP/二产却弱？三解释：")
para(doc, "**其一，三产/消费/旅游主导、强**　三产+6.4%（占75.9%）、社零+12.7%（汽车+125.3%）、旅游、免税——\u201c消费旅游强\u201d。")
para(doc, "**其二，二产/工业弱**　二产-3.1%（规上-9.2%、制造业少）、三亚工业基础弱——\u201c二产短\u201d。")
para(doc, "**其三，GDP+4.8%（低目标6.5%）**　旅游消费好但体量/房地产/固投弱修复、外来——\u201c消费快、总量缓\u201d。")
para(doc, "小结：三亚2025年是\u201c**旅游消费免税强、二产工业弱、GDP低于目标**\u201d：旅游/免税/崖州湾强，工业、规上、基建整体放缓。")

# ---- 四、容易被忽视、但信号很重的细节（15条） ----
heading1(doc, "四、容易被忽视、但信号很重的细节（15条）")
bullet(doc, "**1.社零573亿/+12.7%（汽车+125.3%/粮油+51.6%）**\u201c消费强、免税狂。\u201d")
bullet(doc, "**2.旅游3615万人次/1032亿（+6.8%/+9.5%）**\u201c旅游基石。\u201d")
bullet(doc, "**3.入境游客106万/+41%、入境花费9.34亿美元/+67%**\u201c国际化回暖。\u201d")
bullet(doc, "**4.离岛免税195.5亿（占全省62.5%）、免税城4A**\u201c免税龙头。\u201d")
bullet(doc, "**5.崖州湾：国家实验室、南繁野生稻种质圃（全球最大）**\u201c南繁/深海科创。\u201d")
bullet(doc, "**6.深海技术全国重点实验室、77家种业企业**\u201c海洋+育种。\u201d")
bullet(doc, "**7.货物进出口332.61亿/+7.8%（出口+22.6%）**\u201c外贸出口强。\u201d")
bullet(doc, "**8.服务进出口127.1亿/+50.3%**\u201c服务贸易爆发。\u201d")
bullet(doc, "**9.第二产业-3.1%、规上工业-9.2%（降幅收窄）**\u201c二产/工业弱。\u201d")
bullet(doc, "**10.固定资产投资+0.7%（基础设施+34.8%）**\u201c基建强、地产平。\u201d")
bullet(doc, "**11.居民收入43024元/+3.6%（农村+5.2%）**\u201c增收、城乡缩。\u201d")
bullet(doc, "**12.CPI-0.5%**\u201c低通胀。\u201d")
bullet(doc, "**13.千亿台阶/GDP 1033.88亿**\u201c总量破千亿。\u201d")
bullet(doc, "**14.免税消费、游艇1400艘/国际邮轮（全国第1）**\u201c海洋消费。\u201d")
bullet(doc, "**15.亚沙会/电影节/FE电动方程/环岛赛**\u201c会展体育、热带IP。\u201d")
# ---- 五、2025年GDP目标 vs 实际：对照表 ----
heading1(doc, "五、2025年GDP目标 vs 实际：对照表")
table(doc,
    ["指标", "2025年目标", "2025年实际", "达成评价"],
    [
        ["地区生产总值(GDP)", "增长6.5%左右", "1033.88亿/4.8%", "差1.7pct"],
        ["固定资产投资", "增长8.5%左右", "+0.7%", "大幅不及"],
        ["社会消费品零售总额", "增长10%左右", "573亿/+12.7%", "超额"],
        ["货物进出口", "增长15%以上", "332.61亿/+7.8%", "不及"],
        ["服务进出口", "增长25%以上", "127.1亿/+50.3%", "大幅超额"],
        ["地方一般公共预算收入", "增长5%左右", "（保持增长）", "稳"],
        ["城镇/农村居民收入", "4%/6%左右", "+3.3%/+5.2%", "略低"],
        ["居民消费价格", "2%以内", "-0.5%", "低位"],
    ],
)
para(doc, "注：GDP按不变价。**社零（+12.7%）、服务出口（+50.3%）超额**；**GDP（+4.8%）、固投（+0.7%）、货物（+7.8%）**不及目标。")
para(doc, "拆读：**社零/旅游/免税/服务贸易/崖州湾是亮色**；**二产（-3.1%）、规上（-9.2%）、GDP（+4.8%）**是短板——\u201c旅游消费强、二产工业弱\u201d，是\u201c自贸港旅游城市\u201d样本。")

# ---- 六、2025年增长是谁撑起来的？（结构归因） ----
heading1(doc, "六、2025年增长是谁撑起来的？（结构归因）")
para(doc, "把三亚GDP的4.8%拆开：三次产业分别增3.8%、-3.1%、6.4%（结构10.8：13.3：75.9）。**第三产业（旅游/消费/免税）占比75.9%、+6.4%是绝对引擎**，第二产业（-3.1%）拖累，第一产业（农业）+3.8%。")
para(doc, "2026年三亚强调\u201c国际旅游消费中心、封关运作、崖州湾（南繁/深海）\u201d，聚焦**旅游升级、离岛免税（免税新政）、崖州湾国家实验室、深海/南繁、会展体育（亚沙会）、自贸港封关**——核心是\u201c旅游+免税+科创\u201d。")
para(doc, "**第三产业（服务业/旅游）**：+6.4%（旅游、免税、交通+9.8%、地产+9.9%、会展）——\u201c旅游免税强\u201d。")
para(doc, "**第二产业（工业/制造）**：-3.1%（规上-9.2%、工业基础弱）——\u201c二产短\u201d。")
para(doc, "**第一产业（农业/南繁）**：+3.8%（椰子/芒果/南繁育种、山海）——\u201c南繁农业稳\u201d。")
para(doc, "一句话归因：**2025年三亚增长\u201c靠第三产业（旅游/免税/消费）+崖州湾科创\u201d**，二产工业弱；\u201c旅游+免税\u201d是核心。")

# ---- 七、预算与财政的\u201c含金量\u201d ----
heading1(doc, "七、预算与财政的\u201c含金量\u201d")
para(doc, "2025年三亚**地方一般公共预算收入保持增长（+4.6% 2024）**、争取上级资金194.7亿（支持225个）；支出投向旅游/科教/民生。")
bullet(doc, "财收稳（免税/旅游税、自贸港），原地可用——\u201c财政稳\u201d。")
bullet(doc, "上级资金194.7亿（自贸港/基建/崖州湾）。")
bullet(doc, "金融：QFLP 70亿（全省50%+）、QDLP 10.5亿（30%+）、商品交易额3900亿——自贸港金融。")
para(doc, "**财政含金量小结**：财收稳（+4.6%）、上级支持、金融活；财政对\u201c旅游基建、崖州湾、民生\u201d投入加大。")
# ---- 八、民生底账：人口、收入与城乡 ----
heading1(doc, "八、民生底账：人口、收入与城乡")
para(doc, "2025年三亚**全体居民人均可支配收入43024元（+3.6%）**，其中城镇49225元（+3.3%）、农村25876元（+5.2%），城乡比1.90（缩小）。就业：城镇新增就业4.59万人。")
para(doc, "人口画像：**常住约110万、城镇化率约78%**；旅游/候鸟吸引流动人口、候鸟（冬居）。")
para(doc, "民生投入：新增学位超1万、脱贫人口/乡村振兴、医疗康养——民生平稳。")

# ---- 九、城镇与农村：格局与均衡 ----
heading1(doc, "九、城镇与农村：格局与均衡")
para(doc, "三亚城镇化率约78%；农村（天涯/崖州/育才生态、乡村旅游民宿）；农村收入增速（+5.2%）>城镇（+3.3%），**城乡差距缩小**；和美乡村。")
para(doc, "农业底盘：**芒果/椰子/热带水果、南繁育种（崖州湾）、海洋渔业**——\u201c热带农业+南繁\u201d。")
para(doc, "一句话：\u201c三亚是旅游城市+南繁农业、乡村旅游、城乡融合\u201d。")

# ---- 十、人口流入与流出 ----
heading1(doc, "十、人口流入与流出")
para(doc, "三亚常住约110万（旅游城市）、城镇化高；\u201c旅游/自贸港/崖州湾\u201d吸引（候鸟、人才、东北新海南人）。")
para(doc, "结构观察：**候鸟/流动人口多、旅游淡旺季**；崖州湾创新人才（近千）。")
para(doc, "2026年目标：吸引人才、稳定就业（4.6万）——三亚靠\u201c旅游+自贸港+南繁深海\u201d聚人。")

# ---- 十一、物价与货币环境 ----
heading1(doc, "十一、物价与货币环境")
para(doc, "2025年三亚**CPI-0.5%**——\u201c低通胀、需求结构\u201d。")
bullet(doc, "信贷：QFLP/QDLP、跨境（自贸港）、海南商品交易3900亿——宽信用+金融开放。")
bullet(doc, "消费：社零+12.7%、免税、汽车+125%、金银珠宝+25%——消费强。")
para(doc, "货币环境判断：**宽信用/金融开放、CPI-0.5%**；三亚靠\u201c旅游+免税+自贸港\u201d稳需求（2026 CPI 2%）。")

# ---- 十二、区域一体化：三亚的位置 ----
heading1(doc, "十二、区域一体化：三亚在海南自贸港、三亚经济圈、国际旅游消费中心\u201c三圈\u201d里的位置")
para(doc, "三亚是**海南自贸港核心城市、三亚经济圈（三亚-陵水-乐东-保亭）牵头、国际旅游消费中心引领区**。")
bullet(doc, "**自贸港封关**　2025年12月自贸港全岛封关、两个15%税优、免税——\u201c自贸港门槛\u201d。")
bullet(doc, "**三亚经济圈**　引领琼南、山海联动（陵水/乐东/保亭）、环岛旅游公路。")
bullet(doc, "**国际旅游消费中心**　入境免签、国际航线42、免签扩围、游艇邮轮。")
para(doc, "一句话：**三亚在\u201c自贸港+三亚圈+国际旅游消费\u201d里，最核心是\u201c旅游免税+崖州湾+封关\u201d**；区位、免税、科创是最大优势。")

# ---- 十三、未来5—10年最值得观察的五条主线 ----
heading1(doc, "十三、未来5—10年最值得观察的五条主线")
bullet(doc, "**主线一：自贸港封关（2025底）**\u201c免税/加工增值/两个15%、琼港合作\u201d红利释放。")
bullet(doc, "**主线二：旅游升级/国际化**\u201c入境+41%、国际航线42、游艇邮轮\u201d世界级。")
bullet(doc, "**主线三：崖州湾（南繁/深海）**\u201c国家实验室、种业、深海装备\u201d国家科创。")
bullet(doc, "**主线四：免税消费/体育会展**\u201c免税、亚沙会、电影节\u201d消费与IP。")
bullet(doc, "**主线五：二产补短/人口/产业**\u201c制造、1000亿、聚人\u201d经济韧性。")

# ---- 十四、最终结论 ----
heading1(doc, "十四、最终结论：三亚在\u201c自贸港+旅游免税+崖州湾\u201d里的增长逻辑")
para(doc, "把2025年闭环打穿：**三亚是\u201c热带旅游城市、自贸港核心、国际消费中心\u201d**：GDP 1033.88亿/+4.8%（破千亿）、社零+12.7%、旅游1032亿、免税195.5亿、崖州湾。")
para(doc, "三亚不是\u201c只有沙滩\u201d——它是**旅游免税+崖州湾（南繁深海）+自贸港金融+会展体育**的复合，靠\u201c旅游+免税+科创\u201d驱动；但二产/规上工业弱、GDP低目标。")
para(doc, "一句话结论：**三亚是\u201c国际旅游消费中心、自贸港门面、南繁深海科创\u201d；观察它先看\u201c社零/免税、旅游、服务贸易、崖州湾\u201d，再看\u201c二产、工业、固投\u201d。**它是\u201c三产主导、消费强劲、二产短\u201d的自贸港城市样本。")

# ---- 附录A ----
heading1(doc, "附录A：主要资料来源与核验路径")
bullet(doc, "《2025年三亚市政府工作报告》（2025年1月，2025年目标、2024年回顾破千亿+3.4%）")
bullet(doc, "《2025年三亚市国民经济和社会发展统计公报》（三亚市统计局，2026-03，2025年实际）")
bullet(doc, "《2026年三亚市政府工作报告》（2026年1月，陈希，复盘+2026目标）")
bullet(doc, "三亚市人民政府/统计局（sanya.gov.cn）、海南省政府")

# ---- 附录B：建议建立的年度跟踪仪表盘 ----
heading1(doc, "附录B：建议建立的年度跟踪仪表盘")
bullet(doc, "GDP总量/增速 vs 全年目标。")
bullet(doc, "旅游/游客/入境/花费。")
bullet(doc, "免税/离岛/社零/消费。")
bullet(doc, "货物/服务贸易/外资。")
bullet(doc, "崖州湾/南繁/深海。")
bullet(doc, "固定资产/工业/地产投资。")
bullet(doc, "自贸港政策（免签/封关）。")
bullet(doc, "财收/上级资金/金融。")
bullet(doc, "常住/候鸟/人才。")
bullet(doc, "CPI/会展/体育IP。")

# ---------------- 保存 ----------------
out = "/Users/x/Desktop/content-prod-lab/reports/三亚市_2025年政府工作报告_深度研究_2026-08-14.docx"
doc.save(out)
print("SAVED OK: 三亚市", out)
