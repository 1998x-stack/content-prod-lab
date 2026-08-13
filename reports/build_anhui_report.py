# -*- coding: utf-8 -*-
"""Build 重庆市2025年政府工作报告 深度研究 DOCX, 参照北京/上海/杭州/天津系列版式。"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DARK = RGBColor(0x1F, 0x2A, 0x44)
RED = RGBColor(0xB0, 0x1E, 0x2E)
GRAY = RGBColor(0x59, 0x60, 0x69)
BLUE = RGBColor(0x1F, 0x4E, 0x79)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
HEADER_FILL = "1F2A44"
K = "微软雅黑"


def set_run(run, size=11, bold=False, color=DARK, font=K):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    run.font.name = font
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), font)
    rFonts.set(qn("w:hAnsi"), font)
    rFonts.set(qn("w:eastAsia"), font)


def para(doc, text, size=11, bold=False, color=DARK, align=None,
         space_after=6, space_before=0, font=K):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        set_run(r, size=size, bold=bold or (i % 2 == 1), color=color, font=font)
    return p


def heading1(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(16)
    pf.space_after = Pt(8)
    r = p.add_run(text)
    set_run(r, size=16, bold=True, color=RED)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "B01E2E")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def heading2(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(10)
    pf.space_after = Pt(4)
    r = p.add_run(text)
    set_run(r, size=13, bold=True, color=BLUE)
    return p


def table(doc, headers, rows, widths=None, font_size=9.5):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_run(r, size=font_size, bold=True, color=WHITE)
        tcPr = hdr[i]._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), HEADER_FILL)
        tcPr.append(shd)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(val))
            set_run(r, size=font_size, color=DARK)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    return t


def quote_box(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(8)
    pf.left_indent = Pt(12)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), "B01E2E")
    pBdr.append(left)
    pPr.append(pBdr)
    r = p.add_run(text)
    set_run(r, size=11, color=DARK, font="楷体")
    return p


def bullet(doc, text, size=10.5):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.7)
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        set_run(r, size=size, bold=(i % 2 == 1), color=DARK)
    return p


# ================================================================= 文档主体
doc = Document()
sec = doc.sections[0]
sec.page_width = Cm(21)
sec.page_height = Cm(29.7)
sec.top_margin = Cm(2.4)
sec.bottom_margin = Cm(2.2)
sec.left_margin = Cm(2.5)
sec.right_margin = Cm(2.5)

st = doc.styles["Normal"]
st.font.name = K
st.font.size = Pt(11)
st.element.rPr.rFonts.set(qn("w:eastAsia"), K)


# ---- 封面 ----
para(doc, "安徽省2025年政府工作报告", size=24, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=60, space_after=4)
para(doc, "深度研究与“容易被忽视的细节”分析", size=16, bold=True, color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
para(doc, "从“科技创新、量子信息、新能源汽车与对接长三角”重新理解安徽", size=12, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)
para(doc, "━━━━━━━━━━━━━━━━", color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
para(doc, "研究对象：2025年安徽省政府工作报告及同期财政、国民经济和社会发展统计资料、2026年官方复盘", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
para(doc, "整理时间：2026年8月", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
para(doc, "分析性质：分析性研究 ｜ 非官方解读", size=11, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
doc.add_page_break()

# ---- 目录 ----
catalog = [
    "执行摘要",
    "一、研究口径与阅读方法",
    "二、先看安徽的特殊底盘：科教+汽车+长三角桥头堡",
    "三、最关键的宏观错位：GDP破5.3万亿、制造/汽车强，但投资与地产偏弱",
    "四、容易被忽视、但信号很重的细节（15条）",
    "五、把所有细节连起来：安徽正在经历的“六个换挡”",
    "六、增长暗线：汽车第一省、科技创新/量子与新能源（新质生产力）",
    "七、财政暗线：收入稳、税收强，科教/民生投入大",
    "八、产业暗线：从“农业/家电”到“汽车+高科技+战略性新兴”",
    "九、区域格局：对接长三角、合肥都市圈与皖北",
    "十、人口与城市：超6000万、城镇化加速与老龄化",
    "十一、民营经济：科创、民企与消费",
    "十二、事后验证：用2025年实际执行结果检验报告判断",
    "十三、未来5—10年最值得观察的五条主线",
    "十四、最终结论：安徽在“科技创新+汽车+长三角”里的增长逻辑",
    "附录A：主要资料来源与核验路径",
    "附录B：建议建立的年度跟踪仪表盘",
]
para(doc, "目　录", size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=10, space_after=10)
for item in catalog:
    para(doc, item, size=12, space_after=4)
doc.add_page_break()

# ---- 执行摘要 ----
heading1(doc, "执行摘要")
para(doc, "**核心判断**　如果只看新闻标题，2025年安徽最显眼的是“GDP突破5.3万亿、增长5.5%”、“汽车/新能源汽车产量双双跃居全国第1”和“高技术制造+30.4%、新三样出口+107.8%”。但这份研究真正值得深读的，是这座“科技创新+汽车+绿色/能源”的省份，如何在固定资产投资（-9.2%）、房地产深度调整（-24.5%）的背景下，靠“汽车第一省+战略性新兴（占规上45.4%）+高技术制造+长三角承接”守住增长。")
para(doc, "把2025年初设定的目标（GDP增长5.5%以上）、2025年《国民经济和社会发展统计公报》、以及2026年报告对2025年的复盘放在一起看，安徽呈现清晰暗线：**从“农业/家电/低端制造”的旧底盘，向“新能源汽车+科技创新/量子+高技术/战新+对接长三角”转型**。旧引擎（房地产、一般制造、传统基建）在调整；新引擎（新能源汽车、先进制造、高技术、数字经济、长三角承接）被要求更快补位。")
para(doc, "因此，本报告不按政府工作报告逐段复述，而采用“显性表述—同期数据—制度含义—长期影响”的方式，专门提取容易被忽视、但对判断安徽未来5—10年发展模式更有价值的信号。")
para(doc, "最容易记住的一句话：**安徽是全国“汽车第一省+科技创新（量子/聚变/类脑）+对接长三角”的样本，靠“新能源汽车+先进制造+高技术”实现从“中部农业大省”到“科创/制造强省”的跃迁。**观察安徽，与其看“GDP 5.3万亿”，不如看“新能源汽车产量、高技术/战新占比、科技创新转化与长三角协同”这几张名片。")
heading2(doc, "一页速览：2025年安徽经济的“表与里”")
table(doc,
    ["维度", "表面现象", "更深层信号"],
    [
        ["增长", "GDP 52989亿、+5.5%", "第二产业+5.9%、制造强"],
        ["产业", "规上工业+9.4%、高技术+30.4%", "装备占46.6%、战新占45.4%"],
        ["外贸", "进出口10135.6亿、+17.3%", "出口+17.8%、新三样+107.8%"],
        ["投资", "固定资产投资-9.2%", "地产-24.5%、制造业投资-6.0%"],
        ["财政", "一般公共预算收入4157.3亿、+2.9%", "税收+3.0%"],
        ["消费", "社零23863.1亿、+4.2%", "线上+8.3%"],
        ["人口", "常住6082万、城镇化率63.60%", "城镇化提速、人口负增"],
        ["科创", "汽车产量第一、量子/类脑等", "科技创新高地"],
    ],
    widths=[2.2, 5.2, 8.6])
para(doc, "", size=4, space_after=2)

# ---- 一、研究口径与阅读方法 ----
heading1(doc, "一、研究口径与阅读方法")
heading2(doc, "1.1 本报告的三份核心底稿")
bullet(doc, "**2025年《政府工作报告》**：2025年1月在省十四届人大三次会议上作，给出全年预期目标（GDP增长5.5%以上、实际争取更好结果等）。固定资产投资/进出口/房地产等未全部设具体速率。")
bullet(doc, "**2025年《国民经济和社会发展统计公报》**：安徽省统计局2026年3月19日发布，给出全年实际执行数与增速，是“事后验证”的锚点。")
bullet(doc, "**2026年安徽省政府工作报告及官方复盘**：对2025执行结果的追认，交叉验证“目标—执行—再评价”三段式。")
heading2(doc, "1.2 阅读方法：显性—数据—制度—长期")
para(doc, "每一章按“**显性表述 → 同期数据 → 制度含义 → 长期影响**”四层展开。其中“制度含义”回答“政府为什么这么排优先序”，“长期影响”回答“这对未来5—10年意味着什么”。")
para(doc, "关键判别：**当报告使用的“增长词”和统计公报里的实际数不一致时，数据优先。**例如2025年GDP目标5.5%以上、实际5.5%达标；固定资产投资-9.2%、地产-24.5%。差异反映：安徽“汽车/制造/出口强，房地产投资偏弱”。")

# ---- 二、特殊底盘 ----
heading1(doc, "二、先看安徽的特殊底盘：科教+汽车+长三角桥头堡")
para(doc, "在所有省份里，安徽的“底盘”独特：**科技创新（中科大/量子/类脑）+全国第一汽车大省+对接长三角桥头堡**三合一。常住6082万、城镇化63.6%，是公认的“科创+制造”新贵。")
para(doc, "这决定安徽的多重身份并存：**汽车第一省**（产量/新能源全国第一）、**科技创新**（量子、聚变、类脑、中科大）、**制造/战新**（高技术+30%、战新占45.4%）、**长三角承接**、**农业/家电传统**。")
heading2(doc, "2.1 汽车第一省")
para(doc, "2025年安徽汽车产量368.6万辆、新能源汽车179.4万辆，双双跃居全国第1；汽车出口122.8万辆、全国第一。以比亚迪、奇瑞、蔚来等为龙头的“汽车第一省”崛起。")
heading2(doc, "2.2 科技创新/量子")
para(doc, "中科大、量子信息（“九章”“祖冲之”）、聚变（“人造小太阳”）、类脑等，是安徽科技名片。创新平台+转化（技术合同1.18万亿）构成“科创安徽”。")
heading2(doc, "2.3 对接长三角")
para(doc, "安徽作为长三角的一员已“进群”，承接长三角制造/创新外溢，合肥等与沪苏浙协同。长三角一体化是安徽增长的重要支点。")

# ---- 三、宏观错位 ----
heading1(doc, "三、最关键的宏观错位：GDP破5.3万亿、制造/汽车强，但投资与地产偏弱")
para(doc, "把2025年安徽的宏观面放进一张表，会出现令人惊讶的“错位”：表观增长来自制造/汽车/高技术/出口，而投资/地产偏弱。这个错位，正是读懂安徽的关键。")
heading2(doc, "3.1 “强的一面”")
table(doc,
    ["指标", "2025实绩", "信号"],
    [
        ["GDP", "52989亿、+5.5%", "二产+5.9%、制造强"],
        ["规上工业", "+9.4%", "全国前列"],
        ["高技术制造", "+30.4%", "高技术领跑"],
        ["装备制造", "+18.1%、占46.6%", "高端装备"],
        ["汽车/新能源", "产量全国第一", "汽车第一省"],
        ["出口/新三样", "+17.8%/+107.8%", "出口超强"],
    ],
    widths=[3.2, 5.4, 6.0])
heading2(doc, "3.2 “弱/调整的一面”")
table(doc,
    ["指标", "2025实值", "信号"],
    [
        ["固定资产投资", "-9.2%", "投资回落"],
        ["房地产开发投资", "-24.5%", "地产深度调整"],
        ["商品房销售面积", "-14.5%", "楼市弱"],
        ["制造业投资", "-6.0%", "制造投资降"],
        ["PPI", "-2.9%", "工业通缩"],
    ],
    widths=[3.2, 5.4, 6.0])
para(doc, "**错位结论**　安徽的增长“很强、但也很矛盾”。强的部分（汽车/制造/高技术/出口）与弱的部分（投资/地产/PPI）并存。**真正的焦点是“制造强、投资/地产弱”**：汽车/高技术/出口撑起5.5%增长，但地产-24.5%、固投-9.2%。2026年安徽“稳汽车/科创+扩消费/投资+地产软着陆”是重点。")

# ---- 四、被忽视细节 ----
heading1(doc, "四、容易被忽视、但信号很重的细节（15条）")
details = [
    ("1", "规上工业+9.4%、40个行业32个增长", "制造业全国领先。"),
    ("2", "高技术制造+30.4%、占规上18.1%", "高技术领跑。"),
    ("3", "装备制造+18.1%、占规上46.6%", "高端装备。"),
    ("4", "战略性新兴占规上45.4%、+11.5%", "战新驱动。"),
    ("5", "汽车产量368.6万辆、新能源179.4万辆全国第1", "汽车第一省。"),
    ("6", "汽车出口122.8万辆、+28.7%全国第1", "汽车出海。"),
    ("7", "新三样出口1018.3亿、+107.8%", "新三样爆发。"),
    ("8", "机电产品出口+21.2%、高新出口+16.6%", "出口结构升级。"),
    ("9", "对东盟+38.3%、欧盟+21.6%、一带一路+18%", "多元开放。"),
    ("10", "财政科技投入/支持、全国重点实验室23个", "科创强。"),
    ("11", "技术创新合同成交5591亿、吸纳6180亿", "科创转。"),
    ("12", "常住6082万、城镇化率63.60%（+1.03pct）", "城镇化提速。"),
    ("13", "数字经济核心产业营收1.43万亿、+14.3%", "数字产业。"),
    ("14", "社零23863.1亿、+4.2%、乡村+4.6%", "内需消费。"),
    ("15", "投资/地产偏弱（制造业-6%、地产-24.5%）", "结构换挡。"),
]
for d in details:
    para(doc, "**%s**" % d[0], size=11)
    para(doc, d[1], size=10.5, color=GRAY)

# ---- 五、六个换挡 ----
heading1(doc, "五、把所有细节连起来：安徽正在经历的“六个换挡”")
gear = [
    ("1．动能换挡：从“农业/家电”到“汽车+科创+高技术”", "高技术+30.4%、汽车第一、战新45%。"),
    ("2．产业换挡：从“家电/低端制造”到“新能源车+先进制造”", "汽车/新能源、装备、量子。"),
    ("3．投资换挡：从“基建/地产”到“科创/制造/长三角承接”", "地产-24.5%、科创/制造投资强。"),
    ("4．开放换挡：从“内需”到“出口+新三样+长三角”", "出口+17.8%、新三样+107.8%。"),
    ("5．人口换挡：从“劳务输出”到“长三角/合肥集聚”", "城镇化提速、人口负增。"),
    ("6．动能换挡：从“传统增长”到“科技创新+新能源”", "量子/聚变/类脑、绿色制造。"),
]
for g in gear:
    para(doc, "**%s**　%s" % g, space_after=6)

# ---- 六、增长暗线 ----
heading1(doc, "六、增长暗线：汽车第一省、科技创新/量子与新能源（新质生产力）")
heading2(doc, "6.1 汽车第一省+新质")
para(doc, "安徽汽车产量368.6万辆、新能源179.4万辆全国第一，出口122.8万辆全国第一。比亚迪、奇瑞、蔚来等主导。高技术+30.4%、战新占45.4%，制造/科技三核驱动。")
heading2(doc, "6.2 科技与长三角")
para(doc, "量子、聚变、类脑、中科大是安徽科创名片；长三角一体化让合肥等承接沪苏浙创新外溢。科技+制造+长三角，是安徽跃迁的引擎。")
para(doc, "**这条暗线意味着**：安徽的增长叙事已从“农业/家电”转向“汽车第一省+科技创新+长三角承接”。看安徽，盯住“新能源汽车产量、高技术/战新占比、科技转化、长三角协同”这几组数。")

# ---- 七、财政暗线 ----
heading1(doc, "七、财政暗线：收入稳、税收强，科教/民生投入大")
para(doc, "2025年安徽一般公共预算收入4157.3亿、+2.9%，税收2815.8亿、+3.0%。收入/税收稳，来自汽车/先进制造/科创税源。支出向科教、民生、转型倾斜。")
para(doc, "**制度含义**　安徽财政“收入稳、税收强”，支撑科创转化与长三角承接。关键靠“汽车+高技术”的真实税源。")

# ---- 八、产业暗线 ----
heading1(doc, "八、产业暗线：从“农业/家电”到“汽车+高科技+战略性新兴”")
heading2(doc, "8.1 安徽产业的“表”")
table(doc,
    ["指标", "2025增速/信号", "解读"],
    [
        ["规上工业", "+9.4%", "制造业强"],
        ["高技术制造", "+30.4%", "高技术领跑"],
        ["装备制造", "+18.1%、占46.6%", "高端装备"],
        ["战略性新兴", "占45.4%", "战新主导"],
        ["汽车/新能源", "全国第一", "汽车第一省"],
        ["数字经济", "营收+14.3%", "数字产业"],
        ["出口/新三样", "+17.8%/+107.8%", "出口强"],
    ],
    widths=[4.6, 3.8, 5.0])
heading2(doc, "8.2 从“农业/家电”到“汽车+科创/战新”")
para(doc, "安徽过去以农业、家电（美菱/荣事达）见长，2025年显示“汽车+高技术+战新+科创”成为新引擎。是长三角承接+自身科创的产业跃迁。")

# ---- 九、区域 ----
heading1(doc, "九、区域格局：对接长三角、合肥都市圈与皖北")
para(doc, "安徽深度“对接长三角”，合肥都市圈（合芜蚌等）+皖北承接沪苏浙制造/产业外溢。合肥（科创/汽车）、芜湖（汽车/制造）、蚌埠（制造/转型）、皖北（农业/承接）。")
para(doc, "在长三角一体化下，安徽承接制造/创新/人口，同时向中西部辐射。合肥+都市圈+长三角，是安徽增长布局。")

# ---- 十、人口与城市 ----
heading1(doc, "十、人口与城市：超6000万、城镇化加速与老龄化")
heading2(doc, "10.1 人口总量")
para(doc, "2025年末安徽常住6082万、城镇化率63.60%（+1.03pct），自然增长率约-3.7‰（出生<死亡）。城镇化提速、人口负增并存。")
heading2(doc, "10.2 城市/人才")
para(doc, "合肥是人口/人才集聚地（科创+汽车），带动全省。城镇化每升一个百分点，都对应就业、消费与公共服务新需求。")

# ---- 十一、民营经济 ----
heading1(doc, "十一、民营经济：科创、民企与消费")
heading2(doc, "11.1 民营企业")
para(doc, "安徽民企在汽车（比亚迪/蔚来/奇瑞等）、科创、消费上活跃。民营经济是活力、市场与创新的结合。")
heading2(doc, "11.2 政策/科创")
para(doc, "安徽优化营商环境、支持民营/科技创新。民企、专精特新、长三角合作是安徽未来10年底盘。")

# ---- 十二、事后验证 ----
heading1(doc, "十二、事后验证：用2025年实际执行结果检验报告判断")
heading2(doc, "12.1 年初目标与年末执行对比")
table(doc,
    ["指标", "2025年初目标", "2025实际", "偏差"],
    [
        ["GDP增速", "5.5%以上", "+5.5%", "达标"],
        ["规上工业", "力争", "+9.4%", "超预期"],
        ["进出口", "争取", "+17.3%", "大幅超"],
        ["固定资产投资", "未设", "-9.2%", "回落"],
        ["一般公共预算收入", "未设", "+2.9%", "稳"],
        ["居民收入", "与增长同步", "+5.4%", "同步"],
    ],
    widths=[3.0, 3.0, 3.0, 4.4])
heading2(doc, "12.2 判断验证")
para(doc, "三条核心判断得到支持：**（1）汽车/制造/高技术强（全国第一、高技术+30%），验证；（2）出口/新三样超强（+17.8%/+107.8%），验证；（3）投资/地产偏弱（制造业-6%、地产-24.5%），验证。**")
para(doc, "核心观察：安徽靠“汽车+科创+高技术+出口”守住5.5%增长、制造业汽车反超，但投资/地产偏弱。2026年“稳汽车/科创+扩内需/投资+地产软着陆”是重点。")

# ---- 十三、五条主线 ----
heading1(doc, "十三、未来5—10年最值得观察的五条主线")
lines5 = [
    ("① 汽车第一省/新能源/智能化", "能否从“产量第一”升级为“技术/品牌一流”。"),
    ("② 科技创新/量子/类脑科技", "量子、聚变、类脑能否产业化，孵化独角兽。"),
    ("③ 高技术/战新占比提升", "高技术/战新（45%）能否持续放大。"),
    ("④ 长三角一体化/投资", "对接长三角、承接外溢、补投资/地产。"),
    ("⑤ 绿色/民生/人口", "绿色制造、消费、城镇化、人口负增应对。"),
]
for l in lines5:
    para(doc, "**%s**　%s" % l, space_after=6)

# ---- 十四、结论 ----
heading1(doc, "十四、最终结论：安徽在“科技创新+汽车+长三角”里的增长逻辑")
para(doc, "安徽的2025年，本质上是“**汽车第一省+科技创新+高技术制造+出口**”的答卷：GDP5.3万亿、汽车/新能源全国第一、高技术+30.4%、新三样+107.8%，代价是地产-24.5%、固投-9.2%。")
para(doc, "只要汽车/科创/长三角与世界竞争接住，安徽就继续是“科创+制造”强省；如果投资/地产偏弱、汽车价格战，安徽需平衡“制造强、投资/地产弱”的缺口。")
para(doc, "最稳妥的观察信号：**一盯汽车/新能源/高技术（动能）、二盯科创/量子（长长期）、三盯出口/长三角（开放）、四盯投资/地产（约束）、五盯人口/民生（底座）。**安徽，是“长三角+汽车+科创”的新样本。")

# ---- 附录A ----
heading1(doc, "附录A：主要资料来源与核验路径")
bullet(doc, "安徽省2025年《政府工作报告》（2025年1月）——目标来源。")
bullet(doc, "《2025年安徽省国民经济和社会发展统计公报》（安徽省统计局，2026-03-19）——GDP、工业、外贸、人口实值。")
bullet(doc, "安徽省发改委/科技厅/汽车、量子专题——汽车、科创。")
bullet(doc, "2026年安徽省政府工作报告——2025执行复盘。")
bullet(doc, "合肥海关、省财政厅——外贸/财政。")
heading2(doc, "核验说明")
para(doc, "本报告所有增速、总量、占比均以统计公报/官方发布为基准。涉“量子/类脑/聚变具体产值”“长三角”等以官方口径为准。")

# ---- 附录B ----
heading1(doc, "附录B：建议建立的年度跟踪仪表盘")
para(doc, "建议把下面10个“测脉搏”指标做成年度跟踪表：")
dash = [
    ["#", "指标", "2025基值", "为什么重要"],
    ["1", "GDP增速", "+5.5%", "总量与方向"],
    ["2", "规上工业增速", "+9.4%", "制造底盘"],
    ["3", "高技术制造增速", "+30.4%", "新质生产力"],
    ["4", "新能源汽车产量", "179.4万辆全国第1", "汽车第一省"],
    ["5", "出口/新三样增速", "+17.8%/+107.8%", "出口/出海"],
    ["6", "固定资产投资/地产", "-9.2%/-24.5%", "投资结构"],
    ["7", "一般公共预算收入/税收", "+2.9%/+3.0%", "财政质量"],
    ["8", "常住人口/城镇化率", "6082万/63.60%", "人口与城市"],
    ["9", "社零增速", "+4.2%", "内需消费"],
    ["10", "居民人均可支配收入增速", "+5.4%", "民生获得感"],
]
table(doc, dash[0], dash[1:], widths=[0.8, 4.5, 3.6, 4.4])
para(doc, "把这10个指标连起来看，任何一个“新引擎（3/4/5）向上、旧引擎（6）修复”，都说明安徽在真正换挡。")

# ========================================= 保存
out = "/Users/x/Desktop/content-prod-lab/reports/安徽省_2025年政府工作报告_深度研究_2026-08-13.docx"
doc.save(out)
print("SAVED:", out)
print("PARAGRAPHS:", len(doc.paragraphs))
print("TABLES:", len(doc.tables))
